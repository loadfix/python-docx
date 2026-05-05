"""Generator for the `docx/bibliography-authoring` corpus manifest.

Builds a fresh document carrying one `<b:Sources>/<b:Source>` entry
(``smith2020``) in the bibliography part plus one inline `<w:sdt>` citation
reference in the body that points at that tag.

Usage::

    python corpus-manifests/scripts/gen_bibliography_authoring.py <output-path>

The output path may be omitted; it defaults to
``fixtures/docx/bibliography-authoring.docx`` relative to the current
working directory (matching the corpus-repo convention).

The generator deliberately uses only the public python-docx API so that the
manifest doubles as an executable contract for the authoring surface.
"""

from __future__ import annotations

import pathlib
import sys

from docx import Document


def build_document():
    document = Document()

    # -- one primary source, reached from the body by tag --
    document.add_citation(
        "smith2020",
        title="Distributed Systems",
        author="Smith, John",
        year=2020,
        city="London",
        publisher="Acme",
    )
    # -- a second source, exercising source_type + field aliasing --
    document.add_citation(
        "einstein1905",
        source_type="JournalArticle",
        title="Zur Elektrodynamik bewegter Koerper",
        author="Einstein, Albert",
        year=1905,
    )

    p = document.add_paragraph("As argued in ")
    p.add_citation_reference("smith2020")
    p.add_run(", ... and again by ")
    p.add_citation_reference("einstein1905")
    p.add_run(".")

    return document


def main(argv: "list[str]") -> int:
    out = pathlib.Path(
        argv[1] if len(argv) > 1 else "fixtures/docx/bibliography-authoring.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(str(out))
    print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
