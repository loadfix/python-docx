# pyright: reportPrivateUsage=false

"""Unit test suite for the docx.alt_chunk module."""

from __future__ import annotations

import io

import pytest

from docx import Document
from docx.alt_chunk import AltChunk
from docx.document import Document as DocumentCls
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.alt_chunk import AltChunkPart, _ext_for_content_type


class DescribeDocumentAddAltChunk:
    """Unit-test suite for `Document.add_alt_chunk`."""

    def it_returns_an_AltChunk_proxy(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk("<p>hello</p>")

        assert isinstance(alt_chunk, AltChunk)

    def it_appends_a_w_altChunk_element_to_the_body(self):
        document: DocumentCls = Document()

        document.add_alt_chunk(b"<p>hi</p>")

        body = document._element.body
        assert len(body.altChunk_lst) == 1

    def it_creates_a_relationship_with_aFChunk_reltype(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk("<p>hi</p>")

        # -- the rId on the altChunk element resolves to an AltChunkPart --
        assert alt_chunk.rId is not None
        assert isinstance(alt_chunk.part, AltChunkPart)
        # -- and the relationship type is aFChunk --
        document_part = document._part
        assert document_part.rels[alt_chunk.rId].reltype == RT.A_F_CHUNK

    def it_encodes_str_content_as_utf_8_bytes(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk("café", content_type="text/plain")

        assert alt_chunk.content == "café".encode("utf-8")

    def it_defaults_the_content_type_to_text_html(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk("<p>hi</p>")

        assert alt_chunk.content_type == "text/html"

    def it_accepts_a_custom_content_type(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk(b"{\\rtf1}", content_type="application/rtf")

        assert alt_chunk.content_type == "application/rtf"

    def it_omits_altChunkPr_when_match_src_is_not_requested(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk("<p>hi</p>")

        # -- no w:altChunkPr child when match_src is left None --
        assert alt_chunk._element.altChunkPr is None
        assert alt_chunk.match_src is None

    def it_writes_altChunkPr_matchSrc_when_match_src_true(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk("<p>hi</p>", match_src=True)

        pr = alt_chunk._element.altChunkPr
        assert pr is not None
        assert pr.matchSrc is not None
        assert pr.matchSrc.val is True
        assert alt_chunk.match_src is True

    def it_writes_matchSrc_off_when_match_src_false(self):
        document: DocumentCls = Document()

        alt_chunk = document.add_alt_chunk("<p>hi</p>", match_src=False)

        pr = alt_chunk._element.altChunkPr
        assert pr is not None
        assert pr.matchSrc is not None
        assert pr.matchSrc.val is False
        assert alt_chunk.match_src is False


class DescribeAltChunkMatchSrcSetter:
    """Unit-test suite for the `AltChunk.match_src` setter."""

    def it_can_toggle_match_src_on_and_off(self):
        document: DocumentCls = Document()
        alt_chunk = document.add_alt_chunk("<p>hi</p>")

        alt_chunk.match_src = True
        assert alt_chunk.match_src is True

        alt_chunk.match_src = False
        assert alt_chunk.match_src is False

        alt_chunk.match_src = None
        assert alt_chunk.match_src is None
        # -- wrapper altChunkPr is removed when it becomes empty --
        assert alt_chunk._element.altChunkPr is None


class DescribeDocumentAltChunks:
    """Unit-test suite for `Document.alt_chunks`."""

    def it_returns_an_empty_list_when_there_are_no_altChunks(self):
        document: DocumentCls = Document()

        assert document.alt_chunks == []

    def it_lists_one_proxy_per_altChunk_in_document_order(self):
        document: DocumentCls = Document()
        document.add_alt_chunk("<p>first</p>")
        document.add_alt_chunk("<p>second</p>", content_type="text/html")

        chunks = document.alt_chunks

        assert len(chunks) == 2
        assert all(isinstance(ch, AltChunk) for ch in chunks)
        assert chunks[0].content == b"<p>first</p>"
        assert chunks[1].content == b"<p>second</p>"

    def it_round_trips_through_save_and_open(self, tmp_path):
        document: DocumentCls = Document()
        document.add_alt_chunk("<p>hello</p>", content_type="text/html")
        path = tmp_path / "roundtrip.docx"
        document.save(str(path))

        reopened: DocumentCls = Document(str(path))

        chunks = reopened.alt_chunks
        assert len(chunks) == 1
        assert chunks[0].content_type == "text/html"
        assert chunks[0].content == b"<p>hello</p>"

    @pytest.mark.parametrize(
        ("content_type", "payload"),
        [
            ("text/html", b"<p>html</p>"),
            ("application/xhtml+xml", b"<p xmlns='http://www.w3.org/1999/xhtml'/>"),
            ("application/rtf", b"{\\rtf1 rtf}"),
            ("text/plain", "plain café".encode("utf-8")),
            ("message/rfc822", b"From: a@b\r\n\r\nmhtml"),
            (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document.main+xml",
                b"<?xml version='1.0'?><w:document xmlns:w='http://schemas.open"
                b"xmlformats.org/wordprocessingml/2006/main'><w:body/></w:document>",
            ),
        ],
    )
    def it_round_trips_every_supported_content_type(
        self, content_type, payload, tmp_path
    ):
        document: DocumentCls = Document()
        document.add_alt_chunk(payload, content_type=content_type)
        path = tmp_path / "chunk.docx"
        document.save(str(path))

        reopened: DocumentCls = Document(str(path))

        chunks = reopened.alt_chunks
        assert len(chunks) == 1
        assert chunks[0].content_type == content_type
        assert chunks[0].content == payload

    def it_round_trips_altChunkPr_matchSrc(self, tmp_path):
        document: DocumentCls = Document()
        document.add_alt_chunk(
            "<p>x</p>", content_type="text/html", match_src=True
        )
        path = tmp_path / "matchsrc.docx"
        document.save(str(path))

        reopened: DocumentCls = Document(str(path))

        chunks = reopened.alt_chunks
        assert len(chunks) == 1
        assert chunks[0].match_src is True


class DescribeDocumentAddHtmlChunk:
    """Unit-test suite for `Document.add_html_chunk` (R14-5)."""

    def it_sets_content_type_to_xhtml(self):
        document: DocumentCls = Document()

        chunk = document.add_html_chunk("<p>hello</p>")

        assert chunk.content_type == "application/xhtml+xml"
        assert chunk.content == b"<p>hello</p>"

    def it_encodes_string_as_utf_8(self):
        document: DocumentCls = Document()

        chunk = document.add_html_chunk("<p>café</p>")

        assert chunk.content == "<p>café</p>".encode("utf-8")

    def it_passes_match_src_through(self):
        document: DocumentCls = Document()

        chunk = document.add_html_chunk("<p>x</p>", match_src=True)

        assert chunk.match_src is True

    def it_round_trips_through_a_BytesIO_stream(self):
        document: DocumentCls = Document()
        document.add_html_chunk("<p>html</p>")

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened: DocumentCls = Document(buf)

        chunks = reopened.alt_chunks
        assert len(chunks) == 1
        assert chunks[0].content_type == "application/xhtml+xml"
        assert chunks[0].content == b"<p>html</p>"


class DescribeDocumentAddTextChunk:
    """Unit-test suite for `Document.add_text_chunk` (R14-5)."""

    def it_sets_content_type_to_text_plain(self):
        document: DocumentCls = Document()

        chunk = document.add_text_chunk("hello world")

        assert chunk.content_type == "text/plain"
        assert chunk.content == b"hello world"

    def it_defaults_to_utf_8_encoding(self):
        document: DocumentCls = Document()

        chunk = document.add_text_chunk("café")

        assert chunk.content == "café".encode("utf-8")

    def it_accepts_a_custom_encoding(self):
        document: DocumentCls = Document()

        chunk = document.add_text_chunk("café", encoding="latin-1")

        assert chunk.content == "café".encode("latin-1")

    def it_round_trips_through_a_BytesIO_stream(self):
        document: DocumentCls = Document()
        document.add_text_chunk("plain text")

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened: DocumentCls = Document(buf)

        chunks = reopened.alt_chunks
        assert len(chunks) == 1
        assert chunks[0].content_type == "text/plain"
        assert chunks[0].content == b"plain text"


class DescribeDocumentAddRtfChunk:
    """Unit-test suite for `Document.add_rtf_chunk` (R14-5)."""

    def it_sets_content_type_to_application_rtf(self):
        document: DocumentCls = Document()

        chunk = document.add_rtf_chunk(b"{\\rtf1 hi}")

        assert chunk.content_type == "application/rtf"
        assert chunk.content == b"{\\rtf1 hi}"

    def it_round_trips_through_a_BytesIO_stream(self):
        document: DocumentCls = Document()
        rtf_bytes = b"{\\rtf1\\ansi\\deff0 Roundtrip}"
        document.add_rtf_chunk(rtf_bytes, match_src=True)

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened: DocumentCls = Document(buf)

        chunks = reopened.alt_chunks
        assert len(chunks) == 1
        assert chunks[0].content_type == "application/rtf"
        assert chunks[0].content == rtf_bytes
        assert chunks[0].match_src is True


class DescribeDocumentAddMhtmlChunk:
    """Unit-test suite for `Document.add_mhtml_chunk` (R14-5)."""

    def it_sets_content_type_to_message_rfc822(self):
        document: DocumentCls = Document()

        chunk = document.add_mhtml_chunk(b"From: a@b\r\n\r\nbody")

        assert chunk.content_type == "message/rfc822"

    def it_round_trips_through_a_BytesIO_stream(self):
        document: DocumentCls = Document()
        mhtml_bytes = b"MIME-Version: 1.0\r\nContent-Type: text/html\r\n\r\n<p>hi</p>"
        document.add_mhtml_chunk(mhtml_bytes)

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened: DocumentCls = Document(buf)

        chunks = reopened.alt_chunks
        assert len(chunks) == 1
        assert chunks[0].content_type == "message/rfc822"
        assert chunks[0].content == mhtml_bytes


class DescribeAltChunkPart:
    """Unit-test suite for `docx.parts.alt_chunk.AltChunkPart`."""

    def it_picks_a_partname_extension_from_the_content_type(self):
        assert _ext_for_content_type("text/html") == ".html"
        assert _ext_for_content_type("application/rtf") == ".rtf"
        assert _ext_for_content_type("text/rtf") == ".rtf"
        assert _ext_for_content_type("application/xhtml+xml") == ".xhtml"
        assert _ext_for_content_type("text/plain") == ".txt"
        assert _ext_for_content_type("application/msword") == ".doc"
        assert _ext_for_content_type("message/rfc822") == ".mht"
        assert _ext_for_content_type("multipart/related") == ".mht"
        assert (
            _ext_for_content_type(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document.main+xml"
            )
            == ".docx"
        )
        assert _ext_for_content_type("weird/thing") == ".bin"

    def it_can_be_loaded_from_blob(self):
        # -- simulate the PartFactory.load path --
        from docx.opc.packuri import PackURI

        part = AltChunkPart.load(
            PackURI("/word/afchunk1.html"),
            "text/html",
            b"<p>x</p>",
            None,  # type: ignore[arg-type]
        )
        assert part.blob == b"<p>x</p>"
        assert part.content_type == "text/html"
