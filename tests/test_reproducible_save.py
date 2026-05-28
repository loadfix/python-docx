"""End-to-end tests for ``Document.save(..., reproducible=True)``.

Guards the 2026.05.2 reproducible-save invariant against the two bugs
W8-B addressed:

1. ``w:rsidR`` must not be minted onto elements that didn't carry one
   in the source — those attributes are session-scoped churn markers
   that have no place in a content-deterministic artefact.

2. A fresh ``Document()`` must expose the Word-2024 namespace set at
   runtime (``w15``, ``w16``, ``w16cex``, ``w16cid``, ``w16du``,
   ``w16sdtdh``, ``w16sdtfl``, ``w16se``, ``cx``–``cx8``, ``aink``,
   ``am3d``, ``oel``). The source of truth is the unzipped template
   under ``src/docx/templates/default-docx-template/``; this test
   catches drift between that tree and the zipped blob that
   :func:`docx.api.Document` actually loads.

Also covers the issue #150 acceptance contract — the unified
``reproducible=`` keyword shared with the sibling pptx / xlsx / vsdx
parents must produce byte-identical output for byte-identical inputs.
"""

from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path

import pytest

from docx import Document


# ---------------------------------------------------------------------------
# Test helpers
# ---------------------------------------------------------------------------

_FIXTURE_CORPUS = Path("/home/ben/code/ooxml-reference-corpus/fixtures/docx")


def _rsidR_count(docx_bytes_or_path) -> int:
    """Return total count of ``w:rsidR="..."`` attributes in ``word/document.xml``."""
    with zipfile.ZipFile(docx_bytes_or_path) as z:
        data = z.read("word/document.xml").decode("utf-8")
    return len(re.findall(r'\sw:rsidR="', data))


def _r_rsidR_count(docx_bytes_or_path) -> int:
    """Return count of ``w:rsidR="..."`` attributes on ``<w:r>`` elements only."""
    with zipfile.ZipFile(docx_bytes_or_path) as z:
        data = z.read("word/document.xml").decode("utf-8")
    # Match ``<w:r`` followed by any attributes up to the closing ``>``
    # that contains ``w:rsidR="..."`` in that attribute list. The
    # ``[^>]*`` can match an immediate space + ``w:rsidR`` (zero
    # preceding attributes) as well as later occurrences.
    return len(re.findall(r"<w:r[ >][^>]*?\bw:rsidR=\"", data))


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class DescribeReproducibleSave_rsidR_preservation:
    """Reproducible save must not mint rsid-family attributes."""

    @pytest.mark.skipif(
        not (_FIXTURE_CORPUS / "bold-text.office.docx").is_file(),
        reason="ooxml-reference-corpus sibling checkout not available",
    )
    def it_does_not_mint_rsidR_on_runs_that_lacked_it(self, tmp_path):
        src = _FIXTURE_CORPUS / "bold-text.office.docx"
        source_count = _rsidR_count(src)
        source_run_rsidR = _r_rsidR_count(src)

        doc = Document(str(src))
        out = tmp_path / "out.docx"
        doc.save(str(out), reproducible=True)

        # Total rsidR count must match the source — no new rsidR
        # attributes have been synthesised.
        assert _rsidR_count(out) == source_count
        # And specifically, no new rsidR landed on any <w:r>.
        assert _r_rsidR_count(out) == source_run_rsidR

    @pytest.mark.skipif(
        not (_FIXTURE_CORPUS / "bold-text.office.docx").is_file(),
        reason="ooxml-reference-corpus sibling checkout not available",
    )
    def it_is_stable_across_round_trips(self, tmp_path):
        src = _FIXTURE_CORPUS / "bold-text.office.docx"

        out1 = tmp_path / "out1.docx"
        out2 = tmp_path / "out2.docx"
        Document(str(src)).save(str(out1), reproducible=True)
        Document(str(out1)).save(str(out2), reproducible=True)

        assert out1.read_bytes() == out2.read_bytes()

    def it_does_not_mint_rsidR_on_loaded_from_package_in_non_reproducible_mode(
        self, tmp_path
    ):
        # Starting from the bundled template (a loaded-from-package
        # part) and adding content must not retroactively stamp
        # rsidR on paragraphs/runs that lack one. Word itself never
        # retroactively stamps rsid attributes on content authored by
        # another session on plain open+save — mirroring that
        # behaviour is what preserves byte-identical fidelity for
        # ``Document(path).save(out)`` round-trips of
        # Microsoft-Word-authored files.
        doc = Document()
        p = doc.add_paragraph("hello")
        p.add_run("world")

        out = tmp_path / "out.docx"
        doc.save(str(out))

        # Non-reproducible save on a loaded-from-package part must
        # leave rsidR churn out of the newly-emitted runs. The runs
        # in question were added via the API and had no rsidR set;
        # the fidelity policy preserves that.
        assert _r_rsidR_count(out) == 0


class DescribeDefaultTemplateNamespaces:
    """Fresh ``Document()`` exposes the Word-2024 namespace set.

    Guards against the zipped ``default.docx`` drifting out of sync
    with the unzipped ``default-docx-template/`` source tree.
    """

    @pytest.fixture
    def required_word_2024_prefixes(self) -> list[str]:
        return [
            "w15",
            "w16",
            "w16cex",
            "w16cid",
            "w16du",
            "w16sdtdh",
            "w16sdtfl",
            "w16se",
            "cx",
            "cx1",
            "cx2",
            "cx3",
            "cx4",
            "cx5",
            "cx6",
            "cx7",
            "cx8",
            "aink",
            "am3d",
            "oel",
        ]

    def it_declares_every_Word_2024_namespace_prefix(self, required_word_2024_prefixes):
        root = Document().part.element
        nsmap = root.nsmap

        missing = [ns for ns in required_word_2024_prefixes if ns not in nsmap]
        assert not missing, f"default template is missing namespaces: {missing}"

    def it_declares_the_mc_Ignorable_attribute(self):
        from docx.oxml.ns import qn

        root = Document().part.element
        mc_ignorable = root.get(qn("mc:Ignorable"))

        assert mc_ignorable is not None, "mc:Ignorable missing from default template"
        # Must cover at least the Word-2024 extension prefixes.
        for token in ("w14", "w15", "w16", "w16se", "w16cid", "w16du"):
            assert token in mc_ignorable.split(), (
                f"mc:Ignorable missing {token!r}: {mc_ignorable!r}"
            )

    def it_matches_the_unzipped_source_tree(self):
        """Every file in the unzipped template must match the zipped blob."""
        templates_dir = (
            Path(__file__).parent.parent
            / "src"
            / "docx"
            / "templates"
        )
        source_tree = templates_dir / "default-docx-template"
        zipped = templates_dir / "default.docx"

        assert source_tree.is_dir()
        assert zipped.is_file()

        with zipfile.ZipFile(zipped) as z:
            zip_names = {name for name in z.namelist() if not name.endswith("/")}
            for arcname in zip_names:
                src_path = source_tree / arcname
                assert src_path.is_file(), (
                    f"{arcname} in default.docx but not in default-docx-template/"
                )
                assert z.read(arcname) == src_path.read_bytes(), (
                    f"{arcname} differs between default.docx and default-docx-template/"
                )

        source_names = {
            p.relative_to(source_tree).as_posix()
            for p in source_tree.rglob("*")
            if p.is_file()
        }
        extra_on_disk = source_names - zip_names
        assert not extra_on_disk, (
            f"files in default-docx-template/ not in default.docx: {extra_on_disk}"
        )


class DescribeReproducibleSaveAcceptance:
    """Issue #150 acceptance contract for ``Document.save(reproducible=True)``.

    Mirrors the parallel suites under ``python-pptx``, ``python-xlsx``,
    and ``python-vsdx`` so the unified API behaves the same in every
    parent.
    """

    def it_is_byte_identical_across_two_fresh_authoring_runs(self):
        def build() -> bytes:
            doc = Document()
            doc.add_paragraph("Hello")
            buf = io.BytesIO()
            doc.save(buf, reproducible=True)
            return buf.getvalue()

        assert build() == build()

    def it_is_byte_identical_across_load_and_resave_round_trips(self, tmp_path):
        seed = tmp_path / "seed.docx"
        doc = Document()
        doc.add_paragraph("Hello")
        doc.save(str(seed), reproducible=True)

        def reload() -> bytes:
            buf = io.BytesIO()
            Document(str(seed)).save(buf, reproducible=True)
            return buf.getvalue()

        assert reload() == reload()

    def it_stamps_every_zip_member_with_the_fixed_timestamp(self):
        doc = Document()
        doc.add_paragraph("Hello")
        buf = io.BytesIO()
        doc.save(buf, reproducible=True)

        with zipfile.ZipFile(io.BytesIO(buf.getvalue())) as z:
            timestamps = {info.date_time for info in z.infolist()}
        assert timestamps == {(1980, 1, 1, 0, 0, 0)}

    def it_emits_zip_members_in_sorted_order(self):
        doc = Document()
        doc.add_paragraph("Hello")
        buf = io.BytesIO()
        doc.save(buf, reproducible=True)

        with zipfile.ZipFile(io.BytesIO(buf.getvalue())) as z:
            names = z.namelist()
        assert names == sorted(names)

    def but_it_does_not_force_a_fixed_timestamp_when_reproducible_is_False(self):
        doc = Document()
        doc.add_paragraph("Hello")
        buf = io.BytesIO()
        doc.save(buf)  # default: reproducible=False

        with zipfile.ZipFile(io.BytesIO(buf.getvalue())) as z:
            timestamps = {info.date_time for info in z.infolist()}
        assert timestamps != {(1980, 1, 1, 0, 0, 0)}
