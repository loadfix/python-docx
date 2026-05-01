"""Unit-test suite for caption-building helpers."""

from __future__ import annotations

from docx import Document
from docx.captions import new_caption_paragraph
from docx.document import Document as DocumentCls
from docx.fields import Field
from docx.text.paragraph import Paragraph


class DescribeDocument_AddCaption:
    """Unit-test suite for `Document.add_caption`."""

    def it_appends_a_caption_paragraph_to_the_body(self):
        document: DocumentCls = Document()
        start_count = len(document.paragraphs)

        paragraph = document.add_caption("A diagram of the system")

        assert isinstance(paragraph, Paragraph)
        assert len(document.paragraphs) == start_count + 1
        assert document.paragraphs[-1]._p is paragraph._p

    def it_applies_the_caption_style_by_default(self):
        document: DocumentCls = Document()

        paragraph = document.add_caption("A diagram")

        assert paragraph.style is not None
        assert paragraph.style.name == "Caption"

    def it_produces_the_expected_run_and_field_sequence(self):
        document: DocumentCls = Document()

        paragraph = document.add_caption("A diagram of the system")

        # -- expected order: "Figure ", SEQ field(result "1"), ": ", "A diagram..."
        assert paragraph.text == "Figure 1: A diagram of the system"

    def it_builds_a_SEQ_field_targeting_the_Figure_label(self):
        document: DocumentCls = Document()

        paragraph = document.add_caption("A diagram")

        assert len(paragraph.fields) == 1
        field = paragraph.fields[0]
        assert isinstance(field, Field)
        assert field.is_complex is False
        assert field.type == "SEQ"
        assert field.instruction.strip() == "SEQ Figure \\* ARABIC"
        assert field.result_text == "1"

    def it_accepts_a_custom_label(self):
        document: DocumentCls = Document()

        paragraph = document.add_caption("A pricing table", label="Table")

        assert paragraph.text == "Table 1: A pricing table"
        field = paragraph.fields[0]
        assert field.instruction.strip() == "SEQ Table \\* ARABIC"

    def it_accepts_a_custom_style(self):
        document: DocumentCls = Document()
        # -- reuse an existing paragraph style rather than authoring a new one --
        style_name = "Heading 1"

        paragraph = document.add_caption("Custom style", style=style_name)

        assert paragraph.style is not None
        assert paragraph.style.name == style_name

    def it_round_trips_through_document_paragraphs(self):
        document: DocumentCls = Document()

        paragraph = document.add_caption("Round trip")

        retrieved_texts = [p.text for p in document.paragraphs]
        assert "Figure 1: Round trip" in retrieved_texts


class DescribeParagraph_AddCaptionBeforeAfter:
    """Unit-test suite for `Paragraph.add_caption_before` / `_after`."""

    def it_inserts_a_caption_after_the_target_paragraph(self):
        document: DocumentCls = Document()
        anchor = document.add_paragraph("anchor paragraph")

        caption = anchor.add_caption_after("A diagram")

        paragraphs = document.paragraphs
        anchor_idx = paragraphs.index(
            next(p for p in paragraphs if p._p is anchor._p)
        )
        assert paragraphs[anchor_idx + 1]._p is caption._p
        assert caption.text == "Figure 1: A diagram"
        assert caption.style is not None
        assert caption.style.name == "Caption"

    def it_inserts_a_caption_before_the_target_paragraph(self):
        document: DocumentCls = Document()
        anchor = document.add_paragraph("anchor paragraph")

        caption = anchor.add_caption_before("A diagram")

        paragraphs = document.paragraphs
        anchor_idx = paragraphs.index(
            next(p for p in paragraphs if p._p is anchor._p)
        )
        assert paragraphs[anchor_idx - 1]._p is caption._p
        assert caption.text == "Figure 1: A diagram"

    def it_honors_custom_label_and_style_on_add_caption_after(self):
        document: DocumentCls = Document()
        anchor = document.add_paragraph("anchor")

        caption = anchor.add_caption_after(
            "Pricing", label="Table", style="Heading 1"
        )

        assert caption.text == "Table 1: Pricing"
        field = caption.fields[0]
        assert field.instruction.strip() == "SEQ Table \\* ARABIC"
        assert caption.style is not None
        assert caption.style.name == "Heading 1"

    def it_honors_custom_label_and_style_on_add_caption_before(self):
        document: DocumentCls = Document()
        anchor = document.add_paragraph("anchor")

        caption = anchor.add_caption_before(
            "Pricing", label="Table", style="Heading 1"
        )

        assert caption.text == "Table 1: Pricing"
        field = caption.fields[0]
        assert field.instruction.strip() == "SEQ Table \\* ARABIC"
        assert caption.style is not None
        assert caption.style.name == "Heading 1"


class DescribeNewCaptionParagraph:
    """Unit-test suite for the low-level `new_caption_paragraph` helper."""

    def it_populates_an_empty_paragraph_with_the_standard_caption_shape(self):
        document: DocumentCls = Document()
        paragraph = document.add_paragraph()

        result = new_caption_paragraph(paragraph, "A diagram")

        assert result is paragraph
        assert paragraph.text == "Figure 1: A diagram"
        assert paragraph.style is not None
        assert paragraph.style.name == "Caption"

    def it_returns_the_populated_paragraph(self):
        document: DocumentCls = Document()
        paragraph = document.add_paragraph()

        result = new_caption_paragraph(paragraph, "text", label="Table")

        assert result is paragraph
        assert paragraph.fields[0].instruction.strip() == "SEQ Table \\* ARABIC"
