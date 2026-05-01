# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.embedded_objects` module and related access."""

from __future__ import annotations

from typing import cast

import pytest

from docx import types as t
from docx.document import Document
from docx.embedded_objects import EmbeddedObject
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml.document import CT_Document
from docx.oxml.text.paragraph import CT_P
from docx.parts.document import DocumentPart
from docx.parts.embedded_object import EmbeddedObjectPart
from docx.parts.story import StoryPart
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element
from .unitutil.mock import FixtureRequest, instance_mock


OLE_BLOB_EXCEL = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 8 + b"excel-workbook"

OLE_BLOB_PDF = b"%PDF-1.4\n%\xc3\xa1\xc3\xa2 embedded-pdf-bytes"


def _make_embedded_part(
    idx: int = 1, blob: bytes = OLE_BLOB_EXCEL
) -> EmbeddedObjectPart:
    return EmbeddedObjectPart(
        PackURI("/word/embeddings/oleObject%d.bin" % idx), CT.OFC_OLE_OBJECT, blob
    )


class DescribeEmbeddedObject:
    """Unit-test suite for `docx.embedded_objects.EmbeddedObject`."""

    def it_knows_its_prog_id(self, fake_parent: t.ProvidesStoryPart):
        p = cast(
            CT_P,
            element(
                "w:p/w:r/w:object/o:OLEObject"
                "{ProgID=Excel.Sheet.12,Type=Embed,r:id=rId1}"
            ),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]
        embedded_part = _make_embedded_part()

        eo = EmbeddedObject(paragraph, ole_elm, embedded_part)

        assert eo.prog_id == "Excel.Sheet.12"

    def it_returns_None_prog_id_when_attribute_missing(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]

        eo = EmbeddedObject(paragraph, ole_elm, None)

        assert eo.prog_id is None

    def it_knows_its_type(self, fake_parent: t.ProvidesStoryPart):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{Type=Link,r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]

        eo = EmbeddedObject(paragraph, ole_elm, None)

        assert eo.type == "Link"

    def it_returns_None_type_when_attribute_missing(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]

        eo = EmbeddedObject(paragraph, ole_elm, None)

        assert eo.type is None

    def it_knows_its_r_id(self, fake_parent: t.ProvidesStoryPart):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId42}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]

        eo = EmbeddedObject(paragraph, ole_elm, None)

        assert eo.r_id == "rId42"

    def it_returns_None_r_id_when_attribute_missing(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{ProgID=Excel.Sheet.12}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]

        eo = EmbeddedObject(paragraph, ole_elm, None)

        assert eo.r_id is None

    def it_exposes_the_embedded_partname_when_resolved(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]
        embedded_part = _make_embedded_part(idx=7)

        eo = EmbeddedObject(paragraph, ole_elm, embedded_part)

        assert eo.embedded_partname == "/word/embeddings/oleObject7.bin"

    def it_returns_None_embedded_partname_when_unresolved(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]

        eo = EmbeddedObject(paragraph, ole_elm, None)

        assert eo.embedded_partname is None

    def it_provides_access_to_the_raw_blob(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]
        embedded_part = _make_embedded_part(blob=OLE_BLOB_PDF)

        eo = EmbeddedObject(paragraph, ole_elm, embedded_part)

        assert eo.blob == OLE_BLOB_PDF

    def it_returns_empty_blob_when_part_is_unresolved(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]

        eo = EmbeddedObject(paragraph, ole_elm, None)

        assert eo.blob == b""

    def it_provides_access_to_the_paragraph_it_belongs_to(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{r:id=rId1}"),
        )
        paragraph = Paragraph(p, fake_parent)
        ole_elm = p.xpath(".//o:OLEObject")[0]
        embedded_part = _make_embedded_part()

        eo = EmbeddedObject(paragraph, ole_elm, embedded_part)

        assert eo.paragraph is paragraph

    # -- fixtures -----------------------------------------------------------------

    @pytest.fixture
    def fake_parent(self, request: FixtureRequest):
        story_part_ = instance_mock(request, StoryPart)

        class FakeParent:
            @property
            def part(self):
                return story_part_

        return FakeParent()


class DescribeParagraph_embedded_objects:
    """Unit-test suite for `Paragraph.embedded_objects`."""

    def it_returns_empty_list_when_no_object_elements(
        self, request: FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.related_parts = {}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(
            CT_P,
            element('w:p/w:r/w:t"some text"'),
        )
        paragraph = Paragraph(p, FakeParent())

        assert paragraph.embedded_objects == []

    def it_enumerates_each_ole_object_in_order(self, request: FixtureRequest):
        embedded_part_1 = _make_embedded_part(idx=1, blob=OLE_BLOB_EXCEL)
        embedded_part_2 = _make_embedded_part(idx=2, blob=OLE_BLOB_PDF)
        story_part_ = instance_mock(request, StoryPart)
        story_part_.related_parts = {
            "rId5": embedded_part_1,
            "rId6": embedded_part_2,
        }

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(
            CT_P,
            element(
                "w:p/("
                "w:r/w:object/o:OLEObject"
                "{ProgID=Excel.Sheet.12,Type=Embed,r:id=rId5},"
                "w:r/w:object/o:OLEObject"
                "{ProgID=AcroExch.Document.DC,Type=Embed,r:id=rId6}"
                ")"
            ),
        )
        paragraph = Paragraph(p, FakeParent())

        eos = paragraph.embedded_objects

        assert len(eos) == 2
        assert [e.prog_id for e in eos] == [
            "Excel.Sheet.12",
            "AcroExch.Document.DC",
        ]
        assert [e.type for e in eos] == ["Embed", "Embed"]
        assert [e.r_id for e in eos] == ["rId5", "rId6"]
        assert [e.embedded_partname for e in eos] == [
            "/word/embeddings/oleObject1.bin",
            "/word/embeddings/oleObject2.bin",
        ]
        assert [e.blob for e in eos] == [OLE_BLOB_EXCEL, OLE_BLOB_PDF]
        assert all(e.paragraph is paragraph for e in eos)

    def it_returns_an_object_with_empty_blob_when_relationship_is_missing(
        self, request: FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.related_parts = {}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(
            CT_P,
            element(
                "w:p/w:r/w:object/o:OLEObject"
                "{ProgID=Excel.Sheet.12,Type=Embed,r:id=rId99}"
            ),
        )
        paragraph = Paragraph(p, FakeParent())

        eos = paragraph.embedded_objects

        assert len(eos) == 1
        assert eos[0].prog_id == "Excel.Sheet.12"
        assert eos[0].r_id == "rId99"
        assert eos[0].blob == b""
        assert eos[0].embedded_partname is None

    def it_returns_an_object_with_empty_blob_when_related_part_is_wrong_type(
        self, request: FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        # -- rId5 is mapped to something that isn't an EmbeddedObjectPart --
        story_part_.related_parts = {"rId5": "not-a-part"}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(
            CT_P,
            element(
                "w:p/w:r/w:object/o:OLEObject"
                "{ProgID=Excel.Sheet.12,r:id=rId5}"
            ),
        )
        paragraph = Paragraph(p, FakeParent())

        eos = paragraph.embedded_objects

        assert len(eos) == 1
        assert eos[0].blob == b""
        assert eos[0].embedded_partname is None

    def it_handles_ole_object_without_relationship_id(
        self, request: FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.related_parts = {}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(
            CT_P,
            element("w:p/w:r/w:object/o:OLEObject{ProgID=Excel.Sheet.12}"),
        )
        paragraph = Paragraph(p, FakeParent())

        eos = paragraph.embedded_objects

        assert len(eos) == 1
        assert eos[0].prog_id == "Excel.Sheet.12"
        assert eos[0].r_id is None
        assert eos[0].blob == b""
        assert eos[0].embedded_partname is None


class DescribeDocument_embedded_objects:
    """Unit-test suite for `Document.embedded_objects`."""

    def it_returns_empty_list_when_document_has_no_ole_objects(
        self, request: FixtureRequest
    ):
        doc_elm = cast(CT_Document, element("w:document/w:body/w:p"))
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}
        document = Document(doc_elm, document_part_)

        assert document.embedded_objects == []

    def it_collects_embedded_objects_across_all_paragraphs(
        self, request: FixtureRequest
    ):
        embedded_part_1 = _make_embedded_part(idx=1, blob=OLE_BLOB_EXCEL)
        embedded_part_2 = _make_embedded_part(idx=2, blob=OLE_BLOB_PDF)
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {
            "rId1": embedded_part_1,
            "rId2": embedded_part_2,
        }

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                "w:p/w:r/w:object/o:OLEObject"
                "{ProgID=Excel.Sheet.12,r:id=rId1},"
                "w:p,"
                "w:p/w:r/w:object/o:OLEObject"
                "{ProgID=AcroExch.Document.DC,r:id=rId2}"
                ")"
            ),
        )
        document = Document(doc_elm, document_part_)

        eos = document.embedded_objects

        assert len(eos) == 2
        assert {e.prog_id for e in eos} == {
            "Excel.Sheet.12",
            "AcroExch.Document.DC",
        }
        assert {e.embedded_partname for e in eos} == {
            "/word/embeddings/oleObject1.bin",
            "/word/embeddings/oleObject2.bin",
        }
        assert sorted([e.blob for e in eos]) == sorted(
            [OLE_BLOB_EXCEL, OLE_BLOB_PDF]
        )

    def it_matches_paragraph_scoped_enumeration(self, request: FixtureRequest):
        embedded_part_1 = _make_embedded_part(idx=1, blob=OLE_BLOB_EXCEL)
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {"rId1": embedded_part_1}

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/w:r/w:object/o:OLEObject"
                "{ProgID=Excel.Sheet.12,Type=Embed,r:id=rId1}"
            ),
        )
        document = Document(doc_elm, document_part_)

        doc_eos = document.embedded_objects
        para_eos = document.paragraphs[0].embedded_objects

        assert len(doc_eos) == 1
        assert len(para_eos) == 1
        assert doc_eos[0].prog_id == para_eos[0].prog_id
        assert doc_eos[0].r_id == para_eos[0].r_id
        assert doc_eos[0].type == para_eos[0].type
        assert doc_eos[0].embedded_partname == para_eos[0].embedded_partname
        assert doc_eos[0].blob == para_eos[0].blob


class DescribeEmbeddedObjectPart:
    """Unit-test suite for `docx.parts.embedded_object.EmbeddedObjectPart`."""

    def it_exposes_its_blob(self):
        partname = PackURI("/word/embeddings/oleObject1.bin")
        part = EmbeddedObjectPart(partname, CT.OFC_OLE_OBJECT, OLE_BLOB_EXCEL)

        assert part.blob == OLE_BLOB_EXCEL
        assert part.partname == partname
        assert part.content_type == CT.OFC_OLE_OBJECT

    def it_is_constructed_by_the_part_loader(self, request: FixtureRequest):
        from docx.opc.part import PartFactory
        from docx.package import Package

        package_ = instance_mock(request, Package)
        part = PartFactory(
            PackURI("/word/embeddings/oleObject1.bin"),
            CT.OFC_OLE_OBJECT,
            RT.OLE_OBJECT,
            OLE_BLOB_EXCEL,
            package_,
        )

        assert isinstance(part, EmbeddedObjectPart)
        assert part.blob == OLE_BLOB_EXCEL
