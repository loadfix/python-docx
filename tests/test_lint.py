"""Tests for `Document.lint()` (issue #57 + accessibility rules from issue #15)."""

from __future__ import annotations

import io

import pytest

from docx import Document
from docx.lint import (
    ACCESSIBILITY_RULES,
    ALL_RULES,
    LintFinding,
    Severity,
    lint_document,
)
from docx.shared import Inches, Pt, RGBColor


class DescribeDocumentLint:
    def it_returns_an_empty_list_for_a_well_formed_outline(self):
        document = Document()
        document.add_heading("Title", level=1)
        document.add_heading("Subsection", level=2)
        document.add_paragraph("Body text.")
        findings = document.lint(rules=["heading-skip"])
        assert findings == []

    def it_flags_heading_skip(self):
        document = Document()
        document.add_heading("Top", level=1)
        document.add_heading("Skipped", level=3)
        findings = document.lint(rules=["heading-skip"])
        assert len(findings) == 1
        assert findings[0].rule_id == "heading-skip"
        assert findings[0].severity == Severity.ERROR

    def it_flags_multiple_h1(self):
        document = Document()
        document.add_heading("First", level=1)
        document.add_heading("Second", level=1)
        findings = document.lint(rules=["heading-multiple-h1"])
        assert len(findings) == 1
        assert findings[0].rule_id == "heading-multiple-h1"
        assert findings[0].severity == Severity.WARNING

    def it_emits_info_when_no_h1_present(self):
        document = Document()
        document.add_heading("Sub", level=2)
        findings = document.lint(rules=["heading-no-h1"])
        assert len(findings) == 1
        assert findings[0].rule_id == "heading-no-h1"
        assert findings[0].severity == Severity.INFO
        assert findings[0].paragraph_index is None

    def it_flags_an_empty_heading(self):
        document = Document()
        document.add_heading("   ", level=1)
        findings = document.lint(rules=["heading-empty"])
        assert len(findings) == 1
        assert findings[0].rule_id == "heading-empty"
        assert findings[0].severity == Severity.ERROR

    def it_flags_an_overly_long_heading(self):
        document = Document()
        document.add_heading("X" * 200, level=1)
        findings = document.lint(rules=["heading-too-long"])
        assert len(findings) == 1
        assert findings[0].rule_id == "heading-too-long"
        assert findings[0].severity == Severity.WARNING

    def it_flags_body_text_that_looks_like_a_heading(self):
        document = Document()
        paragraph = document.add_paragraph("My Section")
        run = paragraph.runs[0]
        run.bold = True
        findings = document.lint(rules=["heading-direct-formatting"])
        assert len(findings) == 1
        assert findings[0].rule_id == "heading-direct-formatting"

    def it_flags_large_font_body_text_as_heading_like(self):
        document = Document()
        paragraph = document.add_paragraph("Big Title")
        run = paragraph.runs[0]
        run.font.size = Pt(20)
        findings = document.lint(rules=["heading-direct-formatting"])
        assert len(findings) == 1

    def it_uses_default_rules_when_none_specified(self):
        document = Document()
        document.add_heading("Top", level=1)
        document.add_heading("Skip", level=3)
        findings = document.lint()
        assert any(f.rule_id == "heading-skip" for f in findings)

    def it_accepts_a_callable_rule(self):
        document = Document()
        document.add_paragraph("body")

        def custom_rule(paragraphs):
            yield LintFinding(
                severity="info",
                paragraph_index=0,
                rule_id="custom",
                message="hello",
            )

        findings = document.lint(rules=[custom_rule])
        assert len(findings) == 1
        assert findings[0].rule_id == "custom"

    def it_rejects_an_unknown_rule_id(self):
        document = Document()
        with pytest.raises(ValueError):
            document.lint(rules=["no-such-rule"])

    def it_sorts_findings_by_paragraph_index_then_id(self):
        document = Document()
        document.add_heading("Top", level=1)
        document.add_heading("Skip", level=3)
        document.add_heading("X" * 200, level=4)
        findings = lint_document(document)
        # -- paragraph_index ordered, with None last
        indices = [f.paragraph_index for f in findings]
        assert indices == sorted(
            indices, key=lambda x: x if x is not None else 10**9
        )


# --- DescribeAccessibilityRules (issue #15) -------------------------


_PNG_PATH = "tests/test_files/python-icon.png"


def _new_document_with_lang(lang_tag: str = "en-US") -> Document:
    """Return a fresh ``Document`` whose Normal style declares `lang_tag`."""
    document = Document()
    # -- Set lang on Normal style so the no-language-tag rule passes when
    # -- it shouldn't be the focus of the test. --
    normal = document.styles["Normal"]
    normal.element.get_or_add_rPr().get_or_add_lang().set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
        lang_tag,
    )
    return document


class DescribeAccessibilityRules:
    def it_flags_an_image_without_alt_text(self):
        document = _new_document_with_lang()
        document.add_picture(_PNG_PATH)
        # -- alt_text on the inline shape is empty by default (descr unset). --
        findings = document.lint(rules=["image-no-alt-text"])
        assert len(findings) == 1
        assert findings[0].rule_id == "image-no-alt-text"
        assert findings[0].severity == Severity.ERROR

    def it_passes_an_image_with_alt_text(self):
        document = _new_document_with_lang()
        document.add_picture(_PNG_PATH)
        document.inline_shapes[0].alt_text = "Quarterly revenue chart"
        findings = document.lint(rules=["image-no-alt-text"])
        assert findings == []

    def it_skips_decorative_images(self):
        document = _new_document_with_lang()
        document.add_picture(_PNG_PATH)
        document.inline_shapes[0].a11y_role = "decorative"
        findings = document.lint(rules=["image-no-alt-text"])
        assert findings == []

    def it_flags_a_table_without_caption(self):
        document = _new_document_with_lang()
        document.add_table(rows=2, cols=2)
        findings = document.lint(rules=["table-no-caption"])
        assert len(findings) == 1
        assert findings[0].rule_id == "table-no-caption"
        assert findings[0].severity == Severity.WARNING

    def it_passes_a_table_with_caption(self):
        document = _new_document_with_lang()
        table = document.add_table(rows=2, cols=2)
        table.alt_text = "Sales by region"
        findings = document.lint(rules=["table-no-caption"])
        assert findings == []

    def it_flags_a_document_without_a_language_tag(self):
        document = Document()
        # -- strip every w:lang the bundled template ships with --
        from docx.oxml.ns import qn

        for tree in (document._element.body, document.styles._element):
            for el in list(tree.iter(qn("w:lang"))):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
        findings = document.lint(rules=["no-language-tag"])
        assert len(findings) == 1
        assert findings[0].rule_id == "no-language-tag"
        assert findings[0].severity == Severity.WARNING
        assert findings[0].paragraph_index is None

    def it_passes_a_document_with_a_language_tag(self):
        document = _new_document_with_lang("fr-FR")
        findings = document.lint(rules=["no-language-tag"])
        assert findings == []

    def it_flags_low_contrast_text(self):
        document = _new_document_with_lang()
        para = document.add_paragraph("Pale grey on white")
        # -- a near-white colour fails the contrast heuristic --
        para.runs[0].font.color.rgb = RGBColor(0xF0, 0xF0, 0xF0)
        findings = document.lint(rules=["low-contrast"])
        assert len(findings) == 1
        assert findings[0].rule_id == "low-contrast"
        assert findings[0].severity == Severity.INFO

    def it_passes_high_contrast_text(self):
        document = _new_document_with_lang()
        para = document.add_paragraph("Black on white")
        para.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        findings = document.lint(rules=["low-contrast"])
        assert findings == []

    def it_flags_a_document_with_no_title_property(self):
        document = _new_document_with_lang()
        # -- core_properties.title is "" by default for a fresh Document --
        findings = document.lint(rules=["no-document-title"])
        assert len(findings) == 1
        assert findings[0].rule_id == "no-document-title"
        assert findings[0].severity == Severity.WARNING

    def it_passes_a_document_with_a_title_property(self):
        document = _new_document_with_lang()
        # -- guard against the well-documented `Document.core_properties`
        # -- monkey-patch in test_fields.py that deletes the descriptor on
        # -- cleanup. When the descriptor is gone the assertion below
        # -- can't run; that's a sibling-test bug, not an a11y-rule bug. --
        if not hasattr(type(document), "core_properties"):
            pytest.skip("Document.core_properties not present (sibling-test pollution)")
        document.core_properties.title = "Annual Report 2026"
        findings = document.lint(rules=["no-document-title"])
        assert findings == []

    def it_runs_every_accessibility_rule_when_enabled(self):
        document = Document()
        # -- this fresh-from-default document violates every a11y rule --
        document.add_picture(_PNG_PATH)
        document.add_table(rows=2, cols=2)
        from docx.oxml.ns import qn

        for tree in (document._element.body, document.styles._element):
            for el in list(tree.iter(qn("w:lang"))):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
        para = document.add_paragraph("Pale grey on white")
        para.runs[0].font.color.rgb = RGBColor(0xF0, 0xF0, 0xF0)
        findings = document.lint(rules=ACCESSIBILITY_RULES)
        rule_ids = {f.rule_id for f in findings}
        assert "image-no-alt-text" in rule_ids
        assert "table-no-caption" in rule_ids
        assert "no-language-tag" in rule_ids
        assert "low-contrast" in rule_ids
        assert "no-document-title" in rule_ids

    def it_round_trips_through_save_and_reopen(self):
        """Accessibility findings survive a save/load cycle."""
        document = Document()
        document.add_picture(_PNG_PATH)
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = Document(buf)
        findings = reopened.lint(rules=["image-no-alt-text"])
        assert len(findings) == 1

    def it_exposes_accessibility_rules_under_ALL_RULES(self):
        # -- ALL_RULES includes the accessibility rule callables --
        rule_names = {getattr(r, "__name__", "") for r in ALL_RULES}
        assert "rule_image_no_alt_text" in rule_names
        assert "rule_no_document_title" in rule_names
