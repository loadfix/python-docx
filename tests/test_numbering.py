# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.numbering` module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.enum.text import WD_NUMBER_FORMAT
from docx.numbering import (
    AbstractNumberingDefinition,
    Level,
    LevelOverride,
    ListLabelRenderer,
    NumInstance,
    Numbering,
    NumberingDefinition,
)
from docx.oxml.numbering import CT_Numbering
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches

from .unitutil.cxml import element
from .unitutil.mock import instance_mock
from docx.parts.numbering import NumberingPart


class DescribeNumbering:
    """Unit-test suite for `docx.numbering.Numbering`."""

    def it_is_empty_for_a_fresh_numbering_element(self, request: pytest.FixtureRequest):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        part = instance_mock(request, NumberingPart)

        numbering = Numbering(numbering_elm, part)

        assert len(numbering) == 0
        assert list(numbering) == []
        assert numbering.definitions == []

    def it_can_add_a_numbering_definition_from_mapping_specs(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))

        defn = numbering.add_numbering_definition(
            levels=[
                {
                    "format": WD_NUMBER_FORMAT.DECIMAL,
                    "text": "%1.",
                    "indent": Inches(0.5),
                },
                {
                    "format": WD_NUMBER_FORMAT.LOWER_LETTER,
                    "text": "%2)",
                    "indent": Inches(1.0),
                },
            ]
        )

        assert isinstance(defn, NumberingDefinition)
        assert defn.abstract_num_id == 0
        assert len(defn.levels) == 2
        # -- a matching w:num instance was created so the definition is usable --
        assert len(numbering_elm.num_lst) == 1
        assert numbering_elm.num_lst[0].abstractNumId.val == 0

        level_0 = defn.levels[0]
        assert level_0.ilvl == 0
        assert level_0.number_format == WD_NUMBER_FORMAT.DECIMAL
        assert level_0.text == "%1."

        level_1 = defn.levels[1]
        assert level_1.ilvl == 1
        assert level_1.number_format == WD_NUMBER_FORMAT.LOWER_LETTER
        assert level_1.text == "%2)"

    def it_accepts_positional_tuple_level_specs(self, request: pytest.FixtureRequest):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))

        defn = numbering.add_numbering_definition(
            levels=[(WD_NUMBER_FORMAT.UPPER_ROMAN, "%1.")]
        )

        assert defn.levels[0].number_format == WD_NUMBER_FORMAT.UPPER_ROMAN

    def it_accepts_raw_string_format_values(self, request: pytest.FixtureRequest):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))

        defn = numbering.add_numbering_definition(
            levels=[{"format": "bullet", "text": "•", "font": "Symbol"}]
        )

        lvl = defn.levels[0]
        assert lvl.number_format == WD_NUMBER_FORMAT.BULLET
        assert lvl.text == "•"

    def it_supports_multiple_definitions(self, request: pytest.FixtureRequest):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))

        defn_a = numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )
        defn_b = numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.UPPER_LETTER, "text": "%1."}]
        )

        assert defn_a.abstract_num_id == 0
        assert defn_b.abstract_num_id == 1
        assert len(numbering) == 2
        # -- each definition has its own w:num instance --
        assert len(numbering_elm.num_lst) == 2


class DescribeNumberingDefinition:
    """Unit-test suite for `docx.numbering.NumberingDefinition`."""

    def it_exposes_levels_in_order(self, request: pytest.FixtureRequest):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_numbering_definition(
            levels=[
                {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%2."},
                {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%3."},
            ]
        )

        levels = defn.levels

        assert [lvl.ilvl for lvl in levels] == [0, 1, 2]
        assert all(isinstance(lvl, Level) for lvl in levels)

    def it_can_return_a_level_by_ilvl(self, request: pytest.FixtureRequest):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_numbering_definition(
            levels=[
                {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                {"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%2)"},
            ]
        )

        lvl = defn.level(1)

        assert lvl is not None
        assert lvl.ilvl == 1
        assert lvl.number_format == WD_NUMBER_FORMAT.LOWER_LETTER

    def it_returns_None_for_missing_level(self, request: pytest.FixtureRequest):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )

        assert defn.level(5) is None

    def it_creates_a_new_instance_sharing_the_abstract_definition(
        self, request: pytest.FixtureRequest
    ):
        """Closes upstream#25 — `NumberingDefinition.new_instance()`."""
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )

        # -- baseline: one w:num created alongside the abstract definition --
        assert len(numbering_elm.num_lst) == 1
        first_num_id = numbering_elm.num_lst[0].numId

        new_num_id = defn.new_instance()

        assert new_num_id != first_num_id
        assert len(numbering_elm.num_lst) == 2
        # -- the new w:num references the same abstractNumId --
        new_num = numbering_elm.num_having_numId(new_num_id)
        assert new_num.abstractNumId.val == defn.abstract_num_id


class DescribeLevel:
    """Unit-test suite for `docx.numbering.Level`."""

    def it_reports_level_properties_from_the_underlying_lvl(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_numbering_definition(
            levels=[
                {
                    "format": WD_NUMBER_FORMAT.DECIMAL,
                    "text": "%1.",
                    "indent": Inches(0.75),
                }
            ]
        )

        lvl = defn.levels[0]

        assert lvl.ilvl == 0
        assert lvl.number_format == WD_NUMBER_FORMAT.DECIMAL
        assert lvl.text == "%1."
        assert lvl.start == 1
        assert lvl.indent is not None
        assert lvl.indent.inches == pytest.approx(0.75, rel=1e-3)

    def it_returns_None_number_format_for_unknown_value(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(
            CT_Numbering,
            element(
                "w:numbering/w:abstractNum{w:abstractNumId=0}/"
                "w:lvl{w:ilvl=0}/w:numFmt{w:val=chicagoManual}"
            ),
        )
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.definitions[0]
        lvl = defn.level(0)

        assert lvl is not None
        assert lvl.number_format is None


def _make_p(num_id: int, ilvl: int) -> CT_P:
    """Return a freshly built ``w:p`` with the requested numPr values."""
    return cast(
        CT_P,
        element(
            f"w:p/w:pPr/w:numPr/(w:ilvl{{w:val={ilvl}}},w:numId{{w:val={num_id}}})"
        ),
    )


def _make_bare_p() -> CT_P:
    return cast(CT_P, element("w:p"))


def _mk_numbering(
    request: pytest.FixtureRequest, levels_per_def: list[list[dict]]
) -> tuple[CT_Numbering, list[int]]:
    """Build a ``CT_Numbering`` populated with the given definitions.

    Returns the element and a list of numId values, one per definition.
    """
    numbering_elm = cast(CT_Numbering, element("w:numbering"))
    numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
    num_ids: list[int] = []
    for levels in levels_per_def:
        numbering.add_numbering_definition(levels)
        num_ids.append(numbering_elm.num_lst[-1].numId)
    return numbering_elm, num_ids


class DescribeListLabelRenderer:
    """Unit-test suite for `docx.numbering.ListLabelRenderer`."""

    def it_returns_None_for_a_paragraph_without_numPr(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, _ = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]]
        )
        renderer = ListLabelRenderer(numbering_elm)

        assert renderer.label_for(_make_bare_p()) is None

    def it_renders_a_decimal_list(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]]
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(3)]

        labels = [renderer.label_for(p) for p in ps]

        assert labels == ["1.", "2.", "3."]

    def it_renders_decimalZero_with_leading_zero(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]]
        )
        # -- swap in the raw "decimalZero" token that WD_NUMBER_FORMAT doesn't expose --
        lvl = numbering_elm.abstractNum_lst[0].get_lvl(0)
        assert lvl is not None and lvl.numFmt is not None
        lvl.numFmt.set(qn("w:val"), "decimalZero")
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(11)]

        labels = [renderer.label_for(p) for p in ps]
        assert labels[:3] == ["01.", "02.", "03."]
        assert labels[9] == "10."
        assert labels[10] == "11."

    def it_renders_upper_roman(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.UPPER_ROMAN, "text": "%1."}]]
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(4)]

        labels = [renderer.label_for(p) for p in ps]
        assert labels == ["I.", "II.", "III.", "IV."]

    def it_renders_lower_roman(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.LOWER_ROMAN, "text": "%1."}]]
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(4)]

        labels = [renderer.label_for(p) for p in ps]
        assert labels == ["i.", "ii.", "iii.", "iv."]

    def it_renders_upper_letter(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.UPPER_LETTER, "text": "%1)"}]]
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(3)]

        labels = [renderer.label_for(p) for p in ps]
        assert labels == ["A)", "B)", "C)"]

    def it_wraps_upper_letters_past_z(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.UPPER_LETTER, "text": "%1"}]]
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(27)]

        # -- advance 27 times: first 26 are A..Z, 27th is AA --
        labels = [renderer.label_for(p) for p in ps]
        assert labels[0] == "A"
        assert labels[25] == "Z"
        assert labels[26] == "AA"

    def it_renders_lower_letter(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%1)"}]]
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(3)]

        labels = [renderer.label_for(p) for p in ps]
        assert labels == ["a)", "b)", "c)"]

    def it_renders_bullet_text_verbatim(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request,
            [[{"format": WD_NUMBER_FORMAT.BULLET, "text": "•", "font": "Symbol"}]],
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(2)]

        labels = [renderer.label_for(p) for p in ps]
        assert labels == ["•", "•"]

    def it_resets_deeper_counters_on_return_to_shallower_level(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, (num_id,) = _mk_numbering(
            request,
            [
                [
                    {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                    {"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%2)"},
                ]
            ],
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [
            _make_p(num_id, 0),
            _make_p(num_id, 1),
            _make_p(num_id, 1),
            _make_p(num_id, 0),
            _make_p(num_id, 1),
        ]

        labels = [renderer.label_for(p) for p in ps]

        # -- returning to level 0 keeps counter at 2, level 1 restarts --
        assert labels == ["1.", "a)", "b)", "2.", "a)"]

    def it_renders_multi_level_lvlText_patterns(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, (num_id,) = _mk_numbering(
            request,
            [
                [
                    {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                    {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1.%2."},
                ]
            ],
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [
            _make_p(num_id, 0),
            _make_p(num_id, 1),
            _make_p(num_id, 1),
            _make_p(num_id, 0),
            _make_p(num_id, 1),
        ]

        labels = [renderer.label_for(p) for p in ps]

        assert labels == ["1.", "1.1.", "1.2.", "2.", "2.1."]

    def it_handles_skip_level_transitions(self, request: pytest.FixtureRequest):
        numbering_elm, (num_id,) = _mk_numbering(
            request,
            [
                [
                    {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                    {"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%2)"},
                    {"format": WD_NUMBER_FORMAT.LOWER_ROMAN, "text": "%3."},
                ]
            ],
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [
            _make_p(num_id, 0),
            _make_p(num_id, 2),  # skip level 1
            _make_p(num_id, 2),
            _make_p(num_id, 1),  # now walk back up
        ]

        labels = [renderer.label_for(p) for p in ps]

        assert labels == ["1.", "i.", "ii.", "a)"]

    def it_starts_counters_at_the_levels_start_value(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, (num_id,) = _mk_numbering(
            request,
            [[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1.", "start": 5}]],
        )
        renderer = ListLabelRenderer(numbering_elm)
        ps = [_make_p(num_id, 0) for _ in range(2)]

        labels = [renderer.label_for(p) for p in ps]
        assert labels == ["5.", "6."]

    def it_returns_None_when_numbering_element_is_absent(
        self, request: pytest.FixtureRequest
    ):
        renderer = ListLabelRenderer(None)

        assert renderer.label_for(_make_p(1, 0)) is None

    def it_returns_None_for_an_unresolvable_numId(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, _ = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]]
        )
        renderer = ListLabelRenderer(numbering_elm)

        # -- numId 999 has no matching w:num --
        assert renderer.label_for(_make_p(999, 0)) is None

    def it_caches_labels_for_repeat_lookups_of_the_same_element(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]]
        )
        renderer = ListLabelRenderer(numbering_elm)
        p = _make_p(num_id, 0)

        first = renderer.label_for(p)
        second = renderer.label_for(p)
        # -- cached: does not advance the counter --
        assert first == "1." and second == "1."

    def it_resolves_numId_from_an_inherited_paragraph_style(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, (num_id,) = _mk_numbering(
            request, [[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]]
        )

        # -- build a styles element declaring a style that carries w:numPr --
        styles_elm = element(
            "w:styles/w:style{w:type=paragraph,w:styleId=ListBullet}/"
            f"w:pPr/w:numPr/(w:ilvl{{w:val=0}},w:numId{{w:val={num_id}}})"
        )
        renderer = ListLabelRenderer(cast(CT_Numbering, numbering_elm), styles_elm)

        # -- paragraph declares only pStyle, no direct numPr --
        p = cast(
            CT_P,
            element("w:p/w:pPr/w:pStyle{w:val=ListBullet}"),
        )

        assert renderer.label_for(p) == "1."

    def it_can_build_a_label_map_over_many_paragraphs(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm, (num_id,) = _mk_numbering(
            request,
            [
                [
                    {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                    {"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%2)"},
                ]
            ],
        )
        renderer = ListLabelRenderer(numbering_elm)

        ps = [
            _make_p(num_id, 0),
            _make_p(num_id, 1),
            _make_p(num_id, 1),
            _make_bare_p(),  # not in a list
            _make_p(num_id, 0),
        ]

        labels = renderer.label_map(iter(ps))

        assert labels[id(ps[0])] == "1."
        assert labels[id(ps[1])] == "a)"
        assert labels[id(ps[2])] == "b)"
        assert id(ps[3]) not in labels
        assert labels[id(ps[4])] == "2."


class DescribeNumberingAuthoring:
    """Covers the single-level / id-allocator / instance helpers."""

    def it_exposes_AbstractNumberingDefinition_as_an_alias(self):
        assert AbstractNumberingDefinition is NumberingDefinition

    def it_reports_next_num_id_and_next_abstract_num_id(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))

        assert numbering.next_num_id() == 1
        assert numbering.next_abstract_num_id() == 0

        numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )

        # -- after one definition + auto-created instance, next ids advance --
        assert numbering.next_abstract_num_id() == 1
        assert numbering.next_num_id() == 2

    def it_can_add_a_single_level_abstract_definition(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))

        defn = numbering.add_abstract_definition(
            format=WD_NUMBER_FORMAT.DECIMAL, start=3, lvl_text="%1)"
        )

        assert isinstance(defn, NumberingDefinition)
        assert defn.abstract_num_id == 0
        lvl = defn.level(0)
        assert lvl is not None
        assert lvl.number_format == WD_NUMBER_FORMAT.DECIMAL
        assert lvl.start == 3
        assert lvl.text == "%1)"
        # -- add_abstract_definition does *not* auto-create a w:num --
        assert len(numbering_elm.num_lst) == 0

    def it_allocates_a_num_instance_for_an_abstract_definition(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_abstract_definition()

        inst = numbering.add_definition(defn.abstract_num_id)

        assert isinstance(inst, NumInstance)
        assert inst.abstract_num_id == defn.abstract_num_id
        assert inst.num_id == 1
        # -- a second allocation yields a distinct numId --
        assert numbering.add_definition(defn.abstract_num_id).num_id == 2

    def it_raises_when_add_definition_refers_to_an_unknown_abstract(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))

        with pytest.raises(KeyError):
            numbering.add_definition(42)

    def it_attaches_pStyle_to_level0_when_add_definition_gets_a_style_name(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_abstract_definition()

        numbering.add_definition(defn.abstract_num_id, style_name="ListBullet")

        lvl = defn.level(0)
        assert lvl is not None
        assert lvl.element.pStyle_val == "ListBullet"

    def it_exposes_num_instances_list_and_lookup(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_abstract_definition()
        numbering.add_definition(defn.abstract_num_id)
        numbering.add_definition(defn.abstract_num_id)

        instances = numbering.num_instances
        assert len(instances) == 2
        assert all(isinstance(i, NumInstance) for i in instances)
        assert numbering.num_instance(instances[0].num_id) is not None
        assert numbering.num_instance(9999) is None

    def it_can_set_a_level_override_start(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_abstract_definition()
        inst = numbering.add_definition(defn.abstract_num_id)

        override = inst.set_level_override(0, start=7)

        assert isinstance(override, LevelOverride)
        assert override.ilvl == 0
        assert override.start_override == 7
        # -- repeat call stomps, doesn't accumulate --
        inst.set_level_override(0, start=12)
        again = inst.level_override(0)
        assert again is not None
        assert again.start_override == 12

    def it_rejects_out_of_range_ilvl_on_set_level_override(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_abstract_definition()
        inst = numbering.add_definition(defn.abstract_num_id)

        with pytest.raises(ValueError):
            inst.set_level_override(9, start=1)


class DescribeMultiLevelAuthoringRoundTrip:
    """End-to-end: build a list via Document.numbering, save + reload."""

    def it_round_trips_a_multi_level_list_through_save_and_reload(self, tmp_path):
        from io import BytesIO
        from docx import Document

        doc = Document()
        numbering = doc.numbering
        defn = numbering.add_numbering_definition(
            levels=[
                {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                {"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%2)"},
                {"format": WD_NUMBER_FORMAT.LOWER_ROMAN, "text": "%3."},
            ]
        )

        p1 = doc.add_paragraph("First")
        p2 = doc.add_paragraph("Sub")
        p3 = doc.add_paragraph("Sub-sub")
        defn.apply_to(p1, level=0)
        defn.apply_to(p2, level=1)
        defn.apply_to(p3, level=2)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        defs = reloaded.numbering.definitions
        assert len(defs) >= 1
        # -- locate our definition by its level count --
        ours = next(d for d in defs if len(d.levels) == 3)
        assert ours.level(0).number_format == WD_NUMBER_FORMAT.DECIMAL
        assert ours.level(1).number_format == WD_NUMBER_FORMAT.LOWER_LETTER
        assert ours.level(2).number_format == WD_NUMBER_FORMAT.LOWER_ROMAN

        labels = reloaded.list_labels()
        # -- three labelled paragraphs, one per level --
        rendered = list(labels.values())
        assert "1." in rendered
        assert "a)" in rendered
        assert "i." in rendered

    def it_restarts_numbering_when_style_changes_to_a_list_backed_style(
        self, tmp_path
    ):
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        doc = Document()
        numbering = doc.numbering
        defn = numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )
        # -- bind the single-level list to a paragraph style `MyList` --
        style = doc.styles.add_style("MyList", 1)
        pPr = style.element.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
        numPr.append(ilvl)
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(numbering.num_instances[-1].num_id))
        numPr.append(numId_el)
        pPr.append(numPr)

        p = doc.add_paragraph("item")
        instance_count_before = len(doc.numbering.num_instances)

        p.style = "MyList"

        assert len(doc.numbering.num_instances) == instance_count_before + 1
        # -- paragraph now carries the fresh numId, not the style's original --
        assert p._p.pPr.numPr.numId_val == doc.numbering.num_instances[-1].num_id
        # -- level-0 on the new numPr --
        assert p._p.pPr.numPr.ilvl_val == 0

    def it_lets_callers_modify_a_level_after_creation(
        self, request: pytest.FixtureRequest
    ):
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering = Numbering(numbering_elm, instance_mock(request, NumberingPart))
        defn = numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )

        lvl0 = defn.level(0)
        assert lvl0 is not None

        # -- flip the format mid-document --
        lvl0.element.numFmt_val = WD_NUMBER_FORMAT.UPPER_ROMAN
        lvl0.element.lvlText_val = "%1)"
        lvl0.element.start_val = 5

        reread = defn.level(0)
        assert reread is not None
        assert reread.number_format == WD_NUMBER_FORMAT.UPPER_ROMAN
        assert reread.text == "%1)"
        assert reread.start == 5
