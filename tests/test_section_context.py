"""Tests for `Document.section(...)` context manager (issue #79)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.exceptions import NestedSectionError
from docx.shared import Inches


class DescribeDocumentSectionContext:
    def it_yields_a_section_for_landscape_orientation(self):
        document = Document()
        with document.section(orientation="landscape") as section:
            assert section.orientation == WD_ORIENTATION.LANDSCAPE
            document.add_paragraph("inside landscape")

    def it_swaps_page_dimensions_when_flipping_to_landscape(self):
        document = Document()
        prior = document.sections[-1]
        prior.page_width = Inches(8.5)
        prior.page_height = Inches(11)
        with document.section(orientation="landscape") as section:
            assert section.page_width > section.page_height

    def it_applies_named_margin_presets(self):
        document = Document()
        with document.section(margins="narrow") as section:
            assert section.top_margin == Inches(0.5)
            assert section.left_margin == Inches(0.5)

    def it_applies_explicit_margin_dicts(self):
        document = Document()
        with document.section(margins={"top": Inches(2), "left": 1.5}) as section:
            assert section.top_margin == Inches(2)
            assert section.left_margin == Inches(1.5)

    def it_applies_a_named_page_size(self):
        document = Document()
        with document.section(page_size="a4") as section:
            # -- A4 is 210mm x 297mm; tolerate ±0.01" rounding from
            # -- the inches-based preset.
            assert abs(section.page_width - Inches(8.27)) < Inches(0.01)
            assert abs(section.page_height - Inches(11.69)) < Inches(0.01)

    def it_creates_a_new_section_per_with_block(self):
        document = Document()
        before = len(document.sections)
        with document.section(orientation="landscape"):
            pass
        # -- inner section + reverting sentinel ⇒ +2 sections
        assert len(document.sections) == before + 2

    def it_reverts_orientation_after_exit(self):
        document = Document()
        prior_orientation = document.sections[-1].orientation
        with document.section(orientation="landscape"):
            pass
        assert document.sections[-1].orientation == prior_orientation

    def it_raises_NestedSectionError_when_nested(self):
        document = Document()
        with document.section(orientation="landscape"):
            with pytest.raises(NestedSectionError):
                with document.section(orientation="portrait"):
                    pass

    def it_releases_the_lock_after_an_exception_in_the_block(self):
        document = Document()
        with pytest.raises(RuntimeError):
            with document.section(orientation="landscape"):
                raise RuntimeError("boom")
        # -- the lock should be released; opening another section works
        with document.section(orientation="portrait"):
            pass

    def it_applies_columns_as_int(self):
        document = Document()
        with document.section(columns=3) as section:
            assert section.columns.count == 3

    def it_applies_columns_as_dict(self):
        document = Document()
        with document.section(columns={"count": 2, "space": Inches(0.25)}) as section:
            assert section.columns.count == 2

    def it_applies_a_header_string(self):
        document = Document()
        with document.section(header="Top of page") as section:
            assert section.header.paragraphs[0].text == "Top of page"

    def it_rejects_an_unknown_margin_preset(self):
        document = Document()
        with pytest.raises(ValueError):
            with document.section(margins="extra-narrow"):
                pass

    def it_rejects_an_unknown_orientation_string(self):
        document = Document()
        with pytest.raises(ValueError):
            with document.section(orientation="diagonal"):
                pass
