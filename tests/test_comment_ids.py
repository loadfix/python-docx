# pyright: reportPrivateUsage=false

"""Unit tests for R13-2 — docx adoption of ``commentsIds.xml`` +
``commentsExtensible.xml`` via ``ooxml-comments`` 0.4.

Covers the two new ``Comment`` accessors (:attr:`paragraph_id` /
:attr:`durable_id`), the lazy :attr:`Document.comments_ids` /
:attr:`Document.comments_extensible` proxies, and the auto-minting of
both identifier families when :meth:`Document.add_comment` /
:meth:`Comment.reply` creates a new comment. All happy-path tests
round-trip through :class:`io.BytesIO` to make sure the new parts reach
the produced package and parse back correctly.
"""

from __future__ import annotations

import io
import re

import pytest

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.parts.comments_ids import CommentsExtensiblePart, CommentsIdsPart


# -- fixtures ---------------------------------------------------------------


@pytest.fixture
def doc_with_one_comment() -> Document:
    """A minimal document with a single comment and its auto-minted ids."""
    doc = Document()
    p = doc.add_paragraph("hello ")
    doc.add_comment(runs=[p.add_run("world")], text="nit", author="A")
    return doc


# -- Comment.paragraph_id --------------------------------------------------


class DescribeCommentParagraphId:
    """Unit-test suite for ``Comment.paragraph_id`` getter/setter."""

    def it_auto_mints_a_paragraph_id_on_add_comment(self, doc_with_one_comment: Document):
        c = list(doc_with_one_comment.comments)[0]

        # -- 8-char uppercase hex token, matching Word's shape --
        assert re.fullmatch(r"[0-9A-F]{8}", c.paragraph_id)

    def it_mirrors_the_paragraph_id_into_commentsIds_xml(
        self, doc_with_one_comment: Document
    ):
        c = list(doc_with_one_comment.comments)[0]

        entries = list(doc_with_one_comment.comments_ids)

        assert entries == [(c.comment_id, c.paragraph_id)]

    def it_round_trips_paragraph_id_through_save_and_load(
        self, doc_with_one_comment: Document
    ):
        buf = io.BytesIO()
        doc_with_one_comment.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        original = list(doc_with_one_comment.comments)[0]
        loaded = list(reloaded.comments)[0]

        assert loaded.paragraph_id == original.paragraph_id
        assert loaded.paragraph_id != ""

    def it_allows_overwriting_the_paragraph_id(self, doc_with_one_comment: Document):
        c = list(doc_with_one_comment.comments)[0]

        c.paragraph_id = "DEADBEEF"

        assert c.paragraph_id == "DEADBEEF"
        assert (c.comment_id, "DEADBEEF") in list(doc_with_one_comment.comments_ids)

    def it_returns_empty_string_when_no_ids_registry_exists(self):
        # -- the legacy fixture has no commentsIds part; a read should be
        # -- safe and not materialise the part out from under the caller.
        from docx.oxml.comments import CT_Comment, CT_Comments
        from docx.opc.packuri import PackURI
        from docx.comments import Comment
        from docx.parts.comments import CommentsPart
        from tests.unitutil.cxml import element
        from tests.unitutil.mock import instance_mock
        from docx.package import Package
        from unittest.mock import Mock

        package = Mock(spec=Package)
        comments_elm = element("w:comments/w:comment{w:id=7}")
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"), CT.WML_COMMENTS, comments_elm, package
        )
        comment_elm = comments_elm.comment_lst[0]
        comment = Comment(comment_elm, comments_part)

        # -- no ids part is related on a synthetic fixture; read returns "" --
        assert comment.paragraph_id == ""


# -- Comment.durable_id -----------------------------------------------------


class DescribeCommentDurableId:
    """Unit-test suite for ``Comment.durable_id`` getter/setter."""

    def it_auto_mints_a_braced_guid_on_add_comment(self, doc_with_one_comment: Document):
        c = list(doc_with_one_comment.comments)[0]

        # -- ``{XXXXXXXX-XXXX-XXXX-XXXX-XXXXXXXXXXXX}`` uppercase hex. --
        assert re.fullmatch(
            r"\{[0-9A-F]{8}-[0-9A-F]{4}-[0-9A-F]{4}-[0-9A-F]{4}-[0-9A-F]{12}\}",
            c.durable_id,
        )

    def it_mirrors_the_durable_id_into_commentsExtensible_xml(
        self, doc_with_one_comment: Document
    ):
        c = list(doc_with_one_comment.comments)[0]

        entries = list(doc_with_one_comment.comments_extensible)

        # -- one entry whose durableId matches the comment's. --
        assert len(entries) == 1
        assert entries[0] == (c.durable_id, c.durable_id)

    def it_round_trips_durable_id_through_save_and_load(
        self, doc_with_one_comment: Document
    ):
        buf = io.BytesIO()
        doc_with_one_comment.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        original = list(doc_with_one_comment.comments)[0]
        loaded = list(reloaded.comments)[0]

        assert loaded.durable_id == original.durable_id
        assert loaded.durable_id != ""


# -- per-reply auto-mint ----------------------------------------------------


class DescribeAutoMintOnReply:
    """Replies also get fresh ``paragraph_id`` + ``durable_id`` entries."""

    def it_mints_fresh_ids_for_each_reply(self, doc_with_one_comment: Document):
        parent = list(doc_with_one_comment.comments)[0]

        reply = parent.reply(text="r", author="B")

        assert reply.paragraph_id != "" and reply.paragraph_id != parent.paragraph_id
        assert reply.durable_id != "" and reply.durable_id != parent.durable_id

    def it_round_trips_both_parent_and_reply_ids(self, doc_with_one_comment: Document):
        parent = list(doc_with_one_comment.comments)[0]
        reply = parent.reply(text="r", author="B")

        buf = io.BytesIO()
        doc_with_one_comment.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        by_id = {c.comment_id: c for c in reloaded.comments}
        assert by_id[parent.comment_id].paragraph_id == parent.paragraph_id
        assert by_id[parent.comment_id].durable_id == parent.durable_id
        assert by_id[reply.comment_id].paragraph_id == reply.paragraph_id
        assert by_id[reply.comment_id].durable_id == reply.durable_id


# -- Document.comments_ids / Document.comments_extensible --------------------


class DescribeDocumentCommentIdsProxy:
    """Unit-test suite for the lazy ``Document.comments_ids`` accessor."""

    def it_materialises_the_commentsIds_part_on_first_access(self):
        doc = Document()

        # -- no comment has been added yet; accessing the proxy should --
        # -- still create the part so callers can author the registry ahead
        # -- of time.
        _ = doc.comments_ids

        # pyright: ignore[reportPrivateUsage]
        assert doc.part._comments_part.comments_ids_part is not None

    def it_returns_a_CommentIds_proxy(self):
        from ooxml_comments import CommentIds

        doc = Document()

        assert isinstance(doc.comments_ids, CommentIds)

    def it_returns_a_CommentsExtensible_proxy(self):
        from ooxml_comments import CommentsExtensible

        doc = Document()

        assert isinstance(doc.comments_extensible, CommentsExtensible)

    def it_relates_the_new_parts_from_comments_xml(self, doc_with_one_comment: Document):
        # pyright: ignore[reportPrivateUsage]
        comments_part = doc_with_one_comment.part._comments_part
        ids_part = comments_part.comments_ids_part
        ex_part = comments_part.comments_extensible_part

        assert isinstance(ids_part, CommentsIdsPart)
        assert isinstance(ex_part, CommentsExtensiblePart)
        assert ids_part.partname == "/word/commentsIds.xml"
        assert ex_part.partname == "/word/commentsExtensible.xml"
        assert ids_part.content_type == CT.WML_COMMENTS_IDS
        assert ex_part.content_type == CT.WML_COMMENTS_EXTENSIBLE

    def it_round_trips_the_parts_through_save_and_load(
        self, doc_with_one_comment: Document
    ):
        buf = io.BytesIO()
        doc_with_one_comment.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        # pyright: ignore[reportPrivateUsage]
        comments_part = reloaded.part._comments_part
        assert comments_part.comments_ids_part is not None
        assert comments_part.comments_extensible_part is not None


# -- content-type and relationship constants --------------------------------


class DescribeConstants:
    """The docx CT / RT namespaces pick up the shared 0.4 strings."""

    def it_exposes_WML_COMMENTS_IDS_content_type(self):
        assert CT.WML_COMMENTS_IDS == "application/vnd.ms-word.commentsIds+xml"

    def it_exposes_WML_COMMENTS_EXTENSIBLE_content_type(self):
        assert (
            CT.WML_COMMENTS_EXTENSIBLE
            == "application/vnd.ms-word.commentsExtensible+xml"
        )

    def it_exposes_COMMENTS_IDS_relationship_type(self):
        assert RT.COMMENTS_IDS == (
            "http://schemas.microsoft.com/office/2016/09/relationships/commentsIds"
        )

    def it_exposes_COMMENTS_EXTENSIBLE_relationship_type(self):
        assert RT.COMMENTS_EXTENSIBLE == (
            "http://schemas.microsoft.com/office/2018/08/relationships/commentsExtensible"
        )
