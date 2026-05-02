# pyright: reportPrivateUsage=false

"""Unit-test suite for `docx.attachments` + `Document.attachments`."""

from __future__ import annotations

from typing import cast

from docx.attachments import Attachment
from docx.document import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.document import CT_Document
from docx.parts.document import DocumentPart

from .unitutil.cxml import element
from .unitutil.mock import FixtureRequest, instance_mock


def _html_part(blob: bytes = b"<html><body>hi</body></html>") -> Part:
    part = Part(PackURI("/word/afchunk.html"), "text/html", blob)
    return part


class DescribeAttachment:
    def it_knows_its_r_id(self):
        alt = element("w:altChunk{r:id=rId7}")
        att = Attachment(alt, None)
        assert att.r_id == "rId7"

    def it_returns_None_r_id_when_missing(self):
        alt = element("w:altChunk")
        att = Attachment(alt, None)
        assert att.r_id is None

    def it_exposes_blob_and_content_type_when_resolved(self):
        alt = element("w:altChunk{r:id=rId1}")
        part = _html_part()
        att = Attachment(alt, part)
        assert att.blob == b"<html><body>hi</body></html>"
        assert att.content_type == "text/html"
        assert att.partname == "/word/afchunk.html"

    def it_returns_empty_blob_when_unresolved(self):
        alt = element("w:altChunk{r:id=rId99}")
        att = Attachment(alt, None)
        assert att.blob == b""
        assert att.content_type is None
        assert att.partname is None


class DescribeDocument_attachments:
    def it_returns_empty_list_when_no_altChunks(self, request: FixtureRequest):
        doc_elm = cast(CT_Document, element("w:document/w:body/w:p"))
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}
        document = Document(doc_elm, document_part_)

        assert document.attachments == []

    def it_enumerates_each_altChunk(self, request: FixtureRequest):
        part_html = _html_part(b"<html>A</html>")
        part_rtf = Part(PackURI("/word/afchunk.rtf"), "application/rtf", b"{\\rtf1}")
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {
            "rId10": part_html,
            "rId11": part_rtf,
        }
        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                "w:altChunk{r:id=rId10},"
                "w:p,"
                "w:altChunk{r:id=rId11}"
                ")"
            ),
        )
        document = Document(doc_elm, document_part_)

        atts = document.attachments
        assert len(atts) == 2
        assert [a.r_id for a in atts] == ["rId10", "rId11"]
        assert [a.content_type for a in atts] == ["text/html", "application/rtf"]
        assert [a.blob for a in atts] == [b"<html>A</html>", b"{\\rtf1}"]

    def it_handles_unresolved_altChunk(self, request: FixtureRequest):
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}
        doc_elm = cast(
            CT_Document,
            element("w:document/w:body/w:altChunk{r:id=rIdMissing}"),
        )
        document = Document(doc_elm, document_part_)
        atts = document.attachments
        assert len(atts) == 1
        assert atts[0].blob == b""
        assert atts[0].partname is None
        assert atts[0].content_type is None
