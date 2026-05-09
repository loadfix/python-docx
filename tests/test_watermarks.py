# pyright: reportPrivateUsage=false

"""Unit-test suite for R10 Document-level watermark and page-background APIs."""

from __future__ import annotations

from io import BytesIO

import pytest

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from .unitutil.file import test_file


class DescribeDocumentPageBackgroundColor:
    """Document.page_background_color getter/setter."""

    def it_returns_None_when_no_background_is_set(self):
        document = Document()

        assert document.page_background_color is None

    def it_round_trips_a_hex_string(self):
        document = Document()

        document.page_background_color = "4472C4"

        assert document.page_background_color == "4472C4"

    def it_accepts_a_leading_hash(self):
        document = Document()

        document.page_background_color = "#FF8800"

        assert document.page_background_color == "FF8800"

    def it_clears_the_background_when_set_to_None(self):
        document = Document()
        document.page_background_color = "ABCDEF"

        document.page_background_color = None

        assert document.page_background_color is None
        assert document._element.background is None

    def it_raises_on_an_invalid_length(self):
        document = Document()

        with pytest.raises(ValueError, match="6-char hex"):
            document.page_background_color = "ZZZ"

    def it_raises_on_a_non_string(self):
        document = Document()

        with pytest.raises(TypeError, match="hex string"):
            document.page_background_color = 0x4472C4  # type: ignore[assignment]

    def it_also_reflects_changes_via_background_color(self):
        document = Document()

        document.background_color = RGBColor(0x11, 0x22, 0x33)

        assert document.page_background_color == "112233"


class DescribeDocumentAddTextWatermark:
    """Document.add_text_watermark across all sections."""

    def it_adds_a_watermark_with_defaults(self):
        document = Document()

        wm = document.add_text_watermark("CONFIDENTIAL")

        assert wm.type == "text"
        assert wm.text == "CONFIDENTIAL"
        assert len(document.watermarks) == 1
        shape = document.sections[0].header._element.find(".//" + qn("v:shape"))
        assert shape is not None
        fill = shape.find(qn("v:fill"))
        assert fill is not None
        assert fill.get("color") == "#808080"
        # -- diagonal=True is the default --
        assert "rotation:-45" in (shape.get("style") or "")
        # -- default font/size baked in --
        textpath = shape.find(qn("v:textpath"))
        assert textpath is not None
        assert "Calibri" in textpath.get("style", "")
        assert "36" in textpath.get("style", "")

    def it_honors_custom_options(self):
        document = Document()

        wm = document.add_text_watermark(
            "DRAFT",
            font_name="Arial",
            font_size=48,
            color_rgb="#FF0000",
            diagonal=False,
        )

        assert wm.text == "DRAFT"
        shape = document.sections[0].header._element.find(".//" + qn("v:shape"))
        assert shape is not None
        fill = shape.find(qn("v:fill"))
        assert fill is not None
        assert fill.get("color") == "#FF0000"
        assert "rotation:0" in (shape.get("style") or "")
        textpath = shape.find(qn("v:textpath"))
        assert textpath is not None
        assert "Arial" in textpath.get("style", "")
        assert "48" in textpath.get("style", "")

    def it_rejects_a_bad_color_string(self):
        document = Document()

        with pytest.raises(ValueError, match="6-char hex"):
            document.add_text_watermark("X", color_rgb="red")

    def it_round_trips_through_a_save(self):
        document = Document()
        document.add_text_watermark("SECRET", font_size=24, color_rgb="00FF00")

        buf = BytesIO()
        document.save(buf)
        buf.seek(0)
        loaded = Document(buf)

        wms = loaded.watermarks
        assert len(wms) == 1
        assert wms[0].type == "text"
        assert wms[0].text == "SECRET"


class DescribeDocumentAddPictureWatermark:
    """Document.add_picture_watermark across all sections."""

    def it_embeds_a_picture_from_a_path(self):
        document = Document()

        wm = document.add_picture_watermark(test_file("monty-truth.png"))

        assert wm.type == "image"
        assert len(document.watermarks) == 1
        shape = document.sections[0].header._element.find(".//" + qn("v:shape"))
        assert shape is not None
        imagedata = shape.find(qn("v:imagedata"))
        assert imagedata is not None
        rId = imagedata.get(qn("r:id"))
        assert rId is not None and rId.startswith("rId")

    def it_embeds_a_picture_from_a_BytesIO(self):
        with open(test_file("monty-truth.png"), "rb") as f:
            png = f.read()
        document = Document()

        wm = document.add_picture_watermark(BytesIO(png), scale=0.75)

        assert wm.type == "image"

    def it_rejects_a_non_positive_scale(self):
        document = Document()

        with pytest.raises(ValueError, match="scale"):
            document.add_picture_watermark(test_file("monty-truth.png"), scale=0)

    def it_round_trips_a_BytesIO_through_a_save(self):
        with open(test_file("monty-truth.png"), "rb") as f:
            png = f.read()
        document = Document()
        document.add_picture_watermark(BytesIO(png), scale=0.5)

        buf = BytesIO()
        document.save(buf)
        buf.seek(0)
        loaded = Document(buf)

        wms = loaded.watermarks
        assert len(wms) == 1
        assert wms[0].type == "image"


class DescribeDocumentWatermarks:
    """Document.watermarks property and Watermark.remove()."""

    def it_returns_an_empty_list_when_no_watermarks_present(self):
        document = Document()

        assert document.watermarks == []

    def it_removes_a_watermark_from_every_section(self):
        document = Document()
        wm = document.add_text_watermark("REMOVE_ME")
        assert len(document.watermarks) == 1

        wm.remove()

        assert document.watermarks == []
        shape = document.sections[0].header._element.find(".//" + qn("v:shape"))
        assert shape is None
