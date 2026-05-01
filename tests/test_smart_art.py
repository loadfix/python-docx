# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.smart_art` module and related plumbing."""

from __future__ import annotations

from typing import cast

from docx.document import Document
from docx.drawing import Drawing
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.oxml.document import CT_Document
from docx.oxml.drawing import CT_Drawing
from docx.oxml.parser import parse_xml
from docx.oxml.smart_art import CT_DataModel
from docx.parts.document import DocumentPart
from docx.parts.smart_art import DiagramDataPart
from docx.smart_art import SmartArt, SmartArtNode, smart_art_for_drawing

from .unitutil.cxml import element
from .unitutil.mock import FixtureRequest, instance_mock


# -- minimal data-model XML used across several tests --

SIMPLE_DATA_XML = (
    b'<?xml version="1.0"?>\n'
    b'<dgm:dataModel'
    b' xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"'
    b' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">\n'
    b"  <dgm:ptLst>\n"
    b'    <dgm:pt modelId="n1" type="node">\n'
    b'      <dgm:t><a:p><a:r><a:t>Alpha</a:t></a:r></a:p></dgm:t>\n'
    b"    </dgm:pt>\n"
    b'    <dgm:pt modelId="n2" type="node">\n'
    b'      <dgm:t><a:p><a:r><a:t>Beta</a:t></a:r></a:p></dgm:t>\n'
    b"    </dgm:pt>\n"
    b'    <dgm:pt modelId="n3" type="node">\n'
    b'      <dgm:t><a:p><a:r><a:t>Gamma</a:t></a:r></a:p></dgm:t>\n'
    b"    </dgm:pt>\n"
    b"  </dgm:ptLst>\n"
    b"</dgm:dataModel>\n"
)

# -- hierarchical data: n1 is root; n2 and n3 are its children --

HIERARCHY_DATA_XML = (
    b'<?xml version="1.0"?>\n'
    b'<dgm:dataModel'
    b' xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"'
    b' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">\n'
    b"  <dgm:ptLst>\n"
    b'    <dgm:pt modelId="n1" type="node">\n'
    b'      <dgm:t><a:p><a:r><a:t>Root</a:t></a:r></a:p></dgm:t>\n'
    b"    </dgm:pt>\n"
    b'    <dgm:pt modelId="n2" type="node">\n'
    b'      <dgm:t><a:p><a:r><a:t>Child1</a:t></a:r></a:p></dgm:t>\n'
    b"    </dgm:pt>\n"
    b'    <dgm:pt modelId="n3" type="node">\n'
    b'      <dgm:t><a:p><a:r><a:t>Child2</a:t></a:r></a:p></dgm:t>\n'
    b"    </dgm:pt>\n"
    b"  </dgm:ptLst>\n"
    b"  <dgm:cxnLst>\n"
    b'    <dgm:cxn type="parOf" srcId="n1" destId="n2"/>\n'
    b'    <dgm:cxn type="parOf" srcId="n1" destId="n3"/>\n'
    b"  </dgm:cxnLst>\n"
    b"</dgm:dataModel>\n"
)

# -- data containing a presentation node (type="pres") which must be skipped --

MIXED_DATA_XML = (
    b'<?xml version="1.0"?>\n'
    b'<dgm:dataModel'
    b' xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"'
    b' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">\n'
    b"  <dgm:ptLst>\n"
    b'    <dgm:pt modelId="n1" type="node">\n'
    b'      <dgm:t><a:p><a:r><a:t>Only</a:t></a:r></a:p></dgm:t>\n'
    b"    </dgm:pt>\n"
    b'    <dgm:pt modelId="p1" type="pres"/>\n'
    b"  </dgm:ptLst>\n"
    b"</dgm:dataModel>\n"
)


def _make_data_part(
    blob: bytes = SIMPLE_DATA_XML, idx: int = 1
) -> DiagramDataPart:
    element_ = cast(CT_DataModel, parse_xml(blob))
    return DiagramDataPart(
        PackURI("/word/diagrams/data%d.xml" % idx),
        CT.DML_DIAGRAM_DATA,
        element_,
        cast(object, None),  # pyright: ignore[reportArgumentType]
    )


def _smart_art_drawing_cxml(rId: str = "rId4") -> str:
    return (
        "w:drawing/wp:inline/a:graphic/a:graphicData"
        "/dgm:relIds{r:dm=%s,r:lo=rId5,r:qs=rId6,r:cs=rId7}" % rId
    )


class DescribeSmartArtNode:
    """Unit-test suite for `docx.smart_art.SmartArtNode`."""

    def it_exposes_its_text_and_level(self):
        data = cast(CT_DataModel, parse_xml(SIMPLE_DATA_XML))
        pt = data.pt_lst[0]

        node = SmartArtNode(pt, level=2)

        assert node.text == "Alpha"
        assert node.level == 2
        assert node.model_id == "n1"
        assert node.children == []

    def it_stores_given_children(self):
        data = cast(CT_DataModel, parse_xml(SIMPLE_DATA_XML))
        parent_pt = data.pt_lst[0]
        child = SmartArtNode(data.pt_lst[1], level=1)

        node = SmartArtNode(parent_pt, level=0, children=[child])

        assert node.children == [child]
        assert node.children[0].text == "Beta"


class DescribeSmartArt:
    """Unit-test suite for `docx.smart_art.SmartArt` proxy."""

    def it_reports_None_data_partname_when_data_part_is_absent(self):
        relIds = cast("object", element("dgm:relIds{r:dm=rId4}"))

        sa = SmartArt(relIds, None)  # pyright: ignore[reportArgumentType]

        assert sa.data_partname is None
        assert sa.nodes == []
        assert sa.text == ""
        assert sa.dm_rId == "rId4"

    def it_exposes_the_data_partname_when_resolved(self):
        relIds = cast("object", element("dgm:relIds{r:dm=rId4}"))
        data_part = _make_data_part()

        sa = SmartArt(relIds, data_part)  # pyright: ignore[reportArgumentType]

        assert sa.data_partname == "/word/diagrams/data1.xml"

    def it_returns_a_flat_node_list_when_no_connections(self):
        relIds = cast("object", element("dgm:relIds{r:dm=rId4}"))
        data_part = _make_data_part(SIMPLE_DATA_XML)

        sa = SmartArt(relIds, data_part)  # pyright: ignore[reportArgumentType]
        nodes = sa.nodes

        assert [n.text for n in nodes] == ["Alpha", "Beta", "Gamma"]
        assert all(n.level == 0 for n in nodes)
        assert all(n.children == [] for n in nodes)

    def it_reconstructs_hierarchy_from_connections(self):
        relIds = cast("object", element("dgm:relIds{r:dm=rId4}"))
        data_part = _make_data_part(HIERARCHY_DATA_XML)

        sa = SmartArt(relIds, data_part)  # pyright: ignore[reportArgumentType]
        nodes = sa.nodes

        assert len(nodes) == 1
        root = nodes[0]
        assert root.text == "Root"
        assert root.level == 0
        assert [c.text for c in root.children] == ["Child1", "Child2"]
        assert all(c.level == 1 for c in root.children)

    def its_text_concatenates_every_node_with_indent(self):
        relIds = cast("object", element("dgm:relIds{r:dm=rId4}"))
        data_part = _make_data_part(HIERARCHY_DATA_XML)

        sa = SmartArt(relIds, data_part)  # pyright: ignore[reportArgumentType]

        assert sa.text == "Root\n  Child1\n  Child2"

    def it_skips_presentation_nodes(self):
        relIds = cast("object", element("dgm:relIds{r:dm=rId4}"))
        data_part = _make_data_part(MIXED_DATA_XML)

        sa = SmartArt(relIds, data_part)  # pyright: ignore[reportArgumentType]
        nodes = sa.nodes

        assert len(nodes) == 1
        assert nodes[0].text == "Only"


class DescribeSmartArtForDrawing:
    """Unit-test suite for `docx.smart_art.smart_art_for_drawing`."""

    def it_returns_None_when_drawing_has_no_relIds(self, request: FixtureRequest):
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}

        drawing = cast(
            CT_Drawing,
            element("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic"),
        )

        assert smart_art_for_drawing(drawing, document_part_) is None

    def it_returns_SmartArt_even_when_data_part_is_missing(
        self, request: FixtureRequest
    ):
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}

        drawing = cast(CT_Drawing, element(_smart_art_drawing_cxml()))

        sa = smart_art_for_drawing(drawing, document_part_)

        assert isinstance(sa, SmartArt)
        assert sa.data_partname is None
        assert sa.nodes == []

    def it_returns_SmartArt_wired_to_the_resolved_data_part(
        self, request: FixtureRequest
    ):
        data_part = _make_data_part(HIERARCHY_DATA_XML, idx=2)
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {"rId4": data_part}

        drawing = cast(CT_Drawing, element(_smart_art_drawing_cxml()))

        sa = smart_art_for_drawing(drawing, document_part_)

        assert isinstance(sa, SmartArt)
        assert sa.data_partname == "/word/diagrams/data2.xml"
        assert len(sa.nodes) == 1
        assert sa.nodes[0].text == "Root"

    def it_ignores_related_parts_of_the_wrong_type(self, request: FixtureRequest):
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {"rId4": "not-a-diagram-part"}

        drawing = cast(CT_Drawing, element(_smart_art_drawing_cxml()))

        sa = smart_art_for_drawing(drawing, document_part_)

        assert sa is not None
        assert sa.data_partname is None
        assert sa.nodes == []


class DescribeDrawing_smart_art:
    """Unit-test suite for `Drawing.is_smart_art` and `Drawing.smart_art`."""

    def it_knows_it_is_not_smart_art_for_a_plain_picture(
        self, request: FixtureRequest
    ):
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}

        class _Parent:
            @property
            def part(self):
                return document_part_

        drawing_elm = cast(
            CT_Drawing,
            element("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic"),
        )
        drawing = Drawing(drawing_elm, _Parent())  # pyright: ignore[reportArgumentType]

        assert drawing.is_smart_art is False
        assert drawing.smart_art is None

    def it_detects_smart_art_even_when_data_part_unresolvable(
        self, request: FixtureRequest
    ):
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}

        class _Parent:
            @property
            def part(self):
                return document_part_

        drawing_elm = cast(CT_Drawing, element(_smart_art_drawing_cxml()))
        drawing = Drawing(drawing_elm, _Parent())  # pyright: ignore[reportArgumentType]

        assert drawing.is_smart_art is True
        sa = drawing.smart_art
        assert sa is not None
        assert sa.nodes == []

    def it_returns_a_SmartArt_with_nodes_when_data_part_resolves(
        self, request: FixtureRequest
    ):
        data_part = _make_data_part(SIMPLE_DATA_XML)
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {"rId4": data_part}

        class _Parent:
            @property
            def part(self):
                return document_part_

        drawing_elm = cast(CT_Drawing, element(_smart_art_drawing_cxml()))
        drawing = Drawing(drawing_elm, _Parent())  # pyright: ignore[reportArgumentType]

        sa = drawing.smart_art

        assert sa is not None
        assert [n.text for n in sa.nodes] == ["Alpha", "Beta", "Gamma"]


class DescribeDocument_smart_art:
    """Unit-test suite for `Document.smart_art`."""

    def it_is_empty_for_a_document_without_smart_art(
        self, request: FixtureRequest
    ):
        doc_elm = cast(CT_Document, element("w:document/w:body/w:p"))
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}
        document = Document(doc_elm, document_part_)

        assert document.smart_art == []

    def it_collects_every_smart_art_in_the_body(self, request: FixtureRequest):
        data_part_a = _make_data_part(SIMPLE_DATA_XML, idx=1)
        data_part_b = _make_data_part(HIERARCHY_DATA_XML, idx=2)
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {
            "rId4": data_part_a,
            "rId8": data_part_b,
        }

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                "w:p/w:r/" + _smart_art_drawing_cxml("rId4") + ","
                "w:p/w:r/w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic,"
                "w:p/w:r/" + _smart_art_drawing_cxml("rId8") + ")"
            ),
        )
        document = Document(doc_elm, document_part_)

        smart_arts = document.smart_art

        assert len(smart_arts) == 2
        assert smart_arts[0].data_partname == "/word/diagrams/data1.xml"
        assert smart_arts[1].data_partname == "/word/diagrams/data2.xml"
        # -- content from the first --
        assert [n.text for n in smart_arts[0].nodes] == [
            "Alpha",
            "Beta",
            "Gamma",
        ]
        # -- content from the second (hierarchical) --
        assert smart_arts[1].nodes[0].text == "Root"

    def it_still_enumerates_smart_art_when_data_part_unresolved(
        self, request: FixtureRequest
    ):
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/w:r/" + _smart_art_drawing_cxml()
            ),
        )
        document = Document(doc_elm, document_part_)

        smart_arts = document.smart_art

        assert len(smart_arts) == 1
        assert smart_arts[0].data_partname is None
        assert smart_arts[0].nodes == []


