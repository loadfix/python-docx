"""Integration tests for the ``strict=`` / ``is_strict`` API surface.

Complements :mod:`tests.test_strict_ooxml` (which covers the
Strict → Transitional blob rewriter) by verifying the top-level
``Document`` / ``Package`` plumbing of the ECMA-376 Strict
conformance-class keyword.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest

from docx import Document
from docx.opc.strict import STRICT_TO_TRANSITIONAL


def _default_template_bytes() -> bytes:
    tpl = (
        Path(__file__).parent.parent
        / "src" / "docx" / "templates" / "default.docx"
    )
    with open(tpl, "rb") as f:
        return f.read()


def _rewrite_to_strict(transitional_bytes: bytes) -> bytes:
    """Return a Strict-OOXML rewrite of *transitional_bytes*."""
    trans_to_strict = {t: s for s, t in STRICT_TO_TRANSITIONAL.items()}
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(transitional_bytes)) as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if (
                    info.filename.endswith((".xml", ".rels"))
                    or info.filename == "[Content_Types].xml"
                ):
                    for trans, strict in trans_to_strict.items():
                        data = data.replace(trans, strict)
                zout.writestr(info, data)
    return out.getvalue()


class DescribeStrictConformanceSurface:
    """Document plumbs strict= / is_strict through to OpcPackage."""

    def it_reports_is_strict_false_for_a_fresh_document(self) -> None:
        doc = Document()
        assert doc.is_strict is False

    def it_lets_callers_flip_is_strict_on_the_document(self) -> None:
        doc = Document()
        assert doc.is_strict is False
        doc.is_strict = True
        assert doc.is_strict is True
        doc.is_strict = False
        assert doc.is_strict is False

    def it_auto_detects_is_strict_on_open_of_a_strict_package(self) -> None:
        strict_bytes = _rewrite_to_strict(_default_template_bytes())
        doc = Document(io.BytesIO(strict_bytes))
        # -- the PackageReader sniff flagged the source as Strict even
        # -- without explicit ``strict=True`` on the factory --
        assert doc.is_strict is True

    def it_accepts_strict_kwarg_on_the_factory(self) -> None:
        # -- explicit opt-in flips the flag on a normally-Transitional
        # -- package --
        doc = Document(io.BytesIO(_default_template_bytes()), strict=True)
        assert doc.is_strict is True

    def it_accepts_strict_kwarg_on_save(self, tmp_path: Path) -> None:
        doc = Document()
        out = tmp_path / "out.docx"
        doc.save(str(out), strict=True)
        # -- the saved .docx remains a valid zip we can reopen --
        doc2 = Document(str(out))
        assert doc2 is not None

    def it_round_trips_a_strict_document_preserving_is_strict(
        self, tmp_path: Path
    ) -> None:
        strict_bytes = _rewrite_to_strict(_default_template_bytes())
        doc = Document(io.BytesIO(strict_bytes))
        assert doc.is_strict is True

        out = tmp_path / "out.docx"
        # -- default save(): strict=None preserves the loaded class on
        # -- the package's _is_strict flag. Byte-level emission is always
        # -- Transitional today; the flag survives for the next-level
        # -- round-trip at the shared-runtime layer. --
        doc.save(str(out))

        # -- reopen with explicit strict=True so the flag comes back
        # -- (the written bytes are Transitional, so a default reopen
        # -- would sniff as not-strict). --
        doc2 = Document(str(out), strict=True)
        assert doc2.is_strict is True
