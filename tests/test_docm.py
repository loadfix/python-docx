"""Integration test suite for .docm (macro-enabled document) support."""

from __future__ import annotations

import os
import tempfile
import zipfile

import pytest

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT


TESTS_DIR = os.path.dirname(__file__)
DOCM_PATH = os.path.join(TESTS_DIR, "test_files", "macros.docm")


class DescribeDocmSupport:
    """Integration tests for opening and saving .docm files."""

    def it_can_open_a_docm_file(self):
        document = Document(DOCM_PATH)
        assert document is not None
        assert document.has_macros is True

    def it_reports_has_macros_False_for_a_regular_docx(self):
        document = Document()
        assert document.has_macros is False

    def it_preserves_vba_project_on_round_trip(self):
        document = Document(DOCM_PATH)

        with tempfile.NamedTemporaryFile(suffix=".docm", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            document.save(tmp_path)

            # verify the saved file contains vbaProject.bin
            with zipfile.ZipFile(tmp_path, "r") as zf:
                assert "word/vbaProject.bin" in zf.namelist()
                blob = zf.read("word/vbaProject.bin")
                assert blob == b"FAKE_VBA_PROJECT_BINARY_DATA_FOR_TESTING"

            # verify the saved file can be re-opened and still has macros
            document2 = Document(tmp_path)
            assert document2.has_macros is True
        finally:
            os.unlink(tmp_path)

    def it_preserves_the_macro_content_type_on_round_trip(self):
        document = Document(DOCM_PATH)

        with tempfile.NamedTemporaryFile(suffix=".docm", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            document.save(tmp_path)

            with zipfile.ZipFile(tmp_path, "r") as zf:
                content_types_xml = zf.read("[Content_Types].xml").decode("utf-8")
                assert CT.WML_DOCUMENT_MACRO in content_types_xml
                assert CT.WML_VBA_PROJECT in content_types_xml
        finally:
            os.unlink(tmp_path)

    def it_can_read_paragraphs_from_a_docm(self):
        document = Document(DOCM_PATH)
        assert len(document.paragraphs) >= 1
        assert document.paragraphs[0].text == "Document with macros"
