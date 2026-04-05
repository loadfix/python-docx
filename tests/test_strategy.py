# pyright: reportPrivateUsage=false
# pyright: reportUnknownMemberType=false

"""Multi-layered testing strategy exercising all five validation layers.

This module establishes the testing patterns described in issue #43 and provides
at least one example test for each layer, using the comments feature as the
reference implementation.

Layer 1: XML Structure Tests — validates python-docx produces correct OOXML elements
Layer 2: OOXML Schema Validation — validates output against XSD schemas
Layer 3: Round-Trip Tests — write/save/reopen/assert pattern
Layer 4: Reference File Comparison — validates reading of Word-created .docx files
Layer 5: LibreOffice Headless Validation — optional CI validation via conversion
"""

from __future__ import annotations

import os
import tempfile
from typing import cast

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.oxml.ns import qn

from tests.helpers.libreoffice import is_libreoffice_available, validate_with_libreoffice
from tests.helpers.refcmp import compare_xml_structure, ref_docx_exists, ref_docx_path
from tests.helpers.roundtrip import assert_round_trip, save_and_reopen
from tests.helpers.schema import (
    SchemaValidationResult,
    load_bundled_schema,
    validate_docx_xml_parts,
    validate_part_xml,
)
from tests.helpers.validate import (
    validate_content_type_present,
    validate_elements_present,
    validate_ooxml_structure,
    validate_relationship_present,
)
from tests.helpers.xmlparse import parse_docx_xml


# =====================================================================================
# Layer 1: XML Structure Tests
# =====================================================================================


class DescribeLayer1_XMLStructure:
    """Layer 1: Validate that python-docx produces correct OOXML elements."""

    def it_produces_a_comments_part_with_correct_elements(self):
        doc = Document()
        doc.add_paragraph("Test paragraph.")
        run = doc.paragraphs[0].runs[0]
        doc.add_comment(run, text="A test comment.", author="Test Author", initials="TA")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)

            # -- word/comments.xml contains w:comment elements --
            comments_xml = parse_docx_xml(path, "word/comments.xml")
            assert comments_xml is not None, "word/comments.xml should exist"
            comment_elms = comments_xml.findall(qn("w:comment"))
            assert len(comment_elms) >= 1, "should have at least one w:comment element"

            # -- comment has required attributes --
            comment = comment_elms[0]
            assert comment.get(qn("w:id")) is not None, "w:id attribute required"
            assert comment.get(qn("w:author")) == "Test Author"
            assert comment.get(qn("w:initials")) == "TA"
            assert comment.get(qn("w:date")) is not None, "w:date attribute expected"

            # -- document.xml contains comment range markers --
            doc_xml = parse_docx_xml(path, "word/document.xml")
            assert doc_xml is not None
            range_starts = doc_xml.iter(qn("w:commentRangeStart"))
            range_ends = doc_xml.iter(qn("w:commentRangeEnd"))
            assert len(list(range_starts)) >= 1, "should have commentRangeStart marker"
            assert len(list(range_ends)) >= 1, "should have commentRangeEnd marker"
        finally:
            os.unlink(path)

    def it_registers_comments_content_type(self):
        doc = Document()
        doc.add_paragraph("Test.")
        doc.add_comment(doc.paragraphs[0].runs[0], text="Comment")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)

            ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
            assert validate_content_type_present(path, ct), (
                f"Content type '{ct}' should be in [Content_Types].xml"
            )
        finally:
            os.unlink(path)

    def it_registers_comments_relationship(self):
        doc = Document()
        doc.add_paragraph("Test.")
        doc.add_comment(doc.paragraphs[0].runs[0], text="Comment")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)

            rel_type = (
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
            )
            assert validate_relationship_present(path, rel_type), (
                f"Relationship type '{rel_type}' should be in document.xml.rels"
            )
        finally:
            os.unlink(path)


# =====================================================================================
# Layer 2: OOXML Schema Validation
# =====================================================================================


class DescribeLayer2_SchemaValidation:
    """Layer 2: Validate output XML against OOXML schemas."""

    def it_produces_structurally_valid_docx_files(self):
        doc = Document()
        doc.add_paragraph("Hello, World!")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            errors = validate_ooxml_structure(path)
            assert errors == [], f"Structural validation errors: {errors}"
        finally:
            os.unlink(path)

    def it_produces_structurally_valid_docx_with_comments(self):
        doc = Document()
        doc.add_paragraph("Test paragraph.")
        doc.add_comment(doc.paragraphs[0].runs[0], text="Comment", author="Author")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            errors = validate_ooxml_structure(path)
            assert errors == [], f"Structural validation errors: {errors}"
        finally:
            os.unlink(path)

    def it_produces_well_formed_xml_in_all_parts(self):
        doc = Document()
        doc.add_paragraph("Test.")
        doc.add_comment(doc.paragraphs[0].runs[0], text="Comment", author="A")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            results = validate_docx_xml_parts(path)
            for part_name, result in results.items():
                assert result.is_valid, (
                    f"XML part '{part_name}' is malformed: {result.errors}"
                )
        finally:
            os.unlink(path)

    def it_validates_comments_xml_against_schema(self):
        schema = load_bundled_schema("wml-comments")
        if schema is None:
            pytest.skip("Comments schema not available")

        doc = Document()
        doc.add_paragraph("Test.")
        doc.add_comment(doc.paragraphs[0].runs[0], text="A comment", author="Author")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)

            import zipfile

            with zipfile.ZipFile(path) as zf:
                comments_bytes = zf.read("word/comments.xml")

            result = validate_part_xml(comments_bytes, schema)
            assert result.is_valid, f"Schema validation errors: {result.errors}"
        finally:
            os.unlink(path)


# =====================================================================================
# Layer 3: Round-Trip Tests
# =====================================================================================


class DescribeLayer3_RoundTrip:
    """Layer 3: Write → Save → Reopen → Assert pattern."""

    def it_round_trips_a_simple_comment(self):
        def create(doc: DocumentCls) -> dict[str, str]:
            doc.add_paragraph("Annotated text.")
            run = doc.paragraphs[0].runs[0]
            doc.add_comment(run, text="My comment", author="Jane Doe", initials="JD")
            return {"text": "My comment", "author": "Jane Doe", "initials": "JD"}

        def check(doc: DocumentCls, ctx: dict[str, str]) -> None:
            comments = doc.comments
            assert len(comments) >= 1
            comment = comments.get(0)
            assert comment is not None
            assert comment.text == ctx["text"]
            assert comment.author == ctx["author"]
            assert comment.initials == ctx["initials"]

        assert_round_trip(create, check)

    def it_round_trips_threaded_comments(self):
        def create(doc: DocumentCls) -> dict[str, str]:
            doc.add_paragraph("Threaded comment test.")
            run = doc.paragraphs[0].runs[0]
            comment = doc.add_comment(run, text="Parent comment", author="Author A")
            comment.add_reply(text="Reply 1", author="Author B")
            comment.add_reply(text="Reply 2", author="Author C")
            return {"parent": "Parent comment", "reply1": "Reply 1", "reply2": "Reply 2"}

        def check(doc: DocumentCls, ctx: dict[str, str]) -> None:
            comments = doc.comments
            assert len(comments) >= 3
            parent = comments.get(0)
            assert parent is not None
            assert parent.text == ctx["parent"]
            replies = parent.replies
            assert len(replies) == 2
            assert replies[0].text == ctx["reply1"]
            assert replies[1].text == ctx["reply2"]

        assert_round_trip(create, check)

    def it_round_trips_comment_on_specific_text_range(self):
        def create(doc: DocumentCls) -> str:
            para = doc.add_paragraph()
            run1 = para.add_run("Before ")
            run2 = para.add_run("target text")
            run3 = para.add_run(" after")
            doc.add_comment(run2, text="Comment on target", author="Tester")
            return "target text"

        def check(doc: DocumentCls, target_text: str) -> None:
            comments = doc.comments
            assert len(comments) >= 1
            comment = comments.get(0)
            assert comment is not None
            assert comment.text == "Comment on target"
            assert comment.author == "Tester"

        assert_round_trip(create, check)

    def it_round_trips_multiple_comments_by_different_authors(self):
        doc = Document()
        para = doc.add_paragraph("Multiple authors.")
        run = para.runs[0]
        doc.add_comment(run, text="Comment 1", author="Alice", initials="A")
        doc.add_comment(run, text="Comment 2", author="Bob", initials="B")
        doc.add_comment(run, text="Comment 3", author="Carol", initials="C")

        doc2 = save_and_reopen(doc)

        comments = doc2.comments
        assert len(comments) == 3
        authors = {c.author for c in comments}
        assert authors == {"Alice", "Bob", "Carol"}

    def it_round_trips_a_comment_with_multiline_text(self):
        doc = Document()
        doc.add_paragraph("Multi-line comment test.")
        run = doc.paragraphs[0].runs[0]
        doc.add_comment(run, text="Line 1\nLine 2\nLine 3", author="Author")

        doc2 = save_and_reopen(doc)

        comment = doc2.comments.get(0)
        assert comment is not None
        assert len(comment.paragraphs) == 3
        assert "Line 1" in comment.paragraphs[0].text
        assert comment.paragraphs[1].text == "Line 2"
        assert comment.paragraphs[2].text == "Line 3"


# =====================================================================================
# Layer 4: Reference File Comparison
# =====================================================================================


class DescribeLayer4_ReferenceComparison:
    """Layer 4: Read reference .docx files created in Microsoft Word."""

    def it_reads_existing_comments_fixture(self):
        # -- Use the existing acceptance test fixture that has comments --
        fixture_path = os.path.join(
            os.path.dirname(__file__),
            "..",
            "features",
            "steps",
            "test_files",
            "comments-rich-para.docx",
        )
        if not os.path.exists(fixture_path):
            pytest.skip("comments-rich-para.docx fixture not available")

        doc = Document(fixture_path)
        comments = doc.comments
        assert len(comments) > 0, "Reference file should contain comments"

        first_comment = next(iter(comments))
        assert first_comment.author != "", "Comment should have an author"
        assert first_comment.comment_id is not None

    def it_can_compare_xml_structure_of_generated_vs_reference(self):
        """Pattern test: demonstrates how to compare generated output against a reference."""
        # -- Generate a doc --
        doc = Document()
        doc.add_paragraph("Comparison test.")

        fd, gen_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        fd2, ref_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd2)

        try:
            doc.save(gen_path)
            # -- Use the same doc as "reference" for this pattern demonstration --
            doc.save(ref_path)

            # -- Compare word/document.xml structure --
            diffs = compare_xml_structure(
                gen_path,
                ref_path,
                "word/document.xml",
                ignore_attrs={qn("w:id")},
            )
            assert diffs == [], f"Structural differences found: {diffs}"
        finally:
            os.unlink(gen_path)
            os.unlink(ref_path)

    def it_reads_a_reference_comments_doc_when_available(self):
        if not ref_docx_exists("comments-simple"):
            pytest.skip(
                "Reference file 'comments-simple.docx' not yet created. "
                "See tests/ref-docs/README.md for instructions."
            )

        doc = Document(ref_docx_path("comments-simple"))
        comments = doc.comments
        assert len(comments) >= 1


# =====================================================================================
# Layer 5: LibreOffice Headless Validation
# =====================================================================================


class DescribeLayer5_LibreOfficeValidation:
    """Layer 5: Validate .docx files by converting to PDF with LibreOffice."""

    @pytest.mark.libreoffice
    def it_converts_a_basic_document_to_pdf(self):
        if not is_libreoffice_available():
            pytest.skip("LibreOffice not available")

        doc = Document()
        doc.add_paragraph("LibreOffice validation test.")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            pdf_path = validate_with_libreoffice(path)
            assert os.path.exists(pdf_path), "PDF should have been created"
            assert os.path.getsize(pdf_path) > 0, "PDF should not be empty"
        finally:
            os.unlink(path)

    @pytest.mark.libreoffice
    def it_converts_a_document_with_comments_to_pdf(self):
        if not is_libreoffice_available():
            pytest.skip("LibreOffice not available")

        doc = Document()
        doc.add_paragraph("Document with comments.")
        run = doc.paragraphs[0].runs[0]
        doc.add_comment(run, text="Comment for LO test", author="Author")

        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(path)
            pdf_path = validate_with_libreoffice(path)
            assert os.path.exists(pdf_path), "PDF should have been created"
            assert os.path.getsize(pdf_path) > 0, "PDF should not be empty"
        finally:
            os.unlink(path)
