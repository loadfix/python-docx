# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.ink` module and related ink-annotation access."""

from __future__ import annotations

from typing import cast

import pytest

from docx import types as t
from docx.document import Document
from docx.ink import InkAnnotation
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml.document import CT_Document
from docx.oxml.text.paragraph import CT_P
from docx.parts.document import DocumentPart
from docx.parts.ink import InkPart
from docx.parts.story import StoryPart
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element
from .unitutil.mock import FixtureRequest, instance_mock


INK_XML_TWO_TRACES = (
    b'<?xml version="1.0"?>\n'
    b'<inkml:ink xmlns:inkml="http://www.w3.org/2003/InkML">\n'
    b"  <inkml:trace>100 100, 150 150, 200 200</inkml:trace>\n"
    b"  <inkml:trace>300 300, 350 350</inkml:trace>\n"
    b"</inkml:ink>\n"
)

INK_XML_ONE_TRACE = (
    b'<?xml version="1.0"?>\n'
    b'<inkml:ink xmlns:inkml="http://www.w3.org/2003/InkML">\n'
    b"  <inkml:trace>1 2</inkml:trace>\n"
    b"</inkml:ink>\n"
)

INK_XML_EMPTY = (
    b'<?xml version="1.0"?>\n'
    b'<inkml:ink xmlns:inkml="http://www.w3.org/2003/InkML"/>\n'
)


def _make_ink_part(idx: int = 1, blob: bytes = INK_XML_TWO_TRACES) -> InkPart:
    return InkPart(PackURI("/word/ink/ink%d.xml" % idx), CT.INKML, blob)


class DescribeInkAnnotation:
    """Unit-test suite for `docx.ink.InkAnnotation`."""

    def it_knows_its_partname(self, fake_parent: t.ProvidesStoryPart):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        ink_part = _make_ink_part(idx=3)

        annotation = InkAnnotation(paragraph, ink_part)

        assert annotation.partname == "/word/ink/ink3.xml"

    def it_provides_access_to_the_raw_blob(self, fake_parent: t.ProvidesStoryPart):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        ink_part = _make_ink_part(blob=INK_XML_ONE_TRACE)

        annotation = InkAnnotation(paragraph, ink_part)

        assert annotation.blob == INK_XML_ONE_TRACE

    def it_provides_access_to_the_paragraph_it_belongs_to(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        ink_part = _make_ink_part()

        annotation = InkAnnotation(paragraph, ink_part)

        assert annotation.paragraph is paragraph

    @pytest.mark.parametrize(
        ("blob", "expected_count"),
        [
            (INK_XML_TWO_TRACES, 2),
            (INK_XML_ONE_TRACE, 1),
            (INK_XML_EMPTY, 0),
            (b"", 0),
            (b"not xml at all<<<", 0),
        ],
    )
    def it_counts_strokes_from_the_ink_xml(
        self,
        blob: bytes,
        expected_count: int,
        fake_parent: t.ProvidesStoryPart,
    ):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        ink_part = _make_ink_part(blob=blob)

        annotation = InkAnnotation(paragraph, ink_part)

        assert annotation.stroke_count == expected_count

    def it_surfaces_an_InkContent_via_ink_content_when_ooxml_ink_available(
        self, fake_parent: t.ProvidesStoryPart
    ):
        pytest.importorskip("ooxml_ink")
        from ooxml_ink.proxies import InkContent  # noqa: E402

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        ink_part = _make_ink_part(blob=INK_XML_TWO_TRACES)
        annotation = InkAnnotation(paragraph, ink_part)

        content = annotation.ink_content
        assert isinstance(content, InkContent)
        assert content.stroke_count == 2
        assert content.blob == INK_XML_TWO_TRACES

    def it_returns_None_ink_content_on_empty_blob(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        ink_part = _make_ink_part(blob=b"")
        annotation = InkAnnotation(paragraph, ink_part)
        assert annotation.ink_content is None

    def it_returns_None_ink_content_on_malformed_blob(
        self, fake_parent: t.ProvidesStoryPart
    ):
        pytest.importorskip("ooxml_ink")
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        ink_part = _make_ink_part(blob=b"<definitely-not-ink")
        annotation = InkAnnotation(paragraph, ink_part)
        assert annotation.ink_content is None

    # -- fixtures -----------------------------------------------------------------

    @pytest.fixture
    def fake_parent(self, request: FixtureRequest):
        story_part_ = instance_mock(request, StoryPart)

        class FakeParent:
            @property
            def part(self):
                return story_part_

        return FakeParent()


class DescribeParagraph_ink_annotations:
    """Unit-test suite for `Paragraph.ink_annotations`."""

    def it_returns_empty_list_when_no_contentPart_elements(
        self, request: FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.related_parts = {}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(
            CT_P,
            element('w:p/w:r/w:t"some text"'),
        )
        paragraph = Paragraph(p, FakeParent())

        assert paragraph.ink_annotations == []

    def it_enumerates_each_contentPart_as_an_ink_annotation(
        self, request: FixtureRequest
    ):
        ink_part_1 = _make_ink_part(idx=1, blob=INK_XML_TWO_TRACES)
        ink_part_2 = _make_ink_part(idx=2, blob=INK_XML_ONE_TRACE)
        story_part_ = instance_mock(request, StoryPart)
        story_part_.related_parts = {"rId5": ink_part_1, "rId6": ink_part_2}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(
            CT_P,
            element("w:p/(w:r/w:contentPart{r:id=rId5},w:r/w:contentPart{r:id=rId6})"),
        )
        paragraph = Paragraph(p, FakeParent())

        annotations = paragraph.ink_annotations

        assert len(annotations) == 2
        assert [a.partname for a in annotations] == [
            "/word/ink/ink1.xml",
            "/word/ink/ink2.xml",
        ]
        assert [a.stroke_count for a in annotations] == [2, 1]
        assert all(a.paragraph is paragraph for a in annotations)

    def it_skips_contentPart_with_unresolved_or_nonmatching_rel(
        self, request: FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        # -- rId5 is mapped to something that isn't an InkPart --
        story_part_.related_parts = {"rId5": "not-a-part"}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        # -- one rId resolves to the wrong kind of part, one rId is missing entirely --
        p = cast(
            CT_P,
            element("w:p/(w:r/w:contentPart{r:id=rId5},w:r/w:contentPart{r:id=rId9})"),
        )
        paragraph = Paragraph(p, FakeParent())

        assert paragraph.ink_annotations == []

    def it_skips_contentPart_without_a_relationship_id(
        self, request: FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.related_parts = {}

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p/w:r/w:contentPart"))
        paragraph = Paragraph(p, FakeParent())

        assert paragraph.ink_annotations == []


class DescribeDocument_ink_annotations:
    """Unit-test suite for `Document.ink_annotations`."""

    def it_returns_empty_list_when_document_has_no_ink(self, request: FixtureRequest):
        doc_elm = cast(CT_Document, element("w:document/w:body/w:p"))
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {}
        document = Document(doc_elm, document_part_)

        assert document.ink_annotations == []

    def it_collects_ink_annotations_across_all_paragraphs(
        self, request: FixtureRequest
    ):
        ink_part_1 = _make_ink_part(idx=1, blob=INK_XML_TWO_TRACES)
        ink_part_2 = _make_ink_part(idx=2, blob=INK_XML_ONE_TRACE)
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {"rId1": ink_part_1, "rId2": ink_part_2}

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                "w:p/w:r/w:contentPart{r:id=rId1},"
                "w:p,"
                "w:p/w:r/w:contentPart{r:id=rId2}"
                ")"
            ),
        )
        document = Document(doc_elm, document_part_)

        annotations = document.ink_annotations

        assert len(annotations) == 2
        assert {a.partname for a in annotations} == {
            "/word/ink/ink1.xml",
            "/word/ink/ink2.xml",
        }
        # -- stroke counts consistent across both collections --
        assert sorted(a.stroke_count for a in annotations) == [1, 2]

    def it_matches_paragraph_scoped_enumeration(self, request: FixtureRequest):
        ink_part_1 = _make_ink_part(idx=1, blob=INK_XML_TWO_TRACES)
        document_part_ = instance_mock(request, DocumentPart)
        document_part_.related_parts = {"rId1": ink_part_1}

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/w:r/w:contentPart{r:id=rId1}"
            ),
        )
        document = Document(doc_elm, document_part_)

        doc_annotations = document.ink_annotations
        para_annotations = document.paragraphs[0].ink_annotations

        assert len(doc_annotations) == 1
        assert len(para_annotations) == 1
        assert doc_annotations[0].partname == para_annotations[0].partname
        assert doc_annotations[0].stroke_count == para_annotations[0].stroke_count


class DescribeInkPart:
    """Unit-test suite for `docx.parts.ink.InkPart`."""

    def it_exposes_its_blob(self):
        partname = PackURI("/word/ink/ink1.xml")
        ink_part = InkPart(partname, CT.INKML, INK_XML_ONE_TRACE)

        assert ink_part.blob == INK_XML_ONE_TRACE
        assert ink_part.partname == partname
        assert ink_part.content_type == CT.INKML

    def it_is_constructed_by_the_part_loader(self, request: FixtureRequest):
        from docx.opc.part import PartFactory
        from docx.package import Package

        package_ = instance_mock(request, Package)
        part = PartFactory(
            PackURI("/word/ink/ink1.xml"),
            CT.INKML,
            RT.INK,
            INK_XML_ONE_TRACE,
            package_,
        )

        assert isinstance(part, InkPart)
        assert part.blob == INK_XML_ONE_TRACE
