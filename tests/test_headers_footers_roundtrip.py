"""Round-trip tests for the full first/even/odd header & footer surface.

Verifies that
  * per-section ``w:titlePg`` (different-first-page) survives save/reopen
  * the document-level ``w:evenAndOddHeaders`` setting survives save/reopen
  * the three variants of header/footer -- primary (default), first-page,
    and even-page -- each end up in their own ``/word/headerN.xml``
    or ``/word/footerN.xml`` part with independent content
  * ``_Header.is_linked_to_previous`` and ``_Footer.is_linked_to_previous``
    drive section-to-section inheritance correctly, both ways:
        - assigning ``True`` drops the explicit reference and the section
          inherits the previous section's definition
        - assigning ``False`` adds a fresh, empty part
  * unlinked-then-populated headers in a second section survive round-trip
    without leaking content across sections.
"""

from __future__ import annotations

import zipfile

from docx import Document
from docx.enum.section import WD_SECTION

from .helpers.roundtrip import assert_round_trip, save_and_reopen


class DescribeHeadersFootersRoundTrip:
    """Write header/footer variants, save, reopen, assert."""

    def it_round_trips_all_three_header_variants(self):
        def create(doc):
            section = doc.sections[0]
            section.different_first_page_header_footer = True
            section.different_odd_and_even_pages_header_footer = True
            section.header.paragraphs[0].text = "DEFAULT HEADER"
            section.first_page_header.paragraphs[0].text = "FIRST HEADER"
            section.even_page_header.paragraphs[0].text = "EVEN HEADER"

        def check(doc, _):
            section = doc.sections[0]
            assert section.different_first_page_header_footer is True
            assert section.different_odd_and_even_pages_header_footer is True
            assert section.header.paragraphs[0].text == "DEFAULT HEADER"
            assert section.first_page_header.paragraphs[0].text == "FIRST HEADER"
            assert section.even_page_header.paragraphs[0].text == "EVEN HEADER"

        assert_round_trip(create, check)

    def it_round_trips_all_three_footer_variants(self):
        def create(doc):
            section = doc.sections[0]
            section.different_first_page_header_footer = True
            section.different_odd_and_even_pages_header_footer = True
            section.footer.paragraphs[0].text = "DEFAULT FOOTER"
            section.first_page_footer.paragraphs[0].text = "FIRST FOOTER"
            section.even_page_footer.paragraphs[0].text = "EVEN FOOTER"

        def check(doc, _):
            section = doc.sections[0]
            assert section.footer.paragraphs[0].text == "DEFAULT FOOTER"
            assert section.first_page_footer.paragraphs[0].text == "FIRST FOOTER"
            assert section.even_page_footer.paragraphs[0].text == "EVEN FOOTER"

        assert_round_trip(create, check)

    def it_writes_three_distinct_header_parts_for_a_section(self, tmp_path):
        out = tmp_path / "three-hdrs.docx"
        doc = Document()
        section = doc.sections[0]
        section.different_first_page_header_footer = True
        section.different_odd_and_even_pages_header_footer = True
        section.header.paragraphs[0].text = "DEFAULT"
        section.first_page_header.paragraphs[0].text = "FIRST"
        section.even_page_header.paragraphs[0].text = "EVEN"
        doc.save(str(out))

        # -- three separate /word/headerN.xml parts must exist --
        with zipfile.ZipFile(str(out)) as z:
            members = z.namelist()
            hdr_parts = [m for m in members if m.startswith("word/header") and m.endswith(".xml")]
        assert len(hdr_parts) == 3, (
            "expected 3 header parts for default+first+even, got %r" % hdr_parts
        )

    def it_keeps_even_odd_flag_in_settings_xml(self, tmp_path):
        out = tmp_path / "evenodd.docx"
        doc = Document()
        doc.settings.even_and_odd_headers = True
        doc.save(str(out))

        # -- flag must survive and the settings part must carry it --
        doc2 = Document(str(out))
        assert doc2.settings.even_and_odd_headers is True
        assert doc2.sections[0].different_odd_and_even_pages_header_footer is True


class DescribeSectionLinkedToPreviousRoundTrip:
    """Verify linking semantics -- no reference => inherit from prior section."""

    def it_inherits_header_from_prior_section_by_default(self):
        def create(doc):
            s1 = doc.sections[0]
            s1.header.paragraphs[0].text = "S1 HEADER"
            doc.add_section(WD_SECTION.NEW_PAGE)
            # -- s2 defaults to linked (no headerReference) --

        def check(doc, _):
            assert len(doc.sections) == 2
            s1, s2 = doc.sections
            assert s1.header.paragraphs[0].text == "S1 HEADER"
            # -- linked: no explicit reference on s2 --
            assert s2.header.is_linked_to_previous is True
            # -- and the content visible on s2 is inherited from s1 --
            assert s2.header.paragraphs[0].text == "S1 HEADER"

        assert_round_trip(create, check)

    def it_unlinks_the_header_on_assigning_False(self):
        def create(doc):
            s1 = doc.sections[0]
            s1.header.paragraphs[0].text = "S1 HEADER"
            doc.add_section(WD_SECTION.NEW_PAGE)
            s2 = doc.sections[1]
            # -- explicit unlink then populate --
            s2.header.is_linked_to_previous = False
            s2.header.paragraphs[0].text = "S2 HEADER"

        def check(doc, _):
            s1, s2 = doc.sections
            assert s1.header.is_linked_to_previous is False
            assert s2.header.is_linked_to_previous is False
            assert s1.header.paragraphs[0].text == "S1 HEADER"
            assert s2.header.paragraphs[0].text == "S2 HEADER"

        assert_round_trip(create, check)

    def it_relinks_the_header_on_assigning_True(self):
        doc = Document()
        s1 = doc.sections[0]
        s1.header.paragraphs[0].text = "S1 HEADER"
        doc.add_section(WD_SECTION.NEW_PAGE)
        s2 = doc.sections[1]
        s2.header.is_linked_to_previous = False
        s2.header.paragraphs[0].text = "S2 HEADER"

        # -- round-trip, then relink --
        doc2 = save_and_reopen(doc)
        sec1, sec2 = doc2.sections
        assert sec2.header.paragraphs[0].text == "S2 HEADER"

        sec2.header.is_linked_to_previous = True

        # -- after relink, sec2's displayed content tracks sec1 again --
        assert sec2.header.is_linked_to_previous is True
        assert sec2.header.paragraphs[0].text == "S1 HEADER"

        # -- persist and reopen once more to verify the reference is gone --
        doc3 = save_and_reopen(doc2)
        s1b, s2b = doc3.sections
        assert s2b.header.is_linked_to_previous is True
        assert s2b.header.paragraphs[0].text == "S1 HEADER"

    def it_unlinks_a_footer_symmetrically(self):
        def create(doc):
            s1 = doc.sections[0]
            s1.footer.paragraphs[0].text = "S1 FOOTER"
            doc.add_section(WD_SECTION.NEW_PAGE)
            s2 = doc.sections[1]
            s2.footer.is_linked_to_previous = False
            s2.footer.paragraphs[0].text = "S2 FOOTER"

        def check(doc, _):
            s1, s2 = doc.sections
            assert s1.footer.paragraphs[0].text == "S1 FOOTER"
            assert s2.footer.paragraphs[0].text == "S2 FOOTER"
            assert s1.footer.is_linked_to_previous is False
            assert s2.footer.is_linked_to_previous is False

        assert_round_trip(create, check)

    def it_inherits_first_page_header_across_sections(self):
        def create(doc):
            s1 = doc.sections[0]
            s1.different_first_page_header_footer = True
            s1.first_page_header.paragraphs[0].text = "FIRST PAGE"
            doc.add_section(WD_SECTION.NEW_PAGE)
            s2 = doc.sections[1]
            s2.different_first_page_header_footer = True
            # -- leave linked: s2.first_page_header should inherit from s1 --

        def check(doc, _):
            s1, s2 = doc.sections
            assert s2.first_page_header.is_linked_to_previous is True
            assert s2.first_page_header.paragraphs[0].text == "FIRST PAGE"
            assert s1.first_page_header.paragraphs[0].text == "FIRST PAGE"

        assert_round_trip(create, check)
