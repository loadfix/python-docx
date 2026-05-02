# pyright: reportPrivateUsage=false

"""Test suite for the docx.table module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import (
    WD_ALIGN_VERTICAL,
    WD_BORDER_STYLE,
    WD_ROW_HEIGHT,
    WD_SHADING_PATTERN,
    WD_TABLE_ALIGNMENT,
    WD_TABLE_AUTOFIT,
    WD_TABLE_DIRECTION,
    WD_TEXT_DIRECTION,
)
from docx.oxml.parser import parse_xml
from docx.oxml.table import CT_Row, CT_Tbl, CT_TblGridCol, CT_Tc
from docx.parts.document import DocumentPart
from docx.shared import Emu, Inches, Length, Pt, RGBColor, Twips
from docx.table import (
    BorderElement,
    CellBorders,
    CellMargins,
    CellShading,
    Table,
    TableBorders,
    TableCellMargins,
    TableStyleFlags,
    _Cell,
    _Column,
    _Columns,
    _Row,
    _Rows,
)
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element, xml
from .unitutil.file import snippet_seq
from .unitutil.mock import FixtureRequest, Mock, instance_mock, property_mock


class DescribeTable:
    """Unit-test suite for `docx.table._Rows` objects."""

    def it_can_add_a_row(self, document_: Mock):
        snippets = snippet_seq("add-row-col")
        tbl = cast(CT_Tbl, parse_xml(snippets[0]))
        table = Table(tbl, document_)

        row = table.add_row()

        assert table._tbl.xml == snippets[1]
        assert isinstance(row, _Row)
        assert row._tr is table._tbl.tr_lst[-1]
        assert row._parent is table

    def it_copies_cnfStyle_from_the_preceding_row_when_adding_a_row(
        self, document_: Mock
    ):
        """``add_row`` propagates the prior row's ``w:trPr/w:cnfStyle`` so
        conditional banding / firstRow / lastRow formatting flows to the new
        row, matching Word's own behaviour. (upstream#306)
        """
        # -- 1-col table, existing row has a cnfStyle flagging firstRow+oddHBand --
        tbl_cxml = (
            "w:tbl/(w:tblPr,w:tblGrid/w:gridCol{w:w=1000},"
            "w:tr/(w:trPr/w:cnfStyle{w:val=100000000000},"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=1000},w:p)))"
        )
        tbl = cast(CT_Tbl, element(tbl_cxml))
        table = Table(tbl, document_)

        table.add_row()

        new_tr = tbl.tr_lst[-1]
        trPr = new_tr.trPr
        assert trPr is not None
        from docx.oxml.ns import qn

        cnfStyle = trPr.find(qn("w:cnfStyle"))
        assert cnfStyle is not None
        assert cnfStyle.get(qn("w:val")) == "100000000000"

    def it_leaves_new_row_trPr_clean_when_predecessor_has_no_cnfStyle(
        self, document_: Mock
    ):
        """Rows whose predecessor lacks ``cnfStyle`` must not get a spurious
        ``w:trPr`` element. (upstream#306 safety)
        """
        snippets = snippet_seq("add-row-col")
        tbl = cast(CT_Tbl, parse_xml(snippets[0]))
        table = Table(tbl, document_)

        table.add_row()

        new_tr = table._tbl.tr_lst[-1]
        # -- no cnfStyle on predecessor => no trPr on new row --
        assert new_tr.trPr is None

    @pytest.mark.parametrize(
        ("body_cxml", "tbl_idx", "expected_cxml"),
        [
            # --- table removed from body with paragraph sibling ---
            ("w:body/(w:tbl/w:tblPr,w:p)", 0, "w:body/w:p"),
            # --- table removed leaving another table ---
            ("w:body/(w:tbl/w:tblPr,w:tbl/w:tblPr)", 0, "w:body/w:tbl/w:tblPr"),
            # --- second table removed ---
            ("w:body/(w:p,w:tbl/w:tblPr,w:p)", 0, "w:body/(w:p,w:p)"),
        ],
    )
    def it_can_delete_itself(
        self,
        body_cxml: str,
        tbl_idx: int,
        expected_cxml: str,
        document_: Mock,
    ):
        body = element(body_cxml)
        tbl = body.tbl_lst[tbl_idx]
        table = Table(tbl, document_)

        table.delete()

        assert body.xml == xml(expected_cxml)

    def it_can_add_a_column(self, document_: Mock):
        snippets = snippet_seq("add-row-col")
        tbl = cast(CT_Tbl, parse_xml(snippets[0]))
        table = Table(tbl, document_)

        column = table.add_column(Inches(1.5))

        assert table._tbl.xml == snippets[2]
        assert isinstance(column, _Column)
        assert column._gridCol is table._tbl.tblGrid.gridCol_lst[-1]
        assert column._parent is table

    def it_add_column_inserts_a_tc_in_every_row(self, document_: Mock):
        # -- regression for upstream#1102: `add_column` must append a `w:tc`
        # -- in every row in addition to the new `w:gridCol`; otherwise the
        # -- rendered document ends up with a grid-column count that doesn't
        # -- match row cell counts and Word refuses to open the file. --
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/w:p),"
                "w:tr/(w:tc/w:p,w:tc/w:p))"
            ),
        )
        table = Table(tbl, document_)
        original_row_tc_counts = [len(tr.tc_lst) for tr in tbl.tr_lst]

        table.add_column(Inches(1.0))

        new_row_tc_counts = [len(tr.tc_lst) for tr in tbl.tr_lst]
        # -- every row gained exactly one new tc --
        assert new_row_tc_counts == [n + 1 for n in original_row_tc_counts]
        # -- tblGrid gained exactly one new gridCol --
        assert len(tbl.tblGrid.gridCol_lst) == 3
        # -- every row's total tc count now matches the grid column count --
        for tr in tbl.tr_lst:
            assert len(tr.tc_lst) == len(tbl.tblGrid.gridCol_lst)

    def it_provides_access_to_a_cell_by_row_and_col_indices(self, table: Table):
        for row_idx in range(2):
            for col_idx in range(2):
                cell = table.cell(row_idx, col_idx)
                assert isinstance(cell, _Cell)
                tr = table._tbl.tr_lst[row_idx]
                tc = tr.tc_lst[col_idx]
                assert tc is cell._tc

    def it_provides_access_to_the_table_rows(self, table: Table):
        rows = table.rows
        assert isinstance(rows, _Rows)

    def it_provides_access_to_the_table_columns(self, table: Table):
        columns = table.columns
        assert isinstance(columns, _Columns)

    def it_provides_access_to_the_cells_in_a_column(
        self, _cells_: Mock, _column_count_: Mock, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element("w:tbl")), document_)
        _cells_.return_value = [0, 1, 2, 3, 4, 5, 6, 7, 8]
        _column_count_.return_value = 3
        column_idx = 1

        column_cells = table.column_cells(column_idx)

        assert column_cells == [1, 4, 7]

    def it_provides_access_to_the_cells_in_a_row(
        self, _cells_: Mock, _column_count_: Mock, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element("w:tbl")), document_)
        _cells_.return_value = [0, 1, 2, 3, 4, 5, 6, 7, 8]
        _column_count_.return_value = 3

        with pytest.warns(DeprecationWarning, match="row_cells"):
            row_cells = table.row_cells(1)

        assert row_cells == [3, 4, 5]

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:jc{w:val=center}", WD_TABLE_ALIGNMENT.CENTER),
            ("w:tbl/w:tblPr/w:jc{w:val=right}", WD_TABLE_ALIGNMENT.RIGHT),
            ("w:tbl/w:tblPr/w:jc{w:val=left}", WD_TABLE_ALIGNMENT.LEFT),
        ],
    )
    def it_knows_its_alignment_setting(
        self, tbl_cxml: str, expected_value: WD_TABLE_ALIGNMENT | None, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.alignment == expected_value

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", WD_TABLE_ALIGNMENT.LEFT, "w:tbl/w:tblPr/w:jc{w:val=left}"),
            (
                "w:tbl/w:tblPr/w:jc{w:val=left}",
                WD_TABLE_ALIGNMENT.RIGHT,
                "w:tbl/w:tblPr/w:jc{w:val=right}",
            ),
            ("w:tbl/w:tblPr/w:jc{w:val=right}", None, "w:tbl/w:tblPr"),
        ],
    )
    def it_can_change_its_alignment_setting(
        self,
        tbl_cxml: str,
        new_value: WD_TABLE_ALIGNMENT | None,
        expected_cxml: str,
        document_: Mock,
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        table.alignment = new_value
        assert table._tbl.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", True),
            ("w:tbl/w:tblPr/w:tblLayout", True),
            ("w:tbl/w:tblPr/w:tblLayout{w:type=autofit}", True),
            ("w:tbl/w:tblPr/w:tblLayout{w:type=fixed}", False),
        ],
    )
    def it_knows_whether_it_should_autofit(
        self, tbl_cxml: str, expected_value: bool, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.autofit is expected_value

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", True, "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}"),
            ("w:tbl/w:tblPr", False, "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}"),
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}",
                True,
                "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}",
            ),
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}",
                False,
                "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}",
            ),
        ],
    )
    def it_can_change_its_autofit_setting(
        self, tbl_cxml: str, new_value: bool, expected_cxml: str, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        table.autofit = new_value
        assert table._tbl.xml == xml(expected_cxml)

    def it_knows_it_is_the_table_its_children_belong_to(self, table: Table):
        assert table.table is table

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:bidiVisual", WD_TABLE_DIRECTION.RTL),
            ("w:tbl/w:tblPr/w:bidiVisual{w:val=0}", WD_TABLE_DIRECTION.LTR),
            ("w:tbl/w:tblPr/w:bidiVisual{w:val=on}", WD_TABLE_DIRECTION.RTL),
        ],
    )
    def it_knows_its_direction(
        self, tbl_cxml: str, expected_value: WD_TABLE_DIRECTION | None, document_: Mock
    ):
        tbl = cast(CT_Tbl, element(tbl_cxml))
        assert Table(tbl, document_).table_direction == expected_value

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", WD_TABLE_DIRECTION.RTL, "w:tbl/w:tblPr/w:bidiVisual"),
            (
                "w:tbl/w:tblPr/w:bidiVisual",
                WD_TABLE_DIRECTION.LTR,
                "w:tbl/w:tblPr/w:bidiVisual{w:val=0}",
            ),
            (
                "w:tbl/w:tblPr/w:bidiVisual{w:val=0}",
                WD_TABLE_DIRECTION.RTL,
                "w:tbl/w:tblPr/w:bidiVisual",
            ),
            ("w:tbl/w:tblPr/w:bidiVisual{w:val=1}", None, "w:tbl/w:tblPr"),
        ],
    )
    def it_can_change_its_direction(
        self, tbl_cxml: str, new_value: WD_TABLE_DIRECTION, expected_cxml: str, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        table.table_direction = new_value
        assert table._element.xml == xml(expected_cxml)

    def it_exposes_direction_as_an_upstream_compatible_alias(self, document_: Mock):
        """`Table.direction` is an alias of `table_direction` (upstream#1227).

        loadfix renamed the property to ``table_direction``; the alias
        preserves upstream-name compatibility.
        """
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)

        # -- read delegates to table_direction --
        assert table.direction is None

        # -- write delegates to table_direction --
        table.direction = WD_TABLE_DIRECTION.RTL
        assert table.table_direction == WD_TABLE_DIRECTION.RTL
        assert table.direction == WD_TABLE_DIRECTION.RTL
        assert table._element.xml == xml("w:tbl/w:tblPr/w:bidiVisual")

        table.direction = None
        assert table.table_direction is None
        assert table._element.xml == xml("w:tbl/w:tblPr")

    def it_knows_its_table_style(self, part_prop_: Mock, document_part_: Mock, document_: Mock):
        part_prop_.return_value = document_part_
        style_ = document_part_.get_style.return_value
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblStyle{w:val=BarBaz}")), document_)

        style = table.style

        document_part_.get_style.assert_called_once_with("BarBaz", WD_STYLE_TYPE.TABLE)
        assert style is style_

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "style_id", "expected_cxml"),
        [
            ("w:tbl/w:tblPr", "Tbl A", "TblA", "w:tbl/w:tblPr/w:tblStyle{w:val=TblA}"),
            (
                "w:tbl/w:tblPr/w:tblStyle{w:val=TblA}",
                "Tbl B",
                "TblB",
                "w:tbl/w:tblPr/w:tblStyle{w:val=TblB}",
            ),
            ("w:tbl/w:tblPr/w:tblStyle{w:val=TblB}", None, None, "w:tbl/w:tblPr"),
        ],
    )
    def it_can_change_its_table_style(
        self,
        tbl_cxml: str,
        new_value: str | None,
        style_id: str | None,
        expected_cxml: str,
        document_: Mock,
        part_prop_: Mock,
        document_part_: Mock,
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        part_prop_.return_value = document_part_
        document_part_.get_style_id.return_value = style_id

        table.style = new_value

        document_part_.get_style_id.assert_called_once_with(new_value, WD_STYLE_TYPE.TABLE)
        assert table._tbl.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("snippet_idx", "cell_count", "unique_count", "matches"),
        [
            (0, 9, 9, ()),
            (1, 9, 8, ((0, 1),)),
            (2, 9, 8, ((1, 4),)),
            (3, 9, 6, ((0, 1, 3, 4),)),
            (4, 9, 4, ((0, 1), (3, 6), (4, 5, 7, 8))),
        ],
    )
    def it_provides_access_to_its_cells_to_help(
        self,
        snippet_idx: int,
        cell_count: int,
        unique_count: int,
        matches: tuple[tuple[int, ...]],
        document_: Mock,
    ):
        tbl_xml = snippet_seq("tbl-cells")[snippet_idx]
        table = Table(cast(CT_Tbl, parse_xml(tbl_xml)), document_)

        cells = table._cells

        assert len(cells) == cell_count
        assert len(set(cells)) == unique_count
        for matching_idxs in matches:
            comparator_idx = matching_idxs[0]
            for idx in matching_idxs[1:]:
                assert cells[idx] is cells[comparator_idx]

    def it_knows_its_column_count_to_help(self, document_: Mock):
        tbl_cxml = "w:tbl/w:tblGrid/(w:gridCol,w:gridCol,w:gridCol)"
        expected_value = 3
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        column_count = table._column_count

        assert column_count == expected_value

    def it_synthesises_column_count_when_tblGrid_is_missing(self, document_: Mock):
        """A ``w:tbl`` without ``w:tblGrid`` is invalid per ISO-29500 but is
        commonly emitted by LibreOffice and PDF-to-docx converters.

        Prior to the fix, :attr:`Table._column_count` raised
        ``InvalidXmlError``. After the fix, the column count is derived from
        the widest row (honouring ``gridBefore``/``gridAfter``). Closes
        upstream#548.
        """
        # -- 3-column widest row, no tblGrid --
        tbl_cxml = (
            "w:tbl/(w:tblPr,"
            "w:tr/(w:tc/w:p,w:tc/w:p,w:tc/w:p),"
            "w:tr/(w:tc/w:p,w:tc/w:p))"
        )
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        # -- should not raise --
        assert table._column_count == 3

    def it_tolerates_orphan_vMerge_continuation_in_top_row(self, document_: Mock):
        """A vMerge="continue" in the first row has no row above to delegate to.

        Without the hardening, ``Table._cells`` would index ``cells[-col_count]``
        on an empty-or-short list and raise. After the fix it emits a fresh
        cell for the orphan tc.
        """
        tbl_cxml = (
            "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol),"
            "w:tr/(w:tc/w:p,w:tc/(w:tcPr/w:vMerge,w:p)))"
        )
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        cells = table._cells

        assert len(cells) == 2
        # -- both cells should resolve to real _Cell objects --
        assert all(isinstance(c, _Cell) for c in cells)

    def it_resolves_every_continuation_in_a_multi_row_vMerge_span(
        self, document_: Mock
    ):
        """vMerge="restart" + several vMerge (continue) cells all alias to origin."""
        tbl_cxml = (
            "w:tbl/(w:tblGrid/w:gridCol,"
            "w:tr/w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p),"
            "w:tr/w:tc/(w:tcPr/w:vMerge,w:p),"
            "w:tr/w:tc/(w:tcPr/w:vMerge,w:p),"
            "w:tr/w:tc/(w:tcPr/w:vMerge,w:p))"
        )
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        cells = table._cells

        # -- 4 rows * 1 column = 4 grid positions, all the same cell object --
        assert len(cells) == 4
        assert len(set(cells)) == 1

    def it_preserves_nested_table_access_inside_a_merged_origin_cell(
        self, document_: Mock
    ):
        """Nested ``w:tbl`` inside a vMerge origin cell should remain discoverable."""
        tbl_cxml = (
            "w:tbl/(w:tblGrid/w:gridCol,"
            "w:tr/w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p,"
            "w:tbl/(w:tblGrid/w:gridCol,w:tr/w:tc/w:p),"
            "w:p),"
            "w:tr/w:tc/(w:tcPr/w:vMerge,w:p))"
        )
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        # -- _cells should aggregate correctly (both rows point to origin) --
        cells = table._cells
        assert len(cells) == 2
        assert cells[0] is cells[1]
        # -- origin cell should expose its nested table --
        origin = cells[0]
        assert len(origin.tables) == 1

    def it_honours_grid_before_when_flattening_cells(self, document_: Mock):
        """``w:gridBefore`` shifts the row's real cells to later columns.

        Prior to the fix, ``Table._cells`` ignored ``w:gridBefore`` so a row
        that skipped leading columns would leak its cells into preceding
        columns, breaking ``Table.cell(r, c)`` and ``Table.column_cells(c)``
        arithmetic. (upstream#939, #1367)
        """
        # -- 3x3 grid; row 1 has gridBefore=1 so its two cells live at cols 1-2 --
        tbl_cxml = (
            "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol,w:gridCol),"
            "w:tr/(w:tc/w:p,w:tc/w:p,w:tc/w:p),"
            "w:tr/(w:trPr/w:gridBefore{w:val=1},w:tc/w:p,w:tc/w:p))"
        )
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        cells = table._cells

        # -- grid is 3 rows? no: 2 rows * 3 cols = 6 slots --
        assert len(cells) == 6
        # -- real tc elements for row 1 live at grid columns 1 and 2 --
        row1_real_tc_0 = table._tbl.tr_lst[1].tc_lst[0]
        row1_real_tc_1 = table._tbl.tr_lst[1].tc_lst[1]
        assert table.cell(1, 1)._tc is row1_real_tc_0
        assert table.cell(1, 2)._tc is row1_real_tc_1
        # -- column 0 should NOT leak the row-1 real tc into row-1, col-0 --
        assert table.cell(1, 0)._tc is not row1_real_tc_0
        # -- column_cells(2) must pick up the second real tc of row 1 --
        col2 = table.column_cells(2)
        assert col2[1]._tc is row1_real_tc_1

    def it_honours_grid_after_when_flattening_cells(self, document_: Mock):
        """``w:gridAfter`` lets a row end early; omitted slots must not pull
        later-column cells into those positions. (upstream#1334, #1193)
        """
        tbl_cxml = (
            "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol,w:gridCol),"
            "w:tr/(w:tc/w:p,w:tc/w:p,w:tc/w:p),"
            "w:tr/(w:trPr/w:gridAfter{w:val=1},w:tc/w:p,w:tc/w:p))"
        )
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)

        cells = table._cells

        assert len(cells) == 6  # -- 2 rows * 3 cols --
        # -- row 1 real cells live at columns 0 and 1 --
        row1_real_tc_0 = table._tbl.tr_lst[1].tc_lst[0]
        row1_real_tc_1 = table._tbl.tr_lst[1].tc_lst[1]
        assert table.cell(1, 0)._tc is row1_real_tc_0
        assert table.cell(1, 1)._tc is row1_real_tc_1
        # -- column 2 in row 1 is omitted; placeholder comes from row above --
        row0_real_tc_2 = table._tbl.tr_lst[0].tc_lst[2]
        assert table.cell(1, 2)._tc is row0_real_tc_2

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            # -- no tblLayout, no tblW → autofit-to-contents is the OOXML default --
            ("w:tbl/w:tblPr", WD_TABLE_AUTOFIT.AUTOFIT_TO_CONTENTS),
            # -- explicit w:tblLayout=autofit with auto tblW --
            (
                "w:tbl/w:tblPr/(w:tblW{w:type=auto,w:w=0},w:tblLayout{w:type=autofit})",
                WD_TABLE_AUTOFIT.AUTOFIT_TO_CONTENTS,
            ),
            # -- tblLayout=autofit with tblW type=pct means autofit-to-window --
            (
                "w:tbl/w:tblPr/w:tblW{w:type=pct,w:w=5000}",
                WD_TABLE_AUTOFIT.AUTOFIT_TO_WINDOW,
            ),
            # -- tblLayout=fixed always wins --
            (
                "w:tbl/w:tblPr/(w:tblW{w:type=pct,w:w=5000},w:tblLayout{w:type=fixed})",
                WD_TABLE_AUTOFIT.FIXED_WIDTH,
            ),
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}",
                WD_TABLE_AUTOFIT.FIXED_WIDTH,
            ),
        ],
    )
    def it_knows_its_autofit_behavior(
        self, tbl_cxml: str, expected_value: WD_TABLE_AUTOFIT, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.autofit_behavior == expected_value

    def it_can_change_its_autofit_behavior_to_fixed(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)

        table.autofit_behavior = WD_TABLE_AUTOFIT.FIXED_WIDTH

        assert table.autofit_behavior == WD_TABLE_AUTOFIT.FIXED_WIDTH
        assert table._tbl.xml == xml("w:tbl/w:tblPr/w:tblLayout{w:type=fixed}")

    def it_can_change_its_autofit_behavior_to_autofit_to_contents(self, document_: Mock):
        table = Table(
            cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblLayout{w:type=fixed}")),
            document_,
        )

        table.autofit_behavior = WD_TABLE_AUTOFIT.AUTOFIT_TO_CONTENTS

        assert table.autofit_behavior == WD_TABLE_AUTOFIT.AUTOFIT_TO_CONTENTS
        assert table._tbl.xml == xml("w:tbl/w:tblPr/w:tblW{w:type=auto,w:w=0}")

    def it_can_change_its_autofit_behavior_to_autofit_to_window(self, document_: Mock):
        table = Table(
            cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblLayout{w:type=fixed}")),
            document_,
        )

        table.autofit_behavior = WD_TABLE_AUTOFIT.AUTOFIT_TO_WINDOW

        assert table.autofit_behavior == WD_TABLE_AUTOFIT.AUTOFIT_TO_WINDOW
        assert table._tbl.xml == xml("w:tbl/w:tblPr/w:tblW{w:type=pct,w:w=5000}")

    def it_roundtrips_each_autofit_behavior_value(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        for value in (
            WD_TABLE_AUTOFIT.FIXED_WIDTH,
            WD_TABLE_AUTOFIT.AUTOFIT_TO_CONTENTS,
            WD_TABLE_AUTOFIT.AUTOFIT_TO_WINDOW,
            WD_TABLE_AUTOFIT.FIXED_WIDTH,
        ):
            table.autofit_behavior = value
            assert table.autofit_behavior == value

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", True),
            ("w:tbl/w:tblPr/w:tblLayout", True),
            ("w:tbl/w:tblPr/w:tblLayout{w:type=autofit}", True),
            ("w:tbl/w:tblPr/w:tblLayout{w:type=fixed}", False),
        ],
    )
    def it_knows_whether_it_allows_autofit(
        self, tbl_cxml: str, expected_value: bool, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.allow_autofit is expected_value

    @pytest.mark.parametrize(
        ("tbl_cxml", "new_value", "expected_cxml"),
        [
            # -- turning autofit off writes an explicit tblLayout=fixed --
            ("w:tbl/w:tblPr", False, "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}"),
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}",
                True,
                "w:tbl/w:tblPr",
            ),
            # -- turning autofit on removes the tblLayout child entirely --
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}",
                True,
                "w:tbl/w:tblPr",
            ),
            (
                "w:tbl/w:tblPr/w:tblLayout{w:type=autofit}",
                False,
                "w:tbl/w:tblPr/w:tblLayout{w:type=fixed}",
            ),
        ],
    )
    def it_can_change_whether_it_allows_autofit(
        self, tbl_cxml: str, new_value: bool, expected_cxml: str, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        table.allow_autofit = new_value
        assert table._tbl.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_value"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:tblW{w:type=auto,w:w=0}", None),
            ("w:tbl/w:tblPr/w:tblW{w:type=pct,w:w=5000}", None),
            ("w:tbl/w:tblPr/w:tblW{w:type=dxa,w:w=1440}", 914400),
            ("w:tbl/w:tblPr/w:tblW{w:type=dxa,w:w=4680}", 2971800),
        ],
    )
    def it_knows_its_preferred_width(
        self, tbl_cxml: str, expected_value: int | None, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.preferred_width == expected_value

    def it_can_change_its_preferred_width(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)

        table.preferred_width = Inches(1)

        assert table.preferred_width == Inches(1)
        assert table._tbl.xml == xml("w:tbl/w:tblPr/w:tblW{w:type=dxa,w:w=1440}")

    def it_can_clear_its_preferred_width(self, document_: Mock):
        table = Table(
            cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblW{w:type=dxa,w:w=1440}")),
            document_,
        )

        table.preferred_width = None

        assert table.preferred_width is None
        assert table._tbl.xml == xml("w:tbl/w:tblPr")

    def it_can_overwrite_an_existing_preferred_width(self, document_: Mock):
        table = Table(
            cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblW{w:type=auto,w:w=0}")),
            document_,
        )

        table.preferred_width = Inches(2)

        assert table.preferred_width == Inches(2)
        assert table._tbl.xml == xml("w:tbl/w:tblPr/w:tblW{w:type=dxa,w:w=2880}")

    # -- alt_text / alt_description (upstream#1048, #921) ------------

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:tblCaption{w:val=Sales}", "Sales"),
        ],
    )
    def it_knows_its_alt_text(
        self, tbl_cxml: str, expected: str | None, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.alt_text == expected

    def it_can_set_its_alt_text(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        table.alt_text = "Quarterly sales"
        assert table.alt_text == "Quarterly sales"
        assert table._tbl.xml == xml(
            "w:tbl/w:tblPr/w:tblCaption{w:val=Quarterly sales}"
        )

    def it_can_clear_its_alt_text(self, document_: Mock):
        table = Table(
            cast(
                CT_Tbl,
                element("w:tbl/w:tblPr/w:tblCaption{w:val=x}"),
            ),
            document_,
        )
        table.alt_text = None
        assert table.alt_text is None
        assert table._tbl.xml == xml("w:tbl/w:tblPr")

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:tblDescription{w:val=Monthly totals}", "Monthly totals"),
        ],
    )
    def it_knows_its_alt_description(
        self, tbl_cxml: str, expected: str | None, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.alt_description == expected

    def it_can_set_and_clear_its_alt_description(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        table.alt_description = "A long description"
        assert table.alt_description == "A long description"
        assert table._tbl.xml == xml(
            "w:tbl/w:tblPr/w:tblDescription{w:val=A long description}"
        )
        table.alt_description = None
        assert table.alt_description is None

    # -- indent / left_indent (upstream#1144, #586) ------------------

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected"),
        [
            ("w:tbl/w:tblPr", None),
            ("w:tbl/w:tblPr/w:tblInd{w:w=720,w:type=dxa}", Twips(720)),
        ],
    )
    def it_knows_its_indent(
        self, tbl_cxml: str, expected: Length | None, document_: Mock
    ):
        table = Table(cast(CT_Tbl, element(tbl_cxml)), document_)
        assert table.indent == expected
        assert table.left_indent == expected

    def it_can_change_its_indent(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        table.indent = Inches(0.5)
        assert table.indent == Inches(0.5)
        assert table._tbl.xml == xml("w:tbl/w:tblPr/w:tblInd{w:w=720,w:type=dxa}")

    def it_left_indent_is_an_alias_of_indent(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        table.left_indent = Inches(1)
        assert table.indent == Inches(1)
        assert table._tbl.xml == xml("w:tbl/w:tblPr/w:tblInd{w:w=1440,w:type=dxa}")

    def it_can_clear_its_indent(self, document_: Mock):
        table = Table(
            cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblInd{w:w=720,w:type=dxa}")),
            document_,
        )
        table.indent = None
        assert table.indent is None
        assert table._tbl.xml == xml("w:tbl/w:tblPr")

    # -- cell_margins (upstream#1401) --------------------------------

    def it_exposes_cell_margins_as_a_TableCellMargins(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        assert isinstance(table.cell_margins, TableCellMargins)

    def it_reads_None_for_every_cell_margin_edge_when_absent(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        m = table.cell_margins
        assert m.top is None
        assert m.bottom is None
        assert m.start is None
        assert m.end is None

    def it_creates_tblCellMar_on_first_write(self, document_: Mock):
        table = Table(cast(CT_Tbl, element("w:tbl/w:tblPr")), document_)
        table.cell_margins.top = Twips(100)
        assert table.cell_margins.top == Twips(100)
        assert table._tbl.xml == xml(
            "w:tbl/w:tblPr/w:tblCellMar/w:top{w:w=100,w:type=dxa}"
        )

    def it_removes_empty_tblCellMar_when_last_edge_cleared(self, document_: Mock):
        table = Table(
            cast(
                CT_Tbl,
                element("w:tbl/w:tblPr/w:tblCellMar/w:top{w:w=100,w:type=dxa}"),
            ),
            document_,
        )
        table.cell_margins.top = None
        assert table._tbl.xml == xml("w:tbl/w:tblPr")

    # -- merged_cell_ranges (PR#1036) --------------------------------

    def it_returns_empty_list_when_no_merged_cells(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element("w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol),w:tr/(w:tc,w:tc))"),
        )
        table = Table(tbl, document_)
        assert table.merged_cell_ranges() == []

    def it_reports_horizontal_merge_range(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/("
                "w:tblPr,"
                "w:tblGrid/(w:gridCol,w:gridCol,w:gridCol),"
                "w:tr/("
                "w:tc/w:tcPr/w:gridSpan{w:val=2},"
                "w:tc)"
                ")"
            ),
        )
        table = Table(tbl, document_)
        ranges = table.merged_cell_ranges()
        assert ranges == [(0, 1, 0, 2)]

    def it_reports_vertical_merge_range(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/("
                "w:tblPr,"
                "w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/("
                "w:tc/w:tcPr/w:vMerge{w:val=restart},"
                "w:tc),"
                "w:tr/("
                "w:tc/w:tcPr/w:vMerge,"
                "w:tc)"
                ")"
            ),
        )
        table = Table(tbl, document_)
        ranges = table.merged_cell_ranges()
        assert ranges == [(0, 2, 0, 1)]

    # -- split (upstream#481) ----------------------------------------

    def it_can_split_itself_at_a_row_boundary(self, document_: Mock):
        body = element(
            "w:body/w:tbl/("
            "w:tblPr,w:tblGrid/w:gridCol,"
            "w:tr/w:tc/w:p,"
            "w:tr/w:tc/w:p,"
            "w:tr/w:tc/w:p"
            ")"
        )
        tbl = body.tbl_lst[0]
        table = Table(cast(CT_Tbl, tbl), document_)

        paragraph, new_table = table.split(before_row=2)

        # -- paragraph + new tbl were inserted after the original --
        assert isinstance(paragraph, Paragraph)
        assert isinstance(new_table, Table)
        assert list(body) == [tbl, paragraph._p, new_table._tbl]
        # -- original retains rows 0..1, new table has row 2 --
        assert len(table._tbl.tr_lst) == 2
        assert len(new_table._tbl.tr_lst) == 1
        # -- tblPr and tblGrid were copied onto the new table --
        assert new_table._tbl.tblPr is not None
        assert new_table._tbl.tblGrid is not None

    @pytest.mark.parametrize("bad", [0, -1, 3, 99])
    def it_raises_on_out_of_range_split(self, bad: int, document_: Mock):
        body = element(
            "w:body/w:tbl/("
            "w:tblPr,w:tblGrid/w:gridCol,"
            "w:tr/w:tc/w:p,"
            "w:tr/w:tc/w:p,"
            "w:tr/w:tc/w:p"
            ")"
        )
        tbl = body.tbl_lst[0]
        table = Table(cast(CT_Tbl, tbl), document_)
        with pytest.raises(ValueError):
            table.split(before_row=bad)

    def it_can_insert_a_paragraph_before_itself(self, document_: Mock):
        body = element("w:body/(w:p{id=1},w:tbl/(w:tblPr,w:tblGrid),w:p{id=2})")
        tbl = body.tbl_lst[0]
        table = Table(cast(CT_Tbl, tbl), document_)

        paragraph = table.insert_paragraph_before()

        assert isinstance(paragraph, Paragraph)
        # -- paragraph sits between p{id=1} and the table --
        assert list(body) == [body[0], paragraph._p, tbl, body[3]]
        assert paragraph._parent is document_

    def it_can_insert_a_paragraph_after_itself(self, document_: Mock):
        body = element("w:body/(w:p{id=1},w:tbl/(w:tblPr,w:tblGrid),w:p{id=2})")
        tbl = body.tbl_lst[0]
        table = Table(cast(CT_Tbl, tbl), document_)

        paragraph = table.insert_paragraph_after(text="caption")

        assert isinstance(paragraph, Paragraph)
        # -- paragraph sits between the table and p{id=2} --
        assert list(body) == [body[0], tbl, paragraph._p, body[3]]
        assert paragraph.text == "caption"

    def it_can_insert_a_paragraph_after_inside_a_cell(self, document_: Mock):
        tc = element("w:tc/(w:tbl/(w:tblPr,w:tblGrid),w:p{id=end})")
        tbl = tc.tbl_lst[0]
        table = Table(cast(CT_Tbl, tbl), document_)

        paragraph = table.insert_paragraph_after(text="after-table")

        # -- the new paragraph sits between the inner table and the trailing w:p --
        assert list(tc) == [tbl, paragraph._p, tc[2]]
        assert paragraph.text == "after-table"

    def it_can_insert_a_table_before_itself(self, document_: Mock):
        body = element("w:body/(w:p{id=1},w:tbl/(w:tblPr,w:tblGrid))")
        ref_tbl = body.tbl_lst[0]
        table = Table(cast(CT_Tbl, ref_tbl), document_)

        new_table = table.insert_table_before(rows=2, cols=3)

        assert isinstance(new_table, Table)
        # -- the new table sits between the paragraph and the original table --
        assert list(body) == [body[0], new_table._tbl, ref_tbl]
        assert len(new_table._tbl.tr_lst) == 2
        assert new_table._tbl.col_count == 3
        assert new_table._parent is document_

    def it_can_insert_a_table_after_itself(self, document_: Mock):
        body = element("w:body/(w:tbl/(w:tblPr,w:tblGrid),w:p{id=end})")
        ref_tbl = body.tbl_lst[0]
        table = Table(cast(CT_Tbl, ref_tbl), document_)

        new_table = table.insert_table_after(rows=1, cols=1)

        assert isinstance(new_table, Table)
        # -- the new table sits between the original table and the trailing paragraph --
        assert list(body) == [ref_tbl, new_table._tbl, body[2]]

    def it_returns_None_for_formatting_change_when_tblPrChange_absent(
        self, document_: Mock
    ):
        tbl = cast(CT_Tbl, element("w:tbl/(w:tblPr,w:tblGrid)"))
        table = Table(tbl, document_)
        assert table.formatting_change is None

    def it_exposes_its_formatting_change_when_tblPrChange_present(
        self, document_: Mock
    ):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/("
                "w:tblPr/w:tblPrChange{w:id=1,w:author=A}/w:tblPr/"
                "w:tblW{w:w=5000,w:type=dxa}"
                ",w:tblGrid)"
            ),
        )
        table = Table(tbl, document_)
        fc = table.formatting_change
        assert fc is not None
        assert fc.author == "A"
        assert fc.old_properties is not None
        assert fc.old_properties.xpath("./w:tblW")

    # -- Phase C: delete_column / insert_row / clone add_row ---------

    def it_can_delete_a_column(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/w:p,w:tc/w:p),"
                "w:tr/(w:tc/w:p,w:tc/w:p,w:tc/w:p))"
            ),
        )
        table = Table(tbl, document_)

        table.delete_column(1)

        assert len(tbl.tblGrid.gridCol_lst) == 2
        for tr in tbl.tr_lst:
            assert len(tr.tc_lst) == 2

    def it_shrinks_gridSpan_for_a_cell_spanning_the_deleted_column(
        self, document_: Mock
    ):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol,w:gridCol),"
                "w:tr/w:tc/(w:tcPr/w:gridSpan{w:val=3},w:p))"
            ),
        )
        table = Table(tbl, document_)

        table.delete_column(1)

        # -- grid dropped by one, but the merged cell is preserved --
        assert len(tbl.tblGrid.gridCol_lst) == 2
        merged_tc = tbl.tr_lst[0].tc_lst[0]
        assert merged_tc.grid_span == 2

    def it_raises_on_delete_column_out_of_range(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/w:p))"
            ),
        )
        table = Table(tbl, document_)
        with pytest.raises(IndexError):
            table.delete_column(5)
        with pytest.raises(IndexError):
            table.delete_column(-3)

    def it_can_insert_a_row_at_an_index(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/w:p),"
                "w:tr/(w:tc/w:p,w:tc/w:p))"
            ),
        )
        table = Table(tbl, document_)

        new_row = table.insert_row(1)

        assert isinstance(new_row, _Row)
        assert len(tbl.tr_lst) == 3
        # -- the inserted row is at index 1 --
        assert tbl.tr_lst[1] is new_row._tr
        # -- it has one w:tc per gridCol --
        assert len(new_row._tr.tc_lst) == 2

    def it_can_append_via_insert_row_at_end(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p)"
            ),
        )
        table = Table(tbl, document_)

        new_row = table.insert_row(len(tbl.tr_lst))

        assert tbl.tr_lst[-1] is new_row._tr

    def it_raises_on_insert_row_out_of_range(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p)"
            ),
        )
        table = Table(tbl, document_)
        with pytest.raises(IndexError):
            table.insert_row(99)

    def it_can_add_a_row_from_a_source_row(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol),"
                'w:tr/(w:trPr/w:cantSplit,'
                'w:tc/w:p/w:r/w:t"hello",'
                "w:tc/w:p))"
            ),
        )
        table = Table(tbl, document_)
        source = _Row(tbl.tr_lst[0], table)

        new_row = table.add_row(source_row=source)

        assert isinstance(new_row, _Row)
        assert len(tbl.tr_lst) == 2
        # -- trPr/cantSplit preserved on clone --
        assert new_row._tr.trPr is not None
        # -- first cell's text cloned --
        assert new_row._tr.tc_lst[0].xpath("string(.)") == "hello"

    def it_strips_w_id_attributes_when_cloning_a_source_row(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p/w:bookmarkStart{w:id=7,w:name=B1})"
            ),
        )
        table = Table(tbl, document_)
        source = _Row(tbl.tr_lst[0], table)

        new_row = table.add_row(source_row=source)

        from docx.oxml.ns import qn
        # -- no w:id remains anywhere in the clone --
        for desc in new_row._tr.iter():
            assert desc.get(qn("w:id")) is None

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def _cells_(self, request: FixtureRequest):
        return property_mock(request, Table, "_cells")

    @pytest.fixture
    def _column_count_(self, request: FixtureRequest):
        return property_mock(request, Table, "_column_count")

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, Document)

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def part_prop_(self, request: FixtureRequest):
        return property_mock(request, Table, "part")

    @pytest.fixture
    def table(self, document_: Mock):
        tbl_cxml = "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol),w:tr/(w:tc,w:tc),w:tr/(w:tc,w:tc))"
        return Table(cast(CT_Tbl, element(tbl_cxml)), document_)


class Describe_Cell:
    """Unit-test suite for `docx.table._Cell` objects."""

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc", 1),
            ("w:tc/w:tcPr", 1),
            ("w:tc/w:tcPr/w:gridSpan{w:val=1}", 1),
            ("w:tc/w:tcPr/w:gridSpan{w:val=4}", 4),
        ],
    )
    def it_knows_its_grid_span(self, tc_cxml: str, expected_value: int, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.grid_span == expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_text"),
        [
            ("w:tc", ""),
            ('w:tc/w:p/w:r/w:t"foobar"', "foobar"),
            ('w:tc/(w:p/w:r/w:t"foo",w:p/w:r/w:t"bar")', "foo\nbar"),
            ('w:tc/(w:tcPr,w:p/w:r/w:t"foobar")', "foobar"),
            ('w:tc/w:p/w:r/(w:t"fo",w:tab,w:t"ob",w:br,w:t"ar",w:br)', "fo\tob\nar\n"),
        ],
    )
    def it_knows_what_text_it_contains(self, tc_cxml: str, expected_text: str, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        text = cell.text
        assert text == expected_text

    @pytest.mark.parametrize(
        ("tc_cxml", "new_text", "expected_cxml"),
        [
            ("w:tc/w:p", "foobar", 'w:tc/w:p/w:r/w:t"foobar"'),
            (
                "w:tc/w:p",
                "fo\tob\rar\n",
                'w:tc/w:p/w:r/(w:t"fo",w:tab,w:t"ob",w:br,w:t"ar",w:br)',
            ),
            (
                "w:tc/(w:tcPr, w:p, w:tbl, w:p)",
                "foobar",
                'w:tc/(w:tcPr, w:p/w:r/w:t"foobar")',
            ),
        ],
    )
    def it_can_replace_its_content_with_a_string_of_text(
        self, tc_cxml: str, new_text: str, expected_cxml: str, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        cell.text = new_text
        assert cell._tc.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:vAlign{w:val=bottom}", WD_ALIGN_VERTICAL.BOTTOM),
            ("w:tc/w:tcPr/w:vAlign{w:val=top}", WD_ALIGN_VERTICAL.TOP),
        ],
    )
    def it_knows_its_vertical_alignment(
        self, tc_cxml: str, expected_value: WD_ALIGN_VERTICAL | None, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.vertical_alignment == expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "new_value", "expected_cxml"),
        [
            ("w:tc", WD_ALIGN_VERTICAL.TOP, "w:tc/w:tcPr/w:vAlign{w:val=top}"),
            (
                "w:tc/w:tcPr",
                WD_ALIGN_VERTICAL.CENTER,
                "w:tc/w:tcPr/w:vAlign{w:val=center}",
            ),
            (
                "w:tc/w:tcPr/w:vAlign{w:val=center}",
                WD_ALIGN_VERTICAL.BOTTOM,
                "w:tc/w:tcPr/w:vAlign{w:val=bottom}",
            ),
            ("w:tc/w:tcPr/w:vAlign{w:val=center}", None, "w:tc/w:tcPr"),
            ("w:tc", None, "w:tc/w:tcPr"),
            ("w:tc/w:tcPr", None, "w:tc/w:tcPr"),
        ],
    )
    def it_can_change_its_vertical_alignment(
        self, tc_cxml: str, new_value: WD_ALIGN_VERTICAL | None, expected_cxml: str, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        cell.vertical_alignment = new_value
        assert cell._element.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:textDirection{w:val=lrTb}", WD_TEXT_DIRECTION.LR_TB),
            ("w:tc/w:tcPr/w:textDirection{w:val=tbRl}", WD_TEXT_DIRECTION.TB_RL),
            ("w:tc/w:tcPr/w:textDirection{w:val=btLr}", WD_TEXT_DIRECTION.BT_LR),
            ("w:tc/w:tcPr/w:textDirection{w:val=lrTbV}", WD_TEXT_DIRECTION.LR_TB_V),
            ("w:tc/w:tcPr/w:textDirection{w:val=tbRlV}", WD_TEXT_DIRECTION.TB_RL_V),
            ("w:tc/w:tcPr/w:textDirection{w:val=tbLrV}", WD_TEXT_DIRECTION.TB_LR_V),
        ],
    )
    def it_knows_its_text_direction(
        self, tc_cxml: str, expected_value: WD_TEXT_DIRECTION | None, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.text_direction == expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "new_value", "expected_cxml"),
        [
            (
                "w:tc",
                WD_TEXT_DIRECTION.TB_RL,
                "w:tc/w:tcPr/w:textDirection{w:val=tbRl}",
            ),
            (
                "w:tc/w:tcPr",
                WD_TEXT_DIRECTION.BT_LR,
                "w:tc/w:tcPr/w:textDirection{w:val=btLr}",
            ),
            (
                "w:tc/w:tcPr/w:textDirection{w:val=tbRl}",
                WD_TEXT_DIRECTION.BT_LR,
                "w:tc/w:tcPr/w:textDirection{w:val=btLr}",
            ),
            (
                "w:tc/w:tcPr/w:textDirection{w:val=tbRl}",
                None,
                "w:tc/w:tcPr",
            ),
            ("w:tc", None, "w:tc/w:tcPr"),
            ("w:tc/w:tcPr", None, "w:tc/w:tcPr"),
        ],
    )
    def it_can_change_its_text_direction(
        self,
        tc_cxml: str,
        new_value: WD_TEXT_DIRECTION | None,
        expected_cxml: str,
        parent_: Mock,
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        cell.text_direction = new_value
        assert cell._element.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:tcW{w:w=25%,w:type=pct}", None),
            ("w:tc/w:tcPr/w:tcW{w:w=1440,w:type=dxa}", 914400),
        ],
    )
    def it_knows_its_width_in_EMU(self, tc_cxml: str, expected_value: int | None, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.width == expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "new_value", "expected_cxml"),
        [
            ("w:tc", Inches(1), "w:tc/w:tcPr/w:tcW{w:w=1440,w:type=dxa}"),
            (
                "w:tc/w:tcPr/w:tcW{w:w=25%,w:type=pct}",
                Inches(2),
                "w:tc/w:tcPr/w:tcW{w:w=2880,w:type=dxa}",
            ),
        ],
    )
    def it_can_change_its_width(
        self, tc_cxml: str, new_value: Length, expected_cxml: str, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        cell.width = new_value
        assert cell.width == new_value
        assert cell._tc.xml == xml(expected_cxml)

    def it_provides_access_to_the_paragraphs_it_contains(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc/(w:p, w:p)")), parent_)

        paragraphs = cell.paragraphs

        # -- every w:p produces a Paragraph instance --
        assert len(paragraphs) == 2
        assert all(isinstance(p, Paragraph) for p in paragraphs)
        # -- the return value is iterable and indexable --
        assert all(p is paragraphs[idx] for idx, p in enumerate(paragraphs))

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_table_count"),
        [
            ("w:tc", 0),
            ("w:tc/w:tbl", 1),
            ("w:tc/(w:tbl,w:tbl)", 2),
            ("w:tc/(w:p,w:tbl)", 1),
            ("w:tc/(w:tbl,w:tbl,w:p)", 2),
        ],
    )
    def it_provides_access_to_the_tables_it_contains(
        self, tc_cxml: str, expected_table_count: int, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)

        tables = cell.tables

        # --- test len(), iterable, and indexed access
        assert len(tables) == expected_table_count
        assert all(isinstance(t, Table) for t in tables)
        assert all(t is tables[idx] for idx, t in enumerate(tables))

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_cxml"),
        [
            ("w:tc", "w:tc/w:p"),
            ("w:tc/w:p", "w:tc/(w:p, w:p)"),
            ("w:tc/w:tbl", "w:tc/(w:tbl, w:p)"),
        ],
    )
    def it_can_add_a_paragraph(self, tc_cxml: str, expected_cxml: str, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)

        p = cell.add_paragraph()

        assert isinstance(p, Paragraph)
        assert cell._tc.xml == xml(expected_cxml)

    def it_can_add_a_table(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc/w:p")), parent_)

        table = cell.add_table(rows=2, cols=2)

        assert isinstance(table, Table)
        assert cell._element.xml == snippet_seq("new-tbl")[1]

    def it_can_add_a_picture(self):
        # -- upstream#10, upstream-PR#429: _Cell.add_picture mirrors
        #    Paragraph.add_picture semantics via a newly-added paragraph.
        from docx import Document as OpenDocument
        from docx.shape import InlineShape

        document = OpenDocument()
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        shape = cell.add_picture("tests/test_files/python-icon.png")

        assert isinstance(shape, InlineShape)
        # -- the picture is carried on a newly-added paragraph --
        assert any(
            p._p.xpath(".//wp:inline") for p in cell.paragraphs
        )

    def it_accepts_a_PathLike_image_path_for_add_picture(self):
        from pathlib import Path

        from docx import Document as OpenDocument
        from docx.shape import InlineShape

        document = OpenDocument()
        cell = document.add_table(rows=1, cols=1).cell(0, 0)

        shape = cell.add_picture(Path("tests/test_files/python-icon.png"))

        assert isinstance(shape, InlineShape)

    def it_can_merge_itself_with_other_cells(
        self, tc_: Mock, tc_2_: Mock, parent_: Mock, merged_tc_: Mock
    ):
        cell, other_cell = _Cell(tc_, parent_), _Cell(tc_2_, parent_)
        tc_.merge.return_value = merged_tc_

        merged_cell = cell.merge(other_cell)

        assert isinstance(merged_cell, _Cell)
        tc_.merge.assert_called_once_with(other_cell._tc)
        assert merged_cell._tc is merged_tc_
        assert merged_cell._parent is cell._parent

    def it_provides_access_to_its_shading(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc")), parent_)
        shading = cell.shading
        assert isinstance(shading, CellShading)

    # -- is_merge_origin / merge_origin ----------------------------------------

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            # -- plain cell, no merge, no gridSpan --
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:gridSpan{w:val=1}", None),
            # -- horizontal-only merge (gridSpan>1, no vMerge) -> origin --
            ("w:tc/w:tcPr/w:gridSpan{w:val=2}", True),
            ("w:tc/w:tcPr/w:gridSpan{w:val=5}", True),
            # -- vMerge="restart" (with or without gridSpan) -> origin --
            ("w:tc/w:tcPr/w:vMerge{w:val=restart}", True),
            (
                "w:tc/w:tcPr/(w:gridSpan{w:val=2},w:vMerge{w:val=restart})",
                True,
            ),
            # -- vMerge="continue" -> continuation --
            ("w:tc/w:tcPr/w:vMerge", False),
            ("w:tc/w:tcPr/w:vMerge{w:val=continue}", False),
            (
                "w:tc/w:tcPr/(w:gridSpan{w:val=2},w:vMerge{w:val=continue})",
                False,
            ),
        ],
    )
    def it_knows_its_merge_origin_role(
        self, tc_cxml: str, expected_value: bool | None, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.is_merge_origin is expected_value

    def it_returns_self_as_merge_origin_when_not_merged(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc/w:p")), parent_)
        assert cell.merge_origin is cell

    def it_returns_self_as_merge_origin_for_horizontal_only_merge(self, parent_: Mock):
        cell = _Cell(
            cast(CT_Tc, element("w:tc/(w:tcPr/w:gridSpan{w:val=2},w:p)")),
            parent_,
        )
        assert cell.merge_origin is cell

    def it_returns_self_as_merge_origin_for_vMerge_restart(self, parent_: Mock):
        cell = _Cell(
            cast(CT_Tc, element("w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p)")),
            parent_,
        )
        assert cell.merge_origin is cell

    def it_walks_up_to_find_merge_origin_from_continuation(self, parent_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblGrid/w:gridCol,"
                "w:tr/w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p),"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p),"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p))"
            ),
        )
        root_tc = tbl.tr_lst[0].tc_lst[0]
        continuation_tc = tbl.tr_lst[2].tc_lst[0]
        cell = _Cell(continuation_tc, parent_)

        origin = cell.merge_origin

        assert origin._tc is root_tc

    def it_raises_when_merge_origin_is_orphaned(self, parent_: Mock):
        # -- vMerge="continue" in first row (no row above to walk to) --
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblGrid/w:gridCol,"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p))"
            ),
        )
        orphan_tc = tbl.tr_lst[0].tc_lst[0]
        cell = _Cell(orphan_tc, parent_)

        with pytest.raises(ValueError, match="orphan vMerge continuation"):
            cell.merge_origin

    # -- merged-cell read robustness edge cases --------------------------------

    def it_tolerates_orphan_vMerge_continuation_in_first_row(self, parent_: Mock):
        """Cell with vMerge="continue" in the top row (no restart above).

        ``_Row.cells`` should not crash; the orphan cell is surfaced as its
        own cell rather than delegating up a non-existent row.
        """
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/(w:tcPr/w:vMerge,w:p)))"
            ),
        )
        table = Table(tbl, parent_)

        # -- should not raise --
        cells = table.rows[0].cells

        assert len(cells) == 2

    def it_treats_malformed_gridSpan_zero_as_span_one(self, parent_: Mock):
        cell = _Cell(
            cast(CT_Tc, element("w:tc/(w:tcPr/w:gridSpan{w:val=0},w:p)")),
            parent_,
        )
        assert cell.grid_span == 1
        # -- grid_span==1 with no vMerge should be "not merged" --
        assert cell.is_merge_origin is None

    # --- tracked-change revisions ------------------------------------

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc/w:p", False),
            ("w:tc/(w:tcPr,w:p)", False),
            ("w:tc/(w:tcPr/w:cellIns{w:id=1,w:author=A},w:p)", True),
        ],
    )
    def it_knows_whether_it_is_a_tracked_insertion(
        self, tc_cxml: str, expected_value: bool, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.is_tracked_insertion is expected_value

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_value"),
        [
            ("w:tc/w:p", False),
            ("w:tc/(w:tcPr,w:p)", False),
            ("w:tc/(w:tcPr/w:cellDel{w:id=2,w:author=B},w:p)", True),
        ],
    )
    def it_knows_whether_it_is_a_tracked_deletion(
        self, tc_cxml: str, expected_value: bool, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element(tc_cxml)), parent_)
        assert cell.is_tracked_deletion is expected_value

    def it_returns_None_for_formatting_change_when_tcPr_absent(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc/w:p")), parent_)
        assert cell.formatting_change is None

    def it_returns_None_for_formatting_change_when_tcPrChange_absent(
        self, parent_: Mock
    ):
        cell = _Cell(cast(CT_Tc, element("w:tc/(w:tcPr,w:p)")), parent_)
        assert cell.formatting_change is None

    def it_exposes_its_formatting_change_when_tcPrChange_present(
        self, parent_: Mock
    ):
        cell = _Cell(
            cast(
                CT_Tc,
                element(
                    "w:tc/("
                    "w:tcPr/w:tcPrChange{w:id=1,w:author=A}/w:tcPr/w:vAlign{w:val=top}"
                    ",w:p)"
                ),
            ),
            parent_,
        )
        fc = cell.formatting_change
        assert fc is not None
        assert fc.author == "A"
        assert fc.old_properties is not None
        assert fc.old_properties.xpath("./w:vAlign")

    # -- Phase C: split() and add_table(style=) --------------------------

    def it_split_is_a_noop_for_unmerged_cell(self, parent_: Mock):
        cell = _Cell(cast(CT_Tc, element("w:tc/w:p")), parent_)

        result = cell.split()

        assert result == [cell]
        assert cell._tc.grid_span == 1

    def it_splits_a_horizontally_merged_cell_into_sibling_cells(
        self, parent_: Mock
    ):
        tr = element(
            "w:tr/w:tc/(w:tcPr/(w:tcW{w:type=dxa,w:w=3000},"
            "w:gridSpan{w:val=3}),w:p/w:r/w:t\"A\")"
        )
        merged_tc = cast(CT_Tc, tr[0])
        cell = _Cell(merged_tc, parent_)

        result = cell.split()

        assert len(result) == 3
        # -- gridSpan cleared on all 3 cells --
        assert all(c._tc.grid_span == 1 for c in result)
        # -- width redistributed evenly (twips 3000 / 3 = 1000, EMU 1000/1440 in) --
        widths = [c._tc.width for c in result]
        assert all(w is not None for w in widths)
        assert all(int(w) == int(widths[0]) for w in widths)
        # -- first cell retained original content, siblings are empty --
        assert result[0]._tc.xpath("string(.)") == "A"
        for sibling in result[1:]:
            assert sibling._tc.xpath("string(.)") == ""

    def it_clears_vMerge_on_split(self, parent_: Mock):
        cell = _Cell(
            cast(
                CT_Tc,
                element("w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p)"),
            ),
            parent_,
        )

        cell.split()

        tcPr = cell._tc.tcPr
        assert tcPr is not None
        assert tcPr.vMerge is None

    def it_can_add_a_nested_table_with_style(self, parent_: Mock, request: FixtureRequest):
        # -- set up a parent table whose `part.get_style_id` is mocked --
        from docx.parts.document import DocumentPart

        document_part_ = instance_mock(request, DocumentPart)
        document_part_.get_style_id.return_value = "MyStyle"
        parent_table = instance_mock(request, Table)
        parent_table.part = document_part_
        cell = _Cell(cast(CT_Tc, element("w:tc/w:p")), parent_table)

        table = cell.add_table(rows=1, cols=1, style="My Style")

        assert isinstance(table, Table)
        document_part_.get_style_id.assert_called_once()
        # -- tblStyle was written onto the new table --
        assert table._tbl.tblStyle_val == "MyStyle"

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def merged_tc_(self, request: FixtureRequest):
        return instance_mock(request, CT_Tc)

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def tc_(self, request: FixtureRequest):
        return instance_mock(request, CT_Tc)

    @pytest.fixture
    def tc_2_(self, request: FixtureRequest):
        return instance_mock(request, CT_Tc)


class DescribeCellShading:
    """Unit-test suite for `docx.table.CellShading` objects."""

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_color"),
        [
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:shd{w:fill=D9E2F3}", RGBColor(0xD9, 0xE2, 0xF3)),
            ("w:tc/w:tcPr/w:shd{w:val=clear}", None),
        ],
    )
    def it_can_get_the_fill_color(
        self, tc_cxml: str, expected_color: RGBColor | None
    ):
        tc = cast(CT_Tc, element(tc_cxml))
        shading = CellShading(tc)
        assert shading.fill_color == expected_color

    @pytest.mark.parametrize(
        ("tc_cxml", "new_color", "expected_cxml"),
        [
            (
                "w:tc",
                RGBColor(0xD9, 0xE2, 0xF3),
                "w:tc/w:tcPr/w:shd{w:val=clear,w:fill=D9E2F3}",
            ),
            (
                "w:tc/w:tcPr/w:shd{w:fill=FF0000}",
                RGBColor(0x00, 0x00, 0xFF),
                "w:tc/w:tcPr/w:shd{w:val=clear,w:fill=0000FF}",
            ),
            (
                "w:tc/w:tcPr/w:shd{w:val=clear,w:fill=D9E2F3}",
                None,
                "w:tc/w:tcPr/w:shd{w:val=clear}",
            ),
            ("w:tc", None, "w:tc"),
        ],
    )
    def it_can_set_the_fill_color(
        self, tc_cxml: str, new_color: RGBColor | None, expected_cxml: str
    ):
        tc = cast(CT_Tc, element(tc_cxml))
        shading = CellShading(tc)
        shading.fill_color = new_color
        assert tc.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tc_cxml", "expected_pattern"),
        [
            ("w:tc", None),
            ("w:tc/w:tcPr", None),
            ("w:tc/w:tcPr/w:shd{w:val=clear}", WD_SHADING_PATTERN.CLEAR),
            ("w:tc/w:tcPr/w:shd{w:val=solid}", WD_SHADING_PATTERN.SOLID),
        ],
    )
    def it_can_get_the_pattern(
        self, tc_cxml: str, expected_pattern: WD_SHADING_PATTERN | None
    ):
        tc = cast(CT_Tc, element(tc_cxml))
        shading = CellShading(tc)
        assert shading.pattern == expected_pattern

    @pytest.mark.parametrize(
        ("tc_cxml", "new_pattern", "expected_cxml"),
        [
            (
                "w:tc",
                WD_SHADING_PATTERN.CLEAR,
                "w:tc/w:tcPr/w:shd{w:val=clear}",
            ),
            (
                "w:tc/w:tcPr/w:shd{w:val=clear}",
                WD_SHADING_PATTERN.SOLID,
                "w:tc/w:tcPr/w:shd{w:val=solid}",
            ),
            (
                "w:tc/w:tcPr/w:shd{w:val=clear,w:fill=D9E2F3}",
                None,
                "w:tc/w:tcPr/w:shd{w:fill=D9E2F3}",
            ),
            ("w:tc", None, "w:tc"),
        ],
    )
    def it_can_set_the_pattern(
        self,
        tc_cxml: str,
        new_pattern: WD_SHADING_PATTERN | None,
        expected_cxml: str,
    ):
        tc = cast(CT_Tc, element(tc_cxml))
        shading = CellShading(tc)
        shading.pattern = new_pattern
        assert tc.xml == xml(expected_cxml)


class DescribeTableBorders:
    """Unit-test suite for `docx.table.TableBorders` objects."""

    def it_provides_access_to_table_borders(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        borders = table.borders
        assert isinstance(borders, TableBorders)

    def it_can_get_border_properties_when_no_borders_exist(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        borders = table.borders
        assert borders.top.style is None
        assert borders.bottom.style is None
        assert borders.left.style is None
        assert borders.right.style is None
        assert borders.inside_h.style is None
        assert borders.inside_v.style is None

    def it_can_set_a_table_border(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        borders = table.borders
        borders.top.style = WD_BORDER_STYLE.SINGLE
        borders.top.width = Pt(1)
        borders.top.color = RGBColor(0, 0, 0)
        assert borders.top.style == WD_BORDER_STYLE.SINGLE
        assert borders.top.color == RGBColor(0, 0, 0)

    def it_can_use_set_borders_convenience(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        table.set_borders(top=True, bottom=True, inside_h=True)
        borders = table.borders
        assert borders.top.style == WD_BORDER_STYLE.SINGLE
        assert borders.bottom.style == WD_BORDER_STYLE.SINGLE
        assert borders.inside_h.style == WD_BORDER_STYLE.SINGLE
        assert borders.left.style == WD_BORDER_STYLE.NONE
        assert borders.right.style == WD_BORDER_STYLE.NONE
        assert borders.inside_v.style == WD_BORDER_STYLE.NONE

    def it_can_set_borders_with_custom_style(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        table.set_borders(
            top=True,
            bottom=True,
            style=WD_BORDER_STYLE.DOUBLE,
            width=Pt(2),
            color=RGBColor(0xFF, 0, 0),
        )
        borders = table.borders
        assert borders.top.style == WD_BORDER_STYLE.DOUBLE
        assert borders.top.color == RGBColor(0xFF, 0, 0)
        assert borders.bottom.style == WD_BORDER_STYLE.DOUBLE

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, Document)


class DescribeTableStyleFlags:
    """Unit-test suite for `docx.table.TableStyleFlags` objects."""

    def it_is_accessed_via_Table_style_flags(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        assert isinstance(table.style_flags, TableStyleFlags)

    def it_returns_False_for_every_flag_when_tblLook_is_absent(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        flags = table.style_flags
        assert flags.first_row is False
        assert flags.last_row is False
        assert flags.first_column is False
        assert flags.last_column is False
        assert flags.no_horizontal_banding is False
        assert flags.no_vertical_banding is False

    def it_does_not_create_tblLook_when_only_reading(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        _ = table.style_flags.first_row
        assert tbl.tblPr.tblLook is None

    def it_creates_tblLook_on_first_write(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        table.style_flags.first_row = True
        assert tbl.tblPr.tblLook is not None
        assert tbl.tblPr.tblLook.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstRow"
        ) == "1"

    def it_writes_False_as_explicit_0(self, document_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        table.style_flags.first_row = False
        tblLook = tbl.tblPr.tblLook
        assert tblLook is not None
        assert tblLook.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}firstRow"
        ) == "0"

    @pytest.mark.parametrize(
        ("attr", "xml_attr"),
        [
            ("first_row", "firstRow"),
            ("last_row", "lastRow"),
            ("first_column", "firstColumn"),
            ("last_column", "lastColumn"),
            ("no_horizontal_banding", "noHBand"),
            ("no_vertical_banding", "noVBand"),
        ],
    )
    def it_round_trips_each_flag(
        self, attr: str, xml_attr: str, document_: Mock
    ):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr"))
        table = Table(tbl, document_)
        flags = table.style_flags
        setattr(flags, attr, True)
        assert getattr(flags, attr) is True
        setattr(flags, attr, False)
        assert getattr(flags, attr) is False

    def it_reads_flags_from_existing_tblLook(self, document_: Mock):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/w:tblPr/w:tblLook{w:firstRow=1,w:firstColumn=1,w:noVBand=1}"
            ),
        )
        table = Table(tbl, document_)
        flags = table.style_flags
        assert flags.first_row is True
        assert flags.first_column is True
        assert flags.no_vertical_banding is True
        assert flags.last_row is False
        assert flags.no_horizontal_banding is False

    def it_reads_legacy_val_bitmask_when_individual_attrs_absent(
        self, document_: Mock
    ):
        # 0x04A0 = firstRow | firstColumn | noVBand
        tbl = cast(
            CT_Tbl, element("w:tbl/w:tblPr/w:tblLook{w:val=04A0}")
        )
        table = Table(tbl, document_)
        flags = table.style_flags
        assert flags.first_row is True
        assert flags.first_column is True
        assert flags.no_vertical_banding is True
        assert flags.last_row is False

    def it_individual_attrs_override_legacy_val(self, document_: Mock):
        # val bitmask says firstRow=1 (0x0020) but explicit w:firstRow=0
        tbl = cast(
            CT_Tbl,
            element("w:tbl/w:tblPr/w:tblLook{w:val=0020,w:firstRow=0}"),
        )
        table = Table(tbl, document_)
        assert table.style_flags.first_row is False

    def it_setting_a_flag_toggles_an_existing_value(self, document_: Mock):
        tbl = cast(
            CT_Tbl, element("w:tbl/w:tblPr/w:tblLook{w:firstRow=1}")
        )
        table = Table(tbl, document_)
        flags = table.style_flags
        assert flags.first_row is True
        flags.first_row = False
        assert flags.first_row is False

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, Document)


class DescribeCellBorders:
    """Unit-test suite for `docx.table.CellBorders` objects."""

    def it_provides_access_to_cell_borders(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        borders = cell.borders
        assert isinstance(borders, CellBorders)

    def it_can_get_border_properties_when_no_borders_exist(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        borders = cell.borders
        assert borders.top.style is None
        assert borders.bottom.style is None
        assert borders.left.style is None
        assert borders.right.style is None

    def it_can_set_a_cell_border(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        borders = cell.borders
        borders.top.style = WD_BORDER_STYLE.SINGLE
        borders.top.width = Pt(0.5)
        borders.top.color = RGBColor(0, 0, 0)
        # re-read to verify
        borders2 = cell.borders
        assert borders2.top.style == WD_BORDER_STYLE.SINGLE
        assert borders2.top.color == RGBColor(0, 0, 0)

    def it_can_clear_a_cell_border(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        borders = cell.borders
        borders.top.style = WD_BORDER_STYLE.SINGLE
        borders.top.width = Pt(1)
        # now clear
        borders2 = cell.borders
        borders2.top.style = None
        borders2.top.width = None
        borders3 = cell.borders
        assert borders3.top.style is None
        assert borders3.top.width is None

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)


class DescribeCellMargins:
    """Unit-test suite for `docx.table.CellMargins` objects."""

    def it_is_accessible_via_cell_margins(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        margins = cell.margins
        assert isinstance(margins, CellMargins)

    def it_returns_None_for_every_edge_when_no_tcMar_present(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        margins = cell.margins
        assert margins.top is None
        assert margins.bottom is None
        assert margins.start is None
        assert margins.end is None

    def it_does_not_create_any_xml_on_a_pure_read(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        # -- access all edges as reads --
        m = cell.margins
        _ = (m.top, m.bottom, m.start, m.end)
        # -- no tcPr or tcMar should have been created --
        assert cell._tc.xml == xml("w:tc")

    @pytest.mark.parametrize(
        ("edge", "value"),
        [
            ("top", Inches(0.1)),
            ("bottom", Pt(6)),
            ("start", Twips(120)),
            ("end", Inches(0.25)),
        ],
    )
    def it_round_trips_each_edge(self, edge: str, value: Length, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        setattr(cell.margins, edge, value)
        # -- re-read via a fresh proxy --
        assert getattr(cell.margins, edge) == value

    def it_reads_start_from_legacy_w_left(self, parent_: Mock):
        tc = cast(
            CT_Tc,
            element("w:tc/w:tcPr/w:tcMar/w:left{w:w=120,w:type=dxa}"),
        )
        cell = _Cell(tc, parent_)
        assert cell.margins.start == Twips(120)

    def it_reads_end_from_legacy_w_right(self, parent_: Mock):
        tc = cast(
            CT_Tc,
            element("w:tc/w:tcPr/w:tcMar/w:right{w:w=200,w:type=dxa}"),
        )
        cell = _Cell(tc, parent_)
        assert cell.margins.end == Twips(200)

    def it_creates_tcPr_and_tcMar_on_demand(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        cell.margins.top = Twips(100)
        expected = xml("w:tc/w:tcPr/w:tcMar/w:top{w:w=100,w:type=dxa}")
        assert cell._tc.xml == expected

    def it_writes_only_provided_edges_via_set_margins(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        cell.set_margins(top=Twips(10), end=Twips(40))
        expected = xml(
            "w:tc/w:tcPr/w:tcMar/(w:top{w:w=10,w:type=dxa},"
            "w:end{w:w=40,w:type=dxa})"
        )
        assert cell._tc.xml == expected

    def it_leaves_unmentioned_edges_alone_on_set_margins(self, parent_: Mock):
        tc = cast(
            CT_Tc,
            element(
                "w:tc/w:tcPr/w:tcMar/(w:top{w:w=100,w:type=dxa},"
                "w:bottom{w:w=200,w:type=dxa})"
            ),
        )
        cell = _Cell(tc, parent_)
        cell.set_margins(start=Twips(50))
        assert cell.margins.top == Twips(100)
        assert cell.margins.bottom == Twips(200)
        assert cell.margins.start == Twips(50)

    def it_can_clear_a_single_edge_with_None(self, parent_: Mock):
        tc = cast(
            CT_Tc,
            element(
                "w:tc/w:tcPr/w:tcMar/(w:top{w:w=100,w:type=dxa},"
                "w:bottom{w:w=200,w:type=dxa})"
            ),
        )
        cell = _Cell(tc, parent_)
        cell.margins.top = None
        assert cell.margins.top is None
        assert cell.margins.bottom == Twips(200)

    def it_removes_empty_tcMar_when_last_edge_cleared(self, parent_: Mock):
        tc = cast(
            CT_Tc,
            element("w:tc/w:tcPr/w:tcMar/w:top{w:w=100,w:type=dxa}"),
        )
        cell = _Cell(tc, parent_)
        cell.margins.top = None
        # -- empty tcMar should have been pruned, leaving an empty tcPr --
        assert cell._tc.xml == xml("w:tc/w:tcPr")

    def it_can_remove_all_margins(self, parent_: Mock):
        tc = cast(
            CT_Tc,
            element(
                "w:tc/w:tcPr/w:tcMar/(w:top{w:w=10,w:type=dxa},"
                "w:bottom{w:w=20,w:type=dxa})"
            ),
        )
        cell = _Cell(tc, parent_)
        cell.remove_margins()
        assert cell._tc.xml == xml("w:tc/w:tcPr")

    def it_remove_margins_is_a_no_op_without_tcPr(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        cell.remove_margins()
        assert cell._tc.xml == xml("w:tc")

    def it_set_margins_with_no_args_is_a_no_op(self, parent_: Mock):
        tc = cast(CT_Tc, element("w:tc"))
        cell = _Cell(tc, parent_)
        cell.set_margins()
        assert cell._tc.xml == xml("w:tc")

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)


class DescribeBorderElement:
    """Unit-test suite for `docx.table.BorderElement` objects."""

    def it_can_get_the_style(self):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblBorders/w:top{w:val=single}"))
        tblBorders = tbl.tblPr.tblBorders
        border_el = BorderElement(tblBorders.top, tblBorders.get_or_add_top)
        assert border_el.style == WD_BORDER_STYLE.SINGLE

    def it_can_get_the_width(self):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblPr/w:tblBorders/w:top{w:val=single,w:sz=8}"))
        tblBorders = tbl.tblPr.tblBorders
        border_el = BorderElement(tblBorders.top, tblBorders.get_or_add_top)
        # sz=8 means 8 eighths of a point = 1 point
        assert border_el.width == Pt(1)

    def it_can_get_the_color(self):
        tbl = cast(
            CT_Tbl, element("w:tbl/w:tblPr/w:tblBorders/w:top{w:val=single,w:color=FF0000}")
        )
        tblBorders = tbl.tblPr.tblBorders
        border_el = BorderElement(tblBorders.top, tblBorders.get_or_add_top)
        assert border_el.color == RGBColor(0xFF, 0, 0)

    def it_can_get_the_space(self):
        tbl = cast(
            CT_Tbl, element("w:tbl/w:tblPr/w:tblBorders/w:top{w:val=single,w:space=4}")
        )
        tblBorders = tbl.tblPr.tblBorders
        border_el = BorderElement(tblBorders.top, tblBorders.get_or_add_top)
        # `space` is stored as whole points and exposed as a Length (EMU).
        assert border_el.space == Pt(4)

    def it_returns_none_when_no_border_element(self):
        border_el = BorderElement(None, lambda: None)  # type: ignore
        assert border_el.style is None
        assert border_el.width is None
        assert border_el.color is None
        assert border_el.space is None


class Describe_Column:
    """Unit-test suite for `docx.table._Cell` objects."""

    def it_provides_access_to_its_cells(self, _index_prop_: Mock, table_prop_: Mock, table_: Mock):
        table_prop_.return_value = table_
        _index_prop_.return_value = 4
        column = _Column(cast(CT_TblGridCol, element("w:gridCol{w:w=500}")), table_)
        table_.column_cells.return_value = [3, 2, 1]

        cells = column.cells

        table_.column_cells.assert_called_once_with(4)
        assert cells == (3, 2, 1)

    def it_provides_access_to_the_table_it_belongs_to(self, table_: Mock):
        table_.table = table_
        column = _Column(cast(CT_TblGridCol, element("w:gridCol{w:w=500}")), table_)

        assert column.table is table_

    @pytest.mark.parametrize(
        ("gridCol_cxml", "expected_width"),
        [
            ("w:gridCol{w:w=4242}", 2693670),
            ("w:gridCol{w:w=1440}", 914400),
            ("w:gridCol{w:w=2.54cm}", 914400),
            ("w:gridCol{w:w=54mm}", 1944000),
            ("w:gridCol{w:w=12.5pt}", 158750),
            ("w:gridCol", None),
        ],
    )
    def it_knows_its_width_in_EMU(
        self, gridCol_cxml: str, expected_width: int | None, table_: Mock
    ):
        column = _Column(cast(CT_TblGridCol, element(gridCol_cxml)), table_)
        assert column.width == expected_width

    @pytest.mark.parametrize(
        ("gridCol_cxml", "new_value", "expected_cxml"),
        [
            ("w:gridCol", Emu(914400), "w:gridCol{w:w=1440}"),
            ("w:gridCol{w:w=4242}", Inches(0.5), "w:gridCol{w:w=720}"),
            ("w:gridCol{w:w=4242}", None, "w:gridCol"),
            ("w:gridCol", None, "w:gridCol"),
        ],
    )
    def it_can_change_its_width(
        self, gridCol_cxml: str, new_value: Length | None, expected_cxml: str, table_: Mock
    ):
        column = _Column(cast(CT_TblGridCol, element(gridCol_cxml)), table_)

        column.width = new_value

        assert column.width == new_value
        assert column._gridCol.xml == xml(expected_cxml)

    def it_knows_its_index_in_table_to_help(self, table_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol,w:gridCol)"))
        gridCol = tbl.tblGrid.gridCol_lst[1]
        column = _Column(gridCol, table_)
        assert column._index == 1

    def it_exposes_index_as_a_public_alias(self, table_: Mock):
        # -- upstream#112: public `index` alias for the legacy `_index` --
        tbl = cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol,w:gridCol)"))
        gridCol = tbl.tblGrid.gridCol_lst[2]
        column = _Column(gridCol, table_)
        assert column.index == 2
        assert column.index == column._index

    def it_propagates_width_changes_to_every_rows_cell(self, table_: Mock):
        tbl_cxml = (
            "w:tbl/("
            "w:tblPr,"
            "w:tblGrid/(w:gridCol{w:w=1000},w:gridCol{w:w=2000},w:gridCol{w:w=3000}),"
            "w:tr/(w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=1000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=2000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=3000},w:p)),"
            "w:tr/(w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=1000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=2000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=3000},w:p))"
            ")"
        )
        tbl = cast(CT_Tbl, element(tbl_cxml))
        gridCol = tbl.tblGrid.gridCol_lst[1]
        column = _Column(gridCol, table_)

        column.width = Inches(2)

        # -- middle gridCol updated --
        assert tbl.tblGrid.gridCol_lst[1].w == Inches(2)
        # -- every row's middle tc.tcW updated, and outer tcs left unchanged --
        for tr in tbl.tr_lst:
            widths = [tc.width for tc in tr.tc_lst]
            assert widths[0] == Inches(1000 / 1440)  # unchanged
            assert widths[1] == Inches(2)
            assert widths[2] == Inches(3000 / 1440)  # unchanged

    def it_removes_cell_widths_when_width_set_to_None(self, table_: Mock):
        tbl_cxml = (
            "w:tbl/("
            "w:tblPr,"
            "w:tblGrid/(w:gridCol{w:w=1000},w:gridCol{w:w=2000}),"
            "w:tr/(w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=1000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=2000},w:p)),"
            "w:tr/(w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=1000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=2000},w:p))"
            ")"
        )
        tbl = cast(CT_Tbl, element(tbl_cxml))
        gridCol = tbl.tblGrid.gridCol_lst[0]
        column = _Column(gridCol, table_)

        column.width = None

        assert tbl.tblGrid.gridCol_lst[0].w is None
        # -- every row's first tc has had its tcW removed --
        for tr in tbl.tr_lst:
            assert tr.tc_lst[0].width is None
            # -- second column untouched --
            assert tr.tc_lst[1].width == Inches(2000 / 1440)

    def it_does_not_clobber_merged_cell_widths_when_setting_width(self, table_: Mock):
        # -- cell spanning 2 columns (gridSpan=2) should be left alone --
        tbl_cxml = (
            "w:tbl/("
            "w:tblPr,"
            "w:tblGrid/(w:gridCol{w:w=1000},w:gridCol{w:w=2000}),"
            "w:tr/("
            "w:tc/(w:tcPr/(w:tcW{w:type=dxa,w:w=3000},w:gridSpan{w:val=2}),w:p)"
            "),"
            "w:tr/("
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=1000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=2000},w:p)"
            ")"
            ")"
        )
        tbl = cast(CT_Tbl, element(tbl_cxml))
        gridCol = tbl.tblGrid.gridCol_lst[0]
        column = _Column(gridCol, table_)

        column.width = Inches(1)

        # -- gridCol updated --
        assert tbl.tblGrid.gridCol_lst[0].w == Inches(1)
        # -- second row's non-merged first cell updated --
        assert tbl.tr_lst[1].tc_lst[0].width == Inches(1)
        # -- first row's merged cell untouched --
        assert tbl.tr_lst[0].tc_lst[0].width == Inches(3000 / 1440)

    def it_handles_rows_with_grid_before_offsets(self, table_: Mock):
        # -- row with gridBefore=1 so its single tc is at grid_offset=1 --
        tbl_cxml = (
            "w:tbl/("
            "w:tblPr,"
            "w:tblGrid/(w:gridCol{w:w=1000},w:gridCol{w:w=2000}),"
            "w:tr/("
            "w:trPr/w:gridBefore{w:val=1},"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=2000},w:p)"
            "),"
            "w:tr/("
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=1000},w:p),"
            "w:tc/(w:tcPr/w:tcW{w:type=dxa,w:w=2000},w:p)"
            ")"
            ")"
        )
        tbl = cast(CT_Tbl, element(tbl_cxml))
        gridCol = tbl.tblGrid.gridCol_lst[1]
        column = _Column(gridCol, table_)

        column.width = Inches(4)

        assert tbl.tblGrid.gridCol_lst[1].w == Inches(4)
        # -- the lone tc in the first row (at grid-offset=1) updated --
        assert tbl.tr_lst[0].tc_lst[0].width == Inches(4)
        # -- the second row's tc at grid-offset=1 updated --
        assert tbl.tr_lst[1].tc_lst[1].width == Inches(4)
        # -- the second row's tc at grid-offset=0 untouched --
        assert tbl.tr_lst[1].tc_lst[0].width == Inches(1000 / 1440)

    def it_can_delete_itself(self, request: FixtureRequest):
        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/(w:gridCol,w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/w:p,w:tc/w:p))"
            ),
        )
        table = Table(tbl, document_)
        column = table.columns[1]

        column.delete()

        assert len(tbl.tblGrid.gridCol_lst) == 2
        for tr in tbl.tr_lst:
            assert len(tr.tc_lst) == 2

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def _index_prop_(self, request: FixtureRequest):
        return property_mock(request, _Column, "_index")

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_prop_(self, request: FixtureRequest):
        return property_mock(request, _Column, "table")


class Describe_Columns:
    """Unit-test suite for `docx.table._Columns` objects."""

    def it_has_sequence_behaviors(self, table_: Mock):
        columns = _Columns(cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol)")), table_)

        # -- it supports len() --
        assert len(columns) == 2
        # -- it is iterable --
        assert len(tuple(c for c in columns)) == 2
        assert all(type(c) is _Column for c in columns)
        # -- it is indexable --
        assert all(type(columns[i]) is _Column for i in range(2))

    def it_raises_on_indexed_access_out_of_range(self, table_: Mock):
        columns = _Columns(cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol)")), table_)

        with pytest.raises(IndexError):
            columns[2]
        with pytest.raises(IndexError):
            columns[-3]

    def it_provides_access_to_the_table_it_belongs_to(self, table_: Mock):
        columns = _Columns(cast(CT_Tbl, element("w:tbl")), table_)
        table_.table = table_

        assert columns.table is table_

    def it_provides_sliced_access_to_columns(self, table_: Mock):
        # -- regression for upstream#770: `columns[1:]` previously passed the
        # -- slice straight to `_Column.__init__`, returning a single `_Column`
        # -- wrapping a list. It must now return a list of `_Column` objects. --
        columns = _Columns(
            cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol,w:gridCol)")),
            table_,
        )

        tail = columns[1:]

        assert isinstance(tail, list)
        assert len(tail) == 2
        assert all(isinstance(c, _Column) for c in tail)

    @pytest.mark.parametrize(
        ("start", "stop", "step", "expected_len"),
        [
            (0, 2, 1, 2),
            (1, None, 1, 2),
            (None, -1, 1, 2),
            (None, None, 2, 2),
            (5, 10, 1, 0),
        ],
    )
    def it_supports_slice_objects_on_getitem(
        self,
        start: int | None,
        stop: int | None,
        step: int | None,
        expected_len: int,
        table_: Mock,
    ):
        columns = _Columns(
            cast(CT_Tbl, element("w:tbl/w:tblGrid/(w:gridCol,w:gridCol,w:gridCol)")),
            table_,
        )

        result = columns[start:stop:step]

        assert isinstance(result, list)
        assert len(result) == expected_len

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)


class Describe_Row:
    """Unit-test suite for `docx.table._Row` objects."""

    def it_exposes_index_as_a_public_alias(self, parent_: Mock):
        # -- upstream#112: public `index` alias for the legacy `_index` --
        tbl = cast(CT_Tbl, element("w:tbl/(w:tblPr,w:tblGrid,w:tr,w:tr,w:tr)"))
        tr = tbl.tr_lst[2]
        row = _Row(cast(CT_Row, tr), parent_)
        assert row.index == 2
        assert row.index == row._index

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", True),
            ("w:tr/w:trPr", True),
            ("w:tr/w:trPr/w:cantSplit", False),
            ("w:tr/w:trPr/w:cantSplit{w:val=false}", True),
        ],
    )
    def it_knows_whether_it_allows_break_across_pages(
        self, tr_cxml: str, expected_value: bool, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.allow_break_across_pages is expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "new_value", "expected_cxml"),
        [
            ("w:tr", False, "w:tr/w:trPr/w:cantSplit"),
            ("w:tr/w:trPr/w:cantSplit", True, "w:tr/w:trPr"),
        ],
    )
    def it_can_change_whether_it_allows_break_across_pages(
        self, tr_cxml: str, new_value: bool, expected_cxml: str, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        row.allow_break_across_pages = new_value
        assert row._tr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", False),
            ("w:tr/w:trPr", False),
            ("w:tr/w:trPr/w:tblHeader", True),
            ("w:tr/w:trPr/w:tblHeader{w:val=false}", False),
        ],
    )
    def it_knows_whether_it_is_a_header_row(
        self, tr_cxml: str, expected_value: bool, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.is_header is expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "new_value", "expected_cxml"),
        [
            ("w:tr", True, "w:tr/w:trPr/w:tblHeader"),
            ("w:tr/w:trPr/w:tblHeader", False, "w:tr/w:trPr"),
        ],
    )
    def it_can_change_whether_it_is_a_header_row(
        self, tr_cxml: str, new_value: bool, expected_cxml: str, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        row.is_header = new_value
        assert row._tr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", 0),
            ("w:tr/w:trPr", 0),
            ("w:tr/w:trPr/w:gridAfter{w:val=0}", 0),
            ("w:tr/w:trPr/w:gridAfter{w:val=4}", 4),
        ],
    )
    def it_knows_its_grid_cols_after(self, tr_cxml: str, expected_value: int | None, parent_: Mock):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.grid_cols_after == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", 0),
            ("w:tr/w:trPr", 0),
            ("w:tr/w:trPr/w:gridBefore{w:val=0}", 0),
            ("w:tr/w:trPr/w:gridBefore{w:val=3}", 3),
        ],
    )
    def it_knows_its_grid_cols_before(
        self, tr_cxml: str, expected_value: int | None, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.grid_cols_before == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", None),
            ("w:tr/w:trPr", None),
            ("w:tr/w:trPr/w:trHeight", None),
            ("w:tr/w:trPr/w:trHeight{w:val=0}", 0),
            ("w:tr/w:trPr/w:trHeight{w:val=1440}", 914400),
        ],
    )
    def it_knows_its_height(self, tr_cxml: str, expected_value: int | None, parent_: Mock):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.height == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "new_value", "expected_cxml"),
        [
            ("w:tr", Inches(1), "w:tr/w:trPr/w:trHeight{w:val=1440}"),
            ("w:tr/w:trPr", Inches(1), "w:tr/w:trPr/w:trHeight{w:val=1440}"),
            ("w:tr/w:trPr/w:trHeight", Inches(1), "w:tr/w:trPr/w:trHeight{w:val=1440}"),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440}",
                Inches(2),
                "w:tr/w:trPr/w:trHeight{w:val=2880}",
            ),
            ("w:tr/w:trPr/w:trHeight{w:val=2880}", None, "w:tr/w:trPr/w:trHeight"),
            ("w:tr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr/w:trHeight", None, "w:tr/w:trPr/w:trHeight"),
        ],
    )
    def it_can_change_its_height(
        self, tr_cxml: str, new_value: Length | None, expected_cxml: str, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        row.height = new_value
        assert row._tr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tr_cxml", "expected_value"),
        [
            ("w:tr", None),
            ("w:tr/w:trPr", None),
            ("w:tr/w:trPr/w:trHeight{w:val=0, w:hRule=auto}", WD_ROW_HEIGHT.AUTO),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=atLeast}",
                WD_ROW_HEIGHT.AT_LEAST,
            ),
            (
                "w:tr/w:trPr/w:trHeight{w:val=2880, w:hRule=exact}",
                WD_ROW_HEIGHT.EXACTLY,
            ),
        ],
    )
    def it_knows_its_height_rule(
        self, tr_cxml: str, expected_value: WD_ROW_HEIGHT | None, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        assert row.height_rule == expected_value

    @pytest.mark.parametrize(
        ("tr_cxml", "new_value", "expected_cxml"),
        [
            ("w:tr", WD_ROW_HEIGHT.AUTO, "w:tr/w:trPr/w:trHeight{w:hRule=auto}"),
            (
                "w:tr/w:trPr",
                WD_ROW_HEIGHT.AT_LEAST,
                "w:tr/w:trPr/w:trHeight{w:hRule=atLeast}",
            ),
            (
                "w:tr/w:trPr/w:trHeight",
                WD_ROW_HEIGHT.EXACTLY,
                "w:tr/w:trPr/w:trHeight{w:hRule=exact}",
            ),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=exact}",
                WD_ROW_HEIGHT.AUTO,
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=auto}",
            ),
            (
                "w:tr/w:trPr/w:trHeight{w:val=1440, w:hRule=auto}",
                None,
                "w:tr/w:trPr/w:trHeight{w:val=1440}",
            ),
            ("w:tr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr", None, "w:tr/w:trPr"),
            ("w:tr/w:trPr/w:trHeight", None, "w:tr/w:trPr/w:trHeight"),
        ],
    )
    def it_can_change_its_height_rule(
        self, tr_cxml: str, new_value: WD_ROW_HEIGHT | None, expected_cxml: str, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element(tr_cxml)), parent_)
        row.height_rule = new_value
        assert row._tr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("tbl_cxml", "row_idx", "expected_len"),
        [
            # -- cell corresponds to single layout-grid cell --
            ("w:tbl/w:tr/w:tc/w:p", 0, 1),
            # -- cell has a horizontal span --
            ("w:tbl/w:tr/w:tc/(w:tcPr/w:gridSpan{w:val=2},w:p)", 0, 2),
            # -- cell is in latter row of vertical span --
            (
                "w:tbl/(w:tr/w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p),"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p))",
                1,
                1,
            ),
            # -- cell both has horizontal span and is latter row of vertical span --
            (
                "w:tbl/(w:tr/w:tc/(w:tcPr/(w:gridSpan{w:val=2},w:vMerge{w:val=restart}),w:p),"
                "w:tr/w:tc/(w:tcPr/(w:gridSpan{w:val=2},w:vMerge),w:p))",
                1,
                2,
            ),
        ],
    )
    def it_provides_access_to_its_cells(
        self, tbl_cxml: str, row_idx: int, expected_len: int, parent_: Mock
    ):
        tbl = cast(CT_Tbl, element(tbl_cxml))
        tr = tbl.tr_lst[row_idx]
        table = Table(tbl, parent_)
        row = _Row(tr, table)

        cells = row.cells

        assert len(cells) == expected_len
        assert all(type(c) is _Cell for c in cells)

    def it_resolves_continuation_cells_to_the_same_restart_origin(self, parent_: Mock):
        """vMerge="restart" followed by multiple continuations (no val) all
        yield cells that share the same underlying ``w:tc`` origin.
        """
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblGrid/w:gridCol,"
                "w:tr/w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p),"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p),"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p),"
                "w:tr/w:tc/(w:tcPr/w:vMerge,w:p))"
            ),
        )
        table = Table(tbl, parent_)
        root_tc = tbl.tr_lst[0].tc_lst[0]

        # -- each row's single cell should be backed by the restart tc --
        for row_idx in range(4):
            row = table.rows[row_idx]
            cells = row.cells
            assert len(cells) == 1
            assert cells[0]._tc is root_tc

    def it_tolerates_orphan_continuation_in_first_row_iteration(self, parent_: Mock):
        """First-row vMerge="continue" (orphan) should not crash ``_Row.cells``."""
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/(w:tcPr/w:vMerge,w:p)))"
            ),
        )
        table = Table(tbl, parent_)
        row = table.rows[0]

        # -- should not raise --
        cells = row.cells
        assert len(cells) == 2

    def it_resolves_vMerge_chain_ending_at_last_row(self, parent_: Mock):
        """Vertical merge whose continuation hits the final table row."""
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/(w:tc/w:p,w:tc/(w:tcPr/w:vMerge{w:val=restart},w:p)),"
                "w:tr/(w:tc/w:p,w:tc/(w:tcPr/w:vMerge,w:p)))"
            ),
        )
        table = Table(tbl, parent_)
        last_row_cells = table.rows[1].cells
        restart_tc = tbl.tr_lst[0].tc_lst[1]

        # -- last-row second cell delegates up to restart in row 0 --
        assert last_row_cells[1]._tc is restart_tc

    def it_resolves_rectangular_region_spanning_gridSpan_and_vMerge(
        self, parent_: Mock
    ):
        """Horizontal gridSpan combined with vertical merge."""
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblGrid/(w:gridCol,w:gridCol,w:gridCol),"
                "w:tr/(w:tc/(w:tcPr/(w:gridSpan{w:val=2},w:vMerge{w:val=restart}),w:p),"
                "w:tc/w:p),"
                "w:tr/(w:tc/(w:tcPr/(w:gridSpan{w:val=2},w:vMerge),w:p),"
                "w:tc/w:p))"
            ),
        )
        table = Table(tbl, parent_)
        row0 = table.rows[0].cells
        row1 = table.rows[1].cells
        # -- row 0: gridSpan=2 origin, then cell[2]; row should have 3 cells --
        assert len(row0) == 3
        # -- first two cells of row 0 share the spanned tc --
        assert row0[0]._tc is row0[1]._tc
        # -- row 1: continuation resolves to the same spanned tc in row 0 --
        assert len(row1) == 3
        assert row1[0]._tc is row0[0]._tc
        assert row1[1]._tc is row0[0]._tc

    def it_provides_access_to_the_table_it_belongs_to(self, parent_: Mock, table_: Mock):
        parent_.table = table_
        row = _Row(cast(CT_Row, element("w:tr")), parent_)
        assert row.table is table_

    def it_knows_its_index_in_table_to_help(self, parent_: Mock):
        tbl = element("w:tbl/(w:tr,w:tr,w:tr)")
        row = _Row(cast(CT_Row, tbl[1]), parent_)
        assert row._index == 1

    def it_returns_None_for_formatting_change_when_trPr_absent(self, parent_: Mock):
        row = _Row(cast(CT_Row, element("w:tr")), parent_)
        assert row.formatting_change is None

    def it_returns_None_for_formatting_change_when_trPrChange_absent(
        self, parent_: Mock
    ):
        row = _Row(cast(CT_Row, element("w:tr/w:trPr")), parent_)
        assert row.formatting_change is None

    def it_exposes_its_formatting_change_when_trPrChange_present(
        self, parent_: Mock
    ):
        row = _Row(
            cast(
                CT_Row,
                element(
                    "w:tr/w:trPr/w:trPrChange{w:id=1,w:author=A}/w:trPr/w:cantSplit"
                ),
            ),
            parent_,
        )
        fc = row.formatting_change
        assert fc is not None
        assert fc.author == "A"
        assert fc.old_properties is not None
        assert fc.old_properties.xpath("./w:cantSplit")

    # -- Phase C: insert_row_before/after, clone, add_cell/insert_cell ---

    def it_can_insert_a_row_before_itself(self, request: FixtureRequest):
        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p,w:tr/w:tc/w:p)"
            ),
        )
        table = Table(tbl, document_)
        ref_row = table.rows[1]

        new_row = ref_row.insert_row_before()

        assert len(tbl.tr_lst) == 3
        assert tbl.tr_lst[1] is new_row._tr
        assert tbl.tr_lst[2] is ref_row._tr

    def it_can_insert_a_row_after_itself(self, request: FixtureRequest):
        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p,w:tr/w:tc/w:p)"
            ),
        )
        table = Table(tbl, document_)
        ref_row = table.rows[0]

        new_row = ref_row.insert_row_after()

        assert len(tbl.tr_lst) == 3
        assert tbl.tr_lst[0] is ref_row._tr
        assert tbl.tr_lst[1] is new_row._tr

    def it_can_clone_itself(self, request: FixtureRequest):
        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                'w:tr/(w:trPr/w:cantSplit,w:tc/w:p/w:r/w:t"payload"))'
            ),
        )
        table = Table(tbl, document_)
        row = table.rows[0]

        clone = row.clone()

        assert len(tbl.tr_lst) == 2
        # -- clone is the row immediately after the source --
        assert tbl.tr_lst[1] is clone._tr
        # -- payload was deep-copied --
        assert clone._tr.tc_lst[0].xpath("string(.)") == "payload"
        # -- trPr/cantSplit preserved --
        assert clone._tr.trPr is not None

    def it_strips_w_id_when_cloning(self, request: FixtureRequest):
        from docx.oxml.ns import qn

        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p/w:bookmarkStart{w:id=99,w:name=BM})"
            ),
        )
        table = Table(tbl, document_)
        row = table.rows[0]

        clone = row.clone()

        for desc in clone._tr.iter():
            assert desc.get(qn("w:id")) is None

    def it_can_add_a_cell(self, request: FixtureRequest):
        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p)"
            ),
        )
        table = Table(tbl, document_)
        row = table.rows[0]

        cell = row.add_cell()

        assert isinstance(cell, _Cell)
        assert len(row._tr.tc_lst) == 2
        assert row._tr.tc_lst[-1] is cell._tc

    def it_can_insert_a_cell_at_index(self, request: FixtureRequest):
        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/(w:tc/w:p,w:tc/w:p))"
            ),
        )
        table = Table(tbl, document_)
        row = table.rows[0]

        cell = row.insert_cell(1)

        assert len(row._tr.tc_lst) == 3
        assert row._tr.tc_lst[1] is cell._tc

    def it_raises_on_insert_cell_out_of_range(self, request: FixtureRequest):
        document_ = instance_mock(request, Document)
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/w:p)"
            ),
        )
        table = Table(tbl, document_)
        row = table.rows[0]
        with pytest.raises(IndexError):
            row.insert_cell(99)

    # -- apply_shading (upstream#370) --------------------------------

    def it_applies_shading_to_every_cell_in_the_row(self, parent_: Mock):
        tr = cast(CT_Row, element("w:tr/(w:tc,w:tc,w:tc)"))
        row = _Row(tr, parent_)
        row.apply_shading(RGBColor(0xFF, 0x00, 0x00))
        for tc in tr.tc_lst:
            cell = _Cell(tc, parent_)
            assert cell.shading.fill_color == RGBColor(0xFF, 0x00, 0x00)

    def it_accepts_a_hex_string_for_apply_shading(self, parent_: Mock):
        tr = cast(CT_Row, element("w:tr/(w:tc,w:tc)"))
        row = _Row(tr, parent_)
        row.apply_shading("00FF00")
        for tc in tr.tc_lst:
            cell = _Cell(tc, parent_)
            assert cell.shading.fill_color == RGBColor(0x00, 0xFF, 0x00)

    def it_clears_shading_when_passed_None(self, parent_: Mock):
        tr = cast(
            CT_Row,
            element(
                "w:tr/("
                "w:tc/w:tcPr/w:shd{w:val=clear,w:fill=FF0000},"
                "w:tc/w:tcPr/w:shd{w:val=clear,w:fill=FF0000})"
            ),
        )
        row = _Row(tr, parent_)
        row.apply_shading(None)
        for tc in tr.tc_lst:
            cell = _Cell(tc, parent_)
            assert cell.shading.fill_color is None

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def _index_prop_(self, request: FixtureRequest):
        return property_mock(request, _Row, "_index")

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def table_prop_(self, request: FixtureRequest, table_: Mock):
        return property_mock(request, _Row, "table")


class Describe_Rows:
    """Unit-test suite for `docx.table._Rows` objects."""

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected_len"),
        [
            ("w:tbl", 0),
            ("w:tbl/w:tr", 1),
            ("w:tbl/(w:tr,w:tr)", 2),
            ("w:tbl/(w:tr,w:tr,w:tr)", 3),
        ],
    )
    def it_has_sequence_behaviors(self, tbl_cxml: str, expected_len: int, parent_: Mock):
        tbl = cast(CT_Tbl, element(tbl_cxml))
        table = Table(tbl, parent_)
        rows = _Rows(tbl, table)

        # -- it supports len() --
        assert len(rows) == expected_len
        # -- it is iterable --
        assert len(tuple(r for r in rows)) == expected_len
        assert all(type(r) is _Row for r in rows)
        # -- it is indexable --
        assert all(type(rows[i]) is _Row for i in range(expected_len))

    @pytest.mark.parametrize(
        ("tbl_cxml", "out_of_range_idx"),
        [
            ("w:tbl", 0),
            ("w:tbl", 1),
            ("w:tbl", -1),
            ("w:tbl/w:tr", 1),
            ("w:tbl/w:tr", -2),
            ("w:tbl/(w:tr,w:tr,w:tr)", 3),
            ("w:tbl/(w:tr,w:tr,w:tr)", -4),
        ],
    )
    def it_raises_on_indexed_access_out_of_range(
        self, tbl_cxml: str, out_of_range_idx: int, parent_: Mock
    ):
        rows = _Rows(cast(CT_Tbl, element(tbl_cxml)), parent_)

        with pytest.raises(IndexError, match="list index out of range"):
            rows[out_of_range_idx]

    @pytest.mark.parametrize(("start", "end", "expected_len"), [(1, 3, 2), (0, -1, 2)])
    def it_provides_sliced_access_to_rows(
        self, start: int, end: int, expected_len: int, parent_: Mock
    ):
        tbl = cast(CT_Tbl, element("w:tbl/(w:tr,w:tr,w:tr)"))
        rows = _Rows(tbl, parent_)

        slice_of_rows = rows[start:end]

        assert len(slice_of_rows) == expected_len
        for idx, row in enumerate(slice_of_rows):
            assert tbl.tr_lst.index(row._tr) == start + idx
            assert isinstance(row, _Row)

    def it_provides_access_to_the_table_it_belongs_to(self, parent_: Mock):
        tbl = cast(CT_Tbl, element("w:tbl"))
        table = Table(tbl, parent_)
        rows = _Rows(tbl, table)

        assert rows.table is table

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def parent_(self, request: FixtureRequest):
        return instance_mock(request, Document)


class DescribeTable_StableId:
    """Unit-test suite for the `stable_id` accessor on `Table`."""

    def _make_table(self, request: FixtureRequest, tbl: CT_Tbl) -> Table:
        parent_ = instance_mock(request, Document)
        return Table(tbl, parent_)

    def it_returns_a_16_character_hex_string(self, request: FixtureRequest):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                'w:tr/w:tc/w:p/w:r/w:t"alpha")'
            ),
        )
        table = self._make_table(request, tbl)
        result = table.stable_id

        assert isinstance(result, str)
        assert len(result) == 16
        assert all(c in "0123456789abcdef" for c in result)

    def it_returns_the_same_id_on_repeated_access(self, request: FixtureRequest):
        tbl = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                'w:tr/w:tc/w:p/w:r/w:t"alpha")'
            ),
        )
        table = self._make_table(request, tbl)
        assert table.stable_id == table.stable_id

    def it_returns_different_ids_for_tables_with_different_text(
        self, request: FixtureRequest
    ):
        t1 = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                'w:tr/w:tc/w:p/w:r/w:t"alpha")'
            ),
        )
        t2 = cast(
            CT_Tbl,
            element(
                "w:tbl/(w:tblPr,w:tblGrid/w:gridCol,"
                'w:tr/w:tc/w:p/w:r/w:t"beta")'
            ),
        )
        assert (
            self._make_table(request, t1).stable_id
            != self._make_table(request, t2).stable_id
        )

    def it_returns_different_ids_for_sibling_tables_with_same_content(
        self, request: FixtureRequest
    ):
        body = element(
            "w:body/("
            'w:tbl/(w:tblPr,w:tblGrid/w:gridCol,w:tr/w:tc/w:p/w:r/w:t"same"),'
            'w:tbl/(w:tblPr,w:tblGrid/w:gridCol,w:tr/w:tc/w:p/w:r/w:t"same")'
            ")"
        )
        tbl_a = self._make_table(request, cast(CT_Tbl, body[0]))
        tbl_b = self._make_table(request, cast(CT_Tbl, body[1]))
        assert tbl_a.stable_id != tbl_b.stable_id


class Describe_Cell_StableId:
    """Unit-test suite for the `stable_id` accessor on `_Cell`."""

    def _make_cell(self, request: FixtureRequest, tc: CT_Tc) -> _Cell:
        parent_ = instance_mock(request, Table)
        return _Cell(tc, parent_)

    def it_returns_a_16_character_hex_string(self, request: FixtureRequest):
        tc = cast(CT_Tc, element('w:tc/w:p/w:r/w:t"hello"'))
        cell = self._make_cell(request, tc)

        result = cell.stable_id

        assert isinstance(result, str)
        assert len(result) == 16
        assert all(c in "0123456789abcdef" for c in result)

    def it_returns_the_same_id_on_repeated_access(self, request: FixtureRequest):
        tc = cast(CT_Tc, element('w:tc/w:p/w:r/w:t"hello"'))
        cell = self._make_cell(request, tc)
        assert cell.stable_id == cell.stable_id

    def it_returns_different_ids_for_cells_with_different_text(
        self, request: FixtureRequest
    ):
        tc1 = cast(CT_Tc, element('w:tc/w:p/w:r/w:t"alpha"'))
        tc2 = cast(CT_Tc, element('w:tc/w:p/w:r/w:t"beta"'))
        assert (
            self._make_cell(request, tc1).stable_id
            != self._make_cell(request, tc2).stable_id
        )

    def it_returns_different_ids_for_sibling_cells_with_same_text(
        self, request: FixtureRequest
    ):
        tr = element(
            'w:tr/(w:tc/w:p/w:r/w:t"same",w:tc/w:p/w:r/w:t"same")'
        )
        cell_a = self._make_cell(request, cast(CT_Tc, tr[0]))
        cell_b = self._make_cell(request, cast(CT_Tc, tr[1]))
        assert cell_a.stable_id != cell_b.stable_id


class DescribeTable_spans_page_break:
    """Unit-test suite for `Table.spans_page_break` page-break detection."""

    @pytest.mark.parametrize(
        ("tbl_cxml", "expected"),
        [
            ("w:tbl/(w:tblPr,w:tblGrid)", False),
            ("w:tbl/(w:tblPr,w:tblGrid,w:tr/w:tc/w:p/w:r)", False),
            (
                "w:tbl/(w:tblPr,w:tblGrid,w:tr/w:tc/w:p/w:r/w:lastRenderedPageBreak)",
                True,
            ),
        ],
    )
    def it_detects_rendered_page_breaks(
        self, tbl_cxml: str, expected: bool, request: FixtureRequest
    ):
        tbl = cast(CT_Tbl, element(tbl_cxml))
        table = Table(tbl, instance_mock(request, DocumentPart))

        assert table.spans_page_break is expected
