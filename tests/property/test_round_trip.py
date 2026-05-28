"""Round-trip property tests for python-docx.

Each property generates a random valid input, drives the public
authoring API, saves the document to an in-memory buffer, reloads it,
and asserts the resulting document matches the input.

Pattern
-------

The strategies are deliberately narrow — they generate values that
the OOXML spec considers valid (no lone surrogates, no XML control
characters except ``\\t`` / ``\\n`` / ``\\r``). When a property test
fails, Hypothesis's shrinker reports the minimal failing input, which
is usually a 1-2 character payload exposing a real fidelity bug.

To extend: add a new ``@composite`` strategy near the others, then a
new ``def it_round_trips_<thing>`` test that consumes it. Keep
``max_examples`` at the default (100) unless the cost of authoring a
single example is high — table tests use a smaller bound because
each example builds an entire table.
"""

from __future__ import annotations

import io
from typing import Union

import pytest

hypothesis = pytest.importorskip("hypothesis")

from hypothesis import given, settings, strategies as st  # noqa: E402

import docx  # noqa: E402


# ---- strategies -----------------------------------------------------


# XML 1.0 forbids most C0 control codepoints (only \t, \n, \r are
# allowed) and lone surrogates. lxml refuses to set text containing
# any C0 control byte, including \t/\n/\r in some serialization paths,
# so we keep this alphabet conservative: printable Unicode minus
# surrogates and minus all controls.
_OOXML_TEXT_ALPHABET = st.characters(
    blacklist_categories=("Cs", "Cc"),
)


@st.composite
def docx_paragraph_text(draw: st.DrawFn) -> str:
    return draw(st.text(alphabet=_OOXML_TEXT_ALPHABET, min_size=0, max_size=200))


@st.composite
def docx_run_text(draw: st.DrawFn) -> str:
    # Runs are typically shorter; cap at 100 for faster shrinking.
    return draw(st.text(alphabet=_OOXML_TEXT_ALPHABET, min_size=0, max_size=100))


@st.composite
def docx_run_formatting(
    draw: st.DrawFn,
) -> "dict[str, Union[bool, None]]":
    """Generate a dict of bool/None flags for run formatting."""
    triboolean = st.one_of(st.booleans(), st.none())
    return {
        "bold": draw(triboolean),
        "italic": draw(triboolean),
        "underline": draw(triboolean),
    }


@st.composite
def docx_table_shape(
    draw: st.DrawFn,
) -> "tuple[int, int, list[list[str]]]":
    rows = draw(st.integers(min_value=1, max_value=4))
    cols = draw(st.integers(min_value=1, max_value=4))
    cells = [
        [
            draw(st.text(alphabet=_OOXML_TEXT_ALPHABET, min_size=0, max_size=20))
            for _ in range(cols)
        ]
        for _ in range(rows)
    ]
    return rows, cols, cells


# ---- helpers --------------------------------------------------------


def _round_trip(d: "docx.document.Document") -> "docx.document.Document":
    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return docx.Document(buf)


# ---- tests ----------------------------------------------------------


class DescribeParagraphRoundTrip:
    @given(docx_paragraph_text())
    def it_preserves_paragraph_text(self, text: str) -> None:
        d = docx.Document()
        d.add_paragraph(text)
        d2 = _round_trip(d)
        assert [p.text for p in d2.paragraphs] == [text]

    @given(st.lists(docx_paragraph_text(), min_size=0, max_size=8))
    def it_preserves_a_sequence_of_paragraphs(
        self, texts: "list[str]"
    ) -> None:
        d = docx.Document()
        for t in texts:
            d.add_paragraph(t)
        d2 = _round_trip(d)
        assert [p.text for p in d2.paragraphs] == texts


class DescribeRunRoundTrip:
    @given(docx_run_text(), docx_run_formatting())
    def it_preserves_run_bold_italic_underline(
        self, text: str, fmt: "dict[str, Union[bool, None]]"
    ) -> None:
        d = docx.Document()
        p = d.add_paragraph()
        run = p.add_run(text)
        run.bold = fmt["bold"]
        run.italic = fmt["italic"]
        run.underline = fmt["underline"]
        d2 = _round_trip(d)
        out = d2.paragraphs[0].runs[0]
        assert out.text == text
        assert out.bold == fmt["bold"]
        assert out.italic == fmt["italic"]
        assert out.underline == fmt["underline"]


class DescribeTableRoundTrip:
    @settings(max_examples=30)
    @given(docx_table_shape())
    def it_preserves_table_cells(
        self, shape: "tuple[int, int, list[list[str]]]"
    ) -> None:
        rows, cols, cells = shape
        d = docx.Document()
        table = d.add_table(rows=rows, cols=cols)
        for r_idx in range(rows):
            for c_idx in range(cols):
                table.cell(r_idx, c_idx).text = cells[r_idx][c_idx]

        d2 = _round_trip(d)
        out = d2.tables[0]
        assert len(out.rows) == rows
        assert len(out.columns) == cols
        for r_idx in range(rows):
            for c_idx in range(cols):
                assert out.cell(r_idx, c_idx).text == cells[r_idx][c_idx]
