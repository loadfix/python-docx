"""Property-based invariants for python-docx save+reopen.

Wave 6-F — property-based fuzzing.  The goal is *not* to surface new
bugs (wave 5 / wave 6-A handle real drift) but to lock down the basic
round-trip invariants of the authoring library so any future regression
is caught by CI.

Each property is a generator over a value space paired with a
save-and-reopen assertion.  Using hypothesis keeps the strategies
concise and lets us replay any failing example deterministically via
the built-in shrink/replay database.

Seeds and example counts are deliberately conservative so the suite
stays fast on CI (<15s total for the whole module on a cold venv).

Tests use the project's ``Describe`` / ``it_*`` BDD naming convention
so pytest picks them up under the stricter ``python_functions`` filter
in ``pyproject.toml``.
"""

from __future__ import annotations

from io import BytesIO

from hypothesis import HealthCheck, given, seed, settings
from hypothesis import strategies as st

from docx import Document
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# Shared settings / helpers
# ---------------------------------------------------------------------------

# Save+reload is expensive.  50 examples per property is plenty to
# exercise the value space without blowing up CI time; 1s deadline is
# generous for cold-cache runs.
FAST = settings(
    max_examples=50,
    deadline=1000,
    suppress_health_check=[HealthCheck.too_slow, HealthCheck.data_too_large],
)


def _roundtrip_doc(doc: Document) -> Document:
    """Save ``doc`` to an in-memory buffer and reload."""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


# XML 1.0 forbids surrogates (``Cs``) and unassigned codepoints (``Cn``)
# can serialize but round-trip inconsistently; restrict to the safe
# intersection.  Control chars other than \t\n are also stripped by
# Word, so we include only printable text + the two whitespace chars.
XML_SAFE_TEXT = st.text(
    alphabet=st.characters(
        blacklist_categories=("Cs", "Cn", "Cc"),
        whitelist_characters="\t\n",
    ),
    max_size=500,
)


class DescribePropertyInvariants:
    """Round-trip invariants holding across a broad value space."""

    @seed(0xD0C1)
    @FAST
    @given(n=st.integers(min_value=0, max_value=200))
    def it_roundtrips_paragraph_count(self, n: int) -> None:
        """For all N ∈ [0, 200]: doc with N paragraphs reopens with N."""
        doc = Document()
        # A fresh Document() starts with a single empty paragraph we
        # did not author — delete it so our count is unambiguous.
        body = doc.element.body
        for p in list(body.iterchildren()):
            if p.tag.endswith("}p"):
                body.remove(p)
        for i in range(n):
            doc.add_paragraph(f"p{i}")

        reopened = _roundtrip_doc(doc)
        assert len(reopened.paragraphs) == n

    @seed(0xD0C2)
    @FAST
    @given(text=XML_SAFE_TEXT)
    def it_roundtrips_run_text(self, text: str) -> None:
        """Run text survives save+reopen byte-for-byte for XML-safe input."""
        doc = Document()
        doc.add_paragraph().add_run(text)

        reopened = _roundtrip_doc(doc)
        # The empty default paragraph sits at index 0; our run is in
        # the paragraph we appended.
        authored = reopened.paragraphs[-1]
        assert authored.text == text

    @seed(0xD0C3)
    @FAST
    @given(
        r=st.integers(min_value=0, max_value=255),
        g=st.integers(min_value=0, max_value=255),
        b=st.integers(min_value=0, max_value=255),
    )
    def it_roundtrips_font_color_rgb(self, r: int, g: int, b: int) -> None:
        """For all (R, G, B) ∈ [0, 255]^3: run.font.color.rgb round-trips."""
        doc = Document()
        run = doc.add_paragraph().add_run("x")
        run.font.color.rgb = RGBColor(r, g, b)

        reopened = _roundtrip_doc(doc)
        reloaded_run = reopened.paragraphs[-1].runs[0]
        assert reloaded_run.font.color.rgb == RGBColor(r, g, b)

    @seed(0xD0C4)
    @FAST
    @given(bold=st.booleans(), italic=st.booleans())
    def it_roundtrips_bold_and_italic(self, bold: bool, italic: bool) -> None:
        """For all (bold, italic) ∈ {T, F}^2: both flags round-trip."""
        doc = Document()
        run = doc.add_paragraph().add_run("x")
        run.bold = bold
        run.italic = italic

        reopened = _roundtrip_doc(doc)
        reloaded_run = reopened.paragraphs[-1].runs[0]
        # python-docx returns None when the property is unset; we set
        # it explicitly, so we expect True/False back (not a tri-state).
        assert bool(reloaded_run.bold) is bold
        assert bool(reloaded_run.italic) is italic

    @seed(0xD0C5)
    @FAST
    @given(half_pts=st.integers(min_value=1, max_value=2000))
    def it_roundtrips_font_size(self, half_pts: int) -> None:
        """Font size (half-point granularity) round-trips.

        Word stores sizes in half-points via ``<w:sz w:val="N"/>``.  We
        sweep [0.5pt, 1000pt] to catch truncation/clamping bugs.
        """
        pt = Pt(half_pts / 2)

        doc = Document()
        run = doc.add_paragraph().add_run("x")
        run.font.size = pt

        reopened = _roundtrip_doc(doc)
        reloaded_run = reopened.paragraphs[-1].runs[0]
        assert reloaded_run.font.size == pt
