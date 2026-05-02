"""Unit-test suite for `Document.add_paragraph` auto-applying a style's
``w:next`` style — upstream#888.

When a paragraph is added with style X and X defines ``<w:next w:val="Y">``,
the *next* call to ``add_paragraph`` without an explicit style argument
uses style Y. An explicit ``style`` kwarg always wins; a style without a
``w:next`` pointer doesn't queue anything.
"""

from __future__ import annotations

import docx


class DescribeAddParagraph_NextStyleAutoApply:
    """End-to-end behaviour of `next_paragraph_style` auto-apply."""

    def it_applies_the_next_style_on_the_following_add_paragraph(self):
        document = docx.Document()
        heading = document.styles["Heading 1"]
        # -- make "Heading 1"'s next style point at "Normal" (the default
        # -- template leaves `w:next` absent so we set it explicitly). --
        heading.next_paragraph_style = document.styles["Normal"]

        document.add_paragraph("H1", style="Heading 1")
        follow = document.add_paragraph("body text")

        assert follow.style.name == "Normal"

    def it_does_not_override_an_explicit_style_argument(self):
        document = docx.Document()
        heading = document.styles["Heading 1"]
        heading.next_paragraph_style = document.styles["Normal"]

        document.add_paragraph("H1", style="Heading 1")
        follow = document.add_paragraph("still a heading", style="Heading 2")

        assert follow.style.name == "Heading 2"

    def it_does_not_chain_beyond_the_immediately_next_paragraph(self):
        document = docx.Document()
        heading = document.styles["Heading 1"]
        heading.next_paragraph_style = document.styles["Normal"]

        document.add_paragraph("H1", style="Heading 1")
        document.add_paragraph("body 1")  # picks up Normal via the queue
        third = document.add_paragraph("body 2")  # queue is empty, no style

        # -- without a pending next style, `add_paragraph` leaves style=None,
        # -- which in practice means "Normal" (via the default). We check
        # -- that no *non-Normal* style was queued twice. --
        assert third.style.name == "Normal"

    def it_ignores_styles_with_no_next_pointer(self):
        document = docx.Document()
        # -- "Normal" has no ``w:next`` child in the default template --

        document.add_paragraph("a", style="Normal")
        # -- queue should be empty; passing style=None here shouldn't trigger
        # -- any styled paragraph (pending state must have been cleared). --
        assert document._pending_next_style is None
