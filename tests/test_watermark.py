# pyright: reportPrivateUsage=false

"""Unit-test suite for watermark support on Section."""

from __future__ import annotations

from typing import cast

import pytest

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.watermark import CT_VmlShape
from docx.shared import Inches, Pt, RGBColor
from docx.watermark import Watermark

from .unitutil.cxml import element
from .unitutil.file import test_file


class DescribeWatermark:
    """Unit-test suite for docx.watermark.Watermark."""

    def it_reports_text_type_when_textpath_is_present(self):
        shape_cxml = 'v:shape/v:textpath{string=HELLO}'
        shape = cast(CT_VmlShape, element(shape_cxml))

        watermark = Watermark(shape)

        assert watermark.type == "text"

    def it_reports_image_type_when_imagedata_is_present(self):
        shape_cxml = 'v:shape/v:imagedata{r:id=rId42}'
        shape = cast(CT_VmlShape, element(shape_cxml))

        watermark = Watermark(shape)

        assert watermark.type == "image"

    def it_returns_the_textpath_string_for_text_watermark(self):
        shape_cxml = 'v:shape/v:textpath{string=CONFIDENTIAL}'
        shape = cast(CT_VmlShape, element(shape_cxml))

        watermark = Watermark(shape)

        assert watermark.text == "CONFIDENTIAL"

    def it_returns_None_for_text_of_an_image_watermark(self):
        shape_cxml = 'v:shape/v:imagedata{r:id=rId42}'
        shape = cast(CT_VmlShape, element(shape_cxml))

        watermark = Watermark(shape)

        assert watermark.text is None


class DescribeSectionWatermarkAPI:
    """Integration-style tests for the Section watermark API, end-to-end."""

    def it_can_add_a_default_text_watermark(self):
        document = Document()
        section = document.sections[0]

        watermark = section.add_text_watermark("CONFIDENTIAL")

        assert watermark.type == "text"
        assert watermark.text == "CONFIDENTIAL"
        shape = self._shape(section)
        assert shape is not None
        assert shape.get("type") == "#_x0000_t136"
        fill = shape.find(qn("v:fill"))
        assert fill is not None
        assert fill.get("color") == "#C0C0C0"
        textpath = shape.find(qn("v:textpath"))
        assert textpath is not None
        assert 'Calibri' in textpath.get("style", "")
        assert '72' in textpath.get("style", "")
        assert "rotation:-45" in (shape.get("style") or "")

    def it_honors_custom_text_options(self):
        document = Document()
        section = document.sections[0]

        watermark = section.add_text_watermark(
            "DRAFT",
            font="Arial",
            size=Pt(48),
            color=RGBColor(0xFF, 0x00, 0x00),
            layout="horizontal",
        )

        assert watermark.text == "DRAFT"
        shape = self._shape(section)
        assert shape is not None
        fill = shape.find(qn("v:fill"))
        assert fill is not None
        assert fill.get("color") == "#FF0000"
        textpath = shape.find(qn("v:textpath"))
        assert textpath is not None
        assert 'Arial' in textpath.get("style", "")
        assert '48' in textpath.get("style", "")
        assert "rotation:0" in (shape.get("style") or "")

    def it_rejects_invalid_layout_values(self):
        document = Document()
        section = document.sections[0]

        with pytest.raises(ValueError, match="layout"):
            section.add_text_watermark("X", layout="circular")

    def it_replaces_an_existing_watermark(self):
        document = Document()
        section = document.sections[0]

        section.add_text_watermark("FIRST")
        section.add_text_watermark("SECOND")

        watermark = section.watermark
        assert watermark is not None
        assert watermark.text == "SECOND"

        # Only one watermark shape should exist.
        header_elm = section.header._element
        shapes = header_elm.findall(".//" + qn("v:shape"))
        assert len(shapes) == 1

    def it_can_add_an_image_watermark(self):
        document = Document()
        section = document.sections[0]

        watermark = section.add_image_watermark(
            test_file("monty-truth.png"), width=Inches(2)
        )

        assert watermark.type == "image"
        shape = self._shape(section)
        assert shape is not None
        assert shape.get("type") == "#_x0000_t75"
        imagedata = shape.find(qn("v:imagedata"))
        assert imagedata is not None
        rId = imagedata.get(qn("r:id"))
        assert rId and rId.startswith("rId")

        # The relationship should exist on the header part.
        header_part = section.header.part
        assert rId in header_part.rels

    def it_can_read_an_existing_watermark(self):
        document = Document()
        section = document.sections[0]

        section.add_text_watermark("SECRET")

        watermark = section.watermark
        assert watermark is not None
        assert watermark.text == "SECRET"

    def it_returns_None_for_watermark_when_none_present(self):
        document = Document()
        section = document.sections[0]

        # No header definition yet -> linked-to-previous -> no watermark.
        assert section.watermark is None

    def it_can_remove_a_watermark(self):
        document = Document()
        section = document.sections[0]

        section.add_text_watermark("REMOVE_ME")
        section.remove_watermark()

        assert section.watermark is None
        header_elm = section.header._element
        assert header_elm.find(".//" + qn("v:shape")) is None

    def it_ignores_remove_when_no_watermark_present(self):
        document = Document()
        section = document.sections[0]

        # Should not raise even though nothing was added.
        section.remove_watermark()

        assert section.watermark is None

    # -- fixtures/helpers -------------------------------------------------------------

    @staticmethod
    def _shape(section):
        header_elm = section.header._element
        return header_elm.find(".//" + qn("v:shape"))
