# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.drawing` module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.drawing import Drawing, GroupShape, Picture, WordprocessingShape
from docx.enum.shape import WD_DRAWING_TYPE, WD_SHAPE
from docx.image.image import Image
from docx.oxml.drawing import CT_Drawing, CT_GroupShape, CT_WordprocessingShape
from docx.parts.document import DocumentPart
from docx.parts.image import ImagePart
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element
from .unitutil.mock import FixtureRequest, Mock, instance_mock


class DescribeDrawing:
    """Unit-test suite for `docx.drawing.Drawing` objects."""

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic", True),
            ("w:drawing/wp:anchor/a:graphic/a:graphicData/pic:pic", True),
            ("w:drawing/wp:inline/a:graphic/a:graphicData/a:grpSp", False),
            ("w:drawing/wp:anchor/a:graphic/a:graphicData/a:chart", False),
        ],
    )
    def it_knows_when_it_contains_a_Picture(
        self, cxml: str, expected_value: bool, document_part_: Mock
    ):
        drawing = Drawing(cast(CT_Drawing, element(cxml)), document_part_)
        assert drawing.has_picture == expected_value

    def it_provides_access_to_the_image_in_a_Picture_drawing(
        self, document_part_: Mock, image_part_: Mock, image_: Mock
    ):
        image_part_.image = image_
        document_part_.part.related_parts = {"rId1": image_part_}
        cxml = (
            "w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip{r:embed=rId1}"
        )
        drawing = Drawing(cast(CT_Drawing, element(cxml)), document_part_)

        image = drawing.image

        assert image is image_

    def but_it_raises_when_the_drawing_does_not_contain_a_Picture(self, document_part_: Mock):
        drawing = Drawing(
            cast(CT_Drawing, element("w:drawing/wp:inline/a:graphic/a:graphicData/a:grpSp")),
            document_part_,
        )

        with pytest.raises(ValueError, match="drawing does not contain a picture"):
            drawing.image

    def it_provides_access_to_text_in_a_text_box(self, document_part_: Mock):
        cxml = (
            "w:drawing/wp:anchor/a:graphic/a:graphicData"
            '/wps:wsp/wps:txbx/w:txbxContent/(w:p/w:r/w:t"Hello",w:p/w:r/w:t"World")'
        )
        drawing = Drawing(cast(CT_Drawing, element(cxml)), document_part_)

        assert drawing.text == "Hello\nWorld"

    def it_returns_empty_text_when_no_text_frames(self, document_part_: Mock):
        drawing = Drawing(
            cast(CT_Drawing, element("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic")),
            document_part_,
        )

        assert drawing.text == ""

    def it_provides_access_to_paragraphs_in_a_text_box(self, document_part_: Mock):
        cxml = (
            "w:drawing/wp:anchor/a:graphic/a:graphicData"
            '/wps:wsp/wps:txbx/w:txbxContent/(w:p/w:r/w:t"Hello",w:p/w:r/w:t"World")'
        )
        drawing = Drawing(cast(CT_Drawing, element(cxml)), document_part_)

        paragraphs = drawing.paragraphs

        assert len(paragraphs) == 2
        assert all(isinstance(p, Paragraph) for p in paragraphs)
        assert paragraphs[0].text == "Hello"
        assert paragraphs[1].text == "World"

    def it_returns_empty_paragraphs_when_no_text_frames(self, document_part_: Mock):
        drawing = Drawing(
            cast(CT_Drawing, element("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic")),
            document_part_,
        )

        assert drawing.paragraphs == []

    @pytest.mark.parametrize(
        ("cxml", "expected_type"),
        [
            (
                "w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic",
                WD_DRAWING_TYPE.PICTURE,
            ),
            (
                "w:drawing/wp:anchor/a:graphic/a:graphicData/pic:pic",
                WD_DRAWING_TYPE.PICTURE,
            ),
            (
                "w:drawing/wp:inline/a:graphic/a:graphicData"
                "/wps:wsp/wps:txbx/w:txbxContent/w:p",
                WD_DRAWING_TYPE.TEXT_BOX,
            ),
            (
                "w:drawing/wp:inline/a:graphic/a:graphicData/c:chart",
                WD_DRAWING_TYPE.CHART,
            ),
            (
                "w:drawing/wp:inline/a:graphic/a:graphicData/wps:wsp",
                WD_DRAWING_TYPE.SHAPE,
            ),
            (
                "w:drawing/wp:inline/a:graphic/a:graphicData/wpg:wgp",
                WD_DRAWING_TYPE.GROUP,
            ),
            (
                "w:drawing/wp:inline/a:graphic/a:graphicData/dgm:relIds",
                WD_DRAWING_TYPE.DIAGRAM,
            ),
        ],
    )
    def it_knows_its_type(
        self, cxml: str, expected_type: WD_DRAWING_TYPE, document_part_: Mock
    ):
        drawing = Drawing(cast(CT_Drawing, element(cxml)), document_part_)

        assert drawing.type == expected_type

    @pytest.mark.parametrize(
        ("cxml", "expected"),
        [
            ("w:drawing/wp:inline/a:graphic/a:graphicData/c:chart", True),
            ("w:drawing/wp:anchor/a:graphic/a:graphicData/c:chart", True),
            ("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic", False),
            ("w:drawing", False),
        ],
    )
    def it_knows_when_it_has_a_chart(
        self, cxml: str, expected: bool, document_part_: Mock
    ):
        drawing = Drawing(cast(CT_Drawing, element(cxml)), document_part_)
        assert drawing.has_chart is expected

    def it_returns_None_chart_when_no_chart_ref(self, document_part_: Mock):
        drawing = Drawing(
            cast(CT_Drawing, element("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic")),
            document_part_,
        )
        assert drawing.chart is None

    def it_returns_None_chart_when_related_part_missing(self, document_part_: Mock):
        document_part_.part.related_parts = {}
        drawing = Drawing(
            cast(
                CT_Drawing,
                element(
                    "w:drawing/wp:inline/a:graphic/a:graphicData"
                    "/c:chart{r:id=rIdX}"
                ),
            ),
            document_part_,
        )
        assert drawing.chart is None

    def it_returns_a_Chart_when_related_part_is_a_ChartPart(self, document_part_: Mock):
        from docx.chart import Chart
        from docx.opc.constants import CONTENT_TYPE as CT
        from docx.opc.packuri import PackURI
        from docx.oxml.chart import CT_ChartSpace
        from docx.package import Package
        from docx.parts.chart import ChartPart

        chartSpace = cast(
            CT_ChartSpace,
            element(
                "c:chartSpace/c:chart/c:plotArea/c:barChart/c:barDir{val=bar}"
            ),
        )
        package = Package()
        chart_part = ChartPart(
            PackURI("/word/charts/chart1.xml"), CT.DML_CHART, chartSpace, package
        )
        document_part_.part.related_parts = {"rId9": chart_part}
        drawing = Drawing(
            cast(
                CT_Drawing,
                element(
                    "w:drawing/wp:inline/a:graphic/a:graphicData"
                    "/c:chart{r:id=rId9}"
                ),
            ),
            document_part_,
        )

        chart = drawing.chart
        assert isinstance(chart, Chart)

    def it_knows_it_is_not_a_group_when_it_wraps_a_picture(self, document_part_: Mock):
        drawing = Drawing(
            cast(CT_Drawing, element("w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic")),
            document_part_,
        )

        assert drawing.is_group is False
        assert drawing.group_shape is None
        assert drawing.group_shapes == []

    def it_knows_it_is_a_group_when_root_is_grpSp(self, document_part_: Mock):
        cxml = (
            "w:drawing/wp:inline/a:graphic/a:graphicData"
            "/wpg:grpSp/(wpg:nvGrpSpPr/wpg:cNvPr{id=1,name=Group 1},wps:wsp,wps:wsp)"
        )
        drawing = Drawing(cast(CT_Drawing, element(cxml)), document_part_)

        assert drawing.is_group is True

        group_shape = drawing.group_shape
        assert group_shape is not None
        assert isinstance(group_shape, GroupShape)
        assert group_shape.name == "Group 1"

        group_shapes = drawing.group_shapes
        assert len(group_shapes) == 1
        assert isinstance(group_shapes[0], GroupShape)

    def it_recognizes_legacy_wgp_as_a_group(self, document_part_: Mock):
        drawing = Drawing(
            cast(
                CT_Drawing,
                element("w:drawing/wp:inline/a:graphic/a:graphicData/wpg:wgp"),
            ),
            document_part_,
        )

        assert drawing.is_group is True
        assert drawing.group_shape is not None

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def image_(self, request: FixtureRequest):
        return instance_mock(request, Image)

    @pytest.fixture
    def image_part_(self, request: FixtureRequest):
        return instance_mock(request, ImagePart)


class DescribeGroupShape:
    """Unit-test suite for `docx.drawing.GroupShape` proxy."""

    def it_reads_name_from_the_grpSp_element(self, document_part_: Mock):
        grpSp = cast(
            CT_GroupShape,
            element("wpg:grpSp/wpg:nvGrpSpPr/wpg:cNvPr{id=1,name=My Group}"),
        )
        group = GroupShape(grpSp, document_part_)

        assert group.name == "My Group"

    def its_name_is_None_when_absent(self, document_part_: Mock):
        grpSp = cast(CT_GroupShape, element("wpg:grpSp"))
        group = GroupShape(grpSp, document_part_)

        assert group.name is None

    def it_provides_nested_shapes_in_document_order(self, document_part_: Mock):
        cxml = (
            "wpg:grpSp"
            "/(wpg:nvGrpSpPr/wpg:cNvPr{id=1,name=g},wps:wsp,pic:pic,wps:wsp)"
        )
        group = GroupShape(cast(CT_GroupShape, element(cxml)), document_part_)

        shapes = group.shapes

        assert len(shapes) == 3
        assert isinstance(shapes[0], WordprocessingShape)
        assert isinstance(shapes[1], Picture)
        assert isinstance(shapes[2], WordprocessingShape)

    def it_returns_nested_groups_as_GroupShape(self, document_part_: Mock):
        cxml = (
            "wpg:grpSp"
            "/(wpg:nvGrpSpPr/wpg:cNvPr{id=1,name=outer}"
            ",wpg:grpSp"
            "/(wpg:nvGrpSpPr/wpg:cNvPr{id=2,name=inner},wps:wsp)"
            ")"
        )
        group = GroupShape(cast(CT_GroupShape, element(cxml)), document_part_)

        shapes = group.shapes
        assert len(shapes) == 1
        inner = shapes[0]
        assert isinstance(inner, GroupShape)
        assert inner.name == "inner"

        inner_shapes = inner.shapes
        assert len(inner_shapes) == 1
        assert isinstance(inner_shapes[0], WordprocessingShape)

    def it_returns_empty_shapes_when_group_has_no_children(self, document_part_: Mock):
        grpSp = cast(
            CT_GroupShape,
            element("wpg:grpSp/wpg:nvGrpSpPr/wpg:cNvPr{id=1,name=g}"),
        )
        group = GroupShape(grpSp, document_part_)

        assert group.shapes == []

    def it_exposes_shape_text_for_textbox_shapes(self, document_part_: Mock):
        cxml = (
            "wpg:grpSp"
            "/(wpg:nvGrpSpPr/wpg:cNvPr{id=1,name=g}"
            ',wps:wsp/wps:txbx/w:txbxContent/(w:p/w:r/w:t"Hi",w:p/w:r/w:t"There"))'
        )
        group = GroupShape(cast(CT_GroupShape, element(cxml)), document_part_)

        shape = group.shapes[0]
        assert isinstance(shape, WordprocessingShape)
        assert shape.text == "Hi\nThere"

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeWordprocessingShape:
    """Unit-test suite for `docx.drawing.WordprocessingShape`."""

    def it_reads_name_from_wps_cNvPr(self, document_part_: Mock):
        wsp = cast(
            CT_WordprocessingShape,
            element("wps:wsp/wps:cNvPr{id=1,name=Rectangle 1}"),
        )

        shape = WordprocessingShape(wsp, document_part_)

        assert shape.name == "Rectangle 1"

    def its_name_is_None_when_absent(self, document_part_: Mock):
        wsp = cast(CT_WordprocessingShape, element("wps:wsp"))

        shape = WordprocessingShape(wsp, document_part_)

        assert shape.name is None

    @pytest.mark.parametrize(
        ("prst", "expected"),
        [
            ("rect", WD_SHAPE.RECTANGLE),
            ("roundRect", WD_SHAPE.ROUNDED_RECTANGLE),
            ("ellipse", WD_SHAPE.OVAL),
            ("rightArrow", WD_SHAPE.ARROW_RIGHT),
            ("wedgeRoundRectCallout", WD_SHAPE.CALLOUT_ROUNDED_RECTANGLE),
        ],
    )
    def it_infers_shape_type_from_prstGeom(
        self, prst: str, expected: WD_SHAPE, document_part_: Mock
    ):
        cxml = "wps:wsp/wps:spPr/a:prstGeom{prst=%s}" % prst
        wsp = cast(CT_WordprocessingShape, element(cxml))

        shape = WordprocessingShape(wsp, document_part_)

        assert shape.shape_type is expected

    def its_shape_type_is_None_when_preset_is_unknown(self, document_part_: Mock):
        wsp = cast(
            CT_WordprocessingShape,
            element("wps:wsp/wps:spPr/a:prstGeom{prst=someUnknownPreset}"),
        )

        shape = WordprocessingShape(wsp, document_part_)

        assert shape.shape_type is None

    def its_shape_type_is_None_when_prstGeom_absent(self, document_part_: Mock):
        wsp = cast(CT_WordprocessingShape, element("wps:wsp"))

        shape = WordprocessingShape(wsp, document_part_)

        assert shape.shape_type is None

    def it_can_set_text(self, document_part_: Mock):
        wsp = cast(CT_WordprocessingShape, element("wps:wsp"))
        shape = WordprocessingShape(wsp, document_part_)

        shape.text = "Hello"

        assert shape.text == "Hello"

    def it_replaces_existing_text_on_assignment(self, document_part_: Mock):
        wsp = cast(
            CT_WordprocessingShape,
            element('wps:wsp/wps:txbx/w:txbxContent/w:p/w:r/w:t"Old"'),
        )
        shape = WordprocessingShape(wsp, document_part_)

        shape.text = "New"

        assert shape.text == "New"

    def it_can_append_a_paragraph_to_its_text_frame(self, document_part_: Mock):
        wsp = cast(CT_WordprocessingShape, element("wps:wsp"))
        shape = WordprocessingShape(wsp, document_part_)

        paragraph = shape.add_paragraph("Hello")

        assert shape.text == "Hello"
        assert paragraph.text == "Hello"

    def it_can_append_paragraphs_to_an_existing_text_frame(
        self, document_part_: Mock
    ):
        wsp = cast(
            CT_WordprocessingShape,
            element('wps:wsp/wps:txbx/w:txbxContent/w:p/w:r/w:t"First"'),
        )
        shape = WordprocessingShape(wsp, document_part_)

        shape.add_paragraph("Second")

        assert shape.text == "First\nSecond"

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeCanvas:
    """Unit-test suite for `docx.drawing.Canvas`."""

    def it_lists_contained_shapes(self, request: FixtureRequest):
        from docx.drawing import Canvas
        from docx.oxml.drawing import CT_WordprocessingCanvas

        story_part_ = instance_mock(request, DocumentPart)

        class FakeParent:
            @property
            def part(self):
                return story_part_

        wpc = cast(
            CT_WordprocessingCanvas,
            element("wpc:wpc/(wps:wsp,wps:wsp)"),
        )
        canvas = Canvas(wpc, FakeParent())  # type: ignore[arg-type]

        assert len(canvas.shapes) == 2

    def it_can_add_a_shape(self, request: FixtureRequest):
        from docx.drawing import Canvas
        from docx.enum.shape import WD_SHAPE
        from docx.oxml.drawing import CT_WordprocessingCanvas
        from docx.shared import Inches

        story_part_ = instance_mock(request, DocumentPart)
        story_part_.next_id = 3

        class FakeParent:
            @property
            def part(self):
                return story_part_

        wpc = cast(CT_WordprocessingCanvas, element("wpc:wpc"))
        canvas = Canvas(wpc, FakeParent())  # type: ignore[arg-type]

        shape = canvas.add_shape(
            WD_SHAPE.ROUNDED_RECTANGLE, Inches(2), Inches(1), text="X"
        )

        assert shape.shape_type is WD_SHAPE.ROUNDED_RECTANGLE
        assert shape.text == "X"
        # -- wsp is a direct child of the wpc --
        assert len(canvas.shapes) == 1
        assert canvas.shapes[0].name == "Rounded Rectangle 3"

    def it_raises_when_shape_type_is_not_a_WD_SHAPE(self, request: FixtureRequest):
        from docx.drawing import Canvas
        from docx.oxml.drawing import CT_WordprocessingCanvas

        story_part_ = instance_mock(request, DocumentPart)

        class FakeParent:
            @property
            def part(self):
                return story_part_

        wpc = cast(CT_WordprocessingCanvas, element("wpc:wpc"))
        canvas = Canvas(wpc, FakeParent())  # type: ignore[arg-type]

        with pytest.raises(TypeError, match="WD_SHAPE"):
            canvas.add_shape("rect")  # type: ignore[arg-type]
