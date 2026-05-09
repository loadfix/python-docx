"""Unit-test suite for the ``docx.math`` proxy-layer re-export.

Exercises the integration between ``python-docx`` and the
``ooxml_math`` proxy layer (0.3.0): constructing equations with
``Fraction`` / ``Var`` / ``Lit`` / ``oMath``, attaching them to a
paragraph, and round-tripping through ``docx.save()`` / reload.
"""

from __future__ import annotations

import io

import pytest

from docx import Document
from docx.math import (
    Delimiter,
    Fraction,
    Lit,
    MathExpr,
    Matrix,
    Radical,
    Sub,
    Sum,
    Sup,
    Var,
    oMath,
)


class DescribeDocxMathProxyReexport:
    """docx.math re-exports the ooxml_math proxy surface verbatim."""

    def it_exposes_MathExpr_and_the_operator_tree_proxies(self):
        # -- smoke: every public name is importable --
        assert issubclass(Fraction, MathExpr)
        assert issubclass(Radical, MathExpr)
        assert issubclass(Sum, MathExpr)
        assert issubclass(Sub, MathExpr)
        assert issubclass(Sup, MathExpr)
        assert issubclass(Matrix, MathExpr)
        assert issubclass(Delimiter, MathExpr)

    def it_re_exports_from_the_same_underlying_classes_as_ooxml_math(self):
        import ooxml_math as shared

        assert Fraction is shared.Fraction
        assert oMath is shared.oMath
        assert Var is shared.Var
        assert Lit is shared.Lit


class DescribeParagraphAddEquationWithMathExpr:
    """Paragraph.add_equation() accepts ooxml_math proxies as well as XML."""

    def it_accepts_an_oMath_wrapped_proxy(self):
        doc = Document()
        p = doc.add_paragraph("eq: ")
        expr = oMath(Fraction(Var("x"), Lit(2)))

        equation = p.add_equation(expr)

        assert equation.text == "x2"
        assert equation.is_display_mode is False

    def it_auto_wraps_a_bare_operator_proxy_in_m_oMath(self):
        doc = Document()
        p = doc.add_paragraph("")

        equation = p.add_equation(Fraction(Var("a"), Lit(3)))

        # Bare Fraction got wrapped in an oMath root --
        assert equation.xml_element.tag.endswith("}oMath")
        assert equation.text == "a3"

    def it_emits_display_mode_when_requested_for_a_bare_oMath(self):
        doc = Document()
        p = doc.add_paragraph("")

        equation = p.add_equation(
            oMath(Fraction(Var("y"), Lit(4))), display_mode=True
        )

        assert equation.is_display_mode is True
        assert equation.text == "y4"

    def it_still_accepts_an_omml_xml_string(self):
        doc = Document()
        p = doc.add_paragraph("")
        xml = (
            '<m:oMath xmlns:m='
            '"http://schemas.openxmlformats.org/officeDocument/2006/math">'
            "<m:r><m:t>z</m:t></m:r></m:oMath>"
        )

        equation = p.add_equation(xml)

        assert equation.text == "z"


class DescribeEquationRoundTrip:
    """Equations built from MathExpr survive a save()/reload() cycle."""

    def it_round_trips_a_fraction_through_the_file_format(self):
        doc = Document()
        p = doc.add_paragraph("eq: ")
        p.add_equation(oMath(Fraction(Var("x"), Lit(2))))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        equations = reloaded.paragraphs[0].equations
        assert len(equations) == 1
        assert equations[0].text == "x2"
        assert reloaded.equations[0].text == "x2"

    def it_round_trips_a_nested_sum_over_a_fraction(self):
        doc = Document()
        p = doc.add_paragraph("")
        expr = oMath(
            Sum(
                body=Fraction(Var("x"), Lit(2)),
                lower=Var("i"),
                upper=Lit("n"),
            )
        )

        p.add_equation(expr)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)
        equations = reloaded.paragraphs[0].equations
        assert len(equations) == 1
        # -- text is a best-effort concatenation of every m:t --
        assert "x" in equations[0].text
        assert "2" in equations[0].text
        assert "i" in equations[0].text
        assert "n" in equations[0].text

    def it_round_trips_a_radical(self):
        doc = Document()
        p = doc.add_paragraph("")

        p.add_equation(oMath(Radical(Var("a"))))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)
        assert reloaded.paragraphs[0].equations[0].text == "a"


class DescribeAddEquationInvalidInput:
    """Paragraph.add_equation() rejects clearly bogus input shapes."""

    def it_raises_for_a_non_math_xml_root(self):
        doc = Document()
        p = doc.add_paragraph("")

        with pytest.raises(ValueError, match="m:oMath or m:oMathPara"):
            p.add_equation(
                '<w:p xmlns:w='
                '"http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
