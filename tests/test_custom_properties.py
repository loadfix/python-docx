# pyright: reportPrivateUsage=false

"""Unit-test suite for the `docx.custom_properties` module."""

from __future__ import annotations

import datetime as dt
from typing import cast

import pytest

from docx.custom_properties import CustomProperties
from docx.oxml.custom_properties import CT_CustomProperties
from docx.oxml.parser import parse_xml

from .unitutil.mock import FixtureRequest, Mock, instance_mock


_EMPTY_PROPERTIES_XML = (
    b'<Properties '
    b'xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
    b'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"/>'
)


def _empty_collection(part_mock: object | None = None) -> CustomProperties:
    elm = cast(CT_CustomProperties, parse_xml(_EMPTY_PROPERTIES_XML))
    return CustomProperties(elm, cast("object", part_mock))  # type: ignore[arg-type]


class DescribeCustomProperties:
    """Unit-test suite for `docx.custom_properties.CustomProperties`."""

    def it_has_zero_length_when_empty(self, part_: Mock):
        cp = _empty_collection(part_)

        assert len(cp) == 0
        assert list(cp) == []

    def it_can_add_and_read_back_a_string_value(self, part_: Mock):
        cp = _empty_collection(part_)

        cp.add("Project", "Alpha")

        assert len(cp) == 1
        assert cp["Project"] == "Alpha"
        assert isinstance(cp["Project"], str)

    def it_can_add_and_read_back_an_int_value(self, part_: Mock):
        cp = _empty_collection(part_)

        cp.add("Year", 2024)

        assert cp["Year"] == 2024
        assert type(cp["Year"]) is int

    def it_can_add_and_read_back_a_float_value(self, part_: Mock):
        cp = _empty_collection(part_)

        cp.add("Ratio", 1.25)

        assert cp["Ratio"] == 1.25
        assert type(cp["Ratio"]) is float

    def it_can_add_and_read_back_a_bool_value(self, part_: Mock):
        cp = _empty_collection(part_)

        cp.add("Released", True)
        cp.add("Archived", False)

        assert cp["Released"] is True
        assert cp["Archived"] is False

    def it_can_add_and_read_back_a_datetime_value(self, part_: Mock):
        cp = _empty_collection(part_)
        moment = dt.datetime(2024, 1, 15, 10, 30, 0)

        cp.add("ReleaseDate", moment)

        assert cp["ReleaseDate"] == moment
        assert isinstance(cp["ReleaseDate"], dt.datetime)

    def it_can_add_and_read_back_a_date_value(self, part_: Mock):
        """`datetime.date` values round-trip as `vt:date` (distinct from `datetime`)."""
        cp = _empty_collection(part_)
        review_date = dt.date(2024, 1, 15)

        cp.add("ReviewDate", review_date)

        retrieved = cp["ReviewDate"]
        assert retrieved == review_date
        assert isinstance(retrieved, dt.date)
        # -- must not be widened to a `datetime` on read --
        assert not isinstance(retrieved, dt.datetime)

    def it_raises_KeyError_on_missing_name(self, part_: Mock):
        cp = _empty_collection(part_)

        with pytest.raises(KeyError):
            _ = cp["Missing"]

    def it_raises_ValueError_when_adding_a_duplicate_name(self, part_: Mock):
        cp = _empty_collection(part_)
        cp.add("Project", "Alpha")

        with pytest.raises(ValueError):
            cp.add("Project", "Beta")

    def it_allows_overwrite_via_setitem(self, part_: Mock):
        cp = _empty_collection(part_)
        cp.add("Project", "Alpha")

        cp["Project"] = "Beta"

        assert cp["Project"] == "Beta"
        assert len(cp) == 1

    def it_supports_setitem_on_a_new_name(self, part_: Mock):
        cp = _empty_collection(part_)

        cp["Project"] = "Gamma"

        assert cp["Project"] == "Gamma"
        assert len(cp) == 1

    def it_supports_containment_checks(self, part_: Mock):
        cp = _empty_collection(part_)
        cp.add("Project", "Alpha")

        assert "Project" in cp
        assert "Missing" not in cp

    def it_supports_iteration_over_names(self, part_: Mock):
        cp = _empty_collection(part_)
        cp.add("A", 1)
        cp.add("B", 2)
        cp.add("C", 3)

        assert list(cp) == ["A", "B", "C"]
        assert cp.names() == ["A", "B", "C"]

    def it_exposes_name_value_pairs_via_items(self, part_: Mock):
        cp = _empty_collection(part_)
        cp.add("A", 1)
        cp.add("B", "two")

        assert cp.items() == [("A", 1), ("B", "two")]

    def it_can_delete_a_property(self, part_: Mock):
        cp = _empty_collection(part_)
        cp.add("A", 1)
        cp.add("B", 2)

        del cp["A"]

        assert "A" not in cp
        assert list(cp) == ["B"]
        assert len(cp) == 1

    def it_raises_KeyError_on_deleting_a_missing_name(self, part_: Mock):
        cp = _empty_collection(part_)

        with pytest.raises(KeyError):
            del cp["Missing"]

    def it_returns_default_from_get_when_missing(self, part_: Mock):
        cp = _empty_collection(part_)

        assert cp.get("Missing") is None
        assert cp.get("Missing", "fallback") == "fallback"

    def it_returns_value_from_get_when_present(self, part_: Mock):
        cp = _empty_collection(part_)
        cp.add("A", 42)

        assert cp.get("A") == 42

    def it_raises_TypeError_on_unsupported_value_type(self, part_: Mock):
        cp = _empty_collection(part_)

        with pytest.raises(TypeError):
            cp.add("Bad", object())

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def part_(self, request: FixtureRequest) -> Mock:
        # -- the collection doesn't need to interact with the part in these tests, so
        # -- a plain Mock is sufficient.
        return instance_mock(request, object)


class DescribeCustomProperties_RoundTrip:
    """Full save/reload integration for custom-property types (regression for #171).

    Exercises the real ``Document`` → ``custom.xml`` pipeline (not just the in-memory
    element wrapper) so that part (re-)serialisation, relationship wiring, and the
    vt-type dispatch all work end-to-end.
    """

    def it_round_trips_a_date_value_as_vt_date(self, tmp_path):
        # -- regression for GitHub issue #171: `datetime.date` must round-trip as
        # -- `vt:date` (not widened to `datetime` on read, not rejected on write).
        from docx import Document

        doc = Document()
        review = dt.date(2026, 5, 5)
        doc.custom_properties["ReviewDate"] = review
        out = tmp_path / "roundtrip-date.docx"
        doc.save(str(out))

        reloaded = Document(str(out))
        value = reloaded.custom_properties["ReviewDate"]

        assert value == review
        assert isinstance(value, dt.date)
        # -- crucially NOT a `datetime` (which is a `date` subclass) --
        assert not isinstance(value, dt.datetime)

    def it_round_trips_a_datetime_value_as_vt_filetime(self, tmp_path):
        from docx import Document

        doc = Document()
        moment = dt.datetime(2026, 5, 5, 12, 30, 0)
        doc.custom_properties["Released"] = moment
        out = tmp_path / "roundtrip-datetime.docx"
        doc.save(str(out))

        reloaded = Document(str(out))
        value = reloaded.custom_properties["Released"]

        assert value == moment
        assert type(value) is dt.datetime

    def it_round_trips_mixed_scalar_types_in_one_document(self, tmp_path):
        """All supported vt-types survive save/reload side-by-side."""
        from docx import Document

        doc = Document()
        doc.custom_properties["Project"] = "Alpha"
        doc.custom_properties["Year"] = 2026
        doc.custom_properties["Ratio"] = 1.25
        doc.custom_properties["Released"] = True
        doc.custom_properties["ReleaseDate"] = dt.datetime(2026, 5, 5, 10, 0, 0)
        doc.custom_properties["ReviewDate"] = dt.date(2026, 5, 5)
        out = tmp_path / "roundtrip-mixed.docx"
        doc.save(str(out))

        r = Document(str(out)).custom_properties

        assert r["Project"] == "Alpha"
        assert r["Year"] == 2026
        assert r["Ratio"] == 1.25
        assert r["Released"] is True
        assert r["ReleaseDate"] == dt.datetime(2026, 5, 5, 10, 0, 0)
        assert r["ReviewDate"] == dt.date(2026, 5, 5)
        # -- date must not have been widened to datetime --
        assert not isinstance(r["ReviewDate"], dt.datetime)
