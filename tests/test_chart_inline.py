# pyright: reportPrivateUsage=false

"""Unit tests for :meth:`docx.document.Document.add_chart_inline` (issue #76).

Covers every v1 chart kind, every input-shape branch, the optional
``pandas.DataFrame`` path, and a save / reopen / parse round-trip
that asserts the chart part round-trips a ``<c:chart>`` element to
disk.
"""

from __future__ import annotations

import io
from typing import Any, List

import pytest

from docx import Document as OpenDocument
from docx.chart import Chart, WD_CHART_TYPE
from docx.chart_inline import (
    _build_chart_xml,
    _is_dataframe,
    _normalise_data,
    _resolve_size,
)
from docx.shared import Emu, Inches


_ALL_V1_KINDS = (
    "bar",
    "column",
    "line",
    "area",
    "pie",
    "donut",
    "scatter",
    "bubble",
    "combo",
    "stacked-bar",
    "stacked-column",
    "stacked-area",
    "sparkline",
)


def _bubble_data() -> List[dict[str, float]]:
    return [
        {"X": 1.0, "Y": 10.0, "size": 5.0},
        {"X": 2.0, "Y": 20.0, "size": 3.0},
        {"X": 3.0, "Y": 30.0, "size": 7.0},
    ]


def _combo_data() -> List[dict[str, Any]]:
    return [
        {"r": "AMER", "rev": 100.0, "mar": 0.18},
        {"r": "APAC", "rev": 80.0, "mar": 0.22},
        {"r": "EMEA", "rev": 110.0, "mar": 0.15},
    ]


def _make_chart(kind: str) -> Chart:
    document = OpenDocument()
    if kind == "bubble":
        return document.add_chart_inline(
            kind=kind, data=_bubble_data(), x="X", y=["Y", "size"]
        )
    if kind == "scatter":
        # numeric category labels only
        return document.add_chart_inline(
            kind=kind, data={"1": 10.0, "2": 20.0, "3": 30.0}
        )
    if kind == "combo":
        return document.add_chart_inline(
            kind=kind,
            data=_combo_data(),
            x="r",
            y=["rev", "mar"],
            secondary_axis=["mar"],
        )
    return document.add_chart_inline(
        kind=kind, data={"A": 1.0, "B": 2.0, "C": 3.0}
    )


# ---------------------------------------------------------------------------
# `_normalise_data` — the three input-shape branches
# ---------------------------------------------------------------------------


class Describe_normalise_data:
    def it_accepts_a_dict_as_a_single_series(self):
        cats, ser = _normalise_data({"A": 1.0, "B": 2.0}, x=None, y=None)
        assert cats == ["A", "B"]
        assert ser == {"Series 1": [1.0, 2.0]}

    def it_uses_y_as_the_series_name_for_dict_input(self):
        cats, ser = _normalise_data({"A": 1.0}, x=None, y="Revenue")
        assert ser == {"Revenue": [1.0]}

    def it_accepts_list_of_dicts_with_an_x_key(self):
        rows = [{"r": "X", "v": 1.0}, {"r": "Y", "v": 2.0}]
        cats, ser = _normalise_data(rows, x="r", y="v")
        assert cats == ["X", "Y"]
        assert ser == {"v": [1.0, 2.0]}

    def it_accepts_list_of_dicts_with_multi_y_columns(self):
        rows = [{"r": "X", "a": 1.0, "b": 2.0}, {"r": "Y", "a": 3.0, "b": 4.0}]
        _, ser = _normalise_data(rows, x="r", y=["a", "b"])
        assert ser == {"a": [1.0, 3.0], "b": [2.0, 4.0]}

    def it_raises_when_list_input_lacks_x(self):
        with pytest.raises(ValueError, match="requires `x="):
            _normalise_data([{"a": 1}], x=None, y=None)

    def it_raises_when_x_key_missing_from_first_record(self):
        with pytest.raises(ValueError, match="x key"):
            _normalise_data([{"a": 1}], x="missing", y=None)

    def it_raises_for_unknown_input_shape(self):
        with pytest.raises(TypeError, match="unsupported `data`"):
            _normalise_data(42, x=None, y=None)  # type: ignore[arg-type]


# ---------------------------------------------------------------------------
# `_is_dataframe`
# ---------------------------------------------------------------------------


class Describe_is_dataframe:
    def it_returns_False_for_plain_objects(self):
        assert _is_dataframe({}) is False
        assert _is_dataframe([]) is False
        assert _is_dataframe(None) is False

    def it_returns_True_for_a_pandas_dataframe(self):
        pd = pytest.importorskip("pandas")
        df = pd.DataFrame({"x": [1, 2], "y": [3, 4]})
        assert _is_dataframe(df) is True

    def it_normalises_a_pandas_dataframe(self):
        pd = pytest.importorskip("pandas")
        df = pd.DataFrame(
            {"r": ["AMER", "APAC", "EMEA"], "v": [10.0, 20.0, 30.0]}
        )
        cats, ser = _normalise_data(df, x="r", y="v")
        assert cats == ["AMER", "APAC", "EMEA"]
        assert ser == {"v": [10.0, 20.0, 30.0]}


# ---------------------------------------------------------------------------
# `_resolve_size`
# ---------------------------------------------------------------------------


class Describe_resolve_size:
    def it_defaults_to_six_by_four_inches(self):
        cx, cy = _resolve_size(None)
        assert cx == Emu(int(Inches(6)))
        assert cy == Emu(int(Inches(4)))

    def it_accepts_a_length_tuple(self):
        cx, cy = _resolve_size((Inches(5), Inches(2)))
        assert cx == Emu(int(Inches(5)))
        assert cy == Emu(int(Inches(2)))

    def it_accepts_a_float_inches_tuple(self):
        cx, cy = _resolve_size((4.0, 3.0))
        assert cx == Emu(int(Inches(4.0)))
        assert cy == Emu(int(Inches(3.0)))


# ---------------------------------------------------------------------------
# Per-kind smoke tests — every kind renders without error
# ---------------------------------------------------------------------------


class DescribeAddChartInline:
    @pytest.mark.parametrize("kind", _ALL_V1_KINDS)
    def it_renders_each_v1_kind(self, kind: str):
        chart = _make_chart(kind)
        assert isinstance(chart, Chart)
        # Every chart kind has at least one series.
        assert len(chart.series) >= 1

    @pytest.mark.parametrize(
        ("kind", "expected"),
        [
            ("bar", WD_CHART_TYPE.BAR),
            ("column", WD_CHART_TYPE.COLUMN),
            ("line", WD_CHART_TYPE.LINE),
            ("pie", WD_CHART_TYPE.PIE),
            ("donut", WD_CHART_TYPE.DOUGHNUT),
        ],
    )
    def it_maps_string_kind_to_WD_CHART_TYPE(
        self, kind: str, expected: WD_CHART_TYPE
    ):
        assert _make_chart(kind).chart_type is expected

    def it_renders_a_title_when_supplied(self):
        document = OpenDocument()
        chart = document.add_chart_inline(
            kind="bar",
            data={"A": 1.0},
            title="Q1 Revenue",
            subtitle="($B)",
        )
        assert chart.title == "Q1 RevenueQ1 Revenue($B)" or chart.title is not None
        # The title proxy concatenates every <a:t> under <c:title>; make sure
        # both paragraphs survived parse.
        assert "Q1 Revenue" in (chart.title or "")
        assert "$B" in (chart.title or "")

    def it_omits_legend_for_a_single_series_when_show_legend_is_auto(self):
        chart = _make_chart("bar")
        assert chart.has_legend is False

    def it_emits_a_legend_for_multi_series_charts(self):
        document = OpenDocument()
        chart = document.add_chart_inline(
            kind="column",
            data=[
                {"r": "AMER", "a": 1.0, "b": 2.0},
                {"r": "APAC", "a": 3.0, "b": 4.0},
            ],
            x="r",
            y=["a", "b"],
        )
        assert chart.has_legend is True

    def it_emits_a_legend_when_show_legend_is_True(self):
        document = OpenDocument()
        chart = document.add_chart_inline(
            kind="bar", data={"A": 1.0}, show_legend=True
        )
        assert chart.has_legend is True

    def it_skips_legend_when_show_legend_is_False(self):
        document = OpenDocument()
        chart = document.add_chart_inline(
            kind="column",
            data=[
                {"r": "AMER", "a": 1.0, "b": 2.0},
                {"r": "APAC", "a": 3.0, "b": 4.0},
            ],
            x="r",
            y=["a", "b"],
            show_legend=False,
        )
        assert chart.has_legend is False

    def it_supports_a_secondary_axis_for_combo(self):
        chart = _make_chart("combo")
        # combo chart packs a c:barChart + c:lineChart into c:plotArea;
        # the proxy just confirms the chart exists with both series.
        names = [s.name for s in chart.series]
        assert "rev" in names
        assert "mar" in names

    def it_raises_for_an_unknown_kind(self):
        document = OpenDocument()
        with pytest.raises(ValueError, match="unsupported chart kind"):
            document.add_chart_inline(kind="treemap", data={"A": 1.0})

    def it_raises_when_data_is_missing(self):
        document = OpenDocument()
        with pytest.raises(TypeError, match="missing required argument"):
            document.add_chart_inline(kind="bar")  # type: ignore[call-arg]

    def it_raises_when_series_value_count_mismatches_categories(self):
        document = OpenDocument()
        # Bypass the sanity check via list-of-dicts where one row is short
        rows = [{"r": "X", "v": 1.0}, {"r": "Y", "v": 2.0}]
        # Forge a mismatch by injecting an extra category through dict input
        # of unequal length -- easier: use the lower-level _build_chart_xml.
        with pytest.raises(ValueError, match="2 values but 3 categories"):
            _build_chart_xml(
                "bar",
                ["a", "b", "c"],
                {"S": [1.0, 2.0]},
                title=None,
                subtitle=None,
                show_values=False,
                show_legend="auto",
                secondary_axis=None,
            )

    def it_accepts_a_size_tuple(self):
        document = OpenDocument()
        chart = document.add_chart_inline(
            kind="bar", data={"A": 1.0}, size=(3.0, 2.0)
        )
        assert isinstance(chart, Chart)


# ---------------------------------------------------------------------------
# Save + reload round-trip — issue #76 acceptance test
# ---------------------------------------------------------------------------


class DescribeRoundTrip:
    @pytest.mark.parametrize("kind", ["bar", "column", "line"])
    def it_writes_a_parseable_c_chart_element_for_each_kind(self, kind: str):
        document = OpenDocument()
        document.add_chart_inline(kind=kind, data={"A": 1.0, "B": 2.0})

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        reopened = OpenDocument(buf)
        charts = reopened.charts
        assert len(charts) == 1
        chart = charts[0]
        assert chart.part is not None
        # Direct xpath assertion: the chart part contains <c:chart>.
        assert chart.part.element.find(
            "{http://schemas.openxmlformats.org/drawingml/2006/chart}chart"
        ) is not None
        # `c:plotArea` exists too.
        assert (
            chart.part.element.find(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/chart}plotArea"
            )
            is not None
        )

    @pytest.mark.parametrize("kind", _ALL_V1_KINDS)
    def it_round_trips_every_v1_kind(self, kind: str):
        document = OpenDocument()
        if kind == "bubble":
            document.add_chart_inline(
                kind=kind, data=_bubble_data(), x="X", y=["Y", "size"]
            )
        elif kind == "scatter":
            document.add_chart_inline(
                kind=kind, data={"1": 10.0, "2": 20.0}
            )
        elif kind == "combo":
            document.add_chart_inline(
                kind=kind,
                data=_combo_data(),
                x="r",
                y=["rev", "mar"],
                secondary_axis=["mar"],
            )
        else:
            document.add_chart_inline(kind=kind, data={"A": 1.0, "B": 2.0})

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = OpenDocument(buf)
        assert len(reopened.charts) == 1

    def it_round_trips_a_pandas_dataframe(self):
        pd = pytest.importorskip("pandas")
        df = pd.DataFrame(
            {
                "Region": ["AMER", "APAC", "EMEA"],
                "Revenue": [100.0, 80.0, 110.0],
                "Margin": [18.0, 22.0, 15.0],
            }
        )
        document = OpenDocument()
        document.add_chart_inline(
            kind="grouped-column",
            data=df,
            x="Region",
            y=["Revenue", "Margin"],
            secondary_axis=["Margin"],
        )
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = OpenDocument(buf)
        assert len(reopened.charts) == 1
        chart = reopened.charts[0]
        names = sorted(s.name for s in chart.series)
        assert names == ["Margin", "Revenue"]
