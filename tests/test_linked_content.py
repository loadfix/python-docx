"""Unit-test suite for the docx.linked_content module (issue #96)."""

from __future__ import annotations

import io
import os
import tempfile
from typing import cast

import pytest

from docx import Document
from docx.linked_content import (
    INCLUDETEXT,
    LINK_KIND_PPTX_SLIDE,
    LINK_KIND_UNKNOWN,
    LINK_KIND_XLSX_CELL,
    LINK_KIND_XLSX_TABLE_COLUMN,
    LinkedTarget,
    ParsedLinkTarget,
    UNRESOLVED_PLACEHOLDER,
    build_includetext_instruction,
    iter_linked_targets,
    parse_includetext_instruction,
    parse_link_target,
    update_document_links,
)
from docx.text.paragraph import Paragraph


class DescribeParseLinkTarget:
    """`parse_link_target()` URL parser."""

    @pytest.mark.parametrize(
        ("url", "kind", "detail"),
        [
            ("revenue.xlsx#RevenueQ1!B5", LINK_KIND_XLSX_CELL, ("RevenueQ1", "B5")),
            (
                "data.xlsm#'Quarter One'!$A$1",
                LINK_KIND_XLSX_CELL,
                ("Quarter One", "$A$1"),
            ),
            (
                "revenue.xlsx#RevenueQ1[Total]",
                LINK_KIND_XLSX_TABLE_COLUMN,
                ("RevenueQ1", "Total"),
            ),
            (
                "data.xltx#Sales[Q1 Revenue]",
                LINK_KIND_XLSX_TABLE_COLUMN,
                ("Sales", "Q1 Revenue"),
            ),
            ("summary.pptx#slide-3", LINK_KIND_PPTX_SLIDE, ("3",)),
            ("deck.pptm#Slide_42", LINK_KIND_PPTX_SLIDE, ("42",)),
            ("foo.xlsx", LINK_KIND_UNKNOWN, ()),
            ("foo.txt#bar!Baz", LINK_KIND_UNKNOWN, ()),
            ("", LINK_KIND_UNKNOWN, ()),
        ],
    )
    def it_parses_known_target_shapes(
        self, url: str, kind: str, detail: tuple
    ):
        parsed = parse_link_target(url)
        assert parsed.kind == kind
        assert parsed.detail == detail
        if "#" in url:
            assert parsed.path == url.split("#", 1)[0]
            assert parsed.fragment == url.split("#", 1)[1]
        else:
            assert parsed.path == url
            assert parsed.fragment == ""

    def it_handles_relative_paths_with_directories(self):
        parsed = parse_link_target("subdir/data.xlsx#Sheet1!A1")
        assert parsed.kind == LINK_KIND_XLSX_CELL
        assert parsed.path == "subdir/data.xlsx"
        assert parsed.detail == ("Sheet1", "A1")

    def it_treats_missing_fragment_as_unknown(self):
        parsed = parse_link_target("revenue.xlsx")
        assert parsed.kind == LINK_KIND_UNKNOWN
        assert parsed.fragment == ""

    def it_returns_unknown_for_non_string(self):
        # -- defensive: parser shouldn't blow up on misuse --
        parsed = parse_link_target(cast(str, None))  # type: ignore[arg-type]
        assert parsed.kind == LINK_KIND_UNKNOWN


class DescribeIncludeTextInstructionRoundtrip:
    """`build_includetext_instruction` / `parse_includetext_instruction`."""

    @pytest.mark.parametrize(
        "url",
        [
            "revenue.xlsx#RevenueQ1!B5",
            "summary.pptx#slide-3",
            "data/sub.xlsx#T[Total]",
            'has"quote.xlsx#S!A1',
            "https://example.com/file.xlsx#Sheet1!Z99",
        ],
    )
    def it_round_trips(self, url: str):
        instr = build_includetext_instruction(url)
        assert instr.startswith(INCLUDETEXT + " ")
        assert parse_includetext_instruction(instr) == url

    def it_rejects_empty_url(self):
        with pytest.raises(ValueError):
            build_includetext_instruction("")

    def it_rejects_non_string_url(self):
        with pytest.raises(TypeError):
            build_includetext_instruction(cast(str, 42))  # type: ignore[arg-type]

    def it_returns_None_for_non_includetext_instruction(self):
        assert parse_includetext_instruction("PAGE") is None
        assert parse_includetext_instruction("") is None
        assert parse_includetext_instruction("INCLUDETEXT") is None

    def it_accepts_unquoted_url_form(self):
        # -- some authoring tools omit the surrounding quotes --
        assert (
            parse_includetext_instruction("INCLUDETEXT bare.xlsx#S!A1")
            == "bare.xlsx#S!A1"
        )


class DescribeParagraphLinkTo:
    """`Paragraph.link_to()` writes an INCLUDETEXT complex field."""

    def it_emits_an_INCLUDETEXT_field(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("revenue.xlsx#RevenueQ1!B5")

        assert isinstance(link, LinkedTarget)
        assert link.field.type == INCLUDETEXT
        assert link.field.is_complex is True

    def it_writes_a_placeholder_when_no_cached_result_is_given(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("revenue.xlsx#RevenueQ1!B5")
        assert link.cached_text == UNRESOLVED_PLACEHOLDER

    def it_uses_caller_supplied_cached_result(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("summary.pptx#slide-3", cached_result="Slide 3")
        assert link.cached_text == "Slide 3"

    def it_marks_the_field_dirty_by_default(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("revenue.xlsx#RevenueQ1!B5")
        assert link.field.is_dirty is True

    def it_can_skip_marking_dirty(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to(
            "revenue.xlsx#RevenueQ1!B5", mark_dirty=False
        )
        assert link.field.is_dirty is False

    def it_exposes_kind_and_url_via_proxy(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("revenue.xlsx#RevenueQ1[Total]")
        assert link.url == "revenue.xlsx#RevenueQ1[Total]"
        assert link.kind == LINK_KIND_XLSX_TABLE_COLUMN
        assert link.path == "revenue.xlsx"
        assert link.fragment == "RevenueQ1[Total]"

    def it_rejects_empty_url(self):
        doc = Document()
        para = doc.add_paragraph()
        with pytest.raises(ValueError):
            para.link_to("")


class DescribeDocumentLinkedTargets:
    """`Document.linked_targets` and `update_links()`."""

    def it_returns_every_linked_target_in_order(self):
        doc = Document()
        for url in (
            "a.xlsx#S1!A1",
            "b.xlsx#T[C]",
            "c.pptx#slide-2",
        ):
            doc.add_paragraph().link_to(url)

        targets = doc.linked_targets
        assert [t.url for t in targets] == [
            "a.xlsx#S1!A1",
            "b.xlsx#T[C]",
            "c.pptx#slide-2",
        ]
        assert [t.kind for t in targets] == [
            LINK_KIND_XLSX_CELL,
            LINK_KIND_XLSX_TABLE_COLUMN,
            LINK_KIND_PPTX_SLIDE,
        ]

    def it_skips_non_INCLUDETEXT_fields(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_complex_field("PAGE", "1")
        para.link_to("revenue.xlsx#RevenueQ1!B5")
        para.add_complex_field("DATE")

        assert len(doc.linked_targets) == 1
        assert doc.linked_targets[0].url == "revenue.xlsx#RevenueQ1!B5"

    def it_round_trips_through_save_and_reload(self):
        doc = Document()
        urls = [
            "revenue.xlsx#RevenueQ1!B5",
            "revenue.xlsx#RevenueQ1[Total]",
            "summary.pptx#slide-3",
        ]
        for url in urls:
            doc.add_paragraph().link_to(url)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        targets = reloaded.linked_targets
        assert len(targets) == len(urls)
        for tgt, url in zip(targets, urls):
            assert tgt.url == url
            # -- placeholder cached text round-trips intact --
            assert tgt.cached_text == UNRESOLVED_PLACEHOLDER

    def it_iter_helper_matches_property(self):
        doc = Document()
        doc.add_paragraph().link_to("revenue.xlsx#RevenueQ1!B5")
        from_property = doc.linked_targets
        from_iter = list(iter_linked_targets(doc))
        assert len(from_property) == len(from_iter) == 1
        assert from_property[0].url == from_iter[0].url

    def it_update_links_skips_missing_files(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("does_not_exist.xlsx#Sheet1!A1")
        before = link.cached_text
        # -- no file at the path; refresh leaves the cached result alone --
        n = doc.update_links()
        assert n == 0
        assert link.cached_text == before

    def it_update_links_leaves_unknown_kinds_alone(self):
        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("foo.txt#anchor")
        before = link.cached_text
        n = doc.update_links()
        assert n == 0
        assert link.cached_text == before


class DescribeXlsxResolution:
    """End-to-end resolution against a real workbook (sibling xlsx)."""

    def it_resolves_an_xlsx_cell_via_sibling_package(self, tmp_path):
        xlsx = pytest.importorskip("xlsx")
        xlsx_path = tmp_path / "revenue.xlsx"
        wb = xlsx.Workbook()
        ws = wb.active
        ws.title = "RevenueQ1"
        ws["B5"] = 4242
        wb.save(str(xlsx_path))

        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("revenue.xlsx#RevenueQ1!B5")
        assert link.cached_text == UNRESOLVED_PLACEHOLDER

        n = doc.update_links(base_dir=str(tmp_path))
        assert n == 1
        assert link.cached_text == "4242"

    def it_resolves_a_quoted_sheet_name(self, tmp_path):
        xlsx = pytest.importorskip("xlsx")
        xlsx_path = tmp_path / "data.xlsx"
        wb = xlsx.Workbook()
        ws = wb.active
        ws.title = "Sheet With Spaces"
        ws["A1"] = "hello"
        wb.save(str(xlsx_path))

        doc = Document()
        para = doc.add_paragraph()
        link = para.link_to("data.xlsx#'Sheet With Spaces'!A1")
        n = doc.update_links(base_dir=str(tmp_path))
        assert n == 1
        assert link.cached_text == "hello"

    def it_returns_None_for_missing_sheet(self, tmp_path):
        xlsx = pytest.importorskip("xlsx")
        xlsx_path = tmp_path / "data.xlsx"
        wb = xlsx.Workbook()
        wb.active["A1"] = 1
        wb.save(str(xlsx_path))

        doc = Document()
        link = doc.add_paragraph().link_to(
            "data.xlsx#NoSuchSheet!A1"
        )
        n = doc.update_links(base_dir=str(tmp_path))
        # -- the file exists but the sheet doesn't; refresh is a no-op --
        assert n == 0
        assert link.cached_text == UNRESOLVED_PLACEHOLDER


class DescribeLinkedTargetProxy:
    """`LinkedTarget` proxy view of an INCLUDETEXT field."""

    def it_exposes_the_underlying_field(self):
        doc = Document()
        link = doc.add_paragraph().link_to("revenue.xlsx#RevenueQ1!B5")
        assert link.field.type == INCLUDETEXT

    def it_returns_the_parsed_target(self):
        doc = Document()
        link = doc.add_paragraph().link_to("summary.pptx#slide-3")
        parsed = link.parsed
        assert isinstance(parsed, ParsedLinkTarget)
        assert parsed.kind == LINK_KIND_PPTX_SLIDE
        assert parsed.detail == ("3",)

    def it_returns_empty_url_for_non_includetext(self):
        # -- LinkedTarget can wrap any field, but only pulls a URL from
        #    INCLUDETEXT instructions; a PAGE field returns "" --
        from docx.fields import Field

        doc = Document()
        para = doc.add_paragraph()
        para.add_complex_field("PAGE", "1")
        field = doc.fields[0]
        assert field.type == "PAGE"
        wrapped = LinkedTarget(field)
        assert wrapped.url == ""
        assert wrapped.kind == LINK_KIND_UNKNOWN

    def it_resolve_returns_None_for_unknown_kind(self):
        doc = Document()
        link = doc.add_paragraph().link_to("foo.txt#anchor")
        assert link.resolve() is None

    def it_refresh_returns_None_when_unresolvable(self):
        doc = Document()
        link = doc.add_paragraph().link_to("missing.xlsx#S!A1")
        assert link.refresh() is None
