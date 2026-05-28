"""Tests for `Document.diff()` semantic compare (issue #75)."""

from __future__ import annotations

import io

import pytest

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.semantic_diff import (
    Change,
    SemanticDiff,
    VALID_LEVELS,
    compute_diff,
)
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _make_doc(paragraphs):
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


# Thirty-paragraph "before" fixture. Numbers 0-29 are used as a stable
# identifier for the assertions below.
_BEFORE = [f"Section {i:02d}: original line content" for i in range(30)]


def _build_before_after_pair():
    """Return ``(old_doc, new_doc)`` with a known set of differences.

    Changes injected against the 30-paragraph baseline:

    * **paragraph_added** — three brand-new lines inserted at fresh
      positions.
    * **paragraph_removed** — one removed line.
    * **paragraph_modified** — seven edited lines (text shifts).
    """
    old = _make_doc(_BEFORE)

    after = list(_BEFORE)

    # -- 7 modified -----------------------------------------------
    modified_indices = [2, 5, 8, 11, 14, 17, 20]
    for idx in modified_indices:
        after[idx] = after[idx] + " (revised)"

    # -- 1 removed ------------------------------------------------
    removed_index = 25
    removed_text = after.pop(removed_index)

    # -- 3 added (inserted at fresh positions) --------------------
    after.insert(3, "NEW line A")
    after.insert(10, "NEW line B")
    after.append("NEW line C")

    new = _make_doc(after)
    return old, new, modified_indices, removed_text


# ---------------------------------------------------------------------------
# Public surface
# ---------------------------------------------------------------------------


class DescribeDocumentDiff:
    def it_returns_a_SemanticDiff_instance(self):
        old = Document()
        new = Document()
        diff = old.diff(new)
        assert isinstance(diff, SemanticDiff)
        assert diff.level == "content"

    def it_defaults_to_content_level(self):
        old = Document()
        new = Document()
        diff = old.diff(new)
        assert diff.level == "content"

    def it_accepts_each_documented_level(self):
        old = Document()
        new = Document()
        for level in VALID_LEVELS:
            diff = old.diff(new, level=level)
            assert diff.level == level

    def it_rejects_an_unknown_level(self):
        old = Document()
        new = Document()
        with pytest.raises(ValueError):
            old.diff(new, level="bananas")

    def it_returns_an_empty_diff_for_identical_documents(self):
        old = _make_doc(["Alpha", "Beta", "Gamma"])
        new = _make_doc(["Alpha", "Beta", "Gamma"])
        diff = old.diff(new)
        assert diff.summary["total_changes"] == 0
        assert not diff
        assert list(diff) == []


class DescribeSemanticDiff_Summary:
    def it_returns_zero_counts_when_unchanged(self):
        old = _make_doc(["Alpha"])
        diff = old.diff(_make_doc(["Alpha"]))
        s = diff.summary
        assert s["paragraphs_added"] == 0
        assert s["paragraphs_removed"] == 0
        assert s["paragraphs_modified"] == 0
        assert s["total_changes"] == 0

    def it_counts_a_pure_addition(self):
        old = _make_doc(["A", "B"])
        new = _make_doc(["A", "B", "C"])
        diff = old.diff(new)
        assert diff.summary["paragraphs_added"] == 1
        assert diff.summary["total_changes"] == 1

    def it_counts_a_pure_removal(self):
        old = _make_doc(["A", "B", "C"])
        new = _make_doc(["A", "B"])
        diff = old.diff(new)
        assert diff.summary["paragraphs_removed"] == 1
        assert diff.summary["total_changes"] == 1

    def it_counts_a_modification_at_content_level(self):
        old = _make_doc(["Hello world"])
        new = _make_doc(["Hello mars"])
        diff = old.diff(new, level="content")
        assert diff.summary["paragraphs_modified"] == 1
        assert diff.summary["paragraphs_added"] == 0
        assert diff.summary["paragraphs_removed"] == 0


class DescribeSemanticDiff_FixtureRoundTrip:
    """Verify the 30-paragraph fixture catches all and only the seeded changes."""

    def it_catches_all_and_only_the_seeded_changes(self):
        old, new, modified_indices, _removed = _build_before_after_pair()
        diff = old.diff(new, level="content")
        s = diff.summary
        assert s["paragraphs_added"] == 3
        assert s["paragraphs_removed"] == 1
        assert s["paragraphs_modified"] == len(modified_indices)
        # -- structure: nothing else should have leaked in --
        assert s["tables_modified"] == 0
        assert s["images_added"] == 0
        assert s["styles_changed"] == 0
        assert s["total_changes"] == 3 + 1 + len(modified_indices)

    def it_each_modification_pairs_old_and_new_text(self):
        old, new, _modified_indices, _removed = _build_before_after_pair()
        diff = old.diff(new, level="content")
        mods = diff.filter("paragraph_modified")
        assert len(mods) == 7
        for change in mods:
            assert change.before
            assert change.after
            assert change.before != change.after
            assert "(revised)" in change.after

    def it_records_the_removed_paragraph_text(self):
        old, new, _modified_indices, removed = _build_before_after_pair()
        diff = old.diff(new, level="content")
        removed_changes = diff.filter("paragraph_removed")
        assert len(removed_changes) == 1
        assert removed_changes[0].before == removed

    def it_records_each_added_paragraph_text(self):
        old, new, _modified_indices, _removed = _build_before_after_pair()
        diff = old.diff(new, level="content")
        added = diff.filter("paragraph_added")
        added_text = sorted(c.after for c in added)
        assert added_text == ["NEW line A", "NEW line B", "NEW line C"]


class DescribeSemanticDiff_StructuralLevel:
    def it_does_not_report_text_modifications_at_structural_level(self):
        old = _make_doc(["Hello world", "Bravo"])
        new = _make_doc(["Hello mars", "Bravo"])
        diff = old.diff(new, level="structural")
        s = diff.summary
        # -- structural level reports the edit as remove + add, not modified --
        assert s["paragraphs_modified"] == 0
        assert s["paragraphs_removed"] == 1
        assert s["paragraphs_added"] == 1


class DescribeSemanticDiff_FormattingLevel:
    def it_reports_a_paragraph_style_change(self):
        old = _make_doc(["Hello world"])
        new = _make_doc(["Hello world"])
        new.paragraphs[0].style = "Heading 1"
        diff = old.diff(new, level="formatting")
        fmt = diff.filter("formatting_changed")
        assert len(fmt) >= 1
        assert any("style" in (c.detail or "") for c in fmt)

    def it_reports_a_run_font_size_change(self):
        old = _make_doc(["Hello world"])
        new = _make_doc(["Hello world"])
        new.paragraphs[0].runs[0].font.size = Pt(24)
        diff = old.diff(new, level="formatting")
        assert any(
            c.kind == "formatting_changed"
            and "run formatting" in (c.detail or "")
            for c in diff.changes
        )

    def it_reports_a_style_added_in_styles_collection(self):
        old = _make_doc(["body"])
        new = _make_doc(["body"])
        new.styles.add_style("MyNewStyle", WD_STYLE_TYPE.PARAGRAPH)
        diff = old.diff(new, level="formatting")
        added = [
            c
            for c in diff.changes
            if c.kind == "style_added" and c.after == "MyNewStyle"
        ]
        assert len(added) == 1

    def it_reports_a_style_removed_from_styles_collection(self):
        old = _make_doc(["body"])
        old.styles.add_style("DropMe", WD_STYLE_TYPE.PARAGRAPH)
        new = _make_doc(["body"])
        diff = old.diff(new, level="formatting")
        removed = [
            c
            for c in diff.changes
            if c.kind == "style_removed" and c.before == "DropMe"
        ]
        assert len(removed) == 1

    def it_does_not_emit_styles_changed_at_content_level(self):
        old = _make_doc(["body"])
        new = _make_doc(["body"])
        new.styles.add_style("NotAtContentLevel", WD_STYLE_TYPE.PARAGRAPH)
        diff = old.diff(new, level="content")
        assert diff.summary["styles_added"] == 0


class DescribeSemanticDiff_Tables:
    def it_reports_a_new_table_as_table_added(self):
        old = Document()
        new = Document()
        new.add_table(rows=2, cols=2)
        diff = old.diff(new)
        added = diff.filter("table_added")
        assert len(added) == 1

    def it_reports_a_dropped_table_as_table_removed(self):
        old = Document()
        old.add_table(rows=2, cols=2)
        new = Document()
        diff = old.diff(new)
        removed = diff.filter("table_removed")
        assert len(removed) == 1

    def it_reports_a_cell_text_change_as_table_modified(self):
        old = Document()
        old_table = old.add_table(rows=2, cols=2)
        old_table.cell(0, 0).text = "alpha"
        old_table.cell(0, 1).text = "beta"
        new = Document()
        new_table = new.add_table(rows=2, cols=2)
        new_table.cell(0, 0).text = "alpha"
        new_table.cell(0, 1).text = "beta CHANGED"
        diff = old.diff(new)
        mods = diff.filter("table_modified")
        assert len(mods) == 1
        assert "cell" in (mods[0].detail or "")


class DescribeSemanticDiff_Images:
    def it_reports_zero_image_changes_when_neither_doc_has_images(self):
        old = _make_doc(["a"])
        new = _make_doc(["a"])
        diff = old.diff(new)
        assert diff.summary["images_added"] == 0
        assert diff.summary["images_removed"] == 0


class DescribeSemanticDiff_ToMarkdown:
    def it_emits_a_markdown_table_with_per_kind_counts(self):
        old, new, modified_indices, _removed = _build_before_after_pair()
        md = old.diff(new).to_markdown()
        assert "Document diff" in md
        assert "| Kind | Count |" in md
        assert "Paragraphs added" in md
        assert "Paragraphs modified" in md
        assert "Total changes" in md
        # -- modified section is rendered with backticks around the locator --
        assert "paragraph[" in md

    def it_caps_per_kind_detail_at_max_per_kind(self):
        # -- generate an old vs. new with many additions
        old = _make_doc(["base"])
        many = ["base"] + [f"line {i}" for i in range(50)]
        new = _make_doc(many)
        md = old.diff(new).to_markdown(max_per_kind=5)
        assert "more elided" in md

    def it_returns_str(self):
        old = Document()
        new = Document()
        md = old.diff(new).to_markdown()
        assert isinstance(md, str)


class DescribeSemanticDiff_ToHtml:
    def it_emits_a_self_contained_html_fragment(self):
        old, new, _modified_indices, _removed = _build_before_after_pair()
        html_out = old.diff(new).to_html()
        assert html_out.startswith('<div class="docx-semantic-diff">')
        assert html_out.endswith("</div>")
        assert "<table>" in html_out
        assert "<ul>" in html_out

    def it_escapes_text_content_to_guard_against_xss(self):
        old = _make_doc(["safe"])
        new = _make_doc(["<script>alert(1)</script>"])
        out = old.diff(new).to_html()
        assert "<script>" not in out
        assert "&lt;script&gt;" in out


class DescribeSemanticDiff_ToWordTrackChanges:
    def it_emits_a_document_summarising_each_change(self):
        old, new, _modified_indices, _removed = _build_before_after_pair()
        diff = old.diff(new)
        out_doc = diff.to_word_track_changes()
        text = "\n".join(p.text for p in out_doc.paragraphs)
        assert "Document diff" in text
        assert "[INS]" in text
        assert "[DEL]" in text
        assert "[~MOD]" in text

    def it_returns_a_document_that_can_be_saved(self):
        old = _make_doc(["x"])
        new = _make_doc(["y"])
        out_doc = old.diff(new).to_word_track_changes()
        buf = io.BytesIO()
        out_doc.save(buf)
        assert buf.tell() > 0


class DescribeChange:
    def it_is_a_frozen_dataclass(self):
        c = Change(kind="paragraph_added", target="paragraph[0]", after="hi")
        with pytest.raises(Exception):
            c.kind = "paragraph_removed"  # type: ignore[misc]

    def it_serialises_via_to_dict(self):
        c = Change(
            kind="paragraph_modified",
            target="paragraph[2]",
            before="old",
            after="new",
        )
        d = c.to_dict()
        assert d == {
            "kind": "paragraph_modified",
            "target": "paragraph[2]",
            "before": "old",
            "after": "new",
            "detail": None,
        }


class DescribeComputeDiffFunction:
    """``compute_diff`` is the module-level driver used by ``Document.diff``."""

    def it_is_callable_directly(self):
        old = _make_doc(["A"])
        new = _make_doc(["A", "B"])
        diff = compute_diff(old, new, level="content")
        assert isinstance(diff, SemanticDiff)
        assert diff.summary["paragraphs_added"] == 1

    def it_validates_level(self):
        old = Document()
        new = Document()
        with pytest.raises(ValueError):
            compute_diff(old, new, level="nope")
