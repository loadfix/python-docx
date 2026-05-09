"""Unit-test suite for the docx.form_fields module."""

from __future__ import annotations

import io
from typing import cast

import pytest

from docx import Document
from docx.form_fields import (
    CheckBoxField,
    CheckboxFormField,
    DropDownListField,
    DropdownFormField,
    FormField,
    TextInputField,
    TextInputFormField,
    WD_FORM_FIELD_TYPE,
)
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element


# --- helper utilities ----------------------------------------------------


def _text_form_field_p() -> CT_P:
    """Build a CT_P containing a text form field via the public API."""
    p = cast(CT_P, element("w:p"))
    paragraph = Paragraph(p, None)  # type: ignore[arg-type]
    paragraph.add_text_form_field(name="Text1", default="Hello", maxlength=20)
    return p


def _checkbox_form_field_p(checked: bool = True) -> CT_P:
    p = cast(CT_P, element("w:p"))
    paragraph = Paragraph(p, None)  # type: ignore[arg-type]
    paragraph.add_checkbox_form_field(name="Agree", checked=checked)
    return p


def _dropdown_form_field_p(default_index: int = 0) -> CT_P:
    p = cast(CT_P, element("w:p"))
    paragraph = Paragraph(p, None)  # type: ignore[arg-type]
    paragraph.add_dropdown_form_field(
        name="Country", options=["US", "UK", "AU"], default_index=default_index
    )
    return p


def _begin_run(p: CT_P):
    return p.xpath("./w:r[w:fldChar[@w:fldCharType='begin' and w:ffData]]")[0]


# --- enum tests ----------------------------------------------------------


class DescribeWD_FORM_FIELD_TYPE:
    """Sanity check for the enum members."""

    @pytest.mark.parametrize(
        "member,value",
        [
            ("TEXT", "text"),
            ("CHECKBOX", "checkbox"),
            ("DROPDOWN", "dropdown"),
        ],
    )
    def it_exposes_the_three_form_field_types(self, member, value):
        assert getattr(WD_FORM_FIELD_TYPE, member).value == value


# --- FormField read tests -----------------------------------------------


class DescribeFormField_Text:
    """Reading a text form field."""

    def it_reports_its_type(self):
        p = _text_form_field_p()
        ff = FormField(_begin_run(p))
        assert ff.type is WD_FORM_FIELD_TYPE.TEXT

    def it_reports_its_name(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        assert ff.name == "Text1"

    def it_reports_enabled_true_by_default(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        assert ff.enabled is True

    def it_reports_calc_on_exit_false_by_default(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        assert ff.calc_on_exit is False

    def it_exposes_a_text_input_view(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        ti = ff.text_input
        assert isinstance(ti, TextInputFormField)
        assert ti.default == "Hello"
        assert ti.max_length == 20

    def it_returns_None_for_other_typed_views(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        assert ff.checkbox is None
        assert ff.dropdown is None

    def its_value_is_the_rendered_result_text(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        assert ff.value == "Hello"


class DescribeFormField_Checkbox:
    """Reading a checkbox form field."""

    def it_reports_its_type(self):
        ff = FormField(_begin_run(_checkbox_form_field_p(checked=True)))
        assert ff.type is WD_FORM_FIELD_TYPE.CHECKBOX

    def it_exposes_a_checkbox_view(self):
        ff = FormField(_begin_run(_checkbox_form_field_p(checked=True)))
        cb = ff.checkbox
        assert isinstance(cb, CheckboxFormField)
        assert cb.checked is True
        assert cb.default is True

    def its_value_toggles_with_checked(self):
        ff_on = FormField(_begin_run(_checkbox_form_field_p(checked=True)))
        ff_off = FormField(_begin_run(_checkbox_form_field_p(checked=False)))
        assert ff_on.value is True
        assert ff_off.value is False

    def it_returns_None_for_other_typed_views(self):
        ff = FormField(_begin_run(_checkbox_form_field_p()))
        assert ff.text_input is None
        assert ff.dropdown is None


class DescribeFormField_Dropdown:
    """Reading a dropdown form field."""

    def it_reports_its_type(self):
        ff = FormField(_begin_run(_dropdown_form_field_p()))
        assert ff.type is WD_FORM_FIELD_TYPE.DROPDOWN

    def it_exposes_a_dropdown_view(self):
        ff = FormField(_begin_run(_dropdown_form_field_p(default_index=1)))
        dd = ff.dropdown
        assert isinstance(dd, DropdownFormField)
        assert dd.options == ["US", "UK", "AU"]
        assert dd.default_index == 1
        assert dd.result_index == 1

    def its_value_is_the_selected_option(self):
        ff = FormField(_begin_run(_dropdown_form_field_p(default_index=2)))
        assert ff.value == "AU"

    def its_value_is_empty_string_when_index_is_out_of_range(self):
        # -- manually craft ddList with no entries; result_index defaults to 0
        p = cast(CT_P, element("w:p"))
        Paragraph(p, None).add_dropdown_form_field(  # type: ignore[arg-type]
            name="Empty", options=[], default_index=0
        )
        ff = FormField(_begin_run(p))
        assert ff.value == ""


class DescribeFormField_HelpAndStatus:
    """Help and status text."""

    def it_reads_helpText_and_statusText(self):
        # -- build an ffData block with helpText and statusText via cxml
        xml_str = (
            "w:p/w:r/w:fldChar{w:fldCharType=begin}"
            "/w:ffData/("
            "w:name{w:val=T1}"
            ",w:enabled"
            ",w:helpText{w:val=Type your name}"
            ",w:statusText{w:val=Name field}"
            ",w:textInput"
            ")"
        )
        p = cast(CT_P, element(xml_str))
        begin_run = p.xpath("./w:r")[0]
        ff = FormField(begin_run)
        assert ff.help_text == "Type your name"
        assert ff.status_text == "Name field"


# --- FormField builders ----------------------------------------------------


class DescribeParagraph_add_text_form_field:
    """`paragraph.add_text_form_field()` structure."""

    def it_emits_the_begin_separate_end_run_sequence_with_ffData(self):
        p = _text_form_field_p()

        runs = p.r_lst
        # -- five runs: begin, instrText, separate, result, end --
        assert len(runs) == 5
        fld_begin = runs[0].xpath("./w:fldChar")[0]
        assert fld_begin.get(qn("w:fldCharType")) == "begin"
        assert fld_begin.xpath("./w:ffData") != []
        assert runs[1].xpath("./w:instrText")[0].text == " FORMTEXT "
        assert runs[2].xpath("./w:fldChar")[0].get(qn("w:fldCharType")) == "separate"
        assert runs[3].xpath("./w:t")[0].text == "Hello"
        assert runs[4].xpath("./w:fldChar")[0].get(qn("w:fldCharType")) == "end"

    def it_returns_a_form_field_proxy(self):
        p = cast(CT_P, element("w:p"))
        ff = Paragraph(p, None).add_text_form_field(  # type: ignore[arg-type]
            name="Text1", default="hi"
        )
        assert isinstance(ff, FormField)
        assert ff.type is WD_FORM_FIELD_TYPE.TEXT
        assert ff.name == "Text1"

    def it_writes_maxlength_zero_when_no_limit(self):
        p = cast(CT_P, element("w:p"))
        Paragraph(p, None).add_text_form_field(  # type: ignore[arg-type]
            name="Text1"
        )
        maxLength = p.xpath(".//w:textInput/w:maxLength")[0]
        assert maxLength.get(qn("w:val")) == "0"


class DescribeParagraph_add_checkbox_form_field:
    """`paragraph.add_checkbox_form_field()` structure."""

    def it_writes_the_expected_ffData_shape(self):
        p = _checkbox_form_field_p(checked=True)

        assert p.xpath(".//w:instrText")[0].text == " FORMCHECKBOX "
        cb = p.xpath(".//w:ffData/w:checkBox")[0]
        assert cb.xpath("./w:default")[0].get(qn("w:val")) == "1"
        assert cb.xpath("./w:checked")[0].get(qn("w:val")) == "1"

    def it_defaults_checked_to_false(self):
        p = cast(CT_P, element("w:p"))
        Paragraph(p, None).add_checkbox_form_field(  # type: ignore[arg-type]
            name="Agree"
        )
        cb = p.xpath(".//w:ffData/w:checkBox")[0]
        assert cb.xpath("./w:default")[0].get(qn("w:val")) == "0"
        assert cb.xpath("./w:checked")[0].get(qn("w:val")) == "0"


class DescribeParagraph_add_dropdown_form_field:
    """`paragraph.add_dropdown_form_field()` structure."""

    def it_writes_the_expected_ffData_shape(self):
        p = _dropdown_form_field_p(default_index=1)

        assert p.xpath(".//w:instrText")[0].text == " FORMDROPDOWN "
        dd = p.xpath(".//w:ffData/w:ddList")[0]
        assert dd.xpath("./w:result")[0].get(qn("w:val")) == "1"
        assert dd.xpath("./w:default")[0].get(qn("w:val")) == "1"
        entries = [le.get(qn("w:val")) for le in dd.xpath("./w:listEntry")]
        assert entries == ["US", "UK", "AU"]

    def its_initial_result_text_is_the_selected_option(self):
        p = _dropdown_form_field_p(default_index=2)
        result_run = p.r_lst[3]
        assert result_run.xpath("./w:t")[0].text == "AU"


# --- iteration --------------------------------------------------------------


class DescribeParagraph_form_fields:
    """`paragraph.form_fields` iteration."""

    def it_yields_one_form_field_per_ffData_begin(self):
        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_text_form_field(name="T1", default="a")
        para.add_checkbox_form_field(name="C1", checked=True)
        para.add_dropdown_form_field(
            name="D1", options=["x", "y"], default_index=0
        )
        fields = para.form_fields
        assert len(fields) == 3
        assert [ff.type for ff in fields] == [
            WD_FORM_FIELD_TYPE.TEXT,
            WD_FORM_FIELD_TYPE.CHECKBOX,
            WD_FORM_FIELD_TYPE.DROPDOWN,
        ]

    def it_ignores_complex_fields_without_ffData(self):
        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        # -- a plain complex field (no ffData) should not appear in form_fields
        para.add_complex_field("PAGE", "1")
        para.add_text_form_field(name="T1", default="a")
        fields = para.form_fields
        assert len(fields) == 1
        assert fields[0].name == "T1"


class DescribeDocument_form_fields:
    """`document.form_fields` iteration over body paragraphs."""

    def it_walks_body_paragraphs(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element("w:document/w:body/(w:p,w:p)"),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        # -- add a form field to each body paragraph
        for paragraph in doc.paragraphs:
            paragraph.add_text_form_field(name=f"FF_{id(paragraph)}")

        fields = doc.form_fields
        assert len(fields) == 2
        assert all(ff.type is WD_FORM_FIELD_TYPE.TEXT for ff in fields)


# --- typed subclasses -----------------------------------------------------


class DescribeFormField_proxy_for:
    """`FormField.proxy_for` dispatches on form-field type."""

    def it_returns_TextInputField_for_a_text_field(self):
        p = _text_form_field_p()
        ff = FormField.proxy_for(_begin_run(p))
        assert isinstance(ff, TextInputField)
        assert isinstance(ff, FormField)  # -- still passes the base check

    def it_returns_CheckBoxField_for_a_checkbox_field(self):
        p = _checkbox_form_field_p(checked=True)
        ff = FormField.proxy_for(_begin_run(p))
        assert isinstance(ff, CheckBoxField)

    def it_returns_DropDownListField_for_a_dropdown_field(self):
        p = _dropdown_form_field_p(default_index=0)
        ff = FormField.proxy_for(_begin_run(p))
        assert isinstance(ff, DropDownListField)


class DescribeParagraph_form_fields_subclasses:
    """`Paragraph.form_fields` returns the right subclass per field."""

    def it_surfaces_typed_subclasses(self):
        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_text_form_field(name="T1")
        para.add_checkbox_form_field(name="C1")
        para.add_dropdown_form_field(
            name="D1", options=["a", "b"], default_index=0
        )
        fields = para.form_fields
        assert isinstance(fields[0], TextInputField)
        assert isinstance(fields[1], CheckBoxField)
        assert isinstance(fields[2], DropDownListField)


# --- unified authoring ----------------------------------------------------


class DescribeParagraph_add_form_field:
    """`Paragraph.add_form_field(kind, ...)` dispatches correctly."""

    @pytest.mark.parametrize(
        "kind,expected_type,expected_cls",
        [
            ("text", WD_FORM_FIELD_TYPE.TEXT, TextInputField),
            ("checkbox", WD_FORM_FIELD_TYPE.CHECKBOX, CheckBoxField),
            ("FORMTEXT", WD_FORM_FIELD_TYPE.TEXT, TextInputField),
            (WD_FORM_FIELD_TYPE.DROPDOWN, WD_FORM_FIELD_TYPE.DROPDOWN,
             DropDownListField),
        ],
    )
    def it_accepts_string_and_enum_kinds(
        self, kind, expected_type, expected_cls
    ):
        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        kwargs = {}
        if expected_type is WD_FORM_FIELD_TYPE.DROPDOWN:
            kwargs["options"] = ["a", "b"]
        ff = para.add_form_field(kind, "N1", **kwargs)
        assert isinstance(ff, expected_cls)
        assert ff.type is expected_type
        assert ff.name == "N1"

    def it_forwards_text_kwargs(self):
        p = cast(CT_P, element("w:p"))
        ff = Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
            "text", "T1", default="d", maxlength=5, type_="number",
            format="0.00",
        )
        ti = ff.text_input
        assert ti is not None
        assert ti.default == "d"
        assert ti.max_length == 5
        assert ti.type == "number"
        assert ti.format == "0.00"

    def it_forwards_checkbox_size_kwargs(self):
        p = cast(CT_P, element("w:p"))
        ff = Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
            "checkbox", "C1", checked=True, size=40,
        )
        cb = ff.checkbox
        assert cb is not None
        assert cb.size == 40
        assert cb.size_auto is False

    def it_requires_options_for_dropdown(self):
        p = cast(CT_P, element("w:p"))
        with pytest.raises(TypeError, match="options"):
            Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
                "dropdown", "D1"
            )

    def it_rejects_unknown_kwargs(self):
        p = cast(CT_P, element("w:p"))
        with pytest.raises(TypeError, match="unexpected keyword"):
            Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
                "text", "T1", bogus=True
            )

    def it_rejects_unknown_kind(self):
        p = cast(CT_P, element("w:p"))
        with pytest.raises(ValueError, match="kind must be one of"):
            Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
                "spinner", "S1"
            )


# --- text input type and format -------------------------------------------


class DescribeTextInputFormField_type:
    """`TextInputFormField.type` and `.format`."""

    def its_type_defaults_to_regular(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        ti = ff.text_input
        assert ti is not None
        assert ti.type == "regular"

    @pytest.mark.parametrize(
        "type_",
        ["regular", "number", "date", "currentTime", "currentDate",
         "calculated"],
    )
    def it_round_trips_each_legal_type(self, type_):
        p = cast(CT_P, element("w:p"))
        Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
            "text", "T1", type_=type_,
        )
        ff = FormField(_begin_run(p))
        ti = ff.text_input
        assert ti is not None
        assert ti.type == type_

    def it_rejects_unknown_types(self):
        p = cast(CT_P, element("w:p"))
        with pytest.raises(ValueError, match="type_ must be one of"):
            Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
                "text", "T1", type_="crazy",
            )


# --- checkbox size --------------------------------------------------------


class DescribeCheckboxFormField_size:
    """`CheckboxFormField.size` / `.size_auto`."""

    def it_emits_sizeAuto_by_default(self):
        p = _checkbox_form_field_p(checked=False)
        assert p.xpath(".//w:checkBox/w:sizeAuto") != []
        ff = FormField(_begin_run(p))
        cb = ff.checkbox
        assert cb is not None
        assert cb.size_auto is True
        assert cb.size is None

    def it_emits_size_when_explicit(self):
        p = cast(CT_P, element("w:p"))
        Paragraph(p, None).add_form_field(  # type: ignore[arg-type]
            "checkbox", "C1", size=48,
        )
        assert p.xpath(".//w:checkBox/w:size") != []
        assert p.xpath(".//w:checkBox/w:sizeAuto") == []
        ff = FormField(_begin_run(p))
        cb = ff.checkbox
        assert cb is not None
        assert cb.size == 48
        assert cb.size_auto is False


# --- current_value alias --------------------------------------------------


class DescribeFormField_current_value:
    """`.current_value` aliases `.value`."""

    def it_returns_the_rendered_result_for_text(self):
        ff = FormField(_begin_run(_text_form_field_p()))
        assert ff.current_value == ff.value == "Hello"

    def it_returns_a_bool_for_checkbox(self):
        ff = FormField(_begin_run(_checkbox_form_field_p(checked=True)))
        assert ff.current_value is True

    def it_returns_the_selected_option_for_dropdown(self):
        ff = FormField(_begin_run(_dropdown_form_field_p(default_index=1)))
        assert ff.current_value == "UK"


# --- round-trip -----------------------------------------------------------


class DescribeFormField_round_trip:
    """Author, save, reopen, verify."""

    def it_preserves_all_three_form_field_types(self):
        doc = Document()
        p1 = doc.add_paragraph("Name: ")
        p1.add_form_field(
            "text", "FullName", default="Anon", maxlength=40, type_="regular",
        )
        p2 = doc.add_paragraph("Subscribe: ")
        p2.add_form_field("checkbox", "Sub", checked=True, size=32)
        p3 = doc.add_paragraph("Pick: ")
        p3.add_form_field(
            "dropdown", "Pick",
            options=["Red", "Green", "Blue"], default_index=2,
        )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reopened = Document(buf)

        fields = reopened.form_fields
        assert [ff.type for ff in fields] == [
            WD_FORM_FIELD_TYPE.TEXT,
            WD_FORM_FIELD_TYPE.CHECKBOX,
            WD_FORM_FIELD_TYPE.DROPDOWN,
        ]
        assert isinstance(fields[0], TextInputField)
        assert isinstance(fields[1], CheckBoxField)
        assert isinstance(fields[2], DropDownListField)

        # -- text --
        t = fields[0].text_input
        assert t is not None
        assert fields[0].name == "FullName"
        assert t.default == "Anon"
        assert t.max_length == 40
        assert fields[0].current_value == "Anon"

        # -- checkbox --
        cb = fields[1].checkbox
        assert cb is not None
        assert cb.checked is True
        assert cb.size == 32
        assert cb.size_auto is False

        # -- dropdown --
        dd = fields[2].dropdown
        assert dd is not None
        assert dd.options == ["Red", "Green", "Blue"]
        assert fields[2].current_value == "Blue"


# --- transition to SDT ----------------------------------------------------


class DescribeFormField_to_sdt:
    """`FormField.to_sdt()` migrates a legacy form field to an SDT."""

    def it_replaces_text_form_field_with_plain_text_sdt(self):
        from docx.content_controls import ContentControlType

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_form_field(
            "text", "FullName", default="Anon", maxlength=20,
        )
        assert len(p.r_lst) == 5  # -- 5-run complex field

        ff = para.form_fields[0]
        sdt = ff.to_sdt()

        # -- 5-run sequence gone, single w:sdt remains --
        assert p.r_lst == []
        assert len(p.xpath("./w:sdt")) == 1
        assert sdt.type is ContentControlType.PLAIN_TEXT
        assert sdt.tag == "FullName"
        # -- seed text made it through --
        assert p.xpath(".//w:sdtContent/w:r/w:t/text()") == ["Anon"]

    def it_replaces_checkbox_form_field_with_checkbox_sdt(self):
        from docx.content_controls import ContentControlType

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_form_field("checkbox", "Agree", checked=True)

        ff = para.form_fields[0]
        sdt = ff.to_sdt()

        assert sdt.type is ContentControlType.CHECKBOX
        assert sdt.tag == "Agree"
        # -- w14:checkbox/w14:checked reflects the legacy checked state --
        checked_els = p.xpath(".//w:sdtPr/w14:checkbox/w14:checked")
        assert len(checked_els) == 1
        assert checked_els[0].get(qn("w14:val")) == "1"

    def it_replaces_dropdown_form_field_with_dropdown_sdt(self):
        from docx.content_controls import ContentControlType

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_form_field(
            "dropdown", "Country",
            options=["US", "UK", "AU"], default_index=1,
        )

        ff = para.form_fields[0]
        sdt = ff.to_sdt()

        assert sdt.type is ContentControlType.DROPDOWN
        assert sdt.tag == "Country"
        # -- one w:listItem per legacy entry, in order --
        items = p.xpath(".//w:sdtPr/w:dropDownList/w:listItem")
        assert [li.get(qn("w:value")) for li in items] == ["US", "UK", "AU"]
        # -- seed text is the selected option --
        assert p.xpath(".//w:sdtContent/w:r/w:t/text()") == ["UK"]

    def it_maps_helpText_to_sdt_alias(self):
        # -- build a form field with helpText via cxml, then migrate --
        xml_str = (
            "w:p/w:r/w:fldChar{w:fldCharType=begin}"
            "/w:ffData/("
            "w:name{w:val=T1}"
            ",w:enabled"
            ",w:helpText{w:val=Type your name}"
            ",w:textInput"
            ")"
        )
        p = cast(CT_P, element(xml_str))
        # -- manually add separate/end runs so to_sdt() can find them
        separate_r = element("w:r/w:fldChar{w:fldCharType=separate}")
        end_r = element("w:r/w:fldChar{w:fldCharType=end}")
        p.append(separate_r)
        p.append(end_r)

        begin_run = p.xpath("./w:r")[0]
        ff = FormField(begin_run)
        sdt = ff.to_sdt()
        assert sdt.title == "Type your name"
