"""Unit-test suite for the table-of-contents helpers."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.fields import Field
from docx.text.paragraph import Paragraph
from docx.toc import (
    _collect_entries,
    _paragraph_heading_level,
    _render_result_text,
    _validate_levels,
    build_toc_instruction,
    populate_toc_paragraph,
)


class DescribeBuildTocInstruction:
    """Unit-test suite for `docx.toc.build_toc_instruction`."""

    def it_builds_the_default_TOC_instruction_for_levels_1_3(self):
        assert build_toc_instruction() == ' TOC \\o "1-3" \\h \\z \\u '

    def it_honors_a_custom_level_range(self):
        assert build_toc_instruction((2, 5)) == ' TOC \\o "2-5" \\h \\z \\u '

    def it_accepts_a_single_level_range(self):
        assert build_toc_instruction((1, 1)) == ' TOC \\o "1-1" \\h \\z \\u '

    def it_raises_on_a_single_integer(self):
        with pytest.raises(ValueError, match="levels must be a 2-tuple"):
            build_toc_instruction(3)  # type: ignore[arg-type]

    def it_raises_when_min_level_is_below_1(self):
        with pytest.raises(ValueError, match="1 <= min_level"):
            build_toc_instruction((0, 3))

    def it_raises_when_max_level_is_above_9(self):
        with pytest.raises(ValueError, match="<= 9"):
            build_toc_instruction((1, 10))

    def it_raises_when_min_exceeds_max(self):
        with pytest.raises(ValueError, match="min_level <= max_level"):
            build_toc_instruction((4, 2))


class DescribeValidateLevels:
    """Unit-test suite for `docx.toc._validate_levels`."""

    def it_returns_a_valid_tuple_unchanged(self):
        assert _validate_levels((1, 3)) == (1, 3)

    def it_accepts_a_list_of_two_ints(self):
        # -- lists are tolerated for caller convenience; the type hint says
        #    tuple, but unpacking works equally on both. --
        assert _validate_levels([1, 3]) == (1, 3)  # type: ignore[arg-type]

    def it_rejects_a_3_element_tuple(self):
        with pytest.raises(ValueError):
            _validate_levels((1, 2, 3))  # type: ignore[arg-type]

    def it_rejects_non_integer_members(self):
        with pytest.raises(ValueError):
            _validate_levels(("1", "3"))  # type: ignore[arg-type]


class DescribeParagraphHeadingLevel:
    """Unit-test suite for the internal `_paragraph_heading_level` helper."""

    def it_returns_the_level_for_a_heading_paragraph(self):
        doc: DocumentCls = Document()
        p = doc.add_heading("hello", level=2)
        assert _paragraph_heading_level(p) == 2

    def it_returns_None_for_a_non_heading_paragraph(self):
        doc: DocumentCls = Document()
        p = doc.add_paragraph("body text")
        assert _paragraph_heading_level(p) is None

    def it_matches_style_names_case_insensitively(self):
        doc: DocumentCls = Document()
        p = doc.add_paragraph("hello")
        # -- Paragraph styles are stored by name; we assign a style whose
        #    name matches the regex regardless of casing to prove the
        #    helper's regex is case-insensitive. --
        p.style = "Heading 5"
        assert _paragraph_heading_level(p) == 5


class DescribeCollectEntries:
    """Unit-test suite for the internal `_collect_entries` helper."""

    def it_collects_only_headings_within_the_range(self):
        doc: DocumentCls = Document()
        doc.add_paragraph("body text")  # skipped: not a heading
        doc.add_heading("H1 text", level=1)
        doc.add_heading("H2 text", level=2)
        doc.add_heading("H4 text", level=4)  # skipped: out of range

        entries = _collect_entries(doc.paragraphs, (1, 3))
        assert entries == [(1, "H1 text"), (2, "H2 text")]

    def it_returns_an_empty_list_when_no_headings_are_in_range(self):
        doc: DocumentCls = Document()
        doc.add_heading("H5 text", level=5)
        assert _collect_entries(doc.paragraphs, (1, 3)) == []


class DescribeRenderResultText:
    """Unit-test suite for the internal `_render_result_text` helper."""

    def it_renders_one_line_per_entry_with_tab_and_index(self):
        result = _render_result_text([(1, "First"), (2, "Second"), (1, "Third")])
        assert result == "First\t1\nSecond\t2\nThird\t3"

    def it_returns_an_empty_string_for_no_entries(self):
        assert _render_result_text([]) == ""


class DescribeDocumentAddTableOfContents:
    """Unit-test suite for `Document.add_table_of_contents`."""

    def it_appends_a_TOC_paragraph_to_the_body(self):
        document: DocumentCls = Document()
        start_count = len(document.paragraphs)

        paragraph = document.add_table_of_contents()

        assert isinstance(paragraph, Paragraph)
        assert len(document.paragraphs) == start_count + 1
        assert document.paragraphs[-1]._p is paragraph._p

    def it_builds_a_TOC_complex_field_with_the_default_instruction(self):
        document: DocumentCls = Document()
        document.add_heading("Intro", level=1)

        paragraph = document.add_table_of_contents()

        assert len(paragraph.fields) == 1
        field = paragraph.fields[0]
        assert isinstance(field, Field)
        assert field.is_complex is True
        assert field.type == "TOC"
        assert field.instruction.strip() == 'TOC \\o "1-3" \\h \\z \\u'

    def it_filters_headings_by_the_level_range(self):
        document: DocumentCls = Document()
        document.add_heading("H1 text", level=1)
        document.add_heading("H2 text", level=2)
        document.add_heading("H3 text", level=3)
        document.add_heading("H4 text", level=4)

        paragraph = document.add_table_of_contents(levels=(1, 3))

        result = paragraph.fields[0].result_text
        assert "H1 text" in result
        assert "H2 text" in result
        assert "H3 text" in result
        assert "H4 text" not in result

    def it_honors_a_custom_level_range(self):
        document: DocumentCls = Document()
        document.add_heading("H1 text", level=1)
        document.add_heading("H2 text", level=2)

        paragraph = document.add_table_of_contents(levels=(2, 3))

        result = paragraph.fields[0].result_text
        assert "H1 text" not in result
        assert "H2 text" in result
        assert paragraph.fields[0].instruction.strip() == 'TOC \\o "2-3" \\h \\z \\u'

    def it_handles_an_empty_document_with_an_empty_TOC(self):
        document: DocumentCls = Document()

        paragraph = document.add_table_of_contents()

        field = paragraph.fields[0]
        assert field.type == "TOC"
        assert field.result_text == ""

    def it_does_not_include_itself_in_the_TOC_preview(self):
        document: DocumentCls = Document()
        document.add_heading("Only heading", level=1)

        paragraph = document.add_table_of_contents()

        # -- exactly one heading entry, not two (the empty TOC paragraph
        #    itself is not reprocessed after being appended) --
        result = paragraph.fields[0].result_text
        assert result.count("Only heading") == 1

    def it_renders_each_entry_as_text_tab_index(self):
        document: DocumentCls = Document()
        document.add_heading("Alpha", level=1)
        document.add_heading("Beta", level=2)

        paragraph = document.add_table_of_contents()

        result = paragraph.fields[0].result_text
        assert result == "Alpha\t1\nBeta\t2"

    def it_raises_for_a_bad_level_range(self):
        document: DocumentCls = Document()
        with pytest.raises(ValueError):
            document.add_table_of_contents(levels=(0, 3))


class DescribeParagraphInsertTableOfContents:
    """Unit-test suite for `Paragraph.insert_table_of_contents_before/after`."""

    def it_inserts_a_TOC_before_the_target_paragraph(self):
        document: DocumentCls = Document()
        document.add_heading("First", level=1)
        document.add_heading("Second", level=2)
        anchor = document.add_paragraph("after TOC")

        toc = anchor.insert_table_of_contents_before()

        paragraphs = document.paragraphs
        anchor_idx = paragraphs.index(
            next(p for p in paragraphs if p._p is anchor._p)
        )
        assert paragraphs[anchor_idx - 1]._p is toc._p
        field = toc.fields[0]
        assert field.type == "TOC"
        assert "First\t1" in field.result_text
        assert "Second\t2" in field.result_text

    def it_inserts_a_TOC_after_the_target_paragraph(self):
        document: DocumentCls = Document()
        anchor = document.add_paragraph("title page")
        document.add_heading("First", level=1)

        toc = anchor.insert_table_of_contents_after()

        paragraphs = document.paragraphs
        anchor_idx = paragraphs.index(
            next(p for p in paragraphs if p._p is anchor._p)
        )
        assert paragraphs[anchor_idx + 1]._p is toc._p
        assert toc.fields[0].type == "TOC"

    def it_scans_headings_added_after_the_anchor(self):
        document: DocumentCls = Document()
        anchor = document.add_paragraph("title")
        document.add_heading("Later heading", level=1)

        toc = anchor.insert_table_of_contents_before()

        assert "Later heading" in toc.fields[0].result_text

    def it_honors_a_custom_level_range_on_insert_after(self):
        document: DocumentCls = Document()
        document.add_heading("H1", level=1)
        document.add_heading("H2", level=2)
        anchor = document.add_paragraph("anchor")

        toc = anchor.insert_table_of_contents_after(levels=(2, 2))

        result = toc.fields[0].result_text
        assert "H1" not in result
        assert "H2" in result


class DescribePopulateTocParagraph:
    """Unit-test suite for the low-level `populate_toc_paragraph` helper."""

    def it_populates_an_empty_paragraph_with_a_TOC_field(self):
        document: DocumentCls = Document()
        document.add_heading("H1", level=1)
        paragraph = document.add_paragraph()

        result = populate_toc_paragraph(paragraph, document.paragraphs)

        assert result is paragraph
        assert len(paragraph.fields) == 1
        assert paragraph.fields[0].type == "TOC"

    def it_returns_the_populated_paragraph(self):
        document: DocumentCls = Document()
        paragraph = document.add_paragraph()

        result = populate_toc_paragraph(
            paragraph, document.paragraphs, levels=(1, 5)
        )

        assert result is paragraph
        assert paragraph.fields[0].instruction.strip() == 'TOC \\o "1-5" \\h \\z \\u'
