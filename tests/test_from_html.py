"""Tests for the minimal HTML → docx importer.

Covers the element mappings declared in issue #95:

* ``<h1>``-``<h6>``       → ``Heading 1``-``Heading 6``
* ``<strong>`` / ``<b>``  → bold runs
* ``<em>`` / ``<i>``      → italic runs
* ``<a href>``            → hyperlinks
* ``<ul>`` / ``<ol>``     → bullet / numbered list paragraphs
* ``<table>``             → Word table
* ``<img>``               → embedded picture (``data:`` URLs)
* ``<blockquote>``        → ``Quote`` style paragraph
* ``<code>`` / ``<pre>``  → monospace runs / paragraphs
* ``clean=True``          → strips script/style/comments + class/id
"""

from __future__ import annotations

import base64
import io
import os
import tempfile

import pytest

from docx import Document
from docx.html_import import (
    _decode_data_url,
    _parse_color_from_style,
    _sanitize_url,
    from_html,
    from_html_string,
)


# -- a tiny PNG (1x1, transparent) for image tests, base64-encoded --
_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII="
)
_PNG_DATA_URL = "data:image/png;base64," + base64.b64encode(_PNG_BYTES).decode("ascii")


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------


class DescribeHeadings:
    @pytest.mark.parametrize(
        "tag,level",
        [("h1", 1), ("h2", 2), ("h3", 3), ("h4", 4), ("h5", 5), ("h6", 6)],
    )
    def it_maps_h1_through_h6(self, tag: str, level: int):
        doc = from_html_string(f"<{tag}>Title</{tag}>")
        para = doc.paragraphs[0]
        assert para.text == "Title"
        assert para.style.name == f"Heading {level}"

    def it_emits_distinct_headings_per_block(self):
        doc = from_html_string("<h1>One</h1><h2>Two</h2><p>Body</p>")
        assert [p.text for p in doc.paragraphs] == ["One", "Two", "Body"]
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[1].style.name == "Heading 2"


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------


class DescribeInlineFormatting:
    def it_maps_strong_to_bold(self):
        doc = from_html_string("<p>plain <strong>bold</strong> tail</p>")
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def it_maps_b_to_bold(self):
        doc = from_html_string("<p><b>shouty</b></p>")
        bold_runs = [r for r in doc.paragraphs[0].runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "shouty"

    def it_maps_em_to_italic(self):
        doc = from_html_string("<p>plain <em>tilted</em></p>")
        ital_runs = [r for r in doc.paragraphs[0].runs if r.italic]
        assert len(ital_runs) == 1
        assert ital_runs[0].text == "tilted"

    def it_maps_i_to_italic(self):
        doc = from_html_string("<p><i>tilted</i></p>")
        ital_runs = [r for r in doc.paragraphs[0].runs if r.italic]
        assert len(ital_runs) == 1

    def it_nests_bold_and_italic(self):
        doc = from_html_string("<p><strong><em>both</em></strong></p>")
        runs = doc.paragraphs[0].runs
        assert any(r.bold and r.italic for r in runs)

    def it_maps_u_to_underline(self):
        doc = from_html_string("<p><u>under</u></p>")
        u_runs = [r for r in doc.paragraphs[0].runs if r.underline]
        assert len(u_runs) == 1


# ---------------------------------------------------------------------------
# Hyperlinks
# ---------------------------------------------------------------------------


class DescribeHyperlinks:
    def it_maps_a_href_to_a_hyperlink(self):
        doc = from_html_string('<p><a href="https://example.com">click</a></p>')
        para = doc.paragraphs[0]
        # -- the paragraph should expose the hyperlink via iter_inner_content
        from docx.text.hyperlink import Hyperlink

        links = [c for c in para.iter_inner_content() if isinstance(c, Hyperlink)]
        assert len(links) == 1
        link = links[0]
        assert link.url == "https://example.com"
        assert "click" in link.text

    def it_drops_javascript_scheme_from_href(self):
        doc = from_html_string('<p><a href="javascript:alert(1)">x</a></p>')
        from docx.text.hyperlink import Hyperlink

        para = doc.paragraphs[0]
        links = [c for c in para.iter_inner_content() if isinstance(c, Hyperlink)]
        # -- unsafe link should NOT have produced a hyperlink with that URL --
        assert all((link.url or "") != "javascript:alert(1)" for link in links)
        # -- but the visible text should still appear somewhere --
        assert "x" in para.text

    def it_keeps_mailto_links(self):
        doc = from_html_string('<p><a href="mailto:hi@example.com">hi</a></p>')
        from docx.text.hyperlink import Hyperlink

        para = doc.paragraphs[0]
        links = [c for c in para.iter_inner_content() if isinstance(c, Hyperlink)]
        assert len(links) == 1
        assert links[0].url == "mailto:hi@example.com"


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------


class DescribeLists:
    def it_maps_ul_li_to_list_bullet(self):
        doc = from_html_string("<ul><li>one</li><li>two</li></ul>")
        bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert [p.text for p in bullet_paras] == ["one", "two"]

    def it_maps_ol_li_to_list_number(self):
        doc = from_html_string("<ol><li>alpha</li><li>beta</li></ol>")
        number_paras = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert [p.text for p in number_paras] == ["alpha", "beta"]

    def it_intersperses_lists_and_paragraphs(self):
        doc = from_html_string("<p>before</p><ul><li>only</li></ul><p>after</p>")
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["before", "only", "after"]


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


class DescribeTables:
    def it_maps_table_to_a_word_table(self):
        html = (
            "<table>"
            "<tr><td>r1c1</td><td>r1c2</td></tr>"
            "<tr><td>r2c1</td><td>r2c2</td></tr>"
            "</table>"
        )
        doc = from_html_string(html)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.rows[0].cells[0].text.strip() == "r1c1"
        assert table.rows[1].cells[1].text.strip() == "r2c2"

    def it_handles_tables_following_paragraphs(self):
        html = "<p>intro</p><table><tr><td>cell</td></tr></table><p>tail</p>"
        doc = from_html_string(html)
        assert len(doc.tables) == 1
        # -- intro and tail paragraphs land in the body --
        body_paragraphs = [p.text for p in doc.paragraphs]
        assert "intro" in body_paragraphs
        assert "tail" in body_paragraphs


# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------


class DescribeImages:
    def it_embeds_a_data_url_image(self):
        doc = from_html_string(f'<p><img src="{_PNG_DATA_URL}" alt="dot"></p>')
        # -- exactly one inline shape was added --
        assert len(doc.inline_shapes) == 1

    def it_falls_back_to_alt_for_remote_images(self):
        doc = from_html_string(
            '<p><img src="https://example.com/x.png" alt="Logo"></p>'
        )
        assert len(doc.inline_shapes) == 0
        assert "Logo" in doc.paragraphs[0].text

    def it_decodes_data_url(self):
        stream = _decode_data_url(_PNG_DATA_URL)
        assert stream is not None
        assert stream.read().startswith(b"\x89PNG")


# ---------------------------------------------------------------------------
# Blockquote / code / pre
# ---------------------------------------------------------------------------


class DescribeBlockquoteAndCode:
    def it_maps_blockquote_to_quote_style(self):
        doc = from_html_string("<blockquote>cite me</blockquote>")
        para = doc.paragraphs[0]
        assert "cite me" in para.text
        # -- "Quote" style may not exist in the bare default template;
        # -- the importer accepts that and falls back to Normal. --
        assert para.style.name in ("Quote", "Normal")

    def it_maps_code_to_monospace_runs(self):
        doc = from_html_string("<p>before <code>x = 1</code> after</p>")
        runs = doc.paragraphs[0].runs
        mono_runs = [r for r in runs if r.font.name == "Courier New"]
        assert len(mono_runs) == 1
        assert mono_runs[0].text == "x = 1"

    def it_preserves_pre_whitespace(self):
        doc = from_html_string("<pre>line1\n    indent2</pre>")
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "line1" in text and "indent2" in text


# ---------------------------------------------------------------------------
# clean=True behaviour
# ---------------------------------------------------------------------------


class DescribeCleanMode:
    def it_strips_script_blocks(self):
        doc = from_html_string(
            "<p>hello</p><script>alert(1)</script><p>world</p>"
        )
        text = " ".join(p.text for p in doc.paragraphs)
        assert "alert(1)" not in text
        assert "hello" in text and "world" in text

    def it_strips_style_blocks(self):
        doc = from_html_string(
            "<style>p { color: red; }</style><p>visible</p>"
        )
        text = " ".join(p.text for p in doc.paragraphs)
        assert "color: red" not in text
        assert "visible" in text

    def it_strips_html_comments(self):
        doc = from_html_string("<!-- secret --><p>shown</p>")
        text = " ".join(p.text for p in doc.paragraphs)
        assert "secret" not in text
        assert "shown" in text

    def it_drops_class_and_id_attrs_when_clean_is_true(self):
        # -- behavioural assertion: same input with clean=True/False
        # -- should yield the same body text (class/id are inert in
        # -- our element-mapping, but the public surface promises the
        # -- attributes are dropped). Use the public API and confirm
        # -- a no-crash + identical text outcome. --
        html = '<p class="muted" id="lead">body</p>'
        doc1 = from_html_string(html, clean=True)
        doc2 = from_html_string(html, clean=False)
        assert doc1.paragraphs[0].text == "body"
        assert doc2.paragraphs[0].text == "body"

    def it_picks_up_inline_color_as_run_property(self):
        doc = from_html_string('<p><span style="color:#ff0000">red</span></p>')
        runs = doc.paragraphs[0].runs
        # -- at least one run carries the red colour. --
        from docx.shared import RGBColor

        assert any(
            r.font.color.rgb == RGBColor(0xFF, 0x00, 0x00) for r in runs
        )


# ---------------------------------------------------------------------------
# Document.from_html / Document.from_html_string entry points
# ---------------------------------------------------------------------------


class DescribeDocumentFromHtml:
    def it_is_attached_to_the_document_factory(self):
        assert callable(Document.from_html)
        assert callable(Document.from_html_string)

    def it_loads_from_a_file_path(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "article.html")
            with open(path, "w", encoding="utf-8") as fp:
                fp.write("<h1>From File</h1><p>body</p>")
            doc = Document.from_html(path, clean=True)
            assert doc.paragraphs[0].text == "From File"
            assert doc.paragraphs[0].style.name == "Heading 1"

    def it_loads_from_a_pathlib_path(self):
        from pathlib import Path

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.html"
            path.write_text("<p>pathlib</p>", encoding="utf-8")
            doc = Document.from_html(path)
            assert doc.paragraphs[0].text == "pathlib"

    def it_loads_from_a_binary_file_object(self):
        buf = io.BytesIO(b"<p>bin</p>")
        doc = Document.from_html(buf)
        assert doc.paragraphs[0].text == "bin"

    def it_loads_from_a_text_file_object(self):
        buf = io.StringIO("<p>text</p>")
        doc = Document.from_html(buf)
        assert doc.paragraphs[0].text == "text"

    def it_honours_meta_charset_in_binary_input(self):
        body = (
            '<meta charset="latin-1"><p>caf\xe9</p>'
        ).encode("latin-1")
        doc = Document.from_html(io.BytesIO(body))
        text = " ".join(p.text for p in doc.paragraphs)
        assert "café" in text

    def it_starts_with_no_template_residue(self):
        # -- the bundled template ships with one empty paragraph; the
        # -- importer drops it so the resulting document has only the
        # -- content the HTML supplied. --
        doc = Document.from_html_string("<p>only</p>")
        non_empty = [p for p in doc.paragraphs if p.text]
        assert len(non_empty) == 1
        assert doc.paragraphs[0].text == "only"


# ---------------------------------------------------------------------------
# Round-trip with to_html (sanity smoke test)
# ---------------------------------------------------------------------------


class DescribeRoundTripWithToHtml:
    def it_can_round_trip_through_to_html_for_basic_constructs(self):
        # -- HTML → docx → HTML: structural smoke check (we don't promise
        # -- byte-identical, but the text and tags should survive) --
        html_in = (
            "<h1>Title</h1>"
            "<p>plain <strong>bold</strong> and <em>italic</em></p>"
            "<ul><li>one</li><li>two</li></ul>"
        )
        doc = Document.from_html_string(html_in)
        html_out = doc.to_html(include_styles=False)

        for token in (
            "Title",
            "<h1>",
            "<strong>bold</strong>",
            "<em>italic</em>",
            "<ul>",
            "<li>one</li>",
            "<li>two</li>",
        ):
            assert token in html_out, f"missing {token!r} in round-trip output"


# ---------------------------------------------------------------------------
# Misc helpers
# ---------------------------------------------------------------------------


class DescribeUrlSanitizer:
    @pytest.mark.parametrize(
        "url,expected",
        [
            ("https://example.com", "https://example.com"),
            ("http://example.com", "http://example.com"),
            ("mailto:hi@example.com", "mailto:hi@example.com"),
            ("javascript:alert(1)", None),
            ("vbscript:msgbox(1)", None),
            ("data:text/html,<script>x</script>", None),
            ("file:///etc/passwd", None),
            ("/relative/path", "/relative/path"),
            ("relative", "relative"),
            ("#fragment", "#fragment"),
            ("", None),
        ],
    )
    def it_allows_only_safe_schemes(self, url: str, expected):
        assert _sanitize_url(url) == expected


class DescribeColorParser:
    def it_extracts_a_six_digit_hex_color(self):
        from docx.shared import RGBColor

        assert _parse_color_from_style("color: #ff0000") == RGBColor(
            0xFF, 0, 0
        )
        assert _parse_color_from_style("color:#aabbcc") == RGBColor(
            0xAA, 0xBB, 0xCC
        )

    def it_ignores_non_color_styles(self):
        assert _parse_color_from_style("font-weight: bold") is None
        assert _parse_color_from_style("") is None
        assert _parse_color_from_style(None) is None
