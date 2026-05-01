"""Unit-test suite for the `docx.theme` module."""

from __future__ import annotations

from typing import cast

import pytest

import docx
from docx.oxml.theme import CT_Theme
from docx.shared import RGBColor
from docx.theme import Theme, ThemeColors, ThemeFonts

from .unitutil.cxml import element


# -- full Office-style theme fixture, trimmed to the slots exercised by Theme --
_SAMPLE_THEME = (
    "a:theme{name=Office Theme}/a:themeElements/("
    "a:clrScheme{name=Office}/("
    "a:dk1/a:sysClr{val=windowText,lastClr=000000},"
    "a:lt1/a:sysClr{val=window,lastClr=FFFFFF},"
    "a:dk2/a:srgbClr{val=44546A},"
    "a:lt2/a:srgbClr{val=E7E6E6},"
    "a:accent1/a:srgbClr{val=5B9BD5},"
    "a:accent2/a:srgbClr{val=ED7D31},"
    "a:accent3/a:srgbClr{val=A5A5A5},"
    "a:accent4/a:srgbClr{val=FFC000},"
    "a:accent5/a:srgbClr{val=4472C4},"
    "a:accent6/a:srgbClr{val=70AD47},"
    "a:hlink/a:srgbClr{val=0563C1},"
    "a:folHlink/a:srgbClr{val=954F72}"
    "),"
    "a:fontScheme{name=Office}/("
    "a:majorFont/("
    "a:latin{typeface=Calibri Light},"
    "a:ea{typeface=MS Gothic},"
    "a:cs{typeface=Arial}"
    "),"
    "a:minorFont/("
    "a:latin{typeface=Calibri},"
    "a:ea{typeface=MS Mincho},"
    "a:cs{typeface=Times New Roman}"
    ")"
    ")"
    ")"
)


class DescribeTheme:
    """Unit-test suite for `docx.theme.Theme`."""

    def it_exposes_its_name(self):
        theme = Theme(element("a:theme{name=Office Theme}"))
        assert theme.name == "Office Theme"

    def it_returns_None_for_a_missing_name(self):
        theme = Theme(element("a:theme"))
        assert theme.name is None

    def it_exposes_a_colors_view(self):
        theme = Theme(cast(CT_Theme, element(_SAMPLE_THEME)))
        assert isinstance(theme.colors, ThemeColors)

    def it_exposes_a_fonts_view(self):
        theme = Theme(cast(CT_Theme, element(_SAMPLE_THEME)))
        assert isinstance(theme.fonts, ThemeFonts)

    def it_is_populated_for_a_default_Document(self):
        # The default python-docx template carries a standard Office theme,
        # so ``Document().theme`` should resolve to a non-None Theme that
        # exposes Calibri/Cambria as the major/minor Latin faces.
        document = docx.Document()
        theme = document.theme
        assert isinstance(theme, Theme)
        assert theme.name == "Office Theme"
        assert theme.fonts.major_latin == "Calibri"
        assert theme.fonts.minor_latin == "Cambria"
        # accent1 from the shipped theme template
        assert theme.colors.accent_1 == RGBColor.from_string("4F81BD")
        # dk1 is sysClr; its lastClr fallback should resolve to black
        assert theme.colors.dark_1 == RGBColor.from_string("000000")


class DescribeThemeColors:
    """Unit-test suite for `docx.theme.ThemeColors`."""

    def it_resolves_each_named_slot(self):
        theme = Theme(cast(CT_Theme, element(_SAMPLE_THEME)))
        colors = theme.colors

        assert colors.name == "Office"
        assert colors.dark_1 == RGBColor.from_string("000000")
        assert colors.light_1 == RGBColor.from_string("FFFFFF")
        assert colors.dark_2 == RGBColor.from_string("44546A")
        assert colors.light_2 == RGBColor.from_string("E7E6E6")
        assert colors.accent_1 == RGBColor.from_string("5B9BD5")
        assert colors.accent_2 == RGBColor.from_string("ED7D31")
        assert colors.accent_3 == RGBColor.from_string("A5A5A5")
        assert colors.accent_4 == RGBColor.from_string("FFC000")
        assert colors.accent_5 == RGBColor.from_string("4472C4")
        assert colors.accent_6 == RGBColor.from_string("70AD47")
        assert colors.hyperlink == RGBColor.from_string("0563C1")
        assert colors.followed_hyperlink == RGBColor.from_string("954F72")

    @pytest.mark.parametrize(
        ("name", "expected"),
        [
            ("dk1", "000000"),
            ("lt1", "FFFFFF"),
            ("accent1", "5B9BD5"),
            ("accent6", "70AD47"),
            ("hlink", "0563C1"),
            ("folHlink", "954F72"),
        ],
    )
    def it_looks_up_by_OOXML_token(self, name: str, expected: str):
        theme = Theme(cast(CT_Theme, element(_SAMPLE_THEME)))
        assert theme.colors[name] == RGBColor.from_string(expected)

    def it_raises_KeyError_for_an_unknown_token(self):
        theme = Theme(cast(CT_Theme, element(_SAMPLE_THEME)))
        with pytest.raises(KeyError):
            _ = theme.colors["bogus"]

    def it_returns_None_for_every_slot_when_clrScheme_is_absent(self):
        theme = Theme(cast(CT_Theme, element("a:theme/a:themeElements")))
        colors = theme.colors
        assert colors.name is None
        assert colors.accent_1 is None
        assert colors.dark_1 is None
        assert colors.hyperlink is None
        assert colors["accent1"] is None

    def it_returns_None_for_a_slot_with_unresolvable_sysClr(self):
        # sysClr without lastClr -> unresolved -> None
        theme = Theme(
            cast(
                CT_Theme,
                element(
                    "a:theme/a:themeElements/a:clrScheme/a:dk1/"
                    "a:sysClr{val=windowText}"
                ),
            )
        )
        assert theme.colors.dark_1 is None

    def it_returns_None_for_an_absent_slot(self):
        theme = Theme(
            cast(
                CT_Theme,
                element(
                    "a:theme/a:themeElements/a:clrScheme/a:accent1/"
                    "a:srgbClr{val=5B9BD5}"
                ),
            )
        )
        colors = theme.colors
        assert colors.accent_1 == RGBColor.from_string("5B9BD5")
        assert colors.accent_2 is None
        assert colors.hyperlink is None


class DescribeThemeFonts:
    """Unit-test suite for `docx.theme.ThemeFonts`."""

    def it_exposes_each_named_slot(self):
        theme = Theme(cast(CT_Theme, element(_SAMPLE_THEME)))
        fonts = theme.fonts

        assert fonts.name == "Office"
        assert fonts.major_latin == "Calibri Light"
        assert fonts.minor_latin == "Calibri"
        assert fonts.major_east_asian == "MS Gothic"
        assert fonts.minor_east_asian == "MS Mincho"
        assert fonts.major_cs == "Arial"
        assert fonts.minor_cs == "Times New Roman"

    def it_returns_None_for_every_slot_when_fontScheme_is_absent(self):
        theme = Theme(cast(CT_Theme, element("a:theme/a:themeElements")))
        fonts = theme.fonts
        assert fonts.name is None
        assert fonts.major_latin is None
        assert fonts.minor_latin is None
        assert fonts.major_east_asian is None
        assert fonts.minor_east_asian is None
        assert fonts.major_cs is None
        assert fonts.minor_cs is None

    def it_returns_empty_string_when_typeface_attribute_is_empty(self):
        # Office themes commonly set ea/cs typeface="" for the Latin-centric
        # slot — the empty string is passed through unchanged. The cxml
        # grammar can't express empty-string attr values, so we parse raw XML.
        from docx.oxml.parser import parse_xml

        xml = (
            "<a:theme xmlns:a="
            "'http://schemas.openxmlformats.org/drawingml/2006/main'>"
            "  <a:themeElements>"
            "    <a:fontScheme>"
            "      <a:majorFont>"
            "        <a:latin typeface='Calibri Light'/>"
            "        <a:ea typeface=''/>"
            "        <a:cs typeface=''/>"
            "      </a:majorFont>"
            "    </a:fontScheme>"
            "  </a:themeElements>"
            "</a:theme>"
        ).replace("'", '"')
        theme = Theme(cast(CT_Theme, parse_xml(xml)))
        fonts = theme.fonts
        assert fonts.major_latin == "Calibri Light"
        assert fonts.major_east_asian == ""
        assert fonts.major_cs == ""
