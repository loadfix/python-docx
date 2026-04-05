# pyright: reportPrivateUsage=false

"""Integration tests demonstrating the comprehensive testing strategy (issue #43).

Establishes test patterns for all five testing layers:
- Layer 1: XML Structure Tests
- Layer 2: OOXML Schema Validation
- Layer 3: Round-Trip Tests
- Layer 4: Reference File Comparison
- Layer 5: LibreOffice Headless Validation
"""

from __future__ import annotations

import os

import pytest

from docx import Document
from docx.document import Document as DocumentObject
from tests.helpers.libreoffice import requires_libreoffice, validate_with_libreoffice
from tests.helpers.ref_docs import compare_xml_structure, ref_docx_exists, ref_docx_path
from tests.helpers.roundtrip import assert_round_trip, round_trip_document
from tests.helpers.schema_validate import validate_ooxml, validate_part_xml
from tests.helpers.xml_structure import (
    assert_content_type_exists,
    assert_part_exists,
    assert_relationship_exists,
    extract_xml_part,
    validate_xml_structure,
)

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(attr: str) -> str:
    """Return the Clark-notation form of a `w:`-namespace attribute name."""
    return f"{{{_W}}}{attr}"


# -- fixtures -------------------------------------------------------------------------------


@pytest.fixture
def comments_docx_path(tmp_path: str) -> str:
    """Create a .docx with comments and return its file path."""
    doc = Document()
    p = doc.add_paragraph("Hello World")
    doc.add_comment(
        runs=p.runs[0],
        text="A test comment",
        author="Test Author",
        initials="TA",
    )
    path = os.path.join(str(tmp_path), "comments-test.docx")
    doc.save(path)
    return path


# ======================================================================================
# Layer 1: XML Structure Tests
# ======================================================================================


class DescribeLayer1_XMLStructure:
    """Layer 1: Validate that python-docx produces correct OOXML elements."""

    def it_includes_comments_xml_part_when_comments_are_added(
        self, comments_docx_path: str
    ):
        assert_part_exists(comments_docx_path, "word/comments.xml")

    def it_includes_comments_content_type(self, comments_docx_path: str):
        assert_content_type_exists(
            comments_docx_path,
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.comments+xml",
        )

    def it_includes_comments_relationship(self, comments_docx_path: str):
        assert_relationship_exists(
            comments_docx_path,
            "word/_rels/document.xml.rels",
            "http://schemas.openxmlformats.org/officeDocument"
            "/2006/relationships/comments",
        )

    def it_produces_w_comment_elements_with_correct_attributes(
        self, comments_docx_path: str
    ):
        root = extract_xml_part(comments_docx_path, "word/comments.xml")
        ns = {"w": _W}

        comments = root.findall("w:comment", ns)
        assert len(comments) == 1

        comment = comments[0]
        assert comment.get(_w("id")) is not None
        assert comment.get(_w("author")) == "Test Author"
        assert comment.get(_w("initials")) == "TA"
        assert comment.get(_w("date")) is not None

    def it_produces_matching_comment_range_markers_in_document_xml(
        self, comments_docx_path: str
    ):
        root = extract_xml_part(comments_docx_path, "word/document.xml")
        ns = {"w": _W}

        range_starts = root.findall(".//w:commentRangeStart", ns)
        range_ends = root.findall(".//w:commentRangeEnd", ns)
        comment_refs = root.findall(".//w:commentReference", ns)

        assert len(range_starts) >= 1
        assert len(range_ends) >= 1
        assert len(comment_refs) >= 1

        # -- start and end IDs should match --
        start_ids = {rs.get(_w("id")) for rs in range_starts}
        end_ids = {re.get(_w("id")) for re in range_ends}
        assert start_ids == end_ids

    def it_can_validate_xml_structure_with_declarative_checks(
        self, comments_docx_path: str
    ):
        checks = [
            {
                "part": "word/comments.xml",
                "xpath": "//w:comment",
                "expected_count": 1,
            },
            {
                "part": "word/document.xml",
                "xpath": "//w:commentRangeStart",
                "expected_min": 1,
            },
        ]
        failures = validate_xml_structure(comments_docx_path, checks)
        assert failures == [], f"XML structure validation failed: {failures}"


# ======================================================================================
# Layer 2: OOXML Schema Validation
# ======================================================================================


class DescribeLayer2_SchemaValidation:
    """Layer 2: Validate output XML against OOXML XSD schemas."""

    def it_can_report_when_schemas_are_not_available(
        self, comments_docx_path: str
    ):
        errors = validate_ooxml(comments_docx_path)
        # -- when schemas haven't been downloaded, we get a clear message --
        if errors and "not available" in errors[0]:
            pytest.skip("OOXML schemas not downloaded — run download_schema.py")
        # -- if schemas are available, no errors expected --
        assert errors == []

    def it_can_validate_a_specific_part(self, comments_docx_path: str):
        errors = validate_part_xml(comments_docx_path, "word/comments.xml")
        if errors and "not available" in errors[0]:
            pytest.skip("OOXML schemas not downloaded")
        assert errors == []


# ======================================================================================
# Layer 3: Round-Trip Tests
# ======================================================================================


class DescribeLayer3_RoundTrip:
    """Layer 3: Write -> save -> re-open -> assert data reads back correctly."""

    def it_round_trips_a_comment_with_text_and_author(self):
        def create(doc: DocumentObject) -> None:
            p = doc.add_paragraph("Some text here")
            doc.add_comment(
                runs=p.runs[0],
                text="Round trip comment",
                author="Jane Doe",
                initials="JD",
            )

        def verify(doc: DocumentObject) -> None:
            comments = doc.comments
            assert len(comments) == 1
            comment = next(iter(comments))
            assert comment.text == "Round trip comment"
            assert comment.author == "Jane Doe"
            assert comment.initials == "JD"
            assert comment.timestamp is not None

        assert_round_trip(create, verify)

    def it_round_trips_multiple_comments_by_different_authors(self):
        def create(doc: DocumentObject) -> None:
            p1 = doc.add_paragraph("First paragraph")
            p2 = doc.add_paragraph("Second paragraph")
            doc.add_comment(
                runs=p1.runs[0], text="Comment A", author="Alice", initials="A"
            )
            doc.add_comment(
                runs=p2.runs[0], text="Comment B", author="Bob", initials="B"
            )

        def verify(doc: DocumentObject) -> None:
            comments = list(doc.comments)
            assert len(comments) == 2

            authors = {c.author for c in comments}
            assert authors == {"Alice", "Bob"}

            texts = {c.text for c in comments}
            assert texts == {"Comment A", "Comment B"}

        assert_round_trip(create, verify)

    def it_round_trips_a_comment_with_multi_paragraph_text(self):
        def create(doc: DocumentObject) -> None:
            p = doc.add_paragraph("Annotated text")
            doc.add_comment(
                runs=p.runs[0],
                text="First para\n\nSecond para",
                author="Author",
            )

        def verify(doc: DocumentObject) -> None:
            comment = next(iter(doc.comments))
            paragraphs = comment.paragraphs
            # -- first para has annotation ref run + text run --
            assert len(paragraphs) == 3
            assert paragraphs[0].text == "First para"
            assert paragraphs[1].text == ""
            assert paragraphs[2].text == "Second para"

        assert_round_trip(create, verify)

    def it_preserves_document_content_alongside_comments(self):
        def create(doc: DocumentObject) -> None:
            p = doc.add_paragraph("Document text")
            doc.add_comment(runs=p.runs[0], text="A note", author="Tester")

        def verify(doc: DocumentObject) -> None:
            paras = doc.paragraphs
            text_paras = [p for p in paras if p.text]
            assert any("Document text" in p.text for p in text_paras)

        assert_round_trip(create, verify)

    def it_round_trips_using_the_convenience_function(self):
        doc = Document()
        p = doc.add_paragraph("Test content")
        doc.add_comment(
            runs=p.runs[0], text="Convenience test", author="Dev"
        )

        doc2 = round_trip_document(doc)

        assert len(doc2.comments) == 1
        assert next(iter(doc2.comments)).text == "Convenience test"


# ======================================================================================
# Layer 4: Reference File Comparison
# ======================================================================================


class DescribeLayer4_ReferenceFiles:
    """Layer 4: Read reference .docx files and verify parsing correctness."""

    def it_can_read_comments_from_a_reference_file(self):
        if not ref_docx_exists("comments-simple"):
            pytest.skip("Reference file comments-simple.docx not available")

        doc = Document(ref_docx_path("comments-simple"))
        comments = list(doc.comments)
        assert len(comments) == 2

        # -- verify first comment --
        c0 = comments[0]
        assert c0.author == "Author A"
        assert c0.initials == "AA"
        assert "Simple comment on a word" in c0.text

        # -- verify second comment --
        c1 = comments[1]
        assert c1.author == "Author B"
        assert c1.initials == "AB"
        assert "Comment on a full paragraph" in c1.text

    def it_can_read_document_paragraphs_from_a_reference_file(self):
        if not ref_docx_exists("comments-simple"):
            pytest.skip("Reference file comments-simple.docx not available")

        doc = Document(ref_docx_path("comments-simple"))
        para_texts = [p.text for p in doc.paragraphs if p.text]
        assert "Hello World" in para_texts
        assert "This is a second paragraph with a comment." in para_texts

    def it_can_compare_xml_structures(self):
        from lxml import etree

        actual = etree.fromstring(
            b"<root><child attr='1'>text</child></root>"
        )
        expected = etree.fromstring(
            b"<root><child attr='1'>text</child></root>"
        )
        differences = compare_xml_structure(actual, expected)
        assert differences == []

    def it_detects_xml_structural_differences(self):
        from lxml import etree

        actual = etree.fromstring(
            b"<root><child attr='1'>text</child></root>"
        )
        expected = etree.fromstring(
            b"<root><child attr='2'>text</child></root>"
        )
        differences = compare_xml_structure(actual, expected)
        assert len(differences) > 0

    def it_can_ignore_specific_attributes_during_comparison(self):
        from lxml import etree

        actual = etree.fromstring(
            b"<root><child id='1' name='a'>text</child></root>"
        )
        expected = etree.fromstring(
            b"<root><child id='999' name='a'>text</child></root>"
        )

        diffs = compare_xml_structure(
            actual, expected, ignore_attrs={"id"}
        )
        assert diffs == []


# ======================================================================================
# Layer 5: LibreOffice Headless Validation
# ======================================================================================


class DescribeLayer5_LibreOfficeValidation:
    """Layer 5: Validate .docx files with LibreOffice headless conversion."""

    @requires_libreoffice
    def it_validates_a_simple_document_with_libreoffice(
        self, comments_docx_path: str
    ):
        success, message = validate_with_libreoffice(comments_docx_path)
        assert success, f"LibreOffice validation failed: {message}"

    @requires_libreoffice
    def it_validates_a_document_with_comments_with_libreoffice(
        self, tmp_path: str
    ):
        doc = Document()
        p1 = doc.add_paragraph("First paragraph")
        p2 = doc.add_paragraph("Second paragraph")
        doc.add_comment(
            runs=p1.runs[0], text="Comment one", author="Alice"
        )
        doc.add_comment(
            runs=p2.runs[0], text="Comment two", author="Bob"
        )

        path = os.path.join(str(tmp_path), "multi-comment.docx")
        doc.save(path)

        success, message = validate_with_libreoffice(path)
        assert success, f"LibreOffice validation failed: {message}"
