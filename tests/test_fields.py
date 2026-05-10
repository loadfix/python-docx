"""Unit-test suite for the docx.fields module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.fields import Field, WD_FIELD_TYPE
from docx.oxml.fields import CT_FldSimple
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R

from .unitutil.cxml import element


class DescribeWD_FIELD_TYPE:
    """Sanity check for the constant set."""

    @pytest.mark.parametrize(
        ("name", "value"),
        [
            ("PAGE", "PAGE"),
            ("NUMPAGES", "NUMPAGES"),
            ("DATE", "DATE"),
            ("TIME", "TIME"),
            ("AUTHOR", "AUTHOR"),
            ("REF", "REF"),
            ("TOC", "TOC"),
            ("SEQ", "SEQ"),
            ("HYPERLINK", "HYPERLINK"),
            ("PAGEREF", "PAGEREF"),
            ("MERGEFIELD", "MERGEFIELD"),
            ("FILENAME", "FILENAME"),
            ("TITLE", "TITLE"),
            ("STYLEREF", "STYLEREF"),
            ("NUMBEREDHEADERS", "NUMBEREDHEADERS"),
        ],
    )
    def it_exposes_common_field_types_as_string_constants(self, name: str, value: str):
        assert getattr(WD_FIELD_TYPE, name) == value


class DescribeField_Simple:
    """Unit-test suite for `Field` wrapping a ``w:fldSimple`` element."""

    def it_is_not_complex(self):
        fldSimple = cast(CT_FldSimple, element('w:fldSimple{w:instr=PAGE}'))
        field = Field.for_simple(fldSimple)
        assert field.is_complex is False

    def it_exposes_the_raw_instruction(self):
        # -- backslashes aren't supported by the cxml attribute grammar; build
        #    the element via OxmlElement and set the w:instr attribute directly.
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "REF bookmark1 \\h")
        field = Field.for_simple(fldSimple)
        assert field.instruction == "REF bookmark1 \\h"

    def it_returns_empty_instruction_when_attr_missing(self):
        # -- w:instr is required, but defensively handle absence --
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        field = Field.for_simple(fldSimple)
        assert field.instruction == ""

    @pytest.mark.parametrize(
        ("instr", "expected_type"),
        [
            ("PAGE", "PAGE"),
            ("NUMPAGES", "NUMPAGES"),
            ("REF bookmark1 \\h", "REF"),
            ("TOC \\o \"1-3\"", "TOC"),
            ("SEQ Table", "SEQ"),
            ("DATE", "DATE"),
            ("TIME \\@ \"h:mm AM/PM\"", "TIME"),
            ("AUTHOR", "AUTHOR"),
            ("HYPERLINK \"https://example.com\"", "HYPERLINK"),
            ("PAGEREF _Toc12345", "PAGEREF"),
        ],
    )
    def it_parses_the_type_from_the_instruction(self, instr: str, expected_type: str):
        from docx.oxml.parser import OxmlElement
        from docx.oxml.ns import qn

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), instr)
        field = Field.for_simple(fldSimple)
        assert field.type == expected_type

    def it_returns_empty_type_for_empty_instruction(self):
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        field = Field.for_simple(fldSimple)
        assert field.type == ""

    def it_exposes_the_result_text(self):
        fldSimple = cast(
            CT_FldSimple,
            element('w:fldSimple{w:instr=PAGE}/w:r/w:t"3"'),
        )
        field = Field.for_simple(fldSimple)
        assert field.result_text == "3"

    def it_returns_empty_result_text_when_no_run(self):
        fldSimple = cast(CT_FldSimple, element('w:fldSimple{w:instr=PAGE}'))
        field = Field.for_simple(fldSimple)
        assert field.result_text == ""


class DescribeField_Complex:
    """Unit-test suite for `Field` wrapping a complex field begin-run."""

    def _build_complex_paragraph(
        self, instr: str, result_text: str | None = None
    ) -> tuple[CT_P, CT_R]:
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field(instr, result_text)
        return p, begin_run

    def it_is_complex(self):
        _, begin_run = self._build_complex_paragraph("PAGE", "3")
        field = Field.for_complex(begin_run)
        assert field.is_complex is True

    def it_reads_the_instruction(self):
        _, begin_run = self._build_complex_paragraph("REF bookmark1 \\h", "See here")
        field = Field.for_complex(begin_run)
        assert field.instruction == "REF bookmark1 \\h"

    def it_reads_the_type(self):
        _, begin_run = self._build_complex_paragraph("NUMPAGES", "10")
        field = Field.for_complex(begin_run)
        assert field.type == "NUMPAGES"

    def it_reads_the_result_text(self):
        _, begin_run = self._build_complex_paragraph("PAGE", "42")
        field = Field.for_complex(begin_run)
        assert field.result_text == "42"

    def it_returns_empty_result_text_when_no_separate_marker(self):
        # -- build a field with only begin/instrText/end; no separate marker --
        p = cast(CT_P, element("w:p"))
        p.add_complex_field("PAGE")  # no result_text => 4 runs
        # -- remove the separate marker to leave only begin/instrText/end --
        from docx.oxml.ns import qn

        seps = p.xpath('.//w:fldChar[@w:fldCharType="separate"]')
        sep_run = seps[0].getparent()
        sep_run.getparent().remove(sep_run)
        begin_run = p.r_lst[0]
        field = Field.for_complex(begin_run)
        assert field.result_text == ""

    def it_returns_empty_result_text_when_omitted(self):
        _, begin_run = self._build_complex_paragraph("PAGE")
        field = Field.for_complex(begin_run)
        assert field.result_text == ""

    def it_concatenates_instruction_split_across_runs(self):
        # -- some producers split the instruction across multiple instrText runs --
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        p = cast(CT_P, element("w:p"))

        r_begin = p.add_r()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r_begin.append(fld_begin)

        r_i1 = p.add_r()
        i1 = OxmlElement("w:instrText")
        i1.text = "REF "
        r_i1.append(i1)

        r_i2 = p.add_r()
        i2 = OxmlElement("w:instrText")
        i2.text = "bookmark1"
        r_i2.append(i2)

        r_sep = p.add_r()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r_sep.append(fld_sep)

        r_end = p.add_r()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end"),
        r_end.append(fld_end)

        field = Field.for_complex(r_begin)
        assert field.instruction == "REF bookmark1"
        assert field.type == "REF"


class DescribeField_resolve:
    """Unit-test suite for `Field.resolve()` cross-reference resolution."""

    def _doc_with_bookmark(
        self, bookmark_name: str, bookmark_text: str
    ):
        """Return a `Document` whose body contains a bookmark wrapping `bookmark_text`.

        The bookmark uses id=0 and sits in a single paragraph alongside a
        ``w:r/w:t`` run containing `bookmark_text`.
        """
        from docx.document import Document
        from docx.oxml.document import CT_Document

        # -- build: <w:body><w:p><w:bookmarkStart .../><w:r><w:t>...</w:t></w:r>
        #    <w:bookmarkEnd .../></w:p></w:body> --
        doc_elm = cast(
            CT_Document,
            element(
                f"w:document/w:body/w:p/("
                f"w:bookmarkStart{{w:id=0,w:name={bookmark_name}}}"
                f",w:r/w:t\"{bookmark_text}\""
                f",w:bookmarkEnd{{w:id=0}}"
                f")"
            ),
        )
        return Document(doc_elm, None)  # type: ignore[arg-type]

    def it_resolves_REF_to_bookmark_text_for_simple_field(self):
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        doc = self._doc_with_bookmark("Ref1", "Hello")
        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "REF Ref1 \\h")
        field = Field.for_simple(fldSimple)

        assert field.resolve(doc) == "Hello"

    def it_resolves_REF_to_bookmark_text_for_complex_field(self):
        doc = self._doc_with_bookmark("Ref1", "Hello")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("REF Ref1 \\h", "stale")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "Hello"

    def it_resolves_heading_style_REF_referencing_a_heading_bookmark(self):
        # -- Word uses names like "_Ref12345" or "_Toc12345" for auto-generated
        #    cross-references; treat them like any other bookmark. --
        doc = self._doc_with_bookmark("_Ref12345", "Chapter 2")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("REF _Ref12345 \\h", "")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "Chapter 2"

    def it_returns_cached_result_for_PAGEREF_when_present(self):
        doc = self._doc_with_bookmark("Ref1", "ignored")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("PAGEREF Ref1 \\h", "42")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "42"

    def it_returns_question_mark_for_PAGEREF_when_no_cached_result(self):
        doc = self._doc_with_bookmark("Ref1", "ignored")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("PAGEREF Ref1 \\h")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "?"

    def it_returns_cached_result_for_unrelated_field_types(self):
        doc = self._doc_with_bookmark("Ref1", "ignored")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("PAGE", "7")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "7"

    def it_returns_cached_result_when_REF_bookmark_not_found(self):
        doc = self._doc_with_bookmark("Ref1", "Hello")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("REF Missing \\h", "stale")
        field = Field.for_complex(begin_run)

        # -- unresolvable references leave the cached result untouched --
        assert field.resolve(doc) == "stale"

    def it_returns_cached_result_when_REF_has_no_bookmark_name(self):
        doc = self._doc_with_bookmark("Ref1", "Hello")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("REF", "stale")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "stale"

    def it_strips_switches_before_looking_up_the_bookmark_name(self):
        # -- verify backslash switches are skipped --
        doc = self._doc_with_bookmark("Ref1", "Hello")
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field('REF \\h Ref1 \\* MERGEFORMAT', "")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "Hello"

    def it_updates_result_text_in_place_for_simple_field(self):
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "REF Ref1")
        r = fldSimple.add_r()
        r.add_t("old")
        field = Field.for_simple(fldSimple)

        field.update_result_text("new value")

        assert field.result_text == "new value"

    def it_updates_result_text_in_place_for_complex_field(self):
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("REF Ref1", "old")
        field = Field.for_complex(begin_run)

        field.update_result_text("new value")

        assert field.result_text == "new value"

    def it_preserves_whitespace_when_updated_text_has_leading_or_trailing_spaces(self):
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("REF Ref1", "old")
        field = Field.for_complex(begin_run)

        field.update_result_text(" padded ")

        assert field.result_text == " padded "

    def it_walks_bookmark_text_across_multiple_runs(self):
        # -- bookmark range spans two runs; concatenate their text --
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/("
                "w:bookmarkStart{w:id=0,w:name=Ref1}"
                ',w:r/w:t"Hello "'
                ',w:r/w:t"World"'
                ",w:bookmarkEnd{w:id=0}"
                ")"
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field("REF Ref1", "")
        field = Field.for_complex(begin_run)

        assert field.resolve(doc) == "Hello World"


class DescribeDocument_resolve_cross_references:
    """Unit-test suite for `Document.resolve_cross_references()`."""

    def it_resolves_REF_and_updates_result_text_in_place(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        # -- body contains: a paragraph with bookmark "Ref1" wrapping "Hello",
        #    then a second paragraph with a REF simple field. --
        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                "w:p/("
                "w:bookmarkStart{w:id=0,w:name=Ref1}"
                ',w:r/w:t"Hello"'
                ",w:bookmarkEnd{w:id=0}"
                "),"
                "w:p/w:fldSimple{w:instr=REF Ref1}/w:r/w:t\"stale\""
                ")"
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        count = doc.resolve_cross_references()

        assert count == 1
        # -- find the fldSimple and confirm its text is now "Hello" --
        fldSimples = doc._element.body.xpath(".//w:fldSimple")  # pyright: ignore[reportPrivateUsage]
        assert len(fldSimples) == 1
        assert Field.for_simple(fldSimples[0]).result_text == "Hello"

    def it_updates_multiple_fields_and_returns_the_count(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                "w:p/("
                "w:bookmarkStart{w:id=0,w:name=A}"
                ',w:r/w:t"AAA"'
                ",w:bookmarkEnd{w:id=0}"
                "),"
                "w:p/("
                "w:bookmarkStart{w:id=1,w:name=B}"
                ',w:r/w:t"BBB"'
                ",w:bookmarkEnd{w:id=1}"
                "),"
                'w:p/w:fldSimple{w:instr=REF A}/w:r/w:t"stale"'
                ','
                'w:p/w:fldSimple{w:instr=REF B}/w:r/w:t"stale"'
                ")"
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        count = doc.resolve_cross_references()

        assert count == 2
        results = [
            Field.for_simple(fs).result_text
            for fs in doc._element.body.xpath(".//w:fldSimple")  # pyright: ignore[reportPrivateUsage]
        ]
        assert results == ["AAA", "BBB"]

    def it_leaves_unresolvable_fields_untouched(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/w:fldSimple{w:instr=REF NoSuchBookmark}/w:r/w:t"stale"'
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        count = doc.resolve_cross_references()

        assert count == 0
        fs = doc._element.body.xpath(".//w:fldSimple")[0]  # pyright: ignore[reportPrivateUsage]
        assert Field.for_simple(fs).result_text == "stale"

    def it_ignores_non_crossreference_fields(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/w:fldSimple{w:instr=PAGE}/w:r/w:t"3"'
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        count = doc.resolve_cross_references()

        assert count == 0
        fs = doc._element.body.xpath(".//w:fldSimple")[0]  # pyright: ignore[reportPrivateUsage]
        assert Field.for_simple(fs).result_text == "3"

    def it_resolves_PAGEREF_with_empty_cached_result_to_question_mark(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        # -- PAGEREF with no cached result_text; expect "?" after resolve --
        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                "w:p/("
                "w:bookmarkStart{w:id=0,w:name=Ref1}"
                ',w:r/w:t"X"'
                ",w:bookmarkEnd{w:id=0}"
                "),"
                "w:p/w:fldSimple{w:instr=PAGEREF Ref1}"
                ")"
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        count = doc.resolve_cross_references()

        assert count == 1
        fs = [
            fs
            for fs in doc._element.body.xpath(".//w:fldSimple")  # pyright: ignore[reportPrivateUsage]
            if fs.get(qn("w:instr")) == "PAGEREF Ref1"
        ][0]
        assert Field.for_simple(fs).result_text == "?"


class DescribeField_DocPropertyResolution:
    """Closes upstream#1482 — resolve DOCPROPERTY / AUTHOR / TITLE / SUBJECT."""

    def _make_doc(self, body_cxml: str):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(CT_Document, element(body_cxml))
        return Document(doc_elm, None)  # type: ignore[arg-type]

    def it_resolves_AUTHOR_from_core_properties(self, request: pytest.FixtureRequest):
        from unittest.mock import PropertyMock

        from docx.opc.coreprops import CoreProperties

        doc = self._make_doc(
            'w:document/w:body/w:p/w:fldSimple{w:instr=AUTHOR}/w:r/w:t"OLD"'
        )
        core = CoreProperties(None)  # type: ignore[arg-type]
        pm = PropertyMock(return_value="Ada Lovelace")
        type(core).author = pm  # pyright: ignore[reportAttributeAccessIssue]
        request.addfinalizer(
            lambda: delattr(type(core), "author")
            if "author" in type(core).__dict__
            else None
        )

        fs = doc._element.body.xpath(".//w:fldSimple")[0]  # pyright: ignore[reportPrivateUsage]

        class _FakePart:
            core_properties = core

        doc._part = _FakePart()  # type: ignore[assignment]
        assert Field.for_simple(fs).resolve(doc) == "Ada Lovelace"

    def it_resolves_DOCPROPERTY_Author(self):
        from unittest.mock import PropertyMock
        from docx.opc.coreprops import CoreProperties

        doc = self._make_doc(
            "w:document/w:body/w:p/w:fldSimple"
            '{w:instr=DOCPROPERTY Author}/w:r/w:t"OLD"'
        )
        core = CoreProperties(None)  # type: ignore[arg-type]
        type(core).author = PropertyMock(return_value="Grace Hopper")

        class _FakePart:
            core_properties = core

        doc._part = _FakePart()  # type: ignore[assignment]
        fs = doc._element.body.xpath(".//w:fldSimple")[0]  # pyright: ignore[reportPrivateUsage]
        assert Field.for_simple(fs).resolve(doc) == "Grace Hopper"

    def it_resolves_DOCPROPERTY_custom_name(self):
        doc = self._make_doc(
            "w:document/w:body/w:p/w:fldSimple"
            '{w:instr=DOCPROPERTY Project}/w:r/w:t"OLD"'
        )

        class _FakeCustom:
            def get(self, name, default=None):
                return {"Project": "Apollo 11"}.get(name, default)

        class _FakeCore:
            pass

        class _FakePart:
            core_properties = _FakeCore()
            custom_properties = _FakeCustom()

        doc._part = _FakePart()  # type: ignore[assignment]
        fs = doc._element.body.xpath(".//w:fldSimple")[0]  # pyright: ignore[reportPrivateUsage]
        assert Field.for_simple(fs).resolve(doc) == "Apollo 11"

    def it_parses_quoted_docproperty_name(self):
        from docx.fields import _parse_docproperty_name

        assert _parse_docproperty_name('DOCPROPERTY "Some Multi Word"') == (
            "Some Multi Word"
        )
        assert _parse_docproperty_name("DOCPROPERTY Title") == "Title"
        assert _parse_docproperty_name("DOCPROPERTY") is None
        assert _parse_docproperty_name(
            'DOCPROPERTY "Simple" \\* MERGEFORMAT'
        ) == "Simple"


class DescribeField_MarkDirty:
    """Unit-test suite for `Field.mark_dirty` / `Field.is_dirty`."""

    def it_marks_a_simple_field_dirty(self):
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "PAGE")
        field = Field.for_simple(fldSimple)

        field.mark_dirty()

        assert fldSimple.get(qn("w:dirty")) == "true"
        assert field.is_dirty is True

    def it_reads_an_unset_simple_dirty_flag_as_false(self):
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "PAGE")
        field = Field.for_simple(fldSimple)
        assert field.is_dirty is False

    def it_marks_a_complex_field_dirty_on_the_begin_marker(self):
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field(" TOC ", "cached")
        field = Field.for_complex(begin_run)

        field.mark_dirty()

        begin_fldChar = begin_run.find(qn("w:fldChar"))
        assert begin_fldChar is not None
        assert begin_fldChar.get(qn("w:dirty")) == "true"
        assert field.is_dirty is True

    def it_reads_an_unset_complex_dirty_flag_as_false(self):
        p = cast(CT_P, element("w:p"))
        begin_run = p.add_complex_field(" TOC ", "cached")
        assert Field.for_complex(begin_run).is_dirty is False

    @pytest.mark.parametrize(
        ("value", "expected"),
        [("true", True), ("1", True), ("on", True), ("false", False), ("0", False)],
    )
    def it_parses_known_dirty_values_consistently(
        self, value: str, expected: bool
    ):
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "PAGE")
        fldSimple.set(qn("w:dirty"), value)
        assert Field.for_simple(fldSimple).is_dirty is expected


class DescribeField_evaluate:
    """Unit-test suite for `Field.evaluate()` complex-field evaluation."""

    def _simple(self, instr: str, cached: str = "") -> Field:
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), instr)
        if cached:
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = cached
            r.append(t)
            fldSimple.append(r)
        return Field.for_simple(fldSimple)

    # -- MERGEFIELD ----------------------------------------------------------

    def it_evaluates_MERGEFIELD_from_context(self):
        field = self._simple("MERGEFIELD firstname", cached="<<firstname>>")
        assert field.evaluate({"firstname": "Ada"}) == "Ada"

    def it_falls_back_to_cached_result_when_MERGEFIELD_key_missing(self):
        field = self._simple("MERGEFIELD firstname", cached="<<firstname>>")
        assert field.evaluate({}) == "<<firstname>>"

    def it_evaluates_MERGEFIELD_with_quoted_multiword_name(self):
        field = self._simple('MERGEFIELD "Full Name"')
        assert field.evaluate({"Full Name": "Ada Lovelace"}) == "Ada Lovelace"

    # -- IF ------------------------------------------------------------------

    def it_evaluates_IF_with_equal_match(self):
        field = self._simple('IF "yes" = "yes" "match" "nope"')
        assert field.evaluate({}) == "match"

    def it_evaluates_IF_with_equal_mismatch(self):
        field = self._simple('IF "yes" = "no" "match" "nope"')
        assert field.evaluate({}) == "nope"

    def it_evaluates_IF_with_nested_MERGEFIELD(self):
        field = self._simple('IF {MERGEFIELD status} = "active" "OK" "FAIL"')
        assert field.evaluate({"status": "active"}) == "OK"
        assert field.evaluate({"status": "cancelled"}) == "FAIL"

    @pytest.mark.parametrize(
        ("op", "lhs", "rhs", "expected"),
        [
            ("<>", "a", "b", "yes"),
            ("<>", "a", "a", "no"),
            ("!=", "a", "a", "no"),
            ("<", "1", "2", "yes"),
            (">", "2", "1", "yes"),
            ("<=", "2", "2", "yes"),
            (">=", "2", "3", "no"),
        ],
    )
    def it_supports_the_common_comparison_operators(
        self, op: str, lhs: str, rhs: str, expected: str
    ):
        field = self._simple(f'IF "{lhs}" {op} "{rhs}" "yes" "no"')
        assert field.evaluate({}) == expected

    def it_uses_numeric_comparison_when_both_sides_parse_as_numbers(self):
        field = self._simple('IF "10" > "9" "big" "small"')
        assert field.evaluate({}) == "big"

    def it_returns_empty_false_branch_when_only_true_text_given(self):
        field = self._simple('IF "a" = "b" "match"')
        assert field.evaluate({}) == ""

    def it_returns_cached_result_when_IF_is_malformed(self):
        field = self._simple("IF", cached="x")
        assert field.evaluate({}) == "x"

    # -- HYPERLINK -----------------------------------------------------------

    def it_returns_cached_display_text_for_HYPERLINK_when_present(self):
        field = self._simple('HYPERLINK "https://example.com"', cached="click me")
        assert field.evaluate({}) == "click me"

    def it_returns_the_url_for_HYPERLINK_when_no_cached_text(self):
        field = self._simple('HYPERLINK "https://example.com"')
        assert field.evaluate({}) == "https://example.com"

    # -- runtime-dynamic -----------------------------------------------------

    @pytest.mark.parametrize(
        "instr", ["PAGE", "NUMPAGES", "DATE", "TIME"]
    )
    def it_returns_cached_result_for_runtime_dynamic_fields(self, instr: str):
        field = self._simple(instr, cached="7")
        assert field.evaluate({}) == "7"

    @pytest.mark.parametrize(
        "instr", ["PAGE", "NUMPAGES", "DATE", "TIME"]
    )
    def it_returns_question_mark_for_runtime_dynamic_fields_with_no_cache(
        self, instr: str
    ):
        field = self._simple(instr)
        assert field.evaluate({}) == "?"

    # -- formula (=) ---------------------------------------------------------

    @pytest.mark.parametrize(
        ("expr", "expected"),
        [
            ("= 1+2", "3"),
            ("= 2*3", "6"),
            ("= (2+3)*4", "20"),
            ("= 7/2", "3.5"),
            ("= 10 % 3", "1"),
            ("= 2**8", "256"),
        ],
    )
    def it_evaluates_arithmetic_formula_fields(self, expr: str, expected: str):
        field = self._simple(expr)
        assert field.evaluate({}) == expected

    def it_substitutes_MERGEFIELD_in_formula(self):
        field = self._simple("= {MERGEFIELD qty} * 10")
        assert field.evaluate({"qty": 4}) == "40"

    def it_returns_cached_for_formula_with_disallowed_chars(self):
        field = self._simple('= __import__("os").system("ls")', cached="stale")
        assert field.evaluate({}) == "stale"

    # -- pass-through --------------------------------------------------------

    def it_returns_cached_result_for_unknown_field_types(self):
        field = self._simple("TOC \\o \"1-3\"", cached="toc-preview")
        assert field.evaluate({}) == "toc-preview"


class DescribeDocument_evaluate_fields:
    """Unit-test suite for `Document.evaluate_fields()`."""

    def it_updates_MERGEFIELD_cached_text_in_place(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/w:fldSimple{w:instr=MERGEFIELD name}/w:r/w:t"stale"'
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        count = doc.evaluate_fields({"name": "Ada"})
        assert count == 1
        fs = doc._element.body.xpath(".//w:fldSimple")[0]  # pyright: ignore[reportPrivateUsage]
        assert Field.for_simple(fs).result_text == "Ada"

    def it_returns_zero_when_nothing_changed(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/w:fldSimple{w:instr=MERGEFIELD name}/w:r/w:t"Ada"'
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        count = doc.evaluate_fields({"name": "Ada"})
        assert count == 0

    def it_evaluates_IF_formula_and_MERGEFIELD_together(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document
        from docx.oxml.parser import OxmlElement

        doc_elm = cast(
            CT_Document,
            element("w:document/w:body/(w:p,w:p,w:p)"),
        )
        ps = doc_elm.body.xpath(".//w:p")
        # -- merge field --
        fs1 = OxmlElement("w:fldSimple")
        fs1.set(qn("w:instr"), "MERGEFIELD name")
        ps[0].append(fs1)
        # -- IF --
        fs2 = OxmlElement("w:fldSimple")
        fs2.set(qn("w:instr"), 'IF {MERGEFIELD status} = "active" "yes" "no"')
        ps[1].append(fs2)
        # -- formula --
        fs3 = OxmlElement("w:fldSimple")
        fs3.set(qn("w:instr"), "= 2+3")
        ps[2].append(fs3)

        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        count = doc.evaluate_fields({"name": "Ada", "status": "active"})
        assert count == 3
        fs_list = doc._element.body.xpath(".//w:fldSimple")  # pyright: ignore[reportPrivateUsage]
        assert [Field.for_simple(f).result_text for f in fs_list] == [
            "Ada", "yes", "5"
        ]

    def it_passes_document_context_for_property_fields(
        self, request: pytest.FixtureRequest
    ):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/w:fldSimple{w:instr=AUTHOR}/w:r/w:t"stale"'
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        # -- stub core_properties.author --
        class _Props:
            author = "Jane Doe"
            title = None
            subject = None
            keywords = None
            comments = None
            last_modified_by = None

        doc.__class__.core_properties = property(  # type: ignore[assignment]
            lambda self: _Props()
        )
        try:
            count = doc.evaluate_fields({})
            assert count == 1
            fs = doc._element.body.xpath(".//w:fldSimple")[0]  # pyright: ignore[reportPrivateUsage]
            assert Field.for_simple(fs).result_text == "Jane Doe"
        finally:
            del doc.__class__.core_properties  # type: ignore[attr-defined]

    def it_accepts_a_missing_context_argument(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/w:fldSimple{w:instr=PAGE}/w:r/w:t"7"'
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        count = doc.evaluate_fields()
        # -- PAGE with cached "7" stays "7" --
        assert count == 0


class DescribeParseFieldInstruction:
    """Unit-test suite for `docx.fields.parse_field_instruction`."""

    def it_parses_a_bare_field_name(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction("PAGE")
        assert parsed.name == "PAGE"
        assert parsed.args == []
        assert parsed.switches == {}

    def it_parses_mergefield_with_format_switch(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction("MERGEFIELD FirstName \\* MERGEFORMAT")
        assert parsed.name == "MERGEFIELD"
        assert parsed.args == ["FirstName"]
        assert parsed.switches == {"*": "MERGEFORMAT"}

    def it_unquotes_string_literal_arguments(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction('HYPERLINK "https://example.com"')
        assert parsed.name == "HYPERLINK"
        assert parsed.args == ["https://example.com"]

    def it_records_flag_switches_with_empty_string(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction("REF heading1 \\h \\p")
        assert parsed.name == "REF"
        assert parsed.args == ["heading1"]
        assert parsed.switches == {"H": "", "P": ""}

    def it_keeps_argument_taking_switch_arguments(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction('TIME \\@ "h:mm AM/PM"')
        assert parsed.name == "TIME"
        assert parsed.switches == {"@": "h:mm AM/PM"}

    def it_handles_multiple_positional_arguments(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction("STYLEREF \"Heading 1\" \\n")
        assert parsed.name == "STYLEREF"
        assert parsed.args == ["Heading 1"]
        assert parsed.switches == {"N": ""}

    def it_returns_empty_for_empty_instruction(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction("")
        assert parsed.name == ""
        assert parsed.args == []
        assert parsed.switches == {}

    def it_tolerates_leading_and_trailing_whitespace(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction("  PAGE  ")
        assert parsed.name == "PAGE"

    def it_handles_toc_with_quoted_level_range(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction('TOC \\o "1-3" \\h \\z \\u')
        assert parsed.name == "TOC"
        assert parsed.switches == {"O": "", "H": "", "Z": "", "U": ""}
        # -- `\o "1-3"` is a flag switch + separate positional quoted string
        #    per our parser's rule; callers can resolve semantically if
        #    needed. "1-3" ends up in args.
        assert "1-3" in parsed.args

    def it_preserves_nested_field_group_tokens(self):
        from docx.fields import parse_field_instruction

        parsed = parse_field_instruction('IF {MERGEFIELD x} = "a" "yes" "no"')
        assert parsed.name == "IF"
        # -- the nested {MERGEFIELD x} group is kept as one atomic arg --
        assert "{MERGEFIELD x}" in parsed.args
        assert "yes" in parsed.args
        assert "no" in parsed.args


class DescribeField_R3_9_Aliases:
    """`Field.result` and `Field.field_type` R3-9 aliases."""

    def it_exposes_result_as_alias_for_result_text(self):
        fldSimple = cast(
            CT_FldSimple,
            element('w:fldSimple{w:instr=PAGE}/w:r/w:t"7"'),
        )
        field = Field.for_simple(fldSimple)
        assert field.result == "7"
        assert field.result == field.result_text

    def it_exposes_field_type_as_alias_for_type(self):
        fldSimple = cast(
            CT_FldSimple,
            element('w:fldSimple{w:instr=MERGEFIELD}'),
        )
        field = Field.for_simple(fldSimple)
        assert field.field_type == "MERGEFIELD"
        assert field.field_type == field.type


class DescribeParagraph_add_field:
    """`paragraph.add_field(instruction, cached_result=None)` complex-field emitter."""

    def it_emits_a_complex_field_sequence(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        field = para.add_field("MERGEFIELD FirstName")
        assert field.is_complex is True
        # -- begin/separate/end markers present --
        begin = p.xpath(".//w:fldChar[@w:fldCharType='begin']")
        separate = p.xpath(".//w:fldChar[@w:fldCharType='separate']")
        end = p.xpath(".//w:fldChar[@w:fldCharType='end']")
        assert len(begin) == 1
        assert len(separate) == 1
        assert len(end) == 1
        # -- instrText between begin and separate carries the instruction --
        instrText = p.xpath(".//w:instrText")
        assert len(instrText) == 1
        assert instrText[0].text == "MERGEFIELD FirstName"

    def it_writes_the_cached_result_when_provided(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        field = para.add_field("AUTHOR", cached_result="Jane Doe")
        assert field.result_text == "Jane Doe"

    def it_omits_result_run_when_cached_result_is_None(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_field("PAGE")
        # -- exactly three runs: begin, instrText, separate, end (no result) --
        runs = p.xpath("./w:r")
        assert len(runs) == 4

    def it_returns_a_Field_with_correct_type_and_instruction(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        field = para.add_field("MERGEFIELD FirstName \\* MERGEFORMAT")
        assert field.field_type == "MERGEFIELD"
        assert field.instruction == "MERGEFIELD FirstName \\* MERGEFORMAT"


class DescribeDocument_fields:
    """`document.fields` collection accessor (R3-9)."""

    def it_walks_body_paragraphs(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element("w:document/w:body/(w:p,w:p)"),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        for paragraph in doc.paragraphs:
            paragraph.add_field("PAGE", cached_result="1")

        fields = doc.fields
        assert len(fields) == 2
        assert all(f.field_type == "PAGE" for f in fields)

    def it_includes_simple_and_complex_fields(self):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element("w:document/w:body/(w:p,w:p)"),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        paragraphs = list(doc.paragraphs)
        paragraphs[0].add_simple_field("DATE", "2026-05-09")
        paragraphs[1].add_field("TIME", "10:00")

        fields = doc.fields
        assert len(fields) == 2
        assert fields[0].is_complex is False
        assert fields[1].is_complex is True


class DescribeRun_parent_field:
    """`run.parent_field` lookup (R3-9)."""

    def it_returns_None_for_a_plain_run(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        r = para.add_run("plain text")
        assert r.parent_field is None

    def it_finds_the_enclosing_complex_field_for_interior_runs(self):
        from docx.text.paragraph import Paragraph
        from docx.text.run import Run

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_field("MERGEFIELD FirstName", cached_result="Jane")

        # -- find the result-text run (the one between separate and end) --
        r_elements = p.xpath("./w:r")
        # -- runs: begin, instrText, separate, result, end --
        result_run_el = r_elements[3]
        run = Run(cast(CT_R, result_run_el), para)
        field = run.parent_field
        assert field is not None
        assert field.field_type == "MERGEFIELD"
        assert field.instruction == "MERGEFIELD FirstName"

    def it_returns_None_for_a_run_after_the_end_marker(self):
        from docx.text.paragraph import Paragraph
        from docx.text.run import Run

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_field("PAGE", cached_result="1")
        # -- a run appended *after* the end marker is outside the field --
        trailing = para.add_run("after")
        assert trailing.parent_field is None

    def it_returns_None_for_a_run_before_the_begin_marker(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        leading = para.add_run("before")
        para.add_field("PAGE")
        assert leading.parent_field is None


class DescribeField_RoundTrip:
    """End-to-end create -> serialise -> parse round-trip for complex fields."""

    def it_round_trips_a_mergefield(self):
        from docx.text.paragraph import Paragraph
        from lxml import etree

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_field("MERGEFIELD LastName \\* MERGEFORMAT", cached_result="Smith")

        # -- serialise and re-parse --
        xml = etree.tostring(p, pretty_print=False).decode()
        reparsed = etree.fromstring(xml.encode())

        instrText = reparsed.xpath(
            ".//*[local-name()='instrText']"
        )
        assert len(instrText) == 1
        assert instrText[0].text == "MERGEFIELD LastName \\* MERGEFORMAT"
        # -- three fldChar markers preserved --
        markers = reparsed.xpath(
            ".//*[local-name()='fldChar']"
        )
        kinds = [
            m.get(qn("w:fldCharType"))
            for m in markers
        ]
        assert kinds == ["begin", "separate", "end"]


class DescribeBuildCrossReferenceInstruction:
    """`build_cross_reference_instruction()` field-code builder."""

    def it_builds_a_minimal_REF_instruction(self):
        from docx.fields import build_cross_reference_instruction

        assert build_cross_reference_instruction("REF", "heading1") == "REF heading1"

    def it_appends_the_hyperlink_switch(self):
        from docx.fields import build_cross_reference_instruction

        assert (
            build_cross_reference_instruction(
                "REF", "heading1", insert_as_hyperlink=True
            )
            == "REF heading1 \\h"
        )

    def it_appends_paragraph_number_and_relative_position_switches(self):
        from docx.fields import build_cross_reference_instruction

        instr = build_cross_reference_instruction(
            "REF",
            "heading1",
            insert_as_hyperlink=True,
            insert_paragraph_number=True,
            insert_relative_position=True,
        )
        assert instr == "REF heading1 \\h \\r \\p"

    def it_quotes_target_names_containing_spaces(self):
        from docx.fields import build_cross_reference_instruction

        assert (
            build_cross_reference_instruction("PAGEREF", "my bookmark")
            == 'PAGEREF "my bookmark"'
        )

    def it_uppercases_the_ref_type(self):
        from docx.fields import build_cross_reference_instruction

        assert (
            build_cross_reference_instruction("pageref", "Ref1")
            == "PAGEREF Ref1"
        )

    @pytest.mark.parametrize(
        "ref_type",
        ["REF", "PAGEREF", "NOTEREF", "SEQREF", "STYLEREF"],
    )
    def it_accepts_every_cross_reference_type(self, ref_type: str):
        from docx.fields import build_cross_reference_instruction

        instr = build_cross_reference_instruction(ref_type, "Tgt")
        assert instr.startswith(f"{ref_type} ")

    def it_raises_on_empty_ref_type(self):
        from docx.fields import build_cross_reference_instruction

        with pytest.raises(ValueError):
            build_cross_reference_instruction("", "Ref1")

    def it_raises_on_empty_target_name(self):
        from docx.fields import build_cross_reference_instruction

        with pytest.raises(ValueError):
            build_cross_reference_instruction("REF", "")

    def it_appends_extra_raw_switches(self):
        from docx.fields import build_cross_reference_instruction

        instr = build_cross_reference_instruction(
            "REF",
            "Ref1",
            insert_as_hyperlink=True,
            extra_switches=["\\n"],
        )
        assert instr == "REF Ref1 \\h \\n"


class DescribeParagraph_add_cross_reference:
    """`paragraph.add_cross_reference()` cross-reference emitter."""

    def it_emits_a_REF_field_with_hyperlink_switch(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        xref = para.add_cross_reference(
            "REF", "Ref1", insert_as_hyperlink=True, cached_result="Chapter 1"
        )
        assert xref.ref_type == "REF"
        assert xref.target_name == "Ref1"
        assert xref.insert_as_hyperlink is True
        assert xref.insert_paragraph_number is False
        assert xref.insert_relative_position is False
        assert xref.result_text == "Chapter 1"

    def it_emits_a_PAGEREF_field(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        xref = para.add_cross_reference(
            "PAGEREF", "Ref1", insert_as_hyperlink=True, cached_result="42"
        )
        assert xref.ref_type == "PAGEREF"
        assert xref.result_text == "42"
        assert xref.instruction == "PAGEREF Ref1 \\h"

    def it_emits_a_NOTEREF_field(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        xref = para.add_cross_reference(
            "NOTEREF", "footnote_Ref1", cached_result="1"
        )
        assert xref.ref_type == "NOTEREF"
        assert xref.target_name == "footnote_Ref1"
        assert xref.result_text == "1"

    def it_emits_all_three_ref_switches(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        xref = para.add_cross_reference(
            "REF",
            "Ref1",
            insert_as_hyperlink=True,
            insert_paragraph_number=True,
            insert_relative_position=True,
        )
        assert xref.insert_as_hyperlink is True
        assert xref.insert_paragraph_number is True
        assert xref.insert_relative_position is True
        assert xref.instruction == "REF Ref1 \\h \\r \\p"

    def it_returns_a_CrossReference_subclass_of_Field(self):
        from docx.fields import CrossReference, Field
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        xref = para.add_cross_reference("REF", "Ref1")
        assert isinstance(xref, CrossReference)
        assert isinstance(xref, Field)
        assert xref.is_complex is True

    def it_omits_result_run_when_cached_result_is_None(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_cross_reference("REF", "Ref1")
        # -- four runs: begin, instrText, separate, end (no result) --
        runs = p.xpath("./w:r")
        assert len(runs) == 4


class DescribeCrossReference:
    """`CrossReference` proxy behaviours."""

    def _doc_with_bookmark(self, name: str, text: str):
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                f"w:document/w:body/w:p/("
                f"w:bookmarkStart{{w:id=0,w:name={name}}}"
                f",w:r/w:t\"{text}\""
                f",w:bookmarkEnd{{w:id=0}}"
                f")"
            ),
        )
        return Document(doc_elm, None)  # type: ignore[arg-type]

    def it_returns_empty_target_name_for_bare_instruction(self):
        from docx.fields import CrossReference
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "REF")
        xref = CrossReference("simple", fldSimple)
        assert xref.target_name == ""

    def it_parses_switches_case_insensitively(self):
        from docx.fields import CrossReference
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "REF Ref1 \\H")
        xref = CrossReference("simple", fldSimple)
        assert xref.insert_as_hyperlink is True

    def it_resolves_target_bookmark_to_a_Bookmark_proxy(self):
        from docx.bookmarks import Bookmark

        doc = self._doc_with_bookmark("Ref1", "Chapter 1")
        # -- add a cross-reference field pointing at the same bookmark --
        paragraph = list(doc.paragraphs)[0]
        xref = paragraph.add_cross_reference("REF", "Ref1", insert_as_hyperlink=True)

        bm = xref.target_bookmark(doc)
        assert bm is not None
        assert isinstance(bm, Bookmark)
        assert bm.name == "Ref1"

    def it_returns_None_target_bookmark_when_missing(self):
        doc = self._doc_with_bookmark("OtherRef", "x")
        paragraph = list(doc.paragraphs)[0]
        xref = paragraph.add_cross_reference("REF", "MissingRef")

        assert xref.target_bookmark(doc) is None

    def it_resolves_REF_to_bookmark_text_via_inherited_resolve(self):
        doc = self._doc_with_bookmark("Ref1", "Chapter 1")
        paragraph = list(doc.paragraphs)[0]
        xref = paragraph.add_cross_reference(
            "REF", "Ref1", insert_as_hyperlink=True, cached_result="stale"
        )

        assert xref.resolve(doc) == "Chapter 1"

    def it_round_trips_via_serialize_and_reparse(self):
        from docx.fields import CrossReference
        from docx.oxml.parser import parse_xml
        from docx.text.paragraph import Paragraph
        from lxml import etree

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_cross_reference(
            "PAGEREF",
            "heading1",
            insert_as_hyperlink=True,
            cached_result="7",
        )

        # -- serialise with no pretty-print (pretty-print inserts whitespace
        #    between runs which the complex-field walker would emit into the
        #    result text) and reparse via the docx oxml parser so CT_P / CT_R
        #    classes are re-registered and sibling iteration works. --
        xml = etree.tostring(p, pretty_print=False).decode()
        reparsed = cast(CT_P, parse_xml(xml))
        begin_run = reparsed.xpath(
            ".//w:fldChar[@w:fldCharType='begin']/parent::w:r"
        )[0]
        xref = CrossReference("complex", begin_run)
        assert xref.ref_type == "PAGEREF"
        assert xref.target_name == "heading1"
        assert xref.insert_as_hyperlink is True
        assert xref.result_text == "7"


class DescribeField_as_cross_reference:
    """`Field.as_cross_reference` upcast accessor."""

    def it_returns_a_CrossReference_for_REF_fields(self):
        from docx.fields import CrossReference, Field
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "REF Ref1 \\h")
        field = Field.for_simple(fldSimple)

        xref = field.as_cross_reference
        assert isinstance(xref, CrossReference)
        assert xref.ref_type == "REF"

    @pytest.mark.parametrize(
        "ref_type",
        ["REF", "PAGEREF", "NOTEREF", "SEQREF", "STYLEREF"],
    )
    def it_returns_a_CrossReference_for_every_xref_type(self, ref_type: str):
        from docx.fields import Field
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), f"{ref_type} Tgt")
        field = Field.for_simple(fldSimple)

        xref = field.as_cross_reference
        assert xref is not None
        assert xref.ref_type == ref_type

    def it_returns_None_for_non_cross_reference_fields(self):
        from docx.fields import Field
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "PAGE")
        field = Field.for_simple(fldSimple)

        assert field.as_cross_reference is None


# -- Table-of-contents family (R8-1) -------------------------------------


class DescribeParseTocInstruction:
    """`docx.fields.parse_toc_instruction` TOC-aware parser."""

    def it_parses_outline_range_as_switch_argument(self):
        from docx.fields import parse_toc_instruction

        parsed = parse_toc_instruction('TOC \\o "1-3" \\h \\z \\u')
        assert parsed.name == "TOC"
        # -- \o is argument-taking in a TOC context; it picks up "1-3" --
        assert parsed.switches == {"O": "1-3", "H": "", "Z": "", "U": ""}
        # -- nothing spills into args --
        assert parsed.args == []

    def it_records_custom_styles_verbatim_in_t_switch(self):
        from docx.fields import parse_toc_instruction

        parsed = parse_toc_instruction(
            'TOC \\o "1-3" \\t "Quote,1,Intense Quote,2"'
        )
        assert parsed.switches["T"] == "Quote,1,Intense Quote,2"

    def it_treats_bare_n_switch_as_flag(self):
        from docx.fields import parse_toc_instruction

        parsed = parse_toc_instruction('TOC \\o "1-3" \\n')
        assert parsed.switches == {"O": "1-3", "N": ""}

    def it_parses_n_switch_range_when_followed_by_argument(self):
        from docx.fields import parse_toc_instruction

        parsed = parse_toc_instruction('TOC \\o "1-3" \\n "3-5"')
        assert parsed.switches["N"] == "3-5"

    def it_parses_caption_label_c_switch(self):
        from docx.fields import parse_toc_instruction

        parsed = parse_toc_instruction('TOC \\c "Figure"')
        assert parsed.switches == {"C": "Figure"}

    def it_parses_bookmark_b_switch(self):
        from docx.fields import parse_toc_instruction

        parsed = parse_toc_instruction('TOC \\b "Chapter1" \\o "1-3"')
        assert parsed.switches["B"] == "Chapter1"
        assert parsed.switches["O"] == "1-3"

    def it_parses_separator_p_switch(self):
        from docx.fields import parse_toc_instruction

        parsed = parse_toc_instruction('TOC \\o "1-3" \\p "..."')
        assert parsed.switches["P"] == "..."


class DescribeBuildTocFieldInstruction:
    """`docx.fields.build_toc_field_instruction` builder."""

    def it_builds_word_default_toc_instruction(self):
        from docx.fields import build_toc_field_instruction

        assert (
            build_toc_field_instruction()
            == 'TOC \\o "1-3" \\h \\z \\u'
        )

    def it_respects_explicit_heading_range(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(heading_range=(2, 4))
        assert '\\o "2-4"' in instr

    def it_omits_heading_range_switch_when_None(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(heading_range=None)
        assert "\\o" not in instr

    def it_emits_custom_styles_as_comma_list(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(
            custom_styles=[("Quote", 1), ("Intense Quote", 2)],
        )
        assert '\\t "Quote,1,Intense Quote,2"' in instr

    def it_emits_caption_label_switch(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(
            heading_range=None,
            hyperlinks=False,
            hide_in_web=False,
            use_outline_levels=False,
            caption_label="Figure",
        )
        assert instr == 'TOC \\c "Figure"'

    def it_emits_bare_n_switch_for_zero_zero_range(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(
            heading_range=(1, 3), omit_page_numbers_range=(0, 0)
        )
        assert "\\n" in instr
        assert '\\n "' not in instr  # -- bare, no argument --

    def it_emits_n_switch_with_quoted_range(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(
            heading_range=(1, 5), omit_page_numbers_range=(3, 5)
        )
        assert '\\n "3-5"' in instr

    def it_emits_bookmark_name_switch(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(bookmark_name="Chapter1")
        assert '\\b "Chapter1"' in instr

    def it_emits_separator_switch(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(separator="-")
        assert '\\p "-"' in instr

    def it_appends_extra_switches_verbatim(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(extra_switches=["\\w", "\\x"])
        assert instr.endswith("\\w \\x")

    def it_uppercases_field_type(self):
        from docx.fields import build_toc_field_instruction

        instr = build_toc_field_instruction(field_type="toc", heading_range=None)
        assert instr.startswith("TOC")

    def it_accepts_toa_and_tof(self):
        from docx.fields import build_toc_field_instruction

        assert build_toc_field_instruction(
            field_type="TOA",
            heading_range=None,
            hyperlinks=False,
            hide_in_web=False,
            use_outline_levels=False,
        ) == "TOA"
        assert build_toc_field_instruction(
            field_type="TOF",
            heading_range=None,
            hyperlinks=False,
            hide_in_web=False,
            use_outline_levels=False,
        ) == "TOF"

    def it_raises_on_unknown_field_type(self):
        from docx.fields import build_toc_field_instruction

        with pytest.raises(ValueError):
            build_toc_field_instruction(field_type="WEIRD")

    def it_raises_on_bad_heading_range(self):
        from docx.fields import build_toc_field_instruction

        with pytest.raises(ValueError):
            build_toc_field_instruction(heading_range=(5, 2))

    def it_raises_on_bad_omit_page_numbers_range(self):
        from docx.fields import build_toc_field_instruction

        with pytest.raises(ValueError):
            build_toc_field_instruction(omit_page_numbers_range=(5, 2))


class DescribeParagraph_add_toc:
    """`paragraph.add_toc()` TOC field emitter."""

    def it_emits_a_TOC_field_with_word_defaults(self):
        from docx.fields import TocField
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc()
        assert isinstance(toc, TocField)
        assert toc.is_complex is True
        assert toc.heading_range == (1, 3)
        assert toc.hyperlinks_enabled is True
        assert toc.hide_in_web is True
        assert toc.use_outline_levels is True

    def it_marks_the_field_dirty_by_default(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc()
        assert toc.is_dirty is True

    def it_skips_the_dirty_flag_when_asked(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(mark_dirty=False)
        assert toc.is_dirty is False

    def it_respects_heading_range_argument(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(heading_range=(2, 5))
        assert toc.heading_range == (2, 5)

    def it_passes_custom_styles_through(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(
            custom_styles=[("Quote", 1), ("Intense Quote", 2)]
        )
        assert toc.custom_styles == [("Quote", 1), ("Intense Quote", 2)]

    def it_passes_bookmark_name_through(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(bookmark_name="Chapter1")
        assert toc.bookmark_name == "Chapter1"

    def it_passes_separator_through(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(separator="-")
        assert toc.separator == "-"

    def it_emits_omit_page_numbers_bare_n_for_all_levels(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(omit_page_numbers_range=(0, 0))
        assert toc.omit_page_numbers_range == (0, 0)

    def it_emits_omit_page_numbers_range(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(omit_page_numbers_range=(3, 5))
        assert toc.omit_page_numbers_range == (3, 5)

    def it_stores_cached_result_between_separate_and_end(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toc = para.add_toc(cached_result="Heading One\t1\nHeading Two\t2")
        assert "Heading One" in toc.result_text
        assert "Heading Two" in toc.result_text


class DescribeParagraph_add_table_of_figures:
    """`paragraph.add_table_of_figures()` emitter."""

    def it_emits_a_TOC_with_caption_label(self):
        from docx.fields import TableOfFiguresField, TocField
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        tof = para.add_table_of_figures(caption_label="Figure")
        assert isinstance(tof, TableOfFiguresField)
        assert isinstance(tof, TocField)
        assert tof.caption_label == "Figure"
        # -- type is still "TOC" (the \c shape distinguishes it) --
        assert tof.type == "TOC"

    def it_defaults_caption_label_to_Figure(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        tof = para.add_table_of_figures()
        assert tof.caption_label == "Figure"

    def it_accepts_arbitrary_caption_labels(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        tof = para.add_table_of_figures(caption_label="Illustration")
        assert tof.caption_label == "Illustration"


class DescribeParagraph_add_table_of_authorities:
    """`paragraph.add_table_of_authorities()` emitter."""

    def it_emits_a_TOA_field(self):
        from docx.fields import TableOfAuthoritiesField
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toa = para.add_table_of_authorities(category=1)
        assert isinstance(toa, TableOfAuthoritiesField)
        assert toa.type == "TOA"
        assert toa.category == 1

    def it_omits_category_switch_when_category_is_None(self):
        from docx.text.paragraph import Paragraph

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        toa = para.add_table_of_authorities()
        assert toa.category is None
        assert "\\c" not in toa.instruction


class DescribeTocField_properties:
    """`TocField` switch accessors against synthetic fldSimple elements."""

    def _toc(self, instr: str):
        from docx.fields import TocField
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), instr)
        return TocField("simple", fldSimple)

    def it_reads_heading_range_from_o_switch(self):
        toc = self._toc('TOC \\o "1-3" \\h')
        assert toc.heading_range == (1, 3)

    def it_returns_None_heading_range_when_o_absent(self):
        toc = self._toc("TOC \\h")
        assert toc.heading_range is None

    def it_returns_None_heading_range_when_o_malformed(self):
        toc = self._toc('TOC \\o "badrange"')
        assert toc.heading_range is None

    def it_reads_flag_switches(self):
        toc = self._toc('TOC \\o "1-3" \\h \\z \\u')
        assert toc.hyperlinks_enabled is True
        assert toc.hide_in_web is True
        assert toc.use_outline_levels is True

    def it_returns_False_for_missing_flag_switches(self):
        toc = self._toc('TOC \\o "1-3"')
        assert toc.hyperlinks_enabled is False
        assert toc.hide_in_web is False
        assert toc.use_outline_levels is False

    def it_parses_case_insensitive_switch_letters(self):
        toc = self._toc('TOC \\o "1-3" \\H \\Z \\U')
        assert toc.hyperlinks_enabled is True
        assert toc.hide_in_web is True
        assert toc.use_outline_levels is True

    def it_reads_custom_styles_from_t_switch(self):
        toc = self._toc('TOC \\t "Quote,1,Intense Quote,2"')
        assert toc.custom_styles == [("Quote", 1), ("Intense Quote", 2)]

    def it_returns_empty_custom_styles_for_malformed_t(self):
        toc = self._toc('TOC \\t "oddnumberoftokens"')
        assert toc.custom_styles == []

    def it_reads_bookmark_name_from_b_switch(self):
        toc = self._toc('TOC \\b "MyChapter" \\o "1-3"')
        assert toc.bookmark_name == "MyChapter"

    def it_returns_None_bookmark_when_b_absent(self):
        toc = self._toc('TOC \\o "1-3"')
        assert toc.bookmark_name is None

    def it_reads_separator_from_p_switch(self):
        toc = self._toc('TOC \\o "1-3" \\p "..."')
        assert toc.separator == "..."

    def it_reads_caption_label_from_c_switch(self):
        toc = self._toc('TOC \\c "Figure"')
        assert toc.caption_label == "Figure"

    def it_returns_zero_tuple_for_bare_n_switch(self):
        toc = self._toc('TOC \\o "1-3" \\n')
        assert toc.omit_page_numbers_range == (0, 0)

    def it_returns_range_for_n_switch_with_argument(self):
        toc = self._toc('TOC \\o "1-5" \\n "3-5"')
        assert toc.omit_page_numbers_range == (3, 5)

    def it_returns_None_when_n_absent(self):
        toc = self._toc('TOC \\o "1-3"')
        assert toc.omit_page_numbers_range is None


class DescribeTocField_round_trip:
    """`TocField` survives serialise/reparse."""

    def it_round_trips_all_switches_through_xml(self):
        from docx.fields import TocField
        from docx.oxml.parser import parse_xml
        from docx.text.paragraph import Paragraph
        from lxml import etree

        p = cast(CT_P, element("w:p"))
        para = Paragraph(p, None)  # type: ignore[arg-type]
        para.add_toc(
            heading_range=(1, 4),
            hyperlinks=True,
            hide_in_web=True,
            use_outline_levels=False,
            omit_page_numbers_range=(3, 4),
            separator="-",
            custom_styles=[("Quote", 1), ("Intense Quote", 2)],
            bookmark_name="Chapter1",
        )

        xml = etree.tostring(p, pretty_print=False).decode()
        reparsed = cast(CT_P, parse_xml(xml))
        begin_run = reparsed.xpath(
            ".//w:fldChar[@w:fldCharType='begin']/parent::w:r"
        )[0]
        toc = TocField("complex", begin_run)

        assert toc.heading_range == (1, 4)
        assert toc.hyperlinks_enabled is True
        assert toc.hide_in_web is True
        assert toc.use_outline_levels is False
        assert toc.omit_page_numbers_range == (3, 4)
        assert toc.separator == "-"
        assert toc.custom_styles == [("Quote", 1), ("Intense Quote", 2)]
        assert toc.bookmark_name == "Chapter1"


class DescribeField_as_toc:
    """`Field.as_toc` upcast accessor."""

    def it_returns_a_TocField_for_plain_TOC(self):
        from docx.fields import Field, TableOfFiguresField, TocField
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), 'TOC \\o "1-3" \\h')
        toc = Field.for_simple(fldSimple).as_toc
        assert isinstance(toc, TocField)
        assert not isinstance(toc, TableOfFiguresField)

    def it_returns_TableOfFiguresField_for_TOC_with_c_switch(self):
        from docx.fields import Field, TableOfFiguresField
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), 'TOC \\c "Figure"')
        toc = Field.for_simple(fldSimple).as_toc
        assert isinstance(toc, TableOfFiguresField)

    def it_returns_TableOfAuthoritiesField_for_TOA(self):
        from docx.fields import Field, TableOfAuthoritiesField
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), 'TOA \\c "1"')
        toc = Field.for_simple(fldSimple).as_toc
        assert isinstance(toc, TableOfAuthoritiesField)

    def it_returns_TableOfFiguresField_for_bare_TOF(self):
        from docx.fields import Field, TableOfFiguresField
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "TOF")
        toc = Field.for_simple(fldSimple).as_toc
        assert isinstance(toc, TableOfFiguresField)

    def it_returns_None_for_non_TOC_fields(self):
        from docx.fields import Field
        from docx.oxml.ns import qn
        from docx.oxml.parser import OxmlElement

        fldSimple = cast(CT_FldSimple, OxmlElement("w:fldSimple"))
        fldSimple.set(qn("w:instr"), "REF Ref1")
        assert Field.for_simple(fldSimple).as_toc is None


class DescribeTocField_rebuild:
    """`TocField.rebuild` populates the cached result from document headings."""

    def it_populates_cached_result_with_every_heading_in_range(self):
        import docx

        document = docx.Document()
        document.add_heading("Alpha", level=1)
        document.add_heading("Beta", level=2)
        document.add_heading("Gamma", level=3)
        paragraph = document.add_paragraph()

        toc = paragraph.add_toc()
        toc.rebuild()

        assert "Alpha" in toc.result_text
        assert "Beta" in toc.result_text
        assert "Gamma" in toc.result_text

    def it_uses_the_placeholder_for_page_numbers(self):
        import docx

        document = docx.Document()
        document.add_heading("Only", level=1)
        paragraph = document.add_paragraph()
        toc = paragraph.add_toc()
        toc.rebuild()
        # -- default placeholder is "?" --
        assert toc.result_text == "Only\t?"

    def it_respects_heading_range_switch(self):
        import docx

        document = docx.Document()
        document.add_heading("H1", level=1)
        document.add_heading("H2", level=2)
        document.add_heading("H3", level=3)
        paragraph = document.add_paragraph()
        toc = paragraph.add_toc(heading_range=(1, 2))
        toc.rebuild()
        assert "H1" in toc.result_text
        assert "H2" in toc.result_text
        assert "H3" not in toc.result_text

    def it_preserves_the_dirty_flag(self):
        import docx

        document = docx.Document()
        document.add_heading("Alpha", level=1)
        paragraph = document.add_paragraph()
        toc = paragraph.add_toc()
        assert toc.is_dirty is True
        toc.rebuild()
        assert toc.is_dirty is True

    def it_accepts_a_custom_placeholder(self):
        import docx

        document = docx.Document()
        document.add_heading("Alpha", level=1)
        paragraph = document.add_paragraph()
        toc = paragraph.add_toc()
        toc.rebuild(page_number_placeholder="—")
        assert toc.result_text == "Alpha\t—"

    def it_writes_empty_result_when_document_has_no_headings(self):
        import docx

        document = docx.Document()
        document.add_paragraph("just a body paragraph")
        paragraph = document.add_paragraph()
        toc = paragraph.add_toc()
        toc.rebuild()
        assert toc.result_text == ""


class DescribeTableOfFiguresField_rebuild:
    """`TableOfFiguresField.rebuild` populates the cached result from captions."""

    def it_collects_caption_paragraphs_matching_the_label(self):
        import docx

        document = docx.Document()
        document.add_paragraph("Figure 1: Alpha", style="Caption")
        document.add_paragraph("Table 1: Not a figure", style="Caption")
        document.add_paragraph("Figure 2: Beta", style="Caption")
        paragraph = document.add_paragraph()

        tof = paragraph.add_table_of_figures(caption_label="Figure")
        tof.rebuild()

        assert "Figure 1: Alpha" in tof.result_text
        assert "Figure 2: Beta" in tof.result_text
        assert "Table 1" not in tof.result_text


class DescribeDocument_rebuild_tocs:
    """`Document.rebuild_tocs` rebuilds every TOC-family field in the body."""

    def it_rebuilds_every_TOC_and_returns_the_count(self):
        import docx

        document = docx.Document()
        document.add_heading("One", level=1)
        document.add_paragraph("Figure 1: Diagram", style="Caption")
        document.add_heading("Two", level=1)

        p_toc = document.add_paragraph()
        p_toc.add_toc()
        p_tof = document.add_paragraph()
        p_tof.add_table_of_figures(caption_label="Figure")

        count = document.rebuild_tocs()
        assert count == 2

        toc_field = document.fields[0].as_toc
        tof_field = document.fields[1].as_toc
        assert toc_field is not None and tof_field is not None
        assert "One" in toc_field.result_text
        assert "Two" in toc_field.result_text
        assert "Figure 1: Diagram" in tof_field.result_text
