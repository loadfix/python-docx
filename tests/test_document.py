# pyright: reportPrivateUsage=false
# pyright: reportUnknownMemberType=false

"""Unit test suite for the docx.document module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.comments import Comment, Comments
from docx.custom_properties import CustomProperties
from docx.document import Document, _Body
from docx.enum.section import WD_SECTION
from docx.font_table import FontTable
from docx.glossary import Glossary
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.coreprops import CoreProperties
from docx.oxml.document import CT_Body, CT_Document
from docx.parts.document import DocumentPart
from docx.section import Section, Sections
from docx.settings import Settings
from docx.shape import InlineShape, InlineShapes
from docx.shared import Inches, Length, RGBColor
from docx.styles.styles import Styles
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.theme import Theme
from docx.web_settings import WebSettings

from .unitutil.cxml import element, xml
from .unitutil.mock import (
    FixtureRequest,
    Mock,
    class_mock,
    instance_mock,
    method_mock,
    property_mock,
)


class DescribeDocument:
    """Unit-test suite for `docx.document.Document`."""

    def it_can_add_a_comment(
        self,
        document_part_: Mock,
        comments_prop_: Mock,
        comments_: Mock,
        comment_: Mock,
        run_mark_comment_range_: Mock,
    ):
        comment_.comment_id = 42
        comments_.add_comment.return_value = comment_
        comments_prop_.return_value = comments_
        document = Document(cast(CT_Document, element("w:document/w:body/w:p/w:r")), document_part_)
        run = document.paragraphs[0].runs[0]

        comment = document.add_comment(run, "Comment text.")

        # Document.add_comment forwards to Comments.add_comment via kwargs
        # since 2026.05.5 (the `date=` pass-through landed then).
        comments_.add_comment.assert_called_once_with(
            text="Comment text.", author="", initials="", date=None
        )
        run_mark_comment_range_.assert_called_once_with(run, run, 42)
        assert comment is comment_

    def it_forwards_an_explicit_date_to_the_comments_collection(
        self,
        document_part_: Mock,
        comments_prop_: Mock,
        comments_: Mock,
        comment_: Mock,
        run_mark_comment_range_: Mock,
    ):
        import datetime as dt

        comment_.comment_id = 7
        comments_.add_comment.return_value = comment_
        comments_prop_.return_value = comments_
        document = Document(cast(CT_Document, element("w:document/w:body/w:p/w:r")), document_part_)
        run = document.paragraphs[0].runs[0]
        fixed = dt.datetime(2026, 5, 4, tzinfo=dt.timezone.utc)

        document.add_comment(run, "x", author="A", initials="A", date=fixed)

        comments_.add_comment.assert_called_once_with(
            text="x", author="A", initials="A", date=fixed
        )

    @pytest.mark.parametrize(
        ("level", "style"), [(0, "Title"), (1, "Heading 1"), (2, "Heading 2"), (9, "Heading 9")]
    )
    def it_can_add_a_heading(
        self, level: int, style: str, document: Document, add_paragraph_: Mock, paragraph_: Mock
    ):
        add_paragraph_.return_value = paragraph_

        paragraph = document.add_heading("Spam vs. Bacon", level)

        add_paragraph_.assert_called_once_with(document, "Spam vs. Bacon", style)
        assert paragraph is paragraph_

    def it_raises_on_heading_level_out_of_range(self, document: Document):
        with pytest.raises(ValueError, match="level must be in range 0-9, got -1"):
            document.add_heading(level=-1)
        with pytest.raises(ValueError, match="level must be in range 0-9, got 10"):
            document.add_heading(level=10)

    def it_can_add_a_page_break(
        self, document: Document, add_paragraph_: Mock, paragraph_: Mock
    ):
        add_paragraph_.return_value = paragraph_
        paragraph_.add_page_break.return_value = paragraph_

        paragraph = document.add_page_break()

        add_paragraph_.assert_called_once_with(document)
        paragraph_.add_page_break.assert_called_once_with()
        assert paragraph is paragraph_

    @pytest.mark.parametrize(
        ("text", "style"), [("", None), ("", "Heading 1"), ("foo\rbar", "Body Text")]
    )
    def it_can_add_a_paragraph(
        self,
        text: str,
        style: str | None,
        document: Document,
        body_: Mock,
        body_prop_: Mock,
        paragraph_: Mock,
    ):
        body_prop_.return_value = body_
        body_.add_paragraph.return_value = paragraph_

        paragraph = document.add_paragraph(text, style)

        body_.add_paragraph.assert_called_once_with(text, style)
        assert paragraph is paragraph_

    def it_can_add_a_picture(
        self, document: Document, add_paragraph_: Mock, run_: Mock, picture_: Mock
    ):
        path, width, height = "foobar.png", 100, 200
        add_paragraph_.return_value.add_run.return_value = run_
        run_.add_picture.return_value = picture_

        picture = document.add_picture(path, width, height)

        run_.add_picture.assert_called_once_with(
            path, width, height, link=False, save_with_document=True, url=None
        )
        assert picture is picture_

    def it_can_add_a_shape(self):
        from docx import Document as OpenDocument
        from docx.drawing import WordprocessingShape
        from docx.enum.shape import WD_SHAPE
        from docx.shared import Inches

        document = OpenDocument()

        shape = document.add_shape(
            WD_SHAPE.ROUNDED_RECTANGLE, Inches(2), Inches(1), text="Hi"
        )

        assert isinstance(shape, WordprocessingShape)
        assert shape.shape_type is WD_SHAPE.ROUNDED_RECTANGLE
        assert shape.text == "Hi"

    def it_round_trips_a_created_shape(self):
        import io

        from docx import Document as OpenDocument
        from docx.enum.shape import WD_SHAPE
        from docx.shared import Inches

        document = OpenDocument()
        document.add_shape(WD_SHAPE.OVAL, Inches(1), Inches(1), text="Egg")

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = OpenDocument(buf)

        wsp_list = reopened.element.body.xpath(".//wps:wsp")
        assert len(wsp_list) == 1
        prstGeoms = reopened.element.body.xpath(".//a:prstGeom/@prst")
        assert "ellipse" in prstGeoms

    def it_can_add_a_canvas(self):
        from docx import Document as OpenDocument
        from docx.drawing import Canvas
        from docx.enum.shape import WD_SHAPE
        from docx.shared import Inches

        document = OpenDocument()

        canvas = document.add_canvas(Inches(4), Inches(2))
        shape = canvas.add_shape(WD_SHAPE.RECTANGLE, Inches(1), Inches(1))

        assert isinstance(canvas, Canvas)
        assert canvas.shapes[0].shape_type is WD_SHAPE.RECTANGLE
        assert shape.shape_type is WD_SHAPE.RECTANGLE
        # -- the canvas is wrapped in a w:drawing under a new body paragraph --
        wpc_list = document.element.body.xpath(".//wpc:wpc")
        assert len(wpc_list) == 1

    def it_can_add_a_text_box(self):
        from docx import Document as OpenDocument
        from docx.drawing import WordprocessingShape
        from docx.shared import Inches

        document = OpenDocument()

        text_box = document.add_text_box(Inches(3), Inches(1), text="Note")
        text_box.add_paragraph("Second paragraph")

        assert isinstance(text_box, WordprocessingShape)
        assert "Note" in text_box.text
        assert "Second paragraph" in text_box.text
        # -- a wps:txbx was emitted in the body --
        txbx_list = document.element.body.xpath(".//wps:txbx")
        assert len(txbx_list) == 1

    def it_accepts_a_PathLike_image_path_for_add_picture(
        self, document: Document, add_paragraph_: Mock, run_: Mock, picture_: Mock
    ):
        # -- upstream-PR#1168: os.fspath() at entry so pathlib.Path works --
        import os
        from pathlib import Path

        add_paragraph_.return_value.add_run.return_value = run_
        run_.add_picture.return_value = picture_

        picture = document.add_picture(Path("foobar.png"), 100, 200)

        run_.add_picture.assert_called_once_with(
            os.fspath(Path("foobar.png")),
            100,
            200,
            link=False,
            save_with_document=True,
            url=None,
        )
        assert picture is picture_

    def it_returns_empty_charts_for_a_chartless_document(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        assert document.charts == []

    def it_can_add_and_read_back_a_bar_chart(self):
        from docx import Document as OpenDocument
        from docx.chart import Chart as ChartProxy
        from docx.chart import WD_CHART_TYPE

        document = OpenDocument()
        chart = document.add_chart(
            WD_CHART_TYPE.BAR, ["a", "b", "c"], {"Series 1": [1.0, 2.0, 3.0]}
        )
        assert isinstance(chart, ChartProxy)
        assert chart.chart_type is WD_CHART_TYPE.BAR
        assert [s.name for s in chart.series] == ["Series 1"]
        assert chart.series[0].values == [1.0, 2.0, 3.0]
        assert chart.categories == ["a", "b", "c"]

        charts = document.charts
        assert len(charts) == 1
        assert charts[0].chart_type is WD_CHART_TYPE.BAR

    @pytest.mark.parametrize(
        ("chart_type",),
        [
            ("BAR",),
            ("COLUMN",),
            ("LINE",),
            ("PIE",),
        ],
    )
    def it_roundtrips_each_supported_chart_type(self, chart_type: str):
        import io

        from docx import Document as OpenDocument
        from docx.chart import WD_CHART_TYPE

        ct = WD_CHART_TYPE[chart_type]
        document = OpenDocument()
        document.add_chart(ct, ["x", "y"], {"S": [11.0, 22.0]})

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = OpenDocument(buf)

        charts = reopened.charts
        assert len(charts) == 1
        assert charts[0].chart_type is ct
        assert charts[0].categories == ["x", "y"]
        assert charts[0].series[0].name == "S"
        assert charts[0].series[0].values == [11.0, 22.0]

    def it_can_add_multiple_charts_to_one_document(self):
        import io

        from docx import Document as OpenDocument
        from docx.chart import WD_CHART_TYPE

        document = OpenDocument()
        document.add_chart(
            WD_CHART_TYPE.BAR, ["a", "b"], {"S1": [1.0, 2.0]}
        )
        document.add_chart(
            WD_CHART_TYPE.PIE, ["a", "b"], {"S2": [3.0, 4.0]}
        )

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = OpenDocument(buf)

        kinds = [c.chart_type for c in reopened.charts]
        assert kinds == [WD_CHART_TYPE.BAR, WD_CHART_TYPE.PIE]

    @pytest.mark.parametrize(
        ("sentinel_cxml", "start_type", "new_sentinel_cxml"),
        [
            ("w:sectPr", WD_SECTION.EVEN_PAGE, "w:sectPr/w:type{w:val=evenPage}"),
            (
                "w:sectPr/w:type{w:val=evenPage}",
                WD_SECTION.ODD_PAGE,
                "w:sectPr/w:type{w:val=oddPage}",
            ),
            ("w:sectPr/w:type{w:val=oddPage}", WD_SECTION.NEW_PAGE, "w:sectPr"),
        ],
    )
    def it_can_add_a_section(
        self,
        sentinel_cxml: str,
        start_type: WD_SECTION,
        new_sentinel_cxml: str,
        Section_: Mock,
        section_: Mock,
        document_part_: Mock,
    ):
        Section_.return_value = section_
        document = Document(
            cast(CT_Document, element("w:document/w:body/(w:p,%s)" % sentinel_cxml)),
            document_part_,
        )

        section = document.add_section(start_type)

        assert document.element.xml == xml(
            "w:document/w:body/(w:p,w:p/w:pPr/%s,%s)" % (sentinel_cxml, new_sentinel_cxml)
        )
        sectPr = document.element.xpath("w:body/w:sectPr")[0]
        Section_.assert_called_once_with(sectPr, document_part_)
        assert section is section_

    def it_can_add_a_table(
        self,
        document: Document,
        _block_width_prop_: Mock,
        body_prop_: Mock,
        body_: Mock,
        table_: Mock,
    ):
        rows, cols, style = 4, 2, "Light Shading Accent 1"
        body_prop_.return_value = body_
        body_.add_table.return_value = table_
        _block_width_prop_.return_value = width = 42

        table = document.add_table(rows, cols, style)

        body_.add_table.assert_called_once_with(rows, cols, width)
        assert table == table_
        assert table.style == style

    def it_rolls_back_add_table_when_style_is_invalid(
        self, document_part_: Mock
    ):
        # -- regression for upstream#563: a freshly-added w:tbl must not be
        # -- left in the body when the supplied style does not exist. The
        # -- caller sees the exception and the document is unchanged. --
        from docx.styles.style import BaseStyle

        document = Document(
            cast(
                CT_Document,
                element(
                    "w:document/w:body/w:sectPr/"
                    "(w:pgSz{w:w=12240,w:h=15840},"
                    "w:pgMar{w:top=1440,w:right=1440,w:bottom=1440,w:left=1440})"
                ),
            ),
            document_part_,
        )
        # -- simulate the style lookup raising for an invalid style name --
        document_part_.get_style_id.side_effect = KeyError("no style")
        # -- body initially has no table --
        tbls_before = document._element.body.xpath(".//w:tbl")
        assert tbls_before == []

        with pytest.raises(KeyError):
            document.add_table(2, 2, "Not A Real Style")

        # -- the freshly added w:tbl was rolled back --
        tbls_after = document._element.body.xpath(".//w:tbl")
        assert tbls_after == []

    def it_uses_default_block_width_when_body_has_no_sectPr(
        self, document_part_: Mock
    ):
        # -- regression for upstream#514: a document whose body carries no
        # -- w:sectPr must still compute a sensible block width (used by
        # -- Document.add_table) rather than raising IndexError. --
        document = Document(
            cast(CT_Document, element("w:document/w:body")), document_part_
        )

        width = document._block_width

        # -- US-Letter default: 8.5" page minus 1" left + 1" right = 6.5" --
        assert width == Inches(6.5)

    def it_can_add_a_table_of_contents(self):
        # -- integration test: the TOC helper is thin enough that mocking
        #    every collaborator would just re-exercise plumbing; a full
        #    round-trip is a better signal. --
        from docx import Document as new_document

        document = new_document()
        document.add_heading("Chapter One", level=1)
        document.add_heading("Section 1.1", level=2)
        document.add_heading("Skipped detail", level=4)
        document.add_heading("Chapter Two", level=1)

        paragraph = document.add_table_of_contents(levels=(1, 3))

        assert isinstance(paragraph, Paragraph)
        assert document.paragraphs[-1]._p is paragraph._p
        assert len(paragraph.fields) == 1
        field = paragraph.fields[0]
        assert field.type == "TOC"
        assert field.instruction.strip() == 'TOC \\o "1-3" \\h \\z \\u'
        result = field.result_text
        assert "Chapter One" in result
        assert "Section 1.1" in result
        assert "Chapter Two" in result
        assert "Skipped detail" not in result

    def it_can_save_the_document_to_a_file(self, document_part_: Mock):
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        document.save("foobar.docx")

        document_part_.save.assert_called_once_with(
            "foobar.docx", reproducible=False, password=None
        )

    def it_knows_when_the_document_has_macros(self, document_part_: Mock):
        document = Document(cast(CT_Document, element("w:document")), document_part_)
        document_part_.part_related_by.return_value = Mock()

        assert document.has_macros is True
        document_part_.part_related_by.assert_called_once_with(RT.VBA_PROJECT)

    def it_knows_when_the_document_has_no_macros(self, document_part_: Mock):
        document = Document(cast(CT_Document, element("w:document")), document_part_)
        document_part_.part_related_by.side_effect = KeyError

        assert document.has_macros is False

    def it_provides_access_to_the_comments(self, document_part_: Mock, comments_: Mock):
        document_part_.comments = comments_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.comments is comments_

    def it_provides_access_to_its_core_properties(
        self, document_part_: Mock, core_properties_: Mock
    ):
        document_part_.core_properties = core_properties_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        core_properties = document.core_properties

        assert core_properties is core_properties_

    def it_provides_access_to_its_custom_properties(
        self, document_part_: Mock, custom_properties_: Mock
    ):
        document_part_.custom_properties = custom_properties_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        custom_properties = document.custom_properties

        assert custom_properties is custom_properties_

    def it_provides_access_to_its_font_table(
        self, document_part_: Mock, font_table_: Mock
    ):
        document_part_.font_table = font_table_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.font_table is font_table_

    def and_font_table_is_None_when_the_document_has_no_font_table_part(
        self, document_part_: Mock
    ):
        document_part_.font_table = None
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.font_table is None

    def it_provides_access_to_its_web_settings(
        self, document_part_: Mock, web_settings_: Mock
    ):
        document_part_.web_settings = web_settings_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.web_settings is web_settings_

    def and_web_settings_is_None_when_the_document_has_no_web_settings_part(
        self, document_part_: Mock
    ):
        document_part_.web_settings = None
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.web_settings is None

    def it_provides_access_to_its_theme(
        self, document_part_: Mock, theme_: Mock
    ):
        document_part_.theme = theme_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.theme is theme_

    def and_theme_is_None_when_the_document_has_no_theme_part(
        self, document_part_: Mock
    ):
        document_part_.theme = None
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.theme is None

    def it_provides_access_to_its_glossary(
        self, document_part_: Mock, glossary_: Mock
    ):
        document_part_.glossary = glossary_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.glossary is glossary_

    def and_glossary_is_None_when_the_document_has_no_glossary_part(
        self, document_part_: Mock
    ):
        document_part_.glossary = None
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.glossary is None

    def it_provides_access_to_its_inline_shapes(self, document_part_: Mock, inline_shapes_: Mock):
        document_part_.inline_shapes = inline_shapes_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.inline_shapes is inline_shapes_

    def it_can_iterate_the_inner_content_of_the_document(
        self, body_prop_: Mock, body_: Mock, document_part_: Mock
    ):
        body_prop_.return_value = body_
        body_.iter_inner_content.return_value = iter((1, 2, 3))
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert list(document.iter_inner_content()) == [1, 2, 3]

    def it_provides_access_to_its_paragraphs(
        self, document: Document, body_prop_: Mock, body_: Mock, paragraphs_: Mock
    ):
        body_prop_.return_value = body_
        body_.paragraphs = paragraphs_
        paragraphs = document.paragraphs
        assert paragraphs is paragraphs_

    def it_provides_access_to_its_sections(
        self, document_part_: Mock, Sections_: Mock, sections_: Mock
    ):
        document_elm = cast(CT_Document, element("w:document"))
        Sections_.return_value = sections_
        document = Document(document_elm, document_part_)

        sections = document.sections

        Sections_.assert_called_once_with(document_elm, document_part_)
        assert sections is sections_

    def it_provides_access_to_its_settings(self, document_part_: Mock, settings_: Mock):
        document_part_.settings = settings_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.settings is settings_

    def it_delegates_footnote_properties_to_settings(
        self, document_part_: Mock, settings_: Mock
    ):
        document_part_.settings = settings_
        settings_.footnote_properties = "fp-sentinel"
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.footnote_properties == "fp-sentinel"

    def it_delegates_add_footnote_properties_to_settings(
        self, document_part_: Mock, settings_: Mock
    ):
        document_part_.settings = settings_
        settings_.add_footnote_properties.return_value = "added-fp"
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        result = document.add_footnote_properties()

        settings_.add_footnote_properties.assert_called_once_with()
        assert result == "added-fp"

    def it_delegates_endnote_properties_to_settings(
        self, document_part_: Mock, settings_: Mock
    ):
        document_part_.settings = settings_
        settings_.endnote_properties = "ep-sentinel"
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.endnote_properties == "ep-sentinel"

    def it_delegates_add_endnote_properties_to_settings(
        self, document_part_: Mock, settings_: Mock
    ):
        document_part_.settings = settings_
        settings_.add_endnote_properties.return_value = "added-ep"
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        result = document.add_endnote_properties()

        settings_.add_endnote_properties.assert_called_once_with()
        assert result == "added-ep"

    def it_provides_access_to_its_styles(self, document_part_: Mock, styles_: Mock):
        document_part_.styles = styles_
        document = Document(cast(CT_Document, element("w:document")), document_part_)

        assert document.styles is styles_

    def it_provides_access_to_its_tables(
        self, document: Document, body_prop_: Mock, body_: Mock, tables_: Mock
    ):
        body_prop_.return_value = body_
        body_.tables = tables_

        assert document.tables is tables_

    def it_provides_access_to_the_document_part(self, document_part_: Mock):
        document = Document(cast(CT_Document, element("w:document")), document_part_)
        assert document.part is document_part_

    def it_provides_access_to_the_document_body(
        self, _Body_: Mock, body_: Mock, document_part_: Mock
    ):
        _Body_.return_value = body_
        document_elm = cast(CT_Document, element("w:document/w:body"))
        body_elm = document_elm[0]
        document = Document(document_elm, document_part_)

        body = document._body

        _Body_.assert_called_once_with(body_elm, document)
        assert body is body_

    def it_can_accept_all_tracked_changes(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                'w:p/(w:r/w:t"keep ",w:ins{w:id=1,w:author=A}/w:r/w:t"added"),'
                'w:p/(w:del{w:id=2,w:author=B}/w:r/w:delText"removed",w:r/w:t" end")'
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.accept_all_changes()

        assert count == 2
        assert document_elm.xpath(".//w:ins") == []
        assert document_elm.xpath(".//w:del") == []
        # Insertion flattened, deletion gone
        paragraphs = document_elm.xpath(".//w:p")
        assert "".join(t.text for t in paragraphs[0].xpath(".//w:t")) == "keep added"
        assert "".join(t.text for t in paragraphs[1].xpath(".//w:t")) == " end"

    def it_can_reject_all_tracked_changes(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                'w:p/(w:r/w:t"keep ",w:ins{w:id=1,w:author=A}/w:r/w:t"added"),'
                'w:p/(w:del{w:id=2,w:author=B}/w:r/w:delText"removed",w:r/w:t" end")'
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.reject_all_changes()

        assert count == 2
        assert document_elm.xpath(".//w:ins") == []
        assert document_elm.xpath(".//w:del") == []
        assert document_elm.xpath(".//w:delText") == []
        # Insertion removed, deletion restored
        paragraphs = document_elm.xpath(".//w:p")
        assert "".join(t.text for t in paragraphs[0].xpath(".//w:t")) == "keep "
        assert "".join(t.text for t in paragraphs[1].xpath(".//w:t")) == "removed end"

    def it_resolves_formatting_changes_on_accept(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/w:pPr/("
                "w:jc{w:val=center},"
                "w:pPrChange{w:id=5,w:author=C}/w:pPr/w:jc{w:val=left}"
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.accept_all_changes()

        assert count == 1
        # accept: pPrChange removed, current formatting (center) preserved
        assert document_elm.xpath(".//w:pPrChange") == []
        jc = document_elm.xpath(".//w:pPr/w:jc")
        assert len(jc) == 1
        assert jc[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        ) == "center"

    def it_resolves_formatting_changes_on_reject(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/w:pPr/("
                "w:jc{w:val=center},"
                "w:pPrChange{w:id=5,w:author=C}/w:pPr/w:jc{w:val=left}"
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.reject_all_changes()

        assert count == 1
        # reject: pPrChange gone, prior formatting (left) restored
        assert document_elm.xpath(".//w:pPrChange") == []
        jc = document_elm.xpath(".//w:pPr/w:jc")
        assert len(jc) == 1
        assert jc[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        ) == "left"

    def it_can_accept_move_revisions_across_the_document(self, document_part_: Mock):
        # -- source side (moveFrom) is removed entirely, destination (moveTo) is
        # -- unwrapped so its runs survive as live text --
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                'w:p/(w:r/w:t"before ",'
                'w:moveFrom{w:id=1,w:author=A,w:name=m1}/w:r/w:delText"moved"),'
                'w:p/(w:r/w:t"dest: ",'
                'w:moveTo{w:id=2,w:author=B,w:name=m1}/w:r/w:t"moved")'
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.accept_all_changes()

        assert count == 2
        assert document_elm.xpath(".//w:moveFrom") == []
        assert document_elm.xpath(".//w:moveTo") == []
        paragraphs = document_elm.xpath(".//w:p")
        assert "".join(t.text for t in paragraphs[0].xpath(".//w:t")) == "before "
        assert "".join(t.text for t in paragraphs[1].xpath(".//w:t")) == "dest: moved"

    def it_can_reject_move_revisions_across_the_document(self, document_part_: Mock):
        # -- source side (moveFrom) is unwrapped and its delText becomes t,
        # -- destination (moveTo) and its content are removed --
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                'w:p/(w:r/w:t"before ",'
                'w:moveFrom{w:id=1,w:author=A,w:name=m1}/w:r/w:delText"moved"),'
                'w:p/(w:r/w:t"dest: ",'
                'w:moveTo{w:id=2,w:author=B,w:name=m1}/w:r/w:t"moved")'
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.reject_all_changes()

        assert count == 2
        assert document_elm.xpath(".//w:moveFrom") == []
        assert document_elm.xpath(".//w:moveTo") == []
        assert document_elm.xpath(".//w:delText") == []
        paragraphs = document_elm.xpath(".//w:p")
        assert "".join(t.text for t in paragraphs[0].xpath(".//w:t")) == "before moved"
        assert "".join(t.text for t in paragraphs[1].xpath(".//w:t")) == "dest: "

    def it_can_accept_a_cellIns_revision_keeping_the_cell(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:tbl/("
                "w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/("
                "w:tc/(w:tcPr/w:cellIns{w:id=1,w:author=A},w:p),"
                "w:tc/w:p)"
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.accept_all_changes()

        assert count == 1
        assert document_elm.xpath(".//w:cellIns") == []
        # -- both cells survive --
        assert len(document_elm.xpath(".//w:tc")) == 2

    def it_can_reject_a_cellIns_revision_removing_the_cell(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:tbl/("
                "w:tblPr,w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/("
                "w:tc/(w:tcPr/w:cellIns{w:id=1,w:author=A},w:p),"
                "w:tc/w:p)"
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.reject_all_changes()

        assert count == 1
        assert document_elm.xpath(".//w:cellIns") == []
        # -- the inserted cell is removed, the sibling cell survives --
        assert len(document_elm.xpath(".//w:tc")) == 1

    def it_can_accept_a_cellDel_revision_removing_the_cell(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:tbl/("
                "w:tblPr,w:tblGrid/(w:gridCol,w:gridCol),"
                "w:tr/("
                "w:tc/(w:tcPr/w:cellDel{w:id=1,w:author=A},w:p),"
                "w:tc/w:p)"
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.accept_all_changes()

        assert count == 1
        assert document_elm.xpath(".//w:cellDel") == []
        assert len(document_elm.xpath(".//w:tc")) == 1

    def it_can_reject_a_cellDel_revision_keeping_the_cell(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:tbl/("
                "w:tblPr,w:tblGrid/w:gridCol,"
                "w:tr/w:tc/(w:tcPr/w:cellDel{w:id=1,w:author=A},w:p)"
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.reject_all_changes()

        assert count == 1
        assert document_elm.xpath(".//w:cellDel") == []
        assert len(document_elm.xpath(".//w:tc")) == 1

    def it_resolves_tcPrChange_trPrChange_and_tblPrChange(
        self, document_part_: Mock
    ):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:tbl/("
                "w:tblPr/w:tblPrChange{w:id=1,w:author=A}/w:tblPr,"
                "w:tblGrid/w:gridCol,"
                "w:tr/("
                "w:trPr/w:trPrChange{w:id=2,w:author=A}/w:trPr,"
                "w:tc/(w:tcPr/w:tcPrChange{w:id=3,w:author=A}/w:tcPr,w:p)"
                ")"
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        count = document.accept_all_changes()

        assert count == 3
        assert document_elm.xpath(".//w:tblPrChange") == []
        assert document_elm.xpath(".//w:trPrChange") == []
        assert document_elm.xpath(".//w:tcPrChange") == []

    def it_renders_revision_marks_text_joined_by_blank_lines(self, document_part_: Mock):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                'w:p/(w:r/w:t"first ",w:ins{w:id=1,w:author=A}/w:r/w:t"added"),'
                'w:p/(w:del{w:id=2,w:author=B}/w:r/w:delText"removed",w:r/w:t" end"),'
                'w:p/w:r/w:t"plain"'
                ")"
            ),
        )
        document = Document(document_elm, document_part_)

        rendered = document.revision_marks_text()

        assert rendered == "first [+added+]\n\n[-removed-] end\n\nplain"

    def it_supports_custom_markers_in_document_revision_marks(
        self, document_part_: Mock
    ):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/(w:ins{w:id=1,w:author=A}/w:r/w:t"x",'
                'w:del{w:id=2,w:author=B}/w:r/w:delText"y")'
            ),
        )
        document = Document(document_elm, document_part_)

        rendered = document.revision_marks_text(
            open_ins="<+", close_ins="+>", open_del="<-", close_del="->"
        )

        assert rendered == "<+x+><-y->"

    def it_returns_None_background_color_when_no_background_element(
        self, document_part_: Mock
    ):
        document = Document(cast(CT_Document, element("w:document/w:body")), document_part_)

        assert document.background_color is None

    def it_reads_background_color_from_the_background_element(
        self, document_part_: Mock
    ):
        document = Document(
            cast(
                CT_Document,
                element("w:document/(w:background{w:color=FF0000},w:body)"),
            ),
            document_part_,
        )

        assert document.background_color == RGBColor(0xFF, 0x00, 0x00)

    def it_writes_a_background_element_when_color_is_set(self, document_part_: Mock):
        document = Document(
            cast(CT_Document, element("w:document/w:body")), document_part_
        )

        document.background_color = RGBColor(0x12, 0x34, 0x56)

        assert document.element.xml == xml(
            "w:document/(w:background{w:color=123456},w:body)"
        )

    def it_replaces_an_existing_background_color(self, document_part_: Mock):
        document = Document(
            cast(
                CT_Document,
                element("w:document/(w:background{w:color=FF0000},w:body)"),
            ),
            document_part_,
        )

        document.background_color = RGBColor(0x00, 0xFF, 0x00)

        # -- only one w:background child, with the updated color --
        backgrounds = document.element.xpath("w:background")
        assert len(backgrounds) == 1
        assert document.background_color == RGBColor(0x00, 0xFF, 0x00)

    def it_removes_the_background_element_when_color_set_to_None(
        self, document_part_: Mock
    ):
        document = Document(
            cast(
                CT_Document,
                element("w:document/(w:background{w:color=FF0000},w:body)"),
            ),
            document_part_,
        )

        document.background_color = None

        assert document.element.xml == xml("w:document/w:body")
        assert document.background_color is None

    def it_is_a_noop_to_set_background_color_to_None_when_absent(
        self, document_part_: Mock
    ):
        document = Document(
            cast(CT_Document, element("w:document/w:body")), document_part_
        )

        document.background_color = None

        assert document.element.xml == xml("w:document/w:body")

    def it_determines_block_width_to_help(
        self, document: Document, sections_prop_: Mock, section_: Mock
    ):
        sections_prop_.return_value = [None, section_]
        section_.page_width = 6000
        section_.left_margin = 1500
        section_.right_margin = 1000

        width = document._block_width

        assert isinstance(width, Length)
        assert width == 3500

    # -- upstream#504: "Table Grid" style is shipped in default.docx --------

    def it_ships_the_table_grid_style_in_the_default_template(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        style = document.styles["Table Grid"]
        assert style.name == "Table Grid"
        # -- and it is usable as a table style on a fresh document --
        table = document.add_table(rows=1, cols=1, style="Table Grid")
        assert table.style.name == "Table Grid"

    # -- upstream#379: context-manager + close ------------------------------

    def it_supports_the_context_manager_protocol(self):
        from docx import Document as OpenDocument

        with OpenDocument() as document:
            assert isinstance(document, Document)
            paragraph = document.add_paragraph("hello")
            assert paragraph.text == "hello"

    def it_exposes_close_as_a_safe_no_op(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        # -- calling close() repeatedly is safe --
        document.close()
        document.close()
        # -- the document remains usable afterward (close is a lifecycle
        # -- affordance, not a teardown) --
        assert document.add_paragraph("still works").text == "still works"

    # -- upstream#1025: tracked-change writer -------------------------------

    def it_can_wrap_new_paragraphs_in_w_ins_via_context_manager(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        with document.tracked_changes(author="Alice"):
            paragraph = document.add_paragraph("review me")

        ins = paragraph._p.xpath("./w:ins")
        assert len(ins) == 1
        assert ins[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author"
        ) == "Alice"
        assert ins[0].xpath("./w:r/w:t")[0].text == "review me"

    def it_does_not_wrap_when_no_context_is_active(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        paragraph = document.add_paragraph("no tracking")

        assert paragraph._p.xpath("./w:ins") == []

    def it_allows_explicit_track_author_kwarg_without_context(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        paragraph = document.add_paragraph("via kwarg", track_author="Bob")

        ins = paragraph._p.xpath("./w:ins")
        assert len(ins) == 1
        assert ins[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author"
        ) == "Bob"

    def it_wraps_add_run_in_w_ins_under_active_context(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        paragraph = document.add_paragraph()  # empty, no ins yet
        with document.tracked_changes(author="Carol"):
            paragraph.add_run("inline addition")

        ins = paragraph._p.xpath("./w:ins")
        assert len(ins) == 1
        assert ins[0].xpath("./w:r/w:t")[0].text == "inline addition"
        assert ins[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author"
        ) == "Carol"

    def it_allocates_distinct_change_ids_across_multiple_wraps(self):
        from docx import Document as OpenDocument

        document = OpenDocument()
        with document.tracked_changes(author="Dan"):
            p1 = document.add_paragraph("first")
            p2 = document.add_paragraph("second")

        id_attr = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"
        ids = [
            p1._p.xpath("./w:ins")[0].get(id_attr),
            p2._p.xpath("./w:ins")[0].get(id_attr),
        ]
        assert len(set(ids)) == 2

    # -- Phase C: cross-document table copy (upstream#612, #270) -----------

    def it_can_add_a_cross_document_table_copy(self):
        """``add_table_copy`` deep-copies another document's ``w:tbl`` into this body."""
        import docx as _docx

        source = _docx.Document()
        src_table = source.add_table(rows=2, cols=2)
        src_table.cell(0, 0).text = "hello"
        src_table.cell(1, 1).text = "world"

        dest = _docx.Document()
        initial_table_count = len(dest.tables)

        copied = dest.add_table_copy(src_table)

        assert len(dest.tables) == initial_table_count + 1
        # -- copied is a distinct XML subtree --
        assert copied._tbl is not src_table._tbl
        # -- text preserved --
        assert copied.cell(0, 0).text == "hello"
        assert copied.cell(1, 1).text == "world"

    def it_rewires_embedded_images_across_documents(self):
        """``a:blip/@r:embed`` refs must be rewritten to point at the dest's rIds."""
        import docx as _docx
        from docx.opc.constants import RELATIONSHIP_TYPE as _RT
        from docx.oxml.ns import qn as _qn
        from docx.shared import Inches as _Inches

        source = _docx.Document()
        src_table = source.add_table(rows=1, cols=1)
        src_cell = src_table.cell(0, 0)
        src_cell.paragraphs[0].add_run().add_picture(
            "tests/test_files/monty-truth.png", width=_Inches(1)
        )

        # -- record the src-side rId --
        src_blips = src_table._tbl.xpath(".//a:blip")
        src_rids = [b.get(_qn("r:embed")) for b in src_blips]
        assert src_rids and all(src_rids)

        dest = _docx.Document()
        copied = dest.add_table_copy(src_table)

        dest_blips = copied._tbl.xpath(".//a:blip")
        dest_rids = [b.get(_qn("r:embed")) for b in dest_blips]
        assert dest_rids and all(dest_rids)
        # -- each dest rId must resolve to an image part in the dest's rels --
        for rid in dest_rids:
            dest_part = dest.part.related_parts[rid]
            assert dest_part.content_type.startswith("image/")
        # -- dest now has an image rel that source (may) not (and vice versa
        # -- we can assert the rId is *present* at least) --
        image_rels = [
            r for r in dest.part.rels.values() if r.reltype == _RT.IMAGE
        ]
        assert len(image_rels) >= 1

    def it_add_table_from_is_an_alias(self):
        import docx as _docx

        source = _docx.Document()
        src_table = source.add_table(rows=1, cols=1)
        src_table.cell(0, 0).text = "aliased"

        dest = _docx.Document()
        copied = dest.add_table_from(src_table)

        assert copied.cell(0, 0).text == "aliased"

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def add_paragraph_(self, request: FixtureRequest):
        return method_mock(request, Document, "add_paragraph")

    @pytest.fixture
    def _Body_(self, request: FixtureRequest):
        return class_mock(request, "docx.document._Body")

    @pytest.fixture
    def body_(self, request: FixtureRequest):
        return instance_mock(request, _Body)

    @pytest.fixture
    def _block_width_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "_block_width")

    @pytest.fixture
    def body_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "_body")

    @pytest.fixture
    def comment_(self, request: FixtureRequest):
        return instance_mock(request, Comment)

    @pytest.fixture
    def comments_(self, request: FixtureRequest):
        return instance_mock(request, Comments)

    @pytest.fixture
    def comments_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "comments")

    @pytest.fixture
    def core_properties_(self, request: FixtureRequest):
        return instance_mock(request, CoreProperties)

    @pytest.fixture
    def custom_properties_(self, request: FixtureRequest):
        return instance_mock(request, CustomProperties)

    @pytest.fixture
    def document(self, document_part_: Mock) -> Document:
        document_elm = cast(CT_Document, element("w:document"))
        return Document(document_elm, document_part_)

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def font_table_(self, request: FixtureRequest):
        return instance_mock(request, FontTable)

    @pytest.fixture
    def inline_shapes_(self, request: FixtureRequest):
        return instance_mock(request, InlineShapes)

    @pytest.fixture
    def paragraph_(self, request: FixtureRequest):
        return instance_mock(request, Paragraph)

    @pytest.fixture
    def paragraphs_(self, request: FixtureRequest):
        return instance_mock(request, list)

    @pytest.fixture
    def picture_(self, request: FixtureRequest):
        return instance_mock(request, InlineShape)

    @pytest.fixture
    def run_(self, request: FixtureRequest):
        return instance_mock(request, Run)

    @pytest.fixture
    def run_mark_comment_range_(self, request: FixtureRequest):
        return method_mock(request, Run, "mark_comment_range")

    @pytest.fixture
    def Section_(self, request: FixtureRequest):
        return class_mock(request, "docx.document.Section")

    @pytest.fixture
    def section_(self, request: FixtureRequest):
        return instance_mock(request, Section)

    @pytest.fixture
    def Sections_(self, request: FixtureRequest):
        return class_mock(request, "docx.document.Sections")

    @pytest.fixture
    def sections_(self, request: FixtureRequest):
        return instance_mock(request, Sections)

    @pytest.fixture
    def sections_prop_(self, request: FixtureRequest):
        return property_mock(request, Document, "sections")

    @pytest.fixture
    def settings_(self, request: FixtureRequest):
        return instance_mock(request, Settings)

    @pytest.fixture
    def styles_(self, request: FixtureRequest):
        return instance_mock(request, Styles)

    @pytest.fixture
    def table_(self, request: FixtureRequest):
        return instance_mock(request, Table)

    @pytest.fixture
    def tables_(self, request: FixtureRequest):
        return instance_mock(request, list)

    @pytest.fixture
    def web_settings_(self, request: FixtureRequest):
        return instance_mock(request, WebSettings)

    @pytest.fixture
    def theme_(self, request: FixtureRequest):
        return instance_mock(request, Theme)

    @pytest.fixture
    def glossary_(self, request: FixtureRequest):
        return instance_mock(request, Glossary)


class Describe_Body:
    """Unit-test suite for `docx.document._Body`."""

    @pytest.mark.parametrize(
        ("cxml", "expected_cxml"),
        [
            ("w:body", "w:body"),
            ("w:body/w:p", "w:body"),
            ("w:body/w:sectPr", "w:body/w:sectPr"),
            ("w:body/(w:p, w:sectPr)", "w:body/w:sectPr"),
        ],
    )
    def it_can_clear_itself_of_all_content_it_holds(
        self, cxml: str, expected_cxml: str, document_: Mock
    ):
        body = _Body(cast(CT_Body, element(cxml)), document_)

        _body = body.clear_content()

        assert body._body.xml == xml(expected_cxml)
        assert _body is body

    def it_can_add_a_block_level_content_control(self, document_: Mock):
        from docx.content_controls import ContentControl, ContentControlType

        body = _Body(cast(CT_Body, element("w:body")), document_)

        cc = body.add_content_control(
            ContentControlType.RICH_TEXT, tag="T", title="Title"
        )

        assert isinstance(cc, ContentControl)
        assert cc.tag == "T"
        assert cc.title == "Title"
        # -- the sdt was appended to the body --
        assert len(body._body.xpath("./w:sdt")) == 1

    def it_inserts_block_level_sdt_before_trailing_sectPr(self, document_: Mock):
        from docx.content_controls import ContentControlType

        body = _Body(cast(CT_Body, element("w:body/(w:p,w:sectPr)")), document_)

        body.add_content_control(ContentControlType.RICH_TEXT, tag="T")

        # -- verify order: w:p, w:sdt, w:sectPr --
        children = [c.tag.rsplit("}", 1)[-1] for c in body._body]
        assert children == ["p", "sdt", "sectPr"]

    def it_lists_block_level_content_controls(self, document_: Mock):
        from docx.content_controls import ContentControlType

        body = _Body(cast(CT_Body, element("w:body")), document_)
        body.add_content_control(ContentControlType.RICH_TEXT, tag="A")
        body.add_content_control(ContentControlType.RICH_TEXT, tag="B")

        assert [cc.tag for cc in body.content_controls] == ["A", "B"]

    def it_returns_type_specific_proxies_for_extended_types(self, document_: Mock):
        from docx.content_controls import (
            BuildingBlockControl,
            ContentControlType,
            DateControl,
            DropDownListControl,
            RepeatingSectionControl,
            RichTextControl,
        )

        body = _Body(cast(CT_Body, element("w:body")), document_)
        body.add_content_control(ContentControlType.RICH_TEXT, tag="R")
        body.add_content_control(ContentControlType.DATE, tag="D")
        body.add_content_control(ContentControlType.DROPDOWN, tag="DD")
        body.add_content_control(ContentControlType.REPEATING_SECTION, tag="RS")
        body.add_content_control(ContentControlType.BUILDING_BLOCK, tag="BB")

        ccs = body.content_controls
        assert isinstance(ccs[0], RichTextControl)
        assert isinstance(ccs[1], DateControl)
        assert isinstance(ccs[2], DropDownListControl)
        assert isinstance(ccs[3], RepeatingSectionControl)
        assert isinstance(ccs[4], BuildingBlockControl)

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, Document)


class DescribeDocument_Text:
    """Closes upstream#252 / upstream#72 — `Document.text` body concatenation."""

    def it_joins_paragraphs_with_newlines(self, request: FixtureRequest):
        doc_elm = cast(
            CT_Document,
            element(
                'w:document/w:body/('
                'w:p/w:r/w:t"one",'
                'w:p/w:r/w:t"two",'
                'w:p/w:r/w:t"three"'
                ')'
            ),
        )
        part_ = instance_mock(request, DocumentPart)
        doc = Document(doc_elm, part_)

        assert doc.text == "one\ntwo\nthree"

    def it_returns_empty_string_for_empty_body(self, request: FixtureRequest):
        doc_elm = cast(CT_Document, element("w:document/w:body"))
        part_ = instance_mock(request, DocumentPart)
        doc = Document(doc_elm, part_)
        assert doc.text == ""


class DescribeDocument_IterInnerContentSdtFlat:
    """Closes upstream#1280 — flatten ``w:sdt`` blocks when requested."""

    def it_flattens_sdt_paragraphs_when_include_sdt_flat(
        self, request: FixtureRequest
    ):
        doc_elm = cast(
            CT_Document,
            element(
                'w:document/w:body/('
                'w:p/w:r/w:t"outer1",'
                "w:sdt/w:sdtContent/w:p/w:r/w:t\"inner\","
                'w:p/w:r/w:t"outer2"'
                ')'
            ),
        )
        part_ = instance_mock(request, DocumentPart)
        doc = Document(doc_elm, part_)

        default_texts = [
            item.text
            for item in doc.iter_inner_content()
            if isinstance(item, Paragraph)
        ]
        # -- default iteration skips the sdt (non-w:p/w:tbl sibling) --
        assert default_texts == ["outer1", "outer2"]

        flat_texts = [
            item.text
            for item in doc.iter_inner_content(include_sdt_flat=True)
            if isinstance(item, Paragraph)
        ]
        assert flat_texts == ["outer1", "inner", "outer2"]
