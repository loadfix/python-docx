# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.footnotes` module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.enum.text import WD_FOOTNOTE_POSITION, WD_FOOTNOTE_RESTART, WD_NUMBER_FORMAT
from docx.footnotes import Footnote, FootnoteProperties, Footnotes
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.oxml.footnotes import CT_Footnote, CT_Footnotes, CT_FtnDocProps
from docx.oxml.ns import qn
from docx.oxml.text.run import CT_R
from docx.package import Package
from docx.parts.footnotes import FootnotesPart
from docx.text.run import Run

from .unitutil.cxml import element, xml
from .unitutil.mock import FixtureRequest, Mock, instance_mock


class DescribeFootnotes:
    """Unit-test suite for `docx.footnotes.Footnotes` objects."""

    @pytest.mark.parametrize(
        ("cxml", "count"),
        [
            # -- empty footnotes (only separators) --
            (
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator})",
                0,
            ),
            # -- one user footnote --
            (
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator}"
                ",w:footnote{w:id=2})",
                1,
            ),
            # -- two user footnotes --
            (
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator}"
                ",w:footnote{w:id=2},w:footnote{w:id=3})",
                2,
            ),
        ],
    )
    def it_knows_how_many_footnotes_it_contains(self, cxml: str, count: int, package_: Mock):
        footnotes_elm = cast(CT_Footnotes, element(cxml))
        footnotes_part = FootnotesPart(
            PackURI("/word/footnotes.xml"), CT.WML_FOOTNOTES, footnotes_elm, package_
        )
        footnotes = Footnotes(footnotes_elm, footnotes_part)

        assert len(footnotes) == count

    def it_is_iterable_over_user_footnotes(self, package_: Mock):
        footnotes_elm = cast(
            CT_Footnotes,
            element(
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator}"
                ",w:footnote{w:id=2},w:footnote{w:id=3})"
            ),
        )
        footnotes_part = FootnotesPart(
            PackURI("/word/footnotes.xml"), CT.WML_FOOTNOTES, footnotes_elm, package_
        )
        footnotes = Footnotes(footnotes_elm, footnotes_part)

        footnote_iter = iter(footnotes)

        fn1 = next(footnote_iter)
        assert type(fn1) is Footnote
        assert fn1.footnote_id == 2
        fn2 = next(footnote_iter)
        assert type(fn2) is Footnote
        assert fn2.footnote_id == 3
        with pytest.raises(StopIteration):
            next(footnote_iter)

    def it_can_add_a_footnote(self, package_: Mock):
        footnotes_elm = cast(
            CT_Footnotes,
            element(
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator})"
            ),
        )
        footnotes_part = FootnotesPart(
            PackURI("/word/footnotes.xml"), CT.WML_FOOTNOTES, footnotes_elm, package_
        )
        footnotes = Footnotes(footnotes_elm, footnotes_part)

        # -- create a run to anchor the footnote reference --
        para_elm = element("w:p/w:r")
        r_elm = cast(CT_R, para_elm[0])
        run = Run(r_elm, footnotes_part)

        footnote = footnotes.add(run)

        # -- a Footnote is returned --
        assert isinstance(footnote, Footnote)
        assert footnote.footnote_id == 2
        # -- the footnote part is linked --
        assert footnote.part is footnotes_part
        # -- the footnote has a single paragraph with FootnoteText style --
        assert len(footnote.paragraphs) == 1
        assert footnote.paragraphs[0]._p.style == "FootnoteText"
        # -- a footnoteReference was inserted into the run --
        ref_elms = r_elm.xpath("./w:footnoteReference")
        assert len(ref_elms) == 1
        assert ref_elms[0].get(qn("w:id")) == "2"
        # -- the run has FootnoteReference character style --
        assert r_elm.style == "FootnoteReference"

    def it_can_add_a_footnote_with_text(self, package_: Mock):
        footnotes_elm = cast(
            CT_Footnotes,
            element(
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator})"
            ),
        )
        footnotes_part = FootnotesPart(
            PackURI("/word/footnotes.xml"), CT.WML_FOOTNOTES, footnotes_elm, package_
        )
        footnotes = Footnotes(footnotes_elm, footnotes_part)

        para_elm = element("w:p/w:r")
        r_elm = cast(CT_R, para_elm[0])
        run = Run(r_elm, footnotes_part)

        footnote = footnotes.add(run, text="This is a footnote.")

        # -- the first paragraph has the text after the footnote ref run --
        first_para = footnote.paragraphs[0]
        assert len(first_para._p.r_lst) == 2
        assert first_para._p.r_lst[1].text == "This is a footnote."

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def package_(self, request: FixtureRequest):
        return instance_mock(request, Package)


class DescribeFootnote:
    """Unit-test suite for `docx.footnotes.Footnote`."""

    def it_knows_its_footnote_id(self, footnotes_part_: Mock):
        footnote_elm = cast(CT_Footnote, element("w:footnote{w:id=42}"))
        footnote = Footnote(footnote_elm, footnotes_part_)

        assert footnote.footnote_id == 42

    def it_provides_access_to_the_paragraphs_it_contains(self, footnotes_part_: Mock):
        footnote_elm = cast(
            CT_Footnote,
            element('w:footnote{w:id=2}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")'),
        )
        footnote = Footnote(footnote_elm, footnotes_part_)

        paragraphs = footnote.paragraphs

        assert len(paragraphs) == 2
        assert [para.text for para in paragraphs] == ["First para", "Second para"]

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:footnote{w:id=2}", ""),
            ('w:footnote{w:id=2}/w:p/w:r/w:t"Footnote text."', "Footnote text."),
            (
                'w:footnote{w:id=2}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")',
                "First para\nSecond para",
            ),
            (
                'w:footnote{w:id=2}/(w:p/w:r/w:t"First para",w:p,w:p/w:r/w:t"Second para")',
                "First para\n\nSecond para",
            ),
        ],
    )
    def it_can_summarize_its_content_as_text(
        self, cxml: str, expected_value: str, footnotes_part_: Mock
    ):
        assert Footnote(cast(CT_Footnote, element(cxml)), footnotes_part_).text == expected_value

    def it_can_clear_its_content(self, footnotes_part_: Mock):
        footnote_elm = cast(
            CT_Footnote,
            element('w:footnote{w:id=2}/(w:p/w:r/w:t"First",w:p/w:r/w:t"Second")'),
        )
        footnote = Footnote(footnote_elm, footnotes_part_)
        assert len(footnote.paragraphs) == 2

        result = footnote.clear()

        assert result is footnote
        assert len(footnote.paragraphs) == 1
        p = footnote.paragraphs[0]
        assert p.text == ""
        assert p._p.style == "FootnoteText"
        # -- the paragraph retains the footnoteRef run for the auto-number mark --
        assert len(p._p.r_lst) == 1
        assert p._p.r_lst[0].style == "FootnoteReference"
        assert p._p.r_lst[0][-1].tag == qn("w:footnoteRef")

    def it_can_delete_itself(self):
        # -- build a footnotes element with a user footnote (id=2) --
        footnotes_elm = cast(
            CT_Footnotes,
            element(
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator}"
                ',w:footnote{w:id=2}/w:p/w:r/w:t"Footnote text")'
            ),
        )

        # -- build a document element containing the footnoteReference --
        doc_elm = element("w:document/w:body/w:p/w:r/w:footnoteReference{w:id=2}")
        document_part_ = Mock()
        document_part_.element = doc_elm
        footnotes_part_ = Mock()
        footnotes_part_.part = footnotes_part_
        footnotes_part_._document_part = document_part_

        footnote_elm = footnotes_elm.footnote_lst[2]
        footnote = Footnote(footnote_elm, footnotes_part_)

        footnote.delete()

        # -- the footnote element is removed from the footnotes part --
        assert len(footnotes_elm.footnote_lst) == 2
        assert all(fn.type is not None for fn in footnotes_elm.footnote_lst)
        # -- the footnoteReference run is removed from the document body --
        refs = doc_elm.xpath(".//w:footnoteReference")
        assert len(refs) == 0

    def it_removes_the_ref_run_when_deleting_if_run_becomes_empty(self):
        footnotes_elm = cast(
            CT_Footnotes,
            element(
                "w:footnotes/(w:footnote{w:id=0,w:type=separator}"
                ",w:footnote{w:id=1,w:type=continuationSeparator}"
                ",w:footnote{w:id=2}/w:p)"
            ),
        )

        # -- the run has rPr + footnoteReference; after removing ref, only rPr remains --
        doc_elm = element(
            "w:document/w:body/w:p/w:r/(w:rPr/w:rStyle{w:val=FootnoteReference}"
            ",w:footnoteReference{w:id=2})"
        )
        document_part_ = Mock()
        document_part_.element = doc_elm
        footnotes_part_ = Mock()
        footnotes_part_.part = footnotes_part_
        footnotes_part_._document_part = document_part_

        footnote_elm = footnotes_elm.footnote_lst[2]
        footnote = Footnote(footnote_elm, footnotes_part_)

        footnote.delete()

        # -- the entire run is removed since it only had rPr left --
        runs = doc_elm.xpath(".//w:r")
        assert len(runs) == 0

    def it_can_add_a_paragraph(self, footnotes_part_: Mock):
        footnote_elm = cast(CT_Footnote, element("w:footnote{w:id=2}/w:p"))
        footnote = Footnote(footnote_elm, footnotes_part_)

        paragraph = footnote.add_paragraph("New paragraph text")

        assert len(footnote.paragraphs) == 2
        assert footnote.paragraphs[1].text == "New paragraph text"
        # -- default style is FootnoteText --
        assert paragraph._p.style == "FootnoteText"

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def footnotes_part_(self, request: FixtureRequest):
        return instance_mock(request, FootnotesPart)


class DescribeFootnoteProperties:
    """Unit-test suite for `docx.footnotes.FootnoteProperties`."""

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:footnotePr", None),
            ("w:footnotePr/w:numFmt{w:val=decimal}", WD_NUMBER_FORMAT.ARABIC),
            ("w:footnotePr/w:numFmt{w:val=lowerRoman}", WD_NUMBER_FORMAT.LOWER_ROMAN),
            ("w:footnotePr/w:numFmt{w:val=upperRoman}", WD_NUMBER_FORMAT.UPPER_ROMAN),
            ("w:footnotePr/w:numFmt{w:val=lowerLetter}", WD_NUMBER_FORMAT.LOWER_LETTER),
            ("w:footnotePr/w:numFmt{w:val=upperLetter}", WD_NUMBER_FORMAT.UPPER_LETTER),
            ("w:footnotePr/w:numFmt{w:val=chicago}", WD_NUMBER_FORMAT.CHICAGO),
        ],
    )
    def it_can_get_the_number_format(
        self, cxml: str, expected_value: WD_NUMBER_FORMAT | None
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        assert FootnoteProperties(footnotePr).number_format == expected_value

    @pytest.mark.parametrize(
        ("cxml", "new_value", "expected_cxml"),
        [
            ("w:footnotePr", WD_NUMBER_FORMAT.ARABIC, "w:footnotePr/w:numFmt{w:val=decimal}"),
            (
                "w:footnotePr/w:numFmt{w:val=decimal}",
                WD_NUMBER_FORMAT.LOWER_ROMAN,
                "w:footnotePr/w:numFmt{w:val=lowerRoman}",
            ),
            ("w:footnotePr/w:numFmt{w:val=decimal}", None, "w:footnotePr"),
        ],
    )
    def it_can_set_the_number_format(
        self, cxml: str, new_value: WD_NUMBER_FORMAT | None, expected_cxml: str
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        props = FootnoteProperties(footnotePr)
        props.number_format = new_value
        assert footnotePr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:footnotePr", None),
            ("w:footnotePr/w:numStart{w:val=1}", 1),
            ("w:footnotePr/w:numStart{w:val=5}", 5),
        ],
    )
    def it_can_get_the_start_number(self, cxml: str, expected_value: int | None):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        assert FootnoteProperties(footnotePr).start_number == expected_value

    @pytest.mark.parametrize(
        ("cxml", "new_value", "expected_cxml"),
        [
            ("w:footnotePr", 1, "w:footnotePr/w:numStart{w:val=1}"),
            ("w:footnotePr/w:numStart{w:val=1}", 7, "w:footnotePr/w:numStart{w:val=7}"),
            ("w:footnotePr/w:numStart{w:val=1}", None, "w:footnotePr"),
        ],
    )
    def it_can_set_the_start_number(
        self, cxml: str, new_value: int | None, expected_cxml: str
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        props = FootnoteProperties(footnotePr)
        props.start_number = new_value
        assert footnotePr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:footnotePr", None),
            ("w:footnotePr/w:numRestart{w:val=continuous}", WD_FOOTNOTE_RESTART.CONTINUOUS),
            ("w:footnotePr/w:numRestart{w:val=eachSect}", WD_FOOTNOTE_RESTART.EACH_SECTION),
            ("w:footnotePr/w:numRestart{w:val=eachPage}", WD_FOOTNOTE_RESTART.EACH_PAGE),
        ],
    )
    def it_can_get_the_restart_rule(
        self, cxml: str, expected_value: WD_FOOTNOTE_RESTART | None
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        assert FootnoteProperties(footnotePr).restart_rule == expected_value

    @pytest.mark.parametrize(
        ("cxml", "new_value", "expected_cxml"),
        [
            (
                "w:footnotePr",
                WD_FOOTNOTE_RESTART.EACH_PAGE,
                "w:footnotePr/w:numRestart{w:val=eachPage}",
            ),
            (
                "w:footnotePr/w:numRestart{w:val=continuous}",
                WD_FOOTNOTE_RESTART.EACH_SECTION,
                "w:footnotePr/w:numRestart{w:val=eachSect}",
            ),
            ("w:footnotePr/w:numRestart{w:val=eachPage}", None, "w:footnotePr"),
        ],
    )
    def it_can_set_the_restart_rule(
        self, cxml: str, new_value: WD_FOOTNOTE_RESTART | None, expected_cxml: str
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        props = FootnoteProperties(footnotePr)
        props.restart_rule = new_value
        assert footnotePr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:footnotePr", None),
            ("w:footnotePr/w:pos{w:val=pageBottom}", WD_FOOTNOTE_POSITION.BOTTOM_OF_PAGE),
            ("w:footnotePr/w:pos{w:val=beneathText}", WD_FOOTNOTE_POSITION.BENEATH_TEXT),
        ],
    )
    def it_can_get_the_position(
        self, cxml: str, expected_value: WD_FOOTNOTE_POSITION | None
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        assert FootnoteProperties(footnotePr).position == expected_value

    @pytest.mark.parametrize(
        ("cxml", "new_value", "expected_cxml"),
        [
            (
                "w:footnotePr",
                WD_FOOTNOTE_POSITION.BOTTOM_OF_PAGE,
                "w:footnotePr/w:pos{w:val=pageBottom}",
            ),
            (
                "w:footnotePr/w:pos{w:val=pageBottom}",
                WD_FOOTNOTE_POSITION.BENEATH_TEXT,
                "w:footnotePr/w:pos{w:val=beneathText}",
            ),
            ("w:footnotePr/w:pos{w:val=pageBottom}", None, "w:footnotePr"),
        ],
    )
    def it_can_set_the_position(
        self, cxml: str, new_value: WD_FOOTNOTE_POSITION | None, expected_cxml: str
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        props = FootnoteProperties(footnotePr)
        props.position = new_value
        assert footnotePr.xml == xml(expected_cxml)

    def it_orders_children_in_schema_sequence_when_all_set(self):
        footnotePr = cast(CT_FtnDocProps, element("w:footnotePr"))
        props = FootnoteProperties(footnotePr)

        # -- assign in arbitrary order --
        props.restart_rule = WD_FOOTNOTE_RESTART.EACH_PAGE
        props.start_number = 3
        props.position = WD_FOOTNOTE_POSITION.BOTTOM_OF_PAGE
        props.number_format = WD_NUMBER_FORMAT.LOWER_ROMAN

        expected = (
            "w:footnotePr/(w:pos{w:val=pageBottom}"
            ",w:numFmt{w:val=lowerRoman}"
            ",w:numStart{w:val=3}"
            ",w:numRestart{w:val=eachPage})"
        )
        assert footnotePr.xml == xml(expected)

    # -- R5-3: section-end and document-end positions --------------------------------

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:footnotePr/w:pos{w:val=sectEnd}", WD_FOOTNOTE_POSITION.END_OF_SECTION),
            ("w:footnotePr/w:pos{w:val=docEnd}", WD_FOOTNOTE_POSITION.END_OF_DOCUMENT),
        ],
    )
    def it_can_read_section_end_and_doc_end_positions(
        self, cxml: str, expected_value: WD_FOOTNOTE_POSITION
    ):
        footnotePr = cast(CT_FtnDocProps, element(cxml))
        assert FootnoteProperties(footnotePr).position == expected_value

    @pytest.mark.parametrize(
        ("new_value", "expected_cxml"),
        [
            (
                WD_FOOTNOTE_POSITION.END_OF_SECTION,
                "w:footnotePr/w:pos{w:val=sectEnd}",
            ),
            (
                WD_FOOTNOTE_POSITION.END_OF_DOCUMENT,
                "w:footnotePr/w:pos{w:val=docEnd}",
            ),
        ],
    )
    def it_can_write_section_end_and_doc_end_positions(
        self, new_value: WD_FOOTNOTE_POSITION, expected_cxml: str
    ):
        footnotePr = cast(CT_FtnDocProps, element("w:footnotePr"))
        FootnoteProperties(footnotePr).position = new_value
        assert footnotePr.xml == xml(expected_cxml)

    # -- R5-3: `.numbering_restart` alias for `.restart_rule` ------------------------

    def it_exposes_numbering_restart_as_alias_of_restart_rule(self):
        footnotePr = cast(
            CT_FtnDocProps, element("w:footnotePr/w:numRestart{w:val=eachPage}")
        )
        props = FootnoteProperties(footnotePr)
        assert props.numbering_restart == WD_FOOTNOTE_RESTART.EACH_PAGE
        # -- write through alias, read through canonical property --
        props.numbering_restart = WD_FOOTNOTE_RESTART.EACH_SECTION
        assert props.restart_rule == WD_FOOTNOTE_RESTART.EACH_SECTION
        # -- clearing through alias removes the child --
        props.numbering_restart = None
        assert props.restart_rule is None

    # -- R5-3: separator / continuation-separator / continuation-notice refs ---------

    def it_exposes_None_when_no_separator_refs_present(self):
        footnotePr = cast(CT_FtnDocProps, element("w:footnotePr"))
        props = FootnoteProperties(footnotePr)
        assert props.separator_id is None
        assert props.continuation_separator_id is None
        assert props.continuation_notice_id is None

    def it_reads_all_three_separator_reference_kinds(self):
        footnotePr = cast(
            CT_FtnDocProps,
            element(
                "w:footnotePr/("
                'w:footnote{w:id=0,w:type=separator},'
                'w:footnote{w:id=1,w:type=continuationSeparator},'
                'w:footnote{w:id=2,w:type=continuationNotice}'
                ")"
            ),
        )
        props = FootnoteProperties(footnotePr)
        assert props.separator_id == 0
        assert props.continuation_separator_id == 1
        assert props.continuation_notice_id == 2

    def it_upserts_separator_ref_without_duplicating(self):
        footnotePr = cast(CT_FtnDocProps, element("w:footnotePr"))
        props = FootnoteProperties(footnotePr)
        props.separator_id = 0
        props.separator_id = 5  # -- overwrite, should not duplicate --
        seps = [
            fn for fn in footnotePr.footnote_lst if fn.type == "separator"
        ]
        assert len(seps) == 1
        assert seps[0].id == 5

    def it_clears_separator_ref_when_set_to_None(self):
        footnotePr = cast(
            CT_FtnDocProps,
            element('w:footnotePr/w:footnote{w:id=0,w:type=separator}'),
        )
        props = FootnoteProperties(footnotePr)
        assert props.separator_id == 0
        props.separator_id = None
        assert props.separator_id is None
        assert footnotePr.footnote_lst == []

    def it_writes_all_three_separator_refs_round_trip(self):
        footnotePr = cast(CT_FtnDocProps, element("w:footnotePr"))
        props = FootnoteProperties(footnotePr)
        props.separator_id = 0
        props.continuation_separator_id = 1
        props.continuation_notice_id = 2
        # -- each kind round-trips via its own property --
        assert props.separator_id == 0
        assert props.continuation_separator_id == 1
        assert props.continuation_notice_id == 2
        assert len(footnotePr.footnote_lst) == 3
