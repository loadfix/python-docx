# pyright: reportPrivateUsage=false
# pyright: reportUnknownMemberType=false

"""Unit-test suite for the `docx.outline` module (issue #71)."""

from __future__ import annotations

import json
from typing import cast

import pytest

from docx import Document as _DocumentFactory
from docx.outline import (
    Outline,
    OutlineNode,
    _heading_level,
    _stable_id,
    _word_count,
    build_outline,
    slice_document,
)

from .unitutil.mock import Mock


def _fake_paragraph(style_name: "str | None", text: str = "lorem"):
    paragraph = Mock(name="Paragraph")
    paragraph.text = text
    if style_name is None:
        paragraph.style = None
    else:
        style = Mock(name="ParagraphStyle")
        style.name = style_name
        paragraph.style = style
    return paragraph


class Describe_heading_level:
    """Unit-test suite for `docx.outline._heading_level`."""

    @pytest.mark.parametrize(
        ("name", "expected"),
        [
            ("Heading 1", 1),
            ("heading 1", 1),
            ("HEADING 9", 9),
            ("Heading\t3", 3),
            ("Title", 0),
            ("title", 0),
            ("Normal", None),
            ("Heading 0", None),  # -- only 1..9 + Title --
            ("Heading 10", None),
            ("", None),
        ],
    )
    def it_returns_the_outline_level_for_recognized_styles(
        self, name: str, expected: "int | None"
    ):
        assert _heading_level(_fake_paragraph(name)) == expected

    def it_returns_None_when_paragraph_has_no_style(self):
        assert _heading_level(_fake_paragraph(None)) is None

    def it_returns_None_when_style_name_is_None(self):
        paragraph = _fake_paragraph("Heading 1")
        paragraph.style.name = None
        assert _heading_level(paragraph) is None


class Describe_word_count:
    """Unit-test suite for `docx.outline._word_count`."""

    @pytest.mark.parametrize(
        ("text", "expected"),
        [
            ("", 0),
            ("hello", 1),
            ("hello world", 2),
            ("  hello   world  ", 2),
            ("a\tb\nc", 3),
        ],
    )
    def it_counts_whitespace_delimited_tokens(self, text: str, expected: int):
        assert _word_count(text) == expected


class Describe_stable_id:
    """Unit-test suite for `docx.outline._stable_id`."""

    def it_returns_an_8_char_hex_string(self):
        sid = _stable_id(1, "Introduction", 0)
        assert isinstance(sid, str)
        assert len(sid) == 8
        int(sid, 16)  # -- must be valid hex --

    def it_is_deterministic_for_the_same_inputs(self):
        assert _stable_id(1, "X", 5) == _stable_id(1, "X", 5)

    def it_differs_when_any_input_differs(self):
        a = _stable_id(1, "X", 5)
        assert _stable_id(2, "X", 5) != a
        assert _stable_id(1, "Y", 5) != a
        assert _stable_id(1, "X", 6) != a


class DescribeOutlineNode:
    """Unit-test suite for `docx.outline.OutlineNode`."""

    def it_constructs_with_default_word_count_id_and_children(self):
        node = OutlineNode(level=1, text="Intro", paragraph_index=0)
        assert node.level == 1
        assert node.text == "Intro"
        assert node.paragraph_index == 0
        assert node.id == ""
        assert node.word_count == 0
        assert node.children == []

    def it_walks_itself_then_descendants_in_dfs_order(self):
        leaf = OutlineNode(level=2, text="Sub", paragraph_index=2)
        root = OutlineNode(level=1, text="Top", paragraph_index=0, children=[leaf])

        assert list(root.walk()) == [root, leaf]

    def it_round_trips_through_to_dict(self):
        leaf = OutlineNode(
            level=2, text="Sub", paragraph_index=2, id="abcd1234", word_count=5
        )
        root = OutlineNode(
            level=1,
            text="Top",
            paragraph_index=0,
            id="11112222",
            word_count=10,
            children=[leaf],
        )
        d = root.to_dict()
        assert d == {
            "id": "11112222",
            "heading": "Top",
            "level": 1,
            "paragraph_index": 0,
            "word_count": 10,
            "children": [
                {
                    "id": "abcd1234",
                    "heading": "Sub",
                    "level": 2,
                    "paragraph_index": 2,
                    "word_count": 5,
                    "children": [],
                }
            ],
        }
        # -- and the result must be JSON-serialisable --
        json.dumps(d)


class DescribeOutline:
    """Unit-test suite for `docx.outline.Outline`."""

    def it_walks_all_nodes_dfs(self):
        leaf = OutlineNode(level=2, text="Sub", paragraph_index=2)
        root = OutlineNode(level=1, text="Top", paragraph_index=0, children=[leaf])
        outline = Outline(sections=[root], title="T", total_paragraphs=3)

        assert list(outline.walk()) == [root, leaf]
        assert list(iter(outline)) == [root, leaf]
        assert len(outline) == 2

    def it_exposes_a_to_dict_schema(self):
        outline = Outline(
            sections=[OutlineNode(level=1, text="Intro", paragraph_index=0)],
            title="My Doc",
            total_paragraphs=1,
            total_pages_estimated=2,
        )
        d = outline.to_dict()
        assert d["title"] == "My Doc"
        assert d["total_paragraphs"] == 1
        assert d["total_pages_estimated"] == 2
        assert isinstance(d["sections"], list)
        # -- JSON serialisable --
        json.dumps(d)

    def it_finds_a_node_by_exact_heading_text(self):
        a = OutlineNode(level=1, text="Methodology", paragraph_index=2)
        b = OutlineNode(level=2, text="Approach", paragraph_index=4)
        outline = Outline(sections=[a, b], total_paragraphs=5)

        assert outline.find("Methodology") is a
        assert outline.find("  Methodology  ") is a  # -- strips both sides --
        assert outline.find("Approach") is b
        assert outline.find("Nonexistent") is None


class DescribeBuildOutline:
    """End-to-end tests for `docx.outline.build_outline` against real Document."""

    def it_returns_an_empty_outline_for_a_doc_with_no_headings(self):
        doc = _DocumentFactory()
        doc.add_paragraph("Body text only.")
        outline = build_outline(doc)
        assert outline.sections == []
        # -- title defaults to core_properties.title which is empty string for a
        # -- fresh doc; coerced to None by build_outline --
        assert outline.title is None

    def it_collects_headings_in_document_order(self):
        doc = _DocumentFactory()
        doc.add_heading("Intro", 1)
        doc.add_heading("Methods", 1)
        doc.add_heading("Results", 1)
        outline = build_outline(doc)
        assert [s.text for s in outline.sections] == ["Intro", "Methods", "Results"]
        assert all(s.level == 1 for s in outline.sections)

    def it_nests_deeper_headings_under_their_parent(self):
        doc = _DocumentFactory()
        doc.add_heading("Intro", 1)
        doc.add_heading("Background", 2)
        doc.add_heading("Prior work", 3)
        doc.add_heading("Methods", 1)
        outline = build_outline(doc)
        assert len(outline.sections) == 2
        intro = outline.sections[0]
        assert intro.text == "Intro"
        assert len(intro.children) == 1
        assert intro.children[0].text == "Background"
        assert intro.children[0].children[0].text == "Prior work"
        methods = outline.sections[1]
        assert methods.text == "Methods"
        assert methods.children == []

    def it_treats_Title_as_level_zero(self):
        doc = _DocumentFactory()
        doc.add_heading("Cover", 0)
        doc.add_heading("Intro", 1)
        outline = build_outline(doc)
        assert outline.sections[0].level == 0
        assert outline.sections[0].children[0].level == 1
        assert outline.title == "Cover"

    def it_records_word_count_per_leaf_section(self):
        doc = _DocumentFactory()
        doc.add_heading("Intro", 1)  # -- 1 word --
        doc.add_paragraph("alpha beta gamma")  # -- 3 words --
        doc.add_heading("Methods", 1)  # -- 1 word --
        doc.add_paragraph("just two words")  # -- 3 words --
        outline = build_outline(doc)
        # -- "Intro" + 3 body words = 4 --
        assert outline.sections[0].word_count == 4
        # -- "Methods" + 3 body words = 4 --
        assert outline.sections[1].word_count == 4

    def it_records_paragraph_index_pointing_back_to_paragraphs(self):
        doc = _DocumentFactory()
        doc.add_paragraph("body before")
        doc.add_heading("Intro", 1)
        doc.add_paragraph("body after")
        outline = build_outline(doc)
        node = outline.sections[0]
        # -- Document.paragraphs[node.paragraph_index] is the heading itself --
        assert doc.paragraphs[node.paragraph_index].text == "Intro"

    def it_attaches_stable_ids_per_node(self):
        doc = _DocumentFactory()
        doc.add_heading("Intro", 1)
        doc.add_heading("Methods", 1)
        a = build_outline(doc)
        b = build_outline(doc)
        assert [n.id for n in a.walk()] == [n.id for n in b.walk()]

    def it_records_total_paragraphs(self):
        doc = _DocumentFactory()
        doc.add_heading("A", 1)
        doc.add_paragraph("body")
        doc.add_heading("B", 1)
        outline = build_outline(doc)
        assert outline.total_paragraphs == len(doc.paragraphs)

    def it_falls_back_to_core_properties_title(self, monkeypatch: pytest.MonkeyPatch):
        # -- some preceding tests in the suite stub `Document.core_properties`
        # -- via class-level mocks, so go through a synthetic doc whose
        # -- behaviour we control directly. --
        doc = _DocumentFactory()
        doc.add_heading("Intro", 1)

        class _CP:
            title = "Set Via Core"

        monkeypatch.setattr(
            type(doc), "core_properties", property(lambda self: _CP()), raising=False
        )
        outline = build_outline(doc)
        assert outline.title == "Set Via Core"


class DescribeDocumentOutlineMethod:
    """Integration tests for :meth:`Document.outline`."""

    def it_is_exposed_as_a_Document_method(self):
        doc = _DocumentFactory()
        doc.add_heading("Top", 1)
        outline = doc.outline()
        assert isinstance(outline, Outline)
        assert [n.text for n in outline.sections] == ["Top"]


class DescribeSliceDocument:
    """Unit / integration tests for `docx.outline.slice_document`."""

    def it_slices_from_a_heading_to_the_next_heading(self):
        doc = _DocumentFactory()
        doc.add_heading("Intro", 1)
        doc.add_paragraph("intro body")
        doc.add_heading("Methods", 1)
        doc.add_paragraph("methods body 1")
        doc.add_paragraph("methods body 2")
        doc.add_heading("Results", 1)
        doc.add_paragraph("results body")

        sub = slice_document(doc, start="Methods", end="Results")
        texts = [p.text for p in sub.paragraphs]
        assert texts == ["Methods", "methods body 1", "methods body 2"]

    def it_slices_to_end_when_end_is_None(self):
        doc = _DocumentFactory()
        doc.add_heading("Intro", 1)
        doc.add_paragraph("body")
        doc.add_heading("Final", 1)
        doc.add_paragraph("final body")

        sub = slice_document(doc, start="Final")
        texts = [p.text for p in sub.paragraphs]
        assert texts == ["Final", "final body"]

    def it_accepts_an_OutlineNode_for_start(self):
        doc = _DocumentFactory()
        doc.add_heading("A", 1)
        doc.add_paragraph("a-body")
        doc.add_heading("B", 1)

        outline = build_outline(doc)
        node = outline.find("A")
        assert node is not None
        sub = slice_document(doc, start=node, end="B")
        assert [p.text for p in sub.paragraphs] == ["A", "a-body"]

    def it_raises_ValueError_for_unknown_start(self):
        doc = _DocumentFactory()
        doc.add_heading("Only", 1)
        with pytest.raises(ValueError, match="no heading matches"):
            slice_document(doc, start="Missing")

    def it_raises_ValueError_for_unknown_end(self):
        doc = _DocumentFactory()
        doc.add_heading("Only", 1)
        with pytest.raises(ValueError, match="no heading matches"):
            slice_document(doc, start="Only", end="Missing")

    def it_is_exposed_as_Document_dot_slice(self):
        doc = _DocumentFactory()
        doc.add_heading("A", 1)
        doc.add_paragraph("body")
        doc.add_heading("B", 1)
        sub = doc.slice(start="A", end="B")
        # -- only the "A" section is copied --
        assert [p.text for p in sub.paragraphs] == ["A", "body"]
