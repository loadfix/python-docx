"""Unit-test suite for ``docx.kit.resume`` (issue #63)."""

from __future__ import annotations

from io import BytesIO

import pytest

from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import resume
from docx.kit.resume import (
    STYLES,
    TEMPLATES,
    resume_chronological,
    resume_functional,
    resume_technical,
)


# -- Reusable fixtures ----------------------------------------------------


@pytest.fixture
def base_kwargs() -> dict:
    """Minimal kwargs that satisfy every template factory."""
    return {
        "name": "Ben Hooper",
        "title": "Senior Software Engineer",
        "contact": {
            "email": "ben@example.com",
            "phone": "+61 2 1234 5678",
            "linkedin": "in/benhooper",
            "github": "benh",
        },
        "summary": "15+ years of distributed systems experience.",
    }


@pytest.fixture
def experience() -> list:
    return [
        {
            "company": "Acme Corp",
            "title": "Staff Engineer",
            "start": "2020-03",
            "end": "present",
            "bullets": ["Led the X migration", "Shipped Y to production"],
        },
        {
            "company": "Beta Inc",
            "title": "Senior Engineer",
            "start": "2017-01",
            "end": "2020-02",
            "bullets": ["Refactored Z"],
        },
    ]


@pytest.fixture
def education() -> list:
    return [{"school": "UNSW", "degree": "BE (Hons) Software", "year": 2010}]


def _hyperlink_urls(doc: DocumentCls) -> list:
    """Return every external hyperlink address discovered in ``doc``."""
    urls: list = []
    for para in doc.paragraphs:
        for link in getattr(para, "hyperlinks", []) or []:
            addr = getattr(link, "address", None) or getattr(link, "url", None) or ""
            urls.append(addr)
    return urls


# -- Module-level surface assertions --------------------------------------


class DescribeResumeModule:
    """Module-level invariants — public surface is what the issue called for."""

    def it_exposes_three_template_factories(self):
        assert resume.resume_chronological is resume_chronological
        assert resume.resume_functional is resume_functional
        assert resume.resume_technical is resume_technical

    def it_exposes_the_three_built_in_styles(self):
        assert STYLES == ("modern", "classic", "minimal")

    def it_exposes_the_three_template_names(self):
        assert TEMPLATES == ("chronological", "functional", "technical")


# -- resume_chronological -------------------------------------------------


class DescribeResumeChronological:
    """Behavioural tests for :func:`resume.resume_chronological`."""

    def it_returns_a_Document(self, base_kwargs: dict, experience: list, education: list):
        doc = resume_chronological(
            experience=experience,
            education=education,
            skills=["Python", "Go", "AWS"],
            **base_kwargs,
        )
        assert isinstance(doc, DocumentCls)

    def it_emits_the_name_then_title_at_the_top(self, base_kwargs: dict):
        doc = resume_chronological(**base_kwargs)
        texts = [p.text for p in doc.paragraphs]
        name_idx = next(i for i, t in enumerate(texts) if "Ben Hooper" in t)
        title_idx = next(
            i for i, t in enumerate(texts) if "Senior Software Engineer" in t
        )
        assert name_idx < title_idx

    def it_emits_the_summary_block(self, base_kwargs: dict):
        doc = resume_chronological(**base_kwargs)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "SUMMARY" in all_text.upper()
        assert "15+ years" in all_text

    def it_emits_each_experience_entry_with_company_and_bullets(
        self, base_kwargs: dict, experience: list
    ):
        doc = resume_chronological(experience=experience, **base_kwargs)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in all_text
        assert "Beta Inc" in all_text
        assert "Staff Engineer" in all_text
        assert "Led the X migration" in all_text
        assert "Shipped Y to production" in all_text
        assert "Refactored Z" in all_text

    def it_emits_education_after_experience(
        self, base_kwargs: dict, experience: list, education: list
    ):
        doc = resume_chronological(
            experience=experience, education=education, **base_kwargs
        )
        texts = [p.text for p in doc.paragraphs]
        edu_idx = next(i for i, t in enumerate(texts) if "BE (Hons) Software" in t)
        exp_idx = next(i for i, t in enumerate(texts) if "Acme Corp" in t)
        assert exp_idx < edu_idx

    def it_renders_skills_as_a_comma_joined_paragraph_when_a_flat_list(
        self, base_kwargs: dict
    ):
        doc = resume_chronological(skills=["Python", "Go", "AWS"], **base_kwargs)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Python, Go, AWS" in all_text

    def it_renders_skills_as_categorised_lines_when_a_mapping(
        self, base_kwargs: dict
    ):
        doc = resume_chronological(
            skills={"Languages": ["Python", "Go"], "Cloud": ["AWS"]},
            **base_kwargs,
        )
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Languages:" in all_text
        assert "Cloud:" in all_text
        assert "Python, Go" in all_text

    def it_can_be_saved_to_a_BytesIO(self, base_kwargs: dict, experience: list):
        doc = resume_chronological(experience=experience, **base_kwargs)
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue().startswith(b"PK")  # zip magic

    def it_raises_when_name_is_empty(self):
        with pytest.raises(ValueError, match="name must be a non-empty string"):
            resume_chronological(name="")

    def it_raises_when_style_is_unknown(self, base_kwargs: dict):
        with pytest.raises(ValueError, match="style must be one of"):
            resume_chronological(style="bogus", **base_kwargs)

    @pytest.mark.parametrize("style", STYLES)
    def it_supports_every_built_in_style(
        self, style: str, base_kwargs: dict, experience: list
    ):
        doc = resume_chronological(
            experience=experience, style=style, **base_kwargs
        )
        texts = [p.text for p in doc.paragraphs]
        assert any("Ben Hooper" in t for t in texts)
        # Every style emits at least one Heading-styled paragraph.
        styles_used = {p.style.name for p in doc.paragraphs if p.style is not None}
        assert any(s.startswith("Heading") for s in styles_used)


# -- resume_functional ----------------------------------------------------


class DescribeResumeFunctional:
    """Behavioural tests for :func:`resume.resume_functional`."""

    def it_returns_a_Document(self, base_kwargs: dict):
        doc = resume_functional(
            focus_areas=["Engineering Leadership", "Distributed Systems"],
            **base_kwargs,
        )
        assert isinstance(doc, DocumentCls)

    def it_emits_an_areas_of_expertise_section(self, base_kwargs: dict):
        doc = resume_functional(
            focus_areas=["Engineering Leadership", "Distributed Systems"],
            **base_kwargs,
        )
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AREAS OF EXPERTISE" in all_text.upper()
        assert "Engineering Leadership" in all_text
        assert "Distributed Systems" in all_text

    def it_emits_focus_areas_and_skills_before_experience(
        self, base_kwargs: dict, experience: list
    ):
        doc = resume_functional(
            focus_areas=["Engineering Leadership"],
            skills=["Python", "Go"],
            experience=experience,
            **base_kwargs,
        )
        texts = [p.text for p in doc.paragraphs]
        focus_idx = next(
            i for i, t in enumerate(texts) if "Engineering Leadership" in t
        )
        skills_idx = next(i for i, t in enumerate(texts) if "Python, Go" in t)
        exp_idx = next(i for i, t in enumerate(texts) if "Acme Corp" in t)
        assert focus_idx < exp_idx
        assert skills_idx < exp_idx

    def it_raises_when_name_is_empty(self):
        with pytest.raises(ValueError, match="name must be a non-empty string"):
            resume_functional(name="")

    @pytest.mark.parametrize("style", STYLES)
    def it_supports_every_built_in_style(self, style: str, base_kwargs: dict):
        doc = resume_functional(
            focus_areas=["Engineering Leadership"], style=style, **base_kwargs
        )
        texts = [p.text for p in doc.paragraphs]
        assert any("Ben Hooper" in t for t in texts)


# -- resume_technical -----------------------------------------------------


class DescribeResumeTechnical:
    """Behavioural tests for :func:`resume.resume_technical`."""

    def it_returns_a_Document(self, base_kwargs: dict):
        doc = resume_technical(
            projects=[
                {
                    "name": "monorepo-tool",
                    "role": "Author",
                    "tech": ["Python", "Rust"],
                    "bullets": ["Saved 10s on every CI run"],
                }
            ],
            tech_stack={"Languages": ["Python"], "Cloud": ["AWS"]},
            **base_kwargs,
        )
        assert isinstance(doc, DocumentCls)

    def it_emits_projects_with_tech_and_bullets(self, base_kwargs: dict):
        doc = resume_technical(
            projects=[
                {
                    "name": "monorepo-tool",
                    "tech": ["Python", "Rust", "AWS"],
                    "bullets": ["Saved 10s on every CI run"],
                }
            ],
            **base_kwargs,
        )
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "monorepo-tool" in all_text
        assert "Python, Rust, AWS" in all_text
        assert "Saved 10s on every CI run" in all_text

    def it_renders_project_url_as_a_hyperlink(self, base_kwargs: dict):
        doc = resume_technical(
            projects=[{"name": "p1", "url": "github.com/x/p1"}],
            **base_kwargs,
        )
        urls = _hyperlink_urls(doc)
        assert any("github.com/x/p1" in u for u in urls), (
            f"expected github.com/x/p1 hyperlink; got {urls!r}"
        )

    def it_emits_a_technical_skills_section(self, base_kwargs: dict):
        doc = resume_technical(
            tech_stack={"Cloud": ["AWS"], "Languages": ["Python"]},
            **base_kwargs,
        )
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "TECHNICAL SKILLS" in all_text.upper()
        assert "Cloud:" in all_text
        assert "Languages:" in all_text

    def it_emits_projects_before_experience(
        self, base_kwargs: dict, experience: list
    ):
        doc = resume_technical(
            projects=[{"name": "p1"}], experience=experience, **base_kwargs
        )
        texts = [p.text for p in doc.paragraphs]
        proj_idx = next(i for i, t in enumerate(texts) if "p1" in t)
        exp_idx = next(i for i, t in enumerate(texts) if "Acme Corp" in t)
        assert proj_idx < exp_idx

    def it_raises_when_name_is_empty(self):
        with pytest.raises(ValueError, match="name must be a non-empty string"):
            resume_technical(name="")

    @pytest.mark.parametrize("style", STYLES)
    def it_supports_every_built_in_style(self, style: str, base_kwargs: dict):
        doc = resume_technical(
            projects=[{"name": "p1", "bullets": ["a"]}],
            tech_stack=["Python"],
            style=style,
            **base_kwargs,
        )
        texts = [p.text for p in doc.paragraphs]
        assert any("Ben Hooper" in t for t in texts)


# -- Cross-cutting behaviour ---------------------------------------------


class DescribeContactLine:
    """Contact-line rendering is shared across all three factories."""

    def it_renders_recognised_contact_kinds_as_hyperlinks(self):
        doc = resume_chronological(
            name="Ben Hooper",
            contact={
                "email": "ben@example.com",
                "linkedin": "in/benhooper",
                "github": "benh",
                "website": "acme.com",
                "phone": "+61 2 1234 5678",
            },
        )
        joined = " | ".join(_hyperlink_urls(doc))
        assert "mailto:ben@example.com" in joined
        assert "linkedin.com/in/benhooper" in joined
        assert "github.com/benh" in joined
        assert "https://acme.com" in joined

    def it_passes_through_full_urls_unchanged(self):
        doc = resume_chronological(
            name="Ben Hooper",
            contact={
                "linkedin": "https://linkedin.com/in/someone-else",
                "github": "https://github.com/full",
            },
        )
        joined = " | ".join(_hyperlink_urls(doc))
        assert "https://linkedin.com/in/someone-else" in joined
        assert "https://github.com/full" in joined

    def it_renders_phone_location_and_unknown_keys_verbatim(self):
        doc = resume_chronological(
            name="Ben Hooper",
            contact={
                "phone": "+61 2 1234 5678",
                "location": "Sydney, Australia",
                "twitter": "@benh",
            },
        )
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "+61 2 1234 5678" in all_text
        assert "Sydney, Australia" in all_text
        assert "@benh" in all_text


class DescribeStyleVisuals:
    """Each visual style applies the right colour / heading shape."""

    def it_uses_heading_1_for_classic_section_headings(self):
        doc = resume_chronological(
            name="X",
            experience=[{"company": "C", "title": "T", "bullets": ["b"]}],
            style="classic",
        )
        heading_paras = [
            p for p in doc.paragraphs if p.style and p.style.name == "Heading 1"
        ]
        assert any("Experience" in p.text for p in heading_paras)

    def it_uppercases_modern_section_headings(self):
        doc = resume_chronological(
            name="X",
            experience=[{"company": "C", "title": "T"}],
            style="modern",
        )
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "EXPERIENCE" in all_text

    def it_centres_classic_name_block(self):
        doc = resume_chronological(
            name="Ben Hooper", title="Senior Engineer", style="classic"
        )
        title_paras = [
            p for p in doc.paragraphs if p.style and p.style.name == "Title"
        ]
        assert title_paras
        assert title_paras[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


class DescribeEmptyOptionalSections:
    """Optional sections are dropped cleanly when their inputs are empty."""

    def it_omits_summary_when_not_supplied(self):
        doc = resume_chronological(name="Ben Hooper")
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "SUMMARY" not in all_text.upper()

    def it_omits_experience_when_empty(self):
        doc = resume_chronological(name="Ben Hooper", experience=[])
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "EXPERIENCE" not in all_text.upper()

    def it_omits_skills_when_None(self):
        doc = resume_chronological(name="Ben Hooper", skills=None)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "SKILLS" not in all_text.upper()

    def it_omits_focus_areas_when_empty_in_functional(self):
        doc = resume_functional(name="Ben Hooper", focus_areas=[])
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "AREAS OF EXPERTISE" not in all_text.upper()

    def it_omits_projects_when_empty_in_technical(self):
        doc = resume_technical(name="Ben Hooper", projects=[])
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "PROJECTS" not in all_text.upper()
