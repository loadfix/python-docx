"""Unit-test suite for ``docx.kit.case_study`` helper."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import case_study
from docx.table import Table
from docx.text.paragraph import Paragraph


# -- Shared helpers -------------------------------------------------------


@pytest.fixture
def document() -> DocumentCls:
    return Document()


def _texts(doc: DocumentCls):
    return [p.text for p in doc.paragraphs]


def _full_text(doc: DocumentCls) -> str:
    return "\n".join(_texts(doc))


# -- Minimal-call signature -----------------------------------------------


class DescribeCaseStudy:
    """Unit-test suite for ``case_study.case_study``."""

    def it_returns_paragraphs_and_tables_in_document_order(
        self, document: DocumentCls
    ):
        result = case_study.case_study(
            document,
            title="How ACME cut latency by 80% with FrobnitzPro",
            customer="ACME Corp",
            industry="Manufacturing",
            size="5,000 employees",
            location="Detroit, MI",
            summary="One-paragraph elevator pitch.",
            challenge="ACME's primary challenge.",
            solution="With FrobnitzPro.",
            implementation="The rollout took 6 weeks.",
            results=[
                {"metric": "Latency", "before": "500ms",
                 "after": "100ms", "delta": "-80%"},
            ],
            customer_quote='"FrobnitzPro paid for itself."',
            technologies=["FrobnitzPro 5", "Kubernetes"],
            next_steps="Expand to EU region in Q3.",
        )

        # -- Every entry is either a Paragraph or a Table. --
        for item in result:
            assert isinstance(item, (Paragraph, Table))
        # -- Two tables expected: customer profile + results --
        tables = [item for item in result if isinstance(item, Table)]
        assert len(tables) == 2

    def it_renders_the_title_with_the_Title_style(
        self, document: DocumentCls
    ):
        result = case_study.case_study(
            document,
            title="How ACME won with Frobnitz",
            customer="ACME Corp",
        )

        title_para = result[0]
        assert isinstance(title_para, Paragraph)
        assert title_para.text == "How ACME won with Frobnitz"
        assert title_para.style is not None
        assert title_para.style.name == "Title"
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_renders_the_customer_name_centred_under_the_title(
        self, document: DocumentCls
    ):
        result = case_study.case_study(
            document, title="T", customer="ACME Corp"
        )

        customer_para = result[1]
        assert isinstance(customer_para, Paragraph)
        assert customer_para.text == "ACME Corp"
        assert customer_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_renders_the_3_column_customer_profile_strip(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="ACME",
            industry="Manufacturing",
            size="5,000 employees",
            location="Detroit, MI",
        )

        # -- The customer profile is the first table in the doc. --
        profile = document.tables[0]
        assert len(profile.rows) == 1
        assert len(profile.columns) == 3
        # -- Each cell carries label + newline + value. --
        cells = profile.rows[0].cells
        assert "Industry" in cells[0].text
        assert "Manufacturing" in cells[0].text
        assert "Size" in cells[1].text
        assert "5,000 employees" in cells[1].text
        assert "Location" in cells[2].text
        assert "Detroit, MI" in cells[2].text

    def it_skips_the_profile_strip_when_all_three_facts_missing(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document, title="T", customer="C"
        )

        assert document.tables == []

    def it_renders_the_profile_strip_when_only_one_fact_supplied(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            industry="Manufacturing",
        )

        assert len(document.tables) == 1

    def it_renders_each_narrative_section_with_a_Heading_1(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            summary="S",
            challenge="Ch",
            solution="So",
            implementation="Im",
            next_steps="Ns",
        )
        text = _full_text(document)

        for heading in (
            "Summary",
            "Challenge",
            "Solution",
            "Implementation",
            "Next Steps",
        ):
            assert heading in text

    def it_skips_narrative_sections_when_body_is_empty(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document, title="T", customer="C", summary="S"
        )
        text = _full_text(document)

        assert "Summary" in text
        # -- Other narrative section headings should NOT appear. --
        assert "Challenge" not in text
        assert "Solution" not in text
        assert "Implementation" not in text
        assert "Next Steps" not in text

    def it_splits_a_multi_paragraph_body_string_on_blank_lines(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            challenge="Para one.\n\nPara two.\n\nPara three.",
        )
        text = _full_text(document)

        assert "Para one." in text
        assert "Para two." in text
        assert "Para three." in text

    def it_accepts_a_sequence_body_one_paragraph_per_item(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            solution=["First step.", "Second step.", "Third step."],
        )
        text = _full_text(document)

        assert "First step." in text
        assert "Third step." in text

    def it_renders_the_4_column_results_table(
        self, document: DocumentCls
    ):
        results = [
            {"metric": "Latency", "before": "500ms",
             "after": "100ms", "delta": "-80%"},
            {"metric": "Throughput", "before": "1k/s",
             "after": "5k/s", "delta": "+400%"},
        ]
        case_study.case_study(
            document, title="T", customer="C", results=results
        )

        # -- The results table is the only table when no profile strip. --
        assert len(document.tables) == 1
        results_table = document.tables[0]
        assert len(results_table.columns) == 4
        # -- 1 header + 2 data rows --
        assert len(results_table.rows) == 3

        header = results_table.rows[0].cells
        assert header[0].text == "Metric"
        assert header[1].text == "Before"
        assert header[2].text == "After"
        assert header[3].text == "Delta"

        first_data = results_table.rows[1].cells
        assert first_data[0].text == "Latency"
        assert first_data[1].text == "500ms"
        assert first_data[2].text == "100ms"
        assert first_data[3].text == "-80%"

    def it_renders_blank_cells_for_missing_result_keys(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            results=[{"metric": "Latency"}],
        )
        results_table = document.tables[0]
        data = results_table.rows[1].cells
        assert data[0].text == "Latency"
        assert data[1].text == ""
        assert data[2].text == ""
        assert data[3].text == ""

    def it_skips_the_results_section_when_results_empty(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document, title="T", customer="C", results=[]
        )

        assert "Results" not in _full_text(document)

    def it_renders_the_customer_quote_in_the_Quote_style(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            customer_quote='"FrobnitzPro paid for itself."',
        )

        # -- Locate the quote paragraph (one with that exact text). --
        quote_para = next(
            p for p in document.paragraphs
            if p.text == '"FrobnitzPro paid for itself."'
        )
        assert quote_para.style is not None
        assert quote_para.style.name == "Quote"
        # -- Run italic is set as belt-and-braces. --
        assert any(run.italic for run in quote_para.runs)

    def it_renders_technologies_as_bullets_by_default(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            technologies=["FrobnitzPro 5", "Kubernetes", "PostgreSQL 17"],
        )

        bullet_texts = [
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "List Bullet"
        ]
        assert "FrobnitzPro 5" in bullet_texts
        assert "Kubernetes" in bullet_texts
        assert "PostgreSQL 17" in bullet_texts

    def it_renders_technologies_comma_separated_when_bullets_disabled(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document,
            title="T",
            customer="C",
            technologies=["A", "B", "C"],
            technologies_as_bullets=False,
        )
        text = _full_text(document)

        assert "A, B, C" in text

    def it_skips_the_technologies_section_when_list_empty(
        self, document: DocumentCls
    ):
        case_study.case_study(
            document, title="T", customer="C", technologies=[]
        )

        assert "Technologies" not in _full_text(document)

    def it_appends_a_trailing_page_break_by_default(
        self, document: DocumentCls
    ):
        result = case_study.case_study(
            document, title="T", customer="C"
        )

        # -- Last entry should be the page-break paragraph. --
        last = result[-1]
        assert isinstance(last, Paragraph)
        # -- Page break paragraph has a single empty run carrying a break --
        assert any(
            "br" in run.element.xml.lower() for run in last.runs
        )

    def it_skips_the_page_break_when_disabled(
        self, document: DocumentCls
    ):
        result = case_study.case_study(
            document, title="T", customer="C", page_break=False
        )

        # -- No page-break sentinel paragraph at the tail. --
        last = result[-1]
        # -- Tail is the customer name paragraph (no page-break) --
        assert isinstance(last, Paragraph)
        assert last.text == "C"

    def it_appends_to_an_existing_document_without_clobbering_content(
        self, document: DocumentCls
    ):
        # -- Pre-seed the document with one paragraph. --
        document.add_paragraph("pre-existing content")
        before = len(document.paragraphs)

        case_study.case_study(
            document,
            title="T",
            customer="C",
            summary="S",
        )

        assert document.paragraphs[0].text == "pre-existing content"
        assert len(document.paragraphs) > before

    def it_raises_when_title_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="title must be a non-empty string"):
            case_study.case_study(document, title="", customer="C")

    def it_raises_when_title_is_whitespace_only(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="title must be a non-empty string"):
            case_study.case_study(document, title="   ", customer="C")

    def it_raises_when_customer_is_empty(self, document: DocumentCls):
        with pytest.raises(
            ValueError, match="customer must be a non-empty string"
        ):
            case_study.case_study(document, title="T", customer="")

    def it_raises_when_a_results_entry_is_not_a_mapping(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="results\\[0\\] must be a mapping"):
            case_study.case_study(
                document,
                title="T",
                customer="C",
                results=["not-a-dict"],  # type: ignore[list-item]
            )

    def it_falls_back_to_Normal_when_Quote_style_is_missing(
        self, document: DocumentCls, monkeypatch
    ):
        # -- Simulate a template missing the Quote style. --
        original_resolve = case_study._resolve_style

        def fake_resolve(doc, preferred):
            if preferred == "Quote":
                return "Normal"
            return original_resolve(doc, preferred)

        monkeypatch.setattr(case_study, "_resolve_style", fake_resolve)

        case_study.case_study(
            document, title="T", customer="C", customer_quote="Q"
        )

        quote_para = next(
            p for p in document.paragraphs if p.text == "Q"
        )
        assert quote_para.style is not None
        assert quote_para.style.name == "Normal"
