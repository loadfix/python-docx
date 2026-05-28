"""Unit-test suite for ``docx.kit.templates`` document-template registry."""

from __future__ import annotations

import pytest

from docx.document import Document as DocumentCls
from docx.kit import templates
from docx.kit.templates import brief, coe, rfp_response, white_paper


# -- Shared helpers -------------------------------------------------------


def _texts(document: DocumentCls):
    """Return the text of every paragraph in ``document``."""
    return [p.text for p in document.paragraphs]


def _full_text(document: DocumentCls) -> str:
    return "\n".join(_texts(document))


# -- Brief ----------------------------------------------------------------


class DescribeBrief:
    """Unit-test suite for ``templates.brief``."""

    def it_returns_a_document_with_the_title(self):
        doc = brief(title="Q1 Strategy Brief")

        assert isinstance(doc, DocumentCls)
        assert "Q1 Strategy Brief" in _full_text(doc)

    def it_renders_an_optional_author_subtitle(self):
        doc = brief(title="Q1 Strategy Brief", author="Strategy Team")

        assert "Strategy Team" in _full_text(doc)

    def it_renders_an_optional_date(self):
        doc = brief(
            title="Q1 Strategy Brief",
            author="Strategy Team",
            date="2026-05-29",
        )

        assert "2026-05-29" in _full_text(doc)

    def it_omits_subtitle_when_author_unsupplied(self):
        doc = brief(title="Q1 Strategy Brief")
        # -- Without an author, no subtitle paragraph beyond the title
        # -- should carry author-like text. --
        non_empty = [t for t in _texts(doc) if t]
        assert non_empty[0] == "Q1 Strategy Brief"

    def it_renders_supplied_sections_as_headings_and_paragraphs(self):
        doc = brief(
            title="Q1 Strategy Brief",
            sections=[
                {"heading": "Background", "body": "Context."},
                {"heading": "Recommendation", "body": "Do X."},
                {"heading": "Next Steps", "body": "Plan Y."},
            ],
        )
        text = _full_text(doc)

        assert "Background" in text
        assert "Context." in text
        assert "Recommendation" in text
        assert "Do X." in text
        assert "Next Steps" in text
        assert "Plan Y." in text

    def it_accepts_a_sequence_body_for_multi_paragraph_sections(self):
        doc = brief(
            title="Q1 Strategy Brief",
            sections=[
                {
                    "heading": "Background",
                    "body": ["First paragraph.", "Second paragraph."],
                },
            ],
        )
        text = _full_text(doc)

        assert "First paragraph." in text
        assert "Second paragraph." in text

    def it_tolerates_a_section_without_a_body(self):
        doc = brief(
            title="Q1 Strategy Brief",
            sections=[{"heading": "Background"}],
        )

        assert "Background" in _full_text(doc)

    def it_works_with_no_sections(self):
        doc = brief(title="Q1 Strategy Brief")

        # -- No body sections is fine — the brief is just a title. --
        assert "Q1 Strategy Brief" in _full_text(doc)

    def it_raises_when_title_is_empty(self):
        with pytest.raises(ValueError, match="title is required"):
            brief(title="")

    def it_raises_when_title_is_whitespace_only(self):
        with pytest.raises(ValueError, match="title is required"):
            brief(title="   ")

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(ValueError, match="non-empty 'heading'"):
            brief(
                title="Q1 Strategy Brief",
                sections=[{"body": "no heading"}],
            )

    def it_raises_when_a_section_is_not_a_mapping(self):
        with pytest.raises(ValueError, match="must be a mapping"):
            brief(
                title="Q1 Strategy Brief",
                sections=["not a dict"],  # type: ignore[list-item]
            )


# -- Centre of Excellence -------------------------------------------------


class DescribeCoE:
    """Unit-test suite for ``templates.coe``."""

    def it_returns_a_document_with_the_name_in_the_title(self):
        doc = coe(name="Cloud CoE")

        assert isinstance(doc, DocumentCls)
        assert "Cloud CoE" in _full_text(doc)

    def it_renders_the_charter_subtitle(self):
        doc = coe(name="Cloud CoE")

        assert "Centre of Excellence Charter" in _full_text(doc)

    def it_renders_the_three_required_sections(self):
        doc = coe(name="Cloud CoE")
        text = _full_text(doc)

        # -- The fixed three-section structure: charter, governance,
        # -- services. --
        assert "Charter" in text
        assert "Governance" in text
        assert "Services" in text

    def it_renders_the_supplied_charter_paragraph(self):
        doc = coe(
            name="Cloud CoE",
            charter="Drive cloud adoption across the enterprise.",
        )

        assert "Drive cloud adoption across the enterprise." in _full_text(doc)

    def it_renders_a_placeholder_when_charter_missing(self):
        doc = coe(name="Cloud CoE")

        assert "[State the CoE's mission" in _full_text(doc)

    def it_renders_governance_as_a_bulleted_list(self):
        doc = coe(
            name="Cloud CoE",
            governance=[
                "Steering committee meets monthly.",
                "Decisions documented in ADRs.",
            ],
        )
        text = _full_text(doc)

        assert "Steering committee meets monthly." in text
        assert "Decisions documented in ADRs." in text

    def it_renders_a_placeholder_when_governance_missing(self):
        doc = coe(name="Cloud CoE")

        assert "[Describe steering" in _full_text(doc)

    def it_renders_services_as_a_bulleted_list(self):
        doc = coe(
            name="Cloud CoE",
            services=[
                "Platform engineering.",
                "Architecture review.",
                "Training and enablement.",
            ],
        )
        text = _full_text(doc)

        assert "Platform engineering." in text
        assert "Architecture review." in text
        assert "Training and enablement." in text

    def it_renders_a_placeholder_when_services_missing(self):
        doc = coe(name="Cloud CoE")

        assert "[List the services" in _full_text(doc)

    def it_renders_a_sponsor_metadata_line(self):
        doc = coe(name="Cloud CoE", sponsor="CTO")
        text = _full_text(doc)

        assert "Sponsor" in text
        assert "CTO" in text

    def it_omits_sponsor_metadata_when_unsupplied(self):
        doc = coe(name="Cloud CoE")

        assert "Sponsor:" not in _full_text(doc)

    def it_renders_an_optional_date(self):
        doc = coe(name="Cloud CoE", date="2026-05-29")

        assert "2026-05-29" in _full_text(doc)

    def it_raises_when_name_is_empty(self):
        with pytest.raises(ValueError, match="name is required"):
            coe(name="")

    def it_raises_when_name_is_whitespace_only(self):
        with pytest.raises(ValueError, match="name is required"):
            coe(name="   ")


# -- RFP response ---------------------------------------------------------


class DescribeRfpResponse:
    """Unit-test suite for ``templates.rfp_response``."""

    def it_returns_a_document_with_the_rfp_title_and_company(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
        )
        text = _full_text(doc)

        assert isinstance(doc, DocumentCls)
        assert "Cloud Migration Services RFP" in text
        assert "Acme Corp" in text

    def it_renders_an_optional_contact_metadata_line(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
            contact="Jane Doe, jane@acme.com",
        )
        text = _full_text(doc)

        assert "Contact" in text
        assert "Jane Doe, jane@acme.com" in text

    def it_renders_an_optional_date(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
            date="2026-05-29",
        )

        assert "2026-05-29" in _full_text(doc)

    def it_renders_supplied_sections_as_headings_and_paragraphs(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
            sections=[
                {"heading": "Executive Summary", "body": "We are leaders."},
                {"heading": "Approach", "body": "Phased migration."},
            ],
        )
        text = _full_text(doc)

        assert "Executive Summary" in text
        assert "We are leaders." in text
        assert "Approach" in text
        assert "Phased migration." in text

    def it_renders_a_pricing_heading(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
        )

        assert "Pricing" in _full_text(doc)

    def it_renders_a_placeholder_when_pricing_table_missing(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
        )

        assert "[Insert pricing line-items" in _full_text(doc)

    def it_renders_pricing_table_as_a_four_column_table(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
            pricing_table=[
                {
                    "item": "Discovery",
                    "quantity": 1,
                    "unit_price": "$25k",
                    "total": "$25k",
                },
                {
                    "item": "Migration",
                    "quantity": 1,
                    "unit_price": "$80k",
                    "total": "$80k",
                },
            ],
        )

        # -- Find the pricing table by its header --
        pricing_table = None
        for table in doc.tables:
            cells = [cell.text for cell in table.rows[0].cells]
            if cells and cells[0] == "Item":
                pricing_table = table
                break
        assert pricing_table is not None
        # -- header + two rows --
        assert len(pricing_table.rows) == 3
        header = [c.text for c in pricing_table.rows[0].cells]
        assert header == ["Item", "Quantity", "Unit Price", "Total"]
        first = pricing_table.rows[1].cells
        assert first[0].text == "Discovery"
        assert first[1].text == "1"
        assert first[2].text == "$25k"
        assert first[3].text == "$25k"

    def it_handles_pricing_rows_with_missing_optional_fields(self):
        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
            pricing_table=[{"item": "Bare item"}],
        )

        pricing_table = None
        for table in doc.tables:
            cells = [cell.text for cell in table.rows[0].cells]
            if cells and cells[0] == "Item":
                pricing_table = table
                break
        assert pricing_table is not None
        first = pricing_table.rows[1].cells
        assert first[0].text == "Bare item"
        assert first[1].text == ""
        assert first[2].text == ""
        assert first[3].text == ""

    def it_raises_when_rfp_title_is_empty(self):
        with pytest.raises(ValueError, match="rfp_title is required"):
            rfp_response(rfp_title="", company="Acme Corp")

    def it_raises_when_rfp_title_is_whitespace_only(self):
        with pytest.raises(ValueError, match="rfp_title is required"):
            rfp_response(rfp_title="   ", company="Acme Corp")

    def it_raises_when_company_is_empty(self):
        with pytest.raises(ValueError, match="company is required"):
            rfp_response(rfp_title="Cloud Migration Services RFP", company="")

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(ValueError, match="non-empty 'heading'"):
            rfp_response(
                rfp_title="Cloud Migration Services RFP",
                company="Acme Corp",
                sections=[{"body": "no heading"}],
            )

    def it_raises_when_a_pricing_row_has_no_item(self):
        with pytest.raises(ValueError, match="non-empty 'item'"):
            rfp_response(
                rfp_title="Cloud Migration Services RFP",
                company="Acme Corp",
                pricing_table=[{"unit_price": "$1k"}],
            )

    def it_raises_when_a_pricing_row_is_not_a_mapping(self):
        with pytest.raises(ValueError, match="must be a mapping"):
            rfp_response(
                rfp_title="Cloud Migration Services RFP",
                company="Acme Corp",
                pricing_table=["not a dict"],  # type: ignore[list-item]
            )


# -- White paper ----------------------------------------------------------


class DescribeWhitePaper:
    """Unit-test suite for ``templates.white_paper``."""

    def it_returns_a_document_with_the_title(self):
        doc = white_paper(title="The Future of OOXML")

        assert isinstance(doc, DocumentCls)
        assert "The Future of OOXML" in _full_text(doc)

    def it_renders_an_optional_author_subtitle(self):
        doc = white_paper(title="The Future of OOXML", author="Ben Hooper")

        assert "Ben Hooper" in _full_text(doc)

    def it_renders_an_optional_date(self):
        doc = white_paper(
            title="The Future of OOXML",
            author="Ben Hooper",
            date="2026-05-29",
        )

        assert "2026-05-29" in _full_text(doc)

    def it_renders_an_abstract_heading(self):
        doc = white_paper(title="The Future of OOXML")

        assert "Abstract" in _full_text(doc)

    def it_renders_the_supplied_abstract_paragraph(self):
        doc = white_paper(
            title="The Future of OOXML",
            abstract="OOXML remains the dominant document format.",
        )

        assert "OOXML remains the dominant document format." in _full_text(doc)

    def it_renders_a_placeholder_when_abstract_missing(self):
        doc = white_paper(title="The Future of OOXML")

        assert "[Summarise the white paper's thesis" in _full_text(doc)

    def it_renders_supplied_sections_as_headings_and_paragraphs(self):
        doc = white_paper(
            title="The Future of OOXML",
            sections=[
                {"heading": "Introduction", "body": "OOXML is widely used."},
                {"heading": "Background", "body": "Word, Excel, PowerPoint."},
            ],
        )
        text = _full_text(doc)

        assert "Introduction" in text
        assert "OOXML is widely used." in text
        assert "Background" in text
        assert "Word, Excel, PowerPoint." in text

    def it_renders_references_as_a_numbered_list(self):
        doc = white_paper(
            title="The Future of OOXML",
            references=[
                "Hooper, B. (2026). python-docx fork.",
                "ECMA-376 (2016). Office Open XML File Formats.",
            ],
        )
        text = _full_text(doc)

        assert "References" in text
        assert "Hooper, B. (2026). python-docx fork." in text
        assert (
            "ECMA-376 (2016). Office Open XML File Formats." in text
        )

    def it_omits_the_references_section_when_unsupplied(self):
        doc = white_paper(title="The Future of OOXML")

        # -- A white paper without citations is valid; the section
        # -- should not appear at all (no placeholder). --
        assert "References" not in _full_text(doc)

    def it_omits_the_references_section_when_empty_sequence(self):
        doc = white_paper(title="The Future of OOXML", references=[])

        assert "References" not in _full_text(doc)

    def it_raises_when_title_is_empty(self):
        with pytest.raises(ValueError, match="title is required"):
            white_paper(title="")

    def it_raises_when_title_is_whitespace_only(self):
        with pytest.raises(ValueError, match="title is required"):
            white_paper(title="   ")

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(ValueError, match="non-empty 'heading'"):
            white_paper(
                title="The Future of OOXML",
                sections=[{"body": "no heading"}],
            )


# -- Round-trip integration ----------------------------------------------


class DescribeTemplatesRoundTrip:
    """End-to-end smoke-tests: every factory produces a saveable document."""

    def it_can_save_a_brief_to_a_BytesIO(self):
        from io import BytesIO

        doc = brief(
            title="Q1 Strategy Brief",
            author="Strategy Team",
            sections=[
                {"heading": "Background", "body": "..."},
                {"heading": "Recommendation", "body": "..."},
                {"heading": "Next Steps", "body": "..."},
            ],
        )
        buf = BytesIO()
        doc.save(buf)
        # -- Word .docx is a zip; magic bytes are 'PK' --
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_coe_to_a_BytesIO(self):
        from io import BytesIO

        doc = coe(
            name="Cloud CoE",
            charter="Drive cloud adoption.",
            governance=["Monthly steering."],
            services=["Architecture review."],
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_an_rfp_response_to_a_BytesIO(self):
        from io import BytesIO

        doc = rfp_response(
            rfp_title="Cloud Migration Services RFP",
            company="Acme Corp",
            sections=[
                {"heading": "Executive Summary", "body": "Leaders."},
                {"heading": "Approach", "body": "Phased."},
            ],
            pricing_table=[
                {
                    "item": "Discovery",
                    "quantity": 1,
                    "unit_price": "$25k",
                    "total": "$25k",
                },
                {
                    "item": "Migration",
                    "quantity": 1,
                    "unit_price": "$80k",
                    "total": "$80k",
                },
            ],
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_white_paper_to_a_BytesIO(self):
        from io import BytesIO

        doc = white_paper(
            title="The Future of OOXML",
            author="Ben Hooper",
            abstract="OOXML remains dominant.",
            sections=[
                {"heading": "Introduction", "body": "Widely used."},
                {"heading": "Background", "body": "Word + Excel + PPT."},
            ],
            references=["Hooper, B. (2026)."],
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"


# -- Module surface -------------------------------------------------------


class DescribeTemplatesModule:
    """Module-level surface contracts."""

    def it_exposes_the_four_template_factories(self):
        for name in ("brief", "coe", "rfp_response", "white_paper"):
            assert hasattr(templates, name)
            assert name in templates.__all__

    def it_is_re_exported_from_the_kit_package(self):
        from docx.kit import templates as templates_pkg

        assert templates_pkg is templates
