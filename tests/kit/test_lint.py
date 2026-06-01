"""Unit-test suite for :mod:`docx.kit.lint` (issue #304)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_BREAK
from docx.kit import lint as lint_mod
from docx.shared import Pt
from docx.kit.lint import (
    BUILTIN_RULES,
    DEFAULT_STYLE_EXEMPTIONS,
    Finding,
    LintConfig,
    LintReport,
    Rule,
    lint,
    register_rule,
    registered_rules,
    unregister_rule,
)


@pytest.fixture
def document() -> DocumentCls:
    return Document()


@pytest.fixture(autouse=True)
def _restore_registry():
    """Snapshot and restore the registry around each test.

    Custom-rule tests would otherwise leak rules into sibling tests.
    """

    snapshot = dict(lint_mod._REGISTRY)  # noqa: SLF001 - test isolation
    yield
    lint_mod._REGISTRY.clear()  # noqa: SLF001
    lint_mod._REGISTRY.update(snapshot)  # noqa: SLF001


# ---------------------------------------------------------------------------
# Module-level smoke
# ---------------------------------------------------------------------------


class DescribeBuiltinRegistration:

    def it_registers_twelve_built_in_rules(self):
        for name in BUILTIN_RULES:
            assert name in registered_rules()

    def it_lists_the_rules_in_documented_order(self):
        names = registered_rules()
        # Each built-in must appear; relative order must match the
        # documented BUILTIN_RULES tuple.
        positions = [names.index(r) for r in BUILTIN_RULES]
        assert positions == sorted(positions)

    def it_returns_a_LintReport_when_called_on_a_clean_document(
        self, document: DocumentCls
    ):
        report = lint(document)
        assert isinstance(report, LintReport)
        assert report.document is document


# ---------------------------------------------------------------------------
# multiple-spaces
# ---------------------------------------------------------------------------


class DescribeMultipleSpaces:

    def it_flags_three_consecutive_spaces_in_a_run(
        self, document: DocumentCls
    ):
        document.add_paragraph("hello   world")
        report = lint(document)
        rules = [f.rule for f in report.findings]
        assert "multiple-spaces" in rules

    def it_does_not_flag_two_spaces_by_default(self, document: DocumentCls):
        # Two-space gaps are common after sentence-ending punctuation
        # (and after heading-numbering tokens). Default threshold is 3.
        document.add_paragraph("hello  world")
        report = lint(document)
        assert "multiple-spaces" not in [
            f.rule for f in report.findings
        ]

    def it_does_not_flag_a_single_space(self, document: DocumentCls):
        document.add_paragraph("hello world")
        report = lint(document)
        assert "multiple-spaces" not in [f.rule for f in report.findings]

    def it_collapses_runs_of_spaces_when_autofix_runs(
        self, document: DocumentCls
    ):
        document.add_paragraph("a   b    c")
        report = lint(document)
        applied = report.autofix(rules=["multiple-spaces"])
        assert applied >= 1
        assert document.paragraphs[0].text == "a b c"

    def it_marks_findings_with_autofix_available_true(
        self, document: DocumentCls
    ):
        document.add_paragraph("x   y")
        report = lint(document)
        ms = [f for f in report.findings if f.rule == "multiple-spaces"]
        assert ms and ms[0].autofix_available is True
        assert ms[0].autofix_description

    def it_skips_heading_numbering_double_space_gap(
        self, document: DocumentCls
    ):
        # `4.1  Three-LZA topology` is a deliberate heading template
        # convention — the linter must not flag the two-space gap.
        document.add_heading("4.1  Three-LZA topology", level=2)
        report = lint(document)
        assert "multiple-spaces" not in [f.rule for f in report.findings]

    def it_still_flags_extra_spaces_inside_a_heading(
        self, document: DocumentCls
    ):
        # The numbering gap is fine, but unrelated interior runs of
        # three-or-more spaces are still real defects.
        document.add_heading("4.1  Three-LZA   topology", level=2)
        report = lint(document)
        assert "multiple-spaces" in [f.rule for f in report.findings]

    def it_skips_list_styled_paragraphs_with_hanging_indent(
        self, document: DocumentCls
    ):
        # `List Bullet`-styled paragraphs commonly start with multi-space
        # padding before the bullet glyph. Don't flag those.
        para = document.add_paragraph("    - bullet item")
        para.style = document.styles["List Bullet"]
        report = lint(document)
        assert "multiple-spaces" not in [f.rule for f in report.findings]

    def it_does_not_flag_leading_whitespace_on_unstyled_paragraph(
        self, document: DocumentCls
    ):
        # Even on an ordinary paragraph, leading whitespace is not
        # multi-space (that's `tab-instead-of-indent`'s territory).
        document.add_paragraph("    hello world")
        report = lint(document)
        assert "multiple-spaces" not in [f.rule for f in report.findings]

    def it_threshold_is_configurable(self, document: DocumentCls):
        from docx.kit import lint as lint_mod

        document.add_paragraph("hello  world")  # exactly two spaces
        # By default this is silent (threshold 3).
        assert "multiple-spaces" not in [
            f.rule for f in lint(document).findings
        ]
        # Lowering the threshold to 2 surfaces the finding.
        original = lint_mod.MULTIPLE_SPACES_MIN_RUN
        lint_mod.MULTIPLE_SPACES_MIN_RUN = 2
        try:
            assert "multiple-spaces" in [
                f.rule for f in lint(document).findings
            ]
        finally:
            lint_mod.MULTIPLE_SPACES_MIN_RUN = original

    def it_exposes_run_index_and_space_count_in_details(
        self, document: DocumentCls
    ):
        # Issue #682: callers should not have to regex-parse the message
        # to recover which run was hit or how many spaces are involved.
        document.add_paragraph("hello     world")  # 5 spaces in run 0
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "multiple-spaces"
        )
        assert finding.details["run_index"] == 0
        assert finding.details["space_count"] == 5


# ---------------------------------------------------------------------------
# trailing-whitespace
# ---------------------------------------------------------------------------


class DescribeTrailingWhitespace:

    def it_flags_a_paragraph_ending_in_a_space(self, document: DocumentCls):
        document.add_paragraph("hello ")
        report = lint(document)
        assert "trailing-whitespace" in [f.rule for f in report.findings]

    def it_does_not_flag_a_paragraph_with_no_trailing_whitespace(
        self, document: DocumentCls
    ):
        document.add_paragraph("hello")
        assert "trailing-whitespace" not in [
            f.rule for f in lint(document).findings
        ]

    def it_trims_trailing_whitespace_on_autofix(self, document: DocumentCls):
        document.add_paragraph("hello   ")
        report = lint(document)
        report.autofix(rules=["trailing-whitespace"])
        assert document.paragraphs[0].text == "hello"


# ---------------------------------------------------------------------------
# tab-instead-of-indent
# ---------------------------------------------------------------------------


class DescribeTabInsteadOfIndent:

    def it_flags_a_leading_tab(self, document: DocumentCls):
        para = document.add_paragraph()
        para.add_run("\thello")
        report = lint(document)
        assert "tab-instead-of-indent" in [f.rule for f in report.findings]

    def it_does_not_flag_a_paragraph_with_no_leading_tab(
        self, document: DocumentCls
    ):
        document.add_paragraph("hello")
        assert "tab-instead-of-indent" not in [
            f.rule for f in lint(document).findings
        ]

    def it_strips_the_leading_tab_on_autofix_and_sets_left_indent(
        self, document: DocumentCls
    ):
        from docx.shared import Pt

        para = document.add_paragraph()
        para.add_run("\thello")
        report = lint(document)
        applied = report.autofix(rules=["tab-instead-of-indent"])
        assert applied == 1
        assert document.paragraphs[0].runs[0].text == "hello"
        # Indent must be set so the visual position survives the
        # tab-stripping — one tab == 36pt (Word's default tab-stop).
        assert document.paragraphs[0].paragraph_format.left_indent == Pt(36)

    def it_scales_left_indent_with_the_count_of_leading_tabs(
        self, document: DocumentCls
    ):
        from docx.shared import Pt

        para = document.add_paragraph()
        para.add_run("\t\t\thello")
        report = lint(document)
        assert report.autofix(rules=["tab-instead-of-indent"]) == 1
        # Three tabs should produce three tab-stops of indent (108pt).
        assert document.paragraphs[0].runs[0].text == "hello"
        assert document.paragraphs[0].paragraph_format.left_indent == Pt(108)

    def it_reports_the_count_of_leading_tabs_in_the_finding_message(
        self, document: DocumentCls
    ):
        para = document.add_paragraph()
        para.add_run("\t\thello")
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "tab-instead-of-indent"
        )
        assert "2 literal tab" in finding.message

    def it_advertises_the_substitution_in_the_autofix_description(
        self, document: DocumentCls
    ):
        para = document.add_paragraph()
        para.add_run("\thello")
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "tab-instead-of-indent"
        )
        assert finding.autofix_description is not None
        assert "left-indent" in finding.autofix_description

    def it_skips_heading_paragraphs(self, document: DocumentCls):
        # A leading tab on a heading is almost always the rendered
        # leader between number and title; stripping it is destructive.
        h = document.add_heading("First", level=1)
        h.runs[0].text = "\t" + h.runs[0].text
        report = lint(document)
        assert "tab-instead-of-indent" not in [
            f.rule for f in report.findings
        ]

    def it_skips_list_paragraphs(self, document: DocumentCls):
        # A leading tab on a list item is the list's number/bullet
        # leader; the indent is already controlled by ``w:numPr``.
        try:
            para = document.add_paragraph(style="List Number")
        except KeyError:
            pytest.skip("default template lacks 'List Number' style")
        para.add_run("\thello")
        report = lint(document)
        # The fixture template may not register List Number with a
        # numPr; skip the assertion if no numPr was actually applied.
        if para.list_level is None:
            pytest.skip("List Number style did not produce a numPr")
        assert "tab-instead-of-indent" not in [
            f.rule for f in report.findings
        ]

    def it_preserves_an_existing_left_indent_when_autofixing(
        self, document: DocumentCls
    ):
        from docx.shared import Pt

        para = document.add_paragraph()
        para.paragraph_format.left_indent = Pt(18)
        para.add_run("\thello")
        report = lint(document)
        report.autofix(rules=["tab-instead-of-indent"])
        # Existing 18pt + one tab-stop (36pt) == 54pt.
        assert document.paragraphs[0].paragraph_format.left_indent == Pt(54)


# ---------------------------------------------------------------------------
# mixed-quotes
# ---------------------------------------------------------------------------


class DescribeMixedQuotes:

    def it_flags_a_paragraph_mixing_smart_and_straight(
        self, document: DocumentCls
    ):
        document.add_paragraph("she said “hello” and 'goodbye'")
        report = lint(document)
        mq = [f for f in report.findings if f.rule == "mixed-quotes"]
        assert mq and mq[0].severity == "info"
        assert mq[0].autofix_available is False

    def it_does_not_flag_a_paragraph_with_only_smart_quotes(
        self, document: DocumentCls
    ):
        document.add_paragraph("she said “hello”")
        assert "mixed-quotes" not in [f.rule for f in lint(document).findings]


# ---------------------------------------------------------------------------
# empty-paragraph
# ---------------------------------------------------------------------------


class DescribeEmptyParagraph:

    def it_does_not_flag_a_single_blank_paragraph_between_sentences(
        self, document: DocumentCls
    ):
        document.add_paragraph("first")
        document.add_paragraph("")
        document.add_paragraph("second")
        report = lint(document)
        assert "empty-paragraph" not in [f.rule for f in report.findings]

    def it_flags_consecutive_blank_paragraphs(self, document: DocumentCls):
        document.add_paragraph("first")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("second")
        report = lint(document)
        ep = [f for f in report.findings if f.rule == "empty-paragraph"]
        assert len(ep) == 2

    def it_removes_consecutive_empties_on_autofix_keeping_one(
        self, document: DocumentCls
    ):
        document.add_paragraph("first")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("second")
        report = lint(document)
        report.autofix(rules=["empty-paragraph"])
        # First, one empty, second.
        texts = [p.text for p in document.paragraphs]
        assert texts == ["first", "", "second"]

    # -- structural-content guard (issue #656) ---------------------------

    def it_does_not_flag_a_paragraph_carrying_only_a_page_break(
        self, document: DocumentCls
    ):
        document.add_paragraph("first")
        document.add_paragraph("")  # genuine drift
        para = document.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph("last")

        report = lint(document)
        ep = [f for f in report.findings if f.rule == "empty-paragraph"]
        # Neither the paragraph carrying the page break nor the genuine
        # blank that immediately precedes it should be flagged — the run
        # of consecutive empties is broken by the structural paragraph.
        assert ep == []

    def it_preserves_a_page_break_paragraph_through_autofix(
        self, document: DocumentCls
    ):
        document.add_paragraph("first")
        document.add_paragraph("")
        para = document.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph("last")

        report = lint(document)
        report.autofix(rules=["empty-paragraph"])

        # The page-break-carrying paragraph must still be present.
        assert any(p.has_page_break for p in document.paragraphs)
        # And the document still has four paragraphs (nothing destroyed).
        assert len(document.paragraphs) == 4

    def it_does_not_flag_a_paragraph_carrying_only_a_column_break(
        self, document: DocumentCls
    ):
        document.add_paragraph("first")
        document.add_paragraph("")
        para = document.add_paragraph()
        para.add_run().add_break(WD_BREAK.COLUMN)
        document.add_paragraph("last")

        report = lint(document)
        assert "empty-paragraph" not in [f.rule for f in report.findings]

    def it_does_not_flag_a_paragraph_carrying_only_a_line_break(
        self, document: DocumentCls
    ):
        # Even a soft line-break carries layout intent — never silently
        # delete a `<w:br>` paragraph.
        document.add_paragraph("first")
        document.add_paragraph("")
        para = document.add_paragraph()
        para.add_run().add_break(WD_BREAK.LINE)
        document.add_paragraph("last")

        report = lint(document)
        assert "empty-paragraph" not in [f.rule for f in report.findings]

    def it_does_not_flag_a_paragraph_carrying_a_bookmark_anchor(
        self, document: DocumentCls
    ):
        document.add_paragraph("first")
        document.add_paragraph("")
        anchor = document.add_paragraph()
        # Use the public Bookmarks API to drop an anchor inside the
        # otherwise-empty paragraph.
        document.bookmarks.add("anchor-1", anchor)
        document.add_paragraph("last")

        report = lint(document)
        assert "empty-paragraph" not in [f.rule for f in report.findings]

    def it_autofix_refuses_a_handcrafted_finding_for_a_break_paragraph(
        self, document: DocumentCls
    ):
        # Defence-in-depth: a caller could synthesise a Finding directly
        # via register_rule. The autofix callback must still refuse to
        # delete a paragraph that carries a `<w:br>`.
        document.add_paragraph("first")
        para = document.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph("last")

        forged = Finding(
            rule="empty-paragraph",
            severity="info",
            message="forged",
            paragraph_index=1,
            autofix_available=True,
            autofix_description="forged",
            location="paragraph 1",
        )
        ok = lint_mod._autofix_empty_paragraph(document, forged)
        assert ok is False
        assert len(document.paragraphs) == 3
        assert document.paragraphs[1].has_page_break

    # -- issue #656: tightened-predicate regression coverage -------------

    def it_does_not_flag_a_paragraph_with_a_page_break_as_empty(
        self, document: DocumentCls
    ):
        # Sanity: a single paragraph carrying only a page break, with a
        # genuine drift blank just before it, must not be reported as
        # an empty-paragraph finding — the autofix would silently lose
        # the page break.
        document.add_paragraph("first")
        document.add_paragraph("")
        para = document.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph("last")

        report = lint(document)
        ep = [f for f in report.findings if f.rule == "empty-paragraph"]
        assert ep == []

    def it_does_not_flag_a_paragraph_with_a_bookmark_as_empty(
        self, document: DocumentCls
    ):
        # Inject a bare <w:bookmarkStart>/<w:bookmarkEnd> pair into an
        # otherwise-empty paragraph via direct etree manipulation — the
        # public API requires a registered Bookmarks owner; this test
        # is the spec-level check that the predicate catches the raw
        # XML element regardless of how it got there.
        from lxml import etree

        from docx.oxml.ns import qn

        document.add_paragraph("first")
        document.add_paragraph("")
        anchor = document.add_paragraph()
        bm_start = etree.SubElement(anchor._p, qn("w:bookmarkStart"))
        bm_start.set(qn("w:id"), "0")
        bm_start.set(qn("w:name"), "anchor-1")
        bm_end = etree.SubElement(anchor._p, qn("w:bookmarkEnd"))
        bm_end.set(qn("w:id"), "0")
        document.add_paragraph("last")

        report = lint(document)
        ep = [f for f in report.findings if f.rule == "empty-paragraph"]
        assert ep == []

    def it_does_not_flag_a_paragraph_with_a_section_break_as_empty(
        self, document: DocumentCls
    ):
        # A paragraph carrying a <w:pPr>/<w:sectPr> is a section break;
        # losing it would change the document's section layout.
        document.add_paragraph("first")
        document.add_paragraph("")
        section_para = document.add_paragraph()
        section_para.insert_section_break()
        document.add_paragraph("last")

        report = lint(document)
        ep = [f for f in report.findings if f.rule == "empty-paragraph"]
        assert ep == []

    def it_still_flags_a_truly_empty_paragraph(
        self, document: DocumentCls
    ):
        # Multiple genuinely-blank paragraphs in a row with no
        # structural content still produce findings — the tightening
        # must not regress the rule's primary purpose.
        document.add_paragraph("first")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("last")

        report = lint(document)
        ep = [f for f in report.findings if f.rule == "empty-paragraph"]
        assert len(ep) == 2
        # Every finding emitted by the built-in rule defaults to
        # safe_to_delete=True — these are blank-line drift, no
        # structural content, fine to delete.
        assert all(f.safe_to_delete for f in ep)

    def it_autofix_preserves_paragraphs_with_load_bearing_content(
        self, document: DocumentCls
    ):
        # Build a doc with mixed empties: a genuine blank, a page-break
        # paragraph, a section-break paragraph, and a bookmarked
        # paragraph. Run autofix end-to-end and assert every load-
        # bearing paragraph survives, while genuine blank-line drift
        # is collapsed where appropriate.
        from lxml import etree

        from docx.oxml.ns import qn

        document.add_paragraph("first")
        document.add_paragraph("")  # genuine drift
        document.add_paragraph("")  # genuine drift (this one is removed)
        # page-break paragraph
        page_break_para = document.add_paragraph()
        page_break_para.add_run().add_break(WD_BREAK.PAGE)
        # section-break paragraph
        sect_para = document.add_paragraph()
        sect_para.insert_section_break()
        # bookmark anchor paragraph
        bookmark_para = document.add_paragraph()
        bm_start = etree.SubElement(bookmark_para._p, qn("w:bookmarkStart"))
        bm_start.set(qn("w:id"), "0")
        bm_start.set(qn("w:name"), "load-bearing")
        bm_end = etree.SubElement(bookmark_para._p, qn("w:bookmarkEnd"))
        bm_end.set(qn("w:id"), "0")
        document.add_paragraph("last")

        report = lint(document)
        report.autofix(rules=["empty-paragraph"])

        # The page-break, section-break, and bookmark paragraphs must
        # all still be present.
        assert any(p.has_page_break for p in document.paragraphs)
        assert any(p.has_section_break for p in document.paragraphs)
        # Bookmark anchors survive — find a paragraph carrying a
        # <w:bookmarkStart> child.
        assert any(
            p._p.xpath(".//w:bookmarkStart") for p in document.paragraphs
        )

    def it_round_trips_after_autofix(
        self, document: DocumentCls, tmp_path
    ):
        # Save + reopen and verify section breaks (and other load-
        # bearing structure) survive the autofix → save → reopen path.
        from io import BytesIO

        from docx import Document

        document.add_paragraph("first")
        document.add_paragraph("")
        document.add_paragraph("")
        sect_para = document.add_paragraph()
        sect_para.insert_section_break()
        document.add_paragraph("last")

        report = lint(document)
        report.autofix(rules=["empty-paragraph"])

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        reopened = Document(buffer)

        assert any(p.has_section_break for p in reopened.paragraphs)
        # The "last" body paragraph survives.
        assert any(p.text == "last" for p in reopened.paragraphs)

    def it_marks_handcrafted_finding_safe_to_delete_false_for_breaks(
        self, document: DocumentCls
    ):
        # A caller building a Finding by hand with safe_to_delete=False
        # must have its autofix skipped, and the report should record a
        # one-line preservation note.
        document.add_paragraph("first")
        para = document.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        document.add_paragraph("last")

        report = LintReport(
            document=document,
            findings=[
                Finding(
                    rule="empty-paragraph",
                    severity="info",
                    message="hand-built",
                    paragraph_index=1,
                    autofix_available=True,
                    autofix_description="forged",
                    location="paragraph 1",
                    safe_to_delete=False,
                ),
            ],
        )
        applied = report.autofix(rules=["empty-paragraph"])
        assert applied == 0
        assert len(document.paragraphs) == 3
        assert document.paragraphs[1].has_page_break
        assert len(report.preservation_notes) == 1
        assert "preserved" in report.preservation_notes[0]


# ---------------------------------------------------------------------------
# Finding.safe_to_delete (issue #656)
# ---------------------------------------------------------------------------


class DescribeFindingSafeToDelete:
    """Issue #656 — Finding.safe_to_delete defaults to True."""

    def it_defaults_to_true(self):
        finding = Finding(
            rule="x", severity="info", message="m"
        )
        assert finding.safe_to_delete is True

    def it_can_be_set_false_by_a_caller(self):
        finding = Finding(
            rule="x", severity="info", message="m", safe_to_delete=False
        )
        assert finding.safe_to_delete is False


# ---------------------------------------------------------------------------
# trailing-empty-paragraph
# ---------------------------------------------------------------------------


class DescribeTrailingEmptyParagraph:
    """Issue #677 — surface trailing empty paragraphs that
    ``empty-paragraph`` silently misses."""

    def it_flags_two_or_more_trailing_empties(self, document: DocumentCls):
        document.add_paragraph("body text one")
        document.add_paragraph("body text two")
        document.add_paragraph("")
        document.add_paragraph("")
        report = lint(document)
        tep = [
            f for f in report.findings if f.rule == "trailing-empty-paragraph"
        ]
        # Both trailing empties are flagged (the existing
        # empty-paragraph rule misses the first).
        assert len(tep) == 2

    def it_does_not_flag_a_single_trailing_empty_in_a_normal_document(
        self, document: DocumentCls
    ):
        # A normal Word document carries a single trailing paragraph as
        # the section-properties anchor — flagging that would be noisy.
        document.add_paragraph("body one")
        document.add_paragraph("body two")
        document.add_paragraph("body three")
        document.add_paragraph("body four")
        document.add_paragraph("")
        report = lint(document)
        assert "trailing-empty-paragraph" not in [
            f.rule for f in report.findings
        ]

    def it_does_flag_a_single_trailing_empty_in_a_tiny_document(
        self, document: DocumentCls
    ):
        # When the document is genuinely small (<= 3 paragraphs) a
        # trailing empty is much more likely to be authoring residue.
        document.add_paragraph("only body line")
        document.add_paragraph("")
        report = lint(document)
        assert "trailing-empty-paragraph" in [f.rule for f in report.findings]

    def it_does_not_flag_when_the_last_paragraph_has_content(
        self, document: DocumentCls
    ):
        document.add_paragraph("body one")
        document.add_paragraph("body two")
        report = lint(document)
        assert "trailing-empty-paragraph" not in [
            f.rule for f in report.findings
        ]

    def it_does_not_flag_an_empty_document(self, document: DocumentCls):
        # New Document() carries one empty paragraph by default. The
        # rule has no body content to compare against and should stay
        # quiet — flagging the only paragraph would always be wrong.
        report = lint(document)
        assert "trailing-empty-paragraph" not in [
            f.rule for f in report.findings
        ]

    def it_removes_trailing_empties_on_autofix(self, document: DocumentCls):
        document.add_paragraph("body one")
        document.add_paragraph("body two")
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("")
        report = lint(document)
        applied = report.autofix(rules=["trailing-empty-paragraph"])
        # Three trailing empties removed.
        assert applied == 3
        texts = [p.text for p in document.paragraphs]
        assert texts == ["body one", "body two"]


# ---------------------------------------------------------------------------
# inconsistent-heading-levels
# ---------------------------------------------------------------------------


class DescribeInconsistentHeadingLevels:

    def it_flags_a_skipped_level(self, document: DocumentCls):
        document.add_heading("First", level=1)
        document.add_heading("Skipped", level=3)
        report = lint(document)
        assert "inconsistent-heading-levels" in [
            f.rule for f in report.findings
        ]

    def it_does_not_flag_a_clean_progression(self, document: DocumentCls):
        document.add_heading("a", level=1)
        document.add_heading("b", level=2)
        document.add_heading("c", level=3)
        report = lint(document)
        assert "inconsistent-heading-levels" not in [
            f.rule for f in report.findings
        ]

    def it_marks_the_finding_as_no_autofix(self, document: DocumentCls):
        document.add_heading("First", level=1)
        document.add_heading("Skipped", level=3)
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "inconsistent-heading-levels"
        )
        assert finding.autofix_available is False

    def it_flags_when_first_heading_skips_from_document_root(
        self, document: DocumentCls
    ):
        # Document with no Heading 1 / Title — the very first heading
        # is Heading 2. The implicit pre-document level is 0, so
        # jumping to 2 is a skip and must be flagged.
        document.add_paragraph("Some intro prose without a heading.")
        document.add_heading("Skipped root", level=2)
        document.add_heading("Sub", level=3)
        report = lint(document)
        ihl = [
            f for f in report.findings if f.rule == "inconsistent-heading-levels"
        ]
        assert len(ihl) == 1
        assert "level 0 to level 2" in ihl[0].message

    def it_does_not_flag_when_first_heading_is_a_title(
        self, document: DocumentCls
    ):
        # Title is treated as level 0 by ``_heading_level``; using it
        # before Heading 1 is a clean progression, not a skip.
        document.add_paragraph("intro").style = document.styles["Title"]
        document.add_heading("First", level=1)
        document.add_heading("Sub", level=2)
        report = lint(document)
        assert "inconsistent-heading-levels" not in [
            f.rule for f in report.findings
        ]

    def it_exposes_structured_levels_in_details(self, document: DocumentCls):
        document.add_heading("First", level=1)
        document.add_heading("Skipped", level=3)
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "inconsistent-heading-levels"
        )
        # Issue #678: callers should not have to regex-parse the message.
        assert finding.details["level"] == 3
        assert finding.details["previous_level"] == 1
        assert finding.details["skipped"] == 1

    def it_reports_the_skipped_count_for_multi_level_jumps(
        self, document: DocumentCls
    ):
        document.add_heading("First", level=1)
        document.add_heading("Deep", level=5)
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "inconsistent-heading-levels"
        )
        # 1 -> 5 skipped levels 2, 3, 4 (count 3).
        assert finding.details["skipped"] == 3


class DescribeFindingDetails:
    """Issue #678 — Finding.details defaults to an empty mapping."""

    def it_defaults_to_an_empty_read_only_mapping(self):
        from docx.kit.lint import Finding

        finding = Finding(rule="x", severity="info", message="m")
        assert dict(finding.details) == {}
        # Read-only — assignment must raise.
        try:
            finding.details["k"] = 1  # type: ignore[index]
        except TypeError:
            pass
        else:  # pragma: no cover
            raise AssertionError("details should be read-only")


# ---------------------------------------------------------------------------
# missing-alt-text
# ---------------------------------------------------------------------------


class DescribeMissingAltText:

    def it_flags_an_inline_image_without_alt_text(
        self, document: DocumentCls
    ):
        # Use a real fixture image so the python-docx image-header
        # parser actually accepts it.
        document.add_picture(
            "tests/test_files/python-icon.png"
        )
        report = lint(document)
        ma = [f for f in report.findings if f.rule == "missing-alt-text"]
        assert ma and ma[0].autofix_available is False

    def it_defaults_to_info_severity(self, document: DocumentCls):
        # No core-properties title and no a11y-bearing shapes, so the
        # rule should not pretend the omission is a real defect.
        document.add_picture("tests/test_files/python-icon.png")
        report = lint(document)
        ma = [f for f in report.findings if f.rule == "missing-alt-text"]
        assert ma and ma[0].severity == "info"

    def it_escalates_to_warning_when_document_shows_a11y_intent(
        self, document: DocumentCls
    ):
        # Title set + at least one image with alt text => the author is
        # paying attention to a11y, so a missing alt is much more likely
        # to be a real bug.
        document.core_properties.title = "Quarterly Report"
        document.add_picture("tests/test_files/python-icon.png")
        document.inline_shapes[0].alt_text = "Python logo"
        document.add_picture("tests/test_files/monty-truth.png")
        report = lint(document)
        ma = [f for f in report.findings if f.rule == "missing-alt-text"]
        assert ma and ma[0].severity == "warning"

    def it_skips_shapes_marked_decorative_via_a11y_role(
        self, document: DocumentCls
    ):
        document.add_picture("tests/test_files/python-icon.png")
        # python-docx's a11y_role hooks the "[decorative]" prefix in
        # @descr — the lint rule must honour it.
        document.inline_shapes[0].a11y_role = "decorative"
        report = lint(document)
        assert "missing-alt-text" not in [
            f.rule for f in report.findings
        ]

    def it_skips_shapes_with_office365_decorative_extension(
        self, document: DocumentCls
    ):
        document.add_picture("tests/test_files/python-icon.png")
        # Splice in the Office 365 "Mark as decorative" marker that
        # Word writes when the user ticks the checkbox in the alt-text
        # pane. python-docx exposes the underlying docPr element so we
        # can append the extension via lxml without a private API.
        from lxml import etree

        docPr = document.inline_shapes[0]._inline.docPr
        ext_ns = "http://schemas.microsoft.com/office/drawing/2017/decorative"
        a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
        extLst = etree.SubElement(docPr, f"{{{a_ns}}}extLst")
        ext = etree.SubElement(extLst, f"{{{a_ns}}}ext")
        ext.set("uri", "{C183D7F6-B498-43B3-948B-1728B52AA6E4}")
        dec = etree.SubElement(ext, f"{{{ext_ns}}}decorative")
        dec.set("val", "1")
        report = lint(document)
        assert "missing-alt-text" not in [
            f.rule for f in report.findings
        ]

    def it_collapses_duplicate_images_to_a_single_finding(
        self, document: DocumentCls
    ):
        # The same image inserted three times should produce *one*
        # finding (not three), with a message that mentions the
        # duplicate count.
        for _ in range(3):
            document.add_picture("tests/test_files/python-icon.png")
        report = lint(document)
        ma = [f for f in report.findings if f.rule == "missing-alt-text"]
        assert len(ma) == 1
        assert "3" in ma[0].message  # mentions the repeat count

    def it_does_not_collapse_distinct_images(self, document: DocumentCls):
        document.add_picture("tests/test_files/python-icon.png")
        document.add_picture("tests/test_files/monty-truth.png")
        report = lint(document)
        ma = [f for f in report.findings if f.rule == "missing-alt-text"]
        assert len(ma) == 2


# ---------------------------------------------------------------------------
# mixed-fonts
# ---------------------------------------------------------------------------


class DescribeMixedFonts:

    def it_flags_paragraph_with_two_font_families(
        self, document: DocumentCls
    ):
        # Two sans-serif fonts — visible-but-not-loud clash, severity
        # stays at ``info`` (Issue #680).
        para = document.add_paragraph()
        run_a = para.add_run("hello ")
        run_a.font.name = "Calibri"
        run_b = para.add_run("world")
        run_b.font.name = "Arial"
        report = lint(document)
        mf = [f for f in report.findings if f.rule == "mixed-fonts"]
        assert mf and mf[0].severity == "info"
        assert mf[0].autofix_available is False

    def it_does_not_flag_paragraph_with_a_single_family(
        self, document: DocumentCls
    ):
        para = document.add_paragraph()
        for word in ("hello", " world"):
            r = para.add_run(word)
            r.font.name = "Calibri"
        report = lint(document)
        assert "mixed-fonts" not in [f.rule for f in report.findings]

    def it_escalates_to_warning_when_serif_and_sans_straddle(
        self, document: DocumentCls
    ):
        # Issue #680: Calibri (sans) + Times New Roman (serif) is a
        # visually loud defect — surface as ``warning``, not ``info``.
        para = document.add_paragraph()
        para.add_run("hello ").font.name = "Calibri"
        para.add_run("world").font.name = "Times New Roman"
        report = lint(document)
        mf = [f for f in report.findings if f.rule == "mixed-fonts"]
        assert mf and mf[0].severity == "warning"
        assert mf[0].details["straddles_serif_sans"] is True

    def it_exposes_font_names_and_count_in_details(
        self, document: DocumentCls
    ):
        # Issue #680: structured payload, not regex-the-message.
        para = document.add_paragraph()
        para.add_run("a").font.name = "Calibri"
        para.add_run("b").font.name = "Arial"
        para.add_run("c").font.name = "Verdana"
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "mixed-fonts"
        )
        assert finding.details["font_names"] == ("Arial", "Calibri", "Verdana")
        assert finding.details["count"] == 3
        # Three sans fonts — no serif, so no straddle.
        assert finding.details["straddles_serif_sans"] is False


# ---------------------------------------------------------------------------
# missing-document-title
# ---------------------------------------------------------------------------


class DescribeMissingDocumentTitle:

    def it_stays_silent_for_a_no_title_in_memory_document(
        self, document: DocumentCls
    ):
        # An in-memory Document() has no on-disk filename and therefore
        # no autofix path; the rule suppresses the finding rather than
        # emitting a permanent info-level no-op.
        document.core_properties.title = ""
        report = lint(document)
        mdt = [
            f for f in report.findings if f.rule == "missing-document-title"
        ]
        assert mdt == []

    def it_does_not_flag_when_title_is_set(self, document: DocumentCls):
        document.core_properties.title = "My Doc"
        report = lint(document)
        assert "missing-document-title" not in [
            f.rule for f in report.findings
        ]

    def it_autofixes_from_filename_when_loaded_from_disk(self, tmp_path):
        # Document(path) records the load path automatically — no side-
        # channel attribute setup required by the caller.
        path = tmp_path / "report-final.docx"
        Document().save(str(path))
        doc = Document(str(path))
        doc.core_properties.title = ""
        report = lint(doc)
        mdt = next(
            f for f in report.findings if f.rule == "missing-document-title"
        )
        assert mdt.autofix_available is True
        applied = report.autofix(rules=["missing-document-title"])
        assert applied == 1
        assert doc.core_properties.title == "report-final"

    def it_autofixes_from_an_explicit_source_path_kwarg(self, tmp_path):
        # When the document was loaded from a stream the caller can
        # pass the filename explicitly via ``source_path``.
        from io import BytesIO

        buf = BytesIO()
        Document().save(buf)
        buf.seek(0)
        doc = Document(buf)
        doc.core_properties.title = ""
        # Document(stream) does not auto-set _lint_filename
        assert getattr(doc, "_lint_filename", None) is None
        report = lint(doc, source_path="archive/quarterly-report.docx")
        mdt = next(
            f for f in report.findings if f.rule == "missing-document-title"
        )
        assert mdt.autofix_available is True
        assert "quarterly-report" in (mdt.autofix_description or "")
        applied = report.autofix(rules=["missing-document-title"])
        assert applied == 1
        assert doc.core_properties.title == "quarterly-report"

    def it_source_path_does_not_persist_after_lint(
        self, document: DocumentCls
    ):
        # The lint() call must restore the document's state w.r.t.
        # _lint_filename so back-to-back lint passes are independent.
        assert getattr(document, "_lint_filename", None) is None
        lint(document, source_path="ignored.docx")
        assert getattr(document, "_lint_filename", None) is None

    def it_source_path_accepts_pathlike(self, tmp_path, document: DocumentCls):
        # source_path may be str or os.PathLike; the rule should still
        # extract the right stem.
        document.core_properties.title = ""
        report = lint(document, source_path=tmp_path / "neat.docx")
        mdt = next(
            f for f in report.findings if f.rule == "missing-document-title"
        )
        assert mdt.autofix_available is True
        assert "'neat'" in (mdt.autofix_description or "")


# ---------------------------------------------------------------------------
# over-long-paragraph
# ---------------------------------------------------------------------------


class DescribeOverLongParagraph:

    def it_flags_a_paragraph_longer_than_1000_characters(
        self, document: DocumentCls
    ):
        document.add_paragraph("x" * 1500)
        report = lint(document)
        assert "over-long-paragraph" in [f.rule for f in report.findings]

    def it_does_not_flag_a_short_paragraph(self, document: DocumentCls):
        document.add_paragraph("short")
        report = lint(document)
        assert "over-long-paragraph" not in [
            f.rule for f in report.findings
        ]

    def it_exposes_char_count_and_threshold_in_details(
        self, document: DocumentCls
    ):
        # Issue #682: structured payload, not regex-the-message.
        document.add_paragraph("x" * 1500)
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "over-long-paragraph"
        )
        assert finding.details["char_count"] == 1500
        assert finding.details["threshold"] == 1000
        assert finding.details["char_count"] > finding.details["threshold"]


# ---------------------------------------------------------------------------
# placeholder-text
# ---------------------------------------------------------------------------


class DescribePlaceholderText:

    @pytest.mark.parametrize(
        "snippet",
        [
            "[PLACEHOLDER]",
            "[TBD]",
            "Lorem ipsum dolor sit amet",
            "TODO: complete this section.",
            "TODO - finish later",
            "todo: lowercase still flagged",
            "FIXME this assumption",
            "fixme: lowercase",
            "XXX needs review",
            "TKTK to be filled",
            "Replace this text: <replace me>",
            "Edit here <your text here> please",
            "<insert name> goes here",
            "Hand off [FILL IN] before publish",
            "[ FILL ME ]",
        ],
    )
    def it_flags_a_known_placeholder(
        self, document: DocumentCls, snippet: str
    ):
        document.add_paragraph(f"intro {snippet} outro")
        report = lint(document)
        assert "placeholder-text" in [f.rule for f in report.findings]

    def it_does_not_flag_clean_prose(self, document: DocumentCls):
        document.add_paragraph("This is finished prose.")
        assert "placeholder-text" not in [
            f.rule for f in lint(document).findings
        ]

    @pytest.mark.parametrize(
        "snippet",
        [
            # Substrings that *contain* a placeholder token but are not
            # the token itself: the word-boundary anchors must keep
            # these silent.
            "TODOLIST is a brand of organisers.",
            "The xxxxx anonymisation token shows up here.",
            "MetaTKTK product name with no marker convention.",
            "We will <replace> the value.",  # not "<replace me>"
        ],
    )
    def it_does_not_flag_substrings_that_only_look_like_placeholders(
        self, document: DocumentCls, snippet: str
    ):
        document.add_paragraph(snippet)
        assert "placeholder-text" not in [
            f.rule for f in lint(document).findings
        ]

    @pytest.mark.parametrize(
        "snippet,expected_placeholder,expected_category",
        [
            ("Pick this up TODO: later", "TODO:", "todo-marker"),
            ("Stub Lorem ipsum block", "Lorem ipsum", "lorem-ipsum"),
            ("Marker [TBD] inline", "[TBD]", "bracket-token"),
            ("Author note <replace me>", "<replace me>", "angle-bracket"),
            ("Editor TKTK signal", "TKTK", "to-come"),
        ],
    )
    def it_exposes_placeholder_and_category_in_details(
        self,
        document: DocumentCls,
        snippet: str,
        expected_placeholder: str,
        expected_category: str,
    ):
        # Issue #681: callers should not have to regex-parse the
        # message to recover the matched placeholder. The category tag
        # lets a UI bin findings (e.g. "all 3 todo-markers").
        document.add_paragraph(snippet)
        report = lint(document)
        finding = next(
            f for f in report.findings if f.rule == "placeholder-text"
        )
        assert finding.details["placeholder"] == expected_placeholder
        assert finding.details["category"] == expected_category


# ---------------------------------------------------------------------------
# table-without-header-row
# ---------------------------------------------------------------------------


class DescribeTableWithoutHeaderRow:

    def it_flags_a_table_whose_first_row_is_not_a_header(
        self, document: DocumentCls
    ):
        tbl = document.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Name"
        tbl.rows[0].cells[1].text = "Value"
        report = lint(document)
        twh = [
            f for f in report.findings if f.rule == "table-without-header-row"
        ]
        assert twh
        assert twh[0].severity == "warning"
        assert twh[0].location == "table 0"
        assert twh[0].paragraph_index is None
        assert twh[0].autofix_available is False

    def it_does_not_flag_a_table_with_a_header_row(
        self, document: DocumentCls
    ):
        tbl = document.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Name"
        tbl.rows[0].cells[1].text = "Value"
        # Mark the first row as a header (sets <w:trPr>/<w:tblHeader/>).
        tbl.rows[0].is_header = True
        report = lint(document)
        assert "table-without-header-row" not in [
            f.rule for f in report.findings
        ]

    def it_emits_one_finding_per_table_without_header(
        self, document: DocumentCls
    ):
        # Three tables: first marked as header, other two not.
        tbl_a = document.add_table(rows=1, cols=1)
        tbl_a.rows[0].is_header = True
        document.add_paragraph("between tables")
        document.add_table(rows=2, cols=2)
        document.add_paragraph("between tables")
        document.add_table(rows=1, cols=2)
        report = lint(document)
        twh = [
            f for f in report.findings if f.rule == "table-without-header-row"
        ]
        assert len(twh) == 2
        # Each finding should locate the right table by index.
        locations = sorted(f.location for f in twh)
        assert locations == ["table 1", "table 2"]

    def it_does_not_flag_when_document_has_no_tables(
        self, document: DocumentCls
    ):
        document.add_paragraph("just prose")
        report = lint(document)
        assert "table-without-header-row" not in [
            f.rule for f in report.findings
        ]

    def it_marks_findings_with_autofix_unavailable(
        self, document: DocumentCls
    ):
        document.add_table(rows=2, cols=2)
        report = lint(document)
        twh = next(
            f for f in report.findings if f.rule == "table-without-header-row"
        )
        assert twh.autofix_available is False
        assert twh.autofix_description is None


# ---------------------------------------------------------------------------
# bare-url
# ---------------------------------------------------------------------------


class DescribeBareUrl:

    def it_flags_a_bare_https_url(self, document: DocumentCls):
        document.add_paragraph("See https://example.com for details.")
        report = lint(document)
        bu = [f for f in report.findings if f.rule == "bare-url"]
        assert bu
        assert bu[0].severity == "info"
        assert bu[0].autofix_available is False
        assert "https://example.com" in bu[0].message

    def it_flags_a_bare_http_url(self, document: DocumentCls):
        document.add_paragraph("http://example.org/path")
        report = lint(document)
        assert "bare-url" in [f.rule for f in lint(document).findings]
        bu = [f for f in report.findings if f.rule == "bare-url"]
        assert "http://example.org/path" in bu[0].message

    def it_flags_a_bare_www_url(self, document: DocumentCls):
        document.add_paragraph("Visit www.example.com today")
        report = lint(document)
        assert "bare-url" in [f.rule for f in report.findings]

    def it_strips_trailing_sentence_punctuation_from_the_url(
        self, document: DocumentCls
    ):
        document.add_paragraph("Visit https://example.com.")
        report = lint(document)
        bu = [f for f in report.findings if f.rule == "bare-url"]
        assert bu
        # Trailing period should not be considered part of the URL.
        assert "'https://example.com'" in bu[0].message

    def it_emits_one_finding_per_bare_url_in_a_paragraph(
        self, document: DocumentCls
    ):
        document.add_paragraph(
            "First https://a.example.com and second https://b.example.com."
        )
        report = lint(document)
        bu = [f for f in report.findings if f.rule == "bare-url"]
        assert len(bu) == 2

    def it_does_not_flag_a_paragraph_without_any_urls(
        self, document: DocumentCls
    ):
        document.add_paragraph("Plain prose with no link at all.")
        report = lint(document)
        assert "bare-url" not in [f.rule for f in report.findings]

    def it_does_not_flag_a_url_already_wrapped_in_a_hyperlink(
        self, document: DocumentCls
    ):
        # add_paragraph_with_text auto-detects URLs and wraps them in a
        # <w:hyperlink>, so use that helper to construct the wrapped case.
        para = document.add_paragraph()
        para.add_run("See ")
        para.add_hyperlink(
            url="https://example.com",
            text="https://example.com",
            style=None,
        )
        para.add_run(" for details.")
        report = lint(document)
        assert "bare-url" not in [f.rule for f in report.findings]

    def it_marks_severity_info_and_autofix_unavailable(
        self, document: DocumentCls
    ):
        document.add_paragraph("See https://example.com for details.")
        report = lint(document)
        bu = next(f for f in report.findings if f.rule == "bare-url")
        assert bu.severity == "info"
        assert bu.autofix_available is False
        assert bu.autofix_description is None

    def it_records_the_paragraph_index(self, document: DocumentCls):
        document.add_paragraph("intro")
        document.add_paragraph("See https://example.com for details.")
        report = lint(document)
        bu = next(f for f in report.findings if f.rule == "bare-url")
        assert bu.paragraph_index == 1


# ---------------------------------------------------------------------------
# excessive-font-size-variation
# ---------------------------------------------------------------------------


class DescribeExcessiveFontSizeVariation:

    def _add_sized_paragraph(
        self,
        document: DocumentCls,
        text: str,
        size_pt: int,
        style: str | None = None,
    ):
        para = document.add_paragraph(style=style) if style else document.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(size_pt)
        return para

    def it_flags_when_more_than_four_distinct_sizes_appear_in_body(
        self, document: DocumentCls
    ):
        for i, sz in enumerate((11, 12, 13, 14, 15)):
            self._add_sized_paragraph(document, f"p{i}", sz)
        report = lint(document)
        findings = [
            f
            for f in report.findings
            if f.rule == "excessive-font-size-variation"
        ]
        assert len(findings) == 1
        f = findings[0]
        assert f.severity == "info"
        assert f.paragraph_index is None
        assert f.autofix_available is False
        # Sizes should be listed in ascending order in the message.
        assert "11, 12, 13, 14, 15 pt" in f.message
        assert "5 distinct" in f.message

    def it_does_not_flag_when_at_or_below_threshold(
        self, document: DocumentCls
    ):
        for i, sz in enumerate((11, 12, 13, 14)):
            self._add_sized_paragraph(document, f"p{i}", sz)
        report = lint(document)
        assert "excessive-font-size-variation" not in [
            f.rule for f in report.findings
        ]

    def it_ignores_runs_with_no_explicit_size(
        self, document: DocumentCls
    ):
        # Five paragraphs but only two carry an explicit size; the rest
        # inherit from the default style and must NOT count toward drift.
        self._add_sized_paragraph(document, "a", 12)
        self._add_sized_paragraph(document, "b", 13)
        for word in ("c", "d", "e"):
            document.add_paragraph(word)
        report = lint(document)
        assert "excessive-font-size-variation" not in [
            f.rule for f in report.findings
        ]

    def it_skips_heading_paragraphs(self, document: DocumentCls):
        # Headings carry intentionally different sizes; they shouldn't
        # be counted toward body drift. Body sizes alone stay at four,
        # so no finding should fire even though five distinct sizes
        # exist in the document overall.
        for i, sz in enumerate((11, 12, 13, 14)):
            self._add_sized_paragraph(document, f"body{i}", sz)
        # A heading sized 28pt — would push the count to 5 if not skipped.
        self._add_sized_paragraph(
            document, "Section Title", 28, style="Heading 1"
        )
        report = lint(document)
        assert "excessive-font-size-variation" not in [
            f.rule for f in report.findings
        ]

    def it_deduplicates_repeated_sizes_across_paragraphs(
        self, document: DocumentCls
    ):
        # The same size appearing in many paragraphs must count once.
        for i in range(20):
            self._add_sized_paragraph(document, f"p{i}", 12)
        for i, sz in enumerate((13, 14)):
            self._add_sized_paragraph(document, f"q{i}", sz)
        report = lint(document)
        assert "excessive-font-size-variation" not in [
            f.rule for f in report.findings
        ]

    def it_emits_a_single_document_level_finding(
        self, document: DocumentCls
    ):
        # Even with many size-mismatched paragraphs, only one
        # document-scoped finding should be emitted.
        for i, sz in enumerate((9, 10, 11, 12, 13, 14, 15)):
            self._add_sized_paragraph(document, f"p{i}", sz)
        report = lint(document)
        findings = [
            f
            for f in report.findings
            if f.rule == "excessive-font-size-variation"
        ]
        assert len(findings) == 1
        assert findings[0].location == "document body"


# ---------------------------------------------------------------------------
# Report aggregations
# ---------------------------------------------------------------------------


class DescribeLintReport:

    def it_supports_iteration_and_len(self, document: DocumentCls):
        document.add_paragraph("three   spaces")
        report = lint(document)
        assert len(report) == len(report.findings) == len(list(report))

    def it_summary_lists_per_rule_counts(self, document: DocumentCls):
        document.add_paragraph("three   spaces")
        document.add_paragraph("trailing ")
        report = lint(document)
        out = report.summary()
        assert "multiple-spaces" in out
        assert "trailing-whitespace" in out
        assert "findings" in out

    def it_summary_handles_a_clean_document(self, document: DocumentCls):
        document.core_properties.title = "x"  # so missing-doc-title silent
        # Need to make sure no findings from other rules
        for rule in list(BUILTIN_RULES):
            if rule != "missing-document-title":
                # Disable rules not exercised so the summary is empty
                pass
        # Easier approach: build a totally clean doc — just title set
        report = lint(document)
        # Should at least include the totals line
        assert "findings" in report.summary()

    def it_autofix_applies_every_available_fix_when_called_with_no_args(
        self, document: DocumentCls
    ):
        document.add_paragraph("hello   world ")  # multi-space + trailing
        report = lint(document)
        applied = report.autofix()
        # multi-space + trailing-whitespace == 2 fixes
        assert applied >= 2

    def it_autofix_filters_by_rule(self, document: DocumentCls):
        document.add_paragraph("hello   world ")
        report = lint(document)
        report.autofix(rules=["multiple-spaces"])
        # multi-space gone
        assert "  " not in document.paragraphs[0].text
        # trailing space remains because we didn't ask to fix it
        assert document.paragraphs[0].text.endswith(" ")

    def it_autofix_returns_zero_for_a_clean_document(
        self, document: DocumentCls
    ):
        document.add_paragraph("clean text")
        document.core_properties.title = "x"
        report = lint(document)
        # No findings carry autofix_available=True for a clean doc.
        assert report.autofix() == 0

    def it_autofix_breakdown_returns_per_rule_counts(
        self, document: DocumentCls
    ):
        # Issue #679 — callers want to know which rules contributed,
        # not just a single aggregate.
        document.add_paragraph("hello   world ")  # 1x multi-space + 1x trailing
        document.add_paragraph("foo   bar")  # 1x multi-space
        report = lint(document)
        breakdown = report.autofix_breakdown()
        assert breakdown.get("multiple-spaces") == 2
        assert breakdown.get("trailing-whitespace") == 1

    def it_autofix_breakdown_omits_rules_with_zero_successes(
        self, document: DocumentCls
    ):
        document.add_paragraph("clean text")
        document.core_properties.title = "x"
        report = lint(document)
        # No findings -> empty mapping (not {rule: 0}).
        assert report.autofix_breakdown() == {}

    def it_autofix_breakdown_filters_by_rule(self, document: DocumentCls):
        document.add_paragraph("hello   world ")
        report = lint(document)
        breakdown = report.autofix_breakdown(rules=["multiple-spaces"])
        assert "multiple-spaces" in breakdown
        assert "trailing-whitespace" not in breakdown

    def it_autofix_aggregate_matches_breakdown_total(
        self, document: DocumentCls
    ):
        # Sanity: sum of breakdown values equals the aggregate count
        # autofix() reports. Same selection semantics, same successes.
        document.add_paragraph("a   b ")
        document.add_paragraph("c   d")

        # Snapshot per-rule view first…
        report = lint(document)
        breakdown = report.autofix_breakdown()

        # …then on a fresh doc/report run autofix() and compare totals.
        # (autofix mutates, so the second pass needs a new document.)
        from docx import Document

        doc2 = Document()
        doc2.add_paragraph("a   b ")
        doc2.add_paragraph("c   d")
        report2 = lint(doc2)
        aggregate = report2.autofix()
        assert aggregate == sum(breakdown.values())


# ---------------------------------------------------------------------------
# register_rule / unregister_rule
# ---------------------------------------------------------------------------


class DescribeRegisterRule:

    def it_registers_a_custom_rule(self, document: DocumentCls):
        def check(doc):
            for i, p in enumerate(doc.paragraphs):
                if "FOOBAR" in p.text:
                    yield Finding(
                        rule="no-foobar",
                        severity="error",
                        message="contains FOOBAR",
                        paragraph_index=i,
                    )

        register_rule("no-foobar", check)
        document.add_paragraph("contains FOOBAR here")
        report = lint(document)
        assert "no-foobar" in [f.rule for f in report.findings]

    def it_invokes_the_custom_autofix_callback(
        self, document: DocumentCls
    ):
        def check(doc):
            for i, p in enumerate(doc.paragraphs):
                if "BAD" in p.text:
                    yield Finding(
                        rule="no-bad",
                        severity="warning",
                        message="contains BAD",
                        paragraph_index=i,
                        autofix_available=True,
                        autofix_description="replace BAD with GOOD",
                    )

        def fix(doc, finding):
            p = doc.paragraphs[finding.paragraph_index]
            p.text = p.text.replace("BAD", "GOOD")
            return True

        register_rule("no-bad", check, fix)
        document.add_paragraph("a BAD line")
        report = lint(document)
        report.autofix(rules=["no-bad"])
        assert "GOOD" in document.paragraphs[0].text
        assert "BAD" not in document.paragraphs[0].text

    def it_unregisters_a_rule(self):
        def check(doc):
            return []

        register_rule("temp-rule", check)
        assert "temp-rule" in registered_rules()
        assert unregister_rule("temp-rule") is True
        assert "temp-rule" not in registered_rules()
        # Second call returns False — nothing left to remove.
        assert unregister_rule("temp-rule") is False

    def it_returns_a_Rule_dataclass_from_register(
        self, document: DocumentCls
    ):
        rule = register_rule("noop", lambda d: [])
        assert isinstance(rule, Rule)
        assert rule.name == "noop"
        assert rule.autofix is None

    def it_rejects_invalid_inputs(self):
        with pytest.raises(ValueError):
            register_rule("", lambda d: [])
        with pytest.raises(TypeError):
            register_rule("bad", "not-callable")  # type: ignore[arg-type]
        with pytest.raises(TypeError):
            register_rule(
                "bad", lambda d: [], autofix_callback="not-callable"  # type: ignore[arg-type]
            )


# ---------------------------------------------------------------------------
# Smoke-level cross-cutting tests
# ---------------------------------------------------------------------------


class DescribeIntegration:

    def it_returns_findings_in_document_order(self, document: DocumentCls):
        document.add_paragraph("first   para")
        document.add_paragraph("second   para")
        document.add_paragraph("third   para")
        report = lint(document)
        ms = [f for f in report.findings if f.rule == "multiple-spaces"]
        indices = [f.paragraph_index for f in ms]
        assert indices == sorted(indices)

    def it_applies_fixes_in_reverse_order_so_indices_stay_valid(
        self, document: DocumentCls
    ):
        # Two consecutive empties at index 1 and 2 — autofix removes
        # the second one, the first stays; indices for any later
        # paragraph-scoped findings must still resolve.
        document.add_paragraph("a   b")  # 0 — multi-space
        document.add_paragraph("")  # 1 — empty (kept)
        document.add_paragraph("")  # 2 — empty (removed)
        document.add_paragraph("c   d")  # 3 — multi-space
        report = lint(document)
        report.autofix()
        texts = [p.text for p in document.paragraphs]
        # Only one consecutive empty remains; multi-spaces collapsed.
        assert texts == ["a b", "", "c d"]


# ---------------------------------------------------------------------------
# LintConfig — tunable thresholds and style exemptions
# ---------------------------------------------------------------------------


class DescribeLintConfig:

    def it_uses_default_thresholds_when_no_config_passed(
        self, document: DocumentCls
    ):
        report = lint(document)
        assert isinstance(report.config, LintConfig)
        assert report.config.over_long_threshold == 1000
        assert report.config.multi_space_minimum == 2

    def it_records_the_active_config_on_the_report(
        self, document: DocumentCls
    ):
        cfg = LintConfig(over_long_threshold=500)
        report = lint(document, config=cfg)
        assert report.config is cfg

    def it_rejects_a_non_LintConfig_argument(self, document: DocumentCls):
        with pytest.raises(TypeError):
            lint(document, config={"over_long_threshold": 500})  # type: ignore[arg-type]

    def it_rejects_a_zero_or_negative_over_long_threshold(self):
        with pytest.raises(ValueError):
            LintConfig(over_long_threshold=0)
        with pytest.raises(ValueError):
            LintConfig(over_long_threshold=-100)

    def it_rejects_a_multi_space_minimum_below_two(self):
        with pytest.raises(ValueError):
            LintConfig(multi_space_minimum=1)
        with pytest.raises(ValueError):
            LintConfig(multi_space_minimum=0)

    def it_coerces_iterable_style_exemptions_to_a_frozenset(self):
        cfg = LintConfig(style_exemptions={"Caption", "Quote"})
        assert isinstance(cfg.style_exemptions, frozenset)
        assert "Caption" in cfg.style_exemptions
        # Lists / tuples are also accepted.
        cfg2 = LintConfig(style_exemptions=["Caption"])
        assert cfg2.style_exemptions == frozenset({"Caption"})

    def it_exposes_default_style_exemptions_module_constant(self):
        # Sanity check: defaults cover the documented style families.
        for family in (
            "List Bullet",
            "List Number",
            "List Paragraph",
            "Caption",
            "Footnote Text",
            "Quote",
        ):
            assert family in DEFAULT_STYLE_EXEMPTIONS


# ---------------------------------------------------------------------------
# over-long-paragraph respects LintConfig
# ---------------------------------------------------------------------------


class DescribeOverLongParagraphConfig:

    def it_respects_a_custom_threshold(self, document: DocumentCls):
        # 600 chars is under default 1000 but over a custom 500.
        document.add_paragraph("x" * 600)
        default_report = lint(document)
        assert "over-long-paragraph" not in [
            f.rule for f in default_report.findings
        ]
        custom = lint(document, config=LintConfig(over_long_threshold=500))
        assert "over-long-paragraph" in [f.rule for f in custom.findings]

    def it_can_tighten_the_threshold_below_default(
        self, document: DocumentCls
    ):
        document.add_paragraph("x" * 200)
        report = lint(document, config=LintConfig(over_long_threshold=100))
        ol = [f for f in report.findings if f.rule == "over-long-paragraph"]
        assert len(ol) == 1
        # Message includes the custom threshold so callers can audit.
        assert "100" in ol[0].message

    def it_exempts_list_style_paragraphs_by_default(
        self, document: DocumentCls
    ):
        # A long bulleted explanation is editorially intentional and
        # should not trip the rule with default settings.
        document.add_paragraph("y" * 1500, style="List Bullet")
        report = lint(document)
        assert "over-long-paragraph" not in [
            f.rule for f in report.findings
        ]

    def it_exempts_list_paragraph_style_by_default(
        self, document: DocumentCls
    ):
        document.add_paragraph("y" * 1500, style="List Paragraph")
        report = lint(document)
        assert "over-long-paragraph" not in [
            f.rule for f in report.findings
        ]

    def it_can_disable_style_exemptions(self, document: DocumentCls):
        document.add_paragraph("y" * 1500, style="List Bullet")
        # With an empty exemption set the rule fires on the list para.
        report = lint(document, config=LintConfig(style_exemptions=set()))
        assert "over-long-paragraph" in [f.rule for f in report.findings]

    def it_can_supply_custom_exemptions(self, document: DocumentCls):
        # 'Normal' is the default body style — long Normal paragraphs
        # would normally trip the rule, but we exempt them here.
        document.add_paragraph("z" * 1500)
        report = lint(
            document, config=LintConfig(style_exemptions={"Normal"})
        )
        assert "over-long-paragraph" not in [
            f.rule for f in report.findings
        ]

    def it_reverts_to_defaults_after_a_configured_call(
        self, document: DocumentCls
    ):
        # Calling lint() with a custom config must not leak into the
        # next call — the second call sees the default threshold.
        document.add_paragraph("x" * 600)
        lint(document, config=LintConfig(over_long_threshold=500))
        report = lint(document)
        assert "over-long-paragraph" not in [
            f.rule for f in report.findings
        ]


# ---------------------------------------------------------------------------
# multiple-spaces respects LintConfig
# ---------------------------------------------------------------------------


class DescribeMultipleSpacesConfig:

    def it_can_raise_the_minimum_to_three(self, document: DocumentCls):
        # Two spaces is below the configured minimum of three.
        document.add_paragraph("a  b")
        report = lint(document, config=LintConfig(multi_space_minimum=3))
        assert "multiple-spaces" not in [f.rule for f in report.findings]
        # Three spaces fires.
        document.add_paragraph("a   b")
        report = lint(document, config=LintConfig(multi_space_minimum=3))
        assert "multiple-spaces" in [f.rule for f in report.findings]


# ---------------------------------------------------------------------------
# trailing-whitespace tolerates structural <w:tab/> and <w:br/> elements
# ---------------------------------------------------------------------------


class DescribeTrailingWhitespaceStructural:

    def it_does_not_flag_a_paragraph_ending_in_a_tab_element(
        self, document: DocumentCls
    ):
        # Author typed `hello`, then hit Tab — that becomes a `<w:tab/>`,
        # not a literal '\t' character. The rule should not fire.
        para = document.add_paragraph()
        run = para.add_run("hello")
        run.add_tab()
        report = lint(document)
        assert "trailing-whitespace" not in [
            f.rule for f in report.findings
        ]

    def it_does_not_flag_a_paragraph_ending_in_a_break_element(
        self, document: DocumentCls
    ):
        para = document.add_paragraph()
        run = para.add_run("hello")
        run.add_break()
        report = lint(document)
        assert "trailing-whitespace" not in [
            f.rule for f in report.findings
        ]

    def it_still_flags_a_literal_trailing_space_before_a_tab_element(
        self, document: DocumentCls
    ):
        # `hello ` (literal trailing space) followed by a tab element
        # is still author-typed trailing whitespace.
        para = document.add_paragraph()
        run = para.add_run("hello ")
        run.add_tab()
        report = lint(document)
        assert "trailing-whitespace" in [f.rule for f in report.findings]

    def it_skips_empty_runs_when_walking_from_the_end(
        self, document: DocumentCls
    ):
        # An empty (formatting-only) run at the end shouldn't disturb
        # the check; the rule should still inspect the prior run.
        para = document.add_paragraph()
        para.add_run("hello ")  # ends in literal space
        para.add_run("")  # empty trailing run
        report = lint(document)
        assert "trailing-whitespace" in [f.rule for f in report.findings]
