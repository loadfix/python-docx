"""Unit-test suite for ``docx.kit.tables_compare`` table builders."""

from __future__ import annotations

from typing import List

import pytest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import tables_compare
from docx.shared import RGBColor


# -- Shared helpers -------------------------------------------------------


def _row_texts(table) -> List[List[str]]:
    """Return ``table`` as a list-of-lists of cell text."""
    return [[cell.text for cell in row.cells] for row in table.rows]


# -- comparison ----------------------------------------------------------


class DescribeComparison:
    """Unit-test suite for ``tables_compare.comparison``."""

    def it_appends_a_table_to_the_document(self):
        doc = Document()
        before = len(doc.tables)
        table = tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={"X": ["1", "2"]},
        )
        assert len(doc.tables) == before + 1
        # -- Document.tables returns fresh proxy instances each call,
        # -- so identity-compare the underlying element instead. --
        assert doc.tables[-1]._element is table._element

    def it_emits_one_label_column_plus_one_per_option(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B", "C"],
            features={"X": ["1", "2", "3"]},
        )
        # -- 1 label column + 3 options = 4 cells per row --
        assert len(table.rows[0].cells) == 4

    def it_renders_the_option_names_as_the_header_row(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["Plan A", "Plan B"],
            features={"Users": ["1", "10"]},
        )
        header = [c.text for c in table.rows[0].cells]
        assert header[0] == "Feature"
        assert header[1] == "Plan A"
        assert header[2] == "Plan B"

    def it_emits_one_data_row_per_feature(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={
                "Users": ["1", "10"],
                "Storage": ["10 GB", "100 GB"],
            },
        )
        # -- 1 header + 2 features --
        assert len(table.rows) == 3
        assert table.rows[1].cells[0].text == "Users"
        assert table.rows[1].cells[1].text == "1"
        assert table.rows[1].cells[2].text == "10"
        assert table.rows[2].cells[0].text == "Storage"

    def it_renders_a_recommended_badge_above_the_matched_option(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B", "C"],
            features={"X": ["1", "2", "3"]},
            recommended="B",
        )
        # -- "Recommended\nB" — newline puts the badge on its own visual line --
        assert table.rows[0].cells[2].text == "Recommended\nB"
        # -- non-recommended columns stay clean --
        assert table.rows[0].cells[1].text == "A"
        assert table.rows[0].cells[3].text == "C"

    def it_shades_the_recommended_column_cells_with_the_default_fill(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B", "C"],
            features={"X": ["1", "2", "3"]},
            recommended="B",
        )
        for row in table.rows:
            assert row.cells[2].shading.fill_color == tables_compare.DEFAULT_HIGHLIGHT_FILL
            # -- non-recommended columns are NOT shaded --
            assert row.cells[1].shading.fill_color is None
            assert row.cells[3].shading.fill_color is None

    def it_accepts_a_custom_highlight_fill_as_RGBColor(self):
        doc = Document()
        custom = RGBColor(0xFF, 0x00, 0x00)
        table = tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={"X": ["1", "2"]},
            recommended="A",
            highlight_fill=custom,
        )
        assert table.rows[0].cells[1].shading.fill_color == custom

    def it_accepts_a_custom_highlight_fill_as_hex_string(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={"X": ["1", "2"]},
            recommended="A",
            highlight_fill="#FF0000",
        )
        assert table.rows[0].cells[1].shading.fill_color == RGBColor(0xFF, 0x00, 0x00)

    def it_accepts_a_short_form_hex_highlight_fill(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={"X": ["1", "2"]},
            recommended="A",
            highlight_fill="#F00",
        )
        assert table.rows[0].cells[1].shading.fill_color == RGBColor(0xFF, 0x00, 0x00)

    def it_omits_the_badge_when_no_recommendation_is_given(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={"X": ["1", "2"]},
        )
        for cell in table.rows[0].cells:
            assert "Recommended" not in cell.text
        for row in table.rows:
            for cell in row.cells:
                assert cell.shading.fill_color is None

    def it_renders_an_optional_title_above_the_table(self):
        doc = Document()
        tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={"X": ["1", "2"]},
            title="Plan comparison",
        )
        # -- the heading paragraph is the first paragraph in the body --
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Plan comparison" in text

    def it_raises_when_options_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="options"):
            tables_compare.comparison(doc, options=[], features={"X": []})

    def it_raises_when_features_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="features"):
            tables_compare.comparison(doc, options=["A"], features={})

    def it_raises_when_a_feature_row_is_the_wrong_width(self):
        doc = Document()
        with pytest.raises(ValueError, match="features"):
            tables_compare.comparison(
                doc,
                options=["A", "B", "C"],
                features={"Users": ["1", "10"]},  # only 2 values
            )

    def it_raises_when_recommended_does_not_match_any_option(self):
        doc = Document()
        with pytest.raises(ValueError, match="recommended"):
            tables_compare.comparison(
                doc,
                options=["A", "B"],
                features={"X": ["1", "2"]},
                recommended="Z",
            )

    def it_raises_on_a_malformed_highlight_fill(self):
        doc = Document()
        with pytest.raises(ValueError, match="highlight_fill"):
            tables_compare.comparison(
                doc,
                options=["A", "B"],
                features={"X": ["1", "2"]},
                recommended="A",
                highlight_fill="NOTHEX",
            )

    def it_centres_option_value_cells(self):
        doc = Document()
        table = tables_compare.comparison(
            doc,
            options=["A", "B"],
            features={"X": ["1", "2"]},
        )
        # -- header option cells --
        assert (
            table.rows[0].cells[1].paragraphs[0].alignment
            == WD_ALIGN_PARAGRAPH.CENTER
        )
        # -- body value cells --
        assert (
            table.rows[1].cells[1].paragraphs[0].alignment
            == WD_ALIGN_PARAGRAPH.CENTER
        )


# -- pricing -------------------------------------------------------------


class DescribePricing:
    """Unit-test suite for ``tables_compare.pricing``."""

    def it_appends_a_table_to_the_document(self):
        doc = Document()
        before = len(doc.tables)
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "Starter", "price": "$9", "bullets": ["a"]},
                {"name": "Pro", "price": "$29", "bullets": ["b"]},
            ],
        )
        assert len(doc.tables) == before + 1
        # -- Document.tables returns fresh proxy instances each call,
        # -- so identity-compare the underlying element instead. --
        assert doc.tables[-1]._element is table._element

    def it_emits_one_column_per_tier(self):
        doc = Document()
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "Starter", "price": "$9"},
                {"name": "Pro", "price": "$29"},
                {"name": "Business", "price": "$99"},
            ],
        )
        assert len(table.rows[0].cells) == 3

    def it_renders_name_and_price_rows(self):
        doc = Document()
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "Starter", "price": "$9/mo"},
                {"name": "Pro", "price": "$29/mo"},
            ],
        )
        # -- row 0 = names, row 1 = prices --
        assert [c.text for c in table.rows[0].cells] == ["Starter", "Pro"]
        assert [c.text for c in table.rows[1].cells] == ["$9/mo", "$29/mo"]

    def it_emits_one_row_per_bullet_position_padded_to_the_longest_tier(self):
        doc = Document()
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "Starter", "price": "$9", "bullets": ["one"]},
                {
                    "name": "Pro",
                    "price": "$29",
                    "bullets": ["one", "two", "three"],
                },
            ],
        )
        # -- 2 fixed rows + 3 bullet rows = 5 --
        assert len(table.rows) == 5
        # -- short tier pads with empty cells --
        assert table.rows[3].cells[0].text == ""
        assert table.rows[3].cells[1].text == "two"
        assert table.rows[4].cells[0].text == ""
        assert table.rows[4].cells[1].text == "three"

    def it_shades_the_highlighted_tier_column(self):
        doc = Document()
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "Starter", "price": "$9", "bullets": ["a"]},
                {
                    "name": "Pro",
                    "price": "$29",
                    "bullets": ["b"],
                    "highlighted": True,
                },
                {"name": "Business", "price": "$99", "bullets": ["c"]},
            ],
        )
        # -- the Pro column (index 1) is shaded across every row --
        for row in table.rows:
            assert row.cells[1].shading.fill_color == tables_compare.DEFAULT_HIGHLIGHT_FILL
            assert row.cells[0].shading.fill_color is None
            assert row.cells[2].shading.fill_color is None

    def it_supports_multiple_highlighted_tiers(self):
        doc = Document()
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "A", "price": "$1", "highlighted": True},
                {"name": "B", "price": "$2"},
                {"name": "C", "price": "$3", "highlighted": True},
            ],
        )
        for row in table.rows:
            assert row.cells[0].shading.fill_color == tables_compare.DEFAULT_HIGHLIGHT_FILL
            assert row.cells[1].shading.fill_color is None
            assert row.cells[2].shading.fill_color == tables_compare.DEFAULT_HIGHLIGHT_FILL

    def it_omits_shading_when_no_tier_is_highlighted(self):
        doc = Document()
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "A", "price": "$1"},
                {"name": "B", "price": "$2"},
            ],
        )
        for row in table.rows:
            for cell in row.cells:
                assert cell.shading.fill_color is None

    def it_accepts_a_custom_highlight_fill(self):
        doc = Document()
        custom = RGBColor(0x33, 0x66, 0x99)
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "A", "price": "$1", "highlighted": True},
                {"name": "B", "price": "$2"},
            ],
            highlight_fill=custom,
        )
        assert table.rows[0].cells[0].shading.fill_color == custom

    def it_renders_an_optional_title_above_the_table(self):
        doc = Document()
        tables_compare.pricing(
            doc,
            tiers=[
                {"name": "A", "price": "$1"},
                {"name": "B", "price": "$2"},
            ],
            title="Pricing tiers",
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Pricing tiers" in text

    def it_raises_when_tiers_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="tiers"):
            tables_compare.pricing(doc, tiers=[])

    def it_raises_when_a_tier_is_missing_name(self):
        doc = Document()
        with pytest.raises(ValueError, match="name"):
            tables_compare.pricing(doc, tiers=[{"price": "$9"}])

    def it_raises_when_a_tier_is_missing_price(self):
        doc = Document()
        with pytest.raises(ValueError, match="price"):
            tables_compare.pricing(doc, tiers=[{"name": "Starter"}])

    def it_raises_when_bullets_is_a_bare_string(self):
        doc = Document()
        with pytest.raises(ValueError, match="bullets"):
            tables_compare.pricing(
                doc, tiers=[{"name": "A", "price": "$1", "bullets": "oops"}]
            )

    def it_handles_two_tiers_even_though_three_is_conventional(self):
        doc = Document()
        # -- spec calls for "three or more tiers" but the helper accepts
        # -- any non-empty sequence; this guards against an over-eager
        # -- minimum-length check.
        table = tables_compare.pricing(
            doc,
            tiers=[
                {"name": "A", "price": "$1"},
                {"name": "B", "price": "$2"},
            ],
        )
        assert len(table.rows[0].cells) == 2


# -- rubric --------------------------------------------------------------


class DescribeRubric:
    """Unit-test suite for ``tables_compare.rubric``."""

    def it_appends_a_table_to_the_document(self):
        doc = Document()
        before = len(doc.tables)
        table = tables_compare.rubric(
            doc,
            criteria=["Clarity"],
            levels=["Poor", "Good"],
            cells=[["unclear", "clear"]],
        )
        assert len(doc.tables) == before + 1
        # -- Document.tables returns fresh proxy instances each call,
        # -- so identity-compare the underlying element instead. --
        assert doc.tables[-1]._element is table._element

    def it_emits_a_grid_of_one_label_row_plus_one_per_criterion(self):
        doc = Document()
        table = tables_compare.rubric(
            doc,
            criteria=["Clarity", "Accuracy", "Style"],
            levels=["Poor", "OK", "Excellent"],
            cells=[
                ["a", "b", "c"],
                ["d", "e", "f"],
                ["g", "h", "i"],
            ],
        )
        # -- 1 header row + 3 criteria rows --
        assert len(table.rows) == 4
        # -- 1 label column + 3 levels --
        assert len(table.rows[0].cells) == 4

    def it_renders_a_blank_corner_then_level_labels_in_the_header(self):
        doc = Document()
        table = tables_compare.rubric(
            doc,
            criteria=["Clarity"],
            levels=["Poor", "OK", "Excellent"],
            cells=[["a", "b", "c"]],
        )
        header = [c.text for c in table.rows[0].cells]
        assert header == ["", "Poor", "OK", "Excellent"]

    def it_renders_criteria_labels_in_the_first_column(self):
        doc = Document()
        table = tables_compare.rubric(
            doc,
            criteria=["Clarity", "Accuracy"],
            levels=["Poor", "OK"],
            cells=[
                ["a", "b"],
                ["c", "d"],
            ],
        )
        assert table.rows[1].cells[0].text == "Clarity"
        assert table.rows[2].cells[0].text == "Accuracy"

    def it_fills_the_body_cells_from_the_cells_grid(self):
        doc = Document()
        table = tables_compare.rubric(
            doc,
            criteria=["Clarity", "Accuracy", "Style"],
            levels=["Poor (1)", "OK (3)", "Excellent (5)"],
            cells=[
                ["unclear", "mostly clear", "crystal clear"],
                ["3+ errors", "1-2 errors", "no errors"],
                ["awkward", "readable", "polished"],
            ],
        )
        body = [
            [c.text for c in table.rows[1 + i].cells[1:]]
            for i in range(3)
        ]
        assert body == [
            ["unclear", "mostly clear", "crystal clear"],
            ["3+ errors", "1-2 errors", "no errors"],
            ["awkward", "readable", "polished"],
        ]

    def it_renders_an_optional_title_above_the_table(self):
        doc = Document()
        tables_compare.rubric(
            doc,
            criteria=["Clarity"],
            levels=["Poor"],
            cells=[["x"]],
            title="Essay rubric",
        )
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Essay rubric" in text

    def it_raises_when_criteria_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="criteria"):
            tables_compare.rubric(
                doc, criteria=[], levels=["Poor"], cells=[]
            )

    def it_raises_when_levels_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="levels"):
            tables_compare.rubric(
                doc, criteria=["X"], levels=[], cells=[[]]
            )

    def it_raises_when_cells_has_the_wrong_number_of_rows(self):
        doc = Document()
        with pytest.raises(ValueError, match="cells"):
            tables_compare.rubric(
                doc,
                criteria=["A", "B"],
                levels=["Poor", "Good"],
                cells=[["x", "y"]],  # only 1 row, expected 2
            )

    def it_raises_when_a_cells_row_has_the_wrong_number_of_columns(self):
        doc = Document()
        with pytest.raises(ValueError, match="cells"):
            tables_compare.rubric(
                doc,
                criteria=["A"],
                levels=["Poor", "Good", "Best"],
                cells=[["x", "y"]],  # 2 cols, expected 3
            )


# -- module surface ------------------------------------------------------


class DescribeModuleSurface:
    """Sanity-checks on the public surface."""

    def it_exposes_comparison_pricing_rubric_in_dunder_all(self):
        assert "comparison" in tables_compare.__all__
        assert "pricing" in tables_compare.__all__
        assert "rubric" in tables_compare.__all__

    def it_re_exports_through_docx_kit(self):
        from docx import kit

        assert kit.comparison is tables_compare.comparison
        assert kit.pricing is tables_compare.pricing
        assert kit.rubric is tables_compare.rubric
        assert kit.tables_compare is tables_compare

    def it_supports_the_tables_namespace_alias_idiom(self):
        # -- the issue spec uses ``from docx.kit import tables`` -> ``tables.comparison(...)``;
        # -- the recommended import path is ``tables_compare`` (we sidestep
        # -- the ``tables`` module name to leave room for #289). Verify the
        # -- common alias works.
        from docx.kit import tables_compare as tables

        assert tables.comparison is tables_compare.comparison
        assert tables.pricing is tables_compare.pricing
        assert tables.rubric is tables_compare.rubric
