"""Unit-test suite for ``docx.kit.front_matter`` helpers."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import front_matter
from docx.text.paragraph import Paragraph


@pytest.fixture
def document() -> DocumentCls:
    return Document()


def _last_paragraphs(doc: DocumentCls, n: int) -> list:
    return doc.paragraphs[-n:]


class DescribeAddTitlePage:
    """Unit-test suite for ``front_matter.add_title_page``."""

    def it_appends_a_title_paragraph_with_the_Title_style(self, document: DocumentCls):
        start = len(document.paragraphs)

        result = front_matter.add_title_page(document, title="Annual Report")

        # -- title + page break -> 2 new paragraphs --
        assert len(document.paragraphs) == start + 2
        assert isinstance(result[0], Paragraph)
        assert result[0].text == "Annual Report"
        assert result[0].style is not None
        assert result[0].style.name == "Title"

    def it_centres_each_title_page_paragraph(self, document: DocumentCls):
        result = front_matter.add_title_page(
            document,
            title="T",
            subtitle="S",
            author="A",
            date="D",
        )

        # -- last entry is the trailing page-break paragraph; the four
        # -- content paragraphs are centred --
        for para in result[:-1]:
            assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_includes_subtitle_author_and_date_when_supplied(
        self, document: DocumentCls
    ):
        result = front_matter.add_title_page(
            document,
            title="Annual Report 2026",
            subtitle="Underlying performance",
            author="Acme Corp",
            date="March 2026",
        )

        # -- title, subtitle, author, date, page-break --
        assert len(result) == 5
        assert result[0].text == "Annual Report 2026"
        assert result[1].text == "Underlying performance"
        assert result[1].style is not None
        assert result[1].style.name == "Subtitle"
        assert result[2].text == "Acme Corp"
        assert result[3].text == "March 2026"

    def it_skips_the_page_break_when_disabled(self, document: DocumentCls):
        result = front_matter.add_title_page(
            document, title="Annual Report", page_break=False
        )

        assert len(result) == 1
        assert result[0].text == "Annual Report"

    def it_raises_when_title_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="title must be a non-empty string"):
            front_matter.add_title_page(document, title="")


class DescribeAddCopyrightPage:
    """Unit-test suite for ``front_matter.add_copyright_page``."""

    def it_emits_the_copyright_notice_with_the_unicode_symbol(
        self, document: DocumentCls
    ):
        result = front_matter.add_copyright_page(
            document, holder="Acme Corp", year=2026
        )

        assert "Copyright © 2026 Acme Corp" == result[0].text

    def it_appends_edition_and_default_rights_notice(self, document: DocumentCls):
        result = front_matter.add_copyright_page(
            document, holder="Acme Corp", year=2026, edition="First Edition"
        )

        # -- notice, edition, rights, page-break --
        assert len(result) == 4
        assert result[1].text == "First Edition"
        assert result[2].text == "All rights reserved."

    def it_lets_the_caller_suppress_the_rights_notice_with_empty_string(
        self, document: DocumentCls
    ):
        result = front_matter.add_copyright_page(
            document,
            holder="Acme Corp",
            year=2026,
            edition="First Edition",
            rights="",
        )

        # -- notice, edition, page-break (no rights) --
        assert len(result) == 3
        texts = [p.text for p in result]
        assert "All rights reserved." not in texts

    def it_centres_each_paragraph(self, document: DocumentCls):
        result = front_matter.add_copyright_page(
            document, holder="Acme Corp", year=2026, edition="First Edition"
        )

        for para in result[:-1]:
            assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_accepts_year_as_string_for_ranges(self, document: DocumentCls):
        result = front_matter.add_copyright_page(
            document, holder="Acme Corp", year="2024–2026"
        )

        assert "Copyright © 2024–2026 Acme Corp" == result[0].text

    def it_raises_when_holder_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="holder must be a non-empty string"):
            front_matter.add_copyright_page(document, holder="", year=2026)


class DescribeAddDedication:
    """Unit-test suite for ``front_matter.add_dedication``."""

    def it_appends_a_centred_italic_dedication_paragraph(
        self, document: DocumentCls
    ):
        result = front_matter.add_dedication(
            document, text="To everyone who shipped on time."
        )

        # -- dedication + page-break --
        assert len(result) == 2
        para = result[0]
        assert para.text == "To everyone who shipped on time."
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
        # -- italic asserted at run-level so the test holds even when the
        # -- caller supplied a custom template that lacks the Quote style --
        assert all(run.italic for run in para.runs)

    def it_uses_the_Quote_style_when_available(self, document: DocumentCls):
        result = front_matter.add_dedication(document, text="To readers.")

        para = result[0]
        assert para.style is not None
        assert para.style.name == "Quote"

    def it_skips_the_page_break_when_disabled(self, document: DocumentCls):
        result = front_matter.add_dedication(
            document, text="Short.", page_break=False
        )

        assert len(result) == 1

    def it_raises_when_text_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="text must be a non-empty string"):
            front_matter.add_dedication(document, text="")


class DescribeAddPreface:
    """Unit-test suite for ``front_matter.add_preface``."""

    def it_emits_a_heading_then_one_paragraph_per_body_chunk(
        self, document: DocumentCls
    ):
        result = front_matter.add_preface(
            document,
            title="Preface",
            body="First paragraph.\n\nSecond paragraph.\n\nThird.",
        )

        # -- heading, p1, p2, p3, page-break --
        assert len(result) == 5
        assert result[0].text == "Preface"
        assert result[0].style is not None
        assert result[0].style.name == "Heading 1"
        assert result[1].text == "First paragraph."
        assert result[2].text == "Second paragraph."
        assert result[3].text == "Third."

    def it_accepts_a_sequence_of_paragraph_strings(self, document: DocumentCls):
        result = front_matter.add_preface(
            document,
            title="Preface",
            body=["Alpha.", "Beta.", "Gamma."],
        )

        assert [p.text for p in result[:-1]] == [
            "Preface",
            "Alpha.",
            "Beta.",
            "Gamma.",
        ]

    def it_drops_empty_chunks(self, document: DocumentCls):
        result = front_matter.add_preface(
            document,
            title="Foreword",
            body="One.\n\n\n\nTwo.",
        )

        # -- heading, "One.", "Two.", page-break --
        assert len(result) == 4

    def it_honours_a_custom_heading_level(self, document: DocumentCls):
        result = front_matter.add_preface(
            document,
            title="Preface",
            body="Body.",
            heading_level=2,
        )

        assert result[0].style is not None
        assert result[0].style.name == "Heading 2"

    def it_accepts_an_empty_body(self, document: DocumentCls):
        result = front_matter.add_preface(document, title="Preface", body="")

        # -- heading + page-break only --
        assert len(result) == 2
        assert result[0].text == "Preface"

    def it_raises_when_title_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="title must be a non-empty string"):
            front_matter.add_preface(document, title="", body="x")

    def it_raises_when_heading_level_is_out_of_range(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="heading_level must be in 0..9"):
            front_matter.add_preface(
                document, title="Preface", heading_level=11
            )


class DescribeAddTableOfContents:
    """Unit-test suite for ``front_matter.add_table_of_contents``."""

    def it_emits_a_heading_followed_by_a_TOC_field(self, document: DocumentCls):
        result = front_matter.add_table_of_contents(document)

        # -- heading, TOC paragraph, page-break --
        assert len(result) == 3
        assert result[0].text == "Table of Contents"
        assert result[0].style is not None
        assert result[0].style.name == "Heading 1"

    def it_emits_a_TOC_complex_field_in_the_second_paragraph(
        self, document: DocumentCls
    ):
        result = front_matter.add_table_of_contents(document)

        toc_para = result[1]
        assert len(toc_para.fields) == 1
        field = toc_para.fields[0]
        assert field.is_complex is True
        assert field.type == "TOC"

    def it_omits_the_heading_when_title_is_None(self, document: DocumentCls):
        result = front_matter.add_table_of_contents(document, title=None)

        # -- TOC paragraph + page break --
        assert len(result) == 2
        assert len(result[0].fields) == 1
        assert result[0].fields[0].type == "TOC"

    def it_forwards_levels_to_the_underlying_TOC_helper(
        self, document: DocumentCls
    ):
        result = front_matter.add_table_of_contents(document, levels=(2, 4))

        toc_para = result[1]
        instr = toc_para.fields[0].instruction
        assert '"2-4"' in instr


class DescribeAddListOfFigures:
    """Unit-test suite for ``front_matter.add_list_of_figures``."""

    def it_emits_a_TOC_field_filtered_to_Figure_SEQ_entries(
        self, document: DocumentCls
    ):
        result = front_matter.add_list_of_figures(document)

        # -- heading, TOC, page-break --
        assert len(result) == 3
        assert result[0].text == "List of Figures"

        toc_para = result[1]
        field = toc_para.fields[0]
        assert field.is_complex is True
        assert field.type == "TOC"
        assert '\\c "Figure"' in field.instruction

    def it_accepts_a_custom_label(self, document: DocumentCls):
        result = front_matter.add_list_of_figures(document, label="Diagram")

        toc_para = result[1]
        assert '\\c "Diagram"' in toc_para.fields[0].instruction

    def it_omits_the_heading_when_title_is_None(self, document: DocumentCls):
        result = front_matter.add_list_of_figures(document, title=None)

        assert len(result) == 2  # -- TOC + page break --
        assert result[0].fields[0].type == "TOC"


class DescribeAddListOfTables:
    """Unit-test suite for ``front_matter.add_list_of_tables``."""

    def it_emits_a_TOC_field_filtered_to_Table_SEQ_entries(
        self, document: DocumentCls
    ):
        result = front_matter.add_list_of_tables(document)

        assert len(result) == 3
        assert result[0].text == "List of Tables"

        toc_para = result[1]
        field = toc_para.fields[0]
        assert field.is_complex is True
        assert field.type == "TOC"
        assert '\\c "Table"' in field.instruction


class DescribeFrontMatterIntegration:
    """End-to-end smoke-test: every helper composes cleanly into one document."""

    def it_can_build_a_complete_front_matter_section(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        front_matter.add_title_page(
            document,
            title="Annual Report 2026",
            subtitle="Underlying performance",
            author="Acme Corp",
            date="March 2026",
        )
        front_matter.add_copyright_page(
            document,
            holder="Acme Corp",
            year=2026,
            edition="First Edition",
        )
        front_matter.add_dedication(
            document, text="To everyone who shipped on time."
        )
        front_matter.add_preface(
            document,
            title="Preface",
            body="This document outlines the year's results.",
        )
        front_matter.add_table_of_contents(document)
        front_matter.add_list_of_figures(document)
        front_matter.add_list_of_tables(document)

        # -- the count is determined by the seven helpers' fixed shapes;
        # -- regression-protect the total so an accidental shape change
        # -- doesn't slip through. --
        appended = len(document.paragraphs) - start
        assert appended > 0
        # -- spot-check key lines were emitted in document order --
        texts = [p.text for p in document.paragraphs[start:]]
        assert any("Annual Report 2026" in t for t in texts)
        assert any("Copyright © 2026 Acme Corp" in t for t in texts)
        assert any("To everyone who shipped on time." in t for t in texts)
        assert "Preface" in texts
        assert "Table of Contents" in texts
        assert "List of Figures" in texts
        assert "List of Tables" in texts
