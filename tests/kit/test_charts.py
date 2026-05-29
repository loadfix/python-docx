"""Unit-test suite for ``docx.kit.charts`` helpers (issue #291)."""

from __future__ import annotations

import pytest

# -- Skip the entire module when matplotlib is not installed.  The
# -- helpers themselves are import-safe (they defer the import to call
# -- time) but every test here exercises the call path. --
matplotlib = pytest.importorskip("matplotlib")
matplotlib.use("Agg", force=False)  # headless render
import matplotlib.pyplot as plt  # noqa: E402

from docx import Document  # noqa: E402
from docx.document import Document as DocumentCls  # noqa: E402
from docx.kit import charts  # noqa: E402
from docx.kit.charts import (  # noqa: E402
    add_chart,
    bar_chart,
    line_chart,
    pie_chart,
)
from docx.shape import InlineShape  # noqa: E402
from docx.shared import Inches  # noqa: E402


@pytest.fixture
def document() -> DocumentCls:
    return Document()


@pytest.fixture
def simple_figure():
    fig, ax = plt.subplots()
    ax.plot([1, 2, 3, 4], [10, 20, 15, 25])
    ax.set_title("Demo")
    yield fig
    plt.close(fig)


class DescribeAddChart:
    """Behavioural tests for :func:`docx.kit.charts.add_chart`."""

    def it_returns_an_inline_shape(
        self, document: DocumentCls, simple_figure
    ):
        result = add_chart(document, simple_figure)

        assert isinstance(result, InlineShape)

    def it_appends_a_paragraph_holding_the_picture(
        self, document: DocumentCls, simple_figure
    ):
        start = len(document.paragraphs)

        add_chart(document, simple_figure)

        # -- ``Document.add_picture`` always lays the image in its own
        # -- newly-appended paragraph at the end of the body. --
        assert len(document.paragraphs) == start + 1

    def it_sizes_the_picture_to_the_requested_width_in_inches(
        self, document: DocumentCls, simple_figure
    ):
        shape = add_chart(document, simple_figure, width_in=4.5)

        # -- Within 1 EMU of the requested width.  EMU rounding in the
        # -- conversion is harmless. --
        assert abs(int(shape.width) - int(Inches(4.5))) <= 1

    def it_writes_the_alt_text_to_the_inline_shape(
        self, document: DocumentCls, simple_figure
    ):
        shape = add_chart(
            document, simple_figure, alt_text="Q3 revenue chart"
        )

        assert shape.alt_text == "Q3 revenue chart"

    def it_omits_alt_text_when_not_supplied(
        self, document: DocumentCls, simple_figure
    ):
        shape = add_chart(document, simple_figure)

        assert shape.alt_text is None

    def it_applies_brand_colors_to_line_artists(
        self, document: DocumentCls, simple_figure
    ):
        # -- matplotlib normalises hex strings to RGBA tuples; compare
        # -- via ``matplotlib.colors.to_hex`` to avoid colour-space
        # -- mismatches. --
        from matplotlib.colors import to_hex

        add_chart(
            document,
            simple_figure,
            brand_colors=["#1f4e79", "#2e75b6", "#9dc3e6"],
        )

        ax = simple_figure.axes[0]
        line = ax.get_lines()[0]
        assert to_hex(line.get_color()).lower() == "#1f4e79"

    def it_cycles_brand_colors_across_multiple_artists(
        self, document: DocumentCls
    ):
        from matplotlib.colors import to_hex

        fig, ax = plt.subplots()
        try:
            ax.plot([1, 2, 3], [1, 2, 3], label="a")
            ax.plot([1, 2, 3], [3, 2, 1], label="b")
            ax.plot([1, 2, 3], [2, 2, 2], label="c")

            add_chart(
                document, fig, brand_colors=["#1f4e79", "#2e75b6"]
            )

            colors = [
                to_hex(line.get_color()).lower()
                for line in ax.get_lines()
            ]
            assert colors == ["#1f4e79", "#2e75b6", "#1f4e79"]
        finally:
            plt.close(fig)

    def it_preserves_default_colors_when_no_brand_colors_given(
        self, document: DocumentCls
    ):
        from matplotlib.colors import to_hex

        fig, ax = plt.subplots()
        try:
            ax.plot([1, 2, 3], [1, 2, 3])
            default_color = to_hex(ax.get_lines()[0].get_color()).lower()

            add_chart(document, fig)  # no brand_colors

            color_after = to_hex(ax.get_lines()[0].get_color()).lower()
            assert color_after == default_color
        finally:
            plt.close(fig)

    def it_raises_on_a_non_positive_width(
        self, document: DocumentCls, simple_figure
    ):
        with pytest.raises(ValueError, match="width_in must be > 0"):
            add_chart(document, simple_figure, width_in=0)

    def it_raises_on_a_non_positive_dpi(
        self, document: DocumentCls, simple_figure
    ):
        with pytest.raises(ValueError, match="dpi must be > 0"):
            add_chart(document, simple_figure, dpi=0)


class DescribeBarChart:
    """Behavioural tests for :func:`docx.kit.charts.bar_chart`."""

    def it_returns_an_inline_shape(self, document: DocumentCls):
        result = bar_chart(
            document, x=["Q1", "Q2", "Q3", "Q4"], y=[10, 20, 15, 25]
        )

        assert isinstance(result, InlineShape)

    def it_appends_a_picture_paragraph(self, document: DocumentCls):
        start = len(document.paragraphs)

        bar_chart(
            document, x=["Q1", "Q2"], y=[1, 2], title="Revenue"
        )

        assert len(document.paragraphs) == start + 1

    def it_writes_alt_text_when_supplied(self, document: DocumentCls):
        shape = bar_chart(
            document,
            x=["A", "B"],
            y=[1, 2],
            alt_text="Bar chart of A vs B",
        )

        assert shape.alt_text == "Bar chart of A vs B"

    def it_raises_on_mismatched_x_y_lengths(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="equal length"):
            bar_chart(document, x=["A", "B"], y=[1, 2, 3])


class DescribeLineChart:
    """Behavioural tests for :func:`docx.kit.charts.line_chart`."""

    def it_returns_an_inline_shape(self, document: DocumentCls):
        result = line_chart(
            document, x=[1, 2, 3, 4], y=[10, 20, 15, 25]
        )

        assert isinstance(result, InlineShape)

    def it_writes_alt_text_when_supplied(self, document: DocumentCls):
        shape = line_chart(
            document,
            x=[1, 2, 3],
            y=[1, 2, 3],
            title="Trend",
            alt_text="Quarterly trend",
        )

        assert shape.alt_text == "Quarterly trend"

    def it_raises_on_mismatched_x_y_lengths(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="equal length"):
            line_chart(document, x=[1, 2], y=[1, 2, 3])


class DescribePieChart:
    """Behavioural tests for :func:`docx.kit.charts.pie_chart`."""

    def it_returns_an_inline_shape(self, document: DocumentCls):
        result = pie_chart(
            document, labels=["A", "B", "C"], values=[30, 40, 30]
        )

        assert isinstance(result, InlineShape)

    def it_writes_alt_text_when_supplied(self, document: DocumentCls):
        shape = pie_chart(
            document,
            labels=["A", "B"],
            values=[40, 60],
            title="Mix",
            alt_text="Product mix",
        )

        assert shape.alt_text == "Product mix"

    def it_raises_on_mismatched_labels_values_lengths(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="equal length"):
            pie_chart(
                document, labels=["A", "B"], values=[1, 2, 3]
            )


class DescribeKitReExport:
    """The charts module is re-exported from ``docx.kit``."""

    def it_re_exports_the_charts_submodule(self):
        from docx import kit

        assert hasattr(kit, "charts")
        assert kit.charts is charts
        assert "charts" in kit.__all__
