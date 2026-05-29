"""Unit-test suite for ``docx.kit.coe`` Correction of Error helper."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import coe as coe_module
from docx.kit.coe import coe
from docx.table import Table
from docx.text.paragraph import Paragraph


# -- Fixtures -------------------------------------------------------------


@pytest.fixture
def document() -> DocumentCls:
    return Document()


def _kwargs(**overrides):
    """Return a complete keyword-argument bundle with ``overrides`` applied."""
    base = dict(
        title="DB-2026-05-29 — primary failover delay",
        incident_date="2026-05-29",
        severity="Sev2",
        duration="47 minutes",
        customer_impact="20% of users saw 5xx errors during the failover window.",
        summary="One-paragraph summary of what happened.",
        timeline=[
            ("14:32 UTC", "Heartbeat alert fires"),
            ("14:34 UTC", "On-call paged"),
            ("14:42 UTC", "Failover initiated"),
            ("14:55 UTC", "Failover failed; rollback initiated"),
            ("15:19 UTC", "Service restored"),
        ],
        five_whys=[
            ("Why did the service fail?", "The primary replica fell behind."),
            ("Why did the primary fail?", "Disk hit 100% util."),
            ("Why did the disk fill?", "A rogue analytics query backfilled to it."),
            ("Why did the query run on primary?",
             "Routing rule misconfigured 6 weeks ago."),
            ("Why was it not caught?",
             "We don't alert on routing rule changes."),
        ],
        contributing_factors=[
            "Routing rule misconfigured",
            "Lack of canary on rule changes",
        ],
        action_items=[
            {"item": "Add canary on routing rule changes",
             "owner": "SRE", "due": "2026-06-15"},
            {"item": "Backfill alert on primary disk util",
             "owner": "DBA", "due": "2026-06-08"},
        ],
        lessons_learned=[
            "Always canary routing changes",
            "Alert on every disk reaching > 80% utilisation",
        ],
    )
    base.update(overrides)
    return base


def _full_text(document: DocumentCls) -> str:
    return "\n".join(p.text for p in document.paragraphs)


# -- Module re-export -----------------------------------------------------


class DescribeCoeModuleReexport:
    """The ``docx.kit.coe`` module is re-exported from ``docx.kit``."""

    def it_is_importable_as_a_module_attribute(self):
        from docx import kit

        assert kit.coe is coe_module

    def it_exposes_the_coe_function_on_the_module(self):
        assert callable(coe_module.coe)
        assert coe_module.__all__ == ["coe"]


# -- Happy path -----------------------------------------------------------


class DescribeCoe:
    """Unit-test suite for ``coe.coe``."""

    def it_appends_the_title_as_the_first_new_paragraph(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        assert isinstance(result[0], Paragraph)
        assert result[0].text == (
            "DB-2026-05-29 — primary failover delay"
        )

    def it_renders_the_title_in_the_Title_style_when_available(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        assert result[0].style is not None
        assert result[0].style.name == "Title"

    def it_renders_the_metadata_block_with_bold_labels(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        # -- result[1..4] are the four metadata paragraphs --
        labels = [
            ("Date", "2026-05-29"),
            ("Severity", "Sev2"),
            ("Duration", "47 minutes"),
            ("Customer Impact",
             "20% of users saw 5xx errors during the failover window."),
        ]
        for offset, (label, value) in enumerate(labels):
            para = result[1 + offset]
            assert para.runs[0].text == f"{label}: "
            assert para.runs[0].bold is True
            assert para.runs[1].text == value

    def it_emits_a_Summary_heading_followed_by_the_summary_paragraph(
        self, document: DocumentCls
    ):
        coe(document, **_kwargs(summary="The DB fell over."))

        text = _full_text(document)
        assert "Summary" in text
        assert "The DB fell over." in text

    def it_renders_the_timeline_as_a_two_column_table(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        tables = [r for r in result if isinstance(r, Table)]
        timeline_table = tables[0]
        assert len(timeline_table.columns) == 2
        # -- 1 header + 5 timeline rows --
        assert len(timeline_table.rows) == 6
        header_cells = timeline_table.rows[0].cells
        assert header_cells[0].text == "Time"
        assert header_cells[1].text == "Event"
        body_cells = timeline_table.rows[1].cells
        assert body_cells[0].text == "14:32 UTC"
        assert body_cells[1].text == "Heartbeat alert fires"

    def it_renders_the_five_whys_as_a_two_column_table(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        tables = [r for r in result if isinstance(r, Table)]
        five_whys_table = tables[1]
        assert len(five_whys_table.columns) == 2
        # -- 1 header + 5 question rows --
        assert len(five_whys_table.rows) == 6
        header_cells = five_whys_table.rows[0].cells
        assert header_cells[0].text == "Question"
        assert header_cells[1].text == "Answer"
        body_cells = five_whys_table.rows[1].cells
        assert body_cells[0].text == "Why did the service fail?"
        assert body_cells[1].text == "The primary replica fell behind."

    def it_renders_contributing_factors_as_a_bulleted_list(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        bullet_paras = [
            r for r in result
            if isinstance(r, Paragraph)
            and r.style is not None
            and r.style.name == "List Bullet"
        ]
        # -- 2 contributing factors + 2 lessons learned --
        assert len(bullet_paras) == 4
        bullet_texts = [p.text for p in bullet_paras]
        assert "Routing rule misconfigured" in bullet_texts
        assert "Lack of canary on rule changes" in bullet_texts

    def it_renders_action_items_as_a_three_column_table(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        tables = [r for r in result if isinstance(r, Table)]
        action_table = tables[2]
        assert len(action_table.columns) == 3
        # -- 1 header + 2 action items --
        assert len(action_table.rows) == 3
        header_cells = action_table.rows[0].cells
        assert header_cells[0].text == "Action Item"
        assert header_cells[1].text == "Owner"
        assert header_cells[2].text == "Due"
        body_cells = action_table.rows[1].cells
        assert body_cells[0].text == "Add canary on routing rule changes"
        assert body_cells[1].text == "SRE"
        assert body_cells[2].text == "2026-06-15"

    def it_renders_lessons_learned_as_a_bulleted_list(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        # -- the last two bullet-styled paragraphs (after action items) --
        bullet_texts = [
            r.text for r in result
            if isinstance(r, Paragraph)
            and r.style is not None
            and r.style.name == "List Bullet"
        ]
        assert "Always canary routing changes" in bullet_texts
        assert "Alert on every disk reaching > 80% utilisation" in bullet_texts

    def it_emits_every_required_section_heading_in_order(
        self, document: DocumentCls
    ):
        coe(document, **_kwargs())

        text = _full_text(document)
        expected_headings = [
            "Summary",
            "Timeline",
            "Five Whys",
            "Contributing Factors",
            "Action Items",
            "Lessons Learned",
        ]
        positions = [text.index(h) for h in expected_headings]
        assert positions == sorted(positions), (
            "section headings must appear in document order"
        )

    def it_appends_a_trailing_page_break_by_default(
        self, document: DocumentCls
    ):
        result = coe(document, **_kwargs())

        # -- last entry of the returned list is the page-break paragraph --
        assert isinstance(result[-1], Paragraph)

    def it_skips_the_trailing_page_break_when_disabled(
        self, document: DocumentCls
    ):
        result_with_break = coe(Document(), **_kwargs())
        result_without_break = coe(
            document, **_kwargs(page_break=False)
        )

        assert len(result_without_break) == len(result_with_break) - 1

    def it_appends_to_existing_content_without_clearing_it(
        self, document: DocumentCls
    ):
        document.add_paragraph("Pre-existing content.")
        prefix_count = len(document.paragraphs)

        coe(document, **_kwargs())

        assert document.paragraphs[0].text == "Pre-existing content."
        assert len(document.paragraphs) > prefix_count


# -- Validation -----------------------------------------------------------


class DescribeCoeValidation:
    """Unit-test suite for ``coe.coe`` argument validation."""

    def it_raises_when_title_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="title is required"):
            coe(document, **_kwargs(title=""))

    def it_raises_when_title_is_whitespace_only(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="title is required"):
            coe(document, **_kwargs(title="   "))

    def it_raises_when_a_timeline_entry_is_not_a_pair(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="timeline\\[0\\]"):
            coe(
                document,
                **_kwargs(timeline=[("only-one-element",)]),
            )

    def it_raises_when_a_five_whys_entry_is_not_a_pair(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="five_whys\\[1\\]"):
            coe(
                document,
                **_kwargs(
                    five_whys=[
                        ("Q1", "A1"),
                        ("Q2", "A2", "extra"),
                    ]
                ),
            )

    def it_raises_when_a_timeline_entry_is_a_bare_string(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="timeline\\[0\\]"):
            coe(
                document,
                **_kwargs(timeline=["not-a-tuple"]),
            )

    def it_raises_when_an_action_item_is_missing_item_key(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="action_items\\[0\\]"):
            coe(
                document,
                **_kwargs(
                    action_items=[
                        {"owner": "SRE", "due": "2026-06-15"},
                    ]
                ),
            )

    def it_raises_when_an_action_item_is_not_a_mapping(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="action_items\\[0\\]"):
            coe(
                document,
                **_kwargs(action_items=["not-a-dict"]),
            )


# -- Style fallback -------------------------------------------------------


class DescribeCoeStyleFallback:
    """Unit-test suite for ``coe.coe`` style fallback behaviour."""

    def it_falls_back_to_Normal_when_Title_style_is_missing(
        self, document: DocumentCls
    ):
        # -- Drop the Title style to force the fallback path --
        document.styles["Title"].delete()

        result = coe(document, **_kwargs())

        assert result[0].style is not None
        assert result[0].style.name == "Normal"

    def it_falls_back_quietly_when_List_Bullet_style_is_missing(
        self, document: DocumentCls
    ):
        document.styles["List Bullet"].delete()

        result = coe(document, **_kwargs())

        # -- bulleted-list paragraphs land as Normal-styled paragraphs --
        normal_bullets = [
            r for r in result
            if isinstance(r, Paragraph)
            and r.text in {
                "Routing rule misconfigured",
                "Lack of canary on rule changes",
                "Always canary routing changes",
                "Alert on every disk reaching > 80% utilisation",
            }
        ]
        assert len(normal_bullets) == 4
        for para in normal_bullets:
            assert para.style is not None
            assert para.style.name == "Normal"
