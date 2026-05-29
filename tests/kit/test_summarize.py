"""Unit-test suite for ``docx.kit.summarize`` (issue #303)."""

from __future__ import annotations

from typing import List

import pytest

from docx import Document
from docx.kit import summarize as summarize_mod
from docx.kit.summarize import (
    as_text,
    count_tokens,
    summarize,
)


# ---------------------------------------------------------------------------
# Helpers


def _doc_with_h1_sections() -> "Document":
    """Three ``Heading 1`` sections of varying length."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This is the introduction. It is short. It sets the scene."
    )
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "We describe our methods here. The methods are quantitative. "
        "We sampled fifty participants. Each participant completed a "
        "survey. The survey took ten minutes."
    )
    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "Our results were striking. We observed a sevenfold increase. "
        "The effect persisted across all subgroups. Statistical "
        "significance was reached at p<0.001. Replication is needed. "
        "We propose three follow-up studies. The implications are "
        "broad. They affect three downstream domains."
    )
    return doc


def _doc_with_h2_only() -> "Document":
    """Two ``Heading 2`` sections, no ``Heading 1``."""
    doc = Document()
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Body text for A. More A content.")
    doc.add_heading("Section B", level=2)
    doc.add_paragraph("Body text for B.")
    return doc


def _doc_without_headings(n_paragraphs: int) -> "Document":
    """A document with ``n_paragraphs`` body paragraphs and no headings."""
    doc = Document()
    for i in range(n_paragraphs):
        doc.add_paragraph(f"Paragraph {i + 1}. " * 3)
    return doc


# ---------------------------------------------------------------------------
# count_tokens


class DescribeCountTokens:
    def it_returns_zero_for_empty_string(self):
        assert count_tokens("") == 0

    def it_uses_a_4_chars_per_token_heuristic(self):
        # -- 16 chars => 4 tokens (16 // 4)
        assert count_tokens("a" * 16) == 4

    def it_rounds_up_for_short_strings(self):
        # -- 5 chars => 2 tokens, not 1
        assert count_tokens("hello") == 2

    def it_returns_at_least_one_for_a_single_character(self):
        assert count_tokens("a") == 1


# ---------------------------------------------------------------------------
# Public API contract


class DescribeSummarizePublicAPI:
    def it_is_exposed_via_docx_kit(self):
        from docx.kit import summarize as via_kit

        assert via_kit is summarize_mod
        assert callable(via_kit.summarize)
        assert callable(via_kit.as_text)
        assert callable(via_kit.count_tokens)

    def it_returns_an_empty_list_for_an_empty_document(self):
        doc = Document()
        assert summarize(doc, max_tokens=100) == []

    def it_returns_an_empty_string_from_as_text_for_an_empty_document(self):
        doc = Document()
        assert as_text(doc, max_tokens=100) == ""

    def it_raises_on_zero_max_tokens(self):
        doc = Document()
        with pytest.raises(ValueError):
            summarize(doc, max_tokens=0)

    def it_raises_on_negative_max_tokens(self):
        doc = Document()
        with pytest.raises(ValueError):
            summarize(doc, max_tokens=-1)

    def it_raises_on_non_int_max_tokens(self):
        doc = Document()
        with pytest.raises(ValueError):
            summarize(doc, max_tokens="500")  # type: ignore[arg-type]

    def it_raises_on_non_positive_chunk_size(self):
        doc = Document()
        with pytest.raises(ValueError):
            summarize(doc, max_tokens=100, chunk_size=0)


# ---------------------------------------------------------------------------
# Section detection


class DescribeSectionDetection:
    def it_splits_by_h1_when_h1_is_present(self):
        doc = _doc_with_h1_sections()
        rows = summarize(doc, max_tokens=500)

        titles = [r["section"] for r in rows]
        assert titles == ["Introduction", "Methods", "Results"]

    def it_falls_back_to_h2_when_no_h1(self):
        doc = _doc_with_h2_only()
        rows = summarize(doc, max_tokens=500)

        titles = [r["section"] for r in rows]
        assert titles == ["Section A", "Section B"]

    def it_falls_back_to_chunks_when_no_headings(self):
        doc = _doc_without_headings(n_paragraphs=45)
        rows = summarize(doc, max_tokens=500, chunk_size=20)

        # -- 45 paragraphs / 20 = three chunks (20 + 20 + 5).
        assert [r["section"] for r in rows] == [
            "Section 1",
            "Section 2",
            "Section 3",
        ]

    def it_routes_through_a_user_supplied_section_predicate(self):
        doc = Document()
        doc.add_paragraph("=== Custom A ===")
        doc.add_paragraph("Body for A.")
        doc.add_paragraph("=== Custom B ===")
        doc.add_paragraph("Body for B.")

        def is_custom_header(p) -> bool:
            return (p.text or "").startswith("=== ")

        rows = summarize(
            doc, max_tokens=500, section_predicate=is_custom_header
        )

        titles = [r["section"] for r in rows]
        assert titles == ["=== Custom A ===", "=== Custom B ==="]

    def it_handles_a_document_that_opens_with_body_then_introduces_headings(self):
        doc = Document()
        doc.add_paragraph("Preamble paragraph.")
        doc.add_heading("First heading", level=1)
        doc.add_paragraph("Body for the first heading.")

        rows = summarize(doc, max_tokens=500)

        # -- The preamble lands in a synthesised "Section 1", the H1
        # -- becomes its own row.
        assert rows[0]["section"] == "Section 1"
        assert rows[1]["section"] == "First heading"


# ---------------------------------------------------------------------------
# Budget allocation


class DescribeBudgetAllocation:
    def it_keeps_total_summary_tokens_at_or_below_max_tokens(self):
        doc = _doc_with_h1_sections()
        budget = 50
        rows = summarize(doc, max_tokens=budget)

        total = sum(int(r["tokens"]) for r in rows)
        assert total <= budget

    def it_allocates_more_tokens_to_longer_sections(self):
        doc = _doc_with_h1_sections()

        # -- Use a passthrough summariser so the returned summary
        # -- equals the full body and tokens reflect the *budget*
        # -- the section was given.
        def passthrough(text: str, max_tokens: int) -> str:
            # -- Truncate at max_tokens * 4 chars so the returned
            # -- string fits the budget under count_tokens.
            return text[: max_tokens * 4]

        rows = summarize(doc, max_tokens=200, summariser=passthrough)
        token_counts = [int(r["tokens"]) for r in rows]

        # -- Methods is longer than Introduction, Results is longer
        # -- than Methods. Budget should reflect that.
        intro, methods, results = token_counts
        assert intro <= methods <= results


# ---------------------------------------------------------------------------
# Default extractive summariser


class DescribeDefaultExtractiveSummariser:
    def it_returns_the_first_sentences_that_fit_in_the_budget(self):
        doc = Document()
        doc.add_heading("Only", level=1)
        doc.add_paragraph(
            "First sentence. Second sentence. Third sentence. "
            "Fourth sentence. Fifth sentence."
        )

        # -- Generous budget — all five sentences should fit.
        rows = summarize(doc, max_tokens=200)
        summary = str(rows[0]["summary"])

        assert "First sentence." in summary
        assert "Second sentence." in summary

    def it_truncates_when_the_first_sentence_is_already_over_budget(self):
        doc = Document()
        doc.add_heading("Only", level=1)
        long_sentence = "x" * 400  # -- 100 tokens at 4 chars each
        doc.add_paragraph(long_sentence)

        rows = summarize(doc, max_tokens=20)
        summary = str(rows[0]["summary"])

        # -- Non-empty (the row kept *something*) but well under the
        # -- raw body length.
        assert summary
        assert len(summary) <= 100


# ---------------------------------------------------------------------------
# Custom summariser injection


class DescribeCustomSummariserInjection:
    def it_passes_each_section_through_the_summariser_callback(self):
        doc = _doc_with_h1_sections()

        calls: List[tuple] = []

        def fake_summariser(text: str, max_tokens: int) -> str:
            calls.append((text[:30], max_tokens))
            return f"FAKE({max_tokens})"

        rows = summarize(doc, max_tokens=300, summariser=fake_summariser)

        # -- One call per non-empty section.
        assert len(calls) == 3
        # -- Each row's summary is the fake summariser's return.
        for row in rows:
            assert str(row["summary"]).startswith("FAKE(")

    def it_records_the_token_count_of_the_summarisers_output(self):
        doc = _doc_with_h1_sections()

        def fixed(text: str, max_tokens: int) -> str:
            return "abcdefgh"  # -- 8 chars => 2 tokens

        rows = summarize(doc, max_tokens=300, summariser=fixed)

        assert all(int(r["tokens"]) == 2 for r in rows)

    def it_treats_a_summariser_returning_None_as_an_empty_summary(self):
        doc = _doc_with_h1_sections()

        def returns_none(text: str, max_tokens: int):
            return None  # type: ignore[return-value]

        rows = summarize(doc, max_tokens=300, summariser=returns_none)

        assert all(r["summary"] == "" for r in rows)
        assert all(int(r["tokens"]) == 0 for r in rows)


# ---------------------------------------------------------------------------
# Custom token counter injection


class DescribeCustomTokenCounterInjection:
    def it_uses_the_caller_supplied_token_counter(self):
        doc = _doc_with_h1_sections()

        # -- A token-per-word counter — simulates a real tokenizer.
        def words(text: str) -> int:
            return len(text.split())

        rows = summarize(doc, max_tokens=20, token_counter=words)

        # -- Row tokens equal the word count of the row's summary.
        for row in rows:
            assert int(row["tokens"]) == len(str(row["summary"]).split())


# ---------------------------------------------------------------------------
# as_text


class DescribeAsText:
    def it_concatenates_section_summaries_into_a_flat_string(self):
        doc = _doc_with_h1_sections()
        text = as_text(doc, max_tokens=500)

        assert "Introduction" in text
        assert "Methods" in text
        assert "Results" in text

    def it_emits_section_headings_as_bold_markdown_prefixes(self):
        doc = _doc_with_h1_sections()
        text = as_text(doc, max_tokens=500)

        # -- The first row's heading is rendered as **Introduction**
        # -- followed by the body on the next line.
        assert "**Introduction**" in text
        assert "**Methods**" in text
        assert "**Results**" in text

    def it_omits_sections_with_empty_summaries(self):
        doc = Document()
        doc.add_heading("Empty section", level=1)
        doc.add_heading("Real section", level=1)
        doc.add_paragraph("Some real content here. Two sentences.")

        text = as_text(doc, max_tokens=200)

        # -- The "Empty section" body is empty so it's omitted from
        # -- the flat-text output.
        assert "Empty section" not in text
        assert "**Real section**" in text

    def it_passes_summariser_through_to_summarize(self):
        doc = _doc_with_h1_sections()

        def fixed(text: str, max_tokens: int) -> str:
            return "FIXED"

        text = as_text(doc, max_tokens=200, summariser=fixed)

        # -- The "FIXED" string from the custom summariser appears in
        # -- the flat-text output for every section.
        assert text.count("FIXED") == 3


# ---------------------------------------------------------------------------
# Round-trip on an in-memory document


class DescribeEndToEnd:
    def it_summarises_a_realistic_doc_within_budget(self):
        doc = _doc_with_h1_sections()
        rows = summarize(doc, max_tokens=80)

        # -- Three rows, all under the requested budget, with a
        # -- non-empty summary on at least one of them.
        assert len(rows) == 3
        total = sum(int(r["tokens"]) for r in rows)
        assert total <= 80
        assert any(r["summary"] for r in rows)

    def it_returns_dict_keys_in_the_documented_shape(self):
        doc = _doc_with_h1_sections()
        rows = summarize(doc, max_tokens=200)

        for row in rows:
            assert set(row.keys()) == {"section", "summary", "tokens"}
            assert isinstance(row["section"], str)
            assert isinstance(row["summary"], str)
            assert isinstance(row["tokens"], int)
