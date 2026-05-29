"""Unit-test suite for ``docx.kit.runbook``."""

from __future__ import annotations

from io import BytesIO
from typing import List

import pytest

from docx import Document
from docx.kit import runbook
from docx.table import Table
from docx.text.paragraph import Paragraph


# -- Shared fixtures / helpers --------------------------------------------


def _texts(document) -> List[str]:
    return [p.text for p in document.paragraphs]


def _full_text(document) -> str:
    return "\n".join(_texts(document))


def _minimal_kwargs(**overrides):
    """Return a kwargs dict with the eight required-ish fields populated."""
    base = dict(
        title="Database failover runbook",
        purpose="Recover the production primary after a node failure.",
        when_to_use="When the primary's heartbeat times out for >30s.",
        roles=["On-call SRE", "DBA", "Incident commander"],
        prerequisites=[
            "Pager access",
            "VPN connection",
            "Read access to dashboards",
        ],
        procedure=[
            {"step": "Confirm failure", "owner": "On-call SRE",
             "detail": "Check Grafana dashboard."},
            {"step": "Trigger failover", "owner": "DBA",
             "detail": "Run pg_ctl promote replica-1."},
            {"step": "Update DNS", "owner": "On-call SRE",
             "detail": "Promote replica-1 in Route 53."},
            {"step": "Verify", "owner": "Incident commander",
             "detail": "Run smoke tests."},
        ],
    )
    base.update(overrides)
    return base


# -- Required-section happy path -----------------------------------------


class DescribeRunbookSkeleton:
    """The eight required sections of a runbook all render."""

    def it_renders_the_title_as_a_heading(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())

        assert "Database failover runbook" in _full_text(doc)

    def it_renders_the_purpose_section(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())
        text = _full_text(doc)

        assert "Purpose" in text
        assert "Recover the production primary" in text

    def it_renders_the_when_to_use_section(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())
        text = _full_text(doc)

        assert "When to use" in text
        assert "heartbeat times out" in text

    def it_renders_the_roles_as_a_bullet_list(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())
        text = _full_text(doc)

        assert "Roles" in text
        assert "On-call SRE" in text
        assert "DBA" in text
        assert "Incident commander" in text

    def it_renders_the_prerequisites_as_a_checklist(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())
        text = _full_text(doc)

        assert "Prerequisites" in text
        # -- checkbox prefix --
        assert "[ ] Pager access" in text
        assert "[ ] VPN connection" in text
        assert "[ ] Read access to dashboards" in text

    def it_renders_the_procedure_section_heading(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())

        assert "Procedure" in _full_text(doc)


# -- Procedure: numbered list + table ------------------------------------


class DescribeProcedureRendering:
    """The procedure renders both as a numbered list and a 3-column table."""

    def it_renders_each_step_as_a_numbered_list_paragraph(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())
        text = _full_text(doc)

        # -- Each step text appears in the numbered list --
        assert "Confirm failure" in text
        assert "Trigger failover" in text
        assert "Update DNS" in text
        assert "Verify" in text

    def it_renders_the_procedure_as_a_three_column_table(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())

        proc_table = _find_table(doc, "Step")
        assert proc_table is not None

        header = [c.text for c in proc_table.rows[0].cells]
        assert header == ["Step", "Owner", "Detail"]
        # -- header + four steps --
        assert len(proc_table.rows) == 5

    def it_populates_the_owner_column_per_step(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())
        proc_table = _find_table(doc, "Step")

        owners = [proc_table.rows[i].cells[1].text for i in range(1, 5)]
        assert owners == [
            "On-call SRE",
            "DBA",
            "On-call SRE",
            "Incident commander",
        ]

    def it_populates_the_detail_column_per_step(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs())
        proc_table = _find_table(doc, "Step")

        details = [proc_table.rows[i].cells[2].text for i in range(1, 5)]
        assert "Check Grafana dashboard." in details
        assert "Run pg_ctl promote replica-1." in details

    def it_tolerates_a_step_without_owner_or_detail(self):
        doc = Document()
        runbook.runbook(
            doc,
            **_minimal_kwargs(
                procedure=[{"step": "Lone step"}],
            ),
        )

        proc_table = _find_table(doc, "Step")
        assert proc_table is not None
        # -- header + the single step --
        assert len(proc_table.rows) == 2
        first = proc_table.rows[1].cells
        assert first[0].text == "Lone step"
        assert first[1].text == ""
        assert first[2].text == ""


# -- Escalation table ----------------------------------------------------


class DescribeEscalationSection:
    """The optional escalation section renders as a Trigger/Action table."""

    def it_renders_the_escalation_table_when_supplied(self):
        doc = Document()
        runbook.runbook(
            doc,
            **_minimal_kwargs(
                escalation=[
                    ("After 5min", "Page on-call DBA"),
                    ("After 15min", "Page director of engineering"),
                    ("After 30min", "Open incident with vendor"),
                ],
            ),
        )

        text = _full_text(doc)
        assert "Escalation" in text

        esc_table = _find_table(doc, "Trigger")
        assert esc_table is not None
        header = [c.text for c in esc_table.rows[0].cells]
        assert header == ["Trigger", "Action"]
        # -- header + three escalation rows --
        assert len(esc_table.rows) == 4

    def it_preserves_escalation_row_order(self):
        doc = Document()
        runbook.runbook(
            doc,
            **_minimal_kwargs(
                escalation=[
                    ("After 5min", "Page on-call DBA"),
                    ("After 15min", "Page director of engineering"),
                    ("After 30min", "Open incident with vendor"),
                ],
            ),
        )

        esc_table = _find_table(doc, "Trigger")
        triggers = [esc_table.rows[i].cells[0].text for i in range(1, 4)]
        actions = [esc_table.rows[i].cells[1].text for i in range(1, 4)]

        assert triggers == ["After 5min", "After 15min", "After 30min"]
        assert actions[0] == "Page on-call DBA"
        assert actions[2] == "Open incident with vendor"

    def it_omits_the_escalation_section_when_none(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs(escalation=None))

        assert "Escalation" not in _full_text(doc)
        assert _find_table(doc, "Trigger") is None

    def it_omits_the_escalation_section_when_empty_list(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs(escalation=[]))

        assert "Escalation" not in _full_text(doc)


# -- Rollback ------------------------------------------------------------


class DescribeRollbackSection:
    """The optional rollback paragraph renders when supplied."""

    def it_renders_the_rollback_section_when_supplied(self):
        doc = Document()
        runbook.runbook(
            doc,
            **_minimal_kwargs(
                rollback="Revert DNS to old primary, page vendor support.",
            ),
        )
        text = _full_text(doc)

        assert "Rollback" in text
        assert "Revert DNS to old primary" in text

    def it_omits_the_rollback_section_when_none(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs(rollback=None))

        assert "Rollback" not in _full_text(doc)

    def it_omits_the_rollback_section_when_empty_string(self):
        doc = Document()
        runbook.runbook(doc, **_minimal_kwargs(rollback=""))

        assert "Rollback" not in _full_text(doc)


# -- Page break ----------------------------------------------------------


class DescribePageBreakBehaviour:
    """``page_break=True`` (default) appends a trailing page break."""

    def it_appends_a_page_break_by_default(self):
        doc = Document()
        before = len(doc.paragraphs)
        appended = runbook.runbook(doc, **_minimal_kwargs())

        # -- The last appended object should be a page-break Paragraph --
        last = appended[-1]
        assert isinstance(last, Paragraph)
        # -- Page break shows up via run-level break elements; sanity-
        # -- check that the body grew. --
        assert len(doc.paragraphs) > before

    def it_skips_the_page_break_when_disabled(self):
        doc = Document()
        appended = runbook.runbook(
            doc, **_minimal_kwargs(page_break=False)
        )

        # -- Last appended item must NOT be the empty page-break paragraph
        # -- — it should be the rollback or escalation table or proc table. --
        # -- With minimal kwargs and no rollback / escalation, the last
        # -- entry is the procedure table. --
        assert isinstance(appended[-1], Table)


# -- Return value ---------------------------------------------------------


class DescribeReturnValue:
    """``runbook.runbook`` returns paragraphs/tables in document order."""

    def it_returns_a_list_of_paragraphs_and_tables(self):
        doc = Document()
        appended = runbook.runbook(
            doc,
            **_minimal_kwargs(
                escalation=[("After 5min", "Page DBA")],
                rollback="Revert DNS.",
            ),
        )

        # -- Mix of paragraphs and tables --
        assert any(isinstance(item, Paragraph) for item in appended)
        assert any(isinstance(item, Table) for item in appended)

    def it_returns_appended_items_in_document_order(self):
        doc = Document()
        appended = runbook.runbook(doc, **_minimal_kwargs())

        # -- The first appended item is the title heading paragraph --
        first = appended[0]
        assert isinstance(first, Paragraph)
        assert "Database failover runbook" in first.text


# -- Validation errors ----------------------------------------------------


class DescribeValidation:
    """``runbook.runbook`` raises ValueError on invalid input."""

    def it_raises_when_title_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="title is required"):
            runbook.runbook(doc, **_minimal_kwargs(title=""))

    def it_raises_when_title_is_whitespace_only(self):
        doc = Document()
        with pytest.raises(ValueError, match="title is required"):
            runbook.runbook(doc, **_minimal_kwargs(title="   "))

    def it_raises_when_purpose_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="purpose is required"):
            runbook.runbook(doc, **_minimal_kwargs(purpose=""))

    def it_raises_when_when_to_use_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="when_to_use is required"):
            runbook.runbook(doc, **_minimal_kwargs(when_to_use=""))

    def it_raises_when_roles_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="roles must contain"):
            runbook.runbook(doc, **_minimal_kwargs(roles=[]))

    def it_raises_when_prerequisites_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="prerequisites must contain"):
            runbook.runbook(doc, **_minimal_kwargs(prerequisites=[]))

    def it_raises_when_procedure_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="procedure must contain"):
            runbook.runbook(doc, **_minimal_kwargs(procedure=[]))

    def it_raises_when_a_procedure_step_is_missing(self):
        doc = Document()
        with pytest.raises(ValueError, match="non-empty 'step'"):
            runbook.runbook(
                doc,
                **_minimal_kwargs(
                    procedure=[{"owner": "DBA", "detail": "do thing"}],
                ),
            )

    def it_raises_when_a_procedure_entry_is_not_a_mapping(self):
        doc = Document()
        with pytest.raises(ValueError, match="must be a mapping"):
            runbook.runbook(
                doc,
                **_minimal_kwargs(
                    procedure=["not a mapping"],  # type: ignore[list-item]
                ),
            )

    def it_raises_when_an_escalation_entry_is_not_a_tuple(self):
        doc = Document()
        with pytest.raises(ValueError, match="\\(trigger, action\\) tuple"):
            runbook.runbook(
                doc,
                **_minimal_kwargs(
                    escalation=["not a tuple"],  # type: ignore[list-item]
                ),
            )

    def it_raises_when_an_escalation_trigger_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="empty trigger"):
            runbook.runbook(
                doc,
                **_minimal_kwargs(
                    escalation=[("", "Page DBA")],
                ),
            )

    def it_raises_when_an_escalation_action_is_empty(self):
        doc = Document()
        with pytest.raises(ValueError, match="empty action"):
            runbook.runbook(
                doc,
                **_minimal_kwargs(
                    escalation=[("After 5min", "")],
                ),
            )


# -- Round-trip integration ----------------------------------------------


class DescribeRunbookRoundTrip:
    """End-to-end smoke-test: runbook produces a saveable document."""

    def it_can_save_a_runbook_to_a_BytesIO(self):
        doc = Document()
        runbook.runbook(
            doc,
            **_minimal_kwargs(
                escalation=[
                    ("After 5min", "Page on-call DBA"),
                    ("After 15min", "Page director of engineering"),
                ],
                rollback="Revert DNS to old primary, page vendor support.",
            ),
        )
        buf = BytesIO()
        doc.save(buf)

        # -- .docx is a zip; magic bytes are 'PK' --
        assert buf.getvalue()[:2] == b"PK"


# -- Module surface -------------------------------------------------------


class DescribeRunbookModule:
    """Module-level surface contracts."""

    def it_exposes_runbook_as_a_callable(self):
        assert callable(runbook.runbook)
        assert "runbook" in runbook.__all__

    def it_is_re_exported_from_the_kit_package(self):
        from docx.kit import runbook as runbook_pkg

        assert runbook_pkg is runbook


# -- Local helpers --------------------------------------------------------


def _find_table(document, header_first_cell: str):
    """Return the first table whose ``rows[0].cells[0].text`` matches."""
    for table in document.tables:
        cells = table.rows[0].cells
        if cells and cells[0].text == header_first_cell:
            return table
    return None
