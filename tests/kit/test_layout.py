"""Unit-test suite for ``docx.kit.layout`` helpers (issue #286)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.section import WD_SECTION
from docx.kit import layout
from docx.kit.layout import end_multi_column, multi_column
from docx.section import Section
from docx.shared import Inches


@pytest.fixture
def document() -> DocumentCls:
    return Document()


class DescribeMultiColumn:
    """Behavioural tests for :func:`docx.kit.layout.multi_column`."""

    def it_returns_a_section_object(self, document: DocumentCls):
        result = multi_column(document, columns=2)

        assert isinstance(result, Section)

    def it_appends_a_continuous_section_break(self, document: DocumentCls):
        start_count = len(document.sections)

        multi_column(document, columns=2)

        # -- one new section appended at end-of-document --
        assert len(document.sections) == start_count + 1
        # -- new (last) section's start_type is continuous (newspaper-column
        # -- convention — column flow continues on the same page) --
        assert document.sections[-1].start_type == WD_SECTION.CONTINUOUS

    def it_sets_the_column_count_to_the_requested_value(
        self, document: DocumentCls
    ):
        section = multi_column(document, columns=3)

        assert section.columns.count == 3

    def it_defaults_to_two_equal_width_columns(self, document: DocumentCls):
        section = multi_column(document)

        assert section.columns.count == 2
        assert section.columns.equal_width is True

    def it_emits_equal_width_when_requested(self, document: DocumentCls):
        section = multi_column(document, columns=2, equal_width=True)

        assert section.columns.equal_width is True

    def it_writes_the_default_half_inch_gutter(self, document: DocumentCls):
        section = multi_column(document, columns=2)

        # -- 0.5" gutter, expressed as a Length (EMU); compare via .inches --
        assert section.columns.space is not None
        assert pytest.approx(section.columns.space.inches, abs=1e-6) == 0.5

    def it_honours_a_custom_spacing_in_value(self, document: DocumentCls):
        section = multi_column(document, columns=2, spacing_in=0.25)

        assert section.columns.space is not None
        assert pytest.approx(section.columns.space.inches, abs=1e-6) == 0.25

    def it_supports_per_column_widths_in_inches(
        self, document: DocumentCls
    ):
        section = multi_column(
            document,
            columns=2,
            equal_width=False,
            widths_in=[2.0, 4.0],
        )

        # -- equal_width forced to False, two w:col children with the
        # -- requested widths --
        assert section.columns.equal_width is False
        assert len(section.columns) == 2
        assert section.columns[0].width == Inches(2.0)
        assert section.columns[1].width == Inches(4.0)

    def it_raises_when_columns_is_less_than_one(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="columns must be >= 1"):
            multi_column(document, columns=0)

    def it_raises_when_widths_in_length_does_not_match_columns(
        self, document: DocumentCls
    ):
        with pytest.raises(
            ValueError, match="widths_in must have exactly"
        ):
            multi_column(
                document,
                columns=3,
                equal_width=False,
                widths_in=[2.0, 4.0],
            )

    def it_raises_when_widths_in_is_combined_with_equal_width_true(
        self, document: DocumentCls
    ):
        with pytest.raises(
            ValueError, match="mutually exclusive"
        ):
            multi_column(
                document,
                columns=2,
                equal_width=True,
                widths_in=[2.0, 4.0],
            )

    def it_lets_subsequent_paragraphs_belong_to_the_new_section(
        self, document: DocumentCls
    ):
        section = multi_column(document, columns=2)
        document.add_paragraph("flowing across two columns")

        # -- the new paragraph belongs to the multi-column (last) section --
        assert "flowing across two columns" in [
            p.text for p in section.iter_inner_content()  # type: ignore[attr-defined]
        ]


class DescribeEndMultiColumn:
    """Behavioural tests for :func:`docx.kit.layout.end_multi_column`."""

    def it_returns_a_section_object(self, document: DocumentCls):
        multi_column(document, columns=2)
        result = end_multi_column(document)

        assert isinstance(result, Section)

    def it_appends_a_continuous_section_break(self, document: DocumentCls):
        multi_column(document, columns=2)
        start_count = len(document.sections)

        end_multi_column(document)

        assert len(document.sections) == start_count + 1
        # -- the new (last) section's start_type is continuous --
        assert document.sections[-1].start_type == WD_SECTION.CONTINUOUS

    def it_resets_the_column_count_to_one(self, document: DocumentCls):
        multi_column(document, columns=3)

        single = end_multi_column(document)

        assert single.columns.count == 1
        assert single.columns.equal_width is True

    def it_preserves_the_multi_column_geometry_on_the_closed_section(
        self, document: DocumentCls
    ):
        multi_column(document, columns=2)
        # -- writing content while the multi-column section is open --
        document.add_paragraph("in two columns")
        end_multi_column(document)
        document.add_paragraph("back to one column")

        # -- second-to-last section governs the multi-column run --
        closed_section = document.sections[-2]
        assert closed_section.columns.count == 2
        # -- last section is now single-column for what follows --
        assert document.sections[-1].columns.count == 1

    def it_is_safe_to_call_without_a_matching_multi_column_open(
        self, document: DocumentCls
    ):
        # -- redundant single-column break is harmless; helper does not raise --
        result = end_multi_column(document)

        assert isinstance(result, Section)
        assert result.columns.count == 1


class DescribeLayoutPublicSurface:
    """The module is re-exported from ``docx.kit`` so callers can write
    ``from docx.kit import layout`` / ``layout.multi_column(...)`` per the
    issue's example.
    """

    def it_is_reachable_as_docx_kit_layout(self):
        from docx.kit import layout as kit_layout

        assert kit_layout is layout

    def it_exposes_multi_column_and_end_multi_column(self):
        assert hasattr(layout, "multi_column")
        assert hasattr(layout, "end_multi_column")
        assert "multi_column" in layout.__all__
        assert "end_multi_column" in layout.__all__
