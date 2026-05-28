"""Unit-test suite for ``docx.kit.brand`` (issue #90)."""

from __future__ import annotations

import os
from pathlib import Path

import pytest

from docx.kit.brand import (
    BrandAssets,
    BrandAssetsError,
    BrandColors,
    BrandFonts,
    BrandLogos,
    BrandSpacing,
)
from docx.shared import Cm, Inches, Length, Mm, Pt, RGBColor, Twips


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------


_AWS_YAML = """\
name: AWS
colors:
  primary: '#FF9900'
  secondary: '#232F3E'
  accent: '#0073BB'
  background: '#FAFAFA'
fonts:
  heading: 'Amazon Ember Display'
  body: 'Amazon Ember'
logos:
  full_color: 'logos/aws-full-color.png'
  monochrome: 'logos/aws-mono.png'
  reverse: 'logos/aws-reverse.png'
spacing:
  paragraph: 12pt
  section: 24pt
"""


@pytest.fixture
def aws_brand_yaml(tmp_path: Path) -> Path:
    yaml_path = tmp_path / "aws-brand.yaml"
    yaml_path.write_text(_AWS_YAML, encoding="utf-8")
    return yaml_path


@pytest.fixture
def aws_brand(aws_brand_yaml: Path) -> BrandAssets:
    return BrandAssets.load(aws_brand_yaml)


# --------------------------------------------------------------------------
# load() -- end-to-end path
# --------------------------------------------------------------------------


class DescribeBrandAssetsLoad:
    """Acceptance-level tests for :meth:`BrandAssets.load`."""

    def it_returns_a_BrandAssets_instance(self, aws_brand: BrandAssets):
        assert isinstance(aws_brand, BrandAssets)

    def it_records_the_brand_name(self, aws_brand: BrandAssets):
        assert aws_brand.name == "AWS"

    def it_records_the_source_path(
        self, aws_brand: BrandAssets, aws_brand_yaml: Path
    ):
        assert aws_brand.source_path == str(aws_brand_yaml)

    def it_parses_colors_into_RGBColor_instances(self, aws_brand: BrandAssets):
        assert aws_brand.colors.primary == RGBColor(0xFF, 0x99, 0x00)
        assert aws_brand.colors.secondary == RGBColor(0x23, 0x2F, 0x3E)
        assert aws_brand.colors.accent == RGBColor(0x00, 0x73, 0xBB)
        assert aws_brand.colors.background == RGBColor(0xFA, 0xFA, 0xFA)

    def it_exposes_font_names_as_strings(self, aws_brand: BrandAssets):
        assert aws_brand.fonts.heading == "Amazon Ember Display"
        assert aws_brand.fonts.body == "Amazon Ember"

    def it_resolves_relative_logo_paths_against_the_yaml_directory(
        self, aws_brand: BrandAssets, aws_brand_yaml: Path
    ):
        expected_dir = str(aws_brand_yaml.parent)
        assert aws_brand.logos.full_color == os.path.normpath(
            os.path.join(expected_dir, "logos/aws-full-color.png")
        )
        assert os.path.isabs(aws_brand.logos.monochrome)
        assert os.path.isabs(aws_brand.logos.reverse)

    def it_parses_spacing_into_Length_instances(self, aws_brand: BrandAssets):
        assert isinstance(aws_brand.spacing.paragraph, Length)
        assert isinstance(aws_brand.spacing.section, Length)
        assert aws_brand.spacing.paragraph == Pt(12)
        assert aws_brand.spacing.section == Pt(24)

    def it_accepts_a_PathLike_argument(
        self, aws_brand_yaml: Path
    ):
        # Ensure load() accepts both ``str`` and ``pathlib.Path``.
        brand_str = BrandAssets.load(str(aws_brand_yaml))
        brand_path = BrandAssets.load(aws_brand_yaml)

        assert brand_str.name == brand_path.name == "AWS"

    def it_is_an_immutable_dataclass(self, aws_brand: BrandAssets):
        with pytest.raises((AttributeError, TypeError)):
            aws_brand.name = "Other"  # type: ignore[misc]


# --------------------------------------------------------------------------
# from_dict() -- programmatic construction
# --------------------------------------------------------------------------


class DescribeBrandAssetsFromDict:
    """Tests for :meth:`BrandAssets.from_dict`."""

    def it_builds_from_an_in_memory_mapping(self):
        data = {
            "name": "Acme",
            "colors": {"primary": "#FF0000"},
            "fonts": {"heading": "Helvetica"},
            "logos": {"full_color": "/abs/path/logo.png"},
            "spacing": {"paragraph": "10pt"},
        }

        brand = BrandAssets.from_dict(data)

        assert brand.name == "Acme"
        assert brand.colors.primary == RGBColor(0xFF, 0x00, 0x00)
        assert brand.fonts.heading == "Helvetica"
        assert brand.logos.full_color == "/abs/path/logo.png"
        assert brand.spacing.paragraph == Pt(10)

    def it_returns_empty_views_for_omitted_blocks(self):
        brand = BrandAssets.from_dict({})

        assert brand.name is None
        assert brand.colors == BrandColors()
        assert brand.fonts == BrandFonts()
        assert brand.logos == BrandLogos()
        assert brand.spacing == BrandSpacing()
        assert brand.colors.primary is None
        assert brand.fonts.heading is None
        assert brand.logos.full_color is None
        assert brand.spacing.paragraph is None

    def it_anchors_relative_logo_paths_to_base_dir_when_supplied(self):
        brand = BrandAssets.from_dict(
            {"logos": {"full_color": "logo.png"}},
            base_dir="/home/brand",
        )
        assert brand.logos.full_color == os.path.normpath(
            "/home/brand/logo.png"
        )

    def it_passes_relative_paths_through_when_base_dir_is_None(self):
        brand = BrandAssets.from_dict(
            {"logos": {"full_color": "logo.png"}}
        )
        assert brand.logos.full_color == "logo.png"

    def it_keeps_absolute_logo_paths_as_is(self):
        brand = BrandAssets.from_dict(
            {"logos": {"full_color": "/etc/logo.png"}},
            base_dir="/home/brand",
        )
        assert brand.logos.full_color == "/etc/logo.png"

    def it_collects_unknown_color_keys_into_extras(self):
        brand = BrandAssets.from_dict(
            {"colors": {"primary": "#FF0000", "warning": "#F39C12"}}
        )
        assert brand.colors.primary == RGBColor(0xFF, 0x00, 0x00)
        assert brand.colors.extras["warning"] == RGBColor(0xF3, 0x9C, 0x12)
        # Subscript also reaches both named and extras.
        assert brand.colors["primary"] == RGBColor(0xFF, 0x00, 0x00)
        assert brand.colors["warning"] == RGBColor(0xF3, 0x9C, 0x12)

    def it_collects_unknown_logo_keys_into_extras(self):
        brand = BrandAssets.from_dict(
            {"logos": {"full_color": "/a/full.png", "favicon": "/a/fav.ico"}}
        )
        assert brand.logos.full_color == "/a/full.png"
        assert brand.logos.extras["favicon"] == "/a/fav.ico"

    def it_collects_unknown_font_keys_into_extras(self):
        brand = BrandAssets.from_dict(
            {"fonts": {"heading": "Helvetica", "code": "Source Code Pro"}}
        )
        assert brand.fonts.heading == "Helvetica"
        assert brand.fonts.extras["code"] == "Source Code Pro"

    def it_collects_unknown_spacing_keys_into_extras(self):
        brand = BrandAssets.from_dict(
            {"spacing": {"paragraph": "10pt", "gutter": "0.25in"}}
        )
        assert brand.spacing.paragraph == Pt(10)
        assert brand.spacing.extras["gutter"] == Inches(0.25)


# --------------------------------------------------------------------------
# colour parsing
# --------------------------------------------------------------------------


class DescribeColorParsing:
    """Edge cases on the hex-string / RGB-triple colour parser."""

    @pytest.mark.parametrize(
        ("raw", "expected"),
        [
            ("#FF9900", RGBColor(0xFF, 0x99, 0x00)),
            ("FF9900", RGBColor(0xFF, 0x99, 0x00)),
            ("#ff9900", RGBColor(0xFF, 0x99, 0x00)),
            ("ff9900", RGBColor(0xFF, 0x99, 0x00)),
            ("#F90", RGBColor(0xFF, 0x99, 0x00)),
            ("F90", RGBColor(0xFF, 0x99, 0x00)),
        ],
    )
    def it_parses_hex_strings(self, raw: str, expected: RGBColor):
        brand = BrandAssets.from_dict({"colors": {"primary": raw}})
        assert brand.colors.primary == expected

    def it_accepts_a_3_int_list(self):
        brand = BrandAssets.from_dict({"colors": {"primary": [255, 153, 0]}})
        assert brand.colors.primary == RGBColor(0xFF, 0x99, 0x00)

    def it_passes_through_an_RGBColor_instance(self):
        existing = RGBColor(0x12, 0x34, 0x56)
        brand = BrandAssets.from_dict({"colors": {"primary": existing}})
        assert brand.colors.primary == existing

    @pytest.mark.parametrize(
        "bad",
        [
            "#GGHHII",  # not hex
            "#1234",     # neither 3 nor 6 chars
            "12345",     # neither 3 nor 6 chars
            42,          # not a string / list
        ],
    )
    def it_raises_on_malformed_color(self, bad: object):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict({"colors": {"primary": bad}})


# --------------------------------------------------------------------------
# length parsing
# --------------------------------------------------------------------------


class DescribeLengthParsing:
    """Edge cases on the spacing/length string parser."""

    @pytest.mark.parametrize(
        ("raw", "expected"),
        [
            ("12pt", Pt(12)),
            ("12 pt", Pt(12)),
            ("0.5in", Inches(0.5)),
            ("24mm", Mm(24)),
            ("3cm", Cm(3)),
            ("914400emu", Length(914400)),
            ("720twips", Twips(720)),
            ("720twip", Twips(720)),
        ],
    )
    def it_parses_suffixed_strings(self, raw: str, expected: Length):
        brand = BrandAssets.from_dict({"spacing": {"paragraph": raw}})
        assert brand.spacing.paragraph == expected

    def it_treats_bare_numbers_as_points(self):
        brand = BrandAssets.from_dict({"spacing": {"paragraph": 18}})
        assert brand.spacing.paragraph == Pt(18)

    def it_treats_bare_numeric_strings_as_points(self):
        brand = BrandAssets.from_dict({"spacing": {"paragraph": "18"}})
        assert brand.spacing.paragraph == Pt(18)

    def it_passes_a_Length_through_unchanged(self):
        brand = BrandAssets.from_dict({"spacing": {"paragraph": Pt(7)}})
        assert brand.spacing.paragraph == Pt(7)

    @pytest.mark.parametrize(
        "bad",
        ["", "pt", "abc", True, [12]],
    )
    def it_raises_on_malformed_length(self, bad: object):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict({"spacing": {"paragraph": bad}})


# --------------------------------------------------------------------------
# error paths
# --------------------------------------------------------------------------


class DescribeBrandAssetsErrors:
    """Negative cases — malformed manifests should surface clean errors."""

    def it_raises_on_a_non_mapping_top_level(self, tmp_path: Path):
        bad_yaml = tmp_path / "bad.yaml"
        bad_yaml.write_text("- just\n- a\n- list\n", encoding="utf-8")
        with pytest.raises(BrandAssetsError):
            BrandAssets.load(bad_yaml)

    def it_raises_on_invalid_yaml(self, tmp_path: Path):
        bad_yaml = tmp_path / "bad.yaml"
        bad_yaml.write_text(": : :\n  -broken\n", encoding="utf-8")
        with pytest.raises(BrandAssetsError):
            BrandAssets.load(bad_yaml)

    def it_treats_an_empty_yaml_file_as_an_empty_brand(self, tmp_path: Path):
        empty_yaml = tmp_path / "empty.yaml"
        empty_yaml.write_text("", encoding="utf-8")
        brand = BrandAssets.load(empty_yaml)
        assert brand.name is None
        assert brand.colors.primary is None

    def it_raises_on_non_string_name(self):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict({"name": 42})

    def it_raises_on_non_mapping_colors_block(self):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict({"colors": "not-a-mapping"})

    def it_raises_on_non_string_font_value(self):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict({"fonts": {"heading": 12}})

    def it_raises_on_non_string_logo_value(self):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict({"logos": {"full_color": 12}})

    def it_raises_on_empty_logo_path(self):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict({"logos": {"full_color": "  "}})

    def it_raises_on_non_mapping_argument_to_from_dict(self):
        with pytest.raises(BrandAssetsError):
            BrandAssets.from_dict("not-a-dict")  # type: ignore[arg-type]


# --------------------------------------------------------------------------
# subscript access
# --------------------------------------------------------------------------


class DescribeViewSubscript:
    """The four sub-views support ``view[key]`` for both named + extras."""

    def it_raises_KeyError_for_unset_named_field(self):
        brand = BrandAssets.from_dict({"colors": {"secondary": "#000000"}})
        with pytest.raises(KeyError):
            _ = brand.colors["primary"]

    def it_raises_KeyError_for_unknown_extras(self):
        brand = BrandAssets.from_dict({})
        with pytest.raises(KeyError):
            _ = brand.colors["nonsense"]


# -- Tests for validate_brand (#91) ------------------------------------------

from __future__ import annotations

import os
from typing import Any, Dict, List, Mapping

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import brand
from docx.kit.brand import (
    RULE_IDS,
    SEVERITIES,
    BrandFinding,
    validate_brand,
)
from docx.shared import Pt, RGBColor

from ..unitutil.file import test_file


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def document() -> DocumentCls:
    return Document()


@pytest.fixture
def aws_rules() -> Dict[str, Any]:
    """Mapping-shaped brand rules mimicking the AWS palette in the issue body."""
    return {
        "fonts": ["Amazon Ember", "Amazon Ember Display"],
        "colors": {
            "#232F3E": "Squid Ink",
            "#FF9900": "Smile Orange",
        },
        "logos": [
            {"path": "aws-logo.png", "sha1": "0" * 40, "kind": "primary"},
        ],
        "spacing": {
            "line_spacing": 1.15,
            "space_before": 0,
            "space_after": 6,
        },
    }


def _by_rule(findings: List[BrandFinding], rule: str) -> List[BrandFinding]:
    return [f for f in findings if f.rule == rule]


# ---------------------------------------------------------------------------
# Module-level constants and dataclass surface
# ---------------------------------------------------------------------------


class DescribeBrandFinding:
    """Surface tests for the public dataclass + module-level constants."""

    def it_exposes_the_three_severity_tiers_in_ascending_order(self):
        assert SEVERITIES == ("info", "warning", "error")

    def it_exposes_the_five_rule_ids(self):
        assert set(RULE_IDS) == {
            "font-not-on-brand",
            "color-not-on-brand",
            "wrong-logo",
            "heading-style-mismatch",
            "inconsistent-spacing",
        }

    def it_is_a_frozen_dataclass(self):
        f = BrandFinding("info", "paragraph 0", "font-not-on-brand", "...")
        with pytest.raises(Exception):
            f.severity = "error"  # type: ignore[misc]

    def it_carries_severity_location_rule_and_message(self):
        f = BrandFinding("warning", "paragraph 5", "font-not-on-brand", "msg")
        assert f.severity == "warning"
        assert f.location == "paragraph 5"
        assert f.rule == "font-not-on-brand"
        assert f.message == "msg"


# ---------------------------------------------------------------------------
# Rules ingestion: dict / YAML path / BrandAssets-like duck typing
# ---------------------------------------------------------------------------


class DescribeRulesIngestion:
    """``validate_brand`` accepts multiple rule shapes."""

    def it_accepts_a_dict_directly(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        findings = validate_brand(document, aws_rules)
        # An empty document shouldn't trigger any finding.
        assert findings == []

    def it_accepts_a_yaml_file_path(
        self, document: DocumentCls, tmp_path
    ):
        rules_file = tmp_path / "brand.yaml"
        rules_file.write_text(
            "fonts:\n"
            "  - Amazon Ember\n"
            "  - Amazon Ember Display\n"
            "colors:\n"
            "  \"#232F3E\": Squid Ink\n"
            "  \"#FF9900\": Smile Orange\n"
            "spacing:\n"
            "  line_spacing: 1.15\n"
            "  space_before: 0\n"
            "  space_after: 6\n",
            encoding="utf-8",
        )

        # add an off-brand run so we know parsing actually drove a check
        p = document.add_paragraph()
        run = p.add_run("hello")
        run.font.name = "Times New Roman"

        findings = validate_brand(document, str(rules_file))
        font_findings = _by_rule(findings, "font-not-on-brand")
        assert any("Times New Roman" in f.message for f in font_findings)

    def it_accepts_an_os_pathlike(
        self, document: DocumentCls, tmp_path
    ):
        rules_file = tmp_path / "brand.yaml"
        rules_file.write_text("fonts:\n  - Amazon Ember\n", encoding="utf-8")
        # pass the Path directly (os.PathLike instance)
        findings = validate_brand(document, rules_file)
        assert findings == []

    def it_accepts_a_brandassets_like_duck_typed_object(
        self, document: DocumentCls
    ):
        class FakeBrandAssets:
            fonts = ("Amazon Ember",)
            colors = {"#232F3E": "Squid Ink"}
            logos = ()
            spacing = {}

        p = document.add_paragraph()
        p.add_run("text").font.name = "Times New Roman"

        findings = validate_brand(document, FakeBrandAssets())
        assert any(f.rule == "font-not-on-brand" for f in findings)

    def it_raises_when_rules_is_an_unrelated_object(
        self, document: DocumentCls
    ):
        with pytest.raises(TypeError):
            validate_brand(document, 42)  # type: ignore[arg-type]

    def it_treats_colors_supplied_as_a_list_of_hex_strings(
        self, document: DocumentCls
    ):
        rules = {
            "fonts": ["Amazon Ember"],
            "colors": ["#232F3E", "#FF9900"],
        }
        p = document.add_paragraph()
        p.add_run("on-brand").font.color.rgb = RGBColor(0x23, 0x2F, 0x3E)
        findings = validate_brand(document, rules)
        assert _by_rule(findings, "color-not-on-brand") == []


# ---------------------------------------------------------------------------
# font-not-on-brand
# ---------------------------------------------------------------------------


class DescribeFontNotOnBrand:
    """Coverage for the ``font-not-on-brand`` rule."""

    def it_flags_a_run_using_an_off_brand_font(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph()
        p.add_run("hello").font.name = "Times New Roman"

        findings = validate_brand(document, aws_rules)
        font_findings = _by_rule(findings, "font-not-on-brand")
        assert font_findings, findings
        f = font_findings[0]
        assert f.severity == "warning"
        assert "Times New Roman" in f.message
        assert "Amazon Ember" in f.message
        assert f.location.startswith("paragraph ")

    def it_does_not_flag_a_run_using_an_on_brand_font(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph()
        p.add_run("hello").font.name = "Amazon Ember"
        findings = validate_brand(document, aws_rules)
        assert _by_rule(findings, "font-not-on-brand") == []

    def it_skips_runs_whose_font_name_is_unset(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        # No explicit font.name -> run inherits from style; the validator
        # must not synthesise a phantom finding when the run is silent.
        document.add_paragraph("plain")
        findings = validate_brand(document, aws_rules)
        # the default 'Normal' style has no font name set, so no finding
        assert _by_rule(findings, "font-not-on-brand") == []

    def it_emits_a_finding_per_offending_run(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph()
        for text in ("a", "b", "c"):
            run = p.add_run(text)
            run.font.name = "Comic Sans MS"
        findings = validate_brand(document, aws_rules)
        assert len(_by_rule(findings, "font-not-on-brand")) == 3


# ---------------------------------------------------------------------------
# color-not-on-brand
# ---------------------------------------------------------------------------


class DescribeColorNotOnBrand:
    """Coverage for the ``color-not-on-brand`` rule."""

    def it_flags_a_run_with_an_off_brand_text_default_color(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        # Mirror the issue example: explicit black on a brand that
        # mandates Squid Ink.
        for _ in range(13):
            document.add_paragraph()
        target = document.paragraphs[12]
        target.add_run("hi").font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        findings = validate_brand(document, aws_rules)
        color_findings = _by_rule(findings, "color-not-on-brand")
        assert color_findings
        f = color_findings[0]
        assert f.severity == "warning"
        assert f.location == "paragraph 12 run 0"
        assert "text-default" in f.message
        assert "Squid Ink" in f.message
        assert "#232F3E" in f.message

    def it_does_not_flag_an_on_brand_color(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph()
        p.add_run("on-brand").font.color.rgb = RGBColor(0x23, 0x2F, 0x3E)
        findings = validate_brand(document, aws_rules)
        assert _by_rule(findings, "color-not-on-brand") == []

    def it_flags_a_generic_off_brand_color_without_the_text_default_phrasing(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph()
        p.add_run("magenta").font.color.rgb = RGBColor(0xFF, 0x00, 0xFF)
        findings = validate_brand(document, aws_rules)
        f = _by_rule(findings, "color-not-on-brand")[0]
        assert "#FF00FF" in f.message
        assert "text-default" not in f.message

    def it_skips_runs_with_no_color_override(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        document.add_paragraph("plain")
        findings = validate_brand(document, aws_rules)
        assert _by_rule(findings, "color-not-on-brand") == []


# ---------------------------------------------------------------------------
# heading-style-mismatch
# ---------------------------------------------------------------------------


class DescribeHeadingStyleMismatch:
    """Heading paragraphs use a separate rule id when they drift."""

    def it_uses_heading_style_mismatch_for_an_off_brand_heading_font(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph("Heading", style="Heading 1")
        p.runs[0].font.name = "Comic Sans MS"
        findings = validate_brand(document, aws_rules)
        rule_ids = {f.rule for f in findings}
        assert "heading-style-mismatch" in rule_ids
        # ... and *not* the generic font-not-on-brand for the same paragraph
        assert "font-not-on-brand" not in rule_ids

    def it_uses_font_not_on_brand_for_body_paragraphs(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph()
        p.add_run("body").font.name = "Comic Sans MS"
        findings = validate_brand(document, aws_rules)
        rule_ids = {f.rule for f in findings}
        assert "font-not-on-brand" in rule_ids
        assert "heading-style-mismatch" not in rule_ids


# ---------------------------------------------------------------------------
# wrong-logo
# ---------------------------------------------------------------------------


class DescribeWrongLogo:
    """Best-effort logo / competitor-asset detection."""

    def it_flags_a_logo_shaped_image_that_is_not_in_the_brand_palette(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        # Use a small bundled PNG — the python-icon test asset.
        run = document.add_paragraph().add_run()
        run.add_picture(test_file("python-icon.png"))

        findings = validate_brand(document, aws_rules)
        wrong = _by_rule(findings, "wrong-logo")
        assert wrong
        f = wrong[0]
        assert f.severity == "error"
        assert "python-icon.png" in f.location
        assert "logo" in f.message.lower()

    def it_does_not_flag_a_registered_brand_logo_by_filename(
        self, document: DocumentCls
    ):
        rules = {
            "fonts": ["Amazon Ember"],
            "logos": [{"path": "python-icon.png"}],
        }
        run = document.add_paragraph().add_run()
        run.add_picture(test_file("python-icon.png"))

        findings = validate_brand(document, rules)
        assert _by_rule(findings, "wrong-logo") == []

    def it_does_not_flag_an_image_too_large_to_be_a_logo(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        # Synthesise a "big" jpg by reusing the bundled tiff/jpg files.
        # We can't easily fabricate a multi-MB asset in unit tests, so
        # we rely on the size heuristic by directly faking via a stub.
        from docx.kit.brand import _check_inline_shape, _Rules  # type: ignore

        class FakeImage:
            filename = "big-banner.jpg"
            blob = b"x" * (1024 * 1024)  # 1 MB
            sha1 = "deadbeef" * 5

        class FakeShape:
            image = FakeImage()

        findings: List[BrandFinding] = []
        rules = _Rules(
            fonts=("Amazon Ember",),
            colors={},
            logos=(),
            spacing={},
        )
        _check_inline_shape(FakeShape(), rules, findings)
        assert findings == []

    def it_does_not_flag_an_image_with_a_non_logo_extension(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        from docx.kit.brand import _check_inline_shape, _Rules  # type: ignore

        class FakeImage:
            filename = "diagram.tif"
            blob = b"x" * 1000
            sha1 = "feedface" * 5

        class FakeShape:
            image = FakeImage()

        findings: List[BrandFinding] = []
        rules = _Rules(fonts=(), colors={}, logos=(), spacing={})
        _check_inline_shape(FakeShape(), rules, findings)
        assert findings == []

    def it_matches_a_registered_logo_by_sha1(
        self, document: DocumentCls
    ):
        from docx.kit.brand import _check_inline_shape, _Rules  # type: ignore

        sha = "abcd" * 10
        rules = _Rules(
            fonts=(),
            colors={},
            logos=({"sha1": sha},),
            spacing={},
        )

        class FakeImage:
            filename = "renamed-logo.png"
            blob = b"x" * 1000
            sha1 = sha.upper()

        class FakeShape:
            image = FakeImage()

        findings: List[BrandFinding] = []
        _check_inline_shape(FakeShape(), rules, findings)
        assert findings == []


# ---------------------------------------------------------------------------
# inconsistent-spacing
# ---------------------------------------------------------------------------


class DescribeInconsistentSpacing:
    """Coverage for the info-severity ``inconsistent-spacing`` rule."""

    def it_flags_a_paragraph_with_off_brand_line_spacing(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph("body")
        p.paragraph_format.line_spacing = 2.0  # vs. 1.15
        findings = validate_brand(document, aws_rules)
        spacing = _by_rule(findings, "inconsistent-spacing")
        assert spacing
        f = spacing[0]
        assert f.severity == "info"
        assert "line_spacing" in f.message

    def it_flags_a_paragraph_with_off_brand_space_after(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph("body")
        p.paragraph_format.space_after = Pt(24)  # vs. 6
        findings = validate_brand(document, aws_rules)
        spacing = _by_rule(findings, "inconsistent-spacing")
        assert spacing
        assert "space_after" in spacing[0].message

    def it_does_not_flag_when_spacing_matches_the_brand(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        p = document.add_paragraph("body")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        findings = validate_brand(document, aws_rules)
        assert _by_rule(findings, "inconsistent-spacing") == []

    def it_skips_when_brand_has_no_spacing_rules(
        self, document: DocumentCls
    ):
        rules = {"fonts": ["Amazon Ember"]}
        p = document.add_paragraph("body")
        p.paragraph_format.line_spacing = 99
        findings = validate_brand(document, rules)
        assert _by_rule(findings, "inconsistent-spacing") == []


# ---------------------------------------------------------------------------
# Findings shape and document order
# ---------------------------------------------------------------------------


class DescribeFindingsShape:
    """End-to-end-style cross-checks on the returned list."""

    def it_returns_an_empty_list_for_an_on_brand_document(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        # Default empty document — no runs, no inline shapes.
        assert validate_brand(document, aws_rules) == []

    def it_returns_findings_in_paragraph_order(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        # Paragraph 1: bad font; paragraph 3: bad colour.
        document.add_paragraph()
        document.add_paragraph().add_run("a").font.name = "Comic Sans MS"
        document.add_paragraph()
        document.add_paragraph().add_run("b").font.color.rgb = RGBColor(
            0xFF, 0x00, 0xFF
        )

        findings = validate_brand(document, aws_rules)
        # The font finding for paragraph 1 should precede the colour
        # finding for paragraph 3.
        font_idx = next(
            i for i, f in enumerate(findings) if f.rule == "font-not-on-brand"
        )
        color_idx = next(
            i for i, f in enumerate(findings) if f.rule == "color-not-on-brand"
        )
        assert font_idx < color_idx

    def it_returns_a_list_of_BrandFinding_instances(
        self, document: DocumentCls, aws_rules: Dict[str, Any]
    ):
        document.add_paragraph().add_run("x").font.name = "Comic Sans MS"
        findings = validate_brand(document, aws_rules)
        assert findings
        for f in findings:
            assert isinstance(f, BrandFinding)
            assert f.severity in SEVERITIES
            assert f.rule in RULE_IDS


# ---------------------------------------------------------------------------
# YAML subset fallback parser
# ---------------------------------------------------------------------------


class DescribeYamlSubsetParser:
    """The fallback parser handles the schema we document."""

    def it_parses_a_top_level_list_under_a_key(self):
        from docx.kit.brand import _yaml_subset_parse

        data = _yaml_subset_parse("fonts:\n  - Amazon Ember\n  - Helvetica\n")
        assert data == {"fonts": ["Amazon Ember", "Helvetica"]}

    def it_parses_a_nested_mapping(self):
        from docx.kit.brand import _yaml_subset_parse

        data = _yaml_subset_parse(
            "spacing:\n  line_spacing: 1.15\n  space_after: 6\n"
        )
        assert data == {"spacing": {"line_spacing": 1.15, "space_after": 6}}

    def it_parses_a_list_of_single_line_mappings(self):
        from docx.kit.brand import _yaml_subset_parse

        text = (
            "logos:\n"
            "  - path: aws-logo.png\n"
            '    sha1: "abcd1234abcd1234abcd1234abcd1234abcd1234"\n'
        )
        data = _yaml_subset_parse(text)
        assert data["logos"][0]["path"] == "aws-logo.png"
        assert data["logos"][0]["sha1"].startswith("abcd")

    def it_strips_quotes_around_scalars(self):
        from docx.kit.brand import _yaml_subset_parse

        data = _yaml_subset_parse('colors:\n  "#232F3E": Squid Ink\n')
        assert data == {"colors": {"#232F3E": "Squid Ink"}}

    def it_skips_blank_lines_and_comments(self):
        from docx.kit.brand import _yaml_subset_parse

        data = _yaml_subset_parse(
            "# top-level comment\n\nfonts:\n  - Amazon Ember\n\n"
        )
        assert data == {"fonts": ["Amazon Ember"]}
