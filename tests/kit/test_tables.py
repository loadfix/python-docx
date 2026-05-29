"""Unit-test suite for ``docx.kit.tables`` styled-table helpers."""

from __future__ import annotations

import datetime as _dt
import sys
from io import BytesIO
from typing import Any
from unittest.mock import patch

import pytest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import tables
from docx.shared import RGBColor


# -- Shared helpers -------------------------------------------------------


def _new_doc() -> Any:
    return Document()


def _round_trip(document: Any) -> Any:
    """Save ``document`` and load it back so we exercise the writer too."""
    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return Document(buf)


# -- Style resolution -----------------------------------------------------


class DescribeBuiltinStyles:
    def it_lists_the_four_built_in_styles(self):
        assert tables.BUILTIN_STYLES == (
            "modern",
            "zebra",
            "minimal",
            "corporate",
        )

    def it_raises_on_an_unknown_style(self):
        doc = _new_doc()
        with pytest.raises(ValueError, match="unknown style"):
            tables.styled_table(
                doc, headers=["A"], rows=[["x"]], style="rainbow"
            )

    def it_rejects_a_non_string_style(self):
        doc = _new_doc()
        with pytest.raises(ValueError, match="style must be a string"):
            tables.styled_table(
                doc, headers=["A"], rows=[["x"]], style=42  # type: ignore[arg-type]
            )


# -- styled_table ---------------------------------------------------------


class DescribeStyledTable:
    """Unit-test suite for ``tables.styled_table``."""

    def it_returns_the_appended_table(self):
        from docx.table import Table

        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1], ["Beta", 2]],
            style="modern",
        )

        assert isinstance(table, Table)
        assert len(doc.tables) == 1
        # The returned proxy wraps the same w:tbl element.
        assert table._tbl is doc.tables[0]._tbl

    def it_emits_one_header_row_plus_one_row_per_data_entry(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A", "B"],
            rows=[["a1", "b1"], ["a2", "b2"], ["a3", "b3"]],
            style="modern",
        )

        # 1 header + 3 body
        assert len(table.rows) == 4

    def it_writes_the_header_labels_into_the_first_row(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1]],
            style="modern",
        )

        cells = table.rows[0].cells
        assert cells[0].text == "Name"
        assert cells[1].text == "Value"

    def it_writes_the_body_values_into_subsequent_rows(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1], ["Beta", 2]],
            style="modern",
        )

        assert [c.text for c in table.rows[1].cells] == ["Alpha", "1"]
        assert [c.text for c in table.rows[2].cells] == ["Beta", "2"]

    def it_renders_None_as_empty_string(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["X"],
            rows=[[None]],
            style="modern",
        )

        assert table.rows[1].cells[0].text == ""

    def it_renders_booleans_as_yes_no(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Flag"],
            rows=[[True], [False]],
            style="modern",
        )

        assert table.rows[1].cells[0].text == "Yes"
        assert table.rows[2].cells[0].text == "No"

    def it_renders_floats_without_trailing_zeros(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["X"],
            rows=[[3.0], [3.14]],
            style="modern",
        )

        assert table.rows[1].cells[0].text == "3"
        assert table.rows[2].cells[0].text == "3.14"

    def it_renders_dates_in_iso_8601(self):
        doc = _new_doc()
        d = _dt.date(2026, 5, 29)
        dt = _dt.datetime(2026, 5, 29, 12, 30, 45)
        table = tables.styled_table(
            doc,
            headers=["When"],
            rows=[[d], [dt]],
            style="modern",
        )

        assert table.rows[1].cells[0].text == "2026-05-29"
        assert table.rows[2].cells[0].text == "2026-05-29 12:30:45"

    def it_right_aligns_numeric_columns_by_default(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1], ["Beta", 2]],
            style="modern",
        )

        # -- Body row of the numeric column is right-aligned. --
        value_cell = table.rows[1].cells[1]
        assert (
            value_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        )

    def it_left_aligns_text_columns_by_default(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1], ["Beta", 2]],
            style="modern",
        )

        name_cell = table.rows[1].cells[0]
        assert (
            name_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        )

    def it_lets_caller_override_per_column_alignment(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1]],
            style="modern",
            column_alignments=[
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
            ],
        )

        for col in (0, 1):
            assert (
                table.rows[1].cells[col].paragraphs[0].alignment
                == WD_ALIGN_PARAGRAPH.CENTER
            )

    def it_raises_when_column_alignments_length_mismatches(self):
        doc = _new_doc()
        with pytest.raises(ValueError, match="column_alignments has 1"):
            tables.styled_table(
                doc,
                headers=["A", "B"],
                rows=[["x", "y"]],
                style="modern",
                column_alignments=[WD_ALIGN_PARAGRAPH.CENTER],
            )

    def it_raises_on_empty_headers(self):
        doc = _new_doc()
        with pytest.raises(ValueError, match="headers must be a non-empty"):
            tables.styled_table(doc, headers=[], rows=[], style="modern")

    def it_raises_on_row_length_mismatch(self):
        doc = _new_doc()
        with pytest.raises(ValueError, match="rows\\[1\\] has 1 cells"):
            tables.styled_table(
                doc,
                headers=["A", "B"],
                rows=[["x", "y"], ["z"]],
                style="modern",
            )

    def it_applies_modern_style_header_fill(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1]],
            style="modern",
        )

        # -- Modern header fill is the deep blue 0x1F497D. --
        cell = table.rows[0].cells[0]
        assert cell.shading.fill_color == RGBColor(0x1F, 0x49, 0x7D)

    def it_applies_white_text_to_modern_header(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name"],
            rows=[["Alpha"]],
            style="modern",
        )

        run = table.rows[0].cells[0].paragraphs[0].runs[0]
        assert run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF)
        assert run.bold is True

    def it_bands_zebra_rows(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["r0"], ["r1"], ["r2"], ["r3"]],
            style="zebra",
        )

        # -- Body row 0 (table.rows[1]): unshaded. --
        assert table.rows[1].cells[0].shading.fill_color is None
        # -- Body row 1 (table.rows[2]): banded with light grey. --
        assert table.rows[2].cells[0].shading.fill_color == RGBColor(
            0xF2, 0xF2, 0xF2
        )
        # -- Body row 2 (table.rows[3]): unshaded. --
        assert table.rows[3].cells[0].shading.fill_color is None
        # -- Body row 3 (table.rows[4]): banded. --
        assert table.rows[4].cells[0].shading.fill_color == RGBColor(
            0xF2, 0xF2, 0xF2
        )

    def it_skips_banding_in_modern_style(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["r0"], ["r1"], ["r2"]],
            style="modern",
        )

        for row in table.rows[1:]:
            assert row.cells[0].shading.fill_color is None

    def it_underlines_and_bolds_minimal_style_header(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["Name"],
            rows=[["Alpha"]],
            style="minimal",
        )

        run = table.rows[0].cells[0].paragraphs[0].runs[0]
        assert run.bold is True
        assert run.underline is True
        # No header fill.
        assert table.rows[0].cells[0].shading.fill_color is None

    def it_applies_corporate_style_navy_header(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["x"], ["y"]],
            style="corporate",
        )

        assert table.rows[0].cells[0].shading.fill_color == RGBColor(
            0x0B, 0x2D, 0x5C
        )
        # Banding present (light blue, on body row 1).
        assert table.rows[2].cells[0].shading.fill_color == RGBColor(
            0xDC, 0xE6, 0xF2
        )

    def it_lets_caller_override_header_fill_with_rgb(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["x"]],
            style="modern",
            header_fill=RGBColor(0xAA, 0xBB, 0xCC),
        )

        assert table.rows[0].cells[0].shading.fill_color == RGBColor(
            0xAA, 0xBB, 0xCC
        )

    def it_lets_caller_override_header_fill_with_hex(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["x"]],
            style="modern",
            header_fill="A1B2C3",
        )

        assert table.rows[0].cells[0].shading.fill_color == RGBColor(
            0xA1, 0xB2, 0xC3
        )

    def it_suppresses_header_fill_when_caller_passes_None(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["x"]],
            style="modern",
            header_fill=None,
        )

        assert table.rows[0].cells[0].shading.fill_color is None

    def it_suppresses_zebra_banding_when_alt_row_fill_is_None(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["x"], ["y"], ["z"]],
            style="zebra",
            alt_row_fill=None,
        )

        for row in table.rows[1:]:
            assert row.cells[0].shading.fill_color is None

    def it_lets_caller_override_alt_row_fill_with_hex(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc,
            headers=["A"],
            rows=[["x"], ["y"]],
            style="modern",
            alt_row_fill="EEEEEE",
        )

        assert table.rows[2].cells[0].shading.fill_color == RGBColor(
            0xEE, 0xEE, 0xEE
        )

    def it_rejects_a_malformed_colour_override(self):
        doc = _new_doc()
        with pytest.raises(ValueError):
            tables.styled_table(
                doc,
                headers=["A"],
                rows=[["x"]],
                style="modern",
                header_fill=12345,  # not an RGBColor / hex / None
            )

    def it_sets_autofit_True_by_default(self):
        doc = _new_doc()
        table = tables.styled_table(
            doc, headers=["A"], rows=[["x"]], style="modern"
        )

        assert table.autofit is True

    def it_round_trips_through_save_and_load(self):
        doc = _new_doc()
        tables.styled_table(
            doc,
            headers=["Name", "Value"],
            rows=[["Alpha", 1], ["Beta", 2]],
            style="zebra",
        )

        reloaded = _round_trip(doc)
        assert len(reloaded.tables) == 1
        assert reloaded.tables[0].rows[0].cells[0].text == "Name"


# -- from_dataframe (pandas-aware) ---------------------------------------


class DescribeFromDataframe:
    """Unit-test suite for ``tables.from_dataframe``."""

    def it_raises_a_helpful_ImportError_when_pandas_is_missing(self):
        # -- Block pandas import via sys.modules sentinel + builtin
        # -- importer override. --
        original = sys.modules.pop("pandas", None)
        try:
            with patch.dict(sys.modules, {"pandas": None}):
                doc = _new_doc()
                with pytest.raises(ImportError, match="pandas"):
                    tables.from_dataframe(doc, object())
        finally:
            if original is not None:
                sys.modules["pandas"] = original

    def it_raises_when_a_non_DataFrame_is_passed(self):
        pytest.importorskip("pandas")
        doc = _new_doc()
        with pytest.raises(ValueError, match="pandas.DataFrame"):
            tables.from_dataframe(doc, [{"a": 1}])

    def it_renders_a_basic_DataFrame(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame({"Name": ["X", "Y"], "Value": [10, 20]})

        table = tables.from_dataframe(doc, df, style="modern")

        assert [c.text for c in table.rows[0].cells] == ["Name", "Value"]
        assert [c.text for c in table.rows[1].cells] == ["X", "10"]
        assert [c.text for c in table.rows[2].cells] == ["Y", "20"]

    def it_right_aligns_numeric_columns_when_auto_format_is_True(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame({"Name": ["X"], "Value": [10]})

        table = tables.from_dataframe(
            doc, df, style="modern", auto_format=True
        )

        assert (
            table.rows[1].cells[1].paragraphs[0].alignment
            == WD_ALIGN_PARAGRAPH.RIGHT
        )
        assert (
            table.rows[1].cells[0].paragraphs[0].alignment
            == WD_ALIGN_PARAGRAPH.LEFT
        )

    def it_left_aligns_every_column_when_auto_format_is_False(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame({"Name": ["X"], "Value": [10]})

        table = tables.from_dataframe(
            doc, df, style="modern", auto_format=False
        )

        # Both columns should now read as text (str-coerced) and
        # therefore left-align.
        for col in (0, 1):
            assert (
                table.rows[1].cells[col].paragraphs[0].alignment
                == WD_ALIGN_PARAGRAPH.LEFT
            )

    def it_renders_dates_using_iso_format(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame(
            {"When": pd.to_datetime(["2026-05-29", "2026-05-30"])}
        )

        table = tables.from_dataframe(doc, df, style="modern")
        # pandas timestamps subclass datetime so the default formatter
        # falls through the datetime branch.
        assert table.rows[1].cells[0].text.startswith("2026-05-29")
        assert table.rows[2].cells[0].text.startswith("2026-05-30")

    def it_can_include_the_index_column(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame(
            {"Value": [10, 20]}, index=pd.Index(["a", "b"], name="key")
        )

        table = tables.from_dataframe(
            doc, df, style="modern", include_index=True
        )

        # Header has the index name then the data column.
        assert [c.text for c in table.rows[0].cells] == ["key", "Value"]
        # Body row 0: index value + data value.
        assert [c.text for c in table.rows[1].cells] == ["a", "10"]
        assert [c.text for c in table.rows[2].cells] == ["b", "20"]

    def it_omits_the_index_by_default(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame({"Value": [10, 20]})

        table = tables.from_dataframe(doc, df, style="modern")

        assert [c.text for c in table.rows[0].cells] == ["Value"]

    def it_renders_NaN_as_empty_string(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame({"X": [1.0, float("nan"), 3.0]})

        table = tables.from_dataframe(doc, df, style="modern")

        assert table.rows[1].cells[0].text == "1"
        assert table.rows[2].cells[0].text == ""
        assert table.rows[3].cells[0].text == "3"

    def it_applies_zebra_banding(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame({"Name": ["X", "Y", "Z"], "Value": [1, 2, 3]})

        table = tables.from_dataframe(doc, df, style="zebra")

        # Body row 1 (table.rows[2]) is the banded one.
        assert table.rows[2].cells[0].shading.fill_color == RGBColor(
            0xF2, 0xF2, 0xF2
        )

    def it_round_trips_a_DataFrame_through_save_and_load(self):
        pd = pytest.importorskip("pandas")
        doc = _new_doc()
        df = pd.DataFrame({"Name": ["X", "Y"], "Value": [10, 20]})

        tables.from_dataframe(doc, df, style="zebra")
        reloaded = _round_trip(doc)
        assert len(reloaded.tables) == 1
        assert reloaded.tables[0].rows[0].cells[0].text == "Name"
