"""Unit-test suite for ``docx.kit.export`` unified export entry-points."""

from __future__ import annotations

import zipfile

import pytest

from docx.api import Document as DocumentFactoryFn
from docx.kit import export


# -- Shared helpers -------------------------------------------------------


def _make_doc():
    """Return a fresh |Document| with a heading and a body paragraph."""
    document = DocumentFactoryFn()
    document.add_heading("Annual Report 2026", level=1)
    document.add_paragraph("Underlying performance was strong.")
    document.add_paragraph("Revenue grew by 12% year-on-year.")
    return document


# -- to_html --------------------------------------------------------------


class DescribeToHtml:
    """Unit-test suite for :func:`docx.kit.export.to_html`."""

    def it_writes_an_html_file(self, tmp_path):
        out = tmp_path / "out.html"

        export.to_html(_make_doc(), str(out))

        assert out.exists()
        text = out.read_text(encoding="utf-8")
        assert "<html" in text.lower()
        assert "Annual Report 2026" in text

    def it_accepts_a_pathlike(self, tmp_path):
        out = tmp_path / "out.html"

        export.to_html(_make_doc(), out)

        assert out.exists()

    def it_forwards_include_styles(self, tmp_path):
        out = tmp_path / "out.html"

        export.to_html(_make_doc(), str(out), include_styles=False)

        text = out.read_text(encoding="utf-8")
        # -- style block omitted when ``include_styles=False`` --
        assert "<style" not in text.lower()

    def it_rejects_unknown_kwargs(self, tmp_path):
        out = tmp_path / "out.html"

        with pytest.raises(TypeError, match="unexpected keyword"):
            export.to_html(_make_doc(), str(out), bogus=True)


# -- to_md ----------------------------------------------------------------


class DescribeToMd:
    """Unit-test suite for :func:`docx.kit.export.to_md`."""

    def it_writes_a_markdown_file(self, tmp_path):
        out = tmp_path / "out.md"

        export.to_md(_make_doc(), str(out))

        assert out.exists()
        text = out.read_text(encoding="utf-8")
        assert "# Annual Report 2026" in text
        assert "Underlying performance was strong." in text

    def it_accepts_a_pathlike(self, tmp_path):
        out = tmp_path / "out.md"

        export.to_md(_make_doc(), out)

        assert out.exists()

    def it_rejects_unknown_kwargs(self, tmp_path):
        out = tmp_path / "out.md"

        with pytest.raises(TypeError, match="unexpected keyword"):
            export.to_md(_make_doc(), str(out), bogus=True)


# -- to_pdf ---------------------------------------------------------------


class DescribeToPdf:
    """Unit-test suite for :func:`docx.kit.export.to_pdf`."""

    def it_writes_a_pdf_file(self, tmp_path):
        # -- Skips when ``reportlab`` isn't installed; the wrapper
        # -- itself is just a forwarder so all we need to assert is
        # -- that the underlying exporter is invoked. --
        pytest.importorskip("reportlab")
        out = tmp_path / "out.pdf"

        export.to_pdf(_make_doc(), str(out))

        assert out.exists()
        # -- PDF magic bytes --
        assert out.read_bytes().startswith(b"%PDF")

    def it_forwards_the_level_kwarg(self, tmp_path):
        pytest.importorskip("reportlab")
        out = tmp_path / "out.pdf"

        export.to_pdf(_make_doc(), str(out), level="2b")

        assert out.exists()

    def it_rejects_unknown_kwargs(self, tmp_path):
        out = tmp_path / "out.pdf"

        with pytest.raises(TypeError, match="unexpected keyword"):
            export.to_pdf(_make_doc(), str(out), bogus=True)


# -- to_epub --------------------------------------------------------------


class DescribeToEpub:
    """Unit-test suite for :func:`docx.kit.export.to_epub`."""

    def it_writes_an_epub_zip_file(self, tmp_path):
        out = tmp_path / "out.epub"

        export.to_epub(_make_doc(), str(out))

        assert out.exists()
        with zipfile.ZipFile(out, "r") as zf:
            names = zf.namelist()
        assert "mimetype" in names
        assert "META-INF/container.xml" in names
        assert "OEBPS/content.opf" in names
        assert "OEBPS/toc.ncx" in names
        assert "OEBPS/nav.xhtml" in names
        assert "OEBPS/chapter1.xhtml" in names
        assert "OEBPS/styles.css" in names

    def it_writes_the_mimetype_entry_uncompressed_and_first(self, tmp_path):
        out = tmp_path / "out.epub"

        export.to_epub(_make_doc(), str(out))

        with zipfile.ZipFile(out, "r") as zf:
            first = zf.infolist()[0]
            assert first.filename == "mimetype"
            assert first.compress_type == zipfile.ZIP_STORED
            assert zf.read("mimetype") == b"application/epub+zip"

    def it_carries_supplied_title_and_author_in_the_opf(self, tmp_path):
        out = tmp_path / "out.epub"

        export.to_epub(
            _make_doc(),
            str(out),
            title="Annual Report 2026",
            author="Jane Smith",
        )

        with zipfile.ZipFile(out, "r") as zf:
            opf = zf.read("OEBPS/content.opf").decode("utf-8")
        assert "<dc:title>Annual Report 2026</dc:title>" in opf
        assert "<dc:creator>Jane Smith</dc:creator>" in opf

    def it_falls_back_to_default_title_when_unsupplied(self, tmp_path):
        out = tmp_path / "out.epub"

        export.to_epub(_make_doc(), str(out))

        with zipfile.ZipFile(out, "r") as zf:
            opf = zf.read("OEBPS/content.opf").decode("utf-8")
        # -- core props blank in the bundled template -> default. --
        assert "<dc:title>" in opf

    def it_includes_the_document_html_in_the_chapter(self, tmp_path):
        out = tmp_path / "out.epub"

        export.to_epub(_make_doc(), str(out))

        with zipfile.ZipFile(out, "r") as zf:
            chapter = zf.read("OEBPS/chapter1.xhtml").decode("utf-8")
        assert "Annual Report 2026" in chapter
        assert "Underlying performance was strong." in chapter

    def it_emits_xhtml_doctype_in_the_chapter(self, tmp_path):
        out = tmp_path / "out.epub"

        export.to_epub(_make_doc(), str(out))

        with zipfile.ZipFile(out, "r") as zf:
            chapter = zf.read("OEBPS/chapter1.xhtml").decode("utf-8")
        assert "<?xml" in chapter
        assert "DTD XHTML 1.1" in chapter

    def it_emits_a_unique_uuid_identifier_per_call(self, tmp_path):
        a = tmp_path / "a.epub"
        b = tmp_path / "b.epub"

        export.to_epub(_make_doc(), str(a))
        export.to_epub(_make_doc(), str(b))

        with zipfile.ZipFile(a, "r") as zf:
            opf_a = zf.read("OEBPS/content.opf").decode("utf-8")
        with zipfile.ZipFile(b, "r") as zf:
            opf_b = zf.read("OEBPS/content.opf").decode("utf-8")
        # -- both carry urn:uuid:... but the UUIDs differ --
        assert "urn:uuid:" in opf_a and "urn:uuid:" in opf_b
        assert opf_a != opf_b


# -- to (dispatcher) ------------------------------------------------------


class DescribeToDispatcher:
    """Unit-test suite for :func:`docx.kit.export.to` extension dispatcher."""

    def it_dispatches_html_extension_to_to_html(self, tmp_path):
        out = tmp_path / "out.html"

        export.to(_make_doc(), str(out))

        assert out.exists()
        assert "<html" in out.read_text(encoding="utf-8").lower()

    def it_dispatches_htm_extension_to_to_html(self, tmp_path):
        out = tmp_path / "out.htm"

        export.to(_make_doc(), str(out))

        assert out.exists()

    def it_dispatches_md_extension_to_to_md(self, tmp_path):
        out = tmp_path / "out.md"

        export.to(_make_doc(), str(out))

        assert out.exists()
        assert "# Annual Report 2026" in out.read_text(encoding="utf-8")

    def it_dispatches_markdown_extension_to_to_md(self, tmp_path):
        out = tmp_path / "out.markdown"

        export.to(_make_doc(), str(out))

        assert out.exists()

    def it_dispatches_epub_extension_to_to_epub(self, tmp_path):
        out = tmp_path / "out.epub"

        export.to(_make_doc(), str(out))

        assert out.exists()
        with zipfile.ZipFile(out, "r") as zf:
            assert "mimetype" in zf.namelist()

    def it_dispatches_pdf_extension_to_to_pdf(self, tmp_path):
        pytest.importorskip("reportlab")
        out = tmp_path / "out.pdf"

        export.to(_make_doc(), str(out))

        assert out.exists()
        assert out.read_bytes().startswith(b"%PDF")

    def it_is_case_insensitive_on_the_extension(self, tmp_path):
        out = tmp_path / "out.HTML"

        export.to(_make_doc(), str(out))

        assert out.exists()

    def it_raises_for_unsupported_extensions(self, tmp_path):
        out = tmp_path / "out.rtf"

        with pytest.raises(ValueError, match="Unsupported export extension"):
            export.to(_make_doc(), str(out))

    def it_raises_when_no_extension(self, tmp_path):
        out = tmp_path / "out"

        with pytest.raises(ValueError, match="Unsupported export extension"):
            export.to(_make_doc(), str(out))

    def it_forwards_kwargs_to_the_dispatched_exporter(self, tmp_path):
        out = tmp_path / "out.html"

        export.to(_make_doc(), str(out), include_styles=False)

        text = out.read_text(encoding="utf-8")
        assert "<style" not in text.lower()


# -- Module-level re-exports ---------------------------------------------


class DescribeKitExportModule:
    """Verify the :mod:`docx.kit.export` public surface is wired correctly."""

    def it_exposes_the_five_entry_points(self):
        assert callable(export.to_pdf)
        assert callable(export.to_html)
        assert callable(export.to_md)
        assert callable(export.to_epub)
        assert callable(export.to)

    def it_is_re_exported_from_docx_kit(self):
        from docx.kit import export as re_exported

        assert re_exported is export
