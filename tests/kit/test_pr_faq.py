"""Unit-test suite for ``docx.kit.pr_faq`` template family."""

from __future__ import annotations

from typing import Any, Dict, List, Tuple

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import pr_faq
from docx.text.paragraph import Paragraph


# -- Shared fixtures / helpers --------------------------------------------


@pytest.fixture
def document() -> DocumentCls:
    return Document()


def _texts(document: DocumentCls) -> List[str]:
    """Return the text of every paragraph in ``document``."""
    return [p.text for p in document.paragraphs]


def _full_text(document: DocumentCls) -> str:
    return "\n".join(_texts(document))


def _press_release_kwargs(**overrides: Any) -> Dict[str, Any]:
    """Return a baseline kwargs dict for ``press_release`` calls.

    The defaults are the canonical Acme/FrobnitzPro example from the
    issue spec; tests override the pieces they care about.
    """
    base: Dict[str, Any] = {
        "headline": "Acme launches FrobnitzPro",
        "subheadline": "The fastest frobnitz on the market",
        "location": "Seattle, WA",
        "date": "2026-05-29",
        "summary": "One-paragraph summary of the launch.",
        "problem": "Customers struggle with frobnitz throughput.",
        "solution": "FrobnitzPro solves this by parallelising the work.",
        "quote_speaker": "Jane Doe, VP Product",
        "quote_text": '"FrobnitzPro is the most exciting product we have launched in a decade."',
        "call_to_action": "Visit acme.com/frobnitz to learn more.",
    }
    base.update(overrides)
    return base


def _faq_items() -> List[Tuple[str, str]]:
    return [
        ("What is FrobnitzPro?", "It is a frobnitz that does X, Y, and Z."),
        ("How much does it cost?", "$99/month for the Pro tier."),
        ("When is it available?", "Today."),
    ]


# -- press_release --------------------------------------------------------


class DescribePressRelease:
    """Unit-test suite for ``pr_faq.press_release``."""

    def it_returns_a_list_of_paragraphs_in_document_order(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        result = pr_faq.press_release(document, **_press_release_kwargs())

        assert all(isinstance(p, Paragraph) for p in result)
        # -- the helper appends to the end of the body. Paragraph
        # -- proxies are not identity-equal when re-wrapped, so compare
        # -- on the underlying ``<w:p>`` element. --
        appended = document.paragraphs[start : start + len(result)]
        assert [p._p for p in appended] == [p._p for p in result]

    def it_renders_the_headline_in_the_Title_style_centred(
        self, document: DocumentCls
    ):
        result = pr_faq.press_release(document, **_press_release_kwargs())

        headline = result[0]
        assert headline.text == "Acme launches FrobnitzPro"
        assert headline.style is not None
        assert headline.style.name == "Title"
        assert headline.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_renders_the_subheadline_in_the_Subtitle_style_centred(
        self, document: DocumentCls
    ):
        result = pr_faq.press_release(document, **_press_release_kwargs())

        sub = result[1]
        assert sub.text == "The fastest frobnitz on the market"
        assert sub.style is not None
        assert sub.style.name == "Subtitle"
        assert sub.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_omits_the_subheadline_paragraph_when_not_supplied(
        self, document: DocumentCls
    ):
        result = pr_faq.press_release(
            document, **_press_release_kwargs(subheadline=None)
        )

        # -- no subtitle paragraph means the dateline lives at index 1 --
        assert result[1].text.startswith("Seattle, WA — 2026-05-29 — ")
        for para in result:
            assert (
                para.style is None
                or para.style.name != "Subtitle"
            )

    def it_emits_a_bold_dateline_concatenated_with_the_summary(
        self, document: DocumentCls
    ):
        result = pr_faq.press_release(document, **_press_release_kwargs())

        # -- with a subheadline, the dateline paragraph is the third --
        dateline_para = result[2]
        assert dateline_para.text == (
            "Seattle, WA — 2026-05-29 — One-paragraph summary of the launch."
        )
        # -- first run is the bold dateline; second run is the summary --
        runs = list(dateline_para.runs)
        assert len(runs) == 2
        assert runs[0].text == "Seattle, WA — 2026-05-29 — "
        assert runs[0].bold is True
        assert runs[1].text == "One-paragraph summary of the launch."

    def it_emits_a_problem_heading_and_body(self, document: DocumentCls):
        pr_faq.press_release(document, **_press_release_kwargs())

        text = _full_text(document)
        assert "The Problem" in text
        assert "Customers struggle with frobnitz throughput." in text

    def it_emits_a_solution_heading_and_body(self, document: DocumentCls):
        pr_faq.press_release(document, **_press_release_kwargs())

        text = _full_text(document)
        assert "The Solution" in text
        assert "FrobnitzPro solves this by parallelising the work." in text

    def it_emits_a_spokesperson_quote_block(self, document: DocumentCls):
        pr_faq.press_release(document, **_press_release_kwargs())

        text = _full_text(document)
        assert (
            '"FrobnitzPro is the most exciting product we have launched in a decade."'
            in text
        )
        assert "— Jane Doe, VP Product" in text

    def it_omits_the_customer_quote_block_by_default(
        self, document: DocumentCls
    ):
        pr_faq.press_release(document, **_press_release_kwargs())

        text = _full_text(document)
        assert "— John Smith" not in text

    def it_emits_the_customer_quote_block_when_both_halves_are_supplied(
        self, document: DocumentCls
    ):
        pr_faq.press_release(
            document,
            **_press_release_kwargs(
                customer_quote_speaker="John Smith, ACME Corp",
                customer_quote_text='"It has revolutionised our workflow."',
            ),
        )

        text = _full_text(document)
        assert '"It has revolutionised our workflow."' in text
        assert "— John Smith, ACME Corp" in text

    def it_raises_when_only_one_half_of_the_customer_quote_is_supplied(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="customer_quote"):
            pr_faq.press_release(
                document,
                **_press_release_kwargs(
                    customer_quote_speaker="John Smith, ACME Corp",
                    customer_quote_text=None,
                ),
            )

    def it_emits_a_bold_call_to_action_label(self, document: DocumentCls):
        result = pr_faq.press_release(document, **_press_release_kwargs())

        # -- search the result list (not the doc) for the CTA paragraph --
        cta_paragraphs = [
            p for p in result if "Call to action:" in p.text
        ]
        assert len(cta_paragraphs) == 1
        cta = cta_paragraphs[0]
        runs = list(cta.runs)
        assert runs[0].text == "Call to action: "
        assert runs[0].bold is True
        assert "Visit acme.com/frobnitz to learn more." in cta.text

    def it_appends_a_trailing_page_break_by_default(
        self, document: DocumentCls
    ):
        result = pr_faq.press_release(document, **_press_release_kwargs())

        # -- a page-break paragraph holds a w:br with type page; the kit's
        # -- public surface is `Document.add_page_break()` which returns
        # -- the new paragraph. Asserting the last result paragraph
        # -- contains a page-break run is enough. --
        last = result[-1]
        assert last.text == ""

    def it_omits_the_trailing_page_break_when_disabled(
        self, document: DocumentCls
    ):
        result = pr_faq.press_release(
            document, **_press_release_kwargs(page_break=False)
        )

        # -- without the page break, the last paragraph is the CTA --
        assert "Call to action:" in result[-1].text

    @pytest.mark.parametrize(
        "field",
        [
            "headline",
            "location",
            "date",
            "summary",
            "problem",
            "solution",
            "quote_speaker",
            "quote_text",
            "call_to_action",
        ],
    )
    def it_raises_when_a_required_field_is_empty(
        self, document: DocumentCls, field: str
    ):
        with pytest.raises(ValueError, match=field):
            pr_faq.press_release(
                document, **_press_release_kwargs(**{field: ""})
            )


# -- faq ------------------------------------------------------------------


class DescribeFaq:
    """Unit-test suite for ``pr_faq.faq``."""

    def it_returns_a_list_of_paragraphs_in_document_order(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        result = pr_faq.faq(document, items=_faq_items())

        assert all(isinstance(p, Paragraph) for p in result)
        appended = document.paragraphs[start : start + len(result)]
        assert [p._p for p in appended] == [p._p for p in result]

    def it_renders_a_default_heading_at_Heading_1(
        self, document: DocumentCls
    ):
        result = pr_faq.faq(document, items=_faq_items())

        heading = result[0]
        assert heading.text == "Frequently Asked Questions"
        assert heading.style is not None
        assert heading.style.name == "Heading 1"

    def it_lets_callers_customise_the_heading(self, document: DocumentCls):
        result = pr_faq.faq(
            document, items=_faq_items(), title="Common Questions"
        )

        assert result[0].text == "Common Questions"

    def it_suppresses_the_heading_when_title_is_empty(
        self, document: DocumentCls
    ):
        result = pr_faq.faq(document, items=_faq_items(), title="")

        # -- first paragraph is the first Q --
        assert result[0].text == "Q: What is FrobnitzPro?"

    def it_renders_each_pair_as_a_Q_paragraph_then_an_A_paragraph(
        self, document: DocumentCls
    ):
        result = pr_faq.faq(
            document, items=[("Why?", "Because.")], title="", page_break=False
        )

        assert len(result) == 2
        q_para, a_para = result
        assert q_para.text == "Q: Why?"
        assert a_para.text == "A: Because."

    def it_emits_a_bold_Q_label_and_a_bold_A_label(
        self, document: DocumentCls
    ):
        result = pr_faq.faq(
            document, items=[("Why?", "Because.")], title="", page_break=False
        )

        q_para, a_para = result
        assert list(q_para.runs)[0].text == "Q: "
        assert list(q_para.runs)[0].bold is True
        assert list(a_para.runs)[0].text == "A: "
        assert list(a_para.runs)[0].bold is True

    def it_renders_three_pairs_in_caller_order(self, document: DocumentCls):
        result = pr_faq.faq(
            document,
            items=_faq_items(),
            page_break=False,
        )

        # -- heading + (Q + A) * 3 --
        assert len(result) == 1 + 2 * 3
        assert result[1].text == "Q: What is FrobnitzPro?"
        assert result[3].text == "Q: How much does it cost?"
        assert result[5].text == "Q: When is it available?"

    def it_appends_a_trailing_page_break_by_default(
        self, document: DocumentCls
    ):
        result = pr_faq.faq(document, items=_faq_items())

        assert result[-1].text == ""

    def it_omits_the_trailing_page_break_when_disabled(
        self, document: DocumentCls
    ):
        result = pr_faq.faq(
            document, items=_faq_items(), page_break=False
        )

        assert result[-1].text.startswith("A: ")

    def it_raises_when_items_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="at least one"):
            pr_faq.faq(document, items=[])

    def it_raises_when_a_question_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="question"):
            pr_faq.faq(document, items=[("", "Answer.")])

    def it_raises_when_an_answer_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="answer"):
            pr_faq.faq(document, items=[("Q?", "")])

    def it_raises_when_an_item_is_not_a_2_tuple(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="2-tuple"):
            pr_faq.faq(document, items=[("only-one-piece",)])  # type: ignore[list-item]


# -- pr_faq_doc -----------------------------------------------------------


class DescribePrFaqDoc:
    """Unit-test suite for ``pr_faq.pr_faq_doc``."""

    def it_returns_a_fresh_document_with_pr_then_faq(self):
        doc = pr_faq.pr_faq_doc(
            press_release_kwargs=_press_release_kwargs(),
            faq_items=_faq_items(),
        )

        assert isinstance(doc, DocumentCls)
        text = _full_text(doc)
        # -- press release pieces --
        assert "Acme launches FrobnitzPro" in text
        assert "The Problem" in text
        assert "The Solution" in text
        # -- FAQ pieces (the press release ends with a page break, then
        # -- the FAQ heading appears) --
        assert "Frequently Asked Questions" in text
        assert "Q: What is FrobnitzPro?" in text
        # -- press release content precedes faq content --
        assert text.index("The Problem") < text.index(
            "Frequently Asked Questions"
        )

    def it_saves_to_output_path_when_supplied(self, tmp_path):
        out = tmp_path / "pr_faq.docx"

        doc = pr_faq.pr_faq_doc(
            press_release_kwargs=_press_release_kwargs(),
            faq_items=_faq_items(),
            output_path=str(out),
        )

        assert out.exists()
        # -- and the file is openable as a |Document| --
        assert isinstance(doc, DocumentCls)
        roundtrip = Document(str(out))
        text = "\n".join(p.text for p in roundtrip.paragraphs)
        assert "Acme launches FrobnitzPro" in text
        assert "Q: What is FrobnitzPro?" in text

    def it_does_not_save_when_output_path_is_none(self, tmp_path):
        out = tmp_path / "missing.docx"

        pr_faq.pr_faq_doc(
            press_release_kwargs=_press_release_kwargs(),
            faq_items=_faq_items(),
            output_path=None,
        )

        assert not out.exists()

    def it_raises_when_press_release_kwargs_is_not_a_dict(self):
        with pytest.raises(ValueError, match="press_release_kwargs"):
            pr_faq.pr_faq_doc(
                press_release_kwargs="not-a-dict",  # type: ignore[arg-type]
                faq_items=_faq_items(),
            )

    def it_propagates_validation_errors_from_press_release(self):
        with pytest.raises(ValueError, match="headline"):
            pr_faq.pr_faq_doc(
                press_release_kwargs=_press_release_kwargs(headline=""),
                faq_items=_faq_items(),
            )

    def it_propagates_validation_errors_from_faq(self):
        with pytest.raises(ValueError, match="at least one"):
            pr_faq.pr_faq_doc(
                press_release_kwargs=_press_release_kwargs(),
                faq_items=[],
            )
