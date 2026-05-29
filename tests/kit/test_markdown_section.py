"""Unit-test suite for ``docx.kit.markdown_section.add`` (issue #54)."""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import markdown_section
from docx.kit.markdown_section import add
from docx.table import Table
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Helpers


@pytest.fixture
def document() -> DocumentCls:
    return Document()


def _png_bytes() -> bytes:
    """Return the bytes of a 1x1 transparent PNG.

    A real PNG (not a stub) so :meth:`Run.add_picture` can probe DPI /
    dimensions without raising. Hand-rolled via ``zlib`` + ``struct``
    so the test stays stdlib-only — Pillow is not a python-docx
    dependency.
    """
    sig = b"\x89PNG\r\n\x1a\n"

    def _chunk(tag: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", crc)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 6, 0, 0, 0)
    raw = b"\x00\x00\x00\x00\x00"  # filter-byte + RGBA(0,0,0,0)
    idat = zlib.compress(raw)
    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


@pytest.fixture
def png_path(tmp_path: Path) -> Path:
    p = tmp_path / "tiny.png"
    p.write_bytes(_png_bytes())
    return p


# ---------------------------------------------------------------------------
# Public API contract


class DescribeAddPublicAPI:
    def it_is_exposed_via_docx_kit(self):
        from docx.kit import markdown_section as ms

        assert ms is markdown_section
        assert callable(ms.add)

    def it_returns_an_empty_list_for_empty_input(self, document: DocumentCls):
        result = add(document, "")

        assert result == []

    def it_returns_an_empty_list_for_whitespace_only_input(
        self, document: DocumentCls
    ):
        result = add(document, "   \n\n   \n")

        assert result == []

    def it_raises_on_None_input(self, document: DocumentCls):
        with pytest.raises(TypeError):
            add(document, None)  # type: ignore[arg-type]

    def it_returns_paragraphs_in_document_order(self, document: DocumentCls):
        result = add(document, "first\n\nsecond\n\nthird")

        assert len(result) == 3
        assert all(isinstance(p, Paragraph) for p in result)
        assert [p.text for p in result] == ["first", "second", "third"]

    def it_appends_to_the_existing_body(self, document: DocumentCls):
        document.add_paragraph("pre-existing")
        before = len(document.paragraphs)

        result = add(document, "hello")

        assert len(document.paragraphs) == before + 1
        assert result[0].text == "hello"


# ---------------------------------------------------------------------------
# Headings


class DescribeHeadings:
    def it_renders_h1_through_h6(self, document: DocumentCls):
        md = "# H1\n\n## H2\n\n### H3\n\n#### H4\n\n##### H5\n\n###### H6"
        result = add(document, md)

        assert len(result) == 6
        for level, para in enumerate(result, start=1):
            assert para.text == "H%d" % level
            assert para.style.name == "Heading %d" % level

    def it_preserves_inline_formatting_in_a_heading(self, document: DocumentCls):
        result = add(document, "# **Bold** title")

        para = result[0]
        assert para.text == "Bold title"
        assert para.runs[0].bold is True

    def it_clamps_excess_hashes_to_h6_via_paragraph_text(
        self, document: DocumentCls
    ):
        # ``####### x`` (seven hashes) is *not* a valid ATX heading;
        # it should fall through to a regular paragraph.
        result = add(document, "####### too many")

        assert result[0].style.name == "Normal"
        assert "too many" in result[0].text


# ---------------------------------------------------------------------------
# Paragraphs and inline runs


class DescribeParagraphs:
    def it_renders_a_plain_paragraph(self, document: DocumentCls):
        result = add(document, "Just some text.")

        assert len(result) == 1
        assert result[0].text == "Just some text."

    def it_renders_bold_and_italic_runs(self, document: DocumentCls):
        result = add(document, "A **bold** and *italic* mix.")

        para = result[0]
        bold = [r for r in para.runs if r.bold]
        italic = [r for r in para.runs if r.italic]
        assert any(r.text == "bold" for r in bold)
        assert any(r.text == "italic" for r in italic)

    def it_renders_inline_code_with_courier_font(
        self, document: DocumentCls
    ):
        result = add(document, "Run `print()` to log.")

        code_runs = [r for r in result[0].runs if r.font.name == "Courier New"]
        assert any(r.text == "print()" for r in code_runs)

    def it_renders_a_hyperlink(self, document: DocumentCls):
        result = add(document, "See [docs](https://example.com/x).")

        para = result[0]
        # The hyperlink is appended as a Hyperlink object, surfaced via
        # paragraph.text including the link label.
        assert "docs" in para.text
        # And the relationship is registered on the part.
        rels = list(para.part.rels.values())
        assert any(
            getattr(r, "target_ref", "") == "https://example.com/x" for r in rels
        )

    def it_collapses_multi_line_paragraphs_with_a_space(
        self, document: DocumentCls
    ):
        result = add(document, "line one\nline two\nline three")

        assert len(result) == 1
        assert result[0].text == "line one line two line three"


# ---------------------------------------------------------------------------
# Lists


class DescribeLists:
    def it_renders_a_bullet_list(self, document: DocumentCls):
        result = add(document, "- one\n- two\n- three")

        assert len(result) == 3
        for para, expected in zip(result, ("one", "two", "three")):
            assert para.text == expected
            assert para.style.name == "List Bullet"

    def it_accepts_asterisk_and_plus_bullet_markers(
        self, document: DocumentCls
    ):
        result = add(document, "* a\n+ b\n- c")

        assert [p.text for p in result] == ["a", "b", "c"]
        assert all(p.style.name == "List Bullet" for p in result)

    def it_renders_a_numbered_list(self, document: DocumentCls):
        result = add(document, "1. alpha\n2. beta\n3. gamma")

        assert len(result) == 3
        for para, expected in zip(result, ("alpha", "beta", "gamma")):
            assert para.text == expected
            assert para.style.name == "List Number"

    def it_renders_inline_formatting_inside_list_items(
        self, document: DocumentCls
    ):
        result = add(document, "- a **bold** item\n- a *italic* item")

        bold_runs = [r for r in result[0].runs if r.bold]
        italic_runs = [r for r in result[1].runs if r.italic]
        assert any(r.text == "bold" for r in bold_runs)
        assert any(r.text == "italic" for r in italic_runs)


# ---------------------------------------------------------------------------
# Tables


class DescribeTables:
    def it_renders_a_pipe_table(self, document: DocumentCls):
        md = "| Col | Other |\n|-----|-------|\n| 1 | 2 |\n| 3 | 4 |"
        result = add(document, md)

        tables = [r for r in result if isinstance(r, Table)]
        assert len(tables) == 1
        table = tables[0]
        assert len(table.rows) == 3  # header + 2 body rows
        assert len(table.rows[0].cells) == 2
        assert table.rows[0].cells[0].text == "Col"
        assert table.rows[0].cells[1].text == "Other"
        assert table.rows[1].cells[0].text == "1"
        assert table.rows[2].cells[1].text == "4"

    def it_bolds_the_header_row(self, document: DocumentCls):
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        result = add(document, md)

        table = next(r for r in result if isinstance(r, Table))
        header_cell = table.rows[0].cells[0]
        runs = header_cell.paragraphs[0].runs
        assert runs and runs[0].bold is True

    def it_renders_inline_formatting_inside_a_cell(
        self, document: DocumentCls
    ):
        md = "| Header |\n|--------|\n| **bold** cell |"
        result = add(document, md)

        table = next(r for r in result if isinstance(r, Table))
        body_cell = table.rows[1].cells[0]
        runs = body_cell.paragraphs[0].runs
        assert any(r.bold and r.text == "bold" for r in runs)


# ---------------------------------------------------------------------------
# Blockquotes


class DescribeBlockquotes:
    def it_renders_a_single_line_blockquote(self, document: DocumentCls):
        result = add(document, "> a wise quote")

        para = result[0]
        assert para.text == "a wise quote"
        assert para.style.name in {"Intense Quote", "Quote", "Normal"}

    def it_collapses_a_multi_line_blockquote_to_one_paragraph(
        self, document: DocumentCls
    ):
        result = add(document, "> first line\n> second line")

        assert len(result) == 1
        assert "first line" in result[0].text
        assert "second line" in result[0].text


# ---------------------------------------------------------------------------
# Fenced code blocks


class DescribeFencedCodeBlocks:
    def it_renders_a_fenced_code_block_with_courier_font(
        self, document: DocumentCls
    ):
        md = "```python\nprint('hello')\n```"
        result = add(document, md)

        para = result[0]
        runs = para.runs
        assert runs[0].font.name == "Courier New"
        assert "print('hello')" in para.text

    def it_preserves_multi_line_code(self, document: DocumentCls):
        md = "```\nline 1\nline 2\nline 3\n```"
        result = add(document, md)

        run = result[0].runs[0]
        assert "line 1" in run.text
        assert "line 2" in run.text
        assert "line 3" in run.text


# ---------------------------------------------------------------------------
# Horizontal rules


class DescribeHorizontalRules:
    def it_renders_a_dash_hr(self, document: DocumentCls):
        result = add(document, "before\n\n---\n\nafter")

        # before, hr, after — the hr is the middle paragraph.
        assert len(result) == 3
        assert "―" in result[1].text  # em-dash row

    def it_accepts_asterisk_and_underscore_hr_forms(
        self, document: DocumentCls
    ):
        result = add(document, "***\n\n___")

        # Two rules, each rendered as a row of em-dashes.
        assert len(result) == 2
        assert "―" in result[0].text
        assert "―" in result[1].text


# ---------------------------------------------------------------------------
# Inline images


class DescribeInlineImages:
    def it_embeds_a_local_image_when_inline_images_is_true(
        self, document: DocumentCls, png_path: Path
    ):
        md = "An ![alt](%s) image." % str(png_path)
        result = add(document, md, inline_images=True)

        # The picture lands as a relationship on the document part.
        rels = list(document.part.rels.values())
        image_rels = [r for r in rels if "image" in (r.reltype or "").lower()]
        assert len(image_rels) >= 1
        # No bracketed placeholder text when the embed succeeded.
        assert "[image:" not in result[0].text

    def it_falls_back_to_a_placeholder_when_inline_images_is_false(
        self, document: DocumentCls, png_path: Path
    ):
        md = "An ![alt-text](%s) image." % str(png_path)
        result = add(document, md, inline_images=False)

        assert "[image: alt-text]" in result[0].text

    def it_falls_back_to_a_placeholder_for_a_missing_path(
        self, document: DocumentCls
    ):
        result = add(document, "Missing ![alt](does-not-exist.png).")

        assert "[image: alt]" in result[0].text

    def it_falls_back_to_a_placeholder_for_a_remote_url(
        self, document: DocumentCls
    ):
        result = add(document, "Remote ![logo](https://example.com/x.png).")

        assert "[image: logo]" in result[0].text


# ---------------------------------------------------------------------------
# Style prefix


class DescribeStylePrefix:
    def it_uses_builtin_styles_when_no_prefix_is_supplied(
        self, document: DocumentCls
    ):
        result = add(document, "# Title\n\nbody", style_prefix="")

        assert result[0].style.name == "Heading 1"
        assert result[1].style.name == "Normal"

    def it_prefers_the_prefixed_style_when_present(
        self, document: DocumentCls
    ):
        document.styles.add_style("MD Heading 1", 1)  # WD_STYLE_TYPE.PARAGRAPH
        document.styles.add_style("MD Body", 1)

        result = add(document, "# Title\n\nbody")

        assert result[0].style.name == "MD Heading 1"
        assert result[1].style.name == "MD Body"

    def it_falls_back_to_builtin_when_only_some_prefixed_variants_are_defined(
        self, document: DocumentCls
    ):
        document.styles.add_style("MD Heading 1", 1)

        result = add(document, "# Title\n\nbody\n\n## Sub")

        assert result[0].style.name == "MD Heading 1"
        # body and sub fall back to builtins.
        assert result[1].style.name == "Normal"
        assert result[2].style.name == "Heading 2"


# ---------------------------------------------------------------------------
# Round-trip with Document.to_markdown


class DescribeRoundTrip:
    def it_round_trips_a_simple_document(self, document: DocumentCls):
        original = (
            "# Section title\n\n"
            "A paragraph with **bold** and *italic* text.\n\n"
            "- bullet 1\n"
            "- bullet 2\n\n"
            "1. step 1\n"
            "2. step 2\n"
        )
        result = add(document, original)

        # Save to a buffer and reopen — the doc must still load and
        # carry the expected content.
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = Document(buf)

        texts = [p.text for p in reopened.paragraphs]
        assert "Section title" in texts
        assert any("bold" in t and "italic" in t for t in texts)
        assert "bullet 1" in texts
        assert "bullet 2" in texts
        assert "step 1" in texts
        assert "step 2" in texts

        # And to_markdown round-trips back to a document containing
        # the same headings / list markers.
        md = reopened.to_markdown()
        assert "# Section title" in md
        assert "**bold**" in md
        assert "_italic_" in md or "*italic*" in md

    def it_renders_the_issue_example_document(self, document: DocumentCls):
        md = (
            "# Section title\n\n"
            "A paragraph with **bold** and *italic* text.\n\n"
            "- bullet 1\n"
            "- bullet 2\n\n"
            "| Col | Other |\n"
            "|-----|-------|\n"
            "| 1   | 2     |\n"
        )
        result = add(document, md)

        # Heading + paragraph + 2 bullets + table.
        assert sum(1 for r in result if isinstance(r, Paragraph)) == 4
        assert sum(1 for r in result if isinstance(r, Table)) == 1
