"""Unit-test suite for ``docx.kit.code_block`` helpers (issue #293)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import code_block
from docx.kit.code_block import (
    add,
    bash,
    css,
    go,
    html,
    javascript,
    rust,
    sql,
    typescript,
    xml,
    yaml,
)
from docx.kit.code_block import json as code_json
from docx.kit.code_block import python as code_python
from docx.shared import Pt, RGBColor
from docx.table import Table


@pytest.fixture
def document() -> DocumentCls:
    return Document()


# -- Fixtures -----------------------------------------------------------------

PY_SOURCE = (
    "def hello():\n"
    "    print(\"Hello world!\")\n"
)

JSON_SOURCE = '{"name": "ada", "age": 36}'

BASH_SOURCE = "ls -la /tmp\necho ok"


# -----------------------------------------------------------------------------
# Highlighted (pygments-installed) path
# -----------------------------------------------------------------------------


class DescribeAddHighlighted:
    """Behavioural tests for the highlighting path (pygments installed)."""

    def it_returns_a_table_object(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = add(document, PY_SOURCE, lang="python")

        assert isinstance(result, Table)

    def it_appends_a_single_row_table_without_line_numbers(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python")

        assert len(table.rows) == 1
        assert len(table.rows[0].cells) == 1

    def it_adds_a_gutter_column_when_line_numbers_is_true(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python", line_numbers=True)

        assert len(table.rows[0].cells) == 2

    def it_paints_the_body_cell_with_a_shaded_background(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python")

        body = table.rows[0].cells[0]
        assert body.shading.fill_color is not None

    def it_uses_the_theme_background_when_pygments_provides_one(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python", theme="monokai")

        body = table.rows[0].cells[0]
        # -- monokai background is #272822
        assert body.shading.fill_color == RGBColor(0x27, 0x28, 0x22)

    def it_falls_back_to_light_grey_when_theme_has_no_background(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python", theme="default")

        body = table.rows[0].cells[0]
        # -- the `default` style declares #f8f8f8 — a *real* very-light
        # -- grey; either that or our fallback (#F5F5F5) is acceptable
        # -- so we just assert "very pale" by checking each component.
        rgb = body.shading.fill_color
        assert rgb is not None
        assert rgb[0] >= 0xF0 and rgb[1] >= 0xF0 and rgb[2] >= 0xF0

    def it_renders_one_paragraph_per_source_line(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python")

        body = table.rows[0].cells[0]
        # -- Two source lines → two paragraphs.
        assert len(body.paragraphs) == 2

    def it_strips_leading_blank_line_from_a_heredoc(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(
            document,
            "\ndef f():\n    return 1\n",
            lang="python",
        )

        body = table.rows[0].cells[0]
        assert len(body.paragraphs) == 2
        assert "def" in body.paragraphs[0].text

    def it_renders_runs_with_the_monospace_font(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python")

        body = table.rows[0].cells[0]
        runs = body.paragraphs[0].runs
        assert len(runs) >= 1
        assert all(run.font.name == "Consolas" for run in runs)
        assert all(run.font.size == Pt(9) for run in runs)

    def it_honours_a_caller_supplied_monospace_font(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(
            document,
            PY_SOURCE,
            lang="python",
            monospace_font="Cascadia Code",
        )

        runs = table.rows[0].cells[0].paragraphs[0].runs
        assert all(run.font.name == "Cascadia Code" for run in runs)

    def it_colours_keyword_tokens(self, document: DocumentCls):
        pytest.importorskip("pygments")

        table = add(document, PY_SOURCE, lang="python", theme="monokai")

        body = table.rows[0].cells[0]
        # -- The first run on the first line is the `def` keyword;
        # -- pygments + monokai colour Token.Keyword as #66d9ef.
        first_run = body.paragraphs[0].runs[0]
        assert first_run.text == "def"
        assert first_run.font.color.rgb == RGBColor(0x66, 0xD9, 0xEF)

    def it_falls_back_to_text_lexer_for_an_unknown_language(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        # -- Unknown lexer should not raise; just renders as plain text.
        table = add(document, "anything goes\nhere", lang="not-a-real-lang")

        assert isinstance(table, Table)
        body = table.rows[0].cells[0]
        assert len(body.paragraphs) == 2

    def it_falls_back_to_default_theme_when_theme_unknown(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        # -- Unknown theme should not raise.
        table = add(document, PY_SOURCE, lang="python", theme="not-a-theme")

        assert isinstance(table, Table)

    def it_renders_the_line_number_gutter_with_right_justified_numbers(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        table = add(
            document,
            "a = 1\nb = 2\nc = 3",
            lang="python",
            line_numbers=True,
        )

        gutter = table.rows[0].cells[0]
        # -- Three lines → three gutter paragraphs, right-justified to 3 chars.
        assert len(gutter.paragraphs) == 3
        assert gutter.paragraphs[0].text == "  1"
        assert gutter.paragraphs[1].text == "  2"
        assert gutter.paragraphs[2].text == "  3"


# -----------------------------------------------------------------------------
# Soft-fallback path (pygments missing)
# -----------------------------------------------------------------------------


class DescribeAddFallback:
    """Behavioural tests for the soft-fallback (no-pygments) path."""

    def it_does_not_raise_when_pygments_is_missing(
        self, document: DocumentCls, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setattr(code_block, "HAS_PYGMENTS", False)

        # -- Must not raise.
        table = add(document, PY_SOURCE, lang="python")

        assert isinstance(table, Table)

    def it_renders_a_plain_monospace_block_without_pygments(
        self, document: DocumentCls, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setattr(code_block, "HAS_PYGMENTS", False)

        table = add(document, PY_SOURCE, lang="python")

        body = table.rows[0].cells[0]
        # -- One paragraph per source line, monospace font, fallback colour.
        assert len(body.paragraphs) == 2
        for para in body.paragraphs:
            assert all(run.font.name == "Consolas" for run in para.runs)

    def it_still_paints_a_shaded_fallback_background(
        self, document: DocumentCls, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setattr(code_block, "HAS_PYGMENTS", False)

        table = add(document, PY_SOURCE, lang="python")

        body = table.rows[0].cells[0]
        # -- Falls back to #F5F5F5 (very light grey).
        assert body.shading.fill_color == RGBColor(0xF5, 0xF5, 0xF5)

    def it_still_renders_a_line_number_gutter_without_pygments(
        self, document: DocumentCls, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setattr(code_block, "HAS_PYGMENTS", False)

        table = add(
            document,
            BASH_SOURCE,
            lang="bash",
            line_numbers=True,
        )

        assert len(table.rows[0].cells) == 2
        gutter = table.rows[0].cells[0]
        assert len(gutter.paragraphs) == 2

    def it_silently_ignores_theme_when_pygments_is_missing(
        self, document: DocumentCls, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setattr(code_block, "HAS_PYGMENTS", False)

        # -- Unknown theme name, missing pygments → still no exception.
        table = add(
            document,
            PY_SOURCE,
            lang="python",
            theme="solarized-dark",
        )

        assert isinstance(table, Table)


# -----------------------------------------------------------------------------
# Convenience wrappers
# -----------------------------------------------------------------------------


class DescribeConvenienceWrappers:
    """Each wrapper binds `lang` and forwards every other kwarg to `add`."""

    def it_python_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = code_python(document, "x = 1")

        assert isinstance(result, Table)

    def it_bash_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = bash(document, "ls -la")

        assert isinstance(result, Table)

    def it_json_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = code_json(document, JSON_SOURCE)

        assert isinstance(result, Table)

    def it_yaml_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = yaml(document, "name: ada\nage: 36")

        assert isinstance(result, Table)

    def it_sql_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = sql(document, "SELECT 1")

        assert isinstance(result, Table)

    def it_javascript_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = javascript(document, "const x = 1;")

        assert isinstance(result, Table)

    def it_typescript_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = typescript(document, "const x: number = 1;")

        assert isinstance(result, Table)

    def it_rust_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = rust(document, "fn main() {}")

        assert isinstance(result, Table)

    def it_go_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = go(document, "package main")

        assert isinstance(result, Table)

    def it_html_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = html(document, "<p>hi</p>")

        assert isinstance(result, Table)

    def it_css_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = css(document, "body { color: red; }")

        assert isinstance(result, Table)

    def it_xml_wrapper_returns_a_table(self, document: DocumentCls):
        pytest.importorskip("pygments")

        result = xml(document, "<root><a/></root>")

        assert isinstance(result, Table)

    def it_forwards_line_numbers_through_a_wrapper(
        self, document: DocumentCls
    ):
        pytest.importorskip("pygments")

        result = code_json(document, JSON_SOURCE, line_numbers=True)

        assert len(result.rows[0].cells) == 2


# -----------------------------------------------------------------------------
# Module-level surface
# -----------------------------------------------------------------------------


class DescribeModuleSurface:
    def it_exposes_the_HAS_PYGMENTS_flag(self):
        assert hasattr(code_block, "HAS_PYGMENTS")
        assert isinstance(code_block.HAS_PYGMENTS, bool)

    def it_exports_every_language_wrapper_in_dunder_all(self):
        for name in (
            "add",
            "bash",
            "css",
            "go",
            "html",
            "javascript",
            "json",
            "python",
            "rust",
            "sql",
            "typescript",
            "xml",
            "yaml",
        ):
            assert name in code_block.__all__

    def it_round_trips_a_document_with_a_code_block(
        self, document: DocumentCls, tmp_path
    ):
        pytest.importorskip("pygments")

        add(document, PY_SOURCE, lang="python", line_numbers=True)
        out = tmp_path / "out.docx"
        document.save(str(out))

        # -- Re-open and verify the table survived the round-trip.
        reopened = Document(str(out))
        assert len(reopened.tables) >= 1
