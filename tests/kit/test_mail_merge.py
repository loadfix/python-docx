"""Unit-test suite for ``docx.kit.mail_merge`` (issue #67)."""

from __future__ import annotations

import os
from io import BytesIO

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import mail_merge
from docx.kit.mail_merge import merge


@pytest.fixture
def template_doc() -> DocumentCls:
    """A fresh template carrying two token-bearing paragraphs."""
    doc = Document()
    doc.add_paragraph(
        "Dear {first_name},",
        bind_to={"first_name": "__placeholder__"},
    )
    doc.add_paragraph(
        "Welcome to your {role} role at salary {salary}.",
        bind_to={"role": "__r__", "salary": "__s__"},
    )
    return doc


@pytest.fixture
def template_path(tmp_path, template_doc: DocumentCls) -> str:
    """Persist `template_doc` to disk and return the path."""
    path = tmp_path / "offer-template.docx"
    template_doc.save(str(path))
    return str(path)


@pytest.fixture
def records() -> list:
    return [
        {"first_name": "Alice", "role": "Engineer", "salary": "$120k"},
        {"first_name": "Bob", "role": "Manager", "salary": "$140k"},
        {"first_name": "Carol", "role": "Designer", "salary": "$110k"},
    ]


class DescribeMerge:
    """Behavioural tests for :func:`docx.kit.mail_merge.merge`."""

    def it_returns_one_document_per_record(self, template_path: str, records: list):
        docs = merge(template=template_path, records=records)

        assert isinstance(docs, list)
        assert len(docs) == len(records)
        for doc in docs:
            assert isinstance(doc, DocumentCls)

    def it_resolves_each_records_tokens_independently(
        self, template_path: str, records: list
    ):
        docs = merge(template=template_path, records=records)

        # -- each rendered doc carries its own record's substitutions --
        assert "Alice" in docs[0].paragraphs[0].text
        assert "Engineer" in docs[0].paragraphs[1].text
        assert "$120k" in docs[0].paragraphs[1].text

        assert "Bob" in docs[1].paragraphs[0].text
        assert "Manager" in docs[1].paragraphs[1].text

        assert "Carol" in docs[2].paragraphs[0].text
        assert "Designer" in docs[2].paragraphs[1].text

        # -- and crucially, no cross-contamination between rows --
        assert "Bob" not in docs[0].paragraphs[0].text
        assert "Alice" not in docs[1].paragraphs[0].text

    def it_accepts_a_path_template(self, template_path: str, records: list):
        docs = merge(template=template_path, records=records[:1])
        assert "Alice" in docs[0].paragraphs[0].text

    def it_accepts_a_pathlike_template(self, tmp_path, template_doc, records: list):
        path = tmp_path / "via-pathlike.docx"
        template_doc.save(str(path))

        docs = merge(template=path, records=records[:1])  # PosixPath, not str
        assert "Alice" in docs[0].paragraphs[0].text

    def it_accepts_a_filelike_template(self, template_path: str, records: list):
        with open(template_path, "rb") as fh:
            docs = merge(template=fh, records=records[:1])

        assert "Alice" in docs[0].paragraphs[0].text

    def it_accepts_a_preloaded_Document_template(
        self, template_doc: DocumentCls, records: list
    ):
        docs = merge(template=template_doc, records=records[:2])

        assert len(docs) == 2
        assert "Alice" in docs[0].paragraphs[0].text
        assert "Bob" in docs[1].paragraphs[0].text

    def it_isolates_each_record_render_from_subsequent_calls(
        self, template_path: str, records: list
    ):
        # -- a second merge() call on the same template + different
        # -- records produces fresh output that doesn't carry leakage
        # -- from the first call's bound record. --
        first = merge(template=template_path, records=records[:1])
        second = merge(template=template_path, records=[{"first_name": "Eve",
                                                          "role": "CEO",
                                                          "salary": "$200k"}])

        assert "Alice" in first[0].paragraphs[0].text
        assert "Eve" in second[0].paragraphs[0].text
        # -- and the second call's results don't carry "Alice" --
        assert "Alice" not in second[0].paragraphs[0].text

    def it_accepts_an_empty_records_iterable(self, template_path: str):
        docs = merge(template=template_path, records=[])
        assert docs == []

    def it_accepts_a_generator_for_records(self, template_path: str, records: list):
        docs = merge(template=template_path, records=iter(records))
        assert len(docs) == len(records)

    def it_resolves_iteration_index_via_i_token(self, tmp_path, records: list):
        # -- author a template that uses {i} --
        doc = Document()
        doc.add_paragraph("Letter {i}: Hello {first_name}", bind_to={"first_name": "x"})
        path = tmp_path / "with-i.docx"
        doc.save(str(path))

        docs = merge(template=str(path), records=records)

        assert "Letter 0: Hello Alice" in docs[0].paragraphs[0].text
        assert "Letter 1: Hello Bob" in docs[1].paragraphs[0].text
        assert "Letter 2: Hello Carol" in docs[2].paragraphs[0].text

    def it_leaves_unknown_tokens_literal(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Hi {first_name}, ref {missing_field}", bind_to={"x": 1})
        path = tmp_path / "partial.docx"
        doc.save(str(path))

        docs = merge(template=str(path), records=[{"first_name": "Alice"}])

        assert "Alice" in docs[0].paragraphs[0].text
        assert "{missing_field}" in docs[0].paragraphs[0].text


class DescribeMergeOutputDir:
    """Behavioural tests for the ``output_dir=`` direct-to-disk mode."""

    def it_writes_each_doc_to_disk_using_the_filename_template(
        self, template_path: str, records: list, tmp_path
    ):
        out_dir = tmp_path / "out"

        result = merge(
            template=template_path,
            records=records,
            output_dir=str(out_dir),
            filename_template="offer-{first_name}.docx",
        )

        # -- result is a list of paths, in record order --
        assert len(result) == 3
        for path in result:
            assert os.path.isfile(path)
        assert os.path.basename(result[0]) == "offer-Alice.docx"
        assert os.path.basename(result[1]) == "offer-Bob.docx"
        assert os.path.basename(result[2]) == "offer-Carol.docx"

    def it_round_trips_each_written_doc_with_the_resolved_tokens(
        self, template_path: str, records: list, tmp_path
    ):
        out_dir = tmp_path / "out"

        paths = merge(
            template=template_path,
            records=records,
            output_dir=str(out_dir),
            filename_template="offer-{first_name}.docx",
        )

        # -- reload each written file and check the tokens resolved --
        reloaded = Document(paths[0])
        assert "Alice" in reloaded.paragraphs[0].text
        assert "Engineer" in reloaded.paragraphs[1].text

    def it_creates_the_output_dir_if_missing(
        self, template_path: str, records: list, tmp_path
    ):
        out_dir = tmp_path / "deep" / "nested" / "out"
        assert not out_dir.exists()

        merge(
            template=template_path,
            records=records[:1],
            output_dir=str(out_dir),
            filename_template="offer-{first_name}.docx",
        )

        assert out_dir.is_dir()

    def it_accepts_a_pathlike_output_dir(
        self, template_path: str, records: list, tmp_path
    ):
        out_dir = tmp_path / "pathlike-out"

        paths = merge(
            template=template_path,
            records=records[:1],
            output_dir=out_dir,  # PosixPath, not str
            filename_template="offer-{first_name}.docx",
        )

        assert os.path.isfile(paths[0])

    def it_supports_i_in_the_filename_template(
        self, template_path: str, records: list, tmp_path
    ):
        out_dir = tmp_path / "by-index"

        paths = merge(
            template=template_path,
            records=records,
            output_dir=str(out_dir),
            filename_template="row-{i:02d}-{first_name}.docx",
        )

        assert os.path.basename(paths[0]) == "row-00-Alice.docx"
        assert os.path.basename(paths[1]) == "row-01-Bob.docx"
        assert os.path.basename(paths[2]) == "row-02-Carol.docx"

    def it_raises_when_output_dir_supplied_without_filename_template(
        self, template_path: str, records: list, tmp_path
    ):
        with pytest.raises(ValueError, match="filename_template is required"):
            merge(
                template=template_path,
                records=records,
                output_dir=str(tmp_path / "out"),
            )

    def it_raises_when_filename_template_supplied_without_output_dir(
        self, template_path: str, records: list
    ):
        with pytest.raises(ValueError, match="filename_template requires output_dir"):
            merge(
                template=template_path,
                records=records,
                filename_template="offer-{first_name}.docx",
            )

    def it_raises_KeyError_when_filename_field_is_missing_from_a_record(
        self, template_path: str, tmp_path
    ):
        out_dir = tmp_path / "out"

        with pytest.raises(KeyError):
            merge(
                template=template_path,
                records=[{"first_name": "Alice"}],  # no ``role``
                output_dir=str(out_dir),
                filename_template="offer-{first_name}-{role}.docx",
            )


class DescribeMailMergeModule:
    """Sanity tests for the public surface."""

    def it_re_exports_merge_from_the_kit_namespace(self):
        from docx.kit.mail_merge import merge as merge_a
        from docx.kit import mail_merge as mod

        assert mod.merge is merge_a

    def it_lists_merge_in___all__(self):
        assert "merge" in mail_merge.__all__
