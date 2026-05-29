"""Unit-test suite for ``docx.kit.headers`` (issue #288)."""

from __future__ import annotations

from io import BytesIO

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import headers
from docx.kit.headers import cover_page, first_page_banner, running_header
from docx.oxml.ns import qn
from docx.shared import RGBColor

from ..unitutil.file import test_file


@pytest.fixture
def document() -> DocumentCls:
    return Document()


@pytest.fixture
def logo_path() -> str:
    return test_file("python-icon.png")


def _section(doc: DocumentCls):
    return doc.sections[0]


def _first_page_header(doc: DocumentCls):
    return _section(doc).first_page_header


def _running_header_container(doc: DocumentCls):
    return _section(doc).header


def _running_footer_container(doc: DocumentCls):
    return _section(doc).footer


class DescribeCoverPage:
    """Behavioural tests for :func:`docx.kit.headers.cover_page`."""

    def it_returns_the_list_of_appended_paragraphs(self, document: DocumentCls):
        result = cover_page(document, title="Annual Report")

        assert isinstance(result, list)
        assert len(result) >= 1

    def it_renders_the_title_in_the_body(self, document: DocumentCls):
        cover_page(document, title="Annual Report")

        body_text = "\n".join(p.text for p in document.paragraphs)
        assert "Annual Report" in body_text

    def it_centres_the_title_paragraph(self, document: DocumentCls):
        result = cover_page(document, title="Annual Report")

        title_para = next(p for p in result if "Annual Report" in p.text)
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_renders_the_title_in_a_bold_run(self, document: DocumentCls):
        result = cover_page(document, title="Annual Report")

        title_para = next(p for p in result if "Annual Report" in p.text)
        title_run = next(r for r in title_para.runs if r.text == "Annual Report")
        assert title_run.bold is True

    def it_emits_subtitle_when_supplied(self, document: DocumentCls):
        cover_page(document, title="Annual Report", subtitle="FY2026")

        body_text = "\n".join(p.text for p in document.paragraphs)
        assert "FY2026" in body_text

    def it_emits_author_when_supplied(self, document: DocumentCls):
        cover_page(document, title="Annual Report", author="Jane Smith")

        body_text = "\n".join(p.text for p in document.paragraphs)
        assert "Jane Smith" in body_text

    def it_emits_date_when_supplied(self, document: DocumentCls):
        cover_page(document, title="Annual Report", date="2026-05-29")

        body_text = "\n".join(p.text for p in document.paragraphs)
        assert "2026-05-29" in body_text

    def it_includes_the_logo_when_supplied(
        self, document: DocumentCls, logo_path: str
    ):
        result = cover_page(document, title="X", logo=logo_path)

        # The logo is the first-emitted paragraph; check for w:drawing.
        found_drawing = False
        for para in result:
            for run in para.runs:
                if run._r.find(qn("w:drawing")) is not None:
                    found_drawing = True
        assert found_drawing

    def it_omits_the_logo_when_None(self, document: DocumentCls):
        result = cover_page(document, title="X")

        for para in result:
            for run in para.runs:
                assert run._r.find(qn("w:drawing")) is None

    def it_emits_a_trailing_page_break_by_default(self, document: DocumentCls):
        result = cover_page(document, title="Title")

        # Page-break paragraph holds a w:br with type="page".
        last_para = result[-1]
        breaks = [b for r in last_para.runs for b in r._r.findall(qn("w:br"))]
        assert any(b.get(qn("w:type")) == "page" for b in breaks)

    def it_can_suppress_the_trailing_page_break(self, document: DocumentCls):
        result = cover_page(document, title="Title", page_break=False)

        for para in result:
            for run in para.runs:
                breaks = run._r.findall(qn("w:br"))
                assert not any(
                    b.get(qn("w:type")) == "page" for b in breaks
                )

    def it_emits_a_decorative_rule_when_author_or_date_is_present(
        self, document: DocumentCls
    ):
        cover_page(document, title="X", author="A")

        body_text = "\n".join(p.text for p in document.paragraphs)
        assert "—" * 5 in body_text

    def it_does_not_emit_a_rule_when_only_title_subtitle(
        self, document: DocumentCls
    ):
        cover_page(document, title="X", subtitle="Y", page_break=False)

        body_text = "\n".join(p.text for p in document.paragraphs)
        assert "—" * 5 not in body_text

    def it_raises_when_title_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="title must be"):
            cover_page(document, title="")

    def it_returns_paragraphs_in_document_order(self, document: DocumentCls):
        result = cover_page(
            document,
            title="T",
            subtitle="S",
            author="A",
            date="D",
            page_break=False,
        )

        texts = [p.text for p in result]
        # Title, Subtitle, rule, Author, Date — rule sits between
        # subtitle and author.
        assert "T" in texts[0]
        assert "S" in texts[1]
        assert "—" in texts[2]
        assert "A" in texts[3]
        assert "D" in texts[4]


class DescribeFirstPageBanner:
    """Behavioural tests for :func:`docx.kit.headers.first_page_banner`."""

    def it_returns_the_list_of_first_page_header_paragraphs(
        self, document: DocumentCls
    ):
        result = first_page_banner(document, title="Annual Report")

        assert isinstance(result, list)
        assert len(result) >= 2  # title + rule (no logo path supplied)

    def it_enables_distinct_first_page_header_on_the_section(
        self, document: DocumentCls
    ):
        first_page_banner(document, title="Annual Report")

        assert _section(document).different_first_page_header_footer is True

    def it_writes_to_the_first_page_header(self, document: DocumentCls):
        first_page_banner(document, title="Annual Report")

        text = "\n".join(p.text for p in _first_page_header(document).paragraphs)
        assert "Annual Report" in text

    def it_does_not_pollute_the_running_header(self, document: DocumentCls):
        first_page_banner(document, title="Annual Report")

        running_text = "\n".join(
            p.text for p in _running_header_container(document).paragraphs
        )
        # Running header is untouched — no banner content leaked.
        assert "Annual Report" not in running_text

    def it_centres_the_banner_title(self, document: DocumentCls):
        result = first_page_banner(document, title="Annual Report")

        title_para = next(p for p in result if "Annual Report" in p.text)
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_emits_a_horizontal_rule(self, document: DocumentCls):
        first_page_banner(document, title="Annual Report")

        text = "\n".join(p.text for p in _first_page_header(document).paragraphs)
        assert "—" * 5 in text

    def it_colours_the_rule_when_line_color_is_a_hex_string(
        self, document: DocumentCls
    ):
        first_page_banner(
            document, title="X", line_color="#FF8800"
        )

        # Find the rule paragraph (the one whose text is em-dashes only).
        rule_para = next(
            p
            for p in _first_page_header(document).paragraphs
            if p.text and set(p.text) == {"—"}
        )
        rule_run = rule_para.runs[0]
        assert rule_run.font.color.rgb == RGBColor(0xFF, 0x88, 0x00)

    def it_accepts_a_hex_string_without_leading_hash(
        self, document: DocumentCls
    ):
        first_page_banner(document, title="X", line_color="123456")

        rule_para = next(
            p
            for p in _first_page_header(document).paragraphs
            if p.text and set(p.text) == {"—"}
        )
        assert rule_para.runs[0].font.color.rgb == RGBColor(
            0x12, 0x34, 0x56
        )

    def it_accepts_an_RGBColor_instance(self, document: DocumentCls):
        rgb = RGBColor(0xAB, 0xCD, 0xEF)
        first_page_banner(document, title="X", line_color=rgb)

        rule_para = next(
            p
            for p in _first_page_header(document).paragraphs
            if p.text and set(p.text) == {"—"}
        )
        assert rule_para.runs[0].font.color.rgb == rgb

    def it_skips_colour_when_line_color_is_None(self, document: DocumentCls):
        first_page_banner(document, title="X", line_color=None)

        rule_para = next(
            p
            for p in _first_page_header(document).paragraphs
            if p.text and set(p.text) == {"—"}
        )
        # font.color.rgb is None when no explicit colour applied.
        assert rule_para.runs[0].font.color.rgb is None

    def it_includes_the_logo_when_supplied(
        self, document: DocumentCls, logo_path: str
    ):
        first_page_banner(
            document, title="Annual Report", logo=logo_path
        )

        found_drawing = False
        for para in _first_page_header(document).paragraphs:
            for run in para.runs:
                if run._r.find(qn("w:drawing")) is not None:
                    found_drawing = True
        assert found_drawing

    def it_clears_existing_first_page_header_content(
        self, document: DocumentCls
    ):
        # Seed pre-existing content on the first-page header.
        section = _section(document)
        section.different_first_page_header_footer = True
        seeded_header = section.first_page_header
        seeded_header.paragraphs[0].text = "STALE"
        seeded_header.add_paragraph("EXTRA STALE")

        first_page_banner(document, title="FRESH")

        text = "\n".join(p.text for p in _first_page_header(document).paragraphs)
        assert "STALE" not in text
        assert "FRESH" in text

    def it_is_idempotent_on_re_run(self, document: DocumentCls):
        first_page_banner(document, title="ONE")
        first_page_banner(document, title="TWO")

        text = "\n".join(p.text for p in _first_page_header(document).paragraphs)
        assert "ONE" not in text
        assert "TWO" in text

    def it_raises_when_title_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="title must be"):
            first_page_banner(document, title="")

    def it_raises_when_line_color_is_a_non_string_non_color(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="line_color must be"):
            first_page_banner(document, title="X", line_color=42)


class DescribeRunningHeader:
    """Behavioural tests for :func:`docx.kit.headers.running_header`."""

    def it_writes_to_the_primary_header_by_default(
        self, document: DocumentCls
    ):
        running_header(document, left="Annual Report", right="Confidential")

        text = "\n".join(
            p.text for p in _running_header_container(document).paragraphs
        )
        # All cell content reachable as paragraphs in the header.
        full_text = text
        for cell in (
            _running_header_container(document)
            .tables[0]
            .rows[0]
            .cells
        ):
            full_text += "\n".join(p.text for p in cell.paragraphs)
        assert "Annual Report" in full_text
        assert "Confidential" in full_text

    def it_writes_to_the_primary_footer_when_footer_is_True(
        self, document: DocumentCls
    ):
        running_header(
            document, left="L", right="R", footer=True
        )

        cells = (
            _running_footer_container(document)
            .tables[0]
            .rows[0]
            .cells
        )
        cell_texts = [p.text for cell in cells for p in cell.paragraphs]
        assert "L" in "\n".join(cell_texts)
        assert "R" in "\n".join(cell_texts)

    def it_emits_a_three_cell_table_when_two_or_more_cells_are_set(
        self, document: DocumentCls
    ):
        running_header(document, left="L", right="R")

        tables = _running_header_container(document).tables
        assert len(tables) == 1
        assert len(tables[0].rows[0].cells) == 3

    def it_aligns_the_three_cells_at_left_center_right(
        self, document: DocumentCls
    ):
        running_header(document, left="L", center="C", right="R")

        cells = (
            _running_header_container(document).tables[0].rows[0].cells
        )
        assert cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert (
            cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        )
        assert cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def it_renders_a_single_paragraph_when_only_one_cell_is_set(
        self, document: DocumentCls
    ):
        running_header(document, left="Annual Report")

        # No table emitted — single-cell shortcut writes a paragraph.
        assert len(_running_header_container(document).tables) == 0
        text = "\n".join(
            p.text for p in _running_header_container(document).paragraphs
        )
        assert "Annual Report" in text

    def it_aligns_the_single_paragraph_to_the_correct_edge(
        self, document: DocumentCls
    ):
        running_header(document, right="Confidential")

        para = _running_header_container(document).paragraphs[0]
        assert para.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def it_clears_existing_header_content_before_writing(
        self, document: DocumentCls
    ):
        seeded = _running_header_container(document)
        seeded.paragraphs[0].text = "STALE"
        seeded.add_paragraph("EXTRA")

        running_header(document, left="FRESH")

        text = "\n".join(
            p.text for p in _running_header_container(document).paragraphs
        )
        assert "STALE" not in text
        assert "FRESH" in text

    def it_is_idempotent_on_re_run(self, document: DocumentCls):
        running_header(document, left="ONE")
        running_header(document, left="TWO")

        text = "\n".join(
            p.text for p in _running_header_container(document).paragraphs
        )
        assert "ONE" not in text
        assert "TWO" in text

    def it_raises_when_no_cells_are_provided(self, document: DocumentCls):
        with pytest.raises(ValueError, match="at least one of"):
            running_header(document)


class DescribeRoundTrip:
    """End-to-end save/reload — content survives a serialisation cycle."""

    def it_preserves_cover_page_content(
        self, document: DocumentCls, logo_path: str
    ):
        cover_page(
            document,
            title="Annual Report",
            subtitle="FY2026",
            logo=logo_path,
            date="2026-05-29",
            author="Jane Smith",
        )
        buf = BytesIO()
        document.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        body_text = "\n".join(p.text for p in reloaded.paragraphs)
        assert "Annual Report" in body_text
        assert "FY2026" in body_text
        assert "Jane Smith" in body_text
        assert "2026-05-29" in body_text

    def it_preserves_first_page_banner(
        self, document: DocumentCls, logo_path: str
    ):
        first_page_banner(
            document, title="Annual Report", logo=logo_path
        )
        buf = BytesIO()
        document.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        section = reloaded.sections[0]
        assert section.different_first_page_header_footer is True
        text = "\n".join(p.text for p in section.first_page_header.paragraphs)
        assert "Annual Report" in text
        assert "—" * 5 in text

    def it_preserves_running_header_three_cell_layout(
        self, document: DocumentCls
    ):
        running_header(
            document, left="Annual Report", right="Confidential"
        )
        buf = BytesIO()
        document.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        tables = reloaded.sections[0].header.tables
        assert len(tables) == 1
        cell_texts = [
            p.text
            for cell in tables[0].rows[0].cells
            for p in cell.paragraphs
        ]
        assert "Annual Report" in "\n".join(cell_texts)
        assert "Confidential" in "\n".join(cell_texts)


class DescribeReExport:
    """Re-export contract: ``from docx.kit import headers`` works."""

    def it_exposes_the_three_helpers_on_the_module(self):
        assert callable(headers.cover_page)
        assert callable(headers.first_page_banner)
        assert callable(headers.running_header)

    def it_lists_the_helpers_in_module_all(self):
        assert set(headers.__all__) == {
            "cover_page",
            "first_page_banner",
            "running_header",
        }
