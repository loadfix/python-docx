"""Unit-test suite for ``docx.kit.memos`` template factories."""

from __future__ import annotations

import pytest

from docx.document import Document as DocumentCls
from docx.kit import memos


# -- Shared helpers -------------------------------------------------------


def _texts(document: DocumentCls):
    """Return the text of every paragraph in ``document``."""
    return [p.text for p in document.paragraphs]


def _full_text(document: DocumentCls) -> str:
    return "\n".join(_texts(document))


# -- Investment memo ------------------------------------------------------


class DescribeInvestmentMemo:
    """Unit-test suite for ``memos.investment_memo``."""

    def it_returns_a_document_with_the_company_in_the_title(self):
        doc = memos.investment_memo(company="Acme Corp")

        assert isinstance(doc, DocumentCls)
        assert "Investment Memo: Acme Corp" in _full_text(doc)

    def it_renders_the_metadata_block(self):
        doc = memos.investment_memo(
            company="Acme Corp",
            sector="SaaS",
            stage="Series B",
            ask="$25M for 18 months runway",
        )
        text = _full_text(doc)

        assert "Sector" in text
        assert "SaaS" in text
        assert "Stage" in text
        assert "Series B" in text
        assert "Ask" in text
        assert "$25M for 18 months runway" in text

    def it_omits_metadata_lines_when_unsupplied(self):
        doc = memos.investment_memo(company="Acme Corp")
        text = _full_text(doc)

        # -- Without a sector / stage / ask, no metadata-line should
        # -- be emitted at all (the labels become noise without values). --
        assert "Sector:" not in text
        assert "Stage:" not in text
        assert "Ask:" not in text

    def it_renders_an_executive_summary_heading(self):
        doc = memos.investment_memo(company="Acme Corp")

        assert "Executive Summary" in _full_text(doc)

    def it_renders_the_four_SCQA_labels(self):
        doc = memos.investment_memo(company="Acme Corp")
        text = _full_text(doc)

        # -- Each SCQA piece is rendered as "Label: body" --
        assert "Situation:" in text
        assert "Complication:" in text
        assert "Question:" in text
        assert "Answer:" in text

    def it_renders_supplied_SCQA_bodies(self):
        doc = memos.investment_memo(
            company="Acme Corp",
            situation="Acme operates in a $50B market.",
            complication="Despite product-market fit, Acme faces churn.",
            question="Should we lead the Series B?",
            answer="Yes, with terms of $25M at $200M post.",
        )
        text = _full_text(doc)

        assert "Acme operates in a $50B market." in text
        assert "Despite product-market fit, Acme faces churn." in text
        assert "Should we lead the Series B?" in text
        assert "Yes, with terms of $25M at $200M post." in text

    def it_renders_placeholders_when_SCQA_pieces_missing(self):
        doc = memos.investment_memo(company="Acme Corp")
        text = _full_text(doc)

        assert "[Situation:" in text
        assert "[Complication:" in text
        assert "[Question:" in text
        assert "[Answer:" in text

    def it_renders_supplied_body_sections_as_headings_and_paragraphs(self):
        doc = memos.investment_memo(
            company="Acme Corp",
            sections=[
                {"heading": "Market", "body": "TAM is $50B."},
                {"heading": "Team", "body": "Founder/market fit is strong."},
                {"heading": "Risks", "body": "Churn is 8% net."},
            ],
        )
        text = _full_text(doc)

        assert "Market" in text
        assert "TAM is $50B." in text
        assert "Team" in text
        assert "Founder/market fit is strong." in text
        assert "Risks" in text
        assert "Churn is 8% net." in text

    def it_accepts_a_sequence_body_for_multi_paragraph_sections(self):
        doc = memos.investment_memo(
            company="Acme Corp",
            sections=[
                {
                    "heading": "Market",
                    "body": ["First paragraph.", "Second paragraph."],
                },
            ],
        )
        text = _full_text(doc)

        assert "First paragraph." in text
        assert "Second paragraph." in text

    def it_accepts_an_optional_author_and_date(self):
        doc = memos.investment_memo(
            company="Acme Corp",
            author="Investment Team",
            date="2026-05-29",
        )
        text = _full_text(doc)

        assert "Investment Team" in text
        assert "2026-05-29" in text

    def it_raises_when_company_is_empty(self):
        with pytest.raises(ValueError, match="company is required"):
            memos.investment_memo(company="")

    def it_raises_when_company_is_whitespace_only(self):
        with pytest.raises(ValueError, match="company is required"):
            memos.investment_memo(company="   ")

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(
            ValueError, match="non-empty 'heading'"
        ):
            memos.investment_memo(
                company="Acme Corp",
                sections=[{"body": "no heading"}],
            )

    def it_raises_when_a_section_is_not_a_mapping(self):
        with pytest.raises(
            ValueError, match="must be a mapping"
        ):
            memos.investment_memo(
                company="Acme Corp",
                sections=["not a dict"],  # type: ignore[list-item]
            )

    def it_tolerates_a_section_without_a_body(self):
        doc = memos.investment_memo(
            company="Acme Corp",
            sections=[{"heading": "Market"}],
        )

        assert "Market" in _full_text(doc)


# -- Business case --------------------------------------------------------


class DescribeBusinessCase:
    """Unit-test suite for ``memos.business_case``."""

    def it_returns_a_document_with_the_project_in_the_title(self):
        doc = memos.business_case(project="Migration to AWS")

        assert isinstance(doc, DocumentCls)
        assert "Business Case: Migration to AWS" in _full_text(doc)

    def it_renders_the_metadata_block(self):
        doc = memos.business_case(
            project="Migration to AWS",
            sponsor="CTO",
            timeline="Q3 2026 - Q1 2027",
        )
        text = _full_text(doc)

        assert "Sponsor" in text
        assert "CTO" in text
        assert "Timeline" in text
        assert "Q3 2026 - Q1 2027" in text

    def it_renders_the_five_required_sections(self):
        doc = memos.business_case(project="Migration to AWS")
        text = _full_text(doc)

        # -- The fixed five-section structure: objectives, options,
        # -- recommendation, risks, timeline. The executive summary
        # -- precedes all five. --
        assert "Executive Summary" in text
        assert "Objectives" in text
        assert "Options" in text
        assert "Recommendation" in text
        assert "Risks" in text
        assert "Timeline" in text

    def it_renders_the_recommendation_as_a_callout(self):
        doc = memos.business_case(
            project="Migration to AWS",
            recommendation="Replatform",
        )
        text = _full_text(doc)

        # -- The recommendation appears twice: once in the executive
        # -- summary callout, once as the body of the Recommendation
        # -- heading. --
        assert text.count("Replatform") >= 2

    def it_renders_objectives_as_a_bulleted_list(self):
        doc = memos.business_case(
            project="Migration to AWS",
            objectives=[
                "Reduce ops cost 40%",
                "Improve SLA to 99.95%",
            ],
        )
        text = _full_text(doc)

        assert "Reduce ops cost 40%" in text
        assert "Improve SLA to 99.95%" in text

    def it_renders_a_placeholder_when_objectives_missing(self):
        doc = memos.business_case(project="Migration to AWS")

        assert "[Insert measurable objectives" in _full_text(doc)

    def it_renders_options_as_a_four_column_table(self):
        doc = memos.business_case(
            project="Migration to AWS",
            options=[
                {
                    "name": "Status quo",
                    "cost": "$0",
                    "pros": ["No disruption"],
                    "cons": ["High ongoing cost", "Tech debt"],
                },
                {
                    "name": "Rehost",
                    "cost": "$2m",
                    "pros": ["Fast"],
                    "cons": ["Limited gains"],
                },
                {
                    "name": "Replatform",
                    "cost": "$5m",
                    "pros": ["Long-term cost wins", "SLA target met"],
                    "cons": ["Higher upfront cost"],
                },
            ],
        )

        assert len(doc.tables) >= 1
        # -- Locate the options table by its header --
        options_table = None
        for table in doc.tables:
            cells = [cell.text for cell in table.rows[0].cells]
            if cells and cells[0] == "Option":
                options_table = table
                break
        assert options_table is not None
        # -- header + three options --
        assert len(options_table.rows) == 4
        # -- header row order --
        header = [c.text for c in options_table.rows[0].cells]
        assert header == ["Option", "Cost", "Pros", "Cons"]
        # -- first data row --
        first = options_table.rows[1].cells
        assert first[0].text == "Status quo"
        assert first[1].text == "$0"
        assert "No disruption" in first[2].text
        assert "High ongoing cost" in first[3].text
        assert "Tech debt" in first[3].text

    def it_renders_a_placeholder_when_options_missing(self):
        doc = memos.business_case(project="Migration to AWS")

        assert "[Insert at least two options" in _full_text(doc)

    def it_renders_risks_as_a_bulleted_list(self):
        doc = memos.business_case(
            project="Migration to AWS",
            risks=["Skill gap", "Vendor lock-in"],
        )
        text = _full_text(doc)

        assert "Skill gap" in text
        assert "Vendor lock-in" in text

    def it_renders_a_placeholder_when_risks_missing(self):
        doc = memos.business_case(project="Migration to AWS")

        assert "[Identify the key risks" in _full_text(doc)

    def it_renders_the_timeline_body(self):
        doc = memos.business_case(
            project="Migration to AWS",
            timeline="Q3 2026 - Q1 2027",
        )

        assert "Q3 2026 - Q1 2027" in _full_text(doc)

    def it_renders_a_placeholder_when_timeline_missing(self):
        doc = memos.business_case(project="Migration to AWS")

        assert "[Insert milestones" in _full_text(doc)

    def it_renders_a_placeholder_when_recommendation_missing(self):
        doc = memos.business_case(project="Migration to AWS")

        assert "[State the recommended option" in _full_text(doc)

    def it_handles_options_with_missing_pros_or_cons_gracefully(self):
        doc = memos.business_case(
            project="Migration to AWS",
            options=[
                {"name": "Bare option", "cost": "$1m"},
            ],
        )

        # -- Did not raise; the table renders with empty cells. --
        options_table = None
        for table in doc.tables:
            cells = [cell.text for cell in table.rows[0].cells]
            if cells and cells[0] == "Option":
                options_table = table
                break
        assert options_table is not None
        first = options_table.rows[1].cells
        assert first[0].text == "Bare option"
        assert first[2].text == ""
        assert first[3].text == ""

    def it_raises_when_project_is_empty(self):
        with pytest.raises(ValueError, match="project is required"):
            memos.business_case(project="")

    def it_raises_when_project_is_whitespace_only(self):
        with pytest.raises(ValueError, match="project is required"):
            memos.business_case(project="   ")

    def it_raises_when_an_option_has_no_name(self):
        with pytest.raises(ValueError, match="non-empty 'name'"):
            memos.business_case(
                project="Migration to AWS",
                options=[{"cost": "$1m"}],
            )

    def it_raises_when_an_option_is_not_a_mapping(self):
        with pytest.raises(ValueError, match="must be a mapping"):
            memos.business_case(
                project="Migration to AWS",
                options=["not a dict"],  # type: ignore[list-item]
            )

    def it_accepts_a_date_under_the_title(self):
        doc = memos.business_case(
            project="Migration to AWS",
            date="2026-05-29",
        )

        assert "2026-05-29" in _full_text(doc)


# -- Round-trip integration ----------------------------------------------


class DescribeMemosRoundTrip:
    """End-to-end smoke-tests: every factory produces a saveable document."""

    def it_can_save_an_investment_memo_to_a_BytesIO(self):
        from io import BytesIO

        doc = memos.investment_memo(
            company="Acme Corp",
            sector="SaaS",
            stage="Series B",
            ask="$25M for 18 months runway",
            situation="Acme is...",
            complication="Despite PMF, Acme faces churn...",
            question="Should we lead the Series B?",
            answer="Yes, with terms of...",
            sections=[
                {"heading": "Market", "body": "TAM is $50B."},
                {"heading": "Team", "body": "Strong."},
                {"heading": "Risks", "body": "Churn at 8%."},
            ],
        )
        buf = BytesIO()
        doc.save(buf)
        # -- Word .docx is a zip; magic bytes are 'PK' --
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_business_case_to_a_BytesIO(self):
        from io import BytesIO

        doc = memos.business_case(
            project="Migration to AWS",
            sponsor="CTO",
            objectives=["Reduce ops cost 40%", "Improve SLA to 99.95%"],
            options=[
                {
                    "name": "Status quo",
                    "cost": "$0",
                    "pros": ["No disruption"],
                    "cons": ["Tech debt"],
                },
                {
                    "name": "Rehost",
                    "cost": "$2m",
                    "pros": ["Fast"],
                    "cons": ["Limited gains"],
                },
                {
                    "name": "Replatform",
                    "cost": "$5m",
                    "pros": ["Long-term wins"],
                    "cons": ["Upfront cost"],
                },
            ],
            recommendation="Replatform",
            risks=["Skill gap", "Vendor lock-in"],
            timeline="Q3 2026 - Q1 2027",
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"


# -- Module surface -------------------------------------------------------


class DescribeMemosModule:
    """Module-level surface contracts."""

    def it_exposes_investment_memo_and_business_case(self):
        assert hasattr(memos, "investment_memo")
        assert hasattr(memos, "business_case")
        assert "investment_memo" in memos.__all__
        assert "business_case" in memos.__all__

    def it_is_re_exported_from_the_kit_package(self):
        from docx.kit import memos as memos_pkg

        assert memos_pkg is memos
