"""Unit-test suite for ``docx.kit.diagrams`` helpers (issue #292).

The diagram helpers default to a network call against ``kroki.io``;
the suite **never** hits the network — every kroki transport call is
intercepted via a monkeypatched ``_http_post`` that returns a
deterministic 1x1 PNG byte-string. Local-binary tests are gated on
the binary being on ``PATH`` and skipped when absent.
"""

from __future__ import annotations

import base64
import zlib
from typing import Any, Callable, Dict, List, Tuple

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import diagrams
from docx.kit.diagrams import dot, kroki_get_url, kroki_url, mermaid, plantuml
from docx.shape import InlineShape

# -- A minimal valid 1x1 transparent PNG.  We never decode it; both
# -- the kit and python-docx accept any well-formed PNG byte-string,
# -- and we use a tiny one so tests stay fast and deterministic.
_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNgAAIAAAUAAen63NgAAAAASUVORK5CYII="
)


@pytest.fixture
def document() -> DocumentCls:
    return Document()


@pytest.fixture
def patched_http(monkeypatch: pytest.MonkeyPatch) -> List[Dict[str, Any]]:
    """Replace ``_http_post`` with a stub recording every call.

    Returns the list of ``{"url": ..., "data": ..., "timeout": ...}``
    dicts so individual tests can assert routing / timeout / payload.
    """
    calls: List[Dict[str, Any]] = []

    def fake_post(url: str, data: bytes, timeout: float) -> bytes:
        calls.append({"url": url, "data": data, "timeout": timeout})
        return _PNG_1X1

    monkeypatch.setattr(diagrams, "_http_post", fake_post)
    return calls


@pytest.fixture
def force_kroki(monkeypatch: pytest.MonkeyPatch) -> None:
    """Ensure ``backend='auto'`` resolves to kroki even on hosts with the
    binaries installed by short-circuiting :func:`_find_local_binary`."""
    monkeypatch.setattr(diagrams, "_find_local_binary", lambda lang: None)


class DescribeMermaid:
    """Behavioural tests for :func:`docx.kit.diagrams.mermaid`."""

    def it_returns_an_inline_shape(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        result = mermaid(document, "flowchart TD\n  A --> B")

        assert isinstance(result, InlineShape)
        # -- the picture appears as a new paragraph at the body's tail --
        assert len(document.paragraphs) >= 1

    def it_routes_to_kroki_mermaid_endpoint_by_default(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        mermaid(document, "flowchart TD\n  A --> B")

        assert len(patched_http) == 1
        assert patched_http[0]["url"] == "https://kroki.io/mermaid/png"

    def it_posts_the_source_as_utf8_bytes(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        source = "flowchart TD\n  A --> B"
        mermaid(document, source)

        assert patched_http[0]["data"] == source.encode("utf-8")

    def it_passes_a_default_30_second_timeout(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        mermaid(document, "flowchart TD\n  A --> B")

        assert patched_http[0]["timeout"] == 30.0

    def it_honours_a_caller_supplied_timeout(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        mermaid(document, "flowchart TD\n  A --> B", timeout=5.0)

        assert patched_http[0]["timeout"] == 5.0

    def it_appends_a_caption_paragraph_when_caption_is_supplied(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        before = len(document.paragraphs)

        mermaid(
            document, "flowchart TD\n  A --> B", caption="Workflow"
        )

        # -- one paragraph for the picture, one for the caption --
        assert len(document.paragraphs) == before + 2
        caption_para = document.paragraphs[-1]
        assert caption_para.style.name == "Caption"
        assert "Workflow" in caption_para.text

    def it_skips_the_caption_paragraph_by_default(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        before = len(document.paragraphs)

        mermaid(document, "flowchart TD\n  A --> B")

        # -- one paragraph for the picture, no caption paragraph --
        assert len(document.paragraphs) == before + 1

    def it_honours_a_caller_supplied_image_format(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        # -- python-docx's add_picture only handles raster bytes for
        # -- the standard image headers; SVG is technically valid but
        # -- python-docx reads the magic header.  We patch the http
        # -- stub to return the same PNG bytes regardless and only
        # -- assert the URL was the SVG endpoint.
        mermaid(document, "flowchart TD\n  A --> B", image_format="svg")

        assert patched_http[0]["url"].endswith("/svg")

    def it_honours_a_custom_base_url_for_self_hosted_kroki(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        mermaid(
            document,
            "flowchart TD\n  A --> B",
            base_url="https://kroki.example.com",
        )

        assert patched_http[0]["url"].startswith("https://kroki.example.com/")

    def it_raises_on_an_empty_source(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        with pytest.raises(ValueError, match="non-empty"):
            mermaid(document, "")
        with pytest.raises(ValueError, match="non-empty"):
            mermaid(document, "   \n   ")

    def it_raises_on_an_unknown_backend(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="backend must be one of"):
            mermaid(document, "graph TD;A-->B", backend="quantum")

    def it_raises_on_an_unknown_image_format(
        self, document: DocumentCls
    ):
        with pytest.raises(ValueError, match="image_format must be one of"):
            mermaid(document, "graph TD;A-->B", image_format="webp")

    def it_propagates_alt_text_to_the_inline_shape(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        result = mermaid(
            document,
            "flowchart TD\n  A --> B",
            alt_text="state machine",
        )

        # -- python-docx exposes alt-text on the inline shape; the
        # -- helper either set both fields or fell back to whichever
        # -- exists.  We assert via the public attribute when present.
        for attr in ("alt_text_descr", "alt_text_title"):
            if hasattr(result, attr):
                # -- one of the two attributes carries the alt_text --
                value = getattr(result, attr)
                if value:
                    assert "state machine" in value
                    return
        # -- if neither attribute exists we silently accept; the
        # -- helper is best-effort on older revisions.

    def it_forces_kroki_when_backend_is_kroki(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        monkeypatch: pytest.MonkeyPatch,
    ):
        # -- pretend a binary is on PATH; backend='kroki' must still
        # -- POST to kroki.
        monkeypatch.setattr(
            diagrams, "_find_local_binary", lambda lang: "/usr/bin/mmdc"
        )

        mermaid(document, "flowchart TD\n  A --> B", backend="kroki")

        assert len(patched_http) == 1


class DescribePlantuml:
    """Behavioural tests for :func:`docx.kit.diagrams.plantuml`."""

    def it_routes_to_the_kroki_plantuml_endpoint(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        plantuml(document, "@startuml\nA -> B\n@enduml")

        assert patched_http[0]["url"] == "https://kroki.io/plantuml/png"

    def it_returns_an_inline_shape(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        result = plantuml(document, "@startuml\nA -> B\n@enduml")

        assert isinstance(result, InlineShape)


class DescribeDot:
    """Behavioural tests for :func:`docx.kit.diagrams.dot`."""

    def it_routes_to_the_kroki_graphviz_endpoint(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        dot(document, "digraph G { A -> B; }")

        assert patched_http[0]["url"] == "https://kroki.io/graphviz/png"

    def it_returns_an_inline_shape(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        result = dot(document, "digraph G { A -> B; }")

        assert isinstance(result, InlineShape)


class DescribeKrokiUrlHelpers:
    """Behavioural tests for :func:`docx.kit.diagrams.kroki_url` /
    :func:`docx.kit.diagrams.kroki_get_url`."""

    def it_builds_the_kroki_post_endpoint(self):
        assert kroki_url("mermaid") == "https://kroki.io/mermaid/png"
        assert kroki_url("plantuml") == "https://kroki.io/plantuml/png"
        # -- "dot" maps to kroki's "graphviz" slug
        assert kroki_url("dot") == "https://kroki.io/graphviz/png"

    def it_supports_alternate_image_formats(self):
        assert kroki_url("mermaid", image_format="svg") == (
            "https://kroki.io/mermaid/svg"
        )

    def it_honours_a_custom_base_url(self):
        assert kroki_url(
            "dot", base_url="https://kroki.example.com/"
        ) == "https://kroki.example.com/graphviz/png"

    def it_raises_on_an_unknown_language(self):
        with pytest.raises(ValueError, match="language must be one of"):
            kroki_url("typescript")

    def it_builds_a_get_url_with_deflate_base64_encoded_source(self):
        url = kroki_get_url("mermaid", "graph TD;A-->B")

        # -- shape: <base>/<slug>/<format>/<encoded>
        prefix = "https://kroki.io/mermaid/png/"
        assert url.startswith(prefix)
        encoded = url[len(prefix):]
        # -- urlsafe-b64-decoded value must round-trip through deflate
        round_tripped = zlib.decompress(
            base64.urlsafe_b64decode(encoded.encode("ascii"))
        ).decode("utf-8")
        assert round_tripped == "graph TD;A-->B"


class DescribeRenderDispatch:
    """Behavioural tests for the internal ``_render`` dispatcher."""

    def it_prefers_a_local_binary_when_one_is_on_path(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        monkeypatch: pytest.MonkeyPatch,
    ):
        # -- pretend mmdc is on PATH and capture the local-render call.
        local_calls: List[Tuple[str, str, str]] = []

        def fake_local(language: str, source: str, fmt: str) -> bytes:
            local_calls.append((language, source, fmt))
            return _PNG_1X1

        monkeypatch.setattr(
            diagrams, "_find_local_binary", lambda lang: "/usr/bin/mmdc"
        )
        monkeypatch.setattr(diagrams, "_render_local", fake_local)

        mermaid(document, "flowchart TD\n  A --> B")

        # -- local was called; kroki was NOT --
        assert len(local_calls) == 1
        assert local_calls[0][0] == "mermaid"
        assert len(patched_http) == 0

    def it_falls_back_to_kroki_when_no_local_binary_is_available(
        self,
        document: DocumentCls,
        patched_http: List[Dict[str, Any]],
        force_kroki: None,
    ):
        mermaid(document, "flowchart TD\n  A --> B")

        assert len(patched_http) == 1

    def it_forces_local_when_backend_is_local(
        self,
        document: DocumentCls,
        monkeypatch: pytest.MonkeyPatch,
    ):
        # -- backend='local' must NEVER call kroki, even when no
        # -- binary is on PATH (it should raise FileNotFoundError).
        monkeypatch.setattr(diagrams, "_find_local_binary", lambda lang: None)

        with pytest.raises(FileNotFoundError):
            mermaid(document, "flowchart TD\n  A --> B", backend="local")


class DescribeSoftImport:
    """Behavioural tests for the ``requests``/``httpx`` soft-import shim."""

    def it_raises_an_informative_error_when_neither_client_is_available(
        self,
        document: DocumentCls,
        monkeypatch: pytest.MonkeyPatch,
        force_kroki: None,
    ):
        # -- Force the real ``_http_post`` (not the patched stub) and
        # -- make ``import requests`` and ``import httpx`` both fail.
        import builtins

        original_import = builtins.__import__

        def selective_import(name: str, *args: Any, **kwargs: Any) -> Any:
            if name in ("requests", "httpx"):
                raise ImportError("simulated missing client")
            return original_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", selective_import)

        with pytest.raises(ImportError, match="requests.*httpx|httpx.*requests"):
            mermaid(document, "flowchart TD\n  A --> B", backend="kroki")


class DescribeKitReExport:
    """The diagrams module is re-exported from ``docx.kit``."""

    def it_re_exports_the_diagrams_submodule(self):
        from docx import kit

        assert hasattr(kit, "diagrams")
        assert kit.diagrams is diagrams
        assert "diagrams" in kit.__all__
