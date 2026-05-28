"""Unit-test suite for ``docx.kit.invoices`` template factories."""

from __future__ import annotations

from io import BytesIO

import pytest

from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import invoices


# -- Shared helpers -------------------------------------------------------


def _texts(document: DocumentCls):
    """Return the text of every paragraph in ``document``."""
    return [p.text for p in document.paragraphs]


def _full_text(document: DocumentCls) -> str:
    return "\n".join(_texts(document))


def _table_with_first_cell(document: DocumentCls, header_text: str):
    """Return the first table in ``document`` whose ``[0][0]`` cell is ``header_text``."""
    for table in document.tables:
        cells = [cell.text for cell in table.rows[0].cells]
        if cells and cells[0] == header_text:
            return table
    return None


# -- Factory: invoice ----------------------------------------------------


class DescribeInvoice:
    """Unit-test suite for ``invoices.invoice``."""

    def it_returns_a_document_with_a_tax_invoice_header_when_GST_applies(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "quantity": 1, "unit_price": 100}],
        )

        assert isinstance(doc, DocumentCls)
        assert "Tax Invoice" in _full_text(doc)

    def it_falls_back_to_plain_invoice_when_no_GST_applies(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "quantity": 1, "unit_price": 100}],
            default_gst_rate=0,
        )
        text = _full_text(doc)

        assert "Invoice" in text
        # -- and should NOT advertise itself as a tax invoice --
        assert "Tax Invoice" not in text

    def it_renders_the_invoice_number_as_subtitle(self):
        doc = invoices.invoice(
            invoice_number="INV-2026-0042",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
        )

        assert "INV-2026-0042" in _full_text(doc)

    def it_renders_issue_and_due_dates_in_the_header(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            due_date="2026-04-14",
            items=[{"description": "X", "unit_price": 100}],
        )
        text = _full_text(doc)

        assert "Issue Date" in text
        assert "2026-03-15" in text
        assert "Due Date" in text
        assert "2026-04-14" in text

    def it_omits_the_due_date_line_when_unsupplied(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
        )

        assert "Due Date" not in _full_text(doc)

    def it_renders_the_seller_party_block(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            seller={
                "name": "Acme Corp",
                "abn": "12 345 678 901",
                "address": "123 Pitt Street\nSydney NSW 2000",
                "phone": "+61 2 1234 5678",
                "email": "billing@acme.com",
            },
            items=[{"description": "X", "unit_price": 100}],
        )
        text = _full_text(doc)

        assert "From" in text
        assert "Acme Corp" in text
        assert "12 345 678 901" in text
        assert "123 Pitt Street" in text
        assert "Sydney NSW 2000" in text
        assert "+61 2 1234 5678" in text
        assert "billing@acme.com" in text

    def it_renders_the_buyer_party_block(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            buyer={"name": "Beta Pty Ltd", "abn": "98 765 432 109"},
            items=[{"description": "X", "unit_price": 100}],
        )
        text = _full_text(doc)

        assert "Bill To" in text
        assert "Beta Pty Ltd" in text
        assert "98 765 432 109" in text

    def it_renders_a_placeholder_when_a_party_is_missing(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
        )
        text = _full_text(doc)

        assert "[From details]" in text
        assert "[Bill To details]" in text

    def it_renders_unrecognised_party_keys_verbatim(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            seller={"name": "Acme", "tax_office": "Sydney"},
            items=[{"description": "X", "unit_price": 100}],
        )
        text = _full_text(doc)

        # -- Custom key surfaces with title-cased label --
        assert "Tax Office" in text
        assert "Sydney" in text

    def it_renders_a_five_column_table_when_GST_applies(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 40, "unit_price": 250},
                {"description": "Travel", "quantity": 1, "unit_price": 580},
            ],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        # -- 5 cols: Description / Quantity / Unit Price / GST / Line Total --
        header = [c.text for c in items_table.rows[0].cells]
        assert header == [
            "Description",
            "Quantity",
            "Unit Price",
            "GST",
            "Line Total",
        ]
        # -- header row + two data rows --
        assert len(items_table.rows) == 3

    def it_drops_the_gst_column_when_every_line_is_GST_free(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "unit_price": 250, "gst_rate": 0},
            ],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        header = [c.text for c in items_table.rows[0].cells]
        # -- 4 cols (no GST column) --
        assert header == ["Description", "Quantity", "Unit Price", "Line Total"]

    def it_right_aligns_numeric_cells_in_the_line_items_table(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 40, "unit_price": 250},
            ],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        data_row = items_table.rows[1].cells
        # -- Description left-aligned, all the numeric columns right-aligned --
        assert (
            data_row[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        )
        for col in range(1, len(data_row)):
            assert (
                data_row[col].paragraphs[0].alignment
                == WD_ALIGN_PARAGRAPH.RIGHT
            ), f"col {col} should be right-aligned"

    def it_renders_money_values_with_two_decimals_and_currency_prefix(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 40, "unit_price": 250},
            ],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        data_row = items_table.rows[1].cells
        # -- unit price = $250.00, GST = $1,000.00, line total = $10,000.00 --
        assert data_row[2].text == "$250.00"
        assert data_row[3].text == "$1,000.00"
        assert data_row[4].text == "$10,000.00"

    def it_formats_quantity_as_integer_when_whole(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 40, "unit_price": 250},
            ],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        # -- Quantity column reads "40" not "40.0" --
        assert items_table.rows[1].cells[1].text == "40"

    def it_formats_fractional_quantities_with_a_trimmed_decimal(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Hours", "quantity": 1.5, "unit_price": 200},
            ],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        assert items_table.rows[1].cells[1].text == "1.5"

    def it_defaults_quantity_to_one_when_unsupplied(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "Setup fee", "unit_price": 100}],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        assert items_table.rows[1].cells[1].text == "1"

    def it_defaults_GST_to_ten_percent(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 1, "unit_price": 100},
            ],
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        # -- $100 * 0.10 = $10.00 --
        assert items_table.rows[1].cells[3].text == "$10.00"

    def it_renders_the_totals_block_with_subtotal_GST_and_grand_total(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 40, "unit_price": 250},
                {"description": "Travel", "quantity": 1, "unit_price": 580},
            ],
        )

        # -- subtotal = 10000 + 580 = 10580; GST = 1000 + 58 = 1058;
        # -- grand   = 11638 --
        totals = _table_with_first_cell(doc, "Subtotal")
        assert totals is not None
        # -- 3 rows: Subtotal / GST / Total --
        assert len(totals.rows) == 3
        assert totals.rows[0].cells[1].text == "$10,580.00"
        assert totals.rows[1].cells[0].text == "GST"
        assert totals.rows[1].cells[1].text == "$1,058.00"
        assert totals.rows[2].cells[0].text == "Total"
        assert totals.rows[2].cells[1].text == "$11,638.00"

    def it_drops_the_GST_row_when_every_line_is_GST_free(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100, "gst_rate": 0}],
            default_gst_rate=0,
        )

        totals = _table_with_first_cell(doc, "Subtotal")
        assert totals is not None
        # -- only Subtotal + Total when no GST --
        assert len(totals.rows) == 2
        assert totals.rows[1].cells[0].text == "Total"

    def it_supports_an_international_invoice_via_default_gst_rate_zero(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 5, "unit_price": 200},
            ],
            default_gst_rate=0,
        )
        text = _full_text(doc)

        assert "Tax Invoice" not in text
        assert "Invoice" in text
        # -- four-column table (no GST), totals shows Subtotal + Total only --
        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        assert len(items_table.rows[0].cells) == 4
        totals = _table_with_first_cell(doc, "Subtotal")
        assert totals is not None
        assert len(totals.rows) == 2
        assert totals.rows[1].cells[1].text == "$1,000.00"

    def it_supports_a_NZ_15_percent_GST_rate(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 1, "unit_price": 100},
            ],
            default_gst_rate=0.15,
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        assert items_table.rows[1].cells[3].text == "$15.00"
        totals = _table_with_first_cell(doc, "Subtotal")
        assert totals is not None
        assert totals.rows[2].cells[1].text == "$115.00"

    def it_renders_payment_terms_when_supplied(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
            payment_terms="Net 30",
        )
        text = _full_text(doc)

        assert "Payment Terms" in text
        assert "Net 30" in text

    def it_renders_bank_details_when_supplied(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
            bank_details={
                "name": "Acme Corp Pty Ltd",
                "bsb": "062-001",
                "account": "1234 5678",
            },
        )
        text = _full_text(doc)

        assert "Payment Details" in text
        assert "Account Name" in text
        assert "Acme Corp Pty Ltd" in text
        assert "BSB" in text
        assert "062-001" in text
        assert "Account Number" in text
        assert "1234 5678" in text

    def it_renders_optional_notes(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
            notes="Thank you for your business.",
        )
        text = _full_text(doc)

        assert "Notes" in text
        assert "Thank you for your business." in text

    def it_renders_a_placeholder_when_no_items_supplied(self):
        doc = invoices.invoice(
            invoice_number="INV-001",
            issue_date="2026-03-15",
        )

        assert "[No line items supplied" in _full_text(doc)

    def it_raises_when_invoice_number_is_empty(self):
        with pytest.raises(ValueError, match="invoice_number is required"):
            invoices.invoice(invoice_number="", issue_date="2026-03-15")

    def it_raises_when_invoice_number_is_whitespace(self):
        with pytest.raises(ValueError, match="invoice_number is required"):
            invoices.invoice(invoice_number="   ", issue_date="2026-03-15")

    def it_raises_when_issue_date_is_empty(self):
        with pytest.raises(ValueError, match="issue_date is required"):
            invoices.invoice(invoice_number="INV-001", issue_date="")

    def it_raises_when_an_item_lacks_a_description(self):
        with pytest.raises(ValueError, match="non-empty 'description'"):
            invoices.invoice(
                invoice_number="INV-001",
                issue_date="2026-03-15",
                items=[{"unit_price": 100}],
            )

    def it_raises_when_an_item_lacks_unit_price(self):
        with pytest.raises(ValueError, match="missing 'unit_price'"):
            invoices.invoice(
                invoice_number="INV-001",
                issue_date="2026-03-15",
                items=[{"description": "X"}],
            )

    def it_raises_when_unit_price_is_non_numeric(self):
        with pytest.raises(ValueError, match="'unit_price' must be a number"):
            invoices.invoice(
                invoice_number="INV-001",
                issue_date="2026-03-15",
                items=[{"description": "X", "unit_price": "abc"}],
            )

    def it_raises_when_quantity_is_non_numeric(self):
        with pytest.raises(ValueError, match="'quantity' must be a number"):
            invoices.invoice(
                invoice_number="INV-001",
                issue_date="2026-03-15",
                items=[{"description": "X", "quantity": "abc", "unit_price": 100}],
            )

    def it_raises_when_an_item_is_not_a_mapping(self):
        with pytest.raises(ValueError, match="must be a mapping"):
            invoices.invoice(
                invoice_number="INV-001",
                issue_date="2026-03-15",
                items=["not a dict"],  # type: ignore[list-item]
            )

    def it_raises_when_gst_rate_out_of_range(self):
        with pytest.raises(ValueError, match="'gst_rate' must be between 0 and 1"):
            invoices.invoice(
                invoice_number="INV-001",
                issue_date="2026-03-15",
                items=[
                    {"description": "X", "unit_price": 100, "gst_rate": 1.5},
                ],
            )

    def it_raises_when_default_gst_rate_out_of_range(self):
        with pytest.raises(
            ValueError, match="default_gst_rate must be between 0 and 1"
        ):
            invoices.invoice(
                invoice_number="INV-001",
                issue_date="2026-03-15",
                default_gst_rate=2.0,
            )


# -- Factory: quote ------------------------------------------------------


class DescribeQuote:
    """Unit-test suite for ``invoices.quote``."""

    def it_returns_a_document_with_a_quote_header(self):
        doc = invoices.quote(
            quote_number="QU-2026-0099",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
        )

        assert isinstance(doc, DocumentCls)
        text = _full_text(doc)
        assert "Quote" in text
        assert "QU-2026-0099" in text

    def it_renders_the_valid_until_metadata_line(self):
        doc = invoices.quote(
            quote_number="QU-001",
            issue_date="2026-03-15",
            valid_until="2026-04-15",
            items=[{"description": "X", "unit_price": 100}],
        )
        text = _full_text(doc)

        assert "Valid Until" in text
        assert "2026-04-15" in text

    def it_omits_valid_until_when_unsupplied(self):
        doc = invoices.quote(
            quote_number="QU-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
        )

        assert "Valid Until" not in _full_text(doc)

    def it_uses_estimated_total_label_in_the_totals_block(self):
        doc = invoices.quote(
            quote_number="QU-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
        )

        totals = _table_with_first_cell(doc, "Subtotal")
        assert totals is not None
        # -- last row labels grand total "Estimated Total" --
        assert totals.rows[-1].cells[0].text == "Estimated Total"

    def it_renders_the_buyer_under_a_quote_to_label(self):
        doc = invoices.quote(
            quote_number="QU-001",
            issue_date="2026-03-15",
            buyer={"name": "Beta Pty Ltd"},
            items=[{"description": "X", "unit_price": 100}],
        )
        text = _full_text(doc)

        assert "Quote To" in text
        # -- and not the invoice's "Bill To" label --
        assert "Bill To" not in text

    def it_calculates_subtotal_GST_and_total(self):
        doc = invoices.quote(
            quote_number="QU-001",
            issue_date="2026-03-15",
            items=[
                {"description": "Consulting", "quantity": 10, "unit_price": 200},
            ],
        )

        totals = _table_with_first_cell(doc, "Subtotal")
        assert totals is not None
        assert totals.rows[0].cells[1].text == "$2,000.00"
        assert totals.rows[1].cells[1].text == "$200.00"
        assert totals.rows[2].cells[1].text == "$2,200.00"

    def it_supports_a_GST_free_quote(self):
        doc = invoices.quote(
            quote_number="QU-001",
            issue_date="2026-03-15",
            items=[{"description": "X", "unit_price": 100}],
            default_gst_rate=0,
        )

        items_table = _table_with_first_cell(doc, "Description")
        assert items_table is not None
        # -- four columns when no GST --
        assert len(items_table.rows[0].cells) == 4

    def it_raises_when_quote_number_is_empty(self):
        with pytest.raises(ValueError, match="quote_number is required"):
            invoices.quote(quote_number="", issue_date="2026-03-15")

    def it_raises_when_issue_date_is_empty(self):
        with pytest.raises(ValueError, match="issue_date is required"):
            invoices.quote(quote_number="QU-001", issue_date="")


# -- Factory: statement --------------------------------------------------


class DescribeStatement:
    """Unit-test suite for ``invoices.statement``."""

    def it_returns_a_document_with_a_statement_header(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
        )

        assert isinstance(doc, DocumentCls)
        assert "Statement" in _full_text(doc)

    def it_renders_the_period_metadata(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
        )
        text = _full_text(doc)

        assert "Period" in text
        assert "2026-03-01" in text
        assert "2026-03-31" in text

    def it_renders_the_optional_statement_number(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            statement_number="STMT-2026-03",
        )

        assert "STMT-2026-03" in _full_text(doc)

    def it_renders_the_buyer_under_a_statement_to_label(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            buyer={"name": "Beta Pty Ltd"},
        )
        text = _full_text(doc)

        assert "Statement To" in text
        assert "Beta Pty Ltd" in text

    def it_renders_invoices_as_a_four_column_table_when_no_statuses(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            invoices=[
                {"invoice_number": "INV-001", "date": "2026-03-05", "amount": 1000},
                {"invoice_number": "INV-002", "date": "2026-03-15", "amount": 2000},
            ],
        )

        invoices_table = _table_with_first_cell(doc, "Invoice")
        assert invoices_table is not None
        header = [c.text for c in invoices_table.rows[0].cells]
        assert header == ["Invoice", "Date", "Amount", "Balance"]
        assert len(invoices_table.rows) == 3

    def it_renders_a_status_column_when_any_record_has_a_status(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            invoices=[
                {
                    "invoice_number": "INV-001",
                    "date": "2026-03-05",
                    "amount": 1000,
                    "status": "Paid",
                },
                {
                    "invoice_number": "INV-002",
                    "date": "2026-03-15",
                    "amount": 2000,
                },
            ],
        )

        invoices_table = _table_with_first_cell(doc, "Invoice")
        assert invoices_table is not None
        header = [c.text for c in invoices_table.rows[0].cells]
        assert header == ["Invoice", "Date", "Amount", "Balance", "Status"]
        # -- second row has empty status (not supplied), first has "Paid" --
        assert invoices_table.rows[1].cells[4].text == "Paid"
        assert invoices_table.rows[2].cells[4].text == ""

    def it_renders_money_values_with_two_decimals(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            invoices=[
                {"invoice_number": "INV-001", "amount": 1000, "balance": 250},
            ],
        )

        invoices_table = _table_with_first_cell(doc, "Invoice")
        assert invoices_table is not None
        row = invoices_table.rows[1].cells
        assert row[2].text == "$1,000.00"
        assert row[3].text == "$250.00"

    def it_defaults_balance_to_amount_when_unsupplied(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            invoices=[
                {"invoice_number": "INV-001", "amount": 750},
            ],
        )

        invoices_table = _table_with_first_cell(doc, "Invoice")
        assert invoices_table is not None
        # -- balance mirrors amount --
        assert invoices_table.rows[1].cells[3].text == "$750.00"

    def it_renders_the_total_balance_owing_row(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            invoices=[
                {"invoice_number": "INV-001", "amount": 1000, "balance": 0},
                {"invoice_number": "INV-002", "amount": 2000, "balance": 2000},
                {"invoice_number": "INV-003", "amount": 500, "balance": 500},
            ],
        )

        totals = _table_with_first_cell(doc, "Total Balance Owing")
        assert totals is not None
        # -- 0 + 2000 + 500 = 2500 --
        assert totals.rows[0].cells[1].text == "$2,500.00"

    def it_renders_a_placeholder_when_no_invoices_supplied(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
        )

        assert "[No invoices recorded" in _full_text(doc)

    def it_right_aligns_numeric_cells_in_the_invoices_table(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            invoices=[
                {"invoice_number": "INV-001", "amount": 1000},
            ],
        )

        invoices_table = _table_with_first_cell(doc, "Invoice")
        assert invoices_table is not None
        row = invoices_table.rows[1].cells
        # -- Invoice + Date left, Amount + Balance right --
        assert row[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert row[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert row[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert row[3].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def it_accepts_number_as_an_alias_for_invoice_number(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            invoices=[{"number": "INV-005", "amount": 100}],
        )

        invoices_table = _table_with_first_cell(doc, "Invoice")
        assert invoices_table is not None
        assert invoices_table.rows[1].cells[0].text == "INV-005"

    def it_raises_when_period_start_is_empty(self):
        with pytest.raises(ValueError, match="period_start is required"):
            invoices.statement(period_start="", period_end="2026-03-31")

    def it_raises_when_period_end_is_empty(self):
        with pytest.raises(ValueError, match="period_end is required"):
            invoices.statement(period_start="2026-03-01", period_end="")

    def it_raises_when_an_invoice_record_has_no_invoice_number(self):
        with pytest.raises(ValueError, match="missing 'invoice_number'"):
            invoices.statement(
                period_start="2026-03-01",
                period_end="2026-03-31",
                invoices=[{"amount": 100}],
            )

    def it_raises_when_an_invoice_record_has_no_amount(self):
        with pytest.raises(ValueError, match="missing 'amount'"):
            invoices.statement(
                period_start="2026-03-01",
                period_end="2026-03-31",
                invoices=[{"invoice_number": "INV-001"}],
            )

    def it_raises_when_an_invoice_record_amount_is_non_numeric(self):
        with pytest.raises(ValueError, match="'amount' must be a number"):
            invoices.statement(
                period_start="2026-03-01",
                period_end="2026-03-31",
                invoices=[{"invoice_number": "INV-001", "amount": "abc"}],
            )

    def it_raises_when_an_invoice_record_is_not_a_mapping(self):
        with pytest.raises(ValueError, match="must be a mapping"):
            invoices.statement(
                period_start="2026-03-01",
                period_end="2026-03-31",
                invoices=["not a dict"],  # type: ignore[list-item]
            )


# -- Round-trip integration ----------------------------------------------


class DescribeInvoicesRoundTrip:
    """End-to-end smoke-tests: every factory produces a saveable document."""

    def it_can_save_an_invoice_to_a_BytesIO(self):
        doc = invoices.invoice(
            invoice_number="INV-2026-0042",
            issue_date="2026-03-15",
            due_date="2026-04-14",
            seller={
                "name": "Acme Corp",
                "abn": "12 345 678 901",
                "address": "123 Pitt Street\nSydney NSW 2000",
                "phone": "+61 2 1234 5678",
                "email": "billing@acme.com",
            },
            buyer={"name": "Beta Pty Ltd", "abn": "98 765 432 109"},
            items=[
                {"description": "Consulting", "quantity": 40, "unit_price": 250},
                {"description": "Travel", "quantity": 1, "unit_price": 580},
            ],
            payment_terms="Net 30",
            bank_details={
                "name": "Acme Corp Pty Ltd",
                "bsb": "062-001",
                "account": "1234 5678",
            },
            notes="Please quote the invoice number on payment.",
        )
        buf = BytesIO()
        doc.save(buf)
        # -- Word .docx is a zip; magic bytes are 'PK' --
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_quote_to_a_BytesIO(self):
        doc = invoices.quote(
            quote_number="QU-2026-0099",
            issue_date="2026-03-15",
            valid_until="2026-04-15",
            seller={"name": "Acme Corp", "abn": "12 345 678 901"},
            buyer={"name": "Beta Pty Ltd"},
            items=[
                {"description": "Discovery", "quantity": 1, "unit_price": 5000},
                {"description": "Implementation", "quantity": 1, "unit_price": 25000},
            ],
            notes="Quote is exclusive of travel and accommodation.",
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_statement_to_a_BytesIO(self):
        doc = invoices.statement(
            period_start="2026-03-01",
            period_end="2026-03-31",
            statement_number="STMT-2026-03",
            seller={"name": "Acme Corp", "abn": "12 345 678 901"},
            buyer={"name": "Beta Pty Ltd"},
            invoices=[
                {
                    "invoice_number": "INV-2026-0040",
                    "date": "2026-03-05",
                    "amount": 1100,
                    "balance": 0,
                    "status": "Paid",
                },
                {
                    "invoice_number": "INV-2026-0042",
                    "date": "2026-03-15",
                    "amount": 11638,
                    "balance": 11638,
                    "status": "Outstanding",
                },
            ],
            notes="Total outstanding is due by the end of next month.",
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"


# -- Module surface -------------------------------------------------------


class DescribeInvoicesModule:
    """Module-level surface contracts."""

    def it_exposes_invoice_quote_and_statement(self):
        assert hasattr(invoices, "invoice")
        assert hasattr(invoices, "quote")
        assert hasattr(invoices, "statement")
        assert "invoice" in invoices.__all__
        assert "quote" in invoices.__all__
        assert "statement" in invoices.__all__

    def it_exposes_the_default_GST_rate(self):
        # -- Australian GST rate, the canonical AUS default --
        assert invoices.DEFAULT_GST_RATE == 0.10
        assert "DEFAULT_GST_RATE" in invoices.__all__

    def it_is_re_exported_from_the_kit_package(self):
        from docx.kit import invoices as invoices_pkg

        assert invoices_pkg is invoices
