"""Tests for ``Document.iter_all_paragraphs`` / ``iter_all_runs`` /
``iter_all_pictures`` (#662).

These iterators are the public cross-story walker that kit modules
(:mod:`docx.kit.lint`, :mod:`docx.kit.brand`,
:mod:`docx.kit.stable_paragraph_ids`, …) compose against in place of
the previously-private ``docx.search._iter_all_paragraphs`` helper.
"""

from __future__ import annotations

import io
import os

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.shape import FloatingImage, InlineShape


# ---------------------------------------------------------------------------
# Fixtures — build a Document carrying content in every story.
# ---------------------------------------------------------------------------


_TEST_DIR = os.path.dirname(__file__)
# -- Reuse the python-docx test fixture image for inline-picture tests. --
_FIXTURE_IMAGE = os.path.join(
    os.path.dirname(_TEST_DIR), "test_files", "monty-truth.png"
)


def _document_with_every_story() -> DocumentCls:
    """Return a Document with one paragraph in every searchable story."""
    doc = Document()

    # -- body --
    doc.add_paragraph("body needle")

    # -- body table with two cells --
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "table 0,0 needle"
    table.cell(1, 1).text = "table 1,1 needle"

    # -- primary header / footer --
    doc.sections[0].header.paragraphs[0].text = "header needle"
    doc.sections[0].footer.paragraphs[0].text = "footer needle"

    # -- footnote / endnote / comment all anchored from a body run --
    run_fn = doc.add_paragraph("fn host").runs[0]
    doc.footnotes.add(run_fn, "footnote needle")

    run_en = doc.add_paragraph("en host").runs[0]
    doc.endnotes.add(run_en, "endnote needle")

    run_cm = doc.add_paragraph("cm host").runs[0]
    doc.add_comment(run_cm, "comment needle", author="Tester")

    return doc


# ---------------------------------------------------------------------------
# iter_all_paragraphs
# ---------------------------------------------------------------------------


class DescribeIterAllParagraphs:
    """Public Document.iter_all_paragraphs() — closes #662."""

    def it_yields_paragraph_location_pairs_for_every_story(self):
        doc = _document_with_every_story()

        pairs = list(doc.iter_all_paragraphs())
        locations = {loc for _, loc in pairs}

        assert "body" in locations
        assert any(
            loc.startswith("table:0:row:0:col:0") for loc in locations
        )
        assert any(
            loc.startswith("table:0:row:1:col:1") for loc in locations
        )
        assert "header:section0:primary" in locations
        assert "footer:section0:primary" in locations
        assert any(loc.startswith("footnote:") for loc in locations)
        assert any(loc.startswith("endnote:") for loc in locations)
        assert any(loc.startswith("comment:") for loc in locations)

    def it_yields_paragraph_proxies_not_raw_elements(self):
        doc = _document_with_every_story()

        for paragraph, _location in doc.iter_all_paragraphs():
            # -- public Paragraph proxies always expose .text --
            assert hasattr(paragraph, "text")
            assert hasattr(paragraph, "runs")

    def it_can_skip_tables_via_keyword_only_flag(self):
        doc = _document_with_every_story()

        locations = {
            loc for _, loc in doc.iter_all_paragraphs(include_tables=False)
        }

        assert not any(loc.startswith("table:") for loc in locations)
        # -- other stories are untouched --
        assert "body" in locations
        assert "header:section0:primary" in locations

    def it_can_skip_headers_and_footers(self):
        doc = _document_with_every_story()

        locations = {
            loc
            for _, loc in doc.iter_all_paragraphs(
                include_headers_footers=False
            )
        }

        assert not any(loc.startswith("header:") for loc in locations)
        assert not any(loc.startswith("footer:") for loc in locations)
        # -- tables / footnotes / etc. still surface --
        assert any(loc.startswith("table:") for loc in locations)
        assert any(loc.startswith("footnote:") for loc in locations)

    def it_can_skip_footnotes_endnotes_and_comments(self):
        doc = _document_with_every_story()

        locations = {
            loc
            for _, loc in doc.iter_all_paragraphs(
                include_footnotes=False,
                include_endnotes=False,
                include_comments=False,
            )
        }

        assert not any(loc.startswith("footnote:") for loc in locations)
        assert not any(loc.startswith("endnote:") for loc in locations)
        assert not any(loc.startswith("comment:") for loc in locations)
        # -- but body and tables still surface --
        assert "body" in locations
        assert any(loc.startswith("table:") for loc in locations)

    def it_always_yields_the_body_group_even_when_all_flags_are_off(self):
        doc = Document()
        doc.add_paragraph("only body")

        pairs = list(
            doc.iter_all_paragraphs(
                include_tables=False,
                include_headers_footers=False,
                include_footnotes=False,
                include_endnotes=False,
                include_comments=False,
            )
        )

        assert len(pairs) == 1
        assert pairs[0][0].text == "only body"
        assert pairs[0][1] == "body"

    def it_round_trips_through_save_and_reopen(self, tmp_path):
        doc = _document_with_every_story()
        path = tmp_path / "fixture.docx"
        doc.save(str(path))

        reopened = Document(str(path))
        locations = {loc for _, loc in reopened.iter_all_paragraphs()}

        assert "body" in locations
        assert "header:section0:primary" in locations
        assert "footer:section0:primary" in locations

    def it_is_a_lazy_iterator(self):
        """The method must return an iterator, not a materialised list."""
        doc = Document()
        doc.add_paragraph("hello")

        result = doc.iter_all_paragraphs()

        # -- generators / iterators expose __iter__ and __next__ --
        assert hasattr(result, "__iter__")
        assert hasattr(result, "__next__")


# ---------------------------------------------------------------------------
# iter_all_runs
# ---------------------------------------------------------------------------


class DescribeIterAllRuns:
    """Public Document.iter_all_runs() — closes #662."""

    def it_yields_runs_from_every_story(self):
        doc = _document_with_every_story()

        runs_by_location: dict[str, list[str]] = {}
        for run, location in doc.iter_all_runs():
            runs_by_location.setdefault(location, []).append(run.text)

        # -- body, table cells, header, footer, footnote, endnote, comment
        assert any(
            "body needle" in t
            for ts in runs_by_location.values()
            for t in ts
        )
        assert any(
            loc.startswith("header:") for loc in runs_by_location
        )
        assert any(
            loc.startswith("footer:") for loc in runs_by_location
        )
        assert any(
            loc.startswith("table:") for loc in runs_by_location
        )

    def it_descends_into_hyperlinks(self):
        """Confirm we expose runs that ``Paragraph.runs`` would skip."""
        doc = Document()
        paragraph = doc.add_paragraph("intro ")
        paragraph.add_hyperlink(
            url="https://example.com", text="link text", style=None
        )

        all_runs = list(doc.iter_all_runs())
        texts = [r.text for r, _loc in all_runs]

        assert "intro " in texts
        assert "link text" in texts

    def it_honours_the_include_tables_flag(self):
        doc = Document()
        doc.add_paragraph("body run")
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "cell run"

        body_only = list(doc.iter_all_runs(include_tables=False))
        all_with_tables = list(doc.iter_all_runs())

        body_texts = {r.text for r, _ in body_only}
        all_texts = {r.text for r, _ in all_with_tables}

        assert "body run" in body_texts
        assert "cell run" not in body_texts
        assert "cell run" in all_texts


# ---------------------------------------------------------------------------
# iter_all_pictures
# ---------------------------------------------------------------------------


class DescribeIterAllPictures:
    """Public Document.iter_all_pictures() — closes #662."""

    def it_yields_inline_pictures_from_the_body(self):
        if not os.path.exists(_FIXTURE_IMAGE):
            pytest.skip(f"missing test fixture image: {_FIXTURE_IMAGE}")
        doc = Document()
        doc.add_picture(_FIXTURE_IMAGE)

        pictures = list(doc.iter_all_pictures())

        assert len(pictures) == 1
        picture, location = pictures[0]
        assert isinstance(picture, InlineShape)
        assert location == "body"

    def it_yields_pictures_inside_table_cells(self):
        if not os.path.exists(_FIXTURE_IMAGE):
            pytest.skip(f"missing test fixture image: {_FIXTURE_IMAGE}")
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell_para = table.cell(0, 0).paragraphs[0]
        cell_para.add_run().add_picture(_FIXTURE_IMAGE)

        pictures = list(doc.iter_all_pictures())

        assert len(pictures) == 1
        _picture, location = pictures[0]
        assert location.startswith("table:0:row:0:col:0")

    def it_yields_pictures_in_headers(self):
        if not os.path.exists(_FIXTURE_IMAGE):
            pytest.skip(f"missing test fixture image: {_FIXTURE_IMAGE}")
        doc = Document()
        header_para = doc.sections[0].header.paragraphs[0]
        header_para.add_run().add_picture(_FIXTURE_IMAGE)

        pictures = list(doc.iter_all_pictures())

        assert len(pictures) == 1
        _picture, location = pictures[0]
        assert location == "header:section0:primary"

    def it_returns_an_empty_iterator_for_a_blank_document(self):
        doc = Document()
        doc.add_paragraph("no images")

        assert list(doc.iter_all_pictures()) == []

    def it_does_not_double_count_inline_pictures(self):
        """The walker must surface each inline picture exactly once."""
        if not os.path.exists(_FIXTURE_IMAGE):
            pytest.skip(f"missing test fixture image: {_FIXTURE_IMAGE}")
        doc = Document()
        doc.add_picture(_FIXTURE_IMAGE)
        doc.add_picture(_FIXTURE_IMAGE)

        pictures = list(doc.iter_all_pictures())

        assert len(pictures) == 2
        for picture, _loc in pictures:
            assert isinstance(picture, (InlineShape, FloatingImage))


# ---------------------------------------------------------------------------
# Sanity: the new public surface still backs docx.search.search_all_paragraphs.
# ---------------------------------------------------------------------------


class DescribeBackwardsCompatibility:
    """The public iterators must not break the existing search helpers."""

    def it_keeps_search_all_paragraphs_working(self):
        from docx.search import search_all_paragraphs

        doc = _document_with_every_story()

        matches = search_all_paragraphs(doc, "needle")

        # -- one match per seeded story: body + 2 cells + header + footer
        # -- + footnote + endnote + comment = 8 --
        assert len(matches) == 8

    def it_keeps_the_underscore_alias_callable(self):
        """The deprecated ``_iter_all_paragraphs`` private alias must still
        work for any third-party code that imported it before #662."""
        from docx.search import _iter_all_paragraphs

        doc = Document()
        doc.add_paragraph("hello")

        pairs = list(_iter_all_paragraphs(doc))

        # -- body group is always emitted --
        assert pairs[0][1] == "body"
