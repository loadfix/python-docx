"""Unit-test suite for ``docx.kit.back_matter`` helpers."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import back_matter
from docx.text.paragraph import Paragraph


@pytest.fixture
def document() -> DocumentCls:
    return Document()


class DescribeAddAppendix:
    """Unit-test suite for ``back_matter.add_appendix``."""

    def it_emits_a_label_and_title_heading(self, document: DocumentCls):
        result = back_matter.add_appendix(
            document, label="Appendix A", title="Data Tables"
        )

        # -- heading + page-break --
        assert len(result) == 2
        assert isinstance(result[0], Paragraph)
        assert result[0].text == "Appendix A: Data Tables"
        assert result[0].style is not None
        assert result[0].style.name == "Heading 1"

    def it_emits_one_paragraph_per_body_chunk(self, document: DocumentCls):
        result = back_matter.add_appendix(
            document,
            label="Appendix B",
            title="Methodology",
            body="Sampling notes.\n\nResponse rates.\n\nLimitations.",
        )

        # -- heading, p1, p2, p3, page-break --
        assert len(result) == 5
        assert result[1].text == "Sampling notes."
        assert result[2].text == "Response rates."
        assert result[3].text == "Limitations."

    def it_accepts_a_sequence_of_paragraphs(self, document: DocumentCls):
        result = back_matter.add_appendix(
            document,
            label="Appendix C",
            title="Glossary of acronyms",
            body=["RPO", "RTO", "MTBF"],
        )

        assert [p.text for p in result[1:-1]] == ["RPO", "RTO", "MTBF"]

    def it_drops_empty_chunks(self, document: DocumentCls):
        result = back_matter.add_appendix(
            document,
            label="Annex 1",
            title="Resources",
            body="One.\n\n\n\nTwo.",
        )

        assert len(result) == 4

    def it_renders_a_label_only_heading_when_title_is_empty(
        self, document: DocumentCls
    ):
        result = back_matter.add_appendix(document, label="Annex 1", title="")

        assert result[0].text == "Annex 1"

    def it_honours_a_custom_heading_level(self, document: DocumentCls):
        result = back_matter.add_appendix(
            document, label="Appendix A", title="Data", heading_level=2
        )

        assert result[0].style is not None
        assert result[0].style.name == "Heading 2"

    def it_skips_the_page_break_when_disabled(self, document: DocumentCls):
        result = back_matter.add_appendix(
            document,
            label="Appendix A",
            title="Data",
            body="One.",
            page_break=False,
        )

        # -- heading + body, no page-break --
        assert len(result) == 2

    def it_raises_when_label_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="label must be a non-empty string"):
            back_matter.add_appendix(document, label="", title="Data")

    def it_raises_when_heading_level_is_out_of_range(self, document: DocumentCls):
        with pytest.raises(ValueError, match="heading_level must be in 0..9"):
            back_matter.add_appendix(
                document, label="Appendix A", title="Data", heading_level=11
            )


class DescribeAddGlossary:
    """Unit-test suite for ``back_matter.add_glossary``."""

    def it_emits_a_heading_then_a_2_column_table_by_default(
        self, document: DocumentCls
    ):
        start_tables = len(document.tables)

        result = back_matter.add_glossary(
            document,
            entries={
                "API": "Application Programming Interface",
                "OOXML": "Office Open XML",
            },
        )

        # -- heading + page-break (table is reached via document.tables) --
        assert len(result) == 2
        assert result[0].text == "Glossary"

        assert len(document.tables) == start_tables + 1
        table = document.tables[-1]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.cell(0, 0).text == "API"
        assert table.cell(0, 1).text == "Application Programming Interface"
        assert table.cell(1, 0).text == "OOXML"
        assert table.cell(1, 1).text == "Office Open XML"

    def it_renders_terms_in_bold(self, document: DocumentCls):
        back_matter.add_glossary(
            document, entries={"API": "Application Programming Interface"}
        )

        table = document.tables[-1]
        term_para = table.cell(0, 0).paragraphs[0]
        assert all(run.bold for run in term_para.runs if run.text)

    def it_supports_the_list_layout(self, document: DocumentCls):
        result = back_matter.add_glossary(
            document,
            entries={
                "API": "Application Programming Interface",
                "OOXML": "Office Open XML",
            },
            layout="list",
        )

        # -- heading + 2 entries * 2 paras each + page-break = 6 --
        assert len(result) == 6
        assert result[0].text == "Glossary"
        assert result[1].text == "API"
        assert result[2].text == "Application Programming Interface"
        assert result[3].text == "OOXML"
        assert result[4].text == "Office Open XML"

    def it_omits_the_heading_when_title_is_None(self, document: DocumentCls):
        result = back_matter.add_glossary(
            document,
            entries={"API": "Application Programming Interface"},
            title=None,
        )

        # -- only the page-break, the table is not in the paragraph list --
        assert len(result) == 1

    def it_handles_an_empty_entries_mapping(self, document: DocumentCls):
        start_tables = len(document.tables)

        result = back_matter.add_glossary(document, entries={})

        # -- heading + page-break, no table emitted --
        assert len(result) == 2
        assert result[0].text == "Glossary"
        assert len(document.tables) == start_tables

    def it_skips_the_page_break_when_disabled(self, document: DocumentCls):
        result = back_matter.add_glossary(
            document,
            entries={"A": "alpha"},
            layout="list",
            page_break=False,
        )

        # -- heading + 2 list paras --
        assert len(result) == 3

    def it_raises_on_an_unknown_layout(self, document: DocumentCls):
        with pytest.raises(ValueError, match="layout must be one of"):
            back_matter.add_glossary(
                document, entries={"A": "alpha"}, layout="grid"
            )

    def it_raises_when_heading_level_is_out_of_range(self, document: DocumentCls):
        with pytest.raises(ValueError, match="heading_level must be in 0..9"):
            back_matter.add_glossary(
                document, entries={"A": "alpha"}, heading_level=42
            )


class DescribeAddIndex:
    """Unit-test suite for ``back_matter.add_index``."""

    def it_emits_a_heading_and_an_INDEX_complex_field(
        self, document: DocumentCls
    ):
        result = back_matter.add_index(document)

        # -- heading + INDEX paragraph + page-break --
        assert len(result) == 3
        assert result[0].text == "Index"
        assert result[0].style is not None
        assert result[0].style.name == "Heading 1"

        index_para = result[1]
        assert len(index_para.fields) == 1
        field = index_para.fields[0]
        assert field.is_complex is True
        assert field.type == "INDEX"

    def it_defaults_to_two_columns(self, document: DocumentCls):
        result = back_matter.add_index(document)

        instr = result[1].fields[0].instruction
        assert '\\c "2"' in instr

    def it_honours_a_custom_column_count(self, document: DocumentCls):
        result = back_matter.add_index(document, columns=1)

        instr = result[1].fields[0].instruction
        assert '\\c "1"' in instr

    def it_omits_the_heading_when_title_is_None(self, document: DocumentCls):
        result = back_matter.add_index(document, title=None)

        # -- INDEX paragraph + page-break --
        assert len(result) == 2
        assert result[0].fields[0].type == "INDEX"

    def it_skips_the_page_break_when_disabled(self, document: DocumentCls):
        result = back_matter.add_index(document, page_break=False)

        assert len(result) == 2

    def it_raises_when_columns_is_out_of_range(self, document: DocumentCls):
        with pytest.raises(ValueError, match="columns must be in 1..4"):
            back_matter.add_index(document, columns=5)

    def it_raises_when_columns_is_zero(self, document: DocumentCls):
        with pytest.raises(ValueError, match="columns must be in 1..4"):
            back_matter.add_index(document, columns=0)


class DescribeAddBibliography:
    """Unit-test suite for ``back_matter.add_bibliography``."""

    def it_emits_a_heading_and_one_paragraph_per_source(
        self, document: DocumentCls
    ):
        sources = [
            {
                "kind": "book",
                "authors": ["Knuth, D."],
                "title": "TAOCP",
                "year": 1968,
                "publisher": "Addison-Wesley",
            },
            {
                "kind": "book",
                "authors": ["Lamport, L."],
                "title": "LaTeX",
                "year": 1986,
                "publisher": "Addison-Wesley",
            },
        ]

        result = back_matter.add_bibliography(document, sources=sources)

        # -- heading, 2 source paragraphs, page-break --
        assert len(result) == 4
        assert result[0].text == "Bibliography"
        assert "Knuth, D." in result[1].text
        assert "TAOCP" in result[1].text
        assert "1968" in result[1].text
        assert "Addison-Wesley" in result[1].text

    def it_renders_books_in_an_APA_ish_shape(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "book",
                    "authors": ["Knuth, D."],
                    "title": "TAOCP",
                    "year": 1968,
                    "publisher": "Addison-Wesley",
                }
            ],
        )

        text = result[1].text
        assert text == "Knuth, D. (1968). TAOCP. Addison-Wesley."

    def it_joins_multiple_authors_with_an_ampersand(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "book",
                    "authors": ["Aho, A.", "Sethi, R.", "Ullman, J."],
                    "title": "Compilers",
                    "year": 1986,
                    "publisher": "Addison-Wesley",
                }
            ],
        )

        text = result[1].text
        assert "Aho, A., Sethi, R., & Ullman, J." in text

    def it_renders_articles_with_journal_volume_and_pages(
        self, document: DocumentCls
    ):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "article",
                    "authors": ["Codd, E."],
                    "title": "A Relational Model of Data",
                    "year": 1970,
                    "journal": "CACM",
                    "volume": 13,
                    "issue": 6,
                    "pages": "377-387",
                }
            ],
        )

        text = result[1].text
        assert "Codd, E. (1970)." in text
        assert "A Relational Model of Data." in text
        assert "CACM, 13(6), 377-387." in text

    def it_renders_web_sources_with_a_retrieval_url(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "web",
                    "authors": ["Berners-Lee, T."],
                    "title": "Information Management: A Proposal",
                    "year": 1989,
                    "site": "CERN",
                    "url": "https://example.org/proposal",
                }
            ],
        )

        text = result[1].text
        assert "Berners-Lee, T. (1989)." in text
        assert "CERN." in text
        assert "Retrieved from https://example.org/proposal" in text

    def it_renders_reports_with_a_report_number(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "report",
                    "authors": ["Acme Corp"],
                    "title": "Annual Findings",
                    "year": 2026,
                    "number": "AC-2026-04",
                    "publisher": "Acme",
                }
            ],
        )

        text = result[1].text
        assert "(Report No. AC-2026-04)" in text

    def it_uses_n_d_for_missing_year(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "book",
                    "authors": ["Anon."],
                    "title": "Untitled",
                    "publisher": "Self",
                }
            ],
        )

        assert "(n.d.)" in result[1].text

    def it_accepts_a_single_string_for_authors(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "book",
                    "author": "Knuth, D.",
                    "title": "TAOCP",
                    "year": 1968,
                    "publisher": "Addison-Wesley",
                }
            ],
        )

        assert result[1].text.startswith("Knuth, D. (1968).")

    def it_falls_back_to_key_value_for_unknown_kinds(
        self, document: DocumentCls
    ):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {"kind": "podcast", "host": "Alice", "title": "Show Time"},
            ],
        )

        text = result[1].text
        # -- "kind" is filtered; remaining keys are emitted as key=value pairs --
        assert "host=Alice" in text
        assert "title=Show Time" in text
        assert "kind=podcast" not in text

    def it_omits_the_heading_when_title_is_None(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "book",
                    "authors": ["Knuth, D."],
                    "title": "TAOCP",
                    "year": 1968,
                    "publisher": "Addison-Wesley",
                }
            ],
            title=None,
        )

        # -- source paragraph + page-break (no heading) --
        assert len(result) == 2
        assert "Knuth, D." in result[0].text

    def it_accepts_an_empty_sources_list(self, document: DocumentCls):
        result = back_matter.add_bibliography(document, sources=[])

        # -- heading + page-break --
        assert len(result) == 2

    def it_skips_the_page_break_when_disabled(self, document: DocumentCls):
        result = back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "book",
                    "authors": ["X"],
                    "title": "Y",
                    "year": 2020,
                    "publisher": "Z",
                }
            ],
            page_break=False,
        )

        assert len(result) == 2

    def it_raises_when_heading_level_is_out_of_range(self, document: DocumentCls):
        with pytest.raises(ValueError, match="heading_level must be in 0..9"):
            back_matter.add_bibliography(document, sources=[], heading_level=99)


class DescribeBackMatterIntegration:
    """End-to-end smoke test: every helper composes cleanly into one document."""

    def it_can_build_a_complete_back_matter_section(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        back_matter.add_appendix(
            document,
            label="Appendix A",
            title="Data Tables",
            body="See attached spreadsheets.",
        )
        back_matter.add_glossary(
            document,
            entries={
                "API": "Application Programming Interface",
                "OOXML": "Office Open XML",
            },
        )
        back_matter.add_index(document)
        back_matter.add_bibliography(
            document,
            sources=[
                {
                    "kind": "book",
                    "authors": ["Knuth, D."],
                    "title": "TAOCP",
                    "year": 1968,
                    "publisher": "Addison-Wesley",
                }
            ],
        )

        appended = len(document.paragraphs) - start
        assert appended > 0

        texts = [p.text for p in document.paragraphs[start:]]
        assert "Appendix A: Data Tables" in texts
        assert "Glossary" in texts
        assert "Index" in texts
        assert "Bibliography" in texts
        assert any("Knuth, D." in t for t in texts)
