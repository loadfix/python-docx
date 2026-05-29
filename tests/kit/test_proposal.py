"""Unit-test suite for ``docx.kit.proposal`` helpers."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import proposal
from docx.table import Table
from docx.text.paragraph import Paragraph


# -- Shared helpers -------------------------------------------------------


def _texts(document: DocumentCls):
    """Return the text of every paragraph in ``document``."""
    return [p.text for p in document.paragraphs]


def _full_text(document: DocumentCls) -> str:
    return "\n".join(_texts(document))


def _table_with_header(document: DocumentCls, *headers: str):
    """Return the first table whose header row matches ``headers`` exactly."""
    for table in document.tables:
        cells = [cell.text for cell in table.rows[0].cells]
        if tuple(cells) == headers:
            return table
    return None


@pytest.fixture
def document() -> DocumentCls:
    return Document()


# -- Common kwargs the suite reuses to keep tests legible. ----------------


_PROPOSAL_KWARGS = dict(
    title="Implementing Frobnitz at ACME",
    prepared_for="ACME Corp",
    prepared_by="Acme Consulting",
    date="2026-05-29",
    executive_summary="Frobnitz adoption is on the critical path.",
    problem_statement="ACME's current process is manual and brittle.",
    proposed_solution="Acme will deliver Frobnitz in three phases.",
    deliverables=["Discovery report", "Implementation", "30-day support"],
    timeline=[
        ("Week 1-2", "Discovery"),
        ("Week 3-6", "Implementation"),
        ("Week 7-10", "Rollout + support"),
    ],
    pricing=[
        {"item": "Discovery", "qty": 1, "rate": "$15,000", "total": "$15,000"},
        {"item": "Implementation", "qty": 1, "rate": "$60,000", "total": "$60,000"},
        {"item": "Support (30d)", "qty": 1, "rate": "$10,000", "total": "$10,000"},
    ],
    grand_total="$85,000",
    terms=["50% on signing", "50% on go-live", "Net 30"],
    next_steps=["Sign attached SOW", "Kick-off call within 5 business days"],
)


_SOW_KWARGS = dict(
    title="Statement of Work — Frobnitz Implementation",
    parties=("Acme Consulting Pty Ltd", "ACME Corp"),
    effective_date="2026-06-01",
    end_date="2026-08-31",
    scope="Acme Consulting will perform the following.",
    deliverables=[
        "Discovery report by Week 2",
        "Implementation by Week 6",
        "Documentation by Week 8",
    ],
    fees="Total: $85,000. Payable in two instalments.",
    acceptance_criteria=[
        "All test cases pass",
        "Documentation handed off",
        "Knowledge transfer completed",
    ],
)


# -- sales_proposal -------------------------------------------------------


class DescribeSalesProposal:
    """Unit-test suite for ``proposal.sales_proposal``."""

    def it_renders_the_title_in_the_Title_style(self, document: DocumentCls):
        result = proposal.sales_proposal(document, **_PROPOSAL_KWARGS)

        title_para = result[0]
        assert isinstance(title_para, Paragraph)
        assert title_para.text == _PROPOSAL_KWARGS["title"]
        assert title_para.style is not None
        assert title_para.style.name == "Title"
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_renders_the_disclaimer_into_the_document(
        self, document: DocumentCls
    ):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        text = _full_text(document)

        assert "DISCLAIMER" in text
        assert "not legal advice" in text
        assert "starting point" in text

    def it_renders_the_metadata_block(self, document: DocumentCls):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        text = _full_text(document)

        assert "Prepared for: ACME Corp" in text
        assert "Prepared by: Acme Consulting" in text
        assert "Date: 2026-05-29" in text

    def it_renders_each_top_level_heading(self, document: DocumentCls):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        headings = {
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "Heading 1"
        }

        assert "Executive Summary" in headings
        assert "Problem Statement" in headings
        assert "Proposed Solution" in headings
        assert "Deliverables" in headings
        assert "Timeline" in headings
        assert "Pricing" in headings
        assert "Terms" in headings
        assert "Next Steps" in headings

    def it_renders_the_three_free_text_sections(
        self, document: DocumentCls
    ):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        text = _full_text(document)

        assert _PROPOSAL_KWARGS["executive_summary"] in text
        assert _PROPOSAL_KWARGS["problem_statement"] in text
        assert _PROPOSAL_KWARGS["proposed_solution"] in text

    def it_renders_each_deliverable_as_a_list_bullet_paragraph(
        self, document: DocumentCls
    ):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        bullets = [
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "List Bullet"
        ]

        assert "Discovery report" in bullets
        assert "Implementation" in bullets
        assert "30-day support" in bullets

    def it_renders_each_next_step_as_a_list_number_paragraph(
        self, document: DocumentCls
    ):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        numbered = [
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "List Number"
        ]

        assert "Sign attached SOW" in numbered
        assert "Kick-off call within 5 business days" in numbered

    def it_renders_the_timeline_as_a_two_column_table(
        self, document: DocumentCls
    ):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        table = _table_with_header(document, "Period", "Activity")

        assert table is not None
        # -- header + 3 rows --
        assert len(table.rows) == 4
        assert table.rows[1].cells[0].text == "Week 1-2"
        assert table.rows[1].cells[1].text == "Discovery"
        assert table.rows[2].cells[0].text == "Week 3-6"
        assert table.rows[3].cells[0].text == "Week 7-10"

    def it_accepts_timeline_entries_as_mappings(
        self, document: DocumentCls
    ):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["timeline"] = [
            {"period": "Phase 1", "activity": "Discovery"},
            {"period": "Phase 2", "activity": "Build"},
        ]

        proposal.sales_proposal(document, **kwargs)
        table = _table_with_header(document, "Period", "Activity")

        assert table is not None
        assert table.rows[1].cells[0].text == "Phase 1"
        assert table.rows[2].cells[1].text == "Build"

    def it_renders_the_pricing_as_a_four_column_table(
        self, document: DocumentCls
    ):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        table = _table_with_header(document, "Item", "Qty", "Rate", "Total")

        assert table is not None
        # -- header + 3 line items + grand-total row --
        assert len(table.rows) == 5
        assert table.rows[1].cells[0].text == "Discovery"
        assert table.rows[1].cells[2].text == "$15,000"

    def it_renders_a_grand_total_row_when_supplied(
        self, document: DocumentCls
    ):
        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)
        table = _table_with_header(document, "Item", "Qty", "Rate", "Total")

        last = table.rows[-1].cells
        assert last[0].text == "Grand Total"
        assert last[3].text == "$85,000"

    def it_omits_the_grand_total_row_when_None(
        self, document: DocumentCls
    ):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["grand_total"] = None

        proposal.sales_proposal(document, **kwargs)
        table = _table_with_header(document, "Item", "Qty", "Rate", "Total")

        # -- header + 3 line items only --
        assert len(table.rows) == 4

    def it_skips_terms_when_empty(self, document: DocumentCls):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["terms"] = ()

        proposal.sales_proposal(document, **kwargs)
        headings = {
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "Heading 1"
        }

        assert "Terms" not in headings

    def it_skips_next_steps_when_empty(self, document: DocumentCls):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["next_steps"] = ()

        proposal.sales_proposal(document, **kwargs)
        headings = {
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "Heading 1"
        }

        assert "Next Steps" not in headings

    def it_appends_a_trailing_page_break_by_default(
        self, document: DocumentCls
    ):
        result = proposal.sales_proposal(document, **_PROPOSAL_KWARGS)

        # -- last block is the trailing page-break paragraph --
        last = result[-1]
        assert isinstance(last, Paragraph)
        # -- a page break is emitted as a paragraph carrying a w:br
        # -- with type="page" run. The simplest observable signal is
        # -- the paragraph's XML containing the marker. --
        assert "w:type=\"page\"" in last._p.xml or "type=\"page\"" in last._p.xml

    def it_skips_the_page_break_when_disabled(
        self, document: DocumentCls
    ):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["page_break"] = False

        result = proposal.sales_proposal(document, **kwargs)
        last = result[-1]

        # -- last block should NOT carry a page-break run --
        if isinstance(last, Paragraph):
            xml = last._p.xml
            assert "w:type=\"page\"" not in xml and "type=\"page\"" not in xml

    def it_returns_a_list_of_paragraphs_and_tables(
        self, document: DocumentCls
    ):
        result = proposal.sales_proposal(document, **_PROPOSAL_KWARGS)

        # -- every entry is either a Paragraph or a Table --
        assert all(isinstance(b, (Paragraph, Table)) for b in result)
        # -- and at least one of each kind is present --
        assert any(isinstance(b, Paragraph) for b in result)
        assert any(isinstance(b, Table) for b in result)

    def it_appends_to_an_existing_document_without_clobbering(
        self, document: DocumentCls
    ):
        document.add_paragraph("Existing intro paragraph.")
        before = len(document.paragraphs)

        proposal.sales_proposal(document, **_PROPOSAL_KWARGS)

        assert len(document.paragraphs) > before
        # -- existing paragraph survives at the head of the body --
        assert document.paragraphs[0].text == "Existing intro paragraph."

    @pytest.mark.parametrize(
        "field",
        [
            "title",
            "prepared_for",
            "prepared_by",
            "date",
            "executive_summary",
            "problem_statement",
            "proposed_solution",
        ],
    )
    def it_raises_when_required_string_is_empty(
        self, document: DocumentCls, field: str
    ):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs[field] = ""

        with pytest.raises(ValueError, match=field):
            proposal.sales_proposal(document, **kwargs)

    def it_raises_when_pricing_row_is_missing_item(
        self, document: DocumentCls
    ):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["pricing"] = [
            {"qty": 1, "rate": "$10,000", "total": "$10,000"},
        ]

        with pytest.raises(ValueError, match="item"):
            proposal.sales_proposal(document, **kwargs)

    def it_raises_when_pricing_row_is_not_a_mapping(
        self, document: DocumentCls
    ):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["pricing"] = ["not a dict"]

        with pytest.raises(ValueError, match="mapping"):
            proposal.sales_proposal(document, **kwargs)

    def it_handles_qty_as_a_numeric_value(self, document: DocumentCls):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["pricing"] = [
            {"item": "Hours", "qty": 40, "rate": "$200", "total": "$8,000"},
        ]

        proposal.sales_proposal(document, **kwargs)
        table = _table_with_header(document, "Item", "Qty", "Rate", "Total")

        assert table.rows[1].cells[1].text == "40"

    def it_handles_missing_optional_pricing_columns(
        self, document: DocumentCls
    ):
        kwargs = dict(_PROPOSAL_KWARGS)
        kwargs["pricing"] = [{"item": "Discount"}]
        kwargs["grand_total"] = None

        proposal.sales_proposal(document, **kwargs)
        table = _table_with_header(document, "Item", "Qty", "Rate", "Total")

        assert table.rows[1].cells[0].text == "Discount"
        assert table.rows[1].cells[1].text == ""
        assert table.rows[1].cells[2].text == ""
        assert table.rows[1].cells[3].text == ""


# -- sow ------------------------------------------------------------------


class DescribeSow:
    """Unit-test suite for ``proposal.sow``."""

    def it_renders_the_title_in_the_Title_style(self, document: DocumentCls):
        result = proposal.sow(document, **_SOW_KWARGS)

        title_para = result[0]
        assert isinstance(title_para, Paragraph)
        assert title_para.text == _SOW_KWARGS["title"]
        assert title_para.style is not None
        assert title_para.style.name == "Title"

    def it_renders_the_disclaimer(self, document: DocumentCls):
        proposal.sow(document, **_SOW_KWARGS)
        text = _full_text(document)

        assert "DISCLAIMER" in text
        assert "not legal advice" in text

    def it_renders_the_parties_metadata_line(
        self, document: DocumentCls
    ):
        proposal.sow(document, **_SOW_KWARGS)
        text = _full_text(document)

        assert "Parties: Acme Consulting Pty Ltd and ACME Corp" in text

    def it_renders_the_effective_and_end_date_metadata(
        self, document: DocumentCls
    ):
        proposal.sow(document, **_SOW_KWARGS)
        text = _full_text(document)

        assert "Effective Date: 2026-06-01" in text
        assert "End Date: 2026-08-31" in text

    def it_renders_each_section_heading(self, document: DocumentCls):
        proposal.sow(document, **_SOW_KWARGS)
        headings = {
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "Heading 1"
        }

        assert "Scope" in headings
        assert "Deliverables" in headings
        assert "Fees" in headings
        assert "Acceptance Criteria" in headings

    def it_renders_each_deliverable_as_a_bullet(
        self, document: DocumentCls
    ):
        proposal.sow(document, **_SOW_KWARGS)
        bullets = [
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "List Bullet"
        ]

        assert "Discovery report by Week 2" in bullets
        assert "Implementation by Week 6" in bullets
        assert "Documentation by Week 8" in bullets

    def it_renders_each_acceptance_criterion_as_a_bullet(
        self, document: DocumentCls
    ):
        proposal.sow(document, **_SOW_KWARGS)
        bullets = [
            p.text
            for p in document.paragraphs
            if p.style is not None and p.style.name == "List Bullet"
        ]

        assert "All test cases pass" in bullets
        assert "Documentation handed off" in bullets
        assert "Knowledge transfer completed" in bullets

    def it_renders_the_scope_and_fees_text(
        self, document: DocumentCls
    ):
        proposal.sow(document, **_SOW_KWARGS)
        text = _full_text(document)

        assert _SOW_KWARGS["scope"] in text
        assert _SOW_KWARGS["fees"] in text

    def it_appends_a_trailing_page_break_by_default(
        self, document: DocumentCls
    ):
        result = proposal.sow(document, **_SOW_KWARGS)
        last = result[-1]

        assert isinstance(last, Paragraph)
        assert "type=\"page\"" in last._p.xml

    def it_skips_the_page_break_when_disabled(
        self, document: DocumentCls
    ):
        kwargs = dict(_SOW_KWARGS)
        kwargs["page_break"] = False

        result = proposal.sow(document, **kwargs)
        last = result[-1]

        if isinstance(last, Paragraph):
            assert "type=\"page\"" not in last._p.xml

    def it_returns_a_list_of_paragraphs(self, document: DocumentCls):
        result = proposal.sow(document, **_SOW_KWARGS)

        assert all(isinstance(b, (Paragraph, Table)) for b in result)
        assert any(isinstance(b, Paragraph) for b in result)

    @pytest.mark.parametrize(
        "field",
        [
            "title",
            "effective_date",
            "end_date",
            "scope",
            "fees",
        ],
    )
    def it_raises_when_required_string_is_empty(
        self, document: DocumentCls, field: str
    ):
        kwargs = dict(_SOW_KWARGS)
        kwargs[field] = ""

        with pytest.raises(ValueError, match=field):
            proposal.sow(document, **kwargs)

    def it_raises_when_parties_has_fewer_than_two_entries(
        self, document: DocumentCls
    ):
        kwargs = dict(_SOW_KWARGS)
        kwargs["parties"] = ("Only one party",)

        with pytest.raises(ValueError, match="parties"):
            proposal.sow(document, **kwargs)

    def it_raises_when_a_party_name_is_empty(
        self, document: DocumentCls
    ):
        kwargs = dict(_SOW_KWARGS)
        kwargs["parties"] = ("Acme Consulting Pty Ltd", "")

        with pytest.raises(ValueError, match="parties"):
            proposal.sow(document, **kwargs)


# -- Module surface -------------------------------------------------------


class DescribeModuleSurface:
    """Sanity checks on the module's public API surface."""

    def it_exposes_sales_proposal_and_sow(self):
        assert "sales_proposal" in proposal.__all__
        assert "sow" in proposal.__all__
        assert callable(proposal.sales_proposal)
        assert callable(proposal.sow)

    def it_re_exports_the_module_under_docx_kit(self):
        from docx import kit

        assert kit.proposal is proposal

    def it_carries_the_not_legal_advice_disclaimer_in_its_docstring(self):
        # -- mirroring the contracts.py disclaimer convention --
        doc = proposal.__doc__ or ""
        assert "not legal advice" in doc.lower()
