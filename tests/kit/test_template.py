"""Unit-test suite for ``docx.kit.template`` ``.dotx`` placeholder loader."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest

from docx.api import Document as DocumentFactoryFn
from docx.document import Document as DocumentCls
from docx.kit import from_template_dotx
from docx.kit.template import from_template_dotx as helper
from docx.opc.constants import CONTENT_TYPE as CT


# -- Shared helpers -------------------------------------------------------


def _default_template_bytes() -> bytes:
    tpl_path = (
        Path(__file__).parent.parent.parent
        / "src"
        / "docx"
        / "templates"
        / "default.docx"
    )
    return tpl_path.read_bytes()


def _make_dotx_bytes() -> bytes:
    """Return the bundled default template re-stamped as a ``.dotx`` template package."""
    blob = _default_template_bytes()
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(blob), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        CT.WML_DOCUMENT_MAIN.encode("ascii"),
                        CT.WML_TEMPLATE_MAIN.encode("ascii"),
                    )
                zout.writestr(item, data)
    return out.getvalue()


def _make_dotx_with_paragraph(text: str) -> bytes:
    """Return a ``.dotx`` template whose body has a single paragraph with `text`."""
    document = DocumentFactoryFn()
    document.add_paragraph(text)
    buf = io.BytesIO()
    document.save(buf)
    plain_docx = buf.getvalue()
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(plain_docx), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        CT.WML_DOCUMENT_MAIN.encode("ascii"),
                        CT.WML_TEMPLATE_MAIN.encode("ascii"),
                    )
                zout.writestr(item, data)
    return out.getvalue()


def _make_dotx_with_table(cell_text: str) -> bytes:
    """Return a ``.dotx`` template carrying a one-cell table containing `cell_text`."""
    document = DocumentFactoryFn()
    table = document.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = cell_text
    buf = io.BytesIO()
    document.save(buf)
    plain_docx = buf.getvalue()
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(plain_docx), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        CT.WML_DOCUMENT_MAIN.encode("ascii"),
                        CT.WML_TEMPLATE_MAIN.encode("ascii"),
                    )
                zout.writestr(item, data)
    return out.getvalue()


# -- Tests ----------------------------------------------------------------


class DescribeFromTemplateDotx:
    """Unit-test suite for :func:`docx.kit.from_template_dotx`."""

    def it_loads_a_dotx_template(self):
        stream = io.BytesIO(_make_dotx_bytes())

        document = from_template_dotx(stream)

        assert isinstance(document, DocumentCls)

    def it_returns_a_document_with_the_docx_content_type(self):
        stream = io.BytesIO(_make_dotx_bytes())

        document = from_template_dotx(stream)

        # -- swapped from .dotx -> .docx by the underlying
        # -- ``Document.from_template`` --
        assert document.part.content_type == CT.WML_DOCUMENT_MAIN

    def it_substitutes_simple_placeholders_in_paragraph_text(self):
        stream = io.BytesIO(_make_dotx_with_paragraph("Hello [CLIENT]!"))

        document = from_template_dotx(
            stream, placeholders={"[CLIENT]": "ACME Corp"}
        )

        texts = [p.text for p in document.paragraphs]
        assert "Hello ACME Corp!" in texts

    def it_substitutes_multiple_placeholders(self):
        stream = io.BytesIO(
            _make_dotx_with_paragraph("[GREETING] [CLIENT], dated [DATE]")
        )

        document = from_template_dotx(
            stream,
            placeholders={
                "[GREETING]": "Dear",
                "[CLIENT]": "ACME Corp",
                "[DATE]": "2026-05-29",
            },
        )

        texts = [p.text for p in document.paragraphs]
        assert "Dear ACME Corp, dated 2026-05-29" in texts

    def it_substitutes_placeholders_in_table_cells(self):
        stream = io.BytesIO(_make_dotx_with_table("Customer: [CLIENT]"))

        document = from_template_dotx(
            stream, placeholders={"[CLIENT]": "ACME Corp"}
        )

        cell_text = document.tables[0].rows[0].cells[0].text
        assert "Customer: ACME Corp" in cell_text

    def it_preserves_run_formatting_during_substitution(self):
        # -- Build a template whose paragraph has a single bold run
        # -- containing the placeholder. The substitution must keep
        # -- the bold property after rewriting the run text. --
        document = DocumentFactoryFn()
        para = document.add_paragraph()
        run = para.add_run("Hello [CLIENT]!")
        run.bold = True
        buf = io.BytesIO()
        document.save(buf)
        plain_docx = buf.getvalue()

        out = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(plain_docx), "r") as zin:
            with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        data = data.replace(
                            CT.WML_DOCUMENT_MAIN.encode("ascii"),
                            CT.WML_TEMPLATE_MAIN.encode("ascii"),
                        )
                    zout.writestr(item, data)

        stream = io.BytesIO(out.getvalue())

        derived = from_template_dotx(
            stream, placeholders={"[CLIENT]": "ACME"}
        )

        runs = derived.paragraphs[0].runs
        assert any(r.text == "Hello ACME!" and r.bold for r in runs)

    def it_handles_no_placeholders_supplied(self):
        # -- ``placeholders=None`` is a valid call shape; the helper
        # -- just loads the template untouched. --
        stream = io.BytesIO(_make_dotx_with_paragraph("Hello [CLIENT]!"))

        document = from_template_dotx(stream)

        texts = [p.text for p in document.paragraphs]
        assert "Hello [CLIENT]!" in texts

    def it_handles_empty_placeholders_mapping(self):
        stream = io.BytesIO(_make_dotx_with_paragraph("Hello [CLIENT]!"))

        document = from_template_dotx(stream, placeholders={})

        texts = [p.text for p in document.paragraphs]
        assert "Hello [CLIENT]!" in texts

    def it_skips_placeholders_that_dont_match(self):
        stream = io.BytesIO(_make_dotx_with_paragraph("Hello [CLIENT]!"))

        document = from_template_dotx(
            stream, placeholders={"[UNUSED]": "X"}
        )

        texts = [p.text for p in document.paragraphs]
        assert "Hello [CLIENT]!" in texts

    def it_rejects_non_template_packages(self):
        # -- a plain .docx is not a .dotx template — surfaces the
        # -- underlying ``from_template`` error. --
        stream = io.BytesIO(_default_template_bytes())

        with pytest.raises(ValueError, match="not a Word template"):
            from_template_dotx(stream)

    def it_rejects_non_str_placeholder_keys(self):
        stream = io.BytesIO(_make_dotx_with_paragraph("Hello"))

        with pytest.raises(TypeError, match="placeholder keys must be str"):
            # type: ignore[dict-item]
            from_template_dotx(stream, placeholders={123: "X"})  # type: ignore[arg-type]

    def it_rejects_non_str_placeholder_values(self):
        stream = io.BytesIO(_make_dotx_with_paragraph("Hello"))

        with pytest.raises(TypeError, match="placeholder values must be str"):
            from_template_dotx(stream, placeholders={"[X]": 123})  # type: ignore[arg-type]

    def it_is_re_exported_from_docx_kit(self):
        # -- the kit-level re-export keeps the public API stable. --
        from docx.kit import from_template_dotx as re_exported

        assert re_exported is helper

    def it_accepts_a_pathlib_path(self, tmp_path):
        dotx_path = tmp_path / "tpl.dotx"
        dotx_path.write_bytes(_make_dotx_with_paragraph("Hello [CLIENT]!"))

        document = from_template_dotx(
            dotx_path, placeholders={"[CLIENT]": "ACME"}
        )

        texts = [p.text for p in document.paragraphs]
        assert "Hello ACME!" in texts

    def it_saves_as_docx_after_substitution(self, tmp_path):
        stream = io.BytesIO(_make_dotx_with_paragraph("Hello [CLIENT]!"))

        document = from_template_dotx(
            stream, placeholders={"[CLIENT]": "ACME"}
        )
        out_path = tmp_path / "out.docx"
        document.save(str(out_path))

        # -- reload to confirm the saved package round-trips --
        reloaded = DocumentFactoryFn(str(out_path))
        texts = [p.text for p in reloaded.paragraphs]
        assert "Hello ACME!" in texts
