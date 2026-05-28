"""Unit-test suite for ``docx.kit.contracts`` template factories."""

from __future__ import annotations

import pytest

from docx.document import Document as DocumentCls
from docx.kit import contracts


# -- Shared fixtures ------------------------------------------------------


@pytest.fixture
def parties():
    return [
        {"name": "Acme Corp", "address": "123 Pitt St, Sydney NSW 2000"},
        {"name": "Beta Inc", "address": "1 Market St, San Francisco CA"},
    ]


def _texts(document: DocumentCls):
    """Return the concatenated text of every paragraph in `document`."""
    return [p.text for p in document.paragraphs]


def _full_text(document: DocumentCls) -> str:
    return "\n".join(_texts(document))


# -- NDA ------------------------------------------------------------------


class DescribeNda:
    """Unit-test suite for ``contracts.nda``."""

    def it_returns_a_document_with_a_mutual_nda_title_by_default(
        self, parties
    ):
        doc = contracts.nda(parties=parties)

        assert isinstance(doc, DocumentCls)
        assert "Mutual Non-Disclosure Agreement" in _full_text(doc)

    def it_renders_a_one_way_title_when_kind_is_one_way(self, parties):
        doc = contracts.nda(kind="one-way", parties=parties)

        assert "One-Way Non-Disclosure Agreement" in _full_text(doc)

    def it_includes_the_six_required_sections(self, parties):
        doc = contracts.nda(parties=parties)
        text = _full_text(doc)

        # -- Acceptance criteria: definitions, mutual obligations, term,
        # -- return of materials, remedies, governing law. --
        assert "Definitions" in text
        assert "Obligations of Confidentiality" in text
        assert "Term" in text
        assert "Return of Materials" in text
        assert "Remedies" in text
        assert "Governing Law" in text

    def it_includes_the_disclaimer(self, parties):
        doc = contracts.nda(parties=parties)
        text = _full_text(doc)

        assert "DISCLAIMER" in text
        assert "not legal advice" in text

    def it_renders_the_parties_block(self, parties):
        doc = contracts.nda(parties=parties)
        text = _full_text(doc)

        assert "BETWEEN:" in text
        assert "Acme Corp" in text
        assert "Beta Inc" in text
        assert "(1)" in text
        assert "(2)" in text

    def it_renders_the_effective_date_when_supplied(self, parties):
        doc = contracts.nda(
            parties=parties, effective_date="2026-03-01"
        )

        assert "Effective Date: 2026-03-01" in _full_text(doc)

    def it_renders_the_term_in_years(self, parties):
        doc = contracts.nda(parties=parties, term_years=5)

        assert "5 years" in _full_text(doc)

    def it_renders_indefinite_term_when_zero(self, parties):
        doc = contracts.nda(parties=parties, term_years=0)

        assert "indefinitely" in _full_text(doc)

    def it_renders_singular_year_for_one_year_term(self, parties):
        doc = contracts.nda(parties=parties, term_years=1)
        text = _full_text(doc)

        assert "1 year," in text or "1 year " in text

    def it_renders_the_governing_law_jurisdiction(self, parties):
        doc = contracts.nda(
            parties=parties, governing_law="Victoria, Australia"
        )

        assert "Victoria, Australia" in _full_text(doc)

    def it_defaults_governing_law_to_NSW_Australia(self, parties):
        doc = contracts.nda(parties=parties)

        assert "New South Wales, Australia" in _full_text(doc)

    def it_includes_a_signature_block_for_each_party(self, parties):
        doc = contracts.nda(parties=parties)
        text = _full_text(doc)

        assert "Signed" in text
        assert "Signed for and on behalf of Acme Corp" in text
        assert "Signed for and on behalf of Beta Inc" in text
        # -- per party: Name / Title / Date placeholder lines --
        assert text.count("Name: ______________________________") >= 2
        assert text.count("Title: ______________________________") >= 2
        assert text.count("Date: ______________________________") >= 2

    def it_renders_mutual_obligation_phrasing_when_mutual(self, parties):
        doc = contracts.nda(kind="mutual", parties=parties)

        assert "Each Party shall" in _full_text(doc)

    def it_renders_one_way_obligation_phrasing_when_one_way(self, parties):
        doc = contracts.nda(kind="one-way", parties=parties)

        assert "The Receiving Party shall" in _full_text(doc)

    def it_raises_when_kind_is_invalid(self, parties):
        with pytest.raises(ValueError, match="kind must be"):
            contracts.nda(kind="bilateral", parties=parties)

    def it_raises_when_parties_missing(self):
        with pytest.raises(ValueError, match="parties is required"):
            contracts.nda()

    def it_raises_when_only_one_party_supplied(self):
        with pytest.raises(ValueError, match="at least two parties"):
            contracts.nda(parties=[{"name": "Solo"}])

    def it_raises_when_term_is_negative(self, parties):
        with pytest.raises(ValueError, match="term_years must be >= 0"):
            contracts.nda(parties=parties, term_years=-1)

    def it_raises_when_a_party_has_no_name(self):
        with pytest.raises(ValueError, match="non-empty 'name'"):
            contracts.nda(
                parties=[{"name": "Acme"}, {"address": "no name here"}]
            )


# -- MSA ------------------------------------------------------------------


class DescribeMsa:
    """Unit-test suite for ``contracts.msa``."""

    def it_builds_an_MSA_from_client_and_vendor_strings(self):
        doc = contracts.msa(client="Acme Corp", vendor="Beta Inc")

        assert isinstance(doc, DocumentCls)
        text = _full_text(doc)
        assert "Master Services Agreement" in text
        assert "Acme Corp" in text
        assert "Beta Inc" in text

    def it_includes_the_seven_required_sections(self):
        doc = contracts.msa(client="Acme", vendor="Beta")
        text = _full_text(doc)

        # -- services framework, ordering, payment, IP, warranties,
        # -- liability cap, termination --
        assert "Services Framework" in text
        assert "Ordering Process" in text
        assert "Fees and Payment" in text
        assert "Intellectual Property" in text
        assert "Warranties" in text
        assert "Liability" in text
        assert "Termination" in text

    def it_renders_the_payment_terms(self):
        doc = contracts.msa(
            client="Acme", vendor="Beta", payment_terms="Net 14"
        )

        assert "Net 14" in _full_text(doc)

    def it_defaults_payment_terms_to_Net_30(self):
        doc = contracts.msa(client="Acme", vendor="Beta")

        assert "Net 30" in _full_text(doc)

    def it_renders_the_currency(self):
        doc = contracts.msa(
            client="Acme", vendor="Beta", currency="USD"
        )

        assert "USD" in _full_text(doc)

    def it_defaults_currency_to_AUD(self):
        doc = contracts.msa(client="Acme", vendor="Beta")

        assert "AUD" in _full_text(doc)

    def it_references_the_Australian_Consumer_Law(self):
        doc = contracts.msa(client="Acme", vendor="Beta")

        assert "Australian Consumer Law" in _full_text(doc)

    def it_renders_a_one_times_liability_cap_by_default(self):
        doc = contracts.msa(client="Acme", vendor="Beta")
        text = _full_text(doc)

        assert "12 months preceding" in text
        # -- default is 1x — should not contain a multiple prefix --
        assert "2 times" not in text
        assert "3 times" not in text

    def it_renders_a_multi_times_liability_cap_when_supplied(self):
        doc = contracts.msa(
            client="Acme", vendor="Beta", liability_cap_multiple=3
        )

        assert "3 times the total fees" in _full_text(doc)

    def it_includes_the_disclaimer(self):
        doc = contracts.msa(client="Acme", vendor="Beta")

        assert "DISCLAIMER" in _full_text(doc)

    def it_supports_explicit_parties_with_addresses(self, parties):
        doc = contracts.msa(parties=parties)
        text = _full_text(doc)

        assert "Acme Corp of 123 Pitt St" in text
        assert "Beta Inc of 1 Market St" in text

    def it_includes_a_signature_block(self, parties):
        doc = contracts.msa(parties=parties)
        text = _full_text(doc)

        assert "Signed for and on behalf of Acme Corp" in text
        assert "Signed for and on behalf of Beta Inc" in text

    def it_raises_when_neither_parties_nor_client_vendor_supplied(self):
        with pytest.raises(ValueError, match="parties.*client.*vendor"):
            contracts.msa()

    def it_raises_when_only_client_supplied(self):
        with pytest.raises(ValueError, match="parties.*client.*vendor"):
            contracts.msa(client="Acme")

    def it_raises_when_liability_cap_multiple_is_negative(self):
        with pytest.raises(
            ValueError, match="liability_cap_multiple must be >= 0"
        ):
            contracts.msa(
                client="Acme", vendor="Beta", liability_cap_multiple=-1
            )


# -- SOW ------------------------------------------------------------------


class DescribeSow:
    """Unit-test suite for ``contracts.sow``."""

    def it_builds_a_SOW_with_a_parent_MSA_reference(self):
        doc = contracts.sow(
            parent_msa="msa-2026-001",
            client="Acme Corp",
            vendor="Beta Inc",
        )
        text = _full_text(doc)

        assert "Statement of Work" in text
        assert "msa-2026-001" in text

    def it_renders_the_project_name_in_the_title(self):
        doc = contracts.sow(
            client="Acme",
            vendor="Beta",
            project_name="Mobile App Build",
        )

        assert (
            "Statement of Work — Mobile App Build" in _full_text(doc)
        )

    def it_renders_a_numbered_list_of_deliverables(self):
        doc = contracts.sow(
            client="Acme",
            vendor="Beta",
            deliverables=[
                "Discovery report",
                "MVP build",
                "Production handover",
            ],
        )
        text = _full_text(doc)

        assert "D1: Discovery report" in text
        assert "D2: MVP build" in text
        assert "D3: Production handover" in text

    def it_renders_the_milestones_as_a_table(self):
        doc = contracts.sow(
            client="Acme",
            vendor="Beta",
            milestones=[
                {
                    "name": "Kickoff",
                    "date": "2026-03-15",
                    "payment": "AUD 10,000",
                },
                {
                    "name": "MVP",
                    "date": "2026-06-30",
                    "payment": "AUD 50,000",
                },
            ],
        )

        assert len(doc.tables) >= 1
        # -- find the milestones table by its header row --
        milestone_table = None
        for table in doc.tables:
            cells = [cell.text for cell in table.rows[0].cells]
            if cells and cells[0] == "Milestone":
                milestone_table = table
                break
        assert milestone_table is not None
        assert len(milestone_table.rows) == 3  # header + 2 milestones
        assert milestone_table.rows[1].cells[0].text == "Kickoff"
        assert milestone_table.rows[1].cells[1].text == "2026-03-15"
        assert milestone_table.rows[2].cells[0].text == "MVP"

    def it_falls_back_to_a_placeholder_when_milestones_missing(self):
        doc = contracts.sow(client="Acme", vendor="Beta")

        assert "[Insert milestone schedule here.]" in _full_text(doc)

    def it_renders_the_fees_text(self):
        doc = contracts.sow(
            client="Acme",
            vendor="Beta",
            fees="Fixed-price AUD 120,000 ex GST.",
        )

        assert "Fixed-price AUD 120,000 ex GST." in _full_text(doc)

    def it_includes_the_acceptance_criteria_clause(self):
        doc = contracts.sow(client="Acme", vendor="Beta")

        assert "Acceptance Criteria" in _full_text(doc)
        assert "10 business days" in _full_text(doc)

    def it_includes_an_assumptions_clause(self):
        doc = contracts.sow(client="Acme", vendor="Beta")

        assert "Assumptions and Dependencies" in _full_text(doc)

    def it_includes_a_signature_block(self):
        doc = contracts.sow(client="Acme", vendor="Beta")
        text = _full_text(doc)

        assert "Signed for and on behalf of Acme" in text
        assert "Signed for and on behalf of Beta" in text

    def it_renders_a_default_msa_reference_when_parent_missing(self):
        doc = contracts.sow(client="Acme", vendor="Beta")

        # -- without parent_msa, falls back to the generic phrasing --
        assert "Master Services Agreement" in _full_text(doc)

    def it_raises_when_neither_parties_nor_client_vendor_supplied(self):
        with pytest.raises(ValueError, match="parties.*client.*vendor"):
            contracts.sow()


# -- Contractor agreement -------------------------------------------------


class DescribeContractorAgreement:
    """Unit-test suite for ``contracts.contractor_agreement``."""

    def it_builds_a_contractor_agreement_from_two_strings(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
        )
        text = _full_text(doc)

        assert "Independent Contractor Agreement" in text
        assert "Charlie" in text
        assert "Acme Corp" in text

    def it_renders_a_numeric_day_rate_with_currency(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
            day_rate=1500,
        )

        # -- 1,500.00 with thousand separator + AUD --
        assert "AUD 1,500.00 per day" in _full_text(doc)

    def it_accepts_a_string_day_rate_verbatim(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
            day_rate="AUD 1,200 per day for design, AUD 1,800 for "
            "engineering",
        )

        assert "AUD 1,200 per day for design" in _full_text(doc)

    def it_renders_a_placeholder_when_day_rate_missing(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )

        assert "[Insert fee structure here.]" in _full_text(doc)

    def it_includes_the_independent_contractor_status_clause(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )
        text = _full_text(doc)

        assert "Independent Contractor Status" in text
        assert "not as an employee" in text

    def it_includes_the_taxation_and_superannuation_clause(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )
        text = _full_text(doc)

        assert "Taxation, Superannuation, and Insurance" in text
        assert "Superannuation Guarantee" in text

    def it_includes_the_intellectual_property_clause(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )
        text = _full_text(doc)

        assert "Intellectual Property" in text
        # -- moral rights consent is the AUS-specific feature --
        assert "moral rights" in text

    def it_includes_the_confidentiality_clause(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )

        assert "Confidentiality" in _full_text(doc)

    def it_renders_an_open_ended_term_when_term_months_missing(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )

        assert "until terminated" in _full_text(doc)

    def it_renders_a_fixed_term_when_term_months_supplied(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
            term_months=6,
        )

        assert "6 months" in _full_text(doc)

    def it_renders_singular_month_for_one_month_term(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
            term_months=1,
        )

        assert "1 month," in _full_text(doc)

    def it_includes_the_governing_law_clause(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
            governing_law="Queensland, Australia",
        )

        assert "Queensland, Australia" in _full_text(doc)

    def it_includes_the_disclaimer(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )

        assert "DISCLAIMER" in _full_text(doc)

    def it_includes_a_signature_block(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie", engaging_party="Acme Corp"
        )
        text = _full_text(doc)

        assert "Signed for and on behalf of Acme Corp" in text
        assert "Signed for and on behalf of Charlie" in text

    def it_renders_the_services_description(self):
        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
            services_description="Design and build the website.",
        )

        assert "Design and build the website." in _full_text(doc)

    def it_raises_when_neither_parties_nor_individuals_supplied(self):
        with pytest.raises(
            ValueError, match="parties.*contractor.*engaging_party"
        ):
            contracts.contractor_agreement()

    def it_raises_when_term_months_is_zero_or_negative(self):
        with pytest.raises(
            ValueError, match="term_months must be > 0"
        ):
            contracts.contractor_agreement(
                contractor="Charlie",
                engaging_party="Acme Corp",
                term_months=0,
            )


# -- Round-trip integration ----------------------------------------------


class DescribeContractsRoundTrip:
    """End-to-end smoke-tests: every factory produces a saveable document."""

    def it_can_save_an_NDA_to_a_BytesIO(self, parties, tmp_path):
        from io import BytesIO

        doc = contracts.nda(
            parties=parties,
            effective_date="2026-03-01",
            term_years=3,
        )
        buf = BytesIO()
        doc.save(buf)
        # -- Word .docx is a zip; the magic bytes are 'PK' --
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_an_MSA_to_a_BytesIO(self):
        from io import BytesIO

        doc = contracts.msa(
            client="Acme Corp",
            vendor="Beta Inc",
            payment_terms="Net 30",
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_SOW_to_a_BytesIO(self):
        from io import BytesIO

        doc = contracts.sow(
            parent_msa="msa-2026-001",
            client="Acme Corp",
            vendor="Beta Inc",
            project_name="Phase 1",
            deliverables=["A", "B"],
            milestones=[
                {"name": "Kickoff", "date": "2026-03-15", "payment": "AUD 10k"},
            ],
            fees="Fixed-price AUD 120,000 ex GST",
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_contractor_agreement_to_a_BytesIO(self):
        from io import BytesIO

        doc = contracts.contractor_agreement(
            contractor="Charlie",
            engaging_party="Acme Corp",
            day_rate=1500,
            term_months=6,
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"
