"""Unit-test suite for ``docx.kit.exec_summary`` helpers."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import exec_summary
from docx.text.paragraph import Paragraph


@pytest.fixture
def document() -> DocumentCls:
    return Document()


# -- Helper kwargs reused by several 1-pager tests. --
_ONE_PAGER_KWARGS = dict(
    title="Project Frobnitz: Q3 update",
    purpose="Decide whether to ship FrobnitzPro by end of Q3.",
    background="Customers have asked for FrobnitzPro since launch.",
    current_state="Engineering is 80% done; design review pending.",
    proposal="Ship in two waves — beta March 15, GA April 1.",
    risks=["Dependency on team B", "Holiday freeze cuts review window"],
    asks=["Approval to ship", "Reviewer time on Mar 5"],
)


class DescribeOnePager:
    """Unit-test suite for ``exec_summary.one_pager``."""

    def it_appends_the_title_as_Heading_1(self, document: DocumentCls):
        result = exec_summary.one_pager(document, **_ONE_PAGER_KWARGS)

        assert isinstance(result[0], Paragraph)
        assert result[0].text == "Project Frobnitz: Q3 update"
        assert result[0].style is not None
        assert result[0].style.name == "Heading 1"

    def it_emits_each_canonical_section_as_Heading_2(
        self, document: DocumentCls
    ):
        result = exec_summary.one_pager(document, **_ONE_PAGER_KWARGS)

        headings = [
            p.text
            for p in result
            if p.style is not None and p.style.name == "Heading 2"
        ]
        assert headings == [
            "Purpose",
            "Background",
            "Current state",
            "Proposal",
            "Risks",
            "Asks",
        ]

    def it_renders_list_inputs_as_bulleted_paragraphs(
        self, document: DocumentCls
    ):
        result = exec_summary.one_pager(document, **_ONE_PAGER_KWARGS)

        bullet_texts = [
            p.text
            for p in result
            if p.style is not None and p.style.name == "List Bullet"
        ]
        assert "Dependency on team B" in bullet_texts
        assert "Holiday freeze cuts review window" in bullet_texts
        assert "Approval to ship" in bullet_texts
        assert "Reviewer time on Mar 5" in bullet_texts

    def it_accepts_a_string_for_risks_and_asks(self, document: DocumentCls):
        kwargs = dict(_ONE_PAGER_KWARGS)
        kwargs["risks"] = "Single combined risk paragraph."
        kwargs["asks"] = "Single combined ask paragraph."

        result = exec_summary.one_pager(document, **kwargs)

        texts = [p.text for p in result]
        assert "Single combined risk paragraph." in texts
        assert "Single combined ask paragraph." in texts
        # -- a single string should NOT produce List Bullet paragraphs --
        bullet_texts = [
            p.text
            for p in result
            if p.style is not None and p.style.name == "List Bullet"
        ]
        assert "Single combined risk paragraph." not in bullet_texts
        assert "Single combined ask paragraph." not in bullet_texts

    def it_splits_prose_bodies_on_blank_lines(self, document: DocumentCls):
        kwargs = dict(_ONE_PAGER_KWARGS)
        kwargs["background"] = "First para.\n\nSecond para.\n\nThird."

        result = exec_summary.one_pager(document, **kwargs)

        texts = [p.text for p in result]
        assert "First para." in texts
        assert "Second para." in texts
        assert "Third." in texts

    def it_appends_a_page_break_by_default(self, document: DocumentCls):
        start = len(document.paragraphs)

        result = exec_summary.one_pager(document, **_ONE_PAGER_KWARGS)

        # -- there should be N appended paragraphs total --
        assert len(document.paragraphs) - start == len(result)
        # -- the last returned paragraph should match the last doc paragraph --
        assert result[-1].text == document.paragraphs[-1].text

    def it_skips_the_page_break_when_disabled(self, document: DocumentCls):
        with_break = exec_summary.one_pager(
            Document(), **_ONE_PAGER_KWARGS
        )
        without_break = exec_summary.one_pager(
            document, page_break=False, **_ONE_PAGER_KWARGS
        )

        assert len(without_break) == len(with_break) - 1

    def it_returns_paragraphs_in_document_order(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        result = exec_summary.one_pager(document, **_ONE_PAGER_KWARGS)

        # -- the returned list should match the tail of document.paragraphs --
        # -- (Paragraph proxies are re-created on each access, so compare
        # -- on text rather than identity). --
        appended_texts = [p.text for p in document.paragraphs[start:]]
        assert [p.text for p in result] == appended_texts

    def it_raises_when_title_is_empty(self, document: DocumentCls):
        kwargs = dict(_ONE_PAGER_KWARGS)
        kwargs["title"] = ""
        with pytest.raises(ValueError, match="title must be a non-empty string"):
            exec_summary.one_pager(document, **kwargs)

    def it_raises_when_purpose_is_empty(self, document: DocumentCls):
        kwargs = dict(_ONE_PAGER_KWARGS)
        kwargs["purpose"] = "   "
        with pytest.raises(ValueError, match="purpose must be a non-empty string"):
            exec_summary.one_pager(document, **kwargs)

    def it_raises_when_proposal_is_empty(self, document: DocumentCls):
        kwargs = dict(_ONE_PAGER_KWARGS)
        kwargs["proposal"] = ""
        with pytest.raises(ValueError, match="proposal must be a non-empty string"):
            exec_summary.one_pager(document, **kwargs)


class DescribeSixPager:
    """Unit-test suite for ``exec_summary.six_pager``."""

    @pytest.fixture
    def canonical_sections(self):
        # -- Amazon canonical 6-pager order --
        return {
            "Background": "Why this exists.",
            "Goals": "What we will accomplish.",
            "Tenets": ["Customer obsession", "Speed"],
            "State of the business": "Where we are today.",
            "Lessons learned": "What previous launches taught us.",
            "Strategic priorities": ["Foundation", "Growth", "Trust"],
            "Looking forward": "Twelve-month outlook.",
        }

    def it_appends_the_title_as_Heading_1(
        self, document: DocumentCls, canonical_sections
    ):
        result = exec_summary.six_pager(
            document,
            title="FrobnitzPro launch plan",
            sections=canonical_sections,
        )

        assert result[0].text == "FrobnitzPro launch plan"
        assert result[0].style is not None
        assert result[0].style.name == "Heading 1"

    def it_renders_each_section_heading_as_Heading_2(
        self, document: DocumentCls, canonical_sections
    ):
        result = exec_summary.six_pager(
            document, title="T", sections=canonical_sections
        )

        h2_texts = [
            p.text
            for p in result
            if p.style is not None and p.style.name == "Heading 2"
        ]
        assert h2_texts == list(canonical_sections.keys())

    def it_preserves_section_order(
        self, document: DocumentCls, canonical_sections
    ):
        result = exec_summary.six_pager(
            document, title="T", sections=canonical_sections
        )

        seen_headings = [
            p.text
            for p in result
            if p.style is not None and p.style.name == "Heading 2"
        ]
        assert seen_headings == [
            "Background",
            "Goals",
            "Tenets",
            "State of the business",
            "Lessons learned",
            "Strategic priorities",
            "Looking forward",
        ]

    def it_renders_list_section_bodies_as_bulleted_paragraphs(
        self, document: DocumentCls, canonical_sections
    ):
        result = exec_summary.six_pager(
            document, title="T", sections=canonical_sections
        )

        bullet_texts = [
            p.text
            for p in result
            if p.style is not None and p.style.name == "List Bullet"
        ]
        # -- Tenets and Strategic priorities are list-shaped sections --
        for item in (
            "Customer obsession",
            "Speed",
            "Foundation",
            "Growth",
            "Trust",
        ):
            assert item in bullet_texts

    def it_renders_string_section_bodies_as_paragraphs(
        self, document: DocumentCls, canonical_sections
    ):
        result = exec_summary.six_pager(
            document, title="T", sections=canonical_sections
        )

        texts = [p.text for p in result]
        assert "Why this exists." in texts
        assert "Where we are today." in texts
        assert "Twelve-month outlook." in texts

    def it_skips_the_page_break_when_disabled(
        self, document: DocumentCls, canonical_sections
    ):
        with_break = exec_summary.six_pager(
            Document(), title="T", sections=canonical_sections
        )
        without_break = exec_summary.six_pager(
            document,
            title="T",
            sections=canonical_sections,
            page_break=False,
        )

        assert len(without_break) == len(with_break) - 1

    def it_returns_paragraphs_in_document_order(
        self, document: DocumentCls, canonical_sections
    ):
        start = len(document.paragraphs)

        result = exec_summary.six_pager(
            document, title="T", sections=canonical_sections
        )

        appended_texts = [p.text for p in document.paragraphs[start:]]
        assert [p.text for p in result] == appended_texts

    def it_accepts_an_empty_string_body_as_heading_only(
        self, document: DocumentCls
    ):
        result = exec_summary.six_pager(
            document,
            title="T",
            sections={"Solo heading": ""},
        )

        h2_texts = [
            p.text
            for p in result
            if p.style is not None and p.style.name == "Heading 2"
        ]
        assert h2_texts == ["Solo heading"]

    def it_raises_when_title_is_empty(
        self, document: DocumentCls, canonical_sections
    ):
        with pytest.raises(ValueError, match="title must be a non-empty string"):
            exec_summary.six_pager(
                document, title="", sections=canonical_sections
            )

    def it_raises_when_sections_is_empty(self, document: DocumentCls):
        with pytest.raises(
            ValueError, match="sections must contain at least one entry"
        ):
            exec_summary.six_pager(document, title="T", sections={})

    def it_raises_when_a_section_heading_is_blank(
        self, document: DocumentCls
    ):
        with pytest.raises(
            ValueError, match="every section heading must be a non-empty string"
        ):
            exec_summary.six_pager(
                document,
                title="T",
                sections={"Background": "ok", "  ": "blank heading"},
            )


class DescribeExecSummaryIntegration:
    """End-to-end smoke-test: both helpers compose into one document."""

    def it_can_chain_one_pager_and_six_pager_in_a_single_doc(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        op = exec_summary.one_pager(document, **_ONE_PAGER_KWARGS)
        sp = exec_summary.six_pager(
            document,
            title="FrobnitzPro launch plan",
            sections={
                "Background": "Why this exists.",
                "Goals": "What we will accomplish.",
                "Tenets": ["Customer obsession", "Speed"],
            },
        )

        appended_texts = [p.text for p in document.paragraphs[start:]]
        # -- both helpers' returned lists, concatenated, must equal the
        # -- full set of newly-appended paragraphs in order
        # -- (compare on text since Paragraph proxies are re-instantiated
        # -- on each .paragraphs access). --
        assert [p.text for p in op + sp] == appended_texts
