"""Unit-test suite for ``docx.kit.dividers`` helpers (issue #89)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import dividers
from docx.kit.dividers import (
    add_chapter_break,
    add_divider,
    add_fleuron,
    add_three_stars,
)
from docx.shared import Pt
from docx.text.paragraph import Paragraph


@pytest.fixture
def document() -> DocumentCls:
    return Document()


class DescribeAddDivider:
    """Behavioural tests for :func:`docx.kit.dividers.add_divider`."""

    def it_appends_a_centred_paragraph_for_the_default_kind(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        result = add_divider(document)

        assert isinstance(result, Paragraph)
        assert len(document.paragraphs) == start + 1
        assert result.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_renders_the_line_kind_as_an_underlined_run(
        self, document: DocumentCls
    ):
        result = add_divider(document, kind="line")

        # -- one run, holding NBSP padding, with underline=True --
        assert len(result.runs) == 1
        run = result.runs[0]
        assert run.underline is True
        # -- the run text is non-empty whitespace (NBSPs) --
        assert run.text.strip(" ") == ""
        assert len(run.text) > 0

    def it_renders_the_dashed_kind_as_a_row_of_em_dashes(
        self, document: DocumentCls
    ):
        result = add_divider(document, kind="dashed")

        assert "—" in result.text  # EM DASH present
        assert result.text.count("—") >= 3

    def it_renders_the_dots_kind_as_a_row_of_middle_dots(
        self, document: DocumentCls
    ):
        result = add_divider(document, kind="dots")

        assert "·" in result.text  # MIDDLE DOT
        assert result.text.count("·") >= 3

    def it_renders_the_wave_kind_as_a_row_of_tildes(
        self, document: DocumentCls
    ):
        result = add_divider(document, kind="wave")

        assert "∼" in result.text  # TILDE OPERATOR
        assert result.text.count("∼") >= 3

    def it_raises_on_an_unknown_kind(self, document: DocumentCls):
        with pytest.raises(ValueError, match="kind must be one of"):
            add_divider(document, kind="zigzag")


class DescribeAddFleuron:
    """Behavioural tests for :func:`docx.kit.dividers.add_fleuron`."""

    def it_appends_a_centred_paragraph_with_the_default_glyph(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        result = add_fleuron(document)

        assert isinstance(result, Paragraph)
        assert len(document.paragraphs) == start + 1
        assert result.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert result.text == "❦"  # FLORAL HEART

    def it_accepts_a_caller_supplied_glyph(self, document: DocumentCls):
        result = add_fleuron(document, glyph="⁂")  # ASTERISM

        assert result.text == "⁂"

    def it_raises_when_glyph_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="glyph must be a non-empty"):
            add_fleuron(document, glyph="")


class DescribeAddThreeStars:
    """Behavioural tests for :func:`docx.kit.dividers.add_three_stars`."""

    def it_appends_a_centred_three_stars_paragraph(self, document: DocumentCls):
        start = len(document.paragraphs)

        result = add_three_stars(document)

        assert isinstance(result, Paragraph)
        assert len(document.paragraphs) == start + 1
        assert result.alignment == WD_ALIGN_PARAGRAPH.CENTER
        # -- three BLACK FOUR-POINTED STAR glyphs --
        assert result.text.count("✦") == 3

    def it_separates_the_three_glyphs_with_em_spaces(
        self, document: DocumentCls
    ):
        result = add_three_stars(document)

        # -- glyphs joined by U+2003 EM SPACE --
        assert " " in result.text
        assert result.text == "✦ ✦ ✦"

    def it_accepts_a_caller_supplied_glyph(self, document: DocumentCls):
        result = add_three_stars(document, glyph="*")

        assert result.text.count("*") == 3
        assert result.text == "* * *"

    def it_raises_when_glyph_is_empty(self, document: DocumentCls):
        with pytest.raises(ValueError, match="glyph must be a non-empty"):
            add_three_stars(document, glyph="")


class DescribeAddChapterBreak:
    """Behavioural tests for :func:`docx.kit.dividers.add_chapter_break`."""

    def it_returns_three_paragraphs_in_document_order(
        self, document: DocumentCls
    ):
        start = len(document.paragraphs)

        result = add_chapter_break(document)

        assert isinstance(result, list)
        assert len(result) == 3
        assert all(isinstance(p, Paragraph) for p in result)
        assert len(document.paragraphs) == start + 3
        # -- the returned paragraphs match the last three appended,
        # -- in document order.  ``Paragraph`` objects are reconstructed
        # -- on each ``document.paragraphs`` access so we compare by text
        # -- and centred-alignment shape rather than object identity. --
        last_three = document.paragraphs[-3:]
        assert [p.text for p in result] == [p.text for p in last_three]
        assert [p.alignment for p in result] == [
            p.alignment for p in last_three
        ]

    def it_centres_each_appended_paragraph(self, document: DocumentCls):
        result = add_chapter_break(document)

        for para in result:
            assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_uses_a_line_ornament_by_default(self, document: DocumentCls):
        leading, ornament, trailing = add_chapter_break(document)

        # -- leading + trailing are blank gap paragraphs --
        assert leading.text == ""
        assert trailing.text == ""
        # -- the ornament is the underline-line variant: one run, underlined --
        assert len(ornament.runs) == 1
        assert ornament.runs[0].underline is True

    def it_applies_the_default_36pt_spacing_either_side(
        self, document: DocumentCls
    ):
        leading, _ornament, trailing = add_chapter_break(document)

        assert leading.paragraph_format.space_after == Pt(36)
        assert trailing.paragraph_format.space_before == Pt(36)

    def it_honours_a_caller_supplied_spacing(self, document: DocumentCls):
        leading, _ornament, trailing = add_chapter_break(
            document, spacing=Pt(72)
        )

        assert leading.paragraph_format.space_after == Pt(72)
        assert trailing.paragraph_format.space_before == Pt(72)

    def it_dispatches_to_add_fleuron_for_a_fleuron_ornament(
        self, document: DocumentCls
    ):
        _leading, ornament, _trailing = add_chapter_break(
            document, ornament="fleuron"
        )

        assert ornament.text == "❦"  # default fleuron glyph

    def it_forwards_glyph_to_the_fleuron_helper(self, document: DocumentCls):
        _leading, ornament, _trailing = add_chapter_break(
            document, ornament="fleuron", glyph="❧"
        )

        assert ornament.text == "❧"

    def it_dispatches_to_add_three_stars_for_a_stars_ornament(
        self, document: DocumentCls
    ):
        _leading, ornament, _trailing = add_chapter_break(
            document, ornament="stars"
        )

        assert ornament.text.count("✦") == 3

    def it_forwards_glyph_to_the_three_stars_helper(
        self, document: DocumentCls
    ):
        _leading, ornament, _trailing = add_chapter_break(
            document, ornament="stars", glyph="*"
        )

        assert ornament.text == "* * *"

    @pytest.mark.parametrize("kind", ["line", "dashed", "dots", "wave"])
    def it_dispatches_each_divider_kind_through_the_ornament_argument(
        self, document: DocumentCls, kind: str
    ):
        _leading, ornament, _trailing = add_chapter_break(
            document, ornament=kind
        )

        # -- centred ornament paragraph appended in the middle slot --
        assert ornament.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_raises_on_an_unknown_ornament(self, document: DocumentCls):
        with pytest.raises(ValueError, match="ornament must be one of"):
            add_chapter_break(document, ornament="zigzag")


class DescribeKitReExport:
    """The dividers module is re-exported from ``docx.kit``."""

    def it_re_exports_the_dividers_submodule(self):
        from docx import kit

        assert hasattr(kit, "dividers")
        assert kit.dividers is dividers
        assert "dividers" in kit.__all__
