"""Unit-test suite for ``docx.kit.letterhead`` (issue #61)."""

from __future__ import annotations

from io import BytesIO

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.kit import letterhead
from docx.kit.letterhead import set_letterhead
from docx.shared import RGBColor

from ..unitutil.file import test_file


@pytest.fixture
def document() -> DocumentCls:
    return Document()


@pytest.fixture
def logo_path() -> str:
    return test_file("python-icon.png")


def _header(doc: DocumentCls):
    return doc.sections[0].header


def _footer(doc: DocumentCls):
    return doc.sections[0].footer


class DescribeSetLetterhead:
    """Behavioural tests for :func:`docx.kit.letterhead.set_letterhead`."""

    def it_returns_a_dict_with_header_and_footer_paragraph_lists(
        self, document: DocumentCls
    ):
        result = set_letterhead(
            document,
            return_address="123 Main St",
            phone="555-1234",
        )

        assert set(result) == {"header", "footer"}
        assert isinstance(result["header"], list)
        assert isinstance(result["footer"], list)
        assert len(result["header"]) >= 1
        assert len(result["footer"]) >= 1

    def it_writes_to_the_first_section_header_and_footer(
        self, document: DocumentCls
    ):
        set_letterhead(
            document,
            return_address="ACME HQ",
            phone="555-1234",
        )

        header_text = "\n".join(p.text for p in _header(document).paragraphs)
        footer_text = "\n".join(p.text for p in _footer(document).paragraphs)
        assert "ACME HQ" in header_text
        assert "555-1234" in footer_text

    def it_omits_the_logo_when_None(self, document: DocumentCls):
        result = set_letterhead(
            document, return_address="ACME HQ", phone="555-1234"
        )

        # No logo means no inline shapes in the header.
        header = _header(document)
        for para in header.paragraphs:
            for run in para.runs:
                # Runs holding pictures have a w:drawing element child.
                from docx.oxml.ns import qn

                assert run._r.find(qn("w:drawing")) is None
        assert len(result["header"]) >= 1

    def it_includes_the_logo_when_provided(
        self, document: DocumentCls, logo_path: str
    ):
        set_letterhead(
            document, logo=logo_path, return_address="ACME", phone="x"
        )

        from docx.oxml.ns import qn

        # At least one run somewhere in the header has a w:drawing child.
        found_drawing = False
        for para in _header(document).paragraphs:
            for run in para.runs:
                if run._r.find(qn("w:drawing")) is not None:
                    found_drawing = True
        assert found_drawing

    def it_clears_existing_header_content_before_writing(
        self, document: DocumentCls
    ):
        # Seed pre-existing header content.
        header = _header(document)
        header.paragraphs[0].text = "STALE LINE"
        header.add_paragraph("EXTRA STALE")

        set_letterhead(document, return_address="FRESH", phone="x")

        text_blob = "\n".join(p.text for p in _header(document).paragraphs)
        assert "STALE" not in text_blob
        assert "FRESH" in text_blob

    def it_is_idempotent_when_called_twice(self, document: DocumentCls):
        set_letterhead(document, return_address="ONE", phone="111")
        set_letterhead(document, return_address="TWO", phone="222")

        text_blob = "\n".join(p.text for p in _header(document).paragraphs)
        footer_text = "\n".join(p.text for p in _footer(document).paragraphs)
        assert "ONE" not in text_blob
        assert "TWO" in text_blob
        assert "111" not in footer_text
        assert "222" in footer_text

    def it_raises_when_style_is_unknown(self, document: DocumentCls):
        with pytest.raises(ValueError, match="style must be one of"):
            set_letterhead(
                document, return_address="x", phone="x", style="loud"
            )

    def it_emits_modern_classic_minimal_styles_without_error(
        self, document: DocumentCls, logo_path: str
    ):
        for style_name in letterhead.STYLES:
            doc = Document()
            result = set_letterhead(
                doc,
                logo=logo_path,
                return_address="A St\nCity 1234",
                phone="+61 2 1234 5678",
                email="hello@acme.com",
                website="acme.com",
                style=style_name,
                color="primary",
            )
            assert result["header"]
            assert result["footer"]


class DescribeStyles:
    """Per-style visual contract tests."""

    def it_modern_centers_the_footer(
        self, document: DocumentCls, logo_path: str
    ):
        result = set_letterhead(
            document,
            logo=logo_path,
            return_address="ACME",
            phone="555",
            email="x@y.z",
            website="acme.com",
            style="modern",
        )

        assert (
            result["footer"][0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        )

    def it_modern_separates_footer_fields_with_a_middle_dot(
        self, document: DocumentCls
    ):
        set_letterhead(
            document,
            phone="555",
            email="x@y.z",
            website="acme.com",
            style="modern",
        )

        footer_blob = "\n".join(
            p.text for p in _footer(document).paragraphs
        )
        # The middle-dot separator
        assert "·" in footer_blob

    def it_classic_centers_the_address_paragraph(
        self, document: DocumentCls
    ):
        set_letterhead(
            document,
            return_address="123 Main\nCity",
            phone="555",
            style="classic",
        )

        address_paragraphs = [
            p
            for p in _header(document).paragraphs
            if "Main" in p.text or "123" in p.text
        ]
        assert address_paragraphs
        assert address_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_classic_emits_a_horizontal_rule_in_header_and_footer(
        self, document: DocumentCls
    ):
        set_letterhead(
            document,
            return_address="123 Main",
            phone="555",
            style="classic",
        )

        header_blob = "\n".join(
            p.text for p in _header(document).paragraphs
        )
        footer_blob = "\n".join(
            p.text for p in _footer(document).paragraphs
        )
        # Em-dash rule in both
        assert "—" * 5 in header_blob
        assert "—" * 5 in footer_blob

    def it_classic_emits_one_paragraph_per_contact_field(
        self, document: DocumentCls
    ):
        set_letterhead(
            document,
            phone="555",
            email="x@y.z",
            website="acme.com",
            style="classic",
        )

        # 1 rule + 3 contact fields
        non_empty = [
            p for p in _footer(document).paragraphs if p.text.strip()
        ]
        assert len(non_empty) == 4

    def it_minimal_uses_pipe_separators_in_footer(
        self, document: DocumentCls
    ):
        set_letterhead(
            document,
            phone="555",
            email="x@y.z",
            website="acme.com",
            style="minimal",
        )

        footer_blob = "\n".join(
            p.text for p in _footer(document).paragraphs
        )
        assert " | " in footer_blob

    def it_minimal_flattens_address_to_a_single_line(
        self, document: DocumentCls
    ):
        set_letterhead(
            document,
            return_address="Line A\nLine B\nLine C",
            phone="555",
            style="minimal",
        )

        header_blob = "\n".join(
            p.text for p in _header(document).paragraphs
        )
        # Comma-flattened, on one line
        assert "Line A, Line B, Line C" in header_blob


class DescribeColorResolution:
    """Color helper covers named presets, hex, RGBColor, and theme tokens."""

    def it_resolves_a_named_preset(self, document: DocumentCls):
        result = set_letterhead(
            document,
            return_address="ACME",
            phone="x",
            color="primary",
        )

        # The address run carries the colour
        address_para = result["header"][0]
        # Last run holds the visible address text
        text_runs = [r for r in address_para.runs if r.text.strip()]
        assert text_runs
        assert text_runs[-1].font.color.rgb == RGBColor(0x1F, 0x4E, 0x79)

    def it_accepts_a_hex_string(self, document: DocumentCls):
        result = set_letterhead(
            document,
            return_address="ACME",
            phone="x",
            color="#FF8800",
        )

        text_runs = [r for r in result["header"][0].runs if r.text.strip()]
        assert text_runs[-1].font.color.rgb == RGBColor(0xFF, 0x88, 0x00)

    def it_accepts_an_rgbcolor_instance(self, document: DocumentCls):
        rgb = RGBColor(0x12, 0x34, 0x56)
        result = set_letterhead(
            document,
            return_address="ACME",
            phone="x",
            color=rgb,
        )

        text_runs = [r for r in result["header"][0].runs if r.text.strip()]
        assert text_runs[-1].font.color.rgb == rgb

    def it_resolves_a_theme_token_through_the_document_theme(
        self, document: DocumentCls
    ):
        # Default-template documents have a theme; the accent1 slot
        # in the default theme is 4F81BD.
        result = set_letterhead(
            document,
            return_address="ACME",
            phone="x",
            color="accent1",
        )

        text_runs = [r for r in result["header"][0].runs if r.text.strip()]
        assert text_runs[-1].font.color.rgb == RGBColor(0x4F, 0x81, 0xBD)

    def it_falls_back_to_a_named_preset_when_no_theme_is_available(self):
        # _resolve_color with document=None bypasses the theme lookup
        # and should land on the fallback preset.
        rgb = letterhead._resolve_color("accent1", document=None)
        assert rgb == letterhead._NAMED_COLORS["primary"]

    def it_returns_None_for_None_color(self):
        assert letterhead._resolve_color(None) is None

    def it_raises_when_color_is_not_a_string_or_RGBColor(self):
        with pytest.raises(ValueError, match="color must be"):
            letterhead._resolve_color(42)


class DescribeHyperlinks:
    """Email and website are emitted as hyperlinks."""

    def it_renders_email_as_a_mailto_hyperlink(self, document: DocumentCls):
        set_letterhead(
            document, phone="x", email="hello@acme.com", style="modern"
        )

        # Inspect the footer part for an external relationship to mailto:
        footer_part = _footer(document).part
        rels = footer_part.rels
        targets = [r.target_ref for r in rels.values() if r.is_external]
        assert any(t.startswith("mailto:hello@acme.com") for t in targets)

    def it_normalises_a_bare_website_domain_to_https(
        self, document: DocumentCls
    ):
        set_letterhead(document, website="acme.com", style="modern")

        footer_part = _footer(document).part
        targets = [
            r.target_ref
            for r in footer_part.rels.values()
            if r.is_external
        ]
        assert any(t == "https://acme.com" for t in targets)

    def it_passes_through_a_full_url_unchanged(
        self, document: DocumentCls
    ):
        set_letterhead(
            document, website="http://acme.com/x", style="modern"
        )

        footer_part = _footer(document).part
        targets = [
            r.target_ref
            for r in footer_part.rels.values()
            if r.is_external
        ]
        assert any(t == "http://acme.com/x" for t in targets)


class DescribeRoundTrip:
    """End-to-end save/reload test: letterhead survives a serialisation cycle."""

    def it_can_be_saved_and_reloaded(
        self, document: DocumentCls, logo_path: str
    ):
        set_letterhead(
            document,
            logo=logo_path,
            return_address="ACME\nCity 1234",
            phone="+61 2 1234 5678",
            email="hello@acme.com",
            website="acme.com",
            style="classic",
            color="primary",
        )

        buf = BytesIO()
        document.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        # The header should still hold the address and rule.
        header_blob = "\n".join(
            p.text for p in reloaded.sections[0].header.paragraphs
        )
        footer_blob = "\n".join(
            p.text for p in reloaded.sections[0].footer.paragraphs
        )
        assert "ACME" in header_blob
        assert "—" * 5 in header_blob
        assert "+61 2 1234 5678" in footer_blob
        # Hyperlinks become text on the reloaded paragraphs.
        assert "hello@acme.com" in footer_blob
        assert "acme.com" in footer_blob
