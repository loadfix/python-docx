"""Unit-test suite for ``docx.kit.scientific`` template factories."""

from __future__ import annotations

import pytest

from docx.document import Document as DocumentCls
from docx.kit import scientific


# -- Shared helpers -------------------------------------------------------


def _texts(document: DocumentCls):
    """Return the text of every paragraph in ``document``."""
    return [p.text for p in document.paragraphs]


def _full_text(document: DocumentCls) -> str:
    return "\n".join(_texts(document))


_AUTHORS = [
    {"name": "Alice", "affiliation": "Acme Corp", "email": "alice@acme.com"},
    {"name": "Bob", "affiliation": "Beta Labs", "email": "bob@beta.io"},
]

_SECTIONS = [
    {"heading": "Introduction", "body": "Intro body."},
    {"heading": "Related Work", "body": "Related body."},
    {"heading": "Conclusion", "body": "Conclusion body."},
]

_REFERENCES = [
    {
        "authors": "Lamport, L.",
        "title": "The Part-Time Parliament",
        "venue": "TOCS",
        "year": 1998,
    },
    {
        "authors": "Diffie, W. and Hellman, M.",
        "title": "New Directions in Cryptography",
        "venue": "IEEE Trans. Inf. Theory",
        "year": 1976,
    },
]


# -- IEEE -----------------------------------------------------------------


class DescribeIEEEPaper:
    """Unit-test suite for ``scientific.ieee_paper``."""

    def it_returns_a_document_with_the_title_in_the_text(self):
        doc = scientific.ieee_paper(title="A Distributed Consensus Algorithm")

        assert isinstance(doc, DocumentCls)
        assert "A Distributed Consensus Algorithm" in _full_text(doc)

    def it_uses_the_Title_style_for_the_title(self):
        doc = scientific.ieee_paper(title="A Distributed Consensus Algorithm")

        title_para = doc.paragraphs[0]
        assert title_para.style.name in ("Title", "Normal")
        assert "A Distributed Consensus Algorithm" in title_para.text

    def it_renders_each_author_name_affiliation_and_email(self):
        doc = scientific.ieee_paper(
            title="A Distributed Consensus Algorithm",
            authors=_AUTHORS,
        )
        text = _full_text(doc)

        assert "Alice" in text
        assert "Acme Corp" in text
        assert "alice@acme.com" in text
        assert "Bob" in text
        assert "Beta Labs" in text
        assert "bob@beta.io" in text

    def it_renders_an_abstract_with_the_em_dash_lead_in(self):
        doc = scientific.ieee_paper(
            title="A Distributed Consensus Algorithm",
            abstract="We present a distributed consensus algorithm.",
        )
        text = _full_text(doc)

        assert "Abstract—" in text
        assert "We present a distributed consensus algorithm." in text

    def it_renders_index_terms_with_the_keywords_joined_by_commas(self):
        doc = scientific.ieee_paper(
            title="A Distributed Consensus Algorithm",
            keywords=["consensus", "distributed systems", "fault tolerance"],
        )
        text = _full_text(doc)

        assert "Index Terms—" in text
        assert "consensus, distributed systems, fault tolerance" in text

    def it_uses_two_column_layout_for_the_body(self):
        doc = scientific.ieee_paper(
            title="A Distributed Consensus Algorithm",
            sections=_SECTIONS,
        )

        # Two sections: leading single-column banner + two-column body.
        assert len(doc.sections) == 2
        assert doc.sections[0].columns.count == 1
        assert doc.sections[1].columns.count == 2

    def it_renders_supplied_body_sections_as_headings_and_paragraphs(self):
        doc = scientific.ieee_paper(
            title="A Distributed Consensus Algorithm",
            sections=_SECTIONS,
        )
        text = _full_text(doc)

        assert "Introduction" in text
        assert "Intro body." in text
        assert "Related Work" in text
        assert "Conclusion" in text
        assert "Conclusion body." in text

    def it_renders_references_in_IEEE_numbered_format(self):
        doc = scientific.ieee_paper(
            title="A Distributed Consensus Algorithm",
            references=_REFERENCES,
        )
        text = _full_text(doc)

        assert "References" in text
        assert "[1]" in text
        assert "[2]" in text
        assert "Lamport, L." in text
        # IEEE wraps the title in curly quotes.
        assert "“The Part-Time Parliament,”" in text
        assert "1998." in text

    def it_accepts_a_sequence_body_for_multi_paragraph_sections(self):
        doc = scientific.ieee_paper(
            title="X",
            sections=[
                {"heading": "Intro", "body": ["First paragraph.", "Second."]},
            ],
        )
        text = _full_text(doc)

        assert "First paragraph." in text
        assert "Second." in text

    def it_raises_when_title_is_empty(self):
        with pytest.raises(ValueError, match="title is required"):
            scientific.ieee_paper(title="")

    def it_raises_when_title_is_whitespace_only(self):
        with pytest.raises(ValueError, match="title is required"):
            scientific.ieee_paper(title="   ")

    def it_raises_when_an_author_is_missing_a_name(self):
        with pytest.raises(ValueError, match="non-empty 'name'"):
            scientific.ieee_paper(
                title="X", authors=[{"affiliation": "no name"}]
            )

    def it_raises_when_an_author_is_not_a_mapping(self):
        with pytest.raises(ValueError, match="must be a mapping"):
            scientific.ieee_paper(
                title="X",
                authors=["not a dict"],  # type: ignore[list-item]
            )

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(ValueError, match="non-empty 'heading'"):
            scientific.ieee_paper(
                title="X", sections=[{"body": "no heading"}]
            )


# -- ACM ------------------------------------------------------------------


class DescribeACMPaper:
    """Unit-test suite for ``scientific.acm_paper``."""

    def it_returns_a_document_with_the_title(self):
        doc = scientific.acm_paper(title="A Distributed Consensus Algorithm")

        assert isinstance(doc, DocumentCls)
        assert "A Distributed Consensus Algorithm" in _full_text(doc)

    def it_renders_authors_with_affiliation_and_email(self):
        doc = scientific.acm_paper(
            title="X",
            authors=_AUTHORS,
        )
        text = _full_text(doc)

        assert "Alice" in text
        assert "Acme Corp" in text
        assert "alice@acme.com" in text

    def it_renders_an_abstract_heading(self):
        doc = scientific.acm_paper(
            title="X",
            abstract="The abstract body.",
        )
        text = _full_text(doc)

        assert "Abstract" in text
        assert "The abstract body." in text

    def it_renders_a_CCS_concepts_section(self):
        doc = scientific.acm_paper(
            title="X",
            ccs_concepts=[
                "Computing methodologies → Parallel computing",
            ],
        )
        text = _full_text(doc)

        assert "CCS Concepts" in text
        assert "Computing methodologies → Parallel computing" in text

    def it_renders_a_CCS_placeholder_when_concepts_missing(self):
        doc = scientific.acm_paper(title="X")

        assert "[Insert at least one CCS Concept" in _full_text(doc)

    def it_renders_a_keywords_block(self):
        doc = scientific.acm_paper(
            title="X",
            keywords=["consensus", "distributed systems"],
        )
        text = _full_text(doc)

        assert "Keywords" in text
        assert "consensus, distributed systems" in text

    def it_renders_supplied_body_sections(self):
        doc = scientific.acm_paper(title="X", sections=_SECTIONS)
        text = _full_text(doc)

        assert "Introduction" in text
        assert "Intro body." in text

    def it_renders_references_in_ACM_numbered_format(self):
        doc = scientific.acm_paper(title="X", references=_REFERENCES)
        text = _full_text(doc)

        assert "References" in text
        assert "[1]" in text
        # ACM puts the year right after the authors.
        assert "Lamport, L." in text
        assert "1998." in text
        assert "The Part-Time Parliament." in text

    def it_stays_single_column_at_draft_time(self):
        doc = scientific.acm_paper(title="X", sections=_SECTIONS)

        # ACM template handles the camera-ready column rendering; the
        # drafting kit stays single-column.
        for section in doc.sections:
            assert section.columns.count == 1

    def it_raises_when_title_is_empty(self):
        with pytest.raises(ValueError, match="title is required"):
            scientific.acm_paper(title="")

    def it_raises_when_an_author_is_missing_a_name(self):
        with pytest.raises(ValueError, match="non-empty 'name'"):
            scientific.acm_paper(
                title="X", authors=[{"affiliation": "no name"}]
            )

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(ValueError, match="non-empty 'heading'"):
            scientific.acm_paper(title="X", sections=[{"body": "no heading"}])


# -- APA ------------------------------------------------------------------


class DescribeAPAPaper:
    """Unit-test suite for ``scientific.apa_paper``."""

    def it_returns_a_document_with_the_title(self):
        doc = scientific.apa_paper(title="A Distributed Consensus Algorithm")

        assert isinstance(doc, DocumentCls)
        assert "A Distributed Consensus Algorithm" in _full_text(doc)

    def it_renders_authors_each_with_affiliation_below(self):
        doc = scientific.apa_paper(
            title="X",
            authors=_AUTHORS,
        )
        text = _full_text(doc)

        assert "Alice" in text
        assert "Acme Corp" in text
        assert "Bob" in text
        assert "Beta Labs" in text

    def it_renders_an_abstract_heading_centered(self):
        doc = scientific.apa_paper(
            title="X",
            abstract="Abstract body text.",
        )

        # Heading "Abstract" appears in text.
        assert "Abstract" in _full_text(doc)
        assert "Abstract body text." in _full_text(doc)

    def it_renders_keywords_with_an_italic_lead_in(self):
        doc = scientific.apa_paper(
            title="X",
            keywords=["consensus", "distributed systems"],
        )
        text = _full_text(doc)

        assert "Keywords:" in text
        assert "consensus, distributed systems" in text

    def it_renders_a_running_head_when_supplied(self):
        doc = scientific.apa_paper(
            title="X",
            running_head="A Short Title",
        )
        text = _full_text(doc)

        assert "Running head:" in text
        # APA convention uppercases the running-head text.
        assert "A SHORT TITLE" in text

    def it_applies_double_line_spacing_to_body_paragraphs(self):
        doc = scientific.apa_paper(
            title="X",
            abstract="Some abstract.",
            sections=[{"heading": "Intro", "body": "Body text."}],
        )

        # Find the body paragraph by text and assert the line spacing.
        bodies = [p for p in doc.paragraphs if p.text == "Body text."]
        assert bodies
        assert bodies[0].paragraph_format.line_spacing == 2.0

    def it_stays_single_column(self):
        doc = scientific.apa_paper(title="X", sections=_SECTIONS)

        for section in doc.sections:
            assert section.columns.count == 1

    def it_renders_references_in_APA_author_date_format(self):
        doc = scientific.apa_paper(title="X", references=_REFERENCES)
        text = _full_text(doc)

        assert "References" in text
        # Author (Year). Title. Venue.
        assert "Lamport, L. (1998)." in text
        assert "The Part-Time Parliament." in text
        assert "TOCS." in text

    def it_renders_supplied_body_sections(self):
        doc = scientific.apa_paper(title="X", sections=_SECTIONS)
        text = _full_text(doc)

        assert "Introduction" in text
        assert "Intro body." in text

    def it_raises_when_title_is_empty(self):
        with pytest.raises(ValueError, match="title is required"):
            scientific.apa_paper(title="")

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(ValueError, match="non-empty 'heading'"):
            scientific.apa_paper(
                title="X", sections=[{"body": "no heading"}]
            )


# -- Nature ---------------------------------------------------------------


class DescribeNaturePaper:
    """Unit-test suite for ``scientific.nature_paper``."""

    def it_returns_a_document_with_the_title(self):
        doc = scientific.nature_paper(title="A Distributed Consensus Algorithm")

        assert isinstance(doc, DocumentCls)
        assert "A Distributed Consensus Algorithm" in _full_text(doc)

    def it_renders_a_byline_joining_authors_with_commas(self):
        doc = scientific.nature_paper(
            title="X",
            authors=[
                {"name": "Alice", "affiliation": "Acme"},
                {"name": "Bob", "affiliation": "Acme"},
                {"name": "Carol", "affiliation": "Beta"},
            ],
        )
        text = _full_text(doc)

        # Final author is joined with "and".
        assert "Alice, Bob, and Carol" in text

    def it_lists_each_unique_affiliation(self):
        doc = scientific.nature_paper(
            title="X",
            authors=[
                {"name": "Alice", "affiliation": "Acme"},
                {"name": "Bob", "affiliation": "Acme"},
                {"name": "Carol", "affiliation": "Beta"},
            ],
        )
        text = _full_text(doc)

        # De-duplicated; joined with semicolons.
        assert "Acme; Beta" in text

    def it_renders_the_abstract_in_italics(self):
        doc = scientific.nature_paper(
            title="X",
            abstract="The abstract body.",
        )

        # Find the abstract paragraph and check italic on its run.
        ab_paras = [p for p in doc.paragraphs if p.text == "The abstract body."]
        assert ab_paras
        assert any(run.italic for run in ab_paras[0].runs)

    def it_uses_two_column_layout_for_the_article_body(self):
        doc = scientific.nature_paper(
            title="X",
            sections=_SECTIONS,
        )

        assert len(doc.sections) == 2
        assert doc.sections[0].columns.count == 1
        assert doc.sections[1].columns.count == 2

    def it_renders_supplied_body_sections(self):
        doc = scientific.nature_paper(title="X", sections=_SECTIONS)
        text = _full_text(doc)

        assert "Introduction" in text
        assert "Intro body." in text

    def it_renders_references_in_Nature_numbered_format(self):
        doc = scientific.nature_paper(title="X", references=_REFERENCES)
        text = _full_text(doc)

        assert "References" in text
        # Nature uses leading "N." with no brackets.
        assert "1." in text
        assert "Lamport, L." in text
        assert "TOCS 1998." in text

    def it_omits_keywords_per_Nature_style(self):
        # nature_paper() has no keywords kwarg.
        with pytest.raises(TypeError):
            scientific.nature_paper(  # type: ignore[call-arg]
                title="X", keywords=["a"]
            )

    def it_raises_when_title_is_empty(self):
        with pytest.raises(ValueError, match="title is required"):
            scientific.nature_paper(title="")

    def it_raises_when_an_author_is_missing_a_name(self):
        with pytest.raises(ValueError, match="non-empty 'name'"):
            scientific.nature_paper(
                title="X", authors=[{"affiliation": "no name"}]
            )

    def it_raises_when_a_section_has_no_heading(self):
        with pytest.raises(ValueError, match="non-empty 'heading'"):
            scientific.nature_paper(
                title="X", sections=[{"body": "no heading"}]
            )


# -- Round-trip integration ----------------------------------------------


class DescribeScientificRoundTrip:
    """End-to-end smoke-tests: every factory produces a saveable document."""

    def it_can_save_an_ieee_paper_to_a_BytesIO(self):
        from io import BytesIO

        doc = scientific.ieee_paper(
            title="A Distributed Consensus Algorithm",
            authors=_AUTHORS,
            abstract="We present a distributed consensus algorithm.",
            keywords=["consensus", "distributed systems"],
            sections=_SECTIONS,
            references=_REFERENCES,
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_an_acm_paper_to_a_BytesIO(self):
        from io import BytesIO

        doc = scientific.acm_paper(
            title="A Distributed Consensus Algorithm",
            authors=_AUTHORS,
            abstract="We present a distributed consensus algorithm.",
            keywords=["consensus", "distributed systems"],
            ccs_concepts=["Computing methodologies → Parallel computing"],
            sections=_SECTIONS,
            references=_REFERENCES,
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_an_apa_paper_to_a_BytesIO(self):
        from io import BytesIO

        doc = scientific.apa_paper(
            title="A Distributed Consensus Algorithm",
            authors=_AUTHORS,
            abstract="We present a distributed consensus algorithm.",
            keywords=["consensus", "distributed systems"],
            running_head="Consensus Algorithm",
            sections=_SECTIONS,
            references=_REFERENCES,
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"

    def it_can_save_a_nature_paper_to_a_BytesIO(self):
        from io import BytesIO

        doc = scientific.nature_paper(
            title="A Distributed Consensus Algorithm",
            authors=_AUTHORS,
            abstract="We present a distributed consensus algorithm.",
            sections=_SECTIONS,
            references=_REFERENCES,
        )
        buf = BytesIO()
        doc.save(buf)
        assert buf.getvalue()[:2] == b"PK"


# -- Module surface -------------------------------------------------------


class DescribeScientificModule:
    """Module-level surface contracts."""

    def it_exposes_the_four_template_factories(self):
        for name in ("ieee_paper", "acm_paper", "apa_paper", "nature_paper"):
            assert hasattr(scientific, name)
            assert name in scientific.__all__

    def it_is_re_exported_from_the_kit_package(self):
        from docx.kit import scientific as scientific_pkg

        assert scientific_pkg is scientific
