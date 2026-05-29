"""Unit-test suite for ``docx.kit.stable_paragraph_ids`` (issue #301)."""

from __future__ import annotations

import io
import re

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import stable_paragraph_ids
from docx.kit.stable_paragraph_ids import (
    ensure,
    get,
    id_of,
    iter_with_ids,
    set_id,
)
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


_PARA_ID_QN = qn("w14:paraId")


@pytest.fixture
def document() -> DocumentCls:
    """A fresh Document().  Note: Document() ships with **zero**
    paragraphs in its body — every test that needs one calls
    ``add_paragraph`` first."""
    return Document()


@pytest.fixture
def document_with_table() -> DocumentCls:
    """A document containing two body paragraphs straddling a 2x2
    table whose cells each carry one paragraph (the default
    new-cell paragraph).  Exercises the "walks into table cells"
    contract."""
    doc = Document()
    doc.add_paragraph("Body paragraph 1")
    table = doc.add_table(rows=2, cols=2)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.paragraphs[0].text = f"r{row_idx}c{col_idx}"
    doc.add_paragraph("Body paragraph 2")
    return doc


class DescribeEnsure:
    """Behavioural tests for :func:`ensure`."""

    def it_stamps_an_id_on_every_body_paragraph_lacking_one(
        self, document: DocumentCls
    ):
        document.add_paragraph("First")
        document.add_paragraph("Second")
        document.add_paragraph("Third")

        stamped = ensure(document)

        assert stamped == 3
        for p in document.paragraphs:
            value = p._p.get(_PARA_ID_QN)
            assert value is not None
            assert re.match(r"^[0-9A-F]{8}$", value), value

    def it_walks_into_table_cells_including_nested_paragraphs(
        self, document_with_table: DocumentCls
    ):
        ensure(document_with_table)

        # -- Every body paragraph is stamped.
        for p in document_with_table.paragraphs:
            assert p._p.get(_PARA_ID_QN) is not None
        # -- Every cell paragraph is stamped, too.
        for table in document_with_table.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        assert p._p.get(_PARA_ID_QN) is not None

    def it_is_idempotent_a_second_call_changes_nothing(
        self, document: DocumentCls
    ):
        document.add_paragraph("alpha")
        document.add_paragraph("beta")
        ensure(document)
        before = [p._p.get(_PARA_ID_QN) for p in document.paragraphs]

        stamped = ensure(document)

        after = [p._p.get(_PARA_ID_QN) for p in document.paragraphs]
        assert stamped == 0
        assert before == after

    def it_returns_the_count_of_newly_stamped_paragraphs(
        self, document: DocumentCls
    ):
        document.add_paragraph("one")
        document.add_paragraph("two")
        document.add_paragraph("three")
        # -- pre-stamp only the first paragraph
        document.paragraphs[0]._p.set(_PARA_ID_QN, "DEADBEEF")

        stamped = ensure(document)

        # -- two newly stamped (the existing "DEADBEEF" was preserved)
        assert stamped == 2
        assert document.paragraphs[0]._p.get(_PARA_ID_QN) == "DEADBEEF"

    def it_mints_uppercase_hex_matching_words_own_shape(
        self, document: DocumentCls
    ):
        document.add_paragraph("only")
        ensure(document)
        for p in document.paragraphs:
            value = p._p.get(_PARA_ID_QN)
            # -- exactly 8 hex chars, uppercase
            assert len(value) == 8
            assert value == value.upper()
            int(value, 16)  # parses

    def it_returns_zero_for_an_empty_document(self, document: DocumentCls):
        assert ensure(document) == 0


class DescribeIdOf:
    def it_returns_None_when_no_id_is_set(self, document: DocumentCls):
        document.add_paragraph("alpha")
        assert id_of(document.paragraphs[0]) is None

    def it_returns_the_stamped_id(self, document: DocumentCls):
        document.add_paragraph("alpha")
        ensure(document)
        para = document.paragraphs[0]
        expected = para._p.get(_PARA_ID_QN)
        assert id_of(para) == expected

    def it_treats_the_empty_string_as_no_id(self, document: DocumentCls):
        # -- an empty string would round-trip to an empty attribute value
        # -- which Word's reader treats as absent — match that semantic.
        document.add_paragraph("alpha")
        document.paragraphs[0]._p.set(_PARA_ID_QN, "")
        assert id_of(document.paragraphs[0]) is None


class DescribeSetId:
    def it_sets_the_id_on_a_paragraph(self, document: DocumentCls):
        document.add_paragraph("alpha")
        para = document.paragraphs[0]

        set_id(para, "intro")

        assert para._p.get(_PARA_ID_QN) == "intro"
        assert id_of(para) == "intro"

    def it_overwrites_an_existing_id(self, document: DocumentCls):
        document.add_paragraph("alpha")
        para = document.paragraphs[0]
        set_id(para, "first")

        set_id(para, "second")

        assert id_of(para) == "second"

    def it_accepts_a_word_style_eight_hex_id(self, document: DocumentCls):
        document.add_paragraph("alpha")
        set_id(document.paragraphs[0], "A3F12B4C")
        assert id_of(document.paragraphs[0]) == "A3F12B4C"

    def it_accepts_alphanumeric_underscore_up_to_thirty_two_chars(
        self, document: DocumentCls
    ):
        document.add_paragraph("alpha")
        long_id = "section_1_" + "x" * 22  # exactly 32 chars
        assert len(long_id) == 32
        set_id(document.paragraphs[0], long_id)
        assert id_of(document.paragraphs[0]) == long_id

    @pytest.mark.parametrize(
        "bad",
        [
            "",  # empty
            "x" * 33,  # too long
            "has space",
            "has-dash",
            "has.dot",
            "has/slash",
            "has:colon",
            "uñicode",  # non-ASCII
        ],
    )
    def it_rejects_invalid_ids(self, document: DocumentCls, bad: str):
        document.add_paragraph("alpha")
        with pytest.raises(ValueError, match="paraId"):
            set_id(document.paragraphs[0], bad)

    def it_raises_on_non_str_input(self, document: DocumentCls):
        document.add_paragraph("alpha")
        with pytest.raises(ValueError, match="must be a str"):
            set_id(document.paragraphs[0], 12345)  # type: ignore[arg-type]


class DescribeGet:
    def it_finds_a_paragraph_by_its_id(self, document: DocumentCls):
        document.add_paragraph("alpha")
        document.add_paragraph("beta")
        document.add_paragraph("gamma")
        ensure(document)
        # -- pick the middle paragraph deterministically
        target = document.paragraphs[1]
        target_id = id_of(target)

        result = get(document, target_id)

        assert isinstance(result, Paragraph)
        assert result.text == "beta"

    def it_returns_None_when_no_paragraph_carries_the_id(
        self, document: DocumentCls
    ):
        document.add_paragraph("alpha")
        ensure(document)
        assert get(document, "ZZZ99999") is None

    def it_can_find_a_paragraph_inside_a_table_cell(
        self, document_with_table: DocumentCls
    ):
        # -- pick a known cell paragraph and stamp a known id on it
        cell_para = document_with_table.tables[0].rows[1].cells[0].paragraphs[0]
        set_id(cell_para, "cell_paragraph_target")

        result = get(document_with_table, "cell_paragraph_target")

        assert result is not None
        assert result.text == "r1c0"

    def it_validates_the_id_format(self, document: DocumentCls):
        with pytest.raises(ValueError, match="paraId"):
            get(document, "has space")

    def it_returns_the_first_paragraph_when_an_id_collides(
        self, document: DocumentCls
    ):
        document.add_paragraph("first")
        document.add_paragraph("second")
        # -- forcibly create a collision (Word's invariants forbid this
        # -- but we want predictable behaviour either way)
        for p in document.paragraphs:
            p._p.set(_PARA_ID_QN, "DUP12345")

        result = get(document, "DUP12345")

        assert result.text == "first"


class DescribeIterWithIds:
    def it_yields_id_paragraph_pairs_in_document_order(
        self, document: DocumentCls
    ):
        document.add_paragraph("alpha")
        document.add_paragraph("beta")
        document.add_paragraph("gamma")
        ensure(document)

        pairs = list(iter_with_ids(document))

        assert len(pairs) == 3
        texts_in_order = [p.text for _pid, p in pairs]
        assert texts_in_order == ["alpha", "beta", "gamma"]

    def it_skips_paragraphs_without_an_id(self, document: DocumentCls):
        document.add_paragraph("first")
        document.add_paragraph("second")
        # -- only stamp the second paragraph
        set_id(document.paragraphs[1], "tagged")

        pairs = list(iter_with_ids(document))

        assert len(pairs) == 1
        pid, para = pairs[0]
        assert pid == "tagged"
        assert para.text == "second"

    def it_includes_paragraphs_in_table_cells(
        self, document_with_table: DocumentCls
    ):
        ensure(document_with_table)
        ids_seen = [pid for pid, _p in iter_with_ids(document_with_table)]
        # -- 2 body paragraphs + 4 cell paragraphs = 6
        assert len(ids_seen) == 6
        # -- every id is unique (auto-mint collisions are vanishingly
        # -- improbable but worth catching in CI).
        assert len(set(ids_seen)) == len(ids_seen)


class DescribeRoundTrip:
    """IDs must survive save/load through python-docx unchanged."""

    def it_preserves_caller_supplied_ids_through_save_and_load(
        self, document: DocumentCls
    ):
        document.add_paragraph("introduction")
        document.add_paragraph("body")
        document.add_paragraph("conclusion")
        set_id(document.paragraphs[0], "intro")
        set_id(document.paragraphs[1], "body_text")
        set_id(document.paragraphs[2], "conclusion")

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        reloaded = Document(buffer)

        # -- caller-supplied ids must round-trip verbatim
        assert get(reloaded, "intro") is not None
        assert get(reloaded, "intro").text == "introduction"
        assert get(reloaded, "body_text").text == "body"
        assert get(reloaded, "conclusion").text == "conclusion"

    def it_preserves_auto_minted_ids_through_save_and_load(
        self, document: DocumentCls
    ):
        document.add_paragraph("alpha")
        document.add_paragraph("beta")
        ensure(document)
        before = {p.text: id_of(p) for p in document.paragraphs}

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        reloaded = Document(buffer)

        after = {p.text: id_of(p) for p in reloaded.paragraphs}
        assert before == after


class DescribeModuleSurface:
    """The kit module exposes a small, deliberate public API."""

    def it_exports_the_documented_callables_in_dunder_all(self):
        assert sorted(stable_paragraph_ids.__all__) == sorted(
            ["ensure", "get", "id_of", "iter_with_ids", "set_id"]
        )

    def it_is_re_exported_from_the_kit_package(self):
        from docx import kit

        assert kit.stable_paragraph_ids is stable_paragraph_ids
        assert "stable_paragraph_ids" in kit.__all__
