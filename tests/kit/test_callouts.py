"""Unit-test suite for ``docx.kit.callouts`` helpers (issue #287)."""

from __future__ import annotations

from io import BytesIO

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.kit import callouts
from docx.kit.callouts import (
    box,
    caution,
    example,
    important,
    note,
    tip,
    warning,
)
from docx.shared import RGBColor
from docx.table import Table


@pytest.fixture
def document() -> DocumentCls:
    return Document()


# ----- expected fill / icon / title triples (mirror src/docx/kit/callouts.py) ---

_EXPECTED = {
    "note": (RGBColor(0xDE, 0xEB, 0xF7), "\U0001F4DD", "Note"),
    "warning": (RGBColor(0xFF, 0xE6, 0x99), "⚠", "Warning"),
    "caution": (RGBColor(0xF8, 0xCB, 0xAD), "⛔", "Caution"),
    "tip": (RGBColor(0xE2, 0xEF, 0xDA), "\U0001F4A1", "Tip"),
    "important": (RGBColor(0xE4, 0xDF, 0xEC), "❗", "Important"),
    "example": (RGBColor(0xF2, 0xF2, 0xF2), "\U0001F4D8", "Example"),
    "info": (RGBColor(0xDE, 0xF1, 0xF5), "ℹ", "Info"),
}


class DescribeBox:
    """Behavioural tests for :func:`docx.kit.callouts.box`."""

    def it_appends_a_single_cell_one_row_table(self, document: DocumentCls):
        before = len(document.tables)

        result = box(document, "hello", style="note")

        assert isinstance(result, Table)
        assert len(document.tables) == before + 1
        assert len(result.rows) == 1
        assert len(result.columns) == 1

    def it_shades_the_cell_with_the_style_specific_fill(
        self, document: DocumentCls
    ):
        result = box(document, "hello", style="note")
        cell = result.rows[0].cells[0]

        assert cell.shading.fill_color == _EXPECTED["note"][0]

    def it_falls_back_to_info_for_an_unknown_style(self, document: DocumentCls):
        result = box(document, "hello", style="not-a-style")

        assert result.rows[0].cells[0].shading.fill_color == _EXPECTED["info"][0]

    def it_renders_the_icon_in_front_of_a_titleless_body(
        self, document: DocumentCls
    ):
        result = box(document, "hello", style="note", title=None)

        cell = result.rows[0].cells[0]
        assert cell.paragraphs[0].text == "\U0001F4DD hello"

    def it_emits_a_bold_title_paragraph_above_the_body(
        self, document: DocumentCls
    ):
        result = box(document, "hello", style="note", title="Heads up")

        cell = result.rows[0].cells[0]
        # -- two paragraphs: title (bold), body --
        assert len(cell.paragraphs) == 2
        title_para = cell.paragraphs[0]
        assert title_para.text == "\U0001F4DD Heads up"
        assert title_para.runs[0].bold is True
        assert cell.paragraphs[1].text == "hello"

    def it_supports_a_caller_supplied_icon(self, document: DocumentCls):
        result = box(
            document, "hello", style="info", icon="✅", title=None
        )

        assert result.rows[0].cells[0].paragraphs[0].text == "✅ hello"

    def it_supports_an_empty_icon_to_suppress_the_glyph(
        self, document: DocumentCls
    ):
        result = box(document, "hello", style="info", icon="", title=None)

        assert result.rows[0].cells[0].paragraphs[0].text == "hello"

    def it_supports_a_multi_paragraph_body(self, document: DocumentCls):
        result = box(
            document,
            ["First.", "Second.", "Third."],
            style="note",
            title=None,
        )

        cell = result.rows[0].cells[0]
        assert [p.text for p in cell.paragraphs] == [
            "\U0001F4DD First.",
            "Second.",
            "Third.",
        ]

    def it_keeps_the_title_distinct_from_the_body_paragraphs(
        self, document: DocumentCls
    ):
        result = box(
            document,
            ["First.", "Second."],
            style="warning",
            title="Heads up",
        )

        cell = result.rows[0].cells[0]
        assert [p.text for p in cell.paragraphs] == [
            "⚠ Heads up",
            "First.",
            "Second.",
        ]
        # -- Only the title paragraph carries bold --
        assert cell.paragraphs[0].runs[0].bold is True
        assert cell.paragraphs[1].runs[0].bold in (None, False)

    def it_raises_on_an_empty_iterable_body(self, document: DocumentCls):
        with pytest.raises(ValueError, match="non-empty"):
            box(document, [], style="note")


class DescribeConvenienceHelpers:
    """Each named callout dispatches to ``box`` with the right style."""

    @pytest.mark.parametrize(
        "fn,style",
        [
            (note, "note"),
            (warning, "warning"),
            (tip, "tip"),
            (caution, "caution"),
            (important, "important"),
            (example, "example"),
        ],
    )
    def it_uses_the_expected_fill(
        self, document: DocumentCls, fn, style: str
    ):
        result = fn(document, "body")

        fill, _icon, _title = _EXPECTED[style]
        assert result.rows[0].cells[0].shading.fill_color == fill

    @pytest.mark.parametrize(
        "fn,style",
        [
            (note, "note"),
            (warning, "warning"),
            (tip, "tip"),
            (caution, "caution"),
            (important, "important"),
            (example, "example"),
        ],
    )
    def it_emits_the_default_title_with_the_style_icon(
        self, document: DocumentCls, fn, style: str
    ):
        result = fn(document, "body")

        _fill, icon, title = _EXPECTED[style]
        cell = result.rows[0].cells[0]
        assert cell.paragraphs[0].text == "%s %s" % (icon, title)
        assert cell.paragraphs[0].runs[0].bold is True
        assert cell.paragraphs[1].text == "body"

    @pytest.mark.parametrize(
        "fn", [note, warning, tip, caution, important, example]
    )
    def it_accepts_a_list_of_paragraphs_as_body(
        self, document: DocumentCls, fn
    ):
        result = fn(document, ["one", "two"])

        cell = result.rows[0].cells[0]
        # -- title paragraph + two body paragraphs --
        assert len(cell.paragraphs) == 3
        assert [p.text for p in cell.paragraphs[1:]] == ["one", "two"]

    @pytest.mark.parametrize(
        "fn", [note, warning, tip, caution, important, example]
    )
    def it_honours_title_none_to_suppress_the_title_line(
        self, document: DocumentCls, fn
    ):
        result = fn(document, "body", title=None)

        cell = result.rows[0].cells[0]
        # -- one paragraph: icon + body, no separate bold title --
        assert len(cell.paragraphs) == 1
        assert cell.paragraphs[0].runs[0].bold in (None, False)

    @pytest.mark.parametrize(
        "fn", [note, warning, tip, caution, important, example]
    )
    def it_returns_a_table(self, document: DocumentCls, fn):
        assert isinstance(fn(document, "body"), Table)


class DescribeRoundTrip:
    """A document with every callout style serialises and re-loads cleanly."""

    def it_survives_save_and_reload(self, document: DocumentCls):
        note(document, "n")
        warning(document, "w")
        tip(document, "t")
        caution(document, "c")
        important(document, "i")
        example(document, "e")
        box(document, "info", style="info")

        buf = BytesIO()
        document.save(buf)
        buf.seek(0)

        reloaded = Document(buf)
        # -- seven callout tables emitted in declaration order --
        assert len(reloaded.tables) == 7
        for table in reloaded.tables:
            assert len(table.rows) == 1
            assert len(table.columns) == 1


class DescribeKitReExport:
    """The callouts module is re-exported from ``docx.kit``."""

    def it_re_exports_the_callouts_submodule(self):
        from docx import kit

        assert hasattr(kit, "callouts")
        assert kit.callouts is callouts
        assert "callouts" in kit.__all__

    def it_exports_the_seven_public_helpers(self):
        for name in (
            "box",
            "note",
            "warning",
            "tip",
            "caution",
            "important",
            "example",
        ):
            assert name in callouts.__all__
            assert callable(getattr(callouts, name))
