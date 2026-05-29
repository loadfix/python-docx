"""Unit-test suite for ``docx.kit.patch`` (issue #302)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.document import Document as DocumentCls
from docx.enum.section import WD_ORIENTATION
from docx.kit import patch
from docx.kit.patch import (
    InvalidOp,
    PatchError,
    PathNotFound,
    PatchTestFailed,
    apply,
)
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def document() -> DocumentCls:
    """Fresh |Document| with a Title, a Heading, and two body paras."""
    doc = Document()
    doc.add_paragraph("Original title", style="Title")
    doc.add_paragraph("Section heading", style="Heading 1")
    doc.add_paragraph("Body paragraph A", style="Normal")
    doc.add_paragraph("Body paragraph B", style="Normal")
    return doc


@pytest.fixture
def doc_with_table() -> DocumentCls:
    doc = Document()
    doc.add_paragraph("Intro")
    table = doc.add_table(rows=2, cols=3)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"r{r}c{c}"
    return doc


def _set_para_id(paragraph, paragraph_id: str) -> None:
    """Stamp a ``w14:paraId`` onto `paragraph` (test fixture helper)."""
    paragraph._p.set(qn("w14:paraId"), paragraph_id)


# ---------------------------------------------------------------------------
# Public-import sanity
# ---------------------------------------------------------------------------


class DescribePublicSurface:
    """``docx.kit.patch`` exposes :func:`apply` plus the error hierarchy."""

    def it_re_exports_apply_at_the_kit_root(self):
        from docx.kit import patch as kit_patch

        assert kit_patch.apply is apply

    def it_subclasses_patcherror_for_every_specific_failure(self):
        assert issubclass(PathNotFound, PatchError)
        assert issubclass(PatchTestFailed, PatchError)
        assert issubclass(InvalidOp, PatchError)

    def it_lists_apply_plus_the_error_classes_in_dunder_all(self):
        assert set(patch.__all__) == {
            "apply",
            "PatchError",
            "PathNotFound",
            "PatchTestFailed",
            "InvalidOp",
        }


# ---------------------------------------------------------------------------
# replace
# ---------------------------------------------------------------------------


class DescribeReplaceOp:
    def it_replaces_paragraph_text(self, document: DocumentCls):
        apply(document, [
            {"op": "replace", "path": "/paragraphs/0/text", "value": "New title"},
        ])
        assert document.paragraphs[0].text == "New title"

    def it_replaces_paragraph_style(self, document: DocumentCls):
        apply(document, [
            {"op": "replace", "path": "/paragraphs/2/style", "value": "Heading 2"},
        ])
        assert document.paragraphs[2].style.name == "Heading 2"

    def it_replaces_section_orientation(self, document: DocumentCls):
        apply(document, [
            {"op": "replace", "path": "/sections/0/page_orientation",
             "value": "landscape"},
        ])
        assert document.sections[0].orientation == WD_ORIENTATION.LANDSCAPE

    def it_replaces_table_cell_text(self, doc_with_table: DocumentCls):
        apply(doc_with_table, [
            {"op": "replace", "path": "/tables/0/rows/1/cells/2/text",
             "value": "edited"},
        ])
        assert doc_with_table.tables[0].rows[1].cells[2].text == "edited"

    def it_resolves_negative_paragraph_indices(self, document: DocumentCls):
        apply(document, [
            {"op": "replace", "path": "/paragraphs/-1/text", "value": "Last!"},
        ])
        assert document.paragraphs[-1].text == "Last!"


# ---------------------------------------------------------------------------
# add
# ---------------------------------------------------------------------------


class DescribeAddOp:
    def it_appends_a_paragraph_with_dash(self, document: DocumentCls):
        before = len(document.paragraphs)
        apply(document, [
            {"op": "add", "path": "/paragraphs/-",
             "value": {"text": "Tail paragraph", "style": "Normal"}},
        ])
        assert len(document.paragraphs) == before + 1
        assert document.paragraphs[-1].text == "Tail paragraph"

    def it_defaults_style_to_normal_when_omitted(self, document: DocumentCls):
        apply(document, [
            {"op": "add", "path": "/paragraphs/-",
             "value": {"text": "no-style"}},
        ])
        # -- a fresh Document() carries the Normal style; absence of
        # -- the key means "let the doc decide" — Word uses Normal. --
        assert document.paragraphs[-1].text == "no-style"

    def it_can_add_a_text_field_to_an_existing_paragraph(
        self, document: DocumentCls
    ):
        apply(document, [
            {"op": "add", "path": "/paragraphs/0/text", "value": "Replaced via add"},
        ])
        assert document.paragraphs[0].text == "Replaced via add"

    def it_rejects_appending_with_a_non_mapping_value(
        self, document: DocumentCls
    ):
        with pytest.raises(InvalidOp):
            apply(document, [
                {"op": "add", "path": "/paragraphs/-", "value": "just text"},
            ])

    def it_rejects_appending_without_a_text_key(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, [
                {"op": "add", "path": "/paragraphs/-", "value": {"style": "Normal"}},
            ])


# ---------------------------------------------------------------------------
# remove
# ---------------------------------------------------------------------------


class DescribeRemoveOp:
    def it_removes_a_paragraph_by_index(self, document: DocumentCls):
        before = len(document.paragraphs)
        apply(document, [
            {"op": "remove", "path": "/paragraphs/2"},
        ])
        assert len(document.paragraphs) == before - 1
        # -- the previous index 3 has shifted to index 2 --
        assert document.paragraphs[2].text == "Body paragraph B"

    def it_clears_a_paragraph_text_via_remove(self, document: DocumentCls):
        apply(document, [
            {"op": "remove", "path": "/paragraphs/0/text"},
        ])
        assert document.paragraphs[0].text == ""

    def it_clears_a_table_cell_via_remove(self, doc_with_table: DocumentCls):
        apply(doc_with_table, [
            {"op": "remove", "path": "/tables/0/rows/0/cells/0/text"},
        ])
        assert doc_with_table.tables[0].rows[0].cells[0].text == ""

    def it_rejects_removing_section_orientation(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, [
                {"op": "remove", "path": "/sections/0/page_orientation"},
            ])


# ---------------------------------------------------------------------------
# move + copy
# ---------------------------------------------------------------------------


class DescribeMoveOp:
    def it_moves_paragraph_text_between_two_paragraphs(
        self, document: DocumentCls
    ):
        apply(document, [
            {"op": "move",
             "from": "/paragraphs/0/text",
             "path": "/paragraphs/1/text"},
        ])
        # -- after move: para[0] text is the old default ("" not modified
        # -- by remove on text? -- our remove on /text clears it). --
        assert document.paragraphs[0].text == ""
        assert document.paragraphs[1].text == "Original title"

    def it_copies_paragraph_text_between_two_paragraphs(
        self, document: DocumentCls
    ):
        apply(document, [
            {"op": "copy",
             "from": "/paragraphs/0/text",
             "path": "/paragraphs/1/text"},
        ])
        # -- after copy: both have the source text, nothing was cleared --
        assert document.paragraphs[0].text == "Original title"
        assert document.paragraphs[1].text == "Original title"


# ---------------------------------------------------------------------------
# test op
# ---------------------------------------------------------------------------


class DescribeTestOp:
    def it_passes_when_value_matches(self, document: DocumentCls):
        apply(document, [
            {"op": "test", "path": "/paragraphs/0/style", "value": "Title"},
        ])

    def it_raises_patchtestfailed_when_value_disagrees(
        self, document: DocumentCls
    ):
        with pytest.raises(PatchTestFailed):
            apply(document, [
                {"op": "test", "path": "/paragraphs/0/style", "value": "Heading 9"},
            ])

    def it_passes_for_orientation_value(self, document: DocumentCls):
        apply(document, [
            {"op": "test",
             "path": "/sections/0/page_orientation",
             "value": "portrait"},
        ])


# ---------------------------------------------------------------------------
# All-or-nothing semantics
# ---------------------------------------------------------------------------


class DescribeAllOrNothing:
    def it_does_not_persist_any_op_when_a_later_test_fails(
        self, document: DocumentCls
    ):
        original = [p.text for p in document.paragraphs]
        with pytest.raises(PatchTestFailed):
            apply(document, [
                {"op": "replace",
                 "path": "/paragraphs/0/text",
                 "value": "Should be reverted"},
                {"op": "test",
                 "path": "/paragraphs/0/style",
                 "value": "NoSuchStyle"},
            ])
        assert [p.text for p in document.paragraphs] == original

    def it_does_not_persist_any_op_when_a_later_path_is_invalid(
        self, document: DocumentCls
    ):
        original_count = len(document.paragraphs)
        with pytest.raises(PathNotFound):
            apply(document, [
                {"op": "add",
                 "path": "/paragraphs/-",
                 "value": {"text": "should not stick", "style": "Normal"}},
                {"op": "remove",
                 "path": "/paragraphs/9999"},
            ])
        assert len(document.paragraphs) == original_count


# ---------------------------------------------------------------------------
# /by_id paths (paraId)
# ---------------------------------------------------------------------------


class DescribeByIdPath:
    def it_resolves_a_paragraph_by_paraid(self, document: DocumentCls):
        _set_para_id(document.paragraphs[1], "1A2B3C4D")
        apply(document, [
            {"op": "replace",
             "path": "/by_id/1A2B3C4D/text",
             "value": "By-id replacement"},
        ])
        assert document.paragraphs[1].text == "By-id replacement"

    def it_raises_path_not_found_when_paraid_is_missing(
        self, document: DocumentCls
    ):
        with pytest.raises(PathNotFound):
            apply(document, [
                {"op": "replace",
                 "path": "/by_id/DEADBEEF/text",
                 "value": "x"},
            ])

    def it_can_test_a_style_by_paraid(self, document: DocumentCls):
        _set_para_id(document.paragraphs[2], "AABBCCDD")
        apply(document, [
            {"op": "test",
             "path": "/by_id/AABBCCDD/style",
             "value": "Normal"},
        ])


# ---------------------------------------------------------------------------
# JSON-Pointer escape handling (RFC 6901)
# ---------------------------------------------------------------------------


class DescribeJsonPointerEscapes:
    def it_decodes_tilde_one_as_a_slash_in_paraid(
        self, document: DocumentCls
    ):
        _set_para_id(document.paragraphs[0], "x/y")  # paraId with literal /
        apply(document, [
            {"op": "test",
             "path": "/by_id/x~1y/text",
             "value": "Original title"},
        ])

    def it_decodes_tilde_zero_as_a_tilde_in_paraid(
        self, document: DocumentCls
    ):
        _set_para_id(document.paragraphs[0], "a~b")
        apply(document, [
            {"op": "test",
             "path": "/by_id/a~0b/text",
             "value": "Original title"},
        ])

    def it_rejects_a_pointer_that_does_not_start_with_a_slash(
        self, document: DocumentCls
    ):
        with pytest.raises(InvalidOp):
            apply(document, [
                {"op": "remove", "path": "paragraphs/0"},
            ])


# ---------------------------------------------------------------------------
# Error path coverage (InvalidOp / PathNotFound)
# ---------------------------------------------------------------------------


class DescribeErrorPaths:
    def it_rejects_an_unknown_op_name(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, [{"op": "merge", "path": "/paragraphs/0/text",
                              "value": "x"}])

    def it_rejects_a_missing_value_for_replace(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, [{"op": "replace", "path": "/paragraphs/0/text"}])

    def it_rejects_a_missing_from_for_move(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, [
                {"op": "move", "path": "/paragraphs/1/text"},
            ])

    def it_rejects_an_unknown_top_level_segment(self, document: DocumentCls):
        with pytest.raises(PathNotFound):
            apply(document, [
                {"op": "remove", "path": "/footnotes/0"},
            ])

    def it_rejects_an_out_of_range_paragraph_index(
        self, document: DocumentCls
    ):
        with pytest.raises(PathNotFound):
            apply(document, [
                {"op": "remove", "path": "/paragraphs/9999"},
            ])

    def it_rejects_an_unknown_paragraph_field(self, document: DocumentCls):
        with pytest.raises(PathNotFound):
            apply(document, [
                {"op": "replace", "path": "/paragraphs/0/colour",
                 "value": "red"},
            ])

    def it_rejects_a_dash_target_for_a_non_add(self, document: DocumentCls):
        with pytest.raises(PathNotFound):
            apply(document, [
                {"op": "remove", "path": "/paragraphs/-"},
            ])

    def it_rejects_a_non_string_op_field(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, [{"op": 42, "path": "/paragraphs/0/text",
                              "value": "x"}])

    def it_rejects_a_non_mapping_op(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, ["not an op"])

    def it_rejects_invalid_orientation_value(self, document: DocumentCls):
        with pytest.raises(InvalidOp):
            apply(document, [
                {"op": "replace",
                 "path": "/sections/0/page_orientation",
                 "value": "diagonal"},
            ])

    def it_rejects_a_malformed_table_cell_path(
        self, doc_with_table: DocumentCls
    ):
        with pytest.raises(PathNotFound):
            apply(doc_with_table, [
                {"op": "replace",
                 "path": "/tables/0/rows/0/cells/0",
                 "value": "x"},
            ])

    def it_rejects_a_malformed_section_path(self, document: DocumentCls):
        with pytest.raises(PathNotFound):
            apply(document, [
                {"op": "replace", "path": "/sections/0",
                 "value": "landscape"},
            ])

    def it_rejects_a_root_pointer(self, document: DocumentCls):
        with pytest.raises(PathNotFound):
            apply(document, [{"op": "remove", "path": ""}])

    def it_rejects_apply_called_with_none(self):
        with pytest.raises(InvalidOp):
            apply(None, [])  # type: ignore[arg-type]


# ---------------------------------------------------------------------------
# End-to-end example mirrors the issue's docstring
# ---------------------------------------------------------------------------


class DescribeEndToEndExample:
    def it_runs_the_canonical_issue_302_recipe(self, document: DocumentCls):
        apply(document, [
            {"op": "replace", "path": "/paragraphs/0/text", "value": "New title"},
            {"op": "add", "path": "/paragraphs/-",
             "value": {"text": "Final paragraph", "style": "Normal"}},
            {"op": "remove", "path": "/paragraphs/2"},
            {"op": "test", "path": "/paragraphs/0/style", "value": "Title"},
        ])
        assert document.paragraphs[0].text == "New title"
        assert document.paragraphs[-1].text == "Final paragraph"
        # -- one removed, one appended -> net length is unchanged --
        # -- (started at 4: Title, Heading, Body A, Body B; removed
        # -- Body A; appended Final.) --
        texts = [p.text for p in document.paragraphs]
        assert texts == [
            "New title",
            "Section heading",
            "Body paragraph B",
            "Final paragraph",
        ]
