# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.endnotes` module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.endnotes import Endnote, Endnotes
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.oxml.endnotes import CT_Endnote, CT_Endnotes
from docx.oxml.ns import qn
from docx.oxml.text.run import CT_R
from docx.package import Package
from docx.parts.endnotes import EndnotesPart
from docx.text.run import Run

from .unitutil.cxml import element
from .unitutil.mock import FixtureRequest, Mock, instance_mock


class DescribeEndnotes:
    """Unit-test suite for `docx.endnotes.Endnotes` objects."""

    @pytest.mark.parametrize(
        ("cxml", "count"),
        [
            # -- empty endnotes (only separators) --
            (
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator})",
                0,
            ),
            # -- one user endnote --
            (
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator}"
                ",w:endnote{w:id=2})",
                1,
            ),
            # -- two user endnotes --
            (
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator}"
                ",w:endnote{w:id=2},w:endnote{w:id=3})",
                2,
            ),
        ],
    )
    def it_knows_how_many_endnotes_it_contains(self, cxml: str, count: int, package_: Mock):
        endnotes_elm = cast(CT_Endnotes, element(cxml))
        endnotes_part = EndnotesPart(
            PackURI("/word/endnotes.xml"), CT.WML_ENDNOTES, endnotes_elm, package_
        )
        endnotes = Endnotes(endnotes_elm, endnotes_part)

        assert len(endnotes) == count

    def it_is_iterable_over_user_endnotes(self, package_: Mock):
        endnotes_elm = cast(
            CT_Endnotes,
            element(
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator}"
                ",w:endnote{w:id=2},w:endnote{w:id=3})"
            ),
        )
        endnotes_part = EndnotesPart(
            PackURI("/word/endnotes.xml"), CT.WML_ENDNOTES, endnotes_elm, package_
        )
        endnotes = Endnotes(endnotes_elm, endnotes_part)

        endnote_iter = iter(endnotes)

        en1 = next(endnote_iter)
        assert type(en1) is Endnote
        assert en1.endnote_id == 2
        en2 = next(endnote_iter)
        assert type(en2) is Endnote
        assert en2.endnote_id == 3
        with pytest.raises(StopIteration):
            next(endnote_iter)

    def it_can_add_an_endnote(self, package_: Mock):
        endnotes_elm = cast(
            CT_Endnotes,
            element(
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator})"
            ),
        )
        endnotes_part = EndnotesPart(
            PackURI("/word/endnotes.xml"), CT.WML_ENDNOTES, endnotes_elm, package_
        )
        endnotes = Endnotes(endnotes_elm, endnotes_part)

        # -- create a run to anchor the endnote reference --
        para_elm = element("w:p/w:r")
        r_elm = cast(CT_R, para_elm[0])
        run = Run(r_elm, endnotes_part)

        endnote = endnotes.add(run)

        # -- an Endnote is returned --
        assert isinstance(endnote, Endnote)
        assert endnote.endnote_id == 2
        # -- the endnote part is linked --
        assert endnote.part is endnotes_part
        # -- the endnote has a single paragraph with EndnoteText style --
        assert len(endnote.paragraphs) == 1
        assert endnote.paragraphs[0]._p.style == "EndnoteText"
        # -- an endnoteReference was inserted into the run --
        ref_elms = r_elm.xpath("./w:endnoteReference")
        assert len(ref_elms) == 1
        assert ref_elms[0].get(qn("w:id")) == "2"
        # -- the run has EndnoteReference character style --
        assert r_elm.style == "EndnoteReference"

    def it_can_add_an_endnote_with_text(self, package_: Mock):
        endnotes_elm = cast(
            CT_Endnotes,
            element(
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator})"
            ),
        )
        endnotes_part = EndnotesPart(
            PackURI("/word/endnotes.xml"), CT.WML_ENDNOTES, endnotes_elm, package_
        )
        endnotes = Endnotes(endnotes_elm, endnotes_part)

        para_elm = element("w:p/w:r")
        r_elm = cast(CT_R, para_elm[0])
        run = Run(r_elm, endnotes_part)

        endnote = endnotes.add(run, text="This is an endnote.")

        # -- the first paragraph has the text after the endnote ref run --
        first_para = endnote.paragraphs[0]
        assert len(first_para._p.r_lst) == 2
        assert first_para._p.r_lst[1].text == "This is an endnote."

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def package_(self, request: FixtureRequest):
        return instance_mock(request, Package)


class DescribeEndnote:
    """Unit-test suite for `docx.endnotes.Endnote`."""

    def it_knows_its_endnote_id(self, endnotes_part_: Mock):
        endnote_elm = cast(CT_Endnote, element("w:endnote{w:id=42}"))
        endnote = Endnote(endnote_elm, endnotes_part_)

        assert endnote.endnote_id == 42

    def it_provides_access_to_the_paragraphs_it_contains(self, endnotes_part_: Mock):
        endnote_elm = cast(
            CT_Endnote,
            element('w:endnote{w:id=2}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")'),
        )
        endnote = Endnote(endnote_elm, endnotes_part_)

        paragraphs = endnote.paragraphs

        assert len(paragraphs) == 2
        assert [para.text for para in paragraphs] == ["First para", "Second para"]

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:endnote{w:id=2}", ""),
            ('w:endnote{w:id=2}/w:p/w:r/w:t"Endnote text."', "Endnote text."),
            (
                'w:endnote{w:id=2}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")',
                "First para\nSecond para",
            ),
            (
                'w:endnote{w:id=2}/(w:p/w:r/w:t"First para",w:p,w:p/w:r/w:t"Second para")',
                "First para\n\nSecond para",
            ),
        ],
    )
    def it_can_summarize_its_content_as_text(
        self, cxml: str, expected_value: str, endnotes_part_: Mock
    ):
        assert Endnote(cast(CT_Endnote, element(cxml)), endnotes_part_).text == expected_value

    def it_can_clear_its_content(self, endnotes_part_: Mock):
        endnote_elm = cast(
            CT_Endnote,
            element('w:endnote{w:id=2}/(w:p/w:r/w:t"First",w:p/w:r/w:t"Second")'),
        )
        endnote = Endnote(endnote_elm, endnotes_part_)
        assert len(endnote.paragraphs) == 2

        result = endnote.clear()

        assert result is endnote
        assert len(endnote.paragraphs) == 1
        p = endnote.paragraphs[0]
        assert p.text == ""
        assert p._p.style == "EndnoteText"
        # -- the paragraph retains the endnoteRef run for the auto-number mark --
        assert len(p._p.r_lst) == 1
        assert p._p.r_lst[0].style == "EndnoteReference"
        assert p._p.r_lst[0][-1].tag == qn("w:endnoteRef")

    def it_can_delete_itself(self):
        # -- build an endnotes element with a user endnote (id=2) --
        endnotes_elm = cast(
            CT_Endnotes,
            element(
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator}"
                ',w:endnote{w:id=2}/w:p/w:r/w:t"Endnote text")'
            ),
        )

        # -- build a document element containing the endnoteReference --
        doc_elm = element("w:document/w:body/w:p/w:r/w:endnoteReference{w:id=2}")
        document_part_ = Mock()
        document_part_.element = doc_elm
        endnotes_part_ = Mock()
        endnotes_part_.part = endnotes_part_
        endnotes_part_._document_part = document_part_

        endnote_elm = endnotes_elm.endnote_lst[2]
        endnote = Endnote(endnote_elm, endnotes_part_)

        endnote.delete()

        # -- the endnote element is removed from the endnotes part --
        assert len(endnotes_elm.endnote_lst) == 2
        assert all(en.type is not None for en in endnotes_elm.endnote_lst)
        # -- the endnoteReference run is removed from the document body --
        refs = doc_elm.xpath(".//w:endnoteReference")
        assert len(refs) == 0

    def it_removes_the_ref_run_when_deleting_if_run_becomes_empty(self):
        endnotes_elm = cast(
            CT_Endnotes,
            element(
                "w:endnotes/(w:endnote{w:id=0,w:type=separator}"
                ",w:endnote{w:id=1,w:type=continuationSeparator}"
                ",w:endnote{w:id=2}/w:p)"
            ),
        )

        # -- the run has rPr + endnoteReference; after removing ref, only rPr remains --
        doc_elm = element(
            "w:document/w:body/w:p/w:r/(w:rPr/w:rStyle{w:val=EndnoteReference}"
            ",w:endnoteReference{w:id=2})"
        )
        document_part_ = Mock()
        document_part_.element = doc_elm
        endnotes_part_ = Mock()
        endnotes_part_.part = endnotes_part_
        endnotes_part_._document_part = document_part_

        endnote_elm = endnotes_elm.endnote_lst[2]
        endnote = Endnote(endnote_elm, endnotes_part_)

        endnote.delete()

        # -- the entire run is removed since it only had rPr left --
        runs = doc_elm.xpath(".//w:r")
        assert len(runs) == 0

    def it_can_add_a_paragraph(self, endnotes_part_: Mock):
        endnote_elm = cast(CT_Endnote, element("w:endnote{w:id=2}/w:p"))
        endnote = Endnote(endnote_elm, endnotes_part_)

        paragraph = endnote.add_paragraph("New paragraph text")

        assert len(endnote.paragraphs) == 2
        assert endnote.paragraphs[1].text == "New paragraph text"
        # -- default style is EndnoteText --
        assert paragraph._p.style == "EndnoteText"

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def endnotes_part_(self, request: FixtureRequest):
        return instance_mock(request, EndnotesPart)
