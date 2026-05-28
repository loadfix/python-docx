"""Unit test suite for the ``compatibility=`` save flag (issue #94).

Covers :mod:`docx.compatibility`:

* :data:`COMPATIBILITY_LEVELS` — the canonical label → ``compatibilityMode``
  integer mapping.
* :func:`resolve_compatibility` — argument normalisation.
* :func:`apply_compatibility` — package-level mutation.
* :meth:`docx.document.Document.save(compatibility=...)` — end-to-end
  round-trip: every level produces a saved package whose
  ``settings.xml`` carries the right ``compatibilityMode`` value.

The end-to-end tests open a fresh in-memory ``BytesIO`` so the tests
do not depend on any pre-existing fixture file or filesystem state.
"""

from __future__ import annotations

import io

import pytest

from docx import Document as OpenDocument
from docx.compatibility import (
    COMPATIBILITY_LEVELS,
    apply_compatibility,
    resolve_compatibility,
)
from docx.opc.constants import RELATIONSHIP_TYPE as RT


class DescribeCompatibilityLevels:
    """The canonical ``label → compatibilityMode integer`` mapping."""

    def it_maps_every_supported_word_version(self):
        # -- Word 2003 / 2007 / 2010 / 2013 / 2016 are the five
        # -- File→Save-As compatibility-dropdown labels Word still
        # -- offers as of Word 365 / 2024.
        assert COMPATIBILITY_LEVELS == {
            "Word 2003": 11,
            "Word 2007": 12,
            "Word 2010": 14,
            "Word 2013": 15,
            "Word 2016": 16,
        }


class DescribeResolveCompatibility:
    """`resolve_compatibility` normalises caller input to an integer."""

    def it_returns_None_for_None(self):
        assert resolve_compatibility(None) is None

    @pytest.mark.parametrize(
        ("label", "expected"),
        [
            ("Word 2003", 11),
            ("Word 2007", 12),
            ("Word 2010", 14),
            ("Word 2013", 15),
            ("Word 2016", 16),
        ],
    )
    def it_resolves_each_canonical_label(self, label: str, expected: int):
        assert resolve_compatibility(label) == expected

    @pytest.mark.parametrize("level", [11, 12, 14, 15, 16])
    def it_passes_through_a_valid_int(self, level: int):
        assert resolve_compatibility(level) == level

    @pytest.mark.parametrize(
        "bad_value",
        [
            "Word 95",            # not in the mapping
            "word 2003",          # case-sensitive
            "Word 2003 ",         # trailing whitespace not normalised
            13,                   # 13 was the never-used Office 2008 spot
            17,                   # post-2016 levels are reserved
            0,
            -1,
        ],
    )
    def it_raises_ValueError_for_unsupported_input(self, bad_value):
        with pytest.raises(ValueError):
            resolve_compatibility(bad_value)

    @pytest.mark.parametrize("bad_value", [True, False])
    def it_rejects_bools_to_avoid_silent_int_coercion(self, bad_value: bool):
        # -- ``True``/``False`` are int subclasses; without an explicit
        # -- guard the typo ``compatibility=True`` would resolve to the
        # -- (invalid) value 1 and then raise inside the int branch.
        # -- Reject up front for a clearer error message.
        with pytest.raises(ValueError):
            resolve_compatibility(bad_value)

    def it_raises_TypeError_for_unsupported_types(self):
        with pytest.raises(TypeError):
            resolve_compatibility(12.0)  # type: ignore[arg-type]


class DescribeApplyCompatibility:
    """`apply_compatibility` stamps the setting and filters known features."""

    @pytest.mark.parametrize("level", [11, 12, 14, 15, 16])
    def it_writes_the_compatibilityMode_setting(self, level: int):
        document = OpenDocument()
        apply_compatibility(document, level)
        assert document.settings.compatibility_mode == level

    @pytest.mark.parametrize("bad_level", [0, 13, 17, 100])
    def it_raises_ValueError_for_unsupported_levels(self, bad_level: int):
        document = OpenDocument()
        with pytest.raises(ValueError):
            apply_compatibility(document, bad_level)

    def it_does_not_materialise_a_comments_part_just_to_strip(self):
        # -- A pristine document with no comments should not gain a
        # -- comments part as a side-effect of the strip pass when
        # -- targeting Word 2003. ``_strip_threaded_comments`` looks up
        # -- the part via :meth:`part_related_by` and bails on KeyError.
        document = OpenDocument()
        with pytest.raises(KeyError):
            document._part.part_related_by(RT.COMMENTS)
        apply_compatibility(document, 11)
        with pytest.raises(KeyError):
            document._part.part_related_by(RT.COMMENTS)

    def it_strips_modern_threaded_comment_rels_when_targeting_word_2003(self):
        document = OpenDocument()
        # -- materialise the modern parts so the strip has something
        # -- to remove: touching the lazy properties is enough.
        document.comments_ids
        document.comments_extensible
        comments_part = document._part.part_related_by(RT.COMMENTS)

        # -- pre-condition: both modern rels exist on the comments part
        before = {rel.reltype for rel in comments_part.rels.values()}
        assert RT.COMMENTS_IDS in before
        assert RT.COMMENTS_EXTENSIBLE in before

        apply_compatibility(document, 11)

        after = {rel.reltype for rel in comments_part.rels.values()}
        assert RT.COMMENTS_IDS not in after
        assert RT.COMMENTS_EXTENSIBLE not in after

    def it_strips_modern_threaded_comment_rels_when_targeting_word_2007(self):
        # -- Word 2007 predates the w16 namespaces too.
        document = OpenDocument()
        document.comments_ids
        document.comments_extensible
        comments_part = document._part.part_related_by(RT.COMMENTS)

        apply_compatibility(document, 12)

        after = {rel.reltype for rel in comments_part.rels.values()}
        assert RT.COMMENTS_IDS not in after
        assert RT.COMMENTS_EXTENSIBLE not in after

    @pytest.mark.parametrize("level", [14, 15, 16])
    def it_keeps_modern_threaded_comments_for_word_2010_plus(self, level: int):
        # -- Word 2010+ understands the w16 extensions via mc:Ignorable;
        # -- there is no need to drop them.
        document = OpenDocument()
        document.comments_ids
        document.comments_extensible
        comments_part = document._part.part_related_by(RT.COMMENTS)

        apply_compatibility(document, level)

        after = {rel.reltype for rel in comments_part.rels.values()}
        assert RT.COMMENTS_IDS in after
        assert RT.COMMENTS_EXTENSIBLE in after


class DescribeDocumentSaveCompatibility:
    """End-to-end save / reopen round-trip for every compatibility level."""

    @pytest.mark.parametrize(
        ("label", "expected_val"),
        [
            ("Word 2003", 11),
            ("Word 2007", 12),
            ("Word 2010", 14),
            ("Word 2013", 15),
            ("Word 2016", 16),
        ],
    )
    def it_round_trips_each_level_via_save(self, label: str, expected_val: int):
        document = OpenDocument()
        document.add_paragraph("compat-mode round-trip")

        buf = io.BytesIO()
        document.save(buf, compatibility=label)

        buf.seek(0)
        reopened = OpenDocument(buf)
        assert reopened.settings.compatibility_mode == expected_val

    @pytest.mark.parametrize("raw_int", [11, 12, 14, 15, 16])
    def it_accepts_a_raw_int(self, raw_int: int):
        document = OpenDocument()
        document.add_paragraph("hi")

        buf = io.BytesIO()
        document.save(buf, compatibility=raw_int)

        buf.seek(0)
        reopened = OpenDocument(buf)
        assert reopened.settings.compatibility_mode == raw_int

    def it_does_not_change_the_compatibility_mode_when_compatibility_is_None(
        self,
    ):
        # -- Default behaviour: ``save()`` without ``compatibility=``
        # -- preserves whatever the document already carried. The
        # -- bundled template stamps ``compatibilityMode=14`` (Word
        # -- 2010) on every fresh document; a save with no flag must
        # -- leave that alone, neither downgrading nor clearing it.
        document = OpenDocument()
        document.add_paragraph("default save")
        original_mode = document.settings.compatibility_mode

        buf = io.BytesIO()
        document.save(buf)

        buf.seek(0)
        reopened = OpenDocument(buf)
        assert reopened.settings.compatibility_mode == original_mode

    def it_validates_the_label_before_writing_anything(self):
        document = OpenDocument()
        document.add_paragraph("nope")
        original_mode = document.settings.compatibility_mode
        # -- An invalid label must raise *before* serialisation so the
        # -- caller's destination stream is left untouched.
        buf = io.BytesIO()
        with pytest.raises(ValueError):
            document.save(buf, compatibility="Word 95")
        # -- The settings part must not have been mutated by the
        # -- aborted call either: whatever value was there before is
        # -- still there after.
        assert document.settings.compatibility_mode == original_mode
        # -- ``buf`` should be empty since save aborted before writing.
        assert buf.getvalue() == b""

    def it_strips_threaded_comments_in_the_saved_package_for_word_2003(self):
        # -- Sanity-check the end-to-end strip: after saving with
        # -- ``compatibility="Word 2003"`` the modern threaded-comments
        # -- relationships should be absent in the reopened comments
        # -- part. Materialise both modern parts on the original
        # -- document so there is something concrete to strip.
        document = OpenDocument()
        document.comments_ids
        document.comments_extensible

        buf = io.BytesIO()
        document.save(buf, compatibility="Word 2003")

        buf.seek(0)
        reopened = OpenDocument(buf)
        try:
            comments_part = reopened._part.part_related_by(RT.COMMENTS)
        except KeyError:
            # -- The comments part itself may have been dropped by the
            # -- writer when no actual comments remain; either way the
            # -- threaded extensions are gone, which is the contract.
            return
        rel_types = {rel.reltype for rel in comments_part.rels.values()}
        assert RT.COMMENTS_IDS not in rel_types
        assert RT.COMMENTS_EXTENSIBLE not in rel_types
