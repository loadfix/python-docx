# pyright: reportPrivateUsage=false

"""Unit-test suite for ``docx.kit.chapter`` (issue #87)."""

from __future__ import annotations

import io

import pytest

from docx import Document
from docx.kit import chapter
from docx.oxml.ns import qn
from docx.shared import RGBColor

from .unitutil.file import test_file


class DescribeAddChapterOpener:
    """Behavioural tests for :func:`docx.kit.chapter.add_chapter_opener`."""

    def it_appends_a_section_break_before_the_opener(self):
        document = Document()
        # one default sectPr in the body initially
        body = document.element.body
        assert len(body.findall(qn("w:sectPr"))) + len(
            body.findall(".//" + qn("w:sectPr"))
        ) >= 1

        chapter.add_chapter_opener(document, title="The First Light")

        # a new section break should now be present somewhere in the body
        sect_count = len(body.findall(".//" + qn("w:sectPr"))) + len(
            body.findall(qn("w:sectPr"))
        )
        assert sect_count >= 2

    def it_returns_a_dict_with_each_emitted_paragraph(self):
        document = Document()

        result = chapter.add_chapter_opener(
            document,
            chapter_number="Chapter 1",
            title="The First Light",
            epigraph='"In the beginning..."',
        )

        assert set(result) == {"section", "chapter_number", "title", "epigraph", "image"}
        assert result["chapter_number"] is not None
        assert result["title"] is not None
        assert result["epigraph"] is not None
        assert result["image"] is None

    def it_styles_the_title_paragraph_as_Heading_1(self):
        document = Document()

        result = chapter.add_chapter_opener(document, title="The First Light")

        assert result["title"].style.name == "Heading 1"
        assert result["title"].text == "The First Light"

    def it_centers_the_chapter_number_title_and_epigraph(self):
        document = Document()

        result = chapter.add_chapter_opener(
            document,
            chapter_number="Chapter 1",
            title="The First Light",
            epigraph="quote",
        )

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert result["chapter_number"].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert result["title"].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert result["epigraph"].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_renders_the_epigraph_in_italic(self):
        document = Document()

        result = chapter.add_chapter_opener(
            document, title="The First Light", epigraph="quote"
        )

        assert result["epigraph"].runs[0].italic is True

    def it_resolves_the_named_primary_color(self):
        document = Document()

        result = chapter.add_chapter_opener(
            document, title="The First Light", color="primary"
        )

        rgb = result["title"].runs[0].font.color.rgb
        assert rgb == RGBColor(0x1F, 0x4E, 0x79)

    def it_accepts_a_hex_color_string(self):
        document = Document()

        result = chapter.add_chapter_opener(
            document, title="The First Light", color="ff8800"
        )

        rgb = result["title"].runs[0].font.color.rgb
        assert rgb == RGBColor(0xFF, 0x88, 0x00)

    def it_accepts_an_RGBColor_instance(self):
        document = Document()

        result = chapter.add_chapter_opener(
            document, title="X", color=RGBColor(0x10, 0x20, 0x30)
        )

        rgb = result["title"].runs[0].font.color.rgb
        assert rgb == RGBColor(0x10, 0x20, 0x30)

    def it_rejects_invalid_color_values(self):
        document = Document()

        with pytest.raises(ValueError, match="color"):
            chapter.add_chapter_opener(document, title="X", color=12345)

    def it_skips_chapter_number_when_omitted(self):
        document = Document()

        result = chapter.add_chapter_opener(document, title="X")

        assert result["chapter_number"] is None

    def it_skips_epigraph_when_omitted(self):
        document = Document()

        result = chapter.add_chapter_opener(document, title="X")

        assert result["epigraph"] is None

    def it_can_attach_a_decorative_image(self):
        document = Document()

        result = chapter.add_chapter_opener(
            document, title="X", image=test_file("python-icon.png")
        )

        assert result["image"] is not None
        # the image paragraph should contain an inline drawing
        from docx.oxml.ns import qn

        drawings = result["image"]._element.findall(".//" + qn("w:drawing"))
        assert len(drawings) == 1


class DescribeDropCap:
    """The drop-cap one-shot hook applied to ``Document.add_paragraph``."""

    def it_splits_the_next_paragraph_into_drop_cap_plus_body(self):
        document = Document()

        chapter.add_chapter_opener(document, title="X", drop_cap=True)
        document.add_paragraph("It was a dark and stormy night.")

        # find the framePr-bearing paragraph
        body = document.element.body
        frame_prs = body.findall(".//" + qn("w:framePr"))
        assert len(frame_prs) == 1
        framePr = frame_prs[0]
        assert framePr.get(qn("w:dropCap")) == "drop"
        assert framePr.get(qn("w:lines")) == "3"
        assert framePr.get(qn("w:wrap")) == "around"

        # the drop-cap paragraph should hold exactly one character
        # locate it via its parent ``w:p``
        drop_p = framePr.getparent().getparent()
        # "I" is the leading character of the body text
        text_runs = "".join(t.text or "" for t in drop_p.findall(".//" + qn("w:t")))
        assert text_runs == "I"

    def it_emits_the_remainder_of_text_in_a_following_paragraph(self):
        document = Document()

        chapter.add_chapter_opener(document, title="X", drop_cap=True)
        document.add_paragraph("It was a dark and stormy night.")

        # The combined text of the last two paragraphs should match the original
        last_two = document.paragraphs[-2:]
        combined = "".join(p.text for p in last_two)
        assert combined == "It was a dark and stormy night."

    def it_does_not_apply_a_drop_cap_when_the_flag_is_false(self):
        document = Document()

        chapter.add_chapter_opener(document, title="X", drop_cap=False)
        document.add_paragraph("It was a dark and stormy night.")

        body = document.element.body
        frame_prs = body.findall(".//" + qn("w:framePr"))
        assert len(frame_prs) == 0

    def it_only_fires_the_drop_cap_hook_once(self):
        document = Document()

        chapter.add_chapter_opener(document, title="X", drop_cap=True)
        document.add_paragraph("Alpha is first.")
        document.add_paragraph("Beta is second.")

        body = document.element.body
        frame_prs = body.findall(".//" + qn("w:framePr"))
        # only the first add_paragraph after the opener gets a drop cap
        assert len(frame_prs) == 1

    def it_handles_an_empty_next_paragraph_gracefully(self):
        document = Document()

        chapter.add_chapter_opener(document, title="X", drop_cap=True)
        # Empty body — should not crash, should not emit a drop cap
        document.add_paragraph("")
        document.add_paragraph("Now this is a normal paragraph.")

        body = document.element.body
        frame_prs = body.findall(".//" + qn("w:framePr"))
        assert len(frame_prs) == 0


class DescribeRoundTrip:
    """Save the document, reload, and verify the chapter opener survives."""

    def it_round_trips_a_full_chapter_opener_through_save_and_load(self):
        document = Document()

        chapter.add_chapter_opener(
            document,
            chapter_number="Chapter 1",
            title="The First Light",
            epigraph='"In the beginning..." -- Genesis 1:1',
            drop_cap=True,
            color="primary",
        )
        document.add_paragraph("It was a dark and stormy night.")

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        loaded = Document(buf)

        # Section break preserved
        body = loaded.element.body
        sect_prs = body.findall(".//" + qn("w:sectPr")) + body.findall(qn("w:sectPr"))
        assert len(sect_prs) >= 2

        # Title paragraph survives with Heading 1 style
        titles = [p for p in loaded.paragraphs if p.text == "The First Light"]
        assert len(titles) == 1
        assert titles[0].style.name == "Heading 1"

        # Drop-cap framePr survives
        frame_prs = body.findall(".//" + qn("w:framePr"))
        assert len(frame_prs) == 1
        assert frame_prs[0].get(qn("w:dropCap")) == "drop"
        assert frame_prs[0].get(qn("w:lines")) == "3"

        # Combined body text recovers original
        body_paragraphs = [
            p.text
            for p in loaded.paragraphs
            if p.text and p.text not in ("Chapter 1", "The First Light")
            and not p.text.startswith('"In the beginning')
        ]
        assert "".join(body_paragraphs).endswith("It was a dark and stormy night.")
