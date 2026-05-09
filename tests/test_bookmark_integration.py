# pyright: reportPrivateUsage=false

"""Integration tests for bookmark feature across paragraph and document."""

from __future__ import annotations

from typing import cast

from docx.bookmarks import Bookmark, Bookmarks
from docx.oxml.document import CT_Body, CT_Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element


class DescribeParagraph_add_bookmark:
    """Unit-test suite for `Paragraph.add_bookmark()`."""

    def it_can_add_a_bookmark_wrapping_whole_paragraph(self):
        body = cast(CT_Body, element('w:body/w:p/w:r/w:t"hello"'))
        p_elm = body.p_lst[0]
        para = Paragraph(p_elm, None)  # type: ignore[arg-type]

        bm = para.add_bookmark("test_bm")

        assert isinstance(bm, Bookmark)
        assert bm.name == "test_bm"
        assert bm.bookmark_id == 0
        # -- bookmarkStart is first child (no pPr), bookmarkEnd is last --
        children = list(p_elm)
        assert children[0].tag == qn("w:bookmarkStart")
        assert children[-1].tag == qn("w:bookmarkEnd")

    def it_can_add_a_bookmark_wrapping_whole_paragraph_with_pPr(self):
        body = cast(CT_Body, element('w:body/w:p/(w:pPr,w:r/w:t"hello")'))
        p_elm = body.p_lst[0]
        para = Paragraph(p_elm, None)  # type: ignore[arg-type]

        bm = para.add_bookmark("test_bm")

        assert bm.name == "test_bm"
        children = list(p_elm)
        # -- pPr is first, then bookmarkStart, then run, then bookmarkEnd --
        assert children[0].tag == qn("w:pPr")
        assert children[1].tag == qn("w:bookmarkStart")
        assert children[-1].tag == qn("w:bookmarkEnd")

    def it_can_add_a_bookmark_around_specific_runs(self):
        body = cast(
            CT_Body,
            element('w:body/w:p/(w:r/w:t"aaa",w:r/w:t"bbb",w:r/w:t"ccc")'),
        )
        p_elm = body.p_lst[0]
        para = Paragraph(p_elm, None)  # type: ignore[arg-type]
        runs = para.runs

        bm = para.add_bookmark("mid", start_run=runs[1], end_run=runs[1])

        assert bm.name == "mid"
        # -- bookmarkStart is before the second run, bookmarkEnd is after it --
        children = list(p_elm)
        tags = [c.tag for c in children]
        bs_idx = tags.index(qn("w:bookmarkStart"))
        be_idx = tags.index(qn("w:bookmarkEnd"))
        # bookmarkStart should be right before second w:r
        assert tags[bs_idx + 1] == qn("w:r")
        # bookmarkEnd should be right after that same w:r
        assert be_idx == bs_idx + 2

    def it_allocates_unique_ids(self):
        body = cast(CT_Body, element('w:body/w:p/w:r/w:t"hello"'))
        p_elm = body.p_lst[0]
        para = Paragraph(p_elm, None)  # type: ignore[arg-type]

        bm1 = para.add_bookmark("bm1")
        bm2 = para.add_bookmark("bm2")

        assert bm1.bookmark_id == 0
        assert bm2.bookmark_id == 1

    def it_can_add_a_bookmark_with_only_start_run(self):
        body = cast(
            CT_Body,
            element('w:body/w:p/(w:r/w:t"aaa",w:r/w:t"bbb")'),
        )
        p_elm = body.p_lst[0]
        para = Paragraph(p_elm, None)  # type: ignore[arg-type]
        runs = para.runs

        bm = para.add_bookmark("single", start_run=runs[0])

        assert bm.name == "single"
        assert bm.bookmark_id == 0


class DescribeDocument_bookmarks:
    """Unit-test suite for `Document.bookmarks`."""

    def it_provides_access_to_document_bookmarks(self):
        from docx.document import Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/"
                "(w:bookmarkStart{w:id=0,w:name=bm1},w:bookmarkEnd{w:id=0})"
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        bookmarks = doc.bookmarks

        assert isinstance(bookmarks, Bookmarks)
        assert len(bookmarks) == 1
        bm = next(iter(bookmarks))
        assert bm.name == "bm1"


class DescribeDocument_add_bookmark:
    """Unit-test suite for `Document.add_bookmark(runs, name)`."""

    def it_adds_a_bookmark_spanning_a_single_run(self):
        from docx.document import Document

        doc_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        run = doc.paragraphs[0].runs[0]

        bm = doc.add_bookmark(run, "single")

        assert isinstance(bm, Bookmark)
        assert bm.name == "single"
        assert bm.bookmark_id == 0
        body = doc_elm.body
        assert len(body.xpath(".//w:bookmarkStart")) == 1
        assert len(body.xpath(".//w:bookmarkEnd")) == 1

    def it_adds_a_bookmark_spanning_runs_across_paragraphs(self):
        from docx.document import Document

        doc_elm = cast(
            CT_Document,
            element(
                'w:document/w:body/(w:p/w:r/w:t"aaa",w:p/w:r/w:t"bbb")'
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        first_run = doc.paragraphs[0].runs[0]
        last_run = doc.paragraphs[1].runs[0]

        bm = doc.add_bookmark([first_run, last_run], "spanning")

        assert bm.name == "spanning"
        body = doc_elm.body
        # -- bookmarkStart is a sibling of first_run inside its paragraph,
        #    bookmarkEnd is a sibling of last_run inside the second paragraph --
        p1_children = list(body.p_lst[0])
        p2_children = list(body.p_lst[1])
        assert p1_children[0].tag == qn("w:bookmarkStart")
        assert p2_children[-1].tag == qn("w:bookmarkEnd")

    def it_allocates_unique_ids_across_calls(self):
        from docx.document import Document

        doc_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        run = doc.paragraphs[0].runs[0]

        bm1 = doc.add_bookmark(run, "bm1")
        bm2 = doc.add_bookmark(run, "bm2")

        assert bm1.bookmark_id == 0
        assert bm2.bookmark_id == 1

    def it_raises_on_empty_runs_sequence(self):
        import pytest

        from docx.document import Document

        doc_elm = cast(
            CT_Document,
            element("w:document/w:body"),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]

        with pytest.raises(ValueError, match="non-empty"):
            doc.add_bookmark([], "oops")


class DescribeBookmark_name_setter:
    """Unit-test suite for `Bookmark.name` setter."""

    def it_can_rename_a_bookmark(self):
        body = cast(
            CT_Body,
            element(
                "w:body/w:p/(w:bookmarkStart{w:id=0,w:name=old_name}"
                ",w:bookmarkEnd{w:id=0})"
            ),
        )
        bookmarks = Bookmarks(body)
        bm = next(iter(bookmarks))

        bm.name = "new_name"

        assert bm.name == "new_name"
        # -- underlying w:bookmarkStart/@w:name reflects the rename --
        bookmarkStart = body.xpath(".//w:bookmarkStart")[0]
        assert bookmarkStart.get(qn("w:name")) == "new_name"


class DescribeBookmark_roundtrip:
    """End-to-end round-trip: build doc -> save -> reload -> inspect."""

    def it_preserves_bookmarks_across_save_and_reload(self):
        import io

        from docx import Document

        doc = Document()
        p1 = doc.add_paragraph()
        r1a = p1.add_run("Chapter ")
        r1b = p1.add_run("One")
        p2 = doc.add_paragraph()
        r2a = p2.add_run("continues ")
        r2b = p2.add_run("here.")

        # -- single-run bookmark via Bookmarks.add --
        doc.bookmarks.add("bm_word", r1b)
        # -- multi-run bookmark via Document.add_bookmark --
        doc.add_bookmark([r1a, r1b], "bm_heading")
        # -- cross-paragraph bookmark via Bookmarks.add(start_para, end_para) --
        doc.bookmarks.add("bm_span", p1, p2)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        names = {bm.name for bm in reloaded.bookmarks}
        assert names == {"bm_word", "bm_heading", "bm_span"}

        bm_word = reloaded.bookmarks["bm_word"]
        assert bm_word.text == "One"

        bm_heading = reloaded.bookmarks["bm_heading"]
        assert bm_heading.text == "Chapter One"

        bm_span = reloaded.bookmarks["bm_span"]
        assert bm_span.text == "Chapter Onecontinues here."
        assert len(bm_span.paragraphs) == 2

    def it_roundtrips_a_removal(self):
        import io

        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("hello")
        doc.bookmarks.add("keep", r)
        doc.bookmarks.add("drop", r)
        assert len(doc.bookmarks) == 2

        doc.bookmarks.remove("drop")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        names = {bm.name for bm in reloaded.bookmarks}
        assert names == {"keep"}
