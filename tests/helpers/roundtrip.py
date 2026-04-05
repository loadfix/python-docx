"""Round-trip testing helpers for python-docx.

Provides utilities for the write-save-reopen-assert pattern used to verify that
python-docx can correctly round-trip document content.
"""

from __future__ import annotations

import os
import tempfile
from typing import Callable, TypeVar

from docx import Document
from docx.document import Document as DocumentCls

T = TypeVar("T")


def assert_round_trip(
    create_fn: Callable[[DocumentCls], T],
    assert_fn: Callable[[DocumentCls, T], None],
) -> None:
    """Create a document, save it, re-open it, and run assertions.

    `create_fn` receives a blank Document and should populate it with the content
    under test. It may return any value that will be passed to `assert_fn` as
    context (e.g. expected values).

    `assert_fn` receives the re-opened Document and the context value returned by
    `create_fn`, and should assert that the content survived the round trip.

    The temporary file is automatically cleaned up.
    """
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    try:
        # -- create and save --
        doc = Document()
        context = create_fn(doc)
        doc.save(path)

        # -- re-open and assert --
        doc2 = Document(path)
        assert_fn(doc2, context)
    finally:
        os.unlink(path)


def save_and_reopen(doc: DocumentCls) -> DocumentCls:
    """Save a document to a temp file and re-open it, returning the new Document.

    This is a simpler alternative to `assert_round_trip` when you need more control
    over the test flow. The temporary file is cleaned up automatically.
    """
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    try:
        doc.save(path)
        return Document(path)
    finally:
        os.unlink(path)
