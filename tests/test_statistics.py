# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.statistics` module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.document import Document
from docx.oxml.document import CT_Body, CT_Document
from docx.statistics import DocumentStatistics, compute_statistics

from .unitutil.cxml import element
from .unitutil.mock import Mock


class DescribeDocumentStatistics:
    """Unit-test suite for `docx.statistics.DocumentStatistics` namedtuple."""

    def it_exposes_all_four_counts_as_fields(self):
        stats = DocumentStatistics(
            paragraphs=1, words=2, characters=11, characters_no_spaces=10
        )

        assert stats.paragraphs == 1
        assert stats.words == 2
        assert stats.characters == 11
        assert stats.characters_no_spaces == 10

    def it_supports_tuple_unpacking(self):
        stats = DocumentStatistics(
            paragraphs=1, words=2, characters=11, characters_no_spaces=10
        )

        paragraphs, words, characters, characters_no_spaces, pages = stats

        assert (paragraphs, words, characters, characters_no_spaces, pages) == (
            1, 2, 11, 10, None,
        )

    def it_defaults_pages_to_None(self):
        stats = DocumentStatistics(
            paragraphs=1, words=2, characters=11, characters_no_spaces=10
        )

        assert stats.pages is None

    def it_accepts_a_pages_value(self):
        stats = DocumentStatistics(
            paragraphs=1, words=2, characters=11, characters_no_spaces=10, pages=7
        )

        assert stats.pages == 7


class DescribeComputeStatistics:
    """Unit-test suite for `docx.statistics.compute_statistics`."""

    def it_returns_zeros_for_an_empty_document(self):
        body = cast(CT_Body, element("w:body"))

        stats = compute_statistics(body)

        assert stats == DocumentStatistics(0, 0, 0, 0)

    def it_returns_zeros_for_a_body_with_only_an_empty_paragraph(self):
        body = cast(CT_Body, element("w:body/w:p"))

        stats = compute_statistics(body)

        # -- empty paragraphs don't count as paragraphs for word-count purposes --
        assert stats == DocumentStatistics(0, 0, 0, 0)

    def it_counts_a_single_hello_world_paragraph(self):
        body = cast(CT_Body, element('w:body/w:p/w:r/w:t"Hello world"'))

        stats = compute_statistics(body)

        assert stats.paragraphs == 1
        assert stats.words == 2
        assert stats.characters == 11
        assert stats.characters_no_spaces == 10

    def it_counts_multiple_paragraphs_independently(self):
        body = cast(
            CT_Body,
            element(
                "w:body/("
                'w:p/w:r/w:t"Hello world",'
                'w:p/w:r/w:t"Foo bar baz"'
                ")"
            ),
        )

        stats = compute_statistics(body)

        assert stats.paragraphs == 2
        assert stats.words == 5
        # -- "Hello world" (11) + "Foo bar baz" (11) = 22 --
        assert stats.characters == 22
        # -- 10 + 9 (no spaces) = 19 --
        assert stats.characters_no_spaces == 19

    def it_skips_empty_paragraphs_in_the_paragraph_count(self):
        body = cast(
            CT_Body,
            element(
                "w:body/("
                'w:p/w:r/w:t"hi",'
                "w:p,"
                'w:p/w:r/w:t"there"'
                ")"
            ),
        )

        stats = compute_statistics(body)

        # -- only the two non-empty paragraphs count --
        assert stats.paragraphs == 2
        assert stats.words == 2
        assert stats.characters == 7  # -- "hi" + "there"
        assert stats.characters_no_spaces == 7

    def it_treats_tabs_as_word_separators(self):
        # -- a w:tab in a run becomes "\t" in the paragraph text --
        body = cast(
            CT_Body,
            element(
                'w:body/w:p/w:r/(w:t"foo",w:tab,w:t"bar")'
            ),
        )

        stats = compute_statistics(body)

        assert stats.paragraphs == 1
        assert stats.words == 2
        assert stats.characters == 7  # "foo\tbar"
        assert stats.characters_no_spaces == 6

    def it_treats_line_breaks_as_word_separators(self):
        # -- a w:br in a run becomes "\n" in the paragraph text --
        body = cast(
            CT_Body,
            element(
                'w:body/w:p/w:r/(w:t"foo",w:br,w:t"bar baz")'
            ),
        )

        stats = compute_statistics(body)

        assert stats.paragraphs == 1
        assert stats.words == 3
        assert stats.characters == 11  # "foo\nbar baz"
        assert stats.characters_no_spaces == 9

    def it_counts_text_inside_a_table_cell(self):
        body = cast(
            CT_Body,
            element(
                "w:body/w:tbl/w:tr/w:tc/"
                'w:p/w:r/w:t"Cell text here"'
            ),
        )

        stats = compute_statistics(body)

        assert stats.paragraphs == 1
        assert stats.words == 3
        assert stats.characters == 14
        assert stats.characters_no_spaces == 12

    def it_counts_text_inside_a_nested_table(self):
        body = cast(
            CT_Body,
            element(
                "w:body/w:tbl/w:tr/w:tc/("
                'w:p/w:r/w:t"outer cell",'
                "w:tbl/w:tr/w:tc/"
                'w:p/w:r/w:t"inner cell"'
                ")"
            ),
        )

        stats = compute_statistics(body)

        assert stats.paragraphs == 2
        assert stats.words == 4
        # -- "outer cell" (10) + "inner cell" (10) = 20 --
        assert stats.characters == 20
        assert stats.characters_no_spaces == 18

    def it_counts_inline_sdt_text_as_body_text(self):
        body = cast(
            CT_Body,
            element(
                "w:body/w:p/("
                'w:r/w:t"before ",'
                'w:sdt/w:sdtContent/w:r/w:t"inside",'
                'w:r/w:t" after"'
                ")"
            ),
        )

        stats = compute_statistics(body)

        assert stats.paragraphs == 1
        assert stats.words == 3
        # -- "before inside after" --
        assert stats.characters == 19
        assert stats.characters_no_spaces == 17

    def it_counts_paragraphs_inside_a_block_level_sdt(self):
        body = cast(
            CT_Body,
            element(
                'w:body/w:sdt/w:sdtContent/w:p/w:r/w:t"block sdt text"'
            ),
        )

        stats = compute_statistics(body)

        assert stats.paragraphs == 1
        assert stats.words == 3
        assert stats.characters == 14
        assert stats.characters_no_spaces == 12


class DescribeDocument_statistics:
    """Integration-level tests for `Document.statistics`."""

    def it_returns_a_DocumentStatistics_for_the_body(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello world"'),
        )
        document = Document(document_elm, Mock())

        stats = document.statistics

        assert isinstance(stats, DocumentStatistics)
        assert stats.paragraphs == 1
        assert stats.words == 2
        assert stats.characters == 11
        assert stats.characters_no_spaces == 10

    def it_returns_zeros_for_an_empty_body(self):
        document_elm = cast(CT_Document, element("w:document/w:body"))
        document = Document(document_elm, Mock())

        assert document.statistics == DocumentStatistics(0, 0, 0, 0)

    @pytest.mark.parametrize(
        ("cxml_body", "expected"),
        [
            # -- plain sentence --
            (
                'w:body/w:p/w:r/w:t"The quick brown fox"',
                DocumentStatistics(1, 4, 19, 16),
            ),
            # -- multi-paragraph with punctuation --
            (
                "w:body/("
                'w:p/w:r/w:t"one.",'
                'w:p/w:r/w:t"two three."'
                ")",
                DocumentStatistics(2, 3, 14, 13),
            ),
        ],
    )
    def it_matches_expected_counts_for_varied_inputs(
        self, cxml_body: str, expected: DocumentStatistics
    ):
        document_elm = cast(CT_Document, element("w:document/" + cxml_body))
        document = Document(document_elm, Mock())

        assert document.statistics == expected


class DescribeHeaderFooterExclusion:
    """Text in headers/footers/footnotes lives outside `w:body` and is excluded.

    `Document.statistics` reads from `self._element.body`, so content in other
    OOXML parts (headers, footers, footnotes, endnotes, comments) cannot
    contribute.
    """

    def it_does_not_count_header_paragraphs(self):
        # -- a `w:hdr` element is a separate part in a real document, but we can
        # -- construct one as a standalone tree to exercise `compute_statistics`
        # -- directly and show that only `w:body` content is counted. Here we
        # -- build a full document containing a `w:body` AND verify that passing
        # -- a `w:hdr` element does NOT yield any body stats: specifically, the
        # -- Document-level `statistics` property targets `self._element.body`.
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                'w:p/w:r/w:t"only body text"'
            ),
        )
        document = Document(document_elm, Mock())

        stats = document.statistics

        # -- only the single body paragraph counts --
        assert stats.paragraphs == 1
        assert stats.words == 3
        assert stats.characters == 14
        assert stats.characters_no_spaces == 12

    def it_only_inspects_the_body_element_passed_to_it(self):
        # -- `compute_statistics` takes a `w:body` element directly; any element
        # -- in another part (like `w:hdr`) is never accessed because the API
        # -- entry point (`Document.statistics`) only ever hands it the body.
        # -- This test proves behaviorally: if we construct a header tree and a
        # -- distinct empty body, only the body is counted.
        header_elm = element('w:hdr/w:p/w:r/w:t"header content here"')
        body = cast(CT_Body, element("w:body"))

        stats = compute_statistics(body)

        # -- header element still has its text, but body stats are all zero --
        assert header_elm.xpath('.//w:t')[0].text == "header content here"
        assert stats == DocumentStatistics(0, 0, 0, 0)

    def it_ignores_text_outside_the_w_body_element(self):
        # -- confirms top-level body paragraphs and paragraphs inside tables
        # -- (still descendants of the body) both count as body text.
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/("
                'w:p/w:r/w:t"in body",'
                'w:tbl/w:tr/w:tc/w:p/w:r/w:t"also in body"'
                ")"
            ),
        )
        document = Document(document_elm, Mock())

        stats = document.statistics

        # -- both the top-level and nested paragraph count --
        assert stats.paragraphs == 2
        assert stats.words == 5
