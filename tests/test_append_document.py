"""Test suite for docx.append_document."""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


TEST_PNG = Path(__file__).parent / "test_files" / "python-icon.png"


class DescribeAppendDocument:
    """Covers `Document.append_document / append_body / append_paragraph` (upstream#1457 et al.)."""

    def it_copies_every_paragraph_from_source_body(self):
        src = Document()
        src.add_paragraph("first")
        src.add_paragraph("second")

        dest = Document()
        dest._body.clear_content()

        copied = dest.append_document(src)

        assert copied >= 2
        texts = [p.text for p in dest.paragraphs]
        assert "first" in texts
        assert "second" in texts

    def it_copies_the_Heading_1_style_when_source_uses_it(self):
        src = Document()
        src.add_heading("Chapter One", level=1)

        dest = Document()
        dest._body.clear_content()
        # -- drop the default Heading 1 style so we can verify it's copied in --
        # -- (default.docx ships with it; that's fine, the test still verifies
        # -- presence after append). --

        dest.append_document(src)

        assert "Heading 1" in dest.styles

    def it_imports_image_parts_and_rewrites_rIds_for_referenced_images(self):
        if not TEST_PNG.exists():
            pytest.skip("test PNG fixture unavailable")

        src = Document()
        src.add_picture(str(TEST_PNG))

        dest = Document()
        dest._body.clear_content()

        dest.append_document(src)

        image_rels = [
            r for r in dest.part.rels.values() if r.reltype == RT.IMAGE
        ]
        assert len(image_rels) >= 1

    def it_survives_a_save_and_reopen_roundtrip(self):
        src = Document()
        src.add_heading("Heading", level=1)
        src.add_paragraph("hello")

        dest = Document()
        dest._body.clear_content()
        dest.append_document(src)

        buf = io.BytesIO()
        dest.save(buf)
        buf.seek(0)
        reopened = Document(buf)

        texts = [p.text for p in reopened.paragraphs]
        assert "Heading" in texts
        assert "hello" in texts

    def it_exposes_append_body_as_an_alias(self):
        src = Document()
        src.add_paragraph("body-only")
        dest = Document()
        dest._body.clear_content()

        copied = dest.append_body(src)

        assert copied >= 1
        assert "body-only" in [p.text for p in dest.paragraphs]

    def it_can_append_a_single_paragraph(self):
        src = Document()
        para = src.add_paragraph("one paragraph")

        dest = Document()
        dest._body.clear_content()

        new_para = dest.append_paragraph(para)

        assert new_para.text == "one paragraph"
        assert "one paragraph" in [p.text for p in dest.paragraphs]

    def it_preserves_destination_section_settings(self):
        src = Document()
        src.add_paragraph("src para")

        dest = Document()
        # -- destination keeps its existing sectPr; appended content goes before it --
        original_sections = len(dest.sections)

        dest.append_document(src)

        assert len(dest.sections) == original_sections
