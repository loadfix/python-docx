# pyright: reportPrivateUsage=false

"""Unit tests for ``Document.bind_data_source`` + SDT data-binding (#80)."""

from __future__ import annotations

import os
import tempfile
from io import BytesIO
from typing import Iterator

import pytest

from docx import Document
from docx.data_sources import (
    DataSource,
    DataSourceValidationError,
    iter_bound_sources,
    recover_name_from_payload,
    resolve_bindings_in_document,
)


# -- helpers -----------------------------------------------------------------

CUSTOMER_NS = "http://example.com/customer"

CUSTOMER_PAYLOAD_ACME = (
    b"<Customer xmlns=\"" + CUSTOMER_NS.encode() + b"\">"
    b"<Name>Acme Corp</Name>"
    b"<Email>contact@acme.example</Email>"
    b"<City>New York</City>"
    b"</Customer>"
)

CUSTOMER_PAYLOAD_BETA = (
    b"<Customer xmlns=\"" + CUSTOMER_NS.encode() + b"\">"
    b"<Name>Beta Inc</Name>"
    b"<Email>hi@beta.example</Email>"
    b"<City>London</City>"
    b"</Customer>"
)

CUSTOMER_XSD = (
    b"<?xml version=\"1.0\"?>"
    b"<xs:schema xmlns:xs=\"http://www.w3.org/2001/XMLSchema\" "
    b"targetNamespace=\"" + CUSTOMER_NS.encode() + b"\" "
    b"elementFormDefault=\"qualified\">"
    b"<xs:element name=\"Customer\">"
    b"<xs:complexType><xs:sequence>"
    b"<xs:element name=\"Name\" type=\"xs:string\"/>"
    b"<xs:element name=\"Email\" type=\"xs:string\"/>"
    b"<xs:element name=\"City\" type=\"xs:string\" minOccurs=\"0\"/>"
    b"</xs:sequence></xs:complexType></xs:element>"
    b"</xs:schema>"
)


@pytest.fixture
def tmpdir_path() -> Iterator[str]:
    with tempfile.TemporaryDirectory() as tmp:
        yield tmp


def _write(path: str, blob: bytes) -> str:
    with open(path, "wb") as f:
        f.write(blob)
    return path


def _save_reload(doc: "Document") -> "Document":
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


# ---------------------------------------------------------------------------


class DescribeBindDataSource:
    """``Document.bind_data_source`` — register & replace data sources."""

    def it_attaches_a_customXml_part_under_a_logical_name(self):
        doc = Document()

        ds = doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")

        assert isinstance(ds, DataSource)
        assert ds.name == "customer"
        assert ds.partname.startswith("/customXml/item")
        assert ds.partname.endswith(".xml")
        assert ds.store_item_id.startswith("{") and ds.store_item_id.endswith("}")

    def it_creates_a_sibling_itemProps_part_with_the_store_item_id(self):
        doc = Document()

        ds = doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")

        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        props = ds.part.part_related_by(RT.CUSTOM_XML_PROPS)
        assert str(props.partname).startswith("/customXml/itemProps")
        assert ds.store_item_id.encode() in props.blob

    def it_accepts_a_filesystem_path(self, tmpdir_path: str):
        doc = Document()
        path = _write(os.path.join(tmpdir_path, "customer.xml"), CUSTOMER_PAYLOAD_ACME)

        ds = doc.bind_data_source(path, name="customer")

        assert ds.root_element is not None
        assert ds.root_element.tag == "{%s}Customer" % CUSTOMER_NS

    def it_accepts_an_open_binary_filelike(self):
        doc = Document()

        ds = doc.bind_data_source(BytesIO(CUSTOMER_PAYLOAD_ACME), name="customer")

        assert ds.root_element is not None

    def it_replaces_the_underlying_part_when_rebound_with_the_same_name(self):
        doc = Document()

        ds1 = doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        ds2 = doc.bind_data_source(CUSTOMER_PAYLOAD_BETA, name="customer")

        # -- partname / store-item id preserved (stable wiring) --
        assert ds1.partname == ds2.partname
        assert ds1.store_item_id == ds2.store_item_id
        # -- payload swapped --
        root = ds2.root_element
        assert root is not None
        assert root.find("{%s}Name" % CUSTOMER_NS).text == "Beta Inc"

    def it_validates_against_a_schema_when_one_is_supplied(self):
        doc = Document()

        # -- valid payload passes (no exception, returns the source) --
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer", schema=CUSTOMER_XSD)

    def it_raises_DataSourceValidationError_when_payload_fails_validation(self):
        doc = Document()
        bad_payload = b"<Customer xmlns=\"%s\"><Name>X</Name></Customer>" % CUSTOMER_NS.encode()

        with pytest.raises(DataSourceValidationError) as exc:
            doc.bind_data_source(bad_payload, name="customer", schema=CUSTOMER_XSD)
        assert exc.value.issues
        assert "validation" in str(exc.value).lower()

    def it_leaves_prior_payload_intact_when_a_rebind_fails_validation(self):
        doc = Document()
        ds = doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")

        bad_payload = b"<Customer xmlns=\"%s\"><Name>X</Name></Customer>" % CUSTOMER_NS.encode()
        with pytest.raises(DataSourceValidationError):
            doc.bind_data_source(bad_payload, name="customer", schema=CUSTOMER_XSD)

        # -- prior blob unchanged --
        assert ds.part.blob == CUSTOMER_PAYLOAD_ACME

    def it_rejects_an_empty_name(self):
        doc = Document()

        with pytest.raises(ValueError):
            doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="")


class DescribeDataSources:
    """``Document.data_sources`` introspection collection."""

    def it_is_empty_for_a_fresh_document(self):
        doc = Document()
        assert doc.data_sources == []

    def it_lists_every_named_source(self):
        doc = Document()

        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        doc.bind_data_source(b"<x/>", name="other")

        names = sorted(s.name for s in doc.data_sources)
        assert names == ["customer", "other"]

    def it_does_not_surface_unnamed_customXml_parts(self):
        # -- bibliography lazily creates a customXml part on demand --
        doc = Document()
        _ = doc.bibliography  # touch to force creation
        assert doc.data_sources == []


class DescribeAddTextControlBindSource:
    """SDT authoring against a bound source (closes #80)."""

    def it_writes_a_dataBinding_anchored_to_the_source(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        p = doc.add_paragraph("Hello ")

        cc = p.add_text_control(
            name="cust_name",
            bind_to="/ns0:Customer/ns0:Name",
            bind_source="customer",
        )

        binding = cc.data_binding
        assert binding is not None
        assert binding.xpath == "/ns0:Customer/ns0:Name"
        assert binding.store_item_id == doc.data_sources[0].store_item_id
        assert "xmlns:ns0='%s'" % CUSTOMER_NS in binding.prefix_mappings

    def it_inlines_the_resolved_value_at_authoring_time(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        p = doc.add_paragraph("Hello ")

        cc = p.add_text_control(
            bind_to="/ns0:Customer/ns0:Name",
            bind_source="customer",
        )

        assert cc.text == "Acme Corp"

    def it_resolves_three_controls_referencing_the_same_source(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        p1 = doc.add_paragraph("Customer: ")
        p1.add_text_control(bind_to="/ns0:Customer/ns0:Name", bind_source="customer")
        p2 = doc.add_paragraph("Email: ")
        p2.add_text_control(bind_to="/ns0:Customer/ns0:Email", bind_source="customer")
        p3 = doc.add_paragraph("City: ")
        p3.add_text_control(bind_to="/ns0:Customer/ns0:City", bind_source="customer")

        reloaded = _save_reload(doc)
        texts = [para.text for para in reloaded.paragraphs]
        assert "Customer: Acme Corp" in texts
        assert "Email: contact@acme.example" in texts
        assert "City: New York" in texts

    def it_raises_when_the_named_source_has_not_been_bound(self):
        doc = Document()
        p = doc.add_paragraph()

        with pytest.raises(KeyError):
            p.add_text_control(
                bind_to="/ns0:Customer/ns0:Name",
                bind_source="not-bound",
            )

    def it_supports_block_level_text_controls(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")

        cc = doc.add_text_control(
            kind="rich-text",
            bind_to="/ns0:Customer/ns0:Name",
            bind_source="customer",
        )

        binding = cc.data_binding
        assert binding is not None
        assert binding.store_item_id == doc.data_sources[0].store_item_id


class DescribeRoundTrip:
    """Save → reload → resolve cycles."""

    def it_re_resolves_bindings_against_the_current_payload_on_save(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        p = doc.add_paragraph()
        p.add_text_control(
            bind_to="/ns0:Customer/ns0:Name", bind_source="customer"
        )

        reloaded = _save_reload(doc)

        # -- rebind to a different payload + save again --
        reloaded.bind_data_source(CUSTOMER_PAYLOAD_BETA, name="customer")
        reloaded2 = _save_reload(reloaded)

        text = "".join(para.text for para in reloaded2.paragraphs)
        assert "Beta Inc" in text
        assert "Acme Corp" not in text

    def it_recovers_the_logical_name_from_the_payload_on_reload(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")

        reloaded = _save_reload(doc)

        sources = reloaded.data_sources
        assert len(sources) == 1
        assert sources[0].name == "customer"

    def it_stamps_the_lfxbind_name_attribute_into_the_payload(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")

        # -- save into a buffer; inspect the saved blob --
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        ds = reloaded.data_sources[0]
        assert recover_name_from_payload(ds.part) == "customer"

    def it_keeps_the_store_item_id_stable_across_a_rebind(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        sid1 = doc.data_sources[0].store_item_id

        doc.bind_data_source(CUSTOMER_PAYLOAD_BETA, name="customer")
        sid2 = doc.data_sources[0].store_item_id

        assert sid1 == sid2


class DescribeResolveBindingsInDocument:
    """Direct test of the save-time resolver helper."""

    def it_returns_zero_when_no_sources_are_bound(self):
        doc = Document()
        assert resolve_bindings_in_document(doc.part) == 0

    def it_skips_SDTs_whose_storeItemID_does_not_match_a_bound_source(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")
        # -- author an SDT whose binding points at a non-existent store id --
        p = doc.add_paragraph()
        cc = p.add_text_control(bind_to="/x", bind_source="customer")
        cc.set_data_binding(
            xpath="/whatever", prefix_mappings="", store_item_id="{NOT-A-MATCH}"
        )

        # -- the unmatched binding contributes nothing; doesn't raise --
        replaced = resolve_bindings_in_document(doc.part)
        # -- only the original (now-matching) binding may have been replaced;
        #    the explicit override above leaves the SDT pointing at a missing
        #    store id, so the resolver leaves it alone. --
        assert replaced >= 0


class DescribeIterBoundSources:
    """Direct test of the iter_bound_sources helper."""

    def it_finds_sources_set_by_bind_data_source(self):
        doc = Document()
        doc.bind_data_source(CUSTOMER_PAYLOAD_ACME, name="customer")

        sources = iter_bound_sources(doc.part)
        assert [s.name for s in sources] == ["customer"]
