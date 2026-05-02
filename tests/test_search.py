# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.search` module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.document import Document
from docx.oxml.document import CT_Document
from docx.search import (
    SearchMatch,
    _build_char_map,
    _iter_all_paragraphs,
    replace_in_all_paragraphs,
    replace_in_all_paragraphs_regex,
    replace_in_paragraphs,
    replace_in_paragraphs_regex,
    search_all_paragraphs,
    search_all_paragraphs_regex,
    search_paragraphs,
    search_paragraphs_regex,
)
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element
from .unitutil.mock import Mock

import re


class DescribeSearchMatch:
    """Unit-test suite for `docx.search.SearchMatch` objects."""

    def it_provides_access_to_its_properties(self):
        paragraph_ = Mock(spec=Paragraph)
        match = SearchMatch(
            paragraph=paragraph_,
            paragraph_index=2,
            run_indices=[0, 1],
            start=5,
            end=10,
        )
        assert match.paragraph is paragraph_
        assert match.paragraph_index == 2
        assert match.run_indices == [0, 1]
        assert match.start == 5
        assert match.end == 10

    def it_defaults_location_to_None(self):
        match = SearchMatch(
            paragraph=Mock(spec=Paragraph),
            paragraph_index=0,
            run_indices=[0],
            start=0,
            end=1,
        )
        assert match.location is None

    def it_accepts_an_explicit_location(self):
        match = SearchMatch(
            paragraph=Mock(spec=Paragraph),
            paragraph_index=0,
            run_indices=[0],
            start=0,
            end=1,
            location="footer:section0:primary",
        )
        assert match.location == "footer:section0:primary"


class DescribeSearch:
    """Unit-test suite for `docx.search.search_paragraphs`."""

    def it_finds_text_in_a_single_run(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello world"'),
        )
        doc = Document(document_elm, Mock())
        paragraphs = doc.paragraphs

        matches = search_paragraphs(paragraphs, "world")

        assert len(matches) == 1
        assert matches[0].paragraph_index == 0
        assert matches[0].start == 6
        assert matches[0].end == 11
        assert matches[0].run_indices == [0]

    def it_finds_text_spanning_multiple_runs(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/(w:r/w:t"hel",w:r/w:t"lo world")'),
        )
        doc = Document(document_elm, Mock())
        paragraphs = doc.paragraphs

        matches = search_paragraphs(paragraphs, "hello")

        assert len(matches) == 1
        assert matches[0].run_indices == [0, 1]
        assert matches[0].start == 0
        assert matches[0].end == 5

    def it_finds_multiple_matches_in_one_paragraph(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"foo bar foo"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "foo")

        assert len(matches) == 2
        assert matches[0].start == 0
        assert matches[0].end == 3
        assert matches[1].start == 8
        assert matches[1].end == 11

    def it_finds_matches_across_multiple_paragraphs(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                '(w:p/w:r/w:t"hello"'
                ',w:p/w:r/w:t"world"'
                ',w:p/w:r/w:t"hello again")'
            ),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "hello")

        assert len(matches) == 2
        assert matches[0].paragraph_index == 0
        assert matches[1].paragraph_index == 2

    def it_returns_empty_list_when_no_match(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "xyz")

        assert matches == []

    def it_returns_empty_list_for_empty_search_text(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "")

        assert matches == []

    def it_supports_case_insensitive_search(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello World"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "hello", case_sensitive=False)

        assert len(matches) == 1
        assert matches[0].start == 0
        assert matches[0].end == 5

    def it_supports_case_sensitive_search_by_default(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello World"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "hello")

        assert matches == []

    def it_supports_whole_word_search(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"cat concatenate the cat"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "cat", whole_word=True)

        assert len(matches) == 2
        assert matches[0].start == 0
        assert matches[0].end == 3
        assert matches[1].start == 20
        assert matches[1].end == 23

    def it_handles_paragraph_with_no_runs(self):
        document_elm = cast(
            CT_Document,
            element("w:document/w:body/w:p"),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs(doc.paragraphs, "text")

        assert matches == []

    def it_finds_text_nested_inside_a_hyperlink(self):
        # -- upstream#1370: Find/Replace used to skip runs inside w:hyperlink.
        from docx.oxml.parser import parse_xml

        xml = (
            b'<w:document xmlns:w='
            b'"http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            b' xmlns:r='
            b'"http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            b'<w:body>'
            b'<w:p>'
            b'<w:r><w:t>pre </w:t></w:r>'
            b'<w:hyperlink r:id="rId1">'
            b'<w:r><w:t>target</w:t></w:r>'
            b'</w:hyperlink>'
            b'<w:r><w:t> post</w:t></w:r>'
            b'</w:p>'
            b'</w:body>'
            b'</w:document>'
        )
        doc = Document(cast(CT_Document, parse_xml(xml)), Mock())
        matches = search_paragraphs(doc.paragraphs, "target")
        assert len(matches) == 1
        # -- run-indices now index into all_runs, where the hyperlink run is
        # -- the second element (pre / hyperlink-inner / post). --
        assert matches[0].start == 4
        assert matches[0].end == 10


class DescribeReplace:
    """Unit-test suite for `docx.search.replace_in_paragraphs`."""

    def it_replaces_text_in_a_single_run(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello world"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "world", "there")

        assert count == 1
        assert doc.paragraphs[0].text == "hello there"

    def it_replaces_text_spanning_multiple_runs(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/(w:r/w:t"hel",w:r/w:t"lo world")'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "hello", "hi")

        assert count == 1
        # First run gets the replacement text, second run loses the matched portion.
        assert doc.paragraphs[0].runs[0].text == "hi"
        assert doc.paragraphs[0].runs[1].text == " world"

    def it_replaces_multiple_occurrences(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"foo bar foo"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "foo", "baz")

        assert count == 2
        assert doc.paragraphs[0].text == "baz bar baz"

    def it_replaces_across_multiple_paragraphs(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                '(w:p/w:r/w:t"hello"'
                ',w:p/w:r/w:t"world"'
                ',w:p/w:r/w:t"hello")'
            ),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "hello", "hi")

        assert count == 2
        assert doc.paragraphs[0].text == "hi"
        assert doc.paragraphs[1].text == "world"
        assert doc.paragraphs[2].text == "hi"

    def it_preserves_formatting_of_first_run(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/"
                "(w:r/(w:rPr/w:b,w:t\"hel\")"
                ",w:r/(w:rPr/w:i,w:t\"lo world\"))"
            ),
        )
        doc = Document(document_elm, Mock())

        replace_in_paragraphs(doc.paragraphs, "hello", "hi")

        # First run keeps its bold formatting.
        assert doc.paragraphs[0].runs[0].bold is True
        assert doc.paragraphs[0].runs[0].text == "hi"
        # Second run keeps its italic formatting.
        assert doc.paragraphs[0].runs[1].italic is True
        assert doc.paragraphs[0].runs[1].text == " world"

    def it_handles_replacement_with_longer_text(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hi"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "hi", "hello world")

        assert count == 1
        assert doc.paragraphs[0].text == "hello world"

    def it_handles_replacement_with_empty_text(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello world"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "world", "")

        assert count == 1
        assert doc.paragraphs[0].text == "hello "

    def it_returns_zero_when_no_match(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "xyz", "abc")

        assert count == 0
        assert doc.paragraphs[0].text == "hello"

    def it_returns_zero_for_empty_old_text(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "", "abc")

        assert count == 0

    def it_supports_case_insensitive_replace(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO hello"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(
            doc.paragraphs, "hello", "hi", case_sensitive=False
        )

        assert count == 3
        assert doc.paragraphs[0].text == "hi hi hi"

    def it_supports_whole_word_replace(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"cat concatenate the cat"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(
            doc.paragraphs, "cat", "dog", whole_word=True
        )

        assert count == 2
        assert doc.paragraphs[0].text == "dog concatenate the dog"

    def it_replaces_text_spanning_three_runs(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/"
                '(w:r/w:t"ab",w:r/w:t"cd",w:r/w:t"ef")'
            ),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs(doc.paragraphs, "bcde", "X")

        assert count == 1
        assert doc.paragraphs[0].runs[0].text == "aX"
        assert doc.paragraphs[0].runs[1].text == ""
        assert doc.paragraphs[0].runs[2].text == "f"


class DescribeDocumentSearchAndReplace:
    """Unit-test suite for Document.search() and Document.replace()."""

    def it_exposes_search_on_document(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello world"'),
        )
        doc = Document(document_elm, Mock())

        matches = doc.search("world")

        assert len(matches) == 1
        assert matches[0].start == 6

    def it_exposes_replace_on_document(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello world"'),
        )
        doc = Document(document_elm, Mock())

        count = doc.replace("world", "there")

        assert count == 1
        assert doc.paragraphs[0].text == "hello there"

    def it_passes_options_through_to_search(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO"'),
        )
        doc = Document(document_elm, Mock())

        matches = doc.search("hello", case_sensitive=False)

        assert len(matches) == 2

    def it_passes_options_through_to_replace(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"cat concatenate"'),
        )
        doc = Document(document_elm, Mock())

        count = doc.replace("cat", "dog", whole_word=True)

        assert count == 1
        assert doc.paragraphs[0].text == "dog concatenate"


class DescribeSearchRegex:
    """Unit-test suite for `docx.search.search_paragraphs_regex`."""

    def it_finds_a_simple_regex_match(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"order 12345 shipped"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs_regex(doc.paragraphs, r"\d+")

        assert len(matches) == 1
        assert matches[0].start == 6
        assert matches[0].end == 11
        assert matches[0].run_indices == [0]

    def it_supports_case_insensitive_flag(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO hello"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs_regex(doc.paragraphs, r"hello", re.IGNORECASE)

        assert len(matches) == 3

    def it_accepts_a_compiled_pattern(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO"'),
        )
        doc = Document(document_elm, Mock())
        compiled = re.compile(r"hello", re.IGNORECASE)

        matches = search_paragraphs_regex(doc.paragraphs, compiled)

        assert len(matches) == 2

    def it_finds_multiple_matches_in_one_paragraph(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"a1 b2 c3 d4"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs_regex(doc.paragraphs, r"[a-z]\d")

        assert len(matches) == 4
        assert [m.start for m in matches] == [0, 3, 6, 9]

    def it_finds_matches_across_run_boundaries(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/"
                '(w:r/w:t"foo",w:r/w:t"BAR",w:r/w:t"baz")'
            ),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs_regex(doc.paragraphs, r"ooBARba")

        assert len(matches) == 1
        assert matches[0].start == 1
        assert matches[0].end == 8
        assert matches[0].run_indices == [0, 1, 2]

    def it_returns_correct_match_offsets_across_paragraphs(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/"
                '(w:p/w:r/w:t"hello 42"'
                ',w:p/w:r/w:t"world 99")'
            ),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs_regex(doc.paragraphs, r"\d+")

        assert len(matches) == 2
        assert matches[0].paragraph_index == 0
        assert matches[0].start == 6
        assert matches[0].end == 8
        assert matches[1].paragraph_index == 1
        assert matches[1].start == 6
        assert matches[1].end == 8

    def it_returns_empty_list_when_no_match(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(document_elm, Mock())

        matches = search_paragraphs_regex(doc.paragraphs, r"\d+")

        assert matches == []


class DescribeReplaceRegex:
    """Unit-test suite for `docx.search.replace_in_paragraphs_regex`."""

    def it_replaces_a_simple_regex_match(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"order 12345 shipped"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(doc.paragraphs, r"\d+", "N/A")

        assert count == 1
        assert doc.paragraphs[0].text == "order N/A shipped"

    def it_supports_case_insensitive_flag(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO hello"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(
            doc.paragraphs, r"hello", "hi", re.IGNORECASE
        )

        assert count == 3
        assert doc.paragraphs[0].text == "hi hi hi"

    def it_expands_backreferences(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"foobar"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(doc.paragraphs, r"(foo)bar", r"\1baz")

        assert count == 1
        assert doc.paragraphs[0].text == "foobaz"

    def it_expands_named_backreferences(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Mr. Smith"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(
            doc.paragraphs, r"Mr\. (?P<name>\w+)", r"Dr. \g<name>"
        )

        assert count == 1
        assert doc.paragraphs[0].text == "Dr. Smith"

    def it_accepts_a_compiled_pattern(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO"'),
        )
        doc = Document(document_elm, Mock())
        compiled = re.compile(r"hello", re.IGNORECASE)

        count = replace_in_paragraphs_regex(doc.paragraphs, compiled, "hi")

        assert count == 2
        assert doc.paragraphs[0].text == "hi hi"

    def it_replaces_multiple_matches_in_one_paragraph(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"a1 b2 c3"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(doc.paragraphs, r"[a-z]\d", "X")

        assert count == 3
        assert doc.paragraphs[0].text == "X X X"

    def it_replaces_match_spanning_three_runs(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/"
                '(w:r/w:t"ab",w:r/w:t"cd",w:r/w:t"ef")'
            ),
        )
        doc = Document(document_elm, Mock())

        # Regex matches part of each of the three runs.
        count = replace_in_paragraphs_regex(doc.paragraphs, r"b.de", "X")

        assert count == 1
        assert doc.paragraphs[0].runs[0].text == "aX"
        assert doc.paragraphs[0].runs[1].text == ""
        assert doc.paragraphs[0].runs[2].text == "f"

    def it_preserves_first_run_formatting_on_cross_run_match(self):
        document_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/"
                "(w:r/(w:rPr/w:b,w:t\"hel\")"
                ",w:r/(w:rPr/w:i,w:t\"lo world\"))"
            ),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(doc.paragraphs, r"h\w+o", "hi")

        assert count == 1
        assert doc.paragraphs[0].runs[0].bold is True
        assert doc.paragraphs[0].runs[0].text == "hi"
        assert doc.paragraphs[0].runs[1].italic is True
        assert doc.paragraphs[0].runs[1].text == " world"

    def it_returns_zero_when_no_match(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"hello"'),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(doc.paragraphs, r"\d+", "N/A")

        assert count == 0
        assert doc.paragraphs[0].text == "hello"

    def it_handles_paragraph_with_no_runs(self):
        document_elm = cast(
            CT_Document,
            element("w:document/w:body/w:p"),
        )
        doc = Document(document_elm, Mock())

        count = replace_in_paragraphs_regex(doc.paragraphs, r"\d+", "N/A")

        assert count == 0

    def it_skips_zero_width_matches(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"abc"'),
        )
        doc = Document(document_elm, Mock())

        # ``\b`` is a zero-width assertion. Nothing should be replaced.
        count = replace_in_paragraphs_regex(doc.paragraphs, r"\b", "X")

        assert count == 0
        assert doc.paragraphs[0].text == "abc"


class DescribeDocumentRegex:
    """Unit-test suite for Document.search_regex() and Document.replace_regex()."""

    def it_exposes_search_regex_on_document(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"order 12345 shipped"'),
        )
        doc = Document(document_elm, Mock())

        matches = doc.search_regex(r"\d+")

        assert len(matches) == 1
        assert matches[0].start == 6
        assert matches[0].end == 11

    def it_exposes_replace_regex_on_document(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"foobar"'),
        )
        doc = Document(document_elm, Mock())

        count = doc.replace_regex(r"(foo)bar", r"\1baz")

        assert count == 1
        assert doc.paragraphs[0].text == "foobaz"

    def it_passes_flags_through_to_search_regex(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO"'),
        )
        doc = Document(document_elm, Mock())

        matches = doc.search_regex(r"hello", re.IGNORECASE)

        assert len(matches) == 2

    def it_passes_flags_through_to_replace_regex(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO"'),
        )
        doc = Document(document_elm, Mock())

        count = doc.replace_regex(r"hello", "hi", re.IGNORECASE)

        assert count == 2
        assert doc.paragraphs[0].text == "hi hi"

    def it_accepts_a_compiled_pattern_on_search_regex(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO"'),
        )
        doc = Document(document_elm, Mock())
        compiled = re.compile(r"hello", re.IGNORECASE)

        matches = doc.search_regex(compiled)

        assert len(matches) == 2

    def it_accepts_a_compiled_pattern_on_replace_regex(self):
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"Hello HELLO"'),
        )
        doc = Document(document_elm, Mock())
        compiled = re.compile(r"hello", re.IGNORECASE)

        count = doc.replace_regex(compiled, "hi")

        assert count == 2
        assert doc.paragraphs[0].text == "hi hi"


# ---------------------------------------------------------------------------
# Cross-story search / replace (issue #154)
# ---------------------------------------------------------------------------


def _fixture_document_with_stories():
    """Build a |Document| that has content in every searchable story."""
    import docx as _docx

    doc = _docx.Document()

    # -- body paragraph --
    doc.add_paragraph("needle in body")

    # -- a body table with content in two cells --
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "needle in cell 0,0"
    table.cell(1, 1).text = "needle in cell 1,1"

    # -- primary footer --
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "needle in footer"

    # -- primary header --
    header = doc.sections[0].header
    header.paragraphs[0].text = "needle in header"

    # -- footnote anchored from a run in the body --
    run_fn = doc.add_paragraph("body host").runs[0]
    doc.footnotes.add(run_fn, "needle in footnote")

    # -- endnote anchored from a run in the body --
    run_en = doc.add_paragraph("body host 2").runs[0]
    doc.endnotes.add(run_en, "needle in endnote")

    # -- comment anchored on a run --
    run_cm = doc.add_paragraph("commented host").runs[0]
    doc.add_comment(run_cm, "needle in comment", author="Ben")

    return doc


class DescribeIterAllParagraphs:
    """Unit-test suite for `docx.search._iter_all_paragraphs`."""

    def it_visits_every_story_at_least_once(self):
        doc = _fixture_document_with_stories()

        locations = [loc for _, loc in _iter_all_paragraphs(doc)]

        assert "body" in locations
        assert any(loc.startswith("table:0:row:0:col:0") for loc in locations)
        assert any(loc.startswith("table:0:row:1:col:1") for loc in locations)
        assert "header:section0:primary" in locations
        assert "footer:section0:primary" in locations
        assert any(loc.startswith("footnote:") for loc in locations)
        assert any(loc.startswith("endnote:") for loc in locations)
        assert any(loc.startswith("comment:") for loc in locations)

    def it_skips_linked_headers_and_footers(self):
        import docx as _docx

        doc = _docx.Document()
        # -- leave the primary header/footer linked-to-previous; no definition --
        locations = [loc for _, loc in _iter_all_paragraphs(doc)]

        assert not any(loc.startswith("header:") for loc in locations)
        assert not any(loc.startswith("footer:") for loc in locations)

    def it_does_not_error_on_a_minimal_cxml_document(self):
        """Feeding a Mock-part Document through the iterator must not raise."""
        document_elm = cast(
            CT_Document,
            element('w:document/w:body/w:p/w:r/w:t"only body"'),
        )
        doc = Document(document_elm, Mock())

        pairs = list(_iter_all_paragraphs(doc))

        # -- body is always present --
        assert pairs[0][1] == "body"


class DescribeSearchAllParagraphs:
    """Unit-test suite for `docx.search.search_all_paragraphs`."""

    def it_finds_matches_only_in_body(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("needle body")

        matches = search_all_paragraphs(doc, "needle")

        assert len(matches) == 1
        assert matches[0].location == "body"

    def it_finds_matches_in_a_footer(self):
        import docx as _docx

        doc = _docx.Document()
        doc.sections[0].footer.paragraphs[0].text = "needle footer"

        matches = search_all_paragraphs(doc, "needle")

        assert len(matches) == 1
        assert matches[0].location == "footer:section0:primary"

    def it_finds_matches_in_a_footnote(self):
        import docx as _docx

        doc = _docx.Document()
        run = doc.add_paragraph("host").runs[0]
        doc.footnotes.add(run, "needle in note")

        matches = search_all_paragraphs(doc, "needle")

        assert len(matches) == 1
        assert matches[0].location.startswith("footnote:")

    def it_finds_matches_in_an_endnote(self):
        import docx as _docx

        doc = _docx.Document()
        run = doc.add_paragraph("host").runs[0]
        doc.endnotes.add(run, "needle endnote")

        matches = search_all_paragraphs(doc, "needle")

        assert len(matches) == 1
        assert matches[0].location.startswith("endnote:")

    def it_finds_matches_in_a_comment(self):
        import docx as _docx

        doc = _docx.Document()
        run = doc.add_paragraph("host").runs[0]
        doc.add_comment(run, "needle in comment", author="Ben")

        matches = search_all_paragraphs(doc, "needle")

        assert len(matches) == 1
        assert matches[0].location.startswith("comment:")

    def it_finds_matches_in_a_body_table_cell(self):
        import docx as _docx

        doc = _docx.Document()
        tbl = doc.add_table(rows=1, cols=1)
        tbl.cell(0, 0).text = "needle in cell"

        matches = search_all_paragraphs(doc, "needle")

        assert len(matches) == 1
        assert matches[0].location == "table:0:row:0:col:0"

    def it_aggregates_matches_across_every_story(self):
        doc = _fixture_document_with_stories()

        matches = search_all_paragraphs(doc, "needle")
        locations = [m.location for m in matches]

        # -- one per story we seeded: body, two cells, header, footer, footnote,
        # -- endnote, comment = 8 total --
        assert len(matches) == 8
        assert locations.count("body") == 1
        assert locations.count("header:section0:primary") == 1
        assert locations.count("footer:section0:primary") == 1
        assert sum(1 for loc in locations if loc.startswith("table:")) == 2
        assert sum(1 for loc in locations if loc.startswith("footnote:")) == 1
        assert sum(1 for loc in locations if loc.startswith("endnote:")) == 1
        assert sum(1 for loc in locations if loc.startswith("comment:")) == 1

    def it_returns_empty_list_for_empty_search_text(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("hello")

        assert search_all_paragraphs(doc, "") == []


class DescribeReplaceInAllParagraphs:
    """Unit-test suite for `docx.search.replace_in_all_paragraphs`."""

    def it_updates_text_in_every_story(self):
        doc = _fixture_document_with_stories()

        count = replace_in_all_paragraphs(doc, "needle", "thread")

        # -- one per story as above --
        assert count == 8
        # -- verify no "needle" remains anywhere --
        assert search_all_paragraphs(doc, "needle") == []
        # -- verify "thread" is present everywhere "needle" used to be --
        assert len(search_all_paragraphs(doc, "thread")) == 8

    def it_preserves_stories_that_contain_no_match(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("hello")
        doc.sections[0].footer.paragraphs[0].text = "untouched footer"

        count = replace_in_all_paragraphs(doc, "hello", "hi")

        assert count == 1
        assert doc.paragraphs[0].text == "hi"
        assert doc.sections[0].footer.paragraphs[0].text == "untouched footer"

    def it_returns_zero_for_empty_old_text(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("hello")

        assert replace_in_all_paragraphs(doc, "", "x") == 0


class DescribeSearchAllRegex:
    """Unit-test suite for the `*_regex` cross-story helpers."""

    def it_finds_regex_matches_across_stories(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("body 12")
        doc.sections[0].footer.paragraphs[0].text = "footer 345"

        matches = search_all_paragraphs_regex(doc, r"\d+")

        locs = sorted(m.location for m in matches)
        assert locs == ["body", "footer:section0:primary"]

    def it_replaces_regex_matches_across_stories(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("body 12")
        doc.sections[0].footer.paragraphs[0].text = "footer 345"

        count = replace_in_all_paragraphs_regex(doc, r"\d+", "N")

        assert count == 2
        assert doc.paragraphs[0].text == "body N"
        assert doc.sections[0].footer.paragraphs[0].text == "footer N"


class DescribeDocumentSearchAllAndReplaceAll:
    """Unit-test suite for the Document.*_all convenience methods."""

    def it_exposes_search_all_on_document(self):
        doc = _fixture_document_with_stories()

        matches = doc.search_all("needle")

        assert len(matches) == 8
        # -- every match carries a location --
        assert all(m.location for m in matches)

    def it_exposes_replace_all_on_document(self):
        doc = _fixture_document_with_stories()

        count = doc.replace_all("needle", "thread")

        assert count == 8
        assert doc.search_all("needle") == []

    def it_exposes_search_regex_all_on_document(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("a1 b2")
        doc.sections[0].footer.paragraphs[0].text = "c3"

        matches = doc.search_regex_all(r"\w\d")

        assert len(matches) == 3

    def it_exposes_replace_regex_all_on_document(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("foo bar")
        doc.sections[0].footer.paragraphs[0].text = "foo baz"

        count = doc.replace_regex_all(r"(foo)", r"\1!")

        assert count == 2
        assert doc.paragraphs[0].text == "foo! bar"
        assert doc.sections[0].footer.paragraphs[0].text == "foo! baz"

    def it_preserves_body_only_search_semantics(self):
        """The existing `Document.search()` / `.replace()` must still ignore non-body stories."""
        doc = _fixture_document_with_stories()

        body_only_matches = doc.search("needle")

        # -- only the single body paragraph has "needle"; table cells, header,
        # -- footer, footnote, endnote, and comment content is not visited --
        assert len(body_only_matches) == 1
        assert body_only_matches[0].location is None

    def it_passes_options_through_to_search_all(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("Hello HELLO")

        matches = doc.search_all("hello", case_sensitive=False)

        assert len(matches) == 2

    def it_passes_options_through_to_replace_all(self):
        import docx as _docx

        doc = _docx.Document()
        doc.add_paragraph("cat concatenate cat")

        count = doc.replace_all("cat", "dog", whole_word=True)

        assert count == 2
        assert doc.paragraphs[0].text == "dog concatenate dog"
