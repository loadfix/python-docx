# pyright: reportPrivateUsage=false

"""Unit-test suite for `docx.bibliography` proxy + write-side integration."""

from __future__ import annotations

import io
import zipfile

import pytest

from docx import Document
from docx.bibliography import Bibliography, Source
from docx.oxml.bibliography import new_sources_root
from docx.oxml.ns import qn


class DescribeBibliography:
    """Unit-test suite for the `Bibliography` proxy."""

    def it_iterates_Source_proxies_for_each_b_Source_child(self):
        sources = new_sources_root()
        sources.add_source_from_kwargs("a", title="Book A")
        sources.add_source_from_kwargs("b", title="Book B")
        bib = Bibliography(sources)

        found = list(bib)

        assert [s.tag for s in found] == ["a", "b"]
        assert all(isinstance(s, Source) for s in found)

    def its_len_matches_the_number_of_sources(self):
        sources = new_sources_root()
        sources.add_source_from_kwargs("a")
        sources.add_source_from_kwargs("b")
        bib = Bibliography(sources)

        assert len(bib) == 2

    def it_can_look_up_a_source_by_tag(self):
        sources = new_sources_root()
        sources.add_source_from_kwargs("alpha", title="A")
        sources.add_source_from_kwargs("beta", title="B")
        bib = Bibliography(sources)

        hit = bib.get_by_tag("beta")

        assert hit is not None
        assert hit.tag == "beta"
        assert hit.title == "B"

    def but_it_returns_None_for_an_unknown_tag(self):
        bib = Bibliography(new_sources_root())

        assert bib.get_by_tag("nope") is None

    def it_proxies_selected_style_and_style_name(self):
        bib = Bibliography(new_sources_root())

        bib.selected_style = "/MLA7.XSL"
        bib.style_name = "MLA7"

        assert bib.selected_style == "/MLA7.XSL"
        assert bib.style_name == "MLA7"

    def it_rejects_duplicate_tags(self):
        bib = Bibliography(new_sources_root())
        bib.add_source("dup")

        with pytest.raises(ValueError, match="dup"):
            bib.add_source("dup")


class DescribeDocument_add_citation:
    """Smoke-level integration suite for `Document.add_citation`."""

    def it_creates_a_bibliography_source_reachable_via_bibliography(self):
        doc = Document()

        src = doc.add_citation(
            "smith2020", title="A Book", author="Smith, J.", year=2020
        )

        assert isinstance(src, Source)
        assert doc.bibliography.get_by_tag("smith2020") is not None
        hit = doc.bibliography.get_by_tag("smith2020")
        assert hit is not None
        assert hit.title == "A Book"
        assert hit.year == "2020"

    def it_survives_a_save_reload_roundtrip(self):
        doc = Document()
        doc.add_citation(
            "einstein1905", title="Zur Elektrodynamik", author="Einstein, A.", year=1905
        )
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        sources = list(reloaded.bibliography)

        assert [s.tag for s in sources] == ["einstein1905"]
        assert sources[0].year == "1905"

    def it_emits_a_citation_sdt_for_add_citation_reference(self):
        doc = Document()
        doc.add_citation("smith2020", title="Book")
        p = doc.add_paragraph("See ")
        cc = p.add_citation_reference("smith2020")
        p.add_run(".")

        sdt = cc.element
        # -- must carry <w:citation/> marker --
        sdtPr = sdt.find(qn("w:sdtPr"))
        assert sdtPr is not None
        assert sdtPr.find(qn("w:citation")) is not None
        # -- and a CITATION fieldcode inside sdtContent --
        sdtContent = sdt.find(qn("w:sdtContent"))
        assert sdtContent is not None
        instrs = sdtContent.findall(f".//{qn('w:instrText')}")
        assert len(instrs) == 1
        assert "CITATION" in instrs[0].text
        assert "smith2020" in instrs[0].text

    def it_writes_bibliography_xml_on_save(self):
        doc = Document()
        doc.add_citation(
            "keynes1936",
            title="The General Theory",
            author="Keynes, J.M.",
            year=1936,
        )
        p = doc.add_paragraph()
        p.add_citation_reference("keynes1936")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        with zipfile.ZipFile(buf) as zf:
            names = set(zf.namelist())
            assert "customXml/item1.xml" in names
            assert "customXml/itemProps1.xml" in names
            item1 = zf.read("customXml/item1.xml").decode("utf-8")
            assert "keynes1936" in item1
            assert "The General Theory" in item1

    def it_supports_multiple_sources_in_one_part(self):
        doc = Document()

        doc.add_citation("a", title="A", year=2001)
        doc.add_citation("b", title="B", year=2002)
        doc.add_citation("c", title="C", year=2003)

        tags = [s.tag for s in doc.bibliography]
        assert tags == ["a", "b", "c"]


class DescribeBibliography_source_types:
    """Test suite for the ECMA-376 source-type catalogue."""

    @pytest.mark.parametrize(
        "source_type",
        [
            "Book",
            "JournalArticle",
            "ConferenceProceedings",
            "Report",
            "Misc",
            "InternetSite",
            "Film",
            "SoundRecording",
            "Performance",
            "Art",
            "DocumentFromInternetSite",
            "ElectronicSource",
            "Case",
            "Patent",
            "Interview",
        ],
    )
    def it_accepts_each_ecma_source_type(self, source_type: str):
        from docx.oxml.bibliography import new_sources_root

        bib = Bibliography(new_sources_root())

        src = bib.add_source(f"tag-{source_type}", source_type=source_type)

        assert src.source_type == source_type

    def it_rejects_an_unknown_source_type(self):
        from docx.oxml.bibliography import new_sources_root

        bib = Bibliography(new_sources_root())

        with pytest.raises(ValueError, match="Scroll"):
            bib.add_source("mytag", source_type="Scroll")


class DescribeSource_extra_fields:
    """Test suite for the additional Source read accessors."""

    def it_exposes_publisher_and_city(self):
        doc = Document()

        src = doc.add_citation(
            "smith2020",
            title="A Book",
            author="Smith, J.",
            year=2020,
            publisher="Acme Press",
            city="London",
        )

        assert src.publisher == "Acme Press"
        assert src.city == "London"
        # -- generic accessor for arbitrary fields --
        assert src.field("Title") == "A Book"
        assert src.field("NoSuchField") is None


class DescribeParagraph_add_citation:
    """Test suite for Paragraph.add_citation (plain CITATION complex field)."""

    def it_emits_a_complex_CITATION_field(self):
        doc = Document()
        doc.add_citation("smith2020", title="A Book", year=2020)
        p = doc.add_paragraph("See ")

        field = p.add_citation("smith2020")

        assert field.is_complex
        assert field.type == "CITATION"
        assert "smith2020" in field.instruction

    def it_encodes_pages_prefix_and_suffix_switches(self):
        doc = Document()
        doc.add_citation("smith2020", title="A Book", year=2020)
        p = doc.add_paragraph()

        field = p.add_citation(
            "smith2020", pages="45-48", prefix="cf. ", suffix=", et al."
        )

        assert "\\p" in field.instruction
        assert "45-48" in field.instruction
        assert "\\f" in field.instruction
        assert "cf. " in field.instruction
        assert "\\s" in field.instruction
        assert ", et al." in field.instruction

    def it_survives_save_reload_and_appears_in_document_citations(self):
        doc = Document()
        doc.add_citation("einstein1905", title="Zur Elektrodynamik", year=1905)
        p = doc.add_paragraph("See ")
        p.add_citation("einstein1905", pages="891", prefix="e.g. ")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        citations = reloaded.citations
        assert len(citations) == 1
        assert citations[0].source_tag == "einstein1905"
        assert citations[0].pages == "891"
        assert citations[0].prefix == "e.g. "
        assert citations[0].suffix is None


class DescribeDocument_citations:
    """Test suite for Document.citations — walk every CITATION in the body."""

    def it_returns_an_empty_list_when_no_citations_exist(self):
        doc = Document()

        assert doc.citations == []

    def it_picks_up_sdt_wrapped_and_bare_citations(self):
        doc = Document()
        doc.add_citation("a", title="A")
        doc.add_citation("b", title="B")

        p1 = doc.add_paragraph("See ")
        p1.add_citation_reference("a")
        p2 = doc.add_paragraph("Also ")
        p2.add_citation("b", pages="7")

        citations = doc.citations
        tags = sorted(c.source_tag for c in citations)
        assert tags == ["a", "b"]
        b_cite = next(c for c in citations if c.source_tag == "b")
        assert b_cite.pages == "7"

    def it_skips_non_citation_fields(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_complex_field("PAGE")

        assert doc.citations == []
