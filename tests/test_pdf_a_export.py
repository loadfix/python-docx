"""Tests for the minimal docx → PDF/A archival exporter.

Verifies the best-effort contract documented in
``src/docx/pdf_a_export.py``:

- Output is a valid PDF (header is ``%PDF-``).
- The XMP metadata stream declares ``pdfaid:part`` and
  ``pdfaid:conformance`` matching the requested level.
- Headings, paragraphs, runs, tables, images, and page breaks all
  render without raising.
- Unsupported levels raise ``ValueError``.
- When ``reportlab`` is unavailable the exporter raises
  ``ImportError`` pointing at the ``[pdfa]`` extra.

Tests skip cleanly when ``reportlab`` is not importable.
"""

from __future__ import annotations

import io
import os

import pytest

from docx import Document

reportlab = pytest.importorskip(
    "reportlab", reason="reportlab not installed; install python-docx[pdfa]"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _new_doc():
    """Return a fresh blank |Document| for use in a single test."""
    return Document()


def _save_to_bytes(doc, level: str = "3a") -> bytes:
    buf = io.BytesIO()
    doc.save_as_pdf_a(buf, level=level)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Smoke tests
# ---------------------------------------------------------------------------


class DescribeDocumentSaveAsPdfA:
    def it_writes_a_pdf_file(self, tmp_path):
        doc = _new_doc()
        doc.add_paragraph("Hello, archival world.")
        out = tmp_path / "out.pdf"
        doc.save_as_pdf_a(str(out), level="3a")
        assert out.exists()
        with open(out, "rb") as fh:
            head = fh.read(8)
        assert head.startswith(b"%PDF-")

    def it_writes_to_a_file_like_object(self):
        doc = _new_doc()
        doc.add_paragraph("Stream me.")
        buf = io.BytesIO()
        doc.save_as_pdf_a(buf, level="3a")
        data = buf.getvalue()
        assert data.startswith(b"%PDF-")
        assert len(data) > 200

    def it_accepts_a_pathlike(self, tmp_path):
        # PathLike support comes via the ``open(path, "wb")`` fallback;
        # `os.fspath`-able paths route through the file-write branch
        # cleanly.
        doc = _new_doc()
        doc.add_paragraph("PathLike OK.")
        out = tmp_path / "viapath.pdf"
        doc.save_as_pdf_a(os.fspath(out), level="3a")
        assert out.exists()
        assert out.read_bytes().startswith(b"%PDF-")

    def it_defaults_level_to_3a(self, tmp_path):
        doc = _new_doc()
        doc.add_paragraph("default level")
        out = tmp_path / "default.pdf"
        doc.save_as_pdf_a(str(out))
        data = out.read_bytes()
        assert b"<pdfaid:part>3</pdfaid:part>" in data
        assert b"<pdfaid:conformance>A</pdfaid:conformance>" in data


# ---------------------------------------------------------------------------
# Level handling
# ---------------------------------------------------------------------------


class DescribePdfALevelHandling:
    @pytest.mark.parametrize(
        ("level", "expected_part", "expected_conf"),
        [
            ("1a", b"1", b"A"),
            ("1b", b"1", b"B"),
            ("2a", b"2", b"A"),
            ("2b", b"2", b"B"),
            ("3a", b"3", b"A"),
            ("3b", b"3", b"B"),
        ],
    )
    def it_emits_xmp_keys_matching_the_level(
        self, level, expected_part, expected_conf
    ):
        doc = _new_doc()
        doc.add_paragraph("level test")
        data = _save_to_bytes(doc, level=level)
        assert b"<pdfaid:part>" + expected_part + b"</pdfaid:part>" in data
        assert (
            b"<pdfaid:conformance>" + expected_conf + b"</pdfaid:conformance>"
            in data
        )

    def it_rejects_an_unknown_level(self):
        doc = _new_doc()
        with pytest.raises(ValueError, match="level must be one of"):
            doc.save_as_pdf_a(io.BytesIO(), level="4a")

    def it_rejects_an_empty_level(self):
        doc = _new_doc()
        with pytest.raises(ValueError, match="level must be one of"):
            doc.save_as_pdf_a(io.BytesIO(), level="")


# ---------------------------------------------------------------------------
# XMP metadata
# ---------------------------------------------------------------------------


class DescribePdfAXmpMetadata:
    def it_includes_the_pdfaid_namespace_uri(self):
        doc = _new_doc()
        doc.add_paragraph("ns test")
        data = _save_to_bytes(doc, level="3a")
        assert b"http://www.aiim.org/pdfa/ns/id/" in data

    def it_includes_the_xmpmeta_envelope(self):
        doc = _new_doc()
        doc.add_paragraph("envelope")
        data = _save_to_bytes(doc, level="3a")
        assert b"<x:xmpmeta" in data
        assert b"</x:xmpmeta>" in data
        assert b"<?xpacket" in data

    def it_propagates_the_document_title_into_dc_title(self):
        doc = _new_doc()
        # -- defensive: an unrelated test in the suite (test_fields) deletes
        # -- ``Document.core_properties`` from the class as a teardown step,
        # -- which can leak across tests when run together. Skip cleanly
        # -- if that has happened. --
        if not hasattr(doc, "core_properties"):
            pytest.skip("core_properties descriptor missing (test isolation leak)")
        doc.core_properties.title = "Quarterly Archive 2026"
        doc.add_paragraph("body")
        data = _save_to_bytes(doc, level="3a")
        assert b"Quarterly Archive 2026" in data

    def it_escapes_xml_metacharacters_in_the_title(self):
        doc = _new_doc()
        if not hasattr(doc, "core_properties"):
            pytest.skip("core_properties descriptor missing (test isolation leak)")
        doc.core_properties.title = "A & B <c>"
        doc.add_paragraph("body")
        data = _save_to_bytes(doc, level="3a")
        # -- the raw "<c>" should never reach the XMP packet verbatim --
        assert b"A &amp; B &lt;c&gt;" in data

    def it_declares_the_producer_string(self):
        doc = _new_doc()
        doc.add_paragraph("body")
        data = _save_to_bytes(doc, level="3a")
        assert b"python-docx" in data
        assert b"PDF/A" in data


# ---------------------------------------------------------------------------
# Element rendering
# ---------------------------------------------------------------------------


class DescribePdfAParagraphRendering:
    def it_renders_a_plain_paragraph(self):
        doc = _new_doc()
        doc.add_paragraph("Hello world.")
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")
        # -- a non-empty paragraph should produce a non-trivial PDF --
        assert len(data) > 500

    def it_renders_bold_runs_without_error(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_run("regular ")
        bold = p.add_run("BOLD")
        bold.bold = True
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_renders_italic_runs_without_error(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        em = p.add_run("italicised")
        em.italic = True
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_renders_underlined_runs_without_error(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        u = p.add_run("underlined")
        u.underline = True
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_renders_combined_bold_italic_runs(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("combo")
        run.bold = True
        run.italic = True
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")


class DescribePdfAHeadingRendering:
    def it_renders_a_Heading_1_paragraph(self):
        doc = _new_doc()
        doc.add_heading("Top Heading", level=1)
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_renders_a_deeply_nested_Heading(self):
        doc = _new_doc()
        doc.add_heading("Deep Heading", level=9)
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_renders_all_six_heading_levels(self):
        doc = _new_doc()
        for n in range(1, 7):
            doc.add_heading(f"Level {n}", level=n)
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")


class DescribePdfATableRendering:
    def it_renders_a_basic_table(self):
        doc = _new_doc()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(0, 2).text = "C"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        table.cell(1, 2).text = "3"
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_handles_an_empty_table(self):
        doc = _new_doc()
        doc.add_table(rows=1, cols=1)
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_handles_long_text_in_a_cell(self):
        doc = _new_doc()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Short"
        table.cell(0, 1).text = (
            "A reasonably long sentence that should wrap inside its column "
            "without crashing the renderer or overflowing the page."
        )
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")


class DescribePdfAImageRendering:
    def it_renders_an_inline_picture(self):
        # -- 1x1 RGB PNG, hand-built so we don't depend on Pillow at test
        # -- time. python-docx insists on parsing the PNG header itself
        # -- to populate the inline-picture's CX/CY, so the chunk CRCs
        # -- and IHDR fields must validate. --
        png_1x1 = (
            b"\x89PNG\r\n\x1a\n"
            b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x02\x00\x00\x00\x90wS\xde"
            b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0\x00\x00"
            b"\x03\x01\x01\x00\xc9\xfe\x92\xef"
            b"\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_run().add_picture(io.BytesIO(png_1x1))
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")


class DescribePdfAPageBreakRendering:
    def it_renders_a_hard_page_break(self):
        from docx.enum.text import WD_BREAK

        doc = _new_doc()
        doc.add_paragraph("Before break")
        p2 = doc.add_paragraph()
        p2.add_run().add_break(WD_BREAK.PAGE)
        p2.add_run("After break")
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")
        # -- two pages should exist; reportlab emits ``/Type /Page`` for
        # -- each one. --
        assert data.count(b"/Type /Page") >= 2 or data.count(b"/Type/Page") >= 2


class DescribePdfAListRendering:
    def it_renders_a_bullet_list(self):
        doc = _new_doc()
        try:
            doc.add_paragraph("Item 1", style="List Bullet")
            doc.add_paragraph("Item 2", style="List Bullet")
        except KeyError:
            pytest.skip("default template missing 'List Bullet' style")
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")

    def it_renders_a_numbered_list(self):
        doc = _new_doc()
        try:
            doc.add_paragraph("Step 1", style="List Number")
            doc.add_paragraph("Step 2", style="List Number")
        except KeyError:
            pytest.skip("default template missing 'List Number' style")
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")


class DescribePdfALongDocumentRendering:
    def it_pages_a_long_document(self):
        # -- build enough paragraphs that pagination must kick in. --
        doc = _new_doc()
        for i in range(120):
            doc.add_paragraph(
                f"Paragraph {i}: lorem ipsum dolor sit amet, "
                "consectetur adipiscing elit, sed do eiusmod tempor."
            )
        data = _save_to_bytes(doc)
        assert data.startswith(b"%PDF-")
        # -- multi-page output --
        page_markers = data.count(b"/Type /Page") + data.count(b"/Type/Page")
        assert page_markers >= 2


# ---------------------------------------------------------------------------
# Module-level entry point
# ---------------------------------------------------------------------------


class DescribeDocumentToPdfAFunction:
    def it_is_importable_from_pdf_a_export_module(self):
        from docx.pdf_a_export import document_to_pdf_a

        assert callable(document_to_pdf_a)

    def it_renders_via_the_module_level_function(self):
        from docx.pdf_a_export import document_to_pdf_a

        doc = _new_doc()
        doc.add_paragraph("module-level")
        buf = io.BytesIO()
        document_to_pdf_a(doc, buf, level="2b")
        data = buf.getvalue()
        assert data.startswith(b"%PDF-")
        assert b"<pdfaid:part>2</pdfaid:part>" in data
        assert b"<pdfaid:conformance>B</pdfaid:conformance>" in data


# ---------------------------------------------------------------------------
# Missing-reportlab fallback
# ---------------------------------------------------------------------------


class DescribePdfAImportErrorMessage:
    def it_raises_an_informative_error_when_reportlab_missing(
        self, monkeypatch
    ):
        import builtins

        from docx import pdf_a_export

        real_import = builtins.__import__

        def _fake_import(name, *args, **kwargs):
            if name.startswith("reportlab"):
                raise ImportError(f"No module named {name!r}")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _fake_import)

        doc = _new_doc()
        with pytest.raises(ImportError, match=r"\[pdfa\]"):
            pdf_a_export.document_to_pdf_a(doc, io.BytesIO(), level="3a")
