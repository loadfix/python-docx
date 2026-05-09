"""Unit-test suite for the `docx.glossary` module."""

from __future__ import annotations

from typing import cast

import pytest

from docx.enum.text import WD_BUILDING_BLOCK_GALLERY
from docx.glossary import BuildingBlock, BuildingBlockCategory, Glossary
from docx.oxml.glossary import CT_DocPart, CT_GlossaryDocument

from .unitutil.cxml import element


# -- a compact building block used in a few tests below --
_SAMPLE_BLOCK = (
    "w:docPart/("
    "w:docPartPr/("
    "w:name{w:val=MyBlock},"
    "w:category/(w:name{w:val=General},w:gallery{w:val=quickParts}),"
    "w:description{w:val=sample description},"
    "w:guid{w:val=abc-123-def}"
    "),"
    "w:docPartBody/(w:p,w:p,w:tbl)"
    ")"
)

_SAMPLE_GLOSSARY = (
    "w:glossaryDocument/w:docParts/("
    "w:docPart/(w:docPartPr/w:name{w:val=First}),"
    "w:docPart/(w:docPartPr/w:name{w:val=Second}),"
    "w:docPart/w:docPartPr/w:name{w:val=Third}"
    ")"
)


class DescribeGlossary:
    """Unit-test suite for `docx.glossary.Glossary`."""

    def it_exposes_its_building_blocks(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_SAMPLE_GLOSSARY)))
        blocks = glossary.building_blocks
        assert len(blocks) == 3
        assert all(isinstance(b, BuildingBlock) for b in blocks)
        assert [b.name for b in blocks] == ["First", "Second", "Third"]

    def it_is_iterable_over_building_blocks(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_SAMPLE_GLOSSARY)))
        assert [b.name for b in glossary] == ["First", "Second", "Third"]

    def it_supports_len(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_SAMPLE_GLOSSARY)))
        assert len(glossary) == 3

    def it_returns_zero_len_for_an_empty_docParts(self):
        glossary = Glossary(
            cast(CT_GlossaryDocument, element("w:glossaryDocument/w:docParts"))
        )
        assert len(glossary) == 0
        assert list(glossary) == []

    def it_returns_zero_len_when_docParts_is_absent(self):
        glossary = Glossary(
            cast(CT_GlossaryDocument, element("w:glossaryDocument"))
        )
        assert len(glossary) == 0

    def it_can_look_up_a_building_block_by_name(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_SAMPLE_GLOSSARY)))
        block = glossary["Second"]
        assert isinstance(block, BuildingBlock)
        assert block.name == "Second"

    def it_raises_KeyError_for_an_unknown_name(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_SAMPLE_GLOSSARY)))
        with pytest.raises(KeyError):
            _ = glossary["NoSuchBlock"]


class DescribeBuildingBlock:
    """Unit-test suite for `docx.glossary.BuildingBlock`."""

    def it_exposes_its_name(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        assert block.name == "MyBlock"

    def it_returns_None_for_name_when_docPartPr_is_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        assert block.name is None

    def it_returns_None_for_name_when_w_name_is_absent(self):
        block = BuildingBlock(
            cast(CT_DocPart, element("w:docPart/w:docPartPr"))
        )
        assert block.name is None

    def it_exposes_its_category(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        cat = block.category
        assert isinstance(cat, BuildingBlockCategory)
        assert cat.category_name == "General"
        assert cat.gallery == "quickParts"

    def and_category_returns_a_proxy_with_None_slots_when_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        cat = block.category
        assert isinstance(cat, BuildingBlockCategory)
        assert cat.category_name is None
        assert cat.gallery is None

    def it_exposes_its_description(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        assert block.description == "sample description"

    def it_returns_None_for_description_when_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        assert block.description is None

    def it_exposes_its_guid(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        assert block.guid == "abc-123-def"

    def it_returns_None_for_guid_when_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        assert block.guid is None

    def it_exposes_its_paragraphs(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        paragraphs = block.paragraphs
        assert len(paragraphs) == 2

    def it_returns_empty_paragraphs_when_docPartBody_is_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        assert block.paragraphs == []

    def it_exposes_its_tables(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        tables = block.tables
        assert len(tables) == 1

    def it_returns_empty_tables_when_docPartBody_is_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        assert block.tables == []


class DescribeBuildingBlockCategory:
    """Unit-test suite for `docx.glossary.BuildingBlockCategory`."""

    def it_exposes_its_name_and_gallery(self):
        cat_elm = element(
            "w:category/(w:name{w:val=General},w:gallery{w:val=quickParts})"
        )
        cat = BuildingBlockCategory(cat_elm)  # type: ignore[arg-type]
        assert cat.category_name == "General"
        assert cat.gallery == "quickParts"

    def it_returns_None_when_name_is_absent(self):
        cat_elm = element("w:category/w:gallery{w:val=quickParts}")
        cat = BuildingBlockCategory(cat_elm)  # type: ignore[arg-type]
        assert cat.category_name is None
        assert cat.gallery == "quickParts"

    def it_returns_None_when_gallery_is_absent(self):
        cat_elm = element("w:category/w:name{w:val=General}")
        cat = BuildingBlockCategory(cat_elm)  # type: ignore[arg-type]
        assert cat.category_name == "General"
        assert cat.gallery is None

    def it_returns_None_for_every_slot_when_category_element_is_None(self):
        cat = BuildingBlockCategory(None)
        assert cat.category_name is None
        assert cat.gallery is None

    def it_round_trips_a_known_gallery_through_the_enum(self):
        cat_elm = element("w:category/w:gallery{w:val=quickParts}")
        cat = BuildingBlockCategory(cat_elm)  # type: ignore[arg-type]
        assert cat.gallery_value is WD_BUILDING_BLOCK_GALLERY.QUICK_PARTS

    @pytest.mark.parametrize(
        ("xml_val", "expected"),
        [
            ("coverPg", WD_BUILDING_BLOCK_GALLERY.COVER_PAGES),
            ("hdrs", WD_BUILDING_BLOCK_GALLERY.HEADERS),
            ("ftrs", WD_BUILDING_BLOCK_GALLERY.FOOTERS),
            ("txtBox", WD_BUILDING_BLOCK_GALLERY.TEXT_BOXES),
            ("custom1", WD_BUILDING_BLOCK_GALLERY.CUSTOM_1),
        ],
    )
    def it_maps_common_galleries_to_enum_members(self, xml_val, expected):
        cat_elm = element(f"w:category/w:gallery{{w:val={xml_val}}}")
        cat = BuildingBlockCategory(cat_elm)  # type: ignore[arg-type]
        assert cat.gallery_value is expected

    def it_returns_None_for_gallery_value_when_unknown(self):
        cat_elm = element("w:category/w:gallery{w:val=notARealGallery}")
        cat = BuildingBlockCategory(cat_elm)  # type: ignore[arg-type]
        assert cat.gallery_value is None
        # -- but raw `gallery` still round-trips the literal
        assert cat.gallery == "notARealGallery"

    def it_returns_None_for_gallery_value_when_gallery_is_absent(self):
        cat_elm = element("w:category/w:name{w:val=General}")
        cat = BuildingBlockCategory(cat_elm)  # type: ignore[arg-type]
        assert cat.gallery_value is None

    def it_compares_equal_by_gallery_and_name(self):
        cat_elm_a = element(
            "w:category/(w:name{w:val=General},w:gallery{w:val=quickParts})"
        )
        cat_elm_b = element(
            "w:category/(w:name{w:val=General},w:gallery{w:val=quickParts})"
        )
        a = BuildingBlockCategory(cat_elm_a)  # type: ignore[arg-type]
        b = BuildingBlockCategory(cat_elm_b)  # type: ignore[arg-type]
        assert a == b
        assert hash(a) == hash(b)

    def it_compares_unequal_when_slots_differ(self):
        a_elm = element(
            "w:category/(w:name{w:val=General},w:gallery{w:val=quickParts})"
        )
        b_elm = element(
            "w:category/(w:name{w:val=Other},w:gallery{w:val=quickParts})"
        )
        a = BuildingBlockCategory(a_elm)  # type: ignore[arg-type]
        b = BuildingBlockCategory(b_elm)  # type: ignore[arg-type]
        assert a != b


# -- a mixed-category glossary used for the filter tests below --
_MIXED_GLOSSARY = (
    "w:glossaryDocument/w:docParts/("
    "w:docPart/w:docPartPr/("
    "w:name{w:val=Alpha},"
    "w:category/(w:name{w:val=General},w:gallery{w:val=quickParts})"
    "),"
    "w:docPart/w:docPartPr/("
    "w:name{w:val=Beta},"
    "w:category/(w:name{w:val=General},w:gallery{w:val=quickParts})"
    "),"
    "w:docPart/w:docPartPr/("
    "w:name{w:val=Gamma},"
    "w:category/(w:name{w:val=Built-In},w:gallery{w:val=coverPg})"
    "),"
    "w:docPart/w:docPartPr/("
    "w:name{w:val=Delta},"
    "w:category/(w:name{w:val=Built-In},w:gallery{w:val=hdrs})"
    "),"
    "w:docPart/w:docPartPr/w:name{w:val=Epsilon}"
    ")"
)


class DescribeGlossaryFiltering:
    """Unit-test suite for filter/aggregation methods on `Glossary`."""

    def it_filters_by_gallery_enum(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        blocks = glossary.by_category(
            gallery=WD_BUILDING_BLOCK_GALLERY.QUICK_PARTS
        )
        assert [b.name for b in blocks] == ["Alpha", "Beta"]

    def it_filters_by_gallery_xml_string(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        blocks = glossary.by_category(gallery="coverPg")
        assert [b.name for b in blocks] == ["Gamma"]

    def it_filters_by_category_name(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        blocks = glossary.by_category(category_name="Built-In")
        assert [b.name for b in blocks] == ["Gamma", "Delta"]

    def it_intersects_gallery_and_category_name_filters(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        blocks = glossary.by_category(
            gallery=WD_BUILDING_BLOCK_GALLERY.QUICK_PARTS,
            category_name="General",
        )
        assert [b.name for b in blocks] == ["Alpha", "Beta"]

    def and_an_empty_intersection_returns_an_empty_list(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        blocks = glossary.by_category(
            gallery=WD_BUILDING_BLOCK_GALLERY.QUICK_PARTS,
            category_name="Built-In",
        )
        assert blocks == []

    def it_returns_all_blocks_when_called_with_no_args(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        assert len(glossary.by_category()) == 5

    def it_dedupes_categories(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        cats = glossary.categories
        # -- Alpha + Beta share (General, quickParts); Gamma + Delta have
        # -- distinct categories; Epsilon has no category and is dropped.
        assert len(cats) == 3
        keys = [(c.gallery, c.category_name) for c in cats]
        assert keys == [
            ("quickParts", "General"),
            ("coverPg", "Built-In"),
            ("hdrs", "Built-In"),
        ]

    def it_dedupes_galleries(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element(_MIXED_GLOSSARY)))
        assert glossary.galleries == ["quickParts", "coverPg", "hdrs"]

    def it_returns_empty_aggregates_for_an_empty_glossary(self):
        glossary = Glossary(
            cast(CT_GlossaryDocument, element("w:glossaryDocument"))
        )
        assert glossary.categories == []
        assert glossary.galleries == []
        assert glossary.by_category(category_name="General") == []


# -- a sample with types/behaviors for the new accessors --
_BLOCK_WITH_TYPES_AND_BEHAVIORS = (
    "w:docPart/w:docPartPr/("
    "w:name{w:val=TypedBlock},"
    "w:types/(w:type{w:val=autoTxt},w:type{w:val=toolbar}),"
    "w:behaviors/(w:behavior{w:val=content},w:behavior{w:val=p})"
    ")"
)


class DescribeBuildingBlockTypeAndBehaviors:
    """Unit-test suite for the type/uuid/behaviors accessors on `BuildingBlock`."""

    def it_exposes_its_types(self):
        block = BuildingBlock(
            cast(CT_DocPart, element(_BLOCK_WITH_TYPES_AND_BEHAVIORS))
        )
        assert block.types == ["autoTxt", "toolbar"]

    def it_returns_first_type_as_type_property(self):
        block = BuildingBlock(
            cast(CT_DocPart, element(_BLOCK_WITH_TYPES_AND_BEHAVIORS))
        )
        assert block.type == "autoTxt"

    def it_returns_None_for_type_when_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        assert block.type is None
        assert block.types == []

    def it_exposes_its_behaviors(self):
        block = BuildingBlock(
            cast(CT_DocPart, element(_BLOCK_WITH_TYPES_AND_BEHAVIORS))
        )
        assert block.behaviors == ["content", "p"]

    def it_returns_empty_behaviors_when_absent(self):
        block = BuildingBlock(cast(CT_DocPart, element("w:docPart")))
        assert block.behaviors == []

    def it_aliases_guid_as_uuid(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        assert block.uuid == block.guid == "abc-123-def"

    def it_aliases_paragraphs_as_content_paragraphs(self):
        block = BuildingBlock(cast(CT_DocPart, element(_SAMPLE_BLOCK)))
        assert block.content_paragraphs == block.paragraphs


class DescribeGlossaryWrite:
    """Unit-test suite for mutating methods on `Glossary`."""

    def it_can_add_a_building_block_with_string_content(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        bb = glossary.add_building_block(
            "Snippet", category="Custom", content="Hello, glossary!"
        )
        assert isinstance(bb, BuildingBlock)
        assert bb.name == "Snippet"
        assert bb.category.category_name == "Custom"
        assert bb.category.gallery == "quickParts"  # enum default
        assert bb.guid is not None
        assert bb.guid.startswith("{") and bb.guid.endswith("}")
        assert [p.text for p in bb.paragraphs] == ["Hello, glossary!"]
        assert len(glossary) == 1

    def it_assigns_fresh_guids_to_each_block(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        a = glossary.add_building_block("A", content="x")
        b = glossary.add_building_block("B", content="y")
        assert a.guid != b.guid

    def it_accepts_a_raw_gallery_string(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        bb = glossary.add_building_block("CP", gallery="coverPg", content="Cover")
        assert bb.category.gallery == "coverPg"
        assert bb.category.gallery_value is WD_BUILDING_BLOCK_GALLERY.COVER_PAGES

    def it_accepts_a_gallery_enum(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        bb = glossary.add_building_block(
            "Hdr", gallery=WD_BUILDING_BLOCK_GALLERY.HEADERS, content="H"
        )
        assert bb.category.gallery == "hdrs"

    def it_accepts_None_content_for_an_empty_body(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        bb = glossary.add_building_block("Empty", content=None)
        assert bb.paragraphs == []

    def it_can_remove_a_building_block_by_name(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        glossary.add_building_block("Keep", content="a")
        glossary.add_building_block("Drop", content="b")
        assert glossary.remove_building_block("Drop") is True
        assert [bb.name for bb in glossary] == ["Keep"]

    def it_returns_False_when_removing_an_unknown_name(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        glossary.add_building_block("Only", content="x")
        assert glossary.remove_building_block("Missing") is False
        assert len(glossary) == 1

    def it_returns_False_when_glossary_has_no_docParts(self):
        glossary = Glossary(cast(CT_GlossaryDocument, element("w:glossaryDocument")))
        assert glossary.remove_building_block("Anything") is False

    def it_round_trips_a_building_block_through_save_and_reload(self, tmp_path):
        from docx import Document

        doc = Document()
        g = doc.ensure_glossary()
        g.add_building_block(
            "Round",
            category="MyCat",
            gallery=WD_BUILDING_BLOCK_GALLERY.QUICK_PARTS,
            content="Persisted",
        )
        out = tmp_path / "round.docx"
        doc.save(str(out))
        doc2 = Document(str(out))
        g2 = doc2.glossary
        assert g2 is not None
        assert len(g2) == 1
        bb = g2["Round"]
        assert bb.category.category_name == "MyCat"
        assert bb.category.gallery == "quickParts"
        assert [p.text for p in bb.paragraphs] == ["Persisted"]

    def it_lazy_creates_the_glossary_part_on_ensure_glossary(self):
        from docx import Document

        doc = Document()
        assert doc.glossary is None
        g = doc.ensure_glossary()
        assert g is not None
        # -- a subsequent read-only access now sees the same glossary --
        assert doc.glossary is not None
        assert len(doc.glossary) == 0
