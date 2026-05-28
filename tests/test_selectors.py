"""Unit-test suite for the CSS-selector query language (issue #78).

The selector engine lives in :mod:`docx.selectors`; the public surface
is :meth:`Document.select` / :meth:`Document.select_one`. These tests
build a fixture document end-to-end, then exercise twelve representative
selectors covering every supported combinator and pseudo-class.

The fixture document layout:

* H1 ``"Intro"`` (Heading 1) followed by an intro paragraph and a body
  paragraph carrying a bold run.
* H2 ``"Methods"`` (Heading 2) followed by an intro paragraph.
* H1 ``"Results"`` (Heading 1) followed by an intro paragraph that
  contains a hyperlink and a bookmark.
* A 3x3 ``"Light List"`` table.
* Final body paragraph with a comment anchored on it.
"""

from __future__ import annotations

import pytest

from docx import Document as OpenDocument
from docx.selectors import (
    SelectorSyntaxError,
    compile_selector,
    select,
    select_one,
)
from docx.bookmarks import Bookmark
from docx.comments import Comment
from docx.table import Table, _Cell, _Row
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run


# ---------------------------------------------------------------------------
# Fixture document
# ---------------------------------------------------------------------------


@pytest.fixture(name="doc")
def _doc():
    """Return a populated |Document| matching the layout above."""
    document = OpenDocument()

    # -- H1 "Intro" + supporting paragraphs --
    document.add_heading("Intro", level=1)
    document.add_paragraph("Welcome to the report.")
    body_p = document.add_paragraph("This sentence has a ")
    bold_run = body_p.add_run("bold")
    bold_run.bold = True
    body_p.add_run(" word.")

    # -- H2 "Methods" --
    document.add_heading("Methods", level=2)
    document.add_paragraph("We measured stuff.")

    # -- H1 "Results" + paragraph with hyperlink + bookmark --
    document.add_heading("Results", level=1)
    results_p = document.add_paragraph("See ")
    # add_hyperlink applies the "Hyperlink" character style by default
    # which the brand-new Document fixture doesn't have, so opt out.
    results_p.add_hyperlink(
        "https://example.com/report", text="the website", style=None
    )
    results_p.add_run(".")
    document.add_bookmark(results_p.runs[0], name="results_anchor")

    # -- Light List table (3x3) --
    table_styles = [s.name for s in document.styles]
    style_name = "Light List" if "Light List" in table_styles else None
    table = document.add_table(rows=3, cols=3, style=style_name)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"r{i}c{j}"

    # -- Final body paragraph with a comment --
    final_p = document.add_paragraph("Trailing line.")
    document.add_comment(final_p.runs[0], text="LGTM", author="Alice")

    return document


# ---------------------------------------------------------------------------
# Behaviour tests — each selector
# ---------------------------------------------------------------------------


class DescribeDocumentSelect:
    """End-to-end coverage for :meth:`Document.select`."""

    # 1. exact attribute match on style
    def it_selects_paragraphs_with_an_exact_style_match(self, doc):
        results = doc.select('p[style="Heading 1"]')
        assert len(results) == 2
        assert all(isinstance(p, Paragraph) for p in results)
        assert [p.text for p in results] == ["Intro", "Results"]

    # 2. starts-with attribute selector
    def it_selects_paragraphs_with_a_style_prefix(self, doc):
        results = doc.select('p[style^="Heading "]')
        # Two H1s + one H2 = three matches
        assert [p.text for p in results] == ["Intro", "Methods", "Results"]

    # 3. descendant combinator + boolean attribute
    def it_finds_bold_runs_anywhere_below_a_paragraph(self, doc):
        results = doc.select("p r[bold]")
        assert len(results) == 1
        assert isinstance(results[0], Run)
        assert results[0].text == "bold"

    # 4. compound: tag + attribute + sub-selector
    def it_returns_an_empty_list_when_nothing_matches(self, doc):
        assert doc.select('p[style="Code Block"]') == []

    # 5. table style selector
    def it_selects_tables_by_style_name(self, doc):
        # Only assert when the fixture had the style available; on a
        # bare Document the "Light List" style is not pre-registered.
        if any(s.name == "Light List" for s in doc.styles):
            results = doc.select('tbl[style="Light List"]')
            assert len(results) == 1
            assert isinstance(results[0], Table)
        # Always: a bare ``tbl`` selector returns the one table in body.
        all_tables = doc.select("tbl")
        assert len(all_tables) == 1
        assert isinstance(all_tables[0], Table)

    # 6. nested combinator: descendant + nth-child pseudo
    def it_selects_cells_in_the_second_column_of_every_table(self, doc):
        cells = doc.select("tbl tr td:nth-child(2)")
        assert len(cells) == 3  # one per row
        assert [c.text for c in cells] == ["r0c1", "r1c1", "r2c1"]
        assert all(isinstance(c, _Cell) for c in cells)

    # 7. child combinator
    def it_selects_only_direct_row_children_of_a_table(self, doc):
        rows = doc.select("tbl > tr")
        assert len(rows) == 3
        assert all(isinstance(r, _Row) for r in rows)

    # 8. adjacent-sibling combinator
    def it_selects_the_paragraph_immediately_after_each_heading(self, doc):
        intros = doc.select('p[style^="Heading "] + p')
        # The "Results" heading is followed by the hyperlink paragraph
        # which has plain text "See .". Two intros total: one after the
        # H1 "Intro" and one after the H1 "Results"; the H2 is also a
        # heading so its successor is included.
        assert [p.text.startswith(("Welcome", "We measured", "See"))
                for p in intros] == [True] * len(intros)
        assert len(intros) == 3

    # 9. :first-child / :last-child
    def it_selects_first_and_last_children_of_a_scope(self, doc):
        first_rows = doc.select("tbl tr:first-child")
        last_rows = doc.select("tbl tr:last-child")
        assert len(first_rows) == 1
        assert len(last_rows) == 1
        assert [c.text for c in first_rows[0].cells] == ["r0c0", "r0c1", "r0c2"]
        assert [c.text for c in last_rows[0].cells] == ["r2c0", "r2c1", "r2c2"]

    # 10. :not(...) negation
    def it_negates_a_simple_selector_via_not(self, doc):
        # Body paragraphs that aren't a heading.
        non_headings = doc.select('p:not([style^="Heading "])')
        texts = [p.text for p in non_headings]
        # At minimum the welcome / methods-detail / final lines are in.
        assert "Welcome to the report." in texts
        assert "We measured stuff." in texts
        assert "Trailing line." in texts
        # And no heading text leaked in.
        assert "Intro" not in texts
        assert "Methods" not in texts
        assert "Results" not in texts

    # 11. hyperlink selector
    def it_selects_external_hyperlinks_by_address_prefix(self, doc):
        links = doc.select('hyperlink[address^="https://"]')
        assert len(links) == 1
        assert isinstance(links[0], Hyperlink)
        assert links[0].address == "https://example.com/report"

    # 12. bookmark + comment selectors
    def it_selects_bookmarks_and_comments_by_name_or_author(self, doc):
        bms = doc.select("bookmark[name=results_anchor]")
        assert len(bms) == 1
        assert isinstance(bms[0], Bookmark)
        comments = doc.select("comment[author=Alice]")
        assert len(comments) == 1
        assert isinstance(comments[0], Comment)


class DescribeDocumentSelectOne:
    """``select_one`` is a convenience over ``select``."""

    def it_returns_the_first_match(self, doc):
        first = doc.select_one('p[style^="Heading "]')
        assert isinstance(first, Paragraph)
        assert first.text == "Intro"

    def it_returns_None_when_nothing_matches(self, doc):
        assert doc.select_one('p[style="No Such Style"]') is None


class DescribeSelectorParser:
    """Direct coverage for :func:`compile_selector` error paths."""

    def it_rejects_an_empty_selector(self):
        with pytest.raises(SelectorSyntaxError):
            compile_selector("")

    def it_rejects_an_unknown_element_type(self):
        with pytest.raises(SelectorSyntaxError):
            compile_selector("foo")

    def it_rejects_an_unsupported_pseudo(self):
        with pytest.raises(SelectorSyntaxError):
            compile_selector("p:hover")

    def it_rejects_a_malformed_attribute_clause(self):
        with pytest.raises(SelectorSyntaxError):
            compile_selector("p[style")

    def it_accepts_the_module_level_select_helper(self, doc):
        compiled = compile_selector('p[style="Heading 1"]')
        results = select(doc, compiled)
        assert all(isinstance(p, Paragraph) for p in results)
        assert select_one(doc, compiled).text == "Intro"

    def it_supports_compound_selectors_with_class_and_id_shorthand(self):
        # Just verifies the parser accepts the syntax — there is no
        # semantic meaning attached to ``.foo`` / ``#bar`` in the docx
        # proxy graph, but the selector should round-trip through the
        # parser without error.
        compile_selector("p.heading[level=1]")
        compile_selector("p#section1")

    def it_rejects_nth_child_with_a_zero_or_negative_argument(self):
        with pytest.raises(SelectorSyntaxError):
            compile_selector("p:nth-child(0)")
        with pytest.raises(SelectorSyntaxError):
            compile_selector("p:nth-child(-2)")


class DescribeSelectorPseudoEdgeCases:
    """Targeted regressions for the pseudo-class machinery."""

    def it_handles_nth_child_with_odd_and_even_keywords(self, doc):
        # ``p:nth-child(odd)`` and ``p:nth-child(even)`` partition
        # every candidate paragraph (body + cell paragraphs), so
        # together they must equal the count of a bare ``p`` selector.
        odds = doc.select("p:nth-child(odd)")
        evens = doc.select("p:nth-child(even)")
        all_p = doc.select("p")
        assert len(odds) + len(evens) == len(all_p)
        # No paragraph appears in both lists.
        odd_ids = {id(p._p) for p in odds}
        even_ids = {id(p._p) for p in evens}
        assert odd_ids.isdisjoint(even_ids)
