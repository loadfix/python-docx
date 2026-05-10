# pyright: reportPrivateUsage=false

"""Unit-test suite for `docx.shape.ShapeLocks` and its wiring on
:class:`InlineShape` / :class:`FloatingImage`.

Covers round-trip of every surfaced lock attribute, the aggregate
:meth:`ShapeLocks.lock_all` / :meth:`ShapeLocks.unlock_all` convenience
methods, and the save + reload fidelity that a real-world authoring
flow depends on.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from docx import Document
from docx.shape import FloatingImage, InlineShape, ShapeLocks

_TEST_IMAGE = (
    Path(__file__).resolve().parent / "test_files" / "python-icon.png"
)


# -- attribute -> XML-name pairs tested by the round-trip battery ----
_ATTR_XML_PAIRS = [
    ("no_select", "noSelect"),
    ("no_move", "noMove"),
    ("no_resize", "noResize"),
    ("no_rotate", "noRot"),
    ("no_change_aspect", "noChangeAspect"),
    ("no_edit_points", "noEditPoints"),
    ("no_adjust_handles", "noAdjustHandles"),
    ("no_change_arrowheads", "noChangeArrowheads"),
    ("no_change_shape_type", "noChangeShapeType"),
    ("no_group", "noGrp"),
    ("no_ungroup", "noUngrp"),
    ("no_text_edit", "noTextEdit"),
]


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def document_with_inline_picture() -> tuple[Document, InlineShape]:
    document = Document()
    shape = document.add_picture(str(_TEST_IMAGE))
    return document, shape


@pytest.fixture
def document_with_floating_picture() -> tuple[Document, FloatingImage]:
    document = Document()
    paragraph = document.add_paragraph()
    floating = paragraph.add_floating_image(str(_TEST_IMAGE))
    return document, floating


# ---------------------------------------------------------------------------
# ShapeLocks on InlineShape
# ---------------------------------------------------------------------------


class DescribeInlineShapeLocks:
    """Unit tests for :class:`ShapeLocks` via :attr:`InlineShape.locks`."""

    def it_exposes_a_ShapeLocks_proxy(
        self, document_with_inline_picture: tuple[Document, InlineShape]
    ):
        _, shape = document_with_inline_picture
        assert isinstance(shape.locks, ShapeLocks)

    def it_raises_when_shape_is_not_a_picture(self):
        # -- build an InlineShape wrapping a non-picture drawing (chart uri) --
        from typing import cast

        from docx.oxml.shape import CT_Inline

        inline = cast(
            CT_Inline, CT_Inline.new_chart_inline(shape_id=1, rId="rId1", cx=1, cy=1)
        )
        with pytest.raises(ValueError, match="picture shapes"):
            InlineShape(inline).locks

    @pytest.mark.parametrize("attr,xml_name", _ATTR_XML_PAIRS)
    def it_round_trips_each_lock_attribute_in_memory(
        self,
        document_with_inline_picture: tuple[Document, InlineShape],
        attr: str,
        xml_name: str,
    ):
        _, shape = document_with_inline_picture

        # -- initial state: attribute unset reads False --
        # -- (add_picture pre-writes noChangeAspect="1" so skip that one's
        #    default-False check) --
        if attr != "no_change_aspect":
            assert getattr(shape.locks, attr) is False

        # -- True sets the attribute on the live XML element --
        setattr(shape.locks, attr, True)
        picLocks = shape._cNvPicPr().picLocks
        assert picLocks is not None
        assert picLocks.get(xml_name) == "1"
        assert getattr(shape.locks, attr) is True

        # -- False removes the attribute entirely --
        setattr(shape.locks, attr, False)
        assert picLocks.get(xml_name) is None
        assert getattr(shape.locks, attr) is False

    def it_returns_False_when_no_picLocks_exists(self):
        # -- use the chart-free helper to build an inline picture with an
        #    empty cNvPicPr (no a:picLocks child) --
        from typing import cast

        from docx.oxml.shape import CT_Inline

        inline = cast(
            CT_Inline,
            CT_Inline.new_pic_inline(
                shape_id=1, rId="rId1", filename="f.png", cx=1, cy=1
            ),
        )
        shape = InlineShape(inline)
        assert shape._cNvPicPr().picLocks is None
        for attr, _ in _ATTR_XML_PAIRS:
            assert getattr(shape.locks, attr) is False, attr
        assert shape.locks.locked is False

    def it_lock_all_sets_every_flag(
        self, document_with_inline_picture: tuple[Document, InlineShape]
    ):
        _, shape = document_with_inline_picture

        shape.locks.lock_all()

        for attr, _ in _ATTR_XML_PAIRS:
            assert getattr(shape.locks, attr) is True, attr
        assert shape.locks.locked is True

    def it_unlock_all_clears_every_flag(
        self, document_with_inline_picture: tuple[Document, InlineShape]
    ):
        _, shape = document_with_inline_picture

        shape.locks.lock_all()
        shape.locks.unlock_all()

        for attr, _ in _ATTR_XML_PAIRS:
            assert getattr(shape.locks, attr) is False, attr
        assert shape.locks.locked is False

    def it_locked_setter_is_equivalent_to_lock_all_and_unlock_all(
        self, document_with_inline_picture: tuple[Document, InlineShape]
    ):
        _, shape = document_with_inline_picture

        shape.locks.locked = True
        for attr, _ in _ATTR_XML_PAIRS:
            assert getattr(shape.locks, attr) is True, attr

        shape.locks.locked = False
        for attr, _ in _ATTR_XML_PAIRS:
            assert getattr(shape.locks, attr) is False, attr

    def it_survives_save_and_reload_with_lock_all(
        self, document_with_inline_picture: tuple[Document, InlineShape]
    ):
        document, shape = document_with_inline_picture

        shape.locks.lock_all()

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = Document(buf)
        shape2 = reopened.inline_shapes[0]

        for attr, _ in _ATTR_XML_PAIRS:
            assert getattr(shape2.locks, attr) is True, attr
        assert shape2.locks.locked is True


# ---------------------------------------------------------------------------
# ShapeLocks on FloatingImage
# ---------------------------------------------------------------------------


class DescribeFloatingImageLocks:
    """Unit tests for :attr:`FloatingImage.locks`."""

    def it_exposes_a_ShapeLocks_proxy(
        self, document_with_floating_picture: tuple[Document, FloatingImage]
    ):
        _, floating = document_with_floating_picture
        assert isinstance(floating.locks, ShapeLocks)

    @pytest.mark.parametrize("attr,xml_name", _ATTR_XML_PAIRS)
    def it_round_trips_each_lock_attribute_in_memory(
        self,
        document_with_floating_picture: tuple[Document, FloatingImage],
        attr: str,
        xml_name: str,
    ):
        _, floating = document_with_floating_picture

        setattr(floating.locks, attr, True)
        picLocks = floating._cNvPicPr().picLocks
        assert picLocks is not None
        assert picLocks.get(xml_name) == "1"
        assert getattr(floating.locks, attr) is True

        setattr(floating.locks, attr, False)
        assert picLocks.get(xml_name) is None
        assert getattr(floating.locks, attr) is False

    def it_lock_all_then_save_and_reload_preserves_every_flag(
        self, document_with_floating_picture: tuple[Document, FloatingImage]
    ):
        document, floating = document_with_floating_picture

        floating.locks.lock_all()

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = Document(buf)
        # -- rebuild the FloatingImage by walking the body for the wp:anchor --
        from docx.oxml.ns import qn

        body = reopened.element.body
        anchor = body.find(".//" + qn("wp:anchor"))
        assert anchor is not None
        floating2 = FloatingImage(anchor)

        for attr, _ in _ATTR_XML_PAIRS:
            assert getattr(floating2.locks, attr) is True, attr

    def it_raises_when_anchor_is_not_a_picture(self):
        from typing import cast

        from docx.oxml.ns import qn
        from docx.oxml.parser import parse_xml
        from docx.oxml.shape import CT_Anchor

        # -- construct a minimal anchor with a non-picture uri --
        from docx.oxml.ns import nsdecls

        xml = (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" '
            'simplePos="0" relativeHeight="0" '
            'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="column">'
            '<wp:posOffset>0</wp:posOffset></wp:positionH>\n'
            '  <wp:positionV relativeFrom="paragraph">'
            '<wp:posOffset>0</wp:posOffset></wp:positionV>\n'
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapSquare wrapText="bothSides"/>\n'
            '  <wp:docPr id="1" name="x"/>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="http://example.com/not-a-picture"/>\n'
            '  </a:graphic>\n'
            "</wp:anchor>" % nsdecls("wp", "a", "r")
        )
        anchor = cast(CT_Anchor, parse_xml(xml))
        with pytest.raises(ValueError, match="picture anchors"):
            FloatingImage(anchor).locks
