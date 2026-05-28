"""Unit-test suite for `docx.text.markdown`, `Paragraph.add_markdown`,
and `_Cell.add_markdown` (issue #23)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.text.markdown import (
    InlineRun,
    apply_markdown_to_paragraph,
    tokenize_blocks,
    tokenize_inline,
)


# ---------------------------------------------------------------------------
# Inline tokenizer
# ---------------------------------------------------------------------------


class DescribeTokenizeInline:
    """Unit tests for `docx.text.markdown.tokenize_inline`."""

    def it_returns_a_single_plain_run_for_plain_text(self):
        runs = tokenize_inline("hello world")
        assert runs == [InlineRun(text="hello world")]

    def it_recognises_double_asterisk_bold(self):
        runs = tokenize_inline("see **this** word")
        assert runs == [
            InlineRun(text="see "),
            InlineRun(text="this", bold=True),
            InlineRun(text=" word"),
        ]

    def it_recognises_underscore_italic(self):
        runs = tokenize_inline("emphasise _this_ token")
        assert runs == [
            InlineRun(text="emphasise "),
            InlineRun(text="this", italic=True),
            InlineRun(text=" token"),
        ]

    def it_recognises_inline_code(self):
        runs = tokenize_inline("call `print` now")
        assert any(r.code and r.text == "print" for r in runs)

    def it_recognises_inline_links_and_attaches_url(self):
        runs = tokenize_inline("see [docs](https://x.com/d) please")
        link_runs = [r for r in runs if r.link is not None]
        assert link_runs == [InlineRun(text="docs", link="https://x.com/d")]


# ---------------------------------------------------------------------------
# Block tokenizer
# ---------------------------------------------------------------------------


class DescribeTokenizeBlocks:
    """Unit tests for `docx.text.markdown.tokenize_blocks`."""

    def it_handles_a_single_paragraph(self):
        blocks = tokenize_blocks("Just some text.")
        assert len(blocks) == 1
        assert blocks[0].kind == "para"

    def it_emits_blank_separator_blocks(self):
        blocks = tokenize_blocks("alpha\n\nbeta")
        kinds = [b.kind for b in blocks]
        assert kinds == ["para", "blank", "para"]

    def it_recognises_top_of_input_headings(self):
        blocks = tokenize_blocks("## Section\n\nbody")
        assert blocks[0].kind == "heading"
        assert blocks[0].level == 2
        assert blocks[0].text == "Section"

    def it_only_treats_a_heading_at_top_of_input(self):
        blocks = tokenize_blocks("body\n\n# Not really a heading")
        kinds = [b.kind for b in blocks]
        assert "heading" not in kinds

    def it_recognises_bullet_lists(self):
        blocks = tokenize_blocks("- alpha\n- beta")
        assert [b.kind for b in blocks] == ["bullet", "bullet"]

    def it_recognises_numbered_lists(self):
        blocks = tokenize_blocks("1. first\n2. second")
        assert [b.kind for b in blocks] == ["number", "number"]


# ---------------------------------------------------------------------------
# Paragraph.add_markdown — the eight required rendering shapes
# ---------------------------------------------------------------------------


class DescribeParagraph_AddMarkdown:
    """Eight `add_markdown` rendering shapes covering issue #23 acceptance."""

    def it_renders_plain_text_as_a_single_run(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("hello world")
        assert para.text == "hello world"
        assert len(para.runs) == 1

    def it_renders_bold_and_italic_inline(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("**big** and _small_")
        assert any(r.bold for r in para.runs)
        assert any(r.italic for r in para.runs)

    def it_renders_inline_code_with_a_monospace_font(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("call `func()` now")
        code_runs = [r for r in para.runs if r.font.name == "Consolas"]
        assert len(code_runs) == 1
        assert code_runs[0].text == "func()"

    def it_renders_a_link_with_a_hyperlink_element(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("see [docs](https://example.com)")
        assert len(para.hyperlinks) == 1
        link = para.hyperlinks[0]
        assert link.url == "https://example.com"
        assert link.text == "docs"

    def it_renders_a_bullet_list_as_List_Bullet_paragraphs(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("- AMER\n- APAC\n- EMEA")
        # -- the document body now holds 3 List-Bullet paragraphs
        bullet_paras = [
            p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"
        ]
        assert len(bullet_paras) == 3
        assert [p.text for p in bullet_paras] == ["AMER", "APAC", "EMEA"]

    def it_renders_a_numbered_list_as_List_Number_paragraphs(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("1. first\n2. second")
        number_paras = [
            p for p in doc.paragraphs if p.style and p.style.name == "List Number"
        ]
        assert len(number_paras) == 2

    def it_renders_a_heading_at_top_with_a_Heading_style(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("## Section\n\nbody text")
        assert para.style is not None and para.style.name == "Heading 2"
        assert para.text == "Section"

    def it_handles_blank_separator_and_soft_breaks(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("line one\nstill paragraph one\n\nparagraph two")
        # -- blank-line separator => 3 paragraphs (content + blank + content)
        # -- soft break inside paragraph one => a w:br element
        from docx.oxml.ns import qn

        body_paras = doc.paragraphs
        assert len(body_paras) == 3
        first = body_paras[0]
        brs = first._p.findall(".//" + qn("w:br"))
        assert len(brs) == 1


# ---------------------------------------------------------------------------
# _Cell.add_markdown
# ---------------------------------------------------------------------------


class DescribeCell_AddMarkdown:
    """Tests for `_Cell.add_markdown`."""

    def it_appends_a_paragraph_with_the_parsed_content(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        first = cell.add_markdown("**bold** in a cell")
        # -- cell now has the original empty paragraph plus the markdown one
        assert any(r.bold for r in first.runs)

    def it_renders_a_link_in_a_cell(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.add_markdown("see [docs](https://example.com)")
        # -- find the paragraph with hyperlinks
        links = []
        for p in cell.paragraphs:
            links.extend(p.hyperlinks)
        assert len(links) == 1
        assert links[0].url == "https://example.com"

    def it_returns_the_first_paragraph(self):
        from docx.text.paragraph import Paragraph

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        result = cell.add_markdown("hello")
        assert isinstance(result, Paragraph)
        assert result.text == "hello"


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


class DescribeAddMarkdown_Edges:
    """Edge cases — empty input, type checks, return value."""

    def it_returns_self_for_chaining_on_paragraph(self):
        doc = Document()
        para = doc.add_paragraph()
        result = para.add_markdown("hello")
        assert result is para

    def it_rejects_non_string_input(self):
        doc = Document()
        para = doc.add_paragraph()
        with pytest.raises(TypeError, match="md must be str"):
            apply_markdown_to_paragraph(para, 123)  # type: ignore[arg-type]

    def it_handles_empty_string_input(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_markdown("")
        # -- no content emitted; the original paragraph is left empty
        assert para.text == ""
