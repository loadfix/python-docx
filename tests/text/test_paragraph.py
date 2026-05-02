"""Unit test suite for the docx.text.paragraph module."""

from typing import List, cast

import pytest

from docx import types as t
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.drawing import Drawing
from docx.fields import Field
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.parts.document import DocumentPart
from docx.parts.story import StoryPart
from docx.section import Section
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.parfmt import ParagraphFormat
from docx.text.run import Run

from ..unitutil.cxml import element, xml
from ..unitutil.mock import call, class_mock, instance_mock, method_mock, property_mock


class DescribeParagraph:
    """Unit-test suite for `docx.text.run.Paragraph`."""

    def it_can_add_an_external_hyperlink(self, request: pytest.FixtureRequest):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.relate_to.return_value = "rId7"
        story_part_.get_style_id.return_value = "Hyperlink"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        hyperlink = paragraph.add_hyperlink(url="https://example.com", text="Click here")

        assert isinstance(hyperlink, Hyperlink)
        assert hyperlink.text == "Click here"
        assert len(hyperlink.runs) == 1
        assert len(paragraph.hyperlinks) == 1
        # -- the hyperlink element has the correct rId --
        assert hyperlink._hyperlink.rId == "rId7"
        story_part_.relate_to.assert_called_once()

    def it_can_add_an_internal_hyperlink(self, request: pytest.FixtureRequest):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.get_style_id.return_value = "Hyperlink"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        hyperlink = paragraph.add_hyperlink(anchor="bookmark1", text="Go to section")

        assert isinstance(hyperlink, Hyperlink)
        assert hyperlink.text == "Go to section"
        assert hyperlink.fragment == "bookmark1"
        assert hyperlink._hyperlink.rId is None

    def it_defaults_text_to_url_when_not_provided(self, request: pytest.FixtureRequest):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.relate_to.return_value = "rId7"
        story_part_.get_style_id.return_value = "Hyperlink"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        hyperlink = paragraph.add_hyperlink(url="https://example.com")

        assert hyperlink.text == "https://example.com"

    def it_raises_when_neither_url_nor_anchor_is_provided(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        with pytest.raises(ValueError, match="Either url or anchor must be provided"):
            paragraph.add_hyperlink()

    def it_raises_when_both_url_and_anchor_are_provided(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        with pytest.raises(ValueError, match="Only one of url or anchor"):
            paragraph.add_hyperlink(url="https://example.com", anchor="bookmark1")

    def it_can_add_a_hyperlink_without_style(self, request: pytest.FixtureRequest):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.relate_to.return_value = "rId7"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        hyperlink = paragraph.add_hyperlink(
            url="https://example.com", text="Click", style=None
        )

        assert isinstance(hyperlink, Hyperlink)
        assert hyperlink.text == "Click"
        # -- no rPr/rStyle should be present --
        runs = hyperlink.runs
        assert len(runs) == 1
        assert runs[0]._r.rPr is None

    def it_can_add_a_page_break(self, fake_parent: t.ProvidesStoryPart):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        result = paragraph.add_page_break()

        assert result is paragraph
        assert paragraph.has_page_break is True
        assert len(paragraph.runs) == 1
        assert paragraph._p.xml == xml("w:p/w:r/w:br{w:type=page}")

    def it_can_add_an_inline_content_control(self, fake_parent: t.ProvidesStoryPart):
        from docx.content_controls import ContentControl, ContentControlType

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        cc = paragraph.add_content_control(
            ContentControlType.PLAIN_TEXT, tag="Foo", title="Bar"
        )

        assert isinstance(cc, ContentControl)
        assert cc.tag == "Foo"
        assert cc.title == "Bar"
        assert cc.type is ContentControlType.PLAIN_TEXT
        # -- the sdt was appended to the paragraph --
        assert len(paragraph._p.xpath("./w:sdt")) == 1

    def it_lists_inline_content_controls(self, fake_parent: t.ProvidesStoryPart):
        from docx.content_controls import ContentControlType

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        paragraph.add_content_control(ContentControlType.PLAIN_TEXT, tag="A")
        paragraph.add_content_control(ContentControlType.PLAIN_TEXT, tag="B")

        ccs = paragraph.content_controls

        assert [cc.tag for cc in ccs] == ["A", "B"]

    def it_includes_sdt_text_in_paragraph_text(self, fake_parent: t.ProvidesStoryPart):
        from docx.content_controls import ContentControlType

        p = cast(CT_P, element('w:p/w:r/w:t"Hello "'))
        paragraph = Paragraph(p, fake_parent)
        cc = paragraph.add_content_control(ContentControlType.PLAIN_TEXT, tag="X")
        cc.text = "world"

        assert paragraph.text == "Hello world"

    def it_can_add_a_floating_image(self, request: pytest.FixtureRequest):
        from docx.oxml.shape import CT_Anchor as _CT_Anchor
        from docx.parts.story import StoryPart
        from docx.shape import FloatingImage

        story_part_ = instance_mock(request, StoryPart)
        anchor = _CT_Anchor.new_pic_anchor(1, "rId1", "foo.png", 1000, 2000)
        story_part_.new_pic_anchor.return_value = anchor

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        result = paragraph.add_floating_image("foo.png", 1000, 2000)

        story_part_.new_pic_anchor.assert_called_once_with("foo.png", 1000, 2000)
        assert isinstance(result, FloatingImage)
        assert len(paragraph.floating_images) == 1
        # -- the anchor element is nested inside a w:r/w:drawing --
        assert paragraph._p.xpath(".//w:r/w:drawing/wp:anchor") == [anchor]

    def it_can_add_a_floating_image_with_custom_position(
        self, request: pytest.FixtureRequest
    ):
        from docx.enum.shape import WD_ANCHOR_H, WD_ANCHOR_V, WD_WRAP_TYPE
        from docx.oxml.shape import CT_Anchor as _CT_Anchor
        from docx.parts.story import StoryPart

        story_part_ = instance_mock(request, StoryPart)
        anchor = _CT_Anchor.new_pic_anchor(1, "rId1", "foo.png", 1000, 2000)
        story_part_.new_pic_anchor.return_value = anchor

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        img = paragraph.add_floating_image(
            "foo.png",
            1000,
            2000,
            position={
                "horizontal": 914400,
                "vertical": 457200,
                "h_anchor": WD_ANCHOR_H.PAGE,
                "v_anchor": WD_ANCHOR_V.MARGIN,
                "wrap": WD_WRAP_TYPE.BEHIND,
            },
        )

        assert img.horizontal_anchor == WD_ANCHOR_H.PAGE
        assert img.vertical_anchor == WD_ANCHOR_V.MARGIN
        assert img.horizontal_offset == 914400
        assert img.vertical_offset == 457200
        assert img.wrap_type == WD_WRAP_TYPE.BEHIND

    def it_can_add_a_shape(self, request: pytest.FixtureRequest):
        from docx.drawing import WordprocessingShape
        from docx.enum.shape import WD_SHAPE
        from docx.oxml.ns import qn
        from docx.parts.story import StoryPart
        from docx.shared import Inches

        story_part_ = instance_mock(request, StoryPart)
        story_part_.next_id = 1

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        shape = paragraph.add_shape(
            WD_SHAPE.RECTANGLE, Inches(2), Inches(1), text="Hi"
        )

        assert isinstance(shape, WordprocessingShape)
        assert shape.shape_type is WD_SHAPE.RECTANGLE
        assert shape.name == "Rectangle 1"
        assert shape.text == "Hi"

        # -- drawing is nested inside a w:r/w:drawing --
        drawings = paragraph._p.xpath(".//w:r/w:drawing")
        assert len(drawings) == 1

        # -- extent reflects custom dimensions --
        extent = drawings[0].find(f"{qn('wp:inline')}/{qn('wp:extent')}")
        assert extent is not None
        assert extent.get("cx") == str(int(Inches(2)))
        assert extent.get("cy") == str(int(Inches(1)))

        # -- a:xfrm/a:ext mirrors wp:extent --
        ext = drawings[0].find(
            f"{qn('wp:inline')}/{qn('a:graphic')}/{qn('a:graphicData')}"
            f"/{qn('wps:wsp')}/{qn('wps:spPr')}/{qn('a:xfrm')}/{qn('a:ext')}"
        )
        assert ext is not None
        assert ext.get("cx") == str(int(Inches(2)))
        assert ext.get("cy") == str(int(Inches(1)))

    @pytest.mark.parametrize(
        ("shape_member", "expected_prst"),
        [
            ("RECTANGLE", "rect"),
            ("ROUNDED_RECTANGLE", "roundRect"),
            ("OVAL", "ellipse"),
            ("ARROW_RIGHT", "rightArrow"),
            ("CALLOUT_ROUNDED_RECTANGLE", "wedgeRoundRectCallout"),
        ],
    )
    def it_maps_each_shape_type_to_the_expected_prst(
        self,
        shape_member: str,
        expected_prst: str,
        request: pytest.FixtureRequest,
    ):
        from docx.enum.shape import WD_SHAPE
        from docx.oxml.ns import qn
        from docx.parts.story import StoryPart
        from docx.shared import Inches

        story_part_ = instance_mock(request, StoryPart)
        story_part_.next_id = 1

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        paragraph.add_shape(WD_SHAPE[shape_member], Inches(1), Inches(1))

        prstGeom = paragraph._p.find(
            f"{qn('w:r')}/{qn('w:drawing')}/{qn('wp:inline')}"
            f"/{qn('a:graphic')}/{qn('a:graphicData')}/{qn('wps:wsp')}"
            f"/{qn('wps:spPr')}/{qn('a:prstGeom')}"
        )
        assert prstGeom is not None
        assert prstGeom.get("prst") == expected_prst

    def it_defaults_dimensions_when_omitted(
        self, request: pytest.FixtureRequest
    ):
        from docx.enum.shape import WD_SHAPE
        from docx.oxml.ns import qn
        from docx.parts.story import StoryPart
        from docx.shared import Inches

        story_part_ = instance_mock(request, StoryPart)
        story_part_.next_id = 1

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        paragraph.add_shape(WD_SHAPE.OVAL)

        extent = paragraph._p.find(
            f"{qn('w:r')}/{qn('w:drawing')}/{qn('wp:inline')}/{qn('wp:extent')}"
        )
        assert extent is not None
        assert extent.get("cx") == str(int(Inches(2)))
        assert extent.get("cy") == str(int(Inches(1)))

    def it_raises_when_shape_type_is_not_a_WD_SHAPE(
        self, request: pytest.FixtureRequest
    ):
        from docx.parts.story import StoryPart

        story_part_ = instance_mock(request, StoryPart)

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        with pytest.raises(TypeError, match="WD_SHAPE"):
            paragraph.add_shape("rect")  # type: ignore[arg-type]

    def it_round_trips_a_created_shape_via_drawings(
        self, request: pytest.FixtureRequest
    ):
        from docx.enum.shape import WD_SHAPE
        from docx.parts.story import StoryPart
        from docx.shared import Inches

        story_part_ = instance_mock(request, StoryPart)
        story_part_.next_id = 7

        class FakeParent:
            @property
            def part(self):
                return story_part_

        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, FakeParent())

        created = paragraph.add_shape(
            WD_SHAPE.OVAL, Inches(2), Inches(1), text="Round-trip"
        )

        # -- find the shape via the public drawings API --
        drawings = paragraph.drawings
        assert len(drawings) == 1
        assert drawings[0].text == "Round-trip"

        # -- the created shape's metadata matches --
        assert created.shape_type is WD_SHAPE.OVAL
        assert created.name == "Oval 7"

    @pytest.mark.parametrize(
        ("p_cxml", "count"),
        [
            ("w:p", 0),
            ("w:p/w:r", 0),
            ("w:p/w:r/w:drawing", 0),
            ("w:p/w:r/w:drawing/wp:inline", 0),
            ("w:p/w:r/w:drawing/wp:anchor", 1),
            (
                "w:p/(w:r/w:drawing/wp:anchor,w:r/w:drawing/wp:anchor)",
                2,
            ),
        ],
    )
    def it_provides_access_to_floating_images_it_contains(
        self, p_cxml: str, count: int, fake_parent: t.ProvidesStoryPart
    ):
        from docx.shape import FloatingImage

        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        floating = paragraph.floating_images

        assert len(floating) == count
        assert all(isinstance(f, FloatingImage) for f in floating)

    @pytest.mark.parametrize(
        ("p_cxml", "expected_value"),
        [
            ("w:p", False),
            ("w:p/w:r", False),
            ('w:p/w:r/w:t"foobar"', False),
            ("w:p/w:r/w:br{w:type=page}", True),
            ("w:p/w:r/w:br", False),
            ('w:p/(w:r/w:t"abc",w:r/w:br{w:type=page})', True),
        ],
    )
    def it_knows_whether_it_has_a_page_break(
        self, p_cxml: str, expected_value: bool, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        assert paragraph.has_page_break == expected_value

    @pytest.mark.parametrize(
        ("p_cxml", "expected_cxml"),
        [
            # --- no page breaks: no-op ---
            ("w:p", "w:p"),
            ("w:p/w:r", "w:p/w:r"),
            # --- run with only page break is removed entirely ---
            ("w:p/w:r/w:br{w:type=page}", "w:p"),
            # --- run with text and page break: only br removed ---
            ('w:p/w:r/(w:t"abc",w:br{w:type=page})', 'w:p/w:r/w:t"abc"'),
            # --- multiple page breaks ---
            (
                'w:p/(w:r/w:br{w:type=page},w:r/w:t"abc",w:r/w:br{w:type=page})',
                'w:p/w:r/w:t"abc"',
            ),
            # --- line break (not page) is preserved ---
            ("w:p/w:r/w:br", "w:p/w:r/w:br"),
        ],
    )
    def it_can_clear_page_breaks(
        self, p_cxml: str, expected_cxml: str, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        paragraph.clear_page_breaks()

        assert paragraph._p.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("body_cxml", "p_idx", "expected_cxml"),
        [
            # --- paragraph is removed from body ---
            ("w:body/(w:p,w:p)", 0, "w:body/w:p"),
            # --- last paragraph in body can be removed ---
            ("w:body/w:p", 0, "w:body"),
            # --- paragraph with formatting is removed ---
            ('w:body/(w:p/w:pPr/w:pStyle{w:val=Heading1},w:p/w:r/w:t"keep")', 0,
             'w:body/w:p/w:r/w:t"keep"'),
            # --- middle paragraph removed ---
            ("w:body/(w:p,w:p,w:p)", 1, "w:body/(w:p,w:p)"),
        ],
    )
    def it_can_delete_itself(
        self,
        body_cxml: str,
        p_idx: int,
        expected_cxml: str,
        fake_parent: t.ProvidesStoryPart,
    ):
        body = element(body_cxml)
        p = body[p_idx]
        paragraph = Paragraph(cast(CT_P, p), fake_parent)

        paragraph.delete()

        assert body.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("p_cxml", "expected_value"),
        [
            ("w:p", False),
            ("w:p/w:pPr", False),
            ("w:p/w:pPr/w:sectPr", True),
            ("w:p/w:pPr/w:sectPr/w:type{w:val=continuous}", True),
        ],
    )
    def it_knows_whether_it_has_a_section_break(
        self, p_cxml: str, expected_value: bool
    ):
        paragraph = Paragraph(cast(CT_P, element(p_cxml)), None)
        assert paragraph.has_section_break is expected_value

    @pytest.mark.parametrize(
        ("p_cxml", "start_type", "expected_start_type"),
        [
            ("w:p", WD_SECTION_START.NEW_PAGE, WD_SECTION_START.NEW_PAGE),
            ("w:p", WD_SECTION_START.CONTINUOUS, WD_SECTION_START.CONTINUOUS),
            ("w:p", WD_SECTION_START.ODD_PAGE, WD_SECTION_START.ODD_PAGE),
            ("w:p", WD_SECTION_START.EVEN_PAGE, WD_SECTION_START.EVEN_PAGE),
            # --- replacing existing sectPr type ---
            (
                "w:p/w:pPr/w:sectPr/w:type{w:val=continuous}",
                WD_SECTION_START.ODD_PAGE,
                WD_SECTION_START.ODD_PAGE,
            ),
        ],
    )
    def it_can_insert_a_section_break(
        self,
        p_cxml: str,
        start_type: WD_SECTION_START,
        expected_start_type: WD_SECTION_START,
        part_prop_: DocumentPart,
    ):
        paragraph = Paragraph(cast(CT_P, element(p_cxml)), None)
        section = paragraph.insert_section_break(start_type)
        assert isinstance(section, Section)
        assert section.start_type == expected_start_type
        assert paragraph.has_section_break is True

    def it_inserts_a_section_break_with_default_start_type(self, part_prop_: DocumentPart):
        paragraph = Paragraph(cast(CT_P, element("w:p")), None)
        section = paragraph.insert_section_break()
        assert isinstance(section, Section)
        assert section.start_type == WD_SECTION_START.NEW_PAGE

    def it_does_not_duplicate_sectPr_on_repeated_insert(self, part_prop_: DocumentPart):
        paragraph = Paragraph(cast(CT_P, element("w:p")), None)
        paragraph.insert_section_break(WD_SECTION_START.CONTINUOUS)
        paragraph.insert_section_break(WD_SECTION_START.ODD_PAGE)
        sectPr_elements = paragraph._p.pPr.xpath("w:sectPr")
        assert len(sectPr_elements) == 1
        assert paragraph.has_section_break is True

    @pytest.mark.parametrize(
        ("p_cxml", "expected_has_break_after"),
        [
            ("w:p/w:pPr/w:sectPr", False),
            ("w:p/w:pPr/w:sectPr/w:type{w:val=continuous}", False),
            ("w:p", False),
            ("w:p/w:pPr", False),
        ],
    )
    def it_can_remove_a_section_break(
        self, p_cxml: str, expected_has_break_after: bool
    ):
        paragraph = Paragraph(cast(CT_P, element(p_cxml)), None)
        paragraph.remove_section_break()
        assert paragraph.has_section_break is expected_has_break_after

    @pytest.mark.parametrize(
        ("p_cxml", "expected_value"),
        [
            ("w:p/w:r", False),
            ('w:p/w:r/w:t"foobar"', False),
            ('w:p/w:hyperlink/w:r/(w:t"abc",w:lastRenderedPageBreak,w:t"def")', True),
            ("w:p/w:r/(w:lastRenderedPageBreak, w:lastRenderedPageBreak)", True),
        ],
    )
    def it_knows_whether_it_contains_a_page_break(
        self, p_cxml: str, expected_value: bool, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        assert paragraph.contains_page_break == expected_value

    @pytest.mark.parametrize(
        ("p_cxml", "count"),
        [
            ("w:p", 0),
            ("w:p/w:r", 0),
            ("w:p/w:r/w:drawing", 1),
            ("w:p/(w:r/w:drawing,w:r/w:drawing)", 2),
            ("w:p/(w:r/w:drawing,w:r)", 1),
        ],
    )
    def it_provides_access_to_drawings_it_contains(
        self, p_cxml: str, count: int, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        drawings = paragraph.drawings

        assert len(drawings) == count
        assert all(isinstance(d, Drawing) for d in drawings)

    @pytest.mark.parametrize(
        ("p_cxml", "count"),
        [
            ("w:p", 0),
            ("w:p/w:r", 0),
            ("w:p/w:hyperlink", 1),
            ("w:p/(w:r,w:hyperlink,w:r)", 1),
            ("w:p/(w:r,w:hyperlink,w:r,w:hyperlink)", 2),
            ("w:p/(w:hyperlink,w:r,w:hyperlink,w:r)", 2),
        ],
    )
    def it_provides_access_to_the_hyperlinks_it_contains(
        self, p_cxml: str, count: int, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        hyperlinks = paragraph.hyperlinks

        actual = [type(item).__name__ for item in hyperlinks]
        expected = ["Hyperlink" for _ in range(count)]
        assert actual == expected, f"expected: {expected}, got: {actual}"

    @pytest.mark.parametrize(
        ("p_cxml", "expected"),
        [
            ("w:p", []),
            ("w:p/w:r", ["Run"]),
            ("w:p/w:hyperlink", ["Hyperlink"]),
            ("w:p/(w:r,w:hyperlink,w:r)", ["Run", "Hyperlink", "Run"]),
            ("w:p/(w:hyperlink,w:r,w:hyperlink)", ["Hyperlink", "Run", "Hyperlink"]),
        ],
    )
    def it_can_iterate_its_inner_content_items(
        self, p_cxml: str, expected: List[str], fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        inner_content = paragraph.iter_inner_content()

        actual = [type(item).__name__ for item in inner_content]
        assert actual == expected, f"expected: {expected}, got: {actual}"

    def it_knows_its_paragraph_style(self, style_get_fixture):
        paragraph, style_id_, style_ = style_get_fixture
        style = paragraph.style
        paragraph.part.get_style.assert_called_once_with(style_id_, WD_STYLE_TYPE.PARAGRAPH)
        assert style is style_

    def it_can_change_its_paragraph_style(self, style_set_fixture):
        paragraph, value, expected_xml = style_set_fixture

        paragraph.style = value

        paragraph.part.get_style_id.assert_called_once_with(value, WD_STYLE_TYPE.PARAGRAPH)
        assert paragraph._p.xml == expected_xml

    @pytest.mark.parametrize(
        ("p_cxml", "count"),
        [
            ("w:p", 0),
            ("w:p/w:r", 0),
            ("w:p/w:r/w:lastRenderedPageBreak", 1),
            ("w:p/w:hyperlink/w:r/w:lastRenderedPageBreak", 1),
            (
                "w:p/(w:r/w:lastRenderedPageBreak,w:hyperlink/w:r/w:lastRenderedPageBreak)",
                2,
            ),
            (
                "w:p/(w:hyperlink/w:r/w:lastRenderedPageBreak,w:r,"
                "w:r/w:lastRenderedPageBreak,w:r,w:hyperlink)",
                2,
            ),
        ],
    )
    def it_provides_access_to_the_rendered_page_breaks_it_contains(
        self, p_cxml: str, count: int, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element(p_cxml))
        paragraph = Paragraph(p, fake_parent)

        rendered_page_breaks = paragraph.rendered_page_breaks

        actual = [type(item).__name__ for item in rendered_page_breaks]
        expected = ["RenderedPageBreak" for _ in range(count)]
        assert actual == expected, f"expected: {expected}, got: {actual}"

    @pytest.mark.parametrize(
        ("p_cxml", "expected_value"),
        [
            ("w:p", ""),
            ("w:p/w:r", ""),
            ("w:p/w:r/w:t", ""),
            ('w:p/w:r/w:t"foo"', "foo"),
            ('w:p/w:r/(w:t"foo", w:t"bar")', "foobar"),
            ('w:p/w:r/(w:t"fo ", w:t"bar")', "fo bar"),
            ('w:p/w:r/(w:t"foo", w:tab, w:t"bar")', "foo\tbar"),
            ('w:p/w:r/(w:t"foo", w:br,  w:t"bar")', "foo\nbar"),
            ('w:p/w:r/(w:t"foo", w:cr,  w:t"bar")', "foo\nbar"),
            (
                'w:p/(w:r/w:t"click ",w:hyperlink{r:id=rId6}/w:r/w:t"here",w:r/w:t" for more")',
                "click here for more",
            ),
        ],
    )
    def it_knows_the_text_it_contains(self, p_cxml: str, expected_value: str):
        """Including the text of embedded hyperlinks."""
        paragraph = Paragraph(element(p_cxml), None)
        assert paragraph.text == expected_value

    def it_includes_fldSimple_text_in_paragraph_text(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p_cxml = (
            'w:p/(w:r/w:t"Page ",w:fldSimple{w:instr=PAGE}/w:r/w:t"3",w:r/w:t" of 10")'
        )
        paragraph = Paragraph(cast(CT_P, element(p_cxml)), fake_parent)

        assert paragraph.text == "Page 3 of 10"

    def it_can_add_a_simple_field(self, fake_parent: t.ProvidesStoryPart):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        field = paragraph.add_simple_field("PAGE", "3")

        assert isinstance(field, Field)
        assert field.is_complex is False
        assert field.instruction == "PAGE"
        assert field.type == "PAGE"
        assert field.result_text == "3"
        assert paragraph.text == "3"

    def it_can_add_a_simple_field_without_result_text(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        field = paragraph.add_simple_field("DATE")

        assert field.instruction == "DATE"
        assert field.result_text == ""

    def it_can_add_a_complex_field(self, fake_parent: t.ProvidesStoryPart):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        field = paragraph.add_complex_field("REF bookmark1 \\h", "See here")

        assert isinstance(field, Field)
        assert field.is_complex is True
        assert field.instruction == "REF bookmark1 \\h"
        assert field.type == "REF"
        assert field.result_text == "See here"
        assert paragraph.text == "See here"

    def it_can_add_a_complex_field_without_result_text(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)

        field = paragraph.add_complex_field("PAGE")

        assert field.is_complex is True
        assert field.instruction == "PAGE"
        assert field.result_text == ""

    def it_provides_access_to_its_fields(self, fake_parent: t.ProvidesStoryPart):
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, fake_parent)
        paragraph.add_run("Page ")
        paragraph.add_simple_field("PAGE", "3")
        paragraph.add_run(" of ")
        paragraph.add_complex_field("NUMPAGES", "10")
        paragraph.add_simple_field("DATE", "2026-01-01")

        fields = paragraph.fields

        assert len(fields) == 3
        assert [f.type for f in fields] == ["PAGE", "NUMPAGES", "DATE"]
        assert [f.is_complex for f in fields] == [False, True, False]
        assert [f.result_text for f in fields] == ["3", "10", "2026-01-01"]

    def it_returns_an_empty_field_list_when_none_present(
        self, fake_parent: t.ProvidesStoryPart
    ):
        p = cast(CT_P, element('w:p/w:r/w:t"just text"'))
        paragraph = Paragraph(p, fake_parent)

        assert paragraph.fields == []

    def it_can_resolve_a_REF_field_against_a_document_bookmark(self):
        # -- end-to-end: build a document with a bookmark, add a REF field to
        #    a paragraph in that document, then resolve the field. --
        from docx.document import Document
        from docx.oxml.document import CT_Document

        doc_elm = cast(
            CT_Document,
            element(
                "w:document/w:body/w:p/("
                "w:bookmarkStart{w:id=0,w:name=Ref1}"
                ',w:r/w:t"target text"'
                ",w:bookmarkEnd{w:id=0}"
                ")"
            ),
        )
        doc = Document(doc_elm, None)  # type: ignore[arg-type]
        # -- locate the single paragraph in the body and append a REF field --
        p_elm = cast(CT_P, doc_elm.body.p_lst[0])
        paragraph = Paragraph(p_elm, None)  # type: ignore[arg-type]
        field = paragraph.add_complex_field("REF Ref1 \\h", "stale")

        assert field.resolve(doc) == "target text"

    @pytest.mark.parametrize(
        ("p_cxml", "count"),
        [
            ("w:p", 0),
            ('w:p/w:r/w:t"no changes"', 0),
            ('w:p/w:ins{w:id=1,w:author=A}/w:r/w:t"added"', 1),
            ('w:p/w:del{w:id=2,w:author=B}/w:r/w:delText"removed"', 1),
            (
                'w:p/(w:ins{w:id=1,w:author=A}/w:r/w:t"added"'
                ',w:del{w:id=2,w:author=B}/w:r/w:delText"removed")',
                2,
            ),
            ('w:p/w:moveFrom{w:id=3,w:author=A,w:name=m1}/w:r/w:delText"gone"', 1),
            ('w:p/w:moveTo{w:id=4,w:author=B,w:name=m1}/w:r/w:t"moved"', 1),
            (
                "w:p/("
                'w:ins{w:id=1,w:author=A}/w:r/w:t"i",'
                'w:del{w:id=2,w:author=B}/w:r/w:delText"d",'
                'w:moveFrom{w:id=3,w:author=C,w:name=m1}/w:r/w:delText"mf",'
                'w:moveTo{w:id=4,w:author=D,w:name=m1}/w:r/w:t"mt"'
                ")",
                4,
            ),
        ],
    )
    def it_provides_access_to_tracked_changes(self, p_cxml: str, count: int):
        paragraph = Paragraph(element(p_cxml), None)

        tracked_changes = paragraph.tracked_changes

        assert len(tracked_changes) == count

    def it_wraps_move_revisions_in_MoveRevision_proxies(self):
        from docx.tracked_changes import MoveRevision

        paragraph = Paragraph(
            element(
                "w:p/("
                'w:ins{w:id=1,w:author=A}/w:r/w:t"i",'
                'w:moveFrom{w:id=3,w:author=C,w:name=m1}/w:r/w:delText"mf",'
                'w:moveTo{w:id=4,w:author=D,w:name=m1}/w:r/w:t"mt"'
                ")"
            ),
            None,
        )

        changes = paragraph.tracked_changes

        types = [c.type for c in changes]
        assert types == ["insertion", "move_from", "move_to"]
        assert isinstance(changes[1], MoveRevision)
        assert isinstance(changes[2], MoveRevision)
        assert changes[1].name == "m1"  # type: ignore[attr-defined]
        assert changes[2].name == "m1"  # type: ignore[attr-defined]

    def it_exposes_its_formatting_change_when_pPrChange_present(self):
        p = cast(
            CT_P,
            element(
                "w:p/w:pPr/(w:jc{w:val=center}"
                ",w:pPrChange{w:id=1,w:author=Alice}/w:pPr/w:jc{w:val=left})"
            ),
        )
        paragraph = Paragraph(p, None)

        fc = paragraph.formatting_change

        assert fc is not None
        assert fc.author == "Alice"
        assert fc.old_properties is not None
        assert fc.old_properties.xpath("./w:jc")

    def it_returns_None_for_formatting_change_when_no_pPr(self):
        paragraph = Paragraph(cast(CT_P, element("w:p")), None)
        assert paragraph.formatting_change is None

    def it_returns_None_for_formatting_change_when_no_pPrChange(self):
        paragraph = Paragraph(cast(CT_P, element("w:p/w:pPr/w:jc{w:val=left}")), None)
        assert paragraph.formatting_change is None

    def it_returns_plain_text_for_revision_marks_when_no_changes(self):
        paragraph = Paragraph(cast(CT_P, element('w:p/w:r/w:t"hello"')), None)
        assert paragraph.revision_marks_text() == paragraph.text == "hello"

    def it_renders_tracked_changes_with_default_markers(self):
        paragraph = Paragraph(
            cast(
                CT_P,
                element(
                    'w:p/(w:r/w:t"a ",'
                    'w:ins{w:id=1,w:author=A}/w:r/w:t"B",'
                    'w:del{w:id=2,w:author=B}/w:r/w:delText"c",'
                    'w:r/w:t" d")'
                ),
            ),
            None,
        )

        assert paragraph.revision_marks_text() == "a [+B+][-c-] d"

    def it_honors_custom_revision_markers(self):
        paragraph = Paragraph(
            cast(
                CT_P,
                element(
                    'w:p/(w:ins{w:id=1,w:author=A}/w:r/w:t"new",'
                    'w:del{w:id=2,w:author=B}/w:r/w:delText"old")'
                ),
            ),
            None,
        )

        rendered = paragraph.revision_marks_text(
            open_ins="<I>", close_ins="</I>", open_del="<D>", close_del="</D>"
        )

        assert rendered == "<I>new</I><D>old</D>"

    def it_can_replace_the_text_it_contains(self, text_set_fixture):
        paragraph, text, expected_text = text_set_fixture
        paragraph.text = text
        assert paragraph.text == expected_text

    def it_knows_its_alignment_value(self, alignment_get_fixture):
        paragraph, expected_value = alignment_get_fixture
        assert paragraph.alignment == expected_value

    def it_can_change_its_alignment_value(self, alignment_set_fixture):
        paragraph, value, expected_xml = alignment_set_fixture
        paragraph.alignment = value
        assert paragraph._p.xml == expected_xml

    def it_provides_access_to_its_paragraph_format(self, parfmt_fixture):
        paragraph, ParagraphFormat_, paragraph_format_ = parfmt_fixture
        paragraph_format = paragraph.paragraph_format
        ParagraphFormat_.assert_called_once_with(paragraph._element)
        assert paragraph_format is paragraph_format_

    def it_provides_access_to_the_runs_it_contains(self, runs_fixture):
        paragraph, Run_, r_, r_2_, run_, run_2_ = runs_fixture
        runs = paragraph.runs
        assert Run_.mock_calls == [call(r_, paragraph), call(r_2_, paragraph)]
        assert runs == [run_, run_2_]

    def it_includes_smartTag_wrapped_runs_in_runs_list(self):
        # -- upstream#932, #225: runs inside `w:smartTag` must be visible --
        from docx.oxml.parser import parse_xml

        xml_bytes = (
            b'<w:p xmlns:w='
            b'"http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b'<w:r><w:t>a</w:t></w:r>'
            b'<w:smartTag><w:r><w:t>b</w:t></w:r></w:smartTag>'
            b'<w:r><w:t>c</w:t></w:r>'
            b'</w:p>'
        )
        paragraph = Paragraph(cast(CT_P, parse_xml(xml_bytes)), None)
        assert [r.text for r in paragraph.runs] == ["a", "b", "c"]

    def it_can_add_a_run_to_itself(self, add_run_fixture):
        paragraph, text, style, style_prop_, expected_xml = add_run_fixture
        run = paragraph.add_run(text, style)
        assert paragraph._p.xml == expected_xml
        assert isinstance(run, Run)
        assert run._r is paragraph._p.r_lst[0]
        if style:
            style_prop_.assert_called_once_with(style)

    def it_can_insert_a_paragraph_before_itself(self, insert_before_fixture):
        text, style, paragraph_, add_run_calls = insert_before_fixture
        paragraph = Paragraph(None, None)

        new_paragraph = paragraph.insert_paragraph_before(text, style)

        paragraph._insert_paragraph_before.assert_called_once_with(paragraph)
        assert new_paragraph.add_run.call_args_list == add_run_calls
        assert new_paragraph.style == style
        assert new_paragraph is paragraph_

    def it_updates_section_count_on_insert_and_remove(self, part_prop_: DocumentPart):
        document_elm = element(
            "w:document/w:body/(w:p,w:p,w:sectPr)"
        )
        body = document_elm[0]
        p1 = body[0]
        p2 = body[1]
        paragraph1 = Paragraph(cast(CT_P, p1), None)
        paragraph2 = Paragraph(cast(CT_P, p2), None)
        # --- starts with 1 section (the body sectPr) ---
        assert len(document_elm.sectPr_lst) == 1
        # --- insert section break on paragraph1 ---
        paragraph1.insert_section_break(WD_SECTION_START.CONTINUOUS)
        assert len(document_elm.sectPr_lst) == 2
        # --- insert section break on paragraph2 ---
        paragraph2.insert_section_break(WD_SECTION_START.ODD_PAGE)
        assert len(document_elm.sectPr_lst) == 3
        # --- remove section break from paragraph1 ---
        paragraph1.remove_section_break()
        assert len(document_elm.sectPr_lst) == 2
        # --- remove section break from paragraph2 ---
        paragraph2.remove_section_break()
        assert len(document_elm.sectPr_lst) == 1

    def it_can_remove_its_content_while_preserving_formatting(self, clear_fixture):
        paragraph, expected_xml = clear_fixture
        _paragraph = paragraph.clear()
        assert paragraph._p.xml == expected_xml
        assert _paragraph is paragraph

    def it_inserts_a_paragraph_before_to_help(self, _insert_before_fixture):
        paragraph, body, expected_xml = _insert_before_fixture
        new_paragraph = paragraph._insert_paragraph_before()
        assert isinstance(new_paragraph, Paragraph)
        assert body.xml == expected_xml

    def it_can_insert_a_paragraph_after_itself(self, fake_parent: t.ProvidesStoryPart):
        body = element("w:body/(w:p{id=1},w:p{id=2})")
        p1 = body[0]
        paragraph = Paragraph(cast(CT_P, p1), fake_parent)

        new_paragraph = paragraph.insert_paragraph_after()

        assert isinstance(new_paragraph, Paragraph)
        # -- new paragraph sits between p1 and p2 --
        assert list(body) == [p1, new_paragraph._p, body[2]]
        assert new_paragraph._parent is fake_parent

    def it_can_insert_a_paragraph_after_with_text(
        self, fake_parent: t.ProvidesStoryPart
    ):
        body = element("w:body/w:p{id=1}")
        paragraph = Paragraph(cast(CT_P, body[0]), fake_parent)

        new_paragraph = paragraph.insert_paragraph_after(text="hello")

        assert new_paragraph.text == "hello"
        # -- new paragraph sits after the reference paragraph --
        assert list(body) == [paragraph._p, new_paragraph._p]

    def it_can_insert_a_paragraph_after_with_text_and_style(
        self, request: pytest.FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.get_style_id.return_value = "Heading1"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        body = element("w:body/w:p")
        paragraph = Paragraph(cast(CT_P, body[0]), FakeParent())

        new_paragraph = paragraph.insert_paragraph_after(text="hi", style="Heading 1")

        assert new_paragraph.text == "hi"
        story_part_.get_style_id.assert_called_with("Heading 1", WD_STYLE_TYPE.PARAGRAPH)

    def it_can_insert_a_paragraph_after_inside_a_cell(
        self, fake_parent: t.ProvidesStoryPart
    ):
        # -- the paragraph is inside a w:tc (cell). The new paragraph should also be
        # -- inserted in the cell, not at the body level. --
        tc = element("w:tc/(w:p{id=1},w:p{id=2})")
        p1 = tc[0]
        paragraph = Paragraph(cast(CT_P, p1), fake_parent)

        new_paragraph = paragraph.insert_paragraph_after(text="middle")

        assert list(tc) == [p1, new_paragraph._p, tc[2]]
        assert new_paragraph.text == "middle"

    def it_can_add_a_caption_after_itself(
        self, request: pytest.FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.get_style_id.return_value = "Caption"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        body = element("w:body/(w:p{id=1},w:p{id=2})")
        paragraph = Paragraph(cast(CT_P, body[0]), FakeParent())

        caption = paragraph.add_caption_after("A diagram")

        assert isinstance(caption, Paragraph)
        # -- the caption sits between the two existing paragraphs --
        assert list(body) == [paragraph._p, caption._p, body[2]]
        assert caption.text == "Figure 1: A diagram"
        # -- one SEQ field with the correct instruction and cached "1" --
        assert len(caption.fields) == 1
        field = caption.fields[0]
        assert field.type == "SEQ"
        assert field.instruction.strip() == "SEQ Figure \\* ARABIC"
        assert field.result_text == "1"

    def it_can_add_a_caption_before_itself(
        self, request: pytest.FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.get_style_id.return_value = "Caption"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        body = element("w:body/(w:p{id=1},w:p{id=2})")
        paragraph = Paragraph(cast(CT_P, body[1]), FakeParent())

        caption = paragraph.add_caption_before("A diagram")

        assert isinstance(caption, Paragraph)
        # -- the caption sits between the two existing paragraphs --
        assert list(body) == [body[0], caption._p, paragraph._p]
        assert caption.text == "Figure 1: A diagram"

    def it_honors_custom_label_on_add_caption_after(
        self, request: pytest.FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.get_style_id.return_value = "Caption"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        body = element("w:body/w:p")
        paragraph = Paragraph(cast(CT_P, body[0]), FakeParent())

        caption = paragraph.add_caption_after("Prices", label="Table")

        assert caption.text == "Table 1: Prices"
        assert caption.fields[0].instruction.strip() == "SEQ Table \\* ARABIC"

    def it_honors_custom_style_on_add_caption_before(
        self, request: pytest.FixtureRequest
    ):
        story_part_ = instance_mock(request, StoryPart)
        story_part_.get_style_id.return_value = "MyStyleId"

        class FakeParent:
            @property
            def part(self):
                return story_part_

        body = element("w:body/w:p")
        paragraph = Paragraph(cast(CT_P, body[0]), FakeParent())

        caption = paragraph.add_caption_before("Hello", style="MyStyle")

        # -- pStyle w:val came from get_style_id("MyStyle", ...) --
        story_part_.get_style_id.assert_any_call("MyStyle", WD_STYLE_TYPE.PARAGRAPH)
        pPr = caption._p.pPr
        assert pPr is not None
        assert pPr.style == "MyStyleId"

    def it_can_insert_a_table_before_itself(
        self, fake_parent: t.ProvidesStoryPart
    ):
        body = element("w:body/(w:p{id=1},w:p{id=2})")
        p2 = body[1]
        paragraph = Paragraph(cast(CT_P, p2), fake_parent)

        from docx.table import Table

        table = paragraph.insert_table_before(rows=2, cols=2)

        assert isinstance(table, Table)
        # -- the new table sits between p1 and p2 --
        assert list(body) == [body[0], table._tbl, p2]
        assert table._parent is fake_parent
        # -- structure: w:tblPr, w:tblGrid, 2 rows each with 2 cells --
        assert len(table._tbl.tr_lst) == 2
        assert table._tbl.col_count == 2

    def it_can_insert_a_table_after_itself(
        self, fake_parent: t.ProvidesStoryPart
    ):
        body = element("w:body/(w:p{id=1},w:p{id=2})")
        p1 = body[0]
        paragraph = Paragraph(cast(CT_P, p1), fake_parent)

        from docx.table import Table

        table = paragraph.insert_table_after(rows=1, cols=3)

        assert isinstance(table, Table)
        # -- the new table sits between p1 and p2 --
        assert list(body) == [p1, table._tbl, body[2]]
        assert len(table._tbl.tr_lst) == 1
        assert table._tbl.col_count == 3

    def it_can_insert_a_table_after_itself_inside_a_cell(
        self, fake_parent: t.ProvidesStoryPart
    ):
        tc = element("w:tc/w:p{id=1}")
        p1 = tc[0]
        paragraph = Paragraph(cast(CT_P, p1), fake_parent)

        table = paragraph.insert_table_after(rows=1, cols=1)

        # -- the new table is a sibling of the paragraph in the cell --
        assert list(tc) == [p1, table._tbl]

    # fixtures -------------------------------------------------------

    @pytest.fixture(
        params=[
            ("w:p", None, None, "w:p/w:r"),
            ("w:p", "foobar", None, 'w:p/w:r/w:t"foobar"'),
            ("w:p", None, "Strong", "w:p/w:r"),
            ("w:p", "foobar", "Strong", 'w:p/w:r/w:t"foobar"'),
        ]
    )
    def add_run_fixture(self, request, run_style_prop_):
        before_cxml, text, style, after_cxml = request.param
        paragraph = Paragraph(element(before_cxml), None)
        expected_xml = xml(after_cxml)
        return paragraph, text, style, run_style_prop_, expected_xml

    @pytest.fixture(
        params=[
            ("w:p/w:pPr/w:jc{w:val=center}", WD_ALIGN_PARAGRAPH.CENTER),
            ("w:p", None),
        ]
    )
    def alignment_get_fixture(self, request):
        cxml, expected_alignment_value = request.param
        paragraph = Paragraph(element(cxml), None)
        return paragraph, expected_alignment_value

    @pytest.fixture(
        params=[
            ("w:p", WD_ALIGN_PARAGRAPH.LEFT, "w:p/w:pPr/w:jc{w:val=left}"),
            (
                "w:p/w:pPr/w:jc{w:val=left}",
                WD_ALIGN_PARAGRAPH.CENTER,
                "w:p/w:pPr/w:jc{w:val=center}",
            ),
            ("w:p/w:pPr/w:jc{w:val=left}", None, "w:p/w:pPr"),
            ("w:p", None, "w:p/w:pPr"),
        ]
    )
    def alignment_set_fixture(self, request):
        initial_cxml, new_alignment_value, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, new_alignment_value, expected_xml

    @pytest.fixture(
        params=[
            ("w:p", "w:p"),
            ("w:p/w:pPr", "w:p/w:pPr"),
            ('w:p/w:r/w:t"foobar"', "w:p"),
            ('w:p/(w:pPr, w:r/w:t"foobar")', "w:p/w:pPr"),
        ]
    )
    def clear_fixture(self, request):
        initial_cxml, expected_cxml = request.param
        paragraph = Paragraph(element(initial_cxml), None)
        expected_xml = xml(expected_cxml)
        return paragraph, expected_xml

    @pytest.fixture(
        params=[
            (None, None),
            ("Foo", None),
            (None, "Bar"),
            ("Foo", "Bar"),
        ]
    )
    def insert_before_fixture(self, request, _insert_paragraph_before_, add_run_):
        text, style = request.param
        paragraph_ = _insert_paragraph_before_.return_value
        add_run_calls = [] if text is None else [call(text)]
        paragraph_.style = None
        return text, style, paragraph_, add_run_calls

    @pytest.fixture(params=[("w:body/w:p{id=42}", "w:body/(w:p,w:p{id=42})")])
    def _insert_before_fixture(self, request):
        body_cxml, expected_cxml = request.param
        body = element(body_cxml)
        paragraph = Paragraph(body[0], None)
        expected_xml = xml(expected_cxml)
        return paragraph, body, expected_xml

    @pytest.fixture
    def parfmt_fixture(self, ParagraphFormat_, paragraph_format_):
        paragraph = Paragraph(element("w:p"), None)
        return paragraph, ParagraphFormat_, paragraph_format_

    @pytest.fixture
    def runs_fixture(self, p_, Run_, r_, r_2_, runs_):
        paragraph = Paragraph(p_, None)
        run_, run_2_ = runs_
        return paragraph, Run_, r_, r_2_, run_, run_2_

    @pytest.fixture
    def style_get_fixture(self, part_prop_):
        style_id = "Foobar"
        p_cxml = "w:p/w:pPr/w:pStyle{w:val=%s}" % style_id
        paragraph = Paragraph(element(p_cxml), None)
        style_ = part_prop_.return_value.get_style.return_value
        return paragraph, style_id, style_

    @pytest.fixture(
        params=[
            ("w:p", "Heading 1", "Heading1", "w:p/w:pPr/w:pStyle{w:val=Heading1}"),
            (
                "w:p/w:pPr",
                "Heading 1",
                "Heading1",
                "w:p/w:pPr/w:pStyle{w:val=Heading1}",
            ),
            (
                "w:p/w:pPr/w:pStyle{w:val=Heading1}",
                "Heading 2",
                "Heading2",
                "w:p/w:pPr/w:pStyle{w:val=Heading2}",
            ),
            ("w:p/w:pPr/w:pStyle{w:val=Heading1}", "Normal", None, "w:p/w:pPr"),
            ("w:p", None, None, "w:p/w:pPr"),
        ]
    )
    def style_set_fixture(self, request, part_prop_):
        p_cxml, value, style_id, expected_cxml = request.param
        paragraph = Paragraph(element(p_cxml), None)
        part_prop_.return_value.get_style_id.return_value = style_id
        expected_xml = xml(expected_cxml)
        return paragraph, value, expected_xml

    @pytest.fixture
    def text_set_fixture(self):
        paragraph = Paragraph(element("w:p"), None)
        paragraph.add_run("must not appear in result")
        new_text_value = "foo\tbar\rbaz\n"
        expected_text_value = "foo\tbar\nbaz\n"
        return paragraph, new_text_value, expected_text_value

    # fixture components ---------------------------------------------

    @pytest.fixture
    def add_run_(self, request):
        return method_mock(request, Paragraph, "add_run")

    @pytest.fixture
    def document_part_(self, request):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def _insert_paragraph_before_(self, request):
        return method_mock(request, Paragraph, "_insert_paragraph_before")

    @pytest.fixture
    def p_(self, request, r_, r_2_):
        p = instance_mock(request, CT_P, r_lst=(r_, r_2_))
        # -- `Paragraph.runs` iterates via `iter_r_elements` so the mock must
        # -- expose that generator-producing method. --
        p.iter_r_elements.return_value = iter((r_, r_2_))
        return p

    @pytest.fixture
    def ParagraphFormat_(self, request, paragraph_format_):
        return class_mock(
            request,
            "docx.text.paragraph.ParagraphFormat",
            return_value=paragraph_format_,
        )

    @pytest.fixture
    def paragraph_format_(self, request):
        return instance_mock(request, ParagraphFormat)

    @pytest.fixture
    def part_prop_(self, request, document_part_):
        return property_mock(request, Paragraph, "part", return_value=document_part_)

    @pytest.fixture
    def Run_(self, request, runs_):
        run_, run_2_ = runs_
        return class_mock(request, "docx.text.paragraph.Run", side_effect=[run_, run_2_])

    @pytest.fixture
    def r_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def r_2_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def run_style_prop_(self, request):
        return property_mock(request, Run, "style")

    @pytest.fixture
    def runs_(self, request):
        run_ = instance_mock(request, Run, name="run_")
        run_2_ = instance_mock(request, Run, name="run_2_")
        return run_, run_2_


class DescribeParagraph_ListFeatures:
    """Unit-test suite for the list/numbering API on `Paragraph`."""

    def it_reports_None_list_level_when_no_numPr(self, request: pytest.FixtureRequest):
        paragraph = self._make_paragraph(request, "w:p")
        assert paragraph.list_level is None

    def it_reads_list_level_from_numPr_ilvl(self, request: pytest.FixtureRequest):
        paragraph = self._make_paragraph(
            request,
            "w:p/w:pPr/w:numPr/(w:ilvl{w:val=2},w:numId{w:val=1})",
        )
        assert paragraph.list_level == 2

    def it_can_set_list_level_round_trip(self, request: pytest.FixtureRequest):
        paragraph = self._make_paragraph(request, "w:p")
        paragraph.list_level = 3
        assert paragraph.list_level == 3

    def it_rejects_an_out_of_range_list_level(self, request: pytest.FixtureRequest):
        paragraph = self._make_paragraph(request, "w:p")
        with pytest.raises(ValueError):
            paragraph.list_level = 9

    def it_returns_empty_list_format_when_not_in_a_list(
        self, request: pytest.FixtureRequest
    ):
        paragraph = self._make_paragraph(request, "w:p")

        list_format = paragraph.list_format

        assert list_format.numbering_definition is None
        assert list_format.level is None

    def it_resolves_list_format_from_the_numbering_part(
        self, request: pytest.FixtureRequest
    ):
        from docx.numbering import Numbering
        from docx.enum.text import WD_NUMBER_FORMAT
        from docx.oxml.numbering import CT_Numbering
        from docx.parts.numbering import NumberingPart

        # -- build a real numbering part with one definition, one num --
        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering_part_ = instance_mock(request, NumberingPart)
        numbering_part_.numbering_element = numbering_elm
        numbering = Numbering(numbering_elm, numbering_part_)
        defn = numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )
        num_id = numbering_elm.num_lst[0].numId

        p = cast(
            CT_P,
            element(
                f"w:p/w:pPr/w:numPr/(w:ilvl{{w:val=0}},w:numId{{w:val={num_id}}})"
            ),
        )
        part_ = instance_mock(request, DocumentPart)
        part_.numbering_part = numbering_part_

        class FakeParent:
            @property
            def part(self):
                return part_

        paragraph = Paragraph(p, FakeParent())

        list_format = paragraph.list_format

        assert list_format.level == 0
        assert list_format.numbering_definition is not None
        assert list_format.numbering_definition.abstract_num_id == defn.abstract_num_id

    def it_can_restart_numbering(self, request: pytest.FixtureRequest):
        from docx.numbering import Numbering
        from docx.enum.text import WD_NUMBER_FORMAT
        from docx.oxml.numbering import CT_Numbering
        from docx.oxml.ns import qn
        from docx.parts.numbering import NumberingPart

        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering_part_ = instance_mock(request, NumberingPart)
        numbering_part_.numbering_element = numbering_elm
        numbering = Numbering(numbering_elm, numbering_part_)
        numbering.add_numbering_definition(
            levels=[{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."}]
        )
        original_num_id = numbering_elm.num_lst[0].numId

        p = cast(
            CT_P,
            element(
                f"w:p/w:pPr/w:numPr/"
                f"(w:ilvl{{w:val=0}},w:numId{{w:val={original_num_id}}})"
            ),
        )
        part_ = instance_mock(request, DocumentPart)
        part_.numbering_part = numbering_part_

        class FakeParent:
            @property
            def part(self):
                return part_

        paragraph = Paragraph(p, FakeParent())

        paragraph.restart_numbering(start=1)

        # -- a new w:num was created --
        assert len(numbering_elm.num_lst) == 2
        new_num_id = numbering_elm.num_lst[-1].numId
        assert new_num_id != original_num_id

        # -- the paragraph's numId now points at the new num --
        assert paragraph._p.pPr.numPr.numId_val == new_num_id

        # -- the new num has the startOverride child --
        new_num = numbering_elm.num_lst[-1]
        overrides = new_num.xpath("./w:lvlOverride")
        assert len(overrides) == 1
        assert overrides[0].get(qn("w:ilvl")) == "0"
        start_override = overrides[0].xpath("./w:startOverride")
        assert len(start_override) == 1
        assert start_override[0].get(qn("w:val")) == "1"

    def it_raises_when_restart_numbering_called_on_non_list_paragraph(
        self, request: pytest.FixtureRequest
    ):
        paragraph = self._make_paragraph(request, "w:p")

        with pytest.raises(ValueError):
            paragraph.restart_numbering()

    def it_exposes_numbering_format_for_the_current_level(
        self, request: pytest.FixtureRequest
    ):
        from docx.numbering import Numbering
        from docx.enum.text import WD_NUMBER_FORMAT
        from docx.oxml.numbering import CT_Numbering
        from docx.parts.numbering import NumberingPart

        numbering_elm = cast(CT_Numbering, element("w:numbering"))
        numbering_part_ = instance_mock(request, NumberingPart)
        numbering_part_.numbering_element = numbering_elm
        numbering = Numbering(numbering_elm, numbering_part_)
        numbering.add_numbering_definition(
            levels=[
                {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
                {"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%2)"},
            ]
        )
        num_id = numbering_elm.num_lst[0].numId

        p = cast(
            CT_P,
            element(
                f"w:p/w:pPr/w:numPr/"
                f"(w:ilvl{{w:val=1}},w:numId{{w:val={num_id}}})"
            ),
        )
        part_ = instance_mock(request, DocumentPart)
        part_.numbering_part = numbering_part_

        class FakeParent:
            @property
            def part(self):
                return part_

        paragraph = Paragraph(p, FakeParent())

        fmt = paragraph.numbering_format

        assert fmt is not None
        assert fmt.number_format == WD_NUMBER_FORMAT.LOWER_LETTER
        assert fmt.text == "%2)"

    # -- helpers --

    @staticmethod
    def _make_paragraph(request: pytest.FixtureRequest, cxml: str) -> Paragraph:
        p = cast(CT_P, element(cxml))
        part_ = instance_mock(request, DocumentPart)

        class FakeParent:
            @property
            def part(self):
                return part_

        return Paragraph(p, FakeParent())


class DescribeParagraph_Rsid:
    """Unit-test suite for the RSID accessor on `Paragraph`."""

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:p", None),
            ("w:p{w:rsidR=00FA1B42}", "00FA1B42"),
            ("w:p{w:rsidR=001234AB,w:rsidRPr=005678CD}", "001234AB"),
        ],
    )
    def it_reads_rsid_from_rsidR_attribute(
        self, request: pytest.FixtureRequest, cxml: str, expected_value: str | None
    ):
        p = cast(CT_P, element(cxml))
        part_ = instance_mock(request, DocumentPart)

        class FakeParent:
            @property
            def part(self):
                return part_

        paragraph = Paragraph(p, FakeParent())
        assert paragraph.rsid == expected_value


class DescribeParagraph_StableId:
    """Unit-test suite for the `stable_id` accessor on `Paragraph`."""

    def _make_paragraph(
        self, request: pytest.FixtureRequest, p_element: CT_P
    ) -> Paragraph:
        part_ = instance_mock(request, DocumentPart)

        class FakeParent:
            @property
            def part(self):
                return part_

        return Paragraph(p_element, FakeParent())

    def it_returns_a_16_character_hex_string(self, request: pytest.FixtureRequest):
        p = cast(CT_P, element("w:p/w:r/w:t\"hello\""))
        paragraph = self._make_paragraph(request, p)

        result = paragraph.stable_id

        assert isinstance(result, str)
        assert len(result) == 16
        assert all(c in "0123456789abcdef" for c in result)

    def it_returns_the_same_id_on_repeated_access(
        self, request: pytest.FixtureRequest
    ):
        p = cast(CT_P, element("w:p/w:r/w:t\"hello\""))
        paragraph = self._make_paragraph(request, p)
        assert paragraph.stable_id == paragraph.stable_id

    def it_returns_same_id_for_same_position_and_text(
        self, request: pytest.FixtureRequest
    ):
        body_a = element("w:body/w:p/w:r/w:t\"alpha\"")
        body_b = element("w:body/w:p/w:r/w:t\"alpha\"")
        para_a = self._make_paragraph(request, cast(CT_P, body_a[0]))
        para_b = self._make_paragraph(request, cast(CT_P, body_b[0]))
        assert para_a.stable_id == para_b.stable_id

    def it_returns_different_ids_for_different_text(
        self, request: pytest.FixtureRequest
    ):
        p1 = cast(CT_P, element("w:p/w:r/w:t\"alpha\""))
        p2 = cast(CT_P, element("w:p/w:r/w:t\"beta\""))
        para_a = self._make_paragraph(request, p1)
        para_b = self._make_paragraph(request, p2)
        assert para_a.stable_id != para_b.stable_id

    def it_returns_different_ids_for_siblings_with_same_text(
        self, request: pytest.FixtureRequest
    ):
        body = element(
            "w:body/(w:p/w:r/w:t\"same\",w:p/w:r/w:t\"same\")"
        )
        para_a = self._make_paragraph(request, cast(CT_P, body[0]))
        para_b = self._make_paragraph(request, cast(CT_P, body[1]))
        assert para_a.stable_id != para_b.stable_id

    def it_works_when_paragraph_lacks_rsidR(self, request: pytest.FixtureRequest):
        p = cast(CT_P, element("w:p/w:r/w:t\"hello\""))
        paragraph = self._make_paragraph(request, p)
        # --- no rsidR and still returns a 16-char hex ---
        assert paragraph.rsid is None
        assert len(paragraph.stable_id) == 16

    def it_incorporates_rsidR_when_present(self, request: pytest.FixtureRequest):
        p_no_rsid = cast(CT_P, element("w:p/w:r/w:t\"hello\""))
        p_rsid = cast(CT_P, element("w:p{w:rsidR=00FA1B42}/w:r/w:t\"hello\""))
        para_no_rsid = self._make_paragraph(request, p_no_rsid)
        para_rsid = self._make_paragraph(request, p_rsid)
        assert para_no_rsid.stable_id != para_rsid.stable_id

    def it_changes_when_text_is_edited(self, request: pytest.FixtureRequest):
        p = cast(CT_P, element("w:p/w:r/w:t\"original\""))
        paragraph = self._make_paragraph(request, p)
        original_id = paragraph.stable_id
        paragraph.text = "modified"
        assert paragraph.stable_id != original_id
