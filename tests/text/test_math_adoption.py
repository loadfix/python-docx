"""Adoption-layer tests for ``python-ooxml-math`` 0.4 in ``python-docx``.

Covers the three new surfaces introduced when docx rolled onto
``ooxml_math`` 0.4.0:

* :meth:`Paragraph.math_expressions` — generator over the paragraph's
  OMML expressions as :class:`docx.math.MathExpr` proxies.
* :meth:`Document.iter_math_expressions` — document-wide walk yielding
  the same proxies.
* :meth:`Paragraph.add_math` — insert a math block before the first run
  (or append when no runs exist).

The four deferred 0.4.0 proxies exercised here are :class:`Bar`,
:class:`Box`, :class:`Phantom` and :class:`EqArray`; :class:`BorderBox`
and :class:`GroupChar` are smoke-tested via the re-export assertion to
keep the suite compact.
"""

from __future__ import annotations

import io

from docx import Document
from docx.math import (
    Bar,
    BorderBox,
    Box,
    EqArray,
    Fraction,
    GroupChar,
    Lit,
    MathExpr,
    Phantom,
    Var,
    oMath,
)


class DescribeDocxMathReexportsOf04Proxies:
    """docx.math re-exports the six deferred 0.4.0 proxies."""

    def it_exposes_Bar_Box_BorderBox_Phantom_GroupChar_EqArray(self):
        # -- every 0.4.0 proxy subclasses MathExpr --
        assert issubclass(Bar, MathExpr)
        assert issubclass(Box, MathExpr)
        assert issubclass(BorderBox, MathExpr)
        assert issubclass(Phantom, MathExpr)
        assert issubclass(GroupChar, MathExpr)
        assert issubclass(EqArray, MathExpr)

    def it_re_exports_the_same_classes_as_ooxml_math(self):
        import ooxml_math as shared

        assert Bar is shared.Bar
        assert Box is shared.Box
        assert BorderBox is shared.BorderBox
        assert Phantom is shared.Phantom
        assert GroupChar is shared.GroupChar
        assert EqArray is shared.EqArray


class DescribeParagraphAddMath:
    """Paragraph.add_math inserts an OMML block before the first run."""

    def it_appends_to_an_empty_paragraph_when_no_runs_exist(self):
        doc = Document()
        p = doc.add_paragraph()  # -- empty paragraph, no runs --

        proxy = p.add_math(oMath(Fraction(Var("a"), Lit(2))))

        # -- returned proxy wraps a recognised root --
        assert isinstance(proxy, MathExpr)
        # -- single oMath child on the paragraph --
        roots = p._p.xpath(".//m:oMath")
        assert len(roots) == 1

    def it_inserts_before_the_first_run_when_the_paragraph_has_text(self):
        doc = Document()
        p = doc.add_paragraph("trailing text")

        p.add_math(oMath(Bar(Var("x"))))

        # -- the <m:oMath> element must appear before the first <w:r> --
        children = list(p._p)
        oMath_idx = next(
            i for i, c in enumerate(children) if c.tag.endswith("}oMath")
        )
        first_run_idx = next(
            i for i, c in enumerate(children) if c.tag.endswith("}r")
        )
        assert oMath_idx < first_run_idx

    def it_accepts_a_bare_operator_and_auto_wraps_in_oMath(self):
        doc = Document()
        p = doc.add_paragraph()

        # -- Box is a 0.4.0 bare operator; _make_equation_element wraps it --
        p.add_math(Box(Var("y")))

        roots = p._p.xpath(".//m:oMath")
        assert len(roots) == 1


class DescribeParagraphMathExpressions:
    """Paragraph.math_expressions yields a MathExpr for each OMML element."""

    def it_yields_nothing_when_the_paragraph_has_no_equations(self):
        doc = Document()
        p = doc.add_paragraph("just plain text")

        assert list(p.math_expressions) == []

    def it_yields_one_proxy_per_oMath_in_document_order(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_math(oMath(Fraction(Var("a"), Lit(2))))
        p.add_math(oMath(Phantom(Var("b"))))

        found = list(p.math_expressions)
        assert len(found) == 2
        for item in found:
            assert isinstance(item, MathExpr)


class DescribeDocumentIterMathExpressions:
    """Document.iter_math_expressions walks the whole body."""

    def it_walks_every_paragraph_in_the_body(self):
        doc = Document()
        p1 = doc.add_paragraph()
        p1.add_math(oMath(Fraction(Var("x"), Lit(2))))
        p2 = doc.add_paragraph()
        p2.add_math(oMath(EqArray([Var("a"), Var("b")])))

        found = list(doc.iter_math_expressions())
        assert len(found) == 2
        for item in found:
            assert isinstance(item, MathExpr)


class DescribeMathRoundTrip:
    """Math added via add_math survives save + reload + iteration."""

    def it_round_trips_a_bar_through_save_and_reload(self):
        doc = Document()
        p = doc.add_paragraph("trailing")
        p.add_math(oMath(Bar(Var("x"))))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        # -- paragraph-level generator --
        paragraphs = list(reloaded.paragraphs)
        found_para = list(paragraphs[0].math_expressions)
        assert len(found_para) == 1
        assert isinstance(found_para[0], MathExpr)

        # -- document-level generator --
        found_doc = list(reloaded.iter_math_expressions())
        assert len(found_doc) == 1

    def it_round_trips_an_eqarray_with_multiple_rows(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_math(oMath(EqArray([Var("x"), Fraction(Var("y"), Lit(3))])))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        found = list(reloaded.iter_math_expressions())
        assert len(found) == 1
        assert isinstance(found[0], MathExpr)
