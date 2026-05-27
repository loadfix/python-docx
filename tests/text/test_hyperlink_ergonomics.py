"""Unit tests for the hyperlink ergonomics added in 2026.05.12.

Covers issue #69 (``Paragraph.add_link_to`` + ``tooltip`` plumbing on
``Paragraph.add_hyperlink`` / :class:`Hyperlink`) and issue #70
(``Paragraph.add_url`` + ``Paragraph.add_text_with_links``).
"""

from __future__ import annotations

import io
from typing import cast

import pytest

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.parts.story import StoryPart
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import (
    Paragraph,
    _autolink_text,
    _normalise_url_scheme,
)
from docx.text.run import Run

from ..unitutil.cxml import element
from ..unitutil.mock import instance_mock


def _fake_parent_with_part(request, *, with_style: bool = False):
    """Return a fake parent that exposes a mocked StoryPart."""
    story_part_ = instance_mock(request, StoryPart)
    story_part_.relate_to.return_value = "rId42"
    story_part_.get_style_id.return_value = "Hyperlink" if with_style else None

    class FakeParent:
        @property
        def part(self):
            return story_part_

    return FakeParent(), story_part_


class DescribeParagraph_AddHyperlinkTooltip:
    """``add_hyperlink(tooltip=...)`` writes the ``w:tooltip`` attribute (#69)."""

    def it_writes_the_tooltip_attribute_when_provided(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        hyperlink = paragraph.add_hyperlink(
            url="https://example.com", text="Visit", tooltip="Pop me", style=None
        )

        assert hyperlink.tooltip == "Pop me"
        # -- the underlying element carries the attribute --
        assert hyperlink._hyperlink.tooltip == "Pop me"

    def it_omits_the_tooltip_attribute_when_not_provided(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        hyperlink = paragraph.add_hyperlink(url="https://example.com", style=None)

        assert hyperlink.tooltip is None

    def it_round_trips_the_tooltip_through_save_and_load(self):
        d = Document()
        para = d.add_paragraph()
        para.add_hyperlink(
            url="https://example.com", text="hover me",
            tooltip="hover me", style=None,
        )

        buffer = io.BytesIO()
        d.save(buffer)
        buffer.seek(0)
        d2 = Document(buffer)

        loaded = d2.paragraphs[0].hyperlinks[0]
        assert loaded.tooltip == "hover me"
        assert loaded.url == "https://example.com"

    def it_can_assign_and_clear_the_tooltip_via_the_proxy(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)
        hyperlink = paragraph.add_hyperlink(
            url="https://example.com", style=None
        )

        hyperlink.tooltip = "new tooltip"
        assert hyperlink.tooltip == "new tooltip"

        hyperlink.tooltip = None
        assert hyperlink.tooltip is None


class DescribeParagraph_AddLinkTo:
    """``Paragraph.add_link_to(target)`` polymorphic internal links (#69)."""

    def it_links_to_a_string_anchor(self, request: pytest.FixtureRequest):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        hyperlink = paragraph.add_link_to("section1", style=None)

        assert isinstance(hyperlink, Hyperlink)
        assert hyperlink.fragment == "section1"
        assert hyperlink.text == "section1"
        assert hyperlink._hyperlink.rId is None

    def it_links_to_a_bookmark_object(self):
        d = Document()
        para = d.add_paragraph()
        run = para.add_run("anchor here")
        bookmark = d.add_bookmark(run, "site")

        target_para = d.add_paragraph()
        hyperlink = target_para.add_link_to(bookmark, style=None)

        assert hyperlink.fragment == "site"
        # -- text defaults to the bookmark's contained text --
        assert hyperlink.text == "anchor here"

    def it_uses_an_explicit_text_override_with_a_bookmark_target(self):
        d = Document()
        para = d.add_paragraph()
        run = para.add_run("anchor")
        bookmark = d.add_bookmark(run, "site")

        target_para = d.add_paragraph()
        hyperlink = target_para.add_link_to(bookmark, text="See the site", style=None)

        assert hyperlink.text == "See the site"

    def it_links_to_a_heading_paragraph_creating_a_bookmark(self):
        d = Document()
        heading = d.add_heading("Q1 Review")
        para = d.add_paragraph()

        hyperlink = para.add_link_to(heading, style=None)

        # -- a bookmark was auto-allocated on the heading --
        assert hyperlink.fragment.startswith("_link_Q1_Review")
        # -- the heading paragraph now carries that bookmark --
        names = [
            bs.name for bs in heading._p.xpath(".//w:bookmarkStart")
        ]
        assert hyperlink.fragment in names
        # -- visible text defaults to the heading's text --
        assert hyperlink.text == "Q1 Review"

    def it_reuses_the_first_bookmark_already_on_the_heading(self):
        d = Document()
        heading = d.add_heading("Q1 Review")
        # -- pre-attach a bookmark to the heading paragraph --
        heading.add_bookmark("custom_anchor")

        para = d.add_paragraph()
        hyperlink = para.add_link_to(heading, style=None)

        assert hyperlink.fragment == "custom_anchor"

    def it_can_add_a_tooltip_via_add_link_to(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        hyperlink = paragraph.add_link_to(
            "section1", text="See", tooltip="More info", style=None
        )

        assert hyperlink.tooltip == "More info"

    def it_raises_on_a_non_heading_paragraph_target(self):
        d = Document()
        plain = d.add_paragraph("just text")
        target_para = d.add_paragraph()

        with pytest.raises(ValueError, match="must have a style"):
            target_para.add_link_to(plain, style=None)

    def it_raises_on_an_unsupported_target_type(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        with pytest.raises(TypeError, match="Bookmark, Paragraph, or str"):
            paragraph.add_link_to(42)  # type: ignore[arg-type]


class DescribeNormaliseUrlScheme:
    """``_normalise_url_scheme`` auto-prepends mailto: / tel: / http:// (#70)."""

    def it_leaves_an_already_schemed_url_unchanged(self):
        assert _normalise_url_scheme("https://example.com") == "https://example.com"
        assert _normalise_url_scheme("ftp://example.com") == "ftp://example.com"
        assert _normalise_url_scheme("mailto:a@b.com") == "mailto:a@b.com"

    def it_prepends_mailto_for_an_email_shape(self):
        assert _normalise_url_scheme("alice@example.com") == "mailto:alice@example.com"

    def it_prepends_tel_for_a_telephone_shape(self):
        assert _normalise_url_scheme("+1 555 0100") == "tel:+15550100"
        assert _normalise_url_scheme("(555) 010-0100") == "tel:5550100100"

    def it_prepends_http_for_a_www_shortcut(self):
        assert _normalise_url_scheme("www.example.com") == "http://www.example.com"

    def it_returns_a_bare_token_unchanged_when_unrecognised(self):
        assert _normalise_url_scheme("just-text") == "just-text"


class DescribeParagraph_AddUrl:
    """``Paragraph.add_url`` ergonomic wrapper for external links (#70)."""

    def it_appends_an_http_hyperlink(self):
        d = Document()
        para = d.add_paragraph()

        hyperlink = para.add_url("https://example.com", style=None)

        assert isinstance(hyperlink, Hyperlink)
        assert hyperlink.url == "https://example.com"
        # -- text defaults to the original URL string --
        assert hyperlink.text == "https://example.com"

    def it_auto_prepends_mailto_for_an_email_target(self):
        d = Document()
        para = d.add_paragraph()

        hyperlink = para.add_url("alice@example.com", style=None)

        assert hyperlink.url == "mailto:alice@example.com"
        # -- visible text shows the bare email, not the scheme --
        assert hyperlink.text == "alice@example.com"

    def it_auto_prepends_tel_for_a_phone_number(self):
        d = Document()
        para = d.add_paragraph()

        hyperlink = para.add_url("+1 555 0100", style=None)

        assert hyperlink.url == "tel:+15550100"
        assert hyperlink.text == "+1 555 0100"

    def it_writes_a_tooltip_when_provided(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        hyperlink = paragraph.add_url(
            "https://example.com", tooltip="popup", style=None
        )

        assert hyperlink.tooltip == "popup"

    def it_raises_on_an_empty_url(self, request: pytest.FixtureRequest):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        with pytest.raises(ValueError, match="non-empty"):
            paragraph.add_url("", style=None)


class DescribeParagraph_AddTextWithLinks:
    """``Paragraph.add_text_with_links`` — autolink helper (#70)."""

    def it_returns_only_a_run_for_text_with_no_matches(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        pieces = paragraph.add_text_with_links("just plain text", style=None)

        assert len(pieces) == 1
        assert isinstance(pieces[0], Run)
        assert paragraph.text == "just plain text"

    def it_splits_around_a_url_match(self):
        d = Document()
        para = d.add_paragraph()

        pieces = para.add_text_with_links(
            "See https://example.com for info", style=None
        )

        assert [type(x).__name__ for x in pieces] == [
            "Run", "Hyperlink", "Run"
        ]
        # -- pyright happy on the union narrowing --
        link = pieces[1]
        assert isinstance(link, Hyperlink)
        assert link.url == "https://example.com"
        assert para.text == "See https://example.com for info"

    def it_strips_trailing_punctuation_from_a_url(self):
        d = Document()
        para = d.add_paragraph()

        pieces = para.add_text_with_links(
            "Visit https://example.com.", style=None
        )

        link = pieces[1]
        assert isinstance(link, Hyperlink)
        assert link.url == "https://example.com"
        # -- the trailing "." stays in the plain-text run --
        tail = pieces[2]
        assert isinstance(tail, Run)
        assert tail.text == "."

    def it_emits_a_mailto_for_an_email_match(self):
        d = Document()
        para = d.add_paragraph()

        pieces = para.add_text_with_links(
            "Contact alice@example.com today", style=None
        )

        link = pieces[1]
        assert isinstance(link, Hyperlink)
        assert link.url == "mailto:alice@example.com"
        assert link.text == "alice@example.com"

    def it_handles_text_starting_with_a_match(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        pieces = paragraph.add_text_with_links(
            "https://example.com is cool", style=None
        )

        # -- no leading run, just hyperlink + trailing run --
        assert [type(x).__name__ for x in pieces] == ["Hyperlink", "Run"]

    def it_handles_text_ending_with_a_match(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        pieces = paragraph.add_text_with_links(
            "ping alice@example.com", style=None
        )

        assert [type(x).__name__ for x in pieces] == ["Run", "Hyperlink"]

    def it_finds_multiple_matches(self, request: pytest.FixtureRequest):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        pieces = paragraph.add_text_with_links(
            "See https://a.example or email b@example.com", style=None
        )

        kinds = [type(x).__name__ for x in pieces]
        assert kinds == ["Run", "Hyperlink", "Run", "Hyperlink"]

    def it_normalises_a_www_shortcut_with_http(self):
        d = Document()
        para = d.add_paragraph()

        pieces = para.add_text_with_links(
            "Hit www.example.com please", style=None
        )

        link = pieces[1]
        assert isinstance(link, Hyperlink)
        assert link.url == "http://www.example.com"


class DescribeAutolinkText_HelperDirect:
    """Direct-call coverage of the ``_autolink_text`` helper (#70)."""

    def it_returns_an_empty_list_for_an_empty_string(
        self, request: pytest.FixtureRequest
    ):
        parent, _ = _fake_parent_with_part(request)
        p = cast(CT_P, element("w:p"))
        paragraph = Paragraph(p, parent)

        pieces = _autolink_text(paragraph, "", style=None)

        assert pieces == []
