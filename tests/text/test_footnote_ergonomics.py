"""Unit tests for footnote / endnote authoring ergonomics added in 2026.05.13.

Covers issue #42 (``Paragraph.add_footnote`` / ``Paragraph.add_endnote`` and
the ``Footnotes.numbering`` / ``.restart`` shorthand setters).
"""

from __future__ import annotations

import io

import pytest

from docx import Document
from docx.endnotes import Endnote, Endnotes
from docx.enum.text import WD_FOOTNOTE_RESTART, WD_NUMBER_FORMAT
from docx.footnotes import (
    Footnote,
    Footnotes,
    _resolve_numbering,
    _resolve_restart,
)
from docx.oxml.ns import qn


class DescribeParagraph_AddFootnote:
    """`Paragraph.add_footnote(text)` ergonomic authoring API (#42)."""

    def it_appends_a_footnote_reference_to_the_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("AWS launched 100+ new services in 2025")

        footnote = para.add_footnote("AWS Annual Review 2026, p.42")

        assert isinstance(footnote, Footnote)
        # -- the paragraph gained a trailing run carrying the reference --
        ref_runs = para._p.xpath(".//w:footnoteReference")
        assert len(ref_runs) == 1
        assert ref_runs[0].get(qn("w:id")) == str(footnote.footnote_id)
        # -- the source run is styled FootnoteReference --
        anchor_r = ref_runs[0].getparent()
        assert anchor_r is not None
        assert anchor_r.style == "FootnoteReference"

    def it_seeds_the_footnote_body_with_text_when_provided(self):
        doc = Document()
        para = doc.add_paragraph("Body text")

        footnote = para.add_footnote("Source citation here.")

        # -- the footnote's first paragraph contains the citation text --
        assert footnote.text == "Source citation here."
        # -- and uses the FootnoteText style --
        assert footnote.paragraphs[0]._p.style == "FootnoteText"

    def it_returns_an_empty_footnote_when_text_is_omitted(self):
        doc = Document()
        para = doc.add_paragraph("Body")

        footnote = para.add_footnote()

        assert isinstance(footnote, Footnote)
        assert footnote.text == ""

    def it_assigns_unique_ids_to_successive_footnotes(self):
        doc = Document()
        p1 = doc.add_paragraph("p1")
        p2 = doc.add_paragraph("p2")

        f1 = p1.add_footnote("first")
        f2 = p2.add_footnote("second")

        assert f1.footnote_id != f2.footnote_id
        # -- both are reachable from the document's collection --
        ids = {fn.footnote_id for fn in doc.footnotes}
        assert {f1.footnote_id, f2.footnote_id} <= ids

    def it_round_trips_through_save_and_reload(self):
        doc = Document()
        para = doc.add_paragraph("The quick brown fox.")
        para.add_footnote("Aesop's Fables, see also Reynard.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        # -- the footnote text survives the round-trip --
        assert any(
            fn.text == "Aesop's Fables, see also Reynard."
            for fn in reloaded.footnotes
        )
        # -- the reference is still in the body --
        body_xml = reloaded._body._body.xml  # pyright: ignore[reportPrivateUsage]
        assert "w:footnoteReference" in body_xml

    def it_refuses_to_nest_a_footnote_inside_another_footnote(self):
        doc = Document()
        para = doc.add_paragraph("body")
        footnote = para.add_footnote("outer")
        # -- attempting to add a footnote inside a footnote's own paragraph --
        inner_para = footnote.paragraphs[0]

        with pytest.raises(RuntimeError, match="footnote/endnote paragraphs"):
            inner_para.add_footnote("inner")

    def it_does_not_disturb_existing_runs(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello ")
        para.add_run("world.")

        para.add_footnote("citation")

        # -- the original two runs survive plus a third reference run --
        runs = para._p.xpath("./w:r")
        assert len(runs) == 3
        assert runs[0].text == "Hello "
        assert runs[1].text == "world."
        assert runs[2].xpath("./w:footnoteReference")


class DescribeParagraph_AddEndnote:
    """`Paragraph.add_endnote(text)` ergonomic authoring API (#42)."""

    def it_appends_an_endnote_reference_to_the_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("Body text")

        endnote = para.add_endnote("Source: AWS press release 2026-01-15")

        assert isinstance(endnote, Endnote)
        ref_runs = para._p.xpath(".//w:endnoteReference")
        assert len(ref_runs) == 1
        assert ref_runs[0].get(qn("w:id")) == str(endnote.endnote_id)
        anchor_r = ref_runs[0].getparent()
        assert anchor_r is not None
        assert anchor_r.style == "EndnoteReference"

    def it_seeds_the_endnote_body_with_text(self):
        doc = Document()
        para = doc.add_paragraph("Body")

        endnote = para.add_endnote("End-of-document note text")

        assert endnote.text == "End-of-document note text"
        assert endnote.paragraphs[0]._p.style == "EndnoteText"

    def it_round_trips_through_save_and_reload(self):
        doc = Document()
        doc.add_paragraph("Body").add_endnote("see appendix")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        assert any(en.text == "see appendix" for en in reloaded.endnotes)


class DescribeFootnotesNumberingShorthand:
    """`Footnotes.numbering` and `.restart` shorthand setters (#42)."""

    def it_resolves_arabic_shorthand(self):
        assert _resolve_numbering("1, 2, 3") is WD_NUMBER_FORMAT.DECIMAL
        assert _resolve_numbering("arabic") is WD_NUMBER_FORMAT.DECIMAL

    def it_resolves_roman_shorthand(self):
        assert _resolve_numbering("i, ii, iii") is WD_NUMBER_FORMAT.LOWER_ROMAN

    def it_resolves_chicago_symbols_shorthand(self):
        assert (
            _resolve_numbering("*, dagger, double-dagger")
            is WD_NUMBER_FORMAT.CHICAGO
        )
        assert _resolve_numbering("chicago") is WD_NUMBER_FORMAT.CHICAGO

    def it_passes_enum_members_through_unchanged(self):
        assert (
            _resolve_numbering(WD_NUMBER_FORMAT.UPPER_ROMAN)
            is WD_NUMBER_FORMAT.UPPER_ROMAN
        )

    def it_accepts_raw_ooxml_tokens(self):
        assert _resolve_numbering("upperRoman") is WD_NUMBER_FORMAT.UPPER_ROMAN

    def it_raises_on_unknown_shorthand(self):
        with pytest.raises(ValueError, match="unrecognised numbering shorthand"):
            _resolve_numbering("not-a-format")

    def it_resolves_restart_shorthand(self):
        assert _resolve_restart("section") is WD_FOOTNOTE_RESTART.EACH_SECTION
        assert _resolve_restart("page") is WD_FOOTNOTE_RESTART.EACH_PAGE
        assert _resolve_restart("continuous") is WD_FOOTNOTE_RESTART.CONTINUOUS

    def it_writes_numbering_through_to_settings(self):
        doc = Document()
        # -- precondition: no footnote properties yet --
        assert doc.footnote_properties is None

        doc.footnotes.numbering = "i, ii, iii"

        props = doc.footnote_properties
        assert props is not None
        assert props.number_format is WD_NUMBER_FORMAT.LOWER_ROMAN

    def it_writes_restart_through_to_settings(self):
        doc = Document()

        doc.footnotes.restart = "section"

        assert doc.footnote_properties is not None
        assert (
            doc.footnote_properties.restart_rule
            is WD_FOOTNOTE_RESTART.EACH_SECTION
        )

    def it_clears_numbering_when_set_to_None(self):
        doc = Document()
        doc.footnotes.numbering = "arabic"
        assert doc.footnotes.numbering is WD_NUMBER_FORMAT.DECIMAL

        doc.footnotes.numbering = None

        # -- the footnotePr remains, but numFmt is gone --
        assert doc.footnotes.numbering is None

    def it_reads_numbering_back_after_setting(self):
        doc = Document()
        doc.footnotes.numbering = WD_NUMBER_FORMAT.UPPER_LETTER

        assert doc.footnotes.numbering is WD_NUMBER_FORMAT.UPPER_LETTER

    def it_round_trips_numbering_through_save_and_reload(self):
        doc = Document()
        doc.footnotes.numbering = "i, ii, iii"
        doc.footnotes.restart = "section"

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        assert reloaded.footnotes.numbering is WD_NUMBER_FORMAT.LOWER_ROMAN
        assert reloaded.footnotes.restart is WD_FOOTNOTE_RESTART.EACH_SECTION


class DescribeEndnotesNumberingShorthand:
    """`Endnotes.numbering` and `.restart` shorthand setters (#42)."""

    def it_writes_endnote_numbering_to_settings(self):
        doc = Document()

        doc.endnotes.numbering = "*, dagger, double-dagger"

        assert doc.endnote_properties is not None
        assert doc.endnote_properties.number_format is WD_NUMBER_FORMAT.CHICAGO

    def it_writes_endnote_restart_to_settings(self):
        doc = Document()

        doc.endnotes.restart = "continuous"

        assert doc.endnote_properties is not None
        assert (
            doc.endnote_properties.restart_rule
            is WD_FOOTNOTE_RESTART.CONTINUOUS
        )

    def it_clears_endnote_numbering_when_set_to_None(self):
        doc = Document()
        doc.endnotes.numbering = "arabic"

        doc.endnotes.numbering = None

        assert doc.endnotes.numbering is None

    def it_round_trips_endnote_numbering(self):
        doc = Document()
        doc.endnotes.numbering = "lowerLetter"

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        assert reloaded.endnotes.numbering is WD_NUMBER_FORMAT.LOWER_LETTER


class DescribeMixedFootnotesAndEndnotes:
    """End-to-end: a multi-paragraph doc with mixed footnotes + endnotes."""

    def it_supports_a_three_page_doc_with_mixed_notes(self):
        doc = Document()
        # -- three "pages" worth of paragraphs with mixed annotations --
        p1 = doc.add_paragraph("First page body. ")
        p1.add_footnote("Footnote on page 1")
        p2 = doc.add_paragraph("Second page body. ")
        p2.add_endnote("Endnote referenced on page 2")
        p3 = doc.add_paragraph("Third page body. ")
        p3.add_footnote("Footnote on page 3")
        p3.add_endnote("Second endnote on page 3")

        # -- save / reload and verify everything survives --
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        footnote_texts = sorted(fn.text for fn in reloaded.footnotes)
        endnote_texts = sorted(en.text for en in reloaded.endnotes)
        assert footnote_texts == ["Footnote on page 1", "Footnote on page 3"]
        assert endnote_texts == [
            "Endnote referenced on page 2",
            "Second endnote on page 3",
        ]
        # -- two footnoteReference + two endnoteReference elements in body XML --
        body_xml = reloaded._body._body.xml  # pyright: ignore[reportPrivateUsage]
        assert body_xml.count("w:footnoteReference") == 2
        assert body_xml.count("w:endnoteReference") == 2
