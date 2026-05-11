"""Tests for the minimal docx → HTML exporter.

Covers the element-kind mappings declared in R10-11:
- paragraph → <p> (with heading promotion)
- run → <span> / <strong> / <em> / <u>
- hyperlink → <a href>
- table → <table> / <tr> / <td>
- inline picture → <img src=...>
- list → <ol> / <ul> with <li>
- XSS text escaping
- unsupported-element comment
"""

from __future__ import annotations

import base64
import io

import pytest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


# -- helpers -----------------------------------------------------------------


def _new_doc():
    """Return a fresh blank |Document| for use in a single test."""
    return Document()


# -- paragraph & run --------------------------------------------------------


class DescribeDocumentToHtml:
    def it_wraps_a_paragraph_in_a_p_tag(self):
        doc = _new_doc()
        doc.add_paragraph("Hello world")
        html = doc.to_html(include_styles=False)
        assert "<p>Hello world</p>" in html
        assert html.startswith("<!DOCTYPE html>")
        assert "<meta charset=\"utf-8\">" in html

    def it_promotes_bold_only_run_to_strong(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("bold bit")
        run.bold = True
        html = doc.to_html(include_styles=False)
        assert "<strong>bold bit</strong>" in html

    def it_promotes_italic_only_run_to_em(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("italicized")
        run.italic = True
        html = doc.to_html(include_styles=False)
        assert "<em>italicized</em>" in html

    def it_promotes_underline_only_run_to_u(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("underlined")
        run.underline = True
        html = doc.to_html(include_styles=False)
        assert "<u>underlined</u>" in html

    def it_nests_tags_when_multiple_formats_applied(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("combo")
        run.bold = True
        run.italic = True
        html = doc.to_html(include_styles=False)
        # -- nested <strong><em>…</em></strong> (outer is bold) --
        assert "<strong>" in html and "<em>combo</em>" in html

    def it_renders_alignment_as_inline_css(self):
        doc = _new_doc()
        p = doc.add_paragraph("centered")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        html = doc.to_html(include_styles=False)
        assert 'style="text-align:center"' in html


# -- headings ----------------------------------------------------------------


class DescribeHeadingExport:
    def it_promotes_Heading_1_paragraph_to_h1(self):
        doc = _new_doc()
        doc.add_heading("Top Heading", level=1)
        html = doc.to_html(include_styles=False)
        assert "<h1>Top Heading</h1>" in html

    def it_promotes_Heading_3_paragraph_to_h3(self):
        doc = _new_doc()
        doc.add_heading("Deeper", level=3)
        html = doc.to_html(include_styles=False)
        assert "<h3>Deeper</h3>" in html

    def it_caps_heading_levels_above_six_at_h6(self):
        doc = _new_doc()
        doc.add_heading("Deep", level=9)
        html = doc.to_html(include_styles=False)
        assert "<h6>Deep</h6>" in html


# -- hyperlinks --------------------------------------------------------------


class DescribeHyperlinkExport:
    def it_renders_a_hyperlink_as_an_a_tag(self):
        doc = _new_doc()
        p = doc.add_paragraph("before ")
        # -- pass style=None so we don't require the "Hyperlink" character
        # -- style (not present in the default template) --
        p.add_hyperlink("https://example.com/path", "click me", style=None)
        p.add_run(" after")
        html = doc.to_html(include_styles=False)
        assert '<a href="https://example.com/path">' in html
        assert ">click me</a>" in html


# -- tables ------------------------------------------------------------------


class DescribeTableExport:
    def it_renders_a_table_as_table_tr_td(self):
        doc = _new_doc()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A1"
        table.rows[0].cells[1].text = "A2"
        table.rows[1].cells[0].text = "B1"
        table.rows[1].cells[1].text = "B2"
        html = doc.to_html(include_styles=False)
        assert "<table" in html and "border-collapse:collapse" in html
        assert "<tr>" in html
        assert "A1" in html and "B2" in html
        assert html.count("<td") == 4

    def it_includes_cell_borders_as_inline_css(self):
        doc = _new_doc()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "x"
        html = doc.to_html(include_styles=False)
        assert "border:1px solid #000" in html


# -- images ------------------------------------------------------------------


class DescribeInlineImageExport:
    def it_embeds_images_as_base64_data_urls_by_default(self, tmp_path):
        # -- minimal 1x1 PNG so python-docx can ingest it --
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
            "+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        img_path = tmp_path / "pixel.png"
        img_path.write_bytes(png_bytes)

        doc = _new_doc()
        doc.add_picture(str(img_path))
        html = doc.to_html(include_styles=False, embed_images=True)
        assert '<img src="data:image/png;base64,' in html

    def it_emits_cid_placeholders_when_embed_images_false(self, tmp_path):
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
            "+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        img_path = tmp_path / "pixel.png"
        img_path.write_bytes(png_bytes)

        doc = _new_doc()
        doc.add_picture(str(img_path))
        html = doc.to_html(include_styles=False, embed_images=False)
        assert '<img src="cid:rId' in html


# -- lists -------------------------------------------------------------------


class DescribeListExport:
    def it_renders_decimal_numbering_as_ol(self):
        doc = _new_doc()
        defn = doc.numbering.add_abstract_definition(format="decimal")
        p1 = doc.add_paragraph("first")
        p2 = doc.add_paragraph("second")
        defn.apply_to(p1, level=0)
        defn.apply_to(p2, level=0)
        html = doc.to_html(include_styles=False)
        assert "<ol>" in html
        assert "<li>first</li>" in html
        assert "<li>second</li>" in html
        assert "</ol>" in html

    def it_renders_bullet_numbering_as_ul(self):
        doc = _new_doc()
        defn = doc.numbering.add_abstract_definition(format="bullet", lvl_text="•")
        p1 = doc.add_paragraph("a")
        p2 = doc.add_paragraph("b")
        defn.apply_to(p1, level=0)
        defn.apply_to(p2, level=0)
        html = doc.to_html(include_styles=False)
        assert "<ul>" in html
        assert "<li>a</li>" in html
        assert "</ul>" in html


# -- XSS escape --------------------------------------------------------------


class DescribeXssEscape:
    def it_escapes_script_tags_in_text_content(self):
        doc = _new_doc()
        doc.add_paragraph("<script>alert('xss')</script>")
        html = doc.to_html(include_styles=False)
        # -- the raw tag must not appear; escaped form must --
        assert "<script>" not in html.split("<body>")[1]
        assert "&lt;script&gt;" in html

    def it_escapes_quotes_in_hyperlink_urls(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_hyperlink('https://example.com/?q="malicious"', "x", style=None)
        html = doc.to_html(include_styles=False)
        assert '"malicious"' not in html  # raw quotes must be escaped
        assert "&quot;malicious&quot;" in html


# -- scheme allow-list on hyperlinks ----------------------------------------


class DescribeHyperlinkSchemeFiltering:
    """Unsafe URL schemes in hyperlinks are rewritten to ``"#"`` so
    ``Document.to_html()`` cannot carry stored XSS into web renderers."""

    @pytest.mark.parametrize(
        "malicious_url",
        [
            "javascript:alert(1)",
            "JaVaScRiPt:alert(1)",
            "data:text/html,<script>alert(1)</script>",
            "vbscript:msgbox(1)",
            "file:///etc/passwd",
            "jar:http://evil.com/x.jar!/",
        ],
    )
    def it_rewrites_unsafe_schemes_to_hash(self, malicious_url):
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_hyperlink(malicious_url, "click", style=None)
        html = doc.to_html(include_styles=False)
        # -- the payload scheme must not appear in the final href --
        assert "javascript:" not in html.lower()
        assert "data:text/html" not in html.lower()
        assert "vbscript:" not in html.lower()
        assert "file://" not in html.lower()
        assert "jar:" not in html.lower()
        # -- hyperlink text must still render (anchor is preserved) --
        assert "click" in html

    @pytest.mark.parametrize(
        "safe_url",
        [
            "http://ok.example",
            "https://ok.example/path?x=1",
            "mailto:x@y.z",
            "#anchor",
            "",
        ],
    )
    def it_preserves_safe_urls(self, safe_url):
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_hyperlink(safe_url, "click", style=None)
        html = doc.to_html(include_styles=False)
        if safe_url:
            # -- html.escape may encode the anchor body but the substring
            # -- stays intact; "http" / "https" / "mailto" / "#" survive. --
            if safe_url.startswith("#"):
                assert 'href="#anchor"' in html
            elif safe_url.startswith("mailto:"):
                assert 'href="mailto:x@y.z"' in html
            else:
                assert f'href="{safe_url}"' in html

    def it_exposes_the_sanitizer_as_a_module_level_helper(self):
        # -- callers may want to reuse the allow-list without touching the
        # -- exporter plumbing. Pin the public-ish helper name. --
        from docx.html_export import _sanitize_href

        assert _sanitize_href("javascript:alert(1)") == "#"
        assert _sanitize_href("https://ok") == "https://ok"
        assert _sanitize_href("#a") == "#a"
        assert _sanitize_href("") == ""
        assert _sanitize_href("/relative/path") == "/relative/path"
        assert _sanitize_href("no-scheme-just-text") == "no-scheme-just-text"


# -- style block -------------------------------------------------------------


class DescribeStyleBlock:
    def it_emits_a_style_block_when_include_styles_true(self):
        doc = _new_doc()
        doc.add_paragraph("p")
        html = doc.to_html(include_styles=True)
        assert "<style>" in html

    def it_omits_style_block_when_include_styles_false(self):
        doc = _new_doc()
        doc.add_paragraph("p")
        html = doc.to_html(include_styles=False)
        assert "<style>" not in html


# -- unsupported element comment --------------------------------------------


class DescribeUnsupportedElement:
    def it_emits_a_comment_for_a_simple_field(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_simple_field("PAGE", "1")
        html = doc.to_html(include_styles=False)
        assert "<!-- unsupported: w:fldSimple -->" in html
