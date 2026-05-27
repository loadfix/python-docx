"""Unit-test suite for the fluent chainable builder API (issue #77).

The fluent surface adds three things to python-docx:

* ``Document.h1``..``Document.h6`` and ``Document.p`` shortcuts that
  thinly wrap :meth:`Document.add_heading` / :meth:`Document.add_paragraph`
  and return the freshly-appended :class:`~docx.text.paragraph.Paragraph`.
* ``Paragraph.bold()``, ``Paragraph.italic()``, ``Paragraph.underline()``,
  ``Paragraph.align()``, and ``Paragraph.color()`` — each mutates the
  paragraph (or its runs) via the existing verbose surface and returns
  ``self`` for chaining.
* ``Run.color()`` — same shape on the :class:`~docx.text.run.Run`
  proxy.

The verbose surface (``run.bold = True``, ``paragraph.alignment =
WD_ALIGN_PARAGRAPH.CENTER``, ``run.font.color.rgb = ...``) remains the
ground truth; the fluent layer is sugar that calls into it. These
tests therefore exercise the fluent shape end-to-end against a fresh
``Document()`` so any drift between the sugar and the underlying
verbose surface fails the suite immediately.
"""

from __future__ import annotations

import pytest

from docx import Document as OpenDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run


class DescribeDocumentFluentShortcuts:
    """Document.h1..h6 / Document.p shortcuts."""

    @pytest.mark.parametrize(
        ("level", "method_name", "expected_style"),
        [
            (1, "h1", "Heading 1"),
            (2, "h2", "Heading 2"),
            (3, "h3", "Heading 3"),
            (4, "h4", "Heading 4"),
            (5, "h5", "Heading 5"),
            (6, "h6", "Heading 6"),
        ],
    )
    def it_appends_a_heading_paragraph_with_the_matching_style(
        self, level: int, method_name: str, expected_style: str
    ):
        document = OpenDocument()
        method = getattr(document, method_name)

        paragraph = method("hello world")

        assert isinstance(paragraph, Paragraph)
        assert paragraph.text == "hello world"
        assert paragraph.style.name == expected_style

    def it_returns_the_appended_paragraph_so_callers_can_chain(self):
        document = OpenDocument()

        result = document.h2("Section A").bold().align("center")

        assert isinstance(result, Paragraph)
        assert result.text == "Section A"
        # -- bold + alignment landed on the verbose surface --
        assert result.runs[0].bold is True
        assert result.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_supports_an_empty_text_argument_on_each_heading_shortcut(self):
        document = OpenDocument()
        for method_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            paragraph = getattr(document, method_name)()
            assert isinstance(paragraph, Paragraph)
            assert paragraph.text == ""

    def it_appends_a_body_paragraph_via_p(self):
        document = OpenDocument()

        paragraph = document.p("just a body paragraph")

        assert isinstance(paragraph, Paragraph)
        assert paragraph.text == "just a body paragraph"
        # -- no Heading style applied; .style.name comes from the default --
        assert paragraph.style.name in ("Normal", "Default Paragraph Font") or (
            paragraph.style.name and not paragraph.style.name.startswith("Heading")
        )

    def it_appends_an_empty_body_paragraph_when_called_with_no_args(self):
        document = OpenDocument()

        paragraph = document.p()

        assert isinstance(paragraph, Paragraph)
        assert paragraph.text == ""

    def it_supports_the_full_chain_documented_in_the_issue(self):
        document = OpenDocument()

        result = (
            document.h1("Q1 Review")
            .bold()
            .align("center")
        )
        document.p("Revenue grew 8.7% YoY").bold().align("center")

        # -- the heading came back as a Paragraph --
        assert isinstance(result, Paragraph)
        # -- two paragraphs landed on the body --
        bodies = [p for p in document.paragraphs if p.text]
        assert [p.text for p in bodies] == ["Q1 Review", "Revenue grew 8.7% YoY"]
        # -- both inherited bold + center alignment --
        for p in bodies:
            assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
            assert all(r.bold is True for r in p.runs)


class DescribeParagraphFluentHelpers:
    """Paragraph.bold / italic / underline / align / color."""

    def it_returns_self_from_bold_for_chaining(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        result = paragraph.bold()

        assert result is paragraph

    def it_sets_bold_on_every_run_when_paragraph_has_runs(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("first")
        paragraph.add_run(" second")
        paragraph.add_run(" third")

        paragraph.bold()

        assert [r.bold for r in paragraph.runs] == [True, True, True]

    def it_can_disable_bold_when_called_with_false(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        paragraph.bold(False)

        assert paragraph.runs[0].bold is False

    def it_falls_through_to_pPr_rPr_when_paragraph_has_no_runs(self):
        document = OpenDocument()
        paragraph = document.add_paragraph()
        # -- no runs added; bold() should land on the paragraph mark --
        assert paragraph.runs == []

        paragraph.bold()

        assert paragraph.font.bold is True

    def it_returns_self_from_italic_for_chaining(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        result = paragraph.italic()

        assert result is paragraph
        assert paragraph.runs[0].italic is True

    def it_returns_self_from_underline_for_chaining(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        result = paragraph.underline()

        assert result is paragraph
        assert paragraph.runs[0].underline is True

    @pytest.mark.parametrize(
        ("alias", "expected"),
        [
            ("left", WD_ALIGN_PARAGRAPH.LEFT),
            ("center", WD_ALIGN_PARAGRAPH.CENTER),
            ("centre", WD_ALIGN_PARAGRAPH.CENTER),
            ("right", WD_ALIGN_PARAGRAPH.RIGHT),
            ("justify", WD_ALIGN_PARAGRAPH.JUSTIFY),
            ("both", WD_ALIGN_PARAGRAPH.JUSTIFY),
            ("CENTER", WD_ALIGN_PARAGRAPH.CENTER),
            ("Distribute", WD_ALIGN_PARAGRAPH.DISTRIBUTE),
        ],
    )
    def it_maps_string_aliases_to_WD_ALIGN_PARAGRAPH_members(
        self, alias: str, expected: WD_ALIGN_PARAGRAPH
    ):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        result = paragraph.align(alias)

        assert result is paragraph
        assert paragraph.alignment == expected

    def it_accepts_a_WD_ALIGN_PARAGRAPH_member_directly(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        paragraph.align(WD_ALIGN_PARAGRAPH.RIGHT)

        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def it_raises_ValueError_on_an_unknown_alignment_string(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        with pytest.raises(ValueError, match="unknown alignment"):
            paragraph.align("sideways")

    def it_raises_TypeError_on_a_non_string_non_enum_alignment(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        with pytest.raises(TypeError, match="WD_PARAGRAPH_ALIGNMENT"):
            paragraph.align(42)  # type: ignore[arg-type]

    def it_sets_color_on_every_run_from_a_hex_string(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("first")
        paragraph.add_run(" second")

        result = paragraph.color("#FF8800")

        assert result is paragraph
        for run in paragraph.runs:
            assert run.font.color.rgb == RGBColor(0xFF, 0x88, 0x00)

    def it_accepts_color_strings_without_a_leading_hash(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        paragraph.color("3C2F80")

        assert paragraph.runs[0].font.color.rgb == RGBColor(0x3C, 0x2F, 0x80)

    def it_accepts_an_RGBColor_instance(self):
        document = OpenDocument()
        paragraph = document.add_paragraph("hello")

        paragraph.color(RGBColor(0x10, 0x20, 0x30))

        assert paragraph.runs[0].font.color.rgb == RGBColor(0x10, 0x20, 0x30)

    def it_lands_color_on_the_paragraph_mark_when_no_runs_are_present(self):
        document = OpenDocument()
        paragraph = document.add_paragraph()
        assert paragraph.runs == []

        paragraph.color("#112233")

        assert paragraph.font.color.rgb == RGBColor(0x11, 0x22, 0x33)


class DescribeRunFluentHelpers:
    """Run.color() — the only fluent run-level helper that doesn't shadow
    the existing tri-state ``bold`` / ``italic`` / ``underline``
    properties on :class:`~docx.text.run.Run`. Run-level chainable bold
    / italic / underline are intentionally omitted; chain through the
    paragraph or use the existing verbose form ``run.font.bold = True``."""

    def it_returns_self_from_color_for_chaining(self):
        document = OpenDocument()
        run = document.add_paragraph("hello").runs[0]

        result = run.color("#0066CC")

        assert result is run
        assert run.font.color.rgb == RGBColor(0x00, 0x66, 0xCC)

    def it_accepts_the_3_char_short_hex_form(self):
        document = OpenDocument()
        run = document.add_paragraph("hello").runs[0]

        run.color("#F0A")

        assert run.font.color.rgb == RGBColor(0xFF, 0x00, 0xAA)

    def it_accepts_an_RGBColor_instance_directly(self):
        document = OpenDocument()
        run = document.add_paragraph("hello").runs[0]

        run.color(RGBColor(0x80, 0x80, 0x80))

        assert run.font.color.rgb == RGBColor(0x80, 0x80, 0x80)

    def it_round_trips_through_save_and_reopen(self):
        import io

        document = OpenDocument()
        document.h1("Q1 Review").bold().align("center").color("#3366FF")
        document.p("Revenue grew 8.7% YoY").italic().color("AABBCC")

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = OpenDocument(buf)

        body = [p for p in reopened.paragraphs if p.text]
        assert [p.text for p in body] == ["Q1 Review", "Revenue grew 8.7% YoY"]
        assert body[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert body[0].runs[0].bold is True
        assert body[0].runs[0].font.color.rgb == RGBColor(0x33, 0x66, 0xFF)
        assert body[1].runs[0].italic is True
        assert body[1].runs[0].font.color.rgb == RGBColor(0xAA, 0xBB, 0xCC)
