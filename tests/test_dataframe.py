# pyright: reportPrivateUsage=false

"""Unit tests for :meth:`docx.document.Document.add_dataframe` (issue #40).

Covers:

* Five DataFrame shapes (small numeric, large numeric, mixed types,
  date columns, NaN-containing).
* All four built-in styles (``executive`` / ``minimal`` / ``boxed`` /
  ``striped``).
* Number-format DSL (Python format-specs plus the date-token DSL).
* Total-row aggregators (``sum`` / ``mean`` / ``count`` / ``none`` /
  per-column mapping).
* The "pandas missing" code path — exercised by patching
  ``_is_dataframe`` and ``_require_pandas``.
* Save / reload byte round-trip.
"""

from __future__ import annotations

import datetime as dt
import io

import pytest

from docx import Document as OpenDocument
from docx.dataframe import (
    _BUILTIN_STYLES,
    _aggregate,
    _date_spec_to_strftime,
    _format_value,
    _looks_like_date_format,
    _resolve_alignment,
    _resolve_style,
    _resolve_total_spec,
)
from docx.shared import RGBColor

pd = pytest.importorskip("pandas")


# ---------------------------------------------------------------------------
# Five DataFrame fixtures (issue #40 acceptance: five shapes)
# ---------------------------------------------------------------------------


def _small_numeric() -> "pd.DataFrame":
    return pd.DataFrame(
        {
            "Region": ["AMER", "APAC", "EMEA"],
            "Revenue": [1234.5, 987.6, 654.3],
        }
    )


def _large_numeric(n: int = 50) -> "pd.DataFrame":
    return pd.DataFrame(
        {
            "i": list(range(n)),
            "x": [float(i) * 1.5 for i in range(n)],
            "y": [float(i) ** 0.5 for i in range(n)],
        }
    )


def _mixed_types() -> "pd.DataFrame":
    return pd.DataFrame(
        {
            "Region": ["AMER", "APAC", "EMEA", "LATAM"],
            "Revenue": [1234.5, 987.6, 654.3, 321.0],
            "Growth": [0.087, 0.121, -0.034, 0.05],
            "Active": [True, True, False, True],
            "Headcount": [42, 17, 31, 9],
        }
    )


def _date_columns() -> "pd.DataFrame":
    dates = pd.to_datetime(["2024-01-15", "2024-02-15", "2024-03-15"])
    return pd.DataFrame(
        {
            "Period": dates,
            "Bookings": [100.0, 150.0, 175.0],
        }
    )


def _with_nans() -> "pd.DataFrame":
    return pd.DataFrame(
        {
            "Region": ["AMER", "APAC", "EMEA"],
            "Revenue": [1234.5, float("nan"), 654.3],
            "Growth": [0.087, 0.121, float("nan")],
        }
    )


# ---------------------------------------------------------------------------
# Number-format DSL
# ---------------------------------------------------------------------------


class Describe_format_value:
    def it_renders_a_dollar_amount(self):
        assert _format_value(1234.5, "$,.1f") == "$1,234.5"

    def it_renders_a_percentage(self):
        assert _format_value(0.087, ".1%") == "8.7%"

    def it_renders_an_integer_with_a_thousands_separator(self):
        assert _format_value(12345, ",d") == "12,345"

    def it_blanks_a_nan_value(self):
        assert _format_value(float("nan"), ".1f") == ""

    def it_blanks_a_pandas_NaT(self):
        assert _format_value(pd.NaT, "MMM YYYY") == ""

    def it_blanks_None(self):
        assert _format_value(None, "$,.1f") == ""

    def it_falls_back_to_str_for_non_numeric(self):
        assert _format_value("AMER", "$,.1f") == "AMER"

    def it_renders_a_date_with_the_DSL(self):
        d = dt.date(2024, 5, 28)
        assert _format_value(d, "MMM YYYY") == "May 2024"

    def it_renders_a_full_datetime(self):
        x = dt.datetime(2024, 5, 28, 14, 30, 45)
        assert _format_value(x, "YYYY-MM-DD HH:mm:ss") == "2024-05-28 14:30:45"

    def it_renders_a_pandas_Timestamp(self):
        ts = pd.Timestamp("2024-05-28")
        assert _format_value(ts, "MMMM YYYY") == "May 2024"

    def it_uses_default_iso_date_when_no_spec(self):
        assert _format_value(dt.date(2024, 5, 28), None) == "2024-05-28"

    def it_renders_an_int_float_default(self):
        assert _format_value(42.0, None) == "42"


class Describe_date_dsl:
    def it_translates_the_full_date_DSL(self):
        assert _date_spec_to_strftime("MMMM YYYY") == "%B %Y"

    def it_handles_collisions_between_MMM_and_MMMM(self):
        assert _date_spec_to_strftime("MMM YYYY") == "%b %Y"

    def it_translates_HH_mm_ss(self):
        assert _date_spec_to_strftime("HH:mm:ss") == "%H:%M:%S"

    def it_recognises_a_date_format(self):
        assert _looks_like_date_format("MMM YYYY")
        assert _looks_like_date_format("YYYY")
        assert not _looks_like_date_format(",.1f")
        assert not _looks_like_date_format("$,.1f")


# ---------------------------------------------------------------------------
# Aggregation
# ---------------------------------------------------------------------------


class Describe_aggregate:
    def it_sums_numeric_values(self):
        assert _aggregate([1, 2, 3], "sum") == 6

    def it_means_numeric_values(self):
        assert _aggregate([1, 2, 3, 4], "mean") == 2.5

    def it_counts_non_null_values(self):
        assert _aggregate([1, 2, None, 4], "count") == 3

    def it_returns_blank_for_none(self):
        assert _aggregate([1, 2, 3], "none") == ""

    def it_skips_NaN_in_sum(self):
        assert _aggregate([1.0, float("nan"), 3.0], "sum") == 4.0

    def it_returns_blank_for_all_nan(self):
        assert _aggregate([float("nan")], "sum") == ""

    def it_raises_on_unknown_op(self):
        with pytest.raises(ValueError):
            _aggregate([1, 2], "median")


# ---------------------------------------------------------------------------
# Style resolution
# ---------------------------------------------------------------------------


class Describe_resolve_style:
    @pytest.mark.parametrize("name", _BUILTIN_STYLES)
    def it_resolves_each_built_in_style(self, name):
        spec = _resolve_style(name)
        assert spec.name == name

    def it_raises_on_unknown_style(self):
        with pytest.raises(ValueError):
            _resolve_style("rainbow")

    def it_executive_has_a_header_fill(self):
        assert _resolve_style("executive").header_fill is not None

    def it_minimal_has_no_fills(self):
        spec = _resolve_style("minimal")
        assert spec.header_fill is None
        assert spec.alt_row_fill is None
        assert spec.monospace_numbers
        assert spec.header_underline_only

    def it_boxed_has_full_grid(self):
        assert _resolve_style("boxed").border == "all"

    def it_striped_has_no_borders_but_alt_rows(self):
        spec = _resolve_style("striped")
        assert spec.border == "none"
        assert spec.alt_row_fill is not None


class Describe_resolve_alignment:
    def it_accepts_left(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert _resolve_alignment("left") == WD_ALIGN_PARAGRAPH.LEFT

    def it_accepts_right(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert _resolve_alignment("right") == WD_ALIGN_PARAGRAPH.RIGHT

    def it_accepts_center_and_centre(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        assert _resolve_alignment("center") == WD_ALIGN_PARAGRAPH.CENTER
        assert _resolve_alignment("centre") == WD_ALIGN_PARAGRAPH.CENTER

    def it_raises_on_unknown_alignment(self):
        with pytest.raises(ValueError):
            _resolve_alignment("upside-down")


# ---------------------------------------------------------------------------
# Total-row spec resolution
# ---------------------------------------------------------------------------


class Describe_resolve_total_spec:
    def it_returns_None_when_disabled(self):
        df = _small_numeric()
        dtypes = {c: df[c].dtype for c in df.columns}
        assert _resolve_total_spec(False, list(df.columns), dtypes) is None

    def it_sums_numeric_only_by_default(self):
        df = _mixed_types()
        dtypes = {c: df[c].dtype for c in df.columns}
        spec = _resolve_total_spec(True, list(df.columns), dtypes)
        assert spec["Revenue"] == "sum"
        assert spec["Region"] == "none"

    def it_supports_a_per_column_mapping(self):
        df = _mixed_types()
        dtypes = {c: df[c].dtype for c in df.columns}
        spec = _resolve_total_spec(
            {"Revenue": "sum", "Growth": "mean"}, list(df.columns), dtypes
        )
        assert spec["Revenue"] == "sum"
        assert spec["Growth"] == "mean"
        assert spec["Region"] == "none"

    def it_raises_on_unknown_column(self):
        df = _small_numeric()
        dtypes = {c: df[c].dtype for c in df.columns}
        with pytest.raises(ValueError):
            _resolve_total_spec(
                {"DoesNotExist": "sum"}, list(df.columns), dtypes
            )

    def it_raises_on_unknown_aggregator(self):
        df = _small_numeric()
        dtypes = {c: df[c].dtype for c in df.columns}
        with pytest.raises(ValueError):
            _resolve_total_spec("median", list(df.columns), dtypes)


# ---------------------------------------------------------------------------
# `Document.add_dataframe` — end-to-end behavior
# ---------------------------------------------------------------------------


class DescribeDocument_add_dataframe:
    def it_returns_a_Table_proxy_for_the_smallest_input(self):
        from docx.table import Table

        doc = OpenDocument()
        table = doc.add_dataframe(_small_numeric(), style="boxed")
        assert isinstance(table, Table)
        assert len(table.rows) == 4  # header + 3 data rows
        assert len(table.columns) == 2

    def it_writes_a_header_row_with_the_column_names(self):
        doc = OpenDocument()
        df = _small_numeric()
        table = doc.add_dataframe(df, style="executive")
        header_cells = [c.text for c in table.rows[0].cells]
        assert header_cells == ["Region", "Revenue"]

    def it_writes_each_data_row_in_order(self):
        doc = OpenDocument()
        table = doc.add_dataframe(_small_numeric(), style="boxed")
        # default float rendering -- non-int floats use repr()
        cells = [[c.text for c in row.cells] for row in table.rows[1:]]
        assert cells[0][0] == "AMER"
        assert cells[2][0] == "EMEA"

    def it_applies_per_column_number_formats(self):
        doc = OpenDocument()
        table = doc.add_dataframe(
            _mixed_types(),
            style="executive",
            number_format={"Revenue": "$,.1f", "Growth": ".1%"},
            show_total_row=False,
        )
        # column 1 = Revenue, row 1 = first data row
        assert table.rows[1].cells[1].text == "$1,234.5"
        assert table.rows[1].cells[2].text == "8.7%"

    def it_renders_date_columns_with_the_DSL(self):
        doc = OpenDocument()
        df = _date_columns()
        table = doc.add_dataframe(
            df,
            style="minimal",
            number_format={"Period": "MMM YYYY"},
        )
        assert table.rows[1].cells[0].text == "Jan 2024"
        assert table.rows[2].cells[0].text == "Feb 2024"
        assert table.rows[3].cells[0].text == "Mar 2024"

    def it_supports_total_row_sum(self):
        doc = OpenDocument()
        df = _small_numeric()
        table = doc.add_dataframe(
            df,
            style="executive",
            number_format={"Revenue": "$,.1f"},
            show_total_row=True,
        )
        # header + 3 rows + total
        assert len(table.rows) == 5
        total_cells = [c.text for c in table.rows[-1].cells]
        assert total_cells[0] == ""
        assert total_cells[1] == "$2,876.4"

    def it_supports_total_row_mean_via_string(self):
        doc = OpenDocument()
        df = _small_numeric()
        table = doc.add_dataframe(
            df,
            style="executive",
            number_format={"Revenue": "$,.1f"},
            show_total_row="mean",
        )
        # mean of 1234.5, 987.6, 654.3 ≈ 958.8
        total = table.rows[-1].cells[1].text
        assert total.startswith("$") and "958.8" in total

    def it_supports_total_row_count(self):
        doc = OpenDocument()
        df = _mixed_types()
        table = doc.add_dataframe(
            df, style="boxed", show_total_row="count"
        )
        # Headcount column total = 4
        assert table.rows[-1].cells[-1].text == "4"

    def it_supports_total_row_per_column_mapping(self):
        doc = OpenDocument()
        df = _mixed_types()
        table = doc.add_dataframe(
            df,
            style="executive",
            number_format={"Revenue": "$,.0f", "Growth": ".1%"},
            show_total_row={"Revenue": "sum", "Growth": "mean"},
        )
        cells = [c.text for c in table.rows[-1].cells]
        # Revenue sum = 3197.4 -> "$3,197"
        assert cells[1] == "$3,197"
        # Growth mean ≈ 0.05975 -> "6.0%"
        assert cells[2].endswith("%")

    def it_rejects_an_unknown_style(self):
        doc = OpenDocument()
        with pytest.raises(ValueError):
            doc.add_dataframe(_small_numeric(), style="rainbow")

    def it_rejects_a_non_dataframe_argument(self):
        doc = OpenDocument()
        with pytest.raises(TypeError):
            doc.add_dataframe([1, 2, 3], style="boxed")

    def it_rejects_an_empty_columns_dataframe(self):
        doc = OpenDocument()
        empty = pd.DataFrame()
        with pytest.raises(ValueError):
            doc.add_dataframe(empty, style="boxed")

    @pytest.mark.parametrize("style", _BUILTIN_STYLES)
    def it_handles_every_built_in_style(self, style):
        from docx.table import Table

        doc = OpenDocument()
        table = doc.add_dataframe(_small_numeric(), style=style)
        assert isinstance(table, Table)
        assert len(table.rows) == 4

    def it_handles_a_large_DataFrame(self):
        doc = OpenDocument()
        df = _large_numeric(50)
        table = doc.add_dataframe(
            df, style="striped", number_format={"x": ",.2f", "y": ",.3f"}
        )
        # header + 50 data rows
        assert len(table.rows) == 51

    def it_handles_NaN_cells_as_blanks(self):
        doc = OpenDocument()
        df = _with_nans()
        table = doc.add_dataframe(df, style="boxed", show_total_row=False)
        # second row, Revenue col = NaN -> ""
        assert table.rows[2].cells[1].text == ""
        # third row, Growth col = NaN -> ""
        assert table.rows[3].cells[2].text == ""

    def it_supports_alternating_rows_override(self):
        # alternating_rows=True on a "boxed" style (which has no alt fill
        # by default) should switch tints on
        doc = OpenDocument()
        table = doc.add_dataframe(
            _large_numeric(4), style="boxed", alternating_rows=True
        )
        # Row 1 (idx 0 of data) -> no alt; row 2 -> alt; row 3 -> none; row 4 -> alt.
        # We just check that at least the second data row has a fill set.
        from docx.shared import RGBColor

        assert isinstance(table.rows[2].cells[0].shading.fill_color, RGBColor)

    def it_accepts_a_hex_string_header_color(self):
        doc = OpenDocument()
        table = doc.add_dataframe(
            _small_numeric(),
            style="executive",
            header_color="336699",
            header_text_color="FFFFFF",
        )
        # header cell should carry the override
        assert (
            table.rows[0].cells[0].shading.fill_color
            == RGBColor(0x33, 0x66, 0x99)
        )

    def it_writes_an_alignment_override(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = OpenDocument()
        df = _mixed_types()
        table = doc.add_dataframe(
            df,
            style="boxed",
            align={"Revenue": "left", "Region": "center"},
        )
        # Revenue forced to LEFT, Region forced to CENTER
        rev_para = table.rows[1].cells[1].paragraphs[0]
        assert rev_para.alignment == WD_ALIGN_PARAGRAPH.LEFT
        reg_para = table.rows[1].cells[0].paragraphs[0]
        assert reg_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_round_trips_through_save_and_reopen(self):
        # Save the document containing the styled table to a buffer,
        # reopen it, and assert the table still carries the right cells.
        doc = OpenDocument()
        doc.add_dataframe(
            _mixed_types(),
            style="executive",
            number_format={"Revenue": "$,.0f"},
            show_total_row=True,
        )
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reopened = OpenDocument(buf)
        assert len(reopened.tables) == 1
        table = reopened.tables[0]
        # header + 4 data rows + total row
        assert len(table.rows) == 6
        # Revenue total (sum of 1234.5 + 987.6 + 654.3 + 321.0 = 3197.4) -> "$3,197"
        assert table.rows[-1].cells[1].text == "$3,197"

    def it_raises_ImportError_when_pandas_missing(self, monkeypatch):
        # Simulate pandas missing by patching the helpers used inside
        # add_dataframe so the `_is_dataframe(df)` sniff returns False
        # *and* the lazy `_require_pandas()` reraises.
        from docx import dataframe as df_mod

        def _no_pandas(_obj):
            return False

        def _raise(*_args, **_kw):
            raise ImportError("pretend pandas is missing")

        monkeypatch.setattr(df_mod, "_is_dataframe", _no_pandas)
        monkeypatch.setattr(df_mod, "_require_pandas", _raise)

        doc = OpenDocument()
        with pytest.raises(ImportError):
            doc.add_dataframe(_small_numeric(), style="boxed")
