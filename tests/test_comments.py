# pyright: reportPrivateUsage=false

"""Unit test suite for the `docx.comments` module."""

from __future__ import annotations

import datetime as dt
from typing import cast

import pytest

from docx.comments import Comment, Comments
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.packuri import PackURI
from docx.oxml.comments import CT_Comment, CT_Comments
from docx.oxml.ns import qn
from docx.package import Package
from docx.parts.comments import CommentsPart

from .unitutil.cxml import element
from .unitutil.mock import FixtureRequest, Mock, instance_mock


class DescribeComments:
    """Unit-test suite for `docx.comments.Comments` objects."""

    @pytest.mark.parametrize(
        ("cxml", "count"),
        [
            ("w:comments", 0),
            ("w:comments/w:comment", 1),
            ("w:comments/(w:comment,w:comment,w:comment)", 3),
        ],
    )
    def it_knows_how_many_comments_it_contains(self, cxml: str, count: int, package_: Mock):
        comments_elm = cast(CT_Comments, element(cxml))
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        assert len(comments) == count

    def it_is_iterable_over_the_comments_it_contains(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments/(w:comment,w:comment)"))
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        comment_iter = iter(comments)

        comment1 = next(comment_iter)
        assert type(comment1) is Comment, "expected a `Comment` object"
        comment2 = next(comment_iter)
        assert type(comment2) is Comment, "expected a `Comment` object"
        with pytest.raises(StopIteration):
            next(comment_iter)

    def it_can_get_a_comment_by_id(self, package_: Mock):
        comments_elm = cast(
            CT_Comments,
            element("w:comments/(w:comment{w:id=1},w:comment{w:id=2},w:comment{w:id=3})"),
        )
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        comment = comments.get(2)

        assert type(comment) is Comment, "expected a `Comment` object"
        assert comment._comment_elm is comments_elm.comment_lst[1]

    def but_it_returns_None_when_no_comment_with_that_id_exists(self, package_: Mock):
        comments_elm = cast(
            CT_Comments,
            element("w:comments/(w:comment{w:id=1},w:comment{w:id=2},w:comment{w:id=3})"),
        )
        comments = Comments(
            comments_elm,
            CommentsPart(
                PackURI("/word/comments.xml"),
                CT.WML_COMMENTS,
                comments_elm,
                package_,
            ),
        )

        comment = comments.get(4)

        assert comment is None, "expected None when no comment with that id exists"

    def it_can_add_a_new_comment(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        now_before = dt.datetime.now(dt.timezone.utc).replace(microsecond=0)
        comments = Comments(comments_elm, comments_part)

        comment = comments.add_comment()

        now_after = dt.datetime.now(dt.timezone.utc).replace(microsecond=0)
        # -- a comment is unconditionally added, and returned for any further adjustment --
        assert isinstance(comment, Comment)
        # -- it is "linked" to the comments part so it can add images and hyperlinks, etc. --
        assert comment.part is comments_part
        # -- comment numbering starts at 0, and is incremented for each new comment --
        assert comment.comment_id == 0
        # -- author is a required attribut, but is the empty string by default --
        assert comment.author == ""
        # -- initials is an optional attribute, but defaults to the empty string, same as Word --
        assert comment.initials == ""
        # -- timestamp is also optional, but defaults to now-UTC --
        assert comment.timestamp is not None
        assert now_before <= comment.timestamp <= now_after
        # -- by default, a new comment contains a single empty paragraph --
        assert [p.text for p in comment.paragraphs] == [""]
        # -- that paragraph has the "CommentText" style, same as Word applies --
        comment_elm = comment._comment_elm
        assert len(comment_elm.p_lst) == 1
        p = comment_elm.p_lst[0]
        assert p.style == "CommentText"
        # -- and that paragraph contains a single run with the necessary annotation reference --
        assert len(p.r_lst) == 1
        r = comment_elm.p_lst[0].r_lst[0]
        assert r.style == "CommentReference"
        assert r[-1].tag == qn("w:annotationRef")

    def and_it_can_add_text_to_the_comment_when_adding_it(self, comments: Comments, package_: Mock):
        comment = comments.add_comment(text="para 1\n\npara 2")

        assert len(comment.paragraphs) == 3
        assert [p.text for p in comment.paragraphs] == ["para 1", "", "para 2"]
        assert all(p._p.style == "CommentText" for p in comment.paragraphs)

    def and_it_sets_the_author_and_their_initials_when_adding_a_comment_when_provided(
        self, comments: Comments, package_: Mock
    ):
        comment = comments.add_comment(author="Steve Canny", initials="SJC")

        assert comment.author == "Steve Canny"
        assert comment.initials == "SJC"

    def and_it_accepts_an_explicit_tz_aware_date(
        self, comments: Comments, package_: Mock
    ):
        when = dt.datetime(2024, 1, 15, 10, 0, 0, tzinfo=dt.timezone.utc)

        comment = comments.add_comment(author="A", date=when)

        assert comment.timestamp == when

    def and_it_accepts_an_explicit_date_on_add_reply(
        self, comments: Comments, package_: Mock
    ):
        parent = comments.add_comment(author="A")
        when = dt.datetime(2024, 2, 1, 9, 0, 0, tzinfo=dt.timezone.utc)

        reply = parent.add_reply(author="B", date=when)

        assert reply.timestamp == when

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def comments(self, package_: Mock) -> Comments:
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        return Comments(comments_elm, comments_part)

    @pytest.fixture
    def package_(self, request: FixtureRequest):
        return instance_mock(request, Package)


class DescribeComment:
    """Unit-test suite for `docx.comments.Comment`."""

    def it_knows_its_comment_id(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42}"))
        comment = Comment(comment_elm, comments_part_)

        assert comment.comment_id == 42

    def it_knows_its_author(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:author=Steve Canny}"))
        comment = Comment(comment_elm, comments_part_)

        assert comment.author == "Steve Canny"

    def it_knows_the_initials_of_its_author(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:initials=SJC}"))
        comment = Comment(comment_elm, comments_part_)

        assert comment.initials == "SJC"

    def it_knows_the_date_and_time_it_was_authored(self, comments_part_: Mock):
        comment_elm = cast(
            CT_Comment,
            element("w:comment{w:id=42,w:date=2023-10-01T12:34:56Z}"),
        )
        comment = Comment(comment_elm, comments_part_)

        assert comment.timestamp == dt.datetime(2023, 10, 1, 12, 34, 56, tzinfo=dt.timezone.utc)

    @pytest.mark.parametrize(
        ("cxml", "expected_value"),
        [
            ("w:comment{w:id=42}", ""),
            ('w:comment{w:id=42}/w:p/w:r/w:t"Comment text."', "Comment text."),
            (
                'w:comment{w:id=42}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")',
                "First para\nSecond para",
            ),
            (
                'w:comment{w:id=42}/(w:p/w:r/w:t"First para",w:p,w:p/w:r/w:t"Second para")',
                "First para\n\nSecond para",
            ),
        ],
    )
    def it_can_summarize_its_content_as_text(
        self, cxml: str, expected_value: str, comments_part_: Mock
    ):
        assert Comment(cast(CT_Comment, element(cxml)), comments_part_).text == expected_value

    def it_provides_access_to_the_paragraphs_it_contains(self, comments_part_: Mock):
        comment_elm = cast(
            CT_Comment,
            element('w:comment{w:id=42}/(w:p/w:r/w:t"First para",w:p/w:r/w:t"Second para")'),
        )
        comment = Comment(comment_elm, comments_part_)

        paragraphs = comment.paragraphs

        assert len(paragraphs) == 2
        assert [para.text for para in paragraphs] == ["First para", "Second para"]

    def it_can_update_the_comment_author(self, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:author=Old Author}"))
        comment = Comment(comment_elm, comments_part_)

        comment.author = "New Author"

        assert comment.author == "New Author"

    @pytest.mark.parametrize(
        "initials",
        [
            # -- valid initials --
            "XYZ",
            # -- empty string is valid
            "",
            # -- None is valid, removes existing initials
            None,
        ],
    )
    def it_can_update_the_comment_initials(self, initials: str | None, comments_part_: Mock):
        comment_elm = cast(CT_Comment, element("w:comment{w:id=42,w:initials=ABC}"))
        comment = Comment(comment_elm, comments_part_)

        comment.initials = initials

        assert comment.initials == initials

    def it_can_add_a_reply_to_a_comment(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        comments = Comments(comments_elm, comments_part)
        parent = comments.add_comment(text="Parent comment", author="Author A", initials="AA")

        reply = parent.add_reply(text="Reply text", author="Author B", initials="BB")

        assert isinstance(reply, Comment)
        assert reply.text == "Reply text"
        assert reply.author == "Author B"
        assert reply.initials == "BB"
        assert reply.timestamp is not None
        assert reply.comment_id != parent.comment_id
        # -- verify the reply is linked to the parent via paraIdParent --
        assert reply._comment_elm.paraIdParent == parent._comment_elm.paraId

    def it_can_list_replies_to_a_comment(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        comments = Comments(comments_elm, comments_part)
        parent = comments.add_comment(text="Parent", author="A")
        parent.add_reply(text="Reply 1", author="B")
        parent.add_reply(text="Reply 2", author="C")
        # -- add an unrelated comment --
        comments.add_comment(text="Other comment", author="D")

        replies = parent.replies

        assert len(replies) == 2
        assert replies[0].text == "Reply 1"
        assert replies[0].author == "B"
        assert replies[1].text == "Reply 2"
        assert replies[1].author == "C"

    def and_it_returns_empty_list_when_no_replies(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        comments = Comments(comments_elm, comments_part)
        parent = comments.add_comment(text="Parent", author="A")

        assert parent.replies == []

    def it_can_add_a_reply_with_no_text(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        comments = Comments(comments_elm, comments_part)
        parent = comments.add_comment(author="A")

        reply = parent.add_reply(author="B")

        assert isinstance(reply, Comment)
        assert [p.text for p in reply.paragraphs] == [""]

    def it_can_add_a_multiline_reply(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        comments = Comments(comments_elm, comments_part)
        parent = comments.add_comment(author="A")

        reply = parent.add_reply(text="line 1\nline 2", author="B")

        assert len(reply.paragraphs) == 2
        assert [p.text for p in reply.paragraphs] == ["line 1", "line 2"]

    # -- Word 2013+ commentsExtended (resolve/reopen + parent chain) -------------

    def it_starts_unresolved_for_a_newly_added_comment(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("hi")
        comment = doc.add_comment(runs=[p.add_run(" there")], text="body", author="A")

        assert comment.is_resolved is False

    def it_can_be_resolved_and_reopened(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("hi")
        comment = doc.add_comment(runs=[p.add_run(" there")], text="body", author="A")

        comment.resolve()
        assert comment.is_resolved is True
        comment.reopen()
        assert comment.is_resolved is False

    def it_creates_the_extended_part_on_first_resolve(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("hi")
        comment = doc.add_comment(runs=[p.add_run(" there")], text="body", author="A")
        comments_part = doc.part._comments_part  # type: ignore[attr-defined]

        assert comments_part.comments_extended_part is None
        comment.resolve()
        assert comments_part.comments_extended_part is not None

    def it_preserves_the_parent_linkage_across_resolve_reopen(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("hi")
        parent = doc.add_comment(runs=[p.add_run(" there")], text="parent", author="A")
        reply = parent.reply(text="reply", author="B")

        reply.resolve()
        reply.reopen()

        ex_part = doc.part._comments_part.comments_extended_part  # type: ignore[attr-defined]
        assert ex_part is not None
        reply_para_id = reply._comment_elm.paraId
        assert reply_para_id is not None
        entry = ex_part.element.get_commentEx_by_paraId(reply_para_id)
        assert entry is not None
        # -- parent link survives the reopen --
        assert entry.paraIdParent == parent._comment_elm.paraId

    def it_finds_the_parent_comment_via_paraIdParent(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("hi")
        parent = doc.add_comment(runs=[p.add_run(" there")], text="parent", author="A")
        reply = parent.reply(text="reply", author="B")

        resolved = reply.parent_comment

        assert resolved is not None
        assert resolved.comment_id == parent.comment_id

    def it_returns_None_for_parent_comment_when_the_comment_is_a_root(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph("hi")
        root = doc.add_comment(runs=[p.add_run(" there")], text="root", author="A")

        assert root.parent_comment is None

    def it_exposes_reply_as_an_alias_for_add_reply(self, package_: Mock):
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        comments = Comments(comments_elm, comments_part)
        parent = comments.add_comment(text="p", author="A")

        reply = parent.reply(text="r", author="B")

        assert isinstance(reply, Comment)
        assert reply.text == "r"
        assert reply.author == "B"
        assert reply._comment_elm.paraIdParent == parent._comment_elm.paraId

    def it_round_trips_extended_state_through_save_and_load(self, tmp_path):
        import io

        from docx import Document

        doc = Document()
        p = doc.add_paragraph("hi")
        parent = doc.add_comment(runs=[p.add_run(" there")], text="parent", author="A")
        reply = parent.reply(text="child", author="B")
        parent.resolve()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reloaded = Document(buf)

        loaded = list(reloaded.comments)
        assert len(loaded) == 2
        by_id = {c.comment_id: c for c in loaded}
        assert by_id[parent.comment_id].is_resolved is True
        assert by_id[reply.comment_id].is_resolved is False
        # -- parent-chain survives round-trip --
        assert by_id[reply.comment_id].parent_comment is not None
        assert by_id[reply.comment_id].parent_comment.comment_id == parent.comment_id  # type: ignore[union-attr]
        # -- replies survive round-trip --
        assert [r.comment_id for r in by_id[parent.comment_id].replies] == [
            reply.comment_id
        ]

    def it_raises_when_setting_resolved_state_without_a_paraId(self, package_: Mock):
        # -- simulate a legacy comment parsed without a paraId attribute --
        comment_elm = cast(CT_Comment, element("w:comment{w:id=5}"))
        comments_elm = cast(CT_Comments, element("w:comments"))
        comments_elm.append(comment_elm)
        comments_part = CommentsPart(
            PackURI("/word/comments.xml"),
            CT.WML_COMMENTS,
            comments_elm,
            package_,
        )
        comment = Comment(comment_elm, comments_part)

        with pytest.raises(ValueError, match="paraId"):
            comment.resolve()

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def comments_part_(self, request: FixtureRequest):
        return instance_mock(request, CommentsPart)

    @pytest.fixture
    def package_(self, request: FixtureRequest):
        return instance_mock(request, Package)
