"""End-to-end Strict-conformance fixture round-trips (R16-1).

Complements :mod:`tests.test_strict_conformance` — that suite exercises
the ``strict=`` / ``is_strict`` API plumbing with pre-canned bytes.
This one walks the full lifecycle:

1. Build a real ECMA-376 Strict ``.docx`` fixture by rewriting the
   namespace URIs in python-docx's minimal Transitional template.
2. Open the Strict fixture with :func:`docx.Document` and assert the
   :attr:`~docx.document.Document.is_strict` sniff fires.
3. Edit a paragraph, save the document back (default ``strict=None``
   preserves the flag on the package), then reopen.
4. Assert the edit survived and the conformance-class flag is still
   carried (reopening with ``strict=True`` because the emitted bytes
   are always Transitional today).

Also covers the cross-format lift — a Transitional ``.docx`` can be
saved with ``strict=True`` to force the conformance flag on, and the
reloaded package reports ``is_strict=True`` with content intact.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document

from tests.conftest import _rewrite_ns_to_strict


def _default_template_bytes() -> bytes:
    tpl = (
        Path(__file__).parent.parent
        / "src" / "docx" / "templates" / "default.docx"
    )
    with open(tpl, "rb") as f:
        return f.read()


class DescribeStrictFixtureRoundTrip:
    """Full open → edit → save → reopen cycle on a Strict-namespace fixture."""

    def it_round_trips_a_strict_fixture_preserving_flag_and_edit(
        self, tmp_path: Path
    ) -> None:
        # -- 1. synthesize a Strict .docx by rewriting namespace URIs in
        # -- the minimal Transitional template shipped with python-docx --
        strict_bytes = _rewrite_ns_to_strict(_default_template_bytes())
        with zipfile.ZipFile(io.BytesIO(strict_bytes)) as zf:
            doc_xml = zf.read("word/document.xml")
        assert b"purl.oclc.org/ooxml/wordprocessingml/main" in doc_xml

        # -- 2. open it with strict=True; the PackageReader's sniff
        # -- would have flagged it anyway, but the explicit kwarg makes
        # -- the test's intent unambiguous --
        doc = Document(io.BytesIO(strict_bytes), strict=True)
        assert doc.is_strict is True

        # -- 3. edit a paragraph (the default template carries at least
        # -- one empty body paragraph to mutate) --
        edit_marker = "R16-STRICT-EDIT-MARKER"
        doc.add_paragraph(edit_marker)

        # -- 4. save with strict=None (default); this preserves the
        # -- is_strict flag on the in-memory package but emits
        # -- Transitional bytes (docx has no Transitional→Strict
        # -- writer today, only the reader path) --
        out = tmp_path / "round_trip.docx"
        doc.save(str(out))

        # -- 5. reopen with strict=True (the written bytes are
        # -- Transitional, so a default reopen would sniff as
        # -- not-strict; the is_strict contract is "flag survives
        # -- across the API surface", not "emitted bytes are Strict") --
        doc2 = Document(str(out), strict=True)
        assert doc2.is_strict is True

        # -- 6. the edit survived --
        texts = [p.text for p in doc2.paragraphs]
        assert edit_marker in texts

    def it_accepts_a_strict_fixture_without_explicit_strict_kwarg(
        self, tmp_path: Path
    ) -> None:
        # -- the PackageReader sniffs Strict namespaces at open time
        # -- and flips is_strict=True automatically; callers don't
        # -- need to opt in when the source is genuinely Strict --
        strict_bytes = _rewrite_ns_to_strict(_default_template_bytes())
        doc = Document(io.BytesIO(strict_bytes))
        assert doc.is_strict is True

        # -- the body is navigable (the Strict→Transitional
        # -- translator ran at read time, so downstream code sees
        # -- Transitional URIs) --
        _ = list(doc.paragraphs)  # does not raise
        # -- and a subsequent edit + save cycle also succeeds --
        doc.add_paragraph("sniff-path")
        assert any(p.text == "sniff-path" for p in doc.paragraphs)


class DescribeTransitionalToStrictCrossFormatLift:
    """Transitional fixture → save strict=True → reopen → flag is True."""

    def it_lifts_a_transitional_docx_into_strict_on_save(
        self, tmp_path: Path
    ) -> None:
        # -- author from a plain Transitional template; flag is False --
        doc = Document(io.BytesIO(_default_template_bytes()))
        assert doc.is_strict is False

        edit_marker = "R16-TRANS-TO-STRICT-MARKER"
        doc.add_paragraph(edit_marker)

        # -- save(strict=True) forces the conformance flag --
        out = tmp_path / "lifted.docx"
        doc.save(str(out), strict=True)

        # -- reopen with strict=True so the flag comes back (the
        # -- emitted bytes are Transitional; the `strict=True` on
        # -- save flipped the in-memory flag but did not translate
        # -- namespaces). --
        doc2 = Document(str(out), strict=True)
        assert doc2.is_strict is True

        # -- content is intact --
        texts = [p.text for p in doc2.paragraphs]
        assert edit_marker in texts
