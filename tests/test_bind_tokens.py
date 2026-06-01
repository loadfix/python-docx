# pyright: reportPrivateUsage=false

"""Tests for the smart-placeholder bind-token system landed for issue #68.

Covers the contracts the feature ships with:

1. ``Document.add_paragraph(text, bind_to=record)`` resolves dotted-path
   tokens (``{customer.name}``, ``{customer.address.line1}``) at save
   time against the bound record.
2. ``Document.bind(record=other)`` re-binds a different record so the
   *same* saved document re-resolves cleanly on the next save.
3. The token source string is preserved in a fork-scoped
   ``<lfxbind:src>`` child element so ``load -> bind -> save`` cycles
   re-resolve against the new record instead of carrying a stale
   literal forward.
4. ``{date:short}`` / ``{date:'MMM d, yyyy'}`` resolve via the date
   formatter; ``{property:Title}`` reads from the document's core /
   custom properties; ``{i}`` reflects the current iteration index.
5. Unknown tokens / mis-cased tokens are left literal — a stray
   ``{Foo}`` in user prose must never get silently swallowed.
"""

from __future__ import annotations

import datetime as dt
import io

from docx import Document
from docx.bind_tokens import (
    LFXBIND_NS,
    apply_bind_tokens,
    get_bound_record,
    get_token_source,
    has_persisted_marker,
    has_token,
    render,
    reseat_token_source,
    set_bound_record,
)


# -- helpers ----------------------------------------------------------------


def _roundtrip(doc):
    """Save ``doc`` to a BytesIO and reload it, returning the new |Document|."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf), buf


def _first_run(paragraph):
    return paragraph.runs[0]._r


class _Record:
    """Tiny attribute-access record fixture."""

    def __init__(self, **kwargs):
        for k, v in kwargs.items():
            setattr(self, k, v)


# -- render() ---------------------------------------------------------------


class DescribeRenderHelper:
    def it_substitutes_a_dotted_path_token(self):
        record = _Record(customer=_Record(name="Acme"))
        assert render("Dear {customer.name}", record=record) == "Dear Acme"

    def it_substitutes_a_nested_dotted_path_token(self):
        record = {"customer": {"address": {"line1": "1 Loop"}}}
        assert (
            render("Address: {customer.address.line1}", record=record)
            == "Address: 1 Loop"
        )

    def it_substitutes_multiple_tokens_in_one_string(self):
        record = {"customer": {"name": "Acme"}, "product": {"tier": "Pro"}}
        out = render("Dear {customer.name}, you have {product.tier}", record=record)
        assert out == "Dear Acme, you have Pro"

    def it_leaves_unknown_path_tokens_literal(self):
        record = {"customer": {"name": "Acme"}}
        # -- ``foo.bar`` is unresolvable; left literal --
        assert (
            render("{customer.name} {foo.bar}", record=record)
            == "Acme {foo.bar}"
        )

    def it_leaves_tokens_literal_when_record_is_None(self):
        assert render("Dear {customer.name}", record=None) == "Dear {customer.name}"

    def it_renders_None_value_as_empty_string_via_unresolved_token(self):
        # -- a missing key produces a *literal* token, not "None" --
        record = {"customer": {}}
        assert render("[{customer.name}]", record=record) == "[{customer.name}]"

    def it_resolves_date_short_alias(self):
        out = render("{date:short}", today=dt.date(2026, 1, 9))
        assert out == "2026-01-09"

    def it_resolves_date_iso_alias(self):
        out = render("{date:iso}", today=dt.date(2026, 1, 9))
        assert out == "2026-01-09"

    def it_resolves_date_long_alias(self):
        out = render("{date:long}", today=dt.date(2026, 1, 9))
        assert out == "January 9, 2026"

    def it_resolves_date_medium_alias(self):
        out = render("{date:medium}", today=dt.date(2026, 1, 9))
        assert out == "Jan 9, 2026"

    def it_resolves_date_with_a_quoted_babel_format(self):
        out = render("{date:'MMM d, yyyy'}", today=dt.date(2026, 1, 9))
        assert out == "Jan 9, 2026"

    def it_resolves_date_with_yyyy_MM_dd(self):
        out = render("{date:'yyyy-MM-dd'}", today=dt.date(2026, 1, 9))
        assert out == "2026-01-09"

    def it_resolves_iteration_token(self):
        assert render("Row {i}", iteration=3) == "Row 3"

    def it_leaves_iteration_token_literal_when_no_iteration_supplied(self):
        assert render("Row {i}", iteration=None) == "Row {i}"

    def it_resolves_property_against_a_property_map(self):
        out = render("{property:Title}", properties={"Title": "Q1 Plan"})
        assert out == "Q1 Plan"

    def it_resolves_property_case_insensitively(self):
        out = render("{property:title}", properties={"Title": "Q1 Plan"})
        assert out == "Q1 Plan"

    def it_leaves_property_token_literal_when_unknown(self):
        out = render("{property:Subject}", properties={"Title": "X"})
        assert out == "{property:Subject}"


class DescribeHasToken:
    def it_detects_a_dotted_path_token(self):
        assert has_token("Hi {customer.name}") is True

    def it_detects_a_date_token(self):
        assert has_token("today: {date:short}") is True

    def it_returns_false_on_plain_text(self):
        assert has_token("plain prose") is False

    def it_returns_false_on_empty_braces(self):
        # -- {} is not a recognised token shape --
        assert has_token("set = {} for empty") is False

    def it_returns_false_when_brace_content_has_a_space(self):
        # -- ``{Foo bar}`` is not a recognised token shape --
        assert has_token("see {Foo bar} below") is False


# -- end-to-end via Document API -------------------------------------------


class DescribeAddParagraphBindTo:
    def it_resolves_dotted_path_tokens_on_save(self):
        record = {
            "customer": {"name": "Acme Corp"},
            "product": {"tier": "Pro"},
            "dates": {"expiry": "2026-12-31"},
        }
        doc = Document()
        doc.add_paragraph(
            "Dear {customer.name}, your {product.tier} subscription expires"
            " on {dates.expiry}.",
            bind_to=record,
        )

        reloaded, _ = _roundtrip(doc)

        assert reloaded.paragraphs[0].text == (
            "Dear Acme Corp, your Pro subscription expires on 2026-12-31."
        )

    def it_resolves_nested_dotted_paths(self):
        record = {"customer": {"address": {"line1": "1 Infinite Loop"}}}
        doc = Document()
        doc.add_paragraph("Ship to {customer.address.line1}.", bind_to=record)

        reloaded, _ = _roundtrip(doc)

        assert reloaded.paragraphs[0].text == "Ship to 1 Infinite Loop."

    def it_preserves_the_token_source_in_a_marker_child(self):
        record = {"customer": {"name": "Acme"}}
        doc = Document()
        para = doc.add_paragraph("Dear {customer.name}", bind_to=record)

        # -- before save, source marker is stamped on the run --
        r = _first_run(para)
        assert get_token_source(r) == "Dear {customer.name}"

    def it_preserves_source_through_save_and_reload(self):
        record = {"customer": {"name": "Acme"}}
        doc = Document()
        doc.add_paragraph("Dear {customer.name}", bind_to=record)

        reloaded, _ = _roundtrip(doc)

        ro_run = reloaded.paragraphs[0].runs[0]._r
        assert get_token_source(ro_run) == "Dear {customer.name}"

    def it_re_resolves_after_a_rebind_cycle(self):
        # -- This is the headline acceptance case from issue #68. --
        first = {"customer": {"name": "Acme"}, "dates": {"expiry": "2026-12-31"}}
        second = {"customer": {"name": "Globex"}, "dates": {"expiry": "2027-01-01"}}

        doc = Document()
        doc.add_paragraph(
            "Dear {customer.name}, expires {dates.expiry}.",
            bind_to=first,
        )

        reloaded, _ = _roundtrip(doc)
        assert (
            reloaded.paragraphs[0].text == "Dear Acme, expires 2026-12-31."
        )

        # -- rebind and re-save; the same source string re-resolves --
        reloaded.bind(record=second)
        reloaded2, _ = _roundtrip(reloaded)

        assert (
            reloaded2.paragraphs[0].text == "Dear Globex, expires 2027-01-01."
        )

    def it_returns_self_from_bind(self):
        doc = Document()
        assert doc.bind(record={"a": 1}) is doc
        assert get_bound_record(doc) == {"a": 1}

    def it_resolves_property_token_from_core_properties(self):
        doc = Document()
        doc.core_properties.title = "Q1 Plan"
        doc.add_paragraph("Title: {property:Title}", bind_to={})

        reloaded, _ = _roundtrip(doc)

        assert reloaded.paragraphs[0].text == "Title: Q1 Plan"

    def it_resolves_iteration_token(self):
        doc = Document()
        doc.add_paragraph("Row {i} of {customer.name}", bind_to={"customer": {"name": "Acme"}})
        doc.bind(record={"customer": {"name": "Acme"}}, iteration=4)

        reloaded, _ = _roundtrip(doc)

        assert reloaded.paragraphs[0].text == "Row 4 of Acme"

    def it_leaves_unknown_tokens_literal_in_saved_doc(self):
        doc = Document()
        doc.add_paragraph(
            "Hello {customer.name} and {Foo bar}",
            bind_to={"customer": {"name": "Acme"}},
        )

        reloaded, _ = _roundtrip(doc)

        # -- ``{Foo bar}`` doesn't match the token shape and survives. --
        assert reloaded.paragraphs[0].text == "Hello Acme and {Foo bar}"

    def it_does_not_stamp_a_marker_on_token_less_paragraphs(self):
        doc = Document()
        para = doc.add_paragraph(
            "Plain prose with no tokens.",
            bind_to={"customer": {"name": "Acme"}},
        )

        # -- no token in text => no marker child --
        r = _first_run(para)
        assert get_token_source(r) is None

    def it_is_idempotent_when_resolving(self):
        record = {"customer": {"name": "Acme"}}
        doc = Document()
        doc.add_paragraph("Dear {customer.name}", bind_to=record)

        apply_bind_tokens(doc)
        first_xml = _first_run(doc.paragraphs[0]).xml

        apply_bind_tokens(doc)
        second_xml = _first_run(doc.paragraphs[0]).xml

        assert first_xml == second_xml


class DescribeReseatTokenSource:
    def it_stamps_a_marker_a_subsequent_save_will_re_resolve(self):
        # -- start with a paragraph whose text is a stale literal but no marker --
        doc = Document()
        para = doc.add_paragraph("Dear Acme")
        r = _first_run(para)

        # -- author asserts the run was meant to be a {customer.name} field --
        reseat_token_source(r, "Dear {customer.name}")
        set_bound_record(doc, {"customer": {"name": "Globex"}})

        reloaded, _ = _roundtrip(doc)

        assert reloaded.paragraphs[0].text == "Dear Globex"
        assert (
            get_token_source(reloaded.paragraphs[0].runs[0]._r)
            == "Dear {customer.name}"
        )


class DescribeNamespaceURI:
    def it_uses_the_documented_loadfix_uri(self):
        # -- changing this URI silently breaks every previously-saved doc --
        assert LFXBIND_NS == "https://loadfix.dev/docx/bind-tokens"


# -- issue #733: apply_bind_tokens must be opt-in -------------------------


def _saved_document_xml(doc):
    """Save ``doc`` to BytesIO and return the ``word/document.xml`` payload."""
    import zipfile

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    zf = zipfile.ZipFile(buf)
    return zf.read("word/document.xml").decode("utf-8"), buf


class DescribeApplyBindTokensGate:
    """Guard the unconditional resolver run that broke Word compatibility (#733).

    ``apply_bind_tokens`` previously fired on every ``Document.save``. When
    a run's text contained a brace-quoted identifier like
    ``{customer-code}``, the resolver appended an
    ``<lfxbind:src xmlns:lfxbind="...">…</lfxbind:src>`` element with an
    inline namespace declaration. The fork-internal namespace isn't on
    the document root and isn't wrapped in ``mc:AlternateContent``, so
    Microsoft Word rejects the file as malformed. The fix gates the
    resolver on opt-in: a record must be bound (or a marker must
    already be persisted from a previous bound save) before the
    resolver runs.
    """

    def it_does_not_emit_lfxbind_src_on_a_no_op_round_trip_with_brace_token_prose(
        self,
    ):
        # -- caller has not opted into bind-tokens; prose with brace
        # -- tokens must round-trip cleanly. --
        doc = Document()
        doc.add_paragraph("aws-{customer-code}-{account}-{role}")
        doc.add_paragraph("plain {token} placeholder in code sample")

        xml, _ = _saved_document_xml(doc)

        assert "lfxbind:src" not in xml
        assert "lfxbind" not in xml

    def it_does_not_emit_lfxbind_src_when_no_record_is_bound(self):
        doc = Document()
        doc.add_paragraph("see {customer.name} below")

        xml, _ = _saved_document_xml(doc)

        assert "lfxbind:src" not in xml

    def it_does_not_corrupt_brace_token_prose_in_runs(self):
        # -- a run's text content must survive verbatim through save
        # -- when bind-tokens is not opted into. --
        original = "aws-{customer-code}-{account}-{role}"
        doc = Document()
        doc.add_paragraph(original)

        xml, buf = _saved_document_xml(doc)

        assert original in xml
        # -- and on reload, the paragraph text matches exactly --
        reloaded = Document(buf)
        assert reloaded.paragraphs[0].text == original

    def it_emits_lfxbind_src_when_a_record_has_been_bound(self):
        # -- regression: don't break the bind-tokens feature for users
        # -- who actually use it. --
        doc = Document()
        set_bound_record(doc, {"customer": {"name": "Acme"}})
        doc.add_paragraph("Dear {customer.name}")

        xml, _ = _saved_document_xml(doc)

        assert "lfxbind:src" in xml
        # -- and the prefix is declared on root, not inline on the marker --
        assert "<lfxbind:src xmlns:lfxbind" not in xml

    def it_round_trips_a_document_with_persisted_lfxbind_src_markers(self):
        # -- save once with bind active so the doc carries persisted markers --
        doc = Document()
        doc.add_paragraph("Dear {customer.name}", bind_to={"customer": {"name": "Acme"}})
        _, buf = _saved_document_xml(doc)

        # -- reload (no rebind) and save again. The persisted marker must
        # -- survive: the save-time gate fires on ``has_persisted_marker``
        # -- so previously-bound documents continue to round-trip
        # -- through the resolver. --
        reloaded = Document(buf)
        # -- pre-flight: the loaded doc has the marker --
        assert has_persisted_marker(reloaded) is True

        xml2, _ = _saved_document_xml(reloaded)

        assert "lfxbind:src" in xml2

    def it_declares_lfxbind_on_root_not_inline_when_marker_is_emitted(self):
        # -- the inline form is what triggers Word's namespace-mismatch
        # -- rejection (#733). When we do emit a marker, the prefix
        # -- must be hoisted onto the document root. --
        doc = Document()
        doc.add_paragraph(
            "Dear {customer.name}", bind_to={"customer": {"name": "Acme"}}
        )

        xml, _ = _saved_document_xml(doc)

        # -- inline declaration absent --
        assert "<lfxbind:src xmlns:lfxbind" not in xml
        # -- root declares the prefix --
        body_open = xml.index("<w:body>")
        root_chunk = xml[:body_open]
        assert 'xmlns:lfxbind="https://loadfix.dev/docx/bind-tokens"' in root_chunk


class DescribeHasPersistedMarker:
    def it_returns_false_for_a_fresh_document(self):
        doc = Document()
        assert has_persisted_marker(doc) is False

    def it_returns_false_for_a_document_with_brace_prose_no_binding(self):
        doc = Document()
        doc.add_paragraph("aws-{customer-code}-{role}")
        assert has_persisted_marker(doc) is False

    def it_returns_true_when_a_marker_has_been_stamped(self):
        doc = Document()
        doc.add_paragraph(
            "Dear {customer.name}", bind_to={"customer": {"name": "Acme"}}
        )
        assert has_persisted_marker(doc) is True

    def it_returns_true_after_load_of_a_previously_bound_document(self):
        doc = Document()
        doc.add_paragraph(
            "Dear {customer.name}", bind_to={"customer": {"name": "Acme"}}
        )
        reloaded, _ = _roundtrip(doc)
        assert has_persisted_marker(reloaded) is True


class DescribeWriteSourceMarkerDefensiveGuard:
    """``_write_source_marker`` must suppress writes when root lacks the prefix.

    Belt-and-braces companion to the save-time gate (#733). Even if a
    future code path tries to stamp a marker without first declaring
    the ``lfxbind`` prefix on the document root, we suppress rather
    than emit a Word-incompatible inline ``xmlns:lfxbind``. The
    public stampers (:func:`reseat_token_source` and the bind-feature
    :func:`set_bound_record` pathway) hoist the prefix onto root
    before stamping, so this guard fires only on unexpected paths.
    """

    def it_suppresses_inline_xmlns_when_root_has_no_lfxbind(self):
        # pyright: reportPrivateUsage=false
        from docx.bind_tokens import _write_source_marker

        doc = Document()
        para = doc.add_paragraph("plain text")
        r = _first_run(para)

        # -- root does not declare lfxbind --
        assert doc.element.nsmap.get("lfxbind") is None

        _write_source_marker(r, "Dear {customer.name}")

        # -- guard kicked in: no marker stamped --
        assert get_token_source(r) is None
