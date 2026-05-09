"""Authoring-side tests for ``python-ooxml-signatures`` 0.2 adoption (R3-4).

Exercises :meth:`Document.add_signature_line` — an unsigned signature-line
placeholder emitter that relies on the 0.2 shared-package read API for
round-trip verification. These tests sit beside :mod:`tests.test_signatures`
(read-side) rather than inside it so a future extraction of the authoring
bits into ``ooxml_signatures.Signer`` can move this file wholesale.
"""

from __future__ import annotations

import io

import pytest

from docx import Document


def _ooxml_signatures_available() -> bool:
    try:
        import ooxml_signatures  # noqa: F401

        return True
    except ImportError:  # pragma: no cover — CI always has the package
        return False


class DescribeDocumentAddSignatureLine:
    def it_adds_a_placeholder_visible_in_signatures(self):
        doc = Document()

        info = doc.add_signature_line("CN=Alice Example, O=Acme")

        assert doc.is_signed is True
        assert len(doc.signatures) == 1
        assert info.signer == "CN=Alice Example, O=Acme"
        assert str(info.partname) == "/_xmlsignatures/sig1.xml"

    def it_round_trips_signer_and_comments_through_BytesIO(self):
        if not _ooxml_signatures_available():
            pytest.skip("python-ooxml-signatures not installed")

        doc = Document()
        doc.add_signature_line(
            "CN=Alice Example, O=Acme",
            signer_title="Chief Example Officer",
            email="alice@acme.test",
        )

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        reloaded = Document(buf)
        assert reloaded.is_signed is True
        sigs = reloaded.signatures
        assert len(sigs) == 1

        sig = sigs[0]
        assert sig.signer == "CN=Alice Example, O=Acme"
        shared = sig.shared_signature
        assert shared is not None
        assert shared.comments is not None
        # -- signer_title + email are encoded into the mdssi:SignatureComments
        # -- "Value" child via ``title=...; email=...``.
        assert "title=Chief Example Officer" in shared.comments
        assert "email=alice@acme.test" in shared.comments

    def it_allocates_sequential_partnames_for_multiple_signers(self):
        doc = Document()
        doc.add_signature_line("CN=Alice")
        doc.add_signature_line("CN=Bob")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        reloaded = Document(buf)
        partnames = sorted(str(s.partname) for s in reloaded.signatures)
        assert partnames == [
            "/_xmlsignatures/sig1.xml",
            "/_xmlsignatures/sig2.xml",
        ]

    def it_emits_unsigned_signature_value(self):
        # -- The placeholder deliberately leaves <SignatureValue> empty so
        # -- the round-tripped blob remains cryptographically invalid; a
        # -- downstream signer is expected to fill it in.
        doc = Document()
        info = doc.add_signature_line("CN=Alice")
        assert b"<SignatureValue></SignatureValue>" in info.blob or (
            b"<SignatureValue/>" in info.blob
        )
