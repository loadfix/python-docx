"""Test suite for the docx.api module."""

import importlib.util
import io
import zipfile
from pathlib import Path

import pytest
from lxml import etree

from docx.api import Document as DocumentFactoryFn
from docx.document import Document as DocumentCls
from docx.exceptions import EncryptedDocumentError
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.exceptions import PackageNotFoundError

from .unitutil.mock import FixtureRequest, Mock, class_mock, function_mock, instance_mock

_OLE_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _make_malformed_docx_bytes() -> bytes:
    """Return a zip-packaged .docx whose `word/document.xml` is truncated mid-tag.

    The rest of the package is valid so recovery mode has something to graft
    the degraded document part onto.
    """
    tpl_path = Path(__file__).parent.parent / "src" / "docx" / "templates" / "default.docx"
    with open(tpl_path, "rb") as f:
        blob = f.read()
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(blob), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = data[: len(data) // 2]  # -- truncate mid-element --
                zout.writestr(item, data)
    return out.getvalue()


def _make_valid_docx_bytes() -> bytes:
    tpl_path = Path(__file__).parent.parent / "src" / "docx" / "templates" / "default.docx"
    with open(tpl_path, "rb") as f:
        return f.read()


def _make_docx_with_empty_custom_properties_part() -> io.BytesIO:
    """Return a ``BytesIO`` of a .docx with an empty ``docProps/custom.xml`` part.

    Mirrors the layout produced by older python-docx releases (and by
    Microsoft Office in some templates) where a zero-property
    custom-properties part is materialised eagerly with a
    ``[Content_Types].xml`` override and a package-root rel pointing
    at it. Used to verify back-compat reading and round-trip fidelity
    in :class:`DescribeCustomPropertiesPartRelPlacement`.
    """
    blob = _make_valid_docx_bytes()
    out = io.BytesIO()
    empty_custom_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<Properties '
        b'xmlns="http://schemas.openxmlformats.org/officeDocument/'
        b'2006/custom-properties" '
        b'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/'
        b'2006/docPropsVTypes"/>'
    )
    custom_props_ct = (
        b'<Override PartName="/docProps/custom.xml" '
        b'ContentType="application/vnd.openxmlformats-officedocument.'
        b'custom-properties+xml"/>'
    )
    custom_props_rel = (
        b'<Relationship Id="rIdC1" '
        b'Type="http://schemas.openxmlformats.org/officeDocument/'
        b'2006/relationships/custom-properties" '
        b'Target="docProps/custom.xml"/>'
    )
    with zipfile.ZipFile(io.BytesIO(blob), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(b"</Types>", custom_props_ct + b"</Types>")
                elif item.filename == "_rels/.rels":
                    data = data.replace(
                        b"</Relationships>", custom_props_rel + b"</Relationships>"
                    )
                zout.writestr(item, data)
            zout.writestr("docProps/custom.xml", empty_custom_xml)
    out.seek(0)
    return out


def _make_empty_document_xml_docx_bytes() -> bytes:
    """Return a valid .docx whose `word/document.xml` is an empty byte string.

    Empty content is unrecoverable even with ``recover=True`` — forces the
    stub-element fallback in ``XmlPart.load``.
    """
    tpl_path = Path(__file__).parent.parent / "src" / "docx" / "templates" / "default.docx"
    with open(tpl_path, "rb") as f:
        blob = f.read()
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(blob), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = b""
                zout.writestr(item, data)
    return out.getvalue()


class DescribeDocument:
    """Unit-test suite for `docx.api.Document` factory function."""

    def it_opens_a_docx_file(self, Package_: Mock, document_: Mock):
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MAIN

        document = DocumentFactoryFn("foobar.docx")

        Package_.open.assert_called_once_with(
            "foobar.docx", recover=False, huge_tree=False, password=None, strict=False
        )
        assert document is document_

    def it_accepts_a_PathLike_docx_path(self, Package_: Mock, document_: Mock):
        # -- upstream-PR#1168: accept os.PathLike (e.g. pathlib.Path) --
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MAIN

        document = DocumentFactoryFn(Path("foobar.docx"))

        # -- os.fspath normalises the PathLike to str before delegating --
        Package_.open.assert_called_once_with(
            "foobar.docx", recover=False, huge_tree=False, password=None, strict=False
        )
        assert document is document_

    def it_opens_the_default_docx_if_none_specified(
        self, _default_docx_stream_: Mock, Package_: Mock, document_: Mock
    ):
        default_stream = io.BytesIO(b"fake-default-bytes")
        _default_docx_stream_.return_value = default_stream
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MAIN

        document = DocumentFactoryFn()

        Package_.open.assert_called_once_with(
            default_stream, recover=False, huge_tree=False, password=None, strict=False
        )
        assert document is document_

    def it_sources_the_default_docx_via_importlib_resources(self):
        # -- PyInstaller / cx_freeze / zipimport safety: must not rely on __file__ path.
        # -- Closes upstream#176, upstream-PR#1310, upstream-PR#177.
        from docx.api import _default_docx_stream

        data_stream = _default_docx_stream()

        assert isinstance(data_stream, io.BytesIO)
        # -- first four bytes of every .docx package are the PK\x03\x04 zip signature --
        assert data_stream.getvalue()[:4] == b"PK\x03\x04"

    def it_produces_a_usable_default_Document_instance(self):
        # -- round-trip sanity check: Document() with no arg yields a real Document --
        document = DocumentFactoryFn()

        assert isinstance(document, DocumentCls)

    def it_strips_metadata_when_include_metadata_is_False(self):
        # -- default template ships with Application, AppVersion, Template, etc.
        # -- baseline: with include_metadata=True (default), those survive --
        document = DocumentFactoryFn(include_metadata=False)

        # -- core properties cleared --
        assert document.core_properties.author == ""
        assert document.core_properties.title == ""
        assert document.core_properties.last_modified_by == ""
        assert document.core_properties.modified is None
        # -- extended properties cleared --
        assert document.extended_properties.application is None
        assert document.extended_properties.app_version is None
        assert document.extended_properties.template is None

    def it_keeps_metadata_by_default(self):
        document = DocumentFactoryFn()

        # -- the bundled template writes a known Application name --
        assert document.extended_properties.application is not None

    def it_opens_a_docm_file(self, Package_: Mock, document_: Mock):
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MACRO

        document = DocumentFactoryFn("foobar.docm")

        Package_.open.assert_called_once_with(
            "foobar.docm", recover=False, huge_tree=False, password=None, strict=False
        )
        assert document is document_

    def it_raises_on_not_a_Word_file(self, Package_: Mock):
        Package_.open.return_value.main_document_part.content_type = "BOGUS"

        with pytest.raises(ValueError, match="file 'foobar.xlsx' is not a Word file,"):
            DocumentFactoryFn("foobar.xlsx")

    def it_raises_EncryptedDocumentError_on_password_protected_path(self, tmp_path):
        encrypted_path = tmp_path / "encrypted.docx"
        encrypted_path.write_bytes(_OLE_SIGNATURE + b"\x00" * 512)

        with pytest.raises(EncryptedDocumentError, match="python-ooxml-crypto"):
            DocumentFactoryFn(str(encrypted_path))

    def it_raises_FileNotFoundError_on_missing_path(self, tmp_path):
        # -- upstream#1410: missing file must raise FileNotFoundError so it
        # -- behaves like a normal filesystem-missing error. --
        missing = str(tmp_path / "no-such-file.docx")

        with pytest.raises(FileNotFoundError):
            DocumentFactoryFn(missing)

    def it_raises_NotADocxError_on_non_zip_file(self, tmp_path):
        # -- upstream#1410: existing file that isn't a zip raises NotADocxError --
        from docx.opc.exceptions import NotADocxError

        plain = tmp_path / "plain.docx"
        plain.write_bytes(b"this is just text, not a zip")

        with pytest.raises(NotADocxError):
            DocumentFactoryFn(str(plain))

    def it_raises_EncryptedDocumentError_on_password_protected_stream(self):
        stream = io.BytesIO(_OLE_SIGNATURE + b"\x00" * 512)

        with pytest.raises(EncryptedDocumentError, match="password-protected"):
            DocumentFactoryFn(stream)

    def it_raises_on_malformed_document_xml_by_default(self):
        stream = io.BytesIO(_make_malformed_docx_bytes())

        with pytest.raises(etree.XMLSyntaxError):
            DocumentFactoryFn(stream)

    def it_opens_malformed_document_in_recover_mode(self):
        stream = io.BytesIO(_make_malformed_docx_bytes())

        document = DocumentFactoryFn(stream, recover=True)

        assert isinstance(document, DocumentCls)
        assert len(document.recovery_warnings) > 0
        assert all(isinstance(w, str) for w in document.recovery_warnings)

    def it_reports_no_warnings_for_valid_document_in_recover_mode(self):
        stream = io.BytesIO(_make_valid_docx_bytes())

        document = DocumentFactoryFn(stream, recover=True)

        assert document.recovery_warnings == []

    def it_recovery_mode_still_raises_for_invalid_zip(self, tmp_path):
        not_a_zip = tmp_path / "bogus.docx"
        not_a_zip.write_bytes(b"this is not a zip file")

        with pytest.raises(PackageNotFoundError):
            DocumentFactoryFn(str(not_a_zip), recover=True)

    def it_raises_PackageNotFoundError_on_zip_missing_content_types(self, tmp_path):
        # -- Regression for issue #172: a zip that happens to be a valid archive
        # -- but lacks `[Content_Types].xml` used to surface a bare
        # -- `KeyError("[Content_Types].xml")` from `zipfile.read`, which leaks
        # -- the internal file name and is hard to match on. Now wrapped in a
        # -- typed `PackageNotFoundError`. --
        bogus = tmp_path / "no-content-types.docx"
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr("word/document.xml", b"<doc/>")
        bogus.write_bytes(buf.getvalue())

        with pytest.raises(PackageNotFoundError, match=r"\[Content_Types\]\.xml"):
            DocumentFactoryFn(str(bogus))

    def it_raises_PackageNotFoundError_on_zip_missing_content_types_in_recover_mode(
        self, tmp_path
    ):
        # -- The wrapping happens at the OPC load boundary, before recovery mode
        # -- gets a chance to kick in. `PackageNotFoundError` must surface even
        # -- when the caller opts into `recover=True`. --
        bogus = tmp_path / "no-content-types.docx"
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr("word/document.xml", b"<doc/>")
        bogus.write_bytes(buf.getvalue())

        with pytest.raises(PackageNotFoundError):
            DocumentFactoryFn(str(bogus), recover=True)

    def it_recovery_mode_still_raises_for_encrypted_docx(self, tmp_path):
        encrypted_path = tmp_path / "encrypted.docx"
        encrypted_path.write_bytes(_OLE_SIGNATURE + b"\x00" * 512)

        with pytest.raises(EncryptedDocumentError):
            DocumentFactoryFn(str(encrypted_path), recover=True)

    def it_defaults_recover_to_False_for_valid_document(self):
        stream = io.BytesIO(_make_valid_docx_bytes())

        document = DocumentFactoryFn(stream)

        assert document.recovery_warnings == []

    def it_falls_back_to_stub_when_document_xml_is_empty(self):
        stream = io.BytesIO(_make_empty_document_xml_docx_bytes())

        document = DocumentFactoryFn(stream, recover=True)

        assert isinstance(document, DocumentCls)
        assert document.paragraphs == []
        assert len(document.recovery_warnings) >= 1

    def it_passes_recover_True_through_to_Package_open(
        self, Package_: Mock, document_: Mock
    ):
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MAIN

        DocumentFactoryFn("foobar.docx", recover=True)

        Package_.open.assert_called_once_with(
            "foobar.docx", recover=True, huge_tree=False, password=None, strict=False
        )

    def it_passes_huge_tree_True_through_to_Package_open(
        self, Package_: Mock, document_: Mock
    ):
        # -- upstream#1086: huge_tree=True must propagate to Package.open --
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MAIN

        DocumentFactoryFn("foobar.docx", huge_tree=True)

        Package_.open.assert_called_once_with(
            "foobar.docx", recover=False, huge_tree=True, password=None, strict=False
        )

    def it_defaults_huge_tree_to_False(self, Package_: Mock, document_: Mock):
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MAIN

        DocumentFactoryFn("foobar.docx")

        Package_.open.assert_called_once_with(
            "foobar.docx", recover=False, huge_tree=False, password=None, strict=False
        )

    def it_passes_password_through_to_Package_open(
        self, Package_: Mock, document_: Mock
    ):
        document_part = Package_.open.return_value.main_document_part
        document_part.document = document_
        document_part.content_type = CT.WML_DOCUMENT_MAIN

        DocumentFactoryFn("protected.docx", password="hunter2")

        Package_.open.assert_called_once_with(
            "protected.docx", recover=False, huge_tree=False, password="hunter2", strict=False
        )

    def it_ships_hanging_indents_on_List_Bullet_and_List_Number(self):
        # -- upstream#1443: default.docx used to omit hanging indents on these
        # -- list styles so Word-rendered bullets collided with paragraph text. --
        document = DocumentFactoryFn()

        for name in ("List Bullet", "List Number"):
            pf = document.styles[name].paragraph_format
            assert pf.left_indent is not None and pf.left_indent > 0, (
                f"style {name!r} has no left_indent"
            )
            assert pf.first_line_indent is not None and pf.first_line_indent < 0, (
                f"style {name!r} has no hanging (negative first-line) indent"
            )

    def it_preserves_List_Bullet_indents_after_round_trip(self):
        # -- upstream#1443: round-trip through save/open must preserve indents --
        document = DocumentFactoryFn()
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)

        reopened = DocumentFactoryFn(buf)

        pf = reopened.styles["List Bullet"].paragraph_format
        assert pf.left_indent is not None and pf.left_indent > 0
        assert pf.first_line_indent is not None and pf.first_line_indent < 0

    # -- fixtures --------------------------------------------------------------------------------

    @pytest.fixture
    def _default_docx_stream_(self, request: FixtureRequest):
        return function_mock(request, "docx.api._default_docx_stream")

    @pytest.fixture
    def document_(self, request: FixtureRequest):
        return instance_mock(request, DocumentCls)

    @pytest.fixture
    def Package_(self, request: FixtureRequest):
        return class_mock(request, "docx.api.Package")


class DescribeCustomPropertiesPartRelPlacement:
    """Regression suite for issues #712 and #721.

    Issue #712: when the custom-properties part **is** materialised, its
    relationship must hang off the package root (``_rels/.rels``), never
    off the main-document part (``word/_rels/document.xml.rels``). When
    the rel lands on the document part, Microsoft Word rejects the
    package as malformed even though python-docx itself can re-open it
    cleanly.

    Issue #721: a fresh ``Document().save()`` (or any save where no
    custom property was set) must **not** materialise an empty
    ``docProps/custom.xml`` part at all. Round-trip fidelity is
    preserved for files that shipped the empty part — a read+save of
    such a file keeps the part — but library-authored saves emit only
    what the user actually populated.
    """

    _CUSTOM_PROPS_TARGET = "docProps/custom.xml"

    def _save_and_inspect(self, document):
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        with zipfile.ZipFile(buf) as z:
            names = z.namelist()
            content_types = z.read("[Content_Types].xml").decode()
            package_root_rels = z.read("_rels/.rels").decode()
            document_rels = z.read("word/_rels/document.xml.rels").decode()
        return buf, names, content_types, package_root_rels, document_rels

    def it_does_not_write_custom_properties_part_for_a_fresh_document(self):
        # -- issue #721: a no-op save of a fresh ``Document()`` should
        # -- not produce a ``docProps/custom.xml`` part, content-type
        # -- override, or rels entry. --
        document = DocumentFactoryFn()

        _, names, content_types, root_rels, doc_rels = self._save_and_inspect(document)

        assert self._CUSTOM_PROPS_TARGET not in names
        assert self._CUSTOM_PROPS_TARGET not in content_types
        assert self._CUSTOM_PROPS_TARGET not in root_rels
        assert self._CUSTOM_PROPS_TARGET not in doc_rels

    def it_does_not_write_custom_properties_part_when_only_core_props_set(self):
        # -- mutating ``core_properties`` is unrelated to custom
        # -- properties; the empty-skip must still apply. --
        document = DocumentFactoryFn()
        document.core_properties.author = "x"

        _, names, content_types, root_rels, doc_rels = self._save_and_inspect(document)

        assert self._CUSTOM_PROPS_TARGET not in names
        assert self._CUSTOM_PROPS_TARGET not in content_types
        assert self._CUSTOM_PROPS_TARGET not in root_rels
        assert self._CUSTOM_PROPS_TARGET not in doc_rels

    def it_writes_custom_properties_part_when_a_custom_property_is_set(self):
        # -- the part must materialise (with the rel on the package
        # -- root, per issue #712) once at least one custom property
        # -- has been added. --
        document = DocumentFactoryFn()
        document.custom_properties.add("Project", "demo")

        _, names, content_types, root_rels, doc_rels = self._save_and_inspect(document)

        assert self._CUSTOM_PROPS_TARGET in names
        assert self._CUSTOM_PROPS_TARGET in content_types
        assert self._CUSTOM_PROPS_TARGET in root_rels
        assert self._CUSTOM_PROPS_TARGET not in doc_rels

    def it_round_trips_a_document_with_custom_properties(self):
        document = DocumentFactoryFn()
        document.custom_properties.add("Project", "demo")
        document.custom_properties.add("Iteration", 7)

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = DocumentFactoryFn(buf)

        assert reopened.custom_properties["Project"] == "demo"
        assert reopened.custom_properties["Iteration"] == 7

    def it_round_trips_through_save_and_reopen(self):
        # -- python-docx must still re-read its own output: save, reopen,
        # -- inspect custom_properties without raising. --
        document = DocumentFactoryFn()
        buf, _, _, _, _ = self._save_and_inspect(document)

        buf.seek(0)
        reopened = DocumentFactoryFn(buf)

        # -- no items but the proxy must be live and iterable. --
        assert list(reopened.custom_properties) == []

    def it_reads_a_document_with_an_empty_custom_properties_part(self):
        # -- back-compat: a docx that has the part declared but empty
        # -- (e.g. saved by an older python-docx, or by Office) must
        # -- still open cleanly. --
        buf = _make_docx_with_empty_custom_properties_part()

        document = DocumentFactoryFn(buf)

        assert list(document.custom_properties) == []

    def it_preserves_a_shipped_empty_custom_properties_part_on_round_trip(self):
        # -- issue #721 round-trip fidelity: a docx that arrived with an
        # -- empty custom-properties part keeps it across read+save so
        # -- byte-equal fixtures don't drift just because python-docx
        # -- opened them. --
        buf_in = _make_docx_with_empty_custom_properties_part()

        document = DocumentFactoryFn(buf_in)
        buf_out = io.BytesIO()
        document.save(buf_out)
        buf_out.seek(0)
        with zipfile.ZipFile(buf_out) as z:
            names_out = z.namelist()

        assert self._CUSTOM_PROPS_TARGET in names_out


class DescribePasswordRoundTrip:
    """Integration tests for encrypted Document open/save via ``python-ooxml-crypto``."""

    def _requires_ooxml_crypto(self):
        import importlib.util

        if importlib.util.find_spec("ooxml_crypto") is None:
            pytest.skip(
                "python-ooxml-crypto is not installed (optional dependency)"
            )

    def it_round_trips_through_a_stream(self):
        self._requires_ooxml_crypto()

        document = DocumentFactoryFn()
        document.add_paragraph("encrypted round-trip body")

        buf = io.BytesIO()
        document.save(buf, password="hunter2")

        # -- the saved bytes are a CFBF (OLE2) container, not a plain zip --
        assert buf.getvalue()[:8] == _OLE_SIGNATURE

        buf.seek(0)
        reopened = DocumentFactoryFn(buf, password="hunter2")

        texts = [p.text for p in reopened.paragraphs]
        assert "encrypted round-trip body" in texts

    def it_round_trips_through_a_path(self, tmp_path):
        self._requires_ooxml_crypto()

        document = DocumentFactoryFn()
        document.add_paragraph("encrypted round-trip body via path")

        out_path = tmp_path / "protected.docx"
        document.save(str(out_path), password="hunter2")

        # -- the saved bytes are a CFBF (OLE2) container, not a plain zip --
        with open(out_path, "rb") as f:
            assert f.read(8) == _OLE_SIGNATURE

        reopened = DocumentFactoryFn(str(out_path), password="hunter2")

        texts = [p.text for p in reopened.paragraphs]
        assert "encrypted round-trip body via path" in texts

    def it_raises_EncryptedDocumentError_with_wrong_password(self, tmp_path):
        self._requires_ooxml_crypto()

        document = DocumentFactoryFn()
        document.add_paragraph("wrong-password reject test")
        out_path = tmp_path / "protected.docx"
        document.save(str(out_path), password="correct")

        with pytest.raises(EncryptedDocumentError, match="password does not match"):
            DocumentFactoryFn(str(out_path), password="incorrect")

    def it_raises_EncryptedDocumentError_when_password_is_missing(self, tmp_path):
        self._requires_ooxml_crypto()

        document = DocumentFactoryFn()
        document.add_paragraph("missing-password reject test")
        out_path = tmp_path / "protected.docx"
        document.save(str(out_path), password="correct")

        with pytest.raises(EncryptedDocumentError, match="password-protected"):
            DocumentFactoryFn(str(out_path))

    def it_rejects_flat_opc_with_password(self, tmp_path):
        # -- flat_opc and password are mutually exclusive: Flat-OPC is not a zip. --
        document = DocumentFactoryFn()

        out_path = tmp_path / "protected.xml"
        with pytest.raises(ValueError, match="mutually exclusive"):
            document.save(str(out_path), flat_opc=True, password="hunter2")


class Describe_docx_pptx_import_order:
    """Round-12 regression: cross-package import order must not corrupt dispatch.

    Before the round-12 xmlchemy fix, importing ``pptx`` before ``docx`` would
    stack pptx's namespace registry into the ``ooxml_xmlchemy`` composite.
    Because the composite iterates sub-registries in reverse order, pptx's
    ``OxmlElement`` factory silently won for ``w:p`` lookups, producing a bare
    ``lxml._Element`` from pptx's ``element_class_lookup`` (which has no
    ``CT_P`` binding). ``Document.add_paragraph`` then crashed with
    ``AttributeError: 'lxml.etree._Element' object has no attribute 'add_r'``.

    The fix adds a module-path fallback to ``_cls_registry`` so ``CT_Body``
    (in ``docx.oxml.document``) routes through the docx registry regardless
    of how many sibling parent libraries are co-loaded.
    """

    def it_works_with_pptx_imported_before_docx(self):
        """Run in a clean subprocess so import order is controllable.

        Asserting import order inside the already-loaded pytest process is
        meaningless (both libraries are imported at collection time). The
        subprocess snapshots a realistic end-user scenario.
        """
        import subprocess
        import sys
        import textwrap

        # -- gate without importing pptx into the parent pytest process;
        # -- importing pptx here would stack its registry globally and
        # -- leak into sibling tests (only the subprocess must load pptx)
        if importlib.util.find_spec("pptx") is None:
            pytest.skip("python-pptx not installed")

        script = textwrap.dedent(
            """
            import pptx  # stack pptx registry first
            import docx
            d = docx.Document()
            p = d.add_paragraph("x")
            assert type(p).__name__ == "Paragraph"
            assert p.text == "x"
            """
        )
        result = subprocess.run(
            [sys.executable, "-c", script], capture_output=True, text=True
        )
        assert result.returncode == 0, (
            f"subprocess failed:\nstdout={result.stdout}\nstderr={result.stderr}"
        )

    def it_works_with_docx_imported_before_pptx(self):
        import subprocess
        import sys
        import textwrap

        # -- gate without importing pptx into the parent pytest process;
        # -- importing pptx here would stack its registry globally and
        # -- leak into sibling tests (only the subprocess must load pptx)
        if importlib.util.find_spec("pptx") is None:
            pytest.skip("python-pptx not installed")

        script = textwrap.dedent(
            """
            import docx  # stack docx registry first
            import pptx
            d = docx.Document()
            p = d.add_paragraph("x")
            assert type(p).__name__ == "Paragraph"
            assert p.text == "x"
            """
        )
        result = subprocess.run(
            [sys.executable, "-c", script], capture_output=True, text=True
        )
        assert result.returncode == 0, (
            f"subprocess failed:\nstdout={result.stdout}\nstderr={result.stderr}"
        )

