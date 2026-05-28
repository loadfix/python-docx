"""Tests for the minimal docx -> Markdown (GFM) exporter.

Covers the element-kind mappings declared in issue #74:

- heading paragraphs -> ``#`` .. ``######``
- bold / italic / inline-code runs -> ``**`` / ``_`` / ``` ` ```
- hyperlinks -> ``[text](url)``
- bullet / numbered lists -> ``- `` / ``1. ``
- tables -> GFM ``| ... |``
- block quotes -> ``> ``
- inline images -> ``![alt](archive-path)``
- page breaks -> ``---``
- footnotes -> ``[^N]`` with end-of-doc ``[^N]: text``
- round-trip a fixture and verify output
"""

from __future__ import annotations

import base64
import io

import pytest

from docx import Document
from docx.markdown_export import (
    _escape_inline,
    _wrap_code,
    document_to_markdown,
)


# -- helpers -----------------------------------------------------------------


def _new_doc():
    """Return a fresh blank |Document|."""
    return Document()


def _png_bytes() -> bytes:
    """Return the bytes of a 1x1 transparent PNG."""
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
        "+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )


# -- public API --------------------------------------------------------------


class DescribeDocumentToMarkdown:
    def it_returns_a_string(self):
        doc = _new_doc()
        doc.add_paragraph("hello")
        out = doc.to_markdown()
        assert isinstance(out, str)
        assert "hello" in out

    def it_terminates_output_with_a_newline(self):
        doc = _new_doc()
        doc.add_paragraph("hello")
        out = doc.to_markdown()
        assert out.endswith("\n")

    def it_returns_empty_string_for_an_empty_document(self):
        doc = _new_doc()
        # -- the default doc template carries one blank paragraph; the
        # -- exporter should not emit anything visible for it.
        out = doc.to_markdown()
        assert out == "" or out.strip() == ""


# -- headings ----------------------------------------------------------------


class DescribeHeadingExport:
    def it_promotes_Heading_1_to_a_single_hash(self):
        doc = _new_doc()
        doc.add_heading("Top", level=1)
        md = doc.to_markdown()
        assert "# Top" in md

    def it_promotes_Heading_3_to_three_hashes(self):
        doc = _new_doc()
        doc.add_heading("Deeper", level=3)
        md = doc.to_markdown()
        assert "### Deeper" in md

    def it_caps_levels_above_six_at_six_hashes(self):
        doc = _new_doc()
        doc.add_heading("Deepest", level=9)
        md = doc.to_markdown()
        assert "###### Deepest" in md
        assert "####### " not in md


# -- inline formatting -------------------------------------------------------


class DescribeBoldItalicCodeRuns:
    def it_emits_bold_runs_as_double_asterisks(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("hello")
        run.bold = True
        md = doc.to_markdown()
        assert "**hello**" in md

    def it_emits_italic_runs_as_underscores(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("hello")
        run.italic = True
        md = doc.to_markdown()
        assert "_hello_" in md

    def it_combines_bold_and_italic_as_nested_markers(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("combo")
        run.bold = True
        run.italic = True
        md = doc.to_markdown()
        assert "**_combo_**" in md

    def it_renders_a_monospace_run_as_inline_code(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        run = p.add_run("x = 1")
        run.font.name = "Consolas"
        md = doc.to_markdown()
        assert "`x = 1`" in md

    def it_widens_the_fence_when_the_code_text_contains_backticks(self):
        # -- direct helper test: input with one ``` `````` produces a
        # -- double-backtick fence so the inner backtick is preserved.
        wrapped = _wrap_code("a`b")
        assert wrapped.startswith("``") and wrapped.endswith("``")
        assert "a`b" in wrapped

    def it_pads_the_fence_when_text_starts_or_ends_with_a_backtick(self):
        wrapped = _wrap_code("`weird`")
        assert wrapped == "`` `weird` ``"


# -- escaping ----------------------------------------------------------------


class DescribeInlineEscape:
    @pytest.mark.parametrize(
        "raw, escaped",
        [
            ("a*b", "a\\*b"),
            ("c_d", "c\\_d"),
            ("[x]", "\\[x\\]"),
            ("backslash \\ here", "backslash \\\\ here"),
            ("{json}", "\\{json\\}"),
            ("tag <em>", "tag \\<em\\>"),
            ("ok text", "ok text"),
        ],
    )
    def it_escapes_GFM_significant_characters(self, raw, escaped):
        assert _escape_inline(raw) == escaped


# -- hyperlinks --------------------------------------------------------------


class DescribeHyperlinkExport:
    def it_renders_a_hyperlink_as_a_markdown_link(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_hyperlink("https://example.com/x", "click here", style=None)
        md = doc.to_markdown()
        assert "[click here](https://example.com/x)" in md

    def it_url_encodes_parens_in_the_link_target(self):
        doc = _new_doc()
        p = doc.add_paragraph()
        p.add_hyperlink("https://en.wikipedia.org/wiki/Foo_(bar)", "wiki", style=None)
        md = doc.to_markdown()
        # -- raw parens in the URL would terminate the link target
        assert "(https://en.wikipedia.org/wiki/Foo_%28bar%29)" in md


# -- lists -------------------------------------------------------------------


class DescribeListExport:
    def it_renders_decimal_numbering_as_ol_lines(self):
        doc = _new_doc()
        defn = doc.numbering.add_abstract_definition(format="decimal")
        p1 = doc.add_paragraph("first")
        p2 = doc.add_paragraph("second")
        defn.apply_to(p1, level=0)
        defn.apply_to(p2, level=0)
        md = doc.to_markdown()
        assert "1. first" in md
        assert "1. second" in md

    def it_renders_bullet_numbering_as_dashes(self):
        doc = _new_doc()
        defn = doc.numbering.add_abstract_definition(format="bullet", lvl_text="-")
        p1 = doc.add_paragraph("a")
        p2 = doc.add_paragraph("b")
        defn.apply_to(p1, level=0)
        defn.apply_to(p2, level=0)
        md = doc.to_markdown()
        assert "- a" in md
        assert "- b" in md

    def it_indents_nested_list_levels_with_two_spaces(self):
        doc = _new_doc()
        # -- multi-level numbering: level 0 + level 1 both bullets --
        defn = doc.numbering.add_numbering_definition(
            levels=[
                ("bullet", "-"),
                ("bullet", "-"),
            ]
        )
        p_outer = doc.add_paragraph("outer")
        p_inner = doc.add_paragraph("inner")
        defn.apply_to(p_outer, level=0)
        defn.apply_to(p_inner, level=1)
        md = doc.to_markdown()
        assert "- outer" in md
        assert "  - inner" in md


# -- tables ------------------------------------------------------------------


class DescribeTableExport:
    def it_renders_a_2x2_table_as_a_GFM_table(self):
        doc = _new_doc()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A1"
        table.rows[0].cells[1].text = "A2"
        table.rows[1].cells[0].text = "B1"
        table.rows[1].cells[1].text = "B2"
        md = doc.to_markdown()
        assert "| A1 | A2 |" in md
        assert "| --- | --- |" in md
        assert "| B1 | B2 |" in md

    def it_escapes_pipe_characters_inside_a_cell(self):
        doc = _new_doc()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "a|b"
        md = doc.to_markdown()
        assert r"a\|b" in md


# -- block quotes ------------------------------------------------------------


class DescribeQuoteExport:
    def it_prefixes_quote_styled_paragraphs_with_a_chevron(self):
        doc = _new_doc()
        p = doc.add_paragraph("a wise quote")
        try:
            p.style = "Quote"
        except KeyError:
            # -- "Quote" style isn't in the default template; create it.
            from docx.enum.style import WD_STYLE_TYPE

            doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
            p.style = "Quote"
        md = doc.to_markdown()
        assert "> a wise quote" in md


# -- images ------------------------------------------------------------------


class DescribeInlineImageExport:
    def it_emits_an_archive_relative_image_path(self, tmp_path):
        img_path = tmp_path / "pixel.png"
        img_path.write_bytes(_png_bytes())
        doc = _new_doc()
        doc.add_picture(str(img_path))
        md = doc.to_markdown()
        # -- the partname is /word/media/imageN.png
        assert "![" in md
        assert "](word/media/" in md


# -- page breaks -------------------------------------------------------------


class DescribePageBreakExport:
    def it_emits_a_horizontal_rule_for_a_hard_page_break(self):
        doc = _new_doc()
        doc.add_paragraph("before")
        # -- add_page_break inserts a paragraph whose only run is a w:br
        # -- with type=page; our exporter renders it as ---.
        doc.add_page_break()
        doc.add_paragraph("after")
        md = doc.to_markdown()
        assert "before" in md
        assert "---" in md
        assert "after" in md
        # -- the rule should appear between the two paragraphs --
        assert md.index("before") < md.index("---") < md.index("after")


# -- footnotes ---------------------------------------------------------------


class DescribeFootnoteExport:
    def it_emits_a_caret_reference_and_an_end_of_doc_definition(self):
        doc = _new_doc()
        p = doc.add_paragraph("Important claim")
        run = p.add_run(".")
        doc.footnotes.add(run, "Source: a study from 2024.")
        md = doc.to_markdown()
        # -- inline marker --
        assert "[^1]" in md
        # -- end-of-doc definition --
        assert "[^1]: Source: a study from 2024." in md

    def it_reuses_the_index_when_the_same_footnote_is_referenced_twice(self):
        doc = _new_doc()
        p1 = doc.add_paragraph("first ref")
        run1 = p1.add_run(".")
        fn = doc.footnotes.add(run1, "Shared note.")
        # -- a second reference to the same footnote_id, inserted manually --
        p2 = doc.add_paragraph("second ref")
        run2 = p2.add_run(".")
        run2._r.insert_footnote_reference(fn.footnote_id)
        md = doc.to_markdown()
        # -- both inline markers should be ``[^1]`` and only one definition emitted --
        assert md.count("[^1]") >= 3  # two inline + one definition
        assert md.count("[^1]: ") == 1


# -- end-to-end fixture ------------------------------------------------------


class DescribeFixtureRoundTrip:
    """Build a doc covering every required Markdown feature, save it, reopen it,
    and verify the to_markdown() output contains all the expected constructs.
    """

    def it_round_trips_a_kitchen_sink_fixture(self, tmp_path):
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # -- headings + paragraphs + inline formatting --
        doc.add_heading("Q1 Review", level=1)
        p = doc.add_paragraph()
        bold_run = p.add_run("Revenue grew 8.7% YoY")
        bold_run.bold = True

        # -- bullet list --
        bullets = doc.numbering.add_abstract_definition(format="bullet", lvl_text="-")
        amer = doc.add_paragraph("AMER: $14.2B")
        apac = doc.add_paragraph("APAC: $8.1B")
        bullets.apply_to(amer, level=0)
        bullets.apply_to(apac, level=0)

        # -- numbered list --
        numbers = doc.numbering.add_abstract_definition(format="decimal")
        n1 = doc.add_paragraph("First step")
        n2 = doc.add_paragraph("Second step")
        numbers.apply_to(n1, level=0)
        numbers.apply_to(n2, level=0)

        # -- hyperlink + italic + inline code --
        p2 = doc.add_paragraph()
        p2.add_hyperlink("https://example.com/report", "full report", style=None)
        p2.add_run(" with ")
        ital = p2.add_run("emphasis")
        ital.italic = True
        p2.add_run(" and ")
        code = p2.add_run("inline_code()")
        code.font.name = "Consolas"

        # -- block quote --
        if "Quote" not in [s.name for s in doc.styles]:
            doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
        q = doc.add_paragraph("a wise quote")
        q.style = "Quote"

        # -- image --
        img_path = tmp_path / "pixel.png"
        img_path.write_bytes(_png_bytes())
        doc.add_picture(str(img_path))

        # -- table --
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "Region"
        tbl.rows[0].cells[1].text = "Revenue"
        tbl.rows[1].cells[0].text = "AMER"
        tbl.rows[1].cells[1].text = "$14.2B"

        # -- page break --
        doc.add_page_break()

        # -- footnote --
        p3 = doc.add_paragraph("Important claim")
        run = p3.add_run(".")
        doc.footnotes.add(run, "Source: 2024 study.")

        # -- save and reopen, then convert --
        out_path = tmp_path / "kitchen-sink.docx"
        doc.save(str(out_path))

        reopened = Document(str(out_path))
        md = reopened.to_markdown()

        # -- the assertions below collectively cover every required
        # -- Markdown feature listed in the issue body. --
        assert "# Q1 Review" in md
        assert "**Revenue grew 8.7% YoY**" in md
        assert "- AMER: $14.2B" in md
        assert "- APAC: $8.1B" in md
        assert "1. First step" in md
        assert "1. Second step" in md
        assert "[full report](https://example.com/report)" in md
        assert "_emphasis_" in md
        assert "`inline_code()`" in md
        assert "> a wise quote" in md
        assert "![" in md and "word/media/" in md
        assert "| Region | Revenue |" in md
        assert "| --- | --- |" in md
        assert "| AMER | $14.2B |" in md
        assert "---" in md  # page-break thematic rule
        assert "[^1]" in md
        assert "[^1]: Source: 2024 study." in md


# -- entry-point parity ------------------------------------------------------


class DescribeEntryPoint:
    def it_exposes_document_to_markdown_as_a_module_helper(self):
        doc = _new_doc()
        doc.add_paragraph("hi")
        # -- both the method and the helper produce the same output --
        assert doc.to_markdown() == document_to_markdown(doc)
