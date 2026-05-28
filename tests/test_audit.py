"""Tests for `Document.audit_styles()` (issue #59)."""

from __future__ import annotations

import pytest

from docx import Document
from docx.audit import StyleAudit, StyleIssue, audit_styles
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


class DescribeStyleAudit:
    def it_returns_a_StyleAudit_instance_with_summary_and_issues(self):
        document = Document()
        document.add_paragraph("body")
        audit = document.audit_styles()
        assert isinstance(audit, StyleAudit)
        assert isinstance(audit.summary, dict)
        assert "total" in audit.summary
        assert isinstance(audit.issues, list)

    def it_flags_mixed_fonts_in_a_paragraph(self):
        document = Document()
        paragraph = document.add_paragraph("hello ")
        paragraph.runs[0].font.name = "Arial"
        run = paragraph.add_run("world")
        run.font.name = "Times New Roman"
        audit = document.audit_styles()
        mixed = audit.by_rule("mixed-fonts")
        assert len(mixed) == 1
        assert mixed[0].severity == "warning"

    def it_flags_a_heading_without_style(self):
        document = Document()
        paragraph = document.add_paragraph("My Big Section")
        paragraph.runs[0].bold = True
        audit = document.audit_styles()
        issues = audit.by_rule("heading-without-style")
        assert len(issues) == 1
        assert issues[0].severity == "error"

    def it_flags_orphan_paragraph_styles(self):
        document = Document()
        styles = document.styles
        new_style = styles.add_style("UnusedAlpha", WD_STYLE_TYPE.PARAGRAPH)
        new_style.font.name = "Arial"
        audit = document.audit_styles()
        orphans = audit.by_rule("orphan-style")
        names = {n for issue in orphans for n in issue.style_names}
        assert "UnusedAlpha" in names

    def it_flags_unstyled_paragraphs(self):
        document = Document()
        document.add_paragraph("a body line with no explicit style")
        audit = document.audit_styles()
        unstyled = audit.by_rule("unstyled-paragraph")
        assert len(unstyled) >= 1
        assert all(i.severity == "info" for i in unstyled)

    def it_does_not_flag_empty_paragraphs_as_unstyled(self):
        document = Document()
        document.add_paragraph("")
        audit = document.audit_styles()
        unstyled = audit.by_rule("unstyled-paragraph")
        assert all(
            i.paragraph_index is None
            or document.paragraphs[i.paragraph_index].text.strip()
            for i in unstyled
        )

    def it_can_iterate_and_len_the_audit(self):
        document = Document()
        document.add_paragraph("body")
        audit = document.audit_styles()
        assert len(audit) == sum(
            1 for _ in audit
        )

    def it_consolidates_styles_by_rewriting_paragraph_styles(self):
        document = Document()
        styles = document.styles
        styles.add_style("AltHeading", WD_STYLE_TYPE.PARAGRAPH)
        para1 = document.add_paragraph("Section A")
        para1.style = "AltHeading"
        para2 = document.add_paragraph("Section B")
        para2.style = "AltHeading"
        audit = document.audit_styles()
        rewritten = audit.consolidate_styles("Heading 1", drop=["AltHeading"])
        assert rewritten == 2
        # -- references rewritten
        assert document.paragraphs[0].style.name == "Heading 1"
        # -- style dropped
        assert "AltHeading" not in styles

    def it_raises_KeyError_when_canonical_style_missing(self):
        document = Document()
        audit = document.audit_styles()
        with pytest.raises(KeyError):
            audit.consolidate_styles("DefinitelyNotAStyle", drop=["Normal"])

    def it_can_be_built_from_the_module_function(self):
        document = Document()
        audit = audit_styles(document)
        assert isinstance(audit, StyleAudit)

    def it_summary_counts_issues_by_rule_id(self):
        document = Document()
        # -- mixed-fonts paragraph
        paragraph = document.add_paragraph("a")
        paragraph.runs[0].font.name = "Arial"
        paragraph.add_run("b").font.name = "Verdana"
        audit = document.audit_styles()
        assert audit.summary.get("mixed-fonts", 0) == 1
        assert audit.summary["total"] >= 1
