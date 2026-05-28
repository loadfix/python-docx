# pyright: reportPrivateUsage=false

"""Unit-test suite for `docx.content_controls` module."""

from __future__ import annotations

from typing import cast

from docx.content_controls import (
    BuildingBlockControl,
    CheckboxControl,
    ComboBoxControl,
    ContentControl,
    ContentControlType,
    DataBinding,
    DateControl,
    DropDownListControl,
    PictureControl,
    PlainTextControl,
    RepeatingSectionControl,
    RichTextControl,
    new_sdt,
)
from docx.oxml.content_controls import CT_Sdt
from docx.oxml.ns import qn

from .unitutil.cxml import element


class DescribeContentControl:
    """Unit-test suite for `docx.content_controls.ContentControl`."""

    def it_knows_its_tag(self):
        sdt = cast(
            CT_Sdt,
            element("w:sdt/w:sdtPr/w:tag{w:val=ABC}"),
        )
        cc = ContentControl(sdt)
        assert cc.tag == "ABC"

    def it_can_set_its_tag(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)
        cc.tag = "New"
        assert cc.tag == "New"

    def it_knows_its_title(self):
        sdt = cast(
            CT_Sdt,
            element("w:sdt/w:sdtPr/w:alias{w:val=Title}"),
        )
        cc = ContentControl(sdt)
        assert cc.title == "Title"

    def it_can_set_its_title(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)
        cc.title = "Hello"
        assert cc.title == "Hello"

    def it_reports_RICH_TEXT_when_no_marker_is_present(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr"))
        cc = ContentControl(sdt)
        assert cc.type is ContentControlType.RICH_TEXT

    def it_reports_PLAIN_TEXT_for_w_text_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:text"))
        cc = ContentControl(sdt)
        assert cc.type is ContentControlType.PLAIN_TEXT

    def it_reports_COMBO_BOX_for_w_comboBox_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:comboBox"))
        cc = ContentControl(sdt)
        assert cc.type is ContentControlType.COMBO_BOX

    def it_reports_DROPDOWN_for_w_dropDownList_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:dropDownList"))
        cc = ContentControl(sdt)
        assert cc.type is ContentControlType.DROPDOWN

    def it_reports_DATE_for_w_date_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:date"))
        cc = ContentControl(sdt)
        assert cc.type is ContentControlType.DATE

    def it_reports_PICTURE_for_w_picture_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:picture"))
        cc = ContentControl(sdt)
        assert cc.type is ContentControlType.PICTURE

    def it_reports_CHECKBOX_for_w14_checkbox_marker(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
                "<w:sdtPr><w14:checkbox/></w:sdtPr>"
                "</w:sdt>"
            ),
        )
        cc = ContentControl(sdt)
        assert cc.type is ContentControlType.CHECKBOX

    def it_reads_checkbox_checked_value(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
                "<w:sdtPr><w14:checkbox>"
                '<w14:checked w14:val="1"/>'
                "</w14:checkbox></w:sdtPr>"
                "</w:sdt>"
            ),
        )
        cc = ContentControl(sdt)
        assert cc.checked is True

    def it_can_round_trip_checkbox_checked(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)
        cc.checked = True
        assert cc.checked is True
        cc.checked = False
        assert cc.checked is False

    # -- data binding ----------------------------------------------------

    def it_returns_None_for_data_binding_when_no_dataBinding_child_present(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr"))
        cc = ContentControl(sdt)
        assert cc.data_binding is None

    def it_returns_None_for_data_binding_when_no_sdtPr_present(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)
        assert cc.data_binding is None

    def it_exposes_data_binding_when_dataBinding_child_is_present(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:sdtPr>"
                "<w:dataBinding"
                " w:prefixMappings=\"xmlns:ns0='http://example.com/ns'\""
                ' w:xpath="/ns0:root/ns0:name"'
                ' w:storeItemID="{12345678-1234-1234-1234-1234567890AB}"/>'
                "</w:sdtPr>"
                "</w:sdt>"
            ),
        )
        cc = ContentControl(sdt)

        db = cc.data_binding

        assert isinstance(db, DataBinding)
        assert db.prefix_mappings == "xmlns:ns0='http://example.com/ns'"
        assert db.xpath == "/ns0:root/ns0:name"
        assert db.store_item_id == "{12345678-1234-1234-1234-1234567890AB}"

    def it_can_set_data_binding_on_sdt_without_sdtPr(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)

        db = cc.set_data_binding(
            xpath="/root/child",
            prefix_mappings="xmlns:a='urn:a'",
            store_item_id="{AAAA-BBBB}",
        )

        assert isinstance(db, DataBinding)
        assert sdt.sdtPr is not None
        assert sdt.sdtPr.dataBinding is not None
        assert cc.data_binding is not None
        assert cc.data_binding.xpath == "/root/child"
        assert cc.data_binding.prefix_mappings == "xmlns:a='urn:a'"
        assert cc.data_binding.store_item_id == "{AAAA-BBBB}"

    def it_can_set_data_binding_with_default_prefix_mappings_and_no_store_id(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr"))
        cc = ContentControl(sdt)

        cc.set_data_binding("/Plain")

        assert cc.data_binding is not None
        assert cc.data_binding.xpath == "/Plain"
        assert cc.data_binding.prefix_mappings == ""
        assert cc.data_binding.store_item_id is None

    def it_overwrites_an_existing_data_binding_on_set(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:sdtPr>"
                '<w:dataBinding w:xpath="/old" w:storeItemID="{OLD}"/>'
                "</w:sdtPr>"
                "</w:sdt>"
            ),
        )
        cc = ContentControl(sdt)

        cc.set_data_binding(xpath="/new", store_item_id="{NEW}")

        assert cc.data_binding is not None
        assert cc.data_binding.xpath == "/new"
        assert cc.data_binding.store_item_id == "{NEW}"
        # -- still a single dataBinding child --
        sdtPr = sdt.sdtPr
        assert sdtPr is not None
        assert len(sdtPr.findall(qn("w:dataBinding"))) == 1

    def it_can_remove_a_data_binding(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:sdtPr>"
                '<w:dataBinding w:xpath="/old"/>'
                "</w:sdtPr>"
                "</w:sdt>"
            ),
        )
        cc = ContentControl(sdt)
        assert cc.data_binding is not None

        cc.remove_data_binding()

        assert cc.data_binding is None
        sdtPr = sdt.sdtPr
        assert sdtPr is not None
        assert sdtPr.find(qn("w:dataBinding")) is None

    def it_silently_ignores_remove_when_no_data_binding_is_present(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr"))
        cc = ContentControl(sdt)
        # -- should be a no-op --
        cc.remove_data_binding()
        assert cc.data_binding is None

    def it_silently_ignores_remove_when_no_sdtPr_is_present(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)
        cc.remove_data_binding()
        assert cc.data_binding is None

    def it_concatenates_text_from_sdtContent(self):
        sdt = cast(
            CT_Sdt,
            element('w:sdt/w:sdtContent/w:p/w:r/w:t"hello"'),
        )
        cc = ContentControl(sdt)
        assert cc.text == "hello"

    def it_can_set_text_on_inline_sdt(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtContent/w:r"))
        cc = ContentControl(sdt)
        cc.text = "value"
        assert cc.text == "value"
        # -- should have a w:r child in sdtContent --
        sdtContent = sdt.sdtContent
        assert sdtContent is not None
        assert sdtContent.find(qn("w:r")) is not None

    def it_can_set_text_on_block_sdt(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtContent/w:p"))
        cc = ContentControl(sdt)
        cc.text = "hi"
        assert cc.text == "hi"
        sdtContent = sdt.sdtContent
        assert sdtContent is not None
        assert sdtContent.find(qn("w:p")) is not None


class DescribeDataBinding:
    """Unit-test suite for `docx.content_controls.DataBinding`."""

    def it_reads_attribute_values_from_the_oxml_element(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:sdtPr>"
                '<w:dataBinding w:prefixMappings="pm" w:xpath="/x" w:storeItemID="ID"/>'
                "</w:sdtPr>"
                "</w:sdt>"
            ),
        )
        sdtPr = sdt.sdtPr
        assert sdtPr is not None
        dataBinding = sdtPr.dataBinding
        assert dataBinding is not None

        db = DataBinding(dataBinding)

        assert db.prefix_mappings == "pm"
        assert db.xpath == "/x"
        assert db.store_item_id == "ID"

    def it_returns_empty_strings_for_missing_xpath_and_prefix_mappings(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:dataBinding"))
        sdtPr = sdt.sdtPr
        assert sdtPr is not None
        dataBinding = sdtPr.dataBinding
        assert dataBinding is not None

        db = DataBinding(dataBinding)

        assert db.prefix_mappings == ""
        assert db.xpath == ""
        assert db.store_item_id is None

    def it_can_round_trip_attribute_values(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:dataBinding"))
        sdtPr = sdt.sdtPr
        assert sdtPr is not None
        dataBinding = sdtPr.dataBinding
        assert dataBinding is not None
        db = DataBinding(dataBinding)

        db.prefix_mappings = "xmlns:n='urn:n'"
        db.xpath = "/n:root"
        db.store_item_id = "{GUID}"

        assert db.prefix_mappings == "xmlns:n='urn:n'"
        assert db.xpath == "/n:root"
        assert db.store_item_id == "{GUID}"

    def it_clears_empty_string_attributes_when_set_to_empty(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:sdtPr>"
                '<w:dataBinding w:prefixMappings="foo" w:xpath="/x"/>'
                "</w:sdtPr>"
                "</w:sdt>"
            ),
        )
        sdtPr = sdt.sdtPr
        assert sdtPr is not None
        dataBinding = sdtPr.dataBinding
        assert dataBinding is not None
        db = DataBinding(dataBinding)

        db.prefix_mappings = ""
        db.xpath = ""

        assert dataBinding.prefixMappings is None
        assert dataBinding.xpath_val is None


class DescribeContentControlType:
    """Unit-test suite for `docx.content_controls.ContentControlType`."""

    def it_has_expected_members(self):
        assert ContentControlType.PLAIN_TEXT
        assert ContentControlType.RICH_TEXT
        assert ContentControlType.CHECKBOX
        assert ContentControlType.COMBO_BOX
        assert ContentControlType.DROPDOWN
        assert ContentControlType.DATE
        assert ContentControlType.PICTURE


class DescribeNewSdt:
    """Unit-test suite for the `new_sdt()` factory."""

    def it_creates_a_block_level_sdt_with_a_paragraph_child(self):
        sdt = new_sdt(ContentControlType.RICH_TEXT, tag="X", title="T", inline=False)
        assert sdt.sdtContent is not None
        assert sdt.sdtContent.find(qn("w:p")) is not None

    def it_creates_an_inline_sdt_with_a_run_child(self):
        sdt = new_sdt(ContentControlType.PLAIN_TEXT, tag="X", inline=True)
        assert sdt.sdtContent is not None
        assert sdt.sdtContent.find(qn("w:r")) is not None

    def it_sets_a_text_marker_for_PLAIN_TEXT(self):
        sdt = new_sdt(ContentControlType.PLAIN_TEXT, inline=True)
        assert sdt.type_marker_tag() == "w:text"

    def it_does_not_set_a_marker_for_RICH_TEXT(self):
        sdt = new_sdt(ContentControlType.RICH_TEXT, inline=False)
        assert sdt.type_marker_tag() is None

    def it_sets_the_tag_val_and_alias_val_when_provided(self):
        sdt = new_sdt(
            ContentControlType.RICH_TEXT, tag="TagVal", title="TitleVal", inline=False
        )
        assert sdt.tag_val == "TagVal"
        assert sdt.alias_val == "TitleVal"

    def it_always_sets_an_id(self):
        sdt = new_sdt(ContentControlType.RICH_TEXT, inline=False)
        assert isinstance(sdt.sdt_id, int)
        assert sdt.sdt_id is not None and sdt.sdt_id > 0


# ---------------------------------------------------------------------------
# Type-specific proxy subclass tests


class DescribeContentControl_proxy_for:
    """Unit-test suite for the ``ContentControl.proxy_for()`` dispatcher."""

    def it_returns_a_RichTextControl_when_no_marker_is_present(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, RichTextControl)

    def it_returns_a_PlainTextControl_for_w_text_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:text"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, PlainTextControl)

    def it_returns_a_DateControl_for_w_date_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:date"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, DateControl)

    def it_returns_a_PictureControl_for_w_picture_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:picture"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, PictureControl)

    def it_returns_a_ComboBoxControl_for_w_comboBox_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:comboBox"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, ComboBoxControl)

    def it_returns_a_DropDownListControl_for_w_dropDownList_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:dropDownList"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, DropDownListControl)

    def it_returns_a_BuildingBlockControl_for_w_docPartObj_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:docPartObj"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, BuildingBlockControl)

    def it_returns_a_BuildingBlockControl_for_w_docPartList_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:docPartList"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, BuildingBlockControl)

    def it_returns_a_RepeatingSectionControl_for_w15_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w15:repeatingSection"))
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, RepeatingSectionControl)

    def it_returns_a_CheckboxControl_for_w14_checkbox_marker(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
                "<w:sdtPr><w14:checkbox/></w:sdtPr>"
                "</w:sdt>"
            ),
        )
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, CheckboxControl)


class DescribeContentControl_lock:
    """Unit-test suite for ``ContentControl.lock``."""

    def it_returns_None_when_no_lock_is_set(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)
        assert cc.lock is None

    def it_reads_the_lock_value(self):
        sdt = cast(
            CT_Sdt,
            element("w:sdt/w:sdtPr/w:lock{w:val=sdtContentLocked}"),
        )
        cc = ContentControl(sdt)
        assert cc.lock == "sdtContentLocked"

    def it_can_set_the_lock_value(self):
        sdt = cast(CT_Sdt, element("w:sdt"))
        cc = ContentControl(sdt)
        cc.lock = "contentLocked"
        assert cc.lock == "contentLocked"

    def it_can_clear_the_lock_by_assigning_None(self):
        sdt = cast(
            CT_Sdt,
            element("w:sdt/w:sdtPr/w:lock{w:val=sdtContentLocked}"),
        )
        cc = ContentControl(sdt)
        cc.lock = None
        assert cc.lock is None


class DescribePlainTextControl:
    """Unit-test suite for :class:`PlainTextControl`."""

    def it_is_returned_for_a_plain_text_sdt(self):
        sdt = new_sdt(ContentControlType.PLAIN_TEXT, inline=True)
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, PlainTextControl)

    def it_reports_multiLine_False_by_default(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:text"))
        cc = PlainTextControl(sdt)
        assert cc.multi_line is False

    def it_reads_multiLine_when_set_to_1(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:text{w:multiLine=1}"))
        cc = PlainTextControl(sdt)
        assert cc.multi_line is True

    def it_can_round_trip_multi_line(self):
        sdt = new_sdt(ContentControlType.PLAIN_TEXT, inline=True)
        cc = PlainTextControl(sdt)
        assert cc.multi_line is False
        cc.multi_line = True
        assert cc.multi_line is True
        cc.multi_line = False
        assert cc.multi_line is False


class DescribeDateControl:
    """Unit-test suite for :class:`DateControl`."""

    def it_parses_fullDate_from_existing_date_marker(self):
        sdt = cast(
            CT_Sdt,
            element("w:sdt/w:sdtPr/w:date{w:fullDate=2026-05-09T00:00:00Z}"),
        )
        cc = DateControl(sdt)
        assert cc.full_date == "2026-05-09T00:00:00Z"

    def it_returns_None_for_fullDate_when_date_marker_absent(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr"))
        cc = DateControl(sdt)
        assert cc.full_date is None

    def it_can_create_and_round_trip_full_date(self):
        sdt = new_sdt(ContentControlType.DATE, inline=False)
        cc = DateControl(sdt)
        cc.full_date = "2026-05-09"
        assert cc.full_date == "2026-05-09"

    def it_can_set_the_date_format(self):
        sdt = new_sdt(ContentControlType.DATE, inline=False)
        cc = DateControl(sdt)
        cc.date_format = "yyyy-MM-dd"
        assert cc.date_format == "yyyy-MM-dd"


class DescribeDropDownListControl:
    """Unit-test suite for :class:`DropDownListControl`."""

    def it_returns_an_empty_items_list_when_no_listItems(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:dropDownList"))
        cc = DropDownListControl(sdt)
        assert cc.items == []

    def it_reads_displayText_from_listItems(self):
        sdt = cast(
            CT_Sdt,
            element(
                "w:sdt/w:sdtPr/w:dropDownList/("
                "w:listItem{w:displayText=A,w:value=a},"
                "w:listItem{w:displayText=B,w:value=b}"
                ")"
            ),
        )
        cc = DropDownListControl(sdt)
        assert cc.items == ["A", "B"]

    def it_can_set_the_items_list(self):
        sdt = new_sdt(ContentControlType.DROPDOWN, inline=False)
        cc = DropDownListControl(sdt)
        cc.items = ["Red", "Green", "Blue"]
        assert cc.items == ["Red", "Green", "Blue"]

    def it_can_append_an_item_with_add_item(self):
        sdt = new_sdt(ContentControlType.DROPDOWN, inline=False)
        cc = DropDownListControl(sdt)
        cc.add_item("Yellow", "y")
        assert cc.items == ["Yellow"]

    def it_round_trips_parse_then_serialize(self):
        from lxml import etree

        sdt = new_sdt(ContentControlType.DROPDOWN, inline=False)
        cc = DropDownListControl(sdt)
        cc.items = ["One", "Two"]
        xml = etree.tostring(sdt)
        assert b"w:listItem" in xml
        assert b'w:displayText="One"' in xml
        assert b'w:displayText="Two"' in xml


class DescribeComboBoxControl:
    """Unit-test suite for :class:`ComboBoxControl`."""

    def it_reads_the_lastValue_attribute(self):
        sdt = cast(
            CT_Sdt,
            element("w:sdt/w:sdtPr/w:comboBox{w:lastValue=Custom}"),
        )
        cc = ComboBoxControl(sdt)
        assert cc.last_value == "Custom"

    def it_can_set_the_lastValue(self):
        sdt = new_sdt(ContentControlType.COMBO_BOX, inline=False)
        cc = ComboBoxControl(sdt)
        cc.last_value = "FreeText"
        assert cc.last_value == "FreeText"

    def it_can_round_trip_items(self):
        sdt = new_sdt(ContentControlType.COMBO_BOX, inline=False)
        cc = ComboBoxControl(sdt)
        cc.items = ["Option 1", "Option 2"]
        assert cc.items == ["Option 1", "Option 2"]


class DescribeCheckboxControl:
    """Unit-test suite for :class:`CheckboxControl`."""

    def it_is_returned_by_the_dispatcher_for_checkbox_marker(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
                "<w:sdtPr><w14:checkbox/></w:sdtPr>"
                "</w:sdt>"
            ),
        )
        assert isinstance(ContentControl.proxy_for(sdt), CheckboxControl)

    def it_surfaces_the_checked_state(self):
        sdt = new_sdt(ContentControlType.CHECKBOX, inline=True)
        cc = CheckboxControl(sdt)
        cc.checked = True
        assert cc.checked is True
        cc.checked = False
        assert cc.checked is False


class DescribePictureControl:
    """Unit-test suite for :class:`PictureControl`."""

    def it_is_returned_by_the_dispatcher_for_picture_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:picture"))
        assert isinstance(ContentControl.proxy_for(sdt), PictureControl)

    def it_can_be_created_through_new_sdt(self):
        sdt = new_sdt(ContentControlType.PICTURE, inline=False)
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, PictureControl)
        assert cc.type is ContentControlType.PICTURE


class DescribeBuildingBlockControl:
    """Unit-test suite for :class:`BuildingBlockControl`."""

    def it_is_returned_by_the_dispatcher_for_docPartObj_marker(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w:docPartObj"))
        assert isinstance(ContentControl.proxy_for(sdt), BuildingBlockControl)

    def it_reads_gallery_and_category(self):
        sdt = cast(
            CT_Sdt,
            element(
                "w:sdt/w:sdtPr/w:docPartObj/("
                "w:docPartGallery{w:val=Cover Pages},"
                "w:docPartCategory{w:val=Built-In},"
                "w:docPartUnique"
                ")"
            ),
        )
        cc = BuildingBlockControl(sdt)
        assert cc.gallery == "Cover Pages"
        assert cc.category == "Built-In"
        assert cc.unique is True

    def it_can_round_trip_gallery_and_category(self):
        sdt = new_sdt(ContentControlType.BUILDING_BLOCK, inline=False)
        cc = BuildingBlockControl(sdt)
        cc.gallery = "Quick Parts"
        cc.category = "General"
        cc.unique = True
        assert cc.gallery == "Quick Parts"
        assert cc.category == "General"
        assert cc.unique is True
        cc.unique = False
        assert cc.unique is False


class DescribeRepeatingSectionControl:
    """Unit-test suite for :class:`RepeatingSectionControl`."""

    def it_is_returned_by_the_dispatcher_for_w15_repeatingSection(self):
        sdt = cast(CT_Sdt, element("w:sdt/w:sdtPr/w15:repeatingSection"))
        assert isinstance(ContentControl.proxy_for(sdt), RepeatingSectionControl)

    def it_reads_section_title(self):
        sdt = cast(
            CT_Sdt,
            element(
                "w:sdt/w:sdtPr/w15:repeatingSection{w15:sectionTitle=Rows}"
            ),
        )
        cc = RepeatingSectionControl(sdt)
        assert cc.section_title == "Rows"

    def it_can_round_trip_section_title(self):
        sdt = new_sdt(ContentControlType.REPEATING_SECTION, inline=False)
        cc = RepeatingSectionControl(sdt)
        cc.section_title = "LineItems"
        assert cc.section_title == "LineItems"

    def it_returns_an_empty_rows_list_when_no_items(self):
        sdt = new_sdt(ContentControlType.REPEATING_SECTION, inline=False)
        cc = RepeatingSectionControl(sdt)
        assert cc.rows == []

    def it_can_append_rows_and_expose_them(self):
        sdt = new_sdt(ContentControlType.REPEATING_SECTION, inline=False)
        cc = RepeatingSectionControl(sdt)
        cc.add_row()
        cc.add_row()
        cc.add_row()
        assert len(cc.rows) == 3
        for row in cc.rows:
            assert isinstance(row, ContentControl)

    def it_parses_preexisting_rows(self):
        from docx.oxml.parser import parse_xml

        sdt = cast(
            CT_Sdt,
            parse_xml(
                '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
                "<w:sdtPr><w15:repeatingSection/></w:sdtPr>"
                "<w:sdtContent>"
                "  <w:sdt><w:sdtPr><w15:repeatingSectionItem/></w:sdtPr>"
                "    <w:sdtContent><w:p/></w:sdtContent></w:sdt>"
                "  <w:sdt><w:sdtPr><w15:repeatingSectionItem/></w:sdtPr>"
                "    <w:sdtContent><w:p/></w:sdtContent></w:sdt>"
                "</w:sdtContent>"
                "</w:sdt>"
            ),
        )
        cc = RepeatingSectionControl(sdt)
        assert len(cc.rows) == 2


class DescribeContentControlType_extended:
    """Unit-test suite for the extended ``ContentControlType`` members."""

    def it_includes_REPEATING_SECTION_and_BUILDING_BLOCK(self):
        assert ContentControlType.REPEATING_SECTION
        assert ContentControlType.BUILDING_BLOCK

    def it_has_values_matching_the_marker_tag_local_names(self):
        assert ContentControlType.REPEATING_SECTION.value == "repeatingSection"
        assert ContentControlType.BUILDING_BLOCK.value == "docPartObj"


# ---------------------------------------------------------------------------
# Ergonomic authoring API (`build_text_control`, `add_text_control`,
# `add_repeating_section`)


class DescribeBuildTextControl:
    """Unit-test suite for :func:`docx.content_controls.build_text_control`."""

    def it_resolves_text_kind_string_to_PLAIN_TEXT(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="x")
        assert sdt.type_marker_tag() == "w:text"

    def it_resolves_rich_text_kind_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("rich-text", name="x")
        assert sdt.type_marker_tag() is None

    def it_resolves_dropdown_kind_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("dropdown", name="x")
        assert sdt.type_marker_tag() == "w:dropDownList"

    def it_resolves_combo_kind_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("combo", name="x")
        assert sdt.type_marker_tag() == "w:comboBox"

    def it_resolves_date_kind_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("date", name="x")
        assert sdt.type_marker_tag() == "w:date"

    def it_resolves_checkbox_kind_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("checkbox", name="x")
        assert sdt.type_marker_tag() == "w14:checkbox"

    def it_resolves_picture_kind_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("picture", name="x")
        assert sdt.type_marker_tag() == "w:picture"

    def it_resolves_repeating_section_kind_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("repeating-section", name="x")
        assert sdt.type_marker_tag() == "w15:repeatingSection"

    def it_accepts_a_ContentControlType_member_directly(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control(ContentControlType.DATE, name="x")
        assert sdt.type_marker_tag() == "w:date"

    def it_raises_ValueError_for_unknown_kind(self):
        from docx.content_controls import build_text_control

        try:
            build_text_control("widget")
        except ValueError as exc:
            assert "widget" in str(exc)
        else:
            raise AssertionError("expected ValueError")

    def it_assigns_the_name_to_the_tag_val(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="customer_name")
        assert sdt.tag_val == "customer_name"

    def it_uses_placeholder_as_alias_when_no_title_supplied(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="x", placeholder="Customer Name")
        assert sdt.alias_val == "Customer Name"

    def it_prefers_explicit_title_over_placeholder(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control(
            "text", name="x", placeholder="Pl", title="MyTitle"
        )
        assert sdt.alias_val == "MyTitle"

    def it_seeds_inline_content_with_value(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="x", value="Acme", inline=True)
        text = sdt.text
        assert text == "Acme"

    def it_seeds_block_content_with_value(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control(
            "rich-text", name="x", value="Hello", inline=False
        )
        assert sdt.text == "Hello"

    def it_seeds_with_placeholder_when_no_value(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control(
            "text", name="x", placeholder="Type here", inline=True
        )
        assert sdt.text == "Type here"

    def it_uses_kind_specific_default_placeholder_when_unset(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="x", inline=True)
        assert "text" in sdt.text.lower()  # "Click or tap here to enter text."

    def it_writes_lock_when_locked_True(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="x", locked=True)
        assert sdt.lock_val == "sdtLocked"

    def it_omits_lock_when_locked_False(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="x", locked=False)
        assert sdt.lock_val is None

    def it_accepts_an_explicit_lock_string(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("text", name="x", locked="contentLocked")
        assert sdt.lock_val == "contentLocked"

    def it_writes_a_data_binding_for_a_bare_property_name(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control(
            "text", name="customer_name", bind_to="CustomerName"
        )
        cc = ContentControl.proxy_for(sdt)
        assert cc.data_binding is not None
        assert "CustomerName" in cc.data_binding.xpath
        assert cc.data_binding.prefix_mappings  # non-empty mappings emitted

    def it_writes_a_data_binding_for_a_verbatim_xpath(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control(
            "text", name="x", bind_to="/ns0:foo[1]/ns0:bar[1]"
        )
        cc = ContentControl.proxy_for(sdt)
        assert cc.data_binding is not None
        assert cc.data_binding.xpath == "/ns0:foo[1]/ns0:bar[1]"

    def it_seeds_a_checkbox_value_as_a_check_state(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("checkbox", name="approved", value=True)
        assert ContentControl.proxy_for(sdt).checked is True

    def it_seeds_dropdown_items_from_items_kw(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control(
            "dropdown", name="priority", items=["Low", "Med", "High"]
        )
        cc = ContentControl.proxy_for(sdt)
        assert cc.items == ["Low", "Med", "High"]  # type: ignore[union-attr]


class DescribeBuildTextControl_RoundTrip:
    """Round-trip fidelity for :func:`build_text_control` output.

    Each kind is built, dropped into a fresh :class:`docx.Document`, the
    document is saved to a stream, and the stream is reopened. The
    expectation is that every kind survives load → save → load
    byte-identically (the conformance harness's contract — relaxed only
    when Word's own emission deliberately drifts).
    """

    @staticmethod
    def _round_trip(populate):
        import io

        from docx import Document

        doc = Document()
        populate(doc)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)
        buf2 = io.BytesIO()
        doc2.save(buf2)
        # -- read once more and confirm bytes don't drift --
        return doc2, buf.getvalue() == buf2.getvalue()

    def it_round_trips_a_text_control(self):
        def populate(doc):
            doc.add_paragraph("intro").add_text_control(
                name="customer", placeholder="Name", value="Acme"
            )

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.tag == "customer"
        assert cc.text == "Acme"

    def it_round_trips_a_rich_text_block_control(self):
        def populate(doc):
            doc.add_text_control(
                kind="rich-text", name="exec_summary", placeholder="…"
            )

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.content_controls[0]
        assert cc.tag == "exec_summary"
        assert cc.type is ContentControlType.RICH_TEXT

    def it_round_trips_a_dropdown_with_items(self):
        def populate(doc):
            doc.add_paragraph().add_text_control(
                kind="dropdown",
                name="priority",
                items=["Low", "Medium", "High"],
                value="Medium",
            )

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.type is ContentControlType.DROPDOWN
        assert cc.items == ["Low", "Medium", "High"]  # type: ignore[union-attr]

    def it_round_trips_a_combo_box(self):
        def populate(doc):
            doc.add_paragraph().add_text_control(
                kind="combo",
                name="region",
                items=["EU", "US"],
            )

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.type is ContentControlType.COMBO_BOX

    def it_round_trips_a_date_control(self):
        def populate(doc):
            doc.add_paragraph().add_text_control(kind="date", name="due")

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.type is ContentControlType.DATE

    def it_round_trips_a_checkbox(self):
        def populate(doc):
            doc.add_paragraph().add_text_control(
                kind="checkbox", name="approved", value=True
            )

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.type is ContentControlType.CHECKBOX
        assert cc.checked is True

    def it_round_trips_a_picture_control(self):
        def populate(doc):
            doc.add_paragraph().add_text_control(kind="picture", name="logo")

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.type is ContentControlType.PICTURE

    def it_round_trips_a_repeating_section(self):
        def populate(doc):
            sec = doc.add_repeating_section(
                name="line_items",
                schema={"description": "text", "quantity": "number"},
            )
            sec.add({"description": "Widget", "quantity": "5"})
            sec.add({"description": "Gadget", "quantity": "3"})

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.content_controls[0]
        assert cc.type is ContentControlType.REPEATING_SECTION
        # -- two repeating-section item rows survived --
        assert len(cc.rows) == 2  # type: ignore[union-attr]

    def it_round_trips_a_locked_control(self):
        def populate(doc):
            doc.add_paragraph().add_text_control(
                name="x", value="locked", locked=True
            )

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.lock == "sdtLocked"

    def it_round_trips_a_bound_control(self):
        def populate(doc):
            doc.add_paragraph().add_text_control(
                name="customer", bind_to="CustomerName"
            )

        doc, identical = self._round_trip(populate)
        assert identical
        cc = doc.paragraphs[0].content_controls[0]
        assert cc.data_binding is not None
        assert "CustomerName" in cc.data_binding.xpath


class DescribeParagraph_add_text_control:
    """Unit-test suite for :meth:`docx.text.paragraph.Paragraph.add_text_control`."""

    def it_appends_an_inline_sdt(self):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("Dear ")
        cc = para.add_text_control(
            name="customer_name", placeholder="Customer", value="Acme"
        )
        assert isinstance(cc, ContentControl)
        assert cc.tag == "customer_name"
        assert cc.text == "Acme"
        assert len(para.content_controls) == 1

    def it_supports_chaining_after_an_add_run(self):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph("Dear ")
        para.add_text_control(name="customer_name", value="Acme")
        para.add_run(",")
        # -- verify run order: <w:r>"Dear "</w:r> <w:sdt> <w:r>","</w:r> --
        children = [c.tag for c in para._p]  # type: ignore[attr-defined]
        from docx.oxml.ns import qn

        assert children.count(qn("w:r")) == 2
        assert qn("w:sdt") in children

    def it_supports_lock_True_to_block_deletion(self):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        cc = para.add_text_control(name="x", value="v", locked=True)
        assert cc.lock == "sdtLocked"

    def it_supports_bind_to_keyword(self):
        from docx import Document

        doc = Document()
        para = doc.add_paragraph()
        cc = para.add_text_control(name="x", bind_to="CustomerName")
        assert cc.data_binding is not None


class DescribeDocument_add_text_control:
    """Unit-test suite for :meth:`docx.document.Document.add_text_control`."""

    def it_appends_a_block_level_sdt(self):
        from docx import Document

        doc = Document()
        cc = doc.add_text_control(
            kind="rich-text", name="exec_summary", placeholder="…"
        )
        assert cc.tag == "exec_summary"
        assert cc in doc.content_controls or doc.content_controls

    def it_defaults_to_rich_text_kind(self):
        from docx import Document

        doc = Document()
        cc = doc.add_text_control(name="x")
        assert cc.type is ContentControlType.RICH_TEXT


class DescribeDocument_add_repeating_section:
    """Unit-test suite for :meth:`Document.add_repeating_section`."""

    def it_creates_a_block_repeating_section_sdt(self):
        from docx import Document

        doc = Document()
        sec = doc.add_repeating_section(name="line_items")
        assert isinstance(sec, RepeatingSectionControl)
        assert sec.tag == "line_items"
        assert sec.type is ContentControlType.REPEATING_SECTION

    def it_writes_a_section_title_when_supplied(self):
        from docx import Document

        doc = Document()
        sec = doc.add_repeating_section(
            name="li", section_title="Line Items"
        )
        assert sec.section_title == "Line Items"

    def it_appends_rows_via_add_with_a_schema(self):
        from docx import Document

        doc = Document()
        sec = doc.add_repeating_section(
            name="li",
            schema={"description": "text", "quantity": "number"},
        )
        sec.add({"description": "Widget", "quantity": "5"})
        sec.add({"description": "Gadget", "quantity": "3"})
        assert len(sec.rows) == 2

    def it_appends_rows_via_add_without_a_schema(self):
        from docx import Document

        doc = Document()
        sec = doc.add_repeating_section(name="li")
        sec.add("Widget")
        sec.add("Gadget")
        assert len(sec.rows) == 2


class DescribeRepeatingSectionControl_add:
    """Unit-test suite for the ergonomic ``add()`` and ``set_schema()``
    methods patched onto :class:`RepeatingSectionControl`.
    """

    def it_seeds_per_field_inner_sdts_when_a_schema_is_set(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("repeating-section", name="li", inline=False)
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, RepeatingSectionControl)
        cc.set_schema({"name": "text", "qty": "number"})  # type: ignore[attr-defined]
        cc.add({"name": "Foo", "qty": "1"})  # type: ignore[attr-defined]
        # -- one outer row plus two inner per-field SDTs in the row's <w:p> --
        rows = cc.rows
        assert len(rows) == 1
        # -- each row contains a paragraph holding two inner sdts --
        row_sdt = rows[0].element
        from docx.oxml.ns import qn

        inner_p = row_sdt.find(qn("w:sdtContent")).find(qn("w:p"))
        inner_sdts = inner_p.findall(qn("w:sdt"))
        assert len(inner_sdts) == 2

    def it_falls_back_to_a_text_run_when_no_schema_is_set(self):
        from docx.content_controls import build_text_control

        sdt = build_text_control("repeating-section", name="li", inline=False)
        cc = ContentControl.proxy_for(sdt)
        assert isinstance(cc, RepeatingSectionControl)
        cc.add("Hello")  # type: ignore[attr-defined]
        rows = cc.rows
        assert len(rows) == 1
        assert "Hello" in rows[0].text
