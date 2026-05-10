"""Unit-test suite for :mod:`docx.latex_math`.

Exercises the minimal LaTeX-to-OMML translator: each supported
construct round-trips through a Document save/reload, every
unsupported construct raises :class:`NotImplementedError`, and
malformed LaTeX raises :class:`~docx.latex_math.LatexMathError`.
"""

from __future__ import annotations

import io
from typing import Optional, cast

import pytest

from docx import Document
from docx.latex_math import LatexMathError, latex_to_omml
from docx.math import (
    Delimiter,
    EqArray,
    Fraction,
    MathExpr,
    Radical,
    Sub,
    SubSup,
    Sup,
    from_element,
)
from docx.oxml.math import CT_OMath


# ---------------------------------------------------------------------------
# latex_to_omml — standalone translator
# ---------------------------------------------------------------------------


class DescribeLatexToOmml:
    """`latex_to_omml` returns a CT_OMath element for each supported form."""

    # -- type guards ------------------------------------------------------

    def it_returns_a_CT_OMath_element(self):
        el = latex_to_omml("x")

        assert isinstance(el, CT_OMath)

    def it_rejects_non_string_input(self):
        with pytest.raises(TypeError, match="must be a str"):
            latex_to_omml(42)  # type: ignore[arg-type]

    # -- leaves -----------------------------------------------------------

    def it_translates_a_single_variable(self):
        el = latex_to_omml("x")

        assert el.text == "x"

    def it_translates_a_numeric_literal(self):
        el = latex_to_omml("42")

        assert el.text == "42"

    def it_translates_a_decimal_literal(self):
        el = latex_to_omml("3.14")

        assert el.text == "3.14"

    # -- binary operators -------------------------------------------------

    @pytest.mark.parametrize("op", ["+", "-", "*", "/"])
    def it_translates_each_supported_binary_operator(self, op: str):
        el = latex_to_omml(f"x {op} y")

        assert el.text == f"x{op}y"

    def it_translates_equality(self):
        el = latex_to_omml("x = y")

        assert el.text == "x=y"

    # -- superscript ------------------------------------------------------

    def it_translates_a_single_char_superscript(self):
        el = latex_to_omml("x^2")

        expr = _proxy(el)
        assert isinstance(expr.children[0], Sup)
        assert expr.text == "x2"

    def it_translates_a_braced_superscript(self):
        el = latex_to_omml("x^{i+1}")

        expr = _proxy(el)
        sup = expr.children[0]
        assert isinstance(sup, Sup)
        assert sup.text == "xi+1"

    # -- subscript --------------------------------------------------------

    def it_translates_a_single_char_subscript(self):
        el = latex_to_omml("x_i")

        expr = _proxy(el)
        assert isinstance(expr.children[0], Sub)

    def it_translates_a_braced_subscript(self):
        el = latex_to_omml("x_{ij}")

        expr = _proxy(el)
        sub = expr.children[0]
        assert isinstance(sub, Sub)
        assert sub.text == "xij"

    def it_translates_simultaneous_sub_and_superscript(self):
        el = latex_to_omml("x_i^2")

        expr = _proxy(el)
        assert isinstance(expr.children[0], SubSup)

    # -- fractions --------------------------------------------------------

    def it_translates_a_simple_fraction(self):
        el = latex_to_omml(r"\frac{a}{b}")

        expr = _proxy(el)
        f = expr.children[0]
        assert isinstance(f, Fraction)

    def it_translates_a_fraction_with_compound_arguments(self):
        el = latex_to_omml(r"\frac{a+b}{c-d}")

        expr = _proxy(el)
        f = expr.children[0]
        assert isinstance(f, Fraction)
        assert "a+b" in f.text and "c-d" in f.text

    # -- radicals ---------------------------------------------------------

    def it_translates_a_square_root(self):
        el = latex_to_omml(r"\sqrt{x}")

        expr = _proxy(el)
        assert isinstance(expr.children[0], Radical)

    def it_translates_a_square_root_of_a_compound_expression(self):
        el = latex_to_omml(r"\sqrt{a^2+b^2}")

        expr = _proxy(el)
        rad = expr.children[0]
        assert isinstance(rad, Radical)

    # -- parentheses ------------------------------------------------------

    def it_translates_parenthesised_expressions_as_delimiters(self):
        el = latex_to_omml(r"(a+b)")

        expr = _proxy(el)
        d = expr.children[0]
        assert isinstance(d, Delimiter)
        assert d.begin == "(" and d.end == ")"

    # -- Greek ------------------------------------------------------------

    def it_translates_lower_case_greek_letters(self):
        el = latex_to_omml(r"\alpha + \beta")

        assert "α" in el.text and "β" in el.text

    def it_translates_upper_case_greek_letters(self):
        el = latex_to_omml(r"\Omega + \Gamma")

        assert "Ω" in el.text and "Γ" in el.text

    def it_translates_pi_as_a_greek_letter(self):
        el = latex_to_omml(r"\pi")

        assert "π" in el.text

    # -- align environment ------------------------------------------------

    def it_translates_a_two_row_align_environment(self):
        el = latex_to_omml(
            r"\begin{align} x + y \\ a - b \end{align}"
        )

        expr = _proxy(el)
        arr = expr.children[0]
        assert isinstance(arr, EqArray)
        assert arr.n_rows == 2

    def it_accepts_the_aligned_environment_alias(self):
        el = latex_to_omml(r"\begin{aligned} x \\ y \end{aligned}")

        expr = _proxy(el)
        arr = expr.children[0]
        assert isinstance(arr, EqArray)
        assert arr.n_rows == 2

    # -- composite — the headline example --------------------------------

    def it_translates_Euler_identity(self):
        el = latex_to_omml(r"e^{i \pi} + 1 = 0")

        expr = _proxy(el)
        # First child is e^{iπ}
        assert isinstance(expr.children[0], Sup)
        assert "π" in el.text
        assert el.text == "eiπ+1=0"

    # -- round-trip -------------------------------------------------------

    @pytest.mark.parametrize(
        "src",
        [
            "x + y",
            r"\frac{1}{2}",
            r"x_i^2",
            r"\sqrt{a+b}",
            r"\alpha \beta \gamma",
            r"(x+1)",
            r"\begin{align} x \\ y \end{align}",
        ],
    )
    def it_round_trips_supported_constructs_through_a_document(self, src: str):
        doc = Document()
        p = doc.add_paragraph()
        p.add_math_from_latex(src)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        reopened = Document(buf)
        eqs = list(reopened.paragraphs[0].equations)
        assert len(eqs) == 1


# ---------------------------------------------------------------------------
# Error paths
# ---------------------------------------------------------------------------


class DescribeLatexToOmmlUnsupported:
    """Unsupported (but well-formed) LaTeX raises :class:`NotImplementedError`."""

    @pytest.mark.parametrize(
        "src",
        [
            r"\int_0^1 x",
            r"\lim_{x \to 0} x",
            r"\sum_{i=1}^n i",
            r"\begin{matrix} a & b \\ c & d \end{matrix}",
            r"\begin{pmatrix} 1 \\ 2 \end{pmatrix}",
            r"x \, y",  # thin-space control symbol — unsupported
            r"\mathbb{R}",
            r"\hat{x}",
        ],
    )
    def it_raises_NotImplementedError_for_unsupported_constructs(
        self, src: str
    ):
        with pytest.raises(NotImplementedError) as excinfo:
            latex_to_omml(src)

        # Every NotImplementedError points back at the supported-subset
        # docstring so users know where to look.
        assert "docx.latex_math" in str(excinfo.value) or "supported" in str(
            excinfo.value
        )

    def it_raises_NotImplementedError_for_stray_punctuation(self):
        # ``,`` isn't one of our supported characters — this is well-formed
        # LaTeX but the translator doesn't model it.
        with pytest.raises(NotImplementedError):
            latex_to_omml("x, y")


class DescribeLatexToOmmlMalformed:
    """Malformed LaTeX raises :class:`LatexMathError`."""

    def it_rejects_unbalanced_opening_brace(self):
        with pytest.raises(LatexMathError):
            latex_to_omml(r"{a+b")

    def it_rejects_unterminated_frac(self):
        with pytest.raises(LatexMathError):
            latex_to_omml(r"\frac{a")

    def it_rejects_frac_with_a_missing_argument_group(self):
        # After the first {a}, the parser expects {den} — gets EOF.
        with pytest.raises(LatexMathError):
            latex_to_omml(r"\frac{a}")

    def it_rejects_lone_double_backslash_outside_an_environment(self):
        with pytest.raises(LatexMathError):
            latex_to_omml(r"a \\ b")

    def it_rejects_a_mismatched_environment_footer(self):
        with pytest.raises(LatexMathError):
            latex_to_omml(
                r"\begin{align} x \end{aligned}"
            )

    def it_rejects_content_after_a_complete_environment(self):
        with pytest.raises(LatexMathError):
            latex_to_omml(
                r"\begin{align} x \end{align} y"
            )


# ---------------------------------------------------------------------------
# Paragraph.add_math_from_latex — wrapper method
# ---------------------------------------------------------------------------


class DescribeParagraphAddMathFromLatex:
    """`Paragraph.add_math_from_latex` wraps `latex_to_omml`."""

    def it_appends_an_oMath_element_to_the_paragraph(self):
        doc = Document()
        p = doc.add_paragraph()

        p.add_math_from_latex(r"\frac{a}{b}")

        assert len(list(p.equations)) == 1

    def it_returns_a_MathExpr_proxy(self):
        doc = Document()
        p = doc.add_paragraph()

        expr = p.add_math_from_latex("x + y")

        assert isinstance(expr, MathExpr)

    def it_propagates_unsupported_errors_from_the_translator(self):
        doc = Document()
        p = doc.add_paragraph()

        with pytest.raises(NotImplementedError):
            p.add_math_from_latex(r"\int x")

    def it_propagates_malformed_errors_from_the_translator(self):
        doc = Document()
        p = doc.add_paragraph()

        with pytest.raises(LatexMathError):
            p.add_math_from_latex(r"\frac{a")


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _proxy(el: CT_OMath) -> object:
    """Return the :class:`~docx.math.oMath` proxy over *el*."""
    # pyright: reportUnknownMemberType=false
    return from_element(el)
