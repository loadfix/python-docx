"""Unit-test suite for `FontTable.add_embedded_font` (upstream#1231, #1307).

These tests exercise the authoring surface for embedded fonts introduced in
1.3.0 — `FontTable.add_embedded_font(path, family=...)` — plus the supporting
`FontTablePart.default`/`FontTablePart.add_font_part` factories and the
`Document.font_table_or_new` shortcut.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import cast

import pytest

import docx
from docx.font_obfuscation import (
    OBFUSCATED_FONT_CONTENT_TYPE,
    deobfuscate_font_bytes,
)
from docx.font_table import FontTable
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.font_table import CT_Fonts
from docx.package import Package
from docx.parts.font_table import FontPart, FontTablePart

from .unitutil.cxml import element


# -- a tiny binary blob stands in for a real TrueType font — python-docx never
# -- parses the bytes, it just packages them, so any blob round-trips cleanly --
_FAKE_FONT_BLOB = (
    b"OTTO\x00\x00\x00\x00" + b"\x00" * 128 + b"python-docx embedded-font fixture"
)


@pytest.fixture
def fake_font_path(tmp_path: Path) -> Path:
    path = tmp_path / "FakeFont.ttf"
    path.write_bytes(_FAKE_FONT_BLOB)
    return path


class DescribeFontTablePart_Default:
    """`FontTablePart.default` returns a newly-created empty part."""

    def it_creates_an_empty_fontTable_part(self):
        package = Package()

        part = FontTablePart.default(package)

        assert isinstance(part, FontTablePart)
        assert part.partname == "/word/fontTable.xml"
        assert part.content_type == CT.WML_FONT_TABLE
        assert len(part.font_table) == 0


class DescribeFontTable_AddEmbeddedFont:
    """End-to-end coverage for `FontTable.add_embedded_font`."""

    def it_creates_a_font_entry_with_an_embedRegular_rel(self, fake_font_path: Path):
        package = Package()
        ft_part = FontTablePart.default(package)

        metadata = ft_part.font_table.add_embedded_font(fake_font_path)

        assert metadata.name == "FakeFont"
        assert metadata.embed_regular is True
        # -- the FontTable sees the new entry --
        assert "FakeFont" in ft_part.font_table

    def it_wires_an_r_font_relationship_to_a_new_FontPart(self, fake_font_path: Path):
        package = Package()
        ft_part = FontTablePart.default(package)

        ft_part.font_table.add_embedded_font(fake_font_path)

        font_rels = [rel for rel in ft_part.rels.values() if rel.reltype == RT.FONT]
        assert len(font_rels) == 1
        target = font_rels[0].target_part
        assert isinstance(target, FontPart)
        assert target.blob == _FAKE_FONT_BLOB

    def it_rejects_unknown_family_variants(self, fake_font_path: Path):
        package = Package()
        ft_part = FontTablePart.default(package)

        with pytest.raises(ValueError, match="family must be one of"):
            ft_part.font_table.add_embedded_font(fake_font_path, family="oblique")  # type: ignore[arg-type]

    def it_supports_all_four_font_variants(self, fake_font_path: Path):
        package = Package()
        ft_part = FontTablePart.default(package)
        font_table = ft_part.font_table

        for family in ("regular", "bold", "italic", "bold_italic"):
            font_table.add_embedded_font(fake_font_path, family=family, name="All")

        metadata = font_table["All"]
        assert metadata.embed_regular is True
        assert metadata.embed_bold is True
        assert metadata.embed_italic is True
        assert metadata.embed_bold_italic is True


class DescribeDocument_FontTableOrNew:
    """`Document.font_table_or_new` materialises a fontTable.xml on demand."""

    def it_creates_the_part_when_missing_and_roundtrips_after_save(
        self, fake_font_path: Path
    ):
        document = docx.Document()

        # -- font_table_or_new always returns a live FontTable (creating one
        # -- on demand if needed); the bundled template ships with a
        # -- fontTable.xml so in this case it reuses the existing part. --
        font_table = document.font_table_or_new
        font_table.add_embedded_font(fake_font_path)

        # -- save then reopen to verify round-trip preservation --
        buf = io.BytesIO()
        document.save(buf)

        buf.seek(0)
        reopened = docx.Document(buf)

        assert reopened.font_table is not None
        assert "FakeFont" in reopened.font_table
        metadata = reopened.font_table["FakeFont"]
        assert metadata.embed_regular is True

        # -- the binary round-trips through the package bytewise --
        font_table_part = reopened.font_table.part
        font_rels = [r for r in font_table_part.rels.values() if r.reltype == RT.FONT]
        assert len(font_rels) == 1
        assert font_rels[0].target_part.blob == _FAKE_FONT_BLOB


class DescribeFontPart:
    """Minimal suite for the binary `FontPart` subclass."""

    def it_round_trips_a_blob_through_load(self):
        part = FontPart.load(
            "/word/fonts/font1.fntdata",  # type: ignore[arg-type]
            CT.X_FONTDATA,
            _FAKE_FONT_BLOB,
            Package(),
        )
        assert part.blob == _FAKE_FONT_BLOB
        assert part.content_type == CT.X_FONTDATA


class DescribeFontTable_EmbedFont:
    """Coverage for `FontTable.embed_font(name, regular=..., ...)` (R5-23)."""

    def it_embeds_a_regular_variant_using_the_obfuscated_content_type(self):
        package = Package()
        ft_part = FontTablePart.default(package)

        metadata = ft_part.font_table.embed_font("Acme", regular=_FAKE_FONT_BLOB)

        assert metadata.name == "Acme"
        assert metadata.embed_regular is True
        # -- the backing FontPart uses Word's obfuscatedFont MIME --
        font_rels = [r for r in ft_part.rels.values() if r.reltype == RT.FONT]
        assert len(font_rels) == 1
        font_part = font_rels[0].target_part
        assert font_part.content_type == OBFUSCATED_FONT_CONTENT_TYPE
        # -- and stored bytes are XOR-obfuscated, not raw --
        assert font_part.blob != _FAKE_FONT_BLOB

    def it_writes_a_fontKey_GUID_on_the_embed_element(self):
        package = Package()
        ft_part = FontTablePart.default(package)

        metadata = ft_part.font_table.embed_font("Acme", regular=_FAKE_FONT_BLOB)

        embed = metadata.element.embedRegular
        assert embed is not None
        assert embed.fontKey is not None
        # -- canonical Word-style braced uppercase GUID --
        assert embed.fontKey.startswith("{") and embed.fontKey.endswith("}")

    def it_deobfuscates_embedded_regular_back_to_the_original_bytes(self):
        package = Package()
        ft_part = FontTablePart.default(package)

        metadata = ft_part.font_table.embed_font("Acme", regular=_FAKE_FONT_BLOB)

        assert metadata.embedded_regular == _FAKE_FONT_BLOB

    def it_can_embed_all_four_variants_in_a_single_call(self):
        package = Package()
        ft_part = FontTablePart.default(package)
        bold_blob = b"BOLD" + b"\x00" * 60
        italic_blob = b"ITAL" + b"\xff" * 60
        bi_blob = b"BOIT" + b"\x11" * 60

        metadata = ft_part.font_table.embed_font(
            "Acme",
            regular=_FAKE_FONT_BLOB,
            bold=bold_blob,
            italic=italic_blob,
            bold_italic=bi_blob,
        )

        assert metadata.embed_regular is True
        assert metadata.embed_bold is True
        assert metadata.embed_italic is True
        assert metadata.embed_bold_italic is True
        # -- each variant gets its own GUID (different rIds, different keys) --
        rel_ids = {
            metadata.element.embedRegular.rId,  # type: ignore[union-attr]
            metadata.element.embedBold.rId,  # type: ignore[union-attr]
            metadata.element.embedItalic.rId,  # type: ignore[union-attr]
            metadata.element.embedBoldItalic.rId,  # type: ignore[union-attr]
        }
        assert len(rel_ids) == 4
        # -- and deobfuscated bytes round-trip --
        assert metadata.embedded_regular == _FAKE_FONT_BLOB
        assert metadata.embedded_bold == bold_blob
        assert metadata.embedded_italic == italic_blob
        assert metadata.embedded_bold_italic == bi_blob

    def it_updates_an_existing_entry_in_place(self):
        package = Package()
        ft_part = FontTablePart.default(package)
        ft_part.font_table.embed_font("Acme", regular=_FAKE_FONT_BLOB)

        # -- second call with the same name should NOT create a duplicate entry --
        ft_part.font_table.embed_font("Acme", bold=b"BOLD" + b"\x00" * 60)

        assert len(ft_part.font_table) == 1
        metadata = ft_part.font_table["Acme"]
        assert metadata.embed_regular is True
        assert metadata.embed_bold is True

    def it_rejects_a_call_with_no_variants(self):
        package = Package()
        ft_part = FontTablePart.default(package)
        with pytest.raises(ValueError, match="at least one"):
            ft_part.font_table.embed_font("Acme")


class DescribeFontTable_FontsDict:
    """`FontTable.fonts` returns a snapshot `{name: FontMetadata}` mapping."""

    def it_keys_by_font_name_in_xml_order(self):
        fonts = cast(
            CT_Fonts,
            element(
                "w:fonts/("
                "w:font{w:name=Arial},"
                "w:font{w:name=Calibri}"
                ")"
            ),
        )
        ft_part = FontTablePart.default(Package())
        # -- swap in the cxml-constructed element so the proxy sees both fonts --
        table = FontTable(fonts, ft_part)

        d = table.fonts

        assert list(d) == ["Arial", "Calibri"]
        assert d["Arial"].name == "Arial"


class DescribeDocument_EmbedFont_RoundTrip:
    """Full docx save/reopen round-trip for an obfuscated-font embed."""

    def it_round_trips_the_embedded_bytes_bit_identically(self):
        document = docx.Document()
        ft = document.font_table_or_new
        ft.embed_font("Acme", regular=_FAKE_FONT_BLOB)

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = docx.Document(buf)

        ft2 = reopened.font_table
        assert ft2 is not None
        metadata = ft2["Acme"]
        # -- and the round-trip preserves the exact bytes through the XOR pair --
        assert metadata.embedded_regular == _FAKE_FONT_BLOB

    def it_preserves_the_obfuscatedFont_content_type_across_save(self):
        document = docx.Document()
        document.font_table_or_new.embed_font("Acme", regular=_FAKE_FONT_BLOB)

        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        reopened = docx.Document(buf)

        ft_part = reopened.font_table.part  # type: ignore[union-attr]
        font_rels = [r for r in ft_part.rels.values() if r.reltype == RT.FONT]
        assert len(font_rels) == 1
        assert font_rels[0].target_part.content_type == OBFUSCATED_FONT_CONTENT_TYPE
        # -- the stored bytes remain obfuscated on disk --
        stored = font_rels[0].target_part.blob
        assert stored != _FAKE_FONT_BLOB
        # -- and the fontKey round-trips on the specific Acme entry --
        acme_elm = ft_part.font_table_element.get_font_by_name("Acme")
        assert acme_elm is not None
        embed = acme_elm.embedRegular
        assert embed is not None and embed.fontKey is not None
        assert deobfuscate_font_bytes(stored, embed.fontKey) == _FAKE_FONT_BLOB


# -- sanity: element helper used above yields a CT_Fonts --
def _sanity_cxml_parses_fonts():
    assert cast(CT_Fonts, element("w:fonts")) is not None
