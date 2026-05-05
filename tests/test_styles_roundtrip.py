"""Regression tests verifying that styles.xml and numbering.xml survive a
round-trip through python-docx without semantic data loss.

Originally motivated by a scare that `Document.save()` was regenerating
`word/styles.xml` from the minimal default template and dropping every style
a Word-authored document loaded. Investigation showed the library actually
preserves the loaded XML; the apparent data loss was an artefact of a
reproducer that used ``grep -c`` (which counts matching *lines*) against the
pretty-printed original and the compact single-line XML python-docx emits.

These tests lock that behaviour in: a fixture with custom styles is loaded
and saved, and the round-tripped XML is compared with the original at both
element-count and canonical-XML levels so any real regression (element
being dropped, attributes lost, namespace munged) fails immediately.
"""

from __future__ import annotations

import zipfile

from lxml import etree

from docx import Document

_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _read_xml(zip_path: str, member: str) -> bytes:
    with zipfile.ZipFile(zip_path) as z:
        return z.read(member)


def _canonicalise(xml_bytes: bytes) -> bytes:
    """Return a whitespace-normalised C14N form for equality comparison."""
    parser = etree.XMLParser(remove_blank_text=True)
    root = etree.fromstring(xml_bytes, parser)
    return etree.tostring(root, method="c14n2")


_STYLES_SRC = "features/steps/test_files/par-known-styles.docx"
_NUMBERING_SRC = "features/steps/test_files/num-having-numbering-part.docx"


class DescribeStylesRoundTrip:
    def it_preserves_every_w_style_element_on_roundtrip(self, tmp_path):
        out = tmp_path / "rt.docx"

        Document(_STYLES_SRC).save(str(out))

        orig = etree.fromstring(_read_xml(_STYLES_SRC, "word/styles.xml"))
        rt = etree.fromstring(_read_xml(str(out), "word/styles.xml"))
        orig_ids = {s.get(f"{{{_NS['w']}}}styleId") for s in orig.findall("w:style", _NS)}
        rt_ids = {s.get(f"{{{_NS['w']}}}styleId") for s in rt.findall("w:style", _NS)}
        assert len(orig_ids) > 1, "fixture must have >1 styles for the test to be meaningful"
        assert orig_ids == rt_ids, f"styleIds dropped on roundtrip: {orig_ids - rt_ids}"

    def it_preserves_docDefaults_and_latentStyles_on_roundtrip(self, tmp_path):
        out = tmp_path / "rt.docx"

        Document(_STYLES_SRC).save(str(out))

        orig = etree.fromstring(_read_xml(_STYLES_SRC, "word/styles.xml"))
        rt = etree.fromstring(_read_xml(str(out), "word/styles.xml"))
        assert (orig.find("w:docDefaults", _NS) is not None) == (
            rt.find("w:docDefaults", _NS) is not None
        )
        orig_latent = orig.find("w:latentStyles", _NS)
        rt_latent = rt.find("w:latentStyles", _NS)
        if orig_latent is not None:
            assert rt_latent is not None
            orig_lsd = len(orig_latent.findall("w:lsdException", _NS))
            rt_lsd = len(rt_latent.findall("w:lsdException", _NS))
            assert orig_lsd == rt_lsd

    def it_preserves_styles_xml_canonical_form(self, tmp_path):
        """A semantic (whitespace-insensitive, C14N) round-trip check.

        Any element/attribute/namespace change — the kind of silent data loss a
        grep-based reproducer would miss — fails this assertion.
        """
        out = tmp_path / "rt.docx"

        Document(_STYLES_SRC).save(str(out))

        assert _canonicalise(_read_xml(_STYLES_SRC, "word/styles.xml")) == _canonicalise(
            _read_xml(str(out), "word/styles.xml")
        )


class DescribeNumberingRoundTrip:
    def it_preserves_abstractNum_and_num_elements_on_roundtrip(self, tmp_path):
        out = tmp_path / "rt.docx"

        Document(_NUMBERING_SRC).save(str(out))

        orig = etree.fromstring(_read_xml(_NUMBERING_SRC, "word/numbering.xml"))
        rt = etree.fromstring(_read_xml(str(out), "word/numbering.xml"))

        orig_ab = {
            a.get(f"{{{_NS['w']}}}abstractNumId")
            for a in orig.findall("w:abstractNum", _NS)
        }
        rt_ab = {
            a.get(f"{{{_NS['w']}}}abstractNumId")
            for a in rt.findall("w:abstractNum", _NS)
        }
        assert orig_ab, "fixture must have abstractNum entries"
        assert orig_ab == rt_ab

        orig_num = {n.get(f"{{{_NS['w']}}}numId") for n in orig.findall("w:num", _NS)}
        rt_num = {n.get(f"{{{_NS['w']}}}numId") for n in rt.findall("w:num", _NS)}
        assert orig_num == rt_num

    def it_preserves_numbering_xml_canonical_form(self, tmp_path):
        out = tmp_path / "rt.docx"

        Document(_NUMBERING_SRC).save(str(out))

        assert _canonicalise(_read_xml(_NUMBERING_SRC, "word/numbering.xml")) == _canonicalise(
            _read_xml(str(out), "word/numbering.xml")
        )


class DescribeOrphanPartPreservation:
    """W8-A: round-tripping a Word-authored file must not destroy its
    optional parts just because python-docx can't statically prove they
    are referenced.

    The 2026.05.4 "word-mimicry phase 3" release introduced drop
    heuristics that were too aggressive — they silently pruned
    ``stylesWithEffects.xml``, ``customXml/*``, ``thumbnail.jpeg``, and
    style-indirectly-referenced ``numbering.xml``. These tests lock in
    the narrower policy introduced in 2026.05.7: such parts are
    preserved verbatim when they shipped in the source package.
    """

    def it_preserves_stylesWithEffects_from_source(self, tmp_path):
        out = tmp_path / "rt.docx"

        Document(_STYLES_SRC).save(str(out))

        orig = _read_xml(_STYLES_SRC, "word/stylesWithEffects.xml")
        rt = _read_xml(str(out), "word/stylesWithEffects.xml")
        # -- just proving presence is enough; the heuristic used to
        # -- drop the part unconditionally. --
        assert orig, "fixture must ship stylesWithEffects.xml"
        assert rt, "round-trip must preserve stylesWithEffects.xml"

    def it_preserves_thumbnail_from_source(self, tmp_path):
        out = tmp_path / "rt.docx"

        Document(_STYLES_SRC).save(str(out))

        orig = _read_xml(_STYLES_SRC, "docProps/thumbnail.jpeg")
        with zipfile.ZipFile(str(out)) as z:
            names = z.namelist()
            assert "docProps/thumbnail.jpeg" in names, (
                f"thumbnail dropped on round-trip; got {names}"
            )
            rt = z.read("docProps/thumbnail.jpeg")
        assert orig == rt, "thumbnail bytes must round-trip verbatim"

    def it_preserves_numbering_for_style_indirect_references(self, tmp_path):
        """Paragraphs using a style whose definition carries <w:numPr>
        (directly or via a basedOn chain) must keep numbering.xml even
        without a direct <w:numPr> in the paragraph."""
        out = tmp_path / "rt.docx"

        Document(_NUMBERING_SRC).save(str(out))

        with zipfile.ZipFile(str(out)) as z:
            names = z.namelist()
        assert "word/numbering.xml" in names, (
            f"numbering.xml dropped despite style-indirect references; got {names}"
        )

    def it_preserves_customXml_parts_from_source(self, tmp_path):
        """customXml parts ship with Word files for purposes (Power BI,
        bibliographies, add-in backing data) that a static XPath for
        <w:dataBinding> can't detect. They must be preserved."""
        src = "tests/test_files/expanded_docx"
        # The expanded docx fixture already has customXml/item1.xml.
        # Build a zip from it and check customXml survives round-trip.
        import os
        from pathlib import Path

        pkg = tmp_path / "src.docx"
        with zipfile.ZipFile(str(pkg), "w") as z:
            for dirpath, _dirs, files in os.walk(src):
                for f in files:
                    full = os.path.join(dirpath, f)
                    arc = os.path.relpath(full, src).replace(os.sep, "/")
                    z.write(full, arc)

        out = tmp_path / "rt.docx"
        Document(str(pkg)).save(str(out))
        with zipfile.ZipFile(str(out)) as z:
            names = z.namelist()

        customxml_members = [n for n in names if n.startswith("customXml/")]
        assert customxml_members, (
            f"all customXml parts dropped on round-trip; got {names}"
        )
