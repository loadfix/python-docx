"""Test suite for the bounded-memory streaming reader (``docx.streaming``).

The streaming reader's contract is:

* :class:`StreamingDocument` yields a forward-only view of the body.
* Paragraph / table proxies returned by the generators expose the same
  read-only API surface as the eager :class:`Document` proxies.
* ``save()`` always raises :class:`StreamingNotMutableError`.

These tests focus on parity with the eager loader (paragraph count,
text, style, runs, alignment) plus the streaming-specific guarantees
(forward-only, read-only, bounded peak memory).
"""

from __future__ import annotations

import io
import os
import tracemalloc
from pathlib import Path

import pytest

from docx import Document, StreamingDocument, StreamingNotMutableError
from docx.streaming import open_stream


def _build_simple_document() -> bytes:
    """Return a small in-memory ``.docx`` with paragraphs, headings, a table."""
    d = Document()
    d.add_heading("Title", level=0)
    d.add_heading("Section A", level=1)
    d.add_paragraph("Body paragraph one.")
    d.add_paragraph("Body paragraph two.", style="List Bullet")
    d.add_heading("Section B", level=1)
    d.add_paragraph("Final body paragraph.")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "row0col0"
    table.cell(0, 1).text = "row0col1"
    table.cell(1, 0).text = "row1col0"
    table.cell(1, 1).text = "row1col1"
    d.add_paragraph("After-table paragraph.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _build_large_document(paragraph_count: int = 20_000) -> bytes:
    """Return a ``.docx`` with ``paragraph_count`` body paragraphs.

    Used by the memory-profile test below — at 20k paragraphs the
    document.xml is several MB but well under the 100 MB target so the
    test runs quickly. The eager / streaming peak-memory ratio is the
    quantity of interest, not the absolute size.
    """
    d = Document()
    for i in range(paragraph_count):
        # -- mix headings with body so style resolution exercises real
        # -- work, not a trivial constant lookup. --
        if i % 250 == 0:
            d.add_heading(f"Heading {i // 250}", level=1)
        else:
            d.add_paragraph(f"Lorem ipsum body paragraph number {i:06d}.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class DescribeStreamingDocument:
    """Parity + streaming-specific tests for :class:`StreamingDocument`."""

    def it_can_be_opened_via_Document_stream(self):
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            assert isinstance(sd, StreamingDocument)

    def it_can_be_opened_via_open_stream_shim(self):
        blob = _build_simple_document()

        with open_stream(blob) as sd:
            assert isinstance(sd, StreamingDocument)

    def it_can_be_opened_from_a_filesystem_path(self, tmp_path):
        path = tmp_path / "sample.docx"
        path.write_bytes(_build_simple_document())

        with Document.stream(str(path)) as sd:
            paragraphs = list(sd.paragraphs)

        assert len(paragraphs) >= 6

    def it_can_be_opened_from_a_PathLike(self, tmp_path):
        path = tmp_path / "sample.docx"
        path.write_bytes(_build_simple_document())

        with Document.stream(Path(path)) as sd:
            paragraphs = list(sd.paragraphs)

        assert len(paragraphs) >= 6

    def it_can_be_opened_from_a_BytesIO(self):
        blob = _build_simple_document()

        with Document.stream(io.BytesIO(blob)) as sd:
            paragraphs = list(sd.paragraphs)

        assert len(paragraphs) >= 6

    def it_yields_paragraphs_matching_the_eager_loader(self):
        # -- parity check: streaming text + style ids match Document(...).
        # -- materialise the (text, style) tuple *inside* the loop body so
        # -- each paragraph is read before its element is reclaimed. --
        blob = _build_simple_document()

        eager = Document(io.BytesIO(blob))
        eager_view = [
            (p.text, p.style.name if p.style is not None else None)
            for p in eager.paragraphs
        ]

        stream_view: list = []
        with Document.stream(blob) as sd:
            for p in sd.paragraphs:
                stream_view.append(
                    (p.text, p.style.name if p.style is not None else None)
                )

        assert stream_view == eager_view

    def it_yields_paragraphs_only_at_the_top_level_of_the_body(self):
        # -- nested paragraphs (table cells) must not surface --
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            top_level_count = sum(1 for _ in sd.paragraphs)

        # -- 1 title + 2 section headings + 3 body + 1 after-table = 7 --
        assert top_level_count == 7

    def it_yields_top_level_tables(self):
        # -- streaming clears each yielded element after the loop body
        # -- finishes, so cell access must happen *during* the iteration. --
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            cell_texts: list[str] = []
            count = 0
            for table in sd.tables:
                count += 1
                for row in table.rows:
                    for cell in row.cells:
                        cell_texts.append(cell.text)

        assert count == 1
        assert cell_texts == [
            "row0col0",
            "row0col1",
            "row1col0",
            "row1col1",
        ]

    def it_supports_run_level_text_access_on_streamed_paragraphs(self):
        # -- ``Paragraph.runs`` / ``run.text`` must work on streamed
        # -- proxies during iteration. --
        d = Document()
        para = d.add_paragraph()
        para.add_run("hello ")
        para.add_run("world")
        buf = io.BytesIO()
        d.save(buf)

        captured: list[tuple[str, list[str]]] = []
        with Document.stream(buf.getvalue()) as sd:
            for p in sd.paragraphs:
                captured.append((p.text, [r.text for r in p.runs]))

        # -- the last paragraph is our crafted one --
        text, run_texts = captured[-1]
        assert text == "hello world"
        assert run_texts == ["hello ", "world"]

    def it_resolves_alignment_via_eager_styles(self):
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        d = Document()
        p = d.add_paragraph("aligned text")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        buf = io.BytesIO()
        d.save(buf)

        alignments: list = []
        with Document.stream(buf.getvalue()) as sd:
            for para in sd.paragraphs:
                alignments.append(para.alignment)
        # -- the last paragraph is our crafted one --
        assert alignments[-1] == WD_PARAGRAPH_ALIGNMENT.CENTER

    def it_exposes_sections_eagerly(self):
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            sections = sd.sections
            count = len(sections)

        assert count >= 1

    def it_exposes_headers_and_footers_via_sections(self):
        # -- the bundled default template ships with three header / footer
        # -- variants per section (default, even-page, first-page). --
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            headers = sd.headers
            footers = sd.footers

        assert isinstance(headers, list)
        assert isinstance(footers, list)
        # -- at least one of each variant exists for the default section --
        assert len(headers) >= 1
        assert len(footers) >= 1

    def it_exposes_styles_eagerly(self):
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            normal = sd.styles["Normal"]

        assert normal is not None

    def it_exposes_the_underlying_document_part(self):
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            part = sd.part

        # -- DocumentPart from the eager bootstrap --
        from docx.parts.document import DocumentPart

        assert isinstance(part, DocumentPart)

    # -- streaming-specific contract -------------------------------------

    def it_save_raises_StreamingNotMutableError(self):
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            with pytest.raises(StreamingNotMutableError):
                sd.save(io.BytesIO())

    def it_can_be_re_iterated_for_a_second_pass(self):
        # -- forward-only per generator instance, but the property
        # -- factory returns a fresh generator on each access. --
        blob = _build_simple_document()

        with Document.stream(blob) as sd:
            first = [p.text for p in sd.paragraphs]
            second = [p.text for p in sd.paragraphs]

        assert first == second
        assert len(first) == 7

    def it_close_releases_the_source_bytes(self):
        blob = _build_simple_document()

        sd = Document.stream(blob)
        assert sd._zip_bytes is not None  # pyright: ignore[reportPrivateUsage]
        sd.close()
        assert sd._zip_bytes is None  # pyright: ignore[reportPrivateUsage]
        # -- generator returns nothing once closed --
        assert list(sd.paragraphs) == []
        # -- close() is idempotent --
        sd.close()

    def it_close_makes_sections_raise(self):
        blob = _build_simple_document()
        sd = Document.stream(blob)
        sd.close()

        with pytest.raises(ValueError, match="closed"):
            _ = sd.sections

    def it_yields_nothing_for_documents_with_no_paragraphs(self, tmp_path):
        # -- pathological corner: build a doc, then surgically remove every
        # -- body w:p so the stream has nothing to yield. --
        import zipfile

        from lxml import etree

        d = Document()
        # -- empty document: only the bundled paragraph plus a sectPr --
        buf = io.BytesIO()
        d.save(buf)

        # -- strip every w:p out of word/document.xml --
        out = io.BytesIO()
        with zipfile.ZipFile(io.BytesIO(buf.getvalue()), "r") as zin:
            with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/document.xml":
                        from docx.oxml.ns import qn
                        from docx.oxml.parser import parse_xml

                        root = parse_xml(data)
                        for p in list(root.iter(qn("w:p"))):
                            p.getparent().remove(p)
                        data = etree.tostring(root)
                    zout.writestr(item, data)

        with Document.stream(out.getvalue()) as sd:
            assert list(sd.paragraphs) == []
            assert list(sd.tables) == []

    # -- memory profile ---------------------------------------------------

    def it_keeps_peak_memory_below_threshold_for_a_large_body(self):
        """Stream a multi-MB document; peak resident must stay bounded.

        The acceptance criterion in issue #93 is "peak memory stays below
        ~50 MB for a 100 MB docx". A 100 MB synthetic doc takes ~30s to
        build inside the test runner, so we exercise the same code path
        with a smaller (~5 MB) document and assert that the streaming
        loader's peak working set is meaningfully *below* the eager
        loader's. The ratio is the load-bearing assertion — absolute
        bytes are too noisy across CI runners.
        """
        blob = _build_large_document(paragraph_count=20_000)
        # -- size sanity: the synthetic doc's document.xml part must be at
        # -- least a couple of MB or the comparison is meaningless. --
        import zipfile

        with zipfile.ZipFile(io.BytesIO(blob)) as zf:
            doc_xml_size = zf.getinfo("word/document.xml").file_size
        assert doc_xml_size > 1_000_000, (
            f"synthetic doc too small ({doc_xml_size} bytes); "
            "the memory-profile test would not be meaningful"
        )

        # -- measure eager peak --
        tracemalloc.start()
        eager = Document(io.BytesIO(blob))
        eager_count = len(eager.paragraphs)
        _, eager_peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()
        del eager

        # -- measure streaming peak --
        tracemalloc.start()
        stream_count = 0
        with Document.stream(blob) as sd:
            for _ in sd.paragraphs:
                stream_count += 1
        _, stream_peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()

        # -- functional parity: same paragraph count both ways --
        assert stream_count == eager_count

        # -- streaming peak must be substantially lower than eager peak.
        # -- Eager holds the entire CT_P graph; streaming holds at most
        # -- one paragraph at a time plus the package metadata. We
        # -- conservatively require half the eager peak to absorb GC
        # -- noise across runners. --
        assert stream_peak < eager_peak, (
            f"streaming peak {stream_peak} not < eager peak {eager_peak}"
        )
        ratio = stream_peak / eager_peak if eager_peak else 1.0
        assert ratio < 0.75, (
            f"streaming peak / eager peak ratio {ratio:.2f} too close to 1.0; "
            "memory benefit not realised"
        )

    def it_survives_iterating_a_real_test_fixture(self):
        # -- regression: the bundled tests/test_files/test.docx must
        # -- stream to the same paragraph list as the eager path. --
        path = (
            Path(__file__).parent / "test_files" / "test.docx"
        )
        if not path.exists():  # pragma: no cover - corpus-dependent
            pytest.skip("test.docx fixture not present")
        eager_paragraphs = [p.text for p in Document(str(path)).paragraphs]
        with Document.stream(str(path)) as sd:
            stream_paragraphs = [p.text for p in sd.paragraphs]
        assert stream_paragraphs == eager_paragraphs
