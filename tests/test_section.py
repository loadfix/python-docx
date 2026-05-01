# pyright: reportPrivateUsage=false

"""Unit test suite for the docx.section module."""

from __future__ import annotations

from typing import cast

import pytest

from docx import Document
from docx.enum.section import WD_HEADER_FOOTER, WD_ORIENTATION, WD_SECTION
from docx.oxml.document import CT_Document
from docx.oxml.section import CT_SectPr
from docx.parts.document import DocumentPart
from docx.parts.hdrftr import FooterPart, HeaderPart
from docx.section import Column, Section, SectionColumns, Sections, _BaseHeaderFooter, _Footer, _Header
from docx.shared import Inches, Length, RGBColor, Twips
from docx.table import Table
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element, xml
from .unitutil.file import test_file
from .unitutil.mock import (
    FixtureRequest,
    Mock,
    call,
    class_mock,
    instance_mock,
    method_mock,
    property_mock,
)


class DescribeSections:
    """Unit-test suite for `docx.section.Sections`."""

    def it_knows_how_many_sections_it_contains(self, document_part_: Mock):
        document_elm = cast(
            CT_Document, element("w:document/w:body/(w:p/w:pPr/w:sectPr, w:sectPr)")
        )
        sections = Sections(document_elm, document_part_)
        assert len(sections) == 2

    def it_can_iterate_over_its_Section_instances(
        self, Section_: Mock, section_: Mock, document_part_: Mock
    ):
        document_elm = cast(
            CT_Document, element("w:document/w:body/(w:p/w:pPr/w:sectPr, w:sectPr)")
        )
        sectPrs = document_elm.xpath("//w:sectPr")
        Section_.return_value = section_
        sections = Sections(document_elm, document_part_)

        section_lst = list(sections)

        assert Section_.call_args_list == [
            call(sectPrs[0], document_part_),
            call(sectPrs[1], document_part_),
        ]
        assert section_lst == [section_, section_]

    def it_can_access_its_Section_instances_by_index(
        self, Section_: Mock, section_: Mock, document_part_: Mock
    ):
        document_elm = cast(
            CT_Document,
            element("w:document/w:body/(w:p/w:pPr/w:sectPr,w:p/w:pPr/w:sectPr,w:sectPr)"),
        )
        sectPrs = document_elm.xpath("//w:sectPr")
        Section_.return_value = section_
        sections = Sections(document_elm, document_part_)

        section_lst = [sections[idx] for idx in range(3)]

        assert Section_.call_args_list == [
            call(sectPrs[0], document_part_),
            call(sectPrs[1], document_part_),
            call(sectPrs[2], document_part_),
        ]
        assert section_lst == [section_, section_, section_]

    def it_can_access_its_Section_instances_by_slice(
        self, Section_: Mock, section_: Mock, document_part_: Mock
    ):
        document_elm = cast(
            CT_Document,
            element("w:document/w:body/(w:p/w:pPr/w:sectPr,w:p/w:pPr/w:sectPr,w:sectPr)"),
        )
        sectPrs = document_elm.xpath("//w:sectPr")
        Section_.return_value = section_
        sections = Sections(document_elm, document_part_)

        section_lst = sections[1:9]

        assert Section_.call_args_list == [
            call(sectPrs[1], document_part_),
            call(sectPrs[2], document_part_),
        ]
        assert section_lst == [section_, section_]

    # -- fixtures---------------------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def Section_(self, request: FixtureRequest):
        return class_mock(request, "docx.section.Section")

    @pytest.fixture
    def section_(self, request: FixtureRequest):
        return instance_mock(request, Section)


class DescribeSection:
    """Unit-test suite for `docx.section.Section`."""

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr", False),
            ("w:sectPr/w:titlePg", True),
            ("w:sectPr/w:titlePg{w:val=0}", False),
            ("w:sectPr/w:titlePg{w:val=1}", True),
            ("w:sectPr/w:titlePg{w:val=true}", True),
        ],
    )
    def it_knows_when_it_displays_a_distinct_first_page_header(
        self, sectPr_cxml: str, expected_value: bool, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)

        different_first_page_header_footer = section.different_first_page_header_footer

        assert different_first_page_header_footer is expected_value

    @pytest.mark.parametrize(
        ("sectPr_cxml", "value", "expected_cxml"),
        [
            ("w:sectPr", True, "w:sectPr/w:titlePg"),
            ("w:sectPr/w:titlePg", False, "w:sectPr"),
            ("w:sectPr/w:titlePg{w:val=1}", True, "w:sectPr/w:titlePg"),
            ("w:sectPr/w:titlePg{w:val=off}", False, "w:sectPr"),
        ],
    )
    def it_can_change_whether_the_document_has_distinct_odd_and_even_headers(
        self, sectPr_cxml: str, value: bool, expected_cxml: str, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        expected_xml = xml(expected_cxml)
        section = Section(sectPr, document_part_)

        section.different_first_page_header_footer = value

        assert sectPr.xml == expected_xml

    def it_exposes_its_formatting_change_when_sectPrChange_present(
        self, document_part_: Mock
    ):
        sectPr = cast(
            CT_SectPr,
            element(
                "w:sectPr/(w:pgSz{w:w=12240,w:h=15840}"
                ",w:sectPrChange{w:id=1,w:author=Alice}/w:sectPr/w:pgSz{w:w=10000,w:h=15000})"
            ),
        )
        section = Section(sectPr, document_part_)

        fc = section.formatting_change

        assert fc is not None
        assert fc.author == "Alice"
        assert fc.old_properties is not None
        assert fc.old_properties.xpath("./w:pgSz")

    def it_returns_None_for_formatting_change_when_no_sectPrChange(
        self, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        assert section.formatting_change is None

    def it_provides_access_to_its_even_page_footer(
        self, document_part_: Mock, _Footer_: Mock, footer_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        _Footer_.return_value = footer_
        section = Section(sectPr, document_part_)

        footer = section.even_page_footer

        _Footer_.assert_called_once_with(sectPr, document_part_, WD_HEADER_FOOTER.EVEN_PAGE)
        assert footer is footer_

    def it_provides_access_to_its_even_page_header(
        self, document_part_: Mock, _Header_: Mock, header_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        _Header_.return_value = header_
        section = Section(sectPr, document_part_)

        header = section.even_page_header

        _Header_.assert_called_once_with(sectPr, document_part_, WD_HEADER_FOOTER.EVEN_PAGE)
        assert header is header_

    def it_provides_access_to_its_first_page_footer(
        self, document_part_: Mock, _Footer_: Mock, footer_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        _Footer_.return_value = footer_
        section = Section(sectPr, document_part_)

        footer = section.first_page_footer

        _Footer_.assert_called_once_with(sectPr, document_part_, WD_HEADER_FOOTER.FIRST_PAGE)
        assert footer is footer_

    def it_provides_access_to_its_first_page_header(
        self, document_part_: Mock, _Header_: Mock, header_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        _Header_.return_value = header_
        section = Section(sectPr, document_part_)

        header = section.first_page_header

        _Header_.assert_called_once_with(sectPr, document_part_, WD_HEADER_FOOTER.FIRST_PAGE)
        assert header is header_

    def it_provides_access_to_its_default_footer(
        self, document_part_: Mock, _Footer_: Mock, footer_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        _Footer_.return_value = footer_
        section = Section(sectPr, document_part_)

        footer = section.footer

        _Footer_.assert_called_once_with(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)
        assert footer is footer_

    def it_provides_access_to_its_default_header(
        self, document_part_: Mock, _Header_: Mock, header_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        _Header_.return_value = header_
        section = Section(sectPr, document_part_)

        header = section.header

        _Header_.assert_called_once_with(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)
        assert header is header_

    def it_can_iterate_its_inner_content(self):
        document = Document(test_file("sct-inner-content.docx"))

        assert len(document.sections) == 3

        inner_content = list(document.sections[0].iter_inner_content())

        assert len(inner_content) == 3
        p = inner_content[0]
        assert isinstance(p, Paragraph)
        assert p.text == "P1"
        t = inner_content[1]
        assert isinstance(t, Table)
        assert t.rows[0].cells[0].text == "T2"
        p = inner_content[2]
        assert isinstance(p, Paragraph)
        assert p.text == "P3"

        inner_content = list(document.sections[1].iter_inner_content())

        assert len(inner_content) == 3
        t = inner_content[0]
        assert isinstance(t, Table)
        assert t.rows[0].cells[0].text == "T4"
        p = inner_content[1]
        assert isinstance(p, Paragraph)
        assert p.text == "P5"
        p = inner_content[2]
        assert isinstance(p, Paragraph)
        assert p.text == "P6"

        inner_content = list(document.sections[2].iter_inner_content())

        assert len(inner_content) == 3
        p = inner_content[0]
        assert isinstance(p, Paragraph)
        assert p.text == "P7"
        p = inner_content[1]
        assert isinstance(p, Paragraph)
        assert p.text == "P8"
        p = inner_content[2]
        assert isinstance(p, Paragraph)
        assert p.text == "P9"

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr", WD_SECTION.NEW_PAGE),
            ("w:sectPr/w:type", WD_SECTION.NEW_PAGE),
            ("w:sectPr/w:type{w:val=continuous}", WD_SECTION.CONTINUOUS),
            ("w:sectPr/w:type{w:val=nextPage}", WD_SECTION.NEW_PAGE),
            ("w:sectPr/w:type{w:val=oddPage}", WD_SECTION.ODD_PAGE),
            ("w:sectPr/w:type{w:val=evenPage}", WD_SECTION.EVEN_PAGE),
            ("w:sectPr/w:type{w:val=nextColumn}", WD_SECTION.NEW_COLUMN),
        ],
    )
    def it_knows_its_start_type(
        self, sectPr_cxml: str, expected_value: WD_SECTION, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)

        start_type = section.start_type

        assert start_type is expected_value

    @pytest.mark.parametrize(
        ("sectPr_cxml", "value", "expected_cxml"),
        [
            (
                "w:sectPr/w:type{w:val=oddPage}",
                WD_SECTION.EVEN_PAGE,
                "w:sectPr/w:type{w:val=evenPage}",
            ),
            ("w:sectPr/w:type{w:val=nextPage}", None, "w:sectPr"),
            ("w:sectPr", None, "w:sectPr"),
            ("w:sectPr/w:type{w:val=continuous}", WD_SECTION.NEW_PAGE, "w:sectPr"),
            ("w:sectPr/w:type", WD_SECTION.NEW_PAGE, "w:sectPr"),
            (
                "w:sectPr/w:type",
                WD_SECTION.NEW_COLUMN,
                "w:sectPr/w:type{w:val=nextColumn}",
            ),
        ],
    )
    def it_can_change_its_start_type(
        self,
        sectPr_cxml: str,
        value: WD_SECTION | None,
        expected_cxml: str,
        document_part_: Mock,
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        expected_xml = xml(expected_cxml)
        section = Section(sectPr, document_part_)

        section.start_type = value

        assert section._sectPr.xml == expected_xml

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr/w:pgSz{w:w=1440}", Inches(1)),
            ("w:sectPr/w:pgSz", None),
            ("w:sectPr", None),
        ],
    )
    def it_knows_its_page_width(
        self, sectPr_cxml: str, expected_value: Length | None, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)

        page_width = section.page_width

        assert page_width == expected_value

    @pytest.mark.parametrize(
        ("value", "expected_cxml"),
        [
            (None, "w:sectPr/w:pgSz"),
            (Inches(4), "w:sectPr/w:pgSz{w:w=5760}"),
        ],
    )
    def it_can_change_its_page_width(
        self,
        value: Length | None,
        expected_cxml: str,
        document_part_: Mock,
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        expected_xml = xml(expected_cxml)
        section = Section(sectPr, document_part_)

        section.page_width = value

        assert section._sectPr.xml == expected_xml

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr/w:pgSz{w:h=2880}", Inches(2)),
            ("w:sectPr/w:pgSz", None),
            ("w:sectPr", None),
        ],
    )
    def it_knows_its_page_height(
        self, sectPr_cxml: str, expected_value: Length | None, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)

        page_height = section.page_height

        assert page_height == expected_value

    @pytest.mark.parametrize(
        ("value", "expected_cxml"),
        [
            (None, "w:sectPr/w:pgSz"),
            (Inches(2), "w:sectPr/w:pgSz{w:h=2880}"),
        ],
    )
    def it_can_change_its_page_height(
        self, value: Length | None, expected_cxml: str, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        expected_xml = xml(expected_cxml)
        section = Section(sectPr, document_part_)

        section.page_height = value

        assert section._sectPr.xml == expected_xml

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr/w:pgSz{w:orient=landscape}", WD_ORIENTATION.LANDSCAPE),
            ("w:sectPr/w:pgSz{w:orient=portrait}", WD_ORIENTATION.PORTRAIT),
            ("w:sectPr/w:pgSz", WD_ORIENTATION.PORTRAIT),
            ("w:sectPr", WD_ORIENTATION.PORTRAIT),
        ],
    )
    def it_knows_its_page_orientation(
        self, sectPr_cxml: str, expected_value: WD_ORIENTATION, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)

        orientation = section.orientation

        assert orientation is expected_value

    @pytest.mark.parametrize(
        ("value", "expected_cxml"),
        [
            (WD_ORIENTATION.LANDSCAPE, "w:sectPr/w:pgSz{w:orient=landscape}"),
            (WD_ORIENTATION.PORTRAIT, "w:sectPr/w:pgSz"),
            (None, "w:sectPr/w:pgSz"),
        ],
    )
    def it_can_change_its_orientation(
        self, value: WD_ORIENTATION | None, expected_cxml: str, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        expected_xml = xml(expected_cxml)
        section = Section(sectPr, document_part_)

        section.orientation = value

        assert section._sectPr.xml == expected_xml

    @pytest.mark.parametrize(
        ("sectPr_cxml", "margin_prop_name", "expected_value"),
        [
            ("w:sectPr/w:pgMar{w:left=120}", "left_margin", 76200),
            ("w:sectPr/w:pgMar{w:right=240}", "right_margin", 152400),
            ("w:sectPr/w:pgMar{w:top=-360}", "top_margin", -228600),
            ("w:sectPr/w:pgMar{w:bottom=480}", "bottom_margin", 304800),
            ("w:sectPr/w:pgMar{w:gutter=600}", "gutter", 381000),
            ("w:sectPr/w:pgMar{w:header=720}", "header_distance", 457200),
            ("w:sectPr/w:pgMar{w:footer=840}", "footer_distance", 533400),
            ("w:sectPr/w:pgMar", "left_margin", None),
            ("w:sectPr", "top_margin", None),
        ],
    )
    def it_knows_its_page_margins(
        self,
        sectPr_cxml: str,
        margin_prop_name: str,
        expected_value: int | None,
        document_part_: Mock,
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)

        value = getattr(section, margin_prop_name)

        assert value == expected_value

    @pytest.mark.parametrize(
        ("sectPr_cxml", "margin_prop_name", "value", "expected_cxml"),
        [
            ("w:sectPr", "left_margin", Inches(1), "w:sectPr/w:pgMar{w:left=1440}"),
            ("w:sectPr", "right_margin", Inches(0.5), "w:sectPr/w:pgMar{w:right=720}"),
            ("w:sectPr", "top_margin", Inches(-0.25), "w:sectPr/w:pgMar{w:top=-360}"),
            (
                "w:sectPr",
                "bottom_margin",
                Inches(0.75),
                "w:sectPr/w:pgMar{w:bottom=1080}",
            ),
            ("w:sectPr", "gutter", Inches(0.25), "w:sectPr/w:pgMar{w:gutter=360}"),
            (
                "w:sectPr",
                "header_distance",
                Inches(1.25),
                "w:sectPr/w:pgMar{w:header=1800}",
            ),
            (
                "w:sectPr",
                "footer_distance",
                Inches(1.35),
                "w:sectPr/w:pgMar{w:footer=1944}",
            ),
            ("w:sectPr", "left_margin", None, "w:sectPr/w:pgMar"),
            (
                "w:sectPr/w:pgMar{w:top=-360}",
                "top_margin",
                Inches(0.6),
                "w:sectPr/w:pgMar{w:top=864}",
            ),
        ],
    )
    def it_can_change_its_page_margins(
        self,
        sectPr_cxml: str,
        margin_prop_name: str,
        value: Length | None,
        expected_cxml: str,
        document_part_: Mock,
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        expected_xml = xml(expected_cxml)
        section = Section(sectPr, document_part_)

        setattr(section, margin_prop_name, value)

        assert section._sectPr.xml == expected_xml

    # -- watermark API ----------------------------------------------------------

    def it_can_add_a_text_watermark_via_its_section(self):
        document = Document()
        section = document.sections[0]

        watermark = section.add_text_watermark("HELLO")

        assert watermark.type == "text"
        assert watermark.text == "HELLO"

    def it_can_add_an_image_watermark_via_its_section(self):
        document = Document()
        section = document.sections[0]

        watermark = section.add_image_watermark(test_file("monty-truth.png"))

        assert watermark.type == "image"

    def it_can_read_back_a_watermark_via_its_section(self):
        document = Document()
        section = document.sections[0]
        section.add_text_watermark("DRAFT")

        assert section.watermark is not None
        assert section.watermark.text == "DRAFT"

    def it_can_remove_a_watermark_via_its_section(self):
        document = Document()
        section = document.sections[0]
        section.add_text_watermark("DRAFT")

        section.remove_watermark()

        assert section.watermark is None

    # -- fixtures-----------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def _Footer_(self, request: FixtureRequest):
        return class_mock(request, "docx.section._Footer")

    @pytest.fixture
    def footer_(self, request: FixtureRequest):
        return instance_mock(request, _Footer)

    @pytest.fixture
    def _Header_(self, request: FixtureRequest):
        return class_mock(request, "docx.section._Header")

    @pytest.fixture
    def header_(self, request: FixtureRequest):
        return instance_mock(request, _Header)


class DescribeSectionColumns:
    """Unit-test suite for `docx.section.SectionColumns`."""

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_count"),
        [
            ("w:sectPr", 1),
            ("w:sectPr/w:cols", 1),
            ("w:sectPr/w:cols{w:num=2}", 2),
            ("w:sectPr/w:cols{w:num=3}", 3),
        ],
    )
    def it_knows_its_column_count(self, sectPr_cxml: str, expected_count: int):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        columns = SectionColumns(sectPr)
        assert columns.count == expected_count

    @pytest.mark.parametrize(
        ("sectPr_cxml", "value", "expected_cxml"),
        [
            ("w:sectPr", 2, "w:sectPr/w:cols{w:num=2}"),
            ("w:sectPr/w:cols{w:num=1}", 3, "w:sectPr/w:cols{w:num=3}"),
        ],
    )
    def it_can_change_its_column_count(
        self, sectPr_cxml: str, value: int, expected_cxml: str
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        columns = SectionColumns(sectPr)
        columns.count = value
        assert sectPr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr", True),
            ("w:sectPr/w:cols", True),
            ("w:sectPr/w:cols{w:equalWidth=1}", True),
            ("w:sectPr/w:cols{w:equalWidth=0}", False),
        ],
    )
    def it_knows_whether_columns_have_equal_width(
        self, sectPr_cxml: str, expected_value: bool
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        columns = SectionColumns(sectPr)
        assert columns.equal_width is expected_value

    @pytest.mark.parametrize(
        ("sectPr_cxml", "value", "expected_cxml"),
        [
            ("w:sectPr", True, "w:sectPr/w:cols{w:equalWidth=1}"),
            ("w:sectPr/w:cols{w:equalWidth=1}", False, "w:sectPr/w:cols{w:equalWidth=0}"),
        ],
    )
    def it_can_change_equal_width(
        self, sectPr_cxml: str, value: bool, expected_cxml: str
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        columns = SectionColumns(sectPr)
        columns.equal_width = value
        assert sectPr.xml == xml(expected_cxml)

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr", None),
            ("w:sectPr/w:cols", None),
            ("w:sectPr/w:cols{w:space=720}", Twips(720)),
        ],
    )
    def it_knows_its_default_space(self, sectPr_cxml: str, expected_value: Length | None):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        columns = SectionColumns(sectPr)
        assert columns.space == expected_value

    @pytest.mark.parametrize(
        ("sectPr_cxml", "value", "expected_cxml"),
        [
            ("w:sectPr", Twips(720), "w:sectPr/w:cols{w:space=720}"),
            ("w:sectPr/w:cols{w:space=720}", None, "w:sectPr/w:cols"),
        ],
    )
    def it_can_change_its_default_space(
        self, sectPr_cxml: str, value: Length | None, expected_cxml: str
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        columns = SectionColumns(sectPr)
        columns.space = value
        assert sectPr.xml == xml(expected_cxml)

    def it_provides_access_to_individual_columns(self):
        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:cols/(w:col{w:w=4320,w:space=720},w:col{w:w=4320})"),
        )
        columns = SectionColumns(sectPr)
        assert len(columns) == 2
        assert columns[0].width == Twips(4320)
        assert columns[0].space == Twips(720)
        assert columns[1].width == Twips(4320)
        assert columns[1].space is None

    def it_can_iterate_over_individual_columns(self):
        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:cols/(w:col{w:w=4320,w:space=720},w:col{w:w=2880})"),
        )
        columns = SectionColumns(sectPr)
        col_list = list(columns)
        assert len(col_list) == 2
        assert col_list[0].width == Twips(4320)
        assert col_list[1].width == Twips(2880)

    def it_returns_zero_length_when_no_cols_element(self):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        columns = SectionColumns(sectPr)
        assert len(columns) == 0
        assert list(columns) == []


class DescribeColumn:
    """Unit-test suite for `docx.section.Column`."""

    @pytest.mark.parametrize(
        ("col_cxml", "expected_width", "expected_space"),
        [
            ("w:col", None, None),
            ("w:col{w:w=4320,w:space=720}", Twips(4320), Twips(720)),
        ],
    )
    def it_knows_its_width_and_space(
        self, col_cxml: str, expected_width: Length | None, expected_space: Length | None
    ):
        from docx.oxml.section import CT_Col

        col_elm = cast(CT_Col, element(col_cxml))
        col = Column(col_elm)
        assert col.width == expected_width
        assert col.space == expected_space

    def it_can_change_its_width(self):
        from docx.oxml.section import CT_Col

        col_elm = cast(CT_Col, element("w:col"))
        col = Column(col_elm)
        col.width = Twips(4320)
        assert col_elm.xml == xml("w:col{w:w=4320}")

    def it_can_change_its_space(self):
        from docx.oxml.section import CT_Col

        col_elm = cast(CT_Col, element("w:col"))
        col = Column(col_elm)
        col.space = Twips(720)
        assert col_elm.xml == xml("w:col{w:space=720}")


class DescribeSection_columns:
    """Unit-test suite for `docx.section.Section.columns`."""

    def it_provides_access_to_section_columns(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:cols{w:num=2,w:space=720}"))
        section = Section(sectPr, document_part_)
        columns = section.columns
        assert isinstance(columns, SectionColumns)
        assert columns.count == 2
        assert columns.space == Twips(720)

    # -- fixtures-----------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeSection_footnote_and_endnote_properties:
    """Unit-test suite for `Section.footnote_properties` / `endnote_properties`."""

    def it_returns_None_when_no_footnotePr_present(self, document_part_: Mock):
        from docx.footnotes import FootnoteProperties

        assert FootnoteProperties is not None  # used for type clarity
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        assert section.footnote_properties is None

    def it_returns_a_FootnoteProperties_when_footnotePr_present(self, document_part_: Mock):
        from docx.enum.text import WD_FOOTNOTE_POSITION
        from docx.footnotes import FootnoteProperties

        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:footnotePr/w:pos{w:val=beneathText}"),
        )
        section = Section(sectPr, document_part_)
        props = section.footnote_properties
        assert isinstance(props, FootnoteProperties)
        assert props.position == WD_FOOTNOTE_POSITION.BENEATH_TEXT

    def it_can_add_footnote_properties(self, document_part_: Mock):
        from docx.footnotes import FootnoteProperties

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)

        props = section.add_footnote_properties()

        assert isinstance(props, FootnoteProperties)
        assert sectPr.xml == xml("w:sectPr/w:footnotePr")

    def it_places_footnotePr_before_type_when_added(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:type{w:val=continuous}"))
        section = Section(sectPr, document_part_)

        section.add_footnote_properties()

        assert sectPr.xml == xml(
            "w:sectPr/(w:footnotePr,w:type{w:val=continuous})"
        )

    def it_can_remove_footnote_properties(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:footnotePr"))
        section = Section(sectPr, document_part_)
        section.remove_footnote_properties()
        assert sectPr.xml == xml("w:sectPr")

    def it_returns_None_when_no_endnotePr_present(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        assert section.endnote_properties is None

    def it_returns_an_EndnoteProperties_when_endnotePr_present(self, document_part_: Mock):
        from docx.endnotes import EndnoteProperties
        from docx.enum.text import WD_ENDNOTE_POSITION

        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:endnotePr/w:pos{w:val=sectEnd}"),
        )
        section = Section(sectPr, document_part_)
        props = section.endnote_properties
        assert isinstance(props, EndnoteProperties)
        assert props.position == WD_ENDNOTE_POSITION.END_OF_SECTION

    def it_can_add_endnote_properties(self, document_part_: Mock):
        from docx.endnotes import EndnoteProperties

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)

        props = section.add_endnote_properties()

        assert isinstance(props, EndnoteProperties)
        assert sectPr.xml == xml("w:sectPr/w:endnotePr")

    def it_can_remove_endnote_properties(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:endnotePr"))
        section = Section(sectPr, document_part_)
        section.remove_endnote_properties()
        assert sectPr.xml == xml("w:sectPr")

    # -- fixtures-----------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class Describe_BaseHeaderFooter:
    """Unit-test suite for `docx.section._BaseHeaderFooter`."""

    @pytest.mark.parametrize(("has_definition", "expected_value"), [(False, True), (True, False)])
    def it_knows_when_its_linked_to_the_previous_header_or_footer(
        self,
        has_definition: bool,
        expected_value: bool,
        header: _BaseHeaderFooter,
        _has_definition_prop_: Mock,
    ):
        _has_definition_prop_.return_value = has_definition
        assert header.is_linked_to_previous is expected_value

    @pytest.mark.parametrize(
        ("has_definition", "value", "drop_calls", "add_calls"),
        [
            (False, True, 0, 0),
            (True, False, 0, 0),
            (True, True, 1, 0),
            (False, False, 0, 1),
        ],
    )
    def it_can_change_whether_it_is_linked_to_previous_header_or_footer(
        self,
        has_definition: bool,
        value: bool,
        drop_calls: int,
        add_calls: int,
        header: _BaseHeaderFooter,
        _has_definition_prop_: Mock,
        _drop_definition_: Mock,
        _add_definition_: Mock,
    ):
        _has_definition_prop_.return_value = has_definition

        header.is_linked_to_previous = value

        assert _drop_definition_.call_args_list == [call(header)] * drop_calls
        assert _add_definition_.call_args_list == [call(header)] * add_calls

    def it_provides_access_to_the_header_or_footer_part_for_BlockItemContainer(
        self, header: _BaseHeaderFooter, _get_or_add_definition_: Mock, header_part_: Mock
    ):
        # ---this override fulfills part of the BlockItemContainer subclass interface---
        _get_or_add_definition_.return_value = header_part_

        header_part = header.part

        _get_or_add_definition_.assert_called_once_with(header)
        assert header_part is header_part_

    def it_provides_access_to_the_hdr_or_ftr_element_to_help(
        self, header: _BaseHeaderFooter, _get_or_add_definition_: Mock, header_part_: Mock
    ):
        hdr = element("w:hdr")
        _get_or_add_definition_.return_value = header_part_
        header_part_.element = hdr

        hdr_elm = header._element

        _get_or_add_definition_.assert_called_once_with(header)
        assert hdr_elm is hdr

    def it_gets_the_definition_when_it_has_one(
        self,
        header: _BaseHeaderFooter,
        _has_definition_prop_: Mock,
        _definition_prop_: Mock,
        header_part_: Mock,
    ):
        _has_definition_prop_.return_value = True
        _definition_prop_.return_value = header_part_

        header_part = header._get_or_add_definition()

        assert header_part is header_part_

    def but_it_gets_the_prior_definition_when_it_is_linked(
        self,
        header: _BaseHeaderFooter,
        _has_definition_prop_: Mock,
        _prior_headerfooter_prop_: Mock,
        prior_headerfooter_: Mock,
        header_part_: Mock,
    ):
        _has_definition_prop_.return_value = False
        _prior_headerfooter_prop_.return_value = prior_headerfooter_
        prior_headerfooter_._get_or_add_definition.return_value = header_part_

        header_part = header._get_or_add_definition()

        prior_headerfooter_._get_or_add_definition.assert_called_once_with()
        assert header_part is header_part_

    def and_it_adds_a_definition_when_it_is_linked_and_the_first_section(
        self,
        header: _BaseHeaderFooter,
        _has_definition_prop_: Mock,
        _prior_headerfooter_prop_: Mock,
        _add_definition_: Mock,
        header_part_: Mock,
    ):
        _has_definition_prop_.return_value = False
        _prior_headerfooter_prop_.return_value = None
        _add_definition_.return_value = header_part_

        header_part = header._get_or_add_definition()

        _add_definition_.assert_called_once_with(header)
        assert header_part is header_part_

    # -- fixture -----------------------------------------------------

    @pytest.fixture
    def _add_definition_(self, request: FixtureRequest):
        return method_mock(request, _BaseHeaderFooter, "_add_definition")

    @pytest.fixture
    def _definition_prop_(self, request: FixtureRequest):
        return property_mock(request, _BaseHeaderFooter, "_definition")

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def _drop_definition_(self, request: FixtureRequest):
        return method_mock(request, _BaseHeaderFooter, "_drop_definition")

    @pytest.fixture
    def _get_or_add_definition_(self, request: FixtureRequest):
        return method_mock(request, _BaseHeaderFooter, "_get_or_add_definition")

    @pytest.fixture
    def _has_definition_prop_(self, request: FixtureRequest):
        return property_mock(request, _BaseHeaderFooter, "_has_definition")

    @pytest.fixture
    def header(self, document_part_: Mock) -> _BaseHeaderFooter:
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        return _BaseHeaderFooter(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)

    @pytest.fixture
    def header_part_(self, request: FixtureRequest):
        return instance_mock(request, HeaderPart)

    @pytest.fixture
    def prior_headerfooter_(self, request: FixtureRequest):
        return instance_mock(request, _BaseHeaderFooter)

    @pytest.fixture
    def _prior_headerfooter_prop_(self, request: FixtureRequest):
        return property_mock(request, _BaseHeaderFooter, "_prior_headerfooter")


class Describe_Footer:
    """Unit-test suite for `docx.section._Footer`."""

    def it_can_add_a_footer_part_to_help(self, document_part_: Mock, footer_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr{r:a=b}"))
        document_part_.add_footer_part.return_value = footer_part_, "rId3"
        footer = _Footer(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)

        footer_part = footer._add_definition()

        document_part_.add_footer_part.assert_called_once_with()
        assert sectPr.xml == xml("w:sectPr{r:a=b}/w:footerReference{w:type=default,r:id=rId3}")
        assert footer_part is footer_part_

    def it_provides_access_to_its_footer_part_to_help(
        self, document_part_: Mock, footer_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:footerReference{w:type=even,r:id=rId3}"))
        document_part_.footer_part.return_value = footer_part_
        footer = _Footer(sectPr, document_part_, WD_HEADER_FOOTER.EVEN_PAGE)

        footer_part = footer._definition

        document_part_.footer_part.assert_called_once_with("rId3")
        assert footer_part is footer_part_

    def it_can_drop_the_related_footer_part_to_help(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr, element("w:sectPr{r:a=b}/w:footerReference{w:type=first,r:id=rId42}")
        )
        footer = _Footer(sectPr, document_part_, WD_HEADER_FOOTER.FIRST_PAGE)

        footer._drop_definition()

        assert sectPr.xml == xml("w:sectPr{r:a=b}")
        document_part_.drop_rel.assert_called_once_with("rId42")

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [("w:sectPr", False), ("w:sectPr/w:footerReference{w:type=default}", True)],
    )
    def it_knows_when_it_has_a_definition_to_help(
        self, sectPr_cxml: str, expected_value: bool, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        footer = _Footer(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)

        has_definition = footer._has_definition

        assert has_definition is expected_value

    def it_provides_access_to_the_prior_Footer_to_help(
        self, request: FixtureRequest, document_part_: Mock, footer_: Mock
    ):
        doc_elm = element("w:document/(w:sectPr,w:sectPr)")
        prior_sectPr, sectPr = cast(CT_SectPr, doc_elm[0]), cast(CT_SectPr, doc_elm[1])
        footer = _Footer(sectPr, document_part_, WD_HEADER_FOOTER.EVEN_PAGE)
        # ---mock must occur after construction of "real" footer---
        _Footer_ = class_mock(request, "docx.section._Footer", return_value=footer_)

        prior_footer = footer._prior_headerfooter

        _Footer_.assert_called_once_with(prior_sectPr, document_part_, WD_HEADER_FOOTER.EVEN_PAGE)
        assert prior_footer is footer_

    def but_it_returns_None_when_its_the_first_footer(self, document_part_: Mock):
        doc_elm = cast(CT_Document, element("w:document/w:sectPr"))
        sectPr = cast(CT_SectPr, doc_elm[0])
        footer = _Footer(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)

        prior_footer = footer._prior_headerfooter

        assert prior_footer is None

    # -- fixtures---------------------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def footer_(self, request: FixtureRequest):
        return instance_mock(request, _Footer)

    @pytest.fixture
    def footer_part_(self, request: FixtureRequest):
        return instance_mock(request, FooterPart)


class Describe_Header:
    """Unit-test suite for `docx.section._Header`."""

    def it_can_add_a_header_part_to_help(self, document_part_: Mock, header_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr{r:a=b}"))
        document_part_.add_header_part.return_value = header_part_, "rId3"
        header = _Header(sectPr, document_part_, WD_HEADER_FOOTER.FIRST_PAGE)

        header_part = header._add_definition()

        document_part_.add_header_part.assert_called_once_with()
        assert sectPr.xml == xml("w:sectPr{r:a=b}/w:headerReference{w:type=first,r:id=rId3}")
        assert header_part is header_part_

    def it_provides_access_to_its_header_part_to_help(
        self, document_part_: Mock, header_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:headerReference{w:type=default,r:id=rId8}"))
        document_part_.header_part.return_value = header_part_
        header = _Header(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)

        header_part = header._definition

        document_part_.header_part.assert_called_once_with("rId8")
        assert header_part is header_part_

    def it_can_drop_the_related_header_part_to_help(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr, element("w:sectPr{r:a=b}/w:headerReference{w:type=even,r:id=rId42}")
        )
        header = _Header(sectPr, document_part_, WD_HEADER_FOOTER.EVEN_PAGE)

        header._drop_definition()

        assert sectPr.xml == xml("w:sectPr{r:a=b}")
        document_part_.drop_header_part.assert_called_once_with("rId42")

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [("w:sectPr", False), ("w:sectPr/w:headerReference{w:type=first}", True)],
    )
    def it_knows_when_it_has_a_header_part_to_help(
        self, sectPr_cxml: str, expected_value: bool, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        header = _Header(sectPr, document_part_, WD_HEADER_FOOTER.FIRST_PAGE)

        has_definition = header._has_definition

        assert has_definition is expected_value

    def it_provides_access_to_the_prior_Header_to_help(
        self, request: FixtureRequest, document_part_: Mock, header_: Mock
    ):
        doc_elm = element("w:document/(w:sectPr,w:sectPr)")
        prior_sectPr, sectPr = cast(CT_SectPr, doc_elm[0]), cast(CT_SectPr, doc_elm[1])
        header = _Header(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)
        # ---mock must occur after construction of "real" header---
        _Header_ = class_mock(request, "docx.section._Header", return_value=header_)

        prior_header = header._prior_headerfooter

        _Header_.assert_called_once_with(prior_sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)
        assert prior_header is header_

    def but_it_returns_None_when_its_the_first_header(self, document_part_: Mock):
        doc_elm = element("w:document/w:sectPr")
        sectPr = cast(CT_SectPr, doc_elm[0])
        header = _Header(sectPr, document_part_, WD_HEADER_FOOTER.PRIMARY)

        prior_header = header._prior_headerfooter

        assert prior_header is None

    # -- fixtures---------------------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)

    @pytest.fixture
    def header_(self, request: FixtureRequest):
        return instance_mock(request, _Header)

    @pytest.fixture
    def header_part_(self, request: FixtureRequest):
        return instance_mock(request, HeaderPart)


class DescribePageBorder:
    """Unit-test suite for `docx.section.PageBorder`."""

    @pytest.mark.parametrize(
        ("sectPr_cxml", "side", "expected"),
        [
            ("w:sectPr", "top", None),
            ("w:sectPr/w:pgBorders", "top", None),
            ("w:sectPr/w:pgBorders/w:top{w:val=single}", "top", "single"),
            ("w:sectPr/w:pgBorders/w:bottom{w:val=dashed}", "bottom", "dashed"),
            ("w:sectPr/w:pgBorders/w:left{w:val=double}", "left", "double"),
            ("w:sectPr/w:pgBorders/w:right{w:val=dotted}", "right", "dotted"),
        ],
    )
    def it_knows_its_style(
        self, sectPr_cxml: str, side: str, expected: str | None, document_part_: Mock
    ):
        from docx.enum.text import WD_BORDER_STYLE
        from docx.section import PageBorder

        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        border = PageBorder(sectPr, side)
        style = border.style
        if expected is None:
            assert style is None
        else:
            assert style == WD_BORDER_STYLE.from_xml(expected)

    def it_can_set_its_style_creating_pgBorders_and_edge(self, document_part_: Mock):
        from docx.enum.text import WD_BORDER_STYLE
        from docx.section import PageBorder

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        border = PageBorder(sectPr, "top")
        border.style = WD_BORDER_STYLE.SINGLE
        assert sectPr.xml == xml("w:sectPr/w:pgBorders/w:top{w:val=single}")

    def it_knows_its_width(self, document_part_: Mock):
        from docx.section import PageBorder
        from docx.shared import Pt

        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:pgBorders/w:top{w:val=single,w:sz=24}")
        )
        border = PageBorder(sectPr, "top")
        # -- 24 eighth-points => 3 points --
        assert border.width == Pt(3)

    def it_can_set_its_width(self, document_part_: Mock):
        from docx.section import PageBorder
        from docx.shared import Pt

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        border = PageBorder(sectPr, "top")
        border.width = Pt(3)
        assert sectPr.xml == xml("w:sectPr/w:pgBorders/w:top{w:sz=24}")

    def it_knows_its_color(self, document_part_: Mock):
        from docx.section import PageBorder

        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:pgBorders/w:top{w:color=FF0000}")
        )
        border = PageBorder(sectPr, "top")
        assert border.color == RGBColor(0xFF, 0x00, 0x00)

    def it_returns_None_color_for_auto(self, document_part_: Mock):
        from docx.section import PageBorder

        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:pgBorders/w:top{w:color=auto}")
        )
        border = PageBorder(sectPr, "top")
        assert border.color is None

    def it_can_set_its_color(self, document_part_: Mock):
        from docx.section import PageBorder

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        border = PageBorder(sectPr, "left")
        border.color = RGBColor(0x12, 0x34, 0x56)
        assert sectPr.xml == xml("w:sectPr/w:pgBorders/w:left{w:color=123456}")

    def it_knows_its_space(self, document_part_: Mock):
        from docx.section import PageBorder
        from docx.shared import Pt

        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:pgBorders/w:top{w:space=24}")
        )
        border = PageBorder(sectPr, "top")
        assert border.space == Pt(24)

    def it_can_set_its_space(self, document_part_: Mock):
        from docx.section import PageBorder
        from docx.shared import Pt

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        border = PageBorder(sectPr, "bottom")
        border.space = Pt(24)
        assert sectPr.xml == xml("w:sectPr/w:pgBorders/w:bottom{w:space=24}")

    def it_can_clear_its_style_to_None(self, document_part_: Mock):
        from docx.section import PageBorder

        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:pgBorders/w:top{w:val=single,w:sz=24}"),
        )
        border = PageBorder(sectPr, "top")
        border.style = None
        assert sectPr.xml == xml("w:sectPr/w:pgBorders/w:top{w:sz=24}")

    # -- fixtures ---------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribePageBorders:
    """Unit-test suite for `docx.section.PageBorders`."""

    def it_provides_access_to_each_edge(self, document_part_: Mock):
        from docx.section import PageBorder, PageBorders

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        borders = PageBorders(sectPr)
        assert isinstance(borders.top, PageBorder)
        assert isinstance(borders.bottom, PageBorder)
        assert isinstance(borders.left, PageBorder)
        assert isinstance(borders.right, PageBorder)

    def it_returns_None_attributes_when_pgBorders_missing(self, document_part_: Mock):
        from docx.section import PageBorders

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        borders = PageBorders(sectPr)
        assert borders.display is None
        assert borders.offset_from is None
        assert borders.top.style is None
        assert borders.top.width is None
        assert borders.top.color is None
        assert borders.top.space is None

    @pytest.mark.parametrize(
        ("xml_val", "enum_member"),
        [
            ("allPages", "ALL_PAGES"),
            ("firstPage", "FIRST_PAGE"),
            ("notFirstPage", "NOT_FIRST_PAGE"),
        ],
    )
    def it_knows_its_display(self, xml_val: str, enum_member: str, document_part_: Mock):
        from docx.enum.section import WD_BORDER_DISPLAY
        from docx.section import PageBorders

        sectPr = cast(
            CT_SectPr, element(f"w:sectPr/w:pgBorders{{w:display={xml_val}}}")
        )
        borders = PageBorders(sectPr)
        assert borders.display is getattr(WD_BORDER_DISPLAY, enum_member)

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("ALL_PAGES", "allPages"),
            ("FIRST_PAGE", "firstPage"),
            ("NOT_FIRST_PAGE", "notFirstPage"),
        ],
    )
    def it_can_set_its_display_round_trip(
        self, enum_member: str, xml_val: str, document_part_: Mock
    ):
        from docx.enum.section import WD_BORDER_DISPLAY
        from docx.section import PageBorders

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        borders = PageBorders(sectPr)
        borders.display = getattr(WD_BORDER_DISPLAY, enum_member)
        assert sectPr.xml == xml(f"w:sectPr/w:pgBorders{{w:display={xml_val}}}")

    @pytest.mark.parametrize(
        ("xml_val", "enum_member"),
        [
            ("text", "TEXT"),
            ("page", "PAGE"),
        ],
    )
    def it_knows_its_offset_from(
        self, xml_val: str, enum_member: str, document_part_: Mock
    ):
        from docx.enum.section import WD_BORDER_OFFSET_FROM
        from docx.section import PageBorders

        sectPr = cast(
            CT_SectPr, element(f"w:sectPr/w:pgBorders{{w:offsetFrom={xml_val}}}")
        )
        borders = PageBorders(sectPr)
        assert borders.offset_from is getattr(WD_BORDER_OFFSET_FROM, enum_member)

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("TEXT", "text"),
            ("PAGE", "page"),
        ],
    )
    def it_can_set_its_offset_from_round_trip(
        self, enum_member: str, xml_val: str, document_part_: Mock
    ):
        from docx.enum.section import WD_BORDER_OFFSET_FROM
        from docx.section import PageBorders

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        borders = PageBorders(sectPr)
        borders.offset_from = getattr(WD_BORDER_OFFSET_FROM, enum_member)
        assert sectPr.xml == xml(f"w:sectPr/w:pgBorders{{w:offsetFrom={xml_val}}}")

    def it_can_clear_its_display(self, document_part_: Mock):
        from docx.section import PageBorders

        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:pgBorders{w:display=allPages}")
        )
        borders = PageBorders(sectPr)
        borders.display = None
        assert sectPr.xml == xml("w:sectPr/w:pgBorders")

    # -- fixtures ---------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeSection_page_borders:
    """Unit-test suite for `docx.section.Section.page_borders` and related API."""

    def it_provides_access_to_page_borders(self, document_part_: Mock):
        from docx.section import PageBorders

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        borders = section.page_borders
        assert isinstance(borders, PageBorders)

    def it_returns_None_style_when_no_pgBorders(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        borders = section.page_borders
        assert borders.top.style is None
        assert borders.bottom.style is None
        assert borders.left.style is None
        assert borders.right.style is None
        assert borders.display is None
        assert borders.offset_from is None

    def it_can_set_a_single_edge_via_set_page_border(self, document_part_: Mock):
        from docx.enum.text import WD_BORDER_STYLE
        from docx.shared import Pt

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.set_page_border(
            "top",
            style=WD_BORDER_STYLE.SINGLE,
            width=Pt(3),
            color=RGBColor(0xFF, 0x00, 0x00),
            space=Pt(24),
        )
        assert sectPr.xml == xml(
            "w:sectPr/w:pgBorders/w:top{w:val=single,w:sz=24,"
            "w:space=24,w:color=FF0000}"
        )

    def it_raises_ValueError_for_invalid_side_in_set_page_border(
        self, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        with pytest.raises(ValueError):
            section.set_page_border("middle")

    def it_can_remove_page_borders(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr,
            element(
                "w:sectPr/w:pgBorders/(w:top{w:val=single},w:bottom{w:val=dashed})"
            ),
        )
        section = Section(sectPr, document_part_)
        section.remove_page_borders()
        assert sectPr.xml == xml("w:sectPr")

    def it_does_nothing_on_remove_when_no_pgBorders(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        # -- should not raise --
        section.remove_page_borders()
        assert sectPr.xml == xml("w:sectPr")

    # -- fixtures ---------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeLineNumbering:
    """Unit-test suite for `docx.section.LineNumbering`."""

    def it_knows_its_count_by(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(CT_LineNumber, element("w:lnNumType{w:countBy=5}"))
        numbering = LineNumbering(lnNumType)
        assert numbering.count_by == 5

    def it_knows_its_start(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(CT_LineNumber, element("w:lnNumType{w:start=3}"))
        numbering = LineNumbering(lnNumType)
        assert numbering.start == 3

    def it_knows_its_distance(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering
        from docx.shared import Twips

        lnNumType = cast(CT_LineNumber, element("w:lnNumType{w:distance=720}"))
        numbering = LineNumbering(lnNumType)
        assert numbering.distance == Twips(720)

    def it_returns_None_attributes_when_unset(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(CT_LineNumber, element("w:lnNumType"))
        numbering = LineNumbering(lnNumType)
        assert numbering.count_by is None
        assert numbering.start is None
        assert numbering.distance is None
        assert numbering.restart is None

    def it_can_set_its_count_by(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(CT_LineNumber, element("w:lnNumType"))
        numbering = LineNumbering(lnNumType)
        numbering.count_by = 2
        assert lnNumType.xml == xml("w:lnNumType{w:countBy=2}")

    def it_can_set_its_start(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(CT_LineNumber, element("w:lnNumType"))
        numbering = LineNumbering(lnNumType)
        numbering.start = 10
        assert lnNumType.xml == xml("w:lnNumType{w:start=10}")

    def it_can_set_its_distance(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering
        from docx.shared import Twips

        lnNumType = cast(CT_LineNumber, element("w:lnNumType"))
        numbering = LineNumbering(lnNumType)
        numbering.distance = Twips(720)
        assert lnNumType.xml == xml("w:lnNumType{w:distance=720}")

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("CONTINUOUS", "continuous"),
            ("NEW_SECTION", "newSection"),
            ("NEW_PAGE", "newPage"),
        ],
    )
    def it_knows_its_restart(self, enum_member: str, xml_val: str):
        from docx.enum.section import WD_LINE_NUMBERING_RESTART
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(
            CT_LineNumber, element(f"w:lnNumType{{w:restart={xml_val}}}")
        )
        numbering = LineNumbering(lnNumType)
        assert numbering.restart is getattr(WD_LINE_NUMBERING_RESTART, enum_member)

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("CONTINUOUS", "continuous"),
            ("NEW_SECTION", "newSection"),
            ("NEW_PAGE", "newPage"),
        ],
    )
    def it_can_set_its_restart_round_trip(self, enum_member: str, xml_val: str):
        from docx.enum.section import WD_LINE_NUMBERING_RESTART
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(CT_LineNumber, element("w:lnNumType"))
        numbering = LineNumbering(lnNumType)
        numbering.restart = getattr(WD_LINE_NUMBERING_RESTART, enum_member)
        assert lnNumType.xml == xml(f"w:lnNumType{{w:restart={xml_val}}}")

    def it_can_clear_its_count_by(self):
        from docx.oxml.section import CT_LineNumber
        from docx.section import LineNumbering

        lnNumType = cast(CT_LineNumber, element("w:lnNumType{w:countBy=2}"))
        numbering = LineNumbering(lnNumType)
        numbering.count_by = None
        assert lnNumType.xml == xml("w:lnNumType")


class DescribeSection_line_numbering:
    """Unit-test suite for `docx.section.Section.line_numbering` and related API."""

    def it_returns_None_when_no_lnNumType(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        assert section.line_numbering is None

    def it_provides_access_to_LineNumbering_when_present(self, document_part_: Mock):
        from docx.section import LineNumbering

        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:lnNumType{w:countBy=1,w:start=1}")
        )
        section = Section(sectPr, document_part_)
        numbering = section.line_numbering
        assert isinstance(numbering, LineNumbering)
        assert numbering.count_by == 1
        assert numbering.start == 1

    def it_can_set_line_numbering_creating_lnNumType(self, document_part_: Mock):
        from docx.enum.section import WD_LINE_NUMBERING_RESTART
        from docx.shared import Twips

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.set_line_numbering(
            count_by=1,
            start=1,
            distance=Twips(360),
            restart=WD_LINE_NUMBERING_RESTART.NEW_PAGE,
        )
        assert sectPr.xml == xml(
            "w:sectPr/w:lnNumType{w:countBy=1,w:start=1,"
            "w:distance=360,w:restart=newPage}"
        )

    def it_returns_a_LineNumbering_from_set_line_numbering(self, document_part_: Mock):
        from docx.section import LineNumbering

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        numbering = section.set_line_numbering(count_by=1)
        assert isinstance(numbering, LineNumbering)
        assert numbering.count_by == 1

    def it_leaves_unchanged_attributes_alone_in_set_line_numbering(
        self, document_part_: Mock
    ):
        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:lnNumType{w:countBy=1,w:start=5}"),
        )
        section = Section(sectPr, document_part_)
        section.set_line_numbering(count_by=2)
        assert sectPr.xml == xml("w:sectPr/w:lnNumType{w:countBy=2,w:start=5}")

    def it_can_remove_line_numbering(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:lnNumType{w:countBy=1}"),
        )
        section = Section(sectPr, document_part_)
        section.remove_line_numbering()
        assert sectPr.xml == xml("w:sectPr")

    def it_does_nothing_on_remove_when_no_lnNumType(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.remove_line_numbering()
        assert sectPr.xml == xml("w:sectPr")

    # -- fixtures ---------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeSection_paper_source:
    """Unit-test suite for `Section.first_page_paper_source` / `other_pages_paper_source`."""

    def it_returns_None_for_both_when_no_paperSrc(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        assert section.first_page_paper_source is None
        assert section.other_pages_paper_source is None

    def it_reads_first_when_only_first_set(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:paperSrc{w:first=3}"))
        section = Section(sectPr, document_part_)
        assert section.first_page_paper_source == 3
        assert section.other_pages_paper_source is None

    def it_reads_other_when_only_other_set(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr/w:paperSrc{w:other=4}"))
        section = Section(sectPr, document_part_)
        assert section.first_page_paper_source is None
        assert section.other_pages_paper_source == 4

    def it_reads_both_when_both_set(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:paperSrc{w:first=1,w:other=2}")
        )
        section = Section(sectPr, document_part_)
        assert section.first_page_paper_source == 1
        assert section.other_pages_paper_source == 2

    def it_can_round_trip_first_page_paper_source(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.first_page_paper_source = 7
        assert sectPr.xml == xml("w:sectPr/w:paperSrc{w:first=7}")
        assert section.first_page_paper_source == 7

    def it_can_round_trip_other_pages_paper_source(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.other_pages_paper_source = 9
        assert sectPr.xml == xml("w:sectPr/w:paperSrc{w:other=9}")
        assert section.other_pages_paper_source == 9

    def it_keeps_both_when_set_sequentially(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.first_page_paper_source = 1
        section.other_pages_paper_source = 2
        assert sectPr.xml == xml("w:sectPr/w:paperSrc{w:first=1,w:other=2}")
        assert section.first_page_paper_source == 1
        assert section.other_pages_paper_source == 2

    def it_removes_paperSrc_when_both_set_to_None(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:paperSrc{w:first=1,w:other=2}")
        )
        section = Section(sectPr, document_part_)
        section.first_page_paper_source = None
        section.other_pages_paper_source = None
        assert sectPr.xml == xml("w:sectPr")

    def it_keeps_paperSrc_when_only_one_set_to_None(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:paperSrc{w:first=1,w:other=2}")
        )
        section = Section(sectPr, document_part_)
        section.first_page_paper_source = None
        assert sectPr.xml == xml("w:sectPr/w:paperSrc{w:other=2}")
        assert section.first_page_paper_source is None
        assert section.other_pages_paper_source == 2

    def it_is_noop_to_set_None_when_no_paperSrc(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.first_page_paper_source = None
        section.other_pages_paper_source = None
        assert sectPr.xml == xml("w:sectPr")

    def it_inserts_paperSrc_in_correct_position(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr/(w:pgSz,w:pgMar,w:cols)"))
        section = Section(sectPr, document_part_)
        section.first_page_paper_source = 1
        assert sectPr.xml == xml(
            "w:sectPr/(w:pgSz,w:pgMar,w:paperSrc{w:first=1},w:cols)"
        )

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeDocumentGrid:
    """Unit-test suite for `docx.section.DocumentGrid`."""

    def it_knows_its_type(self):
        from docx.enum.section import WD_DOC_GRID_TYPE
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid{w:type=lines}"))
        grid = DocumentGrid(docGrid)
        assert grid.type is WD_DOC_GRID_TYPE.LINES

    def it_knows_its_line_pitch(self):
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid{w:linePitch=360}"))
        grid = DocumentGrid(docGrid)
        assert grid.line_pitch == 360

    def it_knows_its_char_space(self):
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid{w:charSpace=100}"))
        grid = DocumentGrid(docGrid)
        assert grid.char_space == 100

    def it_returns_None_attributes_when_unset(self):
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid"))
        grid = DocumentGrid(docGrid)
        assert grid.type is None
        assert grid.line_pitch is None
        assert grid.char_space is None

    def it_can_set_its_line_pitch(self):
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid"))
        grid = DocumentGrid(docGrid)
        grid.line_pitch = 360
        assert docGrid.xml == xml("w:docGrid{w:linePitch=360}")

    def it_can_set_its_char_space(self):
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid"))
        grid = DocumentGrid(docGrid)
        grid.char_space = 50
        assert docGrid.xml == xml("w:docGrid{w:charSpace=50}")

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("DEFAULT", "default"),
            ("LINES", "lines"),
            ("LINES_AND_CHARS", "linesAndChars"),
            ("SNAP_TO_CHARS", "snapToChars"),
        ],
    )
    def it_knows_its_type_for_each_enum_value(self, enum_member: str, xml_val: str):
        from docx.enum.section import WD_DOC_GRID_TYPE
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element(f"w:docGrid{{w:type={xml_val}}}"))
        grid = DocumentGrid(docGrid)
        assert grid.type is getattr(WD_DOC_GRID_TYPE, enum_member)

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("DEFAULT", "default"),
            ("LINES", "lines"),
            ("LINES_AND_CHARS", "linesAndChars"),
            ("SNAP_TO_CHARS", "snapToChars"),
        ],
    )
    def it_can_set_its_type_round_trip(self, enum_member: str, xml_val: str):
        from docx.enum.section import WD_DOC_GRID_TYPE
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid"))
        grid = DocumentGrid(docGrid)
        grid.type = getattr(WD_DOC_GRID_TYPE, enum_member)
        assert docGrid.xml == xml(f"w:docGrid{{w:type={xml_val}}}")

    def it_can_clear_its_line_pitch(self):
        from docx.oxml.section import CT_DocGrid
        from docx.section import DocumentGrid

        docGrid = cast(CT_DocGrid, element("w:docGrid{w:linePitch=360}"))
        grid = DocumentGrid(docGrid)
        grid.line_pitch = None
        assert docGrid.xml == xml("w:docGrid")


class DescribeSection_document_grid:
    """Unit-test suite for `docx.section.Section.document_grid` and related API."""

    def it_returns_None_when_no_docGrid(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        assert section.document_grid is None

    def it_provides_access_to_DocumentGrid_when_present(self, document_part_: Mock):
        from docx.section import DocumentGrid

        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:docGrid{w:linePitch=360,w:charSpace=100}"),
        )
        section = Section(sectPr, document_part_)
        grid = section.document_grid
        assert isinstance(grid, DocumentGrid)
        assert grid.line_pitch == 360
        assert grid.char_space == 100

    def it_can_set_document_grid_creating_docGrid(self, document_part_: Mock):
        from docx.enum.section import WD_DOC_GRID_TYPE

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.set_document_grid(
            type=WD_DOC_GRID_TYPE.LINES_AND_CHARS,
            line_pitch=360,
            char_space=100,
        )
        assert sectPr.xml == xml(
            "w:sectPr/w:docGrid{w:type=linesAndChars,w:linePitch=360,w:charSpace=100}"
        )

    def it_returns_a_DocumentGrid_from_set_document_grid(self, document_part_: Mock):
        from docx.section import DocumentGrid

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        grid = section.set_document_grid(line_pitch=360)
        assert isinstance(grid, DocumentGrid)
        assert grid.line_pitch == 360

    def it_leaves_unchanged_attributes_alone_in_set_document_grid(
        self, document_part_: Mock
    ):
        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:docGrid{w:type=lines,w:linePitch=360}"),
        )
        section = Section(sectPr, document_part_)
        section.set_document_grid(line_pitch=400)
        assert sectPr.xml == xml(
            "w:sectPr/w:docGrid{w:type=lines,w:linePitch=400}"
        )

    def it_can_remove_document_grid(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr,
            element("w:sectPr/w:docGrid{w:linePitch=360}"),
        )
        section = Section(sectPr, document_part_)
        section.remove_document_grid()
        assert sectPr.xml == xml("w:sectPr")

    def it_does_nothing_on_remove_when_no_docGrid(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.remove_document_grid()
        assert sectPr.xml == xml("w:sectPr")

    # -- fixtures ---------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeSection_text_direction:
    """Unit-test suite for `docx.section.Section.text_direction`."""

    def it_returns_None_when_no_textDirection(self, document_part_: Mock):
        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        assert section.text_direction is None

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("LR_TB", "lrTb"),
            ("TB_RL", "tbRl"),
            ("BT_LR", "btLr"),
            ("LR_TB_V", "lrTbV"),
            ("TB_RL_V", "tbRlV"),
            ("TB_LR_V", "tbLrV"),
        ],
    )
    def it_knows_its_text_direction_for_each_enum_value(
        self, enum_member: str, xml_val: str, document_part_: Mock
    ):
        from docx.enum.table import WD_TEXT_DIRECTION

        sectPr = cast(
            CT_SectPr, element(f"w:sectPr/w:textDirection{{w:val={xml_val}}}")
        )
        section = Section(sectPr, document_part_)
        assert section.text_direction is getattr(WD_TEXT_DIRECTION, enum_member)

    @pytest.mark.parametrize(
        ("enum_member", "xml_val"),
        [
            ("LR_TB", "lrTb"),
            ("TB_RL", "tbRl"),
            ("BT_LR", "btLr"),
            ("LR_TB_V", "lrTbV"),
            ("TB_RL_V", "tbRlV"),
            ("TB_LR_V", "tbLrV"),
        ],
    )
    def it_can_set_its_text_direction_round_trip(
        self, enum_member: str, xml_val: str, document_part_: Mock
    ):
        from docx.enum.table import WD_TEXT_DIRECTION

        sectPr = cast(CT_SectPr, element("w:sectPr"))
        section = Section(sectPr, document_part_)
        section.text_direction = getattr(WD_TEXT_DIRECTION, enum_member)
        assert sectPr.xml == xml(f"w:sectPr/w:textDirection{{w:val={xml_val}}}")

    def it_can_clear_its_text_direction(self, document_part_: Mock):
        sectPr = cast(
            CT_SectPr, element("w:sectPr/w:textDirection{w:val=tbRl}")
        )
        section = Section(sectPr, document_part_)
        section.text_direction = None
        assert sectPr.xml == xml("w:sectPr")

    # -- fixtures ---------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)


class DescribeSection_right_to_left:
    """Unit-test suite for `docx.section.Section.right_to_left`."""

    @pytest.mark.parametrize(
        ("sectPr_cxml", "expected_value"),
        [
            ("w:sectPr", False),
            ("w:sectPr/w:bidi", True),
            ("w:sectPr/w:bidi{w:val=1}", True),
            ("w:sectPr/w:bidi{w:val=true}", True),
            ("w:sectPr/w:bidi{w:val=on}", True),
            ("w:sectPr/w:bidi{w:val=0}", False),
            ("w:sectPr/w:bidi{w:val=false}", False),
            ("w:sectPr/w:bidi{w:val=off}", False),
        ],
    )
    def it_knows_whether_it_is_right_to_left(
        self, sectPr_cxml: str, expected_value: bool, document_part_: Mock
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)
        assert section.right_to_left is expected_value

    @pytest.mark.parametrize(
        ("sectPr_cxml", "value", "expected_cxml"),
        [
            ("w:sectPr", True, "w:sectPr/w:bidi"),
            ("w:sectPr/w:bidi", False, "w:sectPr"),
            ("w:sectPr/w:bidi", None, "w:sectPr"),
            ("w:sectPr/w:bidi{w:val=off}", True, "w:sectPr/w:bidi"),
            ("w:sectPr", False, "w:sectPr"),
        ],
    )
    def it_can_change_whether_it_is_right_to_left(
        self,
        sectPr_cxml: str,
        value: bool | None,
        expected_cxml: str,
        document_part_: Mock,
    ):
        sectPr = cast(CT_SectPr, element(sectPr_cxml))
        section = Section(sectPr, document_part_)
        section.right_to_left = value
        assert sectPr.xml == xml(expected_cxml)

    # -- fixtures ---------------------------------------------------------------------

    @pytest.fixture
    def document_part_(self, request: FixtureRequest):
        return instance_mock(request, DocumentPart)
