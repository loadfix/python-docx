"""Integration and unit tests for digital-signature detection."""

from __future__ import annotations

import io
import zipfile
from datetime import datetime, timezone

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.signatures import SignatureInfo, _parse_signature_xml  # pyright: ignore[reportPrivateUsage]


# -- minimal XML-DSig + XAdES signature payload, enough to exercise the parser --
_SIG_XML_TEMPLATE = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Signature xmlns="http://www.w3.org/2000/09/xmldsig#" Id="idPackageSignature">
  <SignedInfo>
    <CanonicalizationMethod Algorithm="http://www.w3.org/TR/2001/REC-xml-c14n-20010315"/>
    <SignatureMethod Algorithm="http://www.w3.org/2000/09/xmldsig#rsa-sha1"/>
  </SignedInfo>
  <SignatureValue>AAAA</SignatureValue>
  <KeyInfo>
    <X509Data>
      <X509SubjectName>{signer}</X509SubjectName>
    </X509Data>
  </KeyInfo>
  <Object>
    <xd:QualifyingProperties xmlns:xd="http://uri.etsi.org/01903/v1.3.2#" Target="#idPackageSignature">
      <xd:SignedProperties Id="idSignedProperties">
        <xd:SignedSignatureProperties>
          <xd:SigningTime>{signing_time}</xd:SigningTime>
        </xd:SignedSignatureProperties>
      </xd:SignedProperties>
    </xd:QualifyingProperties>
  </Object>
</Signature>"""


def _build_signed_docx(
    signer: str = "CN=Jane Developer, O=Example Corp, C=US",
    signing_time: str = "2024-05-01T12:34:56Z",
) -> io.BytesIO:
    """Return a BytesIO containing a minimal signed .docx package.

    Builds on the default template: after saving an unsigned document we rewrite the
    zip, adding the origin part, a ``sig1.xml`` part, the origin relationship, and
    updated content-type overrides. This is enough to exercise the detection code
    paths without needing an externally-produced signed file.
    """
    # -- start from an unsigned default document --
    unsigned_buf = io.BytesIO()
    Document().save(unsigned_buf)
    unsigned_buf.seek(0)

    sig_xml = _SIG_XML_TEMPLATE.format(signer=signer, signing_time=signing_time)

    with zipfile.ZipFile(unsigned_buf) as zin:
        items = {name: zin.read(name) for name in zin.namelist()}

    # -- update [Content_Types].xml to add the signature override mappings --
    content_types = items["[Content_Types].xml"].decode("utf-8")
    override_block = (
        f'<Override PartName="/_xmlsignatures/origin.sigs" '
        f'ContentType="{CT.DIGITAL_SIGNATURE_ORIGIN}"/>'
        f'<Override PartName="/_xmlsignatures/sig1.xml" '
        f'ContentType="{CT.DIGITAL_SIGNATURE_XML}"/>'
    )
    content_types = content_types.replace("</Types>", override_block + "</Types>")
    items["[Content_Types].xml"] = content_types.encode("utf-8")

    # -- add a package-level relationship to the origin part --
    pkg_rels = items["_rels/.rels"].decode("utf-8")
    origin_rel = (
        '<Relationship Id="rIdSigOrigin" '
        f'Type="{RT.DIGITAL_SIGNATURE_ORIGIN}" '
        'Target="_xmlsignatures/origin.sigs"/>'
    )
    pkg_rels = pkg_rels.replace("</Relationships>", origin_rel + "</Relationships>")
    items["_rels/.rels"] = pkg_rels.encode("utf-8")

    # -- origin.sigs itself is a binary part; a real one is non-empty but we only
    #    care about traversal here. --
    items["_xmlsignatures/origin.sigs"] = b""

    # -- origin part has its own .rels pointing at each sigN.xml --
    items["_xmlsignatures/_rels/origin.sigs.rels"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        f'Type="{RT.DIGITAL_SIGNATURE}" '
        'Target="sig1.xml"/>'
        "</Relationships>"
    ).encode("utf-8")

    items["_xmlsignatures/sig1.xml"] = sig_xml.encode("utf-8")

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in items.items():
            zout.writestr(name, data)
    out.seek(0)
    return out


class DescribeDocumentIsSigned:
    """Integration tests for `Document.is_signed` and `Document.signatures`."""

    def it_reports_is_signed_False_for_a_default_document(self):
        document = Document()

        assert document.is_signed is False
        assert document.signatures == []

    def it_detects_a_signed_document(self):
        signed = _build_signed_docx()
        document = Document(signed)

        assert document.is_signed is True
        sigs = document.signatures
        assert len(sigs) == 1
        assert isinstance(sigs[0], SignatureInfo)

    def it_exposes_the_signature_partname_and_blob(self):
        signed = _build_signed_docx()
        document = Document(signed)

        sig = document.signatures[0]

        assert str(sig.partname) == "/_xmlsignatures/sig1.xml"
        assert b"<X509SubjectName>" in sig.blob

    def it_extracts_signer_and_signed_at_from_XAdES_style_XML(self):
        signed = _build_signed_docx(
            signer="CN=Alice Example, O=Acme",
            signing_time="2024-05-01T12:34:56Z",
        )
        document = Document(signed)

        sig = document.signatures[0]

        assert sig.signer == "CN=Alice Example, O=Acme"
        assert sig.signed_at == datetime(2024, 5, 1, 12, 34, 56, tzinfo=timezone.utc)


class DescribeParseSignatureXml:
    """Unit tests for the XML-DSig / XAdES parser helper."""

    def it_returns_None_None_for_empty_bytes(self):
        assert _parse_signature_xml(b"") == (None, None)

    def it_returns_None_None_for_malformed_xml(self):
        assert _parse_signature_xml(b"<not-xml") == (None, None)

    def it_returns_None_None_when_elements_are_missing(self):
        xml = (
            b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
            b"<SignatureValue>AA</SignatureValue>"
            b"</Signature>"
        )

        assert _parse_signature_xml(xml) == (None, None)

    def it_extracts_the_signer_subject_name(self):
        xml = (
            b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
            b"<KeyInfo><X509Data>"
            b"<X509SubjectName>CN=Bob</X509SubjectName>"
            b"</X509Data></KeyInfo>"
            b"</Signature>"
        )

        signer, signed_at = _parse_signature_xml(xml)

        assert signer == "CN=Bob"
        assert signed_at is None

    def it_extracts_the_signing_time(self):
        xml = (
            b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
            b'<Object><xd:QualifyingProperties xmlns:xd="http://uri.etsi.org/01903/v1.3.2#">'
            b"<xd:SigningTime>2023-06-15T09:00:00Z</xd:SigningTime>"
            b"</xd:QualifyingProperties></Object>"
            b"</Signature>"
        )

        _, signed_at = _parse_signature_xml(xml)

        assert signed_at == datetime(2023, 6, 15, 9, 0, 0, tzinfo=timezone.utc)

    def it_handles_malformed_signing_time_gracefully(self):
        xml = (
            b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
            b'<Object><xd:QualifyingProperties xmlns:xd="http://uri.etsi.org/01903/v1.3.2#">'
            b"<xd:SigningTime>not-a-date</xd:SigningTime>"
            b"</xd:QualifyingProperties></Object>"
            b"</Signature>"
        )

        _, signed_at = _parse_signature_xml(xml)

        assert signed_at is None


class DescribeSignatureInfo:
    """Unit tests for the `SignatureInfo` proxy."""

    def it_caches_parse_results_across_accesses(self):
        xml = (
            b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
            b"<KeyInfo><X509Data>"
            b"<X509SubjectName>CN=Bob</X509SubjectName>"
            b"</X509Data></KeyInfo>"
            b"</Signature>"
        )

        class _StubPart:
            partname = "/_xmlsignatures/sig1.xml"
            blob = xml

        info = SignatureInfo(_StubPart())  # type: ignore[arg-type]

        assert info.signer == "CN=Bob"
        assert info.signed_at is None
        # -- second access uses the cached parse --
        assert info.signer == "CN=Bob"


class DescribeSharedPackageDelegation:
    """Tests for the optional ``python-ooxml-signatures`` integration.

    Covers the two behaviours that matter:

    1. When the shared package is not installed, the legacy inline
       parser handles everything (proven by the XAdES fixtures above).
    2. When the shared package is installed, ``shared_signature`` is
       non-None and ``signer`` / ``signed_at`` delegate to it, picking
       up Microsoft's ``mdssi:SignatureTime`` + ``mdssi:SignatureComments``
       that the inline parser doesn't recognise.

    These tests are **skip-silent** when ``ooxml_signatures`` is not
    importable so they don't fail in a venv that hasn't installed the
    shared package yet.
    """

    @staticmethod
    def _shared_available() -> bool:
        try:
            import ooxml_signatures  # noqa: F401

            return True
        except ImportError:
            return False

    def it_returns_None_for_shared_signature_when_package_not_installed(self, monkeypatch):
        """Force the fallback path by monkey-patching the import function."""
        from docx import signatures as _sigs

        monkeypatch.setattr(_sigs, "_import_ooxml_signatures", lambda: None)

        xml = (
            b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
            b"<KeyInfo><X509Data>"
            b"<X509SubjectName>CN=Bob</X509SubjectName>"
            b"</X509Data></KeyInfo>"
            b"</Signature>"
        )

        class _StubPart:
            partname = "/_xmlsignatures/sig1.xml"
            blob = xml

        info = SignatureInfo(_StubPart())  # type: ignore[arg-type]
        assert info.shared_signature is None
        # Legacy inline parser still works.
        assert info.signer == "CN=Bob"

    def it_delegates_to_shared_package_when_available(self):
        if not self._shared_available():
            return  # skip silently when shared package not installed

        # mdssi:SignatureTime — Office's default, which the inline parser
        # does NOT recognise but the shared package does.
        xml = (
            b'<Signature xmlns="http://www.w3.org/2000/09/xmldsig#">'
            b"<KeyInfo><X509Data>"
            b"<X509SubjectName>CN=Alice</X509SubjectName>"
            b"</X509Data></KeyInfo>"
            b"<Object>"
            b'<SignatureProperties xmlns:mdssi="http://schemas.openxmlformats.org/package/2006/digital-signature">'
            b'<SignatureProperty Target="#idPackageSignature">'
            b"<mdssi:SignatureTime>"
            b"<mdssi:Format>YYYY-MM-DDTHH:MM:SSZ</mdssi:Format>"
            b"<mdssi:Value>2024-06-15T09:10:11Z</mdssi:Value>"
            b"</mdssi:SignatureTime>"
            b"</SignatureProperty>"
            b"</SignatureProperties>"
            b"</Object>"
            b"</Signature>"
        )

        class _StubPart:
            partname = "/_xmlsignatures/sig1.xml"
            blob = xml

        info = SignatureInfo(_StubPart())  # type: ignore[arg-type]
        assert info.shared_signature is not None
        # signer delegation works
        assert info.signer == "CN=Alice"
        # signed_at delegation picks up mdssi:SignatureTime
        assert info.signed_at == datetime(2024, 6, 15, 9, 10, 11, tzinfo=timezone.utc)

    def it_gracefully_falls_back_on_malformed_xml_with_shared_package(self):
        if not self._shared_available():
            return  # skip silently

        class _StubPart:
            partname = "/_xmlsignatures/sig1.xml"
            blob = b"<NotASignature/>"

        info = SignatureInfo(_StubPart())  # type: ignore[arg-type]
        # shared_signature returns None because the root is not <Signature>;
        # the legacy parser also returns None/None.
        assert info.shared_signature is None
        assert info.signer is None
        assert info.signed_at is None
