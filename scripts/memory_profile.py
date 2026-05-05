#!/usr/bin/env python3
"""W11-B memory profile — python-docx on W6-D scale fixtures.

Loads each W6-D scale fixture (author it on the fly if not present — the
target scales up to 10k paragraphs are what W6-D defined), then records
RSS memory at four checkpoints:

    1. pre-load      — before ``Document(path)``
    2. post-load     — after parse, before manipulation
    3. post-manipulate — after iterating all paragraphs and doing a
                         touch-edit on each (representative read/write
                         workload)
    4. post-save     — after round-tripping to a BytesIO

Peak RSS is sampled from ``resource.getrusage(RUSAGE_SELF).ru_maxrss``
(platform: Linux — value is in KiB). We also expose the Python-heap
``tracemalloc`` peak for comparison, because ``ru_maxrss`` is
high-water-mark only and never decreases within a single process.

Each fixture runs in a fresh subprocess so ``ru_maxrss`` reflects that
fixture alone, not the cumulative high-water of earlier fixtures.

Output is a markdown report at ``memory_profile_report.md`` plus JSON at
``memory_profile_report.json``.

Usage::

    python scripts/memory_profile.py             # run all scales
    python scripts/memory_profile.py --scale 1k  # single scale
    python scripts/memory_profile.py --child 1k  # internal (subprocess)
"""
from __future__ import annotations

import argparse
import json
import os
import resource
import subprocess
import sys
import tempfile
import tracemalloc
from pathlib import Path

# --- scale fixtures (W6-D targets for python-docx) ----------------------

# (name, paragraph_count). A runs is 4 runs/paragraph, each ~10 chars, so
# total text ~40 chars/paragraph. 10k paragraphs ≈ ~3 MiB of raw XML.
SCALES: list[tuple[str, int]] = [
    ("100p", 100),
    ("1k", 1_000),
    ("5k", 5_000),
    ("10k", 10_000),
]


def build_fixture(path: Path, n_paragraphs: int) -> None:
    """Create a .docx with *n_paragraphs* paragraphs."""
    from docx import Document  # noqa: PLC0415

    doc = Document()
    for i in range(n_paragraphs):
        # 4 runs with mixed formatting — representative of real documents.
        p = doc.add_paragraph()
        r1 = p.add_run(f"Paragraph {i} opening sentence. ")
        r1.bold = True
        p.add_run("A second run of regular text follows it. ")
        r3 = p.add_run("A third emphasised clause. ")
        r3.italic = True
        p.add_run(f"And a closing numeric {i * 3}.")
    doc.save(str(path))


# --- single-scale worker (runs in a subprocess) -------------------------

def _rss_kib() -> int:
    """Resident-set-size high-water-mark in KiB (Linux convention)."""
    return resource.getrusage(resource.RUSAGE_SELF).ru_maxrss


def profile_one(fixture_path: Path) -> dict:
    """Return memory checkpoints for a single fixture."""
    import io  # noqa: PLC0415

    # Sample RSS before the library is imported so we can report the
    # interpreter baseline separately.
    baseline_rss = _rss_kib()

    from docx import Document  # noqa: PLC0415

    tracemalloc.start()

    # 1. pre-load — library imported, fixture not yet parsed
    pre_load_rss = _rss_kib()
    pre_load_tm = tracemalloc.get_traced_memory()[0]

    # 2. post-load
    doc = Document(str(fixture_path))
    post_load_rss = _rss_kib()
    post_load_tm = tracemalloc.get_traced_memory()[0]

    # 3. post-manipulate — touch every paragraph
    n_para = 0
    for para in doc.paragraphs:
        _ = para.text  # force read of every run
        n_para += 1
    # small mutation burst to exercise write-path
    for para in doc.paragraphs[:50]:
        para.add_run(" [touched]")
    post_manip_rss = _rss_kib()
    post_manip_tm = tracemalloc.get_traced_memory()[0]

    # 4. post-save
    buf = io.BytesIO()
    doc.save(buf)
    saved_size = len(buf.getvalue())
    post_save_rss = _rss_kib()
    post_save_tm = tracemalloc.get_traced_memory()[0]
    peak_tm = tracemalloc.get_traced_memory()[1]
    tracemalloc.stop()

    return {
        "fixture": str(fixture_path),
        "fixture_size_bytes": fixture_path.stat().st_size,
        "saved_size_bytes": saved_size,
        "paragraphs_visited": n_para,
        "baseline_rss_kib": baseline_rss,
        "rss_kib": {
            "pre_load": pre_load_rss,
            "post_load": post_load_rss,
            "post_manipulate": post_manip_rss,
            "post_save": post_save_rss,
        },
        "tracemalloc_bytes": {
            "pre_load": pre_load_tm,
            "post_load": post_load_tm,
            "post_manipulate": post_manip_tm,
            "post_save": post_save_tm,
            "peak": peak_tm,
        },
    }


# --- orchestrator -------------------------------------------------------

def run_child(scale_name: str, fixture_dir: Path) -> dict:
    path = fixture_dir / f"scale_{scale_name}.docx"
    if not path.exists():
        count = dict(SCALES)[scale_name]
        build_fixture(path, count)
    return profile_one(path)


def run_all(fixture_dir: Path) -> list[dict]:
    results: list[dict] = []
    for scale_name, count in SCALES:
        path = fixture_dir / f"scale_{scale_name}.docx"
        if not path.exists():
            print(f"[build] {scale_name}: {count} paragraphs -> {path}",
                  file=sys.stderr)
            build_fixture(path, count)
        # run the profile in a fresh subprocess so ru_maxrss is clean
        print(f"[profile] {scale_name}", file=sys.stderr)
        proc = subprocess.run(
            [sys.executable, __file__, "--child", scale_name,
             "--fixture-dir", str(fixture_dir)],
            capture_output=True,
            text=True,
            check=True,
        )
        entry = json.loads(proc.stdout)
        entry["scale"] = scale_name
        entry["paragraphs_target"] = count
        results.append(entry)
    return results


def render_report(results: list[dict]) -> str:
    lines: list[str] = []
    lines.append("# python-docx — W11-B memory profile report\n")
    lines.append(
        "Recorded via `resource.getrusage(RUSAGE_SELF).ru_maxrss` "
        "(high-water RSS, KiB) plus `tracemalloc` peak (Python heap, bytes). "
        "Each scale runs in a fresh subprocess so `ru_maxrss` reflects "
        "that fixture alone.\n"
    )
    lines.append(
        "The `Library-attributable` column is `peak RSS − pre-import "
        "baseline`: it excludes the cold Python interpreter "
        "(~30–40 MiB on Linux) which is not the library's responsibility. "
        "This is what we compare against the on-disk fixture size when "
        "flagging outliers (>5x).\n"
    )
    lines.append(
        "| Scale | Paragraphs | Fixture (KiB) | Baseline RSS (MiB) | "
        "Peak RSS (MiB) | Lib-attributable (MiB) | Peak Py heap (MiB) | "
        "Lib RSS / fixture |"
    )
    lines.append("|---|---|---|---|---|---|---|---|")
    for r in results:
        fsize_kib = r["fixture_size_bytes"] / 1024
        rss = r["rss_kib"]
        baseline_mib = r["baseline_rss_kib"] / 1024
        peak_rss_kib = max(rss.values())
        peak_rss_mib = peak_rss_kib / 1024
        lib_attr_kib = max(0, peak_rss_kib - r["baseline_rss_kib"])
        lib_attr_mib = lib_attr_kib / 1024
        tm_peak_mib = r["tracemalloc_bytes"]["peak"] / (1024 * 1024)
        ratio = (lib_attr_kib * 1024) / r["fixture_size_bytes"] if \
            r["fixture_size_bytes"] else float("inf")
        flag = "  FLAG" if ratio > 5 else ""
        lines.append(
            f"| {r['scale']} | {r['paragraphs_target']:,} | "
            f"{fsize_kib:,.1f} | {baseline_mib:,.1f} | "
            f"{peak_rss_mib:,.1f} | {lib_attr_mib:,.1f} | "
            f"{tm_peak_mib:,.1f} | {ratio:.1f}x{flag} |"
        )
    lines.append("")
    lines.append(
        "## Per-checkpoint RSS deltas (MiB above pre-import baseline)\n"
    )
    lines.append("| Scale | pre-load | post-load | post-manip | post-save |")
    lines.append("|---|---|---|---|---|")
    for r in results:
        base = r["baseline_rss_kib"]
        rss = r["rss_kib"]
        lines.append(
            f"| {r['scale']} | "
            f"{(rss['pre_load'] - base) / 1024:,.1f} | "
            f"{(rss['post_load'] - base) / 1024:,.1f} | "
            f"{(rss['post_manipulate'] - base) / 1024:,.1f} | "
            f"{(rss['post_save'] - base) / 1024:,.1f} |"
        )
    lines.append("")
    lines.append("## Outliers (library-attributable RSS > 5x fixture size)\n")
    outliers = []
    for r in results:
        peak_rss_kib = max(r["rss_kib"].values())
        lib_attr_kib = max(0, peak_rss_kib - r["baseline_rss_kib"])
        if r["fixture_size_bytes"] and \
                (lib_attr_kib * 1024) > 5 * r["fixture_size_bytes"]:
            outliers.append(r)
    if outliers:
        for r in outliers:
            peak_rss_kib = max(r["rss_kib"].values())
            lib_attr_mib = (peak_rss_kib - r["baseline_rss_kib"]) / 1024
            fsize_kib = r["fixture_size_bytes"] / 1024
            ratio = ((peak_rss_kib - r["baseline_rss_kib"]) * 1024) / \
                r["fixture_size_bytes"]
            lines.append(
                f"- **{r['scale']}** ({r['paragraphs_target']:,} paragraphs): "
                f"library-attributable RSS {lib_attr_mib:,.1f} MiB vs "
                f"fixture {fsize_kib:,.1f} KiB on disk ({ratio:.1f}x)"
            )
    else:
        lines.append("_None_\n")
    lines.append("")
    return "\n".join(lines)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument(
        "--fixture-dir",
        default=None,
        help="directory for generated scale fixtures (defaults to tempdir)",
    )
    ap.add_argument(
        "--child",
        default=None,
        help="internal: profile a single scale and emit JSON on stdout",
    )
    ap.add_argument(
        "--report",
        default="memory_profile_report.md",
        help="output markdown report path",
    )
    args = ap.parse_args()

    fixture_dir = (
        Path(args.fixture_dir) if args.fixture_dir
        else Path(tempfile.gettempdir()) / "w11_b_docx_fixtures"
    )
    fixture_dir.mkdir(parents=True, exist_ok=True)

    if args.child:
        result = run_child(args.child, fixture_dir)
        print(json.dumps(result))
        return 0

    results = run_all(fixture_dir)
    Path(args.report).write_text(render_report(results))
    json_path = Path(args.report).with_suffix(".json")
    json_path.write_text(json.dumps(results, indent=2))
    print(f"wrote {args.report} and {json_path}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
