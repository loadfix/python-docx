"""Step implementations for Strict-OOXML packaging behaviour.

The Strict fixture is generated on-the-fly by round-tripping the bundled
default template through a namespace rewrite, so no binary fixture needs
checked into the repository. Closes upstream#1520, upstream#693.
"""

from __future__ import annotations

import io
import os
import zipfile

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.opc.strict import STRICT_TO_TRANSITIONAL

from helpers import saved_docx_path, scratch_dir


_STRICT_WML = b"http://purl.oclc.org/ooxml/wordprocessingml/main"


def _default_template_path() -> str:
    here = os.path.dirname(__file__)
    return os.path.abspath(
        os.path.join(here, "..", "..", "src", "docx", "templates", "default.docx")
    )


def _make_strict_docx_bytes() -> bytes:
    """Return a Strict-OOXML variant of the default template.

    Every Transitional namespace URI inside each XML member of the bundled
    default template is rewritten to its Strict equivalent. Binary members
    pass through untouched.
    """
    transitional_to_strict = {
        trans: strict for strict, trans in STRICT_TO_TRANSITIONAL.items()
    }
    with open(_default_template_path(), "rb") as f:
        blob = f.read()
    out = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(blob), "r") as zin:
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith((".xml", ".rels")):
                    for trans, strict in transitional_to_strict.items():
                        data = data.replace(trans, strict)
                zout.writestr(item, data)
    return out.getvalue()


# given ===================================================


@given("a Strict-OOXML .docx document")
def given_a_strict_docx(context: Context):
    if not os.path.isdir(scratch_dir):
        os.makedirs(scratch_dir)
    path = os.path.join(scratch_dir, "strict-fixture.docx")
    with open(path, "wb") as f:
        f.write(_make_strict_docx_bytes())
    context.docx_path = path
    context.document = Document(path)


# when ====================================================


@when("I save the document to the scratch path")
def when_save_to_scratch(context: Context):
    if not os.path.isdir(scratch_dir):
        os.makedirs(scratch_dir)
    context.saved_path = saved_docx_path
    context.document.save(saved_docx_path)


# then ====================================================


@then("document.paragraphs is iterable")
def then_document_paragraphs_iterable(context: Context):
    # -- success is just being able to walk them without raising --
    texts = [p.text for p in context.document.paragraphs]
    assert isinstance(texts, list)


@then("the saved package contains no Strict namespace URIs")
def then_saved_has_no_strict_uris(context: Context):
    with zipfile.ZipFile(context.saved_path, "r") as zf:
        for name in zf.namelist():
            data = zf.read(name)
            assert _STRICT_WML not in data, (
                f"{name} still contains Strict WML namespace"
            )
