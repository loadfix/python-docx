"""Step implementations for tracked-changes features (read + accept + reject).

Covers:
* trk-read-ins-del — iterating Paragraph.tracked_changes
* trk-format-changes — FormattingChange for rPrChange/pPrChange/sectPrChange
* trk-move-revisions — MoveRevision pairing via @w:name + .peer
* trk-table-changes — cell/row revisions (cellIns, cellDel, tcPrChange, trPrChange,
  tblPrChange)
* trk-rsid — revision-save IDs (rsidRoot + per-run/paragraph rsid)
* trk-marks-text — revision_marks_text() preview output
* trk-accept-insertions — Document.accept_all_changes and per-change
  TrackedChange.accept on w:ins / w:cellIns wrappers
* trk-reject-changes — reject-side behaviour for Document.reject_all_changes()
  and per-change TrackedChange.reject()
"""

from __future__ import annotations

import datetime as dt

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.tracked_changes import MoveRevision, TrackedChange

from helpers import test_docx

# given ====================================================


@given("a document with a tracked insertion and deletion in paragraph 1")
def given_a_document_with_tracked_ins_and_del_in_paragraph_1(context: Context):
    context.document = Document(test_docx("trk-ins-del"))
    context.paragraph = context.document.paragraphs[1]
    context.tracked_changes = context.paragraph.tracked_changes


@given("the trk-ins-del document")
def given_the_trk_ins_del_document(context: Context):
    context.document = Document(test_docx("trk-ins-del"))


@given("the trk-format document")
def given_the_trk_format_document(context: Context):
    context.document = Document(test_docx("trk-format"))


@given("the trk-move document")
def given_the_trk_move_document(context: Context):
    context.document = Document(test_docx("trk-move"))


@given("the trk-table document")
def given_the_trk_table_document(context: Context):
    context.document = Document(test_docx("trk-table"))
    # -- `context.table` is reserved by behave's step-result reporting, use a
    # -- different attribute name --
    context.trk_table = context.document.tables[0]


@given("the trk-rsid document")
def given_the_trk_rsid_document(context: Context):
    context.document = Document(test_docx("trk-rsid"))


@given("the trk-marks document")
def given_the_trk_marks_document(context: Context):
    context.document = Document(test_docx("trk-marks"))


@given("the trk-accept-del document")
def given_the_trk_accept_del_document(context: Context):
    context.document = Document(test_docx("trk-accept-del"))
    context.accept_count = None


# when =====================================================


@when("I select the formatting_change of run {run_idx:d} on paragraph {p_idx:d}")
def when_i_select_formatting_change_of_run_on_paragraph(
    context: Context, run_idx: int, p_idx: int
):
    run = context.document.paragraphs[p_idx].runs[run_idx]
    context.formatting_change = run.formatting_change


@when("I select the formatting_change of paragraph {p_idx:d}")
def when_i_select_formatting_change_of_paragraph(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    context.formatting_change = paragraph.formatting_change


@when("I select the formatting_change of section {s_idx:d}")
def when_i_select_formatting_change_of_section(context: Context, s_idx: int):
    section = context.document.sections[s_idx]
    context.formatting_change = section.formatting_change


@when("I select the tracked_changes of paragraph {p_idx:d}")
def when_i_select_tracked_changes_of_paragraph(context: Context, p_idx: int):
    context.paragraph = context.document.paragraphs[p_idx]
    context.tracked_changes = context.paragraph.tracked_changes


@when(
    "I call paragraph {p_idx:d} revision_marks_text with custom <INS>/<DEL> markers"
)
def when_i_call_paragraph_revision_marks_text_custom(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    context.custom_preview = paragraph.revision_marks_text(
        open_ins="<INS>",
        close_ins="</INS>",
        open_del="<DEL>",
        close_del="</DEL>",
    )


# then =====================================================


@then("paragraph.tracked_changes has {count:d} entry")
@then("paragraph.tracked_changes has {count:d} entries")
def then_paragraph_tracked_changes_has_count(context: Context, count: int):
    actual = len(context.tracked_changes)
    assert actual == count, f"expected {count} tracked change(s), got {actual}"


@then("the tracked-change types are {expected}")
def then_tracked_change_types_are(context: Context, expected: str):
    # -- parse the literal list representation, e.g. "['deletion', 'insertion']" --
    parsed = eval(expected, {"__builtins__": {}}, {})
    actual = [tc.type for tc in context.tracked_changes]
    assert actual == parsed, f"expected types {parsed!r}, got {actual!r}"


@then("the tracked-change authors are {expected}")
def then_tracked_change_authors_are(context: Context, expected: str):
    parsed = eval(expected, {"__builtins__": {}}, {})
    actual = [tc.author for tc in context.tracked_changes]
    assert actual == parsed, f"expected authors {parsed!r}, got {actual!r}"


@then("the tracked-change texts are {expected}")
def then_tracked_change_texts_are(context: Context, expected: str):
    parsed = eval(expected, {"__builtins__": {}}, {})
    actual = [tc.text for tc in context.tracked_changes]
    assert actual == parsed, f"expected texts {parsed!r}, got {actual!r}"


@then('the first tracked change author is "{author}"')
def then_first_tracked_change_author_is(context: Context, author: str):
    actual = context.tracked_changes[0].author
    assert actual == author, f"expected author {author!r}, got {actual!r}"


@then("the first tracked change date is a datetime")
def then_first_tracked_change_date_is_a_datetime(context: Context):
    actual = context.tracked_changes[0].date
    assert isinstance(actual, dt.datetime), (
        f"expected datetime, got {type(actual).__name__}"
    )


@then('tracked_change[{idx:d}].type == "{expected}"')
def then_tracked_change_idx_type_eq(context: Context, idx: int, expected: str):
    actual = context.tracked_changes[idx].type
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then('tracked_change[{idx:d}].name == "{expected}"')
def then_tracked_change_idx_name_eq(context: Context, idx: int, expected: str):
    tc = context.tracked_changes[idx]
    assert isinstance(tc, MoveRevision), (
        f"expected MoveRevision at index {idx}, got {type(tc).__name__}"
    )
    actual = tc.name
    assert actual == expected, f"expected name {expected!r}, got {actual!r}"


@then("paragraph {p_idx:d} has no tracked changes")
def then_paragraph_p_idx_has_no_tracked_changes(context: Context, p_idx: int):
    actual = context.document.paragraphs[p_idx].tracked_changes
    assert actual == [], f"expected [], got {actual!r}"


@then(
    "iterating every paragraph's tracked_changes yields"
    " {count:d} TrackedChange objects"
)
def then_iterating_every_paragraphs_tracked_changes_yields(
    context: Context, count: int
):
    collected: list[TrackedChange] = []
    for paragraph in context.document.paragraphs:
        collected.extend(paragraph.tracked_changes)
    assert len(collected) == count, (
        f"expected {count} TrackedChange objects, got {len(collected)}"
    )
    assert all(isinstance(tc, TrackedChange) for tc in collected), (
        "expected every entry to be a TrackedChange"
    )


# -- formatting change assertions --


@then("the formatting change is not None")
def then_the_formatting_change_is_not_none(context: Context):
    assert context.formatting_change is not None, "expected a FormattingChange"


@then("the formatting change is None")
def then_the_formatting_change_is_none(context: Context):
    assert context.formatting_change is None, (
        f"expected None, got {context.formatting_change!r}"
    )


@then('formatting_change.author == "{author}"')
def then_formatting_change_author_eq(context: Context, author: str):
    actual = context.formatting_change.author
    assert actual == author, f"expected author {author!r}, got {actual!r}"


@then("formatting_change.old_properties is not None")
def then_formatting_change_old_properties_is_not_none(context: Context):
    assert context.formatting_change.old_properties is not None, (
        "expected old_properties to be populated"
    )


# -- move revision assertions --


@then("the first tracked change is a MoveRevision")
def then_the_first_tracked_change_is_a_move_revision(context: Context):
    tc = context.tracked_changes[0]
    assert isinstance(tc, MoveRevision), (
        f"expected MoveRevision, got {type(tc).__name__}"
    )


@then('the peer of the first move revision has type "{expected}"')
def then_peer_of_first_move_has_type(context: Context, expected: str):
    move = context.tracked_changes[0]
    assert isinstance(move, MoveRevision)
    peer = move.peer
    assert peer is not None, "expected a peer"
    assert peer.type == expected, (
        f"expected peer type {expected!r}, got {peer.type!r}"
    )


@then('the peer of the first move revision has name "{expected}"')
def then_peer_of_first_move_has_name(context: Context, expected: str):
    move = context.tracked_changes[0]
    assert isinstance(move, MoveRevision)
    peer = move.peer
    assert peer is not None
    assert peer.name == expected, (
        f"expected peer name {expected!r}, got {peer.name!r}"
    )


# -- table-level assertions --


@then("cell ({r:d}, {c:d}).is_tracked_insertion is {flag}")
def then_cell_is_tracked_insertion(context: Context, r: int, c: int, flag: str):
    expected = {"True": True, "False": False}[flag]
    actual = context.trk_table.cell(r, c).is_tracked_insertion
    assert actual is expected, (
        f"cell({r},{c}).is_tracked_insertion expected {expected}, got {actual}"
    )


@then("cell ({r:d}, {c:d}).is_tracked_deletion is {flag}")
def then_cell_is_tracked_deletion(context: Context, r: int, c: int, flag: str):
    expected = {"True": True, "False": False}[flag]
    actual = context.trk_table.cell(r, c).is_tracked_deletion
    assert actual is expected, (
        f"cell({r},{c}).is_tracked_deletion expected {expected}, got {actual}"
    )


@then('cell ({r:d}, {c:d}).formatting_change.author == "{author}"')
def then_cell_formatting_change_author_eq(
    context: Context, r: int, c: int, author: str
):
    fc = context.trk_table.cell(r, c).formatting_change
    assert fc is not None, f"cell({r},{c}) had no formatting_change"
    assert fc.author == author, f"expected author {author!r}, got {fc.author!r}"


@then('row {row_idx:d} formatting_change.author == "{author}"')
def then_row_formatting_change_author_eq(
    context: Context, row_idx: int, author: str
):
    fc = context.trk_table.rows[row_idx].formatting_change
    assert fc is not None, f"row {row_idx} had no formatting_change"
    assert fc.author == author, f"expected author {author!r}, got {fc.author!r}"


@then('table.formatting_change.author == "{author}"')
def then_table_formatting_change_author_eq(context: Context, author: str):
    fc = context.trk_table.formatting_change
    assert fc is not None, "table had no formatting_change"
    assert fc.author == author, f"expected author {author!r}, got {fc.author!r}"


@then("row {row_idx:d} has no formatting_change")
def then_row_has_no_formatting_change(context: Context, row_idx: int):
    fc = context.trk_table.rows[row_idx].formatting_change
    assert fc is None, f"expected None, got {fc!r}"


# -- rsid assertions --


@then('document.settings.rsid_root == "{expected}"')
def then_document_settings_rsid_root_eq(context: Context, expected: str):
    actual = context.document.settings.rsid_root
    assert actual == expected, (
        f"expected rsid_root {expected!r}, got {actual!r}"
    )


@then("document.settings.rsids == {expected}")
def then_document_settings_rsids_eq(context: Context, expected: str):
    parsed = eval(expected, {"__builtins__": {}}, {})
    actual = context.document.settings.rsids
    assert actual == parsed, f"expected rsids {parsed!r}, got {actual!r}"


@then('paragraph {p_idx:d} rsid == "{expected}"')
def then_paragraph_rsid_eq(context: Context, p_idx: int, expected: str):
    actual = context.document.paragraphs[p_idx].rsid
    assert actual == expected, f"expected rsid {expected!r}, got {actual!r}"


@then("paragraph {p_idx:d} rsid is None")
def then_paragraph_rsid_is_none(context: Context, p_idx: int):
    actual = context.document.paragraphs[p_idx].rsid
    assert actual is None, f"expected None, got {actual!r}"


@then('paragraph {p_idx:d} run {r_idx:d} rsid == "{expected}"')
def then_paragraph_run_rsid_eq(
    context: Context, p_idx: int, r_idx: int, expected: str
):
    actual = context.document.paragraphs[p_idx].runs[r_idx].rsid
    assert actual == expected, f"expected rsid {expected!r}, got {actual!r}"


@then("paragraph {p_idx:d} run {r_idx:d} rsid is None")
def then_paragraph_run_rsid_is_none(context: Context, p_idx: int, r_idx: int):
    actual = context.document.paragraphs[p_idx].runs[r_idx].rsid
    assert actual is None, f"expected None, got {actual!r}"


# -- revision_marks_text assertions --


@then("paragraph {p_idx:d} revision_marks_text() matches paragraph.text")
def then_paragraph_revision_marks_text_matches_text(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    preview = paragraph.revision_marks_text()
    assert preview == paragraph.text, (
        f"expected preview to equal text {paragraph.text!r}, got {preview!r}"
    )


@then('paragraph {p_idx:d} revision_marks_text() == "{expected}"')
def then_paragraph_revision_marks_text_eq(
    context: Context, p_idx: int, expected: str
):
    paragraph = context.document.paragraphs[p_idx]
    actual = paragraph.revision_marks_text()
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then('the custom-marker preview == "{expected}"')
def then_custom_marker_preview_eq(context: Context, expected: str):
    actual = context.custom_preview
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("document.revision_marks_text() ends with the final paragraph's preview")
def then_document_revision_marks_text_ends_with_final(context: Context):
    final_preview = context.document.paragraphs[-1].revision_marks_text()
    full = context.document.revision_marks_text()
    assert full.endswith(final_preview), (
        f"expected document preview to end with {final_preview!r}, got {full!r}"
    )


@then("document.revision_marks_text() contains the insertion-only preview")
def then_document_revision_marks_text_contains_insertion_only(context: Context):
    full = context.document.revision_marks_text()
    assert "[+kindly +]" in full, (
        f"expected insertion marker in document preview, got {full!r}"
    )


# -- accept-insertions steps ================================


@given("the trk-accept-ins document")
def given_the_trk_accept_ins_document(context: Context):
    context.document = Document(test_docx("trk-accept-ins"))


@when("I call document.accept_all_changes()")
def when_i_call_document_accept_all_changes(context: Context):
    context.accept_count = context.document.accept_all_changes()


@when("I accept the first tracked change of paragraph {p_idx:d}")
def when_i_accept_first_tracked_change_of_paragraph(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    paragraph.tracked_changes[0].accept()


@then("the document body has {count:d} w:ins elements")
def then_document_body_has_count_w_ins(context: Context, count: int):
    actual = len(context.document.element.body.xpath(".//w:ins"))
    assert actual == count, f"expected {count} w:ins, got {actual}"


@then("the document body has {count:d} w:del elements")
def then_document_body_has_count_w_del(context: Context, count: int):
    actual = len(context.document.element.body.xpath(".//w:del"))
    assert actual == count, f"expected {count} w:del, got {actual}"


@then("the document body has {count:d} w:cellIns elements")
def then_document_body_has_count_w_cell_ins(context: Context, count: int):
    actual = len(context.document.element.body.xpath(".//w:cellIns"))
    assert actual == count, f"expected {count} w:cellIns, got {actual}"


@then('paragraph {p_idx:d} text == "{expected}"')
def then_paragraph_text_eq(context: Context, p_idx: int, expected: str):
    actual = context.document.paragraphs[p_idx].text
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("paragraph {p_idx:d} has {count:d} direct w:r children")
def then_paragraph_has_count_direct_w_r(context: Context, p_idx: int, count: int):
    actual = len(context.document.paragraphs[p_idx]._p.xpath("./w:r"))
    assert actual == count, f"expected {count} direct w:r, got {actual}"
# -- reject-side steps =====================================


def _strip_elements(document: Document, xpath: str) -> None:
    """Detach every element matching `xpath` from the document body."""
    for elm in document._element.body.xpath(xpath):
        parent = elm.getparent()
        if parent is not None:
            parent.remove(elm)


def _unwrap_elements(document: Document, xpath: str) -> None:
    """Replace every matching element with its children, in place."""
    for elm in document._element.body.xpath(xpath):
        parent = elm.getparent()
        if parent is None:
            continue
        idx = parent.index(elm)
        for i, child in enumerate(list(elm)):
            parent.insert(idx + i, child)
        parent.remove(elm)


@given("the trk-reject document")
def given_the_trk_reject_document(context: Context):
    context.document = Document(test_docx("trk-reject"))


@given("the trk-reject document with only insertions")
def given_the_trk_reject_document_only_insertions(context: Context):
    # -- drop every deletion / move / cell marker so the doc only carries
    # -- insertions. The remaining scaffolding keeps paragraph indices 1 and
    # -- 2 stable (the wrapped "w:ins" in those paragraphs survives). --
    document = Document(test_docx("trk-reject"))
    _strip_elements(document, ".//w:del | .//w:moveFrom")
    # -- moveTo contains w:t; unwrapping keeps the destination content but
    # -- removes the revision wrapper so it no longer counts as an insertion
    _unwrap_elements(document, ".//w:moveTo")
    # -- remove cell markers without disturbing the cells themselves
    _strip_elements(document, ".//w:cellIns | .//w:cellDel")
    context.document = document


@given("the trk-reject document with only deletions")
def given_the_trk_reject_document_only_deletions(context: Context):
    # -- mirror of the insertions-only variant: keep every w:del, strip the
    # -- rest. Stripping (not unwrapping) the w:ins wrappers makes the
    # -- pre-reject content match what a reviewer who only deleted text would
    # -- see, with stable paragraph indices preserved. --
    document = Document(test_docx("trk-reject"))
    _strip_elements(document, ".//w:ins | .//w:moveFrom | .//w:moveTo")
    _strip_elements(document, ".//w:cellIns | .//w:cellDel")
    context.document = document


@when("I call document.reject_all_changes()")
def when_i_call_document_reject_all_changes(context: Context):
    context.reject_count = context.document.reject_all_changes()


@when("I reject the insertion in paragraph {p_idx:d}")
def when_i_reject_the_insertion_in_paragraph(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    for tc in paragraph.tracked_changes:
        if tc.type == "insertion":
            tc.reject()
            return
    raise AssertionError(f"paragraph {p_idx} has no insertion to reject")


@when("I reject the deletion in paragraph {p_idx:d}")
def when_i_reject_the_deletion_in_paragraph(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    for tc in paragraph.tracked_changes:
        if tc.type == "deletion":
            tc.reject()
            return
    raise AssertionError(f"paragraph {p_idx} has no deletion to reject")


@when("I reject every tracked change in paragraph {p_idx:d}")
def when_i_reject_every_tracked_change_in_paragraph(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    # -- snapshot the list; rejecting mutates the paragraph's XML tree and
    # -- invalidates any subsequent iteration over the live proxy list
    for tc in list(paragraph.tracked_changes):
        tc.reject()


@then("the reject count is {count:d}")
def then_the_reject_count_is(context: Context, count: int):
    assert context.reject_count == count, (
        f"expected reject count {count}, got {context.reject_count}"
    )


@then("the document has no {tag} elements")
def then_the_document_has_no_tag_elements(context: Context, tag: str):
    found = context.document._element.body.xpath(f".//{tag}")
    assert not found, (
        f"expected no {tag} elements, found {len(found)}"
    )


@then('paragraph {p_idx:d} text equals "{expected}"')
def then_paragraph_text_equals(context: Context, p_idx: int, expected: str):
    actual = context.document.paragraphs[p_idx].text
    assert actual == expected, (
        f"paragraph {p_idx} text mismatch: expected {expected!r}, got {actual!r}"
    )


@then('paragraph {p_idx:d} revision_marks_text equals "{expected}"')
def then_paragraph_revision_marks_text_equals(
    context: Context, p_idx: int, expected: str
):
    actual = context.document.paragraphs[p_idx].revision_marks_text()
    assert actual == expected, (
        f"paragraph {p_idx} revision_marks_text mismatch:"
        f" expected {expected!r}, got {actual!r}"
    )


@then("paragraph {p_idx:d} has {count:d} tracked change remaining")
@then("paragraph {p_idx:d} has {count:d} tracked changes remaining")
def then_paragraph_has_count_tracked_changes_remaining(
    context: Context, p_idx: int, count: int
):
    actual = len(context.document.paragraphs[p_idx].tracked_changes)
    assert actual == count, (
        f"expected {count} tracked changes remaining, got {actual}"
    )


# -- accept-side steps for tracked deletions -------------------------------------


@when("I call document.accept_all_changes")
def when_i_call_document_accept_all_changes(context: Context):
    # -- accumulate when invoked more than once in a single scenario --
    prior = context.accept_count or 0
    context.accept_count = context.document.accept_all_changes()
    context.accept_count_cumulative = prior + context.accept_count


@when("I accept the only tracked change on paragraph {p_idx:d}")
def when_i_accept_the_only_tracked_change_on_paragraph(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    changes = paragraph.tracked_changes
    assert len(changes) == 1, (
        f"expected exactly 1 tracked change on paragraph {p_idx}, got {len(changes)}"
    )
    changes[0].accept()


@when("I accept every deletion-typed tracked change on paragraph {p_idx:d}")
def when_i_accept_every_deletion_on_paragraph(context: Context, p_idx: int):
    paragraph = context.document.paragraphs[p_idx]
    # -- iterate over a snapshot because accept() mutates the underlying XML --
    for tc in list(paragraph.tracked_changes):
        if tc.type == "deletion":
            tc.accept()


@then("the accept-changes count is {count:d}")
def then_the_accept_changes_count_is(context: Context, count: int):
    assert context.accept_count == count, (
        f"expected accept-changes count {count}, got {context.accept_count}"
    )


@then("the document has {count:d} w:del elements")
def then_the_document_has_n_w_del_elements(context: Context, count: int):
    dels = context.document._element.body.xpath(".//w:del")
    assert len(dels) == count, (
        f"expected {count} w:del elements, got {len(dels)}"
    )


@then('paragraph {p_idx:d} text is ""')
def then_paragraph_p_idx_text_is_empty(context: Context, p_idx: int):
    actual = context.document.paragraphs[p_idx].text
    assert actual == "", f"paragraph {p_idx} text: expected empty, got {actual!r}"


@then('paragraph {p_idx:d} text is "{expected}"')
def then_paragraph_p_idx_text_is(context: Context, p_idx: int, expected: str):
    actual = context.document.paragraphs[p_idx].text
    assert actual == expected, (
        f"paragraph {p_idx} text: expected {expected!r}, got {actual!r}"
    )


@then("paragraph {p_idx:d} still has {count:d} tracked change")
@then("paragraph {p_idx:d} still has {count:d} tracked changes")
def then_paragraph_still_has_n_tracked_changes(
    context: Context, p_idx: int, count: int
):
    actual = len(context.document.paragraphs[p_idx].tracked_changes)
    assert actual == count, (
        f"expected {count} tracked changes remaining, got {actual}"
    )


@then('paragraph {p_idx:d} tracked_change[{idx:d}].author == "{expected}"')
def then_paragraph_tracked_change_author_eq(
    context: Context, p_idx: int, idx: int, expected: str
):
    actual = context.document.paragraphs[p_idx].tracked_changes[idx].author
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then('table {t_idx:d} cell ({r:d}, {c:d}) text == "{expected}"')
def then_table_cell_text_eq(
    context: Context, t_idx: int, r: int, c: int, expected: str
):
    actual = context.document.tables[t_idx].cell(r, c).text
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("table {t_idx:d} cell ({r:d}, {c:d}).is_tracked_insertion is {flag}")
def then_table_cell_is_tracked_insertion(
    context: Context, t_idx: int, r: int, c: int, flag: str
):
    expected = {"True": True, "False": False}[flag]
    actual = context.document.tables[t_idx].cell(r, c).is_tracked_insertion
    assert actual is expected, (
        f"cell({r},{c}).is_tracked_insertion expected {expected}, got {actual}"
    )


@then("the accept_all_changes return value is {count:d}")
def then_accept_all_changes_return_value_is(context: Context, count: int):
    actual = context.accept_count
    assert actual == count, (
        f"expected accept_all_changes() == {count}, got {actual}"
    )


@then("paragraph {p_idx:d} has no w:del children")
def then_paragraph_has_no_w_del_children(context: Context, p_idx: int):
    p = context.document.paragraphs[p_idx]._p
    dels = p.xpath("./w:del")
    assert dels == [], f"expected no direct-child w:del, got {len(dels)}"


@then("paragraph {p_idx:d} has no w:delText descendants")
def then_paragraph_has_no_w_delText_descendants(context: Context, p_idx: int):
    p = context.document.paragraphs[p_idx]._p
    dts = p.xpath(".//w:delText")
    assert dts == [], f"expected no w:delText descendants, got {len(dts)}"


@then("paragraph {p_idx:d} has no w:pPr/w:rPr/w:del marker")
def then_paragraph_has_no_pmark_del(context: Context, p_idx: int):
    p = context.document.paragraphs[p_idx]._p
    pmark = p.xpath("./w:pPr/w:rPr/w:del")
    assert pmark == [], (
        f"expected no w:pPr/w:rPr/w:del on paragraph {p_idx}, got {len(pmark)}"
    )


@then("paragraph {p_idx:d} still has a w:ins child")
def then_paragraph_still_has_a_w_ins_child(context: Context, p_idx: int):
    p = context.document.paragraphs[p_idx]._p
    inss = p.xpath("./w:ins")
    assert len(inss) >= 1, (
        f"expected at least one w:ins direct child, got {len(inss)}"
    )


@then('paragraph {p_idx:d} has {count:d} tracked change of type "{expected}"')
@then('paragraph {p_idx:d} has {count:d} tracked changes of type "{expected}"')
def then_paragraph_has_n_tracked_changes_of_type(
    context: Context, p_idx: int, count: int, expected: str
):
    changes = context.document.paragraphs[p_idx].tracked_changes
    assert len(changes) == count, (
        f"paragraph {p_idx}: expected {count} tracked change(s), got {len(changes)}"
    )
    types = [tc.type for tc in changes]
    assert all(t == expected for t in types), (
        f"paragraph {p_idx}: expected all {expected!r}, got {types!r}"
    )


@then("the first table has {rows:d} row with {cells:d} cell")
@then("the first table has {rows:d} row with {cells:d} cells")
@then("the first table has {rows:d} rows with {cells:d} cell")
@then("the first table has {rows:d} rows with {cells:d} cells")
def then_the_first_table_has_rows_and_cells(
    context: Context, rows: int, cells: int
):
    table = context.document.tables[0]
    actual_rows = len(table.rows)
    assert actual_rows == rows, (
        f"expected {rows} row(s), got {actual_rows}"
    )
    actual_cells = len(table.rows[0].cells)
    assert actual_cells == cells, (
        f"expected {cells} cell(s) in row 0, got {actual_cells}"
    )


@then('the first cell text is "{expected}"')
def then_the_first_cell_text_is(context: Context, expected: str):
    actual = context.document.tables[0].cell(0, 0).text
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("the remaining tracked change in paragraph {p_idx:d} is a deletion")
def then_remaining_tc_is_a_deletion(context: Context, p_idx: int):
    tcs = context.document.paragraphs[p_idx].tracked_changes
    assert len(tcs) == 1, f"expected 1 tracked change, got {len(tcs)}"
    assert tcs[0].type == "deletion", (
        f"expected deletion, got {tcs[0].type!r}"
    )


@then("the remaining tracked change in paragraph {p_idx:d} is an insertion")
def then_remaining_tc_is_an_insertion(context: Context, p_idx: int):
    tcs = context.document.paragraphs[p_idx].tracked_changes
    assert len(tcs) == 1, f"expected 1 tracked change, got {len(tcs)}"
    assert tcs[0].type == "insertion", (
        f"expected insertion, got {tcs[0].type!r}"
    )


@then("row {row_idx:d} of table {tbl_idx:d} has {count:d} cell")
@then("row {row_idx:d} of table {tbl_idx:d} has {count:d} cells")
def then_row_of_table_has_n_cells(
    context: Context, row_idx: int, tbl_idx: int, count: int
):
    table = context.document.tables[tbl_idx]
    actual = len(table.rows[row_idx].cells)
    assert actual == count, (
        f"row {row_idx} of table {tbl_idx}: expected {count} cells, got {actual}"
    )
