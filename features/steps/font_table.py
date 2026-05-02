"""Step implementations for font-table features."""

from __future__ import annotations

import io
import os
from pathlib import Path

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.font_table import FontMetadata
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from helpers import saved_docx_path, test_docx


# -- a tiny synthetic blob stands in for a real TrueType font. Word never
# -- opens embedded fonts on the read path, so python-docx treats them as
# -- opaque bytes — any blob will round-trip. --
_SAMPLE_FONT_BLOB = (
    b"OTTO\x00\x00\x00\x00" + b"\x00" * 128 + b"behave-fixture: embedded-font"
)


# given ===================================================


@given("a document having a font table")
def given_a_document_having_a_font_table(context: Context):
    context.document = Document(test_docx("fnt-table"))


# then ====================================================


@then("document.font_table is not None")
def then_document_font_table_not_none(context: Context):
    assert context.document.font_table is not None


@then("len(document.font_table) is at least {count:d}")
def then_font_table_len_at_least(context: Context, count: int):
    font_table = context.document.font_table
    assert font_table is not None
    actual = len(font_table)
    assert actual >= count, f"expected at least {count}, got {actual}"


@then('"{name}" is in document.font_table')
def then_name_in_font_table(context: Context, name: str):
    font_table = context.document.font_table
    assert font_table is not None
    assert name in font_table, f"{name!r} not in font_table"


@then('document.font_table["{name}"].name == "{expected}"')
def then_font_table_item_name(context: Context, name: str, expected: str):
    font_table = context.document.font_table
    assert font_table is not None
    actual = font_table[name].name
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then('document.font_table["{name}"].panose has length {length:d}')
def then_font_table_item_panose_length(
    context: Context, name: str, length: int
):
    font_table = context.document.font_table
    assert font_table is not None
    panose = font_table[name].panose
    assert panose is not None and len(panose) == length, (
        f"panose was {panose!r}"
    )


@then('document.font_table.get("{name}") is None')
def then_font_table_get_none(context: Context, name: str):
    font_table = context.document.font_table
    assert font_table is not None
    assert font_table.get(name) is None


@then('document.font_table["{name}"] raises KeyError')
def then_font_table_lookup_raises_keyerror(context: Context, name: str):
    font_table = context.document.font_table
    assert font_table is not None
    try:
        font_table[name]
    except KeyError:
        return
    raise AssertionError(f"expected KeyError for {name!r}")


@then("iterating document.font_table yields only FontMetadata objects")
def then_iterating_font_table_yields_fontmetadata(context: Context):
    font_table = context.document.font_table
    assert font_table is not None
    for item in font_table:
        assert isinstance(item, FontMetadata), (
            f"expected FontMetadata, got {type(item).__name__}"
        )


# -- embed-font scenarios --------------------------------------------------


@given("a document with no font table")
def given_a_document_with_no_font_table(context: Context):
    # -- a freshly-created Document actually inherits a fontTable.xml from
    # -- the bundled template, but the *scenario* only cares that we start
    # -- without any embedded-font entries, which is guaranteed. --
    context.document = Document()
    assert "BehaveFontFixture" not in (context.document.font_table or ())


@when("I call document.font_table_or_new.add_embedded_font on a sample font")
def when_add_embedded_font(context: Context):
    tmp_font_path = Path(saved_docx_path).with_name("BehaveFontFixture.ttf")
    tmp_font_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_font_path.write_bytes(_SAMPLE_FONT_BLOB)
    context._font_path = tmp_font_path

    context.document.font_table_or_new.add_embedded_font(
        tmp_font_path, family="regular"
    )


@when("I save and reopen the font-embed document")
def when_save_and_reopen_font_embed(context: Context):
    os.makedirs(os.path.dirname(saved_docx_path), exist_ok=True)
    context.document.save(saved_docx_path)
    context.document = Document(saved_docx_path)


@then("the font table has one embedded-regular entry")
def then_font_table_has_one_embedded_regular_entry(context: Context):
    font_table = context.document.font_table
    assert font_table is not None
    embedded = [m for m in font_table if m.embed_regular]
    assert len(embedded) == 1, f"expected 1 embedded-regular entry, got {len(embedded)}"
    assert embedded[0].name == "BehaveFontFixture"


@then("the embedded font binary matches the original")
def then_embedded_font_binary_matches(context: Context):
    font_table = context.document.font_table
    assert font_table is not None
    font_rels = [
        rel for rel in font_table.part.rels.values() if rel.reltype == RT.FONT
    ]
    assert len(font_rels) == 1, f"expected 1 font rel, got {len(font_rels)}"
    assert font_rels[0].target_part.blob == _SAMPLE_FONT_BLOB
