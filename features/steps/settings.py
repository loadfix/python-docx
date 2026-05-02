"""Step implementations for document settings-related features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.text import WD_VIEW
from docx.settings import CompatFlags, Settings
from docx.shared import Twips

from helpers import test_docx

# given ====================================================


@given("a document having a settings part")
def given_a_document_having_a_settings_part(context: Context):
    context.document = Document(test_docx("doc-word-default-blank"))


@given("a document having no settings part")
def given_a_document_having_no_settings_part(context: Context):
    context.document = Document(test_docx("set-no-settings-part"))


@given("a Settings object {with_or_without} odd and even page headers as settings")
def given_a_Settings_object_with_or_without_odd_and_even_hdrs(
    context: Context, with_or_without: str
):
    testfile_name = {"with": "doc-odd-even-hdrs", "without": "sct-section-props"}[with_or_without]
    context.settings = Document(test_docx(testfile_name)).settings


@given("a Settings object loaded from {testfile}")
def given_a_Settings_object_loaded_from(context: Context, testfile: str):
    context.settings = Document(test_docx(testfile)).settings


# when =====================================================


@when("I assign {bool_val} to settings.odd_and_even_pages_header_footer")
def when_I_assign_value_to_settings_odd_and_even_pages_header_footer(
    context: Context, bool_val: str
):
    context.settings.odd_and_even_pages_header_footer = eval(bool_val)


@when("I assign {value} to settings.compat_flags[{name}]")
def when_I_assign_value_to_settings_compat_flags(context: Context, value: str, name: str):
    flag_name = name.strip("\"'")
    context.settings.compat_flags[flag_name] = eval(value)


@when("I assign {value} to settings.compat_settings[{name}]")
def when_I_assign_value_to_settings_compat_settings(
    context: Context, value: str, name: str
):
    key = name.strip("\"'")
    str_val = value.strip("\"'")
    context.settings.compat_settings[key] = str_val


@when("I assign {value} to settings.view")
def when_I_assign_value_to_settings_view(context: Context, value: str):
    context.settings.view = _parse_view(value)


@when("I assign {value} to settings.zoom_percent")
def when_I_assign_value_to_settings_zoom_percent(context: Context, value: str):
    context.settings.zoom_percent = None if value == "None" else int(value)


@when("I assign {value} to settings.track_revisions")
def when_I_assign_value_to_settings_track_revisions(context: Context, value: str):
    context.settings.track_revisions = eval(value)


@when("I assign {value} to settings.default_tab_stop")
def when_I_assign_value_to_settings_default_tab_stop(context: Context, value: str):
    if value == "None":
        new_value = None
    elif value.startswith("Twips("):
        new_value = Twips(int(value[len("Twips(") : -1]))
    else:
        new_value = int(value)
    context.settings.default_tab_stop = new_value


# then =====================================================


@then("document.settings is a Settings object")
def then_document_settings_is_a_Settings_object(context: Context):
    document = context.document
    assert type(document.settings) is Settings


@then("settings.odd_and_even_pages_header_footer is {bool_val}")
def then_settings_odd_and_even_pages_header_footer_is(context: Context, bool_val: str):
    actual = context.settings.odd_and_even_pages_header_footer
    expected = eval(bool_val)
    assert actual == expected, "settings.odd_and_even_pages_header_footer is %s" % actual


@then("settings.compat_flags[{name}] is {value}")
def then_settings_compat_flags_name_is_value(context: Context, name: str, value: str):
    flag_name = name.strip("\"'")
    actual = context.settings.compat_flags[flag_name]
    expected = eval(value)
    assert actual is expected, (
        f"settings.compat_flags[{flag_name!r}] is {actual!r}, expected {expected!r}"
    )


@then("{name} is in settings.compat_flags")
def then_name_is_in_settings_compat_flags(context: Context, name: str):
    flag_name = name.strip("\"'")
    assert flag_name in context.settings.compat_flags, (
        f"expected {flag_name!r} in settings.compat_flags, got {list(context.settings.compat_flags)!r}"
    )


@then("{name} is not in settings.compat_flags")
def then_name_is_not_in_settings_compat_flags(context: Context, name: str):
    flag_name = name.strip("\"'")
    assert flag_name not in context.settings.compat_flags, (
        f"did not expect {flag_name!r} in settings.compat_flags"
    )


@then("len(settings.compat_flags) is {count}")
def then_len_settings_compat_flags_is(context: Context, count: str):
    actual = len(context.settings.compat_flags)
    expected = int(count)
    assert actual == expected, f"len(settings.compat_flags) = {actual}, expected {expected}"


@then("CompatFlags.names() contains {name}")
def then_CompatFlags_names_contains(context: Context, name: str):
    flag_name = name.strip("\"'")
    names = CompatFlags.names()
    assert flag_name in names, f"{flag_name!r} not in CompatFlags.names()"


@then("settings.compat_settings[{name}] is {value}")
def then_settings_compat_settings_name_is_value(context: Context, name: str, value: str):
    key = name.strip("\"'")
    expected = value.strip("\"'")
    actual = context.settings.compat_settings[key]
    assert actual == expected, (
        f"settings.compat_settings[{key!r}] = {actual!r}, expected {expected!r}"
    )


@then("settings.view is {value}")
def then_settings_view_is(context: Context, value: str):
    expected = _parse_view(value)
    actual = context.settings.view
    assert actual == expected, f"settings.view = {actual!r}, expected {expected!r}"


@then("settings.zoom_percent is {value}")
def then_settings_zoom_percent_is(context: Context, value: str):
    expected = None if value == "None" else int(value)
    actual = context.settings.zoom_percent
    assert actual == expected, f"settings.zoom_percent = {actual!r}, expected {expected!r}"


@then("settings.track_revisions is {value}")
def then_settings_track_revisions_is(context: Context, value: str):
    expected = eval(value)
    actual = context.settings.track_revisions
    assert actual is expected, f"settings.track_revisions = {actual!r}, expected {expected!r}"


@then("settings.compatibility_mode is {value}")
def then_settings_compatibility_mode_is(context: Context, value: str):
    expected = None if value == "None" else int(value)
    actual = context.settings.compatibility_mode
    assert actual == expected, f"settings.compatibility_mode = {actual!r}, expected {expected!r}"


@then("settings.default_tab_stop is {value}")
def then_settings_default_tab_stop_is(context: Context, value: str):
    expected = None if value == "None" else int(value)
    actual = context.settings.default_tab_stop
    assert actual == expected, f"settings.default_tab_stop = {actual!r}, expected {expected!r}"


# helpers =================================================


def _parse_view(value: str) -> WD_VIEW | None:
    if value == "None":
        return None
    if not value.startswith("WD_VIEW."):
        raise ValueError(f"unrecognized WD_VIEW literal: {value!r}")
    return getattr(WD_VIEW, value.split(".", 1)[1])
