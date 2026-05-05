"""Step implementations for field (w:fldSimple / w:fldChar) features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.fields import WD_FIELD_TYPE, Field

from helpers import test_docx

# given ====================================================


@given("a document with a simple DATE field in paragraph 2")
def given_a_document_with_a_simple_date_field(context: Context):
    context.document = Document(test_docx("fld-has-fields"))
    context.paragraph = context.document.paragraphs[2]
    context.field = context.paragraph.fields[0]


@given("a document with a complex PAGE field in paragraph 3")
def given_a_document_with_a_complex_page_field(context: Context):
    context.document = Document(test_docx("fld-has-fields"))
    context.paragraph = context.document.paragraphs[3]
    context.field = context.paragraph.fields[0]


@given('a document with a REF field pointing at the bookmark "FavouriteValue"')
def given_a_document_with_a_ref_field_pointing_at_a_bookmark(context: Context):
    context.document = Document(test_docx("fld-has-fields"))
    context.paragraph = context.document.paragraphs[5]
    context.field = context.paragraph.fields[0]


@given("the fld-has-fields document")
def given_the_fld_has_fields_document(context: Context):
    context.document = Document(test_docx("fld-has-fields"))


@given("the fld-has-fields document with a stale REF result")
def given_the_fld_has_fields_document_with_a_stale_ref_result(context: Context):
    context.document = Document(test_docx("fld-has-fields"))
    # -- overwrite the cached result so resolve_cross_references() detects drift --
    ref_field = context.document.paragraphs[5].fields[0]
    ref_field.update_result_text("stale value")


@given("the WD_FIELD_TYPE constant {name}")
def given_the_wd_field_type_constant(context: Context, name: str):
    context.constant_value = getattr(WD_FIELD_TYPE, name)


@given("a new empty document")
def given_a_new_empty_document(context: Context):
    context.document = Document()
    context.paragraph = context.document.add_paragraph()


@given("a new empty document with an unresolved PAGEREF field")
def given_a_new_empty_document_with_an_unresolved_pageref_field(context: Context):
    context.document = Document()
    context.paragraph = context.document.add_paragraph()
    context.field = context.paragraph.add_complex_field(
        "PAGEREF SomeBookmark \\h", None
    )


# when =====================================================


@when('I call paragraph.add_simple_field("{instr}", "{text}")')
def when_I_call_paragraph_add_simple_field(context: Context, instr: str, text: str):
    context.field = context.paragraph.add_simple_field(instr, text)


@when('I call paragraph.add_complex_field("{instr}", "{text}")')
def when_I_call_paragraph_add_complex_field(context: Context, instr: str, text: str):
    context.field = context.paragraph.add_complex_field(instr, text)


@when('I call field.update_result_text("{text}")')
def when_I_call_field_update_result_text(context: Context, text: str):
    context.field.update_result_text(text)


@when("I call document.resolve_cross_references()")
def when_I_call_document_resolve_cross_references(context: Context):
    context.resolve_return = context.document.resolve_cross_references()


# then =====================================================


@then("paragraph.fields has {count:d} entry")
@then("paragraph.fields has {count:d} entries")
def then_paragraph_fields_has_count_entries(context: Context, count: int):
    actual = len(context.paragraph.fields)
    assert actual == count, f"expected {count} field(s), got {actual}"


@then("the field is a complex field")
def then_the_field_is_a_complex_field(context: Context):
    assert context.field.is_complex is True, "expected a complex field"


@then("the field is not a complex field")
def then_the_field_is_not_a_complex_field(context: Context):
    assert context.field.is_complex is False, "expected a simple field"


@then('field.instruction == "{instruction}"')
def then_field_instruction_eq(context: Context, instruction: str):
    actual = context.field.instruction
    assert actual == instruction, (
        f"expected instruction {instruction!r}, got {actual!r}"
    )


@then('field.type == "{field_type}"')
def then_field_type_eq(context: Context, field_type: str):
    actual = context.field.type
    assert actual == field_type, f"expected type {field_type!r}, got {actual!r}"


@then('field.result_text == "{text}"')
def then_field_result_text_eq(context: Context, text: str):
    actual = context.field.result_text
    assert actual == text, f"expected result_text {text!r}, got {actual!r}"


@then('field.resolve(document) == "{text}"')
def then_field_resolve_eq(context: Context, text: str):
    actual = context.field.resolve(context.document)
    assert actual == text, f"expected resolve() {text!r}, got {actual!r}"


@then('the constant value == "{value}"')
def then_the_constant_value_eq(context: Context, value: str):
    actual = context.constant_value
    assert actual == value, f"expected constant value {value!r}, got {actual!r}"


@then("iterating every paragraph's fields yields {count:d} Field objects")
def then_iterating_every_paragraphs_fields_yields_n_fields(
    context: Context, count: int
):
    collected: list[Field] = []
    for paragraph in context.document.paragraphs:
        collected.extend(paragraph.fields)
    assert len(collected) == count, (
        f"expected {count} Field objects, got {len(collected)}"
    )
    assert all(isinstance(f, Field) for f in collected), (
        "expected every entry to be a Field"
    )


@then("the call returned {count:d}")
def then_the_call_returned_count(context: Context, count: int):
    actual = context.resolve_return
    assert actual == count, f"expected return value {count}, got {actual}"


@then('the REF field in paragraph 5 still reads "{text}"')
def then_the_ref_field_in_paragraph_5_still_reads(context: Context, text: str):
    ref_field = context.document.paragraphs[5].fields[0]
    actual = ref_field.result_text
    assert actual == text, f"expected result_text {text!r}, got {actual!r}"


# -- Field.evaluate / Document.evaluate_fields =============================


@given('a simple MERGEFIELD field with name "{name}"')
def given_a_simple_mergefield(context: Context, name: str):
    context.document = Document()
    p = context.document.add_paragraph()
    context.field = p.add_simple_field(f"MERGEFIELD {name}", "stale")


@given('a simple IF field that tests {{MERGEFIELD status}} equals "{value}"')
def given_a_simple_if_field_with_nested_mergefield(context: Context, value: str):
    context.document = Document()
    p = context.document.add_paragraph()
    instr = 'IF {MERGEFIELD status} = "' + value + '" "yes" "no"'
    context.field = p.add_simple_field(instr, "stale")


@given(
    'a simple HYPERLINK field with url "{url}" and no cached text'
)
def given_a_simple_hyperlink_field(context: Context, url: str):
    context.document = Document()
    p = context.document.add_paragraph()
    context.field = p.add_simple_field(f'HYPERLINK "{url}"', None)


@given('a simple formula field with expression "{expr}"')
def given_a_simple_formula_field(context: Context, expr: str):
    context.document = Document()
    p = context.document.add_paragraph()
    context.field = p.add_simple_field(expr, None)


@given('a simple PAGE field with cached text "{cached}"')
def given_a_simple_page_field_with_cached_text(context: Context, cached: str):
    context.document = Document()
    p = context.document.add_paragraph()
    context.field = p.add_simple_field("PAGE", cached)


@given("a fresh document containing a MERGEFIELD, an IF, and a formula")
def given_a_fresh_document_with_three_fields(context: Context):
    context.document = Document()
    p1 = context.document.add_paragraph()
    p1.add_simple_field("MERGEFIELD name", "stale")
    p2 = context.document.add_paragraph()
    p2.add_simple_field('IF {MERGEFIELD status} = "active" "yes" "no"', "stale")
    p3 = context.document.add_paragraph()
    p3.add_simple_field("= 2+3", "stale")


@when("I call field.evaluate with {context_json}")
def when_I_call_field_evaluate_with(context: Context, context_json: str):
    import json

    ctx = json.loads(context_json) if context_json.strip() != "{}" else {}
    context.evaluated = context.field.evaluate(ctx)


@when("I call document.evaluate_fields with {context_json}")
def when_I_call_document_evaluate_fields_with(context: Context, context_json: str):
    import json

    ctx = json.loads(context_json)
    context.resolve_return = context.document.evaluate_fields(ctx)


@then('the evaluated result is "{expected}"')
def then_the_evaluated_result_is(context: Context, expected: str):
    actual = context.evaluated
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then(
    'the three fields now read "{v1}", "{v2}", and "{v3}" respectively'
)
def then_the_three_fields_now_read(
    context: Context, v1: str, v2: str, v3: str
):
    texts = [p.fields[0].result_text for p in context.document.paragraphs]
    assert texts == [v1, v2, v3], f"expected {[v1, v2, v3]!r}, got {texts!r}"
