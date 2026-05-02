"""Step implementations for equation (OMML) behave features.

Shared by equ-identifier, equ-fraction, equ-superscript, equ-subscript, and
equ-radical feature files. Covers ``docx.equations`` builders
(``build_identifier``, ``build_fraction``, ``build_superscript``,
``build_subscript``, ``build_radical``) plus round-trips through
``Paragraph.add_equation`` and fixture-driven read-back scenarios.

Some ``Given`` phrasings (for example ``a fresh default document``) are
registered elsewhere in the steps package and are re-used here by keeping
wording identical.
"""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.equations import Equation, build_fraction, build_identifier

from helpers import test_docx

# given ====================================================


@given("a document having identifier equations")
def given_a_document_having_identifier_equations(context: Context):
    context.document = Document(test_docx("equ-identifier"))


@given("a document having fraction equations")
def given_a_document_having_fraction_equations(context: Context):
    context.document = Document(test_docx("equ-fraction"))


@given("a fresh document")
def given_a_fresh_document(context: Context):
    context.document = Document()


# when =====================================================


@when("I assign equation = the first equation in the document")
def when_I_assign_equation_first(context: Context):
    context.equation = context.document.equations[0]


@when("I assign equation = the second equation in the document")
def when_I_assign_equation_second(context: Context):
    context.equation = context.document.equations[1]


@when('I assign xml = build_identifier("{text}")')
def when_I_assign_xml_eq_build_identifier(context: Context, text: str):
    context.xml = build_identifier(text)


@when('I assign xml = build_fraction("{num}", "{den}")')
def when_I_assign_xml_eq_build_fraction(context: Context, num: str, den: str):
    context.xml = build_fraction(num, den)


@when('I append an identifier equation for "{text}" to a new paragraph')
def when_I_append_identifier_equation(context: Context, text: str):
    paragraph = context.document.add_paragraph("prefix ")
    context.paragraph = paragraph
    context.equation = paragraph.add_equation(build_identifier(text))


@when(
    'I append a fraction equation with numerator "{num}" and denominator "{den}"'
    " to a new paragraph"
)
def when_I_append_fraction_equation(context: Context, num: str, den: str):
    paragraph = context.document.add_paragraph("prefix ")
    context.paragraph = paragraph
    context.equation = paragraph.add_equation(build_fraction(num, den))


# then =====================================================


@then("document.equations is a list of {count:d} Equation objects")
def then_document_equations_is_list_of_count(context: Context, count: int):
    equations = context.document.equations
    assert isinstance(equations, list), f"expected list, got {type(equations)}"
    assert len(equations) == count, (
        f"expected {count} equations, got {len(equations)}"
    )
    for eq in equations:
        assert isinstance(eq, Equation), (
            f"expected Equation, got {type(eq)}"
        )


@then('equation.text == "{text}"')
def then_equation_text_eq_text(context: Context, text: str):
    actual = context.equation.text
    assert actual == text, f"expected equation.text {text!r}, got {actual!r}"


@then("equation.is_display_mode is False")
def then_equation_is_display_mode_false(context: Context):
    assert context.equation.is_display_mode is False, (
        f"expected is_display_mode False, got {context.equation.is_display_mode}"
    )


@then("equation.is_display_mode is True")
def then_equation_is_display_mode_true(context: Context):
    assert context.equation.is_display_mode is True, (
        f"expected is_display_mode True, got {context.equation.is_display_mode}"
    )


@then('equation.raw_xml contains "{fragment}"')
def then_equation_raw_xml_contains(context: Context, fragment: str):
    raw = context.equation.raw_xml
    assert fragment.encode("utf-8") in raw, (
        f"expected {fragment!r} in raw_xml, got {raw!r}"
    )


@then('xml starts with "{prefix}"')
def then_xml_starts_with(context: Context, prefix: str):
    assert context.xml.startswith(prefix), (
        f"expected xml to start with {prefix!r}, got {context.xml[:60]!r}"
    )


@then('xml ends with "{suffix}"')
def then_xml_ends_with(context: Context, suffix: str):
    assert context.xml.endswith(suffix), (
        f"expected xml to end with {suffix!r}, got ...{context.xml[-60:]!r}"
    )


@then('xml contains "{fragment}"')
def then_xml_contains(context: Context, fragment: str):
    assert fragment in context.xml, (
        f"expected {fragment!r} in xml, got {context.xml!r}"
    )
from docx.equations import Equation, build_subscript, build_superscript
from docx.equations import Equation, build_radical

from helpers import test_docx


# given ===================================================


@given("a document having two superscript equations")
def given_document_having_two_superscript_equations(context: Context):
    context.document = Document(test_docx("equ-superscript"))


@given("a document having chained subscript equations")
def given_document_having_chained_subscript_equations(context: Context):
    context.document = Document(test_docx("equ-subscript"))
@given("a document having a radical-equation fixture")
def given_document_with_radical_fixture(context: Context):
    context.document = Document(test_docx("equ-radical"))


# when ====================================================


@when('I call build_superscript("{base}", "{exponent}")')
def when_call_build_superscript(context: Context, base: str, exponent: str):
    xml_str = build_superscript(base, exponent)
    context.built_xml = xml_str
    context.built_equation = Equation.from_omml_xml(xml_str)


@when('I call build_subscript("{base}", "{subscript}")')
def when_call_build_subscript(context: Context, base: str, subscript: str):
    xml_str = build_subscript(base, subscript)
    context.built_xml = xml_str
    context.built_equation = Equation.from_omml_xml(xml_str)


@when('I add a superscript equation "{base}" "{exponent}" to a new paragraph')
def when_add_superscript_to_new_paragraph(
    context: Context, base: str, exponent: str
):
    context.paragraph = context.document.add_paragraph()
    context.paragraph.add_equation(build_superscript(base, exponent))


@when('I add a subscript equation "{base}" "{sub}" to a new paragraph')
def when_add_subscript_to_new_paragraph(
    context: Context, base: str, sub: str
):
    context.paragraph = context.document.add_paragraph()
    context.paragraph.add_equation(build_subscript(base, sub))


@when('I add a subscript equation "{base}" "{sub}" to the same paragraph')
def when_add_subscript_to_same_paragraph(
    context: Context, base: str, sub: str
):
    context.paragraph.add_equation(build_subscript(base, sub))
@when('I build a radical with expr "{expr}" and no degree')
def when_build_radical_no_degree(context: Context, expr: str):
    context.omml_xml = build_radical(expr)
    context.equation = Equation.from_omml_xml(context.omml_xml)


@when('I build a radical with expr "{expr}" and degree "{degree}"')
def when_build_radical_with_degree(context: Context, expr: str, degree: str):
    context.omml_xml = build_radical(expr, degree)
    context.equation = Equation.from_omml_xml(context.omml_xml)


@when(
    'I append a radical equation with expr "{expr}" and degree "{degree}" '
    "to a new paragraph"
)
def when_append_radical_equation(context: Context, expr: str, degree: str):
    paragraph = context.document.add_paragraph()
    omml = build_radical(expr, degree)
    context.paragraph = paragraph
    context.equation = paragraph.add_equation(omml)


# then ====================================================


@then('the built equation text is "{text}"')
def then_built_equation_text(context: Context, text: str):
    actual = context.built_equation.text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('the built equation raw_xml contains "{fragment}"')
def then_built_equation_raw_xml_contains(context: Context, fragment: str):
    raw = context.built_equation.raw_xml.decode("utf-8")
    assert fragment in raw, f"{fragment!r} not found in {raw!r}"


@then("the built equation is_display_mode is {value}")
def then_built_equation_is_display_mode(context: Context, value: str):
    expected = {"True": True, "False": False}[value]
    actual = context.built_equation.is_display_mode
    assert actual is expected, f"expected {expected!r}, got {actual!r}"
@then('the OMML contains an "{tag}" element')
def then_omml_contains_element(context: Context, tag: str):
    raw = context.equation.raw_xml
    needle = ("<%s" % tag).encode("utf-8")
    assert needle in raw, f"{needle!r} not found in {raw!r}"


@then('the OMML contains an empty "{tag}" element')
def then_omml_contains_empty_element(context: Context, tag: str):
    raw = context.equation.raw_xml
    needle = ("<%s/>" % tag).encode("utf-8")
    assert needle in raw, f"{needle!r} not found in {raw!r}"


@then('the OMML contains a populated "{tag}" element')
def then_omml_contains_populated_element(context: Context, tag: str):
    raw = context.equation.raw_xml
    # -- populated means opening tag followed by at least one child, not self-closing --
    empty = ("<%s/>" % tag).encode("utf-8")
    opened = ("<%s>" % tag).encode("utf-8")
    assert opened in raw, f"{opened!r} not found in {raw!r}"
    assert empty not in raw, f"unexpected empty {tag} in {raw!r}"


@then('the equation text is "{text}"')
def then_equation_text_is(context: Context, text: str):
    actual = context.equation.text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then("the document has {count:d} radical equations")
def then_document_has_radical_equations(context: Context, count: int):
    radical_eqs = [
        eq for eq in context.document.equations if b"<m:rad>" in eq.raw_xml
    ]
    assert len(radical_eqs) == count, (
        f"expected {count} radical equations, got {len(radical_eqs)}"
    )
    context.radical_equations = radical_eqs


@then('the third radical has a nested "{tag}" descendant')
def then_third_radical_has_nested(context: Context, tag: str):
    radical = context.radical_equations[2]
    needle = ("<%s>" % tag).encode("utf-8")
    count = radical.raw_xml.count(needle)
    assert count >= 2, (
        f"expected at least 2 {tag!r} elements (outer + nested), got {count}"
    )


@then("the paragraph has {count:d} equation")
@then("the paragraph has {count:d} equations")
def then_paragraph_has_count_equations(context: Context, count: int):
    equations = context.paragraph.equations
    assert len(equations) == count, (
        f"expected {count} equations on paragraph, got {len(equations)}"
    )


@then("xml contains a bar-type fraction property")
def then_xml_contains_bar_type(context: Context):
    fragment = '<m:type m:val="bar"/>'
    assert fragment in context.xml, (
        f"expected {fragment!r} in xml, got {context.xml!r}"
    )


@then('the appended equation.text == "{text}"')
def then_appended_equation_text(context: Context, text: str):
    actual = context.equation.text
    assert actual == text, (
        f"expected appended equation.text {text!r}, got {actual!r}"
    )


@then("the paragraph has {count:d} equation")
@then("the paragraph has {count:d} equations")
def then_paragraph_has_n_equations(context: Context, count: int):
    actual = len(context.paragraph.equations)
    assert actual == count, f"expected {count} equations, got {actual}"


@then('the paragraph first equation text is "{text}"')
def then_paragraph_first_equation_text(context: Context, text: str):
    actual = context.paragraph.equations[0].text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('the paragraph second equation text is "{text}"')
def then_paragraph_second_equation_text(context: Context, text: str):
    actual = context.paragraph.equations[1].text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('the paragraph first equation raw_xml contains "{fragment}"')
def then_paragraph_first_equation_raw_xml_contains(
    context: Context, fragment: str
):
    raw = context.paragraph.equations[0].raw_xml.decode("utf-8")
    assert fragment in raw, f"{fragment!r} not found in {raw!r}"


@then('every paragraph equation raw_xml contains "{fragment}"')
def then_every_paragraph_equation_raw_xml_contains(
    context: Context, fragment: str
):
    for eq in context.paragraph.equations:
        raw = eq.raw_xml.decode("utf-8")
        assert fragment in raw, f"{fragment!r} not found in {raw!r}"


def _doc_equations(context: Context) -> list:
    return [eq for p in context.document.paragraphs for eq in p.equations]


@then("the document has {count:d} superscript equations")
@then("the document has {count:d} subscript equations")
def then_document_has_n_equations(context: Context, count: int):
    equations = _doc_equations(context)
    assert len(equations) == count, (
        f"expected {count} equations, got {len(equations)}"
    )


@then('the first superscript equation text is "{text}"')
@then('the first subscript equation text is "{text}"')
def then_first_document_equation_text(context: Context, text: str):
    actual = _doc_equations(context)[0].text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('the second superscript equation text is "{text}"')
@then('the second subscript equation text is "{text}"')
def then_second_document_equation_text(context: Context, text: str):
    actual = _doc_equations(context)[1].text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('the third subscript equation text is "{text}"')
def then_third_document_equation_text(context: Context, text: str):
    actual = _doc_equations(context)[2].text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('every superscript equation raw_xml contains "{fragment}"')
@then('every subscript equation raw_xml contains "{fragment}"')
def then_every_document_equation_raw_xml_contains(
    context: Context, fragment: str
):
    for eq in _doc_equations(context):
        raw = eq.raw_xml.decode("utf-8")
        assert fragment in raw, f"{fragment!r} not found in {raw!r}"
@then('the appended equation text is "{text}"')
def then_appended_equation_text(context: Context, text: str):
    actual = context.equation.text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then("the appended equation is not display mode")
def then_appended_equation_not_display_mode(context: Context):
    assert context.equation.is_display_mode is False, (
        "expected inline (non-display) equation"
    )
