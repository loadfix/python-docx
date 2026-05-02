"""Step implementations for structured-document-tag (content control) features."""

from __future__ import annotations

import io

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.content_controls import ContentControl, ContentControlType, DataBinding

from helpers import test_docx

# -- references to fixture constants; keep in sync with _gen_sdt_data_bound.py --
_BOUND_XPATH = "/ns0:order[1]/ns0:customer[1]"
_BOUND_PREFIX_MAPPINGS = "xmlns:ns0='http://example.com/orders'"
_BOUND_STORE_ITEM_ID = "{11111111-2222-3333-4444-555555555555}"
_BOUND_SCHEMA_URI = "http://example.com/orders"


def _type_from_name(name: str) -> ContentControlType:
    """Resolve a :class:`ContentControlType` from its member name."""
    return ContentControlType[name.strip()]


# given ====================================================


@given("a default document")
def given_a_default_document(context: Context):
    context.document = Document()


@given('a default document with a paragraph "{text}"')
def given_a_default_document_with_paragraph(context: Context, text: str):
    context.document = Document()
    context.paragraph = context.document.add_paragraph(text)


@given("a document having {count:d} block-level content controls")
def given_a_document_having_n_block_level_content_controls(context: Context, count: int):
    context.document = Document(test_docx("sdt-has-controls"))
    assert len(context.document.content_controls) == count, (
        f"fixture exposes {len(context.document.content_controls)} block "
        f"controls, expected {count}"
    )


@given("a document containing a data-bound content control")
def given_a_document_containing_a_data_bound_content_control(context: Context):
    context.document = Document(test_docx("sdt-data-bound"))
    # -- locate the single bound control for use by later steps --
    controls = context.document.content_controls
    context.cc = next(cc for cc in controls if cc.data_binding is not None)


# when =====================================================


@when(
    'I assign cc = document.add_content_control({type_name}, tag="{tag}", title="{title}")'
)
def when_add_content_control_with_tag_and_title(
    context: Context, type_name: str, tag: str, title: str
):
    cc_type = _type_from_name(type_name)
    context.cc = context.document.add_content_control(cc_type, tag=tag, title=title)


@when('I assign cc = document.add_content_control({type_name}, tag="{tag}")')
def when_add_content_control_with_tag(context: Context, type_name: str, tag: str):
    cc_type = _type_from_name(type_name)
    context.cc = context.document.add_content_control(cc_type, tag=tag)


@when("I assign cc = document.add_content_control({type_name})")
def when_add_content_control_type_only(context: Context, type_name: str):
    cc_type = _type_from_name(type_name)
    context.cc = context.document.add_content_control(cc_type)


@when(
    'I assign cc = paragraph.add_content_control({type_name}, tag="{tag}")'
)
def when_add_inline_content_control(context: Context, type_name: str, tag: str):
    cc_type = _type_from_name(type_name)
    context.cc = context.paragraph.add_content_control(cc_type, tag=tag)


@when('I assign cc.text = "{value}"')
def when_assign_cc_text(context: Context, value: str):
    context.cc.text = value


@when("I assign cc.checked = True")
def when_assign_cc_checked_true(context: Context):
    context.cc.checked = True


@when("I read the {type_name} control")
def when_read_typed_control(context: Context, type_name: str):
    cc_type = _type_from_name(type_name)
    context.cc = next(
        cc for cc in context.document.content_controls if cc.type is cc_type
    )


@when("I call cc.set_data_binding(xpath, prefix_mappings, store_item_id)")
def when_call_set_data_binding_with_fixture_values(context: Context):
    context.xpath = _BOUND_XPATH
    context.prefix_mappings = _BOUND_PREFIX_MAPPINGS
    context.store_item_id = _BOUND_STORE_ITEM_ID
    context.cc.set_data_binding(
        xpath=context.xpath,
        prefix_mappings=context.prefix_mappings,
        store_item_id=context.store_item_id,
    )


@when("I call cc.set_data_binding with new values")
def when_call_set_data_binding_with_new_values(context: Context):
    context.cc.set_data_binding(
        xpath="/a",
        prefix_mappings="",
        store_item_id="{22222222-3333-4444-5555-666666666666}",
    )


@when("I call cc.remove_data_binding()")
def when_call_remove_data_binding(context: Context):
    context.cc.remove_data_binding()


@when("I save and reload the document")
def when_save_and_reload(context: Context):
    buf = io.BytesIO()
    context.document.save(buf)
    buf.seek(0)
    context.document = Document(buf)
    context.cc = next(
        cc for cc in context.document.content_controls if cc.data_binding is not None
    )


# then =====================================================


@then("cc is a ContentControl object")
def then_cc_is_a_content_control_object(context: Context):
    assert type(context.cc) is ContentControl, (
        f"expected ContentControl, got {type(context.cc).__name__}"
    )


@then("cc.type is ContentControlType.{type_name}")
def then_cc_type_is(context: Context, type_name: str):
    expected = _type_from_name(type_name)
    actual = context.cc.type
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


@then('cc.tag == "{tag}"')
def then_cc_tag_eq(context: Context, tag: str):
    actual = context.cc.tag
    assert actual == tag, f"expected tag {tag!r}, got {actual!r}"


@then('cc.title == "{title}"')
def then_cc_title_eq(context: Context, title: str):
    actual = context.cc.title
    assert actual == title, f"expected title {title!r}, got {actual!r}"


@then("cc.sdt_id is a positive integer")
def then_cc_sdt_id_is_positive_integer(context: Context):
    sdt_id = context.cc.sdt_id
    assert isinstance(sdt_id, int), f"expected int, got {type(sdt_id).__name__}"
    assert sdt_id > 0, f"expected positive id, got {sdt_id}"


@then('cc.text == "{text}"')
def then_cc_text_eq(context: Context, text: str):
    actual = context.cc.text
    assert actual == text, f"expected text {text!r}, got {actual!r}"


@then("cc.text is a single line")
def then_cc_text_is_single_line(context: Context):
    assert "\n" not in context.cc.text, (
        f"expected single-line text, got {context.cc.text!r}"
    )


@then("cc.checked is True")
def then_cc_checked_is_true(context: Context):
    assert context.cc.checked is True, f"expected True, got {context.cc.checked!r}"


@then("cc.data_binding is None")
def then_cc_data_binding_is_none(context: Context):
    assert context.cc.data_binding is None, (
        f"expected None, got {context.cc.data_binding!r}"
    )


@then("cc.data_binding is a DataBinding object")
def then_cc_data_binding_is_a_databinding(context: Context):
    assert type(context.cc.data_binding) is DataBinding, (
        f"expected DataBinding, got {type(context.cc.data_binding).__name__}"
    )


@then("cc.data_binding.xpath is the xpath I supplied")
def then_cc_data_binding_xpath_is_supplied(context: Context):
    actual = context.cc.data_binding.xpath
    assert actual == context.xpath, f"expected {context.xpath!r}, got {actual!r}"


@then("cc.data_binding.prefix_mappings is the prefix_mappings I supplied")
def then_cc_data_binding_prefix_mappings_is_supplied(context: Context):
    actual = context.cc.data_binding.prefix_mappings
    assert actual == context.prefix_mappings, (
        f"expected {context.prefix_mappings!r}, got {actual!r}"
    )


@then("cc.data_binding.store_item_id is the store_item_id I supplied")
def then_cc_data_binding_store_item_id_is_supplied(context: Context):
    actual = context.cc.data_binding.store_item_id
    assert actual == context.store_item_id, (
        f"expected {context.store_item_id!r}, got {actual!r}"
    )


@then('cc.data_binding.xpath == "{xpath}"')
def then_cc_data_binding_xpath_eq(context: Context, xpath: str):
    actual = context.cc.data_binding.xpath
    assert actual == xpath, f"expected {xpath!r}, got {actual!r}"


@then('cc.data_binding.prefix_mappings == "{prefix}"')
def then_cc_data_binding_prefix_mappings_eq(context: Context, prefix: str):
    actual = context.cc.data_binding.prefix_mappings
    assert actual == prefix, f"expected {prefix!r}, got {actual!r}"


@then('cc.data_binding.prefix_mappings == ""')
def then_cc_data_binding_prefix_mappings_empty(context: Context):
    actual = context.cc.data_binding.prefix_mappings
    assert actual == "", f"expected empty string, got {actual!r}"


@then('cc.data_binding.store_item_id == "{store_item_id}"')
def then_cc_data_binding_store_item_id_eq(context: Context, store_item_id: str):
    actual = context.cc.data_binding.store_item_id
    assert actual == store_item_id, f"expected {store_item_id!r}, got {actual!r}"


@then("document.content_controls yields {count:d} ContentControl objects")
def then_document_content_controls_yields_n(context: Context, count: int):
    controls = context.document.content_controls
    assert len(controls) == count, f"expected {count}, got {len(controls)}"
    for cc in controls:
        assert type(cc) is ContentControl, (
            f"expected ContentControl, got {type(cc).__name__}"
        )


@then('the control tags are "{tags}"')
def then_control_tags_are(context: Context, tags: str):
    expected = [t.strip() for t in tags.split(",")]
    actual = [cc.tag for cc in context.document.content_controls]
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then('the control types are "{types}"')
def then_control_types_are(context: Context, types: str):
    expected = [_type_from_name(t) for t in types.split(",")]
    actual = [cc.type for cc in context.document.content_controls]
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("paragraph.content_controls[{idx:d}] == cc")
def then_paragraph_content_controls_idx_eq_cc(context: Context, idx: int):
    actual = context.paragraph.content_controls[idx].element
    expected = context.cc.element
    assert actual is expected, "paragraph content control does not match cc"


@then("the final paragraph has {count:d} inline content control")
def then_final_paragraph_has_inline_content_control(context: Context, count: int):
    last_para = context.document.paragraphs[-1]
    controls = last_para.content_controls
    assert len(controls) == count, f"expected {count}, got {len(controls)}"
    context.cc = controls[0]


@then('the inline control\'s tag == "{tag}"')
def then_inline_controls_tag_eq(context: Context, tag: str):
    actual = context.cc.tag
    assert actual == tag, f"expected {tag!r}, got {actual!r}"


@then('the inline control\'s text == "{text}"')
def then_inline_controls_text_eq(context: Context, text: str):
    actual = context.cc.text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('the bound control\'s data_binding.xpath == "{xpath}"')
def then_bound_data_binding_xpath_eq(context: Context, xpath: str):
    actual = context.cc.data_binding.xpath
    assert actual == xpath, f"expected {xpath!r}, got {actual!r}"


@then('the bound control\'s data_binding.prefix_mappings == "{prefix}"')
def then_bound_data_binding_prefix_mappings_eq(context: Context, prefix: str):
    actual = context.cc.data_binding.prefix_mappings
    assert actual == prefix, f"expected {prefix!r}, got {actual!r}"


@then('the bound control\'s data_binding.store_item_id == "{store_item_id}"')
def then_bound_data_binding_store_item_id_eq(context: Context, store_item_id: str):
    actual = context.cc.data_binding.store_item_id
    assert actual == store_item_id, f"expected {store_item_id!r}, got {actual!r}"


@then("document has a custom XML part whose item_id matches the binding's store_item_id")
def then_document_has_matching_custom_xml_part(context: Context):
    store_item_id = context.cc.data_binding.store_item_id
    matches = [p for p in context.document.custom_xml_parts if p.item_id == store_item_id]
    assert len(matches) == 1, (
        f"expected exactly one custom XML part with item_id {store_item_id!r}, got {len(matches)}"
    )
    context.custom_xml_part = matches[0]


@then('that custom XML part\'s schema_refs include "{uri}"')
def then_custom_xml_part_schema_refs_include(context: Context, uri: str):
    refs = context.custom_xml_part.schema_refs
    assert uri in refs, f"expected {uri!r} in schema_refs, got {refs!r}"


@then("that custom XML part's root_element is not None")
def then_custom_xml_part_root_element_is_not_none(context: Context):
    assert context.custom_xml_part.root_element is not None, (
        "expected a parsed root element, got None"
    )
