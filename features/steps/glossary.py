"""Step implementations for glossary / building-block features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.text import WD_BUILDING_BLOCK_GALLERY
from docx.glossary import BuildingBlock, Glossary

from helpers import test_docx


# -- a small lookup used for parameterised filter scenarios ------------------
_GALLERY_BY_NAME: dict[str, WD_BUILDING_BLOCK_GALLERY] = {
    member.name: member for member in WD_BUILDING_BLOCK_GALLERY
}


def _parse_names(names: str) -> list[str]:
    """Return a list of names from a comma-separated `names` string.

    An empty or whitespace-only string yields an empty list.
    """
    names = names.strip()
    if not names:
        return []
    return [n.strip() for n in names.split(",") if n.strip()]


# given ======================================================================


@given("a document having a glossary part")
def given_a_document_having_a_glossary_part(context: Context):
    context.document = Document(test_docx("glo-has-glossary"))


@given("a document having no glossary part")
def given_a_document_having_no_glossary_part(context: Context):
    context.document = Document(test_docx("doc-default"))


@given("a Glossary object with {count} building blocks")
def given_a_glossary_object_with_count_building_blocks(
    context: Context, count: str
):
    document = Document(test_docx("glo-has-glossary"))
    glossary = document.glossary
    assert glossary is not None, "fixture is missing its glossary part"
    assert len(glossary) == int(count), (
        f"expected {count} building blocks, got {len(glossary)}"
    )
    context.glossary = glossary


@given('a BuildingBlock object named "{name}"')
def given_a_building_block_object_named(context: Context, name: str):
    document = Document(test_docx("glo-has-glossary"))
    glossary = document.glossary
    assert glossary is not None, "fixture is missing its glossary part"
    context.building_block = glossary[name]


# when =======================================================================


@when('I call glossary["{name}"]')
def when_I_call_glossary_getitem(context: Context, name: str):
    context.building_block = context.glossary[name]


@when("I call glossary.by_category({args})")
def when_I_call_by_category(context: Context, args: str):
    """Dispatch ``glossary.by_category(<args>)`` calls in the feature files.

    `args` is the raw argument list (everything between the parens). It
    supports the three shapes exercised by the features:

    * ``gallery=WD_BUILDING_BLOCK_GALLERY.<MEMBER>``
    * ``gallery="<xml-string>"``
    * ``category_name="<name>"``
    * ``gallery=WD_BUILDING_BLOCK_GALLERY.<MEMBER>, category_name="<name>"``
    """
    kwargs: dict[str, object] = {}
    # -- split the top-level args on comma; the inputs don't contain nested
    # -- commas so a plain split is sufficient here.
    for token in (t.strip() for t in args.split(",")):
        key, _, raw_value = token.partition("=")
        key = key.strip()
        raw_value = raw_value.strip()
        if key == "gallery":
            if raw_value.startswith("WD_BUILDING_BLOCK_GALLERY."):
                member = raw_value.split(".", 1)[1]
                kwargs["gallery"] = _GALLERY_BY_NAME[member]
            else:
                kwargs["gallery"] = raw_value.strip('"')
        elif key == "category_name":
            kwargs["category_name"] = raw_value.strip('"')
        else:
            raise AssertionError(f"unexpected keyword argument: {key!r}")
    context.result = context.glossary.by_category(**kwargs)  # type: ignore[arg-type]


# then =======================================================================


@then("document.glossary is a Glossary object")
def then_document_glossary_is_a_glossary_object(context: Context):
    glossary = context.document.glossary
    assert isinstance(glossary, Glossary), f"got {type(glossary).__name__}"


@then("document.glossary is None")
def then_document_glossary_is_None(context: Context):
    assert context.document.glossary is None, (
        f"expected None, got {context.document.glossary!r}"
    )


@then("len(glossary) == {count:d}")
def then_len_glossary_eq_count(context: Context, count: int):
    actual = len(context.glossary)
    assert actual == count, f"expected len(glossary) == {count}, got {actual}"


@then("iterating glossary yields {count:d} BuildingBlock objects")
def then_iterating_glossary_yields_count_building_blocks(
    context: Context, count: int
):
    blocks = list(context.glossary)
    assert len(blocks) == count, f"expected {count} blocks, got {len(blocks)}"
    assert all(isinstance(b, BuildingBlock) for b in blocks), (
        "expected every item to be a BuildingBlock"
    )


@then("glossary.building_blocks names are {names}")
def then_building_blocks_names_are(context: Context, names: str):
    expected = _parse_names(names)
    actual = [b.name for b in context.glossary.building_blocks]
    assert actual == expected, f"expected {expected}, got {actual}"


@then('the result is a BuildingBlock named "{name}"')
def then_the_result_is_a_building_block_named(context: Context, name: str):
    block = context.building_block
    assert isinstance(block, BuildingBlock), f"got {type(block).__name__}"
    assert block.name == name, f"expected name {name!r}, got {block.name!r}"


@then('glossary["{name}"] raises KeyError')
def then_glossary_getitem_raises_key_error(context: Context, name: str):
    try:
        _ = context.glossary[name]
    except KeyError:
        return
    raise AssertionError(f"expected KeyError for {name!r}")


@then('building_block.name == "{name}"')
def then_building_block_name_eq(context: Context, name: str):
    actual = context.building_block.name
    assert actual == name, f"expected name {name!r}, got {actual!r}"


@then('building_block.description == "{description}"')
def then_building_block_description_eq(context: Context, description: str):
    actual = context.building_block.description
    assert actual == description, (
        f"expected description {description!r}, got {actual!r}"
    )


@then("building_block.guid is a non-empty string")
def then_building_block_guid_is_a_non_empty_string(context: Context):
    guid = context.building_block.guid
    assert isinstance(guid, str) and guid, f"expected non-empty string, got {guid!r}"


@then('building_block.category.category_name == "{category_name}"')
def then_building_block_category_name_eq(context: Context, category_name: str):
    actual = context.building_block.category.category_name
    assert actual == category_name, (
        f"expected category_name {category_name!r}, got {actual!r}"
    )


@then('building_block.category.gallery == "{gallery}"')
def then_building_block_category_gallery_eq(context: Context, gallery: str):
    actual = context.building_block.category.gallery
    assert actual == gallery, f"expected gallery {gallery!r}, got {actual!r}"


@then("building_block.category.gallery_value is WD_BUILDING_BLOCK_GALLERY.{member}")
def then_building_block_gallery_value_is_member(context: Context, member: str):
    expected = _GALLERY_BY_NAME[member]
    actual = context.building_block.category.gallery_value
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


@then('building_block.paragraphs[0].text == "{text}"')
def then_building_block_paragraphs_0_text_eq(context: Context, text: str):
    paragraphs = context.building_block.paragraphs
    assert paragraphs, "building_block has no paragraphs"
    actual = paragraphs[0].text
    assert actual == text, f"expected text {text!r}, got {actual!r}"


@then("len(building_block.tables) == {count:d}")
def then_len_building_block_tables_eq(context: Context, count: int):
    actual = len(context.building_block.tables)
    assert actual == count, f"expected {count} tables, got {actual}"


@then("building_block.paragraphs == []")
def then_building_block_paragraphs_is_empty(context: Context):
    actual = context.building_block.paragraphs
    assert actual == [], f"expected empty list, got {actual!r}"


@then("building_block.tables == []")
def then_building_block_tables_is_empty(context: Context):
    actual = context.building_block.tables
    assert actual == [], f"expected empty list, got {actual!r}"


@then("building_block.category.category_name is None")
def then_building_block_category_name_is_None(context: Context):
    actual = context.building_block.category.category_name
    assert actual is None, f"expected None, got {actual!r}"


@then("building_block.category.gallery is None")
def then_building_block_category_gallery_is_None(context: Context):
    actual = context.building_block.category.gallery
    assert actual is None, f"expected None, got {actual!r}"


@then("the result names are {names}")
def then_the_result_names_are(context: Context, names: str):
    expected = _parse_names(names)
    actual = [b.name for b in context.result]
    assert actual == expected, f"expected {expected}, got {actual}"


@then("the result is an empty list")
def then_the_result_is_an_empty_list(context: Context):
    assert context.result == [], f"expected empty list, got {context.result!r}"


@then(
    "glossary.categories has {count:d} entries with keys "
    "(quickParts, General), (coverPg, Built-In), (hdrs, Built-In)"
)
def then_glossary_categories_has_expected_keys(context: Context, count: int):
    cats = context.glossary.categories
    assert len(cats) == count, f"expected {count} categories, got {len(cats)}"
    keys = [(c.gallery, c.category_name) for c in cats]
    expected = [
        ("quickParts", "General"),
        ("coverPg", "Built-In"),
        ("hdrs", "Built-In"),
    ]
    assert keys == expected, f"expected {expected}, got {keys}"


@then('glossary.galleries == ["quickParts", "coverPg", "hdrs"]')
def then_glossary_galleries_eq(context: Context):
    actual = context.glossary.galleries
    expected = ["quickParts", "coverPg", "hdrs"]
    assert actual == expected, f"expected {expected}, got {actual}"
