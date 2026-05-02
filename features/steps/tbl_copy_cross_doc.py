"""Step implementations for tbl-copy-cross-doc.feature."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Inches

from helpers import test_file


# given ===================================================


@given(
    "a source document containing a 2x2 table with a styled header and an embedded PNG"
)
def given_source_document_with_styled_table_and_image(context: Context):
    source = Document()
    table = source.add_table(rows=2, cols=2)
    # -- try to apply a built-in table style (headers use a style reference) --
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        # -- default template may not include Table Grid; scenario still
        # -- asserts cell text & image rewiring, not style presence. --
        pass
    # -- header row --
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    # -- body with embedded image in bottom-right cell --
    table.cell(1, 0).text = "figure"
    picture_cell = table.cell(1, 1)
    picture_cell.paragraphs[0].add_run().add_picture(
        test_file("monty-truth.png"), width=Inches(1)
    )
    context.source_document = source
    context.source_table = table


@given("an empty destination document")
def given_empty_destination_document(context: Context):
    context.destination_document = Document()


# when ====================================================


@when("I call add_table_copy on the destination with the source table")
def when_call_add_table_copy(context: Context):
    context.copied_table = context.destination_document.add_table_copy(
        context.source_table
    )


# then ====================================================


@then("the destination contains one table with the copied cell text")
def then_destination_contains_copied_table(context: Context):
    tables = context.destination_document.tables
    assert len(tables) == 1, f"expected 1 table, got {len(tables)}"
    table = tables[0]
    assert table.cell(0, 0).text == "Name"
    assert table.cell(0, 1).text == "Value"
    assert table.cell(1, 0).text == "figure"


@then("the destination contains at least one image part")
def then_destination_contains_image_part(context: Context):
    rels = [
        r
        for r in context.destination_document.part.rels.values()
        if r.reltype == RT.IMAGE
    ]
    assert len(rels) >= 1, "expected at least one IMAGE relationship on dest"


@then("the copied table's embedded image reference resolves in the destination")
def then_copied_image_reference_resolves(context: Context):
    tbl = context.copied_table._tbl  # pyright: ignore[reportPrivateUsage]
    blips = tbl.xpath(".//a:blip[@r:embed]")
    assert blips, "no a:blip@r:embed found in copied table"
    dest_part = context.destination_document.part
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        assert rid in dest_part.related_parts, (
            f"rId {rid!r} from copied blip is not in destination rels"
        )
        target = dest_part.related_parts[rid]
        assert target.content_type.startswith("image/"), (
            f"rId {rid} resolves to non-image part: {target.content_type}"
        )
