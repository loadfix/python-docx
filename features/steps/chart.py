"""Step implementations for chart-read and chart-create behave features."""

from __future__ import annotations

import ast
import io

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.chart import Chart, WD_CHART_TYPE
from docx.parts.chart import ChartPart

from helpers import test_docx

# given ====================================================


@given("a document having three embedded charts")
def given_a_document_having_three_embedded_charts(context: Context):
    context.document = Document(test_docx("chart-has-charts"))


@given("a document having no charts")
def given_a_document_having_no_charts(context: Context):
    context.document = Document(test_docx("doc-default"))


@given("a document having a chart with no title")
def given_a_document_having_a_chart_with_no_title(context: Context):
    # -- build a document in-memory whose single chart has no c:title --
    document = Document()
    document.add_chart(
        WD_CHART_TYPE.PIE,
        ["A", "B", "C"],
        {"Slices": [1.0, 2.0, 3.0]},
    )
    context.document = document
    context.chart = document.charts[0]


# then =====================================================


@then("document.charts is a list of three Chart objects")
def then_document_charts_is_a_list_of_three_chart_objects(context: Context):
    charts = context.document.charts
    assert isinstance(charts, list), f"expected list, got {type(charts)}"
    assert len(charts) == 3, f"expected 3 charts, got {len(charts)}"
    for idx, chart in enumerate(charts):
        assert isinstance(chart, Chart), (
            f"expected Chart at index {idx}, got {type(chart)}"
        )


@then("iterating document.charts yields Chart objects in document order")
def then_iterating_document_charts_yields_charts_in_order(context: Context):
    charts_iter = iter(context.document.charts)
    expected_types = [
        WD_CHART_TYPE.COLUMN,
        WD_CHART_TYPE.BAR,
        WD_CHART_TYPE.LINE,
    ]
    for expected_type in expected_types:
        chart = next(charts_iter)
        assert isinstance(chart, Chart), f"expected Chart, got {type(chart)}"
        assert chart.chart_type == expected_type, (
            f"expected {expected_type}, got {chart.chart_type}"
        )


@then("charts[{idx}].chart_type == WD_CHART_TYPE.{member}")
def then_charts_idx_chart_type_eq_member(context: Context, idx: str, member: str):
    chart = context.document.charts[int(idx)]
    expected = WD_CHART_TYPE[member]
    assert chart.chart_type == expected, (
        f"expected {expected}, got {chart.chart_type}"
    )


@then('charts[{idx}].title == "{title}"')
def then_charts_idx_title_eq_title(context: Context, idx: str, title: str):
    chart = context.document.charts[int(idx)]
    assert chart.title == title, f"expected title {title!r}, got {chart.title!r}"


@then("chart.title is None")
def then_chart_title_is_None(context: Context):
    assert context.chart.title is None, (
        f"expected chart.title to be None, got {context.chart.title!r}"
    )


@then("[s.name for s in charts[{idx}].series] == {names_expr}")
def then_series_names_eq(context: Context, idx: str, names_expr: str):
    chart = context.document.charts[int(idx)]
    expected = ast.literal_eval(names_expr)
    actual = [s.name for s in chart.series]
    assert actual == expected, f"expected series names {expected}, got {actual}"


@then("charts[{idx}].series[{ser_idx}].values == {values_expr}")
def then_series_values_eq(
    context: Context, idx: str, ser_idx: str, values_expr: str
):
    chart = context.document.charts[int(idx)]
    series = chart.series[int(ser_idx)]
    expected = ast.literal_eval(values_expr)
    assert series.values == expected, (
        f"expected values {expected}, got {series.values}"
    )


@then("charts[{idx}].series[{ser_idx}].categories == {categories_expr}")
def then_series_categories_eq(
    context: Context, idx: str, ser_idx: str, categories_expr: str
):
    chart = context.document.charts[int(idx)]
    series = chart.series[int(ser_idx)]
    expected = ast.literal_eval(categories_expr)
    assert series.categories == expected, (
        f"expected categories {expected}, got {series.categories}"
    )


@then("document.charts is an empty list")
def then_document_charts_is_an_empty_list(context: Context):
    charts = context.document.charts
    assert charts == [], f"expected [], got {charts!r}"


# ==========================================================
# chart-create-bar steps
# ==========================================================


# -- bar-chart create-side series data (kept here so the When step is terse) --
_BAR_CATEGORIES: list[str] = ["Q1", "Q2", "Q3"]
_BAR_SERIES_2: dict[str, list[float]] = {
    "North": [10.0, 20.0, 15.0],
    "South": [7.0, 14.0, 21.0],
}
_BAR_SERIES_1: dict[str, list[float]] = {
    "Only": [5.0, 10.0, 15.0],
}


# given ----------------------------------------------------


@given("the chart-create-bar base document")
def given_the_chart_create_bar_base_document(context: Context):
    # -- remember how many paragraphs the base fixture ships with so the
    # -- "chart paragraph is positioned after the base paragraphs" step can
    # -- reason about it without hard-coding a count.
    context.document = Document(test_docx("chart-create-bar-base"))
    context.base_paragraph_count = len(context.document.paragraphs)


# when -----------------------------------------------------


@when("I add a BAR chart with {n_cats:d} categories and {n_series:d} series")
def when_I_add_a_BAR_chart_with_n_cats_and_n_series(
    context: Context, n_cats: int, n_series: int
):
    # -- the scenarios in this feature are pinned to 3 categories and either
    # -- 1 or 2 series so the expected values stay readable. Guard against
    # -- unexpected combinations instead of silently mis-configuring data.
    assert n_cats == 3, (
        f"chart-create-bar scenarios expect 3 categories, got {n_cats}"
    )
    if n_series == 2:
        series_data = _BAR_SERIES_2
    elif n_series == 1:
        series_data = _BAR_SERIES_1
    else:
        raise AssertionError(
            f"chart-create-bar scenarios expect 1 or 2 series, got {n_series}"
        )
    context.chart = context.document.add_chart(
        WD_CHART_TYPE.BAR, _BAR_CATEGORIES, series_data
    )


# -- "I save and reopen the document" is already registered in footnotes.py; we
# -- rely on behave's global step registry to match it from this feature too.


# then -----------------------------------------------------


@then("document.charts has one entry")
def then_document_charts_has_one_entry(context: Context):
    charts = context.document.charts
    assert len(charts) == 1, f"expected 1 chart, got {len(charts)}"
    assert isinstance(charts[0], Chart), f"expected Chart, got {type(charts[0])}"


@then("document.charts has two entries")
def then_document_charts_has_two_entries(context: Context):
    charts = context.document.charts
    assert len(charts) == 2, f"expected 2 charts, got {len(charts)}"
    for idx, chart in enumerate(charts):
        assert isinstance(chart, Chart), (
            f"expected Chart at index {idx}, got {type(chart)}"
        )


@then("the chart reference sits in the last body paragraph")
def then_the_chart_reference_sits_in_the_last_body_paragraph(context: Context):
    # -- add_chart wraps the drawing in its own paragraph appended to the body --
    body = context.document.element.body
    paragraphs = body.xpath("./w:p")
    assert len(paragraphs) > 0, "document body has no paragraphs"
    last_p = paragraphs[-1]
    chart_refs = last_p.xpath(
        ".//w:drawing/wp:inline/a:graphic/a:graphicData/c:chart"
        " | .//w:drawing/wp:anchor/a:graphic/a:graphicData/c:chart"
    )
    assert len(chart_refs) == 1, (
        f"expected exactly one c:chart ref in last paragraph, got {len(chart_refs)}"
    )


@then("the chart_type of the first chart is WD_CHART_TYPE.{member}")
def then_chart_type_of_first_chart_is_WD_CHART_TYPE_member(
    context: Context, member: str
):
    chart = context.document.charts[0]
# -- chart-create-line steps --------------------------------------------------


@given("a blank chart-create-line base document")
def given_a_blank_chart_create_line_base_document(context: Context):
    context.document = Document(test_docx("chart-create-line-base"))
    # -- sanity: fixture must start with zero charts so the assertions below
    # -- about "document.charts has 1 chart" are meaningful.
    assert context.document.charts == [], (
        f"fixture must have no charts, found {len(context.document.charts)}"
    )


@when(
    "I add a LINE chart with categories {categories_expr} and series {series_expr}"
)
def when_I_add_a_LINE_chart_with_categories_and_series(
    context: Context, categories_expr: str, series_expr: str
):
    categories = ast.literal_eval(categories_expr)
    series_data = ast.literal_eval(series_expr)
    context.added_chart = context.document.add_chart(
        WD_CHART_TYPE.LINE, categories, series_data
    )


@when("I save and reopen the chart-create-line document")
def when_I_save_and_reopen_the_chart_create_line_document(context: Context):
    buf = io.BytesIO()
    context.document.save(buf)
    buf.seek(0)
    context.document = Document(buf)
    # -- the added_chart reference is from the pre-save document; drop it so
    # -- later steps reach for document.charts[...] instead.
    context.added_chart = None


@then("document.charts has {count:d} chart")
@then("document.charts has {count:d} charts")
def then_document_charts_has_n_charts(context: Context, count: int):
    charts = context.document.charts
    assert len(charts) == count, f"expected {count} charts, got {len(charts)}"


@then("the added chart is the last embedded chart in the document")
def then_added_chart_is_last_in_document(context: Context):
    charts = context.document.charts
    assert len(charts) >= 1, "expected at least one chart in document"
    last = charts[-1]
    # -- Compare the underlying chart part; proxy identity is not preserved --
    assert last.part is context.added_chart.part, (
        "added chart does not match document.charts[-1]"
    )


@then("the added chart.chart_type == WD_CHART_TYPE.{member}")
def then_added_chart_chart_type_eq(context: Context, member: str):
    expected = WD_CHART_TYPE[member]
    actual = context.added_chart.chart_type
    assert actual == expected, f"expected {expected}, got {actual}"


@then("the added chart has {count:d} series")
def then_added_chart_has_n_series(context: Context, count: int):
    series = context.added_chart.series
    assert len(series) == count, f"expected {count} series, got {len(series)}"


@then("[s.name for s in added_chart.series] == {names_expr}")
def then_added_chart_series_names_eq(context: Context, names_expr: str):
    expected = ast.literal_eval(names_expr)
    actual = [s.name for s in context.added_chart.series]
    assert actual == expected, f"expected series names {expected}, got {actual}"


@then("added_chart.series[{ser_idx:d}].values == {values_expr}")
def then_added_chart_series_values_eq(
    context: Context, ser_idx: int, values_expr: str
):
    expected = ast.literal_eval(values_expr)
    actual = context.added_chart.series[ser_idx].values
    assert actual == expected, f"expected values {expected}, got {actual}"


@then("added_chart.series[{ser_idx:d}].categories == {categories_expr}")
def then_added_chart_series_categories_eq(
    context: Context, ser_idx: int, categories_expr: str
):
    expected = ast.literal_eval(categories_expr)
    actual = context.added_chart.series[ser_idx].categories
    assert actual == expected, f"expected categories {expected}, got {actual}"


@then("document.charts[{idx:d}].chart_type == WD_CHART_TYPE.{member}")
def then_document_charts_idx_chart_type_eq(
    context: Context, idx: int, member: str
):
    chart = context.document.charts[idx]
    expected = WD_CHART_TYPE[member]
    assert chart.chart_type == expected, (
        f"expected {expected}, got {chart.chart_type}"
    )


@then("charts[{idx:d}].series has {n:d} entries")
def then_charts_idx_series_has_n_entries(context: Context, idx: int, n: int):
    chart = context.document.charts[idx]
    series = chart.series
    assert len(series) == n, (
        f"expected {n} series on chart {idx}, got {len(series)}"
    )


@then('charts[{idx:d}].series[{ser_idx:d}].name == "{name}"')
def then_charts_idx_series_ser_idx_name_eq(
    context: Context, idx: int, ser_idx: int, name: str
):
    chart = context.document.charts[idx]
    series = chart.series[ser_idx]
    assert series.name == name, f"expected name {name!r}, got {series.name!r}"


def _first_chart_part(context: Context) -> ChartPart:
    chart = context.document.charts[0]
    part = chart.part
    assert isinstance(part, ChartPart), f"expected ChartPart, got {type(part)}"
    return part


@then("the chart part XML contains a c:barChart element")
def then_chart_part_xml_contains_c_barChart(context: Context):
    part = _first_chart_part(context)
    bar_charts = part.chartSpace.xpath(".//c:barChart")
    assert len(bar_charts) == 1, (
        f"expected exactly one c:barChart element, got {len(bar_charts)}"
    )


@then('the c:barChart has c:barDir with val "{val}"')
def then_barChart_has_barDir_with_val(context: Context, val: str):
    part = _first_chart_part(context)
    bar_dirs = part.chartSpace.xpath(".//c:barChart/c:barDir/@val")
    assert bar_dirs == [val], f"expected c:barDir val={val!r}, got {bar_dirs!r}"


@then('the c:barChart has c:grouping with val "{val}"')
def then_barChart_has_grouping_with_val(context: Context, val: str):
    part = _first_chart_part(context)
    groupings = part.chartSpace.xpath(".//c:barChart/c:grouping/@val")
    assert groupings == [val], (
        f"expected c:grouping val={val!r}, got {groupings!r}"
    )


@then("the chart part XML contains {n:d} c:ser entries")
def then_chart_part_xml_contains_n_c_ser_entries(context: Context, n: int):
    part = _first_chart_part(context)
    sers = part.chartSpace.xpath(".//c:barChart/c:ser")
    assert len(sers) == n, f"expected {n} c:ser entries, got {len(sers)}"


@then("the chart paragraph is positioned after the base paragraphs")
def then_chart_paragraph_is_positioned_after_base_paragraphs(context: Context):
    # -- add_chart appends a brand-new paragraph carrying the drawing, so the
    # -- body should now have exactly one more paragraph than the base had,
    # -- and the new (last) paragraph is the one containing the chart ref.
    body = context.document.element.body
    paragraphs = body.xpath("./w:p")
    expected = context.base_paragraph_count + 1
    assert len(paragraphs) == expected, (
        f"expected {expected} body paragraphs after add_chart, got {len(paragraphs)}"
    )

    # -- none of the base paragraphs should carry a chart ref --
    for i, p in enumerate(paragraphs[: context.base_paragraph_count]):
        refs = p.xpath(".//c:chart")
        assert len(refs) == 0, (
            f"base paragraph {i} unexpectedly carries a chart ref"
        )
    # -- the newly-appended (last) paragraph should carry exactly one --
    last_refs = paragraphs[-1].xpath(".//c:chart")
    assert len(last_refs) == 1, (
        f"expected exactly one c:chart ref in appended paragraph, got {len(last_refs)}"
    )


@then("every chart has chart_type WD_CHART_TYPE.{member}")
def then_every_chart_has_chart_type_WD_CHART_TYPE_member(
    context: Context, member: str
):
    expected = WD_CHART_TYPE[member]
    charts = context.document.charts
    assert charts, "expected at least one chart"
    for idx, chart in enumerate(charts):
        assert chart.chart_type == expected, (
            f"chart {idx}: expected {expected}, got {chart.chart_type}"
        )
@then("document.charts[{idx:d}].series[{ser_idx:d}].values == {values_expr}")
def then_document_charts_idx_series_values_eq(
    context: Context, idx: int, ser_idx: int, values_expr: str
):
    expected = ast.literal_eval(values_expr)
    actual = context.document.charts[idx].series[ser_idx].values
    assert actual == expected, f"expected values {expected}, got {actual}"


@then(
    "document.charts[{idx:d}].series[{ser_idx:d}].categories == {categories_expr}"
)
def then_document_charts_idx_series_categories_eq(
    context: Context, idx: int, ser_idx: int, categories_expr: str
):
    expected = ast.literal_eval(categories_expr)
    actual = context.document.charts[idx].series[ser_idx].categories
    assert actual == expected, f"expected categories {expected}, got {actual}"


@then("the chart part XML contains a c:lineChart element")
def then_chart_part_xml_contains_lineChart(context: Context):
    # -- `added_chart` may have been cleared by save-and-reopen; fall back to
    # -- the first chart in the document in that case.
    chart = context.added_chart or context.document.charts[0]
    from docx.oxml.ns import qn

    chartSpace = chart.part.element
    matches = chartSpace.findall(f".//{qn('c:lineChart')}")
    assert len(matches) == 1, (
        f"expected exactly 1 c:lineChart element, found {len(matches)}"
    )


@then("the chart part XML contains {count:d} c:ser elements")
def then_chart_part_xml_contains_n_ser(context: Context, count: int):
    chart = context.added_chart or context.document.charts[0]
    from docx.oxml.ns import qn

    chartSpace = chart.part.element
    matches = chartSpace.findall(f".//{qn('c:ser')}")
    assert len(matches) == count, (
        f"expected {count} c:ser elements, found {len(matches)}"
    )
