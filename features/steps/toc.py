"""Step implementations for table-of-contents features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.fields import Field
from docx.text.paragraph import Paragraph

from helpers import test_docx

# given ====================================================


@given("a document having heading paragraphs")
def given_a_document_having_heading_paragraphs(context: Context):
    context.document = Document(test_docx("toc-base"))
    # -- the anchor is paragraph 0, a body paragraph just before the first
    #    heading; scenarios that insert at position use it as a target. --
    context.anchor = context.document.paragraphs[0]


@given("a document having no heading paragraphs")
def given_a_document_having_no_heading_paragraphs(context: Context):
    context.document = Document()


# when =====================================================


@when("I assign toc = document.add_table_of_contents()")
def when_I_assign_toc_eq_document_add_table_of_contents(context: Context):
    context.toc = context.document.add_table_of_contents()


@when("I assign toc = document.add_table_of_contents(levels=(1, 1))")
def when_I_assign_toc_eq_add_table_of_contents_levels_1_1(context: Context):
    context.toc = context.document.add_table_of_contents(levels=(1, 1))


@when("I assign toc = document.add_table_of_contents(levels=(2, 3))")
def when_I_assign_toc_eq_add_table_of_contents_levels_2_3(context: Context):
    context.toc = context.document.add_table_of_contents(levels=(2, 3))


@when("I assign toc = anchor.insert_table_of_contents_before()")
def when_I_assign_toc_eq_anchor_insert_table_of_contents_before(context: Context):
    context.toc = context.anchor.insert_table_of_contents_before()


@when("I assign toc = anchor.insert_table_of_contents_after()")
def when_I_assign_toc_eq_anchor_insert_table_of_contents_after(context: Context):
    context.toc = context.anchor.insert_table_of_contents_after()


@when(
    "I call document.add_table_of_contents(levels={levels}) expecting ValueError"
)
def when_I_call_add_table_of_contents_expecting_ValueError(
    context: Context, levels: str
):
    # -- `levels` arrives as the raw tuple literal "(0, 3)"; eval is safe
    #    here because the scenario table is authored in-repo. --
    parsed = eval(levels)  # noqa: S307 — scenario-controlled literal
    context.raised: Exception | None = None
    try:
        context.document.add_table_of_contents(levels=parsed)
    except ValueError as exc:
        context.raised = exc


# then =====================================================


def _toc_field(paragraph: Paragraph) -> Field:
    """Return the sole complex TOC field carried by `paragraph`."""
    fields = [f for f in paragraph.fields if f.type == "TOC"]
    assert len(fields) == 1, (
        f"expected exactly one TOC field, got {len(fields)}"
    )
    assert fields[0].is_complex, "expected TOC field to be a complex field"
    return fields[0]


@then("toc is a Paragraph object")
def then_toc_is_a_Paragraph_object(context: Context):
    assert type(context.toc) is Paragraph, (
        f"expected a Paragraph object, got {type(context.toc)}"
    )


@then("toc is the last paragraph in the document")
def then_toc_is_the_last_paragraph_in_the_document(context: Context):
    last = context.document.paragraphs[-1]
    assert last._p is context.toc._p, (
        "toc is not the last paragraph in the document body"
    )


@then("toc is the first paragraph in the document")
def then_toc_is_the_first_paragraph_in_the_document(context: Context):
    first = context.document.paragraphs[0]
    assert first._p is context.toc._p, (
        "toc is not the first paragraph in the document body"
    )


@then("toc is the paragraph after the anchor")
def then_toc_is_the_paragraph_after_the_anchor(context: Context):
    paragraphs = list(context.document.paragraphs)
    anchor_idx = next(
        i for i, p in enumerate(paragraphs) if p._p is context.anchor._p
    )
    assert paragraphs[anchor_idx + 1]._p is context.toc._p, (
        "toc does not immediately follow the anchor paragraph"
    )


@then("toc has one complex TOC field")
def then_toc_has_one_complex_TOC_field(context: Context):
    _toc_field(context.toc)


@then("the TOC field instruction is '{instruction}'")
def then_the_TOC_field_instruction_is(context: Context, instruction: str):
    field = _toc_field(context.toc)
    actual = field.instruction
    assert actual == instruction, (
        f"expected TOC instruction {instruction!r}, got {actual!r}"
    )


@then("the TOC preview lists {count:d} entries")
def then_the_TOC_preview_lists_count_entries(context: Context, count: int):
    field = _toc_field(context.toc)
    result = field.result_text
    if count == 0:
        assert result == "", (
            f"expected empty TOC preview, got {result!r}"
        )
        return
    # -- each entry renders as a single run; the paragraph contains one
    #    w:br per entry, so counting runs with text bodies is the robust
    #    cross-check. we rely on result_text's own line structure. --
    lines = [ln for ln in result.split("\n") if ln]
    assert len(lines) == count, (
        f"expected {count} TOC entries, got {len(lines)}: {lines!r}"
    )


@then('the TOC preview contains "{text}"')
def then_the_TOC_preview_contains(context: Context, text: str):
    field = _toc_field(context.toc)
    assert text in field.result_text, (
        f"TOC preview missing {text!r}: {field.result_text!r}"
    )


@then('the TOC preview does not contain "{text}"')
def then_the_TOC_preview_does_not_contain(context: Context, text: str):
    field = _toc_field(context.toc)
    assert text not in field.result_text, (
        f"TOC preview unexpectedly contains {text!r}: {field.result_text!r}"
    )


@then("a ValueError is raised")
def then_a_ValueError_is_raised(context: Context):
    assert isinstance(context.raised, ValueError), (
        f"expected ValueError, got {context.raised!r}"
    )
