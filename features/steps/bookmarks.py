"""Step implementations for bookmark-related features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.bookmarks import Bookmark, Bookmarks

from helpers import test_docx


# given ====================================================


@given("a document having bookmarks")
def given_a_document_having_bookmarks(context: Context):
    context.document = Document(test_docx("bmk-has-bookmarks"))


@given("a fresh document with one paragraph of text")
def given_a_fresh_document_with_one_paragraph(context: Context):
    document = Document()
    paragraph = document.add_paragraph("Hello, world.")
    context.document = document
    context.paragraph = paragraph


@given("a paragraph with three runs")
def given_a_paragraph_with_three_runs(context: Context):
    document = Document()
    paragraph = document.add_paragraph("alpha ")
    paragraph.add_run("bravo ")
    paragraph.add_run("charlie")
    context.document = document
    context.paragraph = paragraph


# when =====================================================


@when('I call document.bookmarks.get("{name}")')
def when_I_call_document_bookmarks_get(context: Context, name: str):
    context.result = context.document.bookmarks.get(name)


@when('I assign bookmark = paragraph.add_bookmark("{name}")')
def when_I_assign_bookmark_eq_add_bookmark(context: Context, name: str):
    context.bookmark = context.paragraph.add_bookmark(name)


@when('I add a bookmark named "{name}" around the middle run')
def when_I_add_a_bookmark_around_middle_run(context: Context, name: str):
    middle = context.paragraph.runs[1]
    context.bookmark = context.paragraph.add_bookmark(
        name, start_run=middle, end_run=middle
    )


@when('I add a bookmark named "{name}" around the second and third runs')
def when_I_add_a_bookmark_around_second_and_third_runs(context: Context, name: str):
    runs = context.paragraph.runs
    context.bookmark = context.paragraph.add_bookmark(
        name, start_run=runs[1], end_run=runs[2]
    )


@when('I delete the bookmark named "{name}"')
def when_I_delete_the_bookmark_named(context: Context, name: str):
    bookmark = context.document.bookmarks.get(name)
    assert bookmark is not None, f"bookmark '{name}' not found"
    context.deleted_bookmark_id = bookmark.bookmark_id
    bookmark.delete()


@when('I rename the bookmark "{old_name}" to "{new_name}"')
def when_I_rename_the_bookmark(context: Context, old_name: str, new_name: str):
    bookmark = context.document.bookmarks.get(old_name)
    assert bookmark is not None, f"bookmark '{old_name}' not found"
    context.renamed_bookmark_id = bookmark.bookmark_id
    # -- `CT_BookmarkStart.name` is a settable RequiredAttribute; assign
    #    through the underlying element to rename. --
    bookmark._bookmarkStart.name = new_name  # pyright: ignore[reportPrivateUsage]


# then =====================================================


@then("document.bookmarks is a Bookmarks object")
def then_document_bookmarks_is_a_Bookmarks_object(context: Context):
    assert type(context.document.bookmarks) is Bookmarks


@then("len(document.bookmarks) == {count:d}")
def then_len_document_bookmarks(context: Context, count: int):
    actual = len(context.document.bookmarks)
    assert actual == count, f"expected len(document.bookmarks) == {count}, got {actual}"


@then("iterating document.bookmarks yields {count:d} Bookmark objects")
def then_iterating_document_bookmarks(context: Context, count: int):
    bm_iter = iter(context.document.bookmarks)
    items = list(bm_iter)
    assert len(items) == count, f"expected {count} bookmarks, got {len(items)}"
    for bm in items:
        assert isinstance(bm, Bookmark), f"expected Bookmark, got {type(bm)}"


@then('the result is a Bookmark object named "{name}"')
def then_the_result_is_a_bookmark_named(context: Context, name: str):
    assert isinstance(context.result, Bookmark), (
        f"expected Bookmark, got {type(context.result)}"
    )
    assert context.result.name == name, (
        f"expected name '{name}', got '{context.result.name}'"
    )


@then("the result is None")
def then_the_result_is_None(context: Context):
    assert context.result is None, f"expected None, got {context.result!r}"


@then('"{name}" in document.bookmarks')
def then_name_in_document_bookmarks(context: Context, name: str):
    assert name in context.document.bookmarks, (
        f"expected '{name}' in document.bookmarks"
    )


@then('"{name}" not in document.bookmarks')
def then_name_not_in_document_bookmarks(context: Context, name: str):
    assert name not in context.document.bookmarks, (
        f"expected '{name}' not in document.bookmarks"
    )


@then('bookmark.name == "{name}"')
def then_bookmark_name_eq(context: Context, name: str):
    actual = context.bookmark.name
    assert actual == name, f"expected name '{name}', got '{actual}'"


@then("bookmark.bookmark_id == {bookmark_id:d}")
def then_bookmark_bookmark_id_eq(context: Context, bookmark_id: int):
    actual = context.bookmark.bookmark_id
    assert actual == bookmark_id, (
        f"expected bookmark_id {bookmark_id}, got {actual}"
    )


@then("the bookmark wraps only the middle run")
def then_the_bookmark_wraps_only_the_middle_run(context: Context):
    # -- the bookmarkStart must be the immediate preceding sibling of the
    #    middle run, and the bookmarkEnd must be the immediate following
    #    sibling of that same run. --
    p = context.paragraph._p
    children = list(p)
    runs = p.xpath("./w:r")
    middle = runs[1]
    middle_idx = children.index(middle)
    prev_tag = children[middle_idx - 1].tag
    next_tag = children[middle_idx + 1].tag
    assert prev_tag.endswith("}bookmarkStart"), (
        f"expected bookmarkStart before middle run, got {prev_tag}"
    )
    assert next_tag.endswith("}bookmarkEnd"), (
        f"expected bookmarkEnd after middle run, got {next_tag}"
    )


@then("the bookmark wraps the last two runs")
def then_the_bookmark_wraps_the_last_two_runs(context: Context):
    p = context.paragraph._p
    children = list(p)
    runs = p.xpath("./w:r")
    start_idx = children.index(runs[1])
    end_idx = children.index(runs[2])
    prev_tag = children[start_idx - 1].tag
    next_tag = children[end_idx + 1].tag
    assert prev_tag.endswith("}bookmarkStart"), (
        f"expected bookmarkStart before second run, got {prev_tag}"
    )
    assert next_tag.endswith("}bookmarkEnd"), (
        f"expected bookmarkEnd after third run, got {next_tag}"
    )


@then('the bookmarkStart and bookmarkEnd for "{name}" are in different paragraphs')
def then_bookmark_spans_multiple_paragraphs(context: Context, name: str):
    bookmark = context.document.bookmarks.get(name)
    assert bookmark is not None, f"bookmark '{name}' not found"
    bookmark_id = str(bookmark.bookmark_id)
    body = context.document._body._body  # pyright: ignore[reportPrivateUsage]
    starts = body.xpath(f".//w:bookmarkStart[@w:id='{bookmark_id}']")
    ends = body.xpath(f".//w:bookmarkEnd[@w:id='{bookmark_id}']")
    assert len(starts) == 1 and len(ends) == 1, (
        f"expected one start and one end for '{name}'"
    )
    start_para = starts[0].getparent()
    end_para = ends[0].getparent()
    assert start_para is not end_para, (
        "expected bookmarkStart and bookmarkEnd to be in different paragraphs"
    )


@then("every bookmark has a unique bookmark_id")
def then_every_bookmark_has_unique_id(context: Context):
    ids = [bm.bookmark_id for bm in context.document.bookmarks]
    assert len(ids) == len(set(ids)), f"ids not unique: {ids!r}"


@then("no bookmarkStart with that id remains in the body")
def then_no_bookmarkStart_with_that_id(context: Context):
    bookmark_id = str(context.deleted_bookmark_id)
    body = context.document._body._body  # pyright: ignore[reportPrivateUsage]
    starts = body.xpath(f".//w:bookmarkStart[@w:id='{bookmark_id}']")
    assert len(starts) == 0, f"expected no bookmarkStart with id {bookmark_id}"


@then("no bookmarkEnd with that id remains in the body")
def then_no_bookmarkEnd_with_that_id(context: Context):
    bookmark_id = str(context.deleted_bookmark_id)
    body = context.document._body._body  # pyright: ignore[reportPrivateUsage]
    ends = body.xpath(f".//w:bookmarkEnd[@w:id='{bookmark_id}']")
    assert len(ends) == 0, f"expected no bookmarkEnd with id {bookmark_id}"


@then('the bookmark "{name}" keeps its original bookmark_id')
def then_bookmark_keeps_original_id(context: Context, name: str):
    bookmark = context.document.bookmarks.get(name)
    assert bookmark is not None, f"bookmark '{name}' not found"
    assert bookmark.bookmark_id == context.renamed_bookmark_id, (
        f"expected id {context.renamed_bookmark_id}, got {bookmark.bookmark_id}"
    )
