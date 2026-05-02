"""Step implementations for packaging-level features.

Covers recover mode, encrypted-docx detection, .docm loading, and digital
signature detection.
"""

from __future__ import annotations

import datetime as dt

from behave import given, then, when
from behave.runner import Context
from lxml import etree

from docx import Document
from docx.exceptions import EncryptedDocumentError

from helpers import test_docx, test_file


# given ===================================================


@given("a malformed .docx package")
def given_a_malformed_docx_package(context: Context):
    context.docx_path = test_docx("pkg-malformed")


@given("an OLE compound file masquerading as a .docx")
def given_an_ole_compound_file(context: Context):
    context.docx_path = test_docx("pkg-encrypted")


@given("a macro-enabled .docm document")
def given_a_macro_enabled_docm_document(context: Context):
    path = test_file("api-demo.docm")
    context.docx_path = path
    context.document = Document(path)


@given("a signed document")
def given_a_signed_document(context: Context):
    context.document = Document(test_docx("pkg-signed"))


# when ====================================================


@when("I call Document(path, recover=True)")
def when_call_document_recover(context: Context):
    context.document = Document(context.docx_path, recover=True)


# then ====================================================


@then("Document(path) raises XMLSyntaxError")
def then_document_raises_xmlsyntaxerror(context: Context):
    try:
        Document(context.docx_path)
    except etree.XMLSyntaxError:
        return
    raise AssertionError("expected XMLSyntaxError")


@then("Document(path) raises EncryptedDocumentError")
def then_document_raises_encrypted(context: Context):
    try:
        Document(context.docx_path)
    except EncryptedDocumentError:
        return
    raise AssertionError("expected EncryptedDocumentError")


@then("Document(path, recover=True) raises EncryptedDocumentError")
def then_document_recover_raises_encrypted(context: Context):
    try:
        Document(context.docx_path, recover=True)
    except EncryptedDocumentError:
        return
    raise AssertionError("expected EncryptedDocumentError with recover=True")


@then("document.recovery_warnings is non-empty")
def then_recovery_warnings_non_empty(context: Context):
    warnings = context.document.recovery_warnings
    assert warnings, f"expected at least one warning, got {warnings!r}"


@then("document.recovery_warnings is empty")
def then_recovery_warnings_empty(context: Context):
    warnings = context.document.recovery_warnings
    assert warnings == [], f"expected empty list, got {warnings!r}"


@then('at least one paragraph text contains "{fragment}"')
def then_paragraph_text_contains(context: Context, fragment: str):
    texts = [p.text for p in context.document.paragraphs]
    joined = "\n".join(texts)
    assert fragment in joined, f"{fragment!r} not found in {joined!r}"


@then("the document loads without error")
def then_document_loads_without_error(context: Context):
    assert context.document is not None


@then("document.has_macros is True")
def then_document_has_macros_true(context: Context):
    assert context.document.has_macros is True


@then("document.has_macros is False")
def then_document_has_macros_false(context: Context):
    assert context.document.has_macros is False


@then("document.is_signed is True")
def then_document_is_signed_true(context: Context):
    assert context.document.is_signed is True


@then("document.is_signed is False")
def then_document_is_signed_false(context: Context):
    assert context.document.is_signed is False


@then("document.signatures has length {count:d}")
def then_signatures_length(context: Context, count: int):
    actual = len(context.document.signatures)
    assert actual == count, f"expected {count}, got {actual}"


@then("document.signatures is empty")
def then_signatures_empty(context: Context):
    assert context.document.signatures == []


@then('signatures[{idx:d}].signer == "{value}"')
def then_signatures_signer(context: Context, idx: int, value: str):
    actual = context.document.signatures[idx].signer
    assert actual == value, f"expected {value!r}, got {actual!r}"


@then("signatures[{idx:d}].signed_at == {timestamp}")
def then_signatures_signed_at(context: Context, idx: int, timestamp: str):
    # -- accept "YYYY-MM-DDTHH:MM:SSZ" as UTC --
    stripped = timestamp.strip()
    if stripped.endswith("Z"):
        stripped = stripped[:-1] + "+00:00"
    expected = dt.datetime.fromisoformat(stripped)
    if expected.tzinfo is None:
        expected = expected.replace(tzinfo=dt.timezone.utc)
    actual = context.document.signatures[idx].signed_at
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then('signatures[{idx:d}].partname == "{value}"')
def then_signatures_partname(context: Context, idx: int, value: str):
    actual = str(context.document.signatures[idx].partname)
    assert actual == value, f"expected {value!r}, got {actual!r}"
