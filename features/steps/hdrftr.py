"""Step implementations for header and footer-related features."""

from behave import given, then, when

from docx import Document

from helpers import test_docx, test_file

# given ====================================================


@given("a _Footer object {with_or_no} footer definition as footer")
def given_a_Footer_object_with_or_no_footer_definition(context, with_or_no):
    section_idx = {"with a": 0, "with no": 1}[with_or_no]
    context.sections = Document(test_docx("hdr-header-footer")).sections
    context.footer = context.sections[section_idx].footer


@given("a _Header object {with_or_no} header definition as header")
def given_a_Header_object_with_or_no_header_definition(context, with_or_no):
    section_idx = {"with a": 0, "with no": 1}[with_or_no]
    context.sections = Document(test_docx("hdr-header-footer")).sections
    context.header = context.sections[section_idx].header


@given("a _Run object from a footer as run")
def given_a_Run_object_from_a_footer_as_run(context):
    footer = Document(test_docx("hdr-header-footer")).sections[0].footer
    context.run = footer.paragraphs[0].add_run()


@given("a _Run object from a header as run")
def given_a_Run_object_from_a_header_as_run(context):
    header = Document(test_docx("hdr-header-footer")).sections[0].header
    context.run = header.paragraphs[0].add_run()


@given("the next _Footer object with no footer definition as footer_2")
def given_the_next_Footer_object_with_no_footer_definition(context):
    context.footer_2 = context.sections[1].footer


@given("the next _Header object with no header definition as header_2")
def given_the_next_Header_object_with_no_header_definition(context):
    context.header_2 = context.sections[1].header


# when =====================================================


@when('I assign "Normal" to footer.paragraphs[0].style')
def when_I_assign_Body_Text_to_footer_style(context):
    context.footer.paragraphs[0].style = "Normal"


@when('I assign "Normal" to header.paragraphs[0].style')
def when_I_assign_Body_Text_to_header_style(context):
    context.header.paragraphs[0].style = "Normal"


@when("I assign {value} to header.is_linked_to_previous")
def when_I_assign_value_to_header_is_linked_to_previous(context, value):
    context.header.is_linked_to_previous = eval(value)


@when("I assign {value} to footer.is_linked_to_previous")
def when_I_assign_value_to_footer_is_linked_to_previous(context, value):
    context.footer.is_linked_to_previous = eval(value)


@when("I call run.add_picture()")
def when_I_call_run_add_picture(context):
    context.run.add_picture(test_file("test.png"))


# then =====================================================


@then("footer.is_linked_to_previous is {value}")
def then_footer_is_linked_to_previous_is_value(context, value):
    actual = context.footer.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "footer.is_linked_to_previous is %s" % actual


@then('footer.paragraphs[0].style.name == "Normal"')
def then_footer_paragraphs_0_style_name_eq_Normal(context):
    actual = context.footer.paragraphs[0].style.name
    expected = "Normal"
    assert actual == expected, "footer.paragraphs[0].style.name is %s" % actual


@then("footer_2.is_linked_to_previous is {value}")
def then_footer_2_is_linked_to_previous_is_value(context, value):
    actual = context.footer_2.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "footer_2.is_linked_to_previous is %s" % actual


@then("footer_2.paragraphs[0].text == footer.paragraphs[0].text")
def then_footer_2_text_eq_footer_text(context):
    actual = context.footer_2.paragraphs[0].text
    expected = context.footer.paragraphs[0].text
    assert actual == expected, "footer_2.paragraphs[0].text == %s" % actual


@then("header.is_linked_to_previous is {value}")
def then_header_is_linked_to_previous_is_value(context, value):
    actual = context.header.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "header.is_linked_to_previous is %s" % actual


@then('header.paragraphs[0].style.name == "Normal"')
def then_header_paragraphs_0_style_name_eq_Normal(context):
    actual = context.header.paragraphs[0].style.name
    expected = "Normal"
    assert actual == expected, "header.paragraphs[0].style.name is %s" % actual


@then("header_2.is_linked_to_previous is {value}")
def then_header_2_is_linked_to_previous_is_value(context, value):
    actual = context.header_2.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, "header_2.is_linked_to_previous is %s" % actual


@then("header_2.paragraphs[0].text == header.paragraphs[0].text")
def then_header_2_text_eq_header_text(context):
    actual = context.header_2.paragraphs[0].text
    expected = context.header.paragraphs[0].text
    assert actual == expected, "header_2.paragraphs[0].text == %s" % actual


# ==========================================================================
# Odd / even / first-page headers & footers
# (hdr-header-footer.feature extension)
# ==========================================================================


@given("a Section from the odd-even-hdrs document as section")
def given_a_Section_from_the_odd_even_hdrs_document(context):
    context.document = Document(test_docx("doc-odd-even-hdrs"))
    context.section = context.document.sections[0]


@when("I assign {value} to section.different_odd_and_even_pages_header_footer")
def when_I_assign_to_different_odd_and_even_pages_hdrftr(context, value):
    context.section.different_odd_and_even_pages_header_footer = eval(value)


# NOTE: `when_I_assign_value_to_section_different_first_page_hdrftr` from
# features/steps/section.py already handles the
# "I assign {bool_val} to section.different_first_page_header_footer" step,
# so we do not re-register it here.


@then("section.different_odd_and_even_pages_header_footer is {value}")
def then_section_different_odd_and_even_pages_header_footer_is(context, value):
    actual = context.section.different_odd_and_even_pages_header_footer
    expected = eval(value)
    assert actual == expected, (
        "section.different_odd_and_even_pages_header_footer is %s" % actual
    )


@then('section.header.paragraphs[0].text is "{text}"')
def then_section_header_paragraphs_0_text_is(context, text):
    actual = context.section.header.paragraphs[0].text
    assert actual == text, f"section.header text is {actual!r}"


@then('section.footer.paragraphs[0].text is "{text}"')
def then_section_footer_paragraphs_0_text_is(context, text):
    actual = context.section.footer.paragraphs[0].text
    assert actual == text, f"section.footer text is {actual!r}"


@then('section.even_page_header.paragraphs[0].text is "{text}"')
def then_section_even_page_header_paragraphs_0_text_is(context, text):
    actual = context.section.even_page_header.paragraphs[0].text
    assert actual == text, f"section.even_page_header text is {actual!r}"


@then('section.even_page_footer.paragraphs[0].text is "{text}"')
def then_section_even_page_footer_paragraphs_0_text_is(context, text):
    actual = context.section.even_page_footer.paragraphs[0].text
    assert actual == text, f"section.even_page_footer text is {actual!r}"


@then('section.first_page_header.paragraphs[0].text is "{text}"')
def then_section_first_page_header_paragraphs_0_text_is(context, text):
    actual = context.section.first_page_header.paragraphs[0].text
    assert actual == text, f"section.first_page_header text is {actual!r}"


@then('section.first_page_footer.paragraphs[0].text is "{text}"')
def then_section_first_page_footer_paragraphs_0_text_is(context, text):
    actual = context.section.first_page_footer.paragraphs[0].text
    assert actual == text, f"section.first_page_footer text is {actual!r}"


@then("section.even_page_header.is_linked_to_previous is {value}")
def then_section_even_page_header_is_linked_to_previous_is(context, value):
    actual = context.section.even_page_header.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, (
        "section.even_page_header.is_linked_to_previous is %s" % actual
    )


@then("section.even_page_footer.is_linked_to_previous is {value}")
def then_section_even_page_footer_is_linked_to_previous_is(context, value):
    actual = context.section.even_page_footer.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, (
        "section.even_page_footer.is_linked_to_previous is %s" % actual
    )


@then("section.first_page_header.is_linked_to_previous is {value}")
def then_section_first_page_header_is_linked_to_previous_is(context, value):
    actual = context.section.first_page_header.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, (
        "section.first_page_header.is_linked_to_previous is %s" % actual
    )


@then("section.first_page_footer.is_linked_to_previous is {value}")
def then_section_first_page_footer_is_linked_to_previous_is(context, value):
    actual = context.section.first_page_footer.is_linked_to_previous
    expected = eval(value)
    assert actual == expected, (
        "section.first_page_footer.is_linked_to_previous is %s" % actual
    )


