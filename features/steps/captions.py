"""Step implementations for caption-paragraph features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document

from helpers import test_docx


# given ===================================================


@given("a document having two captions")
def given_a_document_having_two_captions(context: Context):
    context.document = Document(test_docx("cap-caption"))


@given('a fresh default document with one paragraph "{text}"')
def given_a_fresh_default_document_with_one_paragraph_text(
    context: Context, text: str
):
    context.document = Document()
    context.paragraph = context.document.add_paragraph(text)


# when ====================================================


@when('I call document.add_caption("{text}", label="{label}")')
def when_call_document_add_caption(context: Context, text: str, label: str):
    context.paragraph = context.document.add_caption(text, label=label)


@when('I call paragraph.add_caption_before("{text}", label="{label}")')
def when_call_paragraph_add_caption_before(
    context: Context, text: str, label: str
):
    context.caption = context.paragraph.add_caption_before(text, label=label)


@when('I call paragraph.add_caption_after("{text}", label="{label}")')
def when_call_paragraph_add_caption_after(
    context: Context, text: str, label: str
):
    context.caption = context.paragraph.add_caption_after(text, label=label)


# then ====================================================


@then('the returned paragraph style name is "{name}"')
def then_returned_paragraph_style_name(context: Context, name: str):
    actual = context.paragraph.style.name
    assert actual == name, f"expected {name!r}, got {actual!r}"


@then('the returned paragraph text is "{text}"')
def then_returned_paragraph_text(context: Context, text: str):
    actual = context.paragraph.text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then("the document has {count:d} Caption-styled paragraphs")
def then_document_has_caption_paragraphs(context: Context, count: int):
    captions = [
        p for p in context.document.paragraphs if p.style.name == "Caption"
    ]
    assert len(captions) == count, f"expected {count}, got {len(captions)}"


@then('the caption paragraphs contain "{a}" and "{b}"')
def then_caption_paragraphs_contain_both(context: Context, a: str, b: str):
    captions = [
        p.text for p in context.document.paragraphs if p.style.name == "Caption"
    ]
    joined = "\n".join(captions)
    assert a in joined, f"{a!r} not in {joined!r}"
    assert b in joined, f"{b!r} not in {joined!r}"


@then('the first paragraph text is "{text}"')
def then_first_paragraph_text(context: Context, text: str):
    actual = context.document.paragraphs[0].text
    assert actual == text, f"expected {text!r}, got {actual!r}"


@then('the second paragraph text is "{text}"')
def then_second_paragraph_text(context: Context, text: str):
    actual = context.document.paragraphs[1].text
    assert actual == text, f"expected {text!r}, got {actual!r}"
