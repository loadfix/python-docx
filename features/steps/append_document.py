"""Step implementations for cross-document append."""

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from helpers import test_file

# given ===================================================


@given("a blank destination document and a source document with content")
def given_blank_dest_and_source_with_content(context: Context):
    # -- source: 2 paragraphs + embedded image + explicit Heading 1 style --
    source = Document()
    source.add_heading("Chapter One", level=1)
    source.add_paragraph("This paragraph travels with the source body.")
    source.add_picture(test_file("test.png"))
    context.source = source

    # -- destination: a clean blank document --
    dest = Document()
    dest._body.clear_content()
    context.dest = dest


# when ====================================================


@when("I call dest.append_document(source)")
def when_call_append_document(context: Context):
    context.copied_count = context.dest.append_document(context.source)


# then ====================================================


@then("dest has every paragraph from the source")
def then_dest_has_source_paragraphs(context: Context):
    src_texts = [p.text for p in context.source.paragraphs if p.text]
    dst_texts = [p.text for p in context.dest.paragraphs]
    for text in src_texts:
        assert text in dst_texts, (
            "paragraph %r missing from destination (found %r)" % (text, dst_texts)
        )


@then("dest has the image relationship from the source")
def then_dest_has_image_relationship(context: Context):
    image_rels = [
        r for r in context.dest.part.rels.values() if r.reltype == RT.IMAGE
    ]
    assert image_rels, "expected at least one IMAGE relationship in destination"


@then("dest has the Heading 1 style from the source")
def then_dest_has_heading1_style(context: Context):
    assert "Heading 1" in context.dest.styles, (
        "Heading 1 style missing from destination document"
    )
