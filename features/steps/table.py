# pyright: reportPrivateUsage=false

"""Step implementations for table-related features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.table import (
    WD_ALIGN_VERTICAL,
    WD_BORDER_STYLE,
    WD_ROW_HEIGHT_RULE,
    WD_SHADING_PATTERN,
    WD_TABLE_ALIGNMENT,
    WD_TABLE_AUTOFIT,
    WD_TABLE_DIRECTION,
    WD_TEXT_DIRECTION,
)
from docx.shared import Inches, RGBColor
from docx.table import Table, _Cell, _Column, _Columns, _Row, _Rows

from helpers import test_docx

# given ===================================================


@given("a 2 x 2 table")
def given_a_2x2_table(context: Context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given("a 3x3 table having {span_state}")
def given_a_3x3_table_having_span_state(context: Context, span_state: str):
    table_idx = {
        "only uniform cells": 0,
        "a horizontal span": 1,
        "a vertical span": 2,
        "a combined span": 3,
    }[span_state]
    document = Document(test_docx("tbl-cell-access"))
    context.table_ = document.tables[table_idx]


@given("a _Cell object spanning {count} layout-grid cells")
def given_a_Cell_object_spanning_count_layout_grid_cells(context: Context, count: str):
    document = Document(test_docx("tbl-cell-props"))
    table = document.tables[0]
    context.cell = _Cell(table._tbl.tr_lst[int(count)].tc_lst[0], table)


@given("a _Cell object with {state} vertical alignment as cell")
def given_a_Cell_object_with_vertical_alignment_as_cell(context: Context, state: str):
    table_idx = {
        "inherited": 0,
        "bottom": 1,
        "center": 2,
        "top": 3,
    }[state]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    context.cell = table.cell(0, 0)


@given("a column collection having two columns")
def given_a_column_collection_having_two_columns(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    context.columns = document.tables[0].columns


@given("a row collection having two rows")
def given_a_row_collection_having_two_rows(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    context.rows = document.tables[0].rows


@given("a table")
def given_a_table(context: Context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given("a table cell")
def given_a_table_cell(context: Context):
    table = Document(test_docx("tbl-2x2-table")).tables[0]
    context.cell = table.cell(0, 0)


@given("a document containing three tables")
def given_a_document_containing_three_tables_step(context: Context):
    context.document = Document(test_docx("tbl-multi"))


@given("a detached table")
def given_a_detached_table(context: Context):
    from docx.oxml.table import CT_Tbl
    from docx.shared import Inches

    tbl = CT_Tbl.new_tbl(2, 2, Inches(6))
    context.table_ = Table(tbl, None)


@given("a table cell having a width of {width}")
def given_a_table_cell_having_a_width_of_width(context: Context, width: str):
    table_idx = {"no explicit setting": 0, "1 inch": 1, "2 inches": 2}[width]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    cell = table.cell(0, 0)
    context.cell = cell


@given("a table column having a width of {width_desc}")
def given_a_table_having_a_width_of_width_desc(context: Context, width_desc: str):
    col_idx = {
        "no explicit setting": 0,
        "1440": 1,
    }[width_desc]
    docx_path = test_docx("tbl-col-props")
    document = Document(docx_path)
    context.column = document.tables[0].columns[col_idx]


@given("a table having {alignment} alignment")
def given_a_table_having_alignment_alignment(context: Context, alignment: str):
    table_idx = {
        "inherited": 3,
        "left": 4,
        "right": 5,
        "center": 6,
    }[alignment]
    docx_path = test_docx("tbl-props")
    document = Document(docx_path)
    context.table_ = document.tables[table_idx]


@given("a table having an autofit layout of {autofit}")
def given_a_table_having_an_autofit_layout_of_autofit(context: Context, autofit: str):
    tbl_idx = {
        "no explicit setting": 0,
        "autofit": 1,
        "fixed": 2,
    }[autofit]
    document = Document(test_docx("tbl-props"))
    context.table_ = document.tables[tbl_idx]


@given("a table having {style} style")
def given_a_table_having_style(context: Context, style: str):
    table_idx = {
        "no explicit": 0,
        "Table Grid": 1,
        "Light Shading - Accent 1": 2,
    }[style]
    document = Document(test_docx("tbl-having-applied-style"))
    context.document = document
    context.table_ = document.tables[table_idx]


@given("a table having table direction set {setting}")
def given_a_table_having_table_direction_setting(context: Context, setting: str):
    table_idx = ["to inherit", "right-to-left", "left-to-right"].index(setting)
    document = Document(test_docx("tbl-on-off-props"))
    context.table_ = document.tables[table_idx]


@given("a table having two columns")
def given_a_table_having_two_columns(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    # context.table is used internally by behave, underscore added
    # to distinguish this one
    context.table_ = document.tables[0]


@given("a table having two rows")
def given_a_table_having_two_rows(context: Context):
    docx_path = test_docx("blk-containing-table")
    document = Document(docx_path)
    context.table_ = document.tables[0]


@given("a table row ending with {count} empty grid columns")
def given_a_table_row_ending_with_count_empty_grid_columns(context: Context, count: str):
    document = Document(test_docx("tbl-props"))
    table = document.tables[8]
    context.row = table.rows[int(count)]


@given("a table row having height of {state}")
def given_a_table_row_having_height_of_state(context: Context, state: str):
    table_idx = {"no explicit setting": 0, "2 inches": 2, "3 inches": 3}[state]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    context.row = table.rows[0]


@given("a table row having height rule {state}")
def given_a_table_row_having_height_rule_state(context: Context, state: str):
    table_idx = {"no explicit setting": 0, "automatic": 1, "at least": 2, "exactly": 3}[state]
    document = Document(test_docx("tbl-props"))
    table = document.tables[table_idx]
    context.row = table.rows[0]


@given("a table row starting with {count} empty grid columns")
def given_a_table_row_starting_with_count_empty_grid_columns(context: Context, count: str):
    document = Document(test_docx("tbl-props"))
    table = document.tables[7]
    context.row = table.rows[int(count)]


# when =====================================================


@when("I add a 1.0 inch column to the table")
def when_I_add_a_1_inch_column_to_table(context: Context):
    context.column = context.table_.add_column(Inches(1.0))


@when("I add a 2 x 2 table into the first cell")
def when_I_add_a_2x2_table_into_the_first_cell(context: Context):
    context.table_ = context.cell.add_table(2, 2)


@when("I add a row to the table")
def when_add_row_to_table(context: Context):
    table = context.table_
    context.row = table.add_row()


@when("I assign a string to the cell text attribute")
def when_assign_string_to_cell_text_attribute(context: Context):
    cell = context.cell
    text = "foobar"
    cell.text = text
    context.expected_text = text


@when("I assign {value} to cell.vertical_alignment")
def when_I_assign_value_to_cell_vertical_alignment(context: Context, value: str):
    context.cell.vertical_alignment = eval(value)


@when("I assign {value} to row.height")
def when_I_assign_value_to_row_height(context: Context, value: str):
    new_value = None if value == "None" else int(value)
    context.row.height = new_value


@when("I assign {value} to row.height_rule")
def when_I_assign_value_to_row_height_rule(context: Context, value: str):
    new_value = None if value == "None" else getattr(WD_ROW_HEIGHT_RULE, value)
    context.row.height_rule = new_value


@when("I assign {value_str} to table.alignment")
def when_I_assign_value_to_table_alignment(context: Context, value_str: str):
    value = {
        "None": None,
        "WD_TABLE_ALIGNMENT.LEFT": WD_TABLE_ALIGNMENT.LEFT,
        "WD_TABLE_ALIGNMENT.RIGHT": WD_TABLE_ALIGNMENT.RIGHT,
        "WD_TABLE_ALIGNMENT.CENTER": WD_TABLE_ALIGNMENT.CENTER,
    }[value_str]
    table = context.table_
    table.alignment = value


@when("I assign {value} to table.style")
def when_apply_value_to_table_style(context: Context, value: str):
    table, styles = context.table_, context.document.styles
    if value == "None":
        new_value = None
    elif value.startswith("styles["):
        new_value = styles[value.split("'")[1]]
    else:
        new_value = styles[value]
    table.style = new_value


@when("I assign {value} to table.table_direction")
def when_assign_value_to_table_table_direction(context: Context, value: str):
    new_value = None if value == "None" else getattr(WD_TABLE_DIRECTION, value)
    context.table_.table_direction = new_value


@when("I merge from cell {origin} to cell {other}")
def when_I_merge_from_cell_origin_to_cell_other(context: Context, origin: str, other: str):
    def cell(table: Table, idx: int):
        row, col = idx // 3, idx % 3
        return table.cell(row, col)

    a_idx, b_idx = int(origin) - 1, int(other) - 1
    table = context.table_
    a, b = cell(table, a_idx), cell(table, b_idx)
    a.merge(b)


@when("I set the cell width to {width}")
def when_I_set_the_cell_width_to_width(context: Context, width: str):
    new_value = {"1 inch": Inches(1)}[width]
    context.cell.width = new_value


@when("I set the column width to {width_emu}")
def when_I_set_the_column_width_to_width_emu(context: Context, width_emu: str):
    new_value = None if width_emu == "None" else int(width_emu)
    context.column.width = new_value


@when("I set the table autofit to {setting}")
def when_I_set_the_table_autofit_to_setting(context: Context, setting: str):
    new_value = {"autofit": True, "fixed": False}[setting]
    table = context.table_
    table.autofit = new_value


@when("I delete the second table")
def when_I_delete_the_second_table(context: Context):
    context.document.tables[1].delete()


@when("I delete the detached table")
def when_I_delete_the_detached_table(context: Context):
    context.table_.delete()


@when("I insert a paragraph after the second table")
def when_I_insert_a_paragraph_after_the_second_table(context: Context):
    ref = context.document.tables[1]
    ref.insert_paragraph_after("after-table")


@when("I insert a paragraph before the second table")
def when_I_insert_a_paragraph_before_the_second_table(context: Context):
    ref = context.document.tables[1]
    ref.insert_paragraph_before("before-table")


@when("I insert a 2x2 table after the second table")
def when_I_insert_a_2x2_table_after_the_second_table(context: Context):
    ref = context.document.tables[1]
    ref.insert_table_after(rows=2, cols=2)


# then =====================================================


@then("cell.grid_span is {count}")
def then_cell_grid_span_is_count(context: Context, count: str):
    expected = int(count)
    actual = context.cell.grid_span
    assert actual == expected, f"expected {expected}, got {actual}"


@then("cell.tables[0] is a 2 x 2 table")
def then_cell_tables_0_is_a_2x2_table(context: Context):
    cell = context.cell
    table = cell.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2


@then("cell.vertical_alignment is {value}")
def then_cell_vertical_alignment_is_value(context: Context, value: str):
    expected_value = {
        "None": None,
        "WD_ALIGN_VERTICAL.BOTTOM": WD_ALIGN_VERTICAL.BOTTOM,
        "WD_ALIGN_VERTICAL.CENTER": WD_ALIGN_VERTICAL.CENTER,
    }[value]
    actual_value = context.cell.vertical_alignment
    assert actual_value is expected_value, "cell.vertical_alignment is %s" % actual_value


@then("I can access a collection column by index")
def then_can_access_collection_column_by_index(context: Context):
    columns = context.columns
    for idx in range(2):
        column = columns[idx]
        assert isinstance(column, _Column)


@then("I can access a collection row by index")
def then_can_access_collection_row_by_index(context: Context):
    rows = context.rows
    for idx in range(2):
        row = rows[idx]
        assert isinstance(row, _Row)


@then("I can access the column collection of the table")
def then_can_access_column_collection_of_table(context: Context):
    table = context.table_
    columns = table.columns
    assert isinstance(columns, _Columns)


@then("I can access the row collection of the table")
def then_can_access_row_collection_of_table(context: Context):
    table = context.table_
    rows = table.rows
    assert isinstance(rows, _Rows)


@then("I can iterate over the column collection")
def then_can_iterate_over_column_collection(context: Context):
    columns = context.columns
    actual_count = 0
    for column in columns:
        actual_count += 1
        assert isinstance(column, _Column)
    assert actual_count == 2


@then("I can iterate over the row collection")
def then_can_iterate_over_row_collection(context: Context):
    rows = context.rows
    actual_count = 0
    for row in rows:
        actual_count += 1
        assert isinstance(row, _Row)
    assert actual_count == 2


@then("row.grid_cols_after is {value}")
def then_row_grid_cols_after_is_value(context: Context, value: str):
    expected = int(value)
    actual = context.row.grid_cols_after
    assert actual == expected, "expected %s, got %s" % (expected, actual)


@then("row.grid_cols_before is {value}")
def then_row_grid_cols_before_is_value(context: Context, value: str):
    expected = int(value)
    actual = context.row.grid_cols_before
    assert actual == expected, "expected %s, got %s" % (expected, actual)


@then("row.height is {value}")
def then_row_height_is_value(context: Context, value: str):
    expected_height = None if value == "None" else int(value)
    actual_height = context.row.height
    assert actual_height == expected_height, "expected %s, got %s" % (
        expected_height,
        actual_height,
    )


@then("row.height_rule is {value}")
def then_row_height_rule_is_value(context: Context, value: str):
    expected_rule = None if value == "None" else getattr(WD_ROW_HEIGHT_RULE, value)
    actual_rule = context.row.height_rule
    assert actual_rule == expected_rule, "expected %s, got %s" % (
        expected_rule,
        actual_rule,
    )


@then("table.alignment is {value_str}")
def then_table_alignment_is_value(context: Context, value_str: str):
    value = {
        "None": None,
        "WD_TABLE_ALIGNMENT.LEFT": WD_TABLE_ALIGNMENT.LEFT,
        "WD_TABLE_ALIGNMENT.RIGHT": WD_TABLE_ALIGNMENT.RIGHT,
        "WD_TABLE_ALIGNMENT.CENTER": WD_TABLE_ALIGNMENT.CENTER,
    }[value_str]
    table = context.table_
    assert table.alignment == value, "got %s" % table.alignment


@then("table.cell({row}, {col}).text is {expected_text}")
def then_table_cell_row_col_text_is_text(context: Context, row: str, col: str, expected_text: str):
    table = context.table_
    row_idx, col_idx = int(row), int(col)
    cell_text = table.cell(row_idx, col_idx).text
    assert cell_text == expected_text, "got %s" % cell_text


@then("table.style is styles['{style_name}']")
def then_table_style_is_styles_style_name(context: Context, style_name: str):
    table, styles = context.table_, context.document.styles
    expected_style = styles[style_name]
    assert table.style == expected_style, "got '%s'" % table.style


@then("table.table_direction is {value}")
def then_table_table_direction_is_value(context: Context, value: str):
    expected_value = None if value == "None" else getattr(WD_TABLE_DIRECTION, value)
    actual_value = context.table_.table_direction
    assert actual_value == expected_value, "got '%s'" % actual_value


@then("the cell contains the string I assigned")
def then_cell_contains_string_assigned(context: Context):
    cell, expected_text = context.cell, context.expected_text
    text = cell.paragraphs[0].runs[0].text
    msg = "expected '%s', got '%s'" % (expected_text, text)
    assert text == expected_text, msg


@then("the column cells text is {expected_text}")
def then_the_column_cells_text_is_expected_text(context: Context, expected_text: str):
    table = context.table_
    cells_text = " ".join(c.text for col in table.columns for c in col.cells)
    assert cells_text == expected_text, "got %s" % cells_text


@then("the length of the column collection is 2")
def then_len_of_column_collection_is_2(context: Context):
    columns = context.table_.columns
    assert len(columns) == 2


@then("the length of the row collection is 2")
def then_len_of_row_collection_is_2(context: Context):
    rows = context.table_.rows
    assert len(rows) == 2


@then("the new column has 2 cells")
def then_new_column_has_2_cells(context: Context):
    assert len(context.column.cells) == 2


@then("the new column is 1.0 inches wide")
def then_new_column_is_1_inches_wide(context: Context):
    assert context.column.width == Inches(1)


@then("the new row has 2 cells")
def then_new_row_has_2_cells(context: Context):
    assert len(context.row.cells) == 2


@then("the reported autofit setting is {autofit}")
def then_the_reported_autofit_setting_is_autofit(context: Context, autofit: str):
    expected_value = {"autofit": True, "fixed": False}[autofit]
    table = context.table_
    assert table.autofit is expected_value


@then("the reported column width is {width_emu}")
def then_the_reported_column_width_is_width_emu(context: Context, width_emu: str):
    expected_value = None if width_emu == "None" else int(width_emu)
    assert context.column.width == expected_value, "got %s" % context.column.width


@then("the reported width of the cell is {width}")
def then_the_reported_width_of_the_cell_is_width(context: Context, width: str):
    expected_width = {"None": None, "1 inch": Inches(1)}[width]
    actual_width = context.cell.width
    assert actual_width == expected_width, "expected %s, got %s" % (
        expected_width,
        actual_width,
    )


@then("the row cells text is {encoded_text}")
def then_the_row_cells_text_is_expected_text(context: Context, encoded_text: str):
    expected_text = encoded_text.replace("\\", "\n")
    table = context.table_
    cells_text = " ".join(c.text for row in table.rows for c in row.cells)
    assert cells_text == expected_text, "got %s" % cells_text


@then("the table has {count} columns")
def then_table_has_count_columns(context: Context, count: str):
    column_count = int(count)
    columns = context.table_.columns
    assert len(columns) == column_count


@then("the table has {count} rows")
def then_table_has_count_rows(context: Context, count: str):
    row_count = int(count)
    rows = context.table_.rows
    assert len(rows) == row_count


@then("the width of cell {n_str} is {inches_str} inches")
def then_the_width_of_cell_n_is_x_inches(context: Context, n_str: str, inches_str: str):
    def _cell(table: Table, idx: int):
        row, col = idx // 3, idx % 3
        return table.cell(row, col)

    idx, inches = int(n_str) - 1, float(inches_str)
    cell = _cell(context.table_, idx)
    assert cell.width is not None
    assert cell.width == Inches(inches), "got %s" % cell.width.inches


@then("the width of each cell is {inches} inches")
def then_the_width_of_each_cell_is_inches(context: Context, inches: str):
    table = context.table_
    expected_width = Inches(float(inches))
    for cell in table._cells:
        assert cell.width == expected_width, "got %s" % cell.width.inches


@then("the width of each column is {inches} inches")
def then_the_width_of_each_column_is_inches(context: Context, inches: str):
    table = context.table_
    expected_width = Inches(float(inches))
    for column in table.columns:
        assert column.width == expected_width, "got %s" % column.width.inches


# ---------------------------------------------------------------------------
# Borders ---------------------------------------------------------------------
# ---------------------------------------------------------------------------


def _resolve_border_style(name: str):
    if name == "None":
        return None
    # e.g. "WD_BORDER_STYLE.SINGLE"
    _, member = name.split(".")
    return getattr(WD_BORDER_STYLE, member)


def _resolve_rgb(name: str) -> RGBColor | None:
    if name == "None":
        return None
    return RGBColor.from_string(name)


@given("a table having borders on every edge")
def given_a_table_having_borders_on_every_edge(context: Context):
    document = Document(test_docx("tbl-borders"))
    context.table_ = document.tables[0]


@given("a table having no explicit borders")
def given_a_table_having_no_explicit_borders(context: Context):
    document = Document(test_docx("tbl-borders"))
    context.table_ = document.tables[1]


@given("a cell having a THICK left border")
def given_a_cell_having_a_THICK_left_border(context: Context):
    document = Document(test_docx("tbl-borders"))
    context.cell = document.tables[1].cell(0, 0)


@given("a cell having no explicit borders")
def given_a_cell_having_no_explicit_borders(context: Context):
    document = Document(test_docx("tbl-borders"))
    context.cell = document.tables[1].cell(0, 1)


@when("I assign {style}, {width}, {color} to table.borders.{edge}")
def when_I_assign_triplet_to_table_borders_edge(
    context: Context, style: str, width: str, color: str, edge: str
):
    border = getattr(context.table_.borders, edge)
    border.style = _resolve_border_style(style)
    border.width = None if width == "None" else int(width)
    border.color = _resolve_rgb(color)


@when("I assign {style}, {width}, {color} to cell.borders.{edge}")
def when_I_assign_triplet_to_cell_borders_edge(
    context: Context, style: str, width: str, color: str, edge: str
):
    border = getattr(context.cell.borders, edge)
    border.style = _resolve_border_style(style)
    border.width = None if width == "None" else int(width)
    border.color = _resolve_rgb(color)


@when("I call table.set_borders(top=True, bottom=True, inside_h=True)")
def when_I_call_table_set_borders_subset(context: Context):
    context.table_.set_borders(top=True, bottom=True, inside_h=True)


@when("I assign None to cell.borders.{edge}.style")
def when_I_assign_none_to_cell_borders_edge_style(context: Context, edge: str):
    border = getattr(context.cell.borders, edge)
    border.style = None


@then("table.borders.{edge}.style is {value}")
def then_table_borders_edge_style_is(context: Context, edge: str, value: str):
    actual = getattr(context.table_.borders, edge).style
    expected = _resolve_border_style(value)
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("table.borders.{edge}.width is {value}")
def then_table_borders_edge_width_is(context: Context, edge: str, value: str):
    expected = None if value == "None" else int(value)
    actual = getattr(context.table_.borders, edge).width
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("table.borders.{edge}.color is {value}")
def then_table_borders_edge_color_is(context: Context, edge: str, value: str):
    expected = _resolve_rgb(value)
    actual = getattr(context.table_.borders, edge).color
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("cell.borders.{edge}.style is {value}")
def then_cell_borders_edge_style_is(context: Context, edge: str, value: str):
    actual = getattr(context.cell.borders, edge).style
    expected = _resolve_border_style(value)
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("cell.borders.{edge}.width is {value}")
def then_cell_borders_edge_width_is(context: Context, edge: str, value: str):
    expected = None if value == "None" else int(value)
    actual = getattr(context.cell.borders, edge).width
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("cell.borders.{edge}.color is {value}")
def then_cell_borders_edge_color_is(context: Context, edge: str, value: str):
    expected = _resolve_rgb(value)
    actual = getattr(context.cell.borders, edge).color
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


# ---------------------------------------------------------------------------
# Cell margins ----------------------------------------------------------------
# ---------------------------------------------------------------------------


@given("a cell having explicit margins on every edge")
def given_a_cell_having_explicit_margins(context: Context):
    document = Document(test_docx("tbl-cell-margins"))
    context.cell = document.tables[0].cell(0, 0)


@given("a cell having no explicit margins")
def given_a_cell_having_no_explicit_margins(context: Context):
    document = Document(test_docx("tbl-cell-margins"))
    context.cell = document.tables[0].cell(0, 1)


@when("I assign {value} to cell.margins.{edge}")
def when_I_assign_value_to_cell_margins_edge(context: Context, value: str, edge: str):
    new_value = None if value == "None" else int(value)
    setattr(context.cell.margins, edge, new_value)


@when("I call cell.set_margins(top={top}, end={end})")
def when_I_call_cell_set_margins(context: Context, top: str, end: str):
    kwargs = {}
    if top != "None":
        kwargs["top"] = int(top)
    if end != "None":
        kwargs["end"] = int(end)
    context.cell.set_margins(**kwargs)


@when("I call cell.remove_margins()")
def when_I_call_cell_remove_margins(context: Context):
    context.cell.remove_margins()


@then("cell.margins.{edge} is {value}")
def then_cell_margins_edge_is(context: Context, edge: str, value: str):
    expected = None if value == "None" else int(value)
    actual = getattr(context.cell.margins, edge)
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


# ---------------------------------------------------------------------------
# Cell text direction ---------------------------------------------------------
# ---------------------------------------------------------------------------


def _resolve_text_direction(name: str):
    if name == "None":
        return None
    _, member = name.split(".")
    return getattr(WD_TEXT_DIRECTION, member)


@when("I assign {value} to cell.text_direction")
def when_I_assign_value_to_cell_text_direction(context: Context, value: str):
    context.cell.text_direction = _resolve_text_direction(value)


@then("cell.text_direction is {value}")
def then_cell_text_direction_is(context: Context, value: str):
    expected = _resolve_text_direction(value)
    actual = context.cell.text_direction
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


# ---------------------------------------------------------------------------
# Cell shading ----------------------------------------------------------------
# ---------------------------------------------------------------------------


def _resolve_shading_pattern(name: str):
    if name == "None":
        return None
    _, member = name.split(".")
    return getattr(WD_SHADING_PATTERN, member)


@when("I assign {value} to cell.shading.fill_color")
def when_I_assign_value_to_cell_shading_fill_color(context: Context, value: str):
    context.cell.shading.fill_color = _resolve_rgb(value)


@when("I assign {value} to cell.shading.pattern")
def when_I_assign_value_to_cell_shading_pattern(context: Context, value: str):
    context.cell.shading.pattern = _resolve_shading_pattern(value)


@then("cell.shading.fill_color is {value}")
def then_cell_shading_fill_color_is(context: Context, value: str):
    expected = _resolve_rgb(value)
    actual = context.cell.shading.fill_color
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("cell.shading.pattern is {value}")
def then_cell_shading_pattern_is(context: Context, value: str):
    expected = _resolve_shading_pattern(value)
    actual = context.cell.shading.pattern
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


# ---------------------------------------------------------------------------
# Table style flags -----------------------------------------------------------
# ---------------------------------------------------------------------------


@given("the tbl-banded table without any tblLook flags")
def given_tbl_banded_no_flags(context: Context):
    document = Document(test_docx("tbl-banded"))
    context.table_ = document.tables[0]


@given("the tbl-banded table with only first_row set")
def given_tbl_banded_only_first_row(context: Context):
    document = Document(test_docx("tbl-banded"))
    context.table_ = document.tables[1]


@given("the tbl-banded table with banded rows active")
def given_tbl_banded_rows_active(context: Context):
    document = Document(test_docx("tbl-banded"))
    context.table_ = document.tables[2]


@given("the tbl-banded table with banded rows suppressed")
def given_tbl_banded_rows_suppressed(context: Context):
    document = Document(test_docx("tbl-banded"))
    context.table_ = document.tables[3]


@when("I assign {value} to table.style_flags.{flag}")
def when_I_assign_value_to_table_style_flags(context: Context, value: str, flag: str):
    new_value = {"True": True, "False": False}[value]
    setattr(context.table_.style_flags, flag, new_value)


@then("table.style_flags.{flag} is {value}")
def then_table_style_flags_is(context: Context, flag: str, value: str):
    expected = {"True": True, "False": False}[value]
    actual = getattr(context.table_.style_flags, flag)
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


# ---------------------------------------------------------------------------
# Autofit / allow_autofit / preferred_width ----------------------------------
# ---------------------------------------------------------------------------


def _resolve_autofit(name: str):
    if name == "None":
        return None
    _, member = name.split(".")
    return getattr(WD_TABLE_AUTOFIT, member)


@given("a freshly-created table")
def given_a_freshly_created_table(context: Context):
    context.table_ = Document().add_table(rows=2, cols=2)


@when("I assign {value} to table.autofit_behavior")
def when_I_assign_value_to_table_autofit_behavior(context: Context, value: str):
    context.table_.autofit_behavior = _resolve_autofit(value)


@when("I assign {value} to table.preferred_width")
def when_I_assign_value_to_table_preferred_width(context: Context, value: str):
    new_value = None if value == "None" else int(value)
    context.table_.preferred_width = new_value


@when("I assign {value} to table.allow_autofit")
def when_I_assign_value_to_table_allow_autofit(context: Context, value: str):
    context.table_.allow_autofit = {"True": True, "False": False}[value]


@then("table.autofit_behavior is {value}")
def then_table_autofit_behavior_is(context: Context, value: str):
    expected = _resolve_autofit(value)
    actual = context.table_.autofit_behavior
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("table.preferred_width is {value}")
def then_table_preferred_width_is(context: Context, value: str):
    expected = None if value == "None" else int(value)
    actual = context.table_.preferred_width
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("table.allow_autofit is {value}")
def then_table_allow_autofit_is(context: Context, value: str):
    expected = {"True": True, "False": False}[value]
    actual = context.table_.allow_autofit
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


@then("table.autofit is {value}")
def then_table_autofit_is(context: Context, value: str):
    expected = {"True": True, "False": False}[value]
    actual = context.table_.autofit
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


# ---------------------------------------------------------------------------
# Row properties (allow_break_across_pages, is_header) -----------------------
# ---------------------------------------------------------------------------


@given("a row in a freshly-created table")
def given_a_row_in_freshly_created_table(context: Context):
    table = Document().add_table(rows=2, cols=2)
    context.row = table.rows[0]


@when("I assign {value} to row.allow_break_across_pages")
def when_I_assign_value_to_row_allow_break_across_pages(context: Context, value: str):
    context.row.allow_break_across_pages = {"True": True, "False": False}[value]


@when("I assign {value} to row.is_header")
def when_I_assign_value_to_row_is_header(context: Context, value: str):
    context.row.is_header = {"True": True, "False": False}[value]


@then("row.allow_break_across_pages is {value}")
def then_row_allow_break_across_pages_is(context: Context, value: str):
    expected = {"True": True, "False": False}[value]
    actual = context.row.allow_break_across_pages
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


@then("row.is_header is {value}")
def then_row_is_header_is(context: Context, value: str):
    expected = {"True": True, "False": False}[value]
    actual = context.row.is_header
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


# ---------------------------------------------------------------------------
# Merge origin (raw tc access) -----------------------------------------------
# ---------------------------------------------------------------------------


_SPAN_STATE_TABLE_IDX = {
    "only uniform cells": 0,
    "a horizontal span": 1,
    "a vertical span": 2,
    "a combined span": 3,
}


@given("the raw tc at row {row}, col {col} of the {span_state} fixture table")
def given_raw_tc_at_row_col_of_fixture(
    context: Context, row: str, col: str, span_state: str
):
    document = Document(test_docx("tbl-cell-access"))
    table = document.tables[_SPAN_STATE_TABLE_IDX[span_state]]
    tr = table._tbl.tr_lst[int(row)]
    tc = tr.tc_lst[int(col)]
    context.cell = _Cell(tc, table)


@then("cell.is_merge_origin is {value}")
def then_cell_is_merge_origin_is(context: Context, value: str):
    expected = {"True": True, "False": False, "None": None}[value]
    actual = context.cell.is_merge_origin
    assert actual is expected, f"expected {expected!r}, got {actual!r}"


@then("cell.merge_origin.text is '{expected_text}'")
def then_cell_merge_origin_text_is(context: Context, expected_text: str):
    actual = context.cell.merge_origin.text
    assert actual == expected_text, f"expected {expected_text!r}, got {actual!r}"
@then("the document contains two tables")
def then_the_document_contains_two_tables(context: Context):
    assert len(context.document.tables) == 2


@then("the document contains one table")
def then_the_document_contains_one_table(context: Context):
    assert len(context.document.tables) == 1


@then("the document contains four tables")
def then_the_document_contains_four_tables(context: Context):
    assert len(context.document.tables) == 4


@then('the remaining tables contain text "{first}" and "{second}"')
def then_remaining_tables_contain_text(context: Context, first: str, second: str):
    tables = context.document.tables
    assert len(tables) == 2, f"expected 2 tables, got {len(tables)}"
    table_0_first_cell = tables[0].cell(0, 0).text
    table_1_first_cell = tables[1].cell(0, 0).text
    assert table_0_first_cell == first, f"first table cell was {table_0_first_cell!r}"
    assert table_1_first_cell == second, f"second table cell was {table_1_first_cell!r}"


@then("the inserted table has two rows and two columns")
def then_inserted_table_has_two_rows_and_two_columns(context: Context):
    table = context.table_
    assert len(table.rows) == 2
    assert len(table.columns) == 2


@then('the paragraph after the second table has text "{expected}"')
def then_paragraph_after_second_table_has_text(context: Context, expected: str):
    from docx.table import Table as TableCls
    from docx.text.paragraph import Paragraph as ParagraphCls

    second_tbl = context.document.tables[1]._tbl
    next_el = second_tbl.getnext()
    assert next_el is not None, "expected an element after the second table"
    tag = next_el.tag.rsplit("}", 1)[-1]
    assert tag == "p", f"expected a paragraph after the table, got {tag!r}"
    paragraph = ParagraphCls(next_el, None)
    actual = paragraph.text
    assert actual == expected, f"expected {expected!r}, got {actual!r}"
    # --- sanity check that TableCls is the correct proxy type (unused import guard) ---
    assert TableCls is not None


@then('the paragraph before the second table has text "{expected}"')
def then_paragraph_before_second_table_has_text(context: Context, expected: str):
    from docx.text.paragraph import Paragraph as ParagraphCls

    second_tbl = context.document.tables[1]._tbl
    prev_el = second_tbl.getprevious()
    assert prev_el is not None, "expected an element before the second table"
    tag = prev_el.tag.rsplit("}", 1)[-1]
    assert tag == "p", f"expected a paragraph before the table, got {tag!r}"
    paragraph = ParagraphCls(prev_el, None)
    actual = paragraph.text
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("each table stable_id is a 16-char hex string")
def then_each_table_stable_id_is_16_char_hex(context: Context):
    import re

    hex_re = re.compile(r"^[0-9a-f]{16}$")
    for t in context.document.tables:
        sid = t.stable_id
        assert hex_re.match(sid), f"table stable_id {sid!r} is not 16-char hex"


@then("each cell stable_id is a 16-char hex string")
def then_each_cell_stable_id_is_16_char_hex(context: Context):
    import re

    hex_re = re.compile(r"^[0-9a-f]{16}$")
    for t in context.document.tables:
        for row in t.rows:
            for cell in row.cells:
                sid = cell.stable_id
                assert hex_re.match(sid), f"cell stable_id {sid!r} is not 16-char hex"
