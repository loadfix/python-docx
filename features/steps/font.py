"""Step implementations for font-related features."""

from __future__ import annotations

from behave import given, then, when

from docx import Document
from docx.dml.color import ColorFormat
from docx.enum.dml import MSO_COLOR_TYPE, MSO_THEME_COLOR
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE
from docx.shared import RGBColor

from helpers import test_docx

# given ===================================================


@given("a font")
def given_a_font(context):
    document = Document(test_docx("txt-font-props"))
    context.font = document.paragraphs[0].runs[0].font


@given("a font having {color} highlighting")
def given_a_font_having_color_highlighting(context, color):
    paragraph_index = {
        "no": 0,
        "yellow": 1,
        "bright green": 2,
    }[color]
    document = Document(test_docx("txt-font-highlight-color"))
    context.font = document.paragraphs[paragraph_index].runs[0].font


@given("a font having {type} color")
def given_a_font_having_type_color(context, type):
    run_idx = ["no", "auto", "an RGB", "a theme"].index(type)
    document = Document(test_docx("fnt-color"))
    context.font = document.paragraphs[0].runs[run_idx].font


@given("a font having typeface name {name}")
def given_a_font_having_typeface_name(context, name):
    document = Document(test_docx("txt-font-props"))
    style_name = {
        "not specified": "Normal",
        "Avenir Black": "Having Typeface",
    }[name]
    context.font = document.styles[style_name].font


@given("a font having {underline_type} underline")
def given_a_font_having_type_underline(context, underline_type):
    style_name = {
        "inherited": "Normal",
        "no": "None Underlined",
        "single": "Underlined",
        "double": "Double Underlined",
    }[underline_type]
    document = Document(test_docx("txt-font-props"))
    context.font = document.styles[style_name].font


@given("a font having {vertAlign_state} vertical alignment")
def given_a_font_having_vertAlign_state(context, vertAlign_state):
    style_name = {
        "inherited": "Normal",
        "subscript": "Subscript",
        "superscript": "Superscript",
    }[vertAlign_state]
    document = Document(test_docx("txt-font-props"))
    context.font = document.styles[style_name].font


@given("a font of size {size}")
def given_a_font_of_size(context, size):
    document = Document(test_docx("txt-font-props"))
    style_name = {
        "unspecified": "Normal",
        "14 pt": "Having Typeface",
        "18 pt": "Large Size",
    }[size]
    context.font = document.styles[style_name].font


# when ====================================================


@when("I assign {value} to font.color.rgb")
def when_I_assign_value_to_font_color_rgb(context, value):
    font = context.font
    new_value = None if value == "None" else RGBColor.from_string(value)
    font.color.rgb = new_value


@when("I assign {value} to font.color.theme_color")
def when_I_assign_value_to_font_color_theme_color(context, value):
    font = context.font
    new_value = None if value == "None" else getattr(MSO_THEME_COLOR, value)
    font.color.theme_color = new_value


@when("I assign {value} to font.highlight_color")
def when_I_assign_value_to_font_highlight_color(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(WD_COLOR_INDEX, value)
    font.highlight_color = expected_value


@when("I assign {value} to font.name")
def when_I_assign_value_to_font_name(context, value):
    font = context.font
    value = None if value == "None" else value
    font.name = value


@when("I assign {value} to font.name_far_east")
def when_I_assign_value_to_font_name_far_east(context, value):
    font = context.font
    font.name_far_east = None if value == "None" else value


@when("I assign {value} to font.name_east_asia")
def when_I_assign_value_to_font_name_east_asia(context, value):
    font = context.font
    font.name_east_asia = None if value == "None" else value


@when("I assign {value} to font.size")
def when_I_assign_value_str_to_font_size(context, value):
    value = None if value == "None" else int(value)
    font = context.font
    font.size = value


@when("I assign {value} to font.underline")
def when_I_assign_value_to_font_underline(context, value):
    new_value = {
        "True": True,
        "False": False,
        "None": None,
        "WD_UNDERLINE.SINGLE": WD_UNDERLINE.SINGLE,
        "WD_UNDERLINE.DOUBLE": WD_UNDERLINE.DOUBLE,
    }[value]
    font = context.font
    font.underline = new_value


@when("I assign {value} to font.{sub_super}script")
def when_I_assign_value_to_font_sub_super(context, value, sub_super):
    font = context.font
    name = {
        "sub": "subscript",
        "super": "superscript",
    }[sub_super]
    new_value = {
        "None": None,
        "True": True,
        "False": False,
    }[value]

    setattr(font, name, new_value)


# then =====================================================


@then("font.color is a ColorFormat object")
def then_font_color_is_a_ColorFormat_object(context):
    font = context.font
    assert isinstance(font.color, ColorFormat)


@then("font.color.rgb is {value}")
def then_font_color_rgb_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else RGBColor.from_string(value)
    assert font.color.rgb == expected_value


@then("font.color.theme_color is {value}")
def then_font_color_theme_color_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(MSO_THEME_COLOR, value)
    assert font.color.theme_color == expected_value


@then("font.color.type is {value}")
def then_font_color_type_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(MSO_COLOR_TYPE, value)
    assert font.color.type == expected_value


@then("font.highlight_color is {value}")
def then_font_highlight_color_is_value(context, value):
    font = context.font
    expected_value = None if value == "None" else getattr(WD_COLOR_INDEX, value)
    assert font.highlight_color == expected_value


@then("font.name is {value}")
def then_font_name_is_value(context, value):
    font = context.font
    value = None if value == "None" else value
    assert font.name == value


@then("font.name_far_east is {value}")
def then_font_name_far_east_is_value(context, value):
    font = context.font
    expected = None if value == "None" else value
    assert font.name_far_east == expected, (
        f"font.name_far_east = {font.name_far_east!r}, expected {expected!r}"
    )


@then("font.name_east_asia is {value}")
def then_font_name_east_asia_is_value(context, value):
    font = context.font
    expected = None if value == "None" else value
    assert font.name_east_asia == expected, (
        f"font.name_east_asia = {font.name_east_asia!r}, expected {expected!r}"
    )


@then("font.size is {value}")
def then_font_size_is_value(context, value):
    value = None if value == "None" else int(value)
    font = context.font
    assert font.size == value


@then("font.underline is {value}")
def then_font_underline_is_value(context, value):
    expected_value = {
        "None": None,
        "True": True,
        "False": False,
        "WD_UNDERLINE.DOUBLE": WD_UNDERLINE.DOUBLE,
    }[value]
    font = context.font
    assert font.underline == expected_value


@then("font.{sub_super}script is {value}")
def then_font_sub_super_is_value(context, sub_super, value):
    name = {
        "sub": "subscript",
        "super": "superscript",
    }[sub_super]
    expected_value = {
        "None": None,
        "True": True,
        "False": False,
    }[value]
    font = context.font
    actual_value = getattr(font, name)
    assert actual_value == expected_value


# -- shared parsers ------------------------------------------------------

from docx.enum.text import (  # noqa: E402
    WD_BORDER_STYLE,
    WD_FRAME_DROP_CAP,
    WD_FRAME_H_ALIGN,
    WD_FRAME_H_ANCHOR,
    WD_FRAME_V_ALIGN,
    WD_FRAME_V_ANCHOR,
    WD_FRAME_WRAP,
)
from docx.shared import Pt, Inches  # noqa: E402


def _parse_length(spec):
    spec = spec.strip()
    if spec == "None":
        return None
    if spec.startswith("Pt("):
        return Pt(float(spec[3:-1]))
    if spec.startswith("Inches("):
        return Inches(float(spec[7:-1]))
    return int(spec)


def _parse_rgb(spec):
    spec = spec.strip()
    if spec == "None":
        return None
    return RGBColor.from_string(spec)


def _parse_border_style(spec):
    spec = spec.strip()
    if spec == "None":
        return None
    if spec.startswith("WD_BORDER_STYLE."):
        return getattr(WD_BORDER_STYLE, spec.split(".", 1)[1])
    raise ValueError(spec)


# -- run-border givens ---------------------------------------------------


@given("a run from txt-run-border paragraph {idx:d}")
def given_a_run_from_txt_run_border_paragraph(context, idx):
    document = Document(test_docx("txt-run-border"))
    context.run = document.paragraphs[idx].runs[0]
    context.font = context.run.font


# -- run-border whens ----------------------------------------------------


@when("I set font.border_style to {value}")
def when_I_set_font_border_style(context, value):
    context.font.border_style = _parse_border_style(value)


@when("I set font.border_color to {value}")
def when_I_set_font_border_color(context, value):
    context.font.border_color = _parse_rgb(value)


@when("I set font.border_width to {value}")
def when_I_set_font_border_width(context, value):
    context.font.border_width = _parse_length(value)


@when("I set font.border_space to {value}")
def when_I_set_font_border_space(context, value):
    context.font.border_space = _parse_length(value)


@when("I call font.remove_border()")
def when_I_call_font_remove_border(context):
    context.font.remove_border()


# -- run-border thens ----------------------------------------------------


@then("font.border_style is {value}")
def then_font_border_style_is(context, value):
    expected = _parse_border_style(value)
    assert context.font.border_style == expected, (
        f"got {context.font.border_style!r}, expected {expected!r}"
    )


@then("font.border_color is {value}")
def then_font_border_color_is(context, value):
    expected = _parse_rgb(value)
    assert context.font.border_color == expected, (
        f"got {context.font.border_color!r}, expected {expected!r}"
    )


@then("font.border_width is {value}")
def then_font_border_width_is(context, value):
    expected = _parse_length(value)
    assert context.font.border_width == expected, (
        f"got {context.font.border_width!r}, expected {expected!r}"
    )


@then("font.border_space is {value}")
def then_font_border_space_is(context, value):
    expected = _parse_length(value)
    assert context.font.border_space == expected, (
        f"got {context.font.border_space!r}, expected {expected!r}"
    )


# -- east-asian givens ---------------------------------------------------


@given("a run from txt-east-asian paragraph {idx:d}")
def given_a_run_from_txt_east_asian_paragraph(context, idx):
    document = Document(test_docx("txt-east-asian"))
    context.document = document
    context.paragraph = document.paragraphs[idx]
    context.run = context.paragraph.runs[0]
    context.font = context.run.font


@given("a paragraph format from txt-east-asian paragraph {idx:d}")
def given_a_paragraph_format_from_txt_east_asian_paragraph(context, idx):
    document = Document(test_docx("txt-east-asian"))
    context.document = document
    context.paragraph = document.paragraphs[idx]
    context.paragraph_format = context.paragraph.paragraph_format


# -- east-asian whens ----------------------------------------------------


def _parse_kwargs(spec):
    """Parse 'key1=value1 key2=value2' into a dict with Python-typed values."""
    result = {}
    for pair in spec.split():
        key, _, val = pair.partition("=")
        if val in ("True", "False"):
            result[key] = val == "True"
        elif val == "None":
            result[key] = None
        elif val.startswith("Pt("):
            result[key] = Pt(float(val[3:-1]))
        elif val.startswith("Inches("):
            result[key] = Inches(float(val[7:-1]))
        else:
            try:
                result[key] = int(val)
            except ValueError:
                result[key] = val
    return result


@when("I call font.set_east_asian_layout {kwargs}")
def when_I_call_font_set_east_asian_layout(context, kwargs):
    parsed = _parse_kwargs(kwargs)
    # -- remap bool-typed kwargs that also accept None --
    context.font.set_east_asian_layout(**parsed)


@when("I call font.remove_east_asian_layout()")
def when_I_call_font_remove_east_asian_layout(context):
    context.font.remove_east_asian_layout()


# -- east-asian thens ----------------------------------------------------


@then("the run has no east-asian layout")
def then_run_has_no_east_asian_layout(context):
    layout = context.font.east_asian_layout
    assert layout is None, f"expected None, got {layout!r}"


@then("the run has an east-asian layout")
def then_run_has_an_east_asian_layout(context):
    layout = context.font.east_asian_layout
    assert layout is not None, "expected a layout proxy, got None"


@then("font.east_asian_layout.two_lines_in_one is {value}")
def then_font_east_asian_layout_two_lines(context, value):
    expected = {"None": None, "True": True, "False": False}[value]
    assert context.font.east_asian_layout.two_lines_in_one == expected


@then("font.east_asian_layout.vertical_alignment is {value}")
def then_font_east_asian_layout_vert(context, value):
    expected = {"None": None, "True": True, "False": False}[value]
    assert context.font.east_asian_layout.vertical_alignment == expected


@then("font.east_asian_layout.compressed is {value}")
def then_font_east_asian_layout_compressed(context, value):
    expected = {"None": None, "True": True, "False": False}[value]
    assert context.font.east_asian_layout.compressed == expected


@then("font.east_asian_layout.id is {value:d}")
def then_font_east_asian_layout_id_is(context, value):
    assert context.font.east_asian_layout.id == value


# -- generic font.<property> givens/whens/thens for kerning, spacing,
# -- language, name_far_east, etc. ---------------------------------------


@when("I assign {value} to font.kerning")
def when_I_assign_value_to_font_kerning(context, value):
    context.font.kerning = _parse_length(value)


@when("I assign {value} to font.character_spacing")
def when_I_assign_value_to_font_character_spacing(context, value):
    context.font.character_spacing = _parse_length(value)


@when("I assign {value} to font.right_to_left")
def when_I_assign_value_to_font_right_to_left(context, value):
    val = {"True": True, "False": False, "None": None}[value]
    context.font.right_to_left = val


@when("I assign {value} to font.language")
def when_I_assign_value_to_font_language(context, value):
    context.font.language = None if value == "None" else value


@when("I assign {value} to font.east_asian_language")
def when_I_assign_value_to_font_east_asian_language(context, value):
    context.font.east_asian_language = None if value == "None" else value


@when("I assign {value} to font.bidi_language")
def when_I_assign_value_to_font_bidi_language(context, value):
    context.font.bidi_language = None if value == "None" else value


@when("I call font.remove_language()")
def when_I_call_font_remove_language(context):
    context.font.remove_language()


@when("I assign {value} to font.name_far_east")
def when_I_assign_value_to_font_name_far_east(context, value):
    context.font.name_far_east = None if value == "None" else value


@when("I assign {value} to font.name_east_asia")
def when_I_assign_value_to_font_name_east_asia(context, value):
    context.font.name_east_asia = None if value == "None" else value


@when("I assign {value} to font.shading_color")
def when_I_assign_value_to_font_shading_color(context, value):
    context.font.shading_color = _parse_rgb(value)


@then("font.kerning is {value}")
def then_font_kerning_is(context, value):
    expected = _parse_length(value)
    assert context.font.kerning == expected, (
        f"got {context.font.kerning!r}, expected {expected!r}"
    )


@then("font.character_spacing is {value}")
def then_font_character_spacing_is(context, value):
    expected = _parse_length(value)
    assert context.font.character_spacing == expected, (
        f"got {context.font.character_spacing!r}, expected {expected!r}"
    )


@then("font.right_to_left is {value}")
def then_font_right_to_left_is(context, value):
    expected = {"True": True, "False": False}[value]
    assert context.font.right_to_left is expected


@then("font.language is {value}")
def then_font_language_is(context, value):
    expected = None if value == "None" else value
    assert context.font.language == expected


@then("font.east_asian_language is {value}")
def then_font_east_asian_language_is(context, value):
    expected = None if value == "None" else value
    assert context.font.east_asian_language == expected


@then("font.bidi_language is {value}")
def then_font_bidi_language_is(context, value):
    expected = None if value == "None" else value
    assert context.font.bidi_language == expected


@then("font.name_far_east is {value}")
def then_font_name_far_east_is(context, value):
    expected = None if value == "None" else value
    assert context.font.name_far_east == expected


@then("font.name_east_asia is {value}")
def then_font_name_east_asia_is(context, value):
    expected = None if value == "None" else value
    assert context.font.name_east_asia == expected


@then("font.shading_color is {value}")
def then_font_shading_color_is(context, value):
    expected = _parse_rgb(value)
    assert context.font.shading_color == expected, (
        f"got {context.font.shading_color!r}, expected {expected!r}"
    )


# -- ruby givens/thens ---------------------------------------------------


@given("a run from txt-ruby run {run_idx:d}")
def given_a_run_from_txt_ruby_run(context, run_idx):
    document = Document(test_docx("txt-ruby"))
    context.document = document
    context.run = document.paragraphs[0].runs[run_idx]


@given("the ruby annotation at run {run_idx:d} position {pos:d} in txt-ruby")
def given_the_ruby_annotation(context, run_idx, pos):
    document = Document(test_docx("txt-ruby"))
    context.document = document
    rubies = document.paragraphs[0].runs[run_idx].ruby_annotations
    context.ruby = rubies[pos]


@then("len(run.ruby_annotations) is {count:d}")
def then_len_run_ruby_annotations_is(context, count):
    actual = len(context.run.ruby_annotations)
    assert actual == count, f"got {actual}, expected {count}"


@then("ruby.base_text is {text}")
def then_ruby_base_text_is(context, text):
    expected = "" if text == "<empty>" else text
    assert context.ruby.base_text == expected, (
        f"got {context.ruby.base_text!r}, expected {expected!r}"
    )


@then("ruby.ruby_text is {text}")
def then_ruby_ruby_text_is(context, text):
    expected = "" if text == "<empty>" else text
    assert context.ruby.ruby_text == expected, (
        f"got {context.ruby.ruby_text!r}, expected {expected!r}"
    )


@then("ruby.alignment is {value}")
def then_ruby_alignment_is(context, value):
    expected = None if value == "None" else value
    assert context.ruby.alignment == expected


@then("ruby.language is {value}")
def then_ruby_language_is(context, value):
    expected = None if value == "None" else value
    assert context.ruby.language == expected


@then("run.text contains the ruby base strings")
def then_run_text_contains_the_ruby_base_strings(context):
    # -- run 0 in txt-ruby holds two rubies whose base text is "日本" and "東京" --
    assert context.run.text == "日本東京", (
        f"got {context.run.text!r}"
    )
