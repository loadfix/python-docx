"""Step implementations for hyperlink-related features."""

from __future__ import annotations

from typing import Dict, Tuple

from behave import given, then, when
from behave.runner import Context

from docx import Document

from helpers import test_docx

# given ===================================================


@given("a hyperlink")
def given_a_hyperlink(context: Context):
    document = Document(test_docx("par-hyperlinks"))
    context.hyperlink = document.paragraphs[1].hyperlinks[0]


@given("a hyperlink having a URI fragment")
def given_a_hyperlink_having_a_uri_fragment(context: Context):
    document = Document(test_docx("par-hlink-frags"))
    context.hyperlink = document.paragraphs[1].hyperlinks[0]


@given("a hyperlink having address {address} and fragment {fragment}")
def given_a_hyperlink_having_address_and_fragment(context: Context, address: str, fragment: str):
    paragraph_idxs: Dict[Tuple[str, str], int] = {
        ("''", "linkedBookmark"): 1,
        ("https://foo.com", "''"): 2,
        ("https://foo.com?q=bar", "''"): 3,
        ("http://foo.com/", "intro"): 4,
        ("https://foo.com?q=bar#baz", "''"): 5,
        ("court-exif.jpg", "''"): 7,
    }
    paragraph_idx = paragraph_idxs[(address, fragment)]
    document = Document(test_docx("par-hlink-frags"))
    paragraph = document.paragraphs[paragraph_idx]
    context.hyperlink = paragraph.hyperlinks[0]


@given("a hyperlink having {zero_or_more} rendered page breaks")
def given_a_hyperlink_having_rendered_page_breaks(context: Context, zero_or_more: str):
    paragraph_idx = {
        "no": 1,
        "one": 2,
    }[zero_or_more]
    document = Document(test_docx("par-hyperlinks"))
    paragraph = document.paragraphs[paragraph_idx]
    context.hyperlink = paragraph.hyperlinks[0]


@given("a hyperlink having {one_or_more} runs")
def given_a_hyperlink_having_one_or_more_runs(context: Context, one_or_more: str):
    paragraph_idx, hyperlink_idx = {
        "one": (1, 0),
        "two": (2, 1),
    }[one_or_more]
    document = Document(test_docx("par-hyperlinks"))
    paragraph = document.paragraphs[paragraph_idx]
    context.hyperlink = paragraph.hyperlinks[hyperlink_idx]


# then =====================================================


@then("hyperlink.address is the URL of the hyperlink")
def then_hyperlink_address_is_the_URL_of_the_hyperlink(context: Context):
    actual_value = context.hyperlink.address
    expected_value = "http://yahoo.com/"
    assert actual_value == expected_value, f"expected: {expected_value}, got: {actual_value}"


@then("hyperlink.contains_page_break is {value}")
def then_hyperlink_contains_page_break_is_value(context: Context, value: str):
    actual_value = context.hyperlink.contains_page_break
    expected_value = {"True": True, "False": False}[value]
    assert actual_value == expected_value, f"expected: {expected_value}, got: {actual_value}"


@then("hyperlink.fragment is the URI fragment of the hyperlink")
def then_hyperlink_fragment_is_the_URI_fragment_of_the_hyperlink(context: Context):
    actual_value = context.hyperlink.fragment
    expected_value = "linkedBookmark"
    assert actual_value == expected_value, f"expected: {expected_value}, got: {actual_value}"


@then("hyperlink.runs contains only Run instances")
def then_hyperlink_runs_contains_only_Run_instances(context: Context):
    actual_value = [type(item).__name__ for item in context.hyperlink.runs]
    expected_value = ["Run" for _ in context.hyperlink.runs]
    assert actual_value == expected_value, f"expected: {expected_value}, got: {actual_value}"


@then("hyperlink.runs has length {value}")
def then_hyperlink_runs_has_length(context: Context, value: str):
    actual_value = len(context.hyperlink.runs)
    expected_value = int(value)
    assert actual_value == expected_value, f"expected: {expected_value}, got: {actual_value}"


@then("hyperlink.text is the visible text of the hyperlink")
def then_hyperlink_text_is_the_visible_text_of_the_hyperlink(context: Context):
    actual_value = context.hyperlink.text
    expected_value = "awesome hyperlink"
    assert actual_value == expected_value, f"expected: {expected_value}, got: {actual_value}"


@then("hyperlink.url is {value}")
def then_hyperlink_url_is_value(context: Context, value: str):
    actual_value = context.hyperlink.url
    expected_value = "" if value == "''" else value
    assert actual_value == expected_value, f"expected: {expected_value}, got: {actual_value}"


# -- hyperlink-create steps --------------------------------------------------


@given("a fresh paragraph in a default document")
def given_a_fresh_paragraph_in_a_default_document(context: Context):
    context.document = Document()
    context.paragraph = context.document.add_paragraph("")


@when('I call paragraph.add_hyperlink(url="{url}", text="{text}")')
def when_call_paragraph_add_hyperlink_url_text(
    context: Context, url: str, text: str
):
    # -- pass style=None because the stock default template does not define
    # -- a "Hyperlink" character style. --
    context.hyperlink = context.paragraph.add_hyperlink(
        url=url, text=text, style=None
    )


@when('I call paragraph.add_hyperlink(url="{url}")')
def when_call_paragraph_add_hyperlink_url(context: Context, url: str):
    context.hyperlink = context.paragraph.add_hyperlink(url=url, style=None)


@when('I call paragraph.add_hyperlink(anchor="{anchor}", text="{text}")')
def when_call_paragraph_add_hyperlink_anchor_text(
    context: Context, anchor: str, text: str
):
    context.hyperlink = context.paragraph.add_hyperlink(
        anchor=anchor, text=text, style=None
    )


@then('the returned hyperlink.address is "{value}"')
def then_returned_hyperlink_address(context: Context, value: str):
    actual = context.hyperlink.address
    assert actual == value, f"expected {value!r}, got {actual!r}"


@then('the returned hyperlink.fragment is "{value}"')
def then_returned_hyperlink_fragment(context: Context, value: str):
    actual = context.hyperlink.fragment
    assert actual == value, f"expected {value!r}, got {actual!r}"


@then('the returned hyperlink.text is "{value}"')
def then_returned_hyperlink_text(context: Context, value: str):
    actual = context.hyperlink.text
    assert actual == value, f"expected {value!r}, got {actual!r}"


@then("calling paragraph.add_hyperlink() raises ValueError")
def then_paragraph_add_hyperlink_raises_valueerror(context: Context):
    try:
        context.paragraph.add_hyperlink()
    except ValueError:
        return
    raise AssertionError("expected ValueError")


@then('calling paragraph.add_hyperlink(url="{url}", anchor="{anchor}") raises ValueError')
def then_paragraph_add_hyperlink_url_anchor_raises(
    context: Context, url: str, anchor: str
):
    try:
        context.paragraph.add_hyperlink(url=url, anchor=anchor)
    except ValueError:
        return
    raise AssertionError("expected ValueError")
