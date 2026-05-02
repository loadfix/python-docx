"""Step implementations for legacy form-field-related features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.form_fields import (
    CheckboxFormField,
    DropdownFormField,
    FormField,
    TextInputFormField,
    WD_FORM_FIELD_TYPE,
)

from helpers import saved_docx_path, test_docx


# given ====================================================


@given("a document having 3 legacy form fields")
def given_a_document_having_3_legacy_form_fields(context: Context):
    context.document = Document(test_docx("frm-has-form-fields"))


@given("a freshly created document")
def given_a_freshly_created_document(context: Context):
    context.document = Document()


# when =====================================================


@when("I select the text-input form field")
def when_I_select_the_text_input_form_field(context: Context):
    context.form_field = next(
        ff
        for ff in context.document.form_fields
        if ff.type is WD_FORM_FIELD_TYPE.TEXT
    )


@when("I select the checkbox form field")
def when_I_select_the_checkbox_form_field(context: Context):
    context.form_field = next(
        ff
        for ff in context.document.form_fields
        if ff.type is WD_FORM_FIELD_TYPE.CHECKBOX
    )


@when("I select the dropdown form field")
def when_I_select_the_dropdown_form_field(context: Context):
    context.form_field = next(
        ff
        for ff in context.document.form_fields
        if ff.type is WD_FORM_FIELD_TYPE.DROPDOWN
    )


@when('I append a text form field named "{name}" with default "{default}"')
def when_I_append_a_text_form_field(context: Context, name: str, default: str):
    paragraph = context.document.add_paragraph()
    paragraph.add_text_form_field(name=name, default=default, maxlength=50)


@when('I append a checkbox form field named "{name}" checked {checked}')
def when_I_append_a_checkbox_form_field(context: Context, name: str, checked: str):
    paragraph = context.document.add_paragraph()
    paragraph.add_checkbox_form_field(name=name, checked=(checked == "True"))


@when(
    'I append a dropdown form field named "{name}" with options '
    '["{o1}", "{o2}", "{o3}"] default {default:d}'
)
def when_I_append_a_dropdown_form_field(
    context: Context, name: str, o1: str, o2: str, o3: str, default: int
):
    paragraph = context.document.add_paragraph()
    paragraph.add_dropdown_form_field(
        name=name, options=[o1, o2, o3], default_index=default
    )


@when("I save and re-open the document")
def when_I_save_and_reopen_the_document(context: Context):
    context.document.save(saved_docx_path)
    context.document = Document(saved_docx_path)


# then =====================================================


@then("document.form_fields is a list of {count:d} FormField objects")
def then_document_form_fields_is_a_list_of_count_form_fields(
    context: Context, count: int
):
    fields = context.document.form_fields
    assert isinstance(fields, list), f"expected a list, got {type(fields).__name__}"
    assert len(fields) == count, f"expected {count} form fields, got {len(fields)}"
    for ff in fields:
        assert isinstance(ff, FormField), (
            f"expected FormField, got {type(ff).__name__}"
        )


@then("the form fields are returned in document order")
def then_the_form_fields_are_returned_in_document_order(context: Context):
    types = [ff.type for ff in context.document.form_fields]
    expected = [
        WD_FORM_FIELD_TYPE.TEXT,
        WD_FORM_FIELD_TYPE.CHECKBOX,
        WD_FORM_FIELD_TYPE.DROPDOWN,
    ]
    assert types == expected, f"expected {expected}, got {types}"


@then("form_field.type is WD_FORM_FIELD_TYPE.{member}")
def then_form_field_type_is(context: Context, member: str):
    expected = getattr(WD_FORM_FIELD_TYPE, member)
    actual = context.form_field.type
    assert actual is expected, f"expected {expected}, got {actual}"


@then('form_field.name == "{name}"')
def then_form_field_name_eq(context: Context, name: str):
    actual = context.form_field.name
    assert actual == name, f"expected name '{name}', got '{actual}'"


@then("form_field.enabled is True")
def then_form_field_enabled_is_true(context: Context):
    assert context.form_field.enabled is True


@then('form_field.text_input.default == "{default}"')
def then_form_field_text_input_default_eq(context: Context, default: str):
    ti = context.form_field.text_input
    assert isinstance(ti, TextInputFormField), (
        f"expected TextInputFormField, got {type(ti).__name__}"
    )
    assert ti.default == default, f"expected default '{default}', got '{ti.default}'"


@then("form_field.text_input.max_length == {max_length:d}")
def then_form_field_text_input_max_length_eq(context: Context, max_length: int):
    ti = context.form_field.text_input
    assert ti is not None
    assert ti.max_length == max_length, (
        f"expected max_length {max_length}, got {ti.max_length}"
    )


@then("form_field.text_input is None")
def then_form_field_text_input_is_none(context: Context):
    assert context.form_field.text_input is None


@then("form_field.checkbox.checked is True")
def then_form_field_checkbox_checked_is_true(context: Context):
    cb = context.form_field.checkbox
    assert isinstance(cb, CheckboxFormField), (
        f"expected CheckboxFormField, got {type(cb).__name__}"
    )
    assert cb.checked is True, f"expected checked True, got {cb.checked}"


@then("form_field.checkbox is None")
def then_form_field_checkbox_is_none(context: Context):
    assert context.form_field.checkbox is None


@then('form_field.dropdown.options == ["{o1}", "{o2}", "{o3}"]')
def then_form_field_dropdown_options_eq(
    context: Context, o1: str, o2: str, o3: str
):
    dd = context.form_field.dropdown
    assert isinstance(dd, DropdownFormField), (
        f"expected DropdownFormField, got {type(dd).__name__}"
    )
    expected = [o1, o2, o3]
    assert dd.options == expected, f"expected options {expected}, got {dd.options}"


@then("form_field.dropdown.default_index == {idx:d}")
def then_form_field_dropdown_default_index_eq(context: Context, idx: int):
    dd = context.form_field.dropdown
    assert dd is not None
    assert dd.default_index == idx, (
        f"expected default_index {idx}, got {dd.default_index}"
    )


@then("form_field.dropdown is None")
def then_form_field_dropdown_is_none(context: Context):
    assert context.form_field.dropdown is None


@then('form_field.value == "{value}"')
def then_form_field_value_eq_str(context: Context, value: str):
    actual = context.form_field.value
    assert actual == value, f"expected value '{value}', got '{actual!r}'"


@then("form_field.value is True")
def then_form_field_value_is_true(context: Context):
    assert context.form_field.value is True


@then('the text form field\'s default is "{default}"')
def then_the_text_form_field_default_is(context: Context, default: str):
    text_ff = next(
        ff
        for ff in context.document.form_fields
        if ff.type is WD_FORM_FIELD_TYPE.TEXT
    )
    assert text_ff.text_input is not None
    assert text_ff.text_input.default == default, (
        f"expected default '{default}', got '{text_ff.text_input.default}'"
    )


@then("the checkbox form field is unchecked")
def then_the_checkbox_form_field_is_unchecked(context: Context):
    cb_ff = next(
        ff
        for ff in context.document.form_fields
        if ff.type is WD_FORM_FIELD_TYPE.CHECKBOX
    )
    assert cb_ff.checkbox is not None
    assert cb_ff.checkbox.checked is False, (
        f"expected checked False, got {cb_ff.checkbox.checked}"
    )


@then('the dropdown form field\'s selected value is "{value}"')
def then_the_dropdown_form_field_selected_value_is(context: Context, value: str):
    dd_ff = next(
        ff
        for ff in context.document.form_fields
        if ff.type is WD_FORM_FIELD_TYPE.DROPDOWN
    )
    assert dd_ff.value == value, f"expected value '{value}', got '{dd_ff.value!r}'"
