"""Step implementations for section-related features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.section import (
    WD_BORDER_DISPLAY,
    WD_BORDER_OFFSET_FROM,
    WD_DOC_GRID_TYPE,
    WD_LINE_NUMBERING_RESTART,
    WD_ORIENT,
    WD_SECTION,
)
from docx.enum.table import WD_TEXT_DIRECTION
from docx.enum.text import WD_BORDER_STYLE
from docx.oxml.ns import qn
from docx.section import Column, DocumentGrid, LineNumbering, PageBorder, PageBorders, Section, SectionColumns
from docx.shared import Inches, Pt, RGBColor

from helpers import test_docx

# given ====================================================


@given("a Section object as section")
def given_a_Section_object_as_section(context: Context):
    context.section = Document(test_docx("sct-section-props")).sections[-1]


@given("a Section object of a multi-section document as section")
def given_a_Section_object_of_a_multi_section_document_as_section(context: Context):
    context.section = Document(test_docx("sct-inner-content")).sections[1]


@given("a Section object {with_or_without} a distinct first-page header as section")
def given_a_Section_object_with_or_without_first_page_header(
    context: Context, with_or_without: str
):
    section_idx = {"with": 1, "without": 0}[with_or_without]
    context.section = Document(test_docx("sct-first-page-hdrftr")).sections[section_idx]


@given("a section collection containing 3 sections")
def given_a_section_collection_containing_3_sections(context: Context):
    document = Document(test_docx("doc-access-sections"))
    context.sections = document.sections


@given("a section having known page dimension")
def given_a_section_having_known_page_dimension(context: Context):
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[-1]


@given("a section having known page margins")
def given_a_section_having_known_page_margins(context: Context):
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[0]


@given("a section having start type {start_type}")
def given_a_section_having_start_type(context: Context, start_type: str):
    section_idx = {
        "CONTINUOUS": 0,
        "NEW_PAGE": 1,
        "ODD_PAGE": 2,
        "EVEN_PAGE": 3,
        "NEW_COLUMN": 4,
    }[start_type]
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[section_idx]


@given("a section known to have {orientation} orientation")
def given_a_section_having_known_orientation(context: Context, orientation: str):
    section_idx = {"landscape": 0, "portrait": 1}[orientation]
    document = Document(test_docx("sct-section-props"))
    context.section = document.sections[section_idx]


# when =====================================================


@when("I assign {bool_val} to section.different_first_page_header_footer")
def when_I_assign_value_to_section_different_first_page_hdrftr(context: Context, bool_val: str):
    context.section.different_first_page_header_footer = eval(bool_val)


@when("I set the {margin_side} margin to {inches} inches")
def when_I_set_the_margin_side_length(context: Context, margin_side: str, inches: str):
    prop_name = {
        "left": "left_margin",
        "right": "right_margin",
        "top": "top_margin",
        "bottom": "bottom_margin",
        "gutter": "gutter",
        "header": "header_distance",
        "footer": "footer_distance",
    }[margin_side]
    new_value = Inches(float(inches))
    setattr(context.section, prop_name, new_value)


@when("I set the section orientation to {orientation}")
def when_I_set_the_section_orientation(context: Context, orientation: str):
    new_orientation = {
        "WD_ORIENT.PORTRAIT": WD_ORIENT.PORTRAIT,
        "WD_ORIENT.LANDSCAPE": WD_ORIENT.LANDSCAPE,
        "None": None,
    }[orientation]
    context.section.orientation = new_orientation


@when("I set the section page height to {y} inches")
def when_I_set_the_section_page_height_to_y_inches(context: Context, y: str):
    context.section.page_height = Inches(float(y))


@when("I set the section page width to {x} inches")
def when_I_set_the_section_page_width_to_x_inches(context: Context, x: str):
    context.section.page_width = Inches(float(x))


@when("I set the section start type to {start_type}")
def when_I_set_the_section_start_type_to_start_type(context: Context, start_type: str):
    new_start_type = {
        "None": None,
        "CONTINUOUS": WD_SECTION.CONTINUOUS,
        "EVEN_PAGE": WD_SECTION.EVEN_PAGE,
        "NEW_COLUMN": WD_SECTION.NEW_COLUMN,
        "NEW_PAGE": WD_SECTION.NEW_PAGE,
        "ODD_PAGE": WD_SECTION.ODD_PAGE,
    }[start_type]
    context.section.start_type = new_start_type


# then =====================================================


@then("I can access a section by index")
def then_I_can_access_a_section_by_index(context: Context):
    sections = context.sections
    for idx in range(3):
        section = sections[idx]
        assert isinstance(section, Section)


@then("I can iterate over the sections")
def then_I_can_iterate_over_the_sections(context: Context):
    sections = context.sections
    actual_count = 0
    for section in sections:
        actual_count += 1
        assert isinstance(section, Section)
    assert actual_count == 3


@then("len(sections) is 3")
def then_len_sections_is_3(context: Context):
    sections = context.sections
    assert len(sections) == 3, "expected len(sections) of 3, got %s" % len(sections)


@then("section.different_first_page_header_footer is {bool_val}")
def then_section_different_first_page_header_footer_is(context: Context, bool_val: str):
    actual = context.section.different_first_page_header_footer
    expected = eval(bool_val)
    assert actual == expected, "section.different_first_page_header_footer is %s" % actual


@then("section.even_page_footer is a _Footer object")
def then_section_even_page_footer_is_a_Footer_object(context: Context):
    actual = type(context.section.even_page_footer).__name__
    expected = "_Footer"
    assert actual == expected, "section.even_page_footer is a %s object" % actual


@then("section.even_page_header is a _Header object")
def then_section_even_page_header_is_a_Header_object(context: Context):
    actual = type(context.section.even_page_header).__name__
    expected = "_Header"
    assert actual == expected, "section.even_page_header is a %s object" % actual


@then("section.first_page_footer is a _Footer object")
def then_section_first_page_footer_is_a_Footer_object(context: Context):
    actual = type(context.section.first_page_footer).__name__
    expected = "_Footer"
    assert actual == expected, "section.first_page_footer is a %s object" % actual


@then("section.first_page_header is a _Header object")
def then_section_first_page_header_is_a_Header_object(context: Context):
    actual = type(context.section.first_page_header).__name__
    expected = "_Header"
    assert actual == expected, "section.first_page_header is a %s object" % actual


@then("section.footer is a _Footer object")
def then_section_footer_is_a_Footer_object(context: Context):
    actual = type(context.section.footer).__name__
    expected = "_Footer"
    assert actual == expected, "section.footer is a %s object" % actual


@then("section.header is a _Header object")
def then_section_header_is_a_Header_object(context: Context):
    actual = type(context.section.header).__name__
    expected = "_Header"
    assert actual == expected, "section.header is a %s object" % actual


@then("section.iter_inner_content() produces the paragraphs and tables in section")
def step_impl(context: Context):
    actual = [type(item).__name__ for item in context.section.iter_inner_content()]
    expected = ["Table", "Paragraph", "Paragraph"]
    assert actual == expected, f"expected: {expected}, got: {actual}"


@then("section.{propname}.is_linked_to_previous is True")
def then_section_hdrftr_prop_is_linked_to_previous_is_True(context: Context, propname: str):
    actual = getattr(context.section, propname).is_linked_to_previous
    expected = True
    assert actual == expected, "section.%s.is_linked_to_previous is %s" % (
        propname,
        actual,
    )


@then("the reported {margin_side} margin is {inches} inches")
def then_the_reported_margin_is_inches(context: Context, margin_side: str, inches: str):
    prop_name = {
        "left": "left_margin",
        "right": "right_margin",
        "top": "top_margin",
        "bottom": "bottom_margin",
        "gutter": "gutter",
        "header": "header_distance",
        "footer": "footer_distance",
    }[margin_side]
    expected_value = Inches(float(inches))
    actual_value = getattr(context.section, prop_name)
    assert actual_value == expected_value


@then("the reported page orientation is {orientation}")
def then_the_reported_page_orientation_is_orientation(context: Context, orientation: str):
    expected_value = {
        "WD_ORIENT.LANDSCAPE": WD_ORIENT.LANDSCAPE,
        "WD_ORIENT.PORTRAIT": WD_ORIENT.PORTRAIT,
    }[orientation]
    assert context.section.orientation == expected_value


@then("the reported page width is {x} inches")
def then_the_reported_page_width_is_width(context: Context, x: str):
    assert context.section.page_width == Inches(float(x))


@then("the reported page height is {y} inches")
def then_the_reported_page_height_is_11_inches(context: Context, y: str):
    assert context.section.page_height == Inches(float(y))


@then("the reported section start type is {start_type}")
def then_the_reported_section_start_type_is_type(context: Context, start_type: str):
    expected_start_type = {
        "CONTINUOUS": WD_SECTION.CONTINUOUS,
        "EVEN_PAGE": WD_SECTION.EVEN_PAGE,
        "NEW_COLUMN": WD_SECTION.NEW_COLUMN,
        "NEW_PAGE": WD_SECTION.NEW_PAGE,
        "ODD_PAGE": WD_SECTION.ODD_PAGE,
    }[start_type]
    assert context.section.start_type == expected_start_type


# ==========================================================================
# Page borders  (sct-page-borders.feature)
# ==========================================================================

_BORDER_STYLE_MAP = {
    "SINGLE": WD_BORDER_STYLE.SINGLE,
    "THICK": WD_BORDER_STYLE.THICK,
    "DOUBLE": WD_BORDER_STYLE.DOUBLE,
    "DOTTED": WD_BORDER_STYLE.DOTTED,
}
_BORDER_DISPLAY_MAP = {
    "ALL_PAGES": WD_BORDER_DISPLAY.ALL_PAGES,
    "FIRST_PAGE": WD_BORDER_DISPLAY.FIRST_PAGE,
    "NOT_FIRST_PAGE": WD_BORDER_DISPLAY.NOT_FIRST_PAGE,
}
_BORDER_OFFSET_FROM_MAP = {
    "TEXT": WD_BORDER_OFFSET_FROM.TEXT,
    "PAGE": WD_BORDER_OFFSET_FROM.PAGE,
}


@given("a Section with no page borders as section")
def given_a_Section_with_no_page_borders_as_section(context: Context):
    context.section = Document(test_docx("sct-page-borders")).sections[0]


@given("a Section with all four borders set as section")
def given_a_Section_with_all_four_borders_set_as_section(context: Context):
    context.section = Document(test_docx("sct-page-borders")).sections[1]


@given("a Section with only a top border set as section")
def given_a_Section_with_only_a_top_border_set_as_section(context: Context):
    context.section = Document(test_docx("sct-page-borders")).sections[2]


@when(
    "I call section.set_page_border on {side} with SINGLE 1pt black 12pt"
)
def when_I_call_set_page_border_default_args(context: Context, side: str):
    context.section.set_page_border(
        side,
        style=WD_BORDER_STYLE.SINGLE,
        width=Pt(1),
        color=RGBColor(0, 0, 0),
        space=Pt(12),
    )


@when("I call section.set_page_border on {side} with space {pts:d} pt")
def when_I_call_set_page_border_space_only(context: Context, side: str, pts: int):
    context.section.set_page_border(side, space=Pt(pts))


@when("I clear section.page_borders.{side}.{attr}")
def when_I_clear_section_page_borders_edge_attr(
    context: Context, side: str, attr: str
):
    setattr(getattr(context.section.page_borders, side), attr, None)


@when("I assign {value} to section.page_borders.display")
def when_I_assign_to_page_borders_display(context: Context, value: str):
    context.section.page_borders.display = _BORDER_DISPLAY_MAP[value]


@when("I assign {value} to section.page_borders.offset_from")
def when_I_assign_to_page_borders_offset_from(context: Context, value: str):
    context.section.page_borders.offset_from = _BORDER_OFFSET_FROM_MAP[value]


@when("I call section.remove_page_borders()")
def when_I_call_section_remove_page_borders(context: Context):
    context.section.remove_page_borders()


@then("section.page_borders is a PageBorders object")
def then_section_page_borders_is_a_PageBorders_object(context: Context):
    assert isinstance(context.section.page_borders, PageBorders)


@then("section.page_borders.{side} is a PageBorder object")
def then_section_page_borders_side_is_a_PageBorder_object(context: Context, side: str):
    assert isinstance(getattr(context.section.page_borders, side), PageBorder)


@then("section.page_borders.{side}.{attr} is None")
def then_section_page_borders_side_attr_is_None(
    context: Context, side: str, attr: str
):
    actual = getattr(getattr(context.section.page_borders, side), attr)
    assert actual is None, f"page_borders.{side}.{attr} is {actual!r}"


@then("section.page_borders.{side}.style is {style}")
def then_section_page_borders_side_style_is(context: Context, side: str, style: str):
    actual = getattr(context.section.page_borders, side).style
    expected = _BORDER_STYLE_MAP[style]
    assert actual == expected, f"page_borders.{side}.style is {actual!r}"


@then("section.page_borders.{side}.width is {pts:d} pt")
def then_section_page_borders_side_width_is_pt(
    context: Context, side: str, pts: int
):
    actual = getattr(context.section.page_borders, side).width
    expected = Pt(pts)
    assert actual == expected, f"page_borders.{side}.width is {actual!r}"


@then("section.page_borders.{side}.color is {hex_rgb}")
def then_section_page_borders_side_color_is(
    context: Context, side: str, hex_rgb: str
):
    actual = getattr(context.section.page_borders, side).color
    expected = RGBColor.from_string(hex_rgb)
    assert actual == expected, f"page_borders.{side}.color is {actual!r}"


@then("section.page_borders.{side}.space is {pts:d} pt")
def then_section_page_borders_side_space_is_pt(
    context: Context, side: str, pts: int
):
    actual = getattr(context.section.page_borders, side).space
    expected = Pt(pts)
    assert actual == expected, f"page_borders.{side}.space is {actual!r}"


@then("section.page_borders.display is None")
def then_section_page_borders_display_is_None(context: Context):
    actual = context.section.page_borders.display
    assert actual is None, f"page_borders.display is {actual!r}"


@then("section.page_borders.offset_from is None")
def then_section_page_borders_offset_from_is_None(context: Context):
    actual = context.section.page_borders.offset_from
    assert actual is None, f"page_borders.offset_from is {actual!r}"


@then("section.page_borders.display is {value}")
def then_section_page_borders_display_is(context: Context, value: str):
    expected = _BORDER_DISPLAY_MAP[value]
    actual = context.section.page_borders.display
    assert actual == expected, f"page_borders.display is {actual!r}"


@then("section.page_borders.offset_from is {value}")
def then_section_page_borders_offset_from_is(context: Context, value: str):
    expected = _BORDER_OFFSET_FROM_MAP[value]
    actual = context.section.page_borders.offset_from
    assert actual == expected, f"page_borders.offset_from is {actual!r}"


@then('calling section.set_page_border with side "{side}" raises ValueError')
def then_calling_set_page_border_invalid_side_raises(context: Context, side: str):
    try:
        context.section.set_page_border(side, style=WD_BORDER_STYLE.SINGLE)
    except ValueError:
        return
    raise AssertionError("expected ValueError for invalid side")


# ==========================================================================
# Line numbering  (sct-line-numbering.feature)
# ==========================================================================

_LINE_RESTART_MAP = {
    "CONTINUOUS": WD_LINE_NUMBERING_RESTART.CONTINUOUS,
    "NEW_SECTION": WD_LINE_NUMBERING_RESTART.NEW_SECTION,
    "NEW_PAGE": WD_LINE_NUMBERING_RESTART.NEW_PAGE,
}


@given("a Section with no line numbering as section")
def given_a_Section_with_no_line_numbering_as_section(context: Context):
    context.section = Document(test_docx("sct-line-numbering")).sections[0]


@given("a Section with fully populated line numbering as section")
def given_a_Section_with_fully_populated_line_numbering(context: Context):
    context.section = Document(test_docx("sct-line-numbering")).sections[1]


@given("a Section with count_by only line numbering as section")
def given_a_Section_with_count_by_only_line_numbering(context: Context):
    context.section = Document(test_docx("sct-line-numbering")).sections[2]


@when(
    "I call section.set_line_numbering with count_by {n:d} and restart {restart}"
)
def when_I_call_set_line_numbering_count_by_restart(
    context: Context, n: int, restart: str
):
    context.section.set_line_numbering(
        count_by=n, restart=_LINE_RESTART_MAP[restart]
    )


@when("I call section.set_line_numbering with count_by {n:d} only")
def when_I_call_set_line_numbering_count_by_only(context: Context, n: int):
    context.section.set_line_numbering(count_by=n)


@when("I assign {value} to section.line_numbering.count_by")
def when_I_assign_to_line_numbering_count_by(context: Context, value: str):
    context.section.line_numbering.count_by = int(value)


@when("I assign {value} to section.line_numbering.start")
def when_I_assign_to_line_numbering_start(context: Context, value: str):
    context.section.line_numbering.start = int(value)


@when("I assign {value} to section.line_numbering.restart")
def when_I_assign_to_line_numbering_restart(context: Context, value: str):
    context.section.line_numbering.restart = _LINE_RESTART_MAP[value]


@when("I call section.remove_line_numbering()")
def when_I_call_section_remove_line_numbering(context: Context):
    context.section.remove_line_numbering()


@then("section.line_numbering is None")
def then_section_line_numbering_is_None(context: Context):
    assert context.section.line_numbering is None


@then("section.line_numbering is a LineNumbering object")
def then_section_line_numbering_is_a_LineNumbering_object(context: Context):
    assert isinstance(context.section.line_numbering, LineNumbering)


@then("section.line_numbering.count_by is {value}")
def then_section_line_numbering_count_by_is(context: Context, value: str):
    actual = context.section.line_numbering.count_by
    expected = None if value == "None" else int(value)
    assert actual == expected, f"count_by is {actual!r}"


@then("section.line_numbering.start is {value}")
def then_section_line_numbering_start_is(context: Context, value: str):
    actual = context.section.line_numbering.start
    expected = None if value == "None" else int(value)
    assert actual == expected, f"start is {actual!r}"


@then("section.line_numbering.distance is {value}")
def then_section_line_numbering_distance_is(context: Context, value: str):
    actual = context.section.line_numbering.distance
    if value == "None":
        assert actual is None, f"distance is {actual!r}"
    else:
        # format is "<N> pt"
        pts = int(value.split()[0])
        assert actual == Pt(pts), f"distance is {actual!r}"


@then("section.line_numbering.restart is {value}")
def then_section_line_numbering_restart_is(context: Context, value: str):
    actual = context.section.line_numbering.restart
    if value == "None":
        assert actual is None, f"restart is {actual!r}"
    else:
        expected = _LINE_RESTART_MAP[value]
        assert actual == expected, f"restart is {actual!r}"


# ==========================================================================
# Paper source  (sct-paper-source.feature)
# ==========================================================================


@given("a Section with no paperSrc as section")
def given_a_Section_with_no_paperSrc_as_section(context: Context):
    context.section = Document(test_docx("sct-paper-source")).sections[0]


@given("a Section with first=7 and other=15 paperSrc as section")
def given_a_Section_with_first_7_other_15_paperSrc(context: Context):
    context.section = Document(test_docx("sct-paper-source")).sections[1]


@given("a Section with first=1 only paperSrc as section")
def given_a_Section_with_first_1_only_paperSrc(context: Context):
    context.section = Document(test_docx("sct-paper-source")).sections[2]


@given("a Section with other=2 only paperSrc as section")
def given_a_Section_with_other_2_only_paperSrc(context: Context):
    context.section = Document(test_docx("sct-paper-source")).sections[3]


@when("I assign {value} to section.first_page_paper_source")
def when_I_assign_to_first_page_paper_source(context: Context, value: str):
    context.section.first_page_paper_source = None if value == "None" else int(value)


@when("I assign {value} to section.other_pages_paper_source")
def when_I_assign_to_other_pages_paper_source(context: Context, value: str):
    context.section.other_pages_paper_source = None if value == "None" else int(value)


@then("section.first_page_paper_source is {value}")
def then_section_first_page_paper_source_is(context: Context, value: str):
    actual = context.section.first_page_paper_source
    expected = None if value == "None" else int(value)
    assert actual == expected, f"first_page_paper_source is {actual!r}"


@then("section.other_pages_paper_source is {value}")
def then_section_other_pages_paper_source_is(context: Context, value: str):
    actual = context.section.other_pages_paper_source
    expected = None if value == "None" else int(value)
    assert actual == expected, f"other_pages_paper_source is {actual!r}"


@then("section has no paperSrc element")
def then_section_has_no_paperSrc_element(context: Context):
    sectPr = context.section._sectPr  # pyright: ignore[reportPrivateUsage]
    assert sectPr.find(qn("w:paperSrc")) is None, "w:paperSrc still present"


# ==========================================================================
# Document grid  (sct-document-grid.feature)
# ==========================================================================

_DOC_GRID_TYPE_MAP = {
    "DEFAULT": WD_DOC_GRID_TYPE.DEFAULT,
    "LINES": WD_DOC_GRID_TYPE.LINES,
    "LINES_AND_CHARS": WD_DOC_GRID_TYPE.LINES_AND_CHARS,
    "SNAP_TO_CHARS": WD_DOC_GRID_TYPE.SNAP_TO_CHARS,
}


@given("a Section with default document grid as section")
def given_a_Section_with_default_document_grid(context: Context):
    context.section = Document(test_docx("sct-doc-grid")).sections[0]


@given("a Section with fully populated document grid as section")
def given_a_Section_with_fully_populated_document_grid(context: Context):
    context.section = Document(test_docx("sct-doc-grid")).sections[1]


@given("a Section with no document grid as section")
def given_a_Section_with_no_document_grid(context: Context):
    context.section = Document(test_docx("sct-doc-grid")).sections[2]


@when("I call section.set_document_grid with line_pitch {n:d} only")
def when_I_call_set_document_grid_line_pitch_only(context: Context, n: int):
    context.section.set_document_grid(line_pitch=n)


@when(
    "I call section.set_document_grid with type {type_name} and line_pitch {n:d}"
)
def when_I_call_set_document_grid_type_and_line_pitch(
    context: Context, type_name: str, n: int
):
    context.section.set_document_grid(
        type=_DOC_GRID_TYPE_MAP[type_name], line_pitch=n
    )


@when("I assign {value} to section.document_grid.type")
def when_I_assign_to_document_grid_type(context: Context, value: str):
    context.section.document_grid.type = _DOC_GRID_TYPE_MAP[value]


@when("I assign {value} to section.document_grid.line_pitch")
def when_I_assign_to_document_grid_line_pitch(context: Context, value: str):
    context.section.document_grid.line_pitch = int(value)


@when("I assign {value} to section.document_grid.char_space")
def when_I_assign_to_document_grid_char_space(context: Context, value: str):
    context.section.document_grid.char_space = int(value)


@when("I call section.remove_document_grid()")
def when_I_call_section_remove_document_grid(context: Context):
    context.section.remove_document_grid()


@then("section.document_grid is None")
def then_section_document_grid_is_None(context: Context):
    assert context.section.document_grid is None


@then("section.document_grid is a DocumentGrid object")
def then_section_document_grid_is_a_DocumentGrid_object(context: Context):
    assert isinstance(context.section.document_grid, DocumentGrid)


@then("section.document_grid.type is {value}")
def then_section_document_grid_type_is(context: Context, value: str):
    actual = context.section.document_grid.type
    if value == "None":
        assert actual is None, f"type is {actual!r}"
    else:
        expected = _DOC_GRID_TYPE_MAP[value]
        assert actual == expected, f"type is {actual!r}"


@then("section.document_grid.line_pitch is {value}")
def then_section_document_grid_line_pitch_is(context: Context, value: str):
    actual = context.section.document_grid.line_pitch
    expected = None if value == "None" else int(value)
    assert actual == expected, f"line_pitch is {actual!r}"


@then("section.document_grid.char_space is {value}")
def then_section_document_grid_char_space_is(context: Context, value: str):
    actual = context.section.document_grid.char_space
    expected = None if value == "None" else int(value)
    assert actual == expected, f"char_space is {actual!r}"


# ==========================================================================
# Text direction / RTL  (sct-text-direction.feature)
# ==========================================================================

_TEXT_DIRECTION_MAP = {
    "LR_TB": WD_TEXT_DIRECTION.LR_TB,
    "TB_RL": WD_TEXT_DIRECTION.TB_RL,
    "BT_LR": WD_TEXT_DIRECTION.BT_LR,
    "LR_TB_V": WD_TEXT_DIRECTION.LR_TB_V,
    "TB_RL_V": WD_TEXT_DIRECTION.TB_RL_V,
    "TB_LR_V": WD_TEXT_DIRECTION.TB_LR_V,
}


@given("a Section with default text direction as section")
def given_a_Section_with_default_text_direction(context: Context):
    context.section = Document(test_docx("sct-text-direction")).sections[0]


@given("a Section with TB_RL vertical RTL as section")
def given_a_Section_with_TB_RL_vertical_RTL(context: Context):
    context.section = Document(test_docx("sct-text-direction")).sections[1]


@given("a Section with BT_LR vertical LTR as section")
def given_a_Section_with_BT_LR_vertical_LTR(context: Context):
    context.section = Document(test_docx("sct-text-direction")).sections[2]


@when("I assign {value} to section.text_direction")
def when_I_assign_to_section_text_direction(context: Context, value: str):
    context.section.text_direction = (
        None if value == "None" else _TEXT_DIRECTION_MAP[value]
    )


@when("I assign {value} to section.right_to_left")
def when_I_assign_to_section_right_to_left(context: Context, value: str):
    if value == "None":
        context.section.right_to_left = None
    elif value == "True":
        context.section.right_to_left = True
    elif value == "False":
        context.section.right_to_left = False
    else:
        raise ValueError(f"unsupported value {value!r}")


@then("section.text_direction is {value}")
def then_section_text_direction_is(context: Context, value: str):
    actual = context.section.text_direction
    if value == "None":
        assert actual is None, f"text_direction is {actual!r}"
    else:
        expected = _TEXT_DIRECTION_MAP[value]
        assert actual == expected, f"text_direction is {actual!r}"


@then("section.right_to_left is {value}")
def then_section_right_to_left_is(context: Context, value: str):
    actual = context.section.right_to_left
    expected = {"True": True, "False": False}[value]
    assert actual is expected, f"right_to_left is {actual!r}"


# ==========================================================================
# Multi-column layout  (sct-section.feature)
# ==========================================================================


@given("a Section with a single column as section")
def given_a_Section_with_a_single_column(context: Context):
    context.section = Document(test_docx("sct-multi-column")).sections[0]


@given("a Section with three equal columns as section")
def given_a_Section_with_three_equal_columns(context: Context):
    context.section = Document(test_docx("sct-multi-column")).sections[1]


@given("a Section with two unequal columns as section")
def given_a_Section_with_two_unequal_columns(context: Context):
    context.section = Document(test_docx("sct-multi-column")).sections[2]


@given("a Section with two equal columns as section")
def given_a_Section_with_two_equal_columns(context: Context):
    context.section = Document(test_docx("sct-multi-column")).sections[3]


@when("I assign {n:d} to section.columns.count")
def when_I_assign_to_section_columns_count(context: Context, n: int):
    context.section.columns.count = n


@when("I assign {pts:d} to section.columns.space in pt")
def when_I_assign_to_section_columns_space_pt(context: Context, pts: int):
    context.section.columns.space = Pt(pts)


@when("I assign False to section.columns.equal_width")
def when_I_assign_False_to_section_columns_equal_width(context: Context):
    context.section.columns.equal_width = False


@when("I assign {inches:g} inches to section.columns[{idx:d}].width")
def when_I_assign_inches_to_section_columns_idx_width(
    context: Context, inches: float, idx: int
):
    context.section.columns[idx].width = Inches(inches)


@when("I assign {inches:g} inches to section.columns[{idx:d}].space")
def when_I_assign_inches_to_section_columns_idx_space(
    context: Context, inches: float, idx: int
):
    context.section.columns[idx].space = Inches(inches)


@then("section.columns is a SectionColumns object")
def then_section_columns_is_a_SectionColumns_object(context: Context):
    assert isinstance(context.section.columns, SectionColumns)


@then("section.columns.count is {n:d}")
def then_section_columns_count_is(context: Context, n: int):
    actual = context.section.columns.count
    assert actual == n, f"columns.count is {actual!r}"


@then("section.columns.equal_width is {value}")
def then_section_columns_equal_width_is(context: Context, value: str):
    actual = context.section.columns.equal_width
    expected = {"True": True, "False": False}[value]
    assert actual is expected, f"columns.equal_width is {actual!r}"


@then("section.columns.space is None")
def then_section_columns_space_is_None(context: Context):
    actual = context.section.columns.space
    assert actual is None, f"columns.space is {actual!r}"


@then("section.columns.space is {pts:d} pt")
def then_section_columns_space_is_pt(context: Context, pts: int):
    actual = context.section.columns.space
    assert actual == Pt(pts), f"columns.space is {actual!r}"


@then("len(section.columns) is {n:d}")
def then_len_section_columns_is(context: Context, n: int):
    actual = len(context.section.columns)
    assert actual == n, f"len(columns) is {actual!r}"


@then("iterating section.columns yields {n:d} Column objects")
def then_iterating_section_columns_yields_n_Column_objects(
    context: Context, n: int
):
    actual = [col for col in context.section.columns]
    assert len(actual) == n, f"iterated {len(actual)} items"
    for col in actual:
        assert isinstance(col, Column), f"got {type(col).__name__}"


@then("section.columns[{idx:d}].width is {inches:g} inches")
def then_section_columns_idx_width_is_inches(
    context: Context, idx: int, inches: float
):
    actual = context.section.columns[idx].width
    expected = Inches(inches)
    assert actual == expected, f"columns[{idx}].width is {actual!r}"


@then("section.columns[{idx:d}].space is {inches:g} inches")
def then_section_columns_idx_space_is_inches(
    context: Context, idx: int, inches: float
):
    actual = context.section.columns[idx].space
    expected = Inches(inches)
    assert actual == expected, f"columns[{idx}].space is {actual!r}"
