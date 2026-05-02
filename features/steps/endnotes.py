"""Step implementations for document endnotes-related features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.endnotes import Endnote, EndnoteProperties, Endnotes
from docx.enum.text import (
    WD_ENDNOTE_POSITION,
    WD_FOOTNOTE_RESTART,
    WD_NUMBER_FORMAT,
)

from helpers import test_docx

# given ====================================================


@given("a document having an endnotes part")
def given_a_document_having_an_endnotes_part(context: Context):
    context.document = Document(test_docx("end-has-endnotes"))


@given("a document having no endnotes part")
def given_a_document_having_no_endnotes_part(context: Context):
    context.document = Document(test_docx("doc-default"))


# when =====================================================


@when("I add two endnotes to two different runs")
def when_I_add_two_endnotes_to_two_different_runs(context: Context):
    document = context.document
    p1 = document.add_paragraph("First anchor paragraph.")
    p2 = document.add_paragraph("Second anchor paragraph.")
    en1 = document.endnotes.add(p1.runs[-1], text="First")
    en2 = document.endnotes.add(p2.runs[-1], text="Second")
    context.added_endnotes = [en1, en2]


@when("I assign endnote = document.endnotes.add(run)")
def when_I_assign_endnote_eq_endnotes_add_run(context: Context):
    document = context.document
    paragraph = document.add_paragraph("Body text for endnote anchor.")
    context.run = paragraph.runs[-1]
    context.endnote = document.endnotes.add(context.run)


@when('I assign endnote = document.endnotes.add(run, "An endnote body")')
def when_I_assign_endnote_eq_endnotes_add_run_text(context: Context):
    document = context.document
    paragraph = document.add_paragraph("Body text for endnote anchor.")
    context.run = paragraph.runs[-1]
    context.endnote = document.endnotes.add(context.run, "An endnote body")


@when("I assign endnote = the first user endnote in the document")
def when_I_assign_endnote_eq_first_user_endnote(context: Context):
    context.endnote = next(iter(context.document.endnotes))


@when("I assign paragraph = endnote.add_paragraph()")
def when_I_assign_paragraph_eq_endnote_add_paragraph(context: Context):
    context.paragraph = context.endnote.add_paragraph()


@when("I assign paragraph = endnote.add_paragraph(text, style)")
def when_I_assign_paragraph_eq_endnote_add_paragraph_text_style(context: Context):
    context.para_text = text = "Extended endnote content"
    context.para_style = style = "Normal"
    context.paragraph = context.endnote.add_paragraph(text, style)


@when('I assign endnote.add_paragraph("Extra paragraph")')
def when_I_assign_endnote_add_paragraph_extra(context: Context):
    context.endnote.add_paragraph("Extra paragraph")


@when("I call endnote.clear()")
def when_I_call_endnote_clear(context: Context):
    context.clear_result = context.endnote.clear()


@when("I call endnote.delete()")
def when_I_call_endnote_delete(context: Context):
    context.endnote.delete()


@when("I assign props = document.add_endnote_properties()")
def when_I_assign_props_eq_add_endnote_properties(context: Context):
    context.props = context.document.add_endnote_properties()


@when("I assign props.number_format = None")
def when_I_assign_props_number_format_None(context: Context):
    props = context.document.add_endnote_properties()
    props.number_format = None


@when("I assign props.number_format = WD_NUMBER_FORMAT.{value}")
def when_I_assign_props_number_format(context: Context, value: str):
    props = context.document.add_endnote_properties()
    props.number_format = getattr(WD_NUMBER_FORMAT, value)


@when("I assign props.restart_rule = WD_FOOTNOTE_RESTART.{value}")
def when_I_assign_props_restart_rule(context: Context, value: str):
    props = context.document.add_endnote_properties()
    props.restart_rule = getattr(WD_FOOTNOTE_RESTART, value)


@when("I assign props.position = WD_ENDNOTE_POSITION.{value}")
def when_I_assign_props_position(context: Context, value: str):
    props = context.document.add_endnote_properties()
    props.position = getattr(WD_ENDNOTE_POSITION, value)


@when("I assign props.start_number = {value:d}")
def when_I_assign_props_start_number(context: Context, value: int):
    props = context.document.add_endnote_properties()
    props.start_number = value


# then =====================================================


@then("document.endnotes is an Endnotes object")
def then_document_endnotes_is_an_Endnotes_object(context: Context):
    assert type(context.document.endnotes) is Endnotes


@then("document.endnote_properties is None")
def then_document_endnote_properties_is_None(context: Context):
    assert context.document.endnote_properties is None


@then("document.endnote_properties is an EndnoteProperties object")
def then_document_endnote_properties_is_an_EndnoteProperties_object(context: Context):
    actual = context.document.endnote_properties
    assert type(actual) is EndnoteProperties, f"expected EndnoteProperties, got {type(actual)}"


@then("document.endnote_properties.number_format == WD_NUMBER_FORMAT.{value}")
def then_document_endnote_properties_number_format(context: Context, value: str):
    expected = getattr(WD_NUMBER_FORMAT, value)
    actual = context.document.endnote_properties.number_format
    assert actual == expected, f"expected number_format {expected}, got {actual}"


@then("document.endnote_properties.number_format is None")
def then_document_endnote_properties_number_format_is_None(context: Context):
    actual = context.document.endnote_properties.number_format
    assert actual is None, f"expected number_format None, got {actual}"


@then("document.endnote_properties.restart_rule == WD_FOOTNOTE_RESTART.{value}")
def then_document_endnote_properties_restart_rule(context: Context, value: str):
    expected = getattr(WD_FOOTNOTE_RESTART, value)
    actual = context.document.endnote_properties.restart_rule
    assert actual == expected, f"expected restart_rule {expected}, got {actual}"


@then("document.endnote_properties.position == WD_ENDNOTE_POSITION.{value}")
def then_document_endnote_properties_position(context: Context, value: str):
    expected = getattr(WD_ENDNOTE_POSITION, value)
    actual = context.document.endnote_properties.position
    assert actual == expected, f"expected position {expected}, got {actual}"


@then("document.endnote_properties.start_number == {value:d}")
def then_document_endnote_properties_start_number(context: Context, value: int):
    actual = context.document.endnote_properties.start_number
    assert actual == value, f"expected start_number {value}, got {actual}"


@then("endnote is an Endnote object")
def then_endnote_is_an_Endnote_object(context: Context):
    assert type(context.endnote) is Endnote


@then("endnote.endnote_id == {value:d}")
def then_endnote_endnote_id_eq_value(context: Context, value: int):
    actual = context.endnote.endnote_id
    assert actual == value, f"expected endnote_id {value}, got {actual}"


@then('endnote.text == "{text}"')
def then_endnote_text_eq(context: Context, text: str):
    actual = context.endnote.text
    assert actual == text, f"expected endnote.text '{text}', got '{actual}'"


@then("endnote.clear() returns endnote")
def then_endnote_clear_returns_endnote(context: Context):
    returned = context.endnote.clear()
    assert returned is context.endnote, "endnote.clear() did not return self"


@then("endnote.paragraphs[{idx:d}] == paragraph")
def then_endnote_paragraphs_idx_eq_paragraph(context: Context, idx: int):
    actual = context.endnote.paragraphs[idx]._p
    expected = context.paragraph._p
    assert actual == expected, "paragraphs do not compare equal"


@then('endnote.paragraphs[{idx:d}].style.name == "{style}"')
def then_endnote_paragraphs_idx_style_name_eq(context: Context, idx: int, style: str):
    actual = context.endnote.paragraphs[idx]._p.style
    assert actual == style, f"expected style '{style}', got '{actual}'"


@then('endnote.paragraphs[{idx:d}].text == ""')
def then_endnote_paragraphs_idx_text_eq_empty(context: Context, idx: int):
    actual = context.endnote.paragraphs[idx].text
    assert actual == "", f"expected empty text, got '{actual}'"


@then("iterating document.endnotes yields {count:d} Endnote objects")
def then_iterating_document_endnotes_yields_count(context: Context, count: int):
    endnotes = list(context.document.endnotes)
    assert len(endnotes) == count, (
        f"expected {count} endnotes from iteration, got {len(endnotes)}"
    )
    for en in endnotes:
        assert type(en) is Endnote, f"expected Endnote, got {type(en)}"


@then("iterating document.endnotes yields endnote ids [{ids}]")
def then_iterating_document_endnotes_yields_endnote_ids(context: Context, ids: str):
    expected = [int(x.strip()) for x in ids.split(",")]
    actual = [en.endnote_id for en in context.document.endnotes]
    assert actual == expected, f"expected endnote ids {expected}, got {actual}"


@then("len(document.endnotes) == {count:d}")
def then_len_document_endnotes_eq_count(context: Context, count: int):
    actual = len(context.document.endnotes)
    assert actual == count, f"expected len(document.endnotes) {count}, got {actual}"


@then("len(endnote.paragraphs) == {count:d}")
def then_len_endnote_paragraphs_eq_count(context: Context, count: int):
    actual = len(context.endnote.paragraphs)
    assert actual == count, f"expected len(endnote.paragraphs) {count}, got {actual}"


@then("props is an EndnoteProperties object")
def then_props_is_an_EndnoteProperties_object(context: Context):
    assert type(context.props) is EndnoteProperties, (
        f"expected EndnoteProperties, got {type(context.props)}"
    )


@then("the added endnote ids are [{ids}]")
def then_the_added_endnote_ids_are(context: Context, ids: str):
    expected = [int(x.strip()) for x in ids.split(",")]
    actual = [en.endnote_id for en in context.added_endnotes]
    assert actual == expected, f"expected added endnote ids {expected}, got {actual}"
