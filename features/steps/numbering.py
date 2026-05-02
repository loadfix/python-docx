"""Step implementations for numbering-related features."""

from __future__ import annotations

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.text import WD_NUMBER_FORMAT
from docx.oxml.ns import qn
from docx.shared import Inches

from helpers import test_docx

# given ===================================================


@given("a document having a numbering part")
def given_a_document_having_a_numbering_part(context: Context):
    context.document = Document(test_docx("num-having-numbering-part"))


@given("a document having a custom numbering definition")
def given_a_document_having_a_custom_numbering_definition(context: Context):
    context.document = Document(test_docx("num-defs"))


@given("a fresh default document with one paragraph")
def given_a_fresh_default_document_with_one_paragraph(context: Context):
    context.document = Document()
    context.paragraph = context.document.add_paragraph("A single paragraph.")


# when ====================================================


@when("I get the numbering part from the document")
def when_get_numbering_part_from_document(context: Context):
    document = context.document
    context.numbering_part = document.part.numbering_part


@when("I add a single-level bullet numbering definition")
def when_add_single_level_bullet_definition(context: Context):
    document = context.document
    context.definition = document.numbering.add_numbering_definition(
        [
            {
                "format": WD_NUMBER_FORMAT.BULLET,
                "text": "•",
                "indent": Inches(0.25),
                "font": "Symbol",
            }
        ]
    )


@when("I add a single-level decimal numbering definition")
def when_add_single_level_decimal_definition(context: Context):
    document = context.document
    context.definition = document.numbering.add_numbering_definition(
        [{"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1.", "indent": Inches(0.25)}]
    )


@when("I apply the last numbering definition to the paragraph at level {level:d}")
def when_apply_last_definition_to_paragraph(context: Context, level: int):
    document = context.document
    definition = document.numbering.definitions[-1]
    definition.apply_to(context.paragraph, level=level)


# then =====================================================


@then("the numbering part has the expected numbering definitions")
def then_numbering_part_has_expected_numbering_definitions(context: Context):
    numbering_part = context.numbering_part
    assert len(numbering_part.numbering_definitions) == 10


@then("document.numbering has at least {count:d} definition")
@then("document.numbering has at least {count:d} definitions")
def then_numbering_has_at_least_n_definitions(context: Context, count: int):
    actual = len(context.document.numbering.definitions)
    assert actual >= count, f"expected at least {count}, got {actual}"


@then("the last numbering definition has {count:d} levels")
@then("the last numbering definition has {count:d} level")
def then_last_definition_has_levels(context: Context, count: int):
    definition = context.document.numbering.definitions[-1]
    actual = len(definition.levels)
    assert actual == count, f"expected {count} levels, got {actual}"


_FMT_BY_NAME: dict[str, WD_NUMBER_FORMAT] = {
    member.name: member for member in WD_NUMBER_FORMAT
}


@then("level {ilvl:d} of the last definition has number_format == {fmt}")
def then_level_number_format(context: Context, ilvl: int, fmt: str):
    definition = context.document.numbering.definitions[-1]
    lvl = definition.level(ilvl)
    assert lvl is not None, f"no level {ilvl}"
    expected = _FMT_BY_NAME[fmt.strip()]
    assert lvl.number_format == expected, (
        f"expected {expected!r}, got {lvl.number_format!r}"
    )


@then('level {ilvl:d} of the last definition has text == "{text}"')
def then_level_text(context: Context, ilvl: int, text: str):
    definition = context.document.numbering.definitions[-1]
    lvl = definition.level(ilvl)
    assert lvl is not None, f"no level {ilvl}"
    assert lvl.text == text, f"expected {text!r}, got {lvl.text!r}"


@then("level {ilvl:d} of the last definition has start == {start:d}")
def then_level_start(context: Context, ilvl: int, start: int):
    definition = context.document.numbering.definitions[-1]
    lvl = definition.level(ilvl)
    assert lvl is not None, f"no level {ilvl}"
    assert lvl.start == start, f"expected {start}, got {lvl.start}"


@then("the first three paragraphs have a w:numPr child")
def then_first_three_paragraphs_have_numpr(context: Context):
    for idx, paragraph in enumerate(context.document.paragraphs[:3]):
        children = paragraph._p.xpath(".//w:numPr")
        assert children, f"paragraph {idx} has no w:numPr"


@then("the paragraph has a w:numPr child")
def then_paragraph_has_numpr(context: Context):
    children = context.paragraph._p.xpath(".//w:numPr")
    assert children, "paragraph has no w:numPr child"


@then("applying the definition to the paragraph at level {level:d} raises ValueError")
def then_apply_to_raises_valueerror(context: Context, level: int):
    try:
        context.definition.apply_to(context.paragraph, level=level)
    except ValueError:
        return
    raise AssertionError(f"expected ValueError for level {level}")


# -- list-label rendering steps ---------------------------------------


@given(
    "a document with a decimal-then-letter multi-level numbering and "
    "{count:d} nested paragraphs"
)
def given_decimal_letter_multilevel_document(context: Context, count: int):
    document = Document()
    defn = document.numbering.add_numbering_definition(
        [
            {"format": WD_NUMBER_FORMAT.DECIMAL, "text": "%1."},
            {"format": WD_NUMBER_FORMAT.LOWER_LETTER, "text": "%2)"},
        ]
    )
    # -- 1., a), b), 2., 3., a) when count == 6 --
    assert count >= 4, "scenario expects at least 4 paragraphs"
    pattern = [0, 1, 1, 0, 0, 1]
    context.document = document
    context.definition = defn
    context.paragraphs = []
    for i in range(count):
        ilvl = pattern[i % len(pattern)]
        p = document.add_paragraph(f"item {i}")
        defn.apply_to(p, ilvl)
        context.paragraphs.append(p)


@given("a document with a single-level bullet numbering and {count:d} bullet paragraphs")
def given_bullet_document(context: Context, count: int):
    document = Document()
    defn = document.numbering.add_numbering_definition(
        [
            {
                "format": WD_NUMBER_FORMAT.BULLET,
                "text": "•",
                "indent": Inches(0.25),
                "font": "Symbol",
            }
        ]
    )
    context.document = document
    context.paragraphs = []
    for i in range(count):
        p = document.add_paragraph(f"bullet {i}")
        defn.apply_to(p, 0)
        context.paragraphs.append(p)


@given("a trailing plain paragraph is appended")
def given_trailing_plain_paragraph(context: Context):
    context.trailing_paragraph = context.document.add_paragraph("plain item")


@then("the paragraph's list_label is None")
def then_paragraph_list_label_is_none(context: Context):
    assert context.paragraph.list_label is None, (
        f"expected None, got {context.paragraph.list_label!r}"
    )


@then('the list labels for the {count:d} paragraphs are "{expected}"')
def then_list_labels_match(context: Context, count: int, expected: str):
    paragraphs = context.paragraphs[:count]
    actual = [p.list_label for p in paragraphs]
    expected_labels = [s.strip() for s in expected.split(",")]
    assert len(actual) == len(expected_labels), (
        f"expected {len(expected_labels)} paragraphs, got {len(actual)}"
    )
    assert actual == expected_labels, (
        f"expected {expected_labels!r}, got {actual!r}"
    )


@then("document.list_labels has an entry for each of the {count:d} paragraphs")
def then_document_list_labels_size(context: Context, count: int):
    labels = context.document.list_labels()
    numbered = [p for p in context.paragraphs if p.list_label is not None]
    assert len(numbered) == count, (
        f"expected {count} numbered paragraphs, got {len(numbered)}"
    )
    for p in numbered:
        assert id(p._p) in labels, (
            f"paragraph {p.text!r} missing from document.list_labels"
        )


@then('the list_labels entry for paragraph {index:d} is "{expected}"')
def then_list_labels_entry(context: Context, index: int, expected: str):
    labels = context.document.list_labels()
    p = context.paragraphs[index - 1]
    actual = labels.get(id(p._p))
    assert actual == expected, (
        f"paragraph {index}: expected {expected!r}, got {actual!r}"
    )


@then("document.list_labels has no entry for the trailing paragraph")
def then_document_list_labels_no_trailing(context: Context):
    labels = context.document.list_labels()
    assert id(context.trailing_paragraph._p) not in labels, (
        "trailing plain paragraph unexpectedly appears in list_labels"
    )


# -- expose qn for debugging --
_ = qn  # avoid lint "unused import" when refactoring
