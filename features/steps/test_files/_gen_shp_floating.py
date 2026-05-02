"""Generate `shp-floating.docx` behave fixture.

Builds a document exercising ``Paragraph.add_floating_image``: three floating
pictures with different horizontal/vertical anchors and wrap types, plus a
trailing paragraph with plain text so the anchor does not end up as the very
last body element.

Self-validates after generation by re-opening the file and asserting the
floating-image count, anchors, offsets, and wrap types.

Run::

    python features/steps/test_files/_gen_shp_floating.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.shape import WD_ANCHOR_H, WD_ANCHOR_V, WD_WRAP_TYPE
from docx.shared import Emu, Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "shp-floating.docx")
IMAGE_PATH = os.path.join(HERE, "monty-truth.png")


def build() -> Document:
    document = Document()

    document.add_heading("Floating images fixture", level=1)

    # -- Floating image 1: default placement (column/paragraph, SQUARE) --
    p1 = document.add_paragraph("Paragraph with a default floating image. ")
    p1.add_floating_image(IMAGE_PATH, width=Inches(1.0))

    # -- Floating image 2: page/page anchor with non-zero offsets, TIGHT wrap --
    p2 = document.add_paragraph("Paragraph with a page-anchored floating image. ")
    p2.add_floating_image(
        IMAGE_PATH,
        width=Inches(1.5),
        height=Inches(1.0),
        position={
            "h_anchor": WD_ANCHOR_H.PAGE,
            "v_anchor": WD_ANCHOR_V.PAGE,
            "horizontal": Inches(2),
            "vertical": Inches(3),
            "wrap": WD_WRAP_TYPE.TIGHT,
        },
    )

    # -- Floating image 3: margin anchor with alt_text + title set afterwards --
    p3 = document.add_paragraph("Paragraph with a margin-anchored floating image. ")
    float_img = p3.add_floating_image(
        IMAGE_PATH,
        width=Inches(1.2),
        position={
            "h_anchor": WD_ANCHOR_H.MARGIN,
            "v_anchor": WD_ANCHOR_V.MARGIN,
            "horizontal": Emu(457200),  # 0.5in
            "vertical": Emu(914400),    # 1.0in
            "wrap": WD_WRAP_TYPE.TOP_AND_BOTTOM,
        },
    )
    float_img.alt_text = "Decorative mountain graphic"
    float_img.title = "Mountain"

    document.add_paragraph("Trailing paragraph after floating images.")
    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    p1_floats = paragraphs[1].floating_images
    assert len(p1_floats) == 1, f"p1 expected 1 floating image, got {len(p1_floats)}"
    fi1 = p1_floats[0]
    assert fi1.horizontal_anchor is WD_ANCHOR_H.COLUMN
    assert fi1.vertical_anchor is WD_ANCHOR_V.PARAGRAPH
    assert fi1.wrap_type is WD_WRAP_TYPE.SQUARE

    p2_floats = paragraphs[2].floating_images
    assert len(p2_floats) == 1
    fi2 = p2_floats[0]
    assert fi2.horizontal_anchor is WD_ANCHOR_H.PAGE
    assert fi2.vertical_anchor is WD_ANCHOR_V.PAGE
    assert fi2.wrap_type is WD_WRAP_TYPE.TIGHT
    assert fi2.horizontal_offset == int(Inches(2))
    assert fi2.vertical_offset == int(Inches(3))

    p3_floats = paragraphs[3].floating_images
    assert len(p3_floats) == 1
    fi3 = p3_floats[0]
    assert fi3.horizontal_anchor is WD_ANCHOR_H.MARGIN
    assert fi3.vertical_anchor is WD_ANCHOR_V.MARGIN
    assert fi3.wrap_type is WD_WRAP_TYPE.TOP_AND_BOTTOM
    assert fi3.alt_text == "Decorative mountain graphic"
    assert fi3.title == "Mountain"


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    reopened = Document(OUT_PATH)
    self_validate(reopened)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
