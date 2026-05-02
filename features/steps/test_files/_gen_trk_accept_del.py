"""Build the ``trk-accept-del.docx`` fixture for the tracked-change accept-deletion tests.

Run from anywhere (the output path is resolved relative to this file)::

    python features/steps/test_files/_gen_trk_accept_del.py

*Generator strategy*: the paragraph/table scaffolding is built with python-docx's
public API; the tracked-deletion wrappers (``w:del``, ``w:delText``, ``w:cellDel``,
and the paragraph-mark ``w:pPr/w:rPr/w:del``) are injected as raw ``OxmlElement``
nodes because python-docx does not yet surface authoring methods for tracked
deletions. The fixture targets the **accept** side only.

Fixture layout:

* paragraph 0 -- heading (skipped by tests)
* paragraph 1 -- "Keep A [-gone A-] end A." single-run deletion by Bob
* paragraph 2 -- "Span [-first second third-] tail." deletion spanning three
  ``w:r`` children (a single ``w:del`` holding three ``w:r/w:delText`` blocks)
  by Carol
* paragraph 3 -- "alpha [+added+] [-removed-] omega." a tracked insertion by
  Alice next to a deletion by Bob -- accepting deletions should leave the
  insertion wrapper untouched
* paragraph 4 -- "full paragraph removed" with both the paragraph mark and the
  content tracked for deletion: ``w:pPr/w:rPr/w:del`` AND an enclosing
  ``w:del`` around the run
* paragraph 5 -- "survivor" untouched live text (proves non-deleted content is
  preserved)
* table -- 1x2: cell (0,0) carries a ``w:cellDel`` by Dave, cell (0,1) is plain
  live content "kept cell". Accepting the cellDel removes cell (0,0) from the
  row.

Self-checks before the file is written verify that every tracked-deletion
element is visible via the public read-side API (``Paragraph.tracked_changes``,
``Cell.is_tracked_deletion``) so later accept-side assertions have something to
compare against.
"""

from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-accept-del.docx")


def _make_del(run_id: int, author: str, date: str, texts: list[str]):
    """Return a ``w:del`` element wrapping one ``w:r/w:delText`` per entry in ``texts``."""
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:id"), str(run_id))
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), date)
    for text in texts:
        r = OxmlElement("w:r")
        dt_elm = OxmlElement("w:delText")
        dt_elm.text = text
        dt_elm.set(qn("xml:space"), "preserve")
        r.append(dt_elm)
        wdel.append(r)
    return wdel


def _make_ins(run_id: int, author: str, date: str, text: str):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(run_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    ins.append(r)
    return ins


def _make_cell_marker(tag: str, marker_id: int, author: str, date: str):
    elm = OxmlElement(tag)
    elm.set(qn("w:id"), str(marker_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    return elm


def _mark_paragraph_deletion(p, marker_id: int, author: str, date: str) -> None:
    """Add a ``w:pPr/w:rPr/w:del`` marker so the paragraph mark itself is tracked
    for deletion."""
    pPr = p._p.get_or_add_pPr()
    rPr = OxmlElement("w:rPr")
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:id"), str(marker_id))
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), date)
    rPr.append(wdel)
    pPr.append(rPr)


def build() -> Document:
    document = Document()

    # -- paragraph 0: heading --
    document.add_heading("Tracked deletions -- accept fixture", level=1)

    # -- paragraph 1: single-run deletion between two live runs --
    p1 = document.add_paragraph("Keep A ")
    p1._p.append(_make_del(1, "Bob", "2025-05-01T09:00:00Z", ["gone A"]))
    p1.add_run(" end A.")

    # -- paragraph 2: one deletion wrapping three runs --
    p2 = document.add_paragraph("Span ")
    p2._p.append(
        _make_del(
            2, "Carol", "2025-05-01T09:05:00Z",
            ["first ", "second ", "third"],
        )
    )
    p2.add_run(" tail.")

    # -- paragraph 3: insertion by Alice + deletion by Bob side-by-side. The
    # -- accept-all run should unwrap the insertion and drop the deletion; the
    # -- per-change "accept only deletions" step should leave the w:ins intact
    # -- because we only flip `deletion`-typed changes. --
    p3 = document.add_paragraph("alpha ")
    p3._p.append(_make_ins(3, "Alice", "2025-05-01T09:10:00Z", "added "))
    p3._p.append(_make_del(4, "Bob", "2025-05-01T09:11:00Z", ["removed "]))
    p3.add_run("omega.")

    # -- paragraph 4: whole paragraph tracked for deletion (paragraph mark + content) --
    p4 = document.add_paragraph()
    # -- inject the content as a w:del wrapping a single run with w:delText --
    p4._p.append(
        _make_del(5, "Dave", "2025-05-01T09:15:00Z", ["full paragraph removed"])
    )
    _mark_paragraph_deletion(p4, 6, "Dave", "2025-05-01T09:15:00Z")

    # -- paragraph 5: plain survivor --
    document.add_paragraph("survivor")

    # -- a 1x2 table with a cellDel on cell (0,0) --
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "dropped cell"
    table.cell(0, 1).text = "kept cell"
    tc00 = table.cell(0, 0)._tc
    tcPr00 = tc00.get_or_add_tcPr()
    tcPr00.append(
        _make_cell_marker("w:cellDel", 7, "Dave", "2025-05-01T09:20:00Z")
    )

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- paragraph 1: one deletion by Bob --
    p1_changes = paragraphs[1].tracked_changes
    assert len(p1_changes) == 1, (
        f"p1 expected 1 tracked change, got {len(p1_changes)}"
    )
    assert p1_changes[0].type == "deletion"
    assert p1_changes[0].author == "Bob"
    assert p1_changes[0].text == "gone A"

    # -- paragraph 2: a single deletion whose text is the concatenation of three
    # -- delText children --
    p2_changes = paragraphs[2].tracked_changes
    assert len(p2_changes) == 1, (
        f"p2 expected 1 tracked change, got {len(p2_changes)}"
    )
    assert p2_changes[0].type == "deletion"
    assert p2_changes[0].text == "first second third"

    # -- paragraph 3: insertion + deletion (types list order matches document
    # -- order) --
    p3_changes = paragraphs[3].tracked_changes
    types = [tc.type for tc in p3_changes]
    assert types == ["insertion", "deletion"], (
        f"p3 types mismatch: {types}"
    )

    # -- paragraph 4: content deletion visible as a tracked change. The
    # -- paragraph-mark marker (w:pPr/w:rPr/w:del) is not surfaced through
    # -- Paragraph.tracked_changes (it walks direct children), but xpath can
    # -- confirm it was written. --
    p4_changes = paragraphs[4].tracked_changes
    assert len(p4_changes) == 1, (
        f"p4 expected 1 direct-child deletion, got {len(p4_changes)}"
    )
    assert p4_changes[0].author == "Dave"
    pmark_dels = paragraphs[4]._p.xpath("./w:pPr/w:rPr/w:del")
    assert len(pmark_dels) == 1, (
        "expected a w:pPr/w:rPr/w:del paragraph-mark deletion on p4"
    )

    # -- paragraph 5 is plain --
    assert paragraphs[5].tracked_changes == []

    # -- table cell (0,0) is tracked as deleted --
    table = document.tables[0]
    assert table.cell(0, 0).is_tracked_deletion is True
    assert table.cell(0, 1).is_tracked_deletion is False


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
