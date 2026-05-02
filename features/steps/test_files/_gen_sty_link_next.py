"""Generate ``sty-link-next.docx`` fixture for Style linked/next/redefined props.

Builds a document whose styles collection contains:

* ``Body`` paragraph style linked to ``BodyChar`` character style and vice-versa.
* ``Intro`` paragraph style whose ``w:next`` points at ``Body`` and whose
  ``w:autoRedefine`` flag is set so :attr:`Style.is_redefined` reads |True|.
* ``Solo`` paragraph style with no link/next/autoRedefine for negative coverage.

Run ``python features/steps/test_files/_gen_sty_link_next.py`` to regenerate the
fixture in-place and self-validate the result.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "sty-link-next.docx")


def build() -> str:
    """Create ``sty-link-next.docx`` and return its absolute path."""
    document = Document()
    document.add_paragraph("Style link/next/redefined fixture.")

    styles = document.styles

    body_para = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
    body_char = styles.add_style("BodyChar", WD_STYLE_TYPE.CHARACTER)
    intro = styles.add_style("Intro", WD_STYLE_TYPE.PARAGRAPH)
    styles.add_style("Solo", WD_STYLE_TYPE.PARAGRAPH)

    # -- paragraph<->character link --
    body_para.link_style = body_char
    body_char.link_style = body_para

    # -- next style + auto-redefine --
    intro.next_style = body_para
    intro._element.autoRedefine_val = True

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    """Reload `path` and assert each style property round-trips."""
    document = Document(path)
    styles = document.styles

    body_para = styles["Body"]
    body_char = styles["BodyChar"]
    intro = styles["Intro"]
    solo = styles["Solo"]

    # -- link round-trips both directions --
    assert body_para.link_style is not None
    assert body_para.link_style.name == "BodyChar", (
        f"Body.link_style.name = {body_para.link_style.name!r}"
    )
    assert body_char.link_style is not None
    assert body_char.link_style.name == "Body"

    # -- next round-trips --
    assert intro.next_style is not None
    assert intro.next_style.name == "Body"

    # -- is_redefined reads the autoRedefine flag --
    assert intro.is_redefined is True
    assert body_para.is_redefined is False

    # -- Solo has none of the above set --
    assert solo.link_style is None
    assert solo.next_style is None
    assert solo.is_redefined is False


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
