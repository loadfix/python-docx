"""Generate the `bmk-has-bookmarks.docx` behave fixture.

Running this script (idempotently) produces a .docx containing a small set of
well-formed bookmarks covering the shapes exercised by the `bmk-*` scenarios:

- ``bm_intro``  — wraps the whole first paragraph (single run).
- ``bm_middle`` — wraps a specific run in the middle of a paragraph.
- ``bm_span``   — spans two paragraphs, starting before a run in one paragraph
                  and ending after a run in a later paragraph.

The script self-validates the written document by re-opening it and asserting
the expected bookmark names, ids, and count before considering the fixture
good.
"""

from __future__ import annotations

import os

from docx import Document


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "bmk-has-bookmarks.docx")


def build() -> None:
    document = Document()

    # -- Paragraph 0: bookmarked whole --
    p0 = document.add_paragraph("Introduction section.")
    p0.add_bookmark("bm_intro")

    # -- Paragraph 1: bookmark wrapping a single middle run --
    p1 = document.add_paragraph("The ")
    p1.add_run("middle").bold = True
    p1.add_run(" run is bookmarked.")
    middle_run = p1.runs[1]
    p1.add_bookmark("bm_middle", start_run=middle_run, end_run=middle_run)

    # -- Paragraphs 2-3: a bookmark spanning two paragraphs --
    p2 = document.add_paragraph("Span start ")
    p2.add_run("first-half")
    p3 = document.add_paragraph("Span end ")
    p3.add_run("second-half")

    # -- paragraph.add_bookmark() only spans a single paragraph; build a
    #    cross-paragraph range directly on the run elements. --
    start_run = p2.runs[1]
    end_run = p3.runs[1]
    start_run._r.insert_bookmark_start_before(2, "bm_span")
    end_run._r.insert_bookmark_end_after(2)

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    bookmarks = document.bookmarks

    assert len(bookmarks) == 3, f"expected 3 bookmarks, got {len(bookmarks)}"

    names = sorted(bm.name for bm in bookmarks)
    assert names == ["bm_intro", "bm_middle", "bm_span"], f"got {names!r}"

    ids = sorted(bm.bookmark_id for bm in bookmarks)
    assert len(set(ids)) == 3, f"bookmark ids not unique: {ids!r}"

    intro = bookmarks.get("bm_intro")
    assert intro is not None
    assert intro.name == "bm_intro"

    middle = bookmarks.get("bm_middle")
    assert middle is not None

    span = bookmarks.get("bm_span")
    assert span is not None
    assert span.bookmark_id == 2

    # -- both markers for the cross-paragraph bookmark must be present --
    body = document._body._body  # pyright: ignore[reportPrivateUsage]
    starts = body.xpath(".//w:bookmarkStart[@w:id='2']")
    ends = body.xpath(".//w:bookmarkEnd[@w:id='2']")
    assert len(starts) == 1 and len(ends) == 1, "bm_span markers malformed"


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
