"""Generate heading-structure accessibility fixture .docx files.

Creates three fixtures used by ``features/acc-heading-structure.feature``:

* ``acc-valid-headings.docx`` — clean H1 > H2 > H3 outline (no issues).
* ``acc-missing-h2.docx`` — document that starts at Heading 2 (missing H1).
* ``acc-skipped-level.docx`` — H1 directly followed by H3 (H2 skipped).

Each fixture is re-opened after saving and its heading issues are validated so
that the generator fails loudly if python-docx ever changes behavior in a way
that invalidates the fixtures.
"""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.accessibility import (
    EMPTY_HEADING,
    MULTIPLE_H1,
    NO_H1,
    SKIPPED_LEVEL,
    validate_heading_structure,
)

HERE = os.path.dirname(os.path.abspath(__file__))


def _save(document: Document, name: str) -> str:
    path = os.path.join(HERE, f"{name}.docx")
    document.save(path)
    return path


def build_valid_headings() -> Document:
    document = Document()
    document.add_paragraph("Document Title", style="Heading 1")
    document.add_paragraph("An introduction paragraph.")
    document.add_paragraph("First Section", style="Heading 2")
    document.add_paragraph("Body text for the first section.")
    document.add_paragraph("A Subsection", style="Heading 3")
    document.add_paragraph("More body text.")
    document.add_paragraph("Second Section", style="Heading 2")
    document.add_paragraph("Closing paragraph.")
    return document


def build_missing_h2() -> Document:
    """Build a document whose first heading is Heading 2 (no H1 at all)."""
    document = Document()
    document.add_paragraph("A lead paragraph.")
    document.add_paragraph("First Second-Level Heading", style="Heading 2")
    document.add_paragraph("Body text.")
    document.add_paragraph("Another Second-Level Heading", style="Heading 2")
    document.add_paragraph("More body text.")
    return document


def build_skipped_level() -> Document:
    """Build a document that jumps from Heading 1 directly to Heading 3."""
    document = Document()
    document.add_paragraph("Top-level Heading", style="Heading 1")
    document.add_paragraph("Body text.")
    document.add_paragraph("Skipped-level Heading", style="Heading 3")
    document.add_paragraph("More body text.")
    return document


def _kinds(document: Document) -> list[str]:
    return [issue.kind for issue in validate_heading_structure(document.paragraphs)]


def main() -> int:
    # -- valid: no issues --
    doc = build_valid_headings()
    path = _save(doc, "acc-valid-headings")
    kinds = _kinds(Document(path))
    assert kinds == [], f"acc-valid-headings should have no issues, got {kinds}"

    # -- missing-h2: first heading is H2, flagged as NO_H1 --
    doc = build_missing_h2()
    path = _save(doc, "acc-missing-h2")
    kinds = _kinds(Document(path))
    assert NO_H1 in kinds, f"acc-missing-h2 should report NO_H1, got {kinds}"
    assert SKIPPED_LEVEL not in kinds, (
        f"acc-missing-h2 should not report skipped level, got {kinds}"
    )
    assert MULTIPLE_H1 not in kinds, (
        f"acc-missing-h2 should not report multiple H1, got {kinds}"
    )
    assert EMPTY_HEADING not in kinds, (
        f"acc-missing-h2 should not report empty heading, got {kinds}"
    )

    # -- skipped-level: H1 then H3, flagged as SKIPPED_LEVEL --
    doc = build_skipped_level()
    path = _save(doc, "acc-skipped-level")
    kinds = _kinds(Document(path))
    assert SKIPPED_LEVEL in kinds, (
        f"acc-skipped-level should report skipped level, got {kinds}"
    )
    assert NO_H1 not in kinds, (
        f"acc-skipped-level should not report NO_H1, got {kinds}"
    )

    print("generated acc-valid-headings.docx, acc-missing-h2.docx, acc-skipped-level.docx")
    return 0


if __name__ == "__main__":
    sys.exit(main())
