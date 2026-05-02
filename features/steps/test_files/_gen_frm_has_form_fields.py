"""Generate ``frm-has-form-fields.docx`` via python-docx's public API.

The fixture exercises all three legacy form-field types — a text input, a
checkbox, and a dropdown — each in its own paragraph. The generator uses the
public ``Paragraph.add_*_form_field`` methods only, so no raw OOXML is emitted.

Running this module regenerates the fixture in place and then self-validates by
re-opening the saved file and asserting the form-field collection matches
expectations. Intended to be idempotent.

Usage::

    python features/steps/test_files/_gen_frm_has_form_fields.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.form_fields import WD_FORM_FIELD_TYPE


HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "frm-has-form-fields.docx")


def build() -> None:
    """Build and save the fixture."""
    document = Document()

    # -- paragraph 1: a text input form field with a default and max length --
    p_text = document.add_paragraph("Name: ")
    p_text.add_text_form_field(name="FullName", default="Jane Doe", maxlength=40)

    # -- paragraph 2: a checkbox form field, initially checked --
    p_check = document.add_paragraph("Subscribe? ")
    p_check.add_checkbox_form_field(name="Subscribe", checked=True)

    # -- paragraph 3: a dropdown form field with three options --
    p_drop = document.add_paragraph("Country: ")
    p_drop.add_dropdown_form_field(
        name="Country",
        options=["US", "UK", "AU"],
        default_index=1,
    )

    document.save(OUT_PATH)


def validate() -> None:
    """Re-open the saved fixture and assert expected form-field shape."""
    document = Document(OUT_PATH)
    fields = document.form_fields

    assert len(fields) == 3, f"expected 3 form fields, got {len(fields)}"

    # -- text input --
    text_ff = fields[0]
    assert text_ff.type is WD_FORM_FIELD_TYPE.TEXT
    assert text_ff.name == "FullName"
    assert text_ff.text_input is not None
    assert text_ff.text_input.default == "Jane Doe"
    assert text_ff.text_input.max_length == 40
    assert text_ff.value == "Jane Doe"

    # -- checkbox --
    cb_ff = fields[1]
    assert cb_ff.type is WD_FORM_FIELD_TYPE.CHECKBOX
    assert cb_ff.name == "Subscribe"
    assert cb_ff.checkbox is not None
    assert cb_ff.checkbox.checked is True
    assert cb_ff.value is True

    # -- dropdown --
    dd_ff = fields[2]
    assert dd_ff.type is WD_FORM_FIELD_TYPE.DROPDOWN
    assert dd_ff.name == "Country"
    assert dd_ff.dropdown is not None
    assert dd_ff.dropdown.options == ["US", "UK", "AU"]
    assert dd_ff.dropdown.default_index == 1
    assert dd_ff.value == "UK"


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
