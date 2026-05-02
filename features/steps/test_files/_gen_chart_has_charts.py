"""Generate `chart-has-charts.docx` fixture for the chart-read behave feature.

The generated document contains three charts of different `WD_CHART_TYPE` so
the read-side API (`Document.charts`, `Chart.chart_type`, `Chart.title`,
`Chart.series`) can be exercised end-to-end.

`Document.add_chart` does not populate a chart title, so after the charts are
appended this script injects a minimal `c:title` element into each chart part
so the `Chart.title` scenario has something to read.

Running the script is idempotent: it overwrites the output path. The script
also self-validates by reopening the saved document and asserting the
expected shape of the charts collection.

Usage::

    python features/steps/test_files/_gen_chart_has_charts.py
"""

from __future__ import annotations

import os
from typing import cast

from lxml import etree

from docx import Document
from docx.chart import WD_CHART_TYPE
from docx.oxml.chart import CT_Chart, CT_ChartSpace
from docx.oxml.ns import nsmap, qn

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "chart-has-charts.docx")


# -- (chart_type, title, categories, series_data) tuples for each chart to embed --
CHARTS: list[
    tuple[WD_CHART_TYPE, str, list[str], dict[str, list[float]]]
] = [
    (
        WD_CHART_TYPE.COLUMN,
        "Quarterly Sales",
        ["Q1", "Q2", "Q3", "Q4"],
        {"East": [10.0, 20.0, 15.0, 25.0], "West": [12.0, 18.0, 14.0, 22.0]},
    ),
    (
        WD_CHART_TYPE.BAR,
        "Headcount by Region",
        ["North", "South", "East", "West"],
        {"Employees": [30.0, 45.0, 25.0, 40.0]},
    ),
    (
        WD_CHART_TYPE.LINE,
        "Monthly Revenue",
        ["Jan", "Feb", "Mar"],
        {"2024": [100.0, 110.0, 120.0], "2025": [130.0, 140.0, 150.0]},
    ),
]


def _set_chart_title(chart_elm: CT_Chart, title_text: str) -> None:
    """Insert a `c:title` as the first child of `chart_elm` carrying `title_text`.

    The schema places `c:title` before `c:plotArea` in the `c:chart` element,
    so the new title is inserted at index 0.
    """
    c_uri = nsmap["c"]
    a_uri = nsmap["a"]

    title = etree.SubElement(chart_elm, qn("c:title"))
    tx = etree.SubElement(title, qn("c:tx"))
    rich = etree.SubElement(tx, qn("c:rich"))

    # -- a:bodyPr and a:lstStyle are required children of c:rich per schema --
    etree.SubElement(rich, qn("a:bodyPr"))
    etree.SubElement(rich, qn("a:lstStyle"))

    p = etree.SubElement(rich, qn("a:p"))
    r = etree.SubElement(p, qn("a:r"))
    t = etree.SubElement(r, qn("a:t"))
    t.text = title_text

    # -- required sibling per spec --
    overlay = etree.SubElement(title, qn("c:overlay"))
    overlay.set("val", "0")

    # -- move the just-appended title to be the first child of c:chart --
    chart_elm.remove(title)
    chart_elm.insert(0, title)

    # -- suppress these unused-name warnings; symbols used for namespace lookup --
    assert c_uri and a_uri


def main() -> str:
    document = Document()
    document.add_heading("Charts fixture", level=1)
    document.add_paragraph(
        "This document embeds three charts of different types so the "
        "chart-read API can be exercised by the behave suite."
    )

    for chart_type, title, categories, series_data in CHARTS:
        chart = document.add_chart(chart_type, categories, series_data)
        chartSpace = cast(CT_ChartSpace, chart.part.element)
        chart_elm = chartSpace.chart
        assert chart_elm is not None, "add_chart should produce a c:chart element"
        _set_chart_title(chart_elm, title)

    document.save(OUT_PATH)

    # -- self-validate by reopening and asserting the collection shape --
    reopened = Document(OUT_PATH)
    charts = reopened.charts
    assert len(charts) == len(CHARTS), (
        f"expected {len(CHARTS)} charts in fixture, found {len(charts)}"
    )
    for chart, (expected_type, expected_title, _, expected_series) in zip(
        charts, CHARTS
    ):
        assert chart.chart_type == expected_type, (
            f"expected chart_type {expected_type}, got {chart.chart_type}"
        )
        assert chart.title == expected_title, (
            f"expected title {expected_title!r}, got {chart.title!r}"
        )
        series_names = [s.name for s in chart.series]
        assert series_names == list(expected_series.keys()), (
            f"expected series names {list(expected_series)}, got {series_names}"
        )

    return OUT_PATH


if __name__ == "__main__":
    path = main()
    print(f"wrote {path}")
