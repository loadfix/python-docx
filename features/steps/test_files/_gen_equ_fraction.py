"""Generate ``equ-fraction.docx`` fixture for equation behave scenarios.

Produces a document with a single inline fraction equation built with
``docx.equations.build_fraction``. The numerator is ``a`` and the denominator
is ``b``, yielding the OMML equivalent of ``a / b`` with a horizontal bar.

A second paragraph carries a compound fraction equation whose numerator is
``x + 1`` and whose denominator is ``y``; this exercises the builder with
multi-character identifiers so scenarios can assert on the flattened plain-text
rendering.

The file is self-validating: after saving, it is re-opened and the equation
count, structure, and plain-text rendering are all checked.

Run directly::

    python features/steps/test_files/_gen_equ_fraction.py
"""

from __future__ import annotations

import os
import sys

# -- add src/ to sys.path so this script can be run from anywhere ------------------
THIS_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.abspath(os.path.join(THIS_DIR, "..", "..", ".."))
SRC_DIR = os.path.join(REPO_ROOT, "src")
if SRC_DIR not in sys.path:
    sys.path.insert(0, SRC_DIR)

from docx import Document  # noqa: E402
from docx.equations import build_fraction  # noqa: E402

OUT_PATH = os.path.join(THIS_DIR, "equ-fraction.docx")


def build() -> None:
    document = Document()

    # -- paragraph 0: narrative describing the fixture ----------------------------
    document.add_paragraph("A pair of fraction equations follows.")

    # -- paragraph 1: simple a/b fraction -----------------------------------------
    p1 = document.add_paragraph("The ratio ")
    p1.add_equation(build_fraction("a", "b"))

    # -- paragraph 2: (x+1)/y compound fraction -----------------------------------
    p2 = document.add_paragraph("The compound ratio ")
    p2.add_equation(build_fraction("x+1", "y"))

    # -- validate before saving ---------------------------------------------------
    equations = document.equations
    assert len(equations) == 2, f"expected 2 equations, got {len(equations)}"
    # -- flattened text concatenates <m:t> content for numerator then denominator
    assert equations[0].text == "ab", f"expected 'ab', got {equations[0].text!r}"
    assert equations[1].text == "x+1y", (
        f"expected 'x+1y', got {equations[1].text!r}"
    )
    for equation in equations:
        raw = equation.raw_xml
        assert b"<m:f>" in raw, f"expected <m:f> fraction element, got {raw!r}"
        assert b"<m:num>" in raw and b"<m:den>" in raw, (
            f"expected <m:num>/<m:den> children, got {raw!r}"
        )
    assert all(not e.is_display_mode for e in equations), (
        "fraction equations should be inline (m:oMath), not display-mode"
    )

    document.save(OUT_PATH)

    # -- re-load and re-validate from disk ----------------------------------------
    reloaded = Document(OUT_PATH)
    reloaded_equations = reloaded.equations
    assert len(reloaded_equations) == 2
    assert [e.text for e in reloaded_equations] == ["ab", "x+1y"]
    for equation in reloaded_equations:
        assert b"<m:f>" in equation.raw_xml

    print(f"wrote {OUT_PATH} ({len(reloaded_equations)} equations)")


if __name__ == "__main__":
    build()
