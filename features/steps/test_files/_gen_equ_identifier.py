"""Generate ``equ-identifier.docx`` fixture for equation behave scenarios.

Produces a document containing two inline OMML equations built with
``docx.equations.build_identifier``:

- paragraph 0 introduces the fixture with plain body text
- paragraph 1 ends with an italic "x" identifier equation
- paragraph 2 ends with a Greek chi ("chi") identifier equation

The file is self-validating: after saving, it is re-opened and the equation
count, OMML payload, and plain-text rendering are all checked.

Run directly::

    python features/steps/test_files/_gen_equ_identifier.py
"""

from __future__ import annotations

import os
import sys

# -- add src/ to sys.path so this script can be run from anywhere ------------------
THIS_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.abspath(os.path.join(THIS_DIR, "..", "..", ".."))
SRC_DIR = os.path.join(REPO_ROOT, "src")
if SRC_DIR not in sys.path:
    sys.path.insert(0, SRC_DIR)

from docx import Document  # noqa: E402
from docx.equations import build_identifier  # noqa: E402

OUT_PATH = os.path.join(THIS_DIR, "equ-identifier.docx")


def build() -> None:
    document = Document()

    # -- paragraph 0: plain narrative, no equation --------------------------------
    document.add_paragraph("The following paragraphs contain inline identifiers.")

    # -- paragraph 1: an italic "x" identifier ------------------------------------
    p1 = document.add_paragraph("Let the variable be ")
    p1.add_equation(build_identifier("x"))

    # -- paragraph 2: a Greek chi identifier --------------------------------------
    p2 = document.add_paragraph("And the Greek letter ")
    # -- U+03C7 is the Greek small letter chi; stored here as the literal
    # -- character inside the <m:t> element. --
    p2.add_equation(build_identifier("χ"))

    # -- validate before saving ---------------------------------------------------
    equations = document.equations
    assert len(equations) == 2, f"expected 2 equations, got {len(equations)}"
    assert equations[0].text == "x", f"expected 'x', got {equations[0].text!r}"
    assert equations[1].text == "χ", (
        f"expected Greek chi, got {equations[1].text!r}"
    )
    assert all(not e.is_display_mode for e in equations), (
        "identifier equations should be inline (m:oMath), not display-mode"
    )
    for equation in equations:
        raw = equation.raw_xml
        assert b"<m:r>" in raw and b"<m:t>" in raw, (
            f"expected <m:r>/<m:t> in raw OMML, got {raw!r}"
        )

    document.save(OUT_PATH)

    # -- re-load and re-validate from disk ----------------------------------------
    reloaded = Document(OUT_PATH)
    reloaded_equations = reloaded.equations
    assert len(reloaded_equations) == 2
    assert [e.text for e in reloaded_equations] == ["x", "χ"]

    print(f"wrote {OUT_PATH} ({len(reloaded_equations)} equations)")


if __name__ == "__main__":
    build()
