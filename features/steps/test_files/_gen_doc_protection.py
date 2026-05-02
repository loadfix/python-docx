"""Generate ``doc-protection.docx`` fixture for document-protection scenarios.

Produces a document with :class:`WD_PROTECTION.COMMENTS` mode enabled, enforced,
and a SHA-1-hashed password attached via :meth:`.Settings.enable_protection`.

Run ``python features/steps/test_files/_gen_doc_protection.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_PROTECTION

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "doc-protection.docx")


def build() -> str:
    document = Document()
    document.add_paragraph(
        "This document only allows comments to be added by reviewers."
    )

    document.settings.enable_protection(
        WD_PROTECTION.COMMENTS,
        password="s3cret",
        enforce=True,
    )

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    dp = document.settings.document_protection
    assert dp.mode == WD_PROTECTION.COMMENTS, f"mode was {dp.mode!r}"
    assert dp.enforce is True
    assert dp.password_hash is not None, "password hash missing"
    assert dp.password_salt is not None, "password salt missing"
    assert dp.spin_count == 100000, f"spin_count was {dp.spin_count!r}"
    assert dp.crypto_algorithm_sid == 4, f"sid was {dp.crypto_algorithm_sid!r}"


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
