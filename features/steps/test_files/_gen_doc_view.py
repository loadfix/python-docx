"""Generate ``doc-view.docx`` fixture for ``Settings.view`` scenarios.

Builds a document whose ``w:settings/w:view`` is set to ``outline`` so behave
scenarios can read back a non-default |WD_VIEW| value. ``zoom_percent`` is also
set to exercise a second non-default view-related setting.

Run ``python features/steps/test_files/_gen_doc_view.py`` to regenerate the
fixture in-place and self-validate the result.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_VIEW

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "doc-view.docx")


def build() -> str:
    """Create ``doc-view.docx`` and return its absolute path."""
    document = Document()
    document.add_paragraph("View / zoom fixture.")

    settings = document.settings
    settings.view = WD_VIEW.OUTLINE
    settings.zoom_percent = 175
    settings.track_revisions = True

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    """Reload `path` and assert each view-related setting round-trips."""
    document = Document(path)
    settings = document.settings

    assert settings.view is WD_VIEW.OUTLINE, f"view round-trip failed: {settings.view!r}"
    assert settings.zoom_percent == 175, f"zoom round-trip failed: {settings.zoom_percent!r}"
    assert settings.track_revisions is True


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
