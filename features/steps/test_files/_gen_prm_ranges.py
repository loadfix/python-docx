"""Generate ``prm-ranges.docx`` fixture for permission-range scenarios.

Creates a document with three permission ranges — one everyone-edit-group,
one single-user, one named-group — so scenarios can verify :attr:`~docx.document.Document.permission_ranges`
enumerates them in document order and that :attr:`.PermissionRange.user` and
:attr:`.PermissionRange.edit_group` reflect the underlying ``@w:ed``/``@w:edGrp``
attributes.

Run ``python features/steps/test_files/_gen_prm_ranges.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "prm-ranges.docx")


def build() -> str:
    document = Document()
    p1 = document.add_paragraph("Editable by everyone.")
    p2 = document.add_paragraph("Editable by Alice only.")
    p3 = document.add_paragraph("Editable by the Authors group.")

    p1.add_permission_range(edit_group="everyone")
    p2.add_permission_range(user="alice@example.com")
    p3.add_permission_range(edit_group="Authors")

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    ranges = document.permission_ranges
    assert len(ranges) == 3, f"expected 3 permission ranges, got {len(ranges)}"

    assert ranges[0].edit_group == "everyone"
    assert ranges[0].user is None

    assert ranges[1].user == "alice@example.com"
    assert ranges[1].edit_group is None

    assert ranges[2].edit_group == "Authors"
    assert ranges[2].user is None

    # -- ids should be monotonically assigned starting at 0 --
    assert [r.id for r in ranges] == [0, 1, 2]


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
