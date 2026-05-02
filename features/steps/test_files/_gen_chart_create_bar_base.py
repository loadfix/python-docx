"""Generate ``chart-create-bar-base.docx`` — a tiny docx with no chart.

The chart-create-bar behave feature uses this base document as input so each
scenario can exercise :meth:`docx.document.Document.add_chart` in a ``When``
step, and assert against a freshly-created chart rather than one that was
baked into the fixture.

The generated document has:
- one heading paragraph ("Bar chart create fixture"),
- two plain paragraphs of intro text,
- *no* chart parts or chart references,

so ``Document(...).charts`` returns ``[]`` on open.

Run ``python features/steps/test_files/_gen_chart_create_bar_base.py`` to
(re)generate the fixture. The script self-validates by reopening the saved
document and asserting the expected shape.
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "chart-create-bar-base.docx")


INTRO_PARAGRAPHS = (
    "This document is the starting point for the chart-create-bar behave "
    "scenarios; it contains no chart parts.",
    "Scenarios append a bar chart with Document.add_chart and then read the "
    "chart back through the public API.",
)


def build() -> str:
    document = Document()
    document.add_heading("Bar chart create fixture", level=1)
    for text in INTRO_PARAGRAPHS:
        document.add_paragraph(text)
    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)

    # -- no chart parts yet --
    assert document.charts == [], (
        f"expected no charts in base fixture, got {len(document.charts)}"
    )

    # -- heading + the two intro paragraphs are present --
    paragraphs = document.paragraphs
    assert len(paragraphs) >= 1 + len(INTRO_PARAGRAPHS), (
        f"expected at least {1 + len(INTRO_PARAGRAPHS)} paragraphs, "
        f"got {len(paragraphs)}"
    )
    assert paragraphs[0].text == "Bar chart create fixture", (
        f"expected first paragraph to be the heading, got {paragraphs[0].text!r}"
    )
    expected_texts = list(INTRO_PARAGRAPHS)
    actual_texts = [p.text for p in paragraphs[1 : 1 + len(INTRO_PARAGRAPHS)]]
    assert actual_texts == expected_texts, (
        f"expected intro paragraphs {expected_texts}, got {actual_texts}"
    )


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
