"""Generate ``sct-paper-source.docx`` — fixture for paper-source behave tests.

The ``w:paperSrc`` element carries two independently-settable attributes,
``w:first`` and ``w:other``, exposed on |Section| as
:attr:`first_page_paper_source` and :attr:`other_pages_paper_source`.
Four sections cover the four relevant states:

- **Section 0** — no ``w:paperSrc`` element at all.
- **Section 1** — ``w:first=7``, ``w:other=15``.
- **Section 2** — ``w:first=1`` only.
- **Section 3** — ``w:other=2`` only.

Run this script from the repo root::

    python features/steps/test_files/_gen_sct_paper_source.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "sct-paper-source.docx")


def build() -> None:
    document = Document()
    document.add_paragraph("Section 0 — no paperSrc.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 1 — first=7, other=15.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 2 — first only.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 3 — other only.")

    # -- clear on every section to undo add_section cloning, then populate. ---
    for section in document.sections:
        section.first_page_paper_source = None
        section.other_pages_paper_source = None

    document.sections[1].first_page_paper_source = 7
    document.sections[1].other_pages_paper_source = 15

    document.sections[2].first_page_paper_source = 1

    document.sections[3].other_pages_paper_source = 2

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    sections = document.sections
    assert len(sections) == 4

    assert sections[0].first_page_paper_source is None
    assert sections[0].other_pages_paper_source is None

    assert sections[1].first_page_paper_source == 7
    assert sections[1].other_pages_paper_source == 15

    assert sections[2].first_page_paper_source == 1
    assert sections[2].other_pages_paper_source is None

    assert sections[3].first_page_paper_source is None
    assert sections[3].other_pages_paper_source == 2


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
