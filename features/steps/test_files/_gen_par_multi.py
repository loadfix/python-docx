"""Generate ``par-multi.docx`` — fixture for paragraph/run helper scenarios.

The document contains a small sequence of paragraphs exercising the paragraph,
run, and stable-id helpers:

- Paragraph 0: a Heading 1 with the text ``"Intro"`` (single run).
- Paragraph 1: body paragraph with three runs: ``"Alpha "``, ``"Beta "``,
  ``"Gamma"`` — used by ``Run.delete`` and ``Run.split`` scenarios.
- Paragraph 2: body paragraph with the single run ``"Middle"`` — used by
  ``Paragraph.delete`` and paragraph ``stable_id`` scenarios.
- Paragraph 3: a Heading 2 with the text ``"Outro"``.
- Paragraph 4: body paragraph with a single run ``"Tail"``.

Run directly to (re)generate the fixture at::

    features/steps/test_files/par-multi.docx

Self-validates on save by re-opening and asserting the paragraph/run counts
and the expected run text.
"""

from __future__ import annotations

import os
import sys

from docx import Document


THIS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(THIS_DIR, "par-multi.docx")


def _build() -> Document:
    document = Document()

    # -- P0: Heading 1 --
    document.add_paragraph("Intro", style="Heading 1")

    # -- P1: three runs, used for run.split and run.delete --
    p1 = document.add_paragraph()
    p1.add_run("Alpha ")
    p1.add_run("Beta ")
    p1.add_run("Gamma")

    # -- P2: single run "Middle" (target for paragraph.delete) --
    document.add_paragraph("Middle")

    # -- P3: Heading 2 "Outro" --
    document.add_paragraph("Outro", style="Heading 2")

    # -- P4: single run "Tail" --
    document.add_paragraph("Tail")

    return document


def _validate(path: str) -> None:
    document = Document(path)
    paragraphs = document.paragraphs
    assert len(paragraphs) == 5, f"expected 5 paragraphs, got {len(paragraphs)}"

    assert paragraphs[0].text == "Intro"
    assert paragraphs[0].style.name == "Heading 1"

    assert paragraphs[1].text == "Alpha Beta Gamma"
    assert [r.text for r in paragraphs[1].runs] == ["Alpha ", "Beta ", "Gamma"]

    assert paragraphs[2].text == "Middle"
    assert len(paragraphs[2].runs) == 1

    assert paragraphs[3].text == "Outro"
    assert paragraphs[3].style.name == "Heading 2"

    assert paragraphs[4].text == "Tail"


def main() -> int:
    document = _build()
    document.save(OUT_PATH)
    _validate(OUT_PATH)
    print(f"wrote {OUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
