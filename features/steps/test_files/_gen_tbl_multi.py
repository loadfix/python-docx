"""Generate ``tbl-multi.docx`` — fixture for Table.delete scenarios.

The document contains three 2x2 tables interleaved with labelled paragraphs so
that deleting a table by index can be asserted against the remaining document
structure:

- Paragraph:  ``"Before table 1"``
- Table 0:    cells ``"T0 r0c0"..."T0 r1c1"``
- Paragraph:  ``"Between 1 and 2"``
- Table 1:    cells ``"T1 r0c0"..."T1 r1c1"``
- Paragraph:  ``"Between 2 and 3"``
- Table 2:    cells ``"T2 r0c0"..."T2 r1c1"``
- Paragraph:  ``"After table 3"``

Run directly to (re)generate the fixture at::

    features/steps/test_files/tbl-multi.docx

Self-validates on save by re-opening and asserting that three tables are
present with the expected cell text.
"""

from __future__ import annotations

import os
import sys

from docx import Document


THIS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(THIS_DIR, "tbl-multi.docx")


def _build() -> Document:
    document = Document()

    for i in range(3):
        label = {0: "Before table 1", 1: "Between 1 and 2", 2: "Between 2 and 3"}[i]
        document.add_paragraph(label)
        table = document.add_table(rows=2, cols=2)
        for r in range(2):
            for c in range(2):
                table.cell(r, c).text = f"T{i} r{r}c{c}"

    document.add_paragraph("After table 3")
    return document


def _validate(path: str) -> None:
    document = Document(path)
    tables = document.tables
    assert len(tables) == 3, f"expected 3 tables, got {len(tables)}"
    for i, table in enumerate(tables):
        for r in range(2):
            for c in range(2):
                expected = f"T{i} r{r}c{c}"
                got = table.cell(r, c).text
                assert got == expected, f"table {i} cell ({r},{c}): expected {expected!r}, got {got!r}"


def main() -> int:
    document = _build()
    document.save(OUT_PATH)
    _validate(OUT_PATH)
    print(f"wrote {OUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
