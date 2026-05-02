"""Generate `chart-create-line-base.docx` fixture for chart-create-line behave tests.

The generated document is deliberately minimal: a single heading paragraph and
one body paragraph, with no embedded charts. It serves as the "blank canvas"
starting point for scenarios that exercise `Document.add_chart(WD_CHART_TYPE.LINE,
...)`.

Running the script is idempotent: it overwrites the output path. The script
self-validates by reopening the saved document and asserting it contains no
charts.

Usage::

    python features/steps/test_files/_gen_chart_create_line_base.py
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "chart-create-line-base.docx")


def main() -> str:
    document = Document()
    document.add_heading("Line chart fixture", level=1)
    document.add_paragraph(
        "This document has no charts. Scenarios append a line chart via "
        "Document.add_chart(WD_CHART_TYPE.LINE, ...)."
    )

    document.save(OUT_PATH)

    # -- self-validate: reopen and assert no charts present --
    reopened = Document(OUT_PATH)
    charts = reopened.charts
    assert charts == [], (
        f"expected fixture to contain no charts, found {len(charts)}"
    )

    # -- also confirm the two authored paragraphs survived the round trip --
    paragraphs = reopened.paragraphs
    assert len(paragraphs) >= 2, (
        f"expected at least 2 paragraphs in fixture, found {len(paragraphs)}"
    )

    return OUT_PATH


if __name__ == "__main__":
    path = main()
    print(f"wrote {path}")
