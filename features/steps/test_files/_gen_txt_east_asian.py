"""Generate the ``txt-east-asian.docx`` behave fixture.

The file exercises three East Asian layout and typography surfaces:

- paragraph 0: a run with no ``w:eastAsianLayout`` child, ``name_far_east``
  unset, and neither ``w:kinsoku`` nor ``w:wordWrap`` on the paragraph.
- paragraph 1: a run carrying ``w:rPr/w:rFonts@w:eastAsia="MS Mincho"`` and a
  ``w:eastAsianLayout`` element with ``two_lines_in_one=True``, plus
  ``w:pPr/w:kinsoku@w:val="0"`` (kinsoku explicitly off) on the paragraph.
- paragraph 2: a run with ``w:eastAsianLayout`` set for vertical alignment
  with ``compressed=True``, and the paragraph carries ``w:kinsoku@w:val="1"``
  and ``w:wordWrap@w:val="0"``.
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "txt-east-asian.docx")


def build() -> Document:
    document = Document()

    # -- paragraph 0: no east-asian formatting at all --
    document.add_paragraph("Plain paragraph with no East Asian layout.")

    # -- paragraph 1: combined (two-lines-in-one) layout + east-asian font +
    #    kinsoku explicitly off --
    p1 = document.add_paragraph("Run using MS Mincho with two-lines-in-one.")
    font1 = p1.runs[0].font
    font1.name_far_east = "MS Mincho"
    font1.set_east_asian_layout(two_lines_in_one=True, id=1)
    p1.paragraph_format.kinsoku = False

    # -- paragraph 2: vertical + compressed layout, kinsoku on, word_wrap off --
    p2 = document.add_paragraph("Run with vertical-compressed layout.")
    font2 = p2.runs[0].font
    font2.set_east_asian_layout(vertical_alignment=True, compressed=True, id=2)
    p2.paragraph_format.kinsoku = True
    p2.paragraph_format.word_wrap = False

    return document


def validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- paragraph 0 --
    font0 = paragraphs[0].runs[0].font
    assert font0.name_far_east is None
    assert font0.east_asian_layout is None
    assert paragraphs[0].paragraph_format.kinsoku is None
    assert paragraphs[0].paragraph_format.word_wrap is None

    # -- paragraph 1 --
    font1 = paragraphs[1].runs[0].font
    assert font1.name_far_east == "MS Mincho", font1.name_far_east
    layout1 = font1.east_asian_layout
    assert layout1 is not None
    assert layout1.two_lines_in_one is True
    assert layout1.id == 1
    assert paragraphs[1].paragraph_format.kinsoku is False

    # -- paragraph 2 --
    font2 = paragraphs[2].runs[0].font
    layout2 = font2.east_asian_layout
    assert layout2 is not None
    assert layout2.vertical_alignment is True
    assert layout2.compressed is True
    assert paragraphs[2].paragraph_format.kinsoku is True
    assert paragraphs[2].paragraph_format.word_wrap is False


def main() -> None:
    document = build()
    document.save(OUT_PATH)
    validate(Document(OUT_PATH))
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
