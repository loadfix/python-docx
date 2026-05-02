"""Generate the ``txt-ruby.docx`` behave fixture.

The file contains a single paragraph with two runs so scenarios can select
different ruby shapes by run index:

- run 0: two sequential ``w:ruby`` annotations — ``日本`` (with ruby
  ``にほん`` and ``w:rubyPr/w:rubyAlign@w:val="distributeSpace"`` and
  ``w:rubyPr/w:lid@w:val="ja-JP"``), and ``東京`` with ruby ``とうきょう``
  (no ``w:rubyPr`` children).
- run 1: a ``w:ruby`` whose base and annotation are empty (covers the
  "missing components" path).

Ruby elements are read-only in ``python-docx`` so they are constructed
directly by inserting ``w:ruby`` XML into the run via ``lxml``.
"""

from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "txt-ruby.docx")


def _ruby_xml(
    base: str, ruby: str, align: str | None = None, lang: str | None = None
) -> str:
    parts = []
    if align is not None or lang is not None:
        children = []
        if align is not None:
            children.append(f'<w:rubyAlign w:val="{align}"/>')
        if lang is not None:
            children.append(f'<w:lid w:val="{lang}"/>')
        parts.append("<w:rubyPr>" + "".join(children) + "</w:rubyPr>")
    else:
        parts.append("<w:rubyPr/>")
    parts.append(
        f"<w:rt><w:r><w:t>{ruby}</w:t></w:r></w:rt>"
        f"<w:rubyBase><w:r><w:t>{base}</w:t></w:r></w:rubyBase>"
    )
    return f"<w:ruby {nsdecls('w')}>{''.join(parts)}</w:ruby>"


def build() -> Document:
    document = Document()

    paragraph = document.add_paragraph()
    run0 = paragraph.add_run()
    run0._r.append(
        parse_xml(_ruby_xml("日本", "にほん", align="distributeSpace", lang="ja-JP"))
    )
    run0._r.append(parse_xml(_ruby_xml("東京", "とうきょう")))

    run1 = paragraph.add_run()
    run1._r.append(parse_xml(_ruby_xml("", "")))

    return document


def validate(document: Document) -> None:
    paragraph = document.paragraphs[0]
    run0 = paragraph.runs[0]
    rubies0 = run0.ruby_annotations
    assert len(rubies0) == 2, f"run 0: expected 2 rubies, got {len(rubies0)}"

    assert rubies0[0].base_text == "日本"
    assert rubies0[0].ruby_text == "にほん"
    assert rubies0[0].alignment == "distributeSpace"
    assert rubies0[0].language == "ja-JP"

    assert rubies0[1].base_text == "東京"
    assert rubies0[1].ruby_text == "とうきょう"
    assert rubies0[1].alignment is None
    assert rubies0[1].language is None

    run1 = paragraph.runs[1]
    rubies1 = run1.ruby_annotations
    assert len(rubies1) == 1
    assert rubies1[0].base_text == ""
    assert rubies1[0].ruby_text == ""


def main() -> None:
    document = build()
    document.save(OUT_PATH)
    validate(Document(OUT_PATH))
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
