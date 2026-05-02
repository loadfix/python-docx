"""Build the ``trk-ins-del.docx`` fixture used by the tracked-changes acceptance tests.

Run from anywhere (the output path is resolved relative to this file)::

    python features/steps/test_files/_gen_trk_ins_del.py

*Generator strategy*: this fixture is built with python-docx's public API for the
live paragraph/run scaffolding, then raw ``lxml``/``OxmlElement`` injection for the
tracked-change wrappers (``w:ins``, ``w:del``, ``w:delText``). No public "add tracked
insertion/deletion" method exists yet on |Paragraph|; callers who need authoring
support currently drop down to the oxml layer as this generator does.

Self-checks performed before the file is written:

* paragraph 1 exposes one ``w:ins`` (author ``Alice``) and one ``w:del`` (author
  ``Bob``) via ``Paragraph.tracked_changes``
* paragraph 2 exposes a second insertion by a third author so iteration counts are
  non-trivial
* ``revision_marks_text()`` renders the bracketed preview as expected
"""

from __future__ import annotations

import datetime as dt
import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-ins-del.docx")


def _make_ins(run_id: int, author: str, date: str, text: str):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(run_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    ins.append(r)
    return ins


def _make_del(run_id: int, author: str, date: str, text: str):
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:id"), str(run_id))
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    dt_elm = OxmlElement("w:delText")
    dt_elm.text = text
    dt_elm.set(qn("xml:space"), "preserve")
    r.append(dt_elm)
    wdel.append(r)
    return wdel


def build() -> Document:
    document = Document()

    # -- paragraph 0: title (so tests can rely on fixed indices) --
    document.add_heading("Tracked insertions and deletions", level=1)

    # -- paragraph 1: mixed insertion + deletion in one paragraph --
    p1 = document.add_paragraph("The quick ")
    p1._p.append(
        _make_del(1, "Bob", "2025-04-10T09:00:00Z", "brown")
    )
    p1._p.append(
        _make_ins(2, "Alice", "2025-04-10T09:05:00Z", "nimble")
    )
    p1.add_run(" fox jumps.")

    # -- paragraph 2: a lone insertion by a third author --
    p2 = document.add_paragraph("Goodbye")
    p2._p.append(
        _make_ins(3, "Carol", "2025-04-11T14:30:00Z", ", cruel world")
    )
    p2.add_run(".")

    # -- paragraph 3: plain paragraph with no tracked changes --
    document.add_paragraph("Nothing to see here.")

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- paragraph 1 exposes ins + del in order --
    p1_changes = paragraphs[1].tracked_changes
    assert len(p1_changes) == 2, (
        f"expected 2 tracked changes on paragraph 1, got {len(p1_changes)}"
    )
    types = [tc.type for tc in p1_changes]
    assert types == ["deletion", "insertion"], (
        f"expected ['deletion', 'insertion'], got {types}"
    )
    authors = [tc.author for tc in p1_changes]
    assert authors == ["Bob", "Alice"], f"authors mismatch: {authors}"
    texts = [tc.text for tc in p1_changes]
    assert texts == ["brown", "nimble"], f"texts mismatch: {texts}"
    # -- date must round-trip as datetime --
    assert isinstance(p1_changes[0].date, dt.datetime)

    # -- paragraph 2 has one insertion --
    p2_changes = paragraphs[2].tracked_changes
    assert len(p2_changes) == 1
    assert p2_changes[0].type == "insertion"
    assert p2_changes[0].author == "Carol"

    # -- paragraph 3 has no tracked changes --
    assert paragraphs[3].tracked_changes == []

    # -- revision_marks_text renders brackets --
    preview = paragraphs[1].revision_marks_text()
    assert preview == "The quick [-brown-][+nimble+] fox jumps.", (
        f"unexpected preview: {preview!r}"
    )


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
