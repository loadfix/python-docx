"""Generate ``mmg-enabled.docx`` fixture for mail-merge behave scenarios.

Builds a minimal document whose ``w:settings/w:mailMerge`` element is populated
with a realistic combination of arguments via
:meth:`docx.settings.Settings.enable_mail_merge`. Adds a couple of extra
properties (``active_record``, ``view_merged_data``, ``link_to_query``) on top
so the feature scenarios can cover read-side coverage of every |MailMerge|
property.

Run ``python features/steps/test_files/_gen_mmg_enabled.py`` to regenerate the
fixture in-place and self-validate the result.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import (
    WD_MAIL_MERGE_DATA_TYPE,
    WD_MAIL_MERGE_DESTINATION,
    WD_MAIL_MERGE_TYPE,
)

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "mmg-enabled.docx")


def build() -> str:
    """Create ``mmg-enabled.docx`` and return its absolute path."""
    document = Document()
    document.add_paragraph("Dear <<FirstName>>,")

    mail_merge = document.settings.enable_mail_merge(
        main_document_type=WD_MAIL_MERGE_TYPE.EMAIL,
        destination=WD_MAIL_MERGE_DESTINATION.EMAIL,
        data_type=WD_MAIL_MERGE_DATA_TYPE.SPREADSHEET,
        connect_string="Provider=Microsoft.ACE.OLEDB.12.0;Data Source=contacts.xlsx",
        query="SELECT FirstName, Email FROM [Sheet1$]",
        mail_subject="Quarterly update",
        address_field_name="Email",
    )
    # -- additional fields exercised by the read-side scenarios --
    mail_merge.active_record = 3
    mail_merge.view_merged_data = True
    mail_merge.link_to_query = True

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    """Reload `path` and assert every mail-merge field round-trips."""
    document = Document(path)
    mm = document.settings.mail_merge
    assert mm is not None, "mail-merge element missing from saved fixture"
    assert mm.main_document_type == WD_MAIL_MERGE_TYPE.EMAIL
    assert mm.destination == WD_MAIL_MERGE_DESTINATION.EMAIL
    assert mm.data_type == WD_MAIL_MERGE_DATA_TYPE.SPREADSHEET
    assert mm.connect_string == (
        "Provider=Microsoft.ACE.OLEDB.12.0;Data Source=contacts.xlsx"
    )
    assert mm.query == "SELECT FirstName, Email FROM [Sheet1$]"
    assert mm.mail_subject == "Quarterly update"
    assert mm.address_field_name == "Email"
    assert mm.active_record == 3
    assert mm.view_merged_data is True
    assert mm.link_to_query is True
    # -- fields not written stay at their defaults --
    assert mm.check_errors is None
    assert mm.do_not_suppress_blank_lines is False
    assert mm.mail_as_attachment is False


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
