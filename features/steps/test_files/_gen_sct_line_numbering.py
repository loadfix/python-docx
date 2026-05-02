"""Generate ``sct-line-numbering.docx`` — fixture for line-numbering behave.

Three sections provide coverage for every branch of
:attr:`Section.line_numbering` / :meth:`Section.set_line_numbering` /
:meth:`Section.remove_line_numbering`:

- **Section 0** — no ``w:lnNumType`` element. Reads return ``None``. Scenarios
  use it as the starting state for :meth:`set_line_numbering`.
- **Section 1** — ``w:lnNumType`` populated with ``countBy=1``, ``start=1``,
  ``distance=Pt(20)``, ``restart=NEW_PAGE``. Scenarios assert reads and
  verify ``remove_line_numbering`` drops the element.
- **Section 2** — ``w:lnNumType`` with only ``countBy=5`` set. Exercises
  mixed ``None`` / set attribute reads.

Run this script from the repo root::

    python features/steps/test_files/_gen_sct_line_numbering.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_LINE_NUMBERING_RESTART, WD_SECTION
from docx.shared import Pt


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "sct-line-numbering.docx")


def build() -> None:
    document = Document()

    document.add_paragraph("Section 0 body — no line numbering.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 1 body — every-line numbering from 1.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 2 body — every fifth line only.")

    # -- populate only after all sections exist so clone-on-add_section does
    #    not propagate ``w:lnNumType`` settings into later sections. ---
    document.sections[0].remove_line_numbering()

    document.sections[1].set_line_numbering(
        count_by=1,
        start=1,
        distance=Pt(20),
        restart=WD_LINE_NUMBERING_RESTART.NEW_PAGE,
    )

    document.sections[2].remove_line_numbering()
    document.sections[2].set_line_numbering(count_by=5)

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    sections = document.sections
    assert len(sections) == 3

    assert sections[0].line_numbering is None, (
        "section 0 should have no line_numbering, got "
        f"{sections[0].line_numbering!r}"
    )

    ln1 = sections[1].line_numbering
    assert ln1 is not None, "section 1 line_numbering missing"
    assert ln1.count_by == 1, f"count_by={ln1.count_by!r}"
    assert ln1.start == 1, f"start={ln1.start!r}"
    assert ln1.distance == Pt(20), f"distance={ln1.distance!r}"
    assert ln1.restart == WD_LINE_NUMBERING_RESTART.NEW_PAGE, (
        f"restart={ln1.restart!r}"
    )

    ln2 = sections[2].line_numbering
    assert ln2 is not None, "section 2 line_numbering missing"
    assert ln2.count_by == 5, f"count_by={ln2.count_by!r}"
    assert ln2.start is None, f"start={ln2.start!r}"
    assert ln2.distance is None, f"distance={ln2.distance!r}"
    assert ln2.restart is None, f"restart={ln2.restart!r}"


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
