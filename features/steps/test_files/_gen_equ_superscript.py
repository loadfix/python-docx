"""Generate ``equ-superscript.docx`` fixture for OMML superscript scenarios.

Two paragraphs, each carrying a single ``m:oMath`` produced by
:func:`docx.equations.build_superscript`:

- Paragraph 0: identifier base with a numeric exponent — ``x`` raised to ``2``.
- Paragraph 1: identifier base with an identifier exponent — ``e`` raised to
  ``x``.

The script self-validates by reopening the saved document and asserting the
equation count, flattened text, and that each expression carries an
``m:sSup`` element.

Run ``python features/steps/test_files/_gen_equ_superscript.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.equations import build_superscript

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "equ-superscript.docx")


def build() -> str:
    document = Document()
    document.add_heading("Superscript fixture", level=1)

    # -- Paragraph 0: x ** 2 — identifier base + numeric exponent --
    p0 = document.add_paragraph("Quadratic term: ")
    p0.add_equation(build_superscript("x", "2"))

    # -- Paragraph 1: e ** x — identifier base + identifier exponent --
    p1 = document.add_paragraph("Exponential: ")
    p1.add_equation(build_superscript("e", "x"))

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    # -- every equation in the document body --
    equations = [eq for p in document.paragraphs for eq in p.equations]
    assert len(equations) == 2, f"expected 2 equations, got {len(equations)}"

    quadratic, exponential = equations

    assert quadratic.text == "x2", (
        f"expected 'x2', got {quadratic.text!r}"
    )
    assert b"<m:sSup>" in quadratic.raw_xml, (
        f"expected m:sSup in first equation, got {quadratic.raw_xml!r}"
    )
    assert quadratic.is_display_mode is False

    assert exponential.text == "ex", (
        f"expected 'ex', got {exponential.text!r}"
    )
    assert b"<m:sSup>" in exponential.raw_xml, (
        f"expected m:sSup in second equation, got {exponential.raw_xml!r}"
    )
    assert exponential.is_display_mode is False


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
