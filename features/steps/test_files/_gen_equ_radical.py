"""Generate ``equ-radical.docx`` fixture for OMML radical builder scenarios.

Builds a document containing three radical expressions exercising
:func:`docx.equations.build_radical`:

* paragraph 1 — a square root of an identifier, ``sqrt(x)``.
* paragraph 2 — an nth-root (cube root), ``3-root(y)``.
* paragraph 3 — a radical whose radicand is itself a radical (nested).

Run ``python features/steps/test_files/_gen_equ_radical.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.equations import build_radical

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "equ-radical.docx")


def _nested_radical_xml() -> str:
    """Return an ``m:oMath`` containing a radical whose radicand is a radical.

    The outer ``build_radical`` call puts ``build_radical("y")`` inside its
    ``m:e`` element. We splice the inner ``m:rad`` fragment into the outer
    XML by stripping the inner ``<m:oMath …>…</m:oMath>`` wrapper.
    """
    inner = build_radical("y")
    # -- strip ``<m:oMath xmlns:m="…">`` prefix and ``</m:oMath>`` suffix --
    inner_open_end = inner.index(">") + 1
    inner_body = inner[inner_open_end:-len("</m:oMath>")]
    # -- outer radical: square root of the inner radical element --
    from docx.equations import _omath_open  # type: ignore[attr-defined]

    return (
        "%s<m:rad><m:deg/><m:e>%s</m:e></m:rad></m:oMath>"
        % (_omath_open(), inner_body)
    )


def build() -> Document:
    document = Document()

    # -- paragraph 0: heading / intro so scenarios rely on fixed indices --
    document.add_heading("Radical fixture", level=1)

    # -- paragraph 1: square root of an identifier --
    sqrt_para = document.add_paragraph("square root: ")
    sqrt_para.add_equation(build_radical("x"))

    # -- paragraph 2: cube root (nth root with degree 3) --
    nth_para = document.add_paragraph("cube root: ")
    nth_para.add_equation(build_radical("y", degree_text="3"))

    # -- paragraph 3: nested radical — sqrt of sqrt(y) --
    nested_para = document.add_paragraph("nested: ")
    nested_para.add_equation(_nested_radical_xml())

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    sqrt_eqs = paragraphs[1].equations
    assert len(sqrt_eqs) == 1, (
        f"expected 1 equation on paragraph 1, got {len(sqrt_eqs)}"
    )
    assert b"<m:rad>" in sqrt_eqs[0].raw_xml
    assert sqrt_eqs[0].text == "x"
    # -- no degree run on a plain square root --
    assert b"<m:deg/>" in sqrt_eqs[0].raw_xml

    nth_eqs = paragraphs[2].equations
    assert len(nth_eqs) == 1
    assert b"<m:rad>" in nth_eqs[0].raw_xml
    # -- degree and radicand both contribute text --
    assert nth_eqs[0].text == "3y"

    nested_eqs = paragraphs[3].equations
    assert len(nested_eqs) == 1
    raw = nested_eqs[0].raw_xml
    # -- two m:rad elements signals nesting --
    assert raw.count(b"<m:rad>") == 2, (
        f"expected 2 m:rad elements, got {raw.count(b'<m:rad>')}"
    )
    assert nested_eqs[0].text == "y"


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
