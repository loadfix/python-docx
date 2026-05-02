"""Generate `shp-ole.docx` behave fixture.

Builds a document containing an embedded OLE object (a minimal fake Excel
workbook blob) referenced from a ``<w:object>/<o:OLEObject>`` child of a run.
Also includes a second paragraph whose ``OLEObject`` references a missing
relationship, so the read-side API's tolerance for unresolved references can
be exercised.

python-docx does not support creating embedded OLE objects via the public
API, so this generator builds the ``EmbeddedObjectPart`` objects directly and
injects the ``w:object`` XML into the paragraphs.

Run::

    python features/steps/test_files/_gen_shp_ole.py
"""

from __future__ import annotations

import os

from lxml import etree

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml.ns import nsmap, qn
from docx.parts.embedded_object import EmbeddedObjectPart

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "shp-ole.docx")


# -- 16 bytes of deterministic "fake" OLE payload --
FAKE_OLE_BLOB = b"FAKE-OLE-BLOB-01"


def _inject_ole_object(
    paragraph,
    rId: str,
    prog_id: str,
    ole_type: str = "Embed",
) -> None:
    """Append ``<w:r><w:object><o:OLEObject .../></w:object></w:r>`` to paragraph."""
    p = paragraph._p  # pyright: ignore[reportPrivateUsage]
    r = etree.SubElement(p, qn("w:r"))
    obj = etree.SubElement(r, qn("w:object"))
    ole = etree.SubElement(obj, qn("o:OLEObject"))
    ole.set("Type", ole_type)
    ole.set("ProgID", prog_id)
    ole.set(qn("r:id"), rId)
    # -- force namespace registration by reading --
    assert nsmap["o"] and nsmap["r"]


def build() -> Document:
    document = Document()
    document.add_heading("Embedded OLE objects fixture", level=1)

    p1 = document.add_paragraph("Paragraph with an embedded workbook: ")
    p2 = document.add_paragraph("Paragraph with a linked/unresolved OLE object: ")
    document.add_paragraph("Plain paragraph with no embedded object.")

    # -- build a proper EmbeddedObjectPart and relate it --
    doc_part = document.part
    ole_part = EmbeddedObjectPart(
        PackURI("/word/embeddings/oleObject1.bin"),
        CT.OFC_OLE_OBJECT,
        FAKE_OLE_BLOB,
    )
    rId_ok = doc_part.relate_to(ole_part, RT.OLE_OBJECT)
    _inject_ole_object(p1, rId_ok, prog_id="Excel.Sheet.12", ole_type="Embed")

    # -- second paragraph references a relationship id that does not exist --
    # -- the read-side API should still produce an EmbeddedObject whose blob=b"" --
    _inject_ole_object(p2, "rIdDoesNotExist", prog_id="AcroExch.Document", ole_type="Link")

    return document


def self_validate(document: Document) -> None:
    objects = document.embedded_objects
    assert len(objects) == 2, f"expected 2 embedded objects, got {len(objects)}"

    resolved = [o for o in objects if o.blob]
    unresolved = [o for o in objects if not o.blob]
    assert len(resolved) == 1, f"expected 1 resolved, got {len(resolved)}"
    assert len(unresolved) == 1, f"expected 1 unresolved, got {len(unresolved)}"

    ok = resolved[0]
    assert ok.prog_id == "Excel.Sheet.12", ok.prog_id
    assert ok.type == "Embed", ok.type
    assert ok.blob == FAKE_OLE_BLOB, "blob mismatch"
    assert ok.embedded_partname == "/word/embeddings/oleObject1.bin"
    assert ok.r_id, "r_id should be set"

    missing = unresolved[0]
    assert missing.prog_id == "AcroExch.Document", missing.prog_id
    assert missing.type == "Link", missing.type
    assert missing.blob == b"", "unresolved blob must be empty"
    assert missing.embedded_partname is None


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    reopened = Document(OUT_PATH)
    self_validate(reopened)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
