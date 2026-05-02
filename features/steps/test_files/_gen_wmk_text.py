"""Generate ``wmk-text.docx`` fixture for text-watermark behave scenarios.

Adds a ``DRAFT`` text watermark to the first section via
:meth:`Section.add_text_watermark`, then round-trips the document through
:class:`docx.Document` to confirm the watermark survives a save/load cycle.

Run ``python features/steps/test_files/_gen_wmk_text.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import Pt, RGBColor

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "wmk-text.docx")


def build() -> str:
    document = Document()
    document.add_paragraph("Body paragraph one.")
    document.add_paragraph("Body paragraph two.")

    section = document.sections[0]
    section.add_text_watermark(
        text="DRAFT",
        font="Arial",
        size=Pt(80),
        color=RGBColor(0x80, 0x80, 0x80),
        layout="diagonal",
    )

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    watermark = document.sections[0].watermark
    assert watermark is not None, "watermark missing after round-trip"
    assert watermark.type == "text", f"expected text watermark, got {watermark.type!r}"
    assert watermark.text == "DRAFT", f"expected 'DRAFT', got {watermark.text!r}"


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
