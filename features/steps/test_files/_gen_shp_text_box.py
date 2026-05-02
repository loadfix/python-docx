"""Generate `shp-text-box.docx` behave fixture.

Builds a document with two shapes whose ``wps:wsp`` carries a text frame:

* A single rectangle with a single-paragraph text frame.
* A rectangle whose text frame contains two paragraphs so the behave
  ``Drawing.text`` / ``Drawing.paragraphs`` scenarios can exercise
  multi-paragraph text extraction.

Self-validates after generation.

Run::

    python features/steps/test_files/_gen_shp_text_box.py
"""

from __future__ import annotations

import os

from lxml import etree

from docx import Document
from docx.enum.shape import WD_DRAWING_TYPE, WD_SHAPE
from docx.oxml.ns import qn
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "shp-text-box.docx")


def _append_paragraph(txbxContent, text: str) -> None:
    """Append a ``<w:p>`` with a single run of `text` to the txbxContent."""
    p = etree.SubElement(txbxContent, qn("w:p"))
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = text


def build() -> Document:
    document = Document()
    document.add_heading("Text box fixture", level=1)

    # -- paragraph 1: single-paragraph text-frame rectangle --
    p1 = document.add_paragraph()
    shape1 = p1.add_shape(
        WD_SHAPE.RECTANGLE,
        width=Inches(2.0),
        height=Inches(0.75),
        text="Single line",
    )
    assert shape1.text == "Single line"

    # -- paragraph 2: multi-paragraph text-frame rectangle (rebuild txbxContent) --
    p2 = document.add_paragraph()
    shape2 = p2.add_shape(
        WD_SHAPE.RECTANGLE,
        width=Inches(2.5),
        height=Inches(1.0),
        text="First line",
    )

    wsp = shape2._wsp  # pyright: ignore[reportPrivateUsage]
    txbx = wsp.find(qn("wps:txbx"))
    assert txbx is not None
    txbxContent = txbx.find(qn("w:txbxContent"))
    assert txbxContent is not None
    # -- remove the single default paragraph and append our own --
    for child in list(txbxContent):
        txbxContent.remove(child)
    _append_paragraph(txbxContent, "First line")
    _append_paragraph(txbxContent, "Second line")
    _append_paragraph(txbxContent, "Third line")

    document.add_paragraph("Trailing paragraph.")
    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs
    # -- paragraphs[0] heading, [1]=shape1, [2]=shape2 --
    d1 = paragraphs[1].drawings
    assert len(d1) == 1
    assert d1[0].type is WD_DRAWING_TYPE.TEXT_BOX
    assert d1[0].text == "Single line"
    p_objs = d1[0].paragraphs
    assert len(p_objs) == 1
    assert p_objs[0].text == "Single line"

    d2 = paragraphs[2].drawings
    assert len(d2) == 1
    assert d2[0].type is WD_DRAWING_TYPE.TEXT_BOX
    assert d2[0].text == "First line\nSecond line\nThird line", d2[0].text
    p_objs_2 = d2[0].paragraphs
    assert [p.text for p in p_objs_2] == [
        "First line",
        "Second line",
        "Third line",
    ]


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    reopened = Document(OUT_PATH)
    self_validate(reopened)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
