"""Build the ``trk-rsid.docx`` fixture for the revision-save ID tests.

Run::

    python features/steps/test_files/_gen_trk_rsid.py

*Generator strategy*: paragraph and run scaffolding uses the public
python-docx API. The ``w:rsidR`` attributes on individual paragraphs and runs
are set via raw ``OxmlElement.set()`` (no public setter exists — RSIDs are
normally assigned by Word itself). The document-level ``w:rsids`` /
``w:rsidRoot`` are injected into the settings part the same way.

Self-checks:

* ``document.settings.rsid_root`` returns the expected hex value
* ``document.settings.rsids`` returns the expected list in document order
* paragraph 1 has a non-None ``.rsid``; paragraph 2 has a |None| ``.rsid``
* paragraph 1's first run has a non-None ``.rsid``; paragraph 2's run is |None|
"""

from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-rsid.docx")

ROOT_RSID = "00CAFE00"
RSID_VALUES = ["00A1B2C3", "00DEAD00", "00BEEF00"]


def build() -> Document:
    document = Document()

    # -- paragraph 0: title --
    document.add_heading("Revision save IDs", level=1)

    # -- paragraph 1: carries an rsidR, first run carries its own rsidR --
    p1 = document.add_paragraph("First session.")
    p1._p.set(qn("w:rsidR"), "00A1B2C3")
    p1.runs[0]._r.set(qn("w:rsidR"), "00DEAD00")

    # -- paragraph 2: no rsidR anywhere --
    document.add_paragraph("No session markers here.")

    # -- paragraph 3: paragraph-level rsid set, run has none --
    p3 = document.add_paragraph("Only paragraph rsid.")
    p3._p.set(qn("w:rsidR"), "00BEEF00")

    # -- inject w:rsids on the settings part --
    settings = document.settings._settings
    rsids = settings.rsids
    if rsids is not None:
        # -- remove the empty default so we can replace wholesale --
        settings.remove(rsids)
    settings._add_rsids()
    rsids = settings.rsids
    assert rsids is not None

    rsidRoot = OxmlElement("w:rsidRoot")
    rsidRoot.set(qn("w:val"), ROOT_RSID)
    rsids.append(rsidRoot)
    for val in RSID_VALUES:
        rsid = OxmlElement("w:rsid")
        rsid.set(qn("w:val"), val)
        rsids.append(rsid)

    return document


def self_validate(document: Document) -> None:
    assert document.settings.rsid_root == ROOT_RSID, (
        f"expected rsid_root={ROOT_RSID!r}, got {document.settings.rsid_root!r}"
    )
    assert document.settings.rsids == RSID_VALUES, (
        f"expected rsids={RSID_VALUES!r}, got {document.settings.rsids!r}"
    )

    paragraphs = document.paragraphs
    assert paragraphs[1].rsid == "00A1B2C3"
    assert paragraphs[1].runs[0].rsid == "00DEAD00"
    assert paragraphs[2].rsid is None
    assert paragraphs[2].runs == [] or paragraphs[2].runs[0].rsid is None
    assert paragraphs[3].rsid == "00BEEF00"
    # -- run on paragraph 3 has no rsidR (just a plain run with text) --
    assert paragraphs[3].runs[0].rsid is None


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
