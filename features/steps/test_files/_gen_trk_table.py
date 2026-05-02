"""Build the ``trk-table.docx`` fixture for the cell/row tracked-change tests.

Run::

    python features/steps/test_files/_gen_trk_table.py

*Generator strategy*: the table and cell scaffolding is built with
python-docx's public ``add_table()`` API. The tracked-revision markers
(``w:cellIns``, ``w:cellDel``, ``w:tcPrChange``, ``w:trPrChange``,
``w:tblPrChange``) are injected via raw ``OxmlElement`` because python-docx
does not yet surface authoring methods for cell/row-level revisions.

The fixture contains a 2x2 table laid out as follows:

* cell (0,0) — carries a ``w:tcPrChange`` (Alice): its cell properties were
  revised
* cell (0,1) — carries a ``w:cellIns`` (Bob): the whole cell was inserted
* cell (1,0) — carries a ``w:cellDel`` (Carol): the whole cell was deleted
* cell (1,1) — plain (no tracked revision)
* row 1 — carries a ``w:trPrChange`` (Dave): row properties were revised
* the table — carries a ``w:tblPrChange`` (Eve): table properties were revised
"""

from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-table.docx")


def _make_cell_marker(tag: str, marker_id: int, author: str, date: str):
    elm = OxmlElement(tag)
    elm.set(qn("w:id"), str(marker_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    return elm


def _make_props_change(
    tag: str, change_id: int, author: str, date: str, inner_tag: str
):
    elm = OxmlElement(tag)
    elm.set(qn("w:id"), str(change_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    elm.append(OxmlElement(inner_tag))
    return elm


def build() -> Document:
    document = Document()
    document.add_heading("Tracked table changes", level=1)

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"

    # -- cell (0,0): tcPrChange by Alice --
    tc00 = table.cell(0, 0)._tc
    tcPr00 = tc00.get_or_add_tcPr()
    tcPr00.append(
        _make_props_change(
            "w:tcPrChange", 1, "Alice", "2025-04-10T09:00:00Z", "w:tcPr"
        )
    )

    # -- cell (0,1): cellIns by Bob --
    tc01 = table.cell(0, 1)._tc
    tcPr01 = tc01.get_or_add_tcPr()
    tcPr01.append(_make_cell_marker("w:cellIns", 2, "Bob", "2025-04-10T09:05:00Z"))

    # -- cell (1,0): cellDel by Carol --
    tc10 = table.cell(1, 0)._tc
    tcPr10 = tc10.get_or_add_tcPr()
    tcPr10.append(_make_cell_marker("w:cellDel", 3, "Carol", "2025-04-10T09:10:00Z"))

    # -- row 1: trPrChange by Dave --
    tr1 = table.rows[1]._tr
    trPr1 = tr1.get_or_add_trPr()
    trPr1.append(
        _make_props_change(
            "w:trPrChange", 4, "Dave", "2025-04-10T09:15:00Z", "w:trPr"
        )
    )

    # -- table-level: tblPrChange by Eve --
    tblPr = table._element.tblPr
    tblPr.append(
        _make_props_change(
            "w:tblPrChange", 5, "Eve", "2025-04-10T09:20:00Z", "w:tblPr"
        )
    )

    return document


def self_validate(document: Document) -> None:
    table = document.tables[0]

    # -- cell (0,0) exposes a tcPrChange --
    fc = table.cell(0, 0).formatting_change
    assert fc is not None and fc.author == "Alice"

    # -- cell (0,1) is marked as tracked insertion --
    assert table.cell(0, 1).is_tracked_insertion is True
    assert table.cell(0, 1).is_tracked_deletion is False

    # -- cell (1,0) is marked as tracked deletion --
    assert table.cell(1, 0).is_tracked_deletion is True
    assert table.cell(1, 0).is_tracked_insertion is False

    # -- cell (1,1) has no tracked markers --
    assert table.cell(1, 1).is_tracked_insertion is False
    assert table.cell(1, 1).is_tracked_deletion is False
    assert table.cell(1, 1).formatting_change is None

    # -- row 1 trPrChange --
    row_fc = table.rows[1].formatting_change
    assert row_fc is not None and row_fc.author == "Dave"
    # -- row 0 has no trPrChange --
    assert table.rows[0].formatting_change is None

    # -- table tblPrChange --
    tbl_fc = table.formatting_change
    assert tbl_fc is not None and tbl_fc.author == "Eve"


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
