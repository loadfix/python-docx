"""Generate `shp-ink.docx` behave fixture.

Builds a document whose body contains two paragraphs with ``<w:contentPart>``
run children, each referencing a separate ``word/ink/ink*.xml`` InkML part.

python-docx does not support creating ink annotations via the public API,
so this generator:

1. Builds an empty document via :class:`docx.Document`.
2. Constructs two :class:`docx.parts.ink.InkPart` parts containing valid
   InkML with differing ``inkml:trace`` counts.
3. Relates them to the document part via ``RT.INK``.
4. Injects ``<w:r><w:contentPart r:id="..."/></w:r>`` into two paragraphs.

Self-validates after generation by re-opening the file and asserting the
ink-annotation count, partnames, and stroke counts round-trip.

Run::

    python features/steps/test_files/_gen_shp_ink.py
"""

from __future__ import annotations

import os

from lxml import etree

from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.oxml.ns import nsmap, qn
from docx.parts.ink import InkPart

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "shp-ink.docx")


INK_TWO_TRACES = (
    b'<?xml version="1.0" encoding="UTF-8"?>\n'
    b'<inkml:ink xmlns:inkml="http://www.w3.org/2003/InkML">\n'
    b"  <inkml:trace>100 100, 150 150, 200 200</inkml:trace>\n"
    b"  <inkml:trace>300 300, 350 350, 400 400</inkml:trace>\n"
    b"</inkml:ink>\n"
)

INK_ONE_TRACE = (
    b'<?xml version="1.0" encoding="UTF-8"?>\n'
    b'<inkml:ink xmlns:inkml="http://www.w3.org/2003/InkML">\n'
    b"  <inkml:trace>50 50, 75 75, 100 100</inkml:trace>\n"
    b"</inkml:ink>\n"
)


def _inject_content_part(paragraph, rId: str) -> None:
    """Append a ``<w:r><w:contentPart r:id=rId/></w:r>`` to `paragraph`."""
    p = paragraph._p  # pyright: ignore[reportPrivateUsage]
    r = etree.SubElement(p, qn("w:r"))
    cp = etree.SubElement(r, qn("w:contentPart"))
    cp.set(qn("r:id"), rId)
    # -- suppress linter warnings by forcing namespace registration --
    assert nsmap["r"]


def build() -> Document:
    document = Document()
    document.add_heading("Ink annotations fixture", level=1)

    # -- create two paragraphs that will carry ink refs --
    p1 = document.add_paragraph("Paragraph with first ink annotation: ")
    p2 = document.add_paragraph("Paragraph with second ink annotation: ")
    document.add_paragraph("Plain paragraph with no ink.")

    # -- register two InkPart parts and relate them to the document part --
    doc_part = document.part
    package = doc_part.package

    ink_part_1 = InkPart(PackURI("/word/ink/ink1.xml"), CT.INKML, INK_TWO_TRACES)
    ink_part_2 = InkPart(PackURI("/word/ink/ink2.xml"), CT.INKML, INK_ONE_TRACE)

    # -- the package.iter_parts only finds parts that are reachable through --
    # -- relationships, so simply relating them from the document part is enough --
    rId_1 = doc_part.relate_to(ink_part_1, RT.INK)
    rId_2 = doc_part.relate_to(ink_part_2, RT.INK)

    _inject_content_part(p1, rId_1)
    _inject_content_part(p2, rId_2)

    # -- suppress unused warning --
    assert package is not None
    return document


def self_validate(document: Document) -> None:
    annotations = document.ink_annotations
    assert len(annotations) == 2, f"expected 2 ink annotations, got {len(annotations)}"

    stroke_counts = sorted(a.stroke_count for a in annotations)
    assert stroke_counts == [1, 2], f"got stroke_counts={stroke_counts}"

    partnames = sorted(a.partname for a in annotations)
    assert partnames == ["/word/ink/ink1.xml", "/word/ink/ink2.xml"], partnames

    # -- per-paragraph access --
    paragraphs = document.paragraphs
    assert len(paragraphs[1].ink_annotations) == 1
    assert len(paragraphs[2].ink_annotations) == 1
    assert len(paragraphs[3].ink_annotations) == 0


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    reopened = Document(OUT_PATH)
    self_validate(reopened)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
