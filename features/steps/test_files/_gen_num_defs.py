"""Generate ``num-defs.docx`` fixture for numbering behave scenarios.

Creates a document with a custom three-level numbering definition applied to
three paragraphs, one per indent level. The first level uses a non-default
start value (``5``) so scenarios can assert that ``Level.start`` round-trips.

Run ``python features/steps/test_files/_gen_num_defs.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_NUMBER_FORMAT
from docx.shared import Inches

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "num-defs.docx")


BULLET = "•"


def build() -> str:
    document = Document()
    p1 = document.add_paragraph("First (decimal, starts at 5)")
    p2 = document.add_paragraph("Second (lowerLetter)")
    p3 = document.add_paragraph("Third (bullet)")

    definition = document.numbering.add_numbering_definition(
        [
            {
                "format": WD_NUMBER_FORMAT.DECIMAL,
                "text": "%1.",
                "indent": Inches(0.25),
                "start": 5,
            },
            {
                "format": "lowerLetter",
                "text": "%2)",
                "indent": Inches(0.5),
            },
            {
                "format": "bullet",
                "text": BULLET,
                "indent": Inches(0.75),
                "font": "Symbol",
            },
        ]
    )
    definition.apply_to(p1, level=0)
    definition.apply_to(p2, level=1)
    definition.apply_to(p3, level=2)

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    numbering = document.numbering
    # -- the default template already ships several definitions; the one we
    # -- added is the last in document order. --
    assert len(numbering.definitions) >= 1
    our_def = numbering.definitions[-1]
    assert len(our_def.levels) == 3, f"got {len(our_def.levels)} levels"

    lvl0, lvl1, lvl2 = our_def.levels
    assert lvl0.number_format == WD_NUMBER_FORMAT.DECIMAL
    assert lvl0.text == "%1."
    assert lvl0.start == 5
    assert lvl1.number_format == WD_NUMBER_FORMAT.LOWER_LETTER
    assert lvl1.text == "%2)"
    assert lvl2.number_format == WD_NUMBER_FORMAT.BULLET
    assert lvl2.text == BULLET

    # -- paragraphs should reference the numbering instance --
    paragraph_has_num = [
        p._p.xpath(".//w:numPr") != [] for p in document.paragraphs
    ]
    assert paragraph_has_num[:3] == [True, True, True], paragraph_has_num


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
