"""Generate ``pkg-malformed.docx`` fixture for recover-mode scenarios.

Starts from a minimal, valid ``.docx`` built via :class:`docx.Document`, then
truncates ``word/document.xml`` so it ends mid-tag. The resulting package is
unreadable via the default parser (lxml raises :class:`lxml.etree.XMLSyntaxError`)
but recoverable via ``Document(..., recover=True)``, which discards the
unrecoverable trailing fragment and records the parse failure on
:attr:`~docx.document.Document.recovery_warnings`.

Run ``python features/steps/test_files/_gen_pkg_recover.py`` to regenerate.
"""

from __future__ import annotations

import os
import shutil
import tempfile
import zipfile

from docx import Document
from lxml import etree

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "pkg-malformed.docx")


def _make_base_docx(path: str) -> None:
    """Write a minimal valid ``.docx`` to `path`."""
    document = Document()
    document.add_paragraph("Readable prefix paragraph.")
    document.save(path)


def _corrupt(src: str, dst: str) -> None:
    """Copy `src` to `dst`, truncating ``word/document.xml`` mid-tag."""
    with zipfile.ZipFile(src, "r") as zi, zipfile.ZipFile(
        dst, "w", zipfile.ZIP_DEFLATED
    ) as zo:
        for info in zi.infolist():
            data = zi.read(info.filename)
            if info.filename == "word/document.xml":
                # -- lop off ``</w:body></w:document>`` and append a half-open
                # -- run so lxml reports an unterminated tag --
                end = data.find(b"</w:body>")
                assert end != -1
                data = data[:end] + b"<w:p><w:r><w:t>truncated"
            zo.writestr(info, data)


def build() -> str:
    base_fd, base_path = tempfile.mkstemp(suffix=".docx")
    os.close(base_fd)
    try:
        _make_base_docx(base_path)
        _corrupt(base_path, OUT_PATH)
    finally:
        if os.path.exists(base_path):
            os.remove(base_path)
    return OUT_PATH


def validate(path: str) -> None:
    # -- default mode must raise so the recovery-mode assertion is meaningful --
    try:
        Document(path)
    except etree.XMLSyntaxError:
        pass
    else:  # pragma: no cover - only hit if the corruption didn't take
        raise AssertionError("expected default open to raise XMLSyntaxError")

    document = Document(path, recover=True)
    assert document.recovery_warnings, "expected at least one recovery warning"
    # -- the readable prefix paragraph should still be accessible --
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Readable prefix paragraph." in text, text


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
