"""Generate the `tbl-banded.docx` behave fixture.

Running this script (idempotently) produces a .docx containing four
3x3 tables with different ``w:tblLook`` flag configurations. These
exercise the :class:`docx.table.TableStyleFlags` proxy:

- Table 0 (no explicit flags)  — all flags read as |False|.
- Table 1 (first-row emphasis) — only ``firstRow`` is set.
- Table 2 (banded rows on)     — ``firstRow`` + ``firstColumn`` are set,
                                  and ``noHBand`` is cleared (banded rows
                                  are active).
- Table 3 (banded rows off)    — ``firstRow`` + ``firstColumn`` set, and
                                  ``noHBand`` explicitly set to |True|
                                  (banded rows suppressed).

The script self-validates the written document before considering the
fixture good.
"""

from __future__ import annotations

import os

from docx import Document


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "tbl-banded.docx")


def _clear_tbl_look(table) -> None:
    """Remove ``w:tblLook`` entirely from a table's ``w:tblPr``.

    New tables get a default ``w:tblLook`` with ``firstRow="1"``
    written by python-docx; we remove it so table 0 reads as
    "no flags set".
    """
    tblPr = table._tblPr  # pyright: ignore[reportPrivateUsage]
    tblLook = tblPr.tblLook
    if tblLook is not None:
        tblPr.remove(tblLook)


def _fill(table) -> None:
    for row in range(3):
        for col in range(3):
            table.cell(row, col).text = f"{row}-{col}"


def build() -> None:
    document = Document()

    t0 = document.add_table(rows=3, cols=3)
    _clear_tbl_look(t0)
    _fill(t0)

    document.add_paragraph()

    t1 = document.add_table(rows=3, cols=3)
    _clear_tbl_look(t1)
    t1.style_flags.first_row = True
    _fill(t1)

    document.add_paragraph()

    t2 = document.add_table(rows=3, cols=3)
    _clear_tbl_look(t2)
    t2.style_flags.first_row = True
    t2.style_flags.first_column = True
    # -- noHBand left False: row-banding active --
    _fill(t2)

    document.add_paragraph()

    t3 = document.add_table(rows=3, cols=3)
    _clear_tbl_look(t3)
    t3.style_flags.first_row = True
    t3.style_flags.first_column = True
    t3.style_flags.no_horizontal_banding = True
    _fill(t3)

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    assert len(document.tables) == 4, f"expected 4 tables, got {len(document.tables)}"

    t0, t1, t2, t3 = document.tables

    # -- t0: all flags False --
    for flag in (
        "first_row",
        "last_row",
        "first_column",
        "last_column",
        "no_horizontal_banding",
        "no_vertical_banding",
    ):
        assert getattr(t0.style_flags, flag) is False, f"t0.{flag}"

    # -- t1: only first_row --
    assert t1.style_flags.first_row is True
    assert t1.style_flags.first_column is False
    assert t1.style_flags.no_horizontal_banding is False

    # -- t2: first_row + first_column; banding active (noHBand False) --
    assert t2.style_flags.first_row is True
    assert t2.style_flags.first_column is True
    assert t2.style_flags.no_horizontal_banding is False

    # -- t3: first_row + first_column; banding suppressed (noHBand True) --
    assert t3.style_flags.first_row is True
    assert t3.style_flags.first_column is True
    assert t3.style_flags.no_horizontal_banding is True


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
