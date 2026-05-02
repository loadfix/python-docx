"""Generate ``cap-caption.docx`` fixture for caption-paragraph scenarios.

Appends two captions — one ``Figure`` and one ``Table`` — using
:meth:`Document.add_caption`. The scenarios read these back to verify that a
caption paragraph is styled ``Caption`` and contains the literal
``"{label} N: {text}"`` text that :func:`docx.captions.new_caption_paragraph`
emits.

Run ``python features/steps/test_files/_gen_cap_caption.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "cap-caption.docx")


def build() -> str:
    document = Document()
    document.add_paragraph("Introductory text.")
    document.add_caption("A diagram of the system", label="Figure")
    document.add_caption("A reference table", label="Table")
    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    caption_paragraphs = [
        p for p in document.paragraphs if p.style.name == "Caption"
    ]
    assert len(caption_paragraphs) == 2, (
        f"expected 2 caption paragraphs, got {len(caption_paragraphs)}"
    )
    assert caption_paragraphs[0].text == "Figure 1: A diagram of the system"
    assert caption_paragraphs[1].text == "Table 1: A reference table"


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
