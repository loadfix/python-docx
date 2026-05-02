"""Generate `shp-alt-text.docx` behave fixture.

Builds a document containing both an inline picture and a floating picture,
each carrying an accessibility ``alt_text`` (``descr``) and ``title``
attribute on the ``wp:docPr``. This fixture drives the accessibility
scenarios in ``shp-alt-text.feature``.

Run::

    python features/steps/test_files/_gen_shp_alt_text.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.shape import WD_ANCHOR_H, WD_ANCHOR_V, WD_WRAP_TYPE
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "shp-alt-text.docx")
IMAGE_PATH = os.path.join(HERE, "monty-truth.png")


# -- (image index, paragraph prefix, alt_text, title) --
INLINE_SPECS = [
    ("An inline picture with alt text: ", "A pencil-drawing of a mountain peak", "Mountain peak"),
    # -- inline picture without alt_text so read-side returns None --
    ("An inline picture without alt text: ", None, None),
]


FLOATING_SPECS = [
    (
        "A floating picture with alt text and title: ",
        "Decorative floating mountain",
        "Floating mountain",
    ),
    # -- floating without alt_text to exercise None handling --
    ("A floating picture without alt text: ", None, None),
]


def build() -> Document:
    document = Document()
    document.add_heading("Alt text fixture", level=1)

    # -- inline pictures --
    for prefix, alt, title in INLINE_SPECS:
        p = document.add_paragraph(prefix)
        run = p.add_run()
        shape = run.add_picture(IMAGE_PATH, width=Inches(1.0))
        if alt is not None:
            shape.alt_text = alt
        if title is not None:
            shape.title = title

    # -- floating pictures --
    for prefix, alt, title in FLOATING_SPECS:
        p = document.add_paragraph(prefix)
        float_img = p.add_floating_image(
            IMAGE_PATH,
            width=Inches(1.0),
            position={
                "h_anchor": WD_ANCHOR_H.COLUMN,
                "v_anchor": WD_ANCHOR_V.PARAGRAPH,
                "horizontal": Inches(0.5),
                "vertical": Inches(0.25),
                "wrap": WD_WRAP_TYPE.SQUARE,
            },
        )
        if alt is not None:
            float_img.alt_text = alt
        if title is not None:
            float_img.title = title

    return document


def self_validate(document: Document) -> None:
    # -- inline shapes --
    inlines = document.inline_shapes
    assert len(inlines) == len(INLINE_SPECS), (
        f"expected {len(INLINE_SPECS)} inline shapes, got {len(inlines)}"
    )
    for shape, (_, expected_alt, expected_title) in zip(inlines, INLINE_SPECS):
        assert shape.alt_text == expected_alt, (
            f"inline alt_text mismatch: {shape.alt_text!r} != {expected_alt!r}"
        )
        assert shape.title == expected_title, (
            f"inline title mismatch: {shape.title!r} != {expected_title!r}"
        )

    # -- floating images --
    floating = [
        fi
        for para in document.paragraphs
        for fi in para.floating_images
    ]
    assert len(floating) == len(FLOATING_SPECS), (
        f"expected {len(FLOATING_SPECS)} floating images, got {len(floating)}"
    )
    for fi, (_, expected_alt, expected_title) in zip(floating, FLOATING_SPECS):
        assert fi.alt_text == expected_alt, (
            f"floating alt_text mismatch: {fi.alt_text!r} != {expected_alt!r}"
        )
        assert fi.title == expected_title, (
            f"floating title mismatch: {fi.title!r} != {expected_title!r}"
        )


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    reopened = Document(OUT_PATH)
    self_validate(reopened)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
