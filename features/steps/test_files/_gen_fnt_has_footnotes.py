"""Generate the ``fnt-has-footnotes.docx`` acceptance-test fixture.

Creates a small document with three body paragraphs, each referencing a
user footnote. The resulting file is self-validated below before being
saved — running this script both (re)builds the fixture and asserts the
shape the behave scenarios expect.

Run with::

    uv run python features/steps/test_files/_gen_fnt_has_footnotes.py
"""

from __future__ import annotations

import os

from docx import Document

THIS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(THIS_DIR, "fnt-has-footnotes.docx")


def build() -> None:
    document = Document()

    p1 = document.add_paragraph("The rain in Spain falls mainly in the plain.")
    document.footnotes.add(p1.runs[0], "A common saying about Iberian weather.")

    p2 = document.add_paragraph("Python-docx supports footnotes natively.")
    document.footnotes.add(p2.runs[0], "As of the loadfix fork.")

    p3 = document.add_paragraph("Each footnote has a stable integer id.")
    document.footnotes.add(p3.runs[0], "Ids 0 and 1 are reserved for separators.")

    # -- self-validate before saving so the fixture is known-good at build time --
    assert len(document.paragraphs) == 3, "expected 3 body paragraphs"
    assert len(document.footnotes) == 3, "expected 3 user footnotes"

    fn_by_id = {fn.footnote_id: fn for fn in document.footnotes}
    assert sorted(fn_by_id) == [2, 3, 4], f"expected ids [2, 3, 4], got {sorted(fn_by_id)}"
    assert fn_by_id[2].text == "A common saying about Iberian weather."
    assert fn_by_id[3].text == "As of the loadfix fork."
    assert fn_by_id[4].text == "Ids 0 and 1 are reserved for separators."

    for fn in document.footnotes:
        assert len(fn.paragraphs) == 1, "each footnote should have one paragraph"
        assert fn.paragraphs[0]._p.style == "FootnoteText"  # pyright: ignore[reportPrivateUsage]

    document.save(OUT_PATH)

    # -- round-trip check: reopen the saved file and re-validate --
    reopened = Document(OUT_PATH)
    assert len(reopened.footnotes) == 3, "round-trip footnote count mismatch"
    assert [fn.footnote_id for fn in reopened.footnotes] == [2, 3, 4]


if __name__ == "__main__":
    build()
    print(f"Wrote {OUT_PATH}")
