"""Build the ``trk-reject.docx`` fixture for the tracked-change reject tests.

Run from anywhere (the output path is resolved relative to this file)::

    python features/steps/test_files/_gen_trk_reject.py

*Generator strategy*: python-docx's public API lays down paragraphs, runs, and a
table; raw ``OxmlElement`` injection adds the tracked-change wrappers
(``w:ins``, ``w:del``, ``w:delText``, ``w:moveFrom``, ``w:moveTo``,
``w:cellIns``, ``w:cellDel``) because the authoring side of those elements is
not yet surfaced by the Document API.

Fixture layout (paragraph indices below are stable and the reject tests rely
on them):

* paragraph 0 — heading (no tracked changes)
* paragraph 1 — insertion + deletion interleaved in one paragraph
  ("The quick [-brown-][+nimble+] fox jumps.")
* paragraph 2 — a multi-run insertion by Carol — three ``w:r`` children inside
  the same ``w:ins`` wrapper, so "reject" must remove every one
* paragraph 3 — an entire paragraph's body is wrapped in ``w:del`` so the
  "reject a deletion that removes an entire paragraph" scenario restores every
  word
* paragraph 4 — a plain paragraph with no tracked changes (control)
* paragraph 5/6 — a move revision pair (name ``mv1``): paragraph 5 is the
  ``w:moveFrom`` source, paragraph 6 the ``w:moveTo`` destination
* a 2x2 table follows with:
    * cell (0, 0) — a ``w:cellIns`` marker by Bob (reject => cell removed)
    * cell (0, 1) — plain
    * cell (1, 0) — a ``w:cellDel`` marker by Carol (reject => cell kept)
    * cell (1, 1) — plain

Self-checks round-trip a copy through reject_all_changes() to confirm the
reject semantics hold end-to-end.
"""

from __future__ import annotations

import io
import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-reject.docx")


def _make_run(text: str):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _make_del_run(text: str):
    r = OxmlElement("w:r")
    dt_elm = OxmlElement("w:delText")
    dt_elm.text = text
    dt_elm.set(qn("xml:space"), "preserve")
    r.append(dt_elm)
    return r


def _make_ins(change_id: int, author: str, date: str, text: str):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(change_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    ins.append(_make_run(text))
    return ins


def _make_multi_run_ins(
    change_id: int, author: str, date: str, texts: list[str]
):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(change_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    for text in texts:
        ins.append(_make_run(text))
    return ins


def _make_del(change_id: int, author: str, date: str, text: str):
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:id"), str(change_id))
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), date)
    wdel.append(_make_del_run(text))
    return wdel


def _make_whole_paragraph_del(
    change_id: int, author: str, date: str, texts: list[str]
):
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:id"), str(change_id))
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), date)
    for text in texts:
        wdel.append(_make_del_run(text))
    return wdel


def _make_move_from(change_id: int, author: str, date: str, name: str, text: str):
    elm = OxmlElement("w:moveFrom")
    elm.set(qn("w:id"), str(change_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    elm.set(qn("w:name"), name)
    elm.append(_make_del_run(text))
    return elm


def _make_move_to(change_id: int, author: str, date: str, name: str, text: str):
    elm = OxmlElement("w:moveTo")
    elm.set(qn("w:id"), str(change_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    elm.set(qn("w:name"), name)
    elm.append(_make_run(text))
    return elm


def _make_cell_marker(tag: str, marker_id: int, author: str, date: str):
    elm = OxmlElement(tag)
    elm.set(qn("w:id"), str(marker_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    return elm


def build() -> Document:
    document = Document()

    # -- paragraph 0: title --
    document.add_heading("Tracked-change reject fixture", level=1)

    # -- paragraph 1: mixed insertion + deletion in one paragraph --
    p1 = document.add_paragraph("The quick ")
    p1._p.append(_make_del(1, "Bob", "2025-04-10T09:00:00Z", "brown"))
    p1._p.append(_make_ins(2, "Alice", "2025-04-10T09:05:00Z", "nimble"))
    p1.add_run(" fox jumps.")

    # -- paragraph 2: multi-run insertion by Carol --
    p2 = document.add_paragraph("Hello")
    p2._p.append(
        _make_multi_run_ins(
            3,
            "Carol",
            "2025-04-11T14:30:00Z",
            [", ", "cruel ", "world"],
        )
    )
    p2.add_run(".")

    # -- paragraph 3: whole paragraph body wrapped in w:del --
    p3 = document.add_paragraph()
    p3._p.append(
        _make_whole_paragraph_del(
            4,
            "Dave",
            "2025-04-12T08:15:00Z",
            ["This whole paragraph ", "is being deleted."],
        )
    )

    # -- paragraph 4: plain paragraph (control) --
    document.add_paragraph("Untouched paragraph.")

    # -- paragraph 5: move-source --
    p5 = document.add_paragraph("Source: ")
    p5._p.append(
        _make_move_from(5, "Eve", "2025-04-13T10:00:00Z", "mv1", "moved text")
    )
    p5.add_run(" (was here).")

    # -- paragraph 6: move-destination --
    p6 = document.add_paragraph("Destination: ")
    p6._p.append(
        _make_move_to(6, "Eve", "2025-04-13T10:00:00Z", "mv1", "moved text")
    )
    p6.add_run(" (now here).")

    # -- a 2x2 table with cellIns and cellDel markers --
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "A2"
    table.cell(1, 0).text = "B1"
    table.cell(1, 1).text = "B2"

    tc00 = table.cell(0, 0)._tc
    tcPr00 = tc00.get_or_add_tcPr()
    tcPr00.append(
        _make_cell_marker("w:cellIns", 7, "Bob", "2025-04-14T09:00:00Z")
    )

    tc10 = table.cell(1, 0)._tc
    tcPr10 = tc10.get_or_add_tcPr()
    tcPr10.append(
        _make_cell_marker("w:cellDel", 8, "Carol", "2025-04-14T09:05:00Z")
    )

    return document


def self_validate(document: Document) -> None:
    """Confirm the fixture's reject behavior end-to-end.

    Uses a deep-copied body so the on-disk fixture itself still carries every
    tracked change; the round-trip proves the reject-side semantics before we
    write the file.
    """
    paragraphs = document.paragraphs

    # -- paragraph 1 exposes deletion + insertion --
    p1_changes = paragraphs[1].tracked_changes
    assert len(p1_changes) == 2, (
        f"expected 2 tracked changes on paragraph 1, got {len(p1_changes)}"
    )
    types = sorted(tc.type for tc in p1_changes)
    assert types == ["deletion", "insertion"], (
        f"expected ['deletion', 'insertion'], got {types}"
    )

    # -- paragraph 2 has a single multi-run insertion --
    p2_changes = paragraphs[2].tracked_changes
    assert len(p2_changes) == 1
    assert p2_changes[0].type == "insertion"
    assert p2_changes[0].author == "Carol"
    assert p2_changes[0].text == ", cruel world"

    # -- paragraph 3 has a single deletion wrapping the whole paragraph --
    p3_changes = paragraphs[3].tracked_changes
    assert len(p3_changes) == 1
    assert p3_changes[0].type == "deletion"
    assert p3_changes[0].text == "This whole paragraph is being deleted."

    # -- paragraph 4 is plain --
    assert paragraphs[4].tracked_changes == []

    # -- move pair present --
    assert paragraphs[5].tracked_changes[0].type == "move_from"
    assert paragraphs[6].tracked_changes[0].type == "move_to"

    # -- table cell markers --
    table = document.tables[0]
    assert table.cell(0, 0).is_tracked_insertion is True
    assert table.cell(1, 0).is_tracked_deletion is True

    # -- round-trip: save to an in-memory buffer, reopen, reject all changes,
    # -- and confirm the reject semantics hold end-to-end. This leaves the
    # -- original `document` (with tracked changes intact) free to be written
    # -- to disk by main(). --
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    clone = Document(buf)

    count = clone.reject_all_changes()
    # 2 (paragraph 1) + 1 (paragraph 2) + 1 (paragraph 3) + 2 (move pair) +
    # 2 (cell markers) == 8
    assert count == 8, f"expected 8 resolved changes, got {count}"

    r_paragraphs = clone.paragraphs
    # reject insertion => remove, reject deletion => restore
    assert r_paragraphs[1].text == "The quick brown fox jumps.", (
        f"paragraph 1 mismatch after reject: {r_paragraphs[1].text!r}"
    )
    # reject multi-run insertion => remove every inserted run
    assert r_paragraphs[2].text == "Hello.", (
        f"paragraph 2 mismatch after reject: {r_paragraphs[2].text!r}"
    )
    # reject whole-paragraph deletion => restore every fragment
    assert r_paragraphs[3].text == "This whole paragraph is being deleted.", (
        f"paragraph 3 mismatch after reject: {r_paragraphs[3].text!r}"
    )
    # paragraph 4 unchanged
    assert r_paragraphs[4].text == "Untouched paragraph."
    # reject move: source restored, destination removed
    assert r_paragraphs[5].text == "Source: moved text (was here)."
    assert r_paragraphs[6].text == "Destination:  (now here)."

    # reject cellIns => cell removed, reject cellDel => cell kept
    r_table = clone.tables[0]
    # row 0 now has a single cell (the inserted one was rejected away)
    assert len(r_table.rows[0].cells) == 1, (
        f"expected 1 cell in row 0 after reject, got"
        f" {len(r_table.rows[0].cells)}"
    )
    # row 1 still has both cells (the cellDel was rejected, cell kept)
    assert len(r_table.rows[1].cells) == 2, (
        f"expected 2 cells in row 1 after reject, got"
        f" {len(r_table.rows[1].cells)}"
    )
    # -- original still has all tracked changes intact --
    assert len(document.paragraphs[1].tracked_changes) == 2


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
