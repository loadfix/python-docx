"""Generate ``srh-target-text.docx`` — fixture for search/replace scenarios.

Builds a document containing distinct marker phrases in every searchable story so
that behave scenarios can assert which story a match came from:

- body paragraphs   -> "SEARCH_IN_BODY"
- body table cell   -> "SEARCH_IN_TABLE"
- section header    -> "SEARCH_IN_HEADER"
- section footer    -> "SEARCH_IN_FOOTER"
- footnote content  -> "SEARCH_IN_FOOTNOTE"

Also seeds body paragraphs where one target phrase is broken across multiple runs
so that the run-span behaviour of :class:`docx.search.SearchMatch` can be
demonstrated in scenarios and docs.

Run directly to (re)generate the fixture at::

    features/steps/test_files/srh-target-text.docx

The script self-validates by reopening the file and cross-checking that each
marker is found by ``Document.search_all``.
"""

from __future__ import annotations

import os
import sys

from docx import Document


THIS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(THIS_DIR, "srh-target-text.docx")


def _build_document() -> Document:
    document = Document()

    # -- body paragraphs -----------------------------------------------------
    document.add_heading("Search target fixture", level=0)

    document.add_paragraph(
        "The first body paragraph contains SEARCH_IN_BODY as a whole phrase."
    )

    # -- a paragraph whose target phrase is split across three runs so that
    #    the SearchMatch.run_indices behaviour can be asserted.
    split = document.add_paragraph("")
    split.add_run("Multi-run marker: SEARCH_")
    split.add_run("IN_")
    split.add_run("BODY trails here.")

    # -- second plain occurrence to exercise multi-match return values.
    document.add_paragraph(
        "A second body line: SEARCH_IN_BODY appears here too."
    )

    # -- a regex target ------------------------------------------------------
    document.add_paragraph("Invoice INV-12345 was issued; ref INV-99 is older.")

    # -- a case-insensitive target ------------------------------------------
    document.add_paragraph("Mixed case: search_in_body lowercased is still findable.")

    # -- body table ----------------------------------------------------------
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "SEARCH_IN_TABLE appears in this cell."
    table.cell(1, 1).text = "Plain cell content."

    # -- header / footer (primary definition on first section) --------------
    section = document.sections[0]
    section.header.paragraphs[0].text = (
        "Running head: SEARCH_IN_HEADER goes at the top."
    )
    section.footer.paragraphs[0].text = (
        "Running foot: SEARCH_IN_FOOTER goes at the bottom."
    )

    # -- footnote ------------------------------------------------------------
    anchor_paragraph = document.add_paragraph(
        "A paragraph anchoring a footnote reference "
    )
    anchor_run = anchor_paragraph.add_run("here.")
    document.footnotes.add(
        anchor_run,
        text="Footnote body: SEARCH_IN_FOOTNOTE lives down here.",
    )

    return document


def _validate(path: str) -> None:
    """Re-open the saved document and confirm every marker is discoverable."""
    from docx.search import SearchMatch

    document = Document(path)

    expected_locations = {
        "SEARCH_IN_BODY": "body",
        "SEARCH_IN_TABLE": "table:0:row:1:col:0",
        "SEARCH_IN_HEADER": "header:section0:primary",
        "SEARCH_IN_FOOTER": "footer:section0:primary",
        "SEARCH_IN_FOOTNOTE": "footnote:",  # footnote id varies, prefix check
    }

    for marker, location_prefix in expected_locations.items():
        matches: list[SearchMatch] = document.search_all(marker)
        assert matches, f"marker {marker!r} produced no matches"
        locations = {m.location for m in matches}
        assert any(
            (loc or "").startswith(location_prefix) for loc in locations
        ), f"marker {marker!r} not found under expected story {location_prefix!r}; got {locations!r}"

    # The split-run body paragraph should produce a match spanning >= 2 runs.
    split_run_match = next(
        (
            m
            for m in document.search_all("SEARCH_IN_BODY")
            if len(m.run_indices) >= 2
        ),
        None,
    )
    assert split_run_match is not None, (
        "expected at least one SEARCH_IN_BODY match to span multiple runs"
    )


def main() -> int:
    document = _build_document()
    document.save(OUT_PATH)
    _validate(OUT_PATH)
    print(f"wrote {OUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
