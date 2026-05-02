"""Generate ``end-has-endnotes.docx`` fixture for endnotes behave scenarios.

Produces a document with three user endnotes (ids 2, 3, 4), anchored to three
paragraphs in the body, with varied text content so read/iterate/mutate scenarios
have meaningful data to assert against. Self-validates after building and before
writing the file to disk.

Run directly:

    python features/steps/test_files/_gen_end_has_endnotes.py
"""

from __future__ import annotations

import os
import sys

# -- add src/ to sys.path so this script can be run from anywhere -----------------------
THIS_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.abspath(os.path.join(THIS_DIR, "..", "..", ".."))
SRC_DIR = os.path.join(REPO_ROOT, "src")
if SRC_DIR not in sys.path:
    sys.path.insert(0, SRC_DIR)

from docx import Document  # noqa: E402
from docx.enum.text import (  # noqa: E402
    WD_ENDNOTE_POSITION,
    WD_FOOTNOTE_RESTART,
    WD_NUMBER_FORMAT,
)

OUT_PATH = os.path.join(THIS_DIR, "end-has-endnotes.docx")

ENDNOTE_TEXTS = (
    "First endnote citation.",
    "Second endnote with a longer explanatory comment.",
    "Third endnote referring to appendix A.",
)


def build() -> None:
    document = Document()

    # -- body: three paragraphs, each anchoring one endnote --
    for idx, endnote_text in enumerate(ENDNOTE_TEXTS, start=1):
        para = document.add_paragraph(f"Body paragraph {idx}.")
        run = para.runs[-1]
        document.endnotes.add(run, text=endnote_text)

    # -- document-level endnote properties for read-scenarios to assert against --
    props = document.add_endnote_properties()
    props.number_format = WD_NUMBER_FORMAT.LOWER_ROMAN
    props.restart_rule = WD_FOOTNOTE_RESTART.CONTINUOUS
    props.position = WD_ENDNOTE_POSITION.END_OF_DOCUMENT
    props.start_number = 1

    # -- validate before saving ---------------------------------------------------------
    endnotes = list(document.endnotes)
    assert len(endnotes) == 3, f"expected 3 endnotes, got {len(endnotes)}"
    assert [en.endnote_id for en in endnotes] == [2, 3, 4], (
        f"expected ids [2, 3, 4], got {[en.endnote_id for en in endnotes]}"
    )
    assert [en.text for en in endnotes] == list(ENDNOTE_TEXTS), (
        "endnote text did not round-trip through build"
    )

    ep = document.endnote_properties
    assert ep is not None, "endnote_properties should be present"
    assert ep.number_format == WD_NUMBER_FORMAT.LOWER_ROMAN
    assert ep.restart_rule == WD_FOOTNOTE_RESTART.CONTINUOUS
    assert ep.position == WD_ENDNOTE_POSITION.END_OF_DOCUMENT
    assert ep.start_number == 1

    document.save(OUT_PATH)

    # -- re-load and re-validate from disk ---------------------------------------------
    reloaded = Document(OUT_PATH)
    reloaded_endnotes = list(reloaded.endnotes)
    assert len(reloaded_endnotes) == 3
    assert [en.text for en in reloaded_endnotes] == list(ENDNOTE_TEXTS)

    rp = reloaded.endnote_properties
    assert rp is not None
    assert rp.number_format == WD_NUMBER_FORMAT.LOWER_ROMAN
    assert rp.restart_rule == WD_FOOTNOTE_RESTART.CONTINUOUS
    assert rp.position == WD_ENDNOTE_POSITION.END_OF_DOCUMENT
    assert rp.start_number == 1

    print(f"wrote {OUT_PATH} ({len(reloaded_endnotes)} endnotes)")


if __name__ == "__main__":
    build()
