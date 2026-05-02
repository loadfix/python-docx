"""Generate ``sct-page-borders.docx`` — fixture for page-border behave tests.

The fixture has three sections so scenarios can read, mutate, and remove
borders from known starting states:

- **Section 0** — no ``w:pgBorders`` element. Scenarios verify that reads
  return ``None`` attributes and that :meth:`Section.set_page_border` creates
  the element lazily.
- **Section 1** — all four edges pre-populated with a single 1pt red border
  offset from the page edge, ``display="allPages"``, ``offset_from="page"``.
  Scenarios read these values back and exercise removal.
- **Section 2** — only the ``top`` edge pre-populated with a thick blue
  border; other edges unset. Exercises mixed-state reads and per-edge
  mutation.

Run this script from the repo root::

    python features/steps/test_files/_gen_sct_borders.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import (
    WD_BORDER_DISPLAY,
    WD_BORDER_OFFSET_FROM,
    WD_SECTION,
)
from docx.enum.text import WD_BORDER_STYLE
from docx.shared import Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "sct-page-borders.docx")


def build() -> None:
    document = Document()

    # -- Section 0: no borders at all (default section) ---
    document.add_paragraph("Section 0 body — no page borders.")

    # -- Section 1: all four edges + display + offset_from ---
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 1 body — all four edges.")

    # -- Section 2: only a thick blue top border ---
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 2 body — top-only border.")

    # -- ``add_section`` clones the prior section's sectPr, so we populate
    #    borders only after every section exists — otherwise the borders set
    #    on section N would propagate into section N+1 at the clone step. ---
    section_1 = document.sections[1]
    for side in ("top", "bottom", "left", "right"):
        section_1.set_page_border(
            side,
            style=WD_BORDER_STYLE.SINGLE,
            width=Pt(1),
            color=RGBColor(0xFF, 0x00, 0x00),
            space=Pt(24),
        )
    section_1.page_borders.display = WD_BORDER_DISPLAY.ALL_PAGES
    section_1.page_borders.offset_from = WD_BORDER_OFFSET_FROM.PAGE

    section_2 = document.sections[2]
    section_2.remove_page_borders()  # -- start from a clean slate ---
    section_2.set_page_border(
        "top",
        style=WD_BORDER_STYLE.THICK,
        width=Pt(3),
        color=RGBColor(0x00, 0x00, 0xFF),
        space=Pt(12),
    )

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    sections = document.sections
    assert len(sections) == 3, f"expected 3 sections, got {len(sections)}"

    # -- Section 0 unset ---
    pb0 = sections[0].page_borders
    for side in ("top", "bottom", "left", "right"):
        border = getattr(pb0, side)
        assert border.style is None, f"section 0 {side}.style is {border.style!r}"
        assert border.width is None, f"section 0 {side}.width is {border.width!r}"
        assert border.color is None, f"section 0 {side}.color is {border.color!r}"
        assert border.space is None, f"section 0 {side}.space is {border.space!r}"
    assert pb0.display is None
    assert pb0.offset_from is None

    # -- Section 1 fully populated ---
    pb1 = sections[1].page_borders
    for side in ("top", "bottom", "left", "right"):
        border = getattr(pb1, side)
        assert border.style == WD_BORDER_STYLE.SINGLE, (
            f"section 1 {side}.style is {border.style!r}"
        )
        assert border.width == Pt(1), f"section 1 {side}.width is {border.width!r}"
        assert border.color == RGBColor(0xFF, 0x00, 0x00), (
            f"section 1 {side}.color is {border.color!r}"
        )
        assert border.space == Pt(24), f"section 1 {side}.space is {border.space!r}"
    assert pb1.display == WD_BORDER_DISPLAY.ALL_PAGES
    assert pb1.offset_from == WD_BORDER_OFFSET_FROM.PAGE

    # -- Section 2 top-only ---
    pb2 = sections[2].page_borders
    assert pb2.top.style == WD_BORDER_STYLE.THICK
    assert pb2.top.width == Pt(3)
    assert pb2.top.color == RGBColor(0x00, 0x00, 0xFF)
    assert pb2.top.space == Pt(12)
    for side in ("bottom", "left", "right"):
        border = getattr(pb2, side)
        assert border.style is None, f"section 2 {side}.style is {border.style!r}"


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
