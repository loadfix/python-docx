"""Generate ``doc-background.docx`` fixture for |Document.background_color|.

Builds a document with a non-default page background colour
(``w:document/w:background/@w:color``) so behave scenarios can round-trip the
value without constructing the element manually.

Run ``python features/steps/test_files/_gen_doc_background.py`` to regenerate
the fixture in-place and self-validate the result.
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import RGBColor

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "doc-background.docx")

# -- known fixture colour used by the scenarios --
FIXTURE_COLOR = RGBColor(0xFF, 0xA5, 0x00)  # -- orange --


def build() -> str:
    """Create ``doc-background.docx`` and return its absolute path."""
    document = Document()
    document.add_paragraph("Background-color fixture.")
    document.background_color = FIXTURE_COLOR
    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    """Reload `path` and assert background_color round-trips."""
    document = Document(path)
    actual = document.background_color
    assert actual == FIXTURE_COLOR, (
        f"background_color round-trip failed: {actual!r} != {FIXTURE_COLOR!r}"
    )


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
