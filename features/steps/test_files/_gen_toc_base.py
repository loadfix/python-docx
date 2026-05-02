"""Generate ``toc-base.docx`` — a fixture for the table-of-contents feature tests.

The generated document contains a small set of ``Heading 1``, ``Heading 2``,
and ``Heading 3`` paragraphs with body paragraphs between them. It is a
neutral canvas: no TOC has been added yet, so the behave scenarios can
exercise ``Document.add_table_of_contents()`` and
``Paragraph.insert_table_of_contents_{before,after}()`` from a known starting
state.

Run this script from the repo root::

    python features/steps/test_files/_gen_toc_base.py

After writing the file, the script reloads it and self-validates the
resulting structure so breakage is caught at fixture-build time rather than
during a behave run.
"""

from __future__ import annotations

import os

from docx import Document


OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "toc-base.docx")


# -- (style_name, text) pairs. The anchor paragraph at index 0 is a body
#    paragraph the scenarios use as a target for insert-at-position calls. --
STRUCTURE: list[tuple[str, str]] = [
    ("Normal", "Anchor body paragraph."),
    ("Heading 1", "Chapter One"),
    ("Normal", "Body of chapter one."),
    ("Heading 2", "Section 1.1"),
    ("Normal", "Body of section 1.1."),
    ("Heading 3", "Subsection 1.1.1"),
    ("Normal", "Body of subsection 1.1.1."),
    ("Heading 1", "Chapter Two"),
    ("Heading 2", "Section 2.1"),
    ("Heading 3", "Subsection 2.1.1"),
]


def _build() -> None:
    document = Document()
    for style_name, text in STRUCTURE:
        if style_name == "Normal":
            document.add_paragraph(text)
        else:
            # add_heading with level=N maps to "Heading N" style.
            level = int(style_name.split()[1])
            document.add_heading(text, level=level)
    document.save(OUT_PATH)


def _validate() -> None:
    """Reload the saved file and assert the expected shape."""
    document = Document(OUT_PATH)
    paragraphs = list(document.paragraphs)

    assert len(paragraphs) == len(STRUCTURE), (
        f"expected {len(STRUCTURE)} paragraphs, got {len(paragraphs)}"
    )

    for idx, ((style_name, text), paragraph) in enumerate(zip(STRUCTURE, paragraphs)):
        actual_style = paragraph.style.name if paragraph.style else None
        assert actual_style == style_name, (
            f"paragraph {idx}: expected style {style_name!r}, got {actual_style!r}"
        )
        assert paragraph.text == text, (
            f"paragraph {idx}: expected text {text!r}, got {paragraph.text!r}"
        )

    # -- quick counts so downstream scenarios can rely on them. --
    def _count(level: int) -> int:
        return sum(1 for s, _ in STRUCTURE if s == f"Heading {level}")

    assert _count(1) == 2, "expected 2 Heading 1 paragraphs"
    assert _count(2) == 2, "expected 2 Heading 2 paragraphs"
    assert _count(3) == 2, "expected 2 Heading 3 paragraphs"


if __name__ == "__main__":
    _build()
    _validate()
    print(f"Wrote {OUT_PATH}")
