"""Generate ``sct-doc-grid.docx`` — fixture for East-Asian document-grid tests.

The ``w:docGrid`` element carries ``w:type``, ``w:linePitch`` and
``w:charSpace``. Three sections cover:

- **Section 0** — default docGrid as written by ``python-docx``
  (``linePitch=360``, ``type`` unset). The generator explicitly asserts this
  state rather than assuming it so later regressions in the default template
  are caught here.
- **Section 1** — ``type=linesAndChars``, ``linePitch=312``, ``charSpace=0``.
- **Section 2** — no ``w:docGrid`` element at all (we remove it after
  sections are built).

Run this script from the repo root::

    python features/steps/test_files/_gen_sct_doc_grid.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_DOC_GRID_TYPE, WD_SECTION


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "sct-doc-grid.docx")


def build() -> None:
    document = Document()
    document.add_paragraph("Section 0 — default grid.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 1 — lines-and-chars, linePitch 312.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 2 — no docGrid element.")

    document.sections[1].set_document_grid(
        type=WD_DOC_GRID_TYPE.LINES_AND_CHARS,
        line_pitch=312,
        char_space=0,
    )
    document.sections[2].remove_document_grid()

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    sections = document.sections
    assert len(sections) == 3

    dg0 = sections[0].document_grid
    assert dg0 is not None, "section 0 should have a default docGrid"
    assert dg0.type is None, f"type={dg0.type!r}"
    assert dg0.line_pitch == 360, f"line_pitch={dg0.line_pitch!r}"
    assert dg0.char_space is None, f"char_space={dg0.char_space!r}"

    dg1 = sections[1].document_grid
    assert dg1 is not None
    assert dg1.type == WD_DOC_GRID_TYPE.LINES_AND_CHARS
    assert dg1.line_pitch == 312
    assert dg1.char_space == 0

    assert sections[2].document_grid is None


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
