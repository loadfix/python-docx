"""Generate ``sct-text-direction.docx`` — fixture for text-direction / RTL tests.

Three sections exercise :attr:`Section.text_direction` and
:attr:`Section.right_to_left`:

- **Section 0** — no ``w:textDirection`` and no ``w:bidi`` child. Reads
  return ``None`` / ``False``.
- **Section 1** — ``w:textDirection`` set to ``tbRl`` and ``w:bidi``
  present. Reads return ``WD_TEXT_DIRECTION.TB_RL`` / ``True``.
- **Section 2** — ``w:textDirection`` set to ``btLr`` only; RTL not enabled.

Run this script from the repo root::

    python features/steps/test_files/_gen_sct_text_direction.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TEXT_DIRECTION


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "sct-text-direction.docx")


def build() -> None:
    document = Document()
    document.add_paragraph("Section 0 — default LTR, no textDirection.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 1 — TB_RL vertical, RTL.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 2 — BT_LR vertical, LTR.")

    # -- reset every section before setting, so clone-on-add_section doesn't
    #    leak settings from section N into section N+1. ---
    for section in document.sections:
        section.text_direction = None
        section.right_to_left = False

    document.sections[1].text_direction = WD_TEXT_DIRECTION.TB_RL
    document.sections[1].right_to_left = True

    document.sections[2].text_direction = WD_TEXT_DIRECTION.BT_LR

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    sections = document.sections
    assert len(sections) == 3

    assert sections[0].text_direction is None
    assert sections[0].right_to_left is False

    assert sections[1].text_direction == WD_TEXT_DIRECTION.TB_RL
    assert sections[1].right_to_left is True

    assert sections[2].text_direction == WD_TEXT_DIRECTION.BT_LR
    assert sections[2].right_to_left is False


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
