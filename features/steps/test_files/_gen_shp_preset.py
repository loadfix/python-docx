"""Generate `shp-preset-shape.docx` behave fixture.

Builds a document exercising ``Paragraph.add_shape``: four inline DrawingML
preset shapes (rectangle, rounded rectangle, oval, right arrow) with a fifth
shape that carries text inside its text frame.

Self-validates after generation by re-opening the file and asserting the
shape types and text contents round-trip cleanly.

Run::

    python features/steps/test_files/_gen_shp_preset.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.drawing import WordprocessingShape
from docx.enum.shape import WD_DRAWING_TYPE, WD_SHAPE
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "shp-preset-shape.docx")


_SHAPES: list[tuple[WD_SHAPE, str | None]] = [
    (WD_SHAPE.RECTANGLE, None),
    (WD_SHAPE.ROUNDED_RECTANGLE, None),
    (WD_SHAPE.OVAL, None),
    (WD_SHAPE.ARROW_RIGHT, None),
    (WD_SHAPE.RECTANGLE, "Hello shape"),
]


def build() -> Document:
    document = Document()
    document.add_heading("Preset shapes fixture", level=1)
    document.add_paragraph(
        "Each of the following paragraphs contains a single inline "
        "DrawingML preset shape."
    )

    for shape_type, text in _SHAPES:
        p = document.add_paragraph()
        p.add_shape(
            shape_type,
            width=Inches(1.5),
            height=Inches(0.75),
            text=text,
        )

    document.add_paragraph("Trailing paragraph after the shapes.")
    return document


def self_validate(document: Document) -> None:
    # -- paragraphs[0] is the heading, paragraphs[1] is the intro --
    shape_paragraphs = document.paragraphs[2 : 2 + len(_SHAPES)]
    assert len(shape_paragraphs) == len(_SHAPES)

    for paragraph, (expected_type, expected_text) in zip(
        shape_paragraphs, _SHAPES
    ):
        drawings = paragraph.drawings
        assert len(drawings) == 1, (
            f"expected 1 drawing in paragraph, got {len(drawings)}"
        )
        drawing = drawings[0]
        if expected_text is None:
            assert drawing.type is WD_DRAWING_TYPE.SHAPE, (
                f"expected SHAPE, got {drawing.type}"
            )
        else:
            assert drawing.type is WD_DRAWING_TYPE.TEXT_BOX, (
                f"expected TEXT_BOX, got {drawing.type}"
            )

        # -- look up the wps:wsp element and wrap it --
        wsps = drawing._drawing.xpath(  # pyright: ignore[reportPrivateUsage]
            ".//wps:wsp"
        )
        assert wsps, "expected at least one wps:wsp child"
        wsp = WordprocessingShape(wsps[0], paragraph)
        assert wsp.shape_type is expected_type, (
            f"expected shape_type {expected_type}, got {wsp.shape_type}"
        )
        if expected_text is None:
            assert wsp.text == "", f"expected empty text, got {wsp.text!r}"
        else:
            assert wsp.text == expected_text, (
                f"expected {expected_text!r}, got {wsp.text!r}"
            )


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    reopened = Document(OUT_PATH)
    self_validate(reopened)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
