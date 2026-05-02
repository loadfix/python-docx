"""Generate ``fnt-table.docx`` fixture for font-table read scenarios.

The font-table part is managed by Word; python-docx does not create one on
demand. The default python-docx template ships a pre-populated ``fontTable.xml``
containing the eight fonts Word 2016 writes by default (Calibri, Cambria,
Times New Roman, Arial, Symbol, Courier, and the two MS 明朝/ゴシック East
Asian families), so this generator saves a bare :class:`docx.Document` and
relies on the template round-trip to produce the fixture.

Run ``python features/steps/test_files/_gen_fnt_table.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "fnt-table.docx")


def build() -> str:
    document = Document()
    document.add_paragraph("A document whose font table can be enumerated.")
    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    font_table = document.font_table
    assert font_table is not None, "expected a font_table, got None"
    assert len(font_table) >= 4, f"expected >=4 fonts, got {len(font_table)}"

    # -- Calibri is Word's default body font; always present --
    assert "Calibri" in font_table
    calibri = font_table["Calibri"]
    assert calibri.name == "Calibri"
    # -- PANOSE should be 20 hex characters --
    assert calibri.panose is not None and len(calibri.panose) == 20, calibri.panose

    # -- get() returns None rather than raising for unknown fonts --
    assert font_table.get("NoSuchFont") is None
    # -- membership test rejects non-strings --
    assert 42 not in font_table


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
