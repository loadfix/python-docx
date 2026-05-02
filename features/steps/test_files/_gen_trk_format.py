"""Build the ``trk-format.docx`` fixture used by the formatting-change tests.

Run::

    python features/steps/test_files/_gen_trk_format.py

*Generator strategy*: scaffolding (paragraph + run + section) is built with
python-docx's public API. The tracked-revision wrappers (``w:rPrChange``,
``w:pPrChange``, ``w:sectPrChange``) are injected by hand via ``OxmlElement``
because python-docx does not yet expose authoring methods for formatting
revisions. Accept/reject for these elements is covered by the existing pytest
suite and is surfaced on |Document| via ``accept_all_changes()``; the read-side
|FormattingChange| proxy exercised here is what the acceptance tests pin.

Self-checks:

* paragraph 1's run carries an ``rPrChange`` (author ``Alice``) with an old
  ``w:rPr`` exposing no ``w:b`` element
* paragraph 1 carries a ``pPrChange`` (author ``Bob``) exposing the prior
  (left-aligned) pPr
* the final section carries a ``sectPrChange`` (author ``Carol``)
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-format.docx")


def _make_change(tag: str, change_id: int, author: str, date: str, inner_tag: str):
    change = OxmlElement(tag)
    change.set(qn("w:id"), str(change_id))
    change.set(qn("w:author"), author)
    change.set(qn("w:date"), date)
    change.append(OxmlElement(inner_tag))
    return change


def build() -> Document:
    document = Document()

    # -- paragraph 0: title --
    document.add_heading("Tracked formatting changes", level=1)

    # -- paragraph 1: run is now bold + paragraph is now centered. Prior state
    # -- is the default (no bold, left-aligned). --
    p1 = document.add_paragraph()
    r1 = p1.add_run("Reformatted run.")
    r1.bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- inject an rPrChange recording the prior (non-bold) rPr --
    rPr = r1._r.rPr
    assert rPr is not None
    rPr.append(
        _make_change("w:rPrChange", 1, "Alice", "2025-04-10T09:00:00Z", "w:rPr")
    )

    # -- inject a pPrChange recording the prior (left-aligned) pPr --
    pPr = p1._p.pPr
    assert pPr is not None
    pPr.append(
        _make_change("w:pPrChange", 2, "Bob", "2025-04-10T09:05:00Z", "w:pPr")
    )

    # -- paragraph 2: plain (no formatting change) --
    document.add_paragraph("Unchanged paragraph.")

    # -- inject a sectPrChange on the document's (only) section --
    sectPr = document.sections[0]._sectPr
    sectPr.append(
        _make_change("w:sectPrChange", 3, "Carol", "2025-04-11T14:30:00Z", "w:sectPr")
    )

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- run-level rPrChange --
    run = paragraphs[1].runs[0]
    run_change = run.formatting_change
    assert run_change is not None, "expected rPrChange on paragraph 1 run 0"
    assert run_change.author == "Alice"
    assert run_change.date is not None
    assert run_change.old_properties is not None

    # -- paragraph-level pPrChange --
    p_change = paragraphs[1].formatting_change
    assert p_change is not None, "expected pPrChange on paragraph 1"
    assert p_change.author == "Bob"
    assert p_change.old_properties is not None

    # -- paragraph with no formatting change --
    assert paragraphs[2].formatting_change is None
    assert paragraphs[2].runs == [] or paragraphs[2].runs[0].formatting_change is None

    # -- section-level sectPrChange --
    s_change = document.sections[0].formatting_change
    assert s_change is not None, "expected sectPrChange on section 0"
    assert s_change.author == "Carol"
    assert s_change.old_properties is not None


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
