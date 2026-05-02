"""Generate the ``fnt-has-footnote-pr.docx`` acceptance-test fixture.

Creates a tiny document whose document-level footnote properties are fully
populated: number format, starting number, restart rule, and position. The
behave ``fnt-numbering-props.feature`` scenarios read back each of these
values.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import (
    WD_FOOTNOTE_POSITION,
    WD_FOOTNOTE_RESTART,
    WD_NUMBER_FORMAT,
)

THIS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(THIS_DIR, "fnt-has-footnote-pr.docx")


def build() -> None:
    document = Document()
    document.add_paragraph("Body paragraph for document with footnote properties.")

    props = document.add_footnote_properties()
    props.number_format = WD_NUMBER_FORMAT.LOWER_ROMAN
    props.start_number = 7
    props.restart_rule = WD_FOOTNOTE_RESTART.EACH_SECTION
    props.position = WD_FOOTNOTE_POSITION.BENEATH_TEXT

    # -- self-validate before saving --
    assert document.footnote_properties is not None
    fp = document.footnote_properties
    assert fp.number_format == WD_NUMBER_FORMAT.LOWER_ROMAN
    assert fp.start_number == 7
    assert fp.restart_rule == WD_FOOTNOTE_RESTART.EACH_SECTION
    assert fp.position == WD_FOOTNOTE_POSITION.BENEATH_TEXT

    document.save(OUT_PATH)

    # -- round-trip check --
    reopened = Document(OUT_PATH)
    rp = reopened.footnote_properties
    assert rp is not None, "footnote_properties missing after round-trip"
    assert rp.number_format == WD_NUMBER_FORMAT.LOWER_ROMAN
    assert rp.start_number == 7
    assert rp.restart_rule == WD_FOOTNOTE_RESTART.EACH_SECTION
    assert rp.position == WD_FOOTNOTE_POSITION.BENEATH_TEXT


if __name__ == "__main__":
    build()
    print(f"Wrote {OUT_PATH}")
