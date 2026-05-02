"""Build the ``trk-marks.docx`` fixture for the revision_marks_text() preview tests.

Run::

    python features/steps/test_files/_gen_trk_marks.py

*Generator strategy*: mirrors ``_gen_trk_ins_del.py`` but curates paragraphs whose
bracketed preview output is deterministic and covers every shape of the renderer:

* a paragraph with no tracked changes (renders identically to ``paragraph.text``)
* a paragraph with an insertion only
* a paragraph with a deletion only
* a paragraph with both an insertion and a deletion

No public create API exists for ``w:ins``/``w:del``; we inject via ``OxmlElement``.
"""

from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-marks.docx")


def _make_ins(run_id: int, author: str, date: str, text: str):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(run_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    ins.append(r)
    return ins


def _make_del(run_id: int, author: str, date: str, text: str):
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:id"), str(run_id))
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    dt_elm = OxmlElement("w:delText")
    dt_elm.text = text
    dt_elm.set(qn("xml:space"), "preserve")
    r.append(dt_elm)
    wdel.append(r)
    return wdel


def build() -> Document:
    document = Document()

    # -- paragraph 0: plain, no tracked changes --
    document.add_paragraph("Plain paragraph without tracked changes.")

    # -- paragraph 1: insertion only --
    p1 = document.add_paragraph("Please ")
    p1._p.append(_make_ins(1, "Alice", "2025-04-10T09:00:00Z", "kindly "))
    p1.add_run("consider.")

    # -- paragraph 2: deletion only --
    p2 = document.add_paragraph("Delete ")
    p2._p.append(_make_del(2, "Alice", "2025-04-10T09:05:00Z", "this part "))
    p2.add_run("of the text.")

    # -- paragraph 3: insertion and deletion interleaved --
    p3 = document.add_paragraph("The ")
    p3._p.append(_make_del(3, "Bob", "2025-04-10T09:10:00Z", "old"))
    p3._p.append(_make_ins(4, "Bob", "2025-04-10T09:10:05Z", "new"))
    p3.add_run(" value.")

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- plain paragraph: preview matches text --
    assert paragraphs[0].revision_marks_text() == paragraphs[0].text

    # -- insertion-only preview --
    assert paragraphs[1].revision_marks_text() == "Please [+kindly +]consider."

    # -- deletion-only preview --
    assert (
        paragraphs[2].revision_marks_text() == "Delete [-this part -]of the text."
    )

    # -- interleaved preview --
    assert paragraphs[3].revision_marks_text() == "The [-old-][+new+] value."

    # -- document-level preview joins paragraphs with blank lines --
    expected = "\n\n".join(
        p.revision_marks_text() for p in paragraphs
    )
    assert document.revision_marks_text() == expected

    # -- custom markers override the defaults --
    custom = paragraphs[1].revision_marks_text(
        open_ins="<INS>", close_ins="</INS>",
        open_del="<DEL>", close_del="</DEL>",
    )
    assert custom == "Please <INS>kindly </INS>consider."


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
