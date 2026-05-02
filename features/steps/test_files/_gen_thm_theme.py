"""Generate ``thm-theme.docx`` fixture for theme-proxy read scenarios.

The theme part is managed by Word, so python-docx has no public API for
authoring one. The default python-docx template already ships the stock Word
2007 "Office Theme" ``theme1.xml``, so this generator simply saves a bare
:class:`docx.Document` unchanged; the save round-trip ensures the theme part
survives and is reachable via :attr:`.Document.theme`.

Run ``python features/steps/test_files/_gen_thm_theme.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import RGBColor

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "thm-theme.docx")


def build() -> str:
    document = Document()
    document.add_paragraph("A document carrying Word's default Office theme.")
    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    theme = document.theme
    assert theme is not None, "document.theme is None; expected default Office theme"
    assert theme.name == "Office Theme", f"expected 'Office Theme', got {theme.name!r}"

    colors = theme.colors
    # -- dark_1 is black; light_1 is white in the default Office theme --
    assert colors.dark_1 == RGBColor(0x00, 0x00, 0x00), colors.dark_1
    assert colors.light_1 == RGBColor(0xFF, 0xFF, 0xFF), colors.light_1
    # -- every accent slot resolves to a RGBColor --
    for attr in (
        "accent_1",
        "accent_2",
        "accent_3",
        "accent_4",
        "accent_5",
        "accent_6",
    ):
        value = getattr(colors, attr)
        assert isinstance(value, RGBColor), f"{attr} was {value!r}"

    fonts = theme.fonts
    # -- the default python-docx template's theme1.xml pairs Calibri (major)
    # -- with Cambria (minor) --
    assert fonts.major_latin == "Calibri", fonts.major_latin
    assert fonts.minor_latin == "Cambria", fonts.minor_latin


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
