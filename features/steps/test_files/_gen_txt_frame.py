"""Generate the ``txt-frame.docx`` behave fixture.

A text frame is an absolutely-positioned text container (legacy predecessor
to text boxes). The fixture exposes two paragraph shapes:

- paragraph 0: a plain paragraph with no ``w:pPr/w:framePr`` child; all
  :attr:`ParagraphFormat.frame` access returns |None|.
- paragraph 1: a paragraph carrying a fully specified ``w:framePr`` — width
  3 in, height 1 in, horizontal position 0.5 in anchored to the page,
  vertical position 0.75 in anchored to the margin, wrap ``around``,
  horizontal alignment ``center``.
- paragraph 2: a paragraph carrying a drop-cap frame (``dropCap="drop"``)
  with three lines, anchored to the text.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import (
    WD_FRAME_DROP_CAP,
    WD_FRAME_H_ALIGN,
    WD_FRAME_H_ANCHOR,
    WD_FRAME_V_ANCHOR,
    WD_FRAME_WRAP,
)
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "txt-frame.docx")


def build() -> Document:
    document = Document()

    # -- paragraph 0: no frame --
    document.add_paragraph("Plain paragraph — no frame.")

    # -- paragraph 1: fully-specified frame --
    p1 = document.add_paragraph("Floating framed paragraph.")
    p1.paragraph_format.set_frame(
        width=Inches(3),
        height=Inches(1),
        horizontal_position=Inches(0.5),
        vertical_position=Inches(0.75),
        horizontal_anchor=WD_FRAME_H_ANCHOR.PAGE,
        vertical_anchor=WD_FRAME_V_ANCHOR.MARGIN,
        wrap=WD_FRAME_WRAP.AROUND,
        horizontal_alignment=WD_FRAME_H_ALIGN.CENTER,
    )

    # -- paragraph 2: drop-cap frame --
    p2 = document.add_paragraph("Dropped capital frame paragraph.")
    p2.paragraph_format.set_frame(
        drop_cap=WD_FRAME_DROP_CAP.DROP,
        lines=3,
        horizontal_anchor=WD_FRAME_H_ANCHOR.TEXT,
        vertical_anchor=WD_FRAME_V_ANCHOR.TEXT,
    )

    return document


def validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- paragraph 0 --
    assert paragraphs[0].paragraph_format.frame is None

    # -- paragraph 1 --
    frame1 = paragraphs[1].paragraph_format.frame
    assert frame1 is not None
    assert frame1.width == Inches(3), frame1.width
    assert frame1.height == Inches(1), frame1.height
    assert frame1.horizontal_position == Inches(0.5)
    assert frame1.vertical_position == Inches(0.75)
    assert frame1.horizontal_anchor == WD_FRAME_H_ANCHOR.PAGE
    assert frame1.vertical_anchor == WD_FRAME_V_ANCHOR.MARGIN
    assert frame1.wrap == WD_FRAME_WRAP.AROUND
    assert frame1.horizontal_alignment == WD_FRAME_H_ALIGN.CENTER

    # -- paragraph 2 --
    frame2 = paragraphs[2].paragraph_format.frame
    assert frame2 is not None
    assert frame2.drop_cap == WD_FRAME_DROP_CAP.DROP
    assert frame2.lines == 3
    assert frame2.horizontal_anchor == WD_FRAME_H_ANCHOR.TEXT
    assert frame2.vertical_anchor == WD_FRAME_V_ANCHOR.TEXT


def main() -> None:
    document = build()
    document.save(OUT_PATH)
    validate(Document(OUT_PATH))
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
