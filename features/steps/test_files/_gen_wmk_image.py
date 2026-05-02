"""Generate ``wmk-image.docx`` fixture for image-watermark behave scenarios.

Adds an image watermark using the existing ``monty-truth.png`` test image and
verifies the resulting document exposes an image-type |Watermark| on the first
section.

Run ``python features/steps/test_files/_gen_wmk_image.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import Inches

HERE = os.path.abspath(os.path.dirname(__file__))
IMAGE_PATH = os.path.join(HERE, "monty-truth.png")
OUT_PATH = os.path.join(HERE, "wmk-image.docx")


def build() -> str:
    document = Document()
    document.add_paragraph("Body paragraph one.")
    section = document.sections[0]
    section.add_image_watermark(IMAGE_PATH, width=Inches(3), height=Inches(2))
    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    watermark = document.sections[0].watermark
    assert watermark is not None, "image watermark missing after round-trip"
    assert watermark.type == "image", (
        f"expected image watermark, got {watermark.type!r}"
    )
    assert watermark.text is None


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
