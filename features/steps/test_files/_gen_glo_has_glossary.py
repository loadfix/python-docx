"""Generate ``glo-has-glossary.docx`` for behave fixtures.

The glossary-document part (``word/glossary/document.xml``) is Word-authored
and python-docx's proxy is read-only, so there is no public API for building
one. This script therefore starts from ``doc-default.docx`` and injects a
hand-rolled ``word/glossary/document.xml`` together with the matching
``[Content_Types].xml`` Override and ``word/_rels/document.xml.rels``
Relationship so the package round-trips through python-docx's loader.

Five building blocks are written so the filter and aggregate scenarios have
something meaningful to consume:

    Alpha    — gallery=quickParts, category=General   (body: one paragraph)
    Beta     — gallery=quickParts, category=General   (body: one paragraph)
    Gamma    — gallery=coverPg,    category=Built-In  (body: one paragraph,
                                                       one table)
    Delta    — gallery=hdrs,       category=Built-In  (body: one paragraph)
    Epsilon  — (no category element)                  (body: empty)

The script self-validates by loading the resulting ``.docx`` via python-docx
and asserting the proxy layer surfaces each block correctly before writing
the file.
"""

from __future__ import annotations

import os
import shutil
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_BUILDING_BLOCK_GALLERY

HERE = os.path.dirname(os.path.abspath(__file__))
SOURCE = os.path.join(HERE, "doc-default.docx")
TARGET = os.path.join(HERE, "glo-has-glossary.docx")


GLOSSARY_XML = b"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:glossaryDocument
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docParts>
    <w:docPart>
      <w:docPartPr>
        <w:name w:val="Alpha"/>
        <w:category>
          <w:name w:val="General"/>
          <w:gallery w:val="quickParts"/>
        </w:category>
        <w:description w:val="a quick-parts alpha block"/>
        <w:guid w:val="11111111-1111-1111-1111-111111111111"/>
      </w:docPartPr>
      <w:docPartBody>
        <w:p><w:r><w:t>Alpha body</w:t></w:r></w:p>
      </w:docPartBody>
    </w:docPart>
    <w:docPart>
      <w:docPartPr>
        <w:name w:val="Beta"/>
        <w:category>
          <w:name w:val="General"/>
          <w:gallery w:val="quickParts"/>
        </w:category>
        <w:description w:val="a quick-parts beta block"/>
        <w:guid w:val="22222222-2222-2222-2222-222222222222"/>
      </w:docPartPr>
      <w:docPartBody>
        <w:p><w:r><w:t>Beta body</w:t></w:r></w:p>
      </w:docPartBody>
    </w:docPart>
    <w:docPart>
      <w:docPartPr>
        <w:name w:val="Gamma"/>
        <w:category>
          <w:name w:val="Built-In"/>
          <w:gallery w:val="coverPg"/>
        </w:category>
        <w:description w:val="a cover-page block"/>
        <w:guid w:val="33333333-3333-3333-3333-333333333333"/>
      </w:docPartPr>
      <w:docPartBody>
        <w:p><w:r><w:t>Gamma body</w:t></w:r></w:p>
        <w:tbl>
          <w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>
          <w:tblGrid><w:gridCol w:w="2000"/></w:tblGrid>
          <w:tr>
            <w:tc>
              <w:tcPr><w:tcW w:w="2000" w:type="dxa"/></w:tcPr>
              <w:p><w:r><w:t>cell</w:t></w:r></w:p>
            </w:tc>
          </w:tr>
        </w:tbl>
      </w:docPartBody>
    </w:docPart>
    <w:docPart>
      <w:docPartPr>
        <w:name w:val="Delta"/>
        <w:category>
          <w:name w:val="Built-In"/>
          <w:gallery w:val="hdrs"/>
        </w:category>
        <w:description w:val="a header block"/>
        <w:guid w:val="44444444-4444-4444-4444-444444444444"/>
      </w:docPartPr>
      <w:docPartBody>
        <w:p><w:r><w:t>Delta body</w:t></w:r></w:p>
      </w:docPartBody>
    </w:docPart>
    <w:docPart>
      <w:docPartPr>
        <w:name w:val="Epsilon"/>
      </w:docPartPr>
    </w:docPart>
  </w:docParts>
</w:glossaryDocument>
"""


CONTENT_TYPE_GLOSSARY = (
    "application/vnd.openxmlformats-officedocument"
    ".wordprocessingml.document.glossary+xml"
)
RT_GLOSSARY = (
    "http://schemas.openxmlformats.org/officeDocument/2006/"
    "relationships/glossaryDocument"
)


def _inject_glossary_override(content_types_xml: bytes) -> bytes:
    """Return ``[Content_Types].xml`` with a glossary ``Override`` appended."""
    override = (
        f'<Override PartName="/word/glossary/document.xml"'
        f' ContentType="{CONTENT_TYPE_GLOSSARY}"/>'
    ).encode("utf-8")
    # -- insert before the closing </Types> tag
    close = b"</Types>"
    idx = content_types_xml.rfind(close)
    if idx == -1:
        raise ValueError("[Content_Types].xml missing </Types>")
    return content_types_xml[:idx] + override + content_types_xml[idx:]


def _inject_glossary_relationship(rels_xml: bytes) -> bytes:
    """Return ``document.xml.rels`` with a ``glossaryDocument`` relationship."""
    rel = (
        f'<Relationship Id="rIdGlossary" Type="{RT_GLOSSARY}"'
        f' Target="glossary/document.xml"/>'
    ).encode("utf-8")
    close = b"</Relationships>"
    idx = rels_xml.rfind(close)
    if idx == -1:
        raise ValueError("document.xml.rels missing </Relationships>")
    return rels_xml[:idx] + rel + rels_xml[idx:]


def build() -> None:
    """Write ``glo-has-glossary.docx`` alongside this script."""
    if not os.path.isfile(SOURCE):
        raise FileNotFoundError(f"source fixture missing: {SOURCE}")

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)

    try:
        with zipfile.ZipFile(SOURCE, "r") as src, zipfile.ZipFile(
            tmp_path, "w", zipfile.ZIP_DEFLATED
        ) as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = _inject_glossary_override(data)
                elif item.filename == "word/_rels/document.xml.rels":
                    data = _inject_glossary_relationship(data)
                dst.writestr(item, data)

            # -- new glossary part
            dst.writestr("word/glossary/document.xml", GLOSSARY_XML)

        # -- self-validate via the public API before adopting the file
        _validate(tmp_path)

        shutil.move(tmp_path, TARGET)
        print(f"wrote {TARGET}")
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def _validate(path: str) -> None:
    """Load `path` through python-docx and assert the fixture's shape."""
    document = Document(path)
    glossary = document.glossary
    assert glossary is not None, "document.glossary returned None"
    assert len(glossary) == 5, f"expected 5 building blocks, got {len(glossary)}"

    names = [b.name for b in glossary]
    assert names == ["Alpha", "Beta", "Gamma", "Delta", "Epsilon"], names

    # -- Alpha metadata
    alpha = glossary["Alpha"]
    assert alpha.description == "a quick-parts alpha block", alpha.description
    assert alpha.guid == "11111111-1111-1111-1111-111111111111", alpha.guid
    assert alpha.category.gallery == "quickParts", alpha.category.gallery
    assert alpha.category.category_name == "General", alpha.category.category_name
    assert (
        alpha.category.gallery_value is WD_BUILDING_BLOCK_GALLERY.QUICK_PARTS
    ), alpha.category.gallery_value
    assert len(alpha.paragraphs) == 1
    assert alpha.paragraphs[0].text == "Alpha body"

    # -- Gamma has a table
    gamma = glossary["Gamma"]
    assert len(gamma.tables) == 1, len(gamma.tables)
    assert gamma.category.gallery == "coverPg", gamma.category.gallery

    # -- Delta — header gallery
    delta = glossary["Delta"]
    assert (
        delta.category.gallery_value is WD_BUILDING_BLOCK_GALLERY.HEADERS
    ), delta.category.gallery_value

    # -- Epsilon — no category, empty body
    epsilon = glossary["Epsilon"]
    assert epsilon.category.gallery is None
    assert epsilon.category.category_name is None
    assert epsilon.paragraphs == []
    assert epsilon.tables == []

    # -- filtering and aggregation
    quick = glossary.by_category(gallery=WD_BUILDING_BLOCK_GALLERY.QUICK_PARTS)
    assert [b.name for b in quick] == ["Alpha", "Beta"], [b.name for b in quick]
    built_in = glossary.by_category(category_name="Built-In")
    assert [b.name for b in built_in] == ["Gamma", "Delta"], [
        b.name for b in built_in
    ]
    assert glossary.galleries == ["quickParts", "coverPg", "hdrs"], glossary.galleries
    assert len(glossary.categories) == 3


if __name__ == "__main__":
    build()
