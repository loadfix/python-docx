"""Generate ``equ-subscript.docx`` fixture for OMML subscript scenarios.

Two paragraphs built with :func:`docx.equations.build_subscript`:

- Paragraph 0: identifier base with a numeric subscript — ``a`` indexed by
  ``1``.
- Paragraph 1: two subscript equations *chained* in the same paragraph — an
  indexed element ``a_i`` followed by a second indexed element ``b_j``. The
  scenarios read back both ``m:sSub`` expressions from this paragraph.

The script self-validates by reopening the saved document and asserting the
equation count, flattened text, and that each expression carries an
``m:sSub`` element.

Run ``python features/steps/test_files/_gen_equ_subscript.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document
from docx.equations import build_subscript

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "equ-subscript.docx")


def build() -> str:
    document = Document()
    document.add_heading("Subscript fixture", level=1)

    # -- Paragraph 0: a_1 — identifier base + numeric subscript --
    p0 = document.add_paragraph("First term: ")
    p0.add_equation(build_subscript("a", "1"))

    # -- Paragraph 1: a_i b_j — two subscripts chained in a single paragraph --
    p1 = document.add_paragraph("Tensor product: ")
    p1.add_equation(build_subscript("a", "i"))
    p1.add_run(" ")
    p1.add_equation(build_subscript("b", "j"))

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    paragraphs = document.paragraphs

    # -- Paragraph 0: single subscript equation --
    first_terms = paragraphs[1].equations  # paragraphs[0] is the heading
    assert len(first_terms) == 1, (
        f"expected 1 equation in first paragraph, got {len(first_terms)}"
    )
    assert first_terms[0].text == "a1", (
        f"expected 'a1', got {first_terms[0].text!r}"
    )
    assert b"<m:sSub>" in first_terms[0].raw_xml

    # -- Paragraph 1: two chained subscript equations --
    chained = paragraphs[2].equations
    assert len(chained) == 2, (
        f"expected 2 chained equations in second paragraph, got {len(chained)}"
    )
    assert chained[0].text == "ai", f"expected 'ai', got {chained[0].text!r}"
    assert chained[1].text == "bj", f"expected 'bj', got {chained[1].text!r}"
    for eq in chained:
        assert b"<m:sSub>" in eq.raw_xml
        assert eq.is_display_mode is False


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
