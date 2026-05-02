"""Generate ``doc-odd-even-hdrs.docx`` — fixture for odd/even/first-page headers.

Wires up what was previously an orphan fixture to cover
:attr:`Section.different_first_page_header_footer`,
:attr:`Section.different_odd_and_even_pages_header_footer` and the
``first_page_header/footer`` / ``even_page_header/footer`` accessors.

The fixture contains:

- ``different_odd_and_even_pages_header_footer`` **enabled** at the document
  level (that flag is a document-wide setting).
- A single section with:

  * ``different_first_page_header_footer`` set to ``True``,
  * a primary (odd) header reading ``"ODD HEADER"``,
  * a primary (odd) footer reading ``"ODD FOOTER"``,
  * an even-page header reading ``"EVEN HEADER"``,
  * an even-page footer reading ``"EVEN FOOTER"``,
  * a first-page header reading ``"FIRST HEADER"``,
  * a first-page footer reading ``"FIRST FOOTER"``.

Run this script from the repo root::

    python features/steps/test_files/_gen_sct_odd_even_hdrs.py
"""

from __future__ import annotations

import os

from docx import Document


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "doc-odd-even-hdrs.docx")


def build() -> None:
    document = Document()
    document.add_paragraph("Body paragraph for the section.")

    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.different_odd_and_even_pages_header_footer = True

    # -- primary (odd) header / footer ---
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "ODD HEADER"
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "ODD FOOTER"

    # -- even-page header / footer ---
    section.even_page_header.is_linked_to_previous = False
    section.even_page_header.paragraphs[0].text = "EVEN HEADER"
    section.even_page_footer.is_linked_to_previous = False
    section.even_page_footer.paragraphs[0].text = "EVEN FOOTER"

    # -- first-page header / footer ---
    section.first_page_header.is_linked_to_previous = False
    section.first_page_header.paragraphs[0].text = "FIRST HEADER"
    section.first_page_footer.is_linked_to_previous = False
    section.first_page_footer.paragraphs[0].text = "FIRST FOOTER"

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    sections = document.sections
    assert len(sections) == 1, f"expected 1 section, got {len(sections)}"
    section = sections[0]

    assert section.different_first_page_header_footer is True
    assert section.different_odd_and_even_pages_header_footer is True

    assert section.header.paragraphs[0].text == "ODD HEADER"
    assert section.footer.paragraphs[0].text == "ODD FOOTER"
    assert section.even_page_header.paragraphs[0].text == "EVEN HEADER"
    assert section.even_page_footer.paragraphs[0].text == "EVEN FOOTER"
    assert section.first_page_header.paragraphs[0].text == "FIRST HEADER"
    assert section.first_page_footer.paragraphs[0].text == "FIRST FOOTER"

    # -- each hdr/ftr should have its own definition ---
    assert section.header.is_linked_to_previous is False
    assert section.footer.is_linked_to_previous is False
    assert section.even_page_header.is_linked_to_previous is False
    assert section.even_page_footer.is_linked_to_previous is False
    assert section.first_page_header.is_linked_to_previous is False
    assert section.first_page_footer.is_linked_to_previous is False


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
