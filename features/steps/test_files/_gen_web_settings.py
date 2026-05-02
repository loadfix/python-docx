"""Generate ``web-settings.docx`` fixture for web-settings scenarios.

Toggles the writable boolean flags on :attr:`.Document.web_settings` so
scenarios can assert they round-trip through ``word/webSettings.xml``.

Run ``python features/steps/test_files/_gen_web_settings.py`` to regenerate.
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "web-settings.docx")


def build() -> str:
    document = Document()
    document.add_paragraph("A document with non-default web settings.")

    web = document.web_settings
    assert web is not None
    web.optimize_for_browser = True
    web.allow_png = True
    web.do_not_save_as_single_file = True

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    web = document.web_settings
    assert web is not None, "web_settings missing after round-trip"
    assert web.optimize_for_browser is True
    assert web.allow_png is True
    assert web.do_not_save_as_single_file is True


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
