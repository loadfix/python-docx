"""Build the ``fld-has-fields.docx`` fixture used by the fields acceptance tests.

Run from the repo root (or anywhere; paths are resolved relative to this file)::

    python features/steps/test_files/_gen_fld_has_fields.py

The generator uses python-docx itself so that the fixture tracks the library's
own serialization quirks. A round-trip self-check is performed before the file
is written:

* paragraph 2 exposes exactly one simple field (``DATE``) via ``Paragraph.fields``
* paragraph 3 exposes exactly one complex field (``PAGE``) via ``Paragraph.fields``
* paragraph 5 exposes a ``REF`` field that resolves to the bookmark's text

Regenerate any time the fields or bookmarks APIs change shape.
"""

from __future__ import annotations

import os

from docx import Document
from docx.fields import WD_FIELD_TYPE

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "fld-has-fields.docx")


def build() -> Document:
    document = Document()

    # -- paragraph 0: intro heading (so tests can rely on fixed indices) --
    document.add_heading("Fields fixture", level=1)

    # -- paragraph 1: plain text before any field --
    document.add_paragraph("Plain text before fields.")

    # -- paragraph 2: a simple DATE field with cached rendered result --
    simple_para = document.add_paragraph("Today is ")
    simple_para.add_simple_field(WD_FIELD_TYPE.DATE, "2025-01-02")
    simple_para.add_run(".")

    # -- paragraph 3: a complex PAGE field (begin/separate/end) --
    complex_para = document.add_paragraph("Page number: ")
    complex_para.add_complex_field(WD_FIELD_TYPE.PAGE, "7")
    complex_para.add_run(".")

    # -- paragraph 4: bookmark target that REF will point back to --
    bookmark_para = document.add_paragraph("The quoted value is forty-two.")
    bookmark_run = bookmark_para.runs[0]
    bookmark_para.add_bookmark(
        "FavouriteValue", start_run=bookmark_run, end_run=bookmark_run
    )

    # -- paragraph 5: a REF field pointing at the bookmark above --
    ref_para = document.add_paragraph("As noted earlier: ")
    ref_para.add_complex_field(
        "REF FavouriteValue \\h", "The quoted value is forty-two."
    )
    ref_para.add_run(" (see above).")

    return document


def self_validate(document: Document) -> None:
    """Assert the constructed document exposes the field shapes tests expect."""
    paragraphs = document.paragraphs

    # -- simple field on paragraph 2 --
    simple_fields = paragraphs[2].fields
    assert len(simple_fields) == 1, (
        f"expected 1 simple field on paragraph 2, got {len(simple_fields)}"
    )
    assert simple_fields[0].is_complex is False
    assert simple_fields[0].type == "DATE"
    assert simple_fields[0].result_text == "2025-01-02"

    # -- complex field on paragraph 3 --
    complex_fields = paragraphs[3].fields
    assert len(complex_fields) == 1, (
        f"expected 1 complex field on paragraph 3, got {len(complex_fields)}"
    )
    assert complex_fields[0].is_complex is True
    assert complex_fields[0].type == "PAGE"
    assert complex_fields[0].result_text == "7"

    # -- cross-reference field on paragraph 5 resolves against the bookmark --
    ref_fields = paragraphs[5].fields
    assert len(ref_fields) == 1, f"expected 1 REF field, got {len(ref_fields)}"
    ref_field = ref_fields[0]
    assert ref_field.type == "REF"
    resolved = ref_field.resolve(document)
    assert resolved == "The quoted value is forty-two.", (
        f"REF did not resolve to bookmark text, got {resolved!r}"
    )


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
