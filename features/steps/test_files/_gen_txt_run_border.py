"""Generate the ``txt-run-border.docx`` behave fixture.

Builds a document containing one paragraph per border-state so behave scenarios
can select the run by paragraph index:

- paragraph 0: run with **no** ``w:bdr`` child (all border_* properties return ``None``).
- paragraph 1: run with a fully-populated ``w:bdr`` (style SINGLE, color ``FF0000``,
  width ``Pt(1.5)`` — i.e. 12 eighth-points, space ``Pt(4)``).
- paragraph 2: run with a ``w:bdr`` whose ``w:val`` is ``auto`` for color, so
  :attr:`Font.border_color` reads back as ``None``.

The generator round-trips the file through ``Document(...)`` and asserts the
expected public-API values before writing the fixture.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_BORDER_STYLE
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "txt-run-border.docx")


def build() -> Document:
    document = Document()

    # -- paragraph 0: no border --
    document.add_paragraph("Run without a border.")

    # -- paragraph 1: fully-specified border --
    p1 = document.add_paragraph("Run with a red single border.")
    font = p1.runs[0].font
    font.border_style = WD_BORDER_STYLE.SINGLE
    font.border_color = RGBColor(0xFF, 0x00, 0x00)
    font.border_width = Pt(1.5)
    font.border_space = Pt(4)

    # -- paragraph 2: auto color (tested as None via the public API) --
    p2 = document.add_paragraph("Run with a border using auto color.")
    font2 = p2.runs[0].font
    font2.border_style = WD_BORDER_STYLE.DASHED
    font2.border_width = Pt(1)
    # -- touch the oxml layer to set the "auto" sentinel that the getter
    #    normalises back to None --
    rPr = font2._element.get_or_add_rPr()  # pyright: ignore[reportPrivateUsage]
    bdr = rPr.get_or_add_bdr()
    bdr.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color",
        "auto",
    )

    return document


def validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- paragraph 0: nothing set --
    font0 = paragraphs[0].runs[0].font
    assert font0.border_style is None
    assert font0.border_color is None
    assert font0.border_width is None
    assert font0.border_space is None

    # -- paragraph 1: all four attributes round-trip --
    font1 = paragraphs[1].runs[0].font
    assert font1.border_style == WD_BORDER_STYLE.SINGLE, font1.border_style
    assert font1.border_color == RGBColor(0xFF, 0x00, 0x00), font1.border_color
    assert font1.border_width == Pt(1.5), font1.border_width
    assert font1.border_space == Pt(4), font1.border_space

    # -- paragraph 2: color="auto" → None via the proxy --
    font2 = paragraphs[2].runs[0].font
    assert font2.border_style == WD_BORDER_STYLE.DASHED
    assert font2.border_color is None, font2.border_color


def main() -> None:
    document = build()
    document.save(OUT_PATH)
    # -- re-open to make sure the serialisation round-trips cleanly --
    validate(Document(OUT_PATH))
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
