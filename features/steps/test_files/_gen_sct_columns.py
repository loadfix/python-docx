"""Generate ``sct-multi-column.docx`` — fixture for multi-column layout tests.

The ``w:cols`` element has ``num`` / ``space`` / ``equalWidth`` attributes
plus zero-or-more ``w:col`` children. Four sections exercise the matrix:

- **Section 0** — no ``w:cols`` element (single-column default). Reads
  report ``count == 1``, ``equal_width is True``, ``space is None``.
- **Section 1** — three equal-width columns with a 0.25" gutter between
  them. ``equal_width=True``, ``space=Pt(18)``, no ``w:col`` children.
- **Section 2** — two unequal columns declared as explicit ``w:col``
  children: 2.5" wide left column followed by a 4" right column with 0.5"
  space after each.
- **Section 3** — starts with two equal columns. Scenarios mutate this
  section to verify setters and ``count`` changes.

Run this script from the repo root::

    python features/steps/test_files/_gen_sct_columns.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Emu, Inches, Pt


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "sct-multi-column.docx")


def _add_col(cols_elm, width_emu: int, space_emu: int | None) -> None:
    """Append a ``w:col`` child with the given width and space (in EMU)."""
    col = cols_elm.add_col()
    col.w = Emu(width_emu)
    if space_emu is not None:
        col.space = Emu(space_emu)


def build() -> None:
    document = Document()

    document.add_paragraph("Section 0 — single column.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 1 — three equal columns.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 2 — two unequal columns.")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Section 3 — two equal columns (mutation target).")

    # -- clear before populating, so clone-on-add_section doesn't inherit. ---
    for section in document.sections:
        # remove any w:cols written by ``add_section`` clone
        sectPr = section._sectPr  # pyright: ignore[reportPrivateUsage]
        sectPr._remove_cols()  # pyright: ignore[reportPrivateUsage]

    # -- Section 1: three equal columns, 18pt space --
    s1 = document.sections[1]
    s1.columns.count = 3
    s1.columns.equal_width = True
    s1.columns.space = Pt(18)

    # -- Section 2: two unequal columns declared individually --
    s2 = document.sections[2]
    s2.columns.count = 2
    s2.columns.equal_width = False
    cols2 = s2._sectPr.get_or_add_cols()  # pyright: ignore[reportPrivateUsage]
    _add_col(cols2, int(Inches(2.5)), int(Inches(0.5)))
    _add_col(cols2, int(Inches(4.0)), int(Inches(0.5)))

    # -- Section 3: two equal columns baseline --
    s3 = document.sections[3]
    s3.columns.count = 2
    s3.columns.equal_width = True
    s3.columns.space = Pt(12)

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    sections = document.sections
    assert len(sections) == 4

    # -- Section 0 --
    cols0 = sections[0].columns
    assert cols0.count == 1, f"section 0 count={cols0.count!r}"
    assert cols0.equal_width is True, f"section 0 equal_width={cols0.equal_width!r}"
    assert cols0.space is None, f"section 0 space={cols0.space!r}"
    assert len(cols0) == 0, f"section 0 has {len(cols0)} Column(s)"

    # -- Section 1 --
    cols1 = sections[1].columns
    assert cols1.count == 3, f"section 1 count={cols1.count!r}"
    assert cols1.equal_width is True
    assert cols1.space == Pt(18)
    assert len(cols1) == 0, "section 1 should have no explicit w:col children"

    # -- Section 2 --
    cols2 = sections[2].columns
    assert cols2.count == 2, f"section 2 count={cols2.count!r}"
    assert cols2.equal_width is False
    assert len(cols2) == 2, f"section 2 has {len(cols2)} Column(s)"
    assert cols2[0].width == Inches(2.5), f"col0.width={cols2[0].width!r}"
    assert cols2[0].space == Inches(0.5)
    assert cols2[1].width == Inches(4.0)
    assert cols2[1].space == Inches(0.5)

    # -- Section 3 --
    cols3 = sections[3].columns
    assert cols3.count == 2
    assert cols3.equal_width is True
    assert cols3.space == Pt(12)


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
