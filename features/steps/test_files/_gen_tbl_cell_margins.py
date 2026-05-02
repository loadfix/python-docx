"""Generate the `tbl-cell-margins.docx` behave fixture.

Running this script (idempotently) produces a .docx containing a 2x2
table where the top-left cell carries explicit per-cell margin
overrides on all four edges, and the other three cells are plain
(no ``w:tcMar``). The fixture is used by the `tbl-cell-margins.feature`
scenarios to exercise the ``_Cell.margins`` proxy.

The script self-validates the written document before considering the
fixture good.
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import Twips


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "tbl-cell-margins.docx")


# -- explicit edge values on the first cell (exact twip values so the
# --  EMU/Twip round-trip through XML is lossless). --
TOP = Twips(72)     # 0.05 in
BOTTOM = Twips(72)  # 0.05 in
START = Twips(115)  # ~0.08 in
END = Twips(115)    # ~0.08 in


def build() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"

    margins = table.cell(0, 0).margins
    margins.top = TOP
    margins.bottom = BOTTOM
    margins.start = START
    margins.end = END

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    table = document.tables[0]

    m = table.cell(0, 0).margins
    assert m.top == TOP, f"top: got {m.top!r}"
    assert m.bottom == BOTTOM, f"bottom: got {m.bottom!r}"
    assert m.start == START, f"start: got {m.start!r}"
    assert m.end == END, f"end: got {m.end!r}"

    # -- other cells should have no margin overrides --
    for row, col in ((0, 1), (1, 0), (1, 1)):
        m2 = table.cell(row, col).margins
        assert m2.top is None, f"cell({row},{col}).top: got {m2.top!r}"
        assert m2.bottom is None
        assert m2.start is None
        assert m2.end is None

    # -- the four edges round-trip exactly when expressed in twips --
    assert m.top == TOP and m.bottom == BOTTOM
    assert m.start == START and m.end == END


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
