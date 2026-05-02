"""Generate ``doc-customprops.docx`` fixture for custom-properties behave scenarios.

Populates ``Document.custom_properties`` with one value of each supported
type so the feature scenarios have something to read back:

    Project    -- str
    Priority   -- int
    Budget     -- float
    Approved   -- bool
    Reviewed   -- datetime

Self-validates by reloading the saved document and asserting every property
round-trips through the XML layer.

Run ``python features/steps/test_files/_gen_doc_customprops.py`` to regenerate.
"""

from __future__ import annotations

import datetime as dt
import os

from docx import Document

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "doc-customprops.docx")


def build() -> str:
    document = Document()
    document.add_paragraph("A document with custom properties.")

    props = document.custom_properties
    props["Project"] = "Apollo"
    props["Priority"] = 5
    props["Budget"] = 99.95
    props["Approved"] = True
    props["Reviewed"] = dt.datetime(2024, 3, 1, 12, 0, tzinfo=dt.timezone.utc)

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    document = Document(path)
    props = document.custom_properties
    assert len(props) == 5, f"expected 5 custom properties, got {len(props)}"
    assert props["Project"] == "Apollo"
    assert props["Priority"] == 5
    assert props["Budget"] == 99.95
    assert props["Approved"] is True
    reviewed = props["Reviewed"]
    assert isinstance(reviewed, dt.datetime)
    assert reviewed.year == 2024 and reviewed.month == 3 and reviewed.day == 1
    assert props.names() == ["Project", "Priority", "Budget", "Approved", "Reviewed"]
    assert "Project" in props
    assert "NoSuchProperty" not in props
    assert props.get("Missing") is None


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
