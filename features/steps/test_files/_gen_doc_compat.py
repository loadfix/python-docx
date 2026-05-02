"""Generate ``doc-compat.docx`` fixture for compatibility-settings scenarios.

Builds a document whose ``w:settings/w:compat`` block contains a handful of
well-known flag elements (``growAutofit``, ``doNotBreakWrappedTables``, ...) and
also exercises ``w:compatSetting`` entries via :attr:`Settings.compat_settings`.
The fixture lets the behave scenarios read back flags without having to mutate
a runtime document.

Run ``python features/steps/test_files/_gen_doc_compat.py`` to regenerate the
fixture in-place and self-validate the result.
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "doc-compat.docx")


KNOWN_ON_FLAGS = (
    "growAutofit",
    "doNotBreakWrappedTables",
    "useFELayout",
)


def build() -> str:
    """Create ``doc-compat.docx`` and return its absolute path."""
    document = Document()
    document.add_paragraph("Compat-flags fixture.")

    settings = document.settings

    # -- populate the well-known compat flags that read as True --
    for name in KNOWN_ON_FLAGS:
        settings.compat_flags[name] = True

    # -- upgrade compatibilityMode to 15 (Word 2013+) so fixture surfaces --
    # -- the most common production value --
    settings.compat_settings["compatibilityMode"] = "15"

    # -- add a couple of non-default compatSetting entries --
    settings.compat_settings["differentiateMultirowTableHeaders"] = "1"
    settings.compat_settings["useWord2013TrackBottomHyphenation"] = "1"

    document.save(OUT_PATH)
    return OUT_PATH


def validate(path: str) -> None:
    """Reload `path` and assert every compat flag / setting round-trips."""
    document = Document(path)
    settings = document.settings

    for name in KNOWN_ON_FLAGS:
        assert settings.compat_flags[name] is True, (
            f"compat_flags[{name!r}] expected True, got {settings.compat_flags[name]!r}"
        )

    # -- flags that were never set read as False (not KeyError) --
    assert settings.compat_flags["noTabHangInd"] is False
    assert "noTabHangInd" not in settings.compat_flags

    # -- compatSetting round-trip --
    assert settings.compat_settings["compatibilityMode"] == "15"
    assert settings.compat_settings["differentiateMultirowTableHeaders"] == "1"
    assert settings.compat_settings["useWord2013TrackBottomHyphenation"] == "1"
    assert "differentiateMultirowTableHeaders" in settings.compat_settings


if __name__ == "__main__":
    out = build()
    validate(out)
    print(f"wrote {out}")
