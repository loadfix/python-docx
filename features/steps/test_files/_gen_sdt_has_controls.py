"""Generate `sdt-has-controls.docx` fixture for behave acceptance tests.

Creates a small Word document containing one of each supported content-control
type (rich-text, plain-text, date, checkbox, combo-box, dropdown, picture) plus
a companion inline control, then self-validates the result.

Run directly::

    python features/steps/test_files/_gen_sdt_has_controls.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.content_controls import ContentControlType


_HERE = os.path.dirname(os.path.abspath(__file__))
_OUT = os.path.join(_HERE, "sdt-has-controls.docx")


# -- each tuple is (type, tag, title, initial text) --
_BLOCK_CONTROLS: list[tuple[ContentControlType, str, str, str]] = [
    (ContentControlType.RICH_TEXT, "rich-1", "Rich Text", "Rich text body"),
    (ContentControlType.PLAIN_TEXT, "plain-1", "Plain Text", "Plain text body"),
    (ContentControlType.DATE, "date-1", "Date", "2026-05-02"),
    (ContentControlType.CHECKBOX, "cbx-1", "Checkbox", ""),
    (ContentControlType.COMBO_BOX, "cmb-1", "Combo Box", "Choose..."),
    (ContentControlType.DROPDOWN, "dd-1", "Dropdown", "Pick one"),
    (ContentControlType.PICTURE, "pic-1", "Picture", ""),
]


def _build_document() -> Document:
    document = Document()
    document.add_heading("SDT sample", level=1)

    for cc_type, tag, title, text in _BLOCK_CONTROLS:
        cc = document.add_content_control(cc_type, tag=tag, title=title)
        if text:
            cc.text = text
        if cc_type is ContentControlType.CHECKBOX:
            cc.checked = True

    # -- one inline rich-text control inside a paragraph --
    p = document.add_paragraph("Inline control: ")
    inline = p.add_content_control(
        ContentControlType.RICH_TEXT, tag="inline-rt", title="Inline Rich"
    )
    inline.text = "inline value"

    return document


def _self_validate(document: Document) -> None:
    """Run read-side assertions on the document before it is written to disk."""
    block_controls = document.content_controls
    assert len(block_controls) == 7, (
        f"expected 7 block-level controls, got {len(block_controls)}"
    )

    # -- order + type check --
    expected_types = [t for (t, _, _, _) in _BLOCK_CONTROLS]
    actual_types = [cc.type for cc in block_controls]
    assert actual_types == expected_types, (
        f"types out of order: {actual_types} != {expected_types}"
    )

    # -- tag / title round-trip --
    for cc, (_, tag, title, _) in zip(block_controls, _BLOCK_CONTROLS):
        assert cc.tag == tag, f"tag mismatch: {cc.tag!r} != {tag!r}"
        assert cc.title == title, f"title mismatch: {cc.title!r} != {title!r}"
        assert cc.sdt_id is not None, f"sdt_id missing for {tag!r}"

    # -- checkbox round-trip --
    cbx = next(cc for cc in block_controls if cc.type is ContentControlType.CHECKBOX)
    assert cbx.checked is True, f"checkbox state not round-tripped: {cbx.checked!r}"

    # -- inline control accessible via Paragraph.content_controls --
    last_para = document.paragraphs[-1]
    inline_controls = last_para.content_controls
    assert len(inline_controls) == 1, (
        f"expected 1 inline control in final paragraph, got {len(inline_controls)}"
    )
    inline = inline_controls[0]
    assert inline.type is ContentControlType.RICH_TEXT
    assert inline.tag == "inline-rt"
    assert inline.text == "inline value"


def main() -> None:
    document = _build_document()
    _self_validate(document)
    document.save(_OUT)
    # -- reopen to make sure it round-trips through the package serializer --
    reopened = Document(_OUT)
    _self_validate(reopened)
    print(f"wrote {_OUT}")


if __name__ == "__main__":
    main()
