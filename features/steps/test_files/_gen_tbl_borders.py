"""Generate the `tbl-borders.docx` behave fixture.

Running this script (idempotently) produces a .docx containing two small
tables that exercise the ``Table.borders`` and ``_Cell.borders`` proxies:

- Table 0 (2x2) has explicit table-level borders on all six edges
  (top, bottom, left, right, insideH, insideV) — each a SINGLE style,
  half-point width, black.
- Table 1 (2x2) is plain (no directly-applied borders). The top-left
  cell has a cell-level left border (THICK, 1pt, red) applied so that
  ``_Cell.borders`` round-trips can be verified.

The script self-validates the written document by re-opening it and
asserting the expected border-element shape before considering the
fixture good.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_BORDER_STYLE
from docx.shared import Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "tbl-borders.docx")


def build() -> None:
    document = Document()

    # -- Table 0: all six table-level borders set --
    t0 = document.add_table(rows=2, cols=2)
    t0.cell(0, 0).text = "A"
    t0.cell(0, 1).text = "B"
    t0.cell(1, 0).text = "C"
    t0.cell(1, 1).text = "D"
    for edge in ("top", "bottom", "left", "right", "inside_h", "inside_v"):
        border = getattr(t0.borders, edge)
        border.style = WD_BORDER_STYLE.SINGLE
        border.width = Pt(0.5)
        border.color = RGBColor(0x00, 0x00, 0x00)

    document.add_paragraph()  # keep tables separated

    # -- Table 1: no table borders; cell-level border on (0, 0) --
    t1 = document.add_table(rows=2, cols=2)
    t1.cell(0, 0).text = "E"
    t1.cell(0, 1).text = "F"
    t1.cell(1, 0).text = "G"
    t1.cell(1, 1).text = "H"
    left = t1.cell(0, 0).borders.left
    left.style = WD_BORDER_STYLE.THICK
    left.width = Pt(1)
    left.color = RGBColor(0xFF, 0x00, 0x00)

    document.save(OUT_PATH)


def validate() -> None:
    document = Document(OUT_PATH)
    assert len(document.tables) == 2, f"expected 2 tables, got {len(document.tables)}"

    t0 = document.tables[0]
    for edge in ("top", "bottom", "left", "right", "inside_h", "inside_v"):
        b = getattr(t0.borders, edge)
        assert b.style == WD_BORDER_STYLE.SINGLE, f"{edge} style: got {b.style!r}"
        assert b.width == Pt(0.5), f"{edge} width: got {b.width!r}"
        assert b.color == RGBColor(0, 0, 0), f"{edge} color: got {b.color!r}"

    t1 = document.tables[1]
    # -- no table-level borders --
    for edge in ("top", "bottom", "left", "right", "inside_h", "inside_v"):
        b = getattr(t1.borders, edge)
        assert b.style is None, f"table 1 {edge} style should be None, got {b.style!r}"

    # -- cell (0,0) left border should round-trip --
    left = t1.cell(0, 0).borders.left
    assert left.style == WD_BORDER_STYLE.THICK, f"got {left.style!r}"
    assert left.width == Pt(1), f"got {left.width!r}"
    assert left.color == RGBColor(0xFF, 0x00, 0x00), f"got {left.color!r}"

    # -- cell (0,0) top/bottom/right should be unset --
    cell_borders = t1.cell(0, 0).borders
    for edge in ("top", "bottom", "right"):
        b = getattr(cell_borders, edge)
        assert b.style is None, f"cell (0,0) {edge} should be unset; got {b.style!r}"


def main() -> None:
    build()
    validate()
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
