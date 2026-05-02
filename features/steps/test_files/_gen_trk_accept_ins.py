"""Build the ``trk-accept-ins.docx`` fixture for tracked-change accept-side tests.

Run from anywhere (the output path is resolved relative to this file)::

    python features/steps/test_files/_gen_trk_accept_ins.py

*Generator strategy*: live paragraph/run/table scaffolding is built with
python-docx's public API, then raw ``lxml``/``OxmlElement`` injection is used
for the track-change wrappers (``w:ins``, ``w:del``, ``w:cellIns``). See the
sibling ``_gen_trk_ins_del.py`` for the same pattern.

The fixture exercises the acceptance-side code paths for insertions:

* ``paragraph 1`` — a single ``w:ins`` wrapping a single run
* ``paragraph 2`` — a single ``w:ins`` containing two ``w:r`` children
* ``paragraph 3`` — two separate ``w:ins`` wrappers in the same paragraph
  (used by per-change accept scenarios)
* ``paragraph 4`` — a ``w:ins`` and a ``w:del`` in the same paragraph (used
  by the "accept-all cleans up mixed content" scenario)
* ``paragraph 5`` — plain paragraph, no tracked changes
* a 1x2 table whose second cell carries a ``w:cellIns`` marker

Self-checks run before the file is written confirm the tracked-change layout
is as expected, so a broken generator fails loudly instead of producing a
fixture that silently skips assertions.
"""

from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-accept-ins.docx")


def _make_run(text: str):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _make_ins_with_runs(run_id: int, author: str, date: str, texts: list[str]):
    """Return a ``w:ins`` element wrapping one ``w:r`` per entry in ``texts``."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(run_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date)
    for text in texts:
        ins.append(_make_run(text))
    return ins


def _make_del(run_id: int, author: str, date: str, text: str):
    wdel = OxmlElement("w:del")
    wdel.set(qn("w:id"), str(run_id))
    wdel.set(qn("w:author"), author)
    wdel.set(qn("w:date"), date)
    r = OxmlElement("w:r")
    dt_elm = OxmlElement("w:delText")
    dt_elm.text = text
    dt_elm.set(qn("xml:space"), "preserve")
    r.append(dt_elm)
    wdel.append(r)
    return wdel


def _make_cell_ins(marker_id: int, author: str, date: str):
    elm = OxmlElement("w:cellIns")
    elm.set(qn("w:id"), str(marker_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    return elm


def build() -> Document:
    document = Document()

    # -- paragraph 0: title --
    document.add_heading("Tracked accepted insertions", level=1)

    # -- paragraph 1: "The quick [+nimble+] fox jumps." --
    p1 = document.add_paragraph("The quick ")
    p1._p.append(
        _make_ins_with_runs(1, "Alice", "2025-04-10T09:00:00Z", ["nimble"])
    )
    p1.add_run(" fox jumps.")

    # -- paragraph 2: single w:ins wrapping two w:r children --
    p2 = document.add_paragraph("Red fish ")
    p2._p.append(
        _make_ins_with_runs(
            2, "Alice", "2025-04-10T09:05:00Z", ["blue ", "fish"]
        )
    )

    # -- paragraph 3: two separate w:ins wrappers in the same paragraph --
    p3 = document.add_paragraph("Alpha ")
    p3._p.append(
        _make_ins_with_runs(3, "Bob", "2025-04-10T09:10:00Z", ["beta "])
    )
    p3.add_run("gamma")
    p3._p.append(
        _make_ins_with_runs(4, "Carol", "2025-04-10T09:15:00Z", [" delta"])
    )

    # -- paragraph 4: mixed ins + del (sibling agents cover deletions, but the
    # -- accept-all pass resolves both in one go) --
    p4 = document.add_paragraph("Hello ")
    p4._p.append(
        _make_del(5, "Dave", "2025-04-10T09:20:00Z", "dark")
    )
    p4._p.append(
        _make_ins_with_runs(6, "Alice", "2025-04-10T09:21:00Z", ["bright"])
    )
    p4.add_run(" world.")

    # -- paragraph 5: plain, no tracked changes --
    document.add_paragraph("Nothing to see here.")

    # -- 1x2 table: the second cell is a tracked insertion (w:cellIns) --
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "inserted"
    tc01 = table.cell(0, 1)._tc
    tcPr01 = tc01.get_or_add_tcPr()
    tcPr01.append(
        _make_cell_ins(7, "Bob", "2025-04-10T09:25:00Z")
    )

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    # -- paragraph 1: one insertion, text rendered with markers --
    p1_changes = paragraphs[1].tracked_changes
    assert len(p1_changes) == 1, (
        f"expected 1 tracked change on paragraph 1, got {len(p1_changes)}"
    )
    assert p1_changes[0].type == "insertion"
    assert p1_changes[0].author == "Alice"
    assert p1_changes[0].text == "nimble"
    assert (
        paragraphs[1].revision_marks_text()
        == "The quick [+nimble+] fox jumps."
    )

    # -- paragraph 2: one insertion whose text spans two runs --
    p2_changes = paragraphs[2].tracked_changes
    assert len(p2_changes) == 1
    assert p2_changes[0].type == "insertion"
    assert p2_changes[0].text == "blue fish"
    # -- the insertion wrapper in paragraph 2 contains two w:r children --
    ins_elms = paragraphs[2]._p.xpath("./w:ins")
    assert len(ins_elms) == 1
    inner_runs = ins_elms[0].xpath("./w:r")
    assert len(inner_runs) == 2, (
        f"expected 2 w:r children inside paragraph 2 w:ins, got {len(inner_runs)}"
    )

    # -- paragraph 3: two separate insertions --
    p3_changes = paragraphs[3].tracked_changes
    assert len(p3_changes) == 2
    assert [tc.type for tc in p3_changes] == ["insertion", "insertion"]
    assert [tc.text for tc in p3_changes] == ["beta ", " delta"]
    assert [tc.author for tc in p3_changes] == ["Bob", "Carol"]

    # -- paragraph 4: mixed del + ins in document order --
    p4_changes = paragraphs[4].tracked_changes
    assert [tc.type for tc in p4_changes] == ["deletion", "insertion"], (
        f"unexpected types on paragraph 4: {[tc.type for tc in p4_changes]}"
    )
    assert (
        paragraphs[4].revision_marks_text()
        == "Hello [-dark-][+bright+] world."
    )

    # -- paragraph 5 is clean --
    assert paragraphs[5].tracked_changes == []

    # -- table: cell (0,1) is flagged as tracked insertion --
    table = document.tables[0]
    assert table.cell(0, 0).is_tracked_insertion is False
    assert table.cell(0, 1).is_tracked_insertion is True
    assert table.cell(0, 1).text == "inserted"

    # -- whole-body XPath sanity check: five w:ins, one w:del, one w:cellIns --
    body = document.element.body
    assert len(body.xpath(".//w:ins")) == 5
    assert len(body.xpath(".//w:del")) == 1
    assert len(body.xpath(".//w:cellIns")) == 1


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
