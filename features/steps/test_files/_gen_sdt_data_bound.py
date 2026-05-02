"""Generate `sdt-data-bound.docx` fixture for behave acceptance tests.

Creates a Word document containing a data-bound content control plus a matching
custom XML data part (with sibling properties part declaring a store-item id).
Self-validates before saving.

Run directly::

    python features/steps/test_files/_gen_sdt_data_bound.py
"""

from __future__ import annotations

import os

from docx import Document
from docx.content_controls import ContentControlType
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.parts.custom_xml import CustomXmlPart as CustomXmlDataPart


_HERE = os.path.dirname(os.path.abspath(__file__))
_OUT = os.path.join(_HERE, "sdt-data-bound.docx")

# -- fixture constants referenced by the feature file --
ITEM_ID = "{11111111-2222-3333-4444-555555555555}"
PREFIX_MAPPINGS = "xmlns:ns0='http://example.com/orders'"
XPATH = "/ns0:order[1]/ns0:customer[1]"
ROOT_NS = "http://example.com/orders"

_DATA_BLOB = (
    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    b'<ns0:order xmlns:ns0="http://example.com/orders">\n'
    b"  <ns0:customer>Acme Co.</ns0:customer>\n"
    b"  <ns0:total>42.00</ns0:total>\n"
    b"</ns0:order>\n"
)

_PROPS_BLOB = (
    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    b'<ds:datastoreItem ds:itemID="' + ITEM_ID.encode("utf-8") + b'"'
    b' xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml">\n'
    b'  <ds:schemaRefs>\n'
    b'    <ds:schemaRef ds:uri="http://example.com/orders"/>\n'
    b'  </ds:schemaRefs>\n'
    b"</ds:datastoreItem>\n"
)


def _attach_custom_xml_part(document: Document) -> CustomXmlDataPart:
    """Create and relate a new custom XML data part (plus props part)."""
    document_part = document.part
    package = document_part.package

    # -- data part: /customXml/item2.xml (item1.xml ships with the default template) --
    data_partname = PackURI("/customXml/item2.xml")
    data_part = CustomXmlDataPart(data_partname, CT.XML, _DATA_BLOB)

    # -- props part: /customXml/itemProps2.xml --
    from docx.opc.part import Part

    props_partname = PackURI("/customXml/itemProps2.xml")
    props_part = Part(
        props_partname, CT.OFC_CUSTOM_XML_PROPERTIES, _PROPS_BLOB, package
    )

    # -- relate the data part from the document part (customXml rel), and the
    #    props part from the data part (customXmlProps rel). --
    document_part.relate_to(data_part, RT.CUSTOM_XML)
    data_part.relate_to(props_part, RT.CUSTOM_XML_PROPS)

    return data_part


def _build_document() -> Document:
    document = Document()

    _attach_custom_xml_part(document)

    document.add_heading("Data-bound SDT sample", level=1)

    # -- bound block-level control --
    bound = document.add_content_control(
        ContentControlType.PLAIN_TEXT, tag="customer", title="Customer"
    )
    bound.text = "Acme Co."
    bound.set_data_binding(
        xpath=XPATH, prefix_mappings=PREFIX_MAPPINGS, store_item_id=ITEM_ID
    )

    # -- plus an unbound sibling for negative assertions --
    unbound = document.add_content_control(
        ContentControlType.PLAIN_TEXT, tag="note", title="Note"
    )
    unbound.text = "unbound"

    return document


def _self_validate(document: Document) -> None:
    """Run read-side assertions before writing to disk."""
    # -- at least two custom XML parts (default bibliography + our new one) --
    parts = document.custom_xml_parts
    assert len(parts) >= 1, f"expected at least one custom XML part, got {len(parts)}"
    ours = [p for p in parts if p.item_id == ITEM_ID]
    assert len(ours) == 1, f"expected exactly one matching custom XML part, got {len(ours)}"
    part = ours[0]
    assert ROOT_NS in part.schema_refs, (
        f"expected schema ref {ROOT_NS!r} in {part.schema_refs!r}"
    )
    root = part.root_element
    assert root is not None, "custom XML data part did not parse"

    controls = document.content_controls
    assert len(controls) == 2, f"expected 2 block controls, got {len(controls)}"
    bound, unbound = controls

    binding = bound.data_binding
    assert binding is not None, "expected bound control to have a data binding"
    assert binding.xpath == XPATH, f"xpath mismatch: {binding.xpath!r}"
    assert binding.prefix_mappings == PREFIX_MAPPINGS, (
        f"prefix_mappings mismatch: {binding.prefix_mappings!r}"
    )
    assert binding.store_item_id == ITEM_ID, (
        f"store_item_id mismatch: {binding.store_item_id!r}"
    )
    assert bound.text == "Acme Co."

    assert unbound.data_binding is None, "expected sibling control to be unbound"


def main() -> None:
    document = _build_document()
    _self_validate(document)
    document.save(_OUT)
    # -- reopen to confirm the binding round-trips through the serializer --
    reopened = Document(_OUT)
    _self_validate(reopened)
    print(f"wrote {_OUT}")


if __name__ == "__main__":
    main()
