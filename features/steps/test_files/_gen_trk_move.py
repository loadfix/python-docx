"""Build the ``trk-move.docx`` fixture for the move-revision tests.

Run::

    python features/steps/test_files/_gen_trk_move.py

*Generator strategy*: raw ``OxmlElement`` injection for the ``w:moveFrom`` /
``w:moveTo`` wrappers. No public create API is exposed; the read-side
|MoveRevision| proxy (``.name``, ``.peer``) is what the acceptance tests pin.

The fixture creates a move named ``pair1``:

* paragraph 1 is the *source* — it holds a ``w:moveFrom`` around the text
  "moved text" whose runs use ``w:delText``
* paragraph 2 is the *destination* — it holds a ``w:moveTo`` carrying the same
  runs (with ``w:t``)

Self-checks:

* paragraph 1 exposes a single |MoveRevision| of type ``move_from``
* paragraph 2 exposes a single |MoveRevision| of type ``move_to``
* both share ``@w:name == "pair1"`` and each resolves the other as ``.peer``
"""

from __future__ import annotations

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.tracked_changes import MoveRevision

HERE = os.path.abspath(os.path.dirname(__file__))
OUT_PATH = os.path.join(HERE, "trk-move.docx")


def _make_move_from(change_id: int, author: str, date: str, name: str, text: str):
    elm = OxmlElement("w:moveFrom")
    elm.set(qn("w:id"), str(change_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    elm.set(qn("w:name"), name)
    r = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    del_text.text = text
    del_text.set(qn("xml:space"), "preserve")
    r.append(del_text)
    elm.append(r)
    return elm


def _make_move_to(change_id: int, author: str, date: str, name: str, text: str):
    elm = OxmlElement("w:moveTo")
    elm.set(qn("w:id"), str(change_id))
    elm.set(qn("w:author"), author)
    elm.set(qn("w:date"), date)
    elm.set(qn("w:name"), name)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    elm.append(r)
    return elm


def build() -> Document:
    document = Document()

    # -- paragraph 0: title --
    document.add_heading("Tracked move revisions", level=1)

    # -- paragraph 1: source --
    p1 = document.add_paragraph("Source: ")
    p1._p.append(
        _make_move_from(
            1, "Alice", "2025-04-12T10:00:00Z", "pair1", "moved text"
        )
    )
    p1.add_run(" (was here).")

    # -- paragraph 2: destination --
    p2 = document.add_paragraph("Destination: ")
    p2._p.append(
        _make_move_to(
            2, "Alice", "2025-04-12T10:00:00Z", "pair1", "moved text"
        )
    )
    p2.add_run(" (now here).")

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs

    p1_changes = paragraphs[1].tracked_changes
    assert len(p1_changes) == 1
    move_from = p1_changes[0]
    assert isinstance(move_from, MoveRevision)
    assert move_from.type == "move_from"
    assert move_from.name == "pair1"
    assert move_from.text == "moved text"

    p2_changes = paragraphs[2].tracked_changes
    assert len(p2_changes) == 1
    move_to = p2_changes[0]
    assert isinstance(move_to, MoveRevision)
    assert move_to.type == "move_to"
    assert move_to.name == "pair1"

    # -- peer lookup should find the other side --
    peer = move_from.peer
    assert peer is not None
    assert peer.type == "move_to"
    peer2 = move_to.peer
    assert peer2 is not None
    assert peer2.type == "move_from"


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
