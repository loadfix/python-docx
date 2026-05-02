"""Generate ``chart-create-pie-base.docx`` — a base fixture for pie-chart create tests.

The generated document is a tiny blank .docx containing a single heading
paragraph and a body paragraph, with NO chart embedded. Scenarios in
``features/chart-create-pie.feature`` use this file as a starting canvas
and call ``Document.add_chart(WD_CHART_TYPE.PIE, ...)`` on it.

Running the script is idempotent: it overwrites the output path. The script
also self-validates by reopening the saved document and asserting the
expected shape (no charts, two paragraphs).

Usage::

    python features/steps/test_files/_gen_chart_create_pie_base.py
"""

from __future__ import annotations

import os

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "chart-create-pie-base.docx")


def main() -> str:
    document = Document()
    document.add_heading("Pie chart create fixture", level=1)
    document.add_paragraph(
        "This document is a blank canvas for the chart-create-pie behave "
        "scenarios. It contains no embedded chart on purpose."
    )
    document.save(OUT_PATH)

    # -- self-validate: reopen and assert no charts plus two paragraphs --
    reopened = Document(OUT_PATH)
    assert reopened.charts == [], (
        f"base fixture should have no charts; got {reopened.charts!r}"
    )
    paragraphs = list(reopened.paragraphs)
    assert len(paragraphs) == 2, (
        f"expected 2 paragraphs in base fixture, got {len(paragraphs)}"
    )

    return OUT_PATH


if __name__ == "__main__":
    path = main()
    print(f"wrote {path}")
