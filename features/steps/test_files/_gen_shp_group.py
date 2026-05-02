"""Generate `shp-group.docx` behave fixture.

Produces a document whose body contains a single ``w:drawing`` whose root
DrawingML object is a ``wpg:grpSp`` (group shape). The group holds two
``wps:wsp`` shapes and a nested ``wpg:grpSp`` containing a third shape, to
exercise :class:`docx.drawing.GroupShape` traversal (including recursion).

Since ``Paragraph.add_shape`` creates a single ``wps:wsp``, this fixture
starts from an inline shape drawing and rewrites its
``a:graphicData`` subtree to wrap a group shape instead.

Run::

    python features/steps/test_files/_gen_shp_group.py
"""

from __future__ import annotations

import os

from lxml import etree

from docx import Document
from docx.drawing import GroupShape, WordprocessingShape
from docx.enum.shape import WD_DRAWING_TYPE, WD_SHAPE
from docx.oxml.ns import nsmap, qn
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "shp-group.docx")

_WPG_GRAPHIC_URI = "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"


def _wsp_xml(prst: str, shape_id: int, name: str, text: str | None = None) -> str:
    """Return a minimal ``<wps:wsp>`` XML fragment (no namespace declarations)."""
    txbx = ""
    if text is not None:
        txbx = (
            "<wps:txbx>"
            "<w:txbxContent><w:p><w:r><w:t>%s</w:t></w:r></w:p></w:txbxContent>"
            "</wps:txbx>"
        ) % text
    return (
        '<wps:wsp>'
        f'<wps:cNvPr id="{shape_id}" name="{name}"/>'
        '<wps:cNvSpPr/>'
        '<wps:spPr>'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="914400" cy="457200"/></a:xfrm>'
        f'<a:prstGeom prst="{prst}"><a:avLst/></a:prstGeom>'
        '<a:solidFill><a:srgbClr val="4472C4"/></a:solidFill>'
        '</wps:spPr>'
        f'{txbx}'
        '<wps:bodyPr rot="0" wrap="square" anchor="ctr"/>'
        '</wps:wsp>'
    )


def build() -> Document:
    document = Document()
    document.add_heading("Group shape fixture", level=1)
    document.add_paragraph(
        "The next paragraph contains a grouped DrawingML shape with nested children."
    )

    p = document.add_paragraph()
    # -- create a placeholder inline shape then swap graphicData contents --
    placeholder = p.add_shape(
        WD_SHAPE.RECTANGLE, width=Inches(2), height=Inches(1)
    )
    wsp = placeholder._wsp  # pyright: ignore[reportPrivateUsage]
    graphicData = wsp.getparent()
    assert graphicData is not None

    # -- remove the placeholder wsp and re-point graphicData at wpg: --
    graphicData.remove(wsp)
    graphicData.set("uri", _WPG_GRAPHIC_URI)

    # -- build the group shape XML tree --
    grpSp_xml = (
        '<wpg:grpSp xmlns:wpg="%s" xmlns:wps="%s" xmlns:w="%s" xmlns:a="%s">'
        '<wpg:nvGrpSpPr>'
        '<wpg:cNvPr id="100" name="Outer Group"/>'
        '<wpg:cNvGrpSpPr/>'
        '</wpg:nvGrpSpPr>'
        '<wpg:grpSpPr>'
        '<a:xfrm>'
        '<a:off x="0" y="0"/><a:ext cx="1828800" cy="914400"/>'
        '<a:chOff x="0" y="0"/><a:chExt cx="1828800" cy="914400"/>'
        '</a:xfrm>'
        '</wpg:grpSpPr>'
        '%s%s'
        # -- nested group with a single arrow shape --
        '<wpg:grpSp>'
        '<wpg:nvGrpSpPr>'
        '<wpg:cNvPr id="103" name="Inner Group"/>'
        '<wpg:cNvGrpSpPr/>'
        '</wpg:nvGrpSpPr>'
        '<wpg:grpSpPr>'
        '<a:xfrm>'
        '<a:off x="914400" y="0"/><a:ext cx="914400" cy="914400"/>'
        '<a:chOff x="0" y="0"/><a:chExt cx="914400" cy="914400"/>'
        '</a:xfrm>'
        '</wpg:grpSpPr>'
        '%s'
        '</wpg:grpSp>'
        '</wpg:grpSp>'
    ) % (
        nsmap["wpg"],
        nsmap["wps"],
        nsmap["w"],
        nsmap["a"],
        _wsp_xml("rect", 101, "Rect 1", text="Alpha"),
        _wsp_xml("ellipse", 102, "Oval 1"),
        _wsp_xml("rightArrow", 104, "Arrow 1"),
    )

    grpSp_elm = etree.fromstring(grpSp_xml)
    graphicData.append(grpSp_elm)

    return document


def self_validate(document: Document) -> None:
    paragraphs = document.paragraphs
    # -- the group is in paragraphs[2] (heading, intro, shape paragraph) --
    drawings = paragraphs[2].drawings
    assert len(drawings) == 1, f"expected 1 drawing, got {len(drawings)}"
    drawing = drawings[0]

    assert drawing.is_group is True, "drawing should report is_group=True"
    assert drawing.type is WD_DRAWING_TYPE.GROUP

    group = drawing.group_shape
    assert isinstance(group, GroupShape)
    assert group.name == "Outer Group"

    children = group.shapes
    # -- 2 shapes + 1 nested group --
    assert len(children) == 3, f"expected 3 direct children, got {len(children)}"

    # -- first child: rect with text --
    rect = children[0]
    assert isinstance(rect, WordprocessingShape)
    assert rect.shape_type is WD_SHAPE.RECTANGLE
    assert rect.text == "Alpha"

    # -- second child: oval --
    oval = children[1]
    assert isinstance(oval, WordprocessingShape)
    assert oval.shape_type is WD_SHAPE.OVAL

    # -- third child: nested group with one shape --
    inner = children[2]
    assert isinstance(inner, GroupShape), type(inner)
    assert inner.name == "Inner Group"
    inner_children = inner.shapes
    assert len(inner_children) == 1
    arrow = inner_children[0]
    assert isinstance(arrow, WordprocessingShape)
    assert arrow.shape_type is WD_SHAPE.ARROW_RIGHT

    # -- suppress warnings --
    assert qn("wpg:grpSp")


def main() -> None:
    document = build()
    self_validate(document)
    document.save(OUT_PATH)
    reopened = Document(OUT_PATH)
    self_validate(reopened)
    print(f"wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
