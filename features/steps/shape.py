"""Step implementations for graphical object (shape) related features."""

from __future__ import annotations

import hashlib

from behave import given, then, when

from docx import Document
from docx.drawing import GroupShape, WordprocessingShape
from docx.enum.shape import (
    WD_ANCHOR_H,
    WD_ANCHOR_V,
    WD_DRAWING_TYPE,
    WD_INLINE_SHAPE,
    WD_SHAPE,
    WD_WRAP_TYPE,
)
from docx.shape import FloatingImage, InlineShape
from docx.shared import Inches

from helpers import test_docx, test_file

# given ===================================================


@given("an inline shape collection containing five shapes")
def given_an_inline_shape_collection_containing_five_shapes(context):
    docx_path = test_docx("shp-inline-shape-access")
    document = Document(docx_path)
    context.inline_shapes = document.inline_shapes


@given("an inline shape of known dimensions")
def given_inline_shape_of_known_dimensions(context):
    document = Document(test_docx("shp-inline-shape-access"))
    context.inline_shape = document.inline_shapes[0]


@given("an inline shape known to be {shp_of_type}")
def given_inline_shape_known_to_be_shape_of_type(context, shp_of_type):
    inline_shape_idx = {
        "an embedded picture": 0,
        "a linked picture": 1,
        "a link+embed picture": 2,
        "a smart art diagram": 3,
        "a chart": 4,
    }[shp_of_type]
    docx_path = test_docx("shp-inline-shape-access")
    document = Document(docx_path)
    context.inline_shape = document.inline_shapes[inline_shape_idx]


# when =====================================================


@when("I change the dimensions of the inline shape")
def when_change_dimensions_of_inline_shape(context):
    inline_shape = context.inline_shape
    inline_shape.width = Inches(1)
    inline_shape.height = Inches(0.5)


# then =====================================================


@then("I can access each inline shape by index")
def then_can_access_each_inline_shape_by_index(context):
    inline_shapes = context.inline_shapes
    for idx in range(2):
        inline_shape = inline_shapes[idx]
        assert isinstance(inline_shape, InlineShape)


@then("I can iterate over the inline shape collection")
def then_can_iterate_over_inline_shape_collection(context):
    inline_shapes = context.inline_shapes
    shape_count = 0
    for inline_shape in inline_shapes:
        shape_count += 1
        assert isinstance(inline_shape, InlineShape)
    expected_count = 5
    assert shape_count == expected_count, "expected %d, got %d" % (
        expected_count,
        shape_count,
    )


@then("its inline shape type is {shape_type}")
def then_inline_shape_type_is_shape_type(context, shape_type):
    expected_value = {
        "WD_INLINE_SHAPE.CHART": WD_INLINE_SHAPE.CHART,
        "WD_INLINE_SHAPE.LINKED_PICTURE": WD_INLINE_SHAPE.LINKED_PICTURE,
        "WD_INLINE_SHAPE.PICTURE": WD_INLINE_SHAPE.PICTURE,
        "WD_INLINE_SHAPE.SMART_ART": WD_INLINE_SHAPE.SMART_ART,
    }[shape_type]
    inline_shape = context.inline_shape
    assert inline_shape.type == expected_value


@then("the dimensions of the inline shape match the known values")
def then_dimensions_of_inline_shape_match_known_values(context):
    inline_shape = context.inline_shape
    assert inline_shape.width == 1778000, "got %s" % inline_shape.width
    assert inline_shape.height == 711200, "got %s" % inline_shape.height


@then("the dimensions of the inline shape match the new values")
def then_dimensions_of_inline_shape_match_new_values(context):
    inline_shape = context.inline_shape
    assert inline_shape.width == 914400, "got %s" % inline_shape.width
    assert inline_shape.height == 457200, "got %s" % inline_shape.height


@then("the document contains the inline picture")
def then_the_document_contains_the_inline_picture(context):
    document = context.document
    picture_shape = document.inline_shapes[0]
    blip = picture_shape._inline.graphic.graphicData.pic.blipFill.blip
    rId = blip.embed
    image_part = document.part.related_parts[rId]
    image_sha1 = hashlib.sha1(image_part.blob).hexdigest()
    expected_sha1 = "79769f1e202add2e963158b532e36c2c0f76a70c"
    assert image_sha1 == expected_sha1, "image SHA1 doesn't match, expected %s, got %s" % (
        expected_sha1,
        image_sha1,
    )


@then("the length of the inline shape collection is 5")
def then_len_of_inline_shape_collection_is_5(context):
    inline_shapes = context.inline_shapes
    shape_count = len(inline_shapes)
    assert shape_count == 5, "got %s" % shape_count


@then("the picture has its native width and height")
def then_picture_has_native_width_and_height(context):
    picture = context.picture
    assert picture.width == 1905000, "got %d" % picture.width
    assert picture.height == 2717800, "got %d" % picture.height


@then("picture.height is {inches} inches")
def then_picture_height_is_value(context, inches):
    expected_value = {
        "2.14": 1956816,
        "2.5": 2286000,
    }[inches]
    picture = context.picture
    assert picture.height == expected_value, "got %d" % picture.height


@then("picture.width is {inches} inches")
def then_picture_width_is_value(context, inches):
    expected_value = {
        "1.05": 961402,
        "1.75": 1600200,
    }[inches]
    picture = context.picture
    assert picture.width == expected_value, "got %d" % picture.width


# ===========================================================================
# drawing / shape fork-era steps
# ===========================================================================


# -- enum lookup tables shared across scenarios --
_H_ANCHORS = {
    "PAGE": WD_ANCHOR_H.PAGE,
    "MARGIN": WD_ANCHOR_H.MARGIN,
    "COLUMN": WD_ANCHOR_H.COLUMN,
    "CHARACTER": WD_ANCHOR_H.CHARACTER,
}
_V_ANCHORS = {
    "PAGE": WD_ANCHOR_V.PAGE,
    "MARGIN": WD_ANCHOR_V.MARGIN,
    "PARAGRAPH": WD_ANCHOR_V.PARAGRAPH,
    "LINE": WD_ANCHOR_V.LINE,
}
_WRAPS = {
    "SQUARE": WD_WRAP_TYPE.SQUARE,
    "TIGHT": WD_WRAP_TYPE.TIGHT,
    "THROUGH": WD_WRAP_TYPE.THROUGH,
    "TOP_AND_BOTTOM": WD_WRAP_TYPE.TOP_AND_BOTTOM,
    "BEHIND": WD_WRAP_TYPE.BEHIND,
    "IN_FRONT": WD_WRAP_TYPE.IN_FRONT,
}
_SHAPE_TYPES = {
    "RECTANGLE": WD_SHAPE.RECTANGLE,
    "ROUNDED_RECTANGLE": WD_SHAPE.ROUNDED_RECTANGLE,
    "OVAL": WD_SHAPE.OVAL,
    "ARROW_RIGHT": WD_SHAPE.ARROW_RIGHT,
    "CALLOUT_ROUNDED_RECTANGLE": WD_SHAPE.CALLOUT_ROUNDED_RECTANGLE,
}
_DRAWING_TYPES = {
    "SHAPE": WD_DRAWING_TYPE.SHAPE,
    "TEXT_BOX": WD_DRAWING_TYPE.TEXT_BOX,
    "GROUP": WD_DRAWING_TYPE.GROUP,
    "CHART": WD_DRAWING_TYPE.CHART,
    "DIAGRAM": WD_DRAWING_TYPE.DIAGRAM,
    "PICTURE": WD_DRAWING_TYPE.PICTURE,
}


def _strip_quotes(value: str) -> str:
    if len(value) >= 2 and value[0] == value[-1] and value[0] in ('"', "'"):
        return value[1:-1]
    return value


# -- given ------------------------------------------------------------------


@given("a pristine empty document")
def given_a_pristine_empty_document(context):
    context.document = Document()


@given("a document known to contain three floating images")
def given_doc_three_floating(context):
    context.document = Document(test_docx("shp-floating"))


@given("a document known to contain five inline preset shapes")
def given_doc_five_preset_shapes(context):
    context.document = Document(test_docx("shp-preset-shape"))


@given("a document known to contain a DrawingML group shape")
def given_doc_group_shape(context):
    context.document = Document(test_docx("shp-group"))


@given("a document known to contain shape text frames")
def given_doc_text_box(context):
    context.document = Document(test_docx("shp-text-box"))


@given("a document known to contain two ink annotations")
def given_doc_ink(context):
    context.document = Document(test_docx("shp-ink"))


@given("a document known to contain an embedded OLE object")
def given_doc_ole(context):
    context.document = Document(test_docx("shp-ole"))


@given("a document known to contain an inline picture with alt text")
def given_doc_alt_text(context):
    context.document = Document(test_docx("shp-alt-text"))


# -- when -------------------------------------------------------------------


@when("I add a floating image to the paragraph with no position")
def when_add_floating_no_position(context):
    context.floating_image = context.paragraph.add_floating_image(
        test_file("monty-truth.png"),
        width=Inches(1.0),
    )


@when(
    "I add a floating image anchored {h_anchor}/{v_anchor}"
    " with wrap {wrap}"
)
def when_add_floating_with_position(context, h_anchor, v_anchor, wrap):
    position = {
        "h_anchor": _H_ANCHORS[h_anchor],
        "v_anchor": _V_ANCHORS[v_anchor],
        "horizontal": Inches(1),
        "vertical": Inches(1),
        "wrap": _WRAPS[wrap],
    }
    context.floating_image = context.paragraph.add_floating_image(
        test_file("monty-truth.png"),
        width=Inches(1.0),
        position=position,
    )


@when("I add a preset shape of type {shape_type} to the paragraph")
def when_add_preset_shape(context, shape_type):
    wd_shape = _SHAPE_TYPES[shape_type]
    context.wsp_shape = context.paragraph.add_shape(
        wd_shape, width=Inches(1.0), height=Inches(0.5)
    )


@when('I add a preset shape of type {shape_type} with text "{text}"')
def when_add_preset_shape_with_text(context, shape_type, text):
    wd_shape = _SHAPE_TYPES[shape_type]
    context.wsp_shape = context.paragraph.add_shape(
        wd_shape, width=Inches(1.0), height=Inches(0.5), text=text
    )


@when('I set the shape\'s text to "{text}"')
def when_set_shape_text(context, text):
    context.wsp_shape.text = text


@when("I add an SVG picture to the run")
def when_add_svg_picture(context):
    context.run.add_picture(test_file("test.svg"))


@when('I set the inline shape\'s alt_text to "{value}"')
def when_set_inline_alt_text_value(context, value):
    context.inline_shape.alt_text = value


@when("I set the inline shape's alt_text to None")
def when_set_inline_alt_text_none(context):
    context.inline_shape.alt_text = None


@when('I set the inline shape\'s title to "{value}"')
def when_set_inline_title_value(context, value):
    context.inline_shape.title = value


@when("I set the inline shape's title to None")
def when_set_inline_title_none(context):
    context.inline_shape.title = None


# -- then -------------------------------------------------------------------


@then("the paragraph has {n:d} floating image")
@then("the paragraph has {n:d} floating images")
def then_para_has_n_floating(context, n):
    floats = context.paragraph.floating_images
    assert len(floats) == n, f"expected {n}, got {len(floats)}"
    context.floating_image = floats[0]


@then("the floating image has horizontal anchor {name}")
def then_floating_h_anchor(context, name):
    assert context.floating_image.horizontal_anchor is _H_ANCHORS[name], (
        f"expected {name}, got {context.floating_image.horizontal_anchor}"
    )


@then("the floating image has vertical anchor {name}")
def then_floating_v_anchor(context, name):
    assert context.floating_image.vertical_anchor is _V_ANCHORS[name], (
        f"expected {name}, got {context.floating_image.vertical_anchor}"
    )


@then("the floating image has wrap type {name}")
def then_floating_wrap(context, name):
    assert context.floating_image.wrap_type is _WRAPS[name], (
        f"expected {name}, got {context.floating_image.wrap_type}"
    )


@then("the document has {n:d} floating images across its paragraphs")
def then_doc_has_n_floating(context, n):
    floating = [
        fi for p in context.document.paragraphs for fi in p.floating_images
    ]
    assert len(floating) == n, f"expected {n}, got {len(floating)}"
    context.floating_images = floating


@then("the second floating image has horizontal anchor {name}")
def then_second_floating_h_anchor(context, name):
    fi = context.floating_images[1]
    assert fi.horizontal_anchor is _H_ANCHORS[name]


@then("the second floating image has vertical anchor {name}")
def then_second_floating_v_anchor(context, name):
    fi = context.floating_images[1]
    assert fi.vertical_anchor is _V_ANCHORS[name]


@then("the second floating image has horizontal offset {n:d}")
def then_second_floating_h_offset(context, n):
    fi = context.floating_images[1]
    assert fi.horizontal_offset == n, f"expected {n}, got {fi.horizontal_offset}"


@then("the second floating image has vertical offset {n:d}")
def then_second_floating_v_offset(context, n):
    fi = context.floating_images[1]
    assert fi.vertical_offset == n, f"expected {n}, got {fi.vertical_offset}"


@then('the third floating image has alt text "{value}"')
def then_third_floating_alt_text(context, value):
    fi = context.floating_images[2]
    assert fi.alt_text == value, f"expected {value!r}, got {fi.alt_text!r}"


@then('the third floating image has title "{value}"')
def then_third_floating_title(context, value):
    fi = context.floating_images[2]
    assert fi.title == value, f"expected {value!r}, got {fi.title!r}"


@then("the second floating image position dict has the expected keys")
def then_position_dict(context):
    fi = context.floating_images[1]
    pos = fi.position
    assert set(pos.keys()) == {"h_anchor", "v_anchor", "horizontal", "vertical"}
    assert pos["h_anchor"] is WD_ANCHOR_H.PAGE
    assert pos["v_anchor"] is WD_ANCHOR_V.PAGE
    assert pos["horizontal"] == fi.horizontal_offset
    assert pos["vertical"] == fi.vertical_offset


@then("the paragraph has {n:d} inline drawing")
@then("the paragraph has {n:d} inline drawings")
def then_paragraph_has_n_drawings(context, n):
    drawings = context.paragraph.drawings
    assert len(drawings) == n, f"expected {n}, got {len(drawings)}"
    context.drawing = drawings[0]


@then("the drawing type is {name}")
def then_drawing_type_is(context, name):
    assert context.drawing.type is _DRAWING_TYPES[name], (
        f"expected {name}, got {context.drawing.type}"
    )


@then("the wps:wsp shape type is {name}")
def then_wsp_type(context, name):
    wsp_shape = context.wsp_shape
    assert isinstance(wsp_shape, WordprocessingShape)
    assert wsp_shape.shape_type is _SHAPE_TYPES[name], (
        f"expected {name}, got {wsp_shape.shape_type}"
    )


@then('the wps:wsp shape text is "{value}"')
def then_wsp_text_quoted(context, value):
    actual = context.wsp_shape.text
    assert actual == value, f"expected {value!r}, got {actual!r}"


@then('the wps:wsp shape text is ""')
def then_wsp_text_empty(context):
    actual = context.wsp_shape.text
    assert actual == "", f"expected empty, got {actual!r}"


@then("calling add_shape with a non-WD_SHAPE argument raises TypeError")
def then_add_shape_rejects_non_enum(context):
    try:
        context.paragraph.add_shape("rect")  # intentionally wrong type
    except TypeError:
        return
    raise AssertionError("expected TypeError")


@then("the document's third inline drawing is a SHAPE of type OVAL")
def then_third_drawing_is_oval(context):
    drawings = [
        d
        for p in context.document.paragraphs
        for d in p.drawings
    ]
    drawing = drawings[2]
    assert drawing.type is WD_DRAWING_TYPE.SHAPE
    wsps = drawing._drawing.xpath(  # pyright: ignore[reportPrivateUsage]
        ".//wps:wsp"
    )
    wsp = WordprocessingShape(wsps[0], context.document.paragraphs[0])
    assert wsp.shape_type is WD_SHAPE.OVAL


@then('the document\'s fifth inline drawing is a TEXT_BOX with text "{text}"')
def then_fifth_drawing_is_text_box(context, text):
    drawings = [
        d
        for p in context.document.paragraphs
        for d in p.drawings
    ]
    drawing = drawings[4]
    assert drawing.type is WD_DRAWING_TYPE.TEXT_BOX
    assert drawing.text == text, f"expected {text!r}, got {drawing.text!r}"


# -- group shape scenarios --


def _group_drawing(context):
    for p in context.document.paragraphs:
        for d in p.drawings:
            if d.is_group:
                return d
    raise AssertionError("no group drawing in document")


@then("the grouped drawing reports is_group is True")
def then_grouped_is_group(context):
    drawing = _group_drawing(context)
    assert drawing.is_group is True
    context.group_drawing = drawing


@then("the grouped drawing type is {name}")
def then_grouped_drawing_type(context, name):
    drawing = _group_drawing(context)
    assert drawing.type is _DRAWING_TYPES[name]


@then('the outer group has name "{value}"')
def then_outer_group_name(context, value):
    drawing = _group_drawing(context)
    group = drawing.group_shape
    assert isinstance(group, GroupShape)
    assert group.name == value
    context.outer_group = group


@then("the outer group has {n:d} top-level children")
def then_outer_group_children_count(context, n):
    drawing = _group_drawing(context)
    group = drawing.group_shape
    assert isinstance(group, GroupShape)
    assert len(group.shapes) == n
    context.outer_group = group


@then("the first child is a WordprocessingShape of type {name}")
def then_first_child_wsp(context, name):
    child = context.outer_group.shapes[0]
    assert isinstance(child, WordprocessingShape), type(child)
    assert child.shape_type is _SHAPE_TYPES[name]


@then("the second child is a WordprocessingShape of type {name}")
def then_second_child_wsp(context, name):
    child = context.outer_group.shapes[1]
    assert isinstance(child, WordprocessingShape), type(child)
    assert child.shape_type is _SHAPE_TYPES[name]


@then("the third child is a nested GroupShape")
def then_third_child_nested_group(context):
    child = context.outer_group.shapes[2]
    assert isinstance(child, GroupShape), type(child)
    context.inner_group = child


@then('the nested group has name "{value}"')
def then_nested_group_name(context, value):
    drawing = _group_drawing(context)
    group = drawing.group_shape
    assert isinstance(group, GroupShape)
    inner = next(c for c in group.shapes if isinstance(c, GroupShape))
    assert inner.name == value
    context.inner_group = inner


@then("the nested group has {n:d} top-level child")
@then("the nested group has {n:d} top-level children")
def then_nested_group_children_count(context, n):
    assert len(context.inner_group.shapes) == n


@then("the nested group's first child is a WordprocessingShape of type {name}")
def then_nested_first_child_wsp(context, name):
    child = context.inner_group.shapes[0]
    assert isinstance(child, WordprocessingShape), type(child)
    assert child.shape_type is _SHAPE_TYPES[name]


@then('the first shape inside the group has text "{value}"')
def then_first_group_shape_text(context, value):
    drawing = _group_drawing(context)
    group = drawing.group_shape
    assert isinstance(group, GroupShape)
    wsp = next(c for c in group.shapes if isinstance(c, WordprocessingShape))
    assert wsp.text == value, f"expected {value!r}, got {wsp.text!r}"


# -- text-box content scenarios --


def _drawings_in_document(document):
    return [d for p in document.paragraphs for d in p.drawings]


@then('the first shape text frame has text "{value}"')
def then_first_textbox_text(context, value):
    drawings = _drawings_in_document(context.document)
    text_box_drawings = [d for d in drawings if d.type is WD_DRAWING_TYPE.TEXT_BOX]
    assert text_box_drawings, "expected a text-box drawing"
    assert text_box_drawings[0].text == value, (
        f"expected {value!r}, got {text_box_drawings[0].text!r}"
    )


@then("the first shape text frame has {n:d} paragraph")
@then("the first shape text frame has {n:d} paragraphs")
def then_first_textbox_paragraph_count(context, n):
    drawings = _drawings_in_document(context.document)
    text_box_drawings = [d for d in drawings if d.type is WD_DRAWING_TYPE.TEXT_BOX]
    assert len(text_box_drawings[0].paragraphs) == n, (
        f"expected {n}, got {len(text_box_drawings[0].paragraphs)}"
    )


@then("the second shape text frame has text spanning {n:d} paragraphs")
def then_second_textbox_spans_paragraphs(context, n):
    drawings = _drawings_in_document(context.document)
    text_box_drawings = [d for d in drawings if d.type is WD_DRAWING_TYPE.TEXT_BOX]
    assert len(text_box_drawings) >= 2, "expected at least two text-box drawings"
    text = text_box_drawings[1].text
    assert text.count("\n") == n - 1, (
        f"expected {n} paragraphs worth of text, got {text!r}"
    )


@then("the second shape text frame exposes the expected paragraph texts")
def then_second_textbox_paragraphs(context):
    drawings = _drawings_in_document(context.document)
    text_box_drawings = [d for d in drawings if d.type is WD_DRAWING_TYPE.TEXT_BOX]
    texts = [p.text for p in text_box_drawings[1].paragraphs]
    assert texts == ["First line", "Second line", "Third line"], texts


# -- ink annotation scenarios --


@then("the document exposes {n:d} ink annotations")
def then_doc_exposes_n_ink(context, n):
    annotations = context.document.ink_annotations
    assert len(annotations) == n, f"expected {n}, got {len(annotations)}"


@then(
    'the ink annotations have partnames "{first}" and "{second}"'
)
def then_ink_partnames(context, first, second):
    annotations = context.document.ink_annotations
    partnames = sorted(a.partname for a in annotations)
    assert partnames == sorted([first, second]), f"got {partnames}"


@then('the ink annotation at partname "{partname}" has {n:d} strokes')
def then_ink_strokes(context, partname, n):
    annotations = context.document.ink_annotations
    match = next(a for a in annotations if a.partname == partname)
    assert match.stroke_count == n, f"expected {n}, got {match.stroke_count}"


@then("the first paragraph carrying ink has {n:d} annotation")
def then_first_para_carrying_ink(context, n):
    carriers = [p for p in context.document.paragraphs if p.ink_annotations]
    assert len(carriers) >= 1, "expected at least one carrier paragraph"
    assert len(carriers[0].ink_annotations) == n


@then("the second paragraph carrying ink has {n:d} annotation")
def then_second_para_carrying_ink(context, n):
    carriers = [p for p in context.document.paragraphs if p.ink_annotations]
    assert len(carriers) >= 2, "expected at least two carrier paragraphs"
    assert len(carriers[1].ink_annotations) == n


@then("each annotation's paragraph is the paragraph that contains it")
def then_each_ink_paragraph(context):
    for p in context.document.paragraphs:
        for a in p.ink_annotations:
            # identity can differ between successive accesses; compare the
            # underlying CT_P element instead.
            assert (
                a.paragraph._p  # pyright: ignore[reportPrivateUsage]
                is p._p  # pyright: ignore[reportPrivateUsage]
            )


# -- OLE embedded object scenarios --


@then("the document exposes {n:d} embedded objects")
def then_doc_exposes_n_ole(context, n):
    assert len(context.document.embedded_objects) == n


@then('the resolved embedded object has prog_id "{value}"')
def then_resolved_ole_prog_id(context, value):
    resolved = [o for o in context.document.embedded_objects if o.blob]
    assert resolved[0].prog_id == value


@then('the resolved embedded object has type "{value}"')
def then_resolved_ole_type(context, value):
    resolved = [o for o in context.document.embedded_objects if o.blob]
    assert resolved[0].type == value


@then("the resolved embedded object blob is non-empty")
def then_resolved_ole_blob_non_empty(context):
    resolved = [o for o in context.document.embedded_objects if o.blob]
    assert resolved[0].blob


@then('the resolved embedded object has embedded_partname "{value}"')
def then_resolved_ole_partname(context, value):
    resolved = [o for o in context.document.embedded_objects if o.blob]
    assert resolved[0].embedded_partname == value


@then('the unresolved embedded object has prog_id "{value}"')
def then_unresolved_ole_prog_id(context, value):
    unresolved = [o for o in context.document.embedded_objects if not o.blob]
    assert unresolved[0].prog_id == value


@then('the unresolved embedded object has type "{value}"')
def then_unresolved_ole_type(context, value):
    unresolved = [o for o in context.document.embedded_objects if not o.blob]
    assert unresolved[0].type == value


@then("the unresolved embedded object blob is empty")
def then_unresolved_ole_blob_empty(context):
    unresolved = [o for o in context.document.embedded_objects if not o.blob]
    assert unresolved[0].blob == b""


@then("the unresolved embedded object has embedded_partname None")
def then_unresolved_ole_partname_none(context):
    unresolved = [o for o in context.document.embedded_objects if not o.blob]
    assert unresolved[0].embedded_partname is None


@then("the first paragraph carrying an OLE reference has {n:d} embedded object")
def then_first_para_carrying_ole(context, n):
    carriers = [p for p in context.document.paragraphs if p.embedded_objects]
    assert carriers, "expected at least one carrier paragraph"
    assert len(carriers[0].embedded_objects) == n


@then("the embedded object paragraph attribute is the paragraph that contains it")
def then_ole_paragraph_attr(context):
    for p in context.document.paragraphs:
        for o in p.embedded_objects:
            assert (
                o.paragraph._p  # pyright: ignore[reportPrivateUsage]
                is p._p  # pyright: ignore[reportPrivateUsage]
            )


# -- alt text scenarios --


def _inline_shape(context, idx: int) -> InlineShape:
    return context.document.inline_shapes[idx]


def _floating_images(context) -> list[FloatingImage]:
    return [
        fi for p in context.document.paragraphs for fi in p.floating_images
    ]


@then('the first inline shape alt_text is "{value}"')
def then_first_inline_alt_text_value(context, value):
    actual = _inline_shape(context, 0).alt_text
    assert actual == value, f"expected {value!r}, got {actual!r}"


@then('the first inline shape title is "{value}"')
def then_first_inline_title_value(context, value):
    actual = _inline_shape(context, 0).title
    assert actual == value, f"expected {value!r}, got {actual!r}"


@then("the second inline shape alt_text is None")
def then_second_inline_alt_text_none(context):
    assert _inline_shape(context, 1).alt_text is None


@then("the second inline shape title is None")
def then_second_inline_title_none(context):
    assert _inline_shape(context, 1).title is None


@then('the first floating picture alt_text is "{value}"')
def then_first_floating_alt_text(context, value):
    fi = _floating_images(context)[0]
    assert fi.alt_text == value, f"expected {value!r}, got {fi.alt_text!r}"


@then('the first floating picture title is "{value}"')
def then_first_floating_title(context, value):
    fi = _floating_images(context)[0]
    assert fi.title == value, f"expected {value!r}, got {fi.title!r}"


@then("the second floating picture alt_text is None")
def then_second_floating_alt_text_none(context):
    fi = _floating_images(context)[1]
    assert fi.alt_text is None


@then("the second floating picture title is None")
def then_second_floating_title_none(context):
    fi = _floating_images(context)[1]
    assert fi.title is None


@then('the inline shape alt_text is "{value}"')
def then_inline_alt_text_value(context, value):
    assert context.inline_shape.alt_text == value, (
        f"expected {value!r}, got {context.inline_shape.alt_text!r}"
    )


@then("the inline shape alt_text is None")
def then_inline_alt_text_none(context):
    assert context.inline_shape.alt_text is None


@then('the inline shape title is "{value}"')
def then_inline_title_value(context, value):
    assert context.inline_shape.title == value, (
        f"expected {value!r}, got {context.inline_shape.title!r}"
    )


@then("the inline shape title is None")
def then_inline_title_none(context):
    assert context.inline_shape.title is None


# -- SVG coverage for run-add-picture --


@then('the run\'s inline shape has content type "{mime}"')
def then_run_inline_svg_content_type(context, mime):
    run = context.run
    # -- find the most-recently-added inline shape's image part by rId --
    inline = run._r.xpath(".//w:drawing/wp:inline")[  # pyright: ignore[reportPrivateUsage]
        -1
    ]
    # -- for SVG inputs python-docx stores the SVG as an `asvg:svgBlip`
    # -- extension alongside a raster fallback referenced from `a:blip/@r:embed`.
    # -- check for the SVG rId first; fall back to the plain blip reference.
    svg_rIds = inline.xpath(".//asvg:svgBlip/@r:embed")
    rIds = svg_rIds or inline.xpath(".//a:blip/@r:embed")
    assert rIds, "inline has no blip r:embed"
    image_part = context.document.part.related_parts[rIds[0]]
    assert image_part.content_type == mime, (
        f"expected {mime!r}, got {image_part.content_type!r}"
    )


# -- reference assert for unused helpers (keep linter quiet) --
_ = _strip_quotes  # noqa: F841
