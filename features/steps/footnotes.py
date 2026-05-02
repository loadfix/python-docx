"""Step implementations for document footnote-related features."""

from __future__ import annotations

import io

from behave import given, then, when
from behave.runner import Context

from docx import Document
from docx.enum.text import (
    WD_FOOTNOTE_POSITION,
    WD_FOOTNOTE_RESTART,
    WD_NUMBER_FORMAT,
)
from docx.footnotes import Footnote, FootnoteProperties, Footnotes
from docx.oxml.ns import qn

from helpers import test_docx

# given ====================================================


@given("a document having a footnotes part")
def given_a_document_having_a_footnotes_part(context: Context):
    context.document = Document(test_docx("fnt-has-footnotes"))


@given("a document having no footnotes part")
def given_a_document_having_no_footnotes_part(context: Context):
    context.document = Document(test_docx("doc-default"))


@given("a document having {count} footnotes")
def given_a_document_having_count_footnotes(context: Context, count: str):
    testfile_name = {"0": "doc-default", "3": "fnt-has-footnotes"}[count]
    context.document = Document(test_docx(testfile_name))


@given("a document with footnote properties")
def given_a_document_with_footnote_properties(context: Context):
    context.document = Document(test_docx("fnt-has-footnote-pr"))


@given("a document without footnote properties")
def given_a_document_without_footnote_properties(context: Context):
    context.document = Document(test_docx("doc-default"))


# when =====================================================


@when('I add a paragraph with text "{text}"')
def when_I_add_a_paragraph_with_text(context: Context, text: str):
    context.paragraph = context.document.add_paragraph(text)


@when("I assign footnote = document.footnotes.add(paragraph.runs[0])")
def when_I_assign_footnote_eq_add_no_text(context: Context):
    context.footnote = context.document.footnotes.add(context.paragraph.runs[0])


@when('I assign footnote = document.footnotes.add(paragraph.runs[0], "{text}")')
def when_I_assign_footnote_eq_add_with_text(context: Context, text: str):
    context.footnote = context.document.footnotes.add(context.paragraph.runs[0], text)


@when('I assign fn1 = document.footnotes.add(paragraph.runs[0], "{text}")')
def when_I_assign_fn1(context: Context, text: str):
    context.fn1 = context.document.footnotes.add(context.paragraph.runs[0], text)


@when('I assign fn2 = document.footnotes.add(paragraph.runs[0], "{text}")')
def when_I_assign_fn2(context: Context, text: str):
    context.fn2 = context.document.footnotes.add(context.paragraph.runs[0], text)


@when("I save and reopen the document")
def when_I_save_and_reopen_the_document(context: Context):
    buf = io.BytesIO()
    context.document.save(buf)
    buf.seek(0)
    context.document = Document(buf)


@when("I delete the footnote with id {footnote_id:d}")
def when_I_delete_the_footnote_with_id(context: Context, footnote_id: int):
    footnote = _footnote_with_id(context.document, footnote_id)
    footnote.delete()


@when("I clear the footnote with id {footnote_id:d}")
def when_I_clear_the_footnote_with_id(context: Context, footnote_id: int):
    footnote = _footnote_with_id(context.document, footnote_id)
    footnote.clear()


@when('I call footnote.add_paragraph("{text}") on footnote with id {footnote_id:d}')
def when_I_call_footnote_add_paragraph(context: Context, text: str, footnote_id: int):
    footnote = _footnote_with_id(context.document, footnote_id)
    context.new_paragraph = footnote.add_paragraph(text)


@when("I call document.add_footnote_properties()")
def when_I_call_document_add_footnote_properties(context: Context):
    context.footnote_properties = context.document.add_footnote_properties()


@when("I assign WD_NUMBER_FORMAT.{member} to footnote_properties.number_format")
def when_I_assign_number_format(context: Context, member: str):
    context.footnote_properties.number_format = WD_NUMBER_FORMAT[member]


@when("I assign {value:d} to footnote_properties.start_number")
def when_I_assign_start_number(context: Context, value: int):
    context.footnote_properties.start_number = value


@when("I assign WD_FOOTNOTE_RESTART.{member} to footnote_properties.restart_rule")
def when_I_assign_restart_rule(context: Context, member: str):
    context.footnote_properties.restart_rule = WD_FOOTNOTE_RESTART[member]


@when("I assign WD_FOOTNOTE_POSITION.{member} to footnote_properties.position")
def when_I_assign_position(context: Context, member: str):
    context.footnote_properties.position = WD_FOOTNOTE_POSITION[member]


# then =====================================================


@then("footnote is a Footnote object")
def then_footnote_is_a_Footnote_object(context: Context):
    assert isinstance(context.footnote, Footnote), (
        f"expected Footnote, got {type(context.footnote)}"
    )


@then("footnote.footnote_id == {expected:d}")
def then_footnote_footnote_id_eq(context: Context, expected: int):
    actual = context.footnote.footnote_id
    assert actual == expected, f"expected footnote_id {expected}, got {actual}"


@then('footnote.text == "{expected}"')
def then_footnote_text_eq(context: Context, expected: str):
    actual = context.footnote.text
    assert actual == expected, f"expected text '{expected}', got '{actual}'"


@then("len(footnote.paragraphs) == {count:d}")
def then_len_footnote_paragraphs_eq(context: Context, count: int):
    actual = len(context.footnote.paragraphs)
    assert actual == count, f"expected len(footnote.paragraphs) of {count}, got {actual}"


@then('footnote.paragraphs[{idx:d}].style.name == "{style}"')
def then_footnote_paragraphs_idx_style_name_eq(context: Context, idx: int, style: str):
    actual = context.footnote.paragraphs[idx]._p.style  # pyright: ignore[reportPrivateUsage]
    assert actual == style, f"expected paragraph style '{style}', got '{actual}'"


@then("len(document.footnotes) == {count:d}")
def then_len_document_footnotes_eq(context: Context, count: int):
    actual = len(context.document.footnotes)
    assert actual == count, f"expected len(document.footnotes) of {count}, got {actual}"


@then("the anchor run contains a footnoteReference to footnote.footnote_id")
def then_anchor_run_contains_footnoteReference(context: Context):
    run = context.paragraph.runs[0]
    refs = run._r.xpath("./w:footnoteReference")  # pyright: ignore[reportPrivateUsage]
    assert len(refs) == 1, f"expected exactly 1 footnoteReference in run, got {len(refs)}"
    ref_id = int(refs[0].get(qn("w:id")))
    assert ref_id == context.footnote.footnote_id, (
        f"expected footnoteReference id {context.footnote.footnote_id}, got {ref_id}"
    )


@then("the anchor run has the FootnoteReference character style")
def then_anchor_run_has_FootnoteReference_style(context: Context):
    run = context.paragraph.runs[0]
    style = run._r.style  # pyright: ignore[reportPrivateUsage]
    assert style == "FootnoteReference", f"expected run style 'FootnoteReference', got '{style}'"


@then("fn1.footnote_id == {expected:d}")
def then_fn1_footnote_id_eq(context: Context, expected: int):
    actual = context.fn1.footnote_id
    assert actual == expected, f"expected fn1.footnote_id {expected}, got {actual}"


@then("fn2.footnote_id == {expected:d}")
def then_fn2_footnote_id_eq(context: Context, expected: int):
    actual = context.fn2.footnote_id
    assert actual == expected, f"expected fn2.footnote_id {expected}, got {actual}"


@then('document.footnotes[0].text == "{expected}"')
def then_document_footnotes_0_text_eq(context: Context, expected: str):
    first = next(iter(context.document.footnotes))
    actual = first.text
    assert actual == expected, f"expected text '{expected}', got '{actual}'"


@then("document.footnotes[0].footnote_id == {expected:d}")
def then_document_footnotes_0_footnote_id_eq(context: Context, expected: int):
    first = next(iter(context.document.footnotes))
    actual = first.footnote_id
    assert actual == expected, f"expected footnote_id {expected}, got {actual}"


@then("document.footnotes is a Footnotes object")
def then_document_footnotes_is_a_Footnotes_object(context: Context):
    assert isinstance(context.document.footnotes, Footnotes), (
        f"expected Footnotes, got {type(context.document.footnotes)}"
    )


@then("iterating document.footnotes yields {count:d} Footnote objects")
def then_iterating_document_footnotes_yields_count(context: Context, count: int):
    items = list(context.document.footnotes)
    assert len(items) == count, f"expected {count} footnotes, got {len(items)}"
    for fn in items:
        assert isinstance(fn, Footnote), f"expected Footnote, got {type(fn)}"


@then("the separator and continuation-separator footnotes are not yielded")
def then_separator_and_continuation_separator_not_yielded(context: Context):
    ids = [fn.footnote_id for fn in context.document.footnotes]
    assert 0 not in ids, "separator footnote (id 0) should not be yielded"
    assert 1 not in ids, "continuationSeparator footnote (id 1) should not be yielded"


@then("the yielded footnote ids are [{ids}]")
def then_yielded_footnote_ids_are(context: Context, ids: str):
    expected = [int(s.strip()) for s in ids.split(",")]
    actual = [fn.footnote_id for fn in context.document.footnotes]
    assert actual == expected, f"expected ids {expected}, got {actual}"


@then('footnote with id {footnote_id:d} has text ""')
def then_footnote_with_id_has_empty_text(context: Context, footnote_id: int):
    footnote = _footnote_with_id(context.document, footnote_id)
    actual = footnote.text
    assert actual == "", f"expected empty text, got '{actual}'"


@then('footnote with id {footnote_id:d} has text "{expected}"')
def then_footnote_with_id_has_text(context: Context, footnote_id: int, expected: str):
    footnote = _footnote_with_id(context.document, footnote_id)
    actual = footnote.text
    assert actual == expected, f"expected text '{expected}', got '{actual}'"


@then("each footnote has one paragraph")
def then_each_footnote_has_one_paragraph(context: Context):
    for fn in context.document.footnotes:
        assert len(fn.paragraphs) == 1, (
            f"footnote {fn.footnote_id} has {len(fn.paragraphs)} paragraphs, expected 1"
        )


@then("each footnote paragraph has the FootnoteText style")
def then_each_footnote_paragraph_has_FootnoteText_style(context: Context):
    for fn in context.document.footnotes:
        for p in fn.paragraphs:
            style = p._p.style  # pyright: ignore[reportPrivateUsage]
            assert style == "FootnoteText", (
                f"footnote {fn.footnote_id} paragraph style is '{style}', expected 'FootnoteText'"
            )


@then("no footnoteReference with id {footnote_id:d} remains in the document body")
def then_no_footnoteReference_with_id_remains(context: Context, footnote_id: int):
    body = context.document._element.body  # pyright: ignore[reportPrivateUsage]
    refs = body.xpath(f'.//w:footnoteReference[@w:id="{footnote_id}"]')
    assert len(refs) == 0, (
        f"expected 0 footnoteReference elements with id {footnote_id}, got {len(refs)}"
    )


@then("footnote with id {footnote_id:d} has {count:d} paragraph")
@then("footnote with id {footnote_id:d} has {count:d} paragraphs")
def then_footnote_with_id_has_n_paragraphs(context: Context, footnote_id: int, count: int):
    footnote = _footnote_with_id(context.document, footnote_id)
    actual = len(footnote.paragraphs)
    assert actual == count, (
        f"footnote {footnote_id} has {actual} paragraphs, expected {count}"
    )


@then("footnote with id {footnote_id:d} paragraph has the FootnoteText style")
def then_footnote_with_id_paragraph_has_FootnoteText_style(context: Context, footnote_id: int):
    footnote = _footnote_with_id(context.document, footnote_id)
    style = footnote.paragraphs[0]._p.style  # pyright: ignore[reportPrivateUsage]
    assert style == "FootnoteText", (
        f"footnote {footnote_id} paragraph style is '{style}', expected 'FootnoteText'"
    )


@then('the new paragraph has text "{expected}"')
def then_new_paragraph_has_text(context: Context, expected: str):
    actual = context.new_paragraph.text
    assert actual == expected, f"expected text '{expected}', got '{actual}'"


@then("the new paragraph has the FootnoteText style")
def then_new_paragraph_has_FootnoteText_style(context: Context):
    style = context.new_paragraph._p.style  # pyright: ignore[reportPrivateUsage]
    assert style == "FootnoteText", (
        f"expected paragraph style 'FootnoteText', got '{style}'"
    )


@then("document.footnote_properties is a FootnoteProperties object")
def then_footnote_properties_is_FootnoteProperties(context: Context):
    fp = context.document.footnote_properties
    assert isinstance(fp, FootnoteProperties), (
        f"expected FootnoteProperties, got {type(fp)}"
    )


@then("document.footnote_properties is None")
def then_footnote_properties_is_None(context: Context):
    fp = context.document.footnote_properties
    assert fp is None, f"expected None, got {fp!r}"


@then("document.footnote_properties.number_format == WD_NUMBER_FORMAT.{member}")
def then_number_format_eq(context: Context, member: str):
    actual = context.document.footnote_properties.number_format
    expected = WD_NUMBER_FORMAT[member]
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("document.footnote_properties.number_format is None")
def then_number_format_is_None(context: Context):
    actual = context.document.footnote_properties.number_format
    assert actual is None, f"expected None, got {actual!r}"


@then("document.footnote_properties.start_number == {expected:d}")
def then_start_number_eq(context: Context, expected: int):
    actual = context.document.footnote_properties.start_number
    assert actual == expected, f"expected {expected}, got {actual}"


@then("document.footnote_properties.start_number is None")
def then_start_number_is_None(context: Context):
    actual = context.document.footnote_properties.start_number
    assert actual is None, f"expected None, got {actual!r}"


@then("document.footnote_properties.restart_rule == WD_FOOTNOTE_RESTART.{member}")
def then_restart_rule_eq(context: Context, member: str):
    actual = context.document.footnote_properties.restart_rule
    expected = WD_FOOTNOTE_RESTART[member]
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("document.footnote_properties.restart_rule is None")
def then_restart_rule_is_None(context: Context):
    actual = context.document.footnote_properties.restart_rule
    assert actual is None, f"expected None, got {actual!r}"


@then("document.footnote_properties.position == WD_FOOTNOTE_POSITION.{member}")
def then_position_eq(context: Context, member: str):
    actual = context.document.footnote_properties.position
    expected = WD_FOOTNOTE_POSITION[member]
    assert actual == expected, f"expected {expected!r}, got {actual!r}"


@then("document.footnote_properties.position is None")
def then_position_is_None(context: Context):
    actual = context.document.footnote_properties.position
    assert actual is None, f"expected None, got {actual!r}"


# helpers ===================================================


def _footnote_with_id(document: Document, footnote_id: int) -> Footnote:
    """Return the user footnote with `footnote_id` from `document.footnotes`.

    Raises ``AssertionError`` if no such footnote exists so step failures are
    easy to diagnose.
    """
    for fn in document.footnotes:
        if fn.footnote_id == footnote_id:
            return fn
    raise AssertionError(f"no footnote with id {footnote_id} in document.footnotes")
